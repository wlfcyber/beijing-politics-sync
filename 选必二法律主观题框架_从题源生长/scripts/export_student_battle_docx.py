#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path("/Users/wanglifei/Desktop/北京高考政治/选必二法律主观题框架_从题源生长")
IN_MD = ROOT / "12_final_baodian" / "选必二法律主观题满分宝典_学生战斗版.md"
OUT_DOCX = ROOT / "12_final_baodian" / "选必二法律主观题满分宝典_学生战斗版.docx"


def set_run_font(run, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    set_run_font(r, 9, bold)


def parse_cells(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def is_separator(cells: list[str]) -> bool:
    return all(c and set(c) <= {"-", ":", " "} for c in cells)


def add_table(doc: Document, table_lines: list[str]) -> None:
    rows = [parse_cells(line) for line in table_lines]
    rows = [row for row in rows if not is_separator(row)]
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            set_cell_text(table.cell(r_idx, c_idx), row[c_idx] if c_idx < len(row) else "", bold=(r_idx == 0))


def export() -> None:
    lines = IN_MD.read_text(encoding="utf-8").splitlines()
    doc = Document()

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)

    section = doc.sections[0]
    section.top_margin = Pt(50)
    section.bottom_margin = Pt(50)
    section.left_margin = Pt(50)
    section.right_margin = Pt(50)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue

        if line.startswith("|"):
            table_lines: list[str] = []
            while i < len(lines):
                current = lines[i].rstrip()
                if current.startswith("|"):
                    table_lines.append(current)
                    i += 1
                    continue
                if current == "" and i + 1 < len(lines) and lines[i + 1].startswith("|"):
                    i += 1
                    continue
                break
            add_table(doc, table_lines)
            continue

        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=0)
            for run in p.runs:
                set_run_font(run, 18, True)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=1)
            for run in p.runs:
                set_run_font(run, 15, True)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=2)
            for run in p.runs:
                set_run_font(run, 12, True)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(line[2:].strip())
            set_run_font(r, 10)
        elif line[0].isdigit() and ". " in line[:4]:
            p = doc.add_paragraph(style="List Number")
            r = p.add_run(line.split(". ", 1)[1].strip())
            set_run_font(r, 10)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            r = p.add_run(line[2:].strip())
            set_run_font(r, 10)
            r.italic = True
        else:
            p = doc.add_paragraph()
            r = p.add_run(line)
            set_run_font(r, 10.5)
        i += 1

    props = doc.core_properties
    props.title = "选必二法律主观题满分宝典（学生战斗版）"
    props.subject = "北京高考政治选必二法律主观题"
    props.author = "Codex"
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    export()
