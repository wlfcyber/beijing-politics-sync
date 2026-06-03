#!/usr/bin/env python3
"""Generate guarded baodian markdown and sidecars from framework_v2 artifacts."""

from __future__ import annotations

import csv
import shutil
from datetime import datetime
from pathlib import Path


ROOT = Path("/Users/wanglifei/Desktop/北京高考政治/选必二法律主观题框架_从题源生长")
TEST = ROOT / "10_framework_validation" / "framework_v1_2_question_by_question_test_20260519.csv"
FRAMEWORK = ROOT / "11_final_framework" / "framework_v2.md"
STUDENT = ROOT / "11_final_framework" / "framework_v2_student_one_page.md"
TEACHER = ROOT / "11_final_framework" / "framework_v2_teacher_guide.md"
VALIDATION = ROOT / "11_final_framework" / "framework_v2_validation_summary.md"
EVIDENCE_MAP = ROOT / "11_final_framework" / "framework_v2_evidence_map.csv"

OUT_DIR = ROOT / "12_final_baodian"
OUT_MD = OUT_DIR / "选必二法律主观题满分宝典.md"
OUT_DOCX = OUT_DIR / "选必二法律主观题满分宝典.docx"
OUT_RUNS = OUT_DIR / "question_by_question_framework_runs.csv"
OUT_SENTENCES = OUT_DIR / "full_score_sentence_bank.csv"
OUT_FAILURES = OUT_DIR / "common_failure_paths.md"
OUT_TRIGGERS = OUT_DIR / "material_trigger_bank.csv"


def read_csv(path: Path) -> tuple[list[str], list[dict[str, str]]]:
    with path.open(encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        return list(reader.fieldnames or []), list(reader)


def write_csv(path: Path, fields: list[str], rows: list[dict[str, str]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        writer.writerows(rows)


def backup_if_exists(path: Path) -> None:
    if path.exists():
        backup = path.with_name(f"{path.stem}.pre_guarded_v2_{datetime.now().strftime('%Y%m%d_%H%M%S')}{path.suffix}")
        shutil.copy2(path, backup)


def label_for(row: dict[str, str]) -> str:
    if row["framework_entry_node"].startswith("FWV1_2_GATE"):
        return "boundary_non_core"
    if row["framework_entry_node"].startswith("FWV1_2_OPEN"):
        return "open_container_only"
    if row["evidence_level"] == "reference_only":
        return "reference_only_demo"
    if row["pass_status"] == "PASS":
        return "core_full_score_supported"
    return "formal_open_container_partial"


def node_class_for(label: str) -> str:
    if label == "core_full_score_supported":
        return "core_node"
    if label == "boundary_non_core":
        return "boundary_gate"
    if label in {"open_container_only", "formal_open_container_partial"}:
        return "open_container"
    if label == "reference_only_demo":
        return "reference_only_demo"
    return "uncertain"


def sentence_type_for(label: str) -> str:
    return {
        "core_full_score_supported": "满分句",
        "formal_open_container_partial": "题内参考句",
        "reference_only_demo": "参考答案句",
        "open_container_only": "开放容器题内参考句",
        "boundary_non_core": "边界说明",
    }[label]


def guarded_sentence(label: str, text: str) -> str:
    text = text or ""
    if label == "core_full_score_supported":
        return text
    prefix = {
        "formal_open_container_partial": "【题内参考句，非核心迁移模板】",
        "reference_only_demo": "【reference_only参考句，不能支撑核心节点】",
        "open_container_only": "【开放容器题内参考句，等待同型formal样本】",
        "boundary_non_core": "【边界说明，不生成选必二核心满分句】",
    }[label]
    if text.startswith(prefix):
        return text
    return f"{prefix}{text}"


def generate_sidecars() -> tuple[list[dict[str, str]], list[dict[str, str]]]:
    _, test_rows = read_csv(TEST)
    run_fields = [
        "question_id",
        "year",
        "district",
        "exam_stage",
        "question_no",
        "evidence_level",
        "baodian_label",
        "node_class",
        "sentence_type",
        "framework_entry_node",
        "why_this_entry",
        "ask_text",
        "material_trigger",
        "minimum_judgment_required",
        "legal_knowledge_triggered",
        "rubric_atom_ids_matched",
        "full_score_sentence_generated",
        "complete_answer_generated",
        "pass_status",
        "notes",
    ]
    run_rows = []
    for row in test_rows:
        label = label_for(row)
        node_class = node_class_for(label)
        sentence_type = sentence_type_for(label)
        notes = {
            "core_full_score_supported": "可作为核心满分示范。",
            "formal_open_container_partial": "正式证据开放容器；可讲题，不支撑新核心节点。",
            "reference_only_demo": "reference_only；只作弱示范，不支撑核心节点。",
            "open_container_only": "开放容器；等待更多同型 formal 样本。",
            "boundary_non_core": "边界题；不生成选必二核心满分答案。",
        }[label]
        run_rows.append(
            {
                "question_id": row["question_id"],
                "year": row["year"],
                "district": row["district"],
                "exam_stage": row["exam_stage"],
                "question_no": row["question_no"],
                "evidence_level": row["evidence_level"],
                "baodian_label": label,
                "node_class": node_class,
                "sentence_type": sentence_type,
                "framework_entry_node": row["framework_entry_node"],
                "why_this_entry": row["why_this_entry"],
                "ask_text": row["ask_text"],
                "material_trigger": row["material_trigger"],
                "minimum_judgment_required": row["minimum_judgment_required"],
                "legal_knowledge_triggered": row["legal_knowledge_triggered"],
                "rubric_atom_ids_matched": row["rubric_atom_ids_matched"],
                "full_score_sentence_generated": guarded_sentence(label, row["full_score_sentence_generated"]),
                "complete_answer_generated": row["complete_answer_generated"],
                "pass_status": row["pass_status"],
                "notes": notes,
            }
        )
    write_csv(OUT_RUNS, run_fields, run_rows)

    map_fields, map_rows = read_csv(EVIDENCE_MAP)
    trigger_fields = [
        "node_id",
        "node_class",
        "node_name",
        "ask_trigger",
        "material_trigger",
        "minimum_judgment",
        "supporting_question_ids",
        "supporting_rubric_atom_ids",
    ]
    trigger_rows = [
        {field: row.get(field, "") for field in trigger_fields}
        for row in map_rows
    ]
    write_csv(OUT_TRIGGERS, trigger_fields, trigger_rows)

    sent_fields = [
        "node_id",
        "node_class",
        "node_name",
        "sentence_template",
        "supporting_question_ids",
        "supporting_rubric_atom_ids",
        "usage_condition",
        "do_not_use_when",
    ]
    sent_rows = []
    for row in map_rows:
        node_id = row["node_id"]
        if node_id in {"FWV1_2_GATE", "FWV1_2_OPEN"}:
            usage = "boundary/open only; not a core full-score template"
            ban = "Do not use as ordinary selected-compulsory-2 core answer."
        else:
            usage = row.get("ask_trigger", "")
            ban = row.get("risk", "")
        sent_rows.append(
            {
                "node_id": node_id,
                "node_class": row.get("node_class", "core_node"),
                "node_name": row["node_name"],
                "sentence_template": row.get("answer_pattern", ""),
                "supporting_question_ids": row.get("supporting_question_ids", ""),
                "supporting_rubric_atom_ids": row.get("supporting_rubric_atom_ids", ""),
                "usage_condition": usage,
                "do_not_use_when": ban,
            }
        )
    write_csv(OUT_SENTENCES, sent_fields, sent_rows)
    return run_rows, sent_rows


def generate_failure_paths() -> None:
    md = [
        "# Common Failure Paths",
        "",
        "1. 把综合法治/国家治理题误收进选必二核心：用边界先判挡住。",
        "2. 只写公平正义、法治社会等价值话，不写具体法律规则和材料事实：意义题必须由规则推出。",
        "3. 评析题不先表态：先写支持/不支持/有效/无效。",
        "4. 案例题不区分合同、侵权、消费欺诈、不正当竞争：先定性，再写规则。",
        "5. 程序题乱写协商、上诉、行政诉讼：先看材料程序阶段。",
        "6. reference_only 题当核心证据：只作弱示范。",
        "7. open-container 题当稳定模板：必须等同型 formal 样本重复出现。",
        "",
    ]
    OUT_FAILURES.write_text("\n".join(md), encoding="utf-8")


def section_for_question(row: dict[str, str]) -> list[str]:
    if row["baodian_label"] == "core_full_score_supported":
        scoring_heading = "#### 五、细则对应与满分句"
        sentence_label = "满分句"
    elif row["baodian_label"] == "reference_only_demo":
        scoring_heading = "#### 五、参考答案对应与参考答案句"
        sentence_label = "参考答案句"
    elif row["baodian_label"] == "boundary_non_core":
        scoring_heading = "#### 五、边界说明"
        sentence_label = "边界说明"
    else:
        scoring_heading = "#### 五、细则对应与题内参考句"
        sentence_label = "题内参考句"
    lines = [
        f"### {row['question_id']}",
        "",
        f"- 题源：{row['year']} {row['district']} {row['exam_stage']} 第{row['question_no']}题",
        f"- 证据等级：{row['evidence_level']}",
        f"- 标签：{row['baodian_label']}",
        f"- 节点类别：{row['node_class']}",
        f"- 设问：{row['ask_text']}",
        "",
        "#### 一、框架入口",
        "",
        f"- 入口：{row['framework_entry_node']}",
        f"- 为什么：{row['why_this_entry']}",
        "",
        "#### 二、材料分层",
        "",
        f"- 材料触发：{row['material_trigger']}",
        "",
        "#### 三、最小必要判断",
        "",
        f"- {row['minimum_judgment_required']}",
        "",
        "#### 四、法律知识触发",
        "",
        f"- {row['legal_knowledge_triggered']}",
        "",
        scoring_heading,
        "",
        f"- 细则原子：{row['rubric_atom_ids_matched']}",
        f"- {sentence_label}：{row['full_score_sentence_generated']}",
        "",
        "#### 六、完整考场版答案",
        "",
    ]
    if row["baodian_label"] == "core_full_score_supported":
        lines.append(row["complete_answer_generated"] or row["full_score_sentence_generated"])
    elif row["baodian_label"] == "boundary_non_core":
        lines.append("本题作为边界题处理，不生成选必二核心满分答案。")
    else:
        lines.append("本题暂作开放/弱示范处理，不宣称形成核心满分闭环。")
    lines.extend(
        [
            "",
            "#### 七、易错路径",
            "",
            f"- {row['notes']}",
            "",
            "#### 八、迁移提醒",
            "",
            "- 只有证据等级、设问功能、材料触发和细则原子都匹配时，才可迁移本题路径。",
            "",
        ]
    )
    return lines


def generate_markdown(run_rows: list[dict[str, str]], sent_rows: list[dict[str, str]]) -> None:
    framework_text = FRAMEWORK.read_text(encoding="utf-8")
    student_text = STUDENT.read_text(encoding="utf-8")
    teacher_text = TEACHER.read_text(encoding="utf-8")
    validation_text = VALIDATION.read_text(encoding="utf-8")

    md = [
        "# 选必二法律主观题满分宝典（guarded v2 核心版）",
        "",
        "版本：guarded v2。核心框架可教学；开放、参考、边界题已显式标注，不冒充65题全部核心满分闭环。",
        "",
        "# 第一部分：Guarded v2 核心路径框架",
        "",
        framework_text,
        "",
        "## 学生版速记",
        "",
        student_text,
        "",
        "## 教师版解释",
        "",
        teacher_text,
        "",
        "# 第二部分：核心满分句库与非核心警戒句",
        "",
    ]
    for row in sent_rows:
        md.extend(
            [
                f"## {row['node_id']} {row['node_name']}",
                "",
                f"- 节点类别：{row.get('node_class', '')}",
                f"- 句式模板：{row['sentence_template']}",
                f"- 对应题源：{row['supporting_question_ids']}",
                f"- 对应细则原子：{row['supporting_rubric_atom_ids']}",
                f"- 使用条件：{row['usage_condition']}",
                f"- 禁止乱用：{row['do_not_use_when']}",
                "",
            ]
        )
    md.append("# 第三部分：全部主观题逐题框架运行示范")
    md.append("")
    for row in run_rows:
        md.extend(section_for_question(row))
    md.extend(
        [
            "# 第四部分：框架压测报告",
            "",
            validation_text,
            "",
        ]
    )
    OUT_MD.write_text("\n".join(md), encoding="utf-8")


def generate_docx() -> None:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover
        (OUT_DIR / "DOCX_GENERATION_FAILED.md").write_text(
            f"python-docx unavailable: {exc}\nMarkdown is available at {OUT_MD}\n",
            encoding="utf-8",
        )
        return
    doc = Document()
    text = OUT_MD.read_text(encoding="utf-8")
    for line in text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line)
    doc.save(OUT_DOCX)
    (OUT_DIR / "DOCX_QA_GUARDED_V2.md").write_text(
        "Generated with python-docx from guarded v2 Markdown. Word/PDF visual verification not yet run in this pass.\n",
        encoding="utf-8",
    )


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for path in [OUT_MD, OUT_DOCX, OUT_RUNS, OUT_SENTENCES, OUT_FAILURES, OUT_TRIGGERS]:
        backup_if_exists(path)
    run_rows, sent_rows = generate_sidecars()
    generate_failure_paths()
    generate_markdown(run_rows, sent_rows)
    generate_docx()
    print(f"Wrote {OUT_MD}")
    print(f"Wrote {OUT_RUNS}")
    print(f"Wrote {OUT_SENTENCES}")
    print(f"Wrote {OUT_TRIGGERS}")
    print(f"Wrote {OUT_FAILURES}")


if __name__ == "__main__":
    main()
