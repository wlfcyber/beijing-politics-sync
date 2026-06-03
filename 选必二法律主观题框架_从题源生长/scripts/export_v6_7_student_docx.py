from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path("/Users/wanglifei/Desktop/北京高考政治/选必二法律主观题框架_从题源生长")
MD = ROOT / "12_final_baodian" / "选必二法律主观题满分训练宝典_v6_7_学生使用版_20260521.md"
OUT = ROOT / "12_final_baodian" / "选必二法律主观题满分训练宝典_v6_7_学生使用版_20260521.docx"
REPORT = ROOT / "12_final_baodian" / "DOCX_EXPORT_V6_7_STUDENT_20260521.md"


FONT = "PingFang SC"
BLUE = RGBColor(31, 78, 121)
LIGHT = "EAF2F8"
WARN = "FFF4CE"


def set_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cols = len(table.columns)
    total = 6.4
    if cols == 2:
        widths = [1.75, total - 1.75]
    elif cols == 3:
        widths = [1.6, 1.9, total - 3.5]
    elif cols == 4:
        widths = [1.25, 1.75, 1.85, total - 4.85]
    else:
        widths = [total / max(cols, 1)] * cols
    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def clean_inline(text: str) -> str:
    return text.replace("`", "")


def add_inline(paragraph, text: str, size: float | None = None) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        value = part[2:-2] if bold else part
        run = paragraph.add_run(clean_inline(value))
        set_font(run, size=size, bold=bold)


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.header_distance = Cm(1.0)
    sec.footer_distance = Cm(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color, before, after in [
        ("Title", 22, BLUE, 0, 10),
        ("Heading 1", 16, BLUE, 14, 7),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 11.5, RGBColor(67, 67, 67), 8, 4),
    ]:
        st = styles[name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        if name != "Title":
            st.font.bold = True


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def parse_table(lines: list[str], i: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    while i < len(lines) and is_table_line(lines[i]):
        raw = lines[i].strip().strip("|")
        cells = [c.strip() for c in raw.split("|")]
        if not all(set(c) <= {"-", ":", " "} for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            text = row[c_idx] if c_idx < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, text, size=8.7 if col_count >= 4 else 9.2)
            if r_idx == 0:
                shade_cell(cell, LIGHT)
                for run in p.runs:
                    run.bold = True
            if "不能" in text or "不要" in text or "禁止" in text:
                if r_idx != 0:
                    shade_cell(cell, WARN)
    set_table_width(table)
    doc.add_paragraph()


def build() -> None:
    text = MD.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    style_doc(doc)

    in_code = False
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph(style=None)
            run = p.add_run(line)
            set_font(run, size=9.5)
            shade = OxmlElement("w:shd")
            shade.set(qn("w:fill"), "F4F6F9")
            p._p.get_or_add_pPr().append(shade)
            i += 1
            continue
        if is_table_line(line):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Title")
            add_inline(p, line[2:].strip())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, line[3:].strip())
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[4:].strip())
        elif line.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, line[5:].strip())
        elif re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s+", "", line).strip())
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:].strip())
        else:
            p = doc.add_paragraph()
            add_inline(p, line)
        i += 1

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("选必二法律主观题满分训练宝典 · 学生使用版")
    set_font(run, size=8.5, color=RGBColor(100, 100, 100))
    doc.save(OUT)


def main() -> None:
    build()
    REPORT.write_text(
        "# V6.7 学生使用版 DOCX 导出报告\n\n"
        f"- 输入 Markdown：`{MD}`\n"
        f"- 输出 DOCX：`{OUT}`\n"
        "- 样式预设：compact_reference_guide 取向；中文字体 PingFang SC；表格固定宽度并自动换行。\n"
        "- 后续要求：用 render_docx.py 渲染成 PNG/PDF 后做视觉 QA。\n",
        encoding="utf-8",
    )
    print(OUT)
    print(REPORT)


if __name__ == "__main__":
    main()
