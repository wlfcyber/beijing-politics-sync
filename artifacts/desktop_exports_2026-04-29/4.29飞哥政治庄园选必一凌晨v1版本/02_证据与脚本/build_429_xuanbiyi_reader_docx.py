from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DELIVERY = ROOT / "01_交付件"
SOURCE_MD = DELIVERY / "4.29飞哥政治庄园选必一凌晨v1版本.md"
OUT_DOCX = DELIVERY / "4.29飞哥政治庄园选必一凌晨v1版本_阅读版.docx"

ACCENT = RGBColor(35, 82, 132)
MUTED = RGBColor(90, 96, 104)
ANSWER = RGBColor(28, 98, 73)


def set_east_asia(run, font_name="宋体"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def style_border(paragraph, color="D7E3F1"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def setup_doc(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.85)
    section.bottom_margin = Cm(1.75)
    section.left_margin = Cm(1.95)
    section.right_margin = Cm(1.95)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(4)

    heading_specs = {
        "Heading 1": (16, ACCENT),
        "Heading 2": (14, ACCENT),
        "Heading 3": (12, RGBColor(40, 40, 40)),
        "Heading 4": (10.5, RGBColor(40, 40, 40)),
        "Title": (22, ACCENT),
    }
    for name, (size, color) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)


def parse_markdown(text):
    items = []
    lines = text.splitlines()
    i = 0
    current_case = None

    def flush_case():
        nonlocal current_case
        if current_case:
            items.append(current_case)
            current_case = None

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue

        if line.startswith("# "):
            flush_case()
            items.append({"type": "title", "text": line[2:].strip()})
        elif line.startswith("## "):
            flush_case()
            items.append({"type": "h1", "text": line[3:].strip()})
        elif line.startswith("### "):
            flush_case()
            items.append({"type": "h2", "text": line[4:].strip()})
        elif line.startswith("**术语：") and line.endswith("**"):
            flush_case()
            items.append({"type": "term", "text": line[2:-2].strip()})
        elif re.fullmatch(r"\d+\.", line):
            flush_case()
            current_case = {"type": "case", "index": line[:-1], "fields": []}
        elif line.startswith("- ") and "：" in line:
            label, value = line[2:].split("：", 1)
            if current_case and label in {"完整设问", "细则位置", "来源", "材料触发", "答案句"}:
                current_case["fields"].append((label, value.strip()))
            else:
                flush_case()
                items.append({"type": "bullet", "label": label.strip(), "text": value.strip()})
        elif line.startswith("- "):
            flush_case()
            items.append({"type": "bullet", "text": line[2:].strip()})
        else:
            flush_case()
            items.append({"type": "paragraph", "text": re.sub(r"\*\*(.*?)\*\*", r"\1", line)})
        i += 1

    flush_case()
    return items


def add_intro_bullet(doc, item):
    p = doc.add_paragraph(style="List Bullet")
    if "label" in item:
        r = p.add_run(item["label"] + "：")
        r.bold = True
        set_east_asia(r, "黑体")
        p.add_run(item["text"])
    else:
        p.add_run(item["text"])
    for run in p.runs:
        set_east_asia(run)


def add_term(doc, text):
    p = doc.add_paragraph()
    shade_paragraph(p, "EEF5FC")
    style_border(p, "8FB3D9")
    p.paragraph_format.left_indent = Cm(0.15)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(20, 67, 112)
    set_east_asia(r, "黑体")


def add_labeled_para(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.first_line_indent = Cm(-0.55)
    p.paragraph_format.space_after = Pt(3)

    r1 = p.add_run(label + "：")
    r1.bold = True
    r1.font.size = Pt(10.2)
    r1.font.color.rgb = ACCENT if label != "答案句" else ANSWER
    set_east_asia(r1, "黑体")

    r2 = p.add_run(value)
    r2.font.size = Pt(10.2)
    r2.font.color.rgb = MUTED if label == "细则位置" else RGBColor(30, 30, 30)
    if label == "答案句":
        r2.font.color.rgb = ANSWER
    set_east_asia(r2)


def add_case(doc, item):
    source = ""
    for label, value in item["fields"]:
        if label == "来源":
            source = value
            break

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.05)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    title = f"例{item['index']}"
    if source:
        title += f"｜{source}"
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.8)
    r.font.color.rgb = RGBColor(47, 86, 122)
    set_east_asia(r, "黑体")

    for label, value in item["fields"]:
        add_labeled_para(doc, label, value)


def build_docx():
    items = parse_markdown(SOURCE_MD.read_text(encoding="utf-8"))
    doc = Document()
    setup_doc(doc)

    for item in items:
        kind = item["type"]
        if kind == "title":
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(item["text"] + "（阅读版）")
            set_east_asia(r, "黑体")
        elif kind == "h1":
            doc.add_heading(item["text"], level=1)
        elif kind == "h2":
            doc.add_heading(item["text"], level=2)
        elif kind == "term":
            add_term(doc, item["text"])
        elif kind == "case":
            add_case(doc, item)
        elif kind == "bullet":
            add_intro_bullet(doc, item)
        else:
            p = doc.add_paragraph(item["text"])
            for run in p.runs:
                set_east_asia(run)

    doc.core_properties.title = "4.29飞哥政治庄园选必一凌晨v1版本（阅读版）"
    doc.core_properties.subject = "选择性必修一《当代国际政治与经济》评分术语积累"
    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print(build_docx())
