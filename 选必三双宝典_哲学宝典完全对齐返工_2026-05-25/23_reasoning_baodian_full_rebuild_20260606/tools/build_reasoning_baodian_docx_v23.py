from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
RUN = ROOT / "23_reasoning_baodian_full_rebuild_20260606"
IN_MD = RUN / "delivery" / "选必三逻辑与思维_推理宝典_按类型完整题源细则优秀答案版_20260606.md"
OUT_DOCX = RUN / "delivery" / "选必三逻辑与思维_推理宝典_按类型完整题源细则优秀答案版_20260606.docx"
QA_JSON = RUN / "qa" / "reasoning_baodian_v23_docx_structural_qa.json"


def set_east_asia(run, font_name: str) -> None:
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def set_para_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_style_font(style, font: str, size_pt: float, color: str | None = None, bold: bool | None = None) -> None:
    style.font.name = font
    style.font.size = Pt(size_pt)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)


def ensure_style(doc: Document, name: str, base: str, font: str, size_pt: float, color: str | None = None, bold: bool | None = None):
    styles = doc.styles
    if name in styles:
        style = styles[name]
    else:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles[base]
    set_style_font(style, font, size_pt, color, bold)
    return style


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, "Calibri", 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = doc.styles[style_name]
        set_style_font(style, "Calibri", size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    set_style_font(title, "Calibri", 20, "0B2545", True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    ensure_style(doc, "Entry Heading", "Normal", "Calibri", 11.5, "1F4D78", True).paragraph_format.space_before = Pt(8)
    ensure_style(doc, "Field Label", "Normal", "Calibri", 11, "0B2545", True)
    question = ensure_style(doc, "Question Text", "Normal", "Calibri", 10, "1F2937", None)
    question.paragraph_format.left_indent = Inches(0.18)
    question.paragraph_format.right_indent = Inches(0.05)
    question.paragraph_format.space_before = Pt(0)
    question.paragraph_format.space_after = Pt(2)
    question.paragraph_format.line_spacing = 1.15
    note = ensure_style(doc, "Method Bullet", "Normal", "Calibri", 10.5, "111827", None)
    note.paragraph_format.left_indent = Inches(0.3)
    note.paragraph_format.first_line_indent = Inches(-0.15)
    note.paragraph_format.space_after = Pt(3)


def add_label_value(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"【{label}】")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("0B2545")
    set_east_asia(r, "Microsoft YaHei")
    if value:
        r2 = p.add_run(f" {value}")
        set_east_asia(r2, "Microsoft YaHei")


def add_question_block(doc: Document, lines: list[str]) -> None:
    add_label_value(doc, "题目材料信息", "")
    p = doc.add_paragraph(style="Question Text")
    set_para_shading(p, "F4F6F9")
    for idx, raw in enumerate(lines):
        if idx:
            p.add_run().add_break()
        run = p.add_run(raw)
        set_east_asia(run, "Microsoft YaHei")
        run.font.size = Pt(10)


def add_method_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Method Bullet")
    run = p.add_run("• ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("2E74B5")
    run2 = p.add_run(text)
    set_east_asia(run2, "Microsoft YaHei")


def build_docx() -> None:
    doc = Document()
    configure_doc(doc)

    lines = IN_MD.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    counts = {
        "entries": 0,
        "question_blocks": 0,
        "rubrics": 0,
        "answers": 0,
        "method_sections": 0,
    }

    while i < len(lines):
        line = lines[i]

        if in_code:
            if line.strip().startswith("```"):
                add_question_block(doc, code_lines)
                counts["question_blocks"] += 1
                code_lines = []
                in_code = False
            else:
                code_lines.append(line)
            i += 1
            continue

        if line.startswith("```"):
            in_code = True
            code_lines = []
            i += 1
            continue

        if line.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[2:].strip())
            set_east_asia(r, "Microsoft YaHei")
            i += 1
            continue
        if line and not line.startswith("#") and i < 6:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line.strip())
            r.bold = True
            r.font.color.rgb = RGBColor.from_string("1F4D78")
            set_east_asia(r, "Microsoft YaHei")
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        if line.startswith("#### "):
            counts["entries"] += 1
            p = doc.add_paragraph(style="Entry Heading")
            r = p.add_run(line[5:].strip())
            set_east_asia(r, "Microsoft YaHei")
            i += 1
            continue
        if line == "【本类做题方法】":
            counts["method_sections"] += 1
            add_label_value(doc, "本类做题方法", "")
            i += 1
            continue
        if line.startswith("- "):
            add_method_bullet(doc, line[2:].strip())
            i += 1
            continue

        m = re.match(r"^【([^】]+)】\s*(.*)$", line)
        if m:
            label = m.group(1)
            value = m.group(2).strip()
            if label == "题目材料信息":
                i += 1
                continue
            add_label_value(doc, label, value)
            if label == "细则要点":
                counts["rubrics"] += 1
            if label == "考生优秀答案":
                counts["answers"] += 1
            i += 1
            continue

        if line.strip():
            p = doc.add_paragraph()
            r = p.add_run(line.strip())
            set_east_asia(r, "Microsoft YaHei")
        i += 1

    doc.save(OUT_DOCX)

    qa = {
        **counts,
        "docx": str(OUT_DOCX.relative_to(ROOT.parent)),
        "expected_entries": 83,
        "expected_fields_per_entry": {
            "question_blocks": 83,
            "rubrics": 83,
            "answers": 83,
        },
        "structural_pass": counts["entries"] == 83 and counts["question_blocks"] == 83 and counts["rubrics"] == 83 and counts["answers"] == 83,
    }
    QA_JSON.write_text(json.dumps(qa, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(qa, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    build_docx()
