from __future__ import annotations

import json
from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
RUN = ROOT / "23_reasoning_baodian_full_rebuild_20260606"
DOCX = RUN / "delivery" / "选必三逻辑与思维_推理宝典_按类型完整题源细则优秀答案版_20260606.docx"
OUT = RUN / "qa" / "DOCX_TEXT_EXTRACTION_QA_20260606.json"


def main() -> None:
    with ZipFile(DOCX) as zf:
        bad = zf.testzip()
    doc = Document(str(DOCX))
    text = "\n".join(par.text for par in doc.paragraphs)
    patterns = [
        "【题目材料信息】",
        "【设问/选项】",
        "【细则要点】",
        "【为什么能想到】",
        "【考生优秀答案】",
        "【本类做题方法】",
        "/Users",
        "C:\\",
        "source_extracted",
        "OCR",
        "A-formal",
        "B-choice-signal",
        "题号 |",
        "评标",
        "评分标准",
        "参考答案",
    ]
    result = {
        "zip_bad_file": bad,
        "paragraphs": len(doc.paragraphs),
        "sections": len(doc.sections),
        "counts": {pat: text.count(pat) for pat in patterns},
        "pass": bad is None
        and text.count("【题目材料信息】") == 83
        and text.count("【设问/选项】") == 83
        and text.count("【细则要点】") == 83
        and text.count("【为什么能想到】") == 83
        and text.count("【考生优秀答案】") == 83
        and all(text.count(pat) == 0 for pat in patterns[6:]),
    }
    OUT.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
