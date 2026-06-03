from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


RUN_DIR = Path(__file__).resolve().parents[1]
OUT_DIR = RUN_DIR / "07_docx_pdf"
MD_DIR = RUN_DIR / "05_candidate_md"

BODY_FONT = "Microsoft YaHei"
HEADING_FONT = "黑体"
SIGNATURE_FONT = "楷体"
ASCII_FONT = "Arial"
TITLE_ASCII_FONT = "Times New Roman"
TITLE_COLOR = RGBColor(0x16, 0x3E, 0x63)
BROWN_COLOR = RGBColor(0x5E, 0x3F, 0x23)
HEADING1_COLOR = "1F4E79"
HEADING2_COLOR = "2F6F9F"
HEADING3_COLOR = "3A6278"
LABEL_COLOR = RGBColor(0x21, 0x57, 0x4C)

WATERMARK_PICT_XML = r"""
<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        xmlns:v="urn:schemas-microsoft-com:vml"
        xmlns:o="urn:schemas-microsoft-com:office:office"
        xmlns:w10="urn:schemas-microsoft-com:office:word">
  <v:shapetype id="_x0000_t136" coordsize="21600,21600" o:spt="136" adj="10800"
      path="m@7,l@8,m@5,21600l@6,21600e">
    <v:formulas>
      <v:f eqn="sum #0 0 10800"/>
      <v:f eqn="prod #0 2 1"/>
      <v:f eqn="sum 21600 0 @1"/>
      <v:f eqn="sum 0 0 @2"/>
      <v:f eqn="sum 21600 0 @3"/>
      <v:f eqn="if @0 @3 0"/>
      <v:f eqn="if @0 21600 @1"/>
      <v:f eqn="if @0 0 @2"/>
      <v:f eqn="if @0 @4 21600"/>
      <v:f eqn="mid @5 @6"/>
      <v:f eqn="mid @8 @5"/>
      <v:f eqn="mid @7 @8"/>
      <v:f eqn="mid @6 @7"/>
      <v:f eqn="sum @6 0 @5"/>
    </v:formulas>
    <v:path textpathok="t" o:connecttype="custom"
        o:connectlocs="@9,0;@10,10800;@11,21600;@12,10800"
        o:connectangles="270,180,90,0"/>
    <v:textpath on="t" fitshape="t"/>
    <v:handles><v:h position="#0,bottomRight" xrange="6629,14971"/></v:handles>
    <o:lock v:ext="edit" text="t" shapetype="t"/>
  </v:shapetype>
  <v:shape id="PowerPlusWaterMarkObject_CodexFGZZJT"
      o:spid="_x0000_s1025" type="#_x0000_t136" alt=""
      style="position:absolute;margin-left:0;margin-top:0;width:440pt;height:90pt;rotation:315;z-index:-251658752;mso-wrap-edited:f;mso-width-percent:0;mso-height-percent:0;mso-position-horizontal:center;mso-position-horizontal-relative:page;mso-position-vertical:center;mso-position-vertical-relative:page;mso-width-percent:0;mso-height-percent:0"
      o:allowincell="f" fillcolor="#a6a6a6" stroked="f">
    <v:fill opacity="19661f"/>
    <v:textpath style="font-family:&quot;Microsoft YaHei&quot;,&quot;SimHei&quot;;font-size:42pt;font-weight:bold"
        string="飞哥正志讲堂"/>
    <w10:wrap anchorx="page" anchory="page"/>
    <w10:anchorlock/>
  </v:shape>
</w:pict>
""".strip()


BOOKS = [
    {
        "kind": "thinking",
        "title": "2026北京高考政治选必三《逻辑与思维》思维宝典",
        "subtitle": "三年模拟全触发全链条",
        "md": MD_DIR / "选必三_逻辑与思维_思维宝典_哲学完全对齐候选稿.md",
        "docx": OUT_DIR / "选必三_逻辑与思维_思维宝典_哲学完全对齐版.docx",
        "manual_toc": [
            ("一、科学思维", 4),
            ("追求认识的客观性", 4),
            ("结果具有预见性", 8),
            ("结果具有可检验性", 9),
            ("二、辩证思维", 11),
            ("整体性", 11),
            ("动态性", 14),
            ("分析与综合", 15),
            ("矛盾分析法", 17),
            ("量变与质变", 19),
            ("适度原则", 20),
            ("辩证否定", 21),
            ("认识发展历程", 22),
            ("三、创新思维", 23),
            ("特点与三新", 23),
            ("联想思维", 26),
            ("发散思维与聚合思维", 29),
            ("逆向思维", 32),
            ("超前思维", 33),
        ],
    },
    {
        "kind": "reasoning",
        "title": "2026北京高考政治选必三《逻辑与思维》推理宝典",
        "subtitle": "三年模拟同形推理全链条",
        "md": MD_DIR / "选必三_逻辑与思维_推理宝典_哲学完全对齐候选稿.md",
        "docx": OUT_DIR / "选必三_逻辑与思维_推理宝典_哲学完全对齐版.docx",
        "manual_toc": [
            ("一、充分条件假言推理与判断", 5),
            ("有效式：肯定前件或否定后件，方向要走对", 5),
            ("陷阱：后件为真不能倒推前件", 5),
            ("判断构造：把条件补成能推出结果的前件", 7),
            ("陷阱一：重要条件不等于充分条件", 7),
            ("陷阱二：形式像对，前提未必真", 8),
            ("二、必要条件假言推理与判断", 8),
            ("有效式：少了它不行，结果成立就要回找它", 8),
            ("陷阱：有了必要条件不等于结果必然成", 9),
            ("必要条件的文字信号", 10),
            ("必要条件与结论边界", 11),
            ("必要条件与定义边界", 12),
            ("除非、否则与双重否定", 13),
            ("必要条件与选言边界", 14),
            ("相容选言里不要误排斥", 15),
            ("三、三段论、性质判断周延与换质位", 15),
            ("先看三项：大项、小项、中项", 15),
            ("三段论形式正确，前提仍要真实", 16),
            ("补大前提：把省略规则说出来", 17),
            ("四概念：同一个词前后换了对象", 17),
            ("大项扩大：结论说得比前提更满", 18),
            ("小项扩大：主语范围被偷偷放大", 19),
            ("构造三段论：先定中项桥梁", 19),
            ("有效结构：形式有效还要看前提", 20),
            ("有效形式与真实前提要分开", 21),
            ("周延性：看概念有没有被说尽", 21),
            ("换位边界：不是每句话都能倒过来", 22),
            ("类比与换位换质不要混判", 23),
            ("换质位：变形不能变意思", 24),
            ("构建论证：让结论从共同前提长出来", 25),
            ("综合辨析：换位、三段论、外延分开看", 26),
            ("省略前提：把隐藏的大前提补回来", 27),
            ("四、归纳推理与探求因果联系", 28),
            ("不完全归纳推理", 28),
            ("不完全归纳推理可靠程度", 30),
            ("科学归纳推理与不完全归纳推理", 32),
            ("探求因果联系中的求异法", 32),
            ("科学归纳、求异法与类比推理", 33),
            ("五、类比推理", 33),
            ("类比推理", 33),
            ("六、概念、定义、外延关系与划分", 35),
            ("概念划分规则", 35),
            ("同一律与偷换概念", 36),
            ("概念外延关系错项辨析", 36),
            ("定义构成与属加种差定义", 37),
            ("概念外延中的相容关系与属种关系", 37),
            ("概念外延关系图示辨析", 37),
            ("定义方法与种差加属概念", 38),
            ("同一律与概念确定性", 39),
            ("七、选言推理、联言判断与复合判断", 39),
            ("选言判断", 39),
            ("反对关系与联言判断", 40),
            ("不相容选言推理（有效式）", 40),
            ("不相容选言推理无效式辨析", 41),
            ("复合假言与选言推理链", 41),
            ("选言推理与逆向思维复合辨析", 42),
            ("联言判断类型识别", 43),
            ("联言判断真值条件", 43),
            ("复合判断识别与逻辑规则辨析", 44),
            ("八、真假关系、逻辑规律与关系判断", 45),
            ("矛盾律", 45),
            ("矛盾关系与真假判断", 45),
            ("形式逻辑综合选择题", 46),
            ("关系判断与非传递关系", 47),
            ("排中律与两不可错误", 48),
            ("矛盾律与自相矛盾错误", 48),
            ("逻辑规则综合选择", 49),
            ("枚举概括与同一对象替换", 49),
            ("真假话约束推理", 50),
            ("矛盾律与一致性要求", 51),
        ],
    },
]


def set_east_asia_font(run, font_name: str = BODY_FONT, ascii_font: str = ASCII_FONT):
    run.font.name = ascii_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)


def set_style_font(style, east_asia=BODY_FONT, ascii_font=ASCII_FONT, size=Pt(10.5), color=None, bold=None):
    font = style.font
    font.name = ascii_font
    font.size = size
    if bold is not None:
        font.bold = bold
    if color:
        font.color.rgb = RGBColor.from_string(color)
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def add_page_ref_field(paragraph, bookmark_name: str, fallback_page_no: int):
    begin_run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(fld_begin)

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" PAGEREF {bookmark_name} \\h "
    instr_run._r.append(instr)

    sep_run = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(fld_sep)

    result_run = paragraph.add_run(str(fallback_page_no))
    set_east_asia_font(result_run)
    result_run.font.size = Pt(9.5)

    end_run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_end)


def set_raw_run_font(run_element, size_half_points: int = 19):
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:eastAsia"), BODY_FONT)
    r_fonts.set(qn("w:ascii"), ASCII_FONT)
    r_fonts.set(qn("w:hAnsi"), ASCII_FONT)
    r_pr.append(r_fonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size_half_points))
    r_pr.append(sz)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    r_pr.append(underline)

    run_element.append(r_pr)


def add_bookmark(paragraph, name: str, bookmark_id: int, text: str):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    paragraph._p.append(start)

    run = paragraph.add_run(text)
    set_east_asia_font(run, HEADING_FONT)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(end)
    return run


def add_internal_hyperlink(paragraph, bookmark_name: str, text: str):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    set_raw_run_font(run)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def normalize_toc_styles(docx_path: Path):
    """Use the philosophy benchmark's custom TOC style ids for manual TOC rows."""
    with zipfile.ZipFile(docx_path, "r") as zin:
        payload = {name: zin.read(name) for name in zin.namelist()}

    styles_xml = payload["word/styles.xml"].decode("utf-8")
    for source_id, target_id, target_name, source_name in (
        ("TOC11", "TOC1", "toc 1", "TOC 11"),
        ("TOC21", "TOC2", "toc 2", "TOC 21"),
    ):
        if f'w:styleId="{target_id}"' in styles_xml:
            continue
        match = re.search(rf'<w:style\b[^>]*w:styleId="{source_id}"[\s\S]*?</w:style>', styles_xml)
        if match:
            clone = match.group(0)
            clone = clone.replace(f'w:styleId="{source_id}"', f'w:styleId="{target_id}"', 1)
            clone = clone.replace(f'w:val="{source_name}"', f'w:val="{target_name}"', 1)
            styles_xml = styles_xml.replace("</w:styles>", clone + "</w:styles>", 1)

    for obsolete_id in ("TOC11", "TOC21"):
        styles_xml = re.sub(
            rf'<w:style\b[^>]*w:styleId="{obsolete_id}"[\s\S]*?</w:style>',
            "",
            styles_xml,
        )

    document_xml = payload["word/document.xml"].decode("utf-8")
    for source_id, target_id in (
        ("TOC11", "TOC1"),
        ("TOC21", "TOC2"),
    ):
        document_xml = document_xml.replace(f'w:val="{source_id}"', f'w:val="{target_id}"')

    payload["word/styles.xml"] = styles_xml.encode("utf-8")
    payload["word/document.xml"] = document_xml.encode("utf-8")

    tmp_path = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in payload.items():
            zout.writestr(name, data)
    tmp_path.replace(docx_path)


def freeze_pageref_fields(docx_path: Path):
    """Turn Word-updated PAGEREF results into static TOC page text."""
    with zipfile.ZipFile(docx_path, "r") as zin:
        payload = {name: zin.read(name) for name in zin.namelist()}

    document_xml = payload["word/document.xml"].decode("utf-8")
    run_with_begin = (
        r'<w:r\b[^>]*>(?:(?!</w:r>).)*?'
        r'<w:fldChar\b[^>]*w:fldCharType="begin"[^>]*/>'
        r'(?:(?!</w:r>).)*?</w:r>'
    )
    run_with_pageref = (
        r'<w:r\b[^>]*>(?:(?!</w:r>).)*?'
        r'<w:instrText\b[^>]*>[^<]*PAGEREF[^<]*</w:instrText>'
        r'(?:(?!</w:r>).)*?</w:r>'
    )
    run_with_separate = (
        r'<w:r\b[^>]*>(?:(?!</w:r>).)*?'
        r'<w:fldChar\b[^>]*w:fldCharType="separate"[^>]*/>'
        r'(?:(?!</w:r>).)*?</w:r>'
    )
    run_with_end = (
        r'<w:r\b[^>]*>(?:(?!</w:r>).)*?'
        r'<w:fldChar\b[^>]*w:fldCharType="end"[^>]*/>'
        r'(?:(?!</w:r>).)*?</w:r>'
    )
    field_pattern = re.compile(
        run_with_begin
        + r"\s*"
        + run_with_pageref
        + r"\s*"
        + run_with_separate
        + r"\s*"
        r'(?P<result>.*?)'
        + run_with_end,
        re.DOTALL,
    )

    def keep_result(match: re.Match):
        result = match.group("result")
        return result if "<w:t" in result else ""

    document_xml = field_pattern.sub(keep_result, document_xml)
    payload["word/document.xml"] = document_xml.encode("utf-8")

    settings_xml = payload.get("word/settings.xml")
    if settings_xml:
        settings_text = settings_xml.decode("utf-8")
        settings_text = re.sub(r'<w:updateFields\b[^>]*/>', "", settings_text)
        settings_text = re.sub(r'<w:updateFields\b[^>]*>.*?</w:updateFields>', "", settings_text, flags=re.DOTALL)
        payload["word/settings.xml"] = settings_text.encode("utf-8")

    tmp_path = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in payload.items():
            zout.writestr(name, data)
    tmp_path.replace(docx_path)


def toc_bookmark_map(manual_entries):
    if not manual_entries:
        return {}
    return {title: f"toc_anchor_{i:02d}" for i, (title, _page_no) in enumerate(manual_entries, start=1)}


def add_update_fields_setting(doc: Document):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is not None:
        settings.remove(update)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def configure_section_geometry(section, *, different_first_page: bool):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.15)
    section.right_margin = Cm(2.15)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    section.different_first_page_header_footer = different_first_page


def configure_document(doc: Document, running_label: str):
    section = doc.sections[0]
    configure_section_geometry(section, different_first_page=True)

    styles = doc.styles
    set_style_font(styles["Normal"], BODY_FONT, ASCII_FONT, Pt(10.5))
    styles["Normal"].paragraph_format.line_spacing = None
    styles["Normal"].paragraph_format.space_after = None

    set_style_font(styles["Heading 1"], HEADING_FONT, ASCII_FONT, Pt(18), HEADING1_COLOR, True)
    styles["Heading 1"].paragraph_format.space_before = Pt(24)
    styles["Heading 1"].paragraph_format.space_after = Pt(0)
    styles["Heading 1"].paragraph_format.keep_with_next = True

    set_style_font(styles["Heading 2"], HEADING_FONT, ASCII_FONT, Pt(14), HEADING2_COLOR, True)
    styles["Heading 2"].paragraph_format.space_before = Pt(10)
    styles["Heading 2"].paragraph_format.space_after = Pt(0)
    styles["Heading 2"].paragraph_format.keep_with_next = True

    set_style_font(styles["Heading 3"], HEADING_FONT, ASCII_FONT, Pt(11.5), HEADING3_COLOR, True)
    styles["Heading 3"].paragraph_format.space_before = Pt(10)
    styles["Heading 3"].paragraph_format.space_after = Pt(0)
    styles["Heading 3"].paragraph_format.keep_with_next = True

    set_style_font(styles["TOC Heading"], HEADING_FONT, ASCII_FONT, Pt(16), "163E63", True)
    for style_name, style_id, left_indent in (("toc 1", "TOC1", 0), ("toc 2", "TOC2", 0.74)):
        if style_name in styles:
            toc_style = styles[style_name]
        else:
            toc_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        toc_style._element.set(qn("w:styleId"), style_id)
        name_el = toc_style._element.find(qn("w:name"))
        if name_el is None:
            name_el = OxmlElement("w:name")
            toc_style._element.insert(0, name_el)
        name_el.set(qn("w:val"), style_name)
        set_style_font(toc_style, BODY_FONT, ASCII_FONT, Pt(9.5))
        toc_style.paragraph_format.left_indent = Cm(left_indent)
        toc_style.paragraph_format.space_after = Pt(0)

    # Philosophy benchmark uses a faint diagonal "飞哥正志讲堂" watermark rather than a running header.
    section.header.paragraphs[0].text = ""
    section.header.paragraphs[0].add_run()._r.append(parse_xml(WATERMARK_PICT_XML))
    section.first_page_header.paragraphs[0].text = ""
    section.first_page_footer.paragraphs[0].text = ""

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = footer.add_run("— ")
    set_east_asia_font(r1, BODY_FONT)
    r1.font.size = Pt(9)
    add_field(footer, "PAGE")
    r2 = footer.add_run(" —")
    set_east_asia_font(r2, BODY_FONT)
    r2.font.size = Pt(9)

    add_update_fields_setting(doc)


def add_body_section(doc: Document):
    """Mirror the philosophy benchmark's continuous section break before body."""
    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_section_geometry(body_section, different_first_page=False)
    return body_section


def add_title_block(doc: Document, title: str, subtitle: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(180)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(title)
    set_east_asia_font(r, HEADING_FONT, TITLE_ASCII_FONT)
    title_size = Pt(19) if len(title) > 22 else Pt(22)
    r.font.size = title_size
    r.font.bold = True
    r.font.color.rgb = TITLE_COLOR
    r.add_break()
    r = p.add_run(subtitle)
    set_east_asia_font(r, HEADING_FONT, TITLE_ASCII_FONT)
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = TITLE_COLOR

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("飞哥正志讲堂")
    set_east_asia_font(r, SIGNATURE_FONT, TITLE_ASCII_FONT)
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = BROWN_COLOR


def add_toc(doc: Document, manual_entries=None, bookmarks=None):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("目录")
    set_east_asia_font(r, HEADING_FONT, TITLE_ASCII_FONT)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = TITLE_COLOR
    if manual_entries:
        bookmarks = bookmarks or {}
        for title, page_no in manual_entries:
            toc_style = "toc 1" if re.match(r"^[一二三四五六七八九十]+、", title) else "toc 2"
            p = doc.add_paragraph(style=toc_style)
            p.paragraph_format.tab_stops.add_tab_stop(Cm(16.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            if title in bookmarks:
                add_internal_hyperlink(p, bookmarks[title], title)
            else:
                r = p.add_run(title)
                set_east_asia_font(r, BODY_FONT)
                r.font.size = Pt(10.5)
            p.add_run("\t")
            if title in bookmarks:
                add_page_ref_field(p, bookmarks[title], page_no)
            else:
                r = p.add_run(str(page_no))
                set_east_asia_font(r, BODY_FONT)
                r.font.size = Pt(10.5)
    else:
        p = doc.add_paragraph()
        add_field(p, r'TOC \o "1-2" \h \z')
    doc.add_page_break()


LABEL_RE = re.compile(r"^(【[^】]+】)(.*)$")


def add_body_paragraph(doc: Document, text: str):
    match = LABEL_RE.match(text)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if match:
        label, rest = match.groups()
        if rest and not rest.startswith(" "):
            rest = " " + rest
        r = p.add_run(label)
        set_east_asia_font(r, BODY_FONT)
        r.font.bold = True
        r.font.color.rgb = LABEL_COLOR
        if rest:
            r = p.add_run(rest)
            set_east_asia_font(r, BODY_FONT)
        return
    r = p.add_run(text)
    set_east_asia_font(r, BODY_FONT)


def iter_markdown_blocks(path: Path):
    raw = path.read_text(encoding="utf-8")
    raw = raw.replace("\r\n", "\n")
    lines = raw.split("\n")
    blocks: list[str] = []
    current: list[str] = []
    for line in lines:
        if not line.strip():
            if current:
                blocks.append("\n".join(current).strip())
                current = []
            continue
        if line.startswith("#"):
            if current:
                blocks.append("\n".join(current).strip())
                current = []
            blocks.append(line.strip())
        else:
            current.append(line.rstrip())
    if current:
        blocks.append("\n".join(current).strip())
    return blocks


def emit_heading(doc: Document, text: str, style: str, bookmarks: dict[str, str], bookmark_ids: dict[str, int]):
    p = doc.add_paragraph(style=style)
    if text in bookmarks:
        run = add_bookmark(p, bookmarks[text], bookmark_ids[text], text)
    else:
        run = p.add_run(text)
        set_east_asia_font(run, HEADING_FONT)
    keep_with_next(p)
    return p


def emit_content(doc: Document, md_path: Path, title: str, bookmarks=None):
    bookmarks = bookmarks or {}
    bookmark_ids = {heading: idx for idx, heading in enumerate(bookmarks, start=10)}
    seen_first_title = False
    skip_manual_toc = False
    body_started = False
    for block in iter_markdown_blocks(md_path):
        if block == f"# {title}" and not seen_first_title:
            seen_first_title = True
            continue
        if block in {"三年模拟全触发全链条", "三年模拟同形推理全链条", "同一推理形式考题汇总", "飞哥正志讲堂"}:
            continue
        if block == "## 目录":
            skip_manual_toc = True
            continue
        if skip_manual_toc:
            if not block.startswith("# "):
                continue
            skip_manual_toc = False
            body_started = True
        elif not body_started:
            if not block.startswith("# "):
                continue
            body_started = True
        if block.startswith("# "):
            emit_heading(doc, block[2:].strip(), "Heading 1", bookmarks, bookmark_ids)
        elif block.startswith("## "):
            emit_heading(doc, block[3:].strip(), "Heading 2", bookmarks, bookmark_ids)
        elif block.startswith("### "):
            emit_heading(doc, block[4:].strip(), "Heading 3", bookmarks, bookmark_ids)
        else:
            for part in block.split("\n"):
                if part.strip():
                    add_body_paragraph(doc, part.strip())


def collect_toc_entries(md_path: Path, title: str):
    entries = []
    body_started = False
    for block in iter_markdown_blocks(md_path):
        if block == "## 目录":
            body_started = True
            continue
        if not body_started:
            continue
        if block == f"# {title}":
            continue
        if block.startswith("# "):
            entries.append((block[2:].strip(), 0))
        elif block.startswith("## "):
            entries.append((block[3:].strip(), 0))
    return entries


def emit_front_matter(doc: Document, md_path: Path, title: str):
    seen_preface = False
    for block in iter_markdown_blocks(md_path):
        if block in {
            f"# {title}",
            "三年模拟全触发全链条",
            "三年模拟同形推理全链条",
            "同一推理形式考题汇总",
            "飞哥正志讲堂",
        }:
            continue
        if block == "## 前言":
            seen_preface = True
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("前言")
            set_east_asia_font(r, HEADING_FONT, TITLE_ASCII_FONT)
            r.font.size = Pt(20)
            r.font.bold = True
            r.font.color.rgb = TITLE_COLOR
            continue
        if seen_preface:
            if block == "## 目录" or block.startswith("# "):
                break
            for part in block.split("\n"):
                if part.strip():
                    add_body_paragraph(doc, part.strip())


def build_one(book):
    doc = Document()
    configure_document(doc, book["title"])
    add_title_block(doc, book["title"], book["subtitle"])
    doc.add_page_break()
    emit_front_matter(doc, book["md"], book["title"])
    manual_toc = book.get("manual_toc")
    if manual_toc == "auto":
        manual_toc = collect_toc_entries(book["md"], book["title"])
    bookmarks = toc_bookmark_map(manual_toc)
    add_toc(doc, manual_toc, bookmarks)
    add_body_section(doc)
    emit_content(doc, book["md"], book["title"], bookmarks)
    book["docx"].parent.mkdir(parents=True, exist_ok=True)
    doc.save(book["docx"])
    normalize_toc_styles(book["docx"])
    freeze_pageref_fields(book["docx"])
    return book["docx"]


def main():
    made = [build_one(book) for book in BOOKS]
    for path in made:
        print(path)


if __name__ == "__main__":
    sys.exit(main())
