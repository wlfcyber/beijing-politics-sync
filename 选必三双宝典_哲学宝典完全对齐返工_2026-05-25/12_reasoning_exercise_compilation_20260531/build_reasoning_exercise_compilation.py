#!/usr/bin/env python3
from __future__ import annotations

import csv
import hashlib
import json
import re
import subprocess
import unicodedata
import zipfile
from collections import defaultdict
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET

import fitz
from docx import Document


ROOT = Path("/Users/wanglifei/Desktop/北京高考政治")
RUN = ROOT / "选必三双宝典_哲学宝典完全对齐返工_2026-05-25"
OLD_RUN = ROOT / "选必三双宝典_哲学对齐重做_2026-05-25"
OUT = RUN / "12_reasoning_exercise_compilation_20260531"
FRAMEWORK_MD = RUN / "05_candidate_md/选必三_逻辑与思维_推理宝典_哲学完全对齐候选稿.md"
COVERAGE = RUN / "00_control/QUESTION_COVERAGE_MATRIX.csv"
OLD_AUDIT = OLD_RUN / "07_audit/REASONING_BODY_AUDIT.csv"
SOURCE_ROOTS = [
    Path("/Users/wanglifei/Desktop/2024模拟题"),
    Path("/Users/wanglifei/Desktop/2025模拟题"),
    Path("/Users/wanglifei/Desktop/2026模拟题"),
]


@dataclass
class Entry:
    placement_id: str
    family: str
    node: str
    title: str
    suite: str
    question_no: str
    question_type: str
    prompt_from_candidate: str
    answer_landing_from_candidate: str


def clean_text(text: str) -> str:
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r"[\ue000-\uf8ff\U000f0000-\U0010ffff]", ".", text)
    text = text.replace("\r", "\n")
    text = re.sub(r"[ \t\u3000]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def safe_name(s: str) -> str:
    s = re.sub(r"[^\w\u4e00-\u9fff.-]+", "_", s)
    return s[:120].strip("_") or "source"


def parse_framework() -> list[Entry]:
    lines = FRAMEWORK_MD.read_text(encoding="utf-8").splitlines()
    entries: list[Entry] = []
    family = ""
    node = ""
    pending_title = None
    buf: list[str] = []

    def flush():
        nonlocal pending_title, buf
        if not pending_title:
            return
        title = pending_title
        body = "\n".join(buf)
        m = re.match(r"\d+\.\s*(.+?)\s+第(.+?)（(主观题|选择题)）", title)
        if not m:
            m = re.match(r"\d+\.\s*(.+?)\s+第(.+?)\((主观题|选择题)\)", title)
        if m:
            suite = m.group(1).strip()
            qraw = m.group(2).strip()
            qtype = m.group(3).strip()
        else:
            suite = re.sub(r"^\d+\.\s*", "", title).split(" 第")[0].strip()
            qraw = ""
            qtype = "未识别"
        qno = normalize_question_no(qraw)
        prompt = extract_field(body, "设问")
        answer = extract_field(body, "答案落点")
        entries.append(
            Entry(
                placement_id=f"P{len(entries)+1:04d}",
                family=strip_cn_number(family),
                node=node,
                title=title,
                suite=suite,
                question_no=qno,
                question_type=qtype,
                prompt_from_candidate=prompt,
                answer_landing_from_candidate=answer,
            )
        )
        pending_title = None
        buf = []

    for line in lines:
        if line.startswith("# ") and not line.startswith("##"):
            flush()
            family = line[2:].strip()
        elif line.startswith("## "):
            flush()
            node = line[3:].strip()
        elif line.startswith("### "):
            flush()
            pending_title = line[4:].strip()
        else:
            if pending_title:
                buf.append(line)
    flush()
    return entries


def strip_cn_number(s: str) -> str:
    return re.sub(r"^[一二三四五六七八九十]+、", "", s).strip()


def normalize_question_no(qraw: str) -> str:
    q = qraw
    q = q.replace("题第", "")
    q = q.replace("题", "")
    q = q.replace("问", "")
    q = q.replace("（", "(").replace("）", ")")
    q = q.replace("第", "")
    q = q.strip()
    q = re.sub(r"\s+", "", q)
    return q


def extract_field(body: str, name: str) -> str:
    m = re.search(rf"【{name}】\s*(.*?)(?=\n\s*【|$)", body, flags=re.S)
    return clean_text(m.group(1)) if m else ""


def read_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open(encoding="utf-8-sig", errors="replace", newline="") as f:
        return list(csv.DictReader(f))


def key_suite_q(suite: str, qno: str) -> tuple[str, str]:
    return (suite.replace(" ", ""), normalize_question_no(qno).replace("前半", ""))


def discover_bundles() -> list[dict[str, object]]:
    bundles = []
    for root in SOURCE_ROOTS:
        if not root.exists():
            continue
        for d in root.rglob("*"):
            if not d.is_dir():
                continue
            paper_dir = d / "试卷"
            rubric_dir = d / "细则"
            if paper_dir.exists() or rubric_dir.exists():
                files = list(d.rglob("*"))
                source_files = [
                    p
                    for p in files
                    if p.is_file()
                    and p.suffix.lower() in {".pdf", ".docx", ".doc", ".pptx", ".rtf", ".txt"}
                    and not p.name.startswith("~$")
                ]
                if source_files:
                    bundles.append(
                {
                    "dir": d,
                    "name": d.name,
                    "path": str(d),
                    "paper_files": [p for p in source_files if "试卷" in p.relative_to(d).parts[:-1]],
                    "rubric_files": [p for p in source_files if "细则" in p.relative_to(d).parts[:-1] or "答案" in p.name or "评标" in p.name],
                            "all_files": source_files,
                        }
                    )
    return bundles


def suite_parts(suite: str) -> tuple[str, str, str]:
    m = re.search(r"(202[4-6])", suite)
    year = m.group(1) if m else ""
    district = ""
    for d in ["东城", "西城", "朝阳", "海淀", "丰台", "石景山", "顺义", "通州", "门头沟", "房山", "昌平", "延庆"]:
        if d in suite:
            district = d
            break
    stage = ""
    for st in ["期中", "期末", "一模", "二模"]:
        if st in suite:
            stage = st
            break
    return year, district, stage


def score_bundle(suite: str, bundle: dict[str, object]) -> int:
    year, district, stage = suite_parts(suite)
    hay = str(bundle["dir"])
    name = str(bundle["name"])
    score = 0
    if year and year in hay:
        score += 3
    if district and district in hay:
        score += 8
    if stage and stage in hay:
        score += 8
    if suite.replace(".11", "").replace("2024", "") in hay:
        score += 2
    if "已放弃" in hay:
        score -= 8
    if "分类" in hay:
        score -= 10
    if district and stage and district in name and stage in name:
        score += 5
    return score


def find_bundle(suite: str, bundles: list[dict[str, object]]) -> dict[str, object] | None:
    ranked = sorted(((score_bundle(suite, b), b) for b in bundles), key=lambda x: x[0], reverse=True)
    if ranked and ranked[0][0] >= 12:
        return ranked[0][1]
    return None


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            doc = fitz.open(path)
            pages = []
            for i, page in enumerate(doc, 1):
                pages.append(f"\n--- page {i} ---\n" + page.get_text())
            return clean_text("\n".join(pages))
        if suffix == ".docx":
            doc = Document(path)
            parts: list[str] = []
            parts.extend(p.text for p in doc.paragraphs if p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    cells = [clean_text(c.text) for c in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            return clean_text("\n".join(parts))
        if suffix in {".doc", ".rtf"}:
            return clean_text(run_textutil(path))
        if suffix == ".pptx":
            return clean_text(extract_pptx_text(path))
        if suffix == ".txt":
            return clean_text(path.read_text(encoding="utf-8", errors="replace"))
    except Exception as exc:
        return f"[EXTRACT_ERROR {type(exc).__name__}: {exc}]"
    return ""


def run_textutil(path: Path) -> str:
    cp = subprocess.run(["textutil", "-convert", "txt", "-stdout", str(path)], capture_output=True, text=True, timeout=60)
    if cp.returncode != 0:
        return cp.stderr
    return cp.stdout


def extract_pptx_text(path: Path) -> str:
    ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    out = []
    with zipfile.ZipFile(path) as z:
        names = sorted(n for n in z.namelist() if n.startswith("ppt/slides/slide") and n.endswith(".xml"))
        notes = sorted(n for n in z.namelist() if n.startswith("ppt/notesSlides/notesSlide") and n.endswith(".xml"))
        for name in names + notes:
            root = ET.fromstring(z.read(name))
            texts = [el.text for el in root.findall(".//a:t", ns) if el.text]
            if texts:
                out.append(f"\n--- {name} ---\n" + "\n".join(texts))
    return "\n".join(out)


def cache_text(path: Path, cache_dir: Path) -> tuple[Path, str]:
    h = hashlib.sha1(str(path).encode("utf-8")).hexdigest()[:12]
    cache_path = cache_dir / f"{h}_{safe_name(path.name)}.txt"
    if cache_path.exists():
        return cache_path, clean_text(cache_path.read_text(encoding="utf-8", errors="replace"))
    text = extract_text(path)
    cache_path.write_text(text, encoding="utf-8")
    return cache_path, text


def question_base(qno: str) -> str:
    m = re.match(r"(\d+)", qno)
    return m.group(1) if m else qno


def next_number(qbase: str) -> str:
    if qbase.isdigit():
        return str(int(qbase) + 1)
    return ""


def fullwidth_digits(s: str) -> str:
    return s.translate(str.maketrans("0123456789", "０１２３４５６７８９"))


def find_block(text: str, qno: str, kind: str) -> tuple[str, str]:
    if not text:
        return "", "empty_text"
    qbase = question_base(qno)
    nxt = next_number(qbase)
    candidates = []
    patterns = []
    qmark = r"(?:[\.．、:：\)\）\(\（􀆰]|$)"
    if kind == "rubric":
        q_variants = [qno, qno.replace("(", "（").replace(")", "）"), qbase, fullwidth_digits(qbase)]
        for q in q_variants:
            patterns.append(rf"(?m)(^|\n)\s*{re.escape(q)}(?!\d)\s*{qmark}")
    else:
        for q in [qbase, fullwidth_digits(qbase)]:
            patterns.extend(
                [
                    rf"(?m)(^|\n)\s*{re.escape(q)}(?![0-9０-９])\s*{qmark}",
                    rf"(?m)(^|\n)\s*第\s*{re.escape(q)}\s*题",
                ]
            )
    for pat in patterns:
        for m in re.finditer(pat, text):
            if kind == "paper" and re.match(r"\s*\d+(?:\.\d+)?\s*[%％]", text[m.end() : m.end() + 16]):
                continue
            start = m.start()
            end = len(text)
            if nxt:
                next_pats = [
                    rf"(?m)(^|\n)\s*{re.escape(nxt)}(?![0-9０-９])\s*{qmark}",
                    rf"(?m)(^|\n)\s*{re.escape(fullwidth_digits(nxt))}(?![0-9０-９])\s*{qmark}",
                ]
                nm = None
                for npat in next_pats:
                    nm = re.search(npat, text[m.end() :])
                    if nm:
                        break
                if nm:
                    end = m.end() + nm.start()
            snippet = clean_text(text[start:end])
            if len(snippet) > 80:
                candidates.append(snippet[:6000])
    if candidates:
        candidates.sort(key=lambda s: (score_block(s, qno, kind), min(len(s), 6000)), reverse=True)
        return candidates[0], "direct_question_number"
    # fallback: search around candidate prompt keywords.
    return "", "question_number_not_found"


def contains_qmark_score(s: str, qno: str) -> int:
    score = 0
    for token in ["逻辑", "推理", "判断", "三段论", "归纳", "类比", "答案", "参考答案", "评分"]:
        if token in s:
            score += 1
    if "(" in qno and qno in s:
        score += 4
    return score


def looks_like_answer_block(block: str) -> bool:
    head = block[:500]
    return ("参考答案" in head or "评分标准" in head or "答案" in head[:100]) and not any(
        token in head for token in ["材料一", "材料二", "结合材料", "下列", "问题:"]
    )


def score_block(block: str, qno: str, kind: str) -> int:
    score = contains_qmark_score(block, qno)
    head = block[:600]
    if kind == "paper":
        for token in [
            "材料一",
            "材料二",
            "结合材料",
            "下列",
            "问题:",
            "阅读材料",
            "运用",
            "这说明",
            "对此",
            "据此",
            "下列说法正确",
            "认识正确",
            "分析正确",
            "判断与推理正确",
            "下列表述",
        ]:
            if token in block:
                score += 3
        if all(re.search(rf"(^|\n)\s*{opt}[\.．、􀆰]?", block) for opt in ["A", "B", "C", "D"]) or all(
            re.search(rf"(^|\n)\s*{opt}[\.．、􀆰]?", block) for opt in ["Ａ", "Ｂ", "Ｃ", "Ｄ"]
        ):
            score += 6
        if "参考答案" in head or "评分标准" in head:
            score -= 8
    else:
        for token in ["参考答案", "评分", "细则", "答案", "给分"]:
            if token in block:
                score += 2
    return score


def ocr_pdf_text(path: Path, cache_dir: Path) -> str:
    ocr_dir = cache_dir.parent / "ocr_cache" / hashlib.sha1(str(path).encode("utf-8")).hexdigest()[:12]
    ocr_dir.mkdir(parents=True, exist_ok=True)
    merged = ocr_dir / "merged.ocr.txt"
    if merged.exists():
        return merged.read_text(encoding="utf-8", errors="replace")
    try:
        cp = subprocess.run(
            ["/Users/wanglifei/.local/bin/ocr-vision", "--out", str(ocr_dir), str(path)],
            capture_output=True,
            text=True,
            timeout=180,
        )
        texts = []
        for f in sorted(ocr_dir.glob("*.ocr.txt")):
            texts.append(f.read_text(encoding="utf-8", errors="replace"))
        text = clean_text("\n".join(texts))
        if not text:
            text = clean_text(cp.stdout + "\n" + cp.stderr)
        merged.write_text(text, encoding="utf-8")
        return text
    except Exception as exc:
        merged.write_text(f"[OCR_ERROR {type(exc).__name__}: {exc}]", encoding="utf-8")
        return ""


def best_excerpt(files: Iterable[Path], qno: str, kind: str, cache_dir: Path) -> tuple[str, str, str]:
    best = ("", "", "")
    best_score = -999
    for f in files:
        cp, txt = cache_text(f, cache_dir)
        block, status = find_block(txt, qno, kind)
        if kind == "rubric" and not block and file_name_matches_qno(f, qno):
            block = txt[:6000]
            status = "rubric_file_name_matched_full_file"
        if kind == "rubric" and not block:
            sub = find_subquestion_part_excerpt(txt, qno)
            if sub:
                block = sub
                status = "rubric_subquestion_part_matched"
        if kind == "rubric" and not block:
            answer_key = find_answer_key_excerpt(txt, qno)
            if answer_key:
                block = answer_key
                status = "answer_key_table_from_original"
        if kind == "rubric" and f.suffix.lower() == ".pdf" and not block:
            ocr_text = ocr_pdf_text(f, cache_dir)
            if ocr_text:
                ocr_block, ocr_status = find_block(ocr_text, qno, kind)
                if not ocr_block:
                    ocr_block = find_subquestion_part_excerpt(ocr_text, qno)
                    ocr_status = "subquestion_part_matched"
                if not ocr_block:
                    ocr_block = find_answer_key_excerpt(ocr_text, qno)
                    ocr_status = "answer_key_table_from_original"
                if ocr_block:
                    block = ocr_block
                    status = "ocr_vision_" + ocr_status
        if kind == "paper" and f.suffix.lower() == ".pdf" and (not block or looks_like_answer_block(block) or score_block(block, qno, kind) < 4):
            ocr_text = ocr_pdf_text(f, cache_dir)
            if ocr_text:
                ocr_block, ocr_status = find_block(ocr_text, qno, kind)
                if ocr_block and (not block or score_block(ocr_block, qno, kind) >= score_block(block, qno, kind)):
                    block = ocr_block
                    status = "ocr_vision_" + ocr_status
        block_score = score_block(block, qno, kind) if block else -999
        if kind == "paper":
            path_text = str(f)
            if any(token in path_text for token in ["补充材料", "答案", "评分", "细则", "评标", "讲评"]):
                block_score -= 50
            if "试卷" in f.name:
                block_score += 10
        if status == "rubric_file_name_matched_full_file":
            block_score += 50
        elif status == "rubric_subquestion_part_matched":
            block_score += rubric_logic_score(block)
        elif status.startswith("ocr_vision"):
            block_score += 15
        if block and (block_score > best_score or (block_score == best_score and len(block) > len(best[0]))):
            best = (block, str(f), status)
            best_score = block_score
        elif not best[1]:
            best = ("", str(f), status)
    return best


def best_teacher_answer_excerpt(files: Iterable[Path], qno: str, cache_dir: Path) -> tuple[str, str, str]:
    best = ("", "", "")
    best_score = -999
    for f in files:
        cp, txt = cache_text(f, cache_dir)
        block = find_teacher_answer_excerpt(txt, qno)
        if not block and f.suffix.lower() == ".pdf":
            ocr_text = ocr_pdf_text(f, cache_dir)
            block = find_teacher_answer_excerpt(ocr_text, qno)
        if block:
            score = score_block(block, qno, "rubric") + 30
            if "题号" in block and "答案" in block:
                score += 100
            if re.search(rf"(?m)^\s*{re.escape(question_base(qno))}\s*\n\s*[A-D]\s*$", block):
                score += 80
            if score > best_score:
                best = (block, str(f), "teacher_version_answer_section_from_original")
                best_score = score
    return best


def find_teacher_answer_excerpt(text: str, qno: str) -> str:
    if not text:
        return ""
    qbase = question_base(qno)
    table = find_choice_answer_table(text, qbase)
    if table:
        return table
    anchors = [
        text.find("答案和评分参考"),
        text.find("参考答案"),
        text.find("评分标准"),
        text.find("第一部分共 15 题"),
        text.find("第一部分共15题"),
    ]
    anchors = [a for a in anchors if a != -1]
    if not anchors:
        return ""
    answer_text = text[min(anchors) :]
    table = find_choice_answer_table(answer_text, qbase)
    if table:
        return table
    choice_pat = rf"(?<!\d){re.escape(qbase)}\s*[\.．、]?\s*[A-D](?![A-Za-z])"
    cm = re.search(choice_pat, answer_text)
    if cm:
        start = max(0, cm.start() - 250)
        end = min(len(answer_text), cm.end() + 900)
        return clean_text(answer_text[start:end])
    patterns = [
        rf"(?m)(^|\n)\s*{re.escape(qbase)}\s*[\.．、]\s*[（(]?\s*{extract_part(qno) or ''}",
        rf"(?m)(^|\n)\s*{re.escape(qbase)}(?!\d)\s*[\.．、]?",
    ]
    for pat in patterns:
        for m in re.finditer(pat, answer_text):
            start = m.start()
            nxt = next_number(qbase)
            end = len(answer_text)
            if nxt:
                nm = re.search(rf"(?m)(^|\n)\s*{re.escape(nxt)}(?!\d)\s*[\.．、]", answer_text[m.end() :])
                if nm:
                    end = m.end() + nm.start()
            block = clean_text(answer_text[start:end])
            if len(block) > 40 and ("答案" in block[:400] or "评分" in block[:400] or "分" in block[:300] or len(block) > 100):
                return block[:6000]
    return ""


def find_choice_answer_table(text: str, qbase: str) -> str:
    lines = text.splitlines()
    compact_lines = [clean_text(line) for line in lines]
    for i, line in enumerate(lines):
        if "题号" not in line or "|" not in line:
            continue
        for j in range(i + 1, min(len(lines), i + 5)):
            if "答案" not in lines[j] or "|" not in lines[j]:
                continue
            heads = [c.strip() for c in line.split("|")]
            vals = [c.strip() for c in lines[j].split("|")]
            if heads and heads[0] == "题号":
                heads = heads[1:]
            if vals and vals[0] == "答案":
                vals = vals[1:]
            if qbase in heads:
                idx = heads.index(qbase)
                if idx < len(vals) and re.fullmatch(r"[A-D]", vals[idx]):
                    start = max(0, i - 2)
                    end = min(len(lines), j + 3)
                    return clean_text("\n".join(lines[start:end]))
    for i, cell in enumerate(compact_lines):
        if cell != "题号":
            continue
        j = i + 1
        nums = []
        while j < len(compact_lines):
            cur = compact_lines[j]
            if not cur:
                j += 1
                continue
            if cur == "答案":
                break
            if re.fullmatch(r"\d{1,2}", cur):
                nums.append(cur)
                j += 1
                continue
            break
        if j >= len(compact_lines) or compact_lines[j] != "答案" or len(nums) < 3:
            continue
        vals = []
        k = j + 1
        while k < len(compact_lines) and len(vals) < len(nums):
            cur = compact_lines[k]
            if not cur:
                k += 1
                continue
            if re.fullmatch(r"[A-D]", cur):
                vals.append(cur)
                k += 1
                continue
            break
        if len(vals) >= len(nums) and qbase in nums:
            start = max(0, i - 3)
            end = min(len(lines), k + 2)
            table = "题号 | " + " | ".join(nums) + "\n答案 | " + " | ".join(vals[: len(nums)])
            return clean_text(table + "\n原始表附近：\n" + "\n".join(lines[start:end]))
    for i in range(len(lines)):
        nums = []
        j = i
        expected = None
        while j < len(lines):
            cell = clean_text(lines[j])
            if not re.fullmatch(r"\d{1,2}", cell):
                break
            num = int(cell)
            if expected is not None and num != expected:
                break
            nums.append(str(num))
            expected = num + 1
            j += 1
        if len(nums) < 3:
            continue
        vals = []
        k = j
        while k < len(lines) and len(vals) < len(nums):
            cell = clean_text(lines[k])
            if not cell:
                k += 1
                continue
            if not re.fullmatch(r"[A-D]", cell):
                break
            vals.append(cell)
            k += 1
        if len(vals) >= len(nums) and qbase in nums:
            start = max(0, i - 3)
            end = min(len(lines), k + 2)
            table = "题号 | " + " | ".join(nums) + "\n答案 | " + " | ".join(vals[: len(nums)])
            return clean_text(table + "\n原始表附近：\n" + "\n".join(lines[start:end]))
    for i, cell in enumerate(compact_lines):
        if cell not in {"题号", "参考答案"} and not ("选择题" in cell and "答案" in cell):
            continue
        zone_end = min(len(compact_lines), i + 90)
        for j in range(i + 1, zone_end):
            if compact_lines[j] != qbase:
                continue
            k = j + 1
            while k < zone_end and not compact_lines[k]:
                k += 1
            if k < zone_end and re.fullmatch(r"[A-D]", compact_lines[k]):
                start = max(0, i - 2)
                end = min(len(lines), k + 4)
                table = f"题号 | {qbase}\n答案 | {compact_lines[k]}"
                return clean_text(table + "\n原始表附近：\n" + "\n".join(lines[start:end]))
    return ""


def extract_part(qno: str) -> str:
    m = re.search(r"\((\d+)\)", normalize_question_no(qno))
    return m.group(1) if m else ""


def file_name_matches_qno(path: Path, qno: str) -> bool:
    name = clean_text(path.stem)
    compact = re.sub(r"\s+", "", name)
    q = normalize_question_no(qno)
    q_compact = q.replace("(", "").replace(")", "")
    if "(" in q:
        candidates = {
            q,
            q.replace("(", ".("),
            q.replace("(", "-").replace(")", ""),
            q_compact,
        }
    else:
        candidates = {q, question_base(q), q_compact}
    return any(c and c in compact for c in candidates)


def find_answer_key_excerpt(text: str, qno: str) -> str:
    if not text:
        return ""
    qbase = question_base(qno)
    table = find_choice_answer_table(text, qbase)
    if table:
        return table
    # Common answer-key tables are compact; preserve the surrounding table rather
    # than inventing a reason that the source does not provide.
    table_hits = []
    for token in ["题号", "答案", "参考答案", "第一部分"]:
        idx = text.find(token)
        if idx != -1:
            table_hits.append(idx)
    if table_hits and re.search(rf"(?<!\d){re.escape(qbase)}(?!\d)", text[min(table_hits): min(table_hits) + 4000]):
        start = max(0, min(table_hits) - 200)
        return clean_text(text[start : start + 4000])
    m = re.search(rf"(?m)(^|\n)\s*{re.escape(qbase)}(?!\d)\s*[\.．、]?\s*[A-D]", text)
    if m:
        return clean_text(text[max(0, m.start() - 500) : m.start() + 1500])
    return ""


def find_subquestion_part_excerpt(text: str, qno: str) -> str:
    m = re.search(r"\((\d+)\)", normalize_question_no(qno))
    if not m:
        return ""
    part = m.group(1)
    pat = rf"(?m)(^|\n)\s*[（(]?\s*{part}(?!\d)\s*[）)]?\s*(?:[（(]\s*\d+\s*分\s*[）)]?)?"
    candidates = []
    for hit in re.finditer(pat, text):
        window = text[hit.start() : hit.start() + 2500]
        if any(token in window for token in ["推理", "判断", "三段论", "必要条件", "充分条件", "参考答案", "得分", "分）"]):
            start = text.rfind("--- ppt/slides/", 0, hit.start())
            if start == -1:
                start = max(0, hit.start() - 300)
            end = text.find("--- ppt/slides/", hit.start() + 50)
            if end == -1:
                end = hit.start() + 3500
            candidates.append(clean_text(text[start:end]))
    if not candidates:
        return ""
    candidates.sort(key=lambda s: score_block(s, qno, "rubric") + rubric_logic_score(s), reverse=True)
    return candidates[0][:6000]


def rubric_logic_score(text: str) -> int:
    score = 0
    for token in ["推理属于", "必要条件假言推理", "充分条件假言推理", "三段论推理", "不完全归纳推理", "类比推理", "大项不当扩大", "小项不当扩大", "中项不周延", "换位推理", "选言推理", "排中律", "矛盾律"]:
        if token in text:
            score += 10
    if "逻辑与思维" in text:
        score += 8
    if "哲学" in text[:300] and score == 0:
        score -= 15
    return score


def verified_old_choice_excerpt_from_original(paper_files: Iterable[Path], old_choice_excerpt: str, cache_dir: Path) -> tuple[str, str, str]:
    if not old_choice_excerpt:
        return "", "", ""
    cleaned = re.sub(r"^\s*\d+\s*[\.．、]\s*", "", clean_text(old_choice_excerpt))
    probe = cleaned[:60]
    if len(probe) < 12:
        return "", "", ""
    compact_probe = re.sub(r"\s+", "", probe)
    for f in paper_files:
        _, txt = cache_text(f, cache_dir)
        compact_txt = re.sub(r"\s+", "", txt)
        idx = compact_txt.find(compact_probe)
        if idx != -1:
            return old_choice_excerpt, str(f), "verified_old_choice_excerpt_in_original_paper"
    return "", "", ""


def build_old_audit_index() -> dict[tuple[str, str], list[dict[str, str]]]:
    idx: dict[tuple[str, str], list[dict[str, str]]] = defaultdict(list)
    for r in read_csv(OLD_AUDIT):
        suite = r.get("cov_suite_name") or re.sub(r"\s+Q.*", "", r.get("label", ""))
        qno = r.get("cov_question_no") or r.get("question_id") or ""
        if suite and qno:
            idx[key_suite_q(suite, qno)].append(r)
    return idx


def build_coverage_index() -> dict[tuple[str, str], list[dict[str, str]]]:
    idx: dict[tuple[str, str], list[dict[str, str]]] = defaultdict(list)
    for r in read_csv(COVERAGE):
        idx[key_suite_q(r.get("suite_name", ""), r.get("question_no", ""))].append(r)
    return idx


def summarize_old_rows(rows: list[dict[str, str]]) -> dict[str, str]:
    if not rows:
        return {}
    first = rows[0]
    choice_excerpt = next((r.get("choice_excerpt", "") for r in rows if r.get("choice_excerpt", "")), "")
    correct_answer = next((r.get("correct_answer", "") for r in rows if r.get("correct_answer", "")), "")
    correct_reason = next((r.get("correct_reason", "") for r in rows if r.get("correct_reason", "")), "")
    wrong_traps = next((r.get("wrong_option_traps", "") for r in rows if r.get("wrong_option_traps", "")), "")
    return {
        "old_evidence_level": first.get("evidence_level", ""),
        "old_source_packet": first.get("source_packet", "") or first.get("cov_source_packet", ""),
        "old_choice_excerpt": choice_excerpt,
        "old_correct_answer": correct_answer,
        "old_correct_reason": correct_reason,
        "old_wrong_option_traps": wrong_traps,
        "old_answer_sentences": "\n".join(f"- {r.get('family','')}: {r.get('answer_sentence','')}" for r in rows if r.get("answer_sentence")),
    }


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    cache_dir = OUT / "source_text_cache"
    cache_dir.mkdir(exist_ok=True)
    entries = parse_framework()
    old_idx = build_old_audit_index()
    cov_idx = build_coverage_index()
    bundles = discover_bundles()
    bundle_map = {e.suite: find_bundle(e.suite, bundles) for e in entries}

    with (OUT / "SOURCE_BUNDLE_MAP.csv").open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["suite", "bundle_dir", "paper_files", "rubric_files", "status"])
        for suite in sorted(set(e.suite for e in entries)):
            b = bundle_map.get(suite)
            if b:
                w.writerow(
                    [
                        suite,
                        b["path"],
                        "\n".join(str(p) for p in b["paper_files"]),
                        "\n".join(str(p) for p in b["rubric_files"]),
                        "matched",
                    ]
                )
            else:
                w.writerow([suite, "", "", "", "missing_bundle"])

    rows = []
    paper_cache_by_key: dict[tuple[str, str], tuple[str, str, str]] = {}
    rubric_cache_by_key: dict[tuple[str, str], tuple[str, str, str]] = {}
    for e in entries:
        k = key_suite_q(e.suite, e.question_no)
        old = summarize_old_rows(old_idx.get(k, []))
        cov_rows = cov_idx.get(k, [])
        cov = cov_rows[0] if cov_rows else {}
        b = bundle_map.get(e.suite)
        paper_excerpt = paper_path = paper_status = ""
        rubric_excerpt = rubric_path = rubric_status = ""
        if b:
            if k not in paper_cache_by_key:
                paper_files = b["paper_files"] or [p for p in b["all_files"] if "试卷" in p.name or p.suffix.lower() in {".pdf", ".docx", ".doc"}]
                paper_cache_by_key[k] = best_excerpt(paper_files, e.question_no, "paper", cache_dir)
            paper_excerpt, paper_path, paper_status = paper_cache_by_key[k]
            if k not in rubric_cache_by_key:
                rubric_files = b["rubric_files"] or b["all_files"]
                rubric_cache_by_key[k] = best_excerpt(rubric_files, e.question_no, "rubric", cache_dir)
            rubric_excerpt, rubric_path, rubric_status = rubric_cache_by_key[k]
            teacher_answer_files = list(dict.fromkeys([*b["rubric_files"], *b["paper_files"]]))
            if teacher_answer_files:
                teacher_excerpt, teacher_path, teacher_status = best_teacher_answer_excerpt(teacher_answer_files, e.question_no, cache_dir)
                if teacher_excerpt and (
                    not rubric_excerpt
                    or "recovered_from_old" in rubric_status
                    or rubric_status.startswith("ocr_vision")
                    or e.question_type == "选择题"
                ):
                    rubric_excerpt, rubric_path, rubric_status = teacher_excerpt, teacher_path, teacher_status

        if not paper_excerpt and old.get("old_choice_excerpt"):
            verified_excerpt = verified_old_choice_excerpt_from_original((b or {}).get("paper_files", []), old["old_choice_excerpt"], cache_dir) if b else ("", "", "")
            if verified_excerpt[0]:
                paper_excerpt, paper_path, paper_status = verified_excerpt
            else:
                paper_excerpt = old["old_choice_excerpt"]
                paper_status = "recovered_from_old_audit_choice_excerpt_needs_original_review"
        if not rubric_excerpt and old.get("old_answer_sentences"):
            rubric_excerpt = old["old_answer_sentences"]
            rubric_status = "recovered_from_old_audit_answer_needs_original_review"
        review_status = "source_extracted"
        if "recovered_from_old" in paper_status or "recovered_from_old" in rubric_status:
            review_status = "needs_manual_review"
        if not paper_excerpt or not rubric_excerpt:
            review_status = "needs_manual_review"

        row = {
            **asdict(e),
            "coverage_question_id": cov.get("question_id", ""),
            "coverage_evidence_level": cov.get("evidence_level", "") or old.get("old_evidence_level", ""),
            "coverage_decision_reason": cov.get("decision_reason", ""),
            "paper_path": paper_path,
            "paper_status": paper_status,
            "paper_excerpt": paper_excerpt,
            "rubric_path": rubric_path,
            "rubric_status": rubric_status,
            "rubric_excerpt": rubric_excerpt,
            "old_correct_answer": old.get("old_correct_answer", ""),
            "old_correct_reason": old.get("old_correct_reason", ""),
            "old_wrong_option_traps": old.get("old_wrong_option_traps", ""),
            "review_status": review_status,
        }
        rows.append(row)

    write_framework(entries)
    write_ledger(rows)
    write_markdown(rows)
    write_progress_summary(rows)


def write_framework(entries: list[Entry]) -> None:
    by_family: dict[str, dict[str, list[Entry]]] = defaultdict(lambda: defaultdict(list))
    for e in entries:
        by_family[e.family][e.node].append(e)
    parts = ["# 推理框架抽取稿", "", f"- source: `{FRAMEWORK_MD}`", f"- placements: {len(entries)}", ""]
    for family, nodes in by_family.items():
        parts.append(f"## {family}")
        for node, es in nodes.items():
            parts.append(f"### {node}")
            for e in es:
                parts.append(f"- {e.title} -> {e.suite} {e.question_no} {e.question_type}")
            parts.append("")
    (OUT / "REASONING_FRAMEWORK_EXTRACTED_20260531.md").write_text("\n".join(parts), encoding="utf-8")


def write_ledger(rows: list[dict[str, str]]) -> None:
    fields = list(rows[0].keys()) if rows else []
    with (OUT / "REASONING_EXERCISE_COMPILATION_LEDGER.csv").open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fields)
        w.writeheader()
        w.writerows(rows)


def write_markdown(rows: list[dict[str, str]]) -> None:
    by_family: dict[str, dict[str, list[dict[str, str]]]] = defaultdict(lambda: defaultdict(list))
    for r in rows:
        by_family[r["family"]][r["node"]].append(r)
    parts = [
        "# 选必三《逻辑与思维》推理习题汇编（回源提取稿）",
        "",
        "本稿只做习题汇编：按推理框架放置原题与细则，不做教学扩写。旧推理稿与审计表只作定位索引；凡未能直接从桌面原卷/细则提取的地方，均在状态栏标出。",
        "",
    ]
    for family, nodes in by_family.items():
        parts.append(f"# {family}")
        parts.append("")
        for node, es in nodes.items():
            parts.append(f"## {node}")
            parts.append("")
            for r in es:
                parts.append(f"### {r['title']}")
                parts.append("")
                parts.append(f"- 证据状态：{r['review_status']}；{r['coverage_evidence_level']}")
                parts.append(f"- 原卷：`{r['paper_path']}`")
                parts.append(f"- 细则/答案：`{r['rubric_path']}`")
                if r.get("coverage_decision_reason"):
                    parts.append(f"- 定位说明：{r['coverage_decision_reason']}")
                parts.append("")
                parts.append("【试题原文/原卷摘录】")
                parts.append("")
                parts.append(blockquote_or_fence(r.get("paper_excerpt", "")))
                parts.append("")
                parts.append("【细则/答案原文摘录】")
                parts.append("")
                parts.append(blockquote_or_fence(r.get("rubric_excerpt", "")))
                if r.get("old_correct_answer") or r.get("old_correct_reason") or r.get("old_wrong_option_traps"):
                    parts.append("")
                    parts.append("【旧审计辅助信息（待以原源为准）】")
                    if r.get("old_correct_answer"):
                        parts.append(f"- 答案：{r['old_correct_answer']}")
                    if r.get("old_correct_reason"):
                        parts.append(f"- 正确理由：{r['old_correct_reason']}")
                    if r.get("old_wrong_option_traps"):
                        parts.append(f"- 诱人错项/错因：{r['old_wrong_option_traps']}")
                parts.append("")
    (OUT / "选必三_逻辑与思维_推理习题汇编_回源提取稿.md").write_text("\n".join(parts), encoding="utf-8")


def blockquote_or_fence(s: str) -> str:
    s = clean_text(s)
    if not s:
        return "`MISSING_NEEDS_MANUAL_REVIEW`"
    return "```text\n" + s + "\n```"


def write_progress_summary(rows: list[dict[str, str]]) -> None:
    total = len(rows)
    source = sum(1 for r in rows if r["review_status"] == "source_extracted")
    needs = total - source
    missing_paper = [r for r in rows if not r.get("paper_excerpt")]
    missing_rubric = [r for r in rows if not r.get("rubric_excerpt")]
    summary = [
        "# 推理习题汇编生成摘要",
        "",
        f"- placements: {total}",
        f"- source_extracted: {source}",
        f"- needs_manual_review: {needs}",
        f"- missing_paper_excerpt: {len(missing_paper)}",
        f"- missing_rubric_excerpt: {len(missing_rubric)}",
        "",
        "## 需要人工复核的放置点",
        "",
    ]
    for r in rows:
        if r["review_status"] != "source_extracted":
            summary.append(f"- {r['placement_id']} {r['title']} | paper={r['paper_status']} | rubric={r['rubric_status']}")
    (OUT / "REASONING_EXERCISE_COMPILATION_SUMMARY.md").write_text("\n".join(summary), encoding="utf-8")


if __name__ == "__main__":
    main()
