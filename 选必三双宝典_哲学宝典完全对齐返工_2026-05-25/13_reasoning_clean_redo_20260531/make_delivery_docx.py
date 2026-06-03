#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


RUN = Path(__file__).resolve().parent
MD = RUN / "选必三_逻辑与思维_推理习题汇编_干净重做版_20260531.md"
DOCX = RUN / "delivery" / "选必三_逻辑与思维_推理习题汇编_干净重做版_20260531.docx"


def set_east_asia_font(run, font: str) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def set_doc_defaults(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "PingFang SC"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "PingFang SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    if "SourceBlock" not in [s.name for s in styles]:
        source = styles.add_style("SourceBlock", 1)
    else:
        source = styles["SourceBlock"]
    source.font.name = "Menlo"
    source._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    source.font.size = Pt(9)
    source.paragraph_format.space_after = Pt(0)
    source.paragraph_format.line_spacing = 1.05
    source.paragraph_format.left_indent = Pt(12)
    source.paragraph_format.right_indent = Pt(6)


def add_label_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    set_east_asia_font(run, "PingFang SC")


def main() -> None:
    DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_doc_defaults(doc)
    lines = MD.read_text(encoding="utf-8").splitlines()
    in_code = False
    for line in lines:
        if line.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph(style="SourceBlock")
            p.add_run(line)
            continue
        if not line.strip():
            continue
        if line.startswith("# "):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("【") and line.endswith("】"):
            add_label_paragraph(doc, line)
            continue
        if line.startswith("答案："):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            set_east_asia_font(run, "PingFang SC")
            continue
        p = doc.add_paragraph(line)
        for run in p.runs:
            set_east_asia_font(run, "PingFang SC")
    doc.save(DOCX)


if __name__ == "__main__":
    main()
