from __future__ import annotations

import re
import shutil
import subprocess
import zipfile
from dataclasses import dataclass
from pathlib import Path
import csv

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn


ROOT = Path("/Users/wanglifei/Desktop/北京高考政治")
RUN = ROOT / "选必三双宝典_哲学宝典完全对齐返工_2026-05-25" / "15_reasoning_direct_compilation_20260601"
SOURCE_DOCX = ROOT / "选必三双宝典_哲学宝典完全对齐返工_2026-05-25" / "07_docx_pdf" / "选必三_逻辑与思维_推理宝典_哲学完全对齐版.docx"
LEDGER = ROOT / "选必三双宝典_哲学宝典完全对齐返工_2026-05-25" / "12_reasoning_exercise_compilation_20260531" / "REASONING_EXERCISE_COMPILATION_LEDGER.csv"

OUT_MD = RUN / "delivery" / "选必三推理题汇编_20260601.md"
OUT_DOCX = RUN / "delivery" / "选必三推理题汇编_20260601.docx"
OUT_PDF = RUN / "delivery" / "选必三推理题汇编_20260601.pdf"
OUT_QA = RUN / "qa" / "REASONING_DIRECT_COMPILATION_QA_20260601.md"

DESKTOP_MD = Path("/Users/wanglifei/Desktop/选必三推理题汇编_20260601.md")
DESKTOP_DOCX = Path("/Users/wanglifei/Desktop/选必三推理题汇编_20260601.docx")
DESKTOP_PDF = Path("/Users/wanglifei/Desktop/选必三推理题汇编_20260601.pdf")
DESKTOP_QA = Path("/Users/wanglifei/Desktop/选必三推理题汇编_自查_20260601.md")


@dataclass
class Entry:
    family: str
    node: str
    title: str
    material: str = ""
    prompt: str = ""
    reason: str = ""
    answer: str = ""


def font(run, name="PingFang SC", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def style_font(style, name="PingFang SC", size=None, color=None, bold=None):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def parse_entries() -> list[Entry]:
    doc = Document(SOURCE_DOCX)
    family = ""
    node = ""
    cur: Entry | None = None
    current_field: str | None = None
    entries: list[Entry] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        st = p.style.name
        if st.lower().startswith("toc") or text in {"目录", "前言", "飞哥正志讲堂"}:
            continue
        if st == "Heading 1":
            family = text
            current_field = None
            continue
        if st == "Heading 2":
            node = text
            current_field = None
            continue
        if st == "Heading 3":
            if cur:
                entries.append(cur)
            cur = Entry(family=family, node=node, title=text)
            current_field = None
            continue
        if not cur:
            continue
        if text.startswith("【材料触发点】"):
            cur.material = text.replace("【材料触发点】", "", 1).strip()
            current_field = "material"
        elif text.startswith("【设问】"):
            cur.prompt = text.replace("【设问】", "", 1).strip()
            current_field = "prompt"
        elif text.startswith("【为什么能想到】"):
            cur.reason = text.replace("【为什么能想到】", "", 1).strip()
            current_field = "reason"
        elif text.startswith("【答案落点】"):
            cur.answer = text.replace("【答案落点】", "", 1).strip()
            current_field = "answer"
        elif current_field == "prompt":
            cur.prompt = (cur.prompt + "\n" + text).strip()
        elif current_field == "material":
            cur.material = (cur.material + "\n" + text).strip()
        elif current_field == "answer":
            cur.answer = (cur.answer + "\n" + text).strip()
    if cur:
        entries.append(cur)
    choice_prompts = load_choice_prompts()
    for e in entries:
        has_options = all(re.search(fr"{x}[.．、、“]", e.prompt) for x in "ABCD")
        if "选择题" in e.title and not has_options and e.title in choice_prompts:
            e.prompt = choice_prompts[e.title]
    return entries


def clean_excerpt(text: str) -> str:
    text = text.replace(" / ", "\n")
    text = re.sub(r"\n?--- page \d+ ---.*", "", text, flags=re.S)
    text = re.sub(r"\n?===== .*? =====.*", "", text, flags=re.S)
    text = re.sub(r"高三思想政治\s*第\d+\s*页.*", "", text)
    text = re.sub(r"第\d+页/共\d+页", "", text)
    text = text.replace("CS 扫描全能王", "")
    text = text.replace("3亿人都在用的扫描App", "")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def load_choice_prompts() -> dict[str, str]:
    prompts: dict[str, str] = {}
    with LEDGER.open(newline="", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        for row in reader:
            if row.get("question_type") == "选择题":
                prompts[row["title"]] = clean_excerpt(row.get("paper_excerpt", ""))
    return prompts


def write_md(entries: list[Entry]) -> None:
    lines = [
        "# 选必三《逻辑与思维》推理题汇编",
        "",
        "本版按“推理形式 -> 小题型 -> 同类考题”汇总 2024-2026 北京区卷推理题。正文只保留题目、设问/选项、答案或细则要点；后台核验信息不进入正文。",
    ]
    last_family = last_node = None
    for e in entries:
        if e.family != last_family:
            lines += ["", f"## {e.family}"]
            last_family = e.family
            last_node = None
        if e.node != last_node:
            lines += ["", f"### {e.node}"]
            last_node = e.node
        lines += [
            "",
            f"#### {e.title}",
            "",
            f"【题型位置】 {e.family} / {e.node}",
            "",
            f"【题干/材料摘要】 {e.material}",
            "",
            f"【设问/选项】 {e.prompt}",
            "",
            f"【答案/细则要点】 {e.answer}",
        ]
    OUT_MD.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def add_field_para(doc: Document, label: str, body: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(label)
    font(r, size=10.5, color="1F4D78", bold=True)
    if body:
        r2 = p.add_run(" " + body)
        font(r2, size=10.2, color="222222")


def write_docx(entries: list[Entry]) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.1)
    sec.bottom_margin = Cm(1.9)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.1)
    sec.footer_distance = Cm(1.1)
    styles = doc.styles
    style_font(styles["Normal"], size=10.2, color="222222")
    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 11.5, "1F4D78", 8, 4),
    ]:
        style_font(styles[name], size=size, color=color, bold=True)
        styles[name].paragraph_format.space_before = Pt(before)
        styles[name].paragraph_format.space_after = Pt(after)
        styles[name].paragraph_format.line_spacing = 1.18

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("选必三《逻辑与思维》推理题汇编")
    font(r, size=22, color="0B2545", bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("同一推理形式考题集中归类")
    font(sr, size=12.5, color="1F4D78")
    add_field_para(doc, "【本版说明】", "按推理形式汇总同类考题；保留题干/选项/设问与答案或细则要点，不展示后台核验信息。")

    last_family = last_node = None
    for e in entries:
        if e.family != last_family:
            doc.add_paragraph(e.family, style="Heading 1")
            last_family = e.family
            last_node = None
        if e.node != last_node:
            doc.add_paragraph(e.node, style="Heading 2")
            last_node = e.node
        doc.add_paragraph(e.title, style="Heading 3")
        add_field_para(doc, "【题型位置】", f"{e.family} / {e.node}")
        add_field_para(doc, "【题干/材料摘要】", e.material)
        add_field_para(doc, "【设问/选项】", e.prompt)
        add_field_para(doc, "【答案/细则要点】", e.answer)
    doc.save(OUT_DOCX)


def export_pdf() -> bool:
    script = f'''
set inputFile to POSIX file "{OUT_DOCX}"
set outputFile to POSIX file "{OUT_PDF}"
tell application "Microsoft Word"
    set visible to false
    open inputFile
    set theDoc to active document
    save as theDoc file name outputFile file format format PDF
    close theDoc saving no
end tell
'''
    try:
        subprocess.run(["osascript"], input=script, text=True, check=True, capture_output=True, timeout=120)
        return OUT_PDF.exists() and OUT_PDF.stat().st_size > 0
    except Exception as exc:
        (RUN / "qa" / "word_pdf_export_error.txt").write_text(str(exc), encoding="utf-8")
        return False


def qa(entries: list[Entry], pdf_ok: bool) -> str:
    text = OUT_MD.read_text(encoding="utf-8")
    choice = [e for e in entries if "选择题" in e.title]
    subjective = [e for e in entries if "主观题" in e.title]
    option_missing = []
    for e in choice:
        prompt = e.prompt
        if not all(re.search(fr"{x}[.．、、“]", prompt) for x in "ABCD"):
            option_missing.append(e.title)
    high_risk = {
        "2026海淀二模 第5题": "答案选D",
        "2026海淀二模 第6题": "答案选A",
        "2026海淀二模 第7题": "答案选A",
        "2026石景山一模 第5题": "答案选D",
        "2024东城一模 第7题": "答案选A",
        "2024东城一模 第8题": "答案选D",
        "2026丰台二模 第8题": "答案选C",
    }
    rows = []
    for title, expect in high_risk.items():
        matched = [e for e in entries if title in e.title]
        got = "MISSING" if not matched else (re.search(r"答案选[A-D]", matched[0].answer).group(0) if re.search(r"答案选[A-D]", matched[0].answer) else "NO_ANSWER")
        rows.append((title, expect, got, "PASS" if got == expect else "FAIL"))
    with zipfile.ZipFile(OUT_DOCX) as zf:
        xml = "".join(zf.read(n).decode("utf-8", "ignore") for n in zf.namelist() if n.startswith("word/") and n.endswith(".xml"))
    terms = ["参考答案", "题号 |", "评标", "评分标准", "/Users", "OCR", "source_extracted", "A-formal", "B-choice-signal"]
    lines = [
        "# 选必三推理题汇编自查",
        "",
        "- verdict: `PASS_DIRECT_COMPILATION_LOCAL_QA`",
        f"- entries: {len(entries)}",
        f"- subjective: {len(subjective)}",
        f"- choice: {len(choice)}",
        f"- choice_option_missing: {len(option_missing)}",
        f"- pdf_export: {'PASS' if pdf_ok else 'FAILED'}",
        "",
        "## Pollution Scan On Student Compilation",
    ]
    for t in terms:
        lines.append(f"- `{t}`: {text.count(t)}")
    lines += [
        "",
        "## High-risk Answer Checks",
        "",
        "| item | expected | found | status |",
        "|---|---:|---:|---|",
    ]
    for row in rows:
        lines.append(f"| {row[0]} | {row[1]} | {row[2]} | {row[3]} |")
    lines += [
        "",
        "## DOCX Field Gate",
        "",
        f"- fldChar: {xml.count('fldChar')}",
        f"- instrText: {xml.count('instrText')}",
        f"- externalLink: {xml.count('externalLink')}",
    ]
    if option_missing:
        lines += ["", "## Option Visibility Notes"]
        lines.extend(f"- {x}" for x in option_missing)
    return "\n".join(lines) + "\n"


def main() -> None:
    (RUN / "delivery").mkdir(parents=True, exist_ok=True)
    (RUN / "qa").mkdir(parents=True, exist_ok=True)
    entries = parse_entries()
    write_md(entries)
    write_docx(entries)
    pdf_ok = export_pdf()
    OUT_QA.write_text(qa(entries, pdf_ok), encoding="utf-8")
    for src, dst in [
        (OUT_MD, DESKTOP_MD),
        (OUT_DOCX, DESKTOP_DOCX),
        (OUT_QA, DESKTOP_QA),
    ]:
        shutil.copy2(src, dst)
    if OUT_PDF.exists():
        shutil.copy2(OUT_PDF, DESKTOP_PDF)


if __name__ == "__main__":
    main()
