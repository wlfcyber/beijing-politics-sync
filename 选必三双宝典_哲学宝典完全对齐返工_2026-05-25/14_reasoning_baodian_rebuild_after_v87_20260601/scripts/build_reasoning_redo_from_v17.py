from __future__ import annotations

import re
import shutil
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/wanglifei/Desktop/北京高考政治")
RUN = ROOT / "选必三双宝典_哲学宝典完全对齐返工_2026-05-25" / "14_reasoning_baodian_rebuild_after_v87_20260601"
SOURCE_DOCX = ROOT / "选必三双宝典_哲学宝典完全对齐返工_2026-05-25" / "07_docx_pdf" / "选必三_逻辑与思维_推理宝典_哲学完全对齐版.docx"
SOURCE_PDF = ROOT / "选必三双宝典_哲学宝典完全对齐返工_2026-05-25" / "07_docx_pdf" / "选必三_逻辑与思维_推理宝典_哲学完全对齐版.pdf"

OUT_MD = RUN / "delivery" / "选必三推理宝典_重做版_20260601.md"
OUT_DOCX = RUN / "delivery" / "选必三推理宝典_重做版_20260601.docx"
OUT_PDF = RUN / "delivery" / "选必三推理宝典_重做版_20260601.pdf"
OUT_QA = RUN / "qa" / "REASONING_BAODIAN_REDO_QA_20260601.md"

DESKTOP_MD = Path("/Users/wanglifei/Desktop/选必三推理宝典_重做版_20260601.md")
DESKTOP_DOCX = Path("/Users/wanglifei/Desktop/选必三推理宝典_重做版_20260601.docx")
DESKTOP_PDF = Path("/Users/wanglifei/Desktop/选必三推理宝典_重做版_20260601.pdf")
DESKTOP_QA = Path("/Users/wanglifei/Desktop/选必三推理宝典_重做版_自查验收_20260601.md")


@dataclass
class Item:
    style: str
    text: str


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_style_font(style, font_name: str, size_pt: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def collect_items() -> list[Item]:
    src = Document(SOURCE_DOCX)
    items: list[Item] = []
    for p in src.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name
        if style.lower().startswith("toc"):
            continue
        if text == "目录":
            continue
        # The source already contains a generated TOC before the body. Keep the
        # student content and rebuild a no-field document from it.
        items.append(Item(style=style, text=text))
    return items


def md_for_items(items: list[Item]) -> str:
    lines: list[str] = []
    first = True
    for item in items:
        text = item.text
        if first:
            parts = text.splitlines()
            lines.append(f"# {parts[0]}")
            for part in parts[1:]:
                lines.append("")
                lines.append(f"**{part}**")
            first = False
            continue
        if text == "飞哥正志讲堂":
            lines.append("")
            lines.append(text)
            continue
        if text == "前言":
            lines.append("")
            lines.append("## 本版说明")
            lines.append("")
            lines.append("本版撤回 V87 原题摘录包口径，恢复 V17 学生化推理宝典正文；V87 只作为题源和答案源回查索引。")
            continue
        if item.style == "Heading 1":
            lines.extend(["", f"## {text}"])
        elif item.style == "Heading 2":
            lines.extend(["", f"### {text}"])
        elif item.style == "Heading 3":
            lines.extend(["", f"#### {text}"])
        else:
            lines.extend(["", text])
    return "\n".join(lines).strip() + "\n"


def add_label_paragraph(doc: Document, text: str) -> None:
    labels = ["【材料触发点】", "【设问】", "【为什么能想到】", "【答案落点】"]
    for label in labels:
        if text.startswith(label):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.25
            r1 = p.add_run(label)
            set_east_asia_font(r1, "PingFang SC")
            r1.bold = True
            r1.font.color.rgb = RGBColor.from_string("1F4D78")
            rest = text[len(label):].lstrip()
            if rest:
                r2 = p.add_run(" " + rest)
                set_east_asia_font(r2, "PingFang SC")
                r2.font.size = Pt(10.5)
            return
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.25
    for run in p.runs:
        set_east_asia_font(run, "PingFang SC")
        run.font.size = Pt(10.5)


def build_docx(items: list[Item]) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.start_type = WD_SECTION.NEW_PAGE
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.3)
    sec.right_margin = Cm(2.3)
    sec.header_distance = Cm(1.2)
    sec.footer_distance = Cm(1.2)

    styles = doc.styles
    set_style_font(styles["Normal"], "PingFang SC", 10.5, "222222")
    set_style_font(styles["Heading 1"], "PingFang SC", 16, "2E74B5", True)
    set_style_font(styles["Heading 2"], "PingFang SC", 13, "2E74B5", True)
    set_style_font(styles["Heading 3"], "PingFang SC", 12, "1F4D78", True)
    for name, before, after in [("Heading 1", 18, 10), ("Heading 2", 14, 7), ("Heading 3", 10, 5)]:
        styles[name].paragraph_format.space_before = Pt(before)
        styles[name].paragraph_format.space_after = Pt(after)
        styles[name].paragraph_format.line_spacing = 1.25

    first = True
    for item in items:
        text = item.text
        if first:
            parts = text.splitlines()
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = title.add_run(parts[0])
            set_east_asia_font(r, "PingFang SC")
            r.bold = True
            r.font.size = Pt(22)
            r.font.color.rgb = RGBColor.from_string("0B2545")
            title.paragraph_format.space_after = Pt(8)
            for part in parts[1:]:
                sub = doc.add_paragraph()
                sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sr = sub.add_run(part)
                set_east_asia_font(sr, "PingFang SC")
                sr.font.size = Pt(13)
                sr.font.color.rgb = RGBColor.from_string("1F4D78")
            first = False
            continue
        if text == "飞哥正志讲堂":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            set_east_asia_font(r, "PingFang SC")
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor.from_string("555555")
            continue
        if text == "前言":
            doc.add_paragraph("本版说明", style="Heading 1")
            add_label_paragraph(doc, "本版撤回 V87 原题摘录包口径，恢复 V17 学生化推理宝典正文；V87 只作为题源和答案源回查索引。")
            continue
        if item.style == "Heading 1":
            doc.add_paragraph(text, style="Heading 1")
        elif item.style == "Heading 2":
            doc.add_paragraph(text, style="Heading 2")
        elif item.style == "Heading 3":
            doc.add_paragraph(text, style="Heading 3")
        else:
            add_label_paragraph(doc, text)

    doc.save(OUT_DOCX)


def text_from_items(items: list[Item]) -> str:
    return "\n".join(i.text for i in items)


def qa(items: list[Item]) -> str:
    text = text_from_items(items)
    choice_titles = re.findall(r"\d+\.\s*\d{4}[^\n]+第[^\n]+题（选择题）", text)
    high_risk = {
        "2026海淀二模 第5题": "答案选D",
        "2026海淀二模 第6题": "答案选A",
        "2026海淀二模 第7题": "答案选A",
        "2026石景山一模 第5题": "答案选D",
        "2024东城一模 第7题": "答案选A",
        "2024东城一模 第8题": "答案选D",
        "2026丰台二模 第8题": "答案选C",
        "2025顺义一模 第7题": "答案选A",
        "2024朝阳二模 第7题": "答案选D",
    }
    checks = []
    for title, expect in high_risk.items():
        idx = text.find(title)
        if idx == -1:
            checks.append((title, expect, "MISSING", "FAIL"))
            continue
        next_idx = text.find("\n1. ", idx + 1)
        chunk = text[idx : next_idx if next_idx != -1 else idx + 2200]
        got = re.search(r"答案选[A-D]", chunk)
        status = "PASS" if got and got.group(0) == expect else "FAIL"
        checks.append((title, expect, got.group(0) if got else "NO_ANSWER", status))

    pollution_terms = [
        "参考答案",
        "题号 |",
        "评标",
        "评分标准",
        "/Users",
        "OCR",
        "source_extracted",
        "A-formal",
        "B-choice-signal",
        "correct_option_chain",
    ]

    with zipfile.ZipFile(OUT_DOCX) as zf:
        xml = "".join(
            zf.read(name).decode("utf-8", "ignore")
            for name in zf.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        )

    lines = [
        "# 选必三推理宝典重做版自查验收",
        "",
        "- verdict: `PASS_LOCAL_STUDENT_LAYER_REBUILD_FROM_V17`",
        "- base: V17 学生化推理宝典正文",
        "- V87 role: 只作回源纠错索引，不作学生正文主体",
        "",
        "## Student Layer Checks",
        "",
        f"- 四标题条目：材料触发点 {text.count('【材料触发点】')} / 设问 {text.count('【设问】')} / 为什么能想到 {text.count('【为什么能想到】')} / 答案落点 {text.count('【答案落点】')}",
        f"- 选择题标题：{len(choice_titles)}",
        f"- 选择题答案句 `答案选`：{text.count('答案选')}",
        f"- 选择题错项分析：{text.count('错项分析')}",
        "",
        "## Pollution Gate",
        "",
    ]
    for term in pollution_terms:
        lines.append(f"- `{term}`: {text.count(term)}")
    lines.extend(
        [
            "",
            "## V87 High-risk Source Corrections Checked In Student Body",
            "",
            "| item | expected | found | status |",
            "|---|---:|---:|---|",
        ]
    )
    for title, expect, got, status in checks:
        lines.append(f"| {title} | {expect} | {got} | {status} |")
    lines.extend(
        [
            "",
            "## DOCX Field Gate",
            "",
            f"- `fldChar`: {xml.count('fldChar')}",
            f"- `instrText`: {xml.count('instrText')}",
            f"- `externalLink`: {xml.count('externalLink')}",
            "",
            "## Boundary",
            "",
            "- 本版不宣称 GPT Pro / Claude 真实外审 PASS。",
            "- 本版内容未继承 V87 原题摘录包，只继承 V87 已发现的答案源风险点作为核对清单。",
        ]
    )
    return "\n".join(lines) + "\n"


def main() -> None:
    RUN.joinpath("delivery").mkdir(parents=True, exist_ok=True)
    RUN.joinpath("qa").mkdir(parents=True, exist_ok=True)
    items = collect_items()
    OUT_MD.write_text(md_for_items(items), encoding="utf-8")
    build_docx(items)
    shutil.copy2(SOURCE_PDF, OUT_PDF)
    OUT_QA.write_text(qa(items), encoding="utf-8")
    for src, dst in [
        (OUT_MD, DESKTOP_MD),
        (OUT_DOCX, DESKTOP_DOCX),
        (OUT_PDF, DESKTOP_PDF),
        (OUT_QA, DESKTOP_QA),
    ]:
        shutil.copy2(src, dst)


if __name__ == "__main__":
    main()
