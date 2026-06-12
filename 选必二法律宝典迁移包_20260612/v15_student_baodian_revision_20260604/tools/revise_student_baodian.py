from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path
from tempfile import TemporaryDirectory
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from lxml import etree


SRC = Path("/Users/wanglifei/Downloads/选必二法律与生活_学生宝典.docx")
OUT = Path("/Users/wanglifei/Downloads/选必二法律与生活_学生宝典_v15学生本强化版.docx")


BLUE = RGBColor(47, 104, 175)
DARK = RGBColor(20, 20, 20)
MUTED = RGBColor(90, 90, 90)


def paragraph_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs).strip()


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def set_text(paragraph, text: str, bold: bool = False, color: RGBColor | None = None, size: int | None = None) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_p.addnext(new_para._p)
    new_para._p = new_p
    new_para._element = new_p
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def clone_paragraph_after(paragraph, template, text: str):
    new_p = deepcopy(template._p)
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_p.addnext(new_para._p)
    new_para._p = new_p
    new_para._element = new_p
    set_text(new_para, text)
    return new_para


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def normalize_answer_text(text: str) -> str:
    text = text.strip()
    text = re.sub(r"^参考答案要点[:：]\s*", "", text)
    text = text.replace("；最后", "；最后")
    text = re.sub(r"\s+", " ", text)
    return text


def make_model_answer(anchor: str, a_axis: str, b_axis: str) -> str:
    anchor = normalize_answer_text(anchor)
    a_core = re.sub(r"^A\s*轴[:：]\s*", "", a_axis).split("。")[0].strip()
    b_core = re.sub(r"^B\s*轴[:：]\s*", "", b_axis).split("。")[0].strip()

    # Keep legal substance anchored to Claude's existing points, but turn the
    # output into a student-facing paragraph instead of a build-log field.
    converted = anchor
    replacements = [
        ("先写", "先明确"),
        ("先分", "先分"),
        ("再写", "再结合材料说明"),
        ("再用", "再结合"),
        ("最后写", "最后收束为"),
        ("最后用", "最后结合"),
        ("第一格写", "第一格可写"),
        ("第二格写", "第二格可写"),
        ("第三格写", "第三格可写"),
    ]
    for old, new in replacements:
        converted = converted.replace(old, new)

    lead = "成文范答（仿写版）："
    if a_core or b_core:
        return f"{lead}本题先按 {a_core or '对应法律入口'} 定位，再按 {b_core or '设问动作'} 组织答案。{converted}"
    return f"{lead}{converted}"


def set_compact_style(paragraph) -> None:
    try:
        paragraph.style = "Compact"
    except Exception:
        pass
    for run in paragraph.runs:
        run.font.name = "PingFang SC"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        run.font.size = Pt(10.5)


def add_front_notice(doc: Document) -> None:
    paragraphs = doc.paragraphs
    title_idx = next(
        (i for i, p in enumerate(paragraphs) if paragraph_text(p).startswith("选必二《法律与生活》答题宝典")),
        None,
    )
    if title_idx is None:
        return

    title = paragraphs[title_idx]
    set_text(title, "选必二《法律与生活》答题宝典（学生训练版）", bold=True, color=BLUE, size=17)

    intro = paragraphs[title_idx + 1]
    set_text(
        intro,
        "定位：考场速查框架 + 42题精简训练。先用速查卡判断入口，再用A/B双轴组织答案，最后用题卡练仿写。"
        "本册只保留学习和练题需要的信息，不放制作记录、版本说明和后台审计材料。",
        bold=False,
        color=DARK,
        size=11,
    )

    note = insert_paragraph_after(
        intro,
        "使用边界：本册保留42题训练样例，用于掌握法律题入口和踩分写法；不是2024-2026全量题库，正式练题仍以原卷材料和评分细则为准。",
        style="Block Text",
    )
    for run in note.runs:
        run.bold = True
        run.font.color.rgb = MUTED
        run.font.size = Pt(10.5)


def rewrite_question_entries(doc: Document) -> None:
    paragraphs = doc.paragraphs
    for i, p in enumerate(list(paragraphs)):
        text = paragraph_text(p)
        if not text.startswith("参考答案要点："):
            continue

        a_axis = ""
        b_axis = ""
        for back in range(max(0, i - 5), i):
            t = paragraph_text(paragraphs[back])
            if t.startswith("A 轴："):
                a_axis = t
            elif t.startswith("B 轴："):
                b_axis = t

        anchor = normalize_answer_text(text)
        set_text(p, make_model_answer(text, a_axis, b_axis), size=10)
        set_compact_style(p)

        anchor_para = insert_paragraph_after(p, f"评分锚点：{anchor}", style="Compact")
        set_compact_style(anchor_para)
        for run in anchor_para.runs:
            run.bold = False


def rewrite_closing(doc: Document) -> None:
    for p in doc.paragraphs:
        text = paragraph_text(p)
        if text.startswith("每题 7 项：设问 / 材料核心"):
            set_text(
                p,
                "每题 8 项：设问 / 材料核心 / A 轴定位 / B 轴定位 / 成文范答 / 评分锚点 / 易错预警。先按 A 轴定位读，再按 B 轴形状组织。",
                color=MUTED,
                size=10.5,
            )
            try:
                p.style = "Block Text"
            except Exception:
                pass
            continue
        if text.startswith("结语：本册只重排了"):
            set_text(
                p,
                "结语：用这本书时，先练入口判断，再练成文表达，最后回到原卷材料核对踩分点。不要背空话，也不要把一个案例硬套成别的模块。",
                color=MUTED,
                size=10.5,
            )
            try:
                p.style = "Block Text"
            except Exception:
                pass
            break


def remove_english_toc_heading(doc: Document) -> None:
    for p in list(doc.paragraphs[:3]):
        if paragraph_text(p) == "Table of Contents":
            remove_paragraph(p)
            return


def strip_hidden_english_toc(path: Path) -> None:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with TemporaryDirectory() as td:
        td_path = Path(td)
        with ZipFile(path) as zin:
            zin.extractall(td_path)

        doc_xml = td_path / "word" / "document.xml"
        tree = etree.parse(str(doc_xml))
        root = tree.getroot()
        for p in list(root.xpath(".//w:p", namespaces=ns)):
            text = "".join(p.xpath(".//w:t/text()", namespaces=ns)).strip()
            if text == "Table of Contents":
                parent = p.getparent()
                parent.remove(p)

        tree.write(str(doc_xml), encoding="UTF-8", xml_declaration=True, standalone=True)

        with ZipFile(path, "w", ZIP_DEFLATED) as zout:
            for file in td_path.rglob("*"):
                if file.is_file():
                    zout.write(file, file.relative_to(td_path).as_posix())


def style_runs(doc: Document) -> None:
    for p in doc.paragraphs:
        text = paragraph_text(p)
        for run in p.runs:
            run.font.name = "PingFang SC"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        if text.startswith("成文范答"):
            p.runs[0].bold = True
        elif text.startswith("评分锚点"):
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = BLUE
        elif text.startswith("易错预警"):
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = BLUE


def main() -> None:
    doc = Document(SRC)
    remove_english_toc_heading(doc)
    add_front_notice(doc)
    rewrite_question_entries(doc)
    rewrite_closing(doc)
    style_runs(doc)
    doc.save(OUT)
    strip_hidden_english_toc(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
