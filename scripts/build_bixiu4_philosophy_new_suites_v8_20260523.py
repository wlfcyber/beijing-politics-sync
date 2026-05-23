from __future__ import annotations

import csv
import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RUN = ROOT / "reports" / "bixiu4_philosophy_strict_v8_2026-05-23"
OUT = RUN / "new_9_suite_integration"
DESKTOP = Path.home() / "Desktop" / "5.23哲学宝典65套新增卷子并入v8"

BASE_MD = (
    ROOT
    / "artifacts"
    / "desktop_exports_2026-05-23"
    / "5.23哲学宝典漏题排查与二模补丁v7"
    / "01_学生版Word"
    / "必修四哲学材料-知识触发框架_v7_纯学生版.md"
)


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def write_csv(path: Path, rows: list[dict[str, str]], fields: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def strip_v7_second_mock_appendix(md: str) -> str:
    marker = "\n## 2026二模正式补入条目\n"
    idx = md.find(marker)
    if idx == -1:
        return md
    return md[:idx].rstrip() + "\n"


MAIN_ENTRIES: list[dict[str, str]] = [
    {
        "suite": "2026东城二模",
        "question": "第16题",
        "status": "正式并入",
        "evidence_level": "A_formal_rubric",
        "framework_nodes": "实践与认识；联系观；发展观；矛盾特殊性；价值观；文化育人",
        "student_title": "京彩课堂：为什么“实景+任务+探究”会让思政课真正精彩",
        "prompt": "结合材料，运用《哲学与文化》知识，谈谈“京彩课堂”的精彩所在。",
        "material_trigger": "“京彩课堂”把中轴线、故宫等文化实景变成“实景+任务+探究”的学习实践，让学生在真实场景中体验、探究、验证和深化认识。",
        "why": "设问问“精彩所在”，细则把精彩落到实践育人：学生不是只背抽象理论，而是在真实文化场景中获得认识，再把认识回到育人实践。",
        "answer": "可写实践是认识的来源、动力和检验标准，京彩课堂用真实场景和探究任务把抽象思政理论转化为直观感悟；也可扣住联系、发展、矛盾特殊性和价值观，说明课程把文化资源、学习任务和育人目标统一起来。",
        "evidence": "细则点名“实践与认识辩证关系角度看”，示例写“实践是认识的来源、发展动力及检验真理的唯一标准”，并列联系、发展、矛盾特殊性、价值观、文化育人等角度。",
    },
    {
        "suite": "2026丰台二模",
        "question": "第16题",
        "status": "正式并入",
        "evidence_level": "A_formal_rubric",
        "framework_nodes": "具体问题具体分析；规律与主观能动性；联系观/系统优化；价值观导向；两点论重点论",
        "student_title": "守护湿地：一题同时打到具体分析、规律能动、系统治理和生态价值观",
        "prompt": "结合材料，运用《哲学与文化》知识，谈谈你对“守护湿地，就是守护生态之美与文化根脉”的认识。",
        "material_trigger": "材料把湿地放在生态系统和文化根脉中整体看，从单纯治理局部水域转向统筹流域、生态、文化和人的价值追求。",
        "why": "这不是单纯环保题，而是要求说明为什么湿地治理既有生态意义又有文化意义。材料中的“生态之美”“文化根脉”“系统治理”分别触发规律、联系和价值观。",
        "answer": "可写守护湿地要具体分析湿地生态和文化特点，尊重生态规律并发挥主观能动性；用系统观念统筹流域治理、生态修复和文化保护；坚持正确生态价值观，把生态保护和文化传承统一起来。",
        "evidence": "阅卷细则列出“矛盾特殊性/具体问题具体分析、尊重规律与主观能动性相结合、联系观点/系统优化、正确价值观的导向作用”；替代角度含“两点论与重点论相统一”。",
    },
    {
        "suite": "2026朝阳二模",
        "question": "第16题",
        "status": "正式并入",
        "evidence_level": "A_formal_rubric",
        "framework_nodes": "对立统一；辩证否定；发展观；实践基础；具体问题具体分析；价值观导向；文化双创",
        "student_title": "历史传承与现代创新：不是二选一，而是在对立统一中守正创新",
        "prompt": "结合材料，从哲学角度，说明在城市文化建设中怎样实现历史传承与现代创新的有机结合。",
        "material_trigger": "城市文化建设既要保护历史遗存、延续文脉，又要回应新时代市民文化需求、用现代方式激活公共文化空间。",
        "why": "题干直接给出“历史传承与现代创新”的关系。看到“有机结合”，不能只写继承或只写创新，而要写传承和创新在矛盾统一中相互支撑。",
        "answer": "可写传承与创新是对立统一关系；坚持辩证否定，在保留城市文脉中注入时代内容；立足城市实践和群众需求具体施策，用正确价值观引领城市文化建设，实现优秀传统文化创造性转化、创新性发展。",
        "evidence": "细则表写“传承与创新的关系——对立统一/联系”，“继承——坚持辩证否定/发展的观点”，“创新——立足实践/回应群众需求/具体问题具体分析/一切从实际出发”，并列价值观导向作用。",
    },
    {
        "suite": "2026房山二模",
        "question": "第16题",
        "status": "正式并入",
        "evidence_level": "A_formal_rubric",
        "framework_nodes": "量变质变；规律与主观能动性；系统观念；矛盾对立统一；价值观导向；人民群众；民族精神；文化双创",
        "student_title": "中华工业文化：从精益求精的量变积累，读懂民族精神和现代工业浪漫",
        "prompt": "结合材料，运用《哲学与文化》知识，谈谈如何从中华工业文化读懂中华民族最感动人的浪漫。",
        "material_trigger": "材料从良渚玉器、西周青铜、《考工记》分工写到北斗原子钟和完整工业体系，反复呈现长期积累、极致工艺、系统协作和创新突破。",
        "why": "“最感动人的浪漫”不是空泛赞美，而是把工匠精神、工业系统、民族精神和实践创新串起来。细则明确要求多点作答，不能只放在文化双创下面。",
        "answer": "可写长期精益求精的量变积累推动工艺和工业能力发生质的飞跃；劳动者尊重规律并发挥主观能动性，把“天工”转为日常；系统协作和矛盾统一推动制造成为艺术；正确价值观、工匠精神和民族精神支撑中华工业文化守正创新。",
        "evidence": "细则列出尊重客观规律与发挥主观能动性、系统观念、联系、发展、矛盾、人民群众、中华民族精神、价值观导向，并点名“量的积累促成质变”。",
    },
    {
        "suite": "2026顺义二模",
        "question": "第16题",
        "status": "正式并入",
        "evidence_level": "A_formal_rubric",
        "framework_nodes": "两点论与重点论；矛盾普遍性与特殊性；人民群众；社会存在与社会意识；价值观导向；实践基础；文化创作",
        "student_title": "新大众文艺：三组关系一起平衡，不能只写“文化创新”",
        "prompt": "结合材料，运用《哲学与文化》知识，谈谈对新大众文艺面对“尊重多样与弘扬主流、人民性与艺术性、社会效益与经济效益”之间寻找平衡点这一时代课题的思考。",
        "material_trigger": "新大众文艺由人民大众担当主体、主创、主角，同时面临多样与主流、人民性与艺术性、社会效益与经济效益三组关系。",
        "why": "材料不是只讲文艺形式新，而是在三组矛盾关系中找平衡。看到“既要……又要……”“平衡点”，应想到两点论与重点论、价值判断和人民主体。",
        "answer": "可写新大众文艺要坚持两点论与重点论统一，既尊重多样又弘扬主流；坚持人民群众主体地位，扎根人民实践；用正确价值观引领创作，发挥先进社会意识的推动作用，统一社会效益和经济效益。",
        "evidence": "细则展开“两点论与重点论统一、人民群众是历史的创造者、社会存在与社会意识、价值观导向、实践观点”等作答路径。",
    },
    {
        "suite": "2026朝阳二模",
        "question": "第21题",
        "status": "综合题哲学维度并入",
        "evidence_level": "A_formal_rubric_cross",
        "framework_nodes": "系统优化；整体与部分；联系观；量变质变；对立统一",
        "student_title": "系统思维与战略定力：整体布局加长期坚持",
        "prompt": "综合运用所学，谈谈对系统思维和战略定力的认识。",
        "material_trigger": "材料把多个中国式现代化支柱作为整体推进，并强调方向坚定、目标执着、行动持久。",
        "why": "系统思维回答“怎么整体布局”，战略定力回答“怎么长期坚持”。二者结合，正好对应联系整体观和量变质变的持续积累。",
        "answer": "可写推进中国式现代化要从整体出发统筹各方面任务，处理好部分与整体、当前与长远的关系；长期目标要通过阶段性任务持续积累，久久为功推动质变。",
        "evidence": "细则写“系统思维要求我们用联系、全面、整体的观点看问题”，战略定力体现方向坚定和行动持久，并点到量变质变、对立统一、联系观。",
    },
    {
        "suite": "2026房山二模",
        "question": "第21题",
        "status": "综合题哲学维度并入",
        "evidence_level": "A_formal_rubric_cross",
        "framework_nodes": "矛盾普遍性；联系观；发展观；科学思维方法",
        "student_title": "新征程就是新的长征：承认矛盾，在斗争和发展中前进",
        "prompt": "综合运用所学，谈谈对“新征程就是新的长征”的理解。",
        "material_trigger": "“新的长征”强调现代化道路仍会遇到风险挑战，需要保持历史主动，在解决问题中推进事业发展。",
        "why": "长征类设问不能只写精神口号，还要看到新征程中的困难、风险和斗争任务。这里触发矛盾普遍性、联系和发展观点。",
        "answer": "可写新征程上矛盾具有普遍性，要敢于承认矛盾、分析矛盾、解决矛盾；在联系和发展中把握现代化建设任务，把长征精神转化为现实实践力量。",
        "evidence": "细则点到“矛盾普遍性，承认分析解决矛盾/伟大斗争”，并列“联系/发展、运用科学思维方法”。",
    },
    {
        "suite": "2026顺义二模",
        "question": "第21题",
        "status": "综合题哲学维度并入",
        "evidence_level": "A_formal_rubric_cross",
        "framework_nodes": "规律与主观能动性；实践与认识；联系观；发展观；全面观点",
        "student_title": "先见与远虑：尊重规律、主动谋划，再用实践检验和推进",
        "prompt": "综合运用所学，阐述“先见”与“远虑”的传统理政智慧为何能成为支撑中国式现代化稳步前行、护航民族伟大复兴行稳致远的深厚精神底气与重要实践指引。",
        "material_trigger": "“先见”“远虑”强调在复杂环境中预判趋势、守住底线、谋定后动、久久为功。",
        "why": "传统理政智慧能成为现代化实践指引，是因为它不是空泛经验，而是把尊重规律、主动谋划和实践检验统一起来。",
        "answer": "可写推进现代化要尊重客观规律并发挥主观能动性，用实践发展认识、用认识指导实践，以联系、发展、全面观点谋定后动、久久为功。",
        "evidence": "细则写“必修4哲学、选必3科学思维”，并点名尊重客观规律与发挥主观能动性相统一、认识与实践的辩证发展、联系发展全面观点。",
    },
    {
        "suite": "2026通州一模",
        "question": "第18题",
        "status": "正式并入",
        "evidence_level": "A_formal_rubric",
        "framework_nodes": "硬采分：对立统一；辩证否定；文化双创；立足社会实践。替代/副点：发展观可替代辩证否定；联系观、规律与能动性等只作1分替代槽，不重复给分",
        "student_title": "隆福寺街区改造：“古朴”和“创新”怎样共生共荣",
        "prompt": "结合材料，运用《哲学与文化》知识，阐释隆福寺街区改造中“古朴”与“创新”是如何实现共生共荣的。",
        "material_trigger": "隆福寺没有大拆大建，而是“绣花式”保护和城市更新；一方面保留寺院格局、老字号、非遗技艺等历史记忆，另一方面引入数字科技、AI拍照亭、年轻消费业态和新生活方式。",
        "why": "看到“古朴”与“创新”并存、共生共荣，就不能把保护和更新看成互相排斥。材料要求说明二者怎样在保护中发展、在传承中创新。",
        "answer": "可写古朴与创新是对立统一关系，隆福寺在保护古都肌理的同时融入现代科技，实现二者共生共荣；坚持辩证否定，既保留历史记忆，又注入时代元素；推动优秀传统文化创造性转化、创新性发展；立足社会实践、坚守文化根脉，让老街区焕发生机与活力。",
        "evidence": "通州评标细则写“矛盾对立统一、辩证否定观、双创，每点2分”，并列“立足社会实践1分”，替代角度含以人民为中心、立足时代、规律与主观能动性、联系观点。",
    },
    {
        "suite": "2026通州一模",
        "question": "第21题",
        "status": "综合题哲学维度并入",
        "evidence_level": "A_formal_rubric_cross",
        "framework_nodes": "主链：实践观点；具体问题具体分析；一切从实际出发；新发展理念。示例副点：价值观、系统思维/系统优化",
        "student_title": "中国式现代化是干出来的：实干怎样落到实践观点和系统治理",
        "prompt": "“中国式现代化是干出来的，伟大事业都成于实干。”结合材料，综合运用所学，谈谈你对这一观点的理解。",
        "material_trigger": "材料从党员干部务实为民、航天工程长期攻关、河长制压实责任、高铁建设改善民生等角度说明“实干”如何推动现代化。",
        "why": "设问核心是“干出来”，这直接触发实践观点；材料又不是单点行动，而是制度、科技、治理、民生多方面协同，适合用具体问题具体分析和系统思维收束。",
        "answer": "可写中国式现代化不是空想蓝图，而是在党的领导下通过实践探索、接续奋斗推进的历史进程；要坚持具体问题具体分析、一切从实际出发，把制度优势转化为治理效能；坚持系统思维统筹协调，用实干推动科技攻关、生态治理和民生改善。",
        "evidence": "通州评标示例写“实践是认识的基础和目的”，并明确“可从党的初心使命、制度优势、具体问题具体分析、实践的观点、新发展理念等角度作答”，示例中还点到系统思维、系统优化。",
    },
]


CHOICE_ENTRIES: list[dict[str, str]] = [
    {
        "suite": "2026通州一模",
        "question": "第2题",
        "answer": "B（①③）",
        "status": "选择题速记并入",
        "knowledge": "文化自信；中华文化影响力；文化交流",
        "correct_chain": "外国游客探索中式生活方式，能支撑“中国开放中的文化自信”和“中华文化具有独特魅力、文化软实力提升”。",
        "wrong_option_warning": "②把题意转向一般文化多样性和尊重不同民族文化，离开“中国生活方式吸引外国人”的主旨；④“普遍认可”扩大化。",
    },
    {
        "suite": "2026通州一模",
        "question": "第3题",
        "answer": "B（①④）",
        "status": "选择题速记并入",
        "knowledge": "社会存在与社会意识；辩证否定；科学立法",
        "correct_chain": "耕地保护立法回应现实问题，体现社会意识源于社会存在并反作用于社会发展；在继承既有规范基础上完善草案，可对应辩证否定。",
        "wrong_option_warning": "②把从个别经验到一般规则说成演绎推理，推理类型错；③说上层建筑“引领经济基础变革”过度且关系表述不当。",
    },
    {
        "suite": "2026通州一模",
        "question": "第4题",
        "answer": "A",
        "status": "选择题速记并入",
        "knowledge": "系统优化；创新驱动；产业创新与制度创新协同",
        "correct_chain": "人工智能+、量子科技、具身智能以及产业创新与制度创新协同发力，提示用系统优化的科学思维方法推进创新驱动发展。",
        "wrong_option_warning": "B把科技创新说成基本国策；C把智能经济夸成塑造新质生产力的核心；D硬套主要矛盾，材料没有要求抓“产业创新”这一个主要矛盾。",
    },
    {
        "suite": "2026通州一模",
        "question": "第9题",
        "answer": "B（①④）",
        "status": "选择题速记并入",
        "knowledge": "文化交流互鉴；文化传播；价值观与生活方式对话",
        "correct_chain": "文化出海既是传播中华文化，也是价值观、生活方式的跨文化对话，有助于提升国家文化软实力。",
        "wrong_option_warning": "②把“中华优秀传统文化”说成文化出海的核心载体，范围过窄；③“增强文化认同”主体容易泛化，题意更稳的是影响力和软实力。",
    },
    {
        "suite": "2026通州一模",
        "question": "第12题",
        "answer": "D（③④）",
        "status": "跨书边界：不入必修四正文，只作选必三提示",
        "knowledge": "联想思维；创新思维",
        "correct_chain": "机器人春晚体现联想思维的畅想性和创新思维的独特性，属于《逻辑与思维》更直接的题。",
        "wrong_option_warning": "①说超前思维具有确定性，表述错误；②把艺术表达说成必须依托科技实现突破，绝对化。",
    },
]


BOUNDARY_ROWS: list[dict[str, str]] = [
    {
        "suite": "2026西城二模",
        "question": "第16题",
        "decision": "不进学生正文",
        "reason": "西城评标PDF提取为空，现有证据仅为参考答案方向“可从矛盾普遍性和特殊性、实践、中华优秀传统文化等角度作答”，不足以当正式细则链。",
    },
    {
        "suite": "2026海淀二模",
        "question": "第16题",
        "decision": "不进学生正文",
        "reason": "讲评PDF未提取出正文，仅有“联系、实践与认识”等方向词，暂作审计候选。",
    },
    {
        "suite": "2026海淀二模",
        "question": "第21题",
        "decision": "不进学生正文",
        "reason": "哲学只是综合题并列角度之一，且讲评原文缺失，不能按正式哲学样题收。",
    },
    {
        "suite": "2026石景山二模",
        "question": "第17(3)题",
        "decision": "暂不进学生正文",
        "reason": "细则只要求使用一个哲学观点概括并分析，示例主要展开“联系”，矛盾和实践认识维度未给足示范，留作教研候选。",
    },
    {
        "suite": "2026石景山二模",
        "question": "第20题",
        "decision": "不进学生正文",
        "reason": "哲学只是“系统观”等并列方向，细则未对哲学维度展开。",
    },
    {
        "suite": "2026丰台二模",
        "question": "第22题",
        "decision": "不进学生正文",
        "reason": "细则只到主要矛盾、系统观念、联系、创新、生产力等知识点列举层级，未给完整作答示范。",
    },
    {
        "suite": "2026通州一模",
        "question": "第17(2)题",
        "decision": "不进必修四哲学正文",
        "reason": "设问限定为《逻辑与思维》辩证思维，直接归选必三；可在跨书边界提示中保留。",
    },
    {
        "suite": "2026通州一模",
        "question": "第16题",
        "decision": "不进必修四哲学正文",
        "reason": "设问限定为《政治与法治》，虽然评标出现系统思维字样，但不是必修四哲学主链。",
    },
    {
        "suite": "2026通州一模",
        "question": "第19题",
        "decision": "不进必修四哲学正文",
        "reason": "设问限定为《当代国际政治与经济》。",
    },
]


NEW_SUITES = [
    "2026东城二模",
    "2026丰台二模",
    "2026房山二模",
    "2026朝阳二模",
    "2026海淀二模",
    "2026石景山二模",
    "2026西城二模",
    "2026通州一模",
    "2026顺义二模",
]


def suite_coverage_rows() -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for suite in NEW_SUITES:
        main = [f"{e['question']}({e['status']})" for e in MAIN_ENTRIES if e["suite"] == suite]
        choices = [f"{e['question']}({e['status']})" for e in CHOICE_ENTRIES if e["suite"] == suite]
        boundaries = [f"{e['question']}({e['decision']})" for e in BOUNDARY_ROWS if e["suite"] == suite]
        rows.append(
            {
                "suite": suite,
                "student_body_entries": "；".join(main) if main else "",
                "choice_entries": "；".join(choices) if choices else "",
                "boundary_or_candidate_entries": "；".join(boundaries) if boundaries else "",
                "suite_status": "included_in_body" if main else "handled_as_boundary",
            }
        )
    return rows


def grouped_new_section() -> str:
    groups: list[tuple[str, list[str]]] = [
        (
            "实践与认识、立足实践",
            ["2026东城二模-第16题", "2026顺义二模-第16题", "2026通州一模-第18题", "2026通州一模-第21题", "2026顺义二模-第21题"],
        ),
        (
            "联系观、整体与部分、系统优化",
            ["2026丰台二模-第16题", "2026朝阳二模-第21题", "2026通州一模-第21题", "2026顺义二模-第21题"],
        ),
        (
            "发展观、量变质变、辩证否定",
            ["2026朝阳二模-第16题", "2026房山二模-第16题", "2026通州一模-第18题", "2026房山二模-第21题"],
        ),
        (
            "矛盾观：对立统一、具体问题具体分析、两点论重点论",
            ["2026丰台二模-第16题", "2026朝阳二模-第16题", "2026顺义二模-第16题", "2026通州一模-第18题", "2026房山二模-第21题"],
        ),
        (
            "唯物史观与价值观：人民群众、社会意识、价值导向",
            ["2026丰台二模-第16题", "2026房山二模-第16题", "2026顺义二模-第16题", "2026东城二模-第16题"],
        ),
        (
            "文化传承、文化创新与文化育人",
            ["2026东城二模-第16题", "2026朝阳二模-第16题", "2026房山二模-第16题", "2026顺义二模-第16题", "2026通州一模-第18题"],
        ),
    ]
    by_key = {f"{e['suite']}-{e['question']}": e for e in MAIN_ENTRIES}
    lines: list[str] = [
        "# 2026新增9套卷子：已按哲学框架并入",
        "",
        "说明：这一段放在正文前部，是为了让新增卷子直接进入知识框架，而不是继续留在末尾补丁。每道主观题如对应多个答题点，会在多个框架节点下出现；同一来源的不同节点不是重复刷题，而是提醒学生同一材料可以触发多条原理。",
        "",
    ]
    for group_title, keys in groups:
        lines.extend([f"## {group_title}", ""])
        for key in keys:
            e = by_key[key]
            lines.extend(
                [
                    f"### {e['student_title']}",
                    "",
                    f"**来源题目**：{e['suite']} {e['question']}",
                    "",
                    f"**完整设问**：{e['prompt']}",
                    "",
                    f"**材料触发点**：{e['material_trigger']}",
                    "",
                    f"**为什么能想到这个原理**：{e['why']}",
                    "",
                    f"**答案落点**：{e['answer']}",
                    "",
                ]
            )
    lines.extend(
        [
            "## 新增卷子选择题速记",
            "",
            "说明：这里是选择题正确项与错肢提示，不能反推为主观题评分链。",
            "",
            "| 套卷 | 题号 | 答案 | 归属 | 知识点 | 正确项链条 | 易错提醒 |",
            "|---|---|---|---|---|---|---|",
        ]
    )
    for e in CHOICE_ENTRIES:
        lines.append(
            f"| {e['suite']} | {e['question']} | {e['answer']} | {e['status']} | {e['knowledge']} | {e['correct_chain']} | {e['wrong_option_warning']} |"
        )
    lines.append("")
    return "\n".join(lines)


def audit_md() -> str:
    included = [e for e in MAIN_ENTRIES if e["status"] != "边界排除"]
    lines = [
        "# v8 新增9套卷子并入审计",
        "",
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "## 范围结论",
        "",
        "- 当前必修四哲学题源基数：65套。",
        "- 新增正式范围：2026通州一模 + 8套2026二模。",
        f"- 新增主观/综合哲学维度并入条目：{len(included)}条。",
        f"- 新增选择题速记：{len(CHOICE_ENTRIES)}条，其中通州第12题仅作选必三边界提示。",
        f"- 明确不进入必修四哲学正文的边界项：{len(BOUNDARY_ROWS)}条。",
        "",
        "## 新增主观题并入表",
        "",
        "| 套卷 | 题号 | 状态 | 证据等级 | 框架节点 |",
        "|---|---|---|---|---|",
    ]
    for e in included:
        lines.append(f"| {e['suite']} | {e['question']} | {e['status']} | {e['evidence_level']} | {e['framework_nodes']} |")
    lines.extend(
        [
            "",
            "## 通州一模处理结论",
            "",
            "- 第18题正式进入哲学与文化宝典正文：评标逐字点名“矛盾对立统一、辩证否定观、双创、立足社会实践”等采分点。",
            "- 第21题按综合题哲学维度进入：评标示例点名实践观点、具体问题具体分析、系统思维等，但不把它伪装成单一哲学主观题。",
            "- 第17(2)题是《逻辑与思维》辩证思维，边界排除出必修四正文。",
            "",
            "## 边界排除",
            "",
            "| 套卷 | 题号 | 决定 | 理由 |",
            "|---|---|---|---|",
        ]
    )
    for row in BOUNDARY_ROWS:
        lines.append(f"| {row['suite']} | {row['question']} | {row['decision']} | {row['reason']} |")
    lines.extend(
        [
            "",
            "## Governor 预审",
            "",
            "PASS_WITH_BOUNDARIES：新增9套已经进入v8学生版和审计表；但旧v7中旧56套补漏仍保留原状，本次只对新增9套负责。西城、海淀、石景山、丰台部分B类候选因证据不足未写进学生正文。",
            "",
            "## Confucius 学生可用性预审",
            "",
            "PASS：新增条目均按“材料触发点 -> 为什么能想到 -> 答案落点”呈现。选择题明确标注为速记，不和主观评分链混用。",
            "",
        ]
    )
    return "\n".join(lines)


def add_run(paragraph, text: str, bold: bool = False, size: int | None = None, color: RGBColor | None = None) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    add_run(p, text, bold=bold, size=8)


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)

    table_buffer: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if not table_buffer:
            return
        rows = [r for r in table_buffer if not all(re.fullmatch(r"-+", c.strip()) for c in r)]
        if rows:
            table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
            table.style = "Table Grid"
            for i, row in enumerate(rows):
                for j, value in enumerate(row):
                    set_cell_text(table.cell(i, j), value, bold=(i == 0))
        table_buffer = []

    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if line.startswith("|") and line.endswith("|"):
            parts = [p.strip() for p in line.strip("|").split("|")]
            table_buffer.append(parts)
            continue
        flush_table()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, line[2:].strip(), bold=True, size=18, color=RGBColor(31, 78, 121))
            continue
        if line.startswith("## "):
            p = doc.add_paragraph()
            add_run(p, line[3:].strip(), bold=True, size=14, color=RGBColor(31, 78, 121))
            continue
        if line.startswith("### "):
            p = doc.add_paragraph()
            add_run(p, line[4:].strip(), bold=True, size=12, color=RGBColor(68, 68, 68))
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style=None)
            add_run(p, "• " + line[2:].strip(), size=10)
            continue
        p = doc.add_paragraph()
        m = re.match(r"^\*\*(.+?)\*\*：(.+)$", line)
        if m:
            add_run(p, m.group(1) + "：", bold=True, size=10)
            add_run(p, m.group(2), size=10)
        else:
            add_run(p, line, size=10)
    flush_table()
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    (DESKTOP / "01_学生版Word").mkdir(parents=True, exist_ok=True)
    (DESKTOP / "02_审计与证据").mkdir(parents=True, exist_ok=True)
    (DESKTOP / "03_结构化CSV").mkdir(parents=True, exist_ok=True)

    base = BASE_MD.read_text(encoding="utf-8")
    base = strip_v7_second_mock_appendix(base)
    base = base.replace("# 必修四哲学材料-知识触发框架 v6", "# 必修四哲学材料-知识触发框架 v8（当前65套新增卷子并入版）", 1)
    base = re.sub(r"生成时间：.+", f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}；本版在v7基础上将新增9套卷子按框架并入。", base, count=1)
    final_md = base.replace(
        "## 样例已确认：2026东城一模",
        grouped_new_section() + "\n\n## 样例已确认：2026东城一模",
        1,
    )

    student_md = OUT / "必修四哲学材料-知识触发框架_v8_65套新增卷子并入版.md"
    audit_path = OUT / "new_9_suite_integration_audit.md"
    main_csv = OUT / "new_9_suite_main_entries.csv"
    choice_csv = OUT / "new_9_suite_choice_entries.csv"
    boundary_csv = OUT / "new_9_suite_boundary_exclusions.csv"
    suite_coverage_csv = OUT / "new_9_suite_suite_coverage_matrix.csv"
    docx_path = OUT / "必修四哲学材料-知识触发框架_v8_65套新增卷子并入版.docx"

    write_text(student_md, final_md)
    write_text(audit_path, audit_md())
    write_csv(main_csv, MAIN_ENTRIES, list(MAIN_ENTRIES[0].keys()))
    write_csv(choice_csv, CHOICE_ENTRIES, list(CHOICE_ENTRIES[0].keys()))
    write_csv(boundary_csv, BOUNDARY_ROWS, list(BOUNDARY_ROWS[0].keys()))
    write_csv(suite_coverage_csv, suite_coverage_rows(), ["suite", "student_body_entries", "choice_entries", "boundary_or_candidate_entries", "suite_status"])
    md_to_docx(student_md, docx_path)

    for path in [student_md, docx_path]:
        shutil.copy2(path, DESKTOP / "01_学生版Word" / path.name)
    for path in [audit_path]:
        shutil.copy2(path, DESKTOP / "02_审计与证据" / path.name)
    for path in [main_csv, choice_csv, boundary_csv, suite_coverage_csv]:
        shutil.copy2(path, DESKTOP / "03_结构化CSV" / path.name)

    build_report = OUT / "BUILD_REPORT_V8_NEW_9.md"
    write_text(
        build_report,
        "\n".join(
            [
                "# v8 新增9套卷子并入构建报告",
                "",
                f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                "",
                f"- 学生版 Markdown：`{student_md}`",
                f"- 学生版 Word：`{docx_path}`",
                f"- 审计报告：`{audit_path}`",
                f"- 主观条目CSV：`{main_csv}`",
                f"- 选择题CSV：`{choice_csv}`",
                f"- 边界排除CSV：`{boundary_csv}`",
                f"- 新增9套套卷级覆盖矩阵：`{suite_coverage_csv}`",
                "",
                "状态：新增9套已并入；旧56套旧补漏质量不是本脚本处理范围。",
                "",
            ]
        ),
    )
    shutil.copy2(build_report, DESKTOP / "02_审计与证据" / build_report.name)
    print(student_md)
    print(docx_path)
    print(audit_path)


if __name__ == "__main__":
    main()
