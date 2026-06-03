from __future__ import annotations

import csv
import re
from pathlib import Path

from docx import Document


RUN_DIR = Path("/Users/wanglifei/GaokaoPolitics/beijing-politics-sync/reports/选必一_哲学宝典式重建_2026-05-16/20_review237_verify_final_20260528")
REVIEW_DOCX = Path("/Users/wanglifei/Desktop/宝典全文237条逐条审核_20260527.docx")
BASE_DOCX = Path("/Users/wanglifei/Desktop/选必一_当代国际政治与经济_主观题术语宝典_Claude严格细则修订版_OCR后_20260527.docx")


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip())


def review_paragraphs() -> list[str]:
    return [clean(p.text) for p in Document(str(REVIEW_DOCX)).paragraphs]


def parse_flagged(paras: list[str]) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    current: dict[str, str] | None = None
    bucket = ""
    core = ""
    for para_no, text in enumerate(paras, 1):
        if text.startswith("四、逐条审核"):
            break
        if text.startswith("【") and "】核心答题点：" in text:
            m = re.match(r"【([^】]+)】核心答题点：(.+)", text)
            if m:
                bucket, core = m.group(1), m.group(2)
            continue
        if text.startswith("· 题目："):
            current = {
                "flag_order": str(len(rows) + 1),
                "review_para": str(para_no),
                "bucket": bucket,
                "core": core,
                "question": text.split("：", 1)[1],
            }
            rows.append(current)
            continue
        if current is None:
            continue
        if text.startswith("· 卷面句："):
            current["sentence"] = text.split("：", 1)[1]
        elif text.startswith("· 最长命中片段"):
            current["hit"] = text
        elif text.startswith("· 细则未含的卷面句小句："):
            current["missing_clause"] = text.split("：", 1)[1]
        elif text.startswith("· 细则原文"):
            current["rubric_excerpt"] = text.split("：", 1)[1]
    return rows


def parse_full_review(paras: list[str]) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    start = next(i for i, t in enumerate(paras) if t.startswith("四、逐条审核"))
    bucket = ""
    core = ""
    current: dict[str, str] | None = None
    status_re = re.compile(r"^(✅高度匹配|✅基本匹配|⚠部分匹配|❌细则无支撑片段)\s+(.+)$")
    for para_no, text in enumerate(paras[start + 1 :], start + 2):
        if not text:
            continue
        if text in {"时代背景", "理论", "经济全球化", "政治多极化", "中国", "联合国"}:
            bucket = text
            continue
        if text.startswith("◆ 核心答题点："):
            core = text.removeprefix("◆ 核心答题点：")
            continue
        m = status_re.match(text)
        if m:
            current = {
                "review_id": str(len(rows) + 1),
                "review_para": str(para_no),
                "bucket": bucket,
                "core": core,
                "status": m.group(1),
                "question": m.group(2),
            }
            rows.append(current)
            continue
        if current is None:
            continue
        if text.startswith("卷面句："):
            current["sentence"] = text.split("：", 1)[1]
        elif text.startswith("命中片段"):
            current["hit"] = text
        elif text.startswith("细则未含小句："):
            current["missing_clause"] = text.split("：", 1)[1]
        elif text.startswith("同题组："):
            current["same_group"] = text.split("：", 1)[1]
    return rows


def parse_base_entries() -> list[dict[str, str]]:
    paras = [clean(p.text) for p in Document(str(BASE_DOCX)).paragraphs]
    rows: list[dict[str, str]] = []
    bucket = ""
    core = ""
    current: dict[str, str] | None = None
    q_re = re.compile(r"^\d+\.\s+(.+)$")
    for para_no, text in enumerate(paras, 1):
        if not text:
            continue
        if text in {"时代背景", "理论", "经济全球化", "政治多极化", "中国", "联合国"}:
            bucket = text
            continue
        if text.startswith("核心答题点："):
            core = text.removeprefix("核心答题点：")
            continue
        m = q_re.match(text)
        if m and core:
            current = {
                "base_id": str(len(rows) + 1),
                "base_para": str(para_no),
                "bucket": bucket,
                "core": core,
                "question": m.group(1),
            }
            rows.append(current)
            continue
        if current is None:
            continue
        for label in ["什么时候写", "设问", "得分层次", "卷面句", "同题组"]:
            prefix = f"【{label}】"
            if text.startswith(prefix):
                current[label] = clean(text.removeprefix(prefix))
                break
    return rows


def write_csv(path: Path, rows: list[dict[str, str]]) -> None:
    keys: list[str] = []
    for row in rows:
        for key in row:
            if key not in keys:
                keys.append(key)
    with path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=keys)
        writer.writeheader()
        writer.writerows(rows)


def main() -> None:
    RUN_DIR.mkdir(parents=True, exist_ok=True)
    paras = review_paragraphs()
    flagged = parse_flagged(paras)
    full = parse_full_review(paras)
    base = parse_base_entries()
    write_csv(RUN_DIR / "REVIEW237_FULL_PARSED.csv", full)
    write_csv(RUN_DIR / "REVIEW237_FLAGGED32_PARSED.csv", flagged)
    write_csv(RUN_DIR / "BASE237_ENTRIES_PARSED.csv", base)

    matched = 0
    base_keys = {(r["bucket"], r["question"], r.get("卷面句", "")): r for r in base}
    for row in full:
        key = (row["bucket"], row["question"], row.get("sentence", ""))
        if key in base_keys:
            matched += 1
    summary = [
        "# REVIEW237_PARSE_SUMMARY",
        "",
        f"- review_full_rows: {len(full)}",
        f"- review_flagged_rows: {len(flagged)}",
        f"- base_rows: {len(base)}",
        f"- exact_bucket_question_sentence_matches: {matched}",
    ]
    (RUN_DIR / "REVIEW237_PARSE_SUMMARY.md").write_text("\n".join(summary) + "\n", encoding="utf-8")


if __name__ == "__main__":
    main()

