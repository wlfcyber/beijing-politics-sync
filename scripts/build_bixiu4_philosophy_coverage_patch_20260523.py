from __future__ import annotations

import csv
import re
import shutil
from collections import Counter
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path.cwd()
RUN = ROOT / "reports" / "bixiu4_philosophy_coverage_patch_2026-05-23"
DESKTOP = Path.home() / "Desktop" / "5.23哲学宝典漏题排查与二模补丁v7"

OLD_ROSTER = ROOT / "artifacts" / "desktop_exports_2026-04-29" / "4.29凌晨跑完的结果v6" / "04_过程日志" / "SUITE_ROSTER.csv"
OLD_V6 = ROOT / "artifacts" / "desktop_exports_2026-04-29" / "4.29凌晨跑完的结果v6" / "01_学生版Word" / "必修四哲学材料-知识触发框架_v6.md"
V3_INVENTORY = ROOT / "reports" / "philosophy_v3_reaudit_2026-04-26" / "artifacts" / "framework_entry_inventory.csv"
CLAUDE_REVIEW = RUN / "claudecode_second_mock_philosophy_review.md"

STUDENT = DESKTOP / "01_学生版Word"
AUDIT = DESKTOP / "02_覆盖审计"
CSVOUT = DESKTOP / "03_结构化CSV"
LOGS = DESKTOP / "04_过程日志"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def write_csv(path: Path, rows: list[dict[str, object]], fieldnames: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            writer.writerow(row)


def source_lines(md: str) -> list[tuple[str, str, str]]:
    out: list[tuple[str, str, str]] = []
    for line in md.splitlines():
        if line.startswith("**来源题目**："):
            src = line.split("：", 1)[1].strip()
            suite = src.split("第", 1)[0].strip()
            out.append((suite, qnorm(src), src))
    return out


def qnorm(text: str) -> str:
    m = re.search(r"第\s*(\d+)", text)
    if m:
        return m.group(1)
    m = re.search(r"(\d+)", text)
    return m.group(1) if m else ""


def suite_name_from_inventory(row: dict[str, str]) -> str:
    name = f"{row.get('year', '')}{row.get('district', '')}{row.get('stage', '')}"
    if name == "2024顺义二模":
        return "2024顺义思政二模"
    return name


def trim(text: str, limit: int = 180) -> str:
    text = re.sub(r"\s+", " ", (text or "").strip())
    return text if len(text) <= limit else text[: limit - 1] + "…"


SECOND_MOCK_A: list[dict[str, str]] = [
    {
        "suite_name": "2026东城二模",
        "question_no": "16",
        "evidence_level": "A_formal_rubric",
        "source_path": "reports/选必一_哲学宝典式重建_2026-05-16/08_2026_second_mock_backfill/00_extracted_text/texts/2026东城二模__细则__16.pdf.txt",
        "evidence_excerpt": "细则点名“实践与认识辩证关系角度看”，示例写“实践是认识的来源、发展动力及检验真理的唯一标准”。",
        "knowledge_point": "实践与认识辩证关系",
        "material_trigger": "京彩课堂把中轴线、故宫等文化实景变成“实景+任务+探究”的学习实践。",
        "why": "设问问“精彩所在”，细则把精彩落到实践育人：学生不是只背抽象理论，而是在真实文化场景中体验、探究、验证和深化认识。",
        "answer_landing": "写实践是认识的来源、动力和检验标准，京彩课堂用真实场景和探究任务把抽象思政理论转化为直观感悟，并让认识回到育人实践。",
    },
    {
        "suite_name": "2026丰台二模",
        "question_no": "16",
        "evidence_level": "A_formal_rubric",
        "source_path": "reports/选必一_哲学宝典式重建_2026-05-16/08_2026_second_mock_backfill/00_extracted_text/texts/2026丰台二模__细则__2026丰台区二模主观题阅卷细则(1).pptx.txt",
        "evidence_excerpt": "细则列出“矛盾特殊性/具体问题具体分析、尊重规律与主观能动性相结合、联系观点/系统优化、正确价值观的导向作用”。",
        "knowledge_point": "具体问题具体分析、规律与能动性、系统优化、价值观导向",
        "material_trigger": "湿地治理从单纯“治湖泊”转向系统“治流域”，同时兼顾生态之美和文化根脉。",
        "why": "材料把湿地放在生态系统和文化根脉中整体看，既要按湿地生态规律治理，又要根据不同湿地问题选择路径。",
        "answer_landing": "写守护湿地要具体分析湿地生态和文化特点，尊重生态规律并发挥主观能动性，用系统观念统筹治理流域，同时坚持正确生态价值观。",
    },
    {
        "suite_name": "2026朝阳二模",
        "question_no": "16",
        "evidence_level": "A_formal_rubric",
        "source_path": "reports/选必一_哲学宝典式重建_2026-05-16/08_2026_second_mock_backfill/00_extracted_text/texts/2026朝阳二模__细则__202605朝阳高三政治二模阅卷细则(1).docx.txt",
        "evidence_excerpt": "细则表写“传承与创新的关系：对立统一/联系”；“坚持辩证否定/发展的观点”；“立足实践、具体问题具体分析、一切从实际出发”；“价值观导向作用/社会意识作用”。",
        "knowledge_point": "对立统一、辩证否定、实践基础、具体问题具体分析、价值观导向",
        "material_trigger": "城市文化建设既要保护历史文化古迹、延续文脉，又要回应新时代市民文化需求、创新公共文化服务。",
        "why": "题干直接问“历史传承与现代创新”的有机结合，材料不是二选一，而是要求在统一关系中处理保护与创新。",
        "answer_landing": "写传承与创新是对立统一关系；坚持辩证否定，在守住城市文脉中发展创新；立足城市实践和群众需求具体施策，并以正确价值观引领城市文化软实力建设。",
    },
    {
        "suite_name": "2026房山二模",
        "question_no": "16",
        "evidence_level": "A_formal_rubric",
        "source_path": "reports/选必一_哲学宝典式重建_2026-05-16/08_2026_second_mock_backfill/00_extracted_text/texts/2026房山二模__细则__26房山评标.docx.txt",
        "evidence_excerpt": "细则列出“正确发挥主观能动性与尊重客观规律、系统观念、价值观导向作用、联系、发展、矛盾、人民群众”，并点名“量的积累促成质变”。",
        "knowledge_point": "量变质变、规律与主观能动性、系统观念、矛盾对立统一、价值观导向",
        "material_trigger": "从良渚玉器、青铜纹饰、《考工记》分工，到北斗原子钟和完整工业体系，材料反复呈现长期积累、精益求精和系统协作。",
        "why": "“最感动人的浪漫”不是抽象赞美，而是把极致工艺、工业系统、工匠精神和民族精神串起来，细则明确要求用多条哲学链解释。",
        "answer_landing": "写中华工业文化体现长期量的积累推动质的飞跃；劳动者尊重规律并发挥主观能动性，把天工转为日常；系统协作和对立统一推动制造成为艺术，并以正确价值观和工匠精神支撑创新。",
    },
    {
        "suite_name": "2026顺义二模",
        "question_no": "16",
        "evidence_level": "A_formal_rubric",
        "source_path": "reports/选必一_哲学宝典式重建_2026-05-16/08_2026_second_mock_backfill/00_extracted_text/texts/2026顺义二模__26顺义二模评标.doc.txt",
        "evidence_excerpt": "细则列出“实践观点、人民主体、价值观、矛盾观”，并展开“两点论与重点论统一、人民群众是历史的创造者、社会存在与社会意识、价值观导向”。",
        "knowledge_point": "两点论与重点论、人民群众、社会存在与社会意识、价值观导向、实践基础",
        "material_trigger": "新大众文艺由人民大众担当主体、主创、主角，同时面临尊重多样与弘扬主流、人民性与艺术性、社会效益与经济效益三组关系。",
        "why": "题目不是只问文艺形式新，而是要求在多样与主流、人民性与艺术性、社会效益与经济效益之间找平衡，天然触发矛盾分析和价值判断。",
        "answer_landing": "写新大众文艺要坚持两点论与重点论统一，既尊重多样又弘扬主流；坚持人民群众主体地位，扎根人民实践；用正确价值观引领创作，发挥先进社会意识的推动作用。",
    },
    {
        "suite_name": "2026朝阳二模",
        "question_no": "21",
        "evidence_level": "A_formal_rubric_cross",
        "source_path": "reports/选必一_哲学宝典式重建_2026-05-16/08_2026_second_mock_backfill/00_extracted_text/texts/2026朝阳二模__细则__202605朝阳高三政治二模阅卷细则(1).docx.txt",
        "evidence_excerpt": "细则写“系统思维要求我们用联系、全面、整体的观点看问题”；“战略定力……行动持久（量变与质变、久久为功）”；“对立统一、联系观”。",
        "knowledge_point": "系统优化、整体与部分、量变质变、对立统一",
        "material_trigger": "“四个中国”被确立为中国式现代化战略支柱，要求相互协调、长期推进、一张蓝图绘到底。",
        "why": "系统思维回答“怎么整体布局”，战略定力回答“怎么长期坚持”。二者结合，正好对应联系整体观和量变质变的持续积累。",
        "answer_landing": "写推进中国式现代化要把数字、健康、美丽、平安中国作为有机整体统筹；长期目标要通过阶段性任务持续积累，久久为功实现质变。",
    },
    {
        "suite_name": "2026房山二模",
        "question_no": "21",
        "evidence_level": "A_formal_rubric_cross",
        "source_path": "reports/选必一_哲学宝典式重建_2026-05-16/08_2026_second_mock_backfill/00_extracted_text/texts/2026房山二模__细则__26房山评标.docx.txt",
        "evidence_excerpt": "细则点到“矛盾普遍性，承认分析解决矛盾/伟大斗争”，并列“联系/发展、运用科学思维方法”。",
        "knowledge_point": "矛盾普遍性、联系与发展",
        "material_trigger": "“新征程是新的长征”强调现代化道路仍会遇到风险挑战，需要在解决矛盾中推进事业发展。",
        "why": "长征类设问不能只写精神口号，还要看到新征程中的风险、困难和斗争任务，承认矛盾、分析矛盾、解决矛盾。",
        "answer_landing": "写新征程上矛盾具有普遍性，要敢于斗争、善于斗争，在联系和发展中把握现代化建设任务，把历史精神转化为现实实践力量。",
    },
    {
        "suite_name": "2026顺义二模",
        "question_no": "21",
        "evidence_level": "A_formal_rubric_cross",
        "source_path": "reports/选必一_哲学宝典式重建_2026-05-16/08_2026_second_mock_backfill/00_extracted_text/texts/2026顺义二模__26顺义二模评标.doc.txt",
        "evidence_excerpt": "细则写“必修4哲学、选必3科学思维”，并点名“尊重客观规律与发挥主观能动性相统一、认识与实践的辩证发展、联系、发展、全面的辩证观点”。",
        "knowledge_point": "规律与主观能动性、实践与认识、联系发展全面观点",
        "material_trigger": "“先见”与“远虑”强调在复杂发展环境中预判趋势、守住底线、谋定后动、久久为功。",
        "why": "传统理政智慧能成为现代化实践指引，是因为它不是空泛经验，而是把尊重规律、主动谋划和实践检验统一起来。",
        "answer_landing": "写推进现代化要尊重客观规律并发挥主观能动性，用实践发展认识、用认识指导实践，以联系、发展、全面观点谋定后动、久久为功。",
    },
]


def build_coverage(base_md: str, roster: list[dict[str, str]], inventory: list[dict[str, str]]) -> tuple[list[dict[str, object]], list[dict[str, object]], list[dict[str, object]]]:
    covered_sources = source_lines(base_md)
    source_suite_counts = Counter(suite for suite, _q, _src in covered_sources)
    covered_keys = {(suite, q) for suite, q, _src in covered_sources}

    old_subject_missing: list[dict[str, object]] = []
    old_choice_missing: list[dict[str, object]] = []
    for row in inventory:
        suite = suite_name_from_inventory(row)
        q = qnorm(row.get("question", ""))
        row2: dict[str, object] = dict(row)
        row2["suite_name"] = suite
        row2["question_no_norm"] = q
        row2["covered_by_v6_question_number"] = (suite, q) in covered_keys
        if (suite, q) in covered_keys:
            continue
        nature = row.get("question_nature", "")
        grade = row.get("evidence_grade_initial", "")
        if nature == "subjective" and grade in {"A", "B"}:
            old_subject_missing.append(row2)
        elif nature == "choice" and grade == "C":
            old_choice_missing.append(row2)

    subj_counter = Counter(str(r["suite_name"]) for r in old_subject_missing)
    choice_counter = Counter(str(r["suite_name"]) for r in old_choice_missing)
    coverage_rows: list[dict[str, object]] = []
    for r in roster:
        suite = r["suite_name"]
        coverage_rows.append(
            {
                "suite_id": r.get("suite_id", ""),
                "suite_name": suite,
                "priority_bucket": r.get("priority_bucket", ""),
                "present_in_v6_student_md": source_suite_counts[suite] > 0,
                "v6_student_source_lines": source_suite_counts[suite],
                "old_v3_missing_subjective_rows_A_B": subj_counter[suite],
                "old_v3_missing_choice_rows_C": choice_counter[suite],
                "audit_status": (
                    "NO_SOURCE_LINE_IN_V6"
                    if source_suite_counts[suite] == 0
                    else "PARTIAL_SUBJECTIVE_GAP"
                    if subj_counter[suite]
                    else "CHOICE_GAP_ONLY"
                    if choice_counter[suite]
                    else "NO_OLD_GAP_BY_THIS_AUDIT"
                ),
                "notes": r.get("notes", ""),
            }
        )
    return coverage_rows, old_subject_missing, old_choice_missing


def old_subjective_md(rows: list[dict[str, object]]) -> str:
    lines = [
        "## 旧56套主观哲学漏题补丁",
        "",
        "说明：本段只收 v3 inventory 中证据等级为 A/B、且旧 v6 学生版没有同套同题号来源行的主观哲学点。E 级“待补证据”不写入学生版正文。",
        "",
    ]
    for idx, r in enumerate(rows, 1):
        suite = str(r["suite_name"])
        question = str(r.get("question", ""))
        trigger = str(r.get("trigger", "") or r.get("section", "哲学触发点"))
        lines.extend(
            [
                f"## [旧题补漏 {idx}] {trigger}",
                "",
                f"**来源题目**：{suite} {question}",
                "",
                f"**材料触发点**：{trim(str(r.get('material', '')), 260)}",
                "",
                f"**为什么能想到这个原理**：{trim(str(r.get('logic', '')), 360)}",
                "",
                f"**答案落点**：答题时先点出“{trigger}”，再扣住材料中的具体情境说明它为什么成立，最后回到设问要求，避免只背原理不落材料。",
                "",
            ]
        )
    return "\n".join(lines).rstrip() + "\n"


def second_mock_md(entries: list[dict[str, str]]) -> str:
    lines = [
        "## 2026二模正式补入条目",
        "",
        "说明：以下条目采用 Claude Code 独立审计与 Codex 回源核对后的 A 类题。综合题只收其中明确的必修四哲学维度。",
        "",
    ]
    for idx, r in enumerate(entries, 1):
        lines.extend(
            [
                f"## [2026二模补入 {idx}] {r['knowledge_point']}",
                "",
                f"**来源题目**：{r['suite_name']} 第{r['question_no']}题",
                "",
                f"**材料触发点**：{r['material_trigger']}",
                "",
                f"**为什么能想到这个原理**：{r['why']}",
                "",
                f"**答案落点**：{r['answer_landing']}",
                "",
            ]
        )
    return "\n".join(lines).rstrip() + "\n"


def choice_md(rows: list[dict[str, object]], claude_choices: list[dict[str, str]]) -> str:
    lines = [
        "## 哲学选择题补漏索引",
        "",
        "说明：选择题不全部扩写成主观题模板，先作为“看到什么想到什么”的速查索引。旧 56 套选择题补漏详见 CSV；下表先放 2026 二模 Claude Code 独立筛出的候选。",
        "",
        "| 套卷 | 题号 | 触发知识 | 证据要点 |",
        "|---|---:|---|---|",
    ]
    for r in claude_choices:
        lines.append(f"| {r['suite']} | {r['question_no']} | {r['knowledge']} | {r['evidence']} |")
    lines.extend(
        [
            "",
            f"旧 56 套另有 {len(rows)} 条 C 级选择题触发候选，已经写入 `old_missing_choice_rows_v7.csv`。这些条目只作为选择题触发和错肢识别补丁，不反推主观题给分链。",
            "",
        ]
    )
    return "\n".join(lines)


def parse_claude_choice_table(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    rows: list[dict[str, str]] = []
    for line in read_text(path).splitlines():
        if not line.startswith("| C"):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if len(cells) < 6:
            continue
        rows.append(
            {
                "id": cells[0],
                "suite": cells[1],
                "question_no": cells[2],
                "source": cells[3],
                "evidence": cells[4].replace('"', "'"),
                "knowledge": cells[5],
            }
        )
    return rows


def audit_report_md(coverage_rows: list[dict[str, object]], old_subject: list[dict[str, object]], old_choice: list[dict[str, object]], claude_choices: list[dict[str, str]]) -> str:
    no_source = [r for r in coverage_rows if r["audit_status"] == "NO_SOURCE_LINE_IN_V6"]
    partial = [r for r in coverage_rows if r["audit_status"] == "PARTIAL_SUBJECTIVE_GAP"]
    lines = [
        "# 必修四哲学宝典覆盖审计与补丁报告 v7",
        "",
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "## 结论",
        "",
        f"- 旧 v6 学生版出现来源题目的旧套卷数：{56 - len(no_source)} / 56。",
        f"- 旧 56 套完全未出现在哲学学生版的套卷：{len(no_source)} 套。",
        f"- 旧 56 套中，v3 有 A/B 级主观哲学点但旧 v6 没同套同题号来源行的条目：{len(old_subject)} 条。",
        f"- 旧 56 套中，v3 有 C 级哲学选择题触发但旧 v6 没同套同题号来源行的条目：{len(old_choice)} 条。",
        f"- 2026 二模 A 类正式补入主观/综合哲学条目：{len(SECOND_MOCK_A)} 条；二模选择题候选：{len(claude_choices)} 条。",
        "",
        "## 旧 v6 完全未出现的套卷",
        "",
        "| 套卷 | 优先级 | 旧v3主观漏项 | 旧v3选择漏项 |",
        "|---|---:|---:|---:|",
    ]
    for r in no_source:
        lines.append(f"| {r['suite_name']} | {r['priority_bucket']} | {r['old_v3_missing_subjective_rows_A_B']} | {r['old_v3_missing_choice_rows_C']} |")
    lines.extend(
        [
            "",
            "## 出现过但主观哲学点仍有漏项的套卷",
            "",
            "| 套卷 | 旧v6来源行 | 旧v3主观漏项 |",
            "|---|---:|---:|",
        ]
    )
    for r in partial:
        lines.append(f"| {r['suite_name']} | {r['v6_student_source_lines']} | {r['old_v3_missing_subjective_rows_A_B']} |")
    lines.extend(
        [
            "",
            "## 执行口径",
            "",
            "- 旧题补漏只自动收 A/B 级主观条目；E 级待补证据不进正文。",
            "- 二模按 Claude Code 独立审计分 A/B/C 类；A 类进学生版正文，B 类留在审计报告，C 类进选择题索引。",
            "- 选择题触发只作为选择题速记，不反推主观题评分链。",
        ]
    )
    return "\n".join(lines) + "\n"


def add_run(paragraph, text: str, bold: bool = False) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    add_run(p, text, bold=bold)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for key, value in [("top", "80"), ("bottom", "80"), ("start", "80"), ("end", "80")]:
        node = OxmlElement(f"w:{key}")
        node.set(qn("w:w"), value)
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(54)
    sec.bottom_margin = Pt(54)
    sec.left_margin = Pt(54)
    sec.right_margin = Pt(54)

    lines = read_text(md_path).splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            p = doc.add_heading(level=0)
            add_run(p, line[2:].strip(), bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            p = doc.add_heading(level=1)
            add_run(p, line[3:].strip(), bold=True)
        elif line.startswith("### "):
            p = doc.add_heading(level=2)
            add_run(p, line[4:].strip(), bold=True)
        elif line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|---"):
            header = [c.strip() for c in line.strip("|").split("|")]
            data: list[list[str]] = []
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                data.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            for idx, cell_text in enumerate(header):
                set_cell_text(table.rows[0].cells[idx], cell_text, bold=True)
            for row in data:
                cells = table.add_row().cells
                for idx, cell_text in enumerate(row[: len(header)]):
                    set_cell_text(cells[idx], cell_text)
            continue
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_run(p, line[2:].strip())
        else:
            p = doc.add_paragraph()
            if line.startswith("**") and "**：" in line:
                label, rest = line.split("**：", 1)
                add_run(p, label.replace("**", "") + "：", bold=True)
                add_run(p, rest)
            else:
                add_run(p, line.replace("**", ""))
        i += 1
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main() -> int:
    for p in [RUN, STUDENT, AUDIT, CSVOUT, LOGS]:
        p.mkdir(parents=True, exist_ok=True)

    base_md = read_text(OLD_V6)
    roster = read_csv(OLD_ROSTER)
    inventory = read_csv(V3_INVENTORY)
    coverage_rows, old_subject, old_choice = build_coverage(base_md, roster, inventory)
    claude_choices = parse_claude_choice_table(CLAUDE_REVIEW)

    write_csv(
        CSVOUT / "suite_coverage_audit_v7.csv",
        coverage_rows,
        [
            "suite_id",
            "suite_name",
            "priority_bucket",
            "present_in_v6_student_md",
            "v6_student_source_lines",
            "old_v3_missing_subjective_rows_A_B",
            "old_v3_missing_choice_rows_C",
            "audit_status",
            "notes",
        ],
    )
    write_csv(CSVOUT / "old_missing_subjective_rows_v7.csv", old_subject, list(old_subject[0].keys()) if old_subject else [])
    write_csv(CSVOUT / "old_missing_choice_rows_v7.csv", old_choice, list(old_choice[0].keys()) if old_choice else [])
    write_csv(CSVOUT / "second_mock_philosophy_entries_v7.csv", SECOND_MOCK_A, list(SECOND_MOCK_A[0].keys()))
    write_csv(CSVOUT / "second_mock_choice_candidates_v7.csv", claude_choices, ["id", "suite", "question_no", "source", "evidence", "knowledge"])

    report = audit_report_md(coverage_rows, old_subject, old_choice, claude_choices)
    write_text(AUDIT / "必修四哲学宝典覆盖审计与补丁报告_v7.md", report)
    write_text(RUN / "coverage_patch_report_v7.md", report)

    supplement = "\n\n".join(
        [
            "# 2026-05-23 覆盖审计与二模补丁 v7",
            report,
            old_subjective_md(old_subject),
            second_mock_md(SECOND_MOCK_A),
            choice_md(old_choice, claude_choices),
        ]
    )
    final_md = base_md.rstrip() + "\n\n---\n\n" + supplement
    final_md_path = STUDENT / "必修四哲学材料-知识触发框架_v7_覆盖审计与二模补丁.md"
    final_docx_path = STUDENT / "必修四哲学材料-知识触发框架_v7_覆盖审计与二模补丁.docx"
    write_text(final_md_path, final_md)
    markdown_to_docx(final_md_path, final_docx_path)

    old_student = old_subjective_md(old_subject).replace(
        "说明：本段只收 v3 inventory 中证据等级为 A/B、且旧 v6 学生版没有同套同题号来源行的主观哲学点。E 级“待补证据”不写入学生版正文。",
        "说明：以下为旧卷补入的主观哲学触发点，写法按学生可直接复习的材料触发链整理。",
    )
    second_student = second_mock_md(SECOND_MOCK_A).replace(
        "说明：以下条目采用 Claude Code 独立审计与 Codex 回源核对后的 A 类题。综合题只收其中明确的必修四哲学维度。",
        "说明：以下为 2026 二模中可直接使用的哲学触发条目。综合题只收其中明确的必修四哲学维度。",
    )
    choice_student = choice_md(old_choice, claude_choices).replace(
        "说明：选择题不全部扩写成主观题模板，先作为“看到什么想到什么”的速查索引。旧 56 套选择题补漏详见 CSV；下表先放 2026 二模 Claude Code 独立筛出的候选。",
        "说明：选择题先作为“看到什么想到什么”的速查索引。下表放 2026 二模哲学选择题候选。",
    )
    choice_student = re.sub(r"\n旧 56 套另有 .*?不反推主观题给分链。\n", "\n", choice_student)

    pure_student = "\n\n".join(
        [
            "# 2026-05-23 漏题与二模补丁 v7",
            "说明：以下为可直接给学生看的补丁正文。",
            old_student,
            second_student,
            choice_student,
        ]
    )
    pure_md = base_md.rstrip() + "\n\n---\n\n" + pure_student
    pure_md_path = STUDENT / "必修四哲学材料-知识触发框架_v7_纯学生版.md"
    pure_docx_path = STUDENT / "必修四哲学材料-知识触发框架_v7_纯学生版.docx"
    write_text(pure_md_path, pure_md)
    markdown_to_docx(pure_md_path, pure_docx_path)

    if CLAUDE_REVIEW.exists():
        shutil.copy2(CLAUDE_REVIEW, AUDIT / CLAUDE_REVIEW.name)
    shutil.copy2(OLD_V6, LOGS / "base_v6_used.md")

    build_report = "\n".join(
        [
            "# Build Report",
            "",
            f"- generated_at: {datetime.now().isoformat(timespec='seconds')}",
            f"- old_v6: {OLD_V6}",
            f"- final_md: {final_md_path}",
            f"- final_docx: {final_docx_path}",
            f"- pure_student_md: {pure_md_path}",
            f"- pure_student_docx: {pure_docx_path}",
            f"- old_missing_subjective_rows_A_B: {len(old_subject)}",
            f"- old_missing_choice_rows_C: {len(old_choice)}",
            f"- second_mock_A_entries: {len(SECOND_MOCK_A)}",
            f"- second_mock_choice_candidates: {len(claude_choices)}",
        ]
    )
    write_text(DESKTOP / "BUILD_REPORT.md", build_report)
    write_text(RUN / "BUILD_REPORT.md", build_report)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
