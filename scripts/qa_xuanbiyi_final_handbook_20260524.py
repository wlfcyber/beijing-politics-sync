# -*- coding: utf-8 -*-
from __future__ import annotations

import re
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
BASE = next((ROOT / "reports").glob("*2026-05-16/06_final_handbook"))
STUDENT = BASE / "选必一_当代国际政治与经济_主观题术语宝典_学生版.md"
NAV = BASE / "选必一_当代国际政治与经济_主观题术语宝典_考前导航版.md"
STUDENT_DOCX = BASE / "选必一_当代国际政治与经济_主观题术语宝典_学生版.docx"
NAV_DOCX = BASE / "选必一_当代国际政治与经济_主观题术语宝典_考前导航版.docx"


def main() -> None:
    text = STUDENT.read_text(encoding="utf-8")
    nav_text = NAV.read_text(encoding="utf-8")
    appendix_marker = "# 附：模块边界 / 跨书提示"
    if appendix_marker not in text:
        raise SystemExit("missing appendix marker")

    core_matches = list(re.finditer(r"^## 核心答题点：(.+?)（出现(\d+)次）\s*$", text, re.M))
    count_mismatches: list[tuple[str, int, int]] = []
    for i, match in enumerate(core_matches):
        start = match.end()
        next_core = core_matches[i + 1].start() if i + 1 < len(core_matches) else len(text)
        app = text.find("\n" + appendix_marker, start)
        end = next_core if app == -1 else min(next_core, app)
        body = text[start:end]
        actual = len(re.findall(r"^###\s+", body, re.M))
        listed = int(match.group(2))
        if actual != listed:
            count_mismatches.append((match.group(1), listed, actual))

    main_part, appendix = text.split(appendix_marker, 1)
    main_cases = len(re.findall(r"^###\s+", main_part, re.M))
    boundary_cases = len(re.findall(r"^###\s+", appendix, re.M))

    fields = ["【什么时候写】", "【设问】", "【为什么能想到】", "【卷面句】", "【同题组】"]
    field_counts = {field: text.count(field) for field in fields}
    missing_field_delta = sorted(set(field_counts.values()))

    merged_title_flags: list[str] = []
    for raw_title in re.findall(r"^###\s+(.+?)\s*$", text, re.M):
        title = re.sub(r"^\d+\.\s*", "", raw_title)
        if re.search(r"20\d{2}[^\n]{0,20}Q\d+.*[/／].*20\d{2}[^\n]{0,20}Q\d+", title):
            merged_title_flags.append(title)

    nav_rows = [
        line
        for line in nav_text.splitlines()
        if line.startswith("| ")
        and not line.startswith("|---")
        and not line.startswith("| 核心答题点")
    ]
    nav_boundary_rows = [line for line in nav_rows if "边界提示：" in line]
    nav_core_rows = [line for line in nav_rows if "边界提示：" not in line]

    coverage_needles = {
        "2025东城二模共同利益": "2025东城二模 第20题（同球共济精神）",
        "2025朝阳二模共同利益": "2025朝阳二模 第21题（周边工作新局面）",
        "2026西城一模共同利益": "2026西城一模 第20题第(2)问（中国—东盟自贸区3.0版）",
        "2026石景山期末AI治理": "2026石景山期末 第19题第(2)问（人工智能全球治理）",
        "2024顺义二模Q18不纳入": "2024顺义二模Q18",
    }
    coverage_counts = {name: text.count(needle) for name, needle in coverage_needles.items()}

    print(f"student_md={STUDENT}")
    print(f"nav_md={NAV}")
    print(f"core_headings={len(core_matches)}")
    print(f"main_cases={main_cases}")
    print(f"boundary_cases={boundary_cases}")
    print(f"total_h3={main_cases + boundary_cases}")
    print(f"count_mismatches={len(count_mismatches)}")
    if count_mismatches:
        print(f"count_mismatch_examples={count_mismatches[:10]}")
    print(f"field_counts={field_counts}")
    print(f"field_count_values={missing_field_delta}")
    print(f"nav_core_rows={len(nav_core_rows)}")
    print(f"nav_boundary_rows={len(nav_boundary_rows)}")
    print(f"nav_total_rows={len(nav_rows)}")
    print(f"merged_title_flags={len(merged_title_flags)}")
    if merged_title_flags:
        print(f"merged_title_examples={merged_title_flags[:10]}")
    print(f"coverage_counts={coverage_counts}")

    try:
        from docx import Document

        for label, path in [("student_docx", STUDENT_DOCX), ("nav_docx", NAV_DOCX)]:
            doc = Document(path)
            styles = [p.style.name for p in doc.paragraphs if p.style is not None]
            print(
                f"{label}=size:{path.stat().st_size};"
                f"paragraphs:{len(doc.paragraphs)};"
                f"h1:{styles.count('Heading 1')};"
                f"h2:{styles.count('Heading 2')};"
                f"h3:{styles.count('Heading 3')};"
                f"tables:{len(doc.tables)};"
                f"table_rows:{sum(len(table.rows) for table in doc.tables)}"
            )
    except Exception as exc:
        print(f"docx_structural_check=skipped:{exc}")


if __name__ == "__main__":
    main()
