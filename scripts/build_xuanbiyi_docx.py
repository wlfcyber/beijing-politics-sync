from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path.cwd()
FINAL_DIR = next(ROOT.rglob("06_final_handbook"))

FILES = [
    (
        FINAL_DIR / "选必一_当代国际政治与经济_主观题术语宝典_学生版.md",
        FINAL_DIR / "选必一_当代国际政治与经济_主观题术语宝典_学生版.docx",
    ),
    (
        FINAL_DIR / "选必一_当代国际政治与经济_主观题术语宝典_考前导航版.md",
        FINAL_DIR / "选必一_当代国际政治与经济_主观题术语宝典_考前导航版.docx",
    ),
]


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def clean_inline(text: str) -> tuple[str, bool]:
    text = text.strip()
    bold = False
    if text.startswith("**") and text.endswith("**") and len(text) > 4:
        text = text[2:-2]
        bold = True
    text = text.replace("`", "")
    return text, bold


def add_paragraph(doc: Document, text: str, style: str | None = None, bullet: bool = False) -> None:
    text, bold = clean_inline(text)
    if not text:
        return
    p = doc.add_paragraph(style="List Bullet" if bullet else style)
    if text.startswith("【") and "】" in text:
        label, rest = text.split("】", 1)
        r = p.add_run(label + "】")
        r.bold = True
        set_east_asia_font(r, "黑体")
        r.font.size = Pt(10.5)
        if rest:
            r2 = p.add_run(rest)
            set_east_asia_font(r2, "宋体")
            r2.font.size = Pt(10.5)
    else:
        r = p.add_run(text)
        r.bold = bold
        set_east_asia_font(r, "宋体")
        r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    set_east_asia_font(r, "黑体")
    r.bold = True
    if level == 0:
        r.font.size = Pt(20)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(31, 78, 121)
    elif level == 2:
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(47, 84, 150)
    else:
        r.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(8 if level <= 2 else 4)
    p.paragraph_format.space_after = Pt(5)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if i == 0:
                set_cell_shading(cell, "D9EAF7")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 1 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value.strip())
            set_east_asia_font(run, "宋体")
            run.font.size = Pt(8.8 if len(value) > 45 else 9.5)
            if i == 0:
                run.bold = True
    doc.add_paragraph()


def is_sep(row: str) -> bool:
    return bool(re.fullmatch(r"\|[\s:\-|\u2014]+\|", row.strip()))


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int] | None:
    if start + 1 >= len(lines) or not lines[start].strip().startswith("|") or not is_sep(lines[start + 1]):
        return None
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        if not is_sep(lines[i]):
            parts = [x.strip() for x in lines[i].strip().strip("|").split("|")]
            rows.append(parts)
        i += 1
    return rows, i


def markdown_to_docx(md_path: Path, out_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.75)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        table_result = parse_table(lines, i)
        if table_result:
            rows, i = table_result
            add_table(doc, rows)
            continue
        if line.startswith("# "):
            heading = line[2:].strip()
            if i != 0 and heading in {"理论", "经济全球化", "政治多极化", "中国", "联合国", "附：总说句 / 兜底加分表达", "附：模块边界 / 跨书提示"}:
                doc.add_section(WD_SECTION.NEW_PAGE)
            add_heading(doc, heading, 1 if i != 0 else 0)
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3)
        elif line.startswith("- "):
            add_paragraph(doc, line[2:].strip(), bullet=True)
        else:
            add_paragraph(doc, line)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    for md_path, out_path in FILES:
        markdown_to_docx(md_path, out_path)
        print(out_path)


if __name__ == "__main__":
    main()
