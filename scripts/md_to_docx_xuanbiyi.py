from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ACCENT = RGBColor(31, 78, 121)
LIGHT = "D9EAF7"
PALE = "EEF6FB"


def set_east_asia_font(run, font: str) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_bottom_border(paragraph, color: str = "4F81BD") -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def setup_styles(doc: Document, nav: bool = False) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(4)

    for name, size, color in [
        ("Heading 1", 17, ACCENT),
        ("Heading 2", 14, ACCENT),
        ("Heading 3", 11, RGBColor(46, 46, 46)),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = not nav


def add_label_para(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{label} ")
    set_east_asia_font(r, "Microsoft YaHei")
    r.bold = True
    r.font.color.rgb = ACCENT
    r2 = p.add_run(text)
    set_east_asia_font(r2, "Microsoft YaHei")


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.65)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_east_asia_font(r, "Microsoft YaHei")
    r.font.size = Pt(9.3)


def add_table_row(table, cells: list[str], nav: bool = False, header: bool = False) -> None:
    row = table.add_row()
    for idx, cell_text in enumerate(cells):
        cell = row.cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(cell_text)
        set_east_asia_font(r, "Microsoft YaHei")
        r.font.size = Pt(8 if nav else 8.5)
        if header:
            r.bold = True
            set_cell_shading(cell, LIGHT)


def add_title_page(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(96)
    r = p.add_run(title)
    set_east_asia_font(r, "Microsoft YaHei")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = ACCENT

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(12)
    r2 = p2.add_run(subtitle)
    set_east_asia_font(r2, "Microsoft YaHei")
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor(90, 90, 90)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(24)
    r3 = p3.add_run("飞哥正志讲堂")
    set_east_asia_font(r3, "Microsoft YaHei")
    r3.font.size = Pt(11)
    r3.bold = True

    doc.add_section(WD_SECTION.NEW_PAGE)


def add_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p._element.clear_content()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("选必一《当代国际政治与经济》主观题术语宝典")
    set_east_asia_font(run, "Microsoft YaHei")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(120, 120, 120)


def configure_section(section, nav: bool = False) -> None:
    if nav:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)
        section.top_margin = Cm(1.25)
        section.bottom_margin = Cm(1.15)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)
    else:
        section.top_margin = Cm(1.75)
        section.bottom_margin = Cm(1.65)
        section.left_margin = Cm(1.75)
        section.right_margin = Cm(1.75)


def build_docx(md_path: Path, docx_path: Path, nav: bool = False) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    sec = doc.sections[0]
    configure_section(sec, nav)
    add_footer(sec)
    setup_styles(doc, nav)

    title = "选择性必修一《当代国际政治与经济》主观题术语宝典"
    subtitle = "学生版" if not nav else "考前导航版"
    add_title_page(doc, title, subtitle)

    in_table = False
    table = None
    table_header: list[str] | None = None
    nav_table_rows = 0

    for raw in lines:
        line = raw.strip()
        if not line:
            in_table = False
            table_header = None
            nav_table_rows = 0
            continue
        if line.startswith("# "):
            text = line[2:].strip()
            if text == title or text.startswith("选必一《当代国际政治与经济》主观题术语宝典"):
                continue
            p = doc.add_heading(text, level=1)
            add_bottom_border(p)
            continue
        if line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            if re.match(r"\d+\.\s", line[4:].strip()):
                shade_paragraph(p, PALE)
            continue
        if line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line.strip("*"))
            set_east_asia_font(r, "Microsoft YaHei")
            r.bold = True
            r.font.size = Pt(11)
            continue
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(re.fullmatch(r":?-+:?", c) for c in cells):
                continue
            if not in_table:
                table = doc.add_table(rows=0, cols=len(cells))
                table.style = "Table Grid"
                in_table = True
                nav_table_rows = 0
                if nav and table_header is not None and cells != table_header:
                    add_table_row(table, table_header, nav, header=True)
            is_header = len(table.rows) == 0
            add_table_row(table, cells, nav, header=is_header)
            if is_header:
                table_header = cells[:]
            elif nav:
                nav_table_rows += 1
                if nav_table_rows >= 10:
                    in_table = False
            continue
        in_table = False
        table_header = None
        nav_table_rows = 0
        if line.startswith("- "):
            add_bullet(doc, line[2:].strip())
            continue
        m = re.match(r"^(【[^】]+】)(.*)$", line)
        if m:
            add_label_para(doc, m.group(1), m.group(2).strip())
            continue
        p = doc.add_paragraph()
        r = p.add_run(line)
        set_east_asia_font(r, "Microsoft YaHei")

    for section in doc.sections:
        configure_section(section, nav)
        add_footer(section)
    doc.save(docx_path)


def main() -> None:
    if len(sys.argv) < 3:
        raise SystemExit("usage: md_to_docx_xuanbiyi.py input.md output.docx [--nav]")
    md = Path(sys.argv[1])
    out = Path(sys.argv[2])
    out.parent.mkdir(parents=True, exist_ok=True)
    build_docx(md, out, "--nav" in sys.argv)


if __name__ == "__main__":
    main()
