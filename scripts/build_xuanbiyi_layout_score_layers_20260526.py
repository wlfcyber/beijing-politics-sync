from __future__ import annotations

import csv
import json
import re
import subprocess
import sys
import zipfile
from collections import Counter, defaultdict
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer


ROOT = Path("/Users/wanglifei/GaokaoPolitics/beijing-politics-sync")
RUN = ROOT / "reports/选必一_哲学宝典式重建_2026-05-16"
SOURCE_DIR = RUN / "17_clean_student_opus47_review_20260526"
SOURCE_MD = SOURCE_DIR / "选必一_当代国际政治与经济_主观题术语宝典_学生版_FINAL_纯净送审稿_20260526.md"
OUT_DIR = RUN / "18_layout_score_layers_20260526"
OUT_MD = OUT_DIR / "选必一_当代国际政治与经济_主观题术语宝典_学生版_FINAL_排版分值层次优化版_20260526.md"
OUT_DOCX = OUT_DIR / "选必一_当代国际政治与经济_主观题术语宝典_学生版_FINAL_排版分值层次优化版_20260526.docx"
OUT_PDF = OUT_DIR / "选必一_当代国际政治与经济_主观题术语宝典_学生版_FINAL_排版分值层次优化版_20260526.pdf"
OUT_ZIP = OUT_DIR / "xuanbiyi_layout_score_layers_final_delivery_20260526.zip"
QA = OUT_DIR / "LAYOUT_SCORE_LAYER_QA_20260526.md"
TRACE_CSV = OUT_DIR / "SCORE_LAYER_TRACE_20260526.csv"
UNSCORED_CSV = OUT_DIR / "UNSCORED_SAME_GROUP_BACKLOG_20260526.csv"

BUCKET_ORDER = ["时代背景", "理论", "经济全球化", "政治多极化", "中国", "联合国"]
ACCENT = RGBColor(31, 78, 121)
MUTED = RGBColor(92, 92, 92)
LIGHT = "EAF3F8"
PALE = "F4F9FC"
FONT = "Microsoft YaHei"
PDF_FONT_PATH = "/System/Library/Fonts/STHeiti Light.ttc"
PDF_FONT = "STHeitiLight"

CSV_SOURCES = [
    RUN / "14_full_source_recrawl_20260525/FULL_SOURCE_XUANBIYI_CANDIDATE_SCAN_20260525.csv",
    RUN / "12_full_desktop_extract_20260524/xuanbiyi_full_source_question_candidates.csv",
]

HARD_LAYER_NOTES = {
    ("2026通州期中", "Q20", None): (
        "8分",
        "六层：1.共商共建共享全球治理观；2.时代主题、经济全球化、顺应各国人民愿望等；3.符合《联合国宪章》；4.新秩序、国际关系民主化、多边主义、正确义利观、兼顾利益等任一；5.人类命运共同体；6.中国智慧、中国方案、大国担当。",
    ),
    ("2026通州期末", "Q20", None): (
        "8分",
        "六层：1.共商共建共享全球治理观；2.时代主题、经济全球化、顺应各国人民愿望等；3.符合《联合国宪章》；4.新秩序、国际关系民主化、多边主义、正确义利观、兼顾利益等任一；5.人类命运共同体；6.中国智慧、中国方案、大国担当。",
    ),
    ("2026朝阳期中", "Q17", None): (
        "8分",
        "总分结构：一层3分，两层6分，三层8分；每层按“总述—材料—关系”展开，不能只罗列材料或空写教材词。",
    ),
    ("2025海淀期末", "Q22", None): (
        "9分",
        "短文题：形式3分（标题、分段、层次），知识4分（知识+材料两组），综合与论述2分（至少两个模块、论述充分）。",
    ),
    ("2026朝阳二模", "Q20", "2"): (
        "6分",
        "第(2)问6分；先接第(1)问关于国际关系处理方式的差异，再说明人类命运共同体理念为什么能维护世界和平与发展。",
    ),
}


def set_east_asia_font(run, font: str = FONT) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_bottom_border(paragraph, color: str = "4F81BD") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "7")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def title_parts(title_line: str) -> tuple[str | None, str | None, str | None]:
    m = re.match(r"^### \d+\.\s*(\d{4}[^Q\n]+?)Q(\d+)(?:\((\d+)\))?", title_line.strip())
    if not m:
        return None, None, None
    return m.group(1).strip(), "Q" + m.group(2), m.group(3)


def district_key(suite: str) -> str:
    for district in ["东城", "西城", "海淀", "朝阳", "丰台", "石景山", "门头沟", "房山", "通州", "顺义", "昌平", "延庆"]:
        if district in suite:
            year = re.search(r"\d{4}", suite)
            return (year.group(0) if year else "") + district
    return suite


def score_from_text_for_question(text: str, q: str) -> list[int]:
    qnum = re.sub(r"\D", "", q)
    if not qnum:
        return []
    patterns = [
        rf"{qnum}\s*[\.．、]?\s*[（(]\s*(\d{{1,2}})\s*分\s*[）)]",
        rf"{qnum}\s*题\s*[（(]?\s*(\d{{1,2}})\s*分",
        rf"第\s*{qnum}\s*题\s*[（(]?\s*(\d{{1,2}})\s*分",
    ]
    vals: list[int] = []
    for pat in patterns:
        vals.extend(int(x) for x in re.findall(pat, text))
    return [v for v in vals if 1 <= v <= 16]


def load_score_lookup() -> tuple[dict[tuple[str, str], str], dict[tuple[str, str], str]]:
    exact: dict[tuple[str, str], str] = {}
    fuzzy_values: dict[tuple[str, str], list[int]] = defaultdict(list)
    for csv_path in CSV_SOURCES:
        if not csv_path.exists():
            continue
        with csv_path.open(encoding="utf-8-sig", newline="") as f:
            reader = csv.DictReader(f)
            for row in reader:
                suite = (row.get("suite_guess") or row.get("suite") or "").strip()
                q = (row.get("question") or "").strip()
                if not suite or not q:
                    continue
                ctx = " ".join(row.get(k, "") for k in ["paper_context", "rubric_context", "context", "snippet"])
                vals = score_from_text_for_question(ctx, q)
                if not vals:
                    continue
                counts = Counter(vals)
                score = counts.most_common(1)[0][0]
                exact[(suite, q)] = f"{score}分"
                fuzzy_values[(district_key(suite), q)].extend(vals)
    fuzzy: dict[tuple[str, str], str] = {}
    for key, vals in fuzzy_values.items():
        counts = Counter(vals)
        if counts:
            fuzzy[key] = f"{counts.most_common(1)[0][0]}分"
    return exact, fuzzy


def score_from_prompt(question: str, part: str | None = None) -> str | None:
    nums = [int(x) for x in re.findall(r"[（(](\d{1,2})\s*分[）)]", question)]
    nums = [n for n in nums if 1 <= n <= 16]
    if not nums:
        return None
    if part and len(nums) > 1:
        part_pat = rf"[（(]{re.escape(part)}[）)].{{0,80}}?[（(](\d{{1,2}})\s*分[）)]"
        m = re.search(part_pat, question)
        if m:
            return f"{int(m.group(1))}分"
    if len(nums) == 1:
        return f"{nums[0]}分"
    if sum(nums) <= 16:
        return f"{sum(nums)}分"
    return f"{nums[-1]}分"


def norm_with_map(text: str) -> tuple[str, list[int]]:
    chars: list[str] = []
    raw_indexes: list[int] = []
    for idx, ch in enumerate(text):
        if re.match(r"[\w\u4e00-\u9fff]", ch):
            chars.append(ch.lower())
            raw_indexes.append(idx)
    return "".join(chars), raw_indexes


def load_source_contexts() -> list[dict[str, str]]:
    contexts: list[dict[str, str]] = []
    for csv_path in CSV_SOURCES:
        if not csv_path.exists():
            continue
        with csv_path.open(encoding="utf-8-sig", newline="") as f:
            reader = csv.DictReader(f)
            for row in reader:
                q = (row.get("question") or "").strip()
                if not q:
                    continue
                context = " ".join(row.get(k, "") for k in ["paper_context", "rubric_context", "context", "snippet"])
                if context.strip():
                    contexts.append({"question": q, "context": context})
    return contexts


def source_score_from_context(context: str, question: str, q: str, part: str | None) -> tuple[str | None, str | None]:
    normalized_question, _ = norm_with_map(question)
    if len(normalized_question) < 18:
        return None, None
    normalized_context, raw_map = norm_with_map(context)
    if not normalized_context or not raw_map:
        return None, None

    pos = normalized_context.find(normalized_question)
    if pos < 0:
        head = normalized_question[: min(len(normalized_question), 64)]
        pos = normalized_context.find(head)
        if pos < 0:
            return None, None
        tail = normalized_question[-28:]
        if tail and tail not in normalized_context[max(0, pos - 120) : pos + len(normalized_question) + 420]:
            return None, None

    raw_start = raw_map[pos]
    raw_end = raw_map[min(len(raw_map) - 1, pos + min(len(normalized_question), 96) - 1)]
    full_pos = normalized_context.find(normalized_question)
    if full_pos >= 0:
        raw_start = raw_map[full_pos]
        raw_end = raw_map[min(len(raw_map) - 1, full_pos + len(normalized_question) - 1)]

    after = context[raw_end : min(len(context), raw_end + 120)]
    before = context[max(0, raw_start - 160) : raw_start]
    qnum = re.sub(r"\D", "", q)

    if part:
        # For subquestions, only trust a score immediately following the matched subquestion text.
        m = re.search(r"[（(](\d{1,2})\s*分[）)]", after[:90])
        if m:
            return f"{int(m.group(1))}分", "source_after_subquestion"
        return None, None

    m = re.search(r"[（(](\d{1,2})\s*分[）)]", after[:90])
    if m:
        return f"{int(m.group(1))}分", "source_after_question"
    if qnum:
        for pat in [
            rf"{qnum}\s*[\.．、]?\s*[（(]\s*(\d{{1,2}})\s*分\s*[）)]",
            rf"{qnum}\s*题\s*[（(]?\s*(\d{{1,2}})\s*分",
        ]:
            matches = re.findall(pat, before[-120:])
            if matches:
                return f"{int(matches[-1])}分", "source_q_heading"
    return None, None


def source_score_for_question(
    source_contexts: list[dict[str, str]], question: str, q: str | None, part: str | None
) -> tuple[str | None, str | None]:
    if not q:
        return None, None
    found: list[tuple[str, str]] = []
    for row in source_contexts:
        if row["question"] != q:
            continue
        score, method = source_score_from_context(row["context"], question, q, part)
        if score and method:
            found.append((score, method))
    if not found:
        return None, None
    scores = {score for score, _ in found}
    if len(scores) == 1:
        return found[0]
    # Conflicting scores usually indicate adjacent-question bleed; do not annotate.
    return None, None


def split_examples(md: str) -> list[tuple[int, int, str]]:
    matches = list(re.finditer(r"(?m)^### \d+\. .*$", md))
    out: list[tuple[int, int, str]] = []
    for idx, match in enumerate(matches):
        start = match.start()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(md)
        out.append((start, end, md[start:end]))
    return out


def layer_from_prompt(question: str) -> str | None:
    parts = re.findall(r"[（(]([一二三四五六七八九十\d]+)[）)].{0,100}?[（(](\d{1,2})\s*分[）)]", question)
    if len(parts) >= 2:
        rendered = "；".join(f"第({idx})问{score}分" for idx, score in parts)
        return f"按设问分层：{rendered}。"
    return None


def parse_same_group_items(text: str) -> dict[str, list[str]]:
    grouped: dict[str, list[str]] = {name: [] for name in BUCKET_ORDER}
    grouped["其他"] = []
    for raw in [x.strip(" ；") for x in text.split("；") if x.strip(" ；")]:
        raw = raw.rstrip("。.")
        m = re.match(r"(.+?)（([^（）]+桶)）$", raw)
        if m:
            item, bucket = m.group(1).strip(), m.group(2).replace("桶", "")
            grouped.setdefault(bucket, []).append(item)
        else:
            grouped["其他"].append(raw)
    return grouped


def format_same_group(rest: str, score: str | None) -> str:
    grouped = parse_same_group_items(rest)
    label = f"【同题组】（本题{score}）" if score else "【同题组】"
    lines = [label]
    for bucket in BUCKET_ORDER + ["其他"]:
        items = grouped.get(bucket) or []
        if items:
            lines.append(f"- {bucket}：{'；'.join(items)}")
    return "\n".join(lines)


def enhance_markdown(md: str) -> tuple[str, dict[str, int], list[dict[str, str]]]:
    source_contexts = load_source_contexts()
    parts: list[str] = []
    cursor = 0
    trace_rows: list[dict[str, str]] = []
    stats = {
        "examples": 0,
        "same_group_blocks": 0,
        "score_annotated_same_group": 0,
        "score_from_prompt_or_hard": 0,
        "score_from_source": 0,
        "layer_notes_added": 0,
        "same_group_bucketized": 0,
        "unscored_same_group": 0,
        "same_group_scored_examples": 0,
        "same_group_unscored_examples": 0,
    }
    for start, end, block in split_examples(md):
        parts.append(md[cursor:start])
        cursor = end
        stats["examples"] += 1
        lines = block.rstrip("\n").splitlines()
        title = lines[0].lstrip("# ").strip()
        suite, q, part = title_parts(lines[0])
        question = ""
        for line in lines:
            if line.startswith("【设问】"):
                question = line.replace("【设问】", "", 1).strip()
                break
        score = score_from_prompt(question, part)
        score_method = "prompt" if score else None
        hard = HARD_LAYER_NOTES.get((suite, q, part)) or HARD_LAYER_NOTES.get((suite, q, None))
        layer_note = layer_from_prompt(question)
        if hard:
            score = hard[0]
            score_method = "hard"
            layer_note = hard[1]
        if not score:
            source_score, method = source_score_for_question(source_contexts, question, q, part)
            if source_score:
                score = source_score
                score_method = method
        if score:
            if score_method in {"prompt", "hard"}:
                stats["score_from_prompt_or_hard"] += 1
            elif score_method:
                stats["score_from_source"] += 1
        has_layer = any(line.startswith("【得分层次】") for line in lines)
        original_has_layer = has_layer
        has_same_group = False
        new_lines: list[str] = []
        for line in lines:
            if line.startswith("【同题组】"):
                has_same_group = True
                stats["same_group_blocks"] += 1
                rest = line.replace("【同题组】", "", 1).strip()
                new_lines.append(format_same_group(rest, score))
                stats["same_group_bucketized"] += 1
                if score:
                    stats["score_annotated_same_group"] += 1
                else:
                    stats["unscored_same_group"] += 1
                continue
            new_lines.append(line)
            if line.startswith("【设问】") and layer_note and not has_layer:
                new_lines.append(f"【得分层次】{layer_note}")
                stats["layer_notes_added"] += 1
                has_layer = True
        trace_rows.append(
            {
                "title": title,
                "suite": suite or "",
                "question": q or "",
                "part": part or "",
                "same_group": "YES" if has_same_group else "NO",
                "score": score or "",
                "score_method": score_method or "",
                "layer": "YES" if has_layer else "NO",
                "layer_method": "preexisting" if original_has_layer else ("added" if layer_note else ""),
                "question_text": question,
            }
        )
        if has_same_group:
            if score:
                stats["same_group_scored_examples"] += 1
            else:
                stats["same_group_unscored_examples"] += 1
        parts.append("\n".join(new_lines).rstrip() + "\n\n")
    parts.append(md[cursor:])
    text = "".join(parts)
    text = re.sub(r"\n{3,}", "\n\n", text).strip() + "\n"
    return text, stats, trace_rows


def setup_docx_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.2)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(4)
    for name, size, color in [
        ("Heading 1", 17, ACCENT),
        ("Heading 2", 13.2, ACCENT),
        ("Heading 3", 10.8, RGBColor(45, 45, 45)),
    ]:
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(9)
        style.paragraph_format.space_after = Pt(5)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(82)
    r = p.add_run("选择性必修一《当代国际政治与经济》主观题术语宝典")
    set_east_asia_font(r)
    r.bold = True
    r.font.size = Pt(25)
    r.font.color.rgb = ACCENT
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(12)
    r2 = p2.add_run("学生厚版 · 排版分值层次优化版")
    set_east_asia_font(r2)
    r2.font.size = Pt(13)
    r2.font.color.rgb = MUTED
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(24)
    r3 = p3.add_run("飞哥正志讲堂")
    set_east_asia_font(r3)
    r3.bold = True
    r3.font.size = Pt(11)
    doc.add_section(WD_SECTION.NEW_PAGE)


def configure_docx_section(section) -> None:
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    footer = section.footer
    p = footer.paragraphs[0]
    p._element.clear_content()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("选必一《当代国际政治与经济》主观题术语宝典")
    set_east_asia_font(r)
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(120, 120, 120)


def add_label_paragraph(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2 if label in {"【为什么能想到】", "【同题组】"} else 0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.14
    if text:
        p.paragraph_format.left_indent = Cm(0.18)
        p.paragraph_format.first_line_indent = Cm(-0.18)
    r = p.add_run(label)
    set_east_asia_font(r)
    r.bold = True
    r.font.color.rgb = ACCENT
    r.font.size = Pt(10.3)
    if text:
        r2 = p.add_run(" " + text)
        set_east_asia_font(r2)
        r2.font.size = Pt(10.1)


def add_bullet(doc: Document, text: str, context: str | None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.66)
    p.paragraph_format.first_line_indent = Cm(-0.22)
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = 1.13
    size = 10.0 if context == "why" else 9.8
    m = re.match(r"^([^：:]{2,8})[：:](.+)$", text)
    if m:
        r = p.add_run(m.group(1) + "：")
        set_east_asia_font(r)
        r.bold = True
        r.font.size = Pt(size)
        r.font.color.rgb = ACCENT if context == "why" else MUTED
        r2 = p.add_run(m.group(2).strip())
        set_east_asia_font(r2)
        r2.font.size = Pt(size)
    else:
        r = p.add_run(text)
        set_east_asia_font(r)
        r.font.size = Pt(size)


def build_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    configure_docx_section(doc.sections[0])
    setup_docx_styles(doc)
    add_title_page(doc)
    current_context: str | None = None
    title = "选择性必修一《当代国际政治与经济》主观题术语宝典"
    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            current_context = None if current_context == "same" else current_context
            continue
        if line.strip("*").strip() == "飞哥正志讲堂":
            continue
        if line.startswith("# "):
            text = line[2:].strip()
            if text == title or ("当代国际政治与经济" in text and "主观题术语宝典" in text):
                continue
            p = doc.add_heading(text, level=1)
            add_bottom_border(p)
            current_context = None
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            current_context = None
            continue
        if line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            if re.match(r"\d+\.\s", line[4:].strip()):
                shade_paragraph(p, PALE)
            current_context = None
            continue
        if line.startswith("- "):
            add_bullet(doc, line[2:].strip(), current_context)
            continue
        m = re.match(r"^(【[^】]+】)(.*)$", line)
        if m:
            label = m.group(1)
            rest = m.group(2).strip()
            add_label_paragraph(doc, label, rest)
            if label == "【为什么能想到】":
                current_context = "why"
            elif label == "【同题组】":
                current_context = "same"
            else:
                current_context = None
            continue
        p = doc.add_paragraph()
        r = p.add_run(line)
        set_east_asia_font(r)
        r.font.size = Pt(10.1)
        current_context = None
    for section in doc.sections:
        configure_docx_section(section)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def clean_pdf_inline(text: str) -> str:
    return escape(text).replace("**", "").replace("`", "")


def build_pdf(md_path: Path, pdf_path: Path) -> None:
    pdfmetrics.registerFont(TTFont(PDF_FONT, PDF_FONT_PATH))
    base = ParagraphStyle(
        "base",
        fontName=PDF_FONT,
        fontSize=9.8,
        leading=14.2,
        alignment=TA_LEFT,
        wordWrap="CJK",
        spaceAfter=3,
    )
    title = ParagraphStyle("title", parent=base, fontSize=18, leading=24, alignment=TA_CENTER, spaceAfter=8)
    h1 = ParagraphStyle("h1", parent=base, fontSize=15, leading=20, textColor=colors.HexColor("#111111"), spaceBefore=8, spaceAfter=6)
    h2 = ParagraphStyle("h2", parent=base, fontSize=12.2, leading=16, textColor=colors.HexColor("#1F4E79"), spaceBefore=6, spaceAfter=4)
    h3 = ParagraphStyle("h3", parent=base, fontSize=10.6, leading=14.6, textColor=colors.HexColor("#333333"), spaceBefore=5, spaceAfter=3)
    label_style = ParagraphStyle("label", parent=base, fontSize=9.9, leading=14, spaceBefore=1, spaceAfter=2)
    bullet = ParagraphStyle("bullet", parent=base, fontSize=9.55, leading=13.4, leftIndent=13, firstLineIndent=-8, spaceAfter=2)
    why_bullet = ParagraphStyle("why_bullet", parent=bullet, fontSize=9.75, leading=13.8)
    story = []
    first_h1 = True
    current_context: str | None = None
    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            story.append(Spacer(1, 2))
            continue
        if line.startswith("# "):
            if not first_h1:
                story.append(PageBreak())
            first_h1 = False
            story.append(Paragraph(clean_pdf_inline(line[2:]), title if not story else h1))
            current_context = None
        elif line.startswith("## "):
            story.append(Paragraph(clean_pdf_inline(line[3:]), h2))
            current_context = None
        elif line.startswith("### "):
            story.append(Paragraph(clean_pdf_inline(line[4:]), h3))
            current_context = None
        elif line.startswith("- "):
            style = why_bullet if current_context == "why" else bullet
            story.append(Paragraph("• " + clean_pdf_inline(line[2:]), style))
        else:
            m = re.match(r"^(【[^】]+】)(.*)$", line)
            if m:
                label, rest = m.group(1), m.group(2).strip()
                label_html = f'<font color="#1F4E79"><b>{clean_pdf_inline(label)}</b></font>'
                story.append(Paragraph(label_html + (" " + clean_pdf_inline(rest) if rest else ""), label_style))
                current_context = "why" if label == "【为什么能想到】" else ("same" if label == "【同题组】" else None)
            else:
                story.append(Paragraph(clean_pdf_inline(line), base))
                current_context = None
    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=A4,
        leftMargin=14 * mm,
        rightMargin=14 * mm,
        topMargin=13 * mm,
        bottomMargin=13 * mm,
    )
    doc.build(story)


def render_pdf_samples(pdf_path: Path, out_dir: Path) -> dict[str, object]:
    out_dir.mkdir(parents=True, exist_ok=True)
    for old in out_dir.glob("page_*.png"):
        old.unlink()
    try:
        import fitz

        doc = fitz.open(str(pdf_path))
        pages = [0, 1, 2, max(0, len(doc) // 2 - 1), max(0, len(doc) - 3), max(0, len(doc) - 2), max(0, len(doc) - 1)]
        rendered = []
        for idx in sorted(set(pages)):
            page = doc.load_page(idx)
            pix = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
            out = out_dir / f"page_{idx + 1:03d}.png"
            pix.save(str(out))
            rendered.append(str(out))
        return {"page_count": len(doc), "rendered": rendered}
    except ModuleNotFoundError:
        helper = r'''
import json
import sys
from pathlib import Path
import fitz

pdf_path = Path(sys.argv[1])
out_dir = Path(sys.argv[2])
out_dir.mkdir(parents=True, exist_ok=True)
doc = fitz.open(str(pdf_path))
pages = [0, 1, 2, max(0, len(doc) // 2 - 1), max(0, len(doc) - 3), max(0, len(doc) - 2), max(0, len(doc) - 1)]
rendered = []
for idx in sorted(set(pages)):
    page = doc.load_page(idx)
    pix = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
    out = out_dir / f"page_{idx + 1:03d}.png"
    pix.save(str(out))
    rendered.append(str(out))
print(json.dumps({"page_count": len(doc), "rendered": rendered}, ensure_ascii=False))
'''
        try:
            proc = subprocess.run(
                ["python3", "-c", helper, str(pdf_path), str(out_dir)],
                check=True,
                capture_output=True,
                text=True,
            )
            return json.loads(proc.stdout)
        except Exception:
            from pypdf import PdfReader

            reader = PdfReader(str(pdf_path))
            return {"page_count": len(reader.pages), "rendered": []}


def pollution_counts(text: str) -> dict[str, int]:
    markers = [
        "提醒",
        "提示",
        "边界",
        "使用方法",
        "审计",
        "审核",
        "GPT",
        "Claude",
        "Opus",
        "Codex",
        "debug",
        "TODO",
        "SHA",
        "/Users/",
        "reports/",
    ]
    return {m: text.count(m) for m in markers}


def write_trace_tables(trace_rows: list[dict[str, str]]) -> None:
    fieldnames = [
        "title",
        "suite",
        "question",
        "part",
        "same_group",
        "score",
        "score_method",
        "layer",
        "layer_method",
        "question_text",
    ]
    with TRACE_CSV.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(trace_rows)

    backlog_rows = [
        row
        for row in trace_rows
        if row["same_group"] == "YES" and not row["score"]
    ]
    with UNSCORED_CSV.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(backlog_rows)


def write_qa(stats: dict[str, int], pdf_info: dict[str, object]) -> None:
    text = OUT_MD.read_text(encoding="utf-8")
    counts = pollution_counts(text)
    lines = [
        "# 选必一排版分值层次优化 QA",
        "",
        f"- source_md: `{SOURCE_MD}`",
        f"- output_md: `{OUT_MD}`",
        f"- output_docx: `{OUT_DOCX}`",
        f"- output_pdf: `{OUT_PDF}`",
        f"- trace_csv: `{TRACE_CSV}`",
        f"- unscored_same_group_backlog: `{UNSCORED_CSV}`",
        "",
        "## Requested Fixes",
        "",
        "- `为什么能想到`：Word/PDF 生成器已改为正常字号标签 + 正常字号三段式触发链，不再用小号三行压缩。",
        "- `同题组`：已按桶拆成多行，并在设问原文或已钉死层次能可靠确定分值的题上标注本题分值；不从同名候选表猜分，避免误标。",
        "- `得分层次`：对设问自带分层或已钉死的分层细则，新增学生可读的得分层次行。",
        "- DOCX 结构已生成；本机未发现 `soffice`，DOCX 无法走 LibreOffice PNG 全页渲染，已改用 PDF 抽样渲染做视觉 QA。",
        "",
        "## Structure Counts",
        "",
        f"- examples: {stats['examples']}",
        f"- 同题组 blocks: {stats['same_group_blocks']}",
        f"- 同题组 score annotated: {stats['score_annotated_same_group']}",
        f"- 分值来自设问或钉死层次: {stats['score_from_prompt_or_hard']}",
        f"- 分值来自源文本精确设问匹配: {stats['score_from_source']}",
        f"- 同题组 bucketized: {stats['same_group_bucketized']}",
        f"- 同题组 blocks 保守留空分值: {stats['unscored_same_group']}",
        f"- 同题组例题已标分: {stats['same_group_scored_examples']}",
        f"- 同题组例题保守留空分值: {stats['same_group_unscored_examples']}",
        f"- 得分层次 added: {stats['layer_notes_added']}",
        f"- 为什么能想到: {text.count('【为什么能想到】')}",
        f"- 卷面句: {text.count('【卷面句】')}",
        "",
        "## Remaining Pollution Counts",
        "",
    ]
    for marker, count in counts.items():
        lines.append(f"- {marker}: {count}")
    lines.extend(
        [
            "",
            "## PDF Render QA",
            "",
            f"- page_count: {pdf_info['page_count']}",
            "- rendered_samples:",
        ]
    )
    for item in pdf_info["rendered"]:
        lines.append(f"  - `{item}`")
    QA.write_text("\n".join(lines) + "\n", encoding="utf-8")


def zip_delivery() -> None:
    with zipfile.ZipFile(OUT_ZIP, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in [OUT_MD, OUT_DOCX, OUT_PDF, QA, TRACE_CSV, UNSCORED_CSV]:
            zf.write(path, path.name)


def main() -> None:
    if not SOURCE_MD.exists():
        raise SystemExit(f"missing source: {SOURCE_MD}")
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    source = SOURCE_MD.read_text(encoding="utf-8")
    enhanced, stats, trace_rows = enhance_markdown(source)
    OUT_MD.write_text(enhanced, encoding="utf-8", newline="\n")
    write_trace_tables(trace_rows)
    build_docx(OUT_MD, OUT_DOCX)
    build_pdf(OUT_MD, OUT_PDF)
    pdf_info = render_pdf_samples(OUT_PDF, OUT_DIR / "render_qa_pdf_20260526")
    write_qa(stats, pdf_info)
    zip_delivery()
    print(OUT_MD)
    print(OUT_DOCX)
    print(OUT_PDF)
    print(OUT_ZIP)


if __name__ == "__main__":
    main()
