from __future__ import annotations

import csv
import shutil
from pathlib import Path

from docx import Document


RUN_DIR = Path("/Users/wanglifei/GaokaoPolitics/beijing-politics-sync/reports/选必一_哲学宝典式重建_2026-05-16/20_review237_verify_final_20260528")
BASE_DOCX = Path("/Users/wanglifei/Desktop/选必一_当代国际政治与经济_主观题术语宝典_Claude严格细则修订版_OCR后_20260527.docx")
OUT_DOCX = RUN_DIR / "选必一_当代国际政治与经济_主观题术语宝典_237条逐条核实最终稿_20260528.docx"
OUT_MD = RUN_DIR / "选必一_当代国际政治与经济_主观题术语宝典_237条逐条核实最终稿_20260528.md"


PATCHES = [
    {
        "review_id": "3",
        "field": "卷面句",
        "old": "和平与发展仍是时代主题，面对人工智能与大数据带来的治理缺口、全球南方诉求增多的现实，各方应在和平与发展的国际潮流中加强对话合作，共同推动全球治理民主化。",
        "new": "和平与发展仍是时代主题，人工智能、大数据带来的治理挑战和全球南方国家诉求上升，要求各方在和平合作框架内推进更有效、更民主的全球治理。",
        "reason": "原题材料支撑人工智能、大数据挑战、全球南方发声和民主全球治理，但原句“国际潮流中加强对话合作”表述偏扩展，收紧为材料和开放细则均能支撑的表述。",
    },
    {
        "review_id": "86",
        "field": "卷面句",
        "old": "中拉双方维护以世界贸易组织为核心的多边贸易体制，反对单边主义和保护主义，维护公平、开放、非歧视的国际经贸秩序，为共同现代化创造稳定环境。",
        "new": "中拉双方坚定维护多边贸易体制，维护全球产业链供应链稳定畅通和开放合作的国际环境，有利于扩大双方“发展工程”中的利益汇合点。",
        "reason": "原细则/材料写“坚定维护多边贸易体制、维护全球产业链供应链稳定畅通、维护开放合作的国际环境”，未直接写反对单边主义、保护主义和非歧视秩序。",
    },
    {
        "review_id": "91",
        "field": "什么时候写",
        "old": "欧盟对中国新能源汽车征收临时反补贴税，不只是中欧双边争端，也会扭曲市场竞争、抬高消费成本并削弱自身绿色产业转型。",
        "new": "欧盟对中国新能源汽车征收临时反补贴税，不只是中欧双边争端，也会破坏市场竞争、市场秩序，损害消费者利益，并不利于自身发展。",
        "reason": "原细则写破坏市场竞争/市场秩序、不利于自身发展、损害消费者利益；未写“削弱自身绿色产业转型”。",
    },
    {
        "review_id": "91",
        "field": "卷面句",
        "old": "贸易保护主义和单边措施会破坏公平市场竞争和多边贸易秩序，既损害消费者利益，也不利于实施方自身产业升级，最终阻碍世界经济共同发展。",
        "new": "贸易保护主义和单边措施会破坏市场竞争和市场秩序，既损害消费者利益，也不利于自身发展，最终阻碍世界经济共同发展。",
        "reason": "原细则明确“不利于自身的发展”，未写“实施方自身产业升级”。",
    },
    {
        "review_id": "101",
        "field": "卷面句",
        "old": "中国坚持高水平对外开放、高质量共建“一带一路”，有利于建设开放型世界经济，推动各国在开放中合作、在合作中共赢。",
        "new": "中国坚持高水平对外开放、高质量共建“一带一路”，有利于建设开放型世界经济，促进开放合作、互利共赢和世界经济共同繁荣。",
        "reason": "原材料支撑高水平开放与共建“一带一路”，细则角度写开放型世界经济/开放合作/世界经济共同繁荣；原句末尾较口号化，改为细则用语。",
    },
    {
        "review_id": "103",
        "field": "什么时候写",
        "old": "释放数据价值、繁荣数字经济，需要在开放流动和包容发展中提升各国数字经济参与能力。",
        "new": "释放数据价值、繁荣数字经济，需要在开放合作、包容普惠和安全保障中弥合全球数字鸿沟，让更多国家共享数字经济发展成果。",
        "reason": "原细则给出主权平等、开放合作、包容普惠、安全保障，材料表格写弥合数据鸿沟、释放数据价值、繁荣数字经济；未直接写“数字能力建设”。",
    },
    {
        "review_id": "103",
        "field": "卷面句",
        "old": "各国应坚持开放合作和包容普惠，促进数据资源共享和数字能力建设，让更多国家参与数字经济发展并分享发展成果。",
        "new": "各国应坚持开放合作和包容普惠，助力弥合全球数字鸿沟、释放数据价值、繁荣数字经济，让更多国家共享数字经济发展成果。",
        "reason": "原题材料和细则能支撑弥合数字鸿沟、释放数据价值、繁荣数字经济；“数据资源共享和数字能力建设”证据不足。",
    },
    {
        "review_id": "120",
        "field": "卷面句",
        "old": "面对人工智能与大数据带来的两极分化和全球治理代表性不足，各国应推动构建人类命运共同体，共同应对全球性挑战，推动全球治理民主化。",
        "new": "人工智能、大数据带来的全球治理挑战和代表性不足，说明各国需要以人类命运共同体理念共同应对全球性治理问题，推动全球治理体系更加有效、民主。",
        "reason": "原题材料支撑技术挑战、代表性不足和有效民主治理体系；将“全球性挑战/治理民主化”收紧为原材料可直接支撑的治理问题。",
    },
    {
        "review_id": "145",
        "field": "卷面句",
        "old": "金砖国家、二十国集团、上海合作组织等国际组织和合作机制作用增强，为全球南方国家表达诉求、参与全球治理、推动治理民主化提供了重要平台。",
        "new": "金砖国家、二十国集团、上海合作组织等国际组织和合作机制作用增强，为全球南方国家在国际舞台发挥更大作用、发出全球治理声音提供重要平台。",
        "reason": "原材料写全球南方崛起、相关国际组织作用增强、南方国家要发出声音；“推动治理民主化”属于合理方向但不够贴近材料原句，故收紧。",
    },
    {
        "review_id": "148",
        "field": "卷面句",
        "old": "中方主张通过磋商解决中欧新能源汽车贸易争端，坚持互利共赢，避免以单边关税破坏开放合作的国际经贸环境。",
        "new": "中方主张通过磋商解决中欧新能源汽车贸易争端，坚持互利共赢，推动贸易自由便利并维护市场秩序。",
        "reason": "原细则写和平谈判方式、互利共赢、推动贸易自由便利、维护市场秩序；未直接写“开放合作的国际经贸环境”。",
    },
    {
        "review_id": "155",
        "field": "什么时候写",
        "old": "资料卡列举地缘政治动荡、地区安全问题、经贸科技交流政治化和保护主义上升，推动中国经济开新局需要把发展优势与安全风险统筹起来。",
        "new": "细则强调我国发展的外部环境日趋严峻、风险挑战增多，推动中国经济开新局需要把发展优势与安全风险统筹起来。",
        "reason": "原细则写外部环境日趋严峻、风险挑战增多；未直接写“地缘政治动荡、经贸科技交流政治化”。",
    },
    {
        "review_id": "155",
        "field": "卷面句",
        "old": "面对地缘政治动荡、保护主义上升和经贸科技交流政治化，中国应坚持总体国家安全观，统筹发展与安全，以新安全格局保障新发展格局。",
        "new": "面对外部环境日趋严峻、风险挑战增多，中国应坚持总体国家安全观，统筹发展与安全，以新安全格局保障新发展格局。",
        "reason": "改用原细则明确表述。",
    },
    {
        "review_id": "155A",
        "field": "什么时候写",
        "old": "题目先要求分析“变局”中的经济形势，资料卡集中呈现地缘政治动荡、局部冲突、全球性问题、经贸科技交流政治化、保护主义和逆全球化趋势，因此要先写外部风险压力。",
        "new": "题目先要求分析“变局”中的经济形势，细则明确我国发展的外部环境日趋严峻、风险挑战增多，因此要先写外部风险压力。",
        "reason": "同为2024朝阳一模Q21，原细则支撑“外部环境日趋严峻、风险挑战增多”，未直接支撑地缘政治动荡和经贸科技交流政治化清单。",
    },
    {
        "review_id": "155A",
        "field": "材料信号",
        "old": "材料信号：资料卡集中呈现地缘政治动荡、局部冲突、全球性问题、经贸科技交流政治化、保护主义和逆全球化等多类外部压力；这是典型的\"百年变局\"现场。",
        "new": "材料信号：细则集中呈现我国外部环境日趋严峻、风险挑战增多；这是典型的“变局”中分析经济形势的信号。",
        "reason": "收紧为原细则明确表述。",
    },
    {
        "review_id": "155A",
        "field": "答题动作",
        "old": "答题动作：先写我国外部环境日趋严峻、风险挑战增多，再把地缘政治、保护主义、\"小院高墙\"接为对经贸科技交流的具体冲击，最后承接统筹发展与安全的开新局思路。",
        "new": "答题动作：先写我国外部环境日趋严峻、风险挑战增多，再把外部压力接为统筹发展与安全的必要性，最后承接以新安全格局保障新发展格局的开新局思路。",
        "reason": "删除无法在原细则中确认的“小院高墙”扩展。",
    },
    {
        "review_id": "155A",
        "field": "卷面句",
        "old": "我国发展的外部环境日趋严峻、风险挑战增多，地缘政治动荡、保护主义上升和“小院高墙”影响经贸科技交流，使中国经济开新局必须统筹发展与安全。",
        "new": "我国发展的外部环境日趋严峻、风险挑战增多，使中国经济开新局必须统筹发展与安全。",
        "reason": "删除原细则未直接支撑的具体风险清单。",
    },
    {
        "review_id": "155B",
        "field": "材料信号",
        "old": "材料信号：资料卡把地缘政治动荡、地区安全问题、经贸科技交流政治化、保护主义上升四类风险并列摆出；这些风险既来自政治也来自经贸，单纯抓发展或单纯防风险都不够。",
        "new": "材料信号：细则把我国外部环境日趋严峻、风险挑战增多作为形势判断；这些风险会同时影响发展与安全，单纯抓发展或单纯防风险都不够。",
        "reason": "同为2024朝阳一模Q21，原细则未直接支撑地缘政治动荡和经贸科技交流政治化清单。",
    },
    {
        "review_id": "155B",
        "field": "答题动作",
        "old": "答题动作：先把地缘政治动荡、保护主义上升、经贸科技交流政治化作为外部风险的清单，再用坚持总体国家安全观、统筹发展与安全作为应对总思路，最后落到以新安全格局保障新发展格局推动中国经济开新局。",
        "new": "答题动作：先写我国外部环境日趋严峻、风险挑战增多，再用坚持总体国家安全观、统筹发展与安全作为应对总思路，最后落到以新安全格局保障新发展格局推动中国经济开新局。",
        "reason": "改用2024朝阳一模Q21细则可直接核验的表述。",
    },
    {
        "review_id": "219",
        "field": "卷面句",
        "old": "中国坚持大国担当，是150多个国家和地区的主要贸易伙伴，并通过对外投资为东道国贡献3000多亿美元税收，以自身开放发展为世界经济提供稳定预期和发展确定性。",
        "new": "中国坚持大国担当，是150多个国家和地区的主要贸易伙伴，并通过对外投资为东道国贡献3000多亿美元税收，以互利共赢的开放战略为世界共享发展机遇。",
        "reason": "原材料支撑贸易伙伴、东道国纳税和世界共享机遇；“稳定预期和发展确定性”来自设问方向但不宜放成细则化卷面句结论。",
    },
    {
        "review_id": "220",
        "field": "什么时候写",
        "old": "资料卡显示个别国家泛化国家安全、把经贸科技交流政治化并筑起“小院高墙”，推动中国经济开新局必须把关键核心技术和发展主动权掌握在自己手中。",
        "new": "细则强调面对外部风险挑战，推动中国经济开新局必须坚持独立自主、增强自主创新能力，把关键核心技术和发展主动权掌握在自己手中。",
        "reason": "原细则有独立自主、自主创新、关键核心技术和发展主动权；未直接写“小院高墙”。",
    },
    {
        "review_id": "220",
        "field": "卷面句",
        "old": "面对“小院高墙”和经贸科技交流政治化，中国应坚持独立自主，增强自主创新能力，实现关键核心技术自主可控，把发展主动权牢牢掌握在自己手中。",
        "new": "面对外部风险挑战，中国应坚持独立自主，增强自主创新能力，实现关键核心技术自主可控，把发展主动权牢牢掌握在自己手中。",
        "reason": "改用原细则和参考示例能支撑的表述。",
    },
    {
        "review_id": "220",
        "field": "材料信号",
        "old": "材料信号：资料卡里个别国家泛化国家安全、把经贸科技交流政治化、构筑\"小院高墙\"，对中国关键核心技术形成挤压。",
        "new": "材料信号：细则把外部风险挑战与自主创新、关键核心技术联系起来，说明中国经济开新局必须增强发展主动权。",
        "reason": "同为2024朝阳一模Q21，删除原细则无法确认的“小院高墙”扩展。",
    },
    {
        "review_id": "220",
        "field": "答题动作",
        "old": "答题动作：从\"小院高墙\"的外部压力切入，说明中国必须把关键核心技术和发展主动权掌握在自己手里，再落到坚持独立自主、自力更生、增强发展主动权。",
        "new": "答题动作：从外部风险挑战切入，说明中国必须把关键核心技术和发展主动权掌握在自己手里，再落到坚持独立自主、自力更生、增强发展主动权。",
        "reason": "改用原细则能支撑的外部风险表述。",
    },
    {
        "review_id": "86",
        "field": "材料信号",
        "old": "材料信号：“发展工程”把多边贸易体制、反对单边主义和保护主义、开放非歧视秩序放在同一段。",
        "new": "材料信号：“发展工程”把多边贸易体制、全球产业链供应链稳定畅通和开放合作国际环境放在同一段。",
        "reason": "2026丰台期末细则原文未直接写反对单边主义和保护主义、开放非歧视秩序。",
    },
    {
        "review_id": "86",
        "field": "答题动作",
        "old": "答题动作：先写维护 WTO 多边贸易体制，再写反对单边主义和保护主义，最后落到为共同现代化创造稳定环境。",
        "new": "答题动作：先写维护多边贸易体制，再接全球产业链供应链稳定畅通和开放合作国际环境，最后落到扩大中拉发展工程的利益汇合点。",
        "reason": "改用2026丰台期末细则原文支撑的三项。",
    },
    {
        "review_id": "86B",
        "field": "材料信号",
        "old": "材料信号：“发展工程”强调开放合作环境，并指向反对保护主义、封闭排他等破坏合作的外部因素。",
        "new": "材料信号：“发展工程”强调开放合作环境，并与维护多边贸易体制、维护全球产业链供应链稳定畅通并列。",
        "reason": "同为2026丰台期末Q20，删除原细则未直接支撑的保护主义/封闭排他扩展。",
    },
    {
        "review_id": "86B",
        "field": "答题动作",
        "old": "答题动作：先写维护开放合作的国际经济环境，再接抵制保护主义和封闭排他，最后落到拓展双方现代化合作空间。",
        "new": "答题动作：先写维护开放合作的国际经济环境，再接多边贸易体制和全球产业链供应链稳定畅通，最后落到拓展双方现代化合作空间。",
        "reason": "改为原细则可核验的关联项。",
    },
    {
        "review_id": "86B",
        "field": "卷面句",
        "old": "中拉双方维护开放合作的国际经济环境，有利于抵制保护主义和封闭排他，为双方现代化拓展合作空间。",
        "new": "中拉双方维护开放合作的国际经济环境，有利于稳定“发展工程”的外部条件，为双方现代化拓展合作空间。",
        "reason": "删除未在2026丰台期末细则中直接出现的保护主义/封闭排他。",
    },
    {
        "review_id": "86B",
        "field": "同题组",
        "old": "其他：反对单边主义和保护主义（时代背景与经济全球化交叉处）",
        "new": "其他：维护多边贸易体制和全球产业链供应链稳定畅通",
        "reason": "同题组提示改回该题细则原文支撑项。",
    },
]


def replace_in_docx(doc: Document) -> list[dict[str, str]]:
    results: list[dict[str, str]] = []
    # Remove backend/model wording from the student-facing subtitle.
    for p in doc.paragraphs:
        if "学生厚版 · 严格细则修订版（Claude 2026.5.27 · OCR补全后 · 共104术语/237条目）" in p.text:
            p.text = "学生厚版 · 237条逐条核实最终版（2026.5.28 · 共104术语/237条目）"
            results.append({"review_id": "TITLE", "field": "subtitle", "old": "Claude OCR subtitle", "new": p.text, "applied": "YES", "reason": "学生最终稿去除后台模型标识。"})
            break
    for patch in PATCHES:
        applied = 0
        for p in doc.paragraphs:
            if patch["old"] in p.text:
                p.text = p.text.replace(patch["old"], patch["new"])
                applied += 1
        result = dict(patch)
        result["applied"] = "YES" if applied else "NO"
        result["applied_count"] = str(applied)
        results.append(result)
    return results


def docx_to_md(docx_path: Path, md_path: Path) -> None:
    doc = Document(str(docx_path))
    lines: list[str] = []
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            lines.append("")
            continue
        if idx == 0:
            lines.append(f"# {text}")
        elif text in {"时代背景", "理论", "经济全球化", "政治多极化", "中国", "联合国", "六大要素术语极简速记版", "目录"}:
            lines.append(f"\n## {text}")
        elif text.startswith("核心答题点："):
            lines.append(f"\n### {text}")
        else:
            lines.append(text)
    md_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def build_verification_ledger(patch_results: list[dict[str, str]]) -> None:
    full_rows = list(csv.DictReader((RUN_DIR / "REVIEW237_FULL_PARSED.csv").open(encoding="utf-8")))
    source_rows = {r["review_id"]: r for r in csv.DictReader((RUN_DIR / "REVIEW237_SOURCE_CHECK.csv").open(encoding="utf-8"))}
    cross_rows = {r["review_id"]: r for r in csv.DictReader((RUN_DIR / "REVIEW237_19_CROSSCHECK.csv").open(encoding="utf-8"))}
    patched_ids = {r["review_id"] for r in patch_results if r.get("review_id", "").isdigit() and r.get("applied") == "YES"}
    final_sentence_by_id = {
        r["review_id"]: r["new"]
        for r in patch_results
        if r.get("review_id", "").isdigit() and r.get("field") == "卷面句" and r.get("applied") == "YES"
    }
    flagged_patch_ids = {"3", "86", "91", "101", "103", "120", "145", "148", "155", "219", "220"}
    out: list[dict[str, str]] = []
    for row in full_rows:
        rid = row["review_id"]
        status = row["status"]
        if rid in patched_ids:
            decision = "PATCH_APPLIED"
        elif status in {"✅高度匹配", "✅基本匹配"}:
            decision = "REVIEW_STATUS_ACCEPTED_NO_CHANGE"
        elif rid in flagged_patch_ids:
            decision = "PATCH_REVIEWED_BUT_NOT_APPLIED"
        else:
            decision = "RETAIN_SOURCE_SUPPORTED_OR_REASONABLE"
        out.append({
            "review_id": rid,
            "bucket": row["bucket"],
            "core": row["core"],
            "question": row["question"],
            "review_status": status,
            "final_decision": decision,
            "sentence_exact_in_19": cross_rows.get(rid, {}).get("sentence_exact_in_19", ""),
            "source_suite": source_rows.get(rid, {}).get("suite", ""),
            "source_match_note": source_rows.get(rid, {}).get("missing_clause_source_matches", ""),
            "final_sentence": final_sentence_by_id.get(rid, row.get("sentence", "")),
        })
    with (RUN_DIR / "REVIEW237_FINAL_VERIFICATION_LEDGER.csv").open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=list(out[0].keys()))
        writer.writeheader()
        writer.writerows(out)


def main() -> None:
    shutil.copy2(BASE_DOCX, OUT_DOCX)
    doc = Document(str(OUT_DOCX))
    patch_results = replace_in_docx(doc)
    doc.save(str(OUT_DOCX))
    docx_to_md(OUT_DOCX, OUT_MD)

    with (RUN_DIR / "APPLIED_PATCH_LEDGER_20260528.csv").open("w", newline="", encoding="utf-8") as f:
        keys: list[str] = []
        for row in patch_results:
            for key in row:
                if key not in keys:
                    keys.append(key)
        writer = csv.DictWriter(f, fieldnames=keys)
        writer.writeheader()
        writer.writerows(patch_results)
    build_verification_ledger(patch_results)


if __name__ == "__main__":
    main()
