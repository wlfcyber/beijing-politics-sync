import csv
import re
import shutil
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path


RUN = Path(r"C:\bp_sync_visible\reports\culture_all_suites_rerun_2026-04-29")
DESKTOP = Path(r"C:\Users\Administrator\Desktop\4.29文化线56套重跑结果")
STUDENT = DESKTOP / "01_学生版Word"
AUDIT = DESKTOP / "02_审计证据"
CSVOUT = DESKTOP / "03_结构化CSV"
LOGS = DESKTOP / "04_过程日志"

for p in [STUDENT, AUDIT, CSVOUT, LOGS]:
    p.mkdir(parents=True, exist_ok=True)

SLOT_NAMES = {
    "0": "0. 载体：文化载体、文化符号、展示平台、精神标识",
    "1": "1. 特点：源远流长、博大精深",
    "2": "2. 作用：引领风尚、教育人民；服务社会、推动发展",
    "3": "3. 横向：文化交流与传播",
    "4": "4. 纵向：继承发展、融通资源、立足时代、双创",
    "5": "5. 建设文化强国，树立文化自信",
    "6": "6. 民族精神",
    "7": "7. 坚持习近平文化思想",
}


def read(path):
    return path.read_text(encoding="utf-8", errors="replace") if path.exists() else ""


def clean_cell(s):
    return re.sub(r"\s+", " ", s.replace("<br>", " ")).strip()


def split_md_row(line):
    if not line.strip().startswith("|"):
        return []
    cells = [clean_cell(c) for c in line.strip().strip("|").split("|")]
    if not cells:
        return []
    joined = "".join(cells)
    if set(joined) <= set("-: "):
        return []
    return cells


def slot_from_heading(heading):
    m = re.search(r"###\s*([0-7])\.", heading)
    if m:
        return m.group(1)
    if "载体" in heading:
        return "0"
    if "特点" in heading:
        return "1"
    if "作用" in heading:
        return "2"
    if "横向" in heading or "交流" in heading:
        return "3"
    if "纵向" in heading or "继承" in heading or "创造性" in heading:
        return "4"
    if "文化强国" in heading or "文化自信" in heading:
        return "5"
    if "民族精神" in heading:
        return "6"
    if "习近平文化思想" in heading:
        return "7"
    return ""


def slots_from_text(text):
    slots = set()
    for i in range(8):
        if re.search(rf"(^|[^0-9]){i}([^0-9]|$)", text):
            slots.add(str(i))
    keywords = {
        "0": ["载体", "符号", "展示平台", "精神标识", "文物建筑"],
        "1": ["源远流长", "博大精深", "连续性", "中华文明特点", "中华优秀传统文化特点"],
        "2": ["文化功能", "引领风尚", "教育人民", "服务社会", "推动发展", "精神文化需求", "文化权益"],
        "3": ["文化交流", "文化传播", "文明互鉴", "文化多样性", "交流互鉴", "国际传播"],
        "4": ["继承", "传承", "创造性转化", "创新性发展", "融通不同资源", "综合创新", "立足时代"],
        "5": ["文化强国", "文化自信", "文化认同", "中华文化立场"],
        "6": ["民族精神", "时代精神", "革命文化", "爱国主义", "核心价值观", "家国情怀"],
        "7": ["习近平文化思想"],
    }
    for slot, terms in keywords.items():
        if any(t in text for t in terms):
            slots.add(slot)
    return "、".join(sorted(slots, key=int))


def source_to_suite(source):
    m = re.search(r"(20\d{2}[^第|｜]+?)\s*第", source)
    return clean_cell(m.group(1)) if m else source


def source_to_question(source):
    m = re.search(r"第\s*([0-9一二三四五六七八九十]+(?:[（(][0-9一二三四五六七八九十]+[）)])?)\s*题?", source)
    return f"第{m.group(1)}题" if m else ""


def normalize_suite_label(text):
    text = text or ""
    text = re.sub(r"北京|北京市|高三|思想政治|政治|综合练习|学年度|第二学期|第一学期|（上）|\\(上\\)|上", "", text)
    text = text.replace("一模", "一模").replace("二模", "二模").replace("期中", "期中").replace("期末", "期末")
    return text


def match_suite_id(label, roster):
    label_norm = normalize_suite_label(label)
    year = re.search(r"20\d{2}", label_norm)
    districts = ["海淀", "西城", "东城", "朝阳", "丰台", "石景山", "门头沟", "顺义", "延庆", "房山", "昌平", "通州"]
    district = next((d for d in districts if d in label_norm), "")
    stage = ""
    for s in ["一模", "二模", "期中", "期末"]:
        if s in label_norm:
            stage = s
            break
    if not (year and district and stage):
        return ""
    candidates = [
        r for r in roster
        if r.get("year") == year.group(0) and r.get("district") == district and stage in r.get("suite_name", "")
    ]
    if len(candidates) == 1:
        return candidates[0]["suite_id"]
    return candidates[0]["suite_id"] if candidates else ""


def extract_template_entries():
    text = read(RUN / "USER_CULTURE_TEMPLATE.md")
    rows = []
    current_slot = ""
    current_heading = ""
    for line in text.splitlines():
        if line.startswith("### "):
            current_heading = clean_cell(line.lstrip("# "))
            maybe = slot_from_heading(line)
            if maybe:
                current_slot = maybe
        cells = split_md_row(line)
        if len(cells) >= 4 and cells[:4] != ["来源", "材料信息", "触发知识", "触发逻辑"]:
            source, material, knowledge, logic = cells[:4]
            if not source or source == "---":
                continue
            if source.startswith("框架位置") or source.startswith("年份"):
                continue
            slot = current_slot or slots_from_text(" ".join(cells))
            rows.append(
                {
                    "suite_id": "",
                    "suite_name": source_to_suite(source),
                    "question": source_to_question(source),
                    "question_type": "template_existing",
                    "framework_slot": slot,
                    "evidence_class": "template-prior-reviewed",
                    "evidence_source": source,
                    "material_excerpt": material,
                    "knowledge_point": knowledge,
                    "trigger_logic": logic,
                    "student_answer_landing": logic,
                    "choice_answer": "",
                    "wrong_option_pattern": "",
                    "status": "PASS_TEMPLATE",
                    "blocker": "",
                    "_source_file": "USER_CULTURE_TEMPLATE.md",
                    "_template_heading": current_heading,
                }
            )
    return rows


def read_csv_rows(path):
    try:
        with path.open("r", encoding="utf-8-sig", newline="") as f:
            return list(csv.DictReader(f))
    except FileNotFoundError:
        return []


def worker_entries():
    rows = []
    for path in sorted((RUN / "worker_outputs").glob("*_culture_entries.csv")):
        for r in read_csv_rows(path):
            r = dict(r)
            r["_source_file"] = str(path)
            if not r.get("framework_slot"):
                r["framework_slot"] = slots_from_text(" ".join(str(v) for v in r.values()))
            rows.append(r)
    return rows


def dedupe(rows):
    seen = set()
    out = []
    for r in rows:
        key = (
            r.get("suite_name", ""),
            r.get("question", ""),
            r.get("framework_slot", ""),
            r.get("knowledge_point", "")[:60],
            r.get("material_excerpt", "")[:60],
        )
        if key in seen:
            continue
        seen.add(key)
        out.append(r)
    return out


def write_entries_csv(path, rows):
    fields = [
        "suite_id",
        "suite_name",
        "question",
        "question_type",
        "framework_slot",
        "evidence_class",
        "evidence_source",
        "material_excerpt",
        "knowledge_point",
        "trigger_logic",
        "student_answer_landing",
        "choice_answer",
        "wrong_option_pattern",
        "status",
        "blocker",
        "_source_file",
    ]
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fields, extrasaction="ignore")
        w.writeheader()
        w.writerows(rows)


def write_table_csv(path, rows):
    fields = list(rows[0].keys()) if rows else []
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fields, extrasaction="ignore")
        w.writeheader()
        w.writerows(rows)


def student_entry(r):
    source = r.get("suite_name") or r.get("evidence_source", "")
    question = r.get("question", "")
    title = f"{source} {question}".strip()
    evidence = r.get("evidence_class", "")
    evidence = ""
    lines = [f"### {title}"]
    if evidence:
        lines.append(f"- 证据口径：{evidence}")
    if r.get("material_excerpt"):
        lines.append(f"- 材料抓手：{r['material_excerpt']}")
    if r.get("knowledge_point"):
        lines.append(f"- 触发知识：{r['knowledge_point']}")
    if r.get("trigger_logic"):
        lines.append(f"- 为什么想到这里：{r['trigger_logic']}")
    if r.get("student_answer_landing"):
        lines.append(f"- 答案落点：{r['student_answer_landing']}")
    if r.get("choice_answer"):
        lines.append(f"- 选择题答案：{r['choice_answer']}")
    if r.get("wrong_option_pattern"):
        lines.append(f"- 错肢识别：{r['wrong_option_pattern']}")
    return "\n".join(lines)


def norm_status(r):
    return (r.get("status") or "").strip().upper()


def is_pass_row(r):
    status = norm_status(r)
    return status.startswith("PASS") or status in {"A", "B"} or status == "PASS_TEMPLATE"


def is_module_boundary_row(r):
    status = norm_status(r)
    evidence = (r.get("evidence_class") or "").lower()
    return status in {"MODULE_BOUNDARY", "MODULE-BOUNDARY", "BOUNDARY"} or "module_boundary" in evidence or "module-boundary" in evidence


def is_reference_boundary_row(r):
    evidence = (r.get("evidence_class") or "").lower()
    blocker = (r.get("blocker") or "").lower()
    return "blocked_reference_only" in evidence or "reference-only" in blocker


def is_blocked_row(r):
    return norm_status(r).startswith("BLOCKED")


def is_student_row(r):
    if (r.get("question_type") or "").strip().lower() == "coverage":
        return False
    if (r.get("question") or "").strip().upper() == "COVERAGE":
        return False
    return is_pass_row(r)


def slot_ids(text):
    out = []
    for s in re.findall(r"[0-7]", text or ""):
        if s in SLOT_NAMES and s not in out:
            out.append(s)
    return out


def build_student_md(rows):
    grouped = defaultdict(list)
    for r in rows:
        if not is_student_row(r):
            continue
        slots = [s for s in re.split(r"[、,;/ ]+", r.get("framework_slot", "")) if s in SLOT_NAMES]
        slots = slot_ids(r.get("framework_slot", ""))
        if not slots:
            slots = ["4"]
        for slot in slots:
            grouped[slot].append(r)
    lines = [
        "# 必修四文化材料-知识触发框架（56套重跑版）",
        "",
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "说明：本稿按用户文化模板 0-7 框架组织。主观题区分逐点标分、等级细则支持和跨模块边界；选择题只在有可靠答案键时写正确项/错肢链。",
        "",
        "## 总框架",
        "",
    ]
    for i in range(8):
        lines.append(f"- {SLOT_NAMES[str(i)]}")
    lines.append("")
    for i in range(8):
        slot = str(i)
        lines.append(f"## {SLOT_NAMES[slot]}")
        lines.append("")
        items = grouped.get(slot, [])
        for r in items:
            lines.append(student_entry(r))
            lines.append("")
        if not items:
            lines.append("本轮暂无新增稳定挂点。")
            lines.append("")
    return "\n".join(lines)


def build_choice_md(rows):
    choice_rows = [
        r for r in rows
        if is_student_row(r) and "choice" in (r.get("question_type") or "").lower()
    ]
    choice_rows.sort(key=lambda r: (r.get("suite_id", ""), r.get("question", "")))
    lines = [
        "# 必修四文化选择题错肢总结（56套重跑版）",
        "",
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "说明：只收同时具备题干和可靠答案键的文化选择题；缺答案键、缺完整题干或模块边界的题只留在审计表。",
        "",
    ]
    for r in choice_rows:
        title = f"{r.get('suite_name', '')} 第{r.get('question', '')}题".replace("第第", "第").strip()
        lines.append(f"## {title}")
        if r.get("material_excerpt"):
            lines.append(f"- 材料抓手：{r['material_excerpt']}")
        if r.get("knowledge_point"):
            lines.append(f"- 触发知识：{r['knowledge_point']}")
        if r.get("choice_answer"):
            lines.append(f"- 正确答案：{r['choice_answer']}")
        if r.get("student_answer_landing"):
            lines.append(f"- 正确项落点：{r['student_answer_landing']}")
        if r.get("wrong_option_pattern"):
            lines.append(f"- 错肢识别：{r['wrong_option_pattern']}")
        lines.append("")
    if not choice_rows:
        lines.append("本轮没有符合口径的文化选择题。")
    return "\n".join(lines)


def build_coverage(rows):
    roster = read_csv_rows(RUN / "SUITE_ROSTER.csv")
    roster_by_id = {r["suite_id"]: r for r in roster}
    by_sid = defaultdict(list)
    by_suite = defaultdict(list)
    for r in rows:
        if not r.get("suite_id"):
            sid = match_suite_id(r.get("suite_name") or r.get("evidence_source", ""), roster)
            if sid:
                r["suite_id"] = sid
                r["suite_name"] = roster_by_id[sid]["suite_name"]
        if r.get("suite_id"):
            by_sid[r["suite_id"]].append(r)
        if r.get("suite_name"):
            by_suite[r["suite_name"]].append(r)
    matrix = []
    for s in roster:
        sid = s["suite_id"]
        suite_name = s["suite_name"]
        entries = []
        seen_entry_ids = set()
        for e in by_sid.get(sid, []) + by_suite.get(suite_name, []):
            marker = id(e)
            if marker in seen_entry_ids:
                continue
            seen_entry_ids.add(marker)
            entries.append(e)
        for e in entries:
            ids = slot_ids(e.get("framework_slot", ""))
            if ids:
                e["framework_slot"] = "、".join(ids)
        blocked_entries = [e for e in entries if is_blocked_row(e)]
        blockers = [e.get("blocker", "") for e in blocked_entries]
        pass_entries = [e for e in entries if is_pass_row(e)]
        boundary_entries = [e for e in entries if is_module_boundary_row(e)]
        reference_boundaries = [e for e in blocked_entries if is_reference_boundary_row(e)]
        if pass_entries and blocked_entries:
            status = "processed_with_choice_boundary"
        elif pass_entries:
            status = "processed"
        elif blocked_entries and len(reference_boundaries) == len(blocked_entries):
            status = "source-boundary-reference-only"
        elif blocked_entries and not pass_entries:
            status = "blocked"
        elif boundary_entries:
            status = "module-boundary"
        else:
            status = "module-boundary-or-no-culture-hit"
        framework_slots = "、".join(sorted({slot for e in entries for slot in slot_ids(e.get("framework_slot", ""))}, key=int))
        matrix.append(
            {
                "suite_id": sid,
                "suite_name": suite_name,
                "year": s["year"],
                "district": s["district"],
                "stage": s["stage"],
                "culture_status": status,
                "entry_count": len(entries),
                "framework_slots": "、".join(sorted({slot for e in entries for slot in re.split(r"[、,;/ ]+", e.get("framework_slot", "")) if slot in SLOT_NAMES}, key=int)),
                "blocker": "；".join(b for b in blockers if b),
            }
        )
    return matrix


def md_to_docx(md_path, docx_path):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    for raw in md_path.read_text(encoding="utf-8", errors="replace").splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(re.sub(r"\*\*(.*?)\*\*", r"\1", line))
    doc.save(docx_path)


def main():
    rows = dedupe(extract_template_entries() + worker_entries())
    write_entries_csv(CSVOUT / "culture_56_entries_merged.csv", rows)
    coverage = build_coverage(rows)
    write_table_csv(CSVOUT / "culture_56_coverage_matrix.csv", coverage)
    student_md = build_student_md(rows)
    choice_md = build_choice_md(rows)
    (STUDENT / "必修四文化选择题错肢总结_56套重跑版.md").write_text(choice_md, encoding="utf-8")
    md_to_docx(STUDENT / "必修四文化选择题错肢总结_56套重跑版.md", STUDENT / "必修四文化选择题错肢总结_56套重跑版.docx")
    (STUDENT / "必修四文化材料-知识触发框架_56套重跑版.md").write_text(student_md, encoding="utf-8")
    md_to_docx(STUDENT / "必修四文化材料-知识触发框架_56套重跑版.md", STUDENT / "必修四文化材料-知识触发框架_56套重跑版.docx")

    audit = [
        "# 必修四文化 56 套重跑审计报告",
        "",
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        f"- 合并条目：{len(rows)}",
        f"- 覆盖矩阵：{len(coverage)} 套",
        f"- 覆盖状态：{dict(Counter(r['culture_status'] for r in coverage))}",
        "",
        "## Worker Reports",
        "",
    ]
    for path in sorted((RUN / "worker_reports").glob("*_culture.md")):
        audit.append(f"### {path.name}")
        audit.append("")
        audit.append(read(path))
        audit.append("")
    (AUDIT / "culture_56_audit_report.md").write_text("\n".join(audit), encoding="utf-8")

    for name in ["TASK_BRIEF.md", "ROLE_CONTRACTS.md", "PROGRESS.md", "DECISION_LOG.md", "THREAD_REGISTRY.md", "SUITE_ROSTER.csv", "COVERAGE_MATRIX.csv"]:
        src = RUN / name
        if src.exists():
            shutil.copy2(src, LOGS / name)

    report = f"""# Culture Build Report

生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

- 合并条目：{len(rows)}
- 覆盖矩阵：{len(coverage)}
- 覆盖状态：{dict(Counter(r['culture_status'] for r in coverage))}
- Word 已生成：1
"""
    report = report.replace("Word 宸茬敓鎴愶細1", "Word 宸茬敓鎴愶細2").replace("Word 已生成：1", "Word 已生成：2")
    (DESKTOP / "BUILD_REPORT.md").write_text(report, encoding="utf-8")
    print(report)


if __name__ == "__main__":
    main()
