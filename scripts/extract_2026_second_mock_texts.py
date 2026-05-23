from __future__ import annotations

import argparse
import csv
import re
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from pypdf import PdfReader

try:
    import olefile
except ImportError:  # pragma: no cover - optional old .doc support
    olefile = None


TEXT_EXTS = {".txt", ".md"}
SUPPORTED_EXTS = {".docx", ".doc", ".pdf", ".pptx", ".txt", ".md"}
CFB_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
END_OF_CHAIN = 0xFFFFFFFE


def safe_name(path: Path, root: Path) -> str:
    rel = path.relative_to(root)
    name = "__".join(rel.parts)
    name = re.sub(r'[<>:"/\\|?*\s]+', "_", name)
    return name


def clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip() + ("\n" if text.strip() else "")


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table_idx, table in enumerate(doc.tables, 1):
        parts.append(f"\n[表格 {table_idx}]")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return clean_text("\n".join(parts))


def scan_printable_text(decoded: str) -> str:
    parts: list[str] = []
    buf: list[str] = []
    for ch in decoded:
        cp = ord(ch)
        keep = (
            0x4E00 <= cp <= 0x9FFF
            or 0x3400 <= cp <= 0x4DBF
            or 0x3000 <= cp <= 0x303F
            or 0xFF00 <= cp <= 0xFFEF
            or 0x20 <= cp <= 0x7E
            or ch in "\n\t"
        )
        if keep:
            buf.append(ch)
        elif buf:
            run = "".join(buf).strip()
            if len(run) >= 2:
                parts.append(run)
            buf = []
    if buf:
        run = "".join(buf).strip()
        if len(run) >= 2:
            parts.append(run)
    return "\n".join(parts)


def read_u16(data: bytes, offset: int) -> int:
    return int.from_bytes(data[offset : offset + 2], "little")


def read_u32(data: bytes, offset: int) -> int:
    return int.from_bytes(data[offset : offset + 4], "little")


def read_u64(data: bytes, offset: int) -> int:
    return int.from_bytes(data[offset : offset + 8], "little")


def sector_offset(sector_id: int, sector_size: int) -> int:
    return (sector_id + 1) * sector_size


def read_cfb_chain(data: bytes, start_sector: int, fat: list[int], sector_size: int, size: int | None = None) -> bytes:
    chunks: list[bytes] = []
    sector = start_sector
    seen: set[int] = set()
    while sector not in (END_OF_CHAIN, 0xFFFFFFFF):
        if sector < 0 or sector >= len(fat) or sector in seen:
            break
        seen.add(sector)
        offset = sector_offset(sector, sector_size)
        chunks.append(data[offset : offset + sector_size])
        sector = fat[sector]
        if size is not None and sum(len(chunk) for chunk in chunks) >= size:
            break
    stream = b"".join(chunks)
    return stream[:size] if size is not None else stream


def read_cfb_stream(path: Path, stream_name: str) -> bytes:
    data = path.read_bytes()
    if not data.startswith(CFB_MAGIC):
        raise ValueError("not a compound Word document")

    sector_size = 1 << read_u16(data, 0x1E)
    directory_start = read_u32(data, 0x30)
    fat_sector_count = read_u32(data, 0x2C)
    fat_sector_ids = [read_u32(data, 0x4C + idx * 4) for idx in range(min(fat_sector_count, 109))]
    fat: list[int] = []
    for sector_id in fat_sector_ids:
        if sector_id in (END_OF_CHAIN, 0xFFFFFFFF):
            continue
        offset = sector_offset(sector_id, sector_size)
        sector = data[offset : offset + sector_size]
        fat.extend(read_u32(sector, pos) for pos in range(0, len(sector), 4))

    directory = read_cfb_chain(data, directory_start, fat, sector_size)
    for offset in range(0, len(directory), 128):
        entry = directory[offset : offset + 128]
        if len(entry) < 128:
            continue
        name_len = read_u16(entry, 0x40)
        if name_len < 2:
            continue
        name = entry[: name_len - 2].decode("utf-16-le", errors="ignore")
        object_type = entry[0x42]
        if object_type == 2 and name == stream_name:
            start_sector = read_u32(entry, 0x74)
            stream_size = read_u64(entry, 0x78)
            return read_cfb_chain(data, start_sector, fat, sector_size, stream_size)
    raise ValueError(f"stream not found: {stream_name}")


def extract_doc_payloads(path: Path) -> list[bytes]:
    if olefile is not None:
        try:
            with olefile.OleFileIO(str(path)) as ole:
                if ole.exists("WordDocument"):
                    return [ole.openstream("WordDocument").read()]
        except Exception:  # noqa: BLE001
            pass
    try:
        return [read_cfb_stream(path, "WordDocument")]
    except Exception as exc:  # noqa: BLE001
        if path.read_bytes().startswith(CFB_MAGIC):
            raise RuntimeError(f"could not read WordDocument stream from {path}") from exc
    return [path.read_bytes()]


def extract_doc(path: Path) -> str:
    candidates: list[str] = []
    for payload in extract_doc_payloads(path):
        decoded = payload.decode("utf-16-le", errors="ignore")
        text = scan_printable_text(decoded)
        if text:
            candidates.append(text)
    if not candidates:
        decoded = path.read_bytes().decode("latin-1", errors="ignore")
        candidates.append(scan_printable_text(decoded))
    return clean_text("\n".join(candidates))


def extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts: list[str] = []
    for idx, page in enumerate(reader.pages, 1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001
            text = f"[PDF page extraction error: {exc}]"
        parts.append(f"\n--- PAGE {idx} ---\n{text}")
    return clean_text("\n".join(parts))


def extract_pptx(path: Path) -> str:
    # Avoid requiring python-pptx. Read slide XML text boxes directly.
    ns = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "p": "http://schemas.openxmlformats.org/presentationml/2006/main",
    }
    parts: list[str] = []
    with zipfile.ZipFile(path) as zf:
        slide_names = sorted(
            [n for n in zf.namelist() if re.fullmatch(r"ppt/slides/slide\d+\.xml", n)],
            key=lambda n: int(re.search(r"slide(\d+)\.xml", n).group(1)),  # type: ignore[union-attr]
        )
        for slide_name in slide_names:
            slide_no = int(re.search(r"slide(\d+)\.xml", slide_name).group(1))  # type: ignore[union-attr]
            root = ET.fromstring(zf.read(slide_name))
            texts = [node.text for node in root.findall(".//a:t", ns) if node.text]
            if texts:
                parts.append(f"\n--- SLIDE {slide_no} ---\n" + "\n".join(texts))
    return clean_text("\n".join(parts))


def extract_text(path: Path) -> tuple[str, str]:
    ext = path.suffix.lower()
    if ext == ".docx":
        return extract_docx(path), "docx"
    if ext == ".doc":
        return extract_doc(path), "doc:ole-utf16-scan"
    if ext == ".pdf":
        return extract_pdf(path), "pdf"
    if ext == ".pptx":
        return extract_pptx(path), "pptx"
    if ext in TEXT_EXTS:
        for enc in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return clean_text(path.read_text(encoding=enc)), f"text:{enc}"
            except UnicodeDecodeError:
                continue
        return clean_text(path.read_text(encoding="utf-8", errors="replace")), "text:replace"
    raise ValueError(f"unsupported extension: {path.suffix}")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-root", required=True, type=Path)
    parser.add_argument("--out-dir", required=True, type=Path)
    args = parser.parse_args()

    source_root = args.source_root.resolve()
    out_dir = args.out_dir.resolve()
    text_dir = out_dir / "texts"
    text_dir.mkdir(parents=True, exist_ok=True)

    rows: list[dict[str, str]] = []
    for path in sorted(source_root.rglob("*")):
        if not path.is_file():
            continue
        ext = path.suffix.lower()
        if ext not in SUPPORTED_EXTS:
            rows.append(
                {
                    "status": "skipped",
                    "extractor": "",
                    "chars": "0",
                    "lines": "0",
                    "source": str(path),
                    "text_path": "",
                    "error": f"unsupported extension {path.suffix}",
                }
            )
            continue
        out_path = text_dir / f"{safe_name(path, source_root)}.txt"
        try:
            text, extractor = extract_text(path)
            out_path.write_text(text, encoding="utf-8")
            rows.append(
                {
                    "status": "ok",
                    "extractor": extractor,
                    "chars": str(len(text)),
                    "lines": str(text.count("\n") + (1 if text else 0)),
                    "source": str(path),
                    "text_path": str(out_path),
                    "error": "",
                }
            )
        except Exception as exc:  # noqa: BLE001
            rows.append(
                {
                    "status": "error",
                    "extractor": "",
                    "chars": "0",
                    "lines": "0",
                    "source": str(path),
                    "text_path": str(out_path),
                    "error": repr(exc),
                }
            )

    manifest = out_dir / "manifest.csv"
    with manifest.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(
            f, fieldnames=["status", "extractor", "chars", "lines", "source", "text_path", "error"]
        )
        writer.writeheader()
        writer.writerows(rows)

    ok = sum(1 for r in rows if r["status"] == "ok")
    errors = sum(1 for r in rows if r["status"] == "error")
    skipped = sum(1 for r in rows if r["status"] == "skipped")
    print(f"manifest={manifest}")
    print(f"ok={ok} errors={errors} skipped={skipped}")
    return 1 if errors else 0


if __name__ == "__main__":
    raise SystemExit(main())
