from __future__ import annotations

import csv
import importlib.util
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path("/Users/wanglifei/GaokaoPolitics/beijing-politics-sync")
RUN = ROOT / "reports/选必一_哲学宝典式重建_2026-05-16"
SOURCE_MD = RUN / "18_layout_score_layers_20260526/选必一_当代国际政治与经济_主观题术语宝典_学生版_FINAL_排版分值层次优化版_20260526.md"
OUT_DIR = RUN / "19_error_report_patch_20260527"
OUT_MD = OUT_DIR / "选必一_当代国际政治与经济_主观题术语宝典_学生版_FINAL_错误核对修正版_20260527.md"
OUT_DOCX = OUT_DIR / "选必一_当代国际政治与经济_主观题术语宝典_学生版_FINAL_错误核对修正版_20260527.docx"
OUT_PDF = OUT_DIR / "选必一_当代国际政治与经济_主观题术语宝典_学生版_FINAL_错误核对修正版_20260527.pdf"
OUT_ZIP = OUT_DIR / "xuanbiyi_error_patch_final_delivery_20260527.zip"
QA = OUT_DIR / "ERROR_PATCH_QA_20260527.md"
CORRECTION_LEDGER = OUT_DIR / "ERROR_CORRECTION_LEDGER_20260527.csv"
F_VERIFY = OUT_DIR / "F_CLASS_OCR_VERIFY_REPORT_20260527.md"
TRACE_CSV = OUT_DIR / "SCORE_LAYER_TRACE_20260527.csv"
UNSCORED_CSV = OUT_DIR / "UNSCORED_SAME_GROUP_BACKLOG_20260527.csv"
BASE_SCRIPT = ROOT / "scripts/build_xuanbiyi_layout_score_layers_20260526.py"


TITLE_REPLACEMENTS = [
    ("2026通州期中", "2026通州期末", "A1/A2"),
    ("2026丰台期中", "2026丰台期末", "A3"),
    ("2024顺义二模", "2024顺义思政二模", "A4"),
]

TEXT_REPLACEMENTS = [
    (
        "全球治理倡议正逢其时、指引方向、彰显担当",
        "全球治理倡议正逢其时、指引方向彰显担当",
        "B9",
    ),
    (
        "全球治理倡议正逢其时、指引方向，彰显担当",
        "全球治理倡议正逢其时、指引方向彰显担当",
        "B9",
    ),
    (
        "结合材料二，运用《当代国际政治与经济》知识，阐释中国对维护全球产供链韧性和稳定的政策主张。",
        "假如你是外交部发言人，请结合材料二，运用《当代国际政治与经济》知识，对上述提问做出回应。",
        "B1",
    ),
    (
        "结合材料，运用《当代国际政治与经济》知识，说明“五大工程”如何推动中拉双方在各自现代化征程上并肩前行，共筑中拉命运共同体新篇章。",
        "结合材料，运用《当代国际政治与经济》知识，阐释“五大工程”如何扩大中拉利益汇合点。",
        "B2",
    ),
    (
        "结合材料，运用《当代国际政治与经济》知识，说明\"五大工程\"如何推动中拉双方在各自现代化征程上并肩前行，共筑中拉命运共同体新篇章。",
        "结合材料，运用《当代国际政治与经济》知识，阐释\"五大工程\"如何扩大中拉利益汇合点。",
        "B2",
    ),
    (
        "结合材料，运用《当代国际政治与经济》的知识，分析中国与东盟经贸合作不断提质升级的原因和意义。",
        "结合材料，运用《当代国际政治与经济》的知识，分析中国与东盟经贸合作能够提质升级的原因。",
        "B3",
    ),
    (
        "中国与东盟经贸合作不断提质升级",
        "中国与东盟经贸合作能够提质升级",
        "B3",
    ),
    (
        "提质升级的原因和意义",
        "提质升级的原因",
        "B3",
    ),
    (
        "问“原因和意义”，必须先解释合作为什么能持续升级，再带出对双方和区域的价值。",
        "问“原因”，必须先解释合作为什么能够提质升级。",
        "B3",
    ),
    (
        "问\"原因和意义\"，要把合作能不断升级的根本前提解释清楚，再带出对双方和区域的价值。",
        "问\"原因\"，要把合作能够提质升级的根本前提解释清楚。",
        "B3",
    ),
    (
        "问\"原因和意义\"，要解释为什么经贸合作能持续升级，并指出双方都能在合作中获得新的发展空间。",
        "问\"原因\"，要解释为什么经贸合作能够提质升级。",
        "B3",
    ),
    (
        "问\"原因和意义\"中的方向意义，要把双边经贸升级延伸到对区域和全球开放合作格局的拉动，而不是停在\"互利共赢\"。",
        "问\"原因\"中的方向逻辑，要把双边经贸升级放回区域开放合作格局中解释，而不是停在\"互利共赢\"。",
        "B3",
    ),
    (
        "能够不断提质升级",
        "能够提质升级",
        "B3",
    ),
    (
        "结合材料，运用《当代国际政治与经济》知识，写一篇关于“同球共济”精神和中国行动的小论文。要求：自拟题目、观点明确、逻辑清晰，200字左右。",
        "结合材料，运用《当代国际政治与经济》知识，写一篇小论文。要求：自拟题目、观点明确、逻辑清晰，200字左右。",
        "B4",
    ),
    (
        '结合材料，运用《当代国际政治与经济》知识，写一篇关于"同球共济"精神和中国行动的小论文。',
        "结合材料，运用《当代国际政治与经济》知识，写一篇小论文。",
        "B4",
    ),
    (
        "结合材料，运用所学，分析中国—东盟自贸区3.0版",
        "结合材料，运用《当代国际政治与经济》知识，分析中国—东盟自贸区3.0版",
        "B5",
    ),
    (
        "面对新一轮前沿科技发展浪潮，结合材料，从经济角度分析",
        "结合材料，从经济角度分析",
        "B6",
    ),
    ("知识运用准确、贴切", "知识运用准确", "B7"),
    ("自由贸易试验区/自由贸易港", "自由贸易试验区、自由贸易试验港", "B8"),
    ("自由贸易试验区和自由贸易港", "自由贸易试验区和自由贸易试验港", "B8"),
    ("自由贸易港建设", "自由贸易试验港建设", "B8"),
    ("DeepSeek开源在140多个国家和地区下载登顶", "DeepSeek开源在全球140个国家和地区的应用市场下载排行榜同时登顶", "C2"),
    ("DeepSeek开源在140多国下载登顶", "DeepSeek开源在全球140个国家和地区的应用市场下载排行榜同时登顶", "C2"),
    (
        "题目让学生阐释中国对维护全球产供链韧性和稳定的政策主张，需要点出中国主张背后的最高价值底色。",
        "题目要求以外交部发言人身份回应提问，需要点出中国主张背后的最高价值底色。",
        "B1",
    ),
    (
        "问五大工程如何推动中拉现代化，经济规则层要说明稳定开放的外部经贸秩序。",
        "问“五大工程”如何扩大中拉利益汇合点，经济规则层要说明稳定开放的外部经贸秩序如何成为共同利益。",
        "B2",
    ),
    (
        "设问问五大工程如何推动双方现代化。",
        "设问问“五大工程”如何扩大中拉利益汇合点。",
        "B2",
    ),
    (
        "问“五大工程”如何推动双方现代化，要把工程安排转化为经济全球化中的产业链供应链支撑。",
        "问“五大工程”如何扩大中拉利益汇合点，要把工程安排转化为经济全球化中的产业链供应链支撑。",
        "B2",
    ),
    (
        "设问问\"五大工程\"如何\"共同谱写构建中拉命运共同体新篇章\"。要回答中国与拉美现代化协同的总主线，必须用\"中拉命运共同体\"作为整体目标——把团结／发展／文明／和平／民心五项工程串成一条主线。",
        "设问问\"五大工程\"如何扩大中拉利益汇合点。要把团结／发展／文明／和平／民心五项工程分别接到共同利益、合作领域和机制支撑，不能把材料里的结果句当成设问。",
        "B2",
    ),
    (
        "题目让学生说明\"五大工程\"如何推动中拉双方在各自现代化征程上并肩前行，要找出五项工程的总主线。",
        "题目让学生阐释\"五大工程\"如何扩大中拉利益汇合点，要把五项工程分别接到共同利益、合作领域和机制支撑。",
        "B2",
    ),
    (
        "共同谱写推动构建中拉命运共同体新篇章。",
        "扩大中拉利益汇合点，推动构建中拉命运共同体新篇章。",
        "B2",
    ),
    ("2025西城期末Q20", "2025西城期末Q20(2)", "D4"),
]

LAYER_REPLACEMENTS = [
    (
        "4.新秩序、国际关系民主化、多边主义、正确义利观、兼顾利益等任一；5.人类命运共同体；6.中国智慧、中国方案、大国担当。",
        "4.推动构建国际新秩序、倡导国际关系民主化、践行多边主义、坚持正确义利观、兼顾利益等任意一点给1分、共2分；5.人类命运共同体；6.贡献中国智慧、中国方案、勇于大国担当。",
        "D1/D2",
    ),
    (
        "总分结构：一层3分，两层6分，三层8分；每层按“总述—材料—关系”展开，不能只罗列材料或空写教材词。",
        "总分结构：一层3分，两层6分，三层8分；每一层采用“总—分”结构（处理什么关系—如何处理关系），不能只罗列材料或空写教材词。",
        "D3",
    ),
]

SCORE_MAP = {
    ("2026海淀一模", "Q20", None): ("8分", "f_ocr"),
    ("2026房山二模", "Q20", None): ("8分", "f_ocr"),
    ("2026朝阳期末", "Q20", None): ("8分", "f_ocr"),
    ("2026顺义二模", "Q20", None): ("8分", "f_ocr_doc"),
    ("2026丰台期末", "Q20", None): ("8分", "f_ocr"),
    ("2024海淀期中", "Q21", "2"): ("9分", "f_ocr"),
    ("2024东城二模", "Q20", None): ("7分", "f_ocr"),
    ("2026西城期末", "Q20", None): ("8分", "f_ocr_answer"),
}

LAYER_NOTES = {
    ("2026海淀一模", "Q20", None): "8分；角度1对我国3分，角度2对世界经济3分，角度3对全球治理2分。",
    ("2026房山二模", "Q20", None): "8分；普惠包容的经济全球化、平等有序的世界多极化、共同利益/共同发展、合作共赢/互利共赢、共商共建共享/全球治理观/真正的多边主义、数据要素全球流动各1分，结合材料论述弥合数据鸿沟、释放数据价值、繁荣数字经济2分。",
    ("2026朝阳期末", "Q20", None): "8分；和平发展、政治、经济、我国国家利益四个角度各2分，每个角度按背景1分+目标1分处理。",
    ("2026顺义二模", "Q20", None): "8分；可从共商共建共享、新型国际关系、独立自主的外交政策、世界经济发展的重要推动者、中国智慧的生动实践、让经济全球化更有活力等角度作答，按等级描述给分。",
    ("2024海淀期中", "Q21", "2"): "9分；变4分（世界之变2分、中国之变2分）；不变5分（国家性质、党对外交统一集中领导、国家利益、独立自主基本立场、和平共处五项原则等）。",
    ("2024东城二模", "Q20", None): "7分；等级题，观点、知识运用、论述逻辑与条理共同决定分值。",
    ("2026西城期末", "Q20", None): "8分；中国实践是什么2分，为什么参与该实践3分，实践效果3分。",
}

F_VERIFY_ROWS = [
    ("2026海淀一模Q20", "verified", "试卷OCR第7页：设问为“谈谈中国标准走出国门的意义”；细则OCR第6页：本题8分，三角度8分。"),
    ("2026房山二模Q20", "verified", "试卷OCR第8页：设问为“分析世界数据组织完善全球数据治理，服务全球数字经济发展的智慧”；答案页/细则：本题8分，1+1+1+1+1+1+2。"),
    ("2026朝阳期末Q20", "verified", "试卷OCR第7页和细则OCR第3页：设问为“中国特色大国外交为什么要更有作为”；本题8分，四角度各2分。"),
    ("2026顺义二模Q20", "verified", "试卷OCR第11页：设问为“谈谈对‘今天的中国，明天的世界’的理解”；评标doc可读：本题8分，等级题。"),
    ("2026丰台期末Q20", "verified_with_note", "期末细则PDF第64页为复练试题：设问为“阐释‘五大工程’如何扩大中拉利益汇合点”；本题8分。"),
    ("2024海淀期中Q21(2)", "verified", "试卷OCR第8页：设问为“分析新中国外交的‘变’与‘不变’”；细则OCR显示9分，变4分、不变5分。"),
    ("2024东城二模Q20", "verified", "试卷OCR第10页：设问为“三大倡议如何助力人类走向美好未来”；细则OCR显示本题7分等级题。"),
    ("2026西城期末Q20", "partly_verified", "试卷文件夹未找到试卷文件；参考答案PDF OCR第9页包含题面，细则OCR第4-5页显示本题8分。"),
]


def load_builder_module():
    spec = importlib.util.spec_from_file_location("xuanbiyi_builder_20260526", BASE_SCRIPT)
    if spec is None or spec.loader is None:
        raise RuntimeError("cannot load base builder")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def title_parts(title_line: str) -> tuple[str | None, str | None, str | None]:
    m = re.match(r"^### \d+\.\s*(\d{4}[^Q\n]+?)Q(\d+)(?:\((\d+)\))?", title_line.strip())
    if not m:
        return None, None, None
    return m.group(1).strip(), "Q" + m.group(2), m.group(3)


def split_examples(md: str) -> list[tuple[int, int, str]]:
    matches = list(re.finditer(r"(?m)^### \d+\. .*$", md))
    out: list[tuple[int, int, str]] = []
    for idx, match in enumerate(matches):
        start = match.start()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(md)
        out.append((start, end, md[start:end]))
    return out


def score_for(suite: str | None, q: str | None, part: str | None):
    if not suite or not q:
        return None
    return SCORE_MAP.get((suite, q, part)) or SCORE_MAP.get((suite, q, None))


def layer_for(suite: str | None, q: str | None, part: str | None) -> str | None:
    if not suite or not q:
        return None
    return LAYER_NOTES.get((suite, q, part)) or LAYER_NOTES.get((suite, q, None))


def apply_global_replacements(text: str, ledger: list[dict[str, str]]) -> str:
    for old, new, item in TITLE_REPLACEMENTS + TEXT_REPLACEMENTS + LAYER_REPLACEMENTS:
        count = text.count(old)
        if count:
            text = text.replace(old, new)
        ledger.append({"item": item, "action": "replace", "old": old, "new": new, "count": str(count), "status": "applied" if count else "no_match"})
    return text


def remove_boundary_blocks(text: str, ledger: list[dict[str, str]]) -> str:
    out: list[str] = []
    cursor = 0
    removed = 0
    for start, end, block in split_examples(text):
        out.append(text[cursor:start])
        title = block.splitlines()[0]
        if "2025海淀期末Q22" in title:
            removed += 1
        else:
            out.append(block)
        cursor = end
    out.append(text[cursor:])
    ledger.append({"item": "C1", "action": "remove_block", "old": "2025海淀期末Q22", "new": "removed_from_student_main_body", "count": str(removed), "status": "applied" if removed else "no_match"})
    return "".join(out)


def enhance_blocks(text: str, ledger: list[dict[str, str]]) -> tuple[str, list[dict[str, str]]]:
    parts: list[str] = []
    trace: list[dict[str, str]] = []
    cursor = 0
    scored_blocks = 0
    layer_added = 0
    for start, end, block in split_examples(text):
        parts.append(text[cursor:start])
        cursor = end
        lines = block.rstrip("\n").splitlines()
        suite, q, part = title_parts(lines[0])
        score_tuple = score_for(suite, q, part)
        score = score_tuple[0] if score_tuple else ""
        score_method = score_tuple[1] if score_tuple else ""
        layer_note = layer_for(suite, q, part)
        has_layer = any(line.startswith("【得分层次】") for line in lines)
        new_lines: list[str] = []
        has_same_group = False
        for line in lines:
            if line.startswith("【同题组】"):
                has_same_group = True
                if score and "本题" not in line:
                    line = line.replace("【同题组】", f"【同题组】（本题{score}）", 1)
                    scored_blocks += 1
            new_lines.append(line)
            if line.startswith("【设问】") and layer_note and not has_layer:
                new_lines.append(f"【得分层次】{layer_note}")
                has_layer = True
                layer_added += 1
        question = next((line.replace("【设问】", "", 1).strip() for line in new_lines if line.startswith("【设问】")), "")
        trace.append(
            {
                "title": lines[0].lstrip("# ").strip(),
                "suite": suite or "",
                "question": q or "",
                "part": part or "",
                "same_group": "YES" if has_same_group else "NO",
                "score": score,
                "score_method": score_method,
                "layer": "YES" if has_layer else "NO",
                "layer_method": "f_ocr_added" if layer_note else "",
                "question_text": question,
            }
        )
        parts.append("\n".join(new_lines).rstrip() + "\n\n")
    parts.append(text[cursor:])
    ledger.append({"item": "F", "action": "annotate_scores_from_ocr", "old": "", "new": "SCORE_MAP", "count": str(scored_blocks), "status": "applied"})
    ledger.append({"item": "F", "action": "add_layers_from_ocr", "old": "", "new": "LAYER_NOTES", "count": str(layer_added), "status": "applied"})
    return re.sub(r"\n{3,}", "\n\n", "".join(parts)).strip() + "\n", trace


def write_csv(path: Path, rows: list[dict[str, str]], fieldnames: list[str]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def write_f_verify() -> None:
    lines = ["# F类 OCR 核验报告", ""]
    for item, status, note in F_VERIFY_ROWS:
        lines.append(f"- {item}: {status}。{note}")
    lines.append("")
    lines.append("- OCR目录: `f_class_ocr/`")
    lines.append("- OCR清单: `f_class_ocr/F_CLASS_OCR_MANIFEST_20260527.csv`")
    F_VERIFY.write_text("\n".join(lines) + "\n", encoding="utf-8")


def patch_title_page(builder):
    def add_title_page(doc: Document) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(82)
        r = p.add_run("选择性必修一《当代国际政治与经济》主观题术语宝典")
        builder.set_east_asia_font(r)
        r.bold = True
        r.font.size = Pt(25)
        r.font.color.rgb = builder.ACCENT
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(12)
        r2 = p2.add_run("学生厚版 · 错误核对修正版")
        builder.set_east_asia_font(r2)
        r2.font.size = Pt(13)
        r2.font.color.rgb = builder.MUTED
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_before = Pt(24)
        r3 = p3.add_run("飞哥正志讲堂")
        builder.set_east_asia_font(r3)
        r3.bold = True
        r3.font.size = Pt(11)
        doc.add_section(builder.WD_SECTION.NEW_PAGE)

    builder.add_title_page = add_title_page


def pollution_counts(text: str) -> dict[str, int]:
    markers = ["提醒", "提示", "边界", "使用方法", "审计", "审核", "GPT", "Claude", "Opus", "Codex", "debug", "TODO", "SHA", "/Users/", "reports/"]
    return {m: text.count(m) for m in markers}


def write_qa(ledger_rows: list[dict[str, str]], trace_rows: list[dict[str, str]], pdf_info: dict[str, object]) -> None:
    text = OUT_MD.read_text(encoding="utf-8")
    counts = pollution_counts(text)
    applied = sum(1 for r in ledger_rows if r["status"] == "applied")
    lines = [
        "# 选必一错误核对修正版 QA",
        "",
        f"- source_md: `{SOURCE_MD}`",
        f"- output_md: `{OUT_MD}`",
        f"- output_docx: `{OUT_DOCX}`",
        f"- output_pdf: `{OUT_PDF}`",
        f"- correction_ledger: `{CORRECTION_LEDGER}`",
        f"- f_class_verify: `{F_VERIFY}`",
        "",
        "## Patch Counts",
        "",
        f"- correction actions applied: {applied}",
        f"- C1 removed blocks: {next((r['count'] for r in ledger_rows if r['item']=='C1'), '0')}",
        f"- trace rows: {len(trace_rows)}",
        f"- same-group scored examples: {sum(1 for r in trace_rows if r['same_group']=='YES' and r['score'])}",
        f"- same-group unscored examples: {sum(1 for r in trace_rows if r['same_group']=='YES' and not r['score'])}",
        f"- 得分层次 count: {text.count('【得分层次】')}",
        f"- 为什么能想到 count: {text.count('【为什么能想到】')}",
        f"- 卷面句 count: {text.count('【卷面句】')}",
        "",
        "## Remaining Pollution Counts",
        "",
    ]
    for marker, count in counts.items():
        lines.append(f"- {marker}: {count}")
    lines.extend(["", "## PDF Render QA", "", f"- page_count: {pdf_info['page_count']}", "- rendered_samples:"])
    for item in pdf_info["rendered"]:
        lines.append(f"  - `{item}`")
    QA.write_text("\n".join(lines) + "\n", encoding="utf-8")


def zip_delivery() -> None:
    with zipfile.ZipFile(OUT_ZIP, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in [OUT_MD, OUT_DOCX, OUT_PDF, QA, CORRECTION_LEDGER, F_VERIFY, TRACE_CSV, UNSCORED_CSV, OUT_DIR / "f_class_ocr/F_CLASS_OCR_MANIFEST_20260527.csv"]:
            if path.exists():
                zf.write(path, path.relative_to(OUT_DIR) if path.is_relative_to(OUT_DIR) else path.name)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    source = SOURCE_MD.read_text(encoding="utf-8")
    ledger: list[dict[str, str]] = []
    text = apply_global_replacements(source, ledger)
    text = remove_boundary_blocks(text, ledger)
    text, trace_rows = enhance_blocks(text, ledger)
    OUT_MD.write_text(text, encoding="utf-8", newline="\n")

    write_csv(CORRECTION_LEDGER, ledger, ["item", "action", "old", "new", "count", "status"])
    write_csv(TRACE_CSV, trace_rows, ["title", "suite", "question", "part", "same_group", "score", "score_method", "layer", "layer_method", "question_text"])
    backlog = [r for r in trace_rows if r["same_group"] == "YES" and not r["score"]]
    write_csv(UNSCORED_CSV, backlog, ["title", "suite", "question", "part", "same_group", "score", "score_method", "layer", "layer_method", "question_text"])
    write_f_verify()

    builder = load_builder_module()
    patch_title_page(builder)
    builder.build_docx(OUT_MD, OUT_DOCX)
    builder.build_pdf(OUT_MD, OUT_PDF)
    pdf_info = builder.render_pdf_samples(OUT_PDF, OUT_DIR / "render_qa_pdf_20260527")
    write_qa(ledger, trace_rows, pdf_info)
    zip_delivery()
    for path in [OUT_MD, OUT_DOCX, OUT_PDF, OUT_ZIP, QA, CORRECTION_LEDGER, F_VERIFY]:
        print(path)


if __name__ == "__main__":
    main()
