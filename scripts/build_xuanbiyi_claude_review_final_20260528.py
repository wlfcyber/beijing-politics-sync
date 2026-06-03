from __future__ import annotations

import csv
import shutil
from pathlib import Path

from docx import Document


RUN_DIR = Path("/Users/wanglifei/GaokaoPolitics/beijing-politics-sync/reports/选必一_哲学宝典式重建_2026-05-16/21_claude_review_patch_20260528")
REVIEW_DOCX = Path("/Users/wanglifei/Desktop/Codex宝典修改稿_Claude复核_20260528.docx")
BASE_DOCX = Path("/Users/wanglifei/GaokaoPolitics/beijing-politics-sync/reports/选必一_哲学宝典式重建_2026-05-16/20_review237_verify_final_20260528/选必一_当代国际政治与经济_主观题术语宝典_237条逐条核实最终稿_20260528.docx")
OUT_DOCX = RUN_DIR / "选必一_当代国际政治与经济_主观题术语宝典_237条逐条核实二次复核最终稿_20260528.docx"
OUT_MD = RUN_DIR / "选必一_当代国际政治与经济_主观题术语宝典_237条逐条核实二次复核最终稿_20260528.md"


REVIEW_ITEMS = [
    ("1", "明显改善", "2024朝阳一模Q21", "百年变局与外部风险挑战", "ACCEPT_NO_NEW_PATCH", "Codex 已改为细则原文“外部环境日趋严峻、风险挑战增多”，保留。"),
    ("2", "明显改善", "2026丰台期末Q20", "维护多边贸易体制", "APPLY_TITLE_FAMILY_PATCH", "卷面句已改准，但核心标题仍残留“以世界贸易组织为核心”；改为家族标题“维护多边贸易体制”。"),
    ("3", "明显改善", "2026朝阳期末Q20", "建设开放型世界经济，扩大开放合作", "ACCEPT_NO_NEW_PATCH", "Codex 已收紧为“开放合作、互利共赢和世界经济共同繁荣”，保留。"),
    ("4", "明显改善", "2026西城二模Q19(2)", "坚持开放合作和包容普惠", "ACCEPT_NO_NEW_PATCH", "Codex 已改为“弥合全球数字鸿沟、释放数据价值、繁荣数字经济”，保留。"),
    ("5", "明显改善", "2024朝阳一模Q21", "坚持总体国家安全观，统筹发展与安全", "ACCEPT_NO_NEW_PATCH", "Codex 已改为“外部环境日趋严峻、风险挑战增多”，保留。"),
    ("6", "明显改善", "2026门头沟一模Q21", "中国是负责任大国并勇于担当", "ACCEPT_NO_NEW_PATCH", "Codex 已改为“互利共赢的开放战略”，保留。"),
    ("7", "副作用", "2026丰台期末Q20", "维护以世界贸易组织为核心的多边贸易体制标题", "APPLY_TITLE_FAMILY_PATCH", "同第2项。复核意见成立，但最终标题采用可同时覆盖石景山、丰台两题的“维护多边贸易体制”。"),
    ("8", "持平", "2024海淀二模Q18(1)", "和平与发展仍是时代主题", "REVIEWED_KEEP", "Claude 判定两版均为合理外延，无需新改。"),
    ("9", "持平", "2025东城期末Q20", "反对贸易保护主义", "REVIEWED_KEEP", "Claude 判定 Codex 用“市场竞争和市场秩序”贴近细则，无需新改。"),
    ("10", "持平", "2026丰台期末Q20", "维护开放合作的国际经济环境", "REVIEWED_KEEP", "Claude 判定新版与原版基本同义；上一轮已删去无支撑保护主义扩展，保留。"),
    ("11", "持平", "2024海淀二模Q18(1)", "推动构建人类命运共同体", "REVIEWED_KEEP", "细则只列角度，Claude 判定两版均合理，无需新改。"),
    ("12", "持平", "2024海淀二模Q18(1)", "国际组织", "REVIEWED_KEEP", "细则不展开，Claude 判定两版均属合理外延，无需新改。"),
    ("13", "持平", "2024朝阳一模Q21", "坚持独立自主、自力更生", "REVIEWED_KEEP", "Codex 已删“小院高墙、政治化”，改成“外部风险挑战”；保留。"),
]


PATCHES = [
    {
        "scope": "subtitle",
        "old": "学生厚版 · 237条逐条核实最终版（2026.5.28 · 共104术语/237条目）",
        "new": "学生厚版 · 237条逐条核实二次复核最终版（2026.5.28 · 共104术语/237条目）",
        "reason": "标明二次复核最终版，不在学生稿中出现后台模型名。",
    },
    {
        "scope": "core_heading",
        "old": "核心答题点：维护以世界贸易组织为核心的多边贸易体制（出现2次）",
        "new": "核心答题点：维护多边贸易体制（出现2次）",
        "reason": "同一核心覆盖2026石景山一模“共同维护多边贸易体制”和2026丰台期末“坚定维护多边贸易体制”；家族标题不再绑定丰台细则未明示的“以世界贸易组织为核心”。",
    },
    {
        "scope": "term_family",
        "old": "【术语家族】 “维护多边贸易体制”“维护多边贸易体系”“维护以世界贸易组织为核心的多边贸易体制”在本稿中作为同一采分家族处理；作答时依据材料任选最贴近的一种表述，不再拆成三个核心点。",
        "new": "【术语家族】 “共同维护多边贸易体制”“坚定维护多边贸易体制”“维护多边贸易体制”在本稿中作为同一采分家族处理；2026石景山一模可结合世界贸易组织改革方向、最惠国待遇、非歧视原则展开，2026丰台期末使用细则原词“坚定维护多边贸易体制”。",
        "reason": "补足两题来源差异，避免学生把丰台题误写成“以世界贸易组织为核心”。",
    },
    {
        "scope": "shijingshan_action",
        "old": "答题动作：先点出维护 WTO 多边贸易体制，再接最惠国待遇、非歧视原则，最后落到破除壁垒、稳定亚太开放合作。",
        "new": "答题动作：先点出共同维护多边贸易体制，再接世界贸易组织改革正确方向、最惠国待遇、非歧视原则，最后落到破除壁垒、稳定亚太开放合作。",
        "reason": "改回2026石景山一模题面可直接核验的表述。",
    },
    {
        "scope": "shijingshan_sentence",
        "old": "【卷面句】 中国倡议维护以世界贸易组织为核心的多边贸易体制，坚持最惠国待遇、非歧视等基本原则，破除贸易壁垒，为亚太开放合作提供稳定规则支撑。",
        "new": "【卷面句】 中国倡议共同维护多边贸易体制，坚持世界贸易组织改革的正确方向和最惠国待遇、非歧视等基本原则，破除贸易壁垒，为亚太开放合作提供稳定规则支撑。",
        "reason": "石景山一模原题写“共同维护多边贸易体制”“坚持世界贸易组织改革的正确方向”，不写“以世界贸易组织为核心”。",
    },
    {
        "scope": "same_group_label",
        "old": "维护以世界贸易组织为核心的多边贸易体制",
        "new": "维护多边贸易体制",
        "reason": "统一同题组索引到新的核心家族标题。",
    },
    {
        "scope": "speed_memo_label",
        "old": "通过区域经济合作维护自由贸易和以世界贸易组织为核心的多边贸易体制",
        "new": "通过区域经济合作维护自由贸易和多边贸易体制",
        "reason": "同步六大要素极简速记版，避免正文外残留旧核心标题。",
    },
]


def extract_review_text() -> None:
    doc = Document(str(REVIEW_DOCX))
    lines = ["# CLAUDE_REVIEW_TEXT_EXTRACT", ""]
    for i, p in enumerate(doc.paragraphs, 1):
        text = p.text.strip()
        if text:
            lines.append(f"{i}. {text}")
    (RUN_DIR / "CLAUDE_REVIEW_TEXT_EXTRACT.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_review_ledgers() -> None:
    with (RUN_DIR / "CLAUDE_REVIEW_ITEM_LEDGER.csv").open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["review_id", "bucket", "question", "issue", "decision", "note"])
        writer.writeheader()
        for review_id, bucket, question, issue, decision, note in REVIEW_ITEMS:
            writer.writerow({
                "review_id": review_id,
                "bucket": bucket,
                "question": question,
                "issue": issue,
                "decision": decision,
                "note": note,
            })


def replace_docx(doc: Document) -> list[dict[str, str]]:
    results: list[dict[str, str]] = []
    for patch in PATCHES:
        applied = 0
        for p in doc.paragraphs:
            if patch["old"] in p.text:
                p.text = p.text.replace(patch["old"], patch["new"])
                applied += 1
        row = dict(patch)
        row["applied"] = "YES" if applied else "NO"
        row["applied_count"] = str(applied)
        results.append(row)
    return results


def docx_to_md(docx_path: Path, md_path: Path) -> None:
    doc = Document(str(docx_path))
    lines: list[str] = []
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            lines.append("")
            continue
        if idx == 0:
            lines.append(f"# {text}")
        elif text in {"时代背景", "理论", "经济全球化", "政治多极化", "中国", "联合国", "六大要素术语极简速记版", "目录"}:
            lines.append(f"\n## {text}")
        elif text.startswith("核心答题点："):
            lines.append(f"\n### {text}")
        else:
            lines.append(text)
    md_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_patch_ledger(results: list[dict[str, str]]) -> None:
    with (RUN_DIR / "CLAUDE_REVIEW_APPLIED_PATCH_LEDGER.csv").open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["scope", "old", "new", "reason", "applied", "applied_count"])
        writer.writeheader()
        writer.writerows(results)


def main() -> None:
    RUN_DIR.mkdir(parents=True, exist_ok=True)
    extract_review_text()
    write_review_ledgers()
    shutil.copy2(BASE_DOCX, OUT_DOCX)
    doc = Document(str(OUT_DOCX))
    results = replace_docx(doc)
    doc.save(str(OUT_DOCX))
    docx_to_md(OUT_DOCX, OUT_MD)
    write_patch_ledger(results)


if __name__ == "__main__":
    main()
