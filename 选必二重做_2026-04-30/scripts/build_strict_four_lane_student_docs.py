#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


BASE = Path("/Users/wanglifei/Desktop/北京高考政治/选必二重做_2026-04-30")
SOURCE_DELIVERY = BASE / "delivery"
CLOSURE = BASE / "strict_four_lane_closure_2026-05-04"
OUT = CLOSURE / "delivery"

FRAMEWORK_SRC = SOURCE_DELIVERY / "选必二法律框架踩分逻辑_主观选择分列版_2026-05-04.md"
SITUATION_SRC = SOURCE_DELIVERY / "选必二考过情境细则汇总_主观选择分列版_2026-05-04.md"

FRAMEWORK_MD = OUT / "选必二法律框架踩分逻辑_零基础强补版_2026-05-04.md"
SITUATION_MD = OUT / "选必二考过情境细则汇总_零基础强补版_2026-05-04.md"
FRAMEWORK_DOCX = OUT / "选必二法律框架踩分逻辑_零基础强补版_2026-05-04.docx"
SITUATION_DOCX = OUT / "选必二考过情境细则汇总_零基础强补版_2026-05-04.docx"


FRAMEWORK_PATCH = """## 先学这个：新题争点怎么自己生成

这版最重要的新增内容，是把“照旧题找相似”升级成“看到陌生材料也能自己生成争点”。法律主观题的争点，就是本题希望你证明或否定的那一句话。

### 争点三种锁定法

- 设问反推法：设问问“为什么支持、是否有效、谁担责、如何处理”，把问题改成一句判断句，就是争点。例：“法院为什么支持赔偿”反推为“某行为是否构成侵权并应赔偿”。
- 双方主张法：材料里两方各说一套，争点就在两方冲突处。例：消费者说“虚假宣传应赔”，经营者说“只是广告语”，争点就是宣传内容是否进入交易并构成欺诈或违约。
- 要件模糊法：材料最像某条规则，但有一个条件不确定，那个条件就是争点。例：竞业限制题里最模糊的是“劳动者是否负有保密义务、单位是否给补偿、期限范围是否合理”。

### 主观题三段法

- 规则依据：先写教材规则或法律原则，不先喊价值口号。
- 材料事实对应：把材料中的具体事实贴到规则条件上，至少写出一到两个关键词。
- 结果落点：最后落到有效/无效、归谁、谁赔、怎么救济、为什么有意义。

考场句式：本案中，甲与乙之间形成……关系。判断……应看……。材料中……，说明……，因此……。

### 意义类三层模板

- 法律层：维护当事人合法权益，明确权利义务边界，提供有效救济。
- 价值层：弘扬诚信、公平、友善、敬业、孝亲敬老等价值要求。
- 制度层：定分止争，规范市场秩序，推动良法善治，保护公共利益。

意义题不要脱离本案。先写本案处理了什么争议，再写价值和制度。

反例：只写“弘扬诚信、维护秩序、推动良法善治”，却不交代本案处理了什么争议，这种写法不能踩住主观题核心分。

## 五域定位决策树

先问“本题最核心的法律关系是什么”，不要被材料里的热词带跑。

- 如果核心是买卖、服务、平台下单、退款、格式条款，进“交易关系与消费者保护”。
- 如果核心是管理、考勤、工资、辞退、竞业限制、劳动仲裁，进“劳动用工与职业边界”。
- 如果核心是房、物、共有、相邻、继承、赡养、遗赠扶养，进“财产、相邻、继承与家庭”。
- 如果核心是身体受伤、肖像、名誉、隐私、个人信息，进“人格权与侵权责任”。
- 如果核心是作品、商标、商业秘密、数据、竞争秩序，进“创新竞争与公共救济”里的创新竞争线。
- 如果核心是调解、仲裁、诉讼、举证、执行、司法确认、公益保护，进“创新竞争与公共救济”里的程序公益线。

一道题可能跨域。主域看“设问最终要判什么”，副域只作补充。

第五域一定要硬分流：创新竞争线解决“成果和秩序怎么保护”，程序公益线解决“权利和公共利益怎么救济”。两条线不能共用同一套主观题模板。

## 高频要件卡片

这些不是让你法考化，而是防止新题换个细节就丢分。每张卡只记“触发词 -> 必查条件 -> 落点”。

### 合同成立与违约

- 触发词：报价、承诺、通知、付款、交付、解除、赔偿。
- 必查条件：要约内容是否具体确定；承诺是否到达；合同是否有效；履行是否符合约定。
- 落点：合同成立、合同有效、继续履行、解除、赔偿损失或违约责任。

### 消费欺诈与惩罚性赔偿

- 触发词：虚假宣传、普通消费欺诈三倍赔偿、食品安全十倍赔偿、疗效承诺、隐瞒真实信息。
- 必查条件：经营者信息是否真实全面；是否影响消费者真实选择；是否属于生活消费；是否有欺诈或食品安全问题。
- 落点：退款、赔偿、惩罚性赔偿、维护消费者知情权和公平交易权。

### 格式条款

- 触发词：最终解释权、不得退款、不得转让、概不负责、小字提示。
- 必查条件：是否由一方预先拟定反复使用；是否免除己方责任、加重对方责任、排除对方主要权利；是否尽提示说明义务。
- 落点：有效、无效、可撤销或不成为合同内容。

### 竞业限制

- 触发词：离职后不得入职同行、保密协议、商业秘密、违约金。
- 必查条件：劳动者是否属于高管、高级技术人员或其他负有保密义务人员；期限是否不超过二年，这是竞业限制期限的法定上限；范围、地域、行业是否合理；用人单位是否给予经济补偿。
- 落点：保护商业秘密，但不能不当限制普通劳动者择业权。

### 遗赠扶养与继承

- 触发词：老人、扶养、遗产归谁、遗嘱、居委会、亲属争议。
- 必查条件：是否有遗赠扶养协议，教材题通常从书面协议和实际履行入手；扶养人是否为继承人以外的组织或个人；是否实际履行扶养义务；遗嘱或协议是否真实合法。
- 落点：协议有效、遗产归属、酌情分得遗产、赡养义务和家庭美德。

### 一般侵权

- 触发词：受伤、损失、道歉、赔偿、过错、因果关系。
- 必查条件：行为、损害、因果关系、过错是否成立；受害人是否也有过错。
- 落点：停止侵害、赔礼道歉、赔偿损失或减轻责任。

### 特殊侵权识别卡

- 触发词：动物致害、高空抛物、产品缺陷、环境污染、公共场所受伤。
- 必查条件：是无过错、过错推定还是安全保障义务；是否存在被侵权人故意、重大过失、第三人原因等免责或减责事由。
- 落点：赔偿、补偿、减责、免责或举证责任倒置。

这张卡不要求背完整归责体系，只用来识别特殊规则、免责减责和举证边界。

### 知识产权与不正当竞争

- 触发词：作品、署名、图片、AI作品、商标、包装装潢、商业秘密、数据抓取、混淆。
- 必查条件：是否有受保护对象；是否未经许可使用；是否造成混淆或攀附商誉；是否破坏公平竞争秩序。
- 落点：停止侵权、消除影响、赔偿损失、保护创新和市场秩序。

### 程序救济

- 触发词：调解、仲裁、诉讼、司法确认、举证、执行。
- 必查条件：选择哪条路径；是否自愿；是否有强制执行效力；谁主张谁举证；仲裁与诉讼能否重复选择。
- 落点：化解纠纷、定分止争、保障权利实现。

### 公益保护

- 触发词：生态修复、公共利益、检察机关、社会组织、司法建议、公开道歉。
- 必查条件：损害是否涉及公共利益；起诉主体是否适格；高考选必二主观题以民事公益诉讼为主；是否需要修复、公开道歉或承担公共利益保护责任。
- 落点：保护公共利益、修复生态、督促依法治理。

## 六组最容易混的边界

- 违约责任 vs 侵权责任：合同履行不符合约定写违约；人身、财产权益受损且有行为、损害、因果、过错写侵权；同一材料可能两条线并行。
- 合同无效 vs 可撤销 vs 条款无效：违法、公序良俗等可能无效；欺诈胁迫重大误解可能可撤销；格式条款不公平可能只是该条款无效。
- 消费者保护 vs 普通合同：生活消费场景才强调消费者权利和惩罚性赔偿；明知问题后大量加购不能无限扩张消费者保护。
- 劳动关系 vs 普通合作：看管理、报酬、业务组成和从属性，不只看协议名称。
- 一般侵权 vs 特殊侵权：一般侵权查过错；特殊侵权先识别无过错、过错推定或安全保障义务，再查免责减责。
- 民事公益诉讼 vs 行政机关履职：选必二主观题优先按民事公益诉讼处理污染者、侵权者；如果题面明确出现行政机关不履职，再按题面要求说明，不要把一个程序话术硬套到另一个程序。

## 选择题使用边界

选择题区是识别池，不是主观题模板。主观题不能直接照搬“某项不对”这种排错句。正确用法是：从选择题里记住易错边界，再回到主观题三段法写“规则依据、材料事实对应、结果落点”。

三段法是主观题主线；“主体关系、事实条件、法律后果”只是选择题排错三件套，不要把选择题话术整段带进主观题答题卡。
"""


SITUATION_PATCH = """## 先学这个：怎样用情境表做新题

情境表不是让你背完整旧题，而是让你训练“材料故事 -> 争点 -> 踩分句”的转换。

### 每道情境先读三处

- 起因：谁和谁发生了什么法律关系。
- 经过：哪一个事实引起争议，谁主张什么。
- 结果：题目要你落到责任、效力、归属、程序还是意义。

### 争点一句话

把争点改写成“本题到底要判什么”。例如：不是“直播买手串”，而是“直播承诺是否进入合同内容，交付材质不符是否要按约赔偿”。

### 细则使用法

- 先看细则里的规则词：要约、承诺、格式条款、侵权责任、竞业限制、司法确认等。
- 再看细则里的材料词：直播承诺、下单付款、公开监控、平台派单、签收退款等。
- 最后看细则里的落点词：有效、无效、赔偿、退款、停止侵权、确认归属、强制执行。

考场上的踩分顺序就是：规则词 -> 材料词 -> 落点词。先写规则，再贴材料，最后写本题的具体处理结果。

### 主观题和选择题分开用

主观题情境下面的细则可以训练踩分句。选择题情境只训练排错和边界，不要把选择题排错话术当成主观题大段答案。

## 新题迁移提醒

如果新题与旧题相似，只能说明入口相似；真正得分还要查有没有新要件。特别注意：

- 竞业限制要补查保密人员、期限、范围和经济补偿。
- 遗赠扶养要补查书面形式、主体和实际履行。
- 无过错责任也要补查免责或减责事由，例如受害人故意、重大过失、第三人原因、管理人已经尽到安全保障义务等。
- 消费题要区分违约、欺诈、知情权和惩罚性赔偿。
- 公益题要区分民事公益诉讼和行政公益诉讼。
"""


def normalize_student_language(text: str) -> str:
    text = text.replace("把答案落到责任、效力、归属、程序或价值", "把答案落到本题的具体处理结果")
    text = text.replace("信赖利益", "合理信任")
    text = text.replace(
        "本案重点是录用通知形成的合理信任和诚实信用，不是平等就业权受侵害",
        "本案重点是用人单位发出明确录用通知后应遵守诚实信用原则，不能无故取消录用，不是平等就业权受侵害",
    )
    return text


def normalize_framework_postpatch(text: str) -> str:
    text = normalize_student_language(text)
    text = text.replace(
        "情境节点：消费者权益、虚假宣传、格式条款与平台规则\n\n- 情境逻辑：",
        "情境节点：消费者权益、虚假宣传、格式条款与平台规则（主域：交易关系与消费者保护；在其他域出现时只作副线索）\n\n- 情境逻辑：",
    )
    text = text.replace(
        "- 节点说明：这是高考常见组合考点，不是法律理论分类。它把虚假宣传、格式条款、平台规则等放在同一题型节点，是因为命题常在一个买卖或服务合同中同时考经营者信息优势、消费者权利和责任承担。",
        "- 节点说明：这是高考常见组合考点，不是法律理论分类。主域始终是交易关系与消费者保护；若在人格权或创新竞争材料中出现，只表示材料跨域，主观题仍先按交易、消费者权利和经营者责任处理。",
    )
    text = text.replace(
        "情境节点：继承、遗嘱、遗赠扶养协议与赡养义务\n\n- 情境逻辑：",
        "情境节点：继承、遗嘱、遗赠扶养协议与赡养义务（主域：财产、相邻、继承与家庭；在程序题中只作案例之一）\n\n- 情境逻辑：",
    )
    text = text.replace(
        "- 节点踩分词：法定继承；遗嘱；遗赠扶养协议；扶养义务；赡养义务；遗产归属；尊老敬老；友善美德",
        "- 节点踩分词：法定继承；遗嘱；遗赠扶养协议；扶养义务；赡养义务；遗产归属；尊老敬老；友善美德\n\n- 节点说明：继承题的主线是身份、协议、遗嘱、扶养事实和遗产归属；如果它出现在纠纷解决题中，程序只是外壳，不能把继承规则吞掉。",
    )
    text = text.replace(
        "### 创新竞争与公共救济\n\n作答总逻辑：先看核心是创新成果、竞争秩序、程序救济还是公共利益；知识产权和不正当竞争保护的是创新与市场秩序，程序和公益诉讼负责把权利救济、公共利益保护落地。",
        "### 创新竞争与公共救济\n\n作答总逻辑：先硬分流。创新竞争线解决作品、标识、商业秘密、数据和公平竞争秩序；程序公益线解决调解、仲裁、诉讼、举证、执行和公共利益保护。两条线不能共用同一套主观题模板。",
    )
    return text


def patched_framework() -> str:
    text = FRAMEWORK_SRC.read_text(encoding="utf-8")
    text = text.replace(
        "# 选必二《法律与生活》框架踩分逻辑",
        "# 选必二《法律与生活》框架踩分逻辑",
        1,
    )
    marker = "## 第一部分 主观题框架"
    if marker not in text:
        raise RuntimeError("framework marker not found")
    text = text.replace(marker, FRAMEWORK_PATCH + "\n\n" + marker, 1)
    return normalize_framework_postpatch(text)


def normalize_situation_postpatch(text: str) -> str:
    return normalize_student_language(text)


def patched_situation() -> str:
    text = SITUATION_SRC.read_text(encoding="utf-8")
    marker = "## 第一部分 主观题情境"
    if marker not in text:
        raise RuntimeError("situation marker not found")
    text = text.replace(marker, SITUATION_PATCH + "\n\n" + marker, 1)
    return normalize_situation_postpatch(text)


def set_run_font(run, size: float = 10.5, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_style_font(doc: Document) -> None:
    styles = doc.styles
    for style_name, east_asia, size, bold in [
        ("Normal", "宋体", 10.5, False),
        ("Title", "黑体", 22, True),
        ("Heading 1", "黑体", 16, True),
        ("Heading 2", "黑体", 13, True),
        ("Heading 3", "黑体", 11.5, True),
        ("Heading 4", "黑体", 10.5, True),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
        style.font.size = Pt(size)
        style.font.bold = bold


def add_page_number(section) -> None:
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("第 ")
    set_run_font(run, 9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run2 = p.add_run(" 页")
    set_run_font(run2, 9)


def configure_doc(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(60)
    section.right_margin = Pt(60)
    header = section.header.paragraphs[0]
    header.text = title
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        set_run_font(run, 9, False)
    add_page_number(section)
    set_style_font(doc)


def add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    bold = text.startswith("【") and text.endswith("】")
    set_run_font(run, 10.5, bold)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08


def md_to_docx(md_path: Path, docx_path: Path, title: str) -> None:
    doc = Document()
    configure_doc(doc, title)
    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            add_paragraph(doc, line[2:], "Title")
        elif line.startswith("## "):
            add_paragraph(doc, line[3:], "Heading 1")
        elif line.startswith("### "):
            add_paragraph(doc, line[4:], "Heading 2")
        elif line.startswith("#### "):
            add_paragraph(doc, line[5:], "Heading 3")
        elif line.startswith("- "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.first_line_indent = Pt(-9)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run("• " + line[2:])
            set_run_font(run, 10)
        else:
            add_paragraph(doc, line)
    doc.save(docx_path)


def write_control_files() -> None:
    control = CLOSURE / "00_control"
    control.mkdir(parents=True, exist_ok=True)
    (control / "STRICT_FOUR_LANE_CLOSURE_PLAN.md").write_text(
        """# 选必二严格四线补齐闭环计划

目标：在不覆盖旧产物的前提下，把最新学会性缺口补入两份学生/教师可用 Word，并重新跑四线。

四线定义：
- Codex A：本地证据融合、学生补丁、文档生成、角色账本。
- ClaudeCode B：独立生产线补跑，检查是否仍有学生迁移缺口和结构缺口。
- Claude Opus 4.7 Adaptive：基于最新版文档做教学文本与学会性复审。
- GPT-5.5 Pro：基于最新版文档做内容压测与学生迁移复审。

最终门槛：
- 两份新版 Word 已生成并渲染。
- ClaudeCode、Claude、GPT 均有真实输出或明确阻塞记录。
- Codex 完成采纳/拒绝裁决。
- Governor 与 Confucius 刷新后缺口为 0。
""",
        encoding="utf-8",
    )
    (CLOSURE / "codex_lane" / "agents" / "decision_maker_2026-05-04.md").write_text(
        "# 决策者\n\n当前瓶颈：零基础学生会旧题检索，但不会陌生题争点生成。优先补“争点生成器、要件卡、易混边界、意义模板、选择题使用边界”。\n",
        encoding="utf-8",
    )
    (CLOSURE / "codex_lane" / "agents" / "worker_2026-05-04.md").write_text(
        "# 劳动者\n\n已执行：在两份最新版 Markdown 前部补入学生可直接使用的迁移方法，生成零基础强补版 Markdown 和 Word。\n",
        encoding="utf-8",
    )
    (CLOSURE / "codex_lane" / "agents" / "patcher_2026-05-04.md").write_text(
        "# 补丁者\n\n补丁范围：只补方法层和易错边界，不新增未经本地题源验证的题目细则；要件卡按高中政治教材与常见评分语言压缩，避免法考化。\n",
        encoding="utf-8",
    )
    (CLOSURE / "codex_lane" / "agents" / "automation_checker_2026-05-04.md").write_text(
        "# 自动化检测者\n\n待检：新版 Markdown 是否存在后台词；新版 Word 是否生成；外部三线是否有报告；Governor/Confucius 是否刷新。\n",
        encoding="utf-8",
    )
    (CLOSURE / "codex_lane" / "agents" / "governor_local_2026-05-04.md").write_text(
        "# Codex A 内部 Governor\n\n初判：允许进入外部四线复审。禁止在 GPT、Claude、ClaudeCode 新输出未回本地裁决前宣布最终 PASS。\n",
        encoding="utf-8",
    )


def main() -> int:
    OUT.mkdir(parents=True, exist_ok=True)
    FRAMEWORK_MD.write_text(patched_framework(), encoding="utf-8")
    SITUATION_MD.write_text(patched_situation(), encoding="utf-8")
    md_to_docx(FRAMEWORK_MD, FRAMEWORK_DOCX, "选必二《法律与生活》框架踩分逻辑")
    md_to_docx(SITUATION_MD, SITUATION_DOCX, "选必二《法律与生活》考过情境细则汇总")
    write_control_files()
    print(FRAMEWORK_MD)
    print(FRAMEWORK_DOCX)
    print(SITUATION_MD)
    print(SITUATION_DOCX)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
