#!/usr/bin/env python3
from __future__ import annotations

import importlib.util
import re
import shutil
from collections import Counter, OrderedDict
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


RUN_DIR = Path(__file__).resolve().parents[1]
DESKTOP = Path(r"C:\Users\Administrator\Desktop")
CLAUDE_DOCX = DESKTOP / "选必二法律与生活_满分宝典_v2_0606.docx"
OUT_DIR = RUN_DIR / "outputs"
QA_DIR = RUN_DIR / "qa"
VERSION = "v4_0606"
PROJECT_DOCX = OUT_DIR / f"选必二法律与生活_满分宝典_{VERSION}_精简截图校准版.docx"
DESKTOP_DOCX = DESKTOP / PROJECT_DOCX.name
QA_MD = QA_DIR / f"MANFEN_BAODIAN_{VERSION}_QA.md"


def import_clean_builder():
    path = RUN_DIR / "tools" / "build_two_clean_docs.py"
    spec = importlib.util.spec_from_file_location("clean_builder", path)
    if not spec or not spec.loader:
        raise RuntimeError(f"Cannot import {path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


bt = import_clean_builder()


def iter_blocks(doc: Document):
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def block_text(block) -> str:
    if isinstance(block, Paragraph):
        return block.text.strip()
    parts: list[str] = []
    for row in block.rows:
        cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
        if cells:
            parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def clean_phrase(text: str) -> str:
    text = bt.clean(text)
    replacements = {
        "生命线五环": "五环",
        "知识一仓·练习一线": "知识与例题",
        "经多Agent评审订正": "",
        "飞哥正志讲堂": "",
        "必背知识仓": "知识仓",
        "必背主仓": "知识仓",
        "高频采分钢句": "采分点",
        "高频": "",
        "必背": "",
        "最易混·必须分清": "",
        "必须分清": "",
        "最易混": "",
        "速查": "",
        "落点": "答题点",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text.replace("★", "")
    text = re.sub(r"\s*·\s*·\s*", " · ", text)
    text = re.sub(r"\s+", " ", text).strip(" ·　")
    return text


def parse_claude_framework() -> tuple[list[dict], OrderedDict[str, dict]]:
    doc = Document(CLAUDE_DOCX)
    entries: list[dict] = []
    rings: OrderedDict[str, dict] = OrderedDict()
    current_ring = ""
    pending_entry: dict | None = None
    for block in iter_blocks(doc):
        text = block_text(block)
        if not text:
            continue
        first = text.splitlines()[0].strip()
        ring_match = re.search(r"例题区\s*·\s*(环[①②③④⑤])\s*·\s*([^　\n]+)", first)
        if ring_match:
            current_ring = ring_match.group(1)
            name = clean_phrase(ring_match.group(2))
            rings[current_ring] = {"name": name, "titles": []}
            continue
        title_match = re.match(r"^●\s*(20\d{2}\s*·.+)$", first)
        if title_match:
            title = bt.safe_title(title_match.group(1))
            pending_entry = {"title": title, "ring": current_ring, "coord": ""}
            entries.append(pending_entry)
            if current_ring:
                rings.setdefault(current_ring, {"name": current_ring, "titles": []})["titles"].append(title)
            continue
        if pending_entry and first.startswith("坐标："):
            pending_entry["coord"] = clean_phrase(first.replace("坐标：", "", 1))
            pending_entry = None
    return entries, rings


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    styles = doc.styles
    for name, size in [("Normal", 9.2), ("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 10.5)]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_after = Pt(3 if name == "Normal" else 5)
        style.paragraph_format.line_spacing = 1.12 if name == "Normal" else 1.05
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.bold = True


def add_para(doc: Document, text: str, *, size: float = 9.2, bold: bool = False, indent: float = 0) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = RGBColor(0, 0, 0)


def add_lines(doc: Document, text: str, *, max_lines: int | None = None) -> int:
    lines = bt.student_paragraphs(bt.clean_display_text(text))
    if max_lines is not None:
        lines = lines[:max_lines]
    for line in lines:
        add_para(doc, line)
    return len(lines)


def add_points(doc: Document, points: list[str]) -> int:
    count = 0
    for point in points:
        point = bt.clean_answer_point(point)
        if not point:
            continue
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.space_after = Pt(1.8)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(point)
        r.font.name = "Arial"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(8.8)
        r.font.color.rgb = RGBColor(0, 0, 0)
        count += 1
    return count


def add_question_images(doc: Document, row: dict) -> int:
    paths = [path for path in bt.student_question_images(row) if Path(path).exists()]
    if not paths:
        return 0
    doc.add_paragraph("原题图", style="Heading 3")
    inserted = 0
    for path in paths:
        try:
            doc.add_picture(str(path), width=Inches(6.55))
            doc.paragraphs[-1].paragraph_format.space_after = Pt(3)
            inserted += 1
        except Exception:
            add_para(doc, f"原题图暂无法嵌入：{path.name}", size=8.5)
    return inserted


def add_type_table(doc: Document) -> None:
    rows = [
        ("T1 判责/理由", "分析判决、谁担责、法理依据、是否构成", "主体-关系-规则-事实-结论"),
        ("T2 补表/补链", "完成下表、补充完整、填空", "按表头逐格补齐事实、规则、结论"),
        ("T3 辨析/评析", "辨析、评析观点或行为", "判断-规则-事实-结论"),
        ("T4 诉求/文书", "诉求、起诉状、法律文书", "请求-事实-理由-责任"),
        ("T5 程序/证据", "调解、仲裁、诉讼、证据、管辖", "程序路径-证据责任-法律后果"),
        ("T6 意义/价值", "意义、作用、价值、启示", "主体权益-社会秩序-法治价值"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["题型", "设问触发", "答题骨架"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def build_doc() -> dict:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    QA_DIR.mkdir(parents=True, exist_ok=True)
    framework_entries, rings = parse_claude_framework()
    packets = bt.load_packets()
    row_by_title = {bt.safe_title(row["title"]): row for row in packets}
    framework_titles = [entry["title"] for entry in framework_entries]
    missing_in_source = [title for title in framework_titles if title not in row_by_title]
    missing_in_framework = [bt.safe_title(row["title"]) for row in packets if bt.safe_title(row["title"]) not in framework_titles]
    doc = Document()
    setup_doc(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("选择性必修二《法律与生活》主观题满分宝典")
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("v4 0606 · 五环框架 · 原题图优先 · 2024-2026 北京各区主观分问 · 按评分细则校准")
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph("使用原则", style="Heading 1")
    for line in [
        "先确定主体、关系、规则、事实和结论；价值意义只在题目要求或细则给分时写。",
        "所有答题点均从已校准题源和评分细则中抽取，不另造口号式句子。",
        "同一题若存在正式点分布待深挖，本版保留为候选，不作最终闭合。",
    ]:
        add_para(doc, line)
    doc.add_paragraph("题型", style="Heading 1")
    add_type_table(doc)
    doc.add_paragraph("知识仓", style="Heading 1")
    knowledge_lines = [
        "主体与行为能力：自然人、法人、非法人组织；限制民事行为能力人实施与年龄、智力、精神健康状况不相适应的民事法律行为，未经法定代理人同意或追认的，对该限制民事行为能力人不发生效力。",
        "民事法律关系：主体、客体、内容；答题时把材料中的人和组织转写为消费者、经营者、劳动者、用人单位、权利人、侵权人等法律身份。",
        "人格权与个人信息：生命权、身体权、健康权、姓名权、肖像权、名誉权、隐私权、个人信息权益；个人信息处理应合法、正当、必要、诚信。",
        "合同与格式条款：合同成立看意思表示一致；履行看约定和诚信；格式条款提供方负提示说明义务，免除或减轻自身责任、加重对方责任、排除对方主要权利的相关条款依法处理。",
        "侵权责任：一般侵权看行为、损害、因果、过错；建筑物及搁置物、悬挂物坠落等按过错推定；产品责任、环境污染、高度危险、饲养动物等依法适用特殊归责。",
        "物权与相邻：不动产相邻权利人按有利生产、方便生活、团结互助、公平合理原则处理相邻关系。",
        "知识产权与竞争：区分著作权、专利、商标、商业秘密；不正当竞争关注混淆、虚假或引人误解宣传、商业诋毁、权利滥用等。",
        "劳动、消费、婚姻继承：劳动关系看劳动合同、劳动者义务、用人单位义务和争议处理；消费题看知情、自主选择、公平交易和安全保障；家庭继承题区分赡养、扶养、抚养、遗嘱、遗赠和遗赠扶养协议。",
        "纠纷解决：协商、调解、仲裁、诉讼各有适用范围；证据责任通常是谁主张谁举证，法律另有规定时适用倒置或特殊分配。",
    ]
    for line in knowledge_lines:
        add_para(doc, line)

    total_points = 0
    total_material_lines = 0
    image_rows = 0
    inserted_images = 0
    for ring, info in rings.items():
        doc.add_page_break()
        doc.add_paragraph(f"{ring} {info['name']}", style="Heading 1")
        add_para(doc, f"本环节共 {len(info['titles'])} 题。")
        for title in info["titles"]:
            row = row_by_title.get(title)
            if not row:
                continue
            doc.add_paragraph(title, style="Heading 2")
            coord = next((entry["coord"] for entry in framework_entries if entry["title"] == title), "")
            if coord:
                add_para(doc, f"坐标：{coord}", size=8.8)
            images_added = add_question_images(doc, row)
            if images_added:
                image_rows += 1
                inserted_images += images_added
            doc.add_paragraph("材料", style="Heading 3")
            total_material_lines += add_lines(doc, row.get("material", ""))
            doc.add_paragraph("设问", style="Heading 3")
            add_lines(doc, row.get("prompt", ""))
            doc.add_paragraph("答题点（按细则）", style="Heading 3")
            total_points += add_points(doc, bt.answer_points(row))

    doc.save(PROJECT_DOCX)
    shutil.copy2(PROJECT_DOCX, DESKTOP_DOCX)
    return {
        "framework_count": len(framework_titles),
        "framework_unique": len(set(framework_titles)),
        "source_count": len(packets),
        "missing_in_source": missing_in_source,
        "missing_in_framework": missing_in_framework,
        "ring_counts": {ring: len(info["titles"]) for ring, info in rings.items()},
        "total_points": total_points,
        "total_material_lines": total_material_lines,
        "image_rows": image_rows,
        "inserted_images": inserted_images,
    }


def docx_scan(path: Path) -> dict:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    forbidden = [
        "通用答题句", "最易混", "必须分清", "必背", "Agent", "飞哥", "速记",
        "采分钢句", "答案逐题对齐", "采分点已标注", "不要回避", "模板", "万能",
        "第一判断", "易错边界", "必背句", "答案落点", "表格重建", "工程痕迹",
    ]
    with ZipFile(path) as zf:
        bad_zip = zf.testzip()
        root = ET.fromstring(zf.read("word/document.xml"))
        text = "".join(t.text or "" for t in root.findall(".//w:t", ns))
        fills = Counter(shd.attrib.get(f"{{{ns['w']}}}fill", "") for shd in root.findall(".//w:shd", ns))
    doc = Document(path)
    return {
        "zip": bad_zip or "ok",
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "text_chars": len(text),
        "forbidden_hits": {term: text.count(term) for term in forbidden if text.count(term)},
        "fills": dict(fills),
    }


def main() -> None:
    stats = build_doc()
    scan = docx_scan(PROJECT_DOCX)
    lines = [
        "# 满分宝典 v4 0606 QA",
        "",
        f"- 输出 DOCX：`{PROJECT_DOCX}`",
        f"- 桌面副本：`{DESKTOP_DOCX}`",
        f"- Claude v2 框架题目数：{stats['framework_count']}；唯一题目：{stats['framework_unique']}",
        f"- 已校准题源数：{stats['source_count']}",
        f"- 框架缺源：{len(stats['missing_in_source'])}",
        f"- 题源未入框架：{len(stats['missing_in_framework'])}",
        f"- 细则答题点总数：{stats['total_points']}",
        f"- 材料段落数：{stats['total_material_lines']}",
        f"- 嵌入原题图条目：{stats['image_rows']}",
        f"- 嵌入原题图数量：{stats['inserted_images']}",
        "",
        "## 五环题量",
        "",
    ]
    for ring, count in stats["ring_counts"].items():
        lines.append(f"- {ring}：{count}")
    lines.extend([
        "",
        "## DOCX结构扫描",
        "",
        f"- zip：{scan['zip']}",
        f"- paragraphs：{scan['paragraphs']}",
        f"- tables：{scan['tables']}",
        f"- inline_shapes：{scan['inline_shapes']}",
        f"- text_chars：{scan['text_chars']}",
        f"- forbidden_hits：{scan['forbidden_hits']}",
        f"- fills：{scan['fills']}",
        "",
        "## 边界",
        "",
        "- 本版沿用 Claude v2 的五环题序，但内容以 v45 已校准题源和细则踩分点重建。",
        "- 本版为黑白简洁稿，不使用彩色信息块；只保留必要题型表；有原题图资产的条目先放原题图，再放材料文字、设问和答题点。",
        "- render_docx/LibreOffice 视觉渲染需另行验证。",
    ])
    if stats["missing_in_source"]:
        lines.extend(["", "### 框架缺源"])
        lines.extend(f"- {title}" for title in stats["missing_in_source"])
    if stats["missing_in_framework"]:
        lines.extend(["", "### 题源未入框架"])
        lines.extend(f"- {title}" for title in stats["missing_in_framework"])
    QA_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(PROJECT_DOCX)
    print(DESKTOP_DOCX)
    print(QA_MD)


if __name__ == "__main__":
    main()
