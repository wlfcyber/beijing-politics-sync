from __future__ import annotations

import csv
from collections import Counter
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


BASE = Path(__file__).resolve().parents[1]
QA = BASE / "qa"
QA.mkdir(exist_ok=True)

SOURCE_BASE = Path(
    "/Users/wanglifei/Library/Application Support/Claude/local-agent-mode-sessions/"
    "2c0a1a14-78ce-4ddc-8f72-9d4f74dbff6e/703c2d6a-ade6-4ee8-94a4-92b2abb47bf1/"
    "local_817e4c28-ed12-4153-a042-b05d41e5da2e/outputs"
)
XLSX = SOURCE_BASE / "法律与生活v34_复核汇总表.xlsx"
DOCX = SOURCE_BASE / "法律与生活v34_复核报告.docx"


def extract_workbook() -> list[dict[str, object]]:
    wb = load_workbook(XLSX, data_only=False)
    records: list[dict[str, object]] = []
    for ws in wb.worksheets:
        if ws.title == "总览":
            continue
        headers = [cell.value for cell in ws[1]]
        for row_number, row in enumerate(ws.iter_rows(min_row=2, values_only=True), 2):
            values = list(row)
            if not any(value is not None and str(value).strip() for value in values):
                continue
            record = {
                str(headers[index] or f"col{index + 1}"): values[index]
                if index < len(values)
                else None
                for index in range(len(headers))
            }
            record["sheet"] = ws.title
            record["row"] = row_number
            records.append(record)
    return records


def write_tsv(records: list[dict[str, object]]) -> Path:
    fields = [
        "sheet",
        "row",
        "序号",
        "题目定位",
        "题目定位/模块",
        "来源",
        "维度",
        "原severity",
        "严重程度",
        "v34状态",
        "问题原文（摘录）",
        "问题说明",
        "复核说明（含仍需修改建议）",
        "修改建议",
    ]
    out = QA / "v34_recheck_raw_extract_20260605.tsv"
    with out.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(
            handle, fieldnames=fields, delimiter="\t", extrasaction="ignore"
        )
        writer.writeheader()
        writer.writerows(records)
    return out


def write_summary(records: list[dict[str, object]]) -> Path:
    lines = [
        "# v34复核 raw extract summary\n",
        f"source_xlsx: {XLSX}\n",
        f"source_docx: {DOCX}\n",
        f"rows_total: {len(records)}\n",
    ]
    for sheet in ["①汇编v34复核", "②宝典v34复核", "③v34新问题"]:
        sheet_records = [record for record in records if record["sheet"] == sheet]
        lines.append(f"\n## {sheet}\n")
        lines.append(f"rows: {len(sheet_records)}\n")
        for key in ["v34状态", "原severity", "严重程度", "维度", "来源"]:
            counts = Counter(
                str(record.get(key) or "").strip()
                for record in sheet_records
                if str(record.get(key) or "").strip()
            )
            if counts:
                rendered = ", ".join(
                    f"{label}={count}" for label, count in counts.most_common(30)
                )
                lines.append(f"- {key}: {rendered}\n")
        for record in sheet_records[:18]:
            location = record.get("题目定位") or record.get("题目定位/模块") or ""
            severity = (
                record.get("原severity")
                or record.get("严重程度")
                or ""
            )
            status = record.get("v34状态") or ""
            dimension = record.get("维度") or ""
            note = (
                record.get("复核说明（含仍需修改建议）")
                or record.get("问题说明")
                or ""
            )
            lines.append(
                f"- row {record['row']}: {location} | {dimension} | "
                f"{severity} | {status} | {str(note)[:180]}\n"
            )

    doc = Document(DOCX)
    lines.append("\n## DOCX report structure\n")
    lines.append(f"paragraphs: {len(doc.paragraphs)}\n")
    lines.append(f"tables: {len(doc.tables)}\n")
    for index, paragraph in enumerate(doc.paragraphs[:80], 1):
        text = paragraph.text.strip()
        if text:
            lines.append(f"- p{index}: {text[:220]}\n")
    lines.append(
        "- tables dims: "
        + ", ".join(f"{len(table.rows)}x{len(table.columns)}" for table in doc.tables)
        + "\n"
    )

    out = QA / "v34_recheck_raw_extract_summary_20260605.md"
    out.write_text("".join(lines), encoding="utf-8")
    return out


def main() -> None:
    records = extract_workbook()
    tsv = write_tsv(records)
    summary = write_summary(records)
    print(tsv)
    print(summary)


if __name__ == "__main__":
    main()
