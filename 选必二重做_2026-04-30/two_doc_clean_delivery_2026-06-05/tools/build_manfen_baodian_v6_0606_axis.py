#!/usr/bin/env python3
from __future__ import annotations

import importlib.util
import re
import shutil
from collections import Counter, defaultdict
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


RUN_DIR = Path(__file__).resolve().parents[1]
DESKTOP = Path(r"C:\Users\Administrator\Desktop")
OUT_DIR = RUN_DIR / "outputs"
QA_DIR = RUN_DIR / "qa"
VERSION = "v6_0606"
BASE_NAME = f"选必二法律与生活_满分宝典_{VERSION}_A轴细则版"
PROJECT_DOCX = OUT_DIR / f"{BASE_NAME}.docx"
PROJECT_MD = OUT_DIR / f"{BASE_NAME}.md"
DESKTOP_DOCX = DESKTOP / PROJECT_DOCX.name
QA_MD = QA_DIR / f"MANFEN_BAODIAN_{VERSION}_QA.md"


def import_clean_builder():
    path = RUN_DIR / "tools" / "build_two_clean_docs.py"
    spec = importlib.util.spec_from_file_location("clean_builder", path)
    if not spec or not spec.loader:
        raise RuntimeError(f"Cannot import {path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


bt = import_clean_builder()

AXIS_FOCUS_HEADING = "本轴知识和答题抓手"
EXAMPLES_HEADING = "相关例题（2024-2026）"
ANSWER_HEADING = "答题要点"
SOURCE_IMAGE_HEADING = "原题图"

AXIS_FOCUS_POINTS = {
    "A1": [
        "先判断民事法律关系、民事主体和民事法律行为是否成立，再落到权利义务和法律后果。",
        "民事法律行为有效通常看行为能力、意思表示真实、不违反法律强制性规定和公序良俗。",
        "诉讼时效抗辩须由义务人主动提出，人民法院不得主动适用；赡养费等特定请求不适用诉讼时效。",
        "答题时把事实、法律规则和案件结论对应起来，避免只背概念不落案情。",
    ],
    "A2": [
        "人格权题先确认被保护的具体权利：姓名、肖像、名誉、荣誉、隐私、个人信息等。",
        "个人信息处理要抓合法、正当、必要、诚信原则，并说明告知、同意、目的范围和安全保护。",
        "涉及未成年人、校园、网络传播时，要同时落到特殊保护和行为边界。",
        "答题要把侵害行为、权利类型、责任承担和社会治理意义分层写清。",
    ],
    "A3": [
        "物权题先判断所有权、用益物权、担保物权、共有或相邻关系，再写权利边界。",
        "相邻关系坚持有利生产、方便生活、团结互助、公平合理，不能把权利绝对化。",
        "共有、通行、排水、采光、污染、安防设备等题目，要把便利一方与不得损害他方合法权益一起写。",
        "答题落点通常是停止侵害、排除妨碍、消除危险、恢复原状、赔偿损失或协商治理。",
    ],
    "A4": [
        "合同题先看合同是否成立有效，再看履行、违约、免责和责任承担。",
        "诚信原则、全面履行原则、公平原则常是合同题的核心落点。",
        "不可抗力必须同时具备不能预见、不能避免、不能克服，不能把一般困难直接写成免责。",
        "答题要说明当事人行为、合同义务、违约责任或救济路径之间的对应关系。",
    ],
    "A5": [
        "知识产权题要精准区分著作权、商标权、专利权、商业秘密和不正当竞争。",
        "商标与老字号、包装装潢、混淆行为等题目，要落到诚信经营、公平竞争和市场秩序。",
        "保护知识产权就是保护创新，可结合惩罚性赔偿、司法保护和营商环境展开。",
        "不要用笼统的“合法权益”替代具体权利名称，题目要求具体权利时必须写准。",
    ],
    "A6": [
        "侵权题先确定侵权行为、损害后果、因果关系、主观过错或法定归责原则。",
        "物件脱落坠落、高空抛物、动物致害、环境侵权、网络侵权等要分清责任主体和举证责任。",
        "责任承担可写停止侵害、排除妨碍、消除危险、赔偿损失、赔礼道歉、消除影响等。",
        "答题要避免只写价值口号，必须把具体法条逻辑落到本案事实。",
    ],
    "A7": [
        "婚姻家庭题要抓夫妻人身关系、共同财产、共同债务、父母子女权利义务和家庭文明。",
        "继承题先分法定继承、遗嘱继承、遗赠和遗赠扶养协议，再看效力和顺序。",
        "赡养、抚养、扶养等义务要结合身份关系、法定义务和社会主义核心价值观写。",
        "家庭类题目既要有法律规则，也要说明平等、和睦、文明家风等价值导向。",
    ],
    "A8": [
        "劳动关系题先确认是否建立劳动关系，再看劳动合同、工资、休息休假、社保和劳动保护。",
        "劳动者维权通常按协商、调解、劳动仲裁、诉讼等路径展开，仲裁前置要写清。",
        "竞业限制、加班、工伤、解除合同等题目，要把权利义务和程序条件对应起来。",
        "答题既保护劳动者合法权益，也要兼顾依法用工和劳动关系稳定。",
    ],
    "A9": [
        "消费者权益题先判断经营者、消费者和消费行为，再写安全权、知情权、公平交易权等具体权利。",
        "食品安全惩罚性赔偿、消费者权益保护法三倍赔偿、经营者欺诈假一赔三等要分清适用条件。",
        "平台、直播、广告、预付式消费等题目，要说明经营者责任、平台责任和维权路径。",
        "答题要把违法事实、消费者权利、责任承担和市场秩序联系起来。",
    ],
    "A10": [
        "纠纷解决题先判断适合调解、仲裁、诉讼还是行政/公益诉讼等程序。",
        "诉讼题要抓管辖、当事人、证据、举证责任、判决或调解书的法律效力。",
        "公益诉讼、环境司法等题目，要写清法定主体、社会公共利益和司法确认/督促履行。",
        "程序类答案要把程序选择、事实审查、权利救济和公平正义的实现串起来。",
    ],
}

FORBIDDEN_TEXT_TERMS = [
    "来源：本轴",
    "按细则",
    "评分提醒",
    "半句不算",
    "附中版",
    "该层分",
    "满分不超过",
    "C:/Users/Administrator",
]

SCORING_LEAK_TERMS = [
    "半句不算",
    "附中版",
    "评分提醒",
    "该层分",
    "满分不超过",
    "可酌情给",
    "酌情给",
    "阅卷",
    "给分",
    "采意",
]

STUDENT_POINT_REPLACEMENTS = [
    ("均只能按处理", "均不能完整回应本题要求"),
    ("作答必须出现诚信，才能稳定拿到该层分", "作答必须点明诚信原则"),
    ("写专利权不得分", "写成专利权不符合本题"),
    ("空泛写法治中国建设不得分", "不能空泛写法治中国建设，要落到法治信仰或依法行动"),
]


FORBIDDEN_TERMS = [
    "使用原则",
    "知识仓",
    "题型",
    "五环",
    "环①",
    "环②",
    "环③",
    "环④",
    "环⑤",
    "第一判断",
    "易错边界",
    "必背句",
    "必背",
    "通用答题句",
    "最易混",
    "速查",
    "采分钢句",
    "答案落点",
    "表格重建",
    "工程痕迹",
    "Agent",
    "飞哥",
    "模板",
    "万能",
    "试题分析",
    "逻辑思路",
    "阅卷细则",
    "两条都写更稳",
    "酌情给分",
    "不要只写",
    "不能只写",
]


DROP_POINT_PATTERNS = [
    "试题分析",
    "逻辑思路",
    "阅卷细则",
    "学生表现",
    "教学启示",
    "阅读材料",
    "完成表格",
    "完成下表",
    "请你选择",
    "请你",
    "结合材料，说明",
    "结合材料，分析",
    "运用《法律与生活》",
    "作答要求",
    "两条都写更稳",
    "不要只写",
    "不能只写",
    "不宜简单写成",
    "可酌情给分",
    "酌情给分",
    "只给",
    "不给分",
    "总分不超过",
    "采意",
]


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    for name, size in [("Normal", 9.2), ("Heading 1", 14), ("Heading 2", 11.5), ("Heading 3", 10.2), ("Heading 4", 9.4)]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = name != "Normal"
        style.paragraph_format.space_before = Pt(6 if name.startswith("Heading") else 0)
        style.paragraph_format.space_after = Pt(4 if name.startswith("Heading") else 2)
        style.paragraph_format.line_spacing = 1.08 if name.startswith("Heading") else 1.12


def set_run_font(run, *, size: float = 9.2, bold: bool = False) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_para(doc: Document, text: str, *, size: float = 9.2, bold: bool = False, indent: float = 0) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)


def add_bullet(doc: Document, text: str, *, size: float = 8.8, indent: float = 0.2) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(1.6)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    set_run_font(r, size=size)


def normalize_key(text: str) -> str:
    return re.sub(r"\W+", "", text)


def clean_student_point(point: str) -> str:
    text = bt.clean_answer_point(point)
    for old, new in STUDENT_POINT_REPLACEMENTS:
        text = text.replace(old, new)
    text = re.sub(r"本小题答[^。；;]*半句不算[^。；;]*[。；;]?", "", text)
    text = re.sub(r"附中版[^。；;]*[。；;]?", "", text)
    text = re.sub(r"评分提醒[:：]?", "", text)
    text = re.sub(r"评分上[^。；;]*满分不超过[^。；;]*[。；;]?", "", text)
    text = re.sub(r"可酌情给[^。；;]*[。；;]?", "", text)
    text = re.sub(r"^【[^】]{0,20}】", "", text).strip()
    text = re.sub(r"^\d{1,2}[.．、]\s*", "", text).strip()
    text = re.sub(r"^\d{1,2}[（(]\d+[）)]\s*", "", text).strip()
    text = re.sub(r"^（[^）]*?分[^）]*?）", "", text).strip()
    text = re.sub(r"^\([^)]*?分[^)]*?\)", "", text).strip()
    text = re.sub(r"[（(][^）)]*?\d+(?:\+\d+)?分[^）)]*?[）)]", "", text).strip()
    text = re.sub(r"\d+(?:\+\d+)?=\d+分\s*", "", text).strip()
    text = re.sub(r"\d+(?:\+\d+)?分[:：]?", "", text).strip()
    text = re.sub(r"得\d+分", "", text).strip()
    text = re.sub(r"^答案示例", "", text).strip()
    text = re.sub(r"^细则[:：]", "", text).strip()
    text = text.replace("答案示例", "").replace("细则：", "").replace("细则:", "")
    text = re.sub(r"（[^）]*?共\d+分[^）]*?）", "", text).strip()
    text = re.sub(r"（[^）]*?采意[^）]*?）", "", text).strip()
    text = re.sub(r"，?总分不超过本题满分。?", "。", text).strip()
    text = re.sub(r"；?总分不超过本题满分。?", "。", text).strip()
    text = re.sub(r"\s+", " ", text).strip("；;，, ")
    text = re.sub(r"均只能按(?:\d+分)?处理", "均不能完整回应本题要求", text)
    text = text.strip("；;，,。 ")
    return text


def is_student_point(point: str) -> bool:
    if len(point) < 6:
        return False
    if any(term in point for term in SCORING_LEAK_TERMS):
        return False
    if "不得分" in point:
        return False
    if any(pattern in point for pattern in DROP_POINT_PATTERNS):
        return False
    if point.startswith("如") and any(token in point for token in ("给分", "答", "总分", "正确写出")):
        return False
    if re.fullmatch(r"[（(]?\d+(?:\+\d+)*[，,、]?\d*分[）)]?", point):
        return False
    return True


def student_answer_points(row: dict) -> list[str]:
    manual = row.get("answer_points") or []
    raw_points = list(manual) if manual else bt.rubric_score_points(row, limit=None)
    points: list[str] = []
    seen: set[str] = set()
    for raw in raw_points:
        point = clean_student_point(raw)
        if not is_student_point(point):
            continue
        key = normalize_key(point)
        if not key or key in seen:
            continue
        seen.add(key)
        points.append(point)
    if points:
        return points
    return [point for point in (clean_student_point(raw) for raw in bt.answer_points(row)) if point]


def add_text_lines(doc: Document, text: str) -> int:
    lines = bt.student_paragraphs(bt.clean_display_text(text))
    for line in lines:
        add_para(doc, line)
    return len(lines)


def axis_rows(packets: list[dict]) -> dict[str, list[dict]]:
    by_axis: dict[str, list[dict]] = defaultdict(list)
    for row in packets:
        axis = row.get("a_axis")
        if axis in bt.A_ORDER:
            by_axis[axis].append(row)
    return {axis: sorted(rows, key=bt.sort_key) for axis, rows in by_axis.items()}


def add_axis_points(doc: Document, axis: str, rows: list[dict]) -> int:
    count = 0
    for point in AXIS_FOCUS_POINTS.get(axis, []):
        add_bullet(doc, point)
        count += 1
    return count


def add_question_images(doc: Document, row: dict) -> int:
    paths = [Path(path) for path in bt.student_question_images(row) if Path(path).exists()]
    if not paths:
        return 0
    doc.add_paragraph(SOURCE_IMAGE_HEADING, style="Heading 4")
    inserted = 0
    for path in paths:
        try:
            doc.add_picture(str(path), width=Inches(6.55))
            doc.paragraphs[-1].paragraph_format.space_after = Pt(3)
            inserted += 1
        except Exception:
            add_para(doc, f"原题图暂无法嵌入：{path.name}", size=8.5)
    return inserted


def add_example(doc: Document, row: dict) -> tuple[int, int, int]:
    doc.add_paragraph(bt.safe_title(row["title"]), style="Heading 3")
    image_count = add_question_images(doc, row)
    doc.add_paragraph("材料", style="Heading 4")
    material_lines = add_text_lines(doc, row.get("material", ""))
    doc.add_paragraph("设问", style="Heading 4")
    add_text_lines(doc, row.get("prompt", ""))
    doc.add_paragraph(ANSWER_HEADING, style="Heading 4")
    answer_count = 0
    for point in student_answer_points(row):
        add_bullet(doc, point)
        answer_count += 1
    return image_count, material_lines, answer_count


def build_docx(packets: list[dict]) -> dict:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    QA_DIR.mkdir(parents=True, exist_ok=True)
    by_axis = axis_rows(packets)
    doc = Document()
    setup_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("选择性必修二《法律与生活》主观题满分宝典")
    set_run_font(run, size=18, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("v5 0606 · A轴细则版 · 2024-2026 北京各区主观分问")
    set_run_font(run, size=9.5)
    run.text = "v6 0606 · A轴学生精简版 · 2024-2026 北京各区主观题"

    axis_point_count = 0
    example_count = 0
    image_rows = 0
    image_count = 0
    material_lines = 0
    answer_points = 0
    for axis_index, axis in enumerate(bt.A_ORDER):
        rows = by_axis.get(axis, [])
        if axis_index:
            doc.add_page_break()
        doc.add_paragraph(bt.A_NAMES[axis], style="Heading 1")
        doc.add_paragraph(AXIS_FOCUS_HEADING, style="Heading 2")
        axis_point_count += add_axis_points(doc, axis, rows)
        doc.add_paragraph(EXAMPLES_HEADING, style="Heading 2")
        for row in rows:
            imgs, mats, pts = add_example(doc, row)
            example_count += 1
            image_count += imgs
            image_rows += 1 if imgs else 0
            material_lines += mats
            answer_points += pts
    doc.save(PROJECT_DOCX)
    shutil.copy2(PROJECT_DOCX, DESKTOP_DOCX)
    return {
        "axis_counts": {axis: len(by_axis.get(axis, [])) for axis in bt.A_ORDER},
        "axis_point_count": axis_point_count,
        "example_count": example_count,
        "image_rows": image_rows,
        "image_count": image_count,
        "material_lines": material_lines,
        "answer_points": answer_points,
    }


def build_markdown(packets: list[dict]) -> None:
    by_axis = axis_rows(packets)
    lines = [
        "# 选择性必修二《法律与生活》主观题满分宝典",
        "",
        "v6 0606 · A轴学生精简版 · 2024-2026 北京各区主观题",
        "",
    ]
    for axis in bt.A_ORDER:
        rows = by_axis.get(axis, [])
        lines.extend([f"## {bt.A_NAMES[axis]}", "", f"### {AXIS_FOCUS_HEADING}", ""])
        for point in AXIS_FOCUS_POINTS.get(axis, []):
            lines.append(f"- {point}")
        lines.extend(["", f"### {EXAMPLES_HEADING}", ""])
        for row in rows:
            title = bt.safe_title(row["title"])
            lines.extend([f"#### {title}", ""])
            question_images = [Path(path) for path in bt.student_question_images(row) if Path(path).exists()]
            if question_images:
                lines.append(f"##### {SOURCE_IMAGE_HEADING}")
                for idx, path in enumerate(question_images, 1):
                    try:
                        rel_path = path.relative_to(RUN_DIR).as_posix()
                    except ValueError:
                        rel_path = path.name
                    lines.append(f"![原题图{idx}]({rel_path})")
                lines.append("")
            lines.extend(["##### 材料", bt.clean_display_text(row.get("material", "")), "", "##### 设问", bt.clean_display_text(row.get("prompt", "")), "", f"##### {ANSWER_HEADING}"])
            for point in student_answer_points(row):
                lines.append(f"- {point}")
            lines.append("")
    PROJECT_MD.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def docx_scan(path: Path, packets: list[dict]) -> dict:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(path) as zf:
        bad_zip = zf.testzip()
        root = ET.fromstring(zf.read("word/document.xml"))
        text = "".join(t.text or "" for t in root.findall(".//w:t", ns))
        media = [name for name in zf.namelist() if name.startswith("word/media/")]
        fills = [
            shd.attrib.get(f"{{{ns['w']}}}fill", "")
            for shd in root.findall(".//w:shd", ns)
            if shd.attrib.get(f"{{{ns['w']}}}fill", "") not in ("", "auto", "FFFFFF")
        ]
        headings: list[tuple[str, str]] = []
        for p in root.findall(".//w:p", ns):
            p_text = "".join(t.text or "" for t in p.findall(".//w:t", ns)).strip()
            if not p_text:
                continue
            style = ""
            ppr = p.find("w:pPr", ns)
            if ppr is not None:
                ps = ppr.find("w:pStyle", ns)
                if ps is not None:
                    style = ps.attrib.get(f"{{{ns['w']}}}val", "")
            if style.startswith("Heading"):
                headings.append((style, p_text))
    h1 = [text for style, text in headings if style == "Heading1"]
    h2 = [text for style, text in headings if style == "Heading2"]
    h3 = [text for style, text in headings if style == "Heading3"]
    packet_titles = [bt.safe_title(row["title"]) for row in packets]
    missing_titles = [title for title in packet_titles if title not in h3]
    extra_h3 = [title for title in h3 if title not in packet_titles]
    axis_sequence_ok = h1 == [bt.A_NAMES[axis] for axis in bt.A_ORDER]
    axis_block_issues = []
    cursor = 0
    for axis in bt.A_ORDER:
        axis_name = bt.A_NAMES[axis]
        try:
            axis_index = next(i for i in range(cursor, len(headings)) if headings[i] == ("Heading1", axis_name))
        except StopIteration:
            axis_block_issues.append(f"{axis_name}: missing Heading1")
            continue
        next_axis = next((i for i in range(axis_index + 1, len(headings)) if headings[i][0] == "Heading1"), len(headings))
        block = headings[axis_index:next_axis]
        block_h2 = [t for s, t in block if s == "Heading2"]
        if block_h2[:2] != [AXIS_FOCUS_HEADING, EXAMPLES_HEADING]:
            axis_block_issues.append(f"{axis_name}: Heading2 order {block_h2[:2]}")
        cursor = next_axis
    return {
        "zip": bad_zip or "ok",
        "paragraphs": len(root.findall(".//w:p", ns)),
        "tables": len(root.findall(".//w:tbl", ns)),
        "drawings": len(root.findall(".//w:drawing", ns)),
        "media": len(media),
        "fills": sorted(set(fills)),
        "forbidden_hits": {term: text.count(term) for term in [*FORBIDDEN_TERMS, *FORBIDDEN_TEXT_TERMS] if term in text},
        "h1_count": len(h1),
        "h2_knowledge_count": h2.count(AXIS_FOCUS_HEADING),
        "h2_examples_count": h2.count(EXAMPLES_HEADING),
        "h3_example_count": len(h3),
        "missing_titles": missing_titles,
        "extra_h3": extra_h3,
        "axis_sequence_ok": axis_sequence_ok,
        "axis_block_issues": axis_block_issues,
        "answer_labels": text.count(ANSWER_HEADING),
        "source_image_labels": text.count(SOURCE_IMAGE_HEADING),
        "text_chars": len(text),
    }


def write_qa(packets: list[dict], stats: dict, scan: dict) -> None:
    by_axis = axis_rows(packets)
    lines = [
        "# 满分宝典 v6 0606 QA",
        "",
        f"- 输出 DOCX：`{PROJECT_DOCX}`",
        f"- 桌面 DOCX：`{DESKTOP_DOCX}`",
        f"- 输出 MD：`{PROJECT_MD}`",
        f"- 分问总数：{len(packets)}",
        f"- DOCX 例题标题数：{scan['h3_example_count']}",
        f"- 缺失例题标题：{len(scan['missing_titles'])}",
        f"- 额外例题标题：{len(scan['extra_h3'])}",
        f"- A轴顺序：{'PASS' if scan['axis_sequence_ok'] else 'CHECK'}",
        f"- A轴结构问题：{len(scan['axis_block_issues'])}",
        f"- A轴抓手条目：{stats['axis_point_count']}",
        f"- 每题答题点总数：{stats['answer_points']}",
        f"- 材料段落数：{stats['material_lines']}",
        f"- 嵌入原题图条目：{stats['image_rows']}",
        f"- 嵌入原题图数量：{stats['image_count']}",
        "",
        "## A轴题量",
        "",
    ]
    for axis in bt.A_ORDER:
        lines.append(f"- {bt.A_NAMES[axis]}：{len(by_axis.get(axis, []))}")
    lines.extend([
        "",
        "## DOCX结构扫描",
        "",
        f"- zip：{scan['zip']}",
        f"- paragraphs：{scan['paragraphs']}",
        f"- tables：{scan['tables']}",
        f"- drawings：{scan['drawings']}",
        f"- media：{scan['media']}",
        f"- fills：{scan['fills']}",
        f"- forbidden_hits：{scan['forbidden_hits']}",
        f"- h1_count：{scan['h1_count']}",
        f"- h2_knowledge_count：{scan['h2_knowledge_count']}",
        f"- h2_examples_count：{scan['h2_examples_count']}",
        f"- answer_labels：{scan['answer_labels']}",
        f"- source_image_labels：{scan['source_image_labels']}",
        f"- text_chars：{scan['text_chars']}",
        "",
        "## 边界",
        "",
        "- 本版删除 v4 的使用原则、题型表、知识仓和五环分组，改为 A1-A10 轴结构。",
        "- 每个 A轴后先给本轴知识和答题抓手，再跟 2024-2026 全部相关例题；每题完整答案落点保留在例题区。",
        "- v5 已由 Claude Opus 4.8 Max 网页端判定 FAIL；v6 针对重复、后台评分话术、标题口径和 Markdown 绝对路径做修订。",
    ])
    if scan["missing_titles"]:
        lines.extend(["", "### 缺失例题标题"])
        lines.extend(f"- {title}" for title in scan["missing_titles"])
    if scan["axis_block_issues"]:
        lines.extend(["", "### A轴结构问题"])
        lines.extend(f"- {issue}" for issue in scan["axis_block_issues"])
    QA_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    packets = sorted(bt.load_packets(), key=bt.sort_key)
    stats = build_docx(packets)
    build_markdown(packets)
    scan = docx_scan(PROJECT_DOCX, packets)
    write_qa(packets, stats, scan)
    print(PROJECT_DOCX)
    print(DESKTOP_DOCX)
    print(PROJECT_MD)
    print(QA_MD)
    print(scan)


if __name__ == "__main__":
    main()
