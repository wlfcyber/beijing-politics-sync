from __future__ import annotations

import importlib.util
import re
import shutil
import sys
import zipfile
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TOOLS = ROOT / "tools"
OUTPUTS = ROOT / "outputs"
QA_DIR = ROOT / "qa"
DESKTOP = Path.home() / "Desktop"
V2_DOCX = next(iter(DESKTOP.glob("*v2_0606.docx")), None)
V7_SCRIPT = TOOLS / "build_manfen_baodian_v7_0606_axis.py"

OUT_NAME = "选必二法律与生活_满分宝典_v8_0606_新框架版"
DOCX_OUT = OUTPUTS / f"{OUT_NAME}.docx"
MD_OUT = OUTPUTS / f"{OUT_NAME}.md"
QA_OUT = QA_DIR / "MANFEN_BAODIAN_v8_0606_NEW_FRAMEWORK_QA.md"
DESKTOP_DOCX = DESKTOP / f"{OUT_NAME}.docx"
DESKTOP_MD = DESKTOP / f"{OUT_NAME}.md"

RING_ORDER = ["环①", "环②", "环③", "环④", "环⑤"]
RING_NAMES = {
    "环①": "主体、行为能力、三要素、权利识别",
    "环②": "基础规则主仓",
    "环③": "审判定责、原则、时效、举证",
    "环④": "程序与维权",
    "环⑤": "意义与价值",
}
TYPE_GUIDE = [
    ("T1 判责/理由", "定主体，辨关系，套规则要件，落到支持或驳回、谁担何责。"),
    ("T2 补表/补链", "按表头逐格补齐：原则、权利、要件、事实、结论；只填缺项。"),
    ("T3 辨析/评析", "先判断观点，再引规则，扣材料说理，最后下结论。"),
    ("T4 意义/价值", "分层回答：个人权益、市场秩序、司法法治、社会价值，每层回扣材料。"),
    ("T5 对策/维权", "写主体该做什么：依据、具体做法、救济路径和请求内容。"),
    ("T6 程序/纠错", "指出错处，改正程序，补一句法理；或列适用的解纷路径。"),
]
RING_KNOWLEDGE = {
    "环①": [
        "先确定主体：自然人、法人、非法人组织；再判断是否具有相应民事行为能力。",
        "民事法律关系三要素是主体、客体、内容。客体要叫准：物、行为、智力成果、人身利益等。",
        "权利识别要落到具体名称：生命权、身体权、健康权、姓名权、肖像权、名誉权、隐私权、个人信息权益、所有权、用益物权、担保物权、知识产权等。",
        "合同效力判断按四类区分：有效、效力待定、可撤销、无效。未成年人、欺诈胁迫、重大误解、违反强制性规定等要分清。",
        "答题时不要只写“权利受侵害”，要写清谁的什么权利、被谁以什么行为侵害。"
    ],
    "环②": [
        "合同：看订立、效力、履行、变更解除、违约责任。格式条款重点看提示说明、公平合理、无效条款。",
        "侵权：一般过错责任要证明过错、损害、因果关系；特殊情形区分过错推定和无过错责任。",
        "物权和相邻关系：不动产权利人行使权利不得损害相邻权利人的合法权益，要兼顾通行、采光、排水、通风、安全等利益。",
        "知识产权和不正当竞争：区分作品、商标、专利、商业秘密、混淆、虚假宣传、商业诋毁等具体类型。",
        "劳动：围绕劳动合同、劳动者权利、用人单位义务、劳动争议解决路径展开。",
        "消费者：经营者应真实全面告知、保障安全、公平交易；消费者可要求退换、赔偿、停止侵害等。",
        "婚姻继承：抓住婚姻自由、夫妻财产、父母子女义务、遗嘱和法定继承、赡养扶养等。"
    ],
    "环③": [
        "判责题的核心不是背条文，而是把材料事实对应到法律要件。",
        "民法基本原则可作为理由：平等、自愿、公平、诚信、守法和公序良俗、绿色原则。原则必须和事实相连。",
        "自甘风险适用于自愿参加有风险的文体活动，但组织者未尽安全保障义务的仍可能担责。",
        "举证要看谁主张谁举证，同时关注法律规定的举证责任倒置、过错推定等例外。",
        "诉讼时效一般影响胜诉可能，但赡养费、扶养费、抚养费、不动产物权等特定请求不适用诉讼时效。"
    ],
    "环④": [
        "维权路径通常按协商、调解、仲裁、诉讼排列，但要结合材料判断是否存在仲裁协议、劳动争议前置程序等条件。",
        "程序纠错要写清错在何处、应如何处理、背后的程序价值是什么。",
        "起诉要看主体资格、管辖、诉讼请求、证据。诉讼中可关注调解、回避、举证、执行等环节。",
        "对策题要从当事人、经营者、平台、学校、法院或行政机关等具体主体出发，不写空泛号召。",
        "维权请求要落地：停止侵害、排除妨碍、消除危险、返还财产、恢复原状、赔偿损失、赔礼道歉、消除影响等。"
    ],
    "环⑤": [
        "意义题从个案出发，不能只写口号。先说保护了谁的具体权利，再上升到秩序和法治。",
        "个人层面：保护人格、财产、劳动、消费、婚姻家庭等合法权益。",
        "社会和市场层面：维护交易安全、公平竞争、诚信经营、公共安全和社会稳定。",
        "司法和法治层面：统一裁判尺度，增强规则预期，推动全民守法和依法维权。",
        "价值层面：体现社会主义核心价值观、公序良俗、家庭美德、劳动光荣、知识产权保护等，但必须回扣材料。"
    ],
}

BAD_DOC_TERMS = [
    "A1", "A2", "A3", "A4", "A5", "A6", "A7", "A8", "A9", "A10",
    "A轴", "第一判断", "易错边界", "必背句", "可直接默写", "踩分",
    "采分点", "评分口径", "Agent", "飞哥", "来源：本轴", "附中版",
    "半句不算", "得分不如", "五点中写出",
]
BAD_POINT_TERMS = [
    "评分口径", "采分点", "分）", "分)", "不给分", "只给",
    "得分", "可直接默写", "踩分", "范答", "变通", "作答要求",
    "半句不算", "五点中写出", "按细则", "后台",
]


def load_v7():
    if not V7_SCRIPT.exists():
        raise FileNotFoundError(V7_SCRIPT)
    spec = importlib.util.spec_from_file_location("v7_axis_source", V7_SCRIPT)
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def normalize_title(text: str) -> str:
    return "".join(ch for ch in str(text) if ch.isalnum())


def parse_v2_framework() -> list[dict]:
    if V2_DOCX is None or not V2_DOCX.exists():
        raise FileNotFoundError("Desktop *v2_0606.docx not found")
    doc = Document(str(V2_DOCX))
    paragraphs = [para.text.strip() for para in doc.paragraphs]
    bullet = chr(0x25CF)
    sep = chr(0xFF5C)
    times = chr(0xD7)
    coord_prefix = chr(0x5750) + chr(0x6807) + chr(0xFF1A)
    topic_prefix = chr(0x8003) + chr(0x70B9) + chr(0xFF1A)
    extra_prefix = chr(0x5916) + chr(0x6EA2) + chr(0xFF1A)
    items = []
    for index, text in enumerate(paragraphs):
        if not text.startswith(bullet):
            continue
        title = text[1:].strip()
        coord = paragraphs[index + 1].strip() if index + 1 < len(paragraphs) else ""
        ring = typ = topics = extra = ""
        if coord.startswith(coord_prefix):
            parts = [part.strip() for part in coord[len(coord_prefix):].split(sep)]
            left = parts[0] if parts else ""
            if times in left:
                ring, typ = [part.strip() for part in left.split(times, 1)]
            else:
                ring = left.strip()
            for part in parts[1:]:
                if part.startswith(topic_prefix):
                    topics = part[len(topic_prefix):].strip()
                elif part.startswith(extra_prefix):
                    extra = part[len(extra_prefix):].strip()
        items.append({"title": title, "ring": ring, "type": typ, "topics": topics, "extra": extra})
    return items


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    for name, size, before, after in [
        ("Normal", 9.2, 0, 2),
        ("Heading 1", 14, 10, 4),
        ("Heading 2", 11.5, 8, 3),
        ("Heading 3", 10.2, 6, 2),
        ("Heading 4", 9.4, 4, 2),
        ("List Bullet", 8.9, 0, 1.5),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = name.startswith("Heading")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.08 if name.startswith("Heading") else 1.12


def add_text(doc: Document, text: str, *, size: float = 9.2, bold: bool = False, indent: float = 0) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = RGBColor(0, 0, 0)


def add_bullet(doc: Document, text: str, *, size: float = 8.9, indent: float = 0.18) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0, 0, 0)


def clean_point(text: str, v7) -> str:
    text = v7.clean_student_point(text)
    text = re.sub(r"^[-—、，。\\s]+", "", text).strip()
    text = re.sub(r"\\s+", " ", text).strip()
    return text


def keep_point(text: str) -> bool:
    if not text or len(text) < 4:
        return False
    return not any(term in text for term in BAD_POINT_TERMS)


def answer_points(row: dict, v7) -> list[str]:
    points = []
    seen = set()
    for point in v7.student_answer_points(row):
        cleaned = clean_point(point, v7)
        if not keep_point(cleaned):
            continue
        key = normalize_title(cleaned)
        if key and key not in seen:
            seen.add(key)
            points.append(cleaned)
    if points:
        return points
    for point in v7.bt.answer_points(row):
        cleaned = clean_point(point, v7)
        if keep_point(cleaned):
            points.append(cleaned)
    return points[:8]


def add_question_images(doc: Document, row: dict, v7) -> int:
    paths = [Path(path) for path in v7.bt.student_question_images(row) if Path(path).exists()]
    if not paths:
        return 0
    h = doc.add_paragraph("原题图", style="Heading 4")
    h.paragraph_format.keep_with_next = True
    count = 0
    for path in paths:
        try:
            width = 5.6 if row.get("entry_id") == "E052" else 6.55
            doc.add_picture(str(path), width=Inches(width))
            para = doc.paragraphs[-1]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(3)
            count += 1
        except Exception:
            add_text(doc, f"原题图暂无法嵌入：{path.name}", size=8.5)
    return count


def add_paragraph_blocks(doc: Document, raw: str, v7, *, max_blocks: int | None = None) -> None:
    blocks = v7.bt.student_paragraphs(raw)
    if max_blocks is not None:
        blocks = blocks[:max_blocks]
    for block in blocks:
        block = v7.bt.clean_display_text(block)
        if not block:
            continue
        add_text(doc, block, size=8.9)


def add_opening(doc: Document, total: int, ring_counts: Counter, type_counts: Counter) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("选择性必修二《法律与生活》主观题满分宝典")
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("v8 0606 新框架版｜一核五步 × T1-T6 × 五环知识总仓 × 74题")
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0, 0, 0)

    add_text(doc, "本版按新框架重排：先用五步读题，再按T1-T6识别设问，最后落到五环知识总仓。正文保留材料、设问、原题图和答案要点，删除装饰性颜色和后台化表述。", size=9.3)
    add_text(doc, f"覆盖情况：共{total}题；" + "；".join(f"{ring}{ring_counts.get(ring, 0)}题" for ring in RING_ORDER if ring != "环②" or ring_counts.get(ring, 0)), size=9.0)
    add_text(doc, "一核五步：谁和谁；什么关系；哪条规则；材料事实对应哪个要件；得出什么结论及意义。", size=9.0, bold=True)

    doc.add_paragraph("T1-T6设问骨架", style="Heading 1")
    for label, guide in TYPE_GUIDE:
        add_bullet(doc, f"{label}：{guide}")
    add_text(doc, "题型数量：" + "；".join(f"{typ}{count}题" for typ, count in sorted(type_counts.items())), size=8.8)


def build_doc(groups: dict[str, list[dict]], total: int, ring_counts: Counter, type_counts: Counter, v7) -> tuple[int, list[str]]:
    doc = Document()
    setup_doc(doc)
    add_opening(doc, total, ring_counts, type_counts)
    image_count = 0
    for idx, ring in enumerate(RING_ORDER):
        if idx:
            doc.add_page_break()
        doc.add_paragraph(f"{ring} {RING_NAMES[ring]}", style="Heading 1")
        doc.add_paragraph("知识和答题点", style="Heading 2")
        for item in RING_KNOWLEDGE[ring]:
            add_bullet(doc, item)
        rows = groups.get(ring, [])
        doc.add_paragraph("相关例题（2024-2026）", style="Heading 2")
        if not rows:
            add_text(doc, "本环作为总仓基础，供其他各环调用；本版74题中没有单独归入本环的例题。", size=8.9)
            continue
        for index, item in enumerate(rows, start=1):
            row = item["row"]
            title = v7.bt.safe_title(row.get("title", ""))
            meta = item["meta"]
            doc.add_paragraph(f"{index}. {title}", style="Heading 3")
            topics = meta.get("topics") or "综合考点"
            extra = f"；跨模块提示：{meta['extra']}" if meta.get("extra") else ""
            add_text(doc, f"题型：{meta.get('type') or '未标注'}；考点：{topics}{extra}", size=8.7)
            image_count += add_question_images(doc, row, v7)
            doc.add_paragraph("材料", style="Heading 4")
            add_paragraph_blocks(doc, row.get("material", ""), v7)
            doc.add_paragraph("设问", style="Heading 4")
            add_paragraph_blocks(doc, row.get("prompt", ""), v7)
            doc.add_paragraph("答题要点", style="Heading 4")
            for point in answer_points(row, v7):
                add_bullet(doc, point, size=8.8)
    OUTPUTS.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)
    shutil.copy2(DOCX_OUT, DESKTOP_DOCX)
    text = "\n".join(para.text for para in Document(str(DOCX_OUT)).paragraphs)
    bad_hits = [term for term in BAD_DOC_TERMS if term in text]
    return image_count, bad_hits


def build_markdown(groups: dict[str, list[dict]], total: int, ring_counts: Counter, type_counts: Counter, v7) -> None:
    lines = [
        "# 选择性必修二《法律与生活》主观题满分宝典",
        "",
        "v8 0606 新框架版：一核五步 × T1-T6 × 五环知识总仓 × 74题",
        "",
        "## 一核五步",
        "",
        "谁和谁；什么关系；哪条规则；材料事实对应哪个要件；得出什么结论及意义。",
        "",
        "## T1-T6设问骨架",
        "",
    ]
    for label, guide in TYPE_GUIDE:
        lines.append(f"- {label}：{guide}")
    lines.extend(["", "## 覆盖统计", ""])
    lines.append(f"- 总题数：{total}")
    for ring in RING_ORDER:
        lines.append(f"- {ring} {RING_NAMES[ring]}：{ring_counts.get(ring, 0)}题")
    lines.append("- 题型数量：" + "；".join(f"{typ}{count}题" for typ, count in sorted(type_counts.items())))
    for ring in RING_ORDER:
        lines.extend(["", f"## {ring} {RING_NAMES[ring]}", "", "### 知识和答题点", ""])
        for item in RING_KNOWLEDGE[ring]:
            lines.append(f"- {item}")
        rows = groups.get(ring, [])
        lines.extend(["", "### 相关例题（2024-2026）", ""])
        if not rows:
            lines.append("本环作为总仓基础，供其他各环调用；本版74题中没有单独归入本环的例题。")
            continue
        for index, item in enumerate(rows, start=1):
            row = item["row"]
            meta = item["meta"]
            title = v7.bt.safe_title(row.get("title", ""))
            lines.extend(["", f"#### {index}. {title}", ""])
            topics = meta.get("topics") or "综合考点"
            extra = f"；跨模块提示：{meta['extra']}" if meta.get("extra") else ""
            lines.append(f"题型：{meta.get('type') or '未标注'}；考点：{topics}{extra}")
            image_paths = [str(path) for path in v7.bt.student_question_images(row) if Path(path).exists()]
            if image_paths:
                lines.append("原题图：" + "；".join(image_paths))
            lines.extend(["", "材料："])
            for block in v7.bt.student_paragraphs(row.get("material", "")):
                lines.append(v7.bt.clean_display_text(block))
            lines.extend(["", "设问："])
            for block in v7.bt.student_paragraphs(row.get("prompt", "")):
                lines.append(v7.bt.clean_display_text(block))
            lines.extend(["", "答题要点："])
            for point in answer_points(row, v7):
                lines.append(f"- {point}")
    OUTPUTS.mkdir(parents=True, exist_ok=True)
    MD_OUT.write_text("\n".join(lines), encoding="utf-8")
    shutil.copy2(MD_OUT, DESKTOP_MD)


def create_qa(items: list[dict], groups: dict[str, list[dict]], image_count: int, bad_hits: list[str], matched: int, v7) -> None:
    QA_DIR.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(DOCX_OUT) as zf:
        media_count = len([name for name in zf.namelist() if name.startswith("word/media/")])
    headings = Counter()
    for para in Document(str(DOCX_OUT)).paragraphs:
        if para.style and para.style.name:
            headings[para.style.name] += 1
    lines = [
        "# v8 0606 新框架版 QA",
        "",
        f"- v2坐标题数：{len(items)}",
        f"- v7素材匹配题数：{matched}",
        f"- DOCX路径：{DOCX_OUT}",
        f"- 桌面DOCX：{DESKTOP_DOCX}",
        f"- 图片嵌入数：{image_count}",
        f"- DOCX媒体文件数：{media_count}",
        "- 五环分布：" + "；".join(f"{ring}{len(groups.get(ring, []))}题" for ring in RING_ORDER),
        "- 标题样式计数：" + "；".join(f"{k}:{v}" for k, v in sorted(headings.items())),
        "- A轴/包装词扫描：" + ("无命中" if not bad_hits else "命中 " + "、".join(bad_hits)),
        "",
        "## 说明",
        "",
        "本版只使用v2的新框架坐标组织目录，v7仅作为题源、原题图和答案要点素材池。外部模型最终复核未在本脚本中完成。",
    ]
    QA_OUT.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    v7 = load_v7()
    items = parse_v2_framework()
    meta_by_title = {normalize_title(item["title"]): item for item in items}
    packets = sorted(v7.apply_v7_opus_patch(v7.bt.load_packets()), key=v7.bt.sort_key)
    groups = defaultdict(list)
    missing = []
    for row in packets:
        title = v7.bt.safe_title(row.get("title", ""))
        meta = meta_by_title.get(normalize_title(title))
        if not meta:
            missing.append(title)
            continue
        groups[meta["ring"]].append({"row": row, "meta": meta})
    if missing:
        raise RuntimeError("Unmatched titles: " + "; ".join(missing[:10]))
    ring_counts = Counter(item["ring"] for item in items)
    type_counts = Counter(item["type"] for item in items)
    image_count, bad_hits = build_doc(groups, len(items), ring_counts, type_counts, v7)
    build_markdown(groups, len(items), ring_counts, type_counts, v7)
    create_qa(items, groups, image_count, bad_hits, len(packets) - len(missing), v7)
    print(f"DOCX={DOCX_OUT}")
    print(f"DESKTOP_DOCX={DESKTOP_DOCX}")
    print(f"MD={MD_OUT}")
    print(f"QA={QA_OUT}")
    print(f"QUESTIONS={len(items)} MATCHED={len(packets)-len(missing)} IMAGES={image_count} BAD_HITS={bad_hits}")


if __name__ == "__main__":
    main()
