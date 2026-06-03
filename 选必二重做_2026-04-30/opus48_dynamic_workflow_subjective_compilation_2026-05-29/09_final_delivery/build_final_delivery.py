from __future__ import annotations

import csv
import re
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


RUN_ROOT = Path(__file__).resolve().parents[1]
INPUT_CSV = RUN_ROOT / "04_outputs" / "选必二法律主观题_习题与细则汇编_20260529.csv"
OUT_DIR = RUN_ROOT / "09_final_delivery"
TODAY = "2026-05-29"

FINAL_MD = OUT_DIR / "选必二法律主观题_习题与细则汇编_最终版_20260529.md"
FINAL_DOCX = OUT_DIR / "选必二法律主观题_习题与细则汇编_最终版_20260529.docx"


DOMAIN_BY_QID = {
    qid: domain
    for domain, qids in {
        "知识产权与竞争法": "S05_Q19 S16_Q19 S18_Q19 S27_Q20 S29_Q18 S33_Q20 S36_Q19 S37_Q18 S41_Q18 S48_Q19 S50_Q18 S52_Q18 S61_Q19 S62_Q17".split(),
        "侵权责任与人格权": "S08_Q17 S09_Q17 S28_Q20 S38_Q20 S39_Q18 S42_Q18 S45_Q20 S49_Q17 S51_Q18 S53_Q19 S55_Q18".split(),
        "民事总则与基本原则": "S01_Q19 S02_Q19 S25_Q19 S26_Q19 S30_Q18 S40_Q17 S54_Q18".split(),
        "劳动与就业": "S03_Q17 S22_Q20 S23_Q20 S32_Q19 S35_Q21 S47_Q18".split(),
        "合同": "S07_Q19 S11_Q17 S20_Q18 S56_Q18 S61_Q18".split(),
        "社会争议解决": "S15_Q19 S17_Q19 S21_Q20 S43_Q18 S46_Q18".split(),
        "物权": "S24_Q19 S31_Q19 S44_Q17 S63_Q19".split(),
        "消费者权益": "S13_Q16 S19_Q19 S59_Q18".split(),
        "婚姻家庭与继承": "S10_Q19 S12_Q18 S34_Q21".split(),
    }.items()
    for qid in qids
}

STAGE_ORDER = {"期末": 0, "期中": 1, "一模": 2, "二模": 3}


def read_rows() -> list[dict[str, str]]:
    with INPUT_CSV.open(newline="", encoding="utf-8-sig") as f:
        rows = list(csv.DictReader(f))
    rows.sort(key=lambda r: (int(r["year"]), STAGE_ORDER.get(r["stage"], 99), r["district"], r["question_no"], r["question_id"]))
    return rows


def clean_text(value: str) -> str:
    text = (value or "").strip()
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u2028", "\n").replace("\u2029", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


BACKEND_LINE_PATTERNS = [
    r"rubric_status\s*:",
    r"细则来源文件\s*:",
    r"证据类型\s*:",
    r"提取方式\s*:",
    r"模块判定备注",
    r"目标小问",
    r"anomaly\s*:",
    r"/Users/wanglifei/",
    r"source_packets",
    r"SOURCE_LEDGER",
    r"QUESTION_COVERAGE",
    r"GOVERNOR",
    r"workflow",
    r"formal_rubric",
    r"needs_review",
]


def clean_student_text(value: str) -> str:
    """Keep teaching-facing content while dropping source-audit scaffolding."""
    text = clean_text(value)
    text = re.sub(r"〔[^〕]*(?:needs_review|模块判定|非选必二|rubric_status|formal_rubric|source_packets)[^〕]*〕", "", text)
    text = re.sub(r"（注：[^）]*(?:needs_review|模块判定|rubric_status|formal_rubric)[^）]*）", "", text)
    lines: list[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        if any(re.search(pattern, stripped, flags=re.IGNORECASE) for pattern in BACKEND_LINE_PATTERNS):
            continue
        if re.match(r"[-*]?\s*细则原文（.*?）[:：]\s*$", stripped):
            lines.append("细则原文：")
            continue
        lines.append(line)
    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return text


def strip_markdown_emphasis(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    return text


def source_label(row: dict[str, str]) -> str:
    return f"{row['year']}年{row['district']}{row['stage']} 第{row['question_no']}题"


def question_title(index: int, row: dict[str, str]) -> str:
    domain = DOMAIN_BY_QID.get(row["question_id"], "综合")
    return f"{index:02d}. {row['question_id']}｜{source_label(row)}｜{domain}"


def value_needs_block(value: str) -> bool:
    stripped = value.strip()
    return "\n" in stripped or stripped.startswith(("- ", "• "))


def append_analysis_markdown(lines: list[str], label: str, value: str):
    if not value:
        return
    if value_needs_block(value):
        lines.append(f"**{label}**：")
        for line in value.splitlines():
            clean_line = line.strip()
            if not clean_line:
                continue
            if clean_line.startswith("• "):
                clean_line = "- " + clean_line[2:].strip()
            lines.append(clean_line)
    else:
        lines.append(f"- **{label}**：{value}")


def special_note(row: dict[str, str]) -> str:
    qid = row["question_id"]
    notes = {
        "S16_Q19": "本题按正式评分细则6分整理；原题页未印分值，试卷参考答案页曾标8分，备课时以正式细则为准。",
        "S18_Q19": "本题原题页未印分值，按正式评分细则8分整理。",
        "S33_Q20": "案件三结果④已由同套讲评材料补齐：乙公司赔偿甲经济损失。",
        "S42_Q18": "本题法律小问不依赖图表逐年精确数值；整理时保留题面趋势与第(3)问法律细则。",
        "S62_Q17": "原卷设问写“运用法律知识”，本汇编按不正当竞争与未成年人保护相关法律内容归入《法律与生活》。",
    }
    return notes.get(qid, "")


def build_markdown(rows: list[dict[str, str]]) -> str:
    year_counts = Counter(r["year"] for r in rows)
    domain_counts = Counter(DOMAIN_BY_QID.get(r["question_id"], "综合") for r in rows)
    module_counts = Counter(r["module_decision"] for r in rows)

    lines: list[str] = []
    lines.extend(
        [
            "# 选必二《法律与生活》法律主观题习题与细则汇编（最终版）",
            "",
            f"生成日期：{TODAY}",
            "",
            "本汇编收录 2024—2026 年北京区级模拟题中已核实为《法律与生活》范围的法律主观题 58 道。每题保留“材料、设问、评分细则、作答拆解”，适合课堂训练、作业讲评和二轮专题复习使用。",
            "",
            "## 使用说明",
            "",
            "1. 先遮住评分细则，只看材料和设问独立作答。",
            "2. 对照细则时，优先找“法律关系、规则要件、事实匹配、责任落点、法治价值”五类采分点。",
            "3. 讲评时不要只背结论，要把材料中的关键事实扣回法律规则。",
            "4. 个别题目设问只写“法律知识”或“法治相关知识”，本汇编按题面与正式细则内容归入选必二；正文中只保留对教学有用的整理说明。",
            "",
            "## 三年考法速览",
            "",
            f"- 收录规模：2024年 {year_counts['2024']} 道，2025年 {year_counts['2025']} 道，2026年 {year_counts['2026']} 道。",
            f"- 模块构成：{module_counts.get('法律与生活', 0)} 道为单纯《法律与生活》，{module_counts.get('综合(含法律与生活)', 0)} 道为综合题中的法律小问。",
            "- 高频主域：知识产权与竞争法、侵权责任与人格权、民事总则与基本原则、劳动与就业、合同。",
            "- 2026 年综合题明显增多，常把法律判断嵌入经济、逻辑或社会治理材料中考查。",
            "- 稳定答题骨架：先定法律关系，再抓争点任务，再写规则要件，再扣材料事实，最后落到责任、程序或价值。",
            "",
            "## 主考点分布",
            "",
            "| 主考点域 | 题数 |",
            "| --- | ---: |",
        ]
    )
    for domain, count in domain_counts.most_common():
        lines.append(f"| {domain} | {count} |")

    lines.extend(["", "## 目录", ""])
    by_year: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in rows:
        by_year[row["year"]].append(row)
    index = 1
    for year in sorted(by_year):
        lines.append(f"### {year} 年")
        for row in by_year[year]:
            lines.append(f"- {index:02d}. {row['question_id']}｜{row['district']}{row['stage']}｜第{row['question_no']}题｜{DOMAIN_BY_QID.get(row['question_id'], '综合')}")
            index += 1
        lines.append("")

    lines.append("## 正文")
    lines.append("")

    current_year = ""
    for index, row in enumerate(rows, start=1):
        if row["year"] != current_year:
            current_year = row["year"]
            lines.extend([f"## {current_year} 年", ""])
        lines.extend(
            [
                f"### {question_title(index, row)}",
                "",
                f"**题源**：{source_label(row)}",
                "",
                f"**设问**：{clean_student_text(row['完整设问'])}",
                "",
                "#### 材料",
                "",
                clean_student_text(row["完整材料"]) or "（本题材料见设问。）",
                "",
                "#### 评分细则",
                "",
                clean_student_text(row["对应细则原文"]),
                "",
                "#### 作答拆解",
                "",
            ]
        )
        for label, key in [
            ("法律关系", "法律关系"),
            ("争点与任务", "争点诉讼请求"),
            ("规则要件", "规则要件"),
            ("事实匹配", "事实匹配"),
            ("责任、程序与落点", "责任程序落点"),
            ("法治价值", "法治价值"),
        ]:
            value = clean_student_text(row.get(key, ""))
            append_analysis_markdown(lines, label, value)
        note = special_note(row)
        if note:
            lines.extend(["", f"**整理说明**：{note}"])
        lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None):
    font = run.font
    font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    if size is not None:
        font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if color:
        font.color.rgb = RGBColor.from_string(color.replace("#", ""))


def set_paragraph_spacing(paragraph, before: int = 0, after: int = 6, line: float = 1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.885)
    section.right_margin = Inches(0.885)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name, size, after in [
        ("Intense Quote", 10, 5),
        ("Quote", 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
        style.font.size = Pt(size)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run("选必二《法律与生活》法律主观题习题与细则汇编")
    set_run_font(run, size=9, color="555555")

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("第 ")
    set_run_font(run, size=9, color="555555")
    add_page_number(footer_p)
    run = footer_p.add_run(" 页")
    set_run_font(run, size=9, color="555555")


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_title_page(doc: Document, rows: list[dict[str, str]]):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=10, line=1.15)
    run = p.add_run("选必二《法律与生活》")
    set_run_font(run, size=24, bold=True, color="0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("法律主观题习题与细则汇编")
    set_run_font(run, size=22, bold=True, color="0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"最终版｜{TODAY}")
    set_run_font(run, size=12, color="555555")

    doc.add_paragraph()
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = summary.add_run("2024—2026 北京区级模拟题 · 已核实法律主观题 58 道")
    set_run_font(run, size=12, bold=True, color="1F4D78")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=4, after=14, line=1.25)
    run = p.add_run("每题含材料、设问、评分细则和作答拆解。用于课堂训练时，建议先作答、后对照细则。")
    set_run_font(run, size=11, color="333333")

    add_info_table(
        doc,
        [
            ("收录范围", "2024—2026 年北京区级模拟题中已核实为《法律与生活》范围的法律主观题"),
            ("题目数量", f"{len(rows)} 道"),
            ("整理口径", "正式评分细则优先；个别设问未写书名的题，按题面与细则中的法律内容归入"),
            ("使用方式", "遮住细则独立作答，再按“法律关系、规则要件、事实匹配、责任落点、法治价值”复盘"),
        ],
    )
    doc.add_page_break()


def set_cell_text(cell, text: str, bold: bool = False, fill: str | None = None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line=1.15)
    run = p.add_run(text)
    set_run_font(run, size=10, bold=bold)
    if fill:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        cell._tc.get_or_add_tcPr().append(shading)


def add_info_table(doc: Document, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, (label, value) in enumerate(rows):
        cells = table.rows[row_idx].cells
        set_cell_text(cells[0], label, bold=True, fill="E8EEF5")
        set_cell_text(cells[1], value)
        cells[0].width = Inches(1.181)
        cells[1].width = Inches(5.319)
    doc.add_paragraph()


def add_table_from_markdown(doc: Document, lines: list[str]) -> bool:
    cleaned = [line.strip() for line in lines if line.strip()]
    if len(cleaned) < 2:
        return False
    rows: list[list[str]] = []
    for line in cleaned:
        cells = [strip_markdown_emphasis(c.strip()) for c in line.strip("|").split("|")]
        if all(re.fullmatch(r":?-{2,}:?", c.strip()) for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return False
    width = max(len(r) for r in rows)
    if width > 6:
        return False
    rows = [r + [""] * (width - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            set_cell_text(table.cell(i, j), value, bold=(i == 0), fill=("E8EEF5" if i == 0 else None))
    doc.add_paragraph()
    return True


def add_text_block(doc: Document, text: str, style: str | None = None, rubric: bool = False):
    text = clean_text(text)
    if not text:
        return
    pending_table: list[str] = []

    def flush_table():
        nonlocal pending_table
        if pending_table:
            if not add_table_from_markdown(doc, pending_table):
                for line in pending_table:
                    add_plain_line(doc, line, style=style, rubric=rubric)
            pending_table = []

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if line.strip().startswith("|") and line.strip().endswith("|"):
            pending_table.append(line)
            continue
        flush_table()
        if not line.strip():
            continue
        add_plain_line(doc, line, style=style, rubric=rubric)
    flush_table()


def add_plain_line(doc: Document, line: str, style: str | None = None, rubric: bool = False):
    stripped = strip_markdown_emphasis(line.strip())
    if stripped.startswith("- "):
        stripped = "• " + stripped[2:].strip()
    p = doc.add_paragraph(style=style)
    set_paragraph_spacing(p, after=4 if rubric else 5, line=1.15 if rubric else 1.25)
    if stripped.startswith("• "):
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.1)
    run = p.add_run(stripped)
    set_run_font(run, size=10 if rubric else 11)


def add_doc_sections(doc: Document, rows: list[dict[str, str]]):
    year_counts = Counter(r["year"] for r in rows)
    domain_counts = Counter(DOMAIN_BY_QID.get(r["question_id"], "综合") for r in rows)
    module_counts = Counter(r["module_decision"] for r in rows)

    doc.add_heading("使用说明", level=1)
    for item in [
        "先遮住评分细则，只看材料和设问独立作答。",
        "对照细则时，优先找法律关系、规则要件、事实匹配、责任落点、法治价值。",
        "讲评时把材料中的关键事实扣回法律规则，避免只背结论。",
        "个别设问只写“法律知识”或“法治相关知识”的题，按题面与正式细则内容归入选必二。",
    ]:
        add_plain_line(doc, "• " + item)

    doc.add_heading("三年考法速览", level=1)
    overview_rows = [
        ("收录规模", f"2024年 {year_counts['2024']} 道；2025年 {year_counts['2025']} 道；2026年 {year_counts['2026']} 道。"),
        ("模块构成", f"{module_counts.get('法律与生活', 0)} 道为单纯《法律与生活》；{module_counts.get('综合(含法律与生活)', 0)} 道为综合题中的法律小问。"),
        ("高频主域", "知识产权与竞争法、侵权责任与人格权、民事总则与基本原则、劳动与就业、合同。"),
        ("命题变化", "2026 年综合题明显增多，常把法律判断嵌入经济、逻辑或社会治理材料中考查。"),
        ("答题骨架", "定法律关系；抓争点任务；写规则要件；扣材料事实；落责任、程序或价值。"),
    ]
    add_info_table(doc, overview_rows)

    doc.add_heading("主考点分布", level=1)
    table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(table.cell(0, 0), "主考点域", bold=True, fill="E8EEF5")
    set_cell_text(table.cell(0, 1), "题数", bold=True, fill="E8EEF5")
    for i, (domain, count) in enumerate(domain_counts.most_common(), start=1):
        set_cell_text(table.cell(i, 0), domain)
        set_cell_text(table.cell(i, 1), str(count))
    doc.add_paragraph()

    doc.add_heading("目录", level=1)
    by_year: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in rows:
        by_year[row["year"]].append(row)
    index = 1
    for year in sorted(by_year):
        doc.add_heading(f"{year} 年", level=2)
        for row in by_year[year]:
            add_plain_line(
                doc,
                f"{index:02d}. {row['question_id']}｜{row['district']}{row['stage']}｜第{row['question_no']}题｜{DOMAIN_BY_QID.get(row['question_id'], '综合')}",
                rubric=True,
            )
            index += 1
    doc.add_page_break()

    current_year = ""
    for index, row in enumerate(rows, start=1):
        if row["year"] != current_year:
            if current_year:
                doc.add_page_break()
            current_year = row["year"]
            doc.add_heading(f"{current_year} 年", level=1)
        doc.add_heading(question_title(index, row), level=2)
        add_info_table(
            doc,
            [
                ("题源", source_label(row)),
                ("设问", clean_student_text(row["完整设问"])),
            ],
        )

        doc.add_heading("材料", level=3)
        add_text_block(doc, clean_student_text(row["完整材料"]))

        doc.add_heading("评分细则", level=3)
        add_text_block(doc, clean_student_text(row["对应细则原文"]), style="Quote", rubric=True)

        doc.add_heading("作答拆解", level=3)
        for label, key in [
            ("法律关系", "法律关系"),
            ("争点与任务", "争点诉讼请求"),
            ("规则要件", "规则要件"),
            ("事实匹配", "事实匹配"),
            ("责任、程序与落点", "责任程序落点"),
            ("法治价值", "法治价值"),
        ]:
            value = clean_student_text(row.get(key, ""))
            if value:
                if value_needs_block(value):
                    p = doc.add_paragraph()
                    set_paragraph_spacing(p, after=2, line=1.2)
                    label_run = p.add_run(f"{label}：")
                    set_run_font(label_run, size=10.5, bold=True, color="1F4D78")
                    add_text_block(doc, value, rubric=True)
                else:
                    p = doc.add_paragraph()
                    set_paragraph_spacing(p, after=4, line=1.2)
                    label_run = p.add_run(f"{label}：")
                    set_run_font(label_run, size=10.5, bold=True, color="1F4D78")
                    body_run = p.add_run(strip_markdown_emphasis(value))
                    set_run_font(body_run, size=10.5)
        note = special_note(row)
        if note:
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=2, after=8, line=1.15)
            label_run = p.add_run("整理说明：")
            set_run_font(label_run, size=10, bold=True, color="7A5A00")
            body_run = p.add_run(note)
            set_run_font(body_run, size=10, color="333333")


def build_docx(rows: list[dict[str, str]]):
    doc = Document()
    configure_doc(doc)
    add_title_page(doc, rows)
    add_doc_sections(doc, rows)
    doc.save(FINAL_DOCX)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    rows = read_rows()
    FINAL_MD.write_text(build_markdown(rows), encoding="utf-8")
    build_docx(rows)
    print(FINAL_MD)
    print(FINAL_DOCX)


if __name__ == "__main__":
    main()
