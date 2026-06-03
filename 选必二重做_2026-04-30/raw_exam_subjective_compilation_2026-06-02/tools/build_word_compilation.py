#!/usr/bin/env python3
from __future__ import annotations

import csv
import json
import re
import shutil
import zipfile
from collections import defaultdict
from pathlib import Path

import fitz  # PyMuPDF
from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


RUN_DIR = Path(__file__).resolve().parents[1]
PACKETS = RUN_DIR / "03_source_packets" / "source_packets_final.jsonl"
COVERAGE = RUN_DIR / "00_control" / "COVERAGE_MATRIX.csv"
SOURCE_LEDGER = RUN_DIR / "00_control" / "SOURCE_LEDGER.csv"
OUT_DIR = RUN_DIR / "05_output"
ASSET_DIR = OUT_DIR / "word_assets"
FIG_DIR = ASSET_DIR / "figures"
REPORT_CSV = ASSET_DIR / "word_image_extraction_report.csv"
OUT_DOCX = OUT_DIR / "选必二法律与生活_习题汇编_2024-2026.docx"

FONT_EA = "宋体"
FONT_WEST = "Calibri"
HEADING_BLUE = "2E74B5"
HEADING_DARK = "1F4D78"
TABLE_FILL = "E8EEF5"

VISUAL_RE = re.compile(
    r"(下图|右图|左图|图示|示意图|流程图|漫画|图片|截图|资料卡|下表|表格|完成表|数据|公示|判决书|裁判理由|起诉状)"
)
SCORE_RE = re.compile(
    r"((?:每点|每个要点|每条|一点|答出|写出|指出|说明|分析|判断|理由|观点|内容|等级|一等|二等|三等|赋分|给分|得分|扣分|不超过|最高|共|满分|可得|酌情)[^。\n；;]*?\d+(?:\.\d+)?分|"
    r"\d+(?:\.\d+)?分[^。\n；;]*(?:每点|每个要点|共|满分|酌情|给分|得分|扣分|不超过|最高)[^。\n；;]*)"
)


def set_east_asian(run, font: str = FONT_EA) -> None:
    run.font.name = FONT_WEST
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def set_para_font(paragraph, size_pt: float = 11, bold: bool | None = None, color: str | None = None) -> None:
    for run in paragraph.runs:
        set_east_asian(run)
        run.font.size = Pt(size_pt)
        if bold is not None:
            run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def paragraph_spacing(paragraph, before: int = 0, after: int = 6, line: float = 1.25) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "B8C2CC")


def set_table_fixed_width(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_table_borders(table)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_WEST
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EA)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, HEADING_BLUE, 18, 10),
        ("Heading 2", 13, HEADING_BLUE, 14, 7),
        ("Heading 3", 12, HEADING_DARK, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = FONT_WEST
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EA)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    styles = doc.styles
    if "FG Section Label" not in styles:
        label = styles.add_style("FG Section Label", 1)
        label.font.name = FONT_WEST
        label._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EA)
        label.font.size = Pt(11)
        label.font.bold = True
        label.font.color.rgb = RGBColor.from_string(HEADING_DARK)
        label.paragraph_format.space_before = Pt(6)
        label.paragraph_format.space_after = Pt(2)
        label.paragraph_format.line_spacing = 1.25

    if "FG Score Note" not in styles:
        note = styles.add_style("FG Score Note", 1)
        note.font.name = FONT_WEST
        note._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EA)
        note.font.size = Pt(10)
        note.font.color.rgb = RGBColor.from_string("555555")
        note.paragraph_format.space_before = Pt(2)
        note.paragraph_format.space_after = Pt(8)
        note.paragraph_format.line_spacing = 1.25

    if "FG Caption" not in styles:
        caption = styles.add_style("FG Caption", 1)
        caption.font.name = FONT_WEST
        caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EA)
        caption.font.size = Pt(9)
        caption.font.color.rgb = RGBColor.from_string("555555")
        caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_before = Pt(2)
        caption.paragraph_format.space_after = Pt(6)
        caption.paragraph_format.line_spacing = 1.15


def read_packets() -> list[dict]:
    with PACKETS.open(encoding="utf-8") as f:
        return [json.loads(line) for line in f if line.strip()]


def read_coverage() -> list[dict]:
    with COVERAGE.open(encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def read_source_ledger() -> dict[str, dict]:
    with SOURCE_LEDGER.open(encoding="utf-8-sig", newline="") as f:
        return {row["source_id"]: row for row in csv.DictReader(f)}


def clean_score_line(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip(" ；;。")


def total_scores(text: str) -> list[str]:
    hits = []
    for m in re.finditer(r"[（(]\s*(\d+(?:\.\d+)?)\s*分\s*[）)]", text):
        hits.append(f"{m.group(1)}分")
    for m in re.finditer(r"(?:共|满分|本题|本问)\s*(\d+(?:\.\d+)?)\s*分", text):
        hits.append(f"{m.group(1)}分")
    seen = []
    for h in hits:
        if h not in seen:
            seen.append(h)
    return seen


def score_distribution(entry: dict) -> str:
    rubric = entry.get("rubric", "")
    prompt = entry.get("prompt", "")
    prompt_totals = total_scores(prompt)
    lines = [clean_score_line(x) for x in rubric.splitlines() if clean_score_line(x)]
    detailed: list[str] = []
    for line in lines:
        if SCORE_RE.search(line):
            if re.fullmatch(r"\d+[.．、]?[（(]?\d+(?:\.\d+)?分[）)]?", line):
                continue
            detailed.append(line)
    compact: list[str] = []
    for item in detailed:
        if item not in compact:
            compact.append(item)
    if compact:
        text = "；".join(compact)
        if prompt_totals and prompt_totals[0] not in text:
            text = f"本问共{prompt_totals[0]}；" + text
        return text if len(text) <= 480 else text[:477] + "..."

    totals = prompt_totals or total_scores(rubric)
    if totals:
        return f"本题/本问共{totals[0]}；原细则未细分到逐点分值"
    return "原细则未显式标出分值；需人工按原卷/细则复核"


def normalize(text: str) -> str:
    return re.sub(r"[\s　，。；;：:、（）()\[\]【】《》“”\"'·.．,\-—_]+", "", text or "")


def split_pages(text: str) -> list[tuple[int, str]]:
    parts: list[tuple[int, str]] = []
    current_page = 1
    buf: list[str] = []
    for line in text.splitlines():
        m = re.match(r"^===== .* page (\d+) =====$", line.strip())
        if not m:
            m = re.match(r"^\[PDF_PAGE (\d+)\]$", line.strip())
        if m:
            if buf:
                parts.append((current_page, "\n".join(buf)))
                buf = []
            current_page = int(m.group(1))
        else:
            buf.append(line)
    if buf:
        parts.append((current_page, "\n".join(buf)))
    return parts


def likely_visual(entry: dict) -> bool:
    text = f"{entry.get('material', '')}\n{entry.get('prompt', '')}"
    return bool(VISUAL_RE.search(text))


def question_number(entry: dict) -> str:
    m = re.search(r"第(\d+)题", entry.get("title", ""))
    return m.group(1) if m else str(entry.get("question_no", "")).split("(")[0]


def locate_page(entry: dict, text_path: str | None) -> int | None:
    if not text_path:
        return None
    path = Path(text_path)
    if not path.exists():
        return None
    pages = split_pages(path.read_text(encoding="utf-8", errors="ignore"))
    if not pages:
        return None

    material = normalize(entry.get("material", ""))
    prompt = normalize(entry.get("prompt", ""))
    snippets = [s for s in [material[:35], material[35:70], prompt[:35]] if len(s) >= 8]
    qnumb = question_number(entry)
    best_page, best_score = None, -1
    for page_no, page_text in pages:
        norm = normalize(page_text)
        score = 0
        for snippet in snippets:
            if snippet and snippet in norm:
                score += 100 + min(len(snippet), 35)
            elif snippet:
                overlap = sum(1 for i in range(0, max(0, len(snippet) - 5), 6) if snippet[i : i + 6] in norm)
                score += overlap * 6
        if qnumb and re.search(rf"(^|\D){re.escape(qnumb)}[.．、]", page_text):
            score += 30
        if VISUAL_RE.search(page_text):
            score += 5
        if score > best_score:
            best_page, best_score = page_no, score
    if best_score <= 0:
        return None
    return best_page


def render_pdf_page(pdf_path: Path, page_no: int, dest: Path) -> Path | None:
    if dest.exists():
        return dest
    try:
        pdf = fitz.open(pdf_path)
        if not (1 <= page_no <= len(pdf)):
            return None
        page = pdf[page_no - 1]
        pix = page.get_pixmap(matrix=fitz.Matrix(1.65, 1.65), alpha=False)
        pix.save(dest)
        pdf.close()
        return dest
    except Exception:
        return None


def extract_docx_media(docx_path: Path, dest_dir: Path) -> list[Path]:
    if not docx_path.exists():
        return []
    dest_dir.mkdir(parents=True, exist_ok=True)
    out: list[Path] = []
    try:
        with zipfile.ZipFile(docx_path) as z:
            media = [x for x in z.namelist() if x.startswith("word/media/")]
            for name in media:
                suffix = Path(name).suffix.lower()
                if suffix not in {".png", ".jpg", ".jpeg", ".webp"}:
                    continue
                out_path = dest_dir / f"{docx_path.stem}_{Path(name).name}"
                if not out_path.exists():
                    with z.open(name) as src, out_path.open("wb") as dst:
                        shutil.copyfileobj(src, dst)
                out.append(out_path)
    except zipfile.BadZipFile:
        return []
    return out


def collect_assets(entries: list[dict], source_map: dict[str, dict]) -> dict[str, list[tuple[Path, str]]]:
    if FIG_DIR.exists():
        shutil.rmtree(FIG_DIR)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    assets: dict[str, list[tuple[Path, str]]] = defaultdict(list)
    report_rows: list[dict] = []
    for entry in entries:
        eid = entry["entry_id"]
        note = ""
        if not likely_visual(entry):
            report_rows.append(
                {
                    "entry_id": eid,
                    "title": entry["title"],
                    "visual_signal": "no",
                    "asset_count": "0",
                    "asset_paths": "",
                    "note": "locked text has no explicit image/table cue",
                }
            )
            continue

        qpath = Path(entry.get("question_path", ""))
        ext = qpath.suffix.lower()
        found: list[tuple[Path, str]] = []
        if ext == ".pdf":
            source = source_map.get(entry.get("question_source_id", ""), {})
            page_no = locate_page(entry, source.get("text_path"))
            if page_no:
                dest = FIG_DIR / f"{entry['question_source_id']}_p{page_no}.png"
                rendered = render_pdf_page(qpath, page_no, dest)
                if rendered:
                    found.append((rendered, f"原题图表页摘取：{qpath.name} 第{page_no}页"))
                else:
                    note = f"failed to render PDF page {page_no}"
            else:
                note = "could not locate matching source page from extracted text"
        elif ext == ".docx":
            media = extract_docx_media(qpath, FIG_DIR / "docx_media" / entry["question_source_id"])
            for img in media[:6]:
                found.append((img, f"原题内嵌图片摘取：{qpath.name}"))
            if not media:
                note = "DOCX source has no extractable embedded raster image; table/box may be native text"
            elif len(media) > 6:
                note = f"DOCX source has {len(media)} embedded images; first 6 included for this entry"
        else:
            note = f"unsupported question source extension: {ext}"

        assets[eid] = found
        report_rows.append(
            {
                "entry_id": eid,
                "title": entry["title"],
                "visual_signal": "yes",
                "asset_count": str(len(found)),
                "asset_paths": ";".join(str(p) for p, _ in found),
                "note": note,
            }
        )

    with REPORT_CSV.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=["entry_id", "title", "visual_signal", "asset_count", "asset_paths", "note"],
        )
        writer.writeheader()
        writer.writerows(report_rows)
    return assets


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    paragraph_spacing(p, after=3, line=1.15)
    r = p.add_run(title)
    set_east_asian(r)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("000000")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    paragraph_spacing(p, after=12, line=1.15)
    r = p.add_run(subtitle)
    set_east_asian(r)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("555555")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_text_block(doc: Document, label: str, text: str) -> None:
    doc.add_paragraph(f"【{label}】", style="FG Section Label")
    p = doc.add_paragraph()
    paragraph_spacing(p)
    run = p.add_run(text.strip() if text else "")
    set_east_asian(run)
    run.font.size = Pt(11)


def add_score_note(doc: Document, distribution: str) -> None:
    p = doc.add_paragraph(style="FG Score Note")
    r = p.add_run(f"（分数分布：{distribution}）")
    set_east_asian(r)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("555555")


def add_assets(doc: Document, entry: dict, assets: list[tuple[Path, str]]) -> None:
    if not assets:
        return
    doc.add_paragraph("【图片/图表】", style="FG Section Label")
    for idx, (img, caption) in enumerate(assets, start=1):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            with Image.open(img) as im:
                width_px, height_px = im.size
            max_width = 5.85
            max_height = 8.1
            scaled_height = max_width * (height_px / width_px) if width_px else max_height
            if scaled_height > max_height:
                p.add_run().add_picture(str(img), height=Inches(max_height))
            else:
                p.add_run().add_picture(str(img), width=Inches(max_width))
            cap = doc.add_paragraph(f"图{idx} {caption}", style="FG Caption")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as exc:
            p = doc.add_paragraph(style="FG Score Note")
            r = p.add_run(f"（图片插入失败：{img.name}；{exc}）")
            set_east_asian(r)


def add_coverage_table(doc: Document, coverage: list[dict]) -> None:
    doc.add_heading("覆盖清单", level=1)
    headers = ["年份", "区/统考", "卷型", "是否含选必二题", "抽取设问数量", "待确认"]
    rows = []
    for row in coverage:
        rows.append(
            [
                row["year"],
                row["district_or_exam"],
                row["paper_type"],
                "是" if row["legal_subjective_count"] != "0" else "否",
                row["legal_subjective_count"],
                row["pending_confirm_count"],
            ]
        )
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    widths = [720, 2500, 900, 1750, 1300, 950]
    set_table_fixed_width(table, widths)
    for col, header in enumerate(headers):
        cell = table.cell(0, col)
        set_cell_shading(cell, TABLE_FILL)
        p = cell.paragraphs[0]
        r = p.add_run(header)
        set_east_asian(r)
        r.font.bold = True
        r.font.size = Pt(8.5)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            r = p.add_run(value)
            set_east_asian(r)
            r.font.size = Pt(8.2)
            paragraph_spacing(p, after=0, line=1.0)


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("选必二《法律与生活》习题汇编 2024-2026")
    set_east_asian(r)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("777777")


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    entries = read_packets()
    coverage = read_coverage()
    source_map = read_source_ledger()
    assets = collect_assets(entries, source_map)

    doc = Document()
    style_document(doc)
    add_footer(doc)
    add_title(doc, "选必二《法律与生活》习题汇编（2024-2026）", "Word版：含题图/表格页摘取与细则分数分布括注")
    add_coverage_table(doc, coverage)
    doc.add_page_break()

    doc.add_heading("题目汇编", level=1)
    current_year = None
    current_suite = None
    pending_list: list[str] = []
    for entry in entries:
        year = str(entry["year"])
        suite = f"{entry['district_or_exam']} · {entry['paper_type']}"
        if year != current_year:
            current_year = year
            doc.add_heading(year, level=1)
            current_suite = None
        if suite != current_suite:
            current_suite = suite
            doc.add_heading(suite, level=2)

        doc.add_heading(entry["title"], level=3)
        add_text_block(doc, "材料", entry.get("material", ""))
        add_assets(doc, entry, assets.get(entry["entry_id"], []))
        add_text_block(doc, "设问", entry.get("prompt", ""))
        add_text_block(doc, "细则", entry.get("rubric", ""))
        add_score_note(doc, score_distribution(entry))
        if entry.get("pending_reason"):
            pending_list.append(f"{entry['title']}：{entry['pending_reason']}")
            p = doc.add_paragraph(style="FG Score Note")
            r = p.add_run(f"【待确认理由】{entry['pending_reason']}")
            set_east_asian(r)
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor.from_string("9B1C1C")

    doc.add_page_break()
    doc.add_heading("统计", level=1)
    p = doc.add_paragraph(f"总计设问数：{len(entries)}")
    set_para_font(p)
    doc.add_heading("【待确认】条目清单", level=2)
    if pending_list:
        for item in pending_list:
            p = doc.add_paragraph(item)
            paragraph_spacing(p, after=4)
            set_para_font(p, size_pt=10)
    else:
        p = doc.add_paragraph("无")
        set_para_font(p)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)
    print(REPORT_CSV)


if __name__ == "__main__":
    build_doc()
