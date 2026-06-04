#!/usr/bin/env python3
from __future__ import annotations

import json
import re
import subprocess
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import fitz
from docx import Document

RUN_DIR = Path(__file__).resolve().parents[1]
PACKETS = RUN_DIR / "03_source_packets" / "source_packets_final.jsonl"
OUT_MD = RUN_DIR / "05_output" / "A_THEME_SOURCE_REPAIR_AUDIT_20260604.md"
OUT_JSON = RUN_DIR / "05_output" / "A_THEME_SOURCE_REPAIR_AUDIT_20260604.json"

TARGET_IDS = [
    "E005", "E006", "E015", "E016", "E022", "E023", "E024",
    "E035", "E036", "E053", "E074", "E009", "E043",
]

SEARCH_TERMS = {
    "E005": ["平台经济", "劳动人事争议典型案例", "王某于", "从属性"],
    "E006": ["虚拟数字人", "Ada", "首例涉"],
    "E015": ["法治以", "固根本", "完成下表", "案件一 张某"],
    "E016": ["法律以清晰的边界", "龙某为 13", "钱某", "完成下表"],
    "E022": ["AI绘图", "AI 绘画", "李某是一名", "网友刘某"],
    "E023": ["信用卡透支", "小王", "履行还款义务", "合同约定"],
    "E024": ["了解案件", "拥堵延时指数", "摇摆桥", "参考示例"],
    "E035": ["17岁的小刘", "打赏", "平台监管", "第一百五十七条"],
    "E036": ["孙某", "工作经历", "试用期", "甲公司以孙某"],
    "E053": ["校园欺凌", "惩", "教", "治安管理处罚法"],
    "E074": ["财产制度助力经济社会发展", "某校学习小组", "财产制度改革"],
    "E009": ["一米法庭", "天灾", "物业公司", "司法确认"],
    "E043": ["郭某", "餐厅", "台阶区域", "安全保障义务"],
}

MANUAL_PATH_HINTS = {
    "E009": {
        "rubric": ["2024石景山一模", "其他材料", "2024年石景山一模.pptx"],
    },
}


def desktop_roots() -> list[Path]:
    root = Path.home() / "Desktop"
    return [p for p in root.iterdir() if p.is_dir() and p.name.startswith(("2024", "2025", "2026"))]


def tokenise(text: str) -> list[str]:
    return re.findall(r"[\u4e00-\u9fffA-Za-z0-9]+", text or "")


def file_index() -> list[Path]:
    files: list[Path] = []
    for root in desktop_roots():
        for path in root.rglob("*"):
            if path.is_file() and path.suffix.lower() in {".pdf", ".docx", ".doc", ".pptx", ".ppt"}:
                files.append(path)
    return files


def resolve_path(mac_path: str, files: list[Path], entry_id: str = "", role_hint: str = "") -> Path | None:
    hints = MANUAL_PATH_HINTS.get(entry_id, {}).get(role_hint, [])
    if hints:
        hinted = [p for p in files if all(h in str(p) for h in hints)]
        if hinted:
            return sorted(hinted, key=lambda p: len(str(p)))[0]
    basename = Path(mac_path).name
    suffix = Path(mac_path).suffix.lower()
    hits = [p for p in files if p.name == basename]
    if not hits:
        hits = [p for p in files if p.suffix.lower() == suffix]
    ignore = {"Users", "wanglifei", "Desktop", "模拟题", "试卷", "细则", "各区", "2024", "2025", "2026"}
    toks = [t for t in tokenise(mac_path) if t not in ignore]
    scored: list[tuple[int, str, Path]] = []
    role = "rubric" if "细则" in mac_path else "question" if "试卷" in mac_path else ""
    for path in hits:
        s = str(path)
        score = 0
        for tok in toks:
            if tok in s:
                score += 3
            elif len(tok) >= 2 and tok in path.name:
                score += 2
            elif len(tok) >= 2 and tok in path.parent.name:
                score += 1
        if path.name == basename:
            score += 2
        if any(tok in s for tok in toks if tok not in {"pdf", "docx", "doc", "pptx"}):
            score += 6
        if any(tok in path.name for tok in toks):
            score += 4
        if role == "rubric":
            if "细则" in s:
                score += 12
            if "试卷" in s:
                score -= 4
        elif role == "question":
            if "试卷" in s:
                score += 8
            if "细则" in s:
                score -= 4
        scored.append((score, s, path))
    scored.sort(reverse=True)
    if scored and scored[0][0] > 0:
        return scored[0][2]
    return None


def clean(text: str) -> str:
    text = (text or "").replace("\u3000", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf(path: Path) -> str:
    parts: list[str] = []
    doc = fitz.open(path)
    try:
        for i, page in enumerate(doc):
            text = page.get_text("text")
            if text.strip():
                parts.append(f"\n--- page {i + 1} ---\n{text}")
    finally:
        doc.close()
    return clean("\n".join(parts))


def extract_docx(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for ti, table in enumerate(doc.tables, start=1):
        parts.append(f"[table {ti}]")
        for row in table.rows:
            cells = [clean(cell.text) for cell in row.cells]
            parts.append(" | ".join(cells))
    return clean("\n".join(parts))


def extract_pptx(path: Path) -> str:
    ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    parts: list[str] = []
    with zipfile.ZipFile(path) as zf:
        slides = sorted(
            [name for name in zf.namelist() if re.match(r"ppt/slides/slide\d+\.xml$", name)],
            key=lambda name: int(re.search(r"slide(\d+)\.xml", name).group(1)),
        )
        for slide in slides:
            xml = zf.read(slide)
            root = ET.fromstring(xml)
            texts = [node.text for node in root.findall(".//a:t", ns) if node.text]
            if texts:
                parts.append(f"\n--- {slide} ---\n" + "\n".join(texts))
    return clean("\n".join(parts))


def extract_doc(path: Path) -> str:
    ps = r"""
$ErrorActionPreference = 'Stop'
$OutputEncoding = [System.Text.UTF8Encoding]::new()
[Console]::OutputEncoding = [System.Text.UTF8Encoding]::new()
$path = $args[0]
$word = $null
$doc = $null
try {
  $word = New-Object -ComObject Word.Application
  $word.Visible = $false
  $doc = $word.Documents.Open($path, $false, $true)
  $doc.Content.Text
} finally {
  if ($doc -ne $null) { $doc.Close($false) | Out-Null }
  if ($word -ne $null) { $word.Quit() | Out-Null }
}
"""
    with tempfile.NamedTemporaryFile("w", suffix=".ps1", delete=False, encoding="utf-8") as tmp:
        tmp.write(ps)
        script = tmp.name
    try:
        proc = subprocess.run(
            ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-File", script, str(path)],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=60,
        )
        if proc.returncode != 0:
            return f"[EXTRACT_ERROR] powershell_doc_extract: {proc.stderr.strip()}"
        return clean(proc.stdout)
    finally:
        Path(script).unlink(missing_ok=True)


def extract_text(path: Path | None) -> str:
    if path is None or not path.exists():
        return ""
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return extract_pdf(path)
        if suffix == ".docx":
            return extract_docx(path)
        if suffix == ".pptx":
            return extract_pptx(path)
        if suffix == ".doc":
            return extract_doc(path)
        return ""
    except Exception as exc:
        return f"[EXTRACT_ERROR] {type(exc).__name__}: {exc}"


def section_around(text: str, pos: int) -> str:
    starts = [m.start() for m in re.finditer(r"(?m)(?:^|\n)\s*\d{1,2}[\.．、]", text)]
    before = [s for s in starts if s <= pos]
    after = [s for s in starts if s > pos]
    start = max(0, (before[-1] if before else pos) - 250)
    end = min(len(text), (after[0] if after else pos + 4200) + 250)
    if end - start < 1800:
        start = max(0, pos - 900)
        end = min(len(text), pos + 3600)
    return text[start:end]


def best_snippet(full_text: str, entry: dict) -> str:
    if not full_text:
        return ""
    priority_keys = SEARCH_TERMS.get(entry.get("entry_id", ""), [])
    priority_positions = [full_text.find(k) for k in priority_keys if k and full_text.find(k) >= 0]
    if priority_positions:
        return section_around(full_text, min(priority_positions))
    keys: list[str] = []
    qno = str(entry.get("question_no") or "")
    if qno:
        keys.extend([f"\n{qno}.", f"\n{qno}．", f"\n{qno}、", f"{qno}."])
    prompt = clean(entry.get("prompt", ""))
    if prompt:
        keys.extend([prompt[:18], prompt[:10]])
    material = clean(entry.get("material", ""))
    if material:
        keys.extend([material[:18], material[:10]])
    positions = [full_text.find(k) for k in keys if k and full_text.find(k) >= 0]
    if not positions:
        return full_text[:2500]
    return section_around(full_text, min(positions))


def main() -> None:
    packets = [json.loads(line) for line in PACKETS.read_text(encoding="utf-8").splitlines() if line.strip()]
    files = file_index()
    audit_rows = []
    md = ["# A类厚版待重排题回源审计", "", f"- indexed_source_files: {len(files)}", ""]
    for eid in TARGET_IDS:
        entry = next(row for row in packets if row["entry_id"] == eid)
        q_path = resolve_path(entry.get("question_path", ""), files, eid, "question")
        r_path = resolve_path(entry.get("rubric_path", ""), files, eid, "rubric")
        q_text = extract_text(q_path)
        r_text = extract_text(r_path)
        q_snip = best_snippet(q_text, entry)
        r_snip = best_snippet(r_text, entry)
        row = {
            "entry_id": eid,
            "title": entry["title"],
            "question_path": str(q_path) if q_path else "",
            "rubric_path": str(r_path) if r_path else "",
            "question_extract_len": len(q_text),
            "rubric_extract_len": len(r_text),
            "question_snippet": q_snip,
            "rubric_snippet": r_snip,
            "current_material_len": len(entry.get("material", "")),
            "current_prompt_len": len(entry.get("prompt", "")),
            "current_rubric_len": len(entry.get("rubric", "")),
        }
        audit_rows.append(row)
        md.extend([
            f"## {eid} {entry['title']}",
            "",
            f"- question_path: `{row['question_path']}`",
            f"- rubric_path: `{row['rubric_path']}`",
            f"- extracted_lengths: question={row['question_extract_len']}; rubric={row['rubric_extract_len']}",
            f"- current_lengths: material={row['current_material_len']}; prompt={row['current_prompt_len']}; rubric={row['current_rubric_len']}",
            "",
            "### question snippet",
            "",
            "```text",
            q_snip[:4000],
            "```",
            "",
            "### rubric snippet",
            "",
            "```text",
            r_snip[:4000],
            "```",
            "",
        ])
    OUT_JSON.write_text(json.dumps(audit_rows, ensure_ascii=False, indent=2), encoding="utf-8")
    OUT_MD.write_text("\n".join(md), encoding="utf-8")
    print(OUT_MD)
    print(OUT_JSON)


if __name__ == "__main__":
    main()
