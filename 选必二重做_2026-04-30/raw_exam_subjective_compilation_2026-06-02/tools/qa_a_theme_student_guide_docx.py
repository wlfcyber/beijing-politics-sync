#!/usr/bin/env python3
from __future__ import annotations

import argparse
from collections import Counter
from pathlib import Path

from docx import Document


RUN_DIR = Path(__file__).resolve().parents[1]
OUT_DIR = RUN_DIR / "05_output"
QA_MD = OUT_DIR / "A_THEME_STUDENT_GUIDE_DOCX_QA_20260604.md"

FIVE_LABELS = ["【入口】", "【判定依据】", "【题目材料】", "【设问】", "【细则】", "【答案落点】"]
METHOD_LABELS = [
    "【本章怎么学】",
    "【题源画像】",
    "【一眼判断】",
    "【采分动作】",
    "【考场句型】",
    "【边界提醒】",
    "【本章先背】",
    "【核心链条】",
    "【常见设问翻译】",
    "【采分模板】",
    "【错答清单】",
    "【真题分组导读】",
]
PHASE2_LABELS = [
    "【命题人路径】",
    "【判题四步】",
    "【高频给分件】",
    "【易混边界】",
]
BAD_TOKENS = ["SRC_", "source_id", "entry_E"]
STUDENT_NOISE = [
    "题眼",
    "评分锚点",
    "PAGEPAGE",
    "阅卷总结",
    "学生问题",
    "学生表现",
    "教师教学",
    "改进措施",
    "全球治理倡议",
    "上合组织",
    "全球南方",
]


def latest_a_theme_docx() -> Path:
    candidates = [
        p
        for p in OUT_DIR.glob("*.docx")
        if "_A" in p.name and "20260604" in p.name and "~$" not in p.name
    ]
    if not candidates:
        raise FileNotFoundError("No A-theme 20260604 DOCX found in 05_output")
    return max(candidates, key=lambda p: p.stat().st_mtime)


def count_text(text: str, token: str) -> int:
    return text.count(token)


def count_label_paragraphs(paragraphs, label: str) -> int:
    return sum(1 for paragraph in paragraphs if paragraph.text.strip().startswith(label))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--word-smoke", default="not_run")
    parser.add_argument(
        "--render-status",
        default="not_run",
        help="render_docx status line, e.g. failed: missing LibreOffice/soffice",
    )
    args = parser.parse_args()

    docx = latest_a_theme_docx()
    doc = Document(docx)
    paragraphs = doc.paragraphs
    all_text = "\n".join(p.text for p in paragraphs)
    styles = Counter(p.style.name for p in paragraphs)
    heading2 = styles.get("Heading 2", 0)
    heading3 = styles.get("Heading 3", 0)

    lines = [
        "# A类十主题学生宝典 DOCX QA",
        "",
        "## 1. 文件",
        "",
        f"- DOCX: `05_output/{docx.name}`",
        f"- 桌面副本: `C:\\Users\\Administrator\\Desktop\\{docx.name}`",
        "- 覆盖报告: `05_output/A_THEME_STUDENT_GUIDE_COVERAGE_REPORT_20260604.md`",
        "- 回源审计: `05_output/A_THEME_SOURCE_REPAIR_AUDIT_20260604.md`",
        "- 回源覆盖层: `05_output/A_THEME_SOURCE_REPAIR_OVERRIDES_20260604.json`",
        "- 边界裁决覆盖层: `05_output/A_THEME_PENDING_BOUNDARY_RESOLUTIONS_20260604.json`",
        "- 细则原文QA: `05_output/A_THEME_RUBRIC_SOURCE_INTEGRITY_QA_20260604.md`",
        "- 剩余待核复查: `05_output/A_THEME_REMAINING_PENDING_RECHECK_20260604.md`",
        "- Word COM视觉QA: `05_output/A_THEME_WORD_COM_VISUAL_QA_20260604.md`",
        "- 视觉接触表复查: `05_output/A_THEME_VISUAL_CONTACT_SHEET_REVIEW_20260604.md`",
        "",
        "## 2. 本阶段新增结构",
        "",
        "- A类十主题速查页: 1",
        "- 章首学习页: 10",
        "- 每章新增学习卡片: 本章怎么学 / 题源画像 / 一眼判断 / 采分动作 / 考场句型 / 边界提醒",
        "- 每章二轮加深卡片: 命题人路径 / 判题四步 / 高频给分件 / 易混边界",
        "- 每章新增方法页组件: 本章先背 / 核心链条 / 常见设问翻译 / 采分模板 / 错答清单 / 真题分组导读",
        "- 封面标题: `《法律与生活》A类十主题学生宝典工作稿`",
        "",
        "## 3. 结构核验",
        "",
        "| 项目 | 结果 |",
        "| --- | ---: |",
        f"| DOCX 文件大小 | {docx.stat().st_size} bytes |",
        f"| Word 段落数 | {len(paragraphs)} |",
        f"| Word 表格数 | {len(doc.tables)} |",
        f"| A类 Heading 2 | {heading2}（含 A类十主题速查 + A1-A10） |",
        f"| Heading 3 真题条目 | {heading3} |",
    ]
    for label in METHOD_LABELS:
        lines.append(f"| {label.strip('【】')} | {count_label_paragraphs(paragraphs, label)} |")
    for label in PHASE2_LABELS:
        lines.append(f"| {label.strip('【】')} | {count_label_paragraphs(paragraphs, label)} |")
    lines.extend(
        [
            f"| 回源修复 | {count_text(all_text, '【回源修复】')} |",
            f"| 待回源重排 | {count_text(all_text, '【待回源重排】')} |",
            f"| 边界裁决 | {count_text(all_text, '【边界裁决】')} |",
            f"| 待核边界 | {count_text(all_text, '【待核边界】')} |",
            f"| 末尾待核/待补清单 | {count_text(all_text, '八、待核/待补清单')} |",
            f"| 工程痕迹 `SRC_` / `source_id` / `entry_E` | {sum(count_text(all_text, t) for t in BAD_TOKENS)} |",
            "",
            "## 4. 五件套核验",
            "",
            "| 标签 | 段落数 |",
            "| --- | ---: |",
        ]
    )
    for label in FIVE_LABELS:
        lines.append(f"| {label} | {count_label_paragraphs(paragraphs, label)} |")
    lines.extend(
        [
            "",
            "## 5. 学生正文噪音扫描",
            "",
            "| 词/模式 | 命中 |",
            "| --- | ---: |",
        ]
    )
    for token in BAD_TOKENS + STUDENT_NOISE:
        lines.append(f"| `{token}` | {count_text(all_text, token)} |")
    lines.extend(
        [
            "",
            "## 6. Word / 渲染",
            "",
            f"- Word COM smoke test: `{args.word_smoke}`",
            f"- `render_docx.py --emit_pdf`: {args.render_status}",
            "- 若 render_status 记录 Word COM -> PDF -> PNG，则本阶段另以 `A_THEME_WORD_COM_VISUAL_QA_20260604.md` 和接触表复查作为视觉渲染证据。",
            "",
            "## 7. 状态边界",
            "",
            "- 本稿是 A类十主题学生化加工稿，不是最终闭合稿。",
            "- 覆盖口径仍为 63 套 / 67 道大题 / 74 个分问。",
            "- Word 末尾已单列待核/待补清单。",
            "- 原始 18 个 pending_reason 中，本轮边界裁决清除 16 个，正文仍保留 2 个【待核边界】。",
            "- 未闭合待回源重排/材料风险项为 0。",
            "",
        ]
    )
    QA_MD.write_text("\n".join(lines), encoding="utf-8")
    print(QA_MD)
    print(f"paragraphs={len(paragraphs)} heading2={heading2} heading3={heading3} tables={len(doc.tables)}")
    print("five_labels=" + ", ".join(f"{label}:{count_label_paragraphs(paragraphs, label)}" for label in FIVE_LABELS))
    print("method_labels=" + ", ".join(f"{label}:{count_label_paragraphs(paragraphs, label)}" for label in METHOD_LABELS))
    print("phase2_labels=" + ", ".join(f"{label}:{count_label_paragraphs(paragraphs, label)}" for label in PHASE2_LABELS))
    print("bad_tokens=" + ", ".join(f"{token}:{count_text(all_text, token)}" for token in BAD_TOKENS + STUDENT_NOISE))


if __name__ == "__main__":
    main()
