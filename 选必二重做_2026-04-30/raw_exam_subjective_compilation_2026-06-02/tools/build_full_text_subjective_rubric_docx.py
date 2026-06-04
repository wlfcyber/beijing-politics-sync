#!/usr/bin/env python3
from __future__ import annotations

import csv
import importlib.util
import json
import re
import shutil
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


RUN_DIR = Path(__file__).resolve().parents[1]
PACKETS = RUN_DIR / "03_source_packets" / "source_packets_final.jsonl"
COVERAGE = RUN_DIR / "00_control" / "COVERAGE_MATRIX.csv"
SOURCE_LEDGER = RUN_DIR / "00_control" / "SOURCE_LEDGER.csv"
OUT_DIR = RUN_DIR / "05_output"
OUT_DOCX = OUT_DIR / "选必二法律与生活_主观题与细则_全量汇编.docx"
DESKTOP_COPY = Path.home() / "Desktop" / OUT_DOCX.name
CHECK_MD = OUT_DIR / "FULL_TEXT_DELIVERY_CHECK.md"

FONT_EA = "宋体"
FONT_WEST = "Calibri"
HEADING_BLUE = "2E74B5"
HEADING_DARK = "1F4D78"
TABLE_FILL = "E8EEF5"

FORMAL_INCOMPLETE = {
    "E009": "原始石景山一模细则 PPT 未找到本题带分值分布的正式阅卷细则；该分问保留为边界备注。",
}

EXTRA_RUBRIC_OVERRIDES = {
    "E007": """19.（6分）
【答案】
在本案判决中，人民法院确认居委会提供的遗赠扶养协议，认定居委会的赡养行为，充分肯定居委会对老人养老送终所起的作用。本案判决不仅实现了个案上的公平正义，更倡导全社会积极助力养老，促进“不尽孝者少分或者不分遗产”的司法理念深入人心。判决有利于弘扬尊老、敬老、爱老、助老的中华传统美德，让崇尚和践行社会主义核心价值观成为人们的自觉行动和全体社会的良好风尚。
【细则】
判决内容：居委会尽到扶养义务，遗赠扶养协议合法有效（1分）
社会价值：
维护居委会合法权益/使老人的生活得到保障（1分）
替代：有利于权利义务相统一（1分）
公平正义（2分）
替代：公序良俗/诚信原则/提高法治意识/司法理念深入人心（1分）
弘扬中华传统美德/践行社会主义核心价值观（2分）
替代：孝亲敬老（1分）""",
    "E028": """21．（15分）
（1）小海的观点片面。新中国成立初期尚未完全建立夫妻共同财产制度，在家庭生活中男方经济地位更强，根据权利与义务相对应原则，由男方清偿具有合理性，有利于保护当事各方的利益。随着时代的发展，夫妻双方的经济状况发生变化，具备了“共同偿还”“协议清偿”的条件。良法应当符合国情和实际，符合社会发展的需求。婚姻法与民法典的规定都符合良法的要求。
【细则】
观点（1分）：片面；错误；正确；（认同/同意）。
分析（5分）：
角度：时代背景
知识：从中国实际出发/科学立法符合国情和实际（2分）
变通：科学立法/由社会实际情况决定/立足国情/符合实际（1分）
材料分析：分析1950年和2021年具体社会情况（1分）；对比男女的经济地位（1950）（1分）
角度：男女个体
知识：合理设定权利与义务/法律面前人人平等/公平/尊重保障人权（1分）""",
    "E031": """18．（7分）
小黄与主题乐园签订的格式合同意思表示真实，条款符合公平原则，合同有效。小黄利用审查漏洞不当获利的行为违反合同约定，违背诚信原则。主题乐园可以根据合同约定，采取撤销年卡的方式解除合同。法院判决保护了主题乐园合法权益，有利于营造公平公正市场环境。
【细则】
合同有效（2分）
2分：条款符合公平原则/权利义务相统一，合同有效
1分：格式合同意思表示真实，合同有效
0分：意思表示真实/合同有效

案情分析（3分）
小黄的行为违反合同约定/违约，（1分）违背诚信原则（1分）
主题乐园可以根据合同约定，采取撤销年卡的方式解除合同。（1分）

判决意义（2分）
法院判决保护了主题乐园合法权益
有利于平衡经营者和消费者的关系
有利于营造公平公正市场环境/维护良好的市场秩序
有利于维护社会公平正义
有利于践行诚信的社会主义核心价值观/诚信的道德风尚
（以上任意出现2条，给2分）""",
    "E043": """20.（8 分）
人民法院以事实为根据、以法律为准绳，作出了上述判决。根据民法典规定，宾馆、商场等公共场所经营者、管理者因过错造成他人损害的，应当承担侵权责任。被告对原告的安全保障义务应保持在合理限度内，且相关证据表明事发现场不存在影响原告通行的客观因素。原告是完全民事行为能力人，摔倒是其自身未尽安全注意义务所致。因此，某餐饮公司和某商业管理公司对此并无过错，不应承担赔偿责任。
该判决有利于平衡原被告双方的权利与义务，倡导安全文明出行和自我负责的安全责任意识，明确经营场所、公共场所的经营者、管理者安全保障义务的边界，弘扬社会主义核心价值观。""",
}


@dataclass
class BuildStats:
    source_total: int
    included_sources: int
    converted: int
    ocr: int
    conversion_errors: list[dict]
    suites_total: int
    has_xuanbier: int
    no_xuanbier: int
    subquestions_total: int
    big_questions_total: int
    pending_total: int
    formal_incomplete_total: int


def read_jsonl(path: Path) -> list[dict]:
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def read_csv(path: Path) -> list[dict]:
    with path.open(encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def load_rubric_overrides() -> dict[str, str]:
    script = RUN_DIR / "tools" / "build_image_packet_docx.py"
    spec = importlib.util.spec_from_file_location("image_packet_builder", script)
    if not spec or not spec.loader:
        return {}
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return {
        entry_id: payload["text"]
        for entry_id, payload in getattr(module, "RUBRIC_TEXT_OVERRIDES", {}).items()
        if payload.get("text")
    }


def set_east_asian(run, font: str = FONT_EA) -> None:
    run.font.name = FONT_WEST
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def paragraph_spacing(paragraph, before: int = 0, after: int = 6, line: float = 1.25) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "B8C2CC")


def set_table_fixed_width(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_table_borders(table)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_WEST
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EA)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, HEADING_BLUE, 18, 10),
        ("Heading 2", 13, HEADING_BLUE, 14, 7),
        ("Heading 3", 12, HEADING_DARK, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = FONT_WEST
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EA)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    if "FG Label" not in doc.styles:
        style = doc.styles.add_style("FG Label", 1)
        style.font.name = FONT_WEST
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EA)
        style.font.size = Pt(11)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(HEADING_DARK)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.25


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    paragraph_spacing(p, after=3, line=1.15)
    r = p.add_run("选必二《法律与生活》主观题与细则全量汇编")
    set_east_asian(r)
    r.font.size = Pt(22)
    r.font.bold = True

    p = doc.add_paragraph()
    paragraph_spacing(p, after=12, line=1.15)
    r = p.add_run("纯文字版：只含出处、材料、设问、官方参考答案/评分细则状态")
    set_east_asian(r)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("555555")


def add_paragraph_text(doc: Document, text: str, style: str | None = None, size: float = 11) -> None:
    p = doc.add_paragraph(style=style)
    paragraph_spacing(p)
    for idx, part in enumerate((text or "").split("\n")):
        if idx:
            p.add_run().add_break()
        run = p.add_run(part)
        set_east_asian(run)
        run.font.size = Pt(size)


def add_label_block(doc: Document, label: str, text: str) -> None:
    add_paragraph_text(doc, f"【{label}】", style="FG Label")
    add_paragraph_text(doc, text.strip() if text else "【缺失】")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for idx, head in enumerate(headers):
        hdr[idx].text = head
        set_cell_shading(hdr[idx], TABLE_FILL)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    set_table_fixed_width(table, widths)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph_spacing(paragraph, after=2, line=1.15)
                for run in paragraph.runs:
                    set_east_asian(run)
                    run.font.size = Pt(9)


def q_base(entry: dict) -> str:
    title = entry.get("title", "")
    m = re.search(r"第(\d+)题", title)
    if m:
        return m.group(1)
    m = re.search(r"(\d+)", str(entry.get("question_no", "")))
    return m.group(1) if m else str(entry.get("question_no", ""))


def group_key(entry: dict) -> tuple:
    return (
        str(entry["year"]),
        entry["district_or_exam"],
        entry["paper_type"],
        q_base(entry),
        (entry.get("material") or "").strip(),
    )


def suite_key(year: str, district: str, paper: str) -> str:
    return re.sub(r"\s+", "", f"{year}{district}{paper}".replace("、北京", ""))


def source_status_stats(source_rows: list[dict], coverage_rows: list[dict], packets: list[dict]) -> BuildStats:
    conv = Counter(r["conversion_status"] for r in source_rows)
    status = Counter(r["status"] for r in coverage_rows)
    groups = {group_key(r) for r in packets}
    return BuildStats(
        source_total=len(source_rows),
        included_sources=sum(1 for r in source_rows if r["include_status"] == "include"),
        converted=conv["converted"],
        ocr=conv["ocr"],
        conversion_errors=[r for r in source_rows if r["conversion_status"] == "conversion_error"],
        suites_total=len(coverage_rows),
        has_xuanbier=status["has_xuanbier"],
        no_xuanbier=status["no_xuanbier"],
        subquestions_total=len(packets),
        big_questions_total=len(groups),
        pending_total=sum(1 for r in packets if r.get("pending_reason")),
        formal_incomplete_total=sum(1 for r in packets if r["entry_id"] in FORMAL_INCOMPLETE),
    )


def display_question_list(entries: list[dict]) -> str:
    if not entries:
        return ""
    values = []
    for entry in entries:
        title = entry["title"].split("·")[-1].strip()
        if entry.get("pending_reason"):
            title += "（待确认）"
        if entry["entry_id"] in FORMAL_INCOMPLETE:
            title += "（细则待补）"
        values.append(title)
    return "；".join(values)


def rubric_text(entry: dict, overrides: dict[str, str]) -> str:
    entry_id = entry["entry_id"]
    base = overrides.get(entry_id) or (entry.get("rubric") or "").strip()
    if entry_id in FORMAL_INCOMPLETE:
        reason = FORMAL_INCOMPLETE[entry_id]
        if base:
            return f"【细则缺失/不完整】{reason}\n\n现有参考答案式文本（不作为完整评分细则）：\n{base}"
        return f"【细则缺失/不完整】{reason}"
    return base


def add_validation_section(doc: Document, stats: BuildStats, coverage_rows: list[dict], packets: list[dict]) -> None:
    doc.add_heading("验收清单", level=1)
    add_table(
        doc,
        ["项目", "结果"],
        [
            ["源文件总数", f"扫描登记 {stats.source_total} 个文件；纳入本轮原卷/细则/讲评候选 {stats.included_sources} 个。"],
            ["转换情况", f"文本转换 {stats.converted} 个；OCR {stats.ocr} 个；转换失败 {len(stats.conversion_errors)} 个。"],
            ["套卷覆盖", f"{stats.suites_total} 套；识别出选必二主观题 {stats.has_xuanbier} 套；0题 {stats.no_xuanbier} 套。"],
            ["题量总计", f"{stats.big_questions_total} 道大题，{stats.subquestions_total} 个分问；模块待确认 {stats.pending_total} 个分问。"],
            ["细则匹配", f"细则齐全 {stats.subquestions_total - stats.formal_incomplete_total} 个分问；正式评分细则待补/不完整 {stats.formal_incomplete_total} 个分问。"],
            ["自查核对", "63 套均在覆盖矩阵中逐套记录；0题套卷、模块待确认、OCR/转换异常和细则缺口均列入下方异常清单。"],
        ],
        [2200, 7160],
    )

    by_suite: dict[str, list[dict]] = defaultdict(list)
    for entry in packets:
        by_suite[suite_key(str(entry["year"]), entry["district_or_exam"], entry["paper_type"])].append(entry)

    rows = []
    for cov in coverage_rows:
        key = re.sub(r"\s+", "", cov["suite_id"].replace("、北京", ""))
        entries = by_suite.get(key, [])
        missing = [e["entry_id"] for e in entries if e["entry_id"] in FORMAL_INCOMPLETE]
        pending = [e["entry_id"] for e in entries if e.get("pending_reason")]
        detail = "齐全" if not missing else "待补：" + "、".join(missing)
        if pending:
            detail += "；模块待确认：" + "、".join(pending)
        rows.append(
            [
                cov["year"],
                cov["district_or_exam"],
                cov["paper_type"],
                cov["status"],
                cov["legal_subjective_count"],
                display_question_list(entries),
                detail,
            ]
        )
    doc.add_heading("逐套卷处理清单", level=2)
    add_table(doc, ["年", "区/卷", "卷型", "状态", "题数", "题号列表", "细则/待确认"], rows, [720, 1700, 800, 1200, 650, 2600, 1690])


def add_exception_section(doc: Document, stats: BuildStats, coverage_rows: list[dict], source_rows: list[dict], packets: list[dict]) -> None:
    doc.add_heading("异常清单", level=1)
    doc.add_heading("解析失败文件", level=2)
    if stats.conversion_errors:
        add_table(
            doc,
            ["source_id", "角色", "格式", "路径"],
            [[r["source_id"], r["file_role"], r["file_ext"], r["path"]] for r in stats.conversion_errors],
            [1700, 1300, 700, 5660],
        )
    else:
        add_paragraph_text(doc, "无。")

    doc.add_heading("提取到 0 题的卷", level=2)
    zero_rows = [r for r in coverage_rows if int(r.get("legal_subjective_count") or 0) == 0]
    add_table(
        doc,
        ["年份", "区/卷", "卷型", "状态"],
        [[r["year"], r["district_or_exam"], r["paper_type"], r["status"]] for r in zero_rows],
        [900, 3500, 1200, 3760],
    )

    doc.add_heading("模块待确认区索引", level=2)
    pending = [r for r in packets if r.get("pending_reason")]
    add_table(
        doc,
        ["编号", "出处", "原因"],
        [[r["entry_id"], r["title"], r["pending_reason"]] for r in pending],
        [900, 3100, 5360],
    )

    doc.add_heading("正式评分细则缺失或不完整", level=2)
    add_table(
        doc,
        ["编号", "出处", "缺口"],
        [[r["entry_id"], r["title"], FORMAL_INCOMPLETE[r["entry_id"]]] for r in packets if r["entry_id"] in FORMAL_INCOMPLETE],
        [900, 3100, 5360],
    )

    doc.add_heading("OCR 来源文件", level=2)
    ocr_rows = [r for r in source_rows if r["conversion_status"] == "ocr" and r["include_status"] == "include"]
    add_table(
        doc,
        ["source_id", "角色", "路径"],
        [[r["source_id"], r["file_role"], r["path"]] for r in ocr_rows],
        [1800, 1400, 6160],
    )


def add_entries(doc: Document, title: str, entries: list[dict], overrides: dict[str, str]) -> None:
    doc.add_heading(title, level=1)
    grouped: dict[tuple, list[dict]] = defaultdict(list)
    for entry in entries:
        grouped[group_key(entry)].append(entry)

    current_year = None
    current_suite = None
    for key, group in sorted(grouped.items(), key=lambda item: min(e["entry_id"] for e in item[1])):
        year, district, paper, qno, _ = key
        if year != current_year:
            current_year = year
            current_suite = None
            doc.add_heading(year, level=1)
        suite = f"{district} · {paper}"
        if suite != current_suite:
            current_suite = suite
            doc.add_heading(suite, level=2)
        doc.add_heading(f"{year} · {district} · {paper} · 第{qno}题", level=3)
        if any(e.get("pending_reason") for e in group):
            add_label_block(doc, "模块待确认", "；".join(e["pending_reason"] for e in group if e.get("pending_reason")))

        first = group[0]
        add_label_block(doc, "出处", f"{year}年 {district} {paper} 第{qno}题")
        add_label_block(doc, "材料", first.get("material", ""))
        add_paragraph_text(doc, "【设问与细则】", style="FG Label")
        for idx, entry in enumerate(group, start=1):
            prefix = f"({idx}) " if len(group) > 1 else ""
            add_paragraph_text(doc, f"{prefix}设问：{entry.get('prompt', '').strip()}")
            add_paragraph_text(doc, f"{prefix}细则：{rubric_text(entry, overrides)}")


def write_check(stats: BuildStats, packets: list[dict]) -> None:
    rows = [
        "# FULL_TEXT_DELIVERY_CHECK",
        "",
        f"- Output: `{OUT_DOCX}`",
        f"- Desktop copy: `{DESKTOP_COPY}`",
        f"- Source files scanned: {stats.source_total}",
        f"- Included source/rubric candidates: {stats.included_sources}",
        f"- Coverage suites: {stats.suites_total}",
        f"- Big questions: {stats.big_questions_total}",
        f"- Subquestions: {stats.subquestions_total}",
        f"- Pending module-confirmation subquestions: {stats.pending_total}",
        f"- Formal rubric incomplete/missing subquestions: {stats.formal_incomplete_total}",
        "",
        "## Formal Rubric Incomplete",
    ]
    for entry in packets:
        if entry["entry_id"] in FORMAL_INCOMPLETE:
            rows.append(f"- {entry['entry_id']} {entry['title']}: {FORMAL_INCOMPLETE[entry['entry_id']]}")
    CHECK_MD.write_text("\n".join(rows) + "\n", encoding="utf-8")


def build_doc() -> None:
    packets = read_jsonl(PACKETS)
    coverage_rows = read_csv(COVERAGE)
    source_rows = read_csv(SOURCE_LEDGER)
    overrides = load_rubric_overrides()
    overrides.update(EXTRA_RUBRIC_OVERRIDES)
    stats = source_status_stats(source_rows, coverage_rows, packets)

    doc = Document()
    style_document(doc)
    add_title(doc)
    add_validation_section(doc, stats, coverage_rows, packets)
    add_exception_section(doc, stats, coverage_rows, source_rows, packets)

    confirmed = [r for r in packets if not r.get("pending_reason")]
    pending = [r for r in packets if r.get("pending_reason")]
    add_entries(doc, "正文：已确认选必二主观题", confirmed, overrides)
    add_entries(doc, "待确认区：模块边界需人工复核", pending, overrides)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    shutil.copy2(OUT_DOCX, DESKTOP_COPY)
    write_check(stats, packets)
    print(OUT_DOCX)
    print(CHECK_MD)


if __name__ == "__main__":
    build_doc()
