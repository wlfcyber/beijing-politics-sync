#!/usr/bin/env python3
from __future__ import annotations

import re
import shutil
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import build_student_baodian_docx as old


RUN_DIR = Path(__file__).resolve().parents[1]
OUT_DIR = RUN_DIR / "05_output"
OUT_DOCX = OUT_DIR / "选必二法律与生活_主观题答题宝典_学生重排版_20260604.docx"
DESKTOP_COPY = Path.home() / "Desktop" / OUT_DOCX.name
QA_MD = OUT_DIR / "STUDENT_BAODIAN_POLISHED_LAYOUT_CHECK_20260604.md"

FONT_BODY = "宋体"
FONT_HEAD = "微软雅黑"
FONT_WEST = "Calibri"
BLUE = "17365D"
RED = "C00000"
GOLD = "7F6000"
GRAY = "666666"
LIGHT_BLUE = "EAF2F8"
LIGHT_GOLD = "FFF2CC"
LIGHT_GRAY = "F5F5F5"
BORDER = "B7C9D6"


A_AXIS = [
    (
        "A1 民事行为与时效",
        "未成年人、意思表示、代理、追认、诉讼时效、不适用时效。",
        "先判断行为是否有效、权利还能不能主张，再写保护交易安全或特殊身份利益。",
        "不要把时效写成权利消灭；通常是胜诉权受限。",
    ),
    (
        "A2 人格权与个人信息",
        "名誉、隐私、肖像、姓名、生命健康、个人信息、平台数据。",
        "先写人格利益受法律保护，再用材料事实说明侵害方式、损害和责任。",
        "有网络、消费、平台背景，也不能跳过具体人格权。",
    ),
    (
        "A3 物权、相邻与共有",
        "共有部分、相邻通行、采光、加装电梯、业主共同决定、物业服务。",
        "先找不动产权利边界，再落到有利生产、方便生活、团结互助、公平合理。",
        "审批通过不等于相邻权益必然让位。",
    ),
    (
        "A4 合同与诚信",
        "要约、承诺、履行、违约、格式条款、补充协议、意思表示不明。",
        "先判合同是否成立和内容是否明确，再写全面履行、诚信协商、违约责任。",
        "不要只写诚信口号，要把约定、回复、履行事实压进去。",
    ),
    (
        "A5 知识产权与公平竞争",
        "著作权、商标、专利、商业秘密、虚假宣传、混淆、恶意诉讼。",
        "先定具体权利或竞争秩序，再写创新保护、市场公平和侵权成本。",
        "创新价值必须附着在具体权利和具体行为上。",
    ),
    (
        "A6 侵权责任",
        "损害、过错、因果关系、安全保障义务、污染环境、赔偿。",
        "按要件写：权利受损、违法或过错、因果关系、责任承担。",
        "不能只写有人受损就担责；要看过错、义务和因果。",
    ),
    (
        "A7 婚姻家庭与继承",
        "夫妻财产、赡养、抚养、继承、遗赠扶养协议、家庭义务。",
        "先写法定义务或协议效力，再写保护弱者、弘扬家庭美德和公序良俗。",
        "亲情价值必须落到赡养、扶养、继承等具体义务。",
    ),
    (
        "A8 劳动关系",
        "劳动合同、事实劳动关系、平台用工、解除、竞业限制、社保。",
        "先判是否存在劳动关系或劳动法保护对象，再平衡劳动者权益与有序用工。",
        "劳动法不等于只保护劳动者，也维护企业依法管理。",
    ),
    (
        "A9 消费者权益",
        "经营者、消费者、欺诈、知情权、公平交易、惩罚性赔偿、食品安全。",
        "先确认消费者和经营者关系，再写保护合法权益、规范经营、维护市场秩序。",
        "职业打假类题要看法院支持的是哪一笔、哪一种请求。",
    ),
    (
        "A10 纠纷解决程序",
        "调解、仲裁、诉讼、证据、举证、起诉状、公益诉讼、执行监督。",
        "先写适用的程序和证据要求，再写依法、公正、高效、多元化解纠纷。",
        "只有题干问路径或程序时，程序才是主入口。",
    ),
]


B_AXIS = [
    ("B1 补表/补链", "完成下表、补全示例、填写理由", "事实 -> 规则 -> 结论，短句作答。"),
    ("B2 判责/理由", "为什么判、谁承担责任、裁判依据", "法律关系 -> 规则 -> 材料事实 -> 责任。"),
    ("B3 诉求支持", "诉求能否支持、请求是否成立", "先结论，再逐项匹配权利基础和事实。"),
    ("B4 评析认识", "评析、认识、谈谈理解", "肯定合理处 -> 指出片面处 -> 给出法律立场。"),
    ("B5 意义价值", "意义、价值、原因、作用、启示", "保护谁 -> 约束谁 -> 形成什么秩序。"),
    ("B6 维权路径", "如何维权、准备什么证据、怎么办", "证据 -> 选择程序 -> 提出请求。"),
    ("B7 短答识别", "指出问题、改正、填空、概念判断", "只写关键词和短句，不展开长判决。"),
]


MODEL_BANK = [
    (
        "合同与诚信",
        "合同、报价、中标、仓储、补充协议、履行、解除、尾款。",
        "合同成立先找要约与承诺；合同履行先找约定是否明确、双方是否全面履行。",
        "民事主体应当遵循诚信原则，按照约定全面履行义务；约定不明时，应平等协商补充或解释。",
        "不要把“乙方回复可以”直接写成一切费用都转移，要看双方有没有明确磋商内容。",
    ),
    (
        "侵权责任与安全保障",
        "受伤、死亡、污染、损害、过错、安全提醒、救助、因果关系。",
        "先按侵权四要素拆，再判断是否存在免责、减责或无因果关系。",
        "承担侵权责任通常要有受保护权益受损、违法或过错行为、损害结果以及因果关系。",
        "义务组织者已经合理提醒并及时救助时，不能只凭结果倒推责任。",
    ),
    (
        "人格权与网络平台",
        "个人信息、隐私、名誉、肖像、虚拟数字人、账号、公开传播。",
        "先锁定哪一项人格利益，再写未经同意、超范围使用或造成社会评价降低。",
        "人格权受法律保护，侵权人应停止侵害、消除影响、赔礼道歉并承担赔偿等责任。",
        "网络材料里同时可能有知识产权、竞争秩序和人格权，要按题干重点取舍。",
    ),
    (
        "知识产权与公平竞争",
        "商标、著作权、商业秘密、技术秘密、仿冒、虚假宣传、惩罚性赔偿。",
        "先写权利类型和侵权行为，再写保护创新、提高侵权成本、维护市场公平。",
        "保护知识产权就是保护创新；对恶意侵权依法提高成本，有利于维护公平竞争秩序。",
        "不要只写“创新很重要”，必须说清被保护的是哪项权利、被侵害的行为是什么。",
    ),
    (
        "消费者权益",
        "经营者、消费者、食品安全、欺诈、知情权、退一赔三、惩罚性赔偿。",
        "先确认交易身份和违法事实，再写法院支持或限制请求的理由。",
        "经营者应依法诚信经营，保障消费者知情权、公平交易权和安全权。",
        "连续购买、明知再买的材料，要特别看法院支持的是首单还是全部请求。",
    ),
    (
        "劳动关系与竞业限制",
        "平台派单、人格从属性、经济从属性、组织从属性、竞业限制、解除合同。",
        "事实劳动关系看从属性；竞业限制看人员范围、期限、补偿和商业秘密保护。",
        "依法规范劳动关系，有利于保障劳动者合法权益，也有利于维护用人单位依法用工秩序。",
        "竞业限制不是所有员工都能套，普通岗位一般不能被过度限制。",
    ),
    (
        "婚姻家庭与继承",
        "赡养费、老人、遗赠扶养协议、继承、夫妻共同财产、抚养义务。",
        "先写法定义务或协议效力，再把家庭伦理转化为法律责任和价值引导。",
        "成年子女对父母负有赡养义务，赡养包括经济供养、生活照料和精神慰藉。",
        "不能只写尊老爱老，必须落到谁对谁承担什么义务。",
    ),
    (
        "物权相邻与基层治理",
        "加装电梯、共有部分、业主表决、通行采光、物业、社区治理。",
        "先写权利边界和法定程序，再写相邻关系处理原则和基层治理价值。",
        "不动产权利人应按照有利生产、方便生活、团结互助、公平合理原则处理相邻关系。",
        "不能把多数人同意写成少数人权益自动消失。",
    ),
    (
        "纠纷解决与司法价值",
        "起诉状、证据、回避、调解、仲裁、公益诉讼、执行、诉讼代理人。",
        "先判断程序角色和制度，再写降低成本、实质化解、维护公平正义。",
        "人民法院以事实为依据、以法律为准绳，通过公正司法维护合法权益和社会秩序。",
        "程序题少写抽象法治口号，多写具体制度名称和本案怎么用。",
    ),
    (
        "绿色原则与生态保护",
        "绿色原则、污染环境、生态修复、碳排放、资源节约、环境公益诉讼。",
        "先写民法典绿色原则或侵权责任，再写权利义务配置和生态环境公共利益。",
        "民事活动应有利于节约资源、保护生态环境；破坏生态造成损害的，应依法承担责任。",
        "不要把生态题写成政治经济口号，法律落点是民事权利义务和责任。",
    ),
]


def set_run(run, *, ea: str = FONT_BODY, west: str = FONT_WEST, size: float | None = None,
            bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = west
    run._element.rPr.rFonts.set(qn("w:eastAsia"), ea)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def spacing(p, before: float = 0, after: float = 3, line: float = 1.12) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def shade_paragraph(p, fill: str, border: str | None = None) -> None:
    ppr = p._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    if border:
        pbdr = ppr.find(qn("w:pBdr"))
        if pbdr is None:
            pbdr = OxmlElement("w:pBdr")
            ppr.append(pbdr)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:color"), border)
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "4")
        pbdr.append(left)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_WEST
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.font.size = Pt(10.2)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(3)

    for name, size, color, before, after in [
        ("Heading 1", 15.5, BLUE, 16, 7),
        ("Heading 2", 13, BLUE, 11, 5),
        ("Heading 3", 11.2, "222222", 6, 2),
    ]:
        style = doc.styles[name]
        style.font.name = FONT_WEST
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEAD)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "FG Card" not in doc.styles:
        style = doc.styles.add_style("FG Card", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT_WEST
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style.font.size = Pt(10)
        style.paragraph_format.left_indent = Inches(0.12)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = header.add_run("飞哥正志讲堂｜选必二《法律与生活》主观题答题宝典")
    set_run(r, ea=FONT_HEAD, size=8.5, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("学生重排版 20260604")
    set_run(r, ea=FONT_HEAD, size=8.5, color=GRAY)
    r = footer.add_run("  第 ")
    set_run(r, ea=FONT_HEAD, size=8.5, color=GRAY)
    add_page_field(footer)
    r = footer.add_run(" 页")
    set_run(r, ea=FONT_HEAD, size=8.5, color=GRAY)


def add_page_field(p) -> None:
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)
    set_run(run, ea=FONT_HEAD, size=8.5, color=GRAY)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        set_run(r, ea=FONT_HEAD)
    return p


def add_line(doc: Document, text: str = "", *, size: float = 10.2, color: str | None = None,
             bold: bool = False, align=None, before: float = 0, after: float = 3):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    spacing(p, before=before, after=after)
    if text:
        r = p.add_run(text)
        set_run(r, size=size, bold=bold, color=color)
    return p


def add_box(doc: Document, title: str, body: str, *, fill: str = LIGHT_BLUE, color: str = BLUE) -> None:
    p = doc.add_paragraph(style="FG Card")
    shade_paragraph(p, fill, BORDER)
    spacing(p, after=5)
    r = p.add_run(f"【{title}】")
    set_run(r, ea=FONT_HEAD, size=10.4, bold=True, color=color)
    r = p.add_run(body)
    set_run(r, size=10.2, color="222222")


def add_label(doc: Document, label: str, body: str, *, red: bool = False, fill: str | None = None) -> None:
    p = doc.add_paragraph(style="FG Card")
    if fill:
        shade_paragraph(p, fill)
    spacing(p, after=2)
    r = p.add_run(f"【{label}】")
    set_run(r, ea=FONT_HEAD, size=10, bold=True, color=BLUE if not red else RED)
    r = p.add_run(body)
    set_run(r, size=10, color=RED if red else "222222", bold=red)


def clean_text(text: str) -> str:
    text = old.clean_common_student_text(text or "")
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"高三(年级)?(思想政治|政治).*?共\d+页", "", text)
    return text.strip()


def truncate(text: str, limit: int) -> str:
    text = clean_text(text)
    if len(text) <= limit:
        return text
    return text[:limit].rstrip(" ，。；、") + "..."


def score_atoms(text: str, max_items: int = 2) -> list[str]:
    text = clean_text(text)
    if "具体：" in text:
        text = text.split("具体：", 1)[1]
    if "具体:" in text:
        text = text.split("具体:", 1)[1]
    for marker in ["参考答案：", "【答案】", "答案："]:
        if marker in text:
            text = text.split(marker, 1)[1]
    text = re.sub(r"^\d+[．.、]?\s*（?\d+分）?", "", text)
    raw = re.split(r"(?<=[。！？；])\s*|\n+", text)
    atoms: list[str] = []
    for item in raw:
        item = clean_text(item).strip("；。 ")
        if not item:
            continue
        if any(bad in item for bad in ["阅卷细则", "题号与设问", "参考答案", "示例"]):
            continue
        if any(
            bad in item
            for bad in [
                "结合材料",
                "运用《法律与生活》知识",
                "运用法律相关知识",
                "谈谈",
                "将以上回答补充完整",
                "分析其",
                "阐述该",
                "说明人民法院",
                "说明法院",
                "如何依法",
            ]
        ):
            continue
        if re.search(r"^\d+[．.]?（?\d+分）?\s*[（(]?\d*[）)]?\s*结合", item):
            continue
        if len(item) < 8:
            continue
        atoms.append(item)
        if len(atoms) >= max_items:
            break
    return atoms or ["给分点待回源核对，见全量汇编对应题卡。"]


def card_ref(group: list[dict]) -> str:
    first = group[0]
    return f"{first['year']}{first['district_or_exam']}{first['paper_type']}第{old.BASE.q_base(first)}题"


def group_info(group: list[dict], overrides: dict[str, str]) -> dict:
    first = group[0]
    material = old.display_material(first)
    combined_prompt = " ".join(old.display_prompt(e) for e in group)
    rubric_for_infer = " ".join(old.clean_student_rubric(e, overrides) for e in group)
    a_axis, b_axis = old.group_axis(group, material, combined_prompt, rubric_for_infer)
    return {
        "year": str(first["year"]),
        "district": first["district_or_exam"],
        "paper": first["paper_type"],
        "qno": old.BASE.q_base(first),
        "ref": card_ref(group),
        "material": material,
        "material_hook": truncate(old.first_sentence(material, 95), 95),
        "a_axis": a_axis,
        "b_axis": b_axis,
        "entries": group,
    }


def group_examples(infos: list[dict], axis_prefix: str, limit: int = 3) -> str:
    hits = [info["ref"] for info in infos if info["a_axis"].startswith(axis_prefix)]
    return "；".join(hits[:limit]) if hits else "见真题索引"


def add_cover(doc: Document, stats, infos: list[dict]) -> None:
    p = add_line(doc, "选择性必修二", size=15, color=BLUE, bold=True, after=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = add_line(doc, "《法律与生活》主观题答题宝典", size=25, bold=True, after=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = add_line(doc, "飞哥正志讲堂", size=12.5, color=GRAY, bold=True, after=16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_box(
        doc,
        "本版定位",
        "这不是题源汇编。学生先用 A 线找法律入口，用 B 线找题干动作，再把材料事实写进给分句。完整材料和细则原文保留在全量汇编与附录线索中。",
        fill=LIGHT_GOLD,
        color=GOLD,
    )
    add_label(doc, "总公式", "生活事实 -> 争点 -> 法律翻译 -> 法律结果 -> 必要价值。", red=True)
    add_label(doc, "入册范围", f"正文覆盖 {len(infos)} 组真题题卡，来自 {stats.big_questions_total} 道大题、{stats.subquestions_total} 个分问；边界题和期中无题套卷只在附录说明。")

    add_heading(doc, "目录", 1)
    for item in [
        "一、先背住：法律主观题总路线",
        "二、A 线：看到材料，先归入哪个法律入口",
        "三、B 线：看到设问，先确定答案长相",
        "四、高频模型：十类题的考场句式",
        "五、真题索引：按年份回练，按细则改写",
        "附录：覆盖范围与暂不入正文题卡",
    ]:
        add_line(doc, item, size=10.5, color=BLUE if item[0] in "一二三四五" else GRAY, bold=item[0] in "一二三四五", after=1)
    doc.add_page_break()


def add_route(doc: Document) -> None:
    add_heading(doc, "一、先背住：法律主观题总路线", 1)
    add_box(doc, "考场顺序", "先圈主体，再找争点；先落法律关系，再写规则；先贴材料事实，再补价值。", fill=LIGHT_BLUE)
    for title, body in [
        ("主体", "谁和谁发生法律关系？自然人、法人、经营者、消费者、劳动者、用人单位、业主、法院、仲裁机构各自是什么角色。"),
        ("争点", "争的是合同履行、侵权赔偿、人格利益、劳动关系、消费请求、继承赡养，还是程序路径。"),
        ("规则", "写规则不要悬空。规则句后面必须接材料事实，让阅卷人看到你会把生活事实翻译成法律语言。"),
        ("结果", "结果包括责任、效力、请求是否支持、程序如何走，以及由此保护的权利和秩序。"),
    ]:
        add_label(doc, title, body)
    add_label(doc, "一句话提醒", "法律题最怕两种空：一是只背法条不贴材料，二是只喊法治价值不写具体权利义务。", red=True)
    doc.add_page_break()


def add_axis_sections(doc: Document, infos: list[dict]) -> None:
    add_heading(doc, "二、A 线：看到材料，先归入哪个法律入口", 1)
    add_box(doc, "用法", "材料一出现人、财产、合同、劳动、消费、程序等信号，先在 A 线定入口。入口定错，后面规则和价值都会散。", fill=LIGHT_BLUE)
    for idx, (name, trigger, first_judge, trap) in enumerate(A_AXIS, 1):
        add_heading(doc, f"{idx}. {name}", 2)
        prefix = name.split()[0]
        add_label(doc, "材料信号", trigger)
        add_label(doc, "先判什么", first_judge, red=True)
        add_label(doc, "易错点", trap)
        add_label(doc, "代表题源", group_examples(infos, prefix))

    doc.add_page_break()
    add_heading(doc, "三、B 线：看到设问，先确定答案长相", 1)
    add_box(doc, "用法", "同一法律入口，设问动词不同，答案长相也不同。B 线决定你写判责、评析、意义、路径，还是短答补链。", fill=LIGHT_GOLD, color=GOLD)
    for idx, (name, words, shape) in enumerate(B_AXIS, 1):
        add_heading(doc, f"{idx}. {name}", 2)
        add_label(doc, "题干词", words)
        add_label(doc, "答案形状", shape, red=True)


def add_models(doc: Document, infos: list[dict]) -> None:
    doc.add_page_break()
    add_heading(doc, "四、高频模型：十类题的考场句式", 1)
    add_box(doc, "训练方法", "每个模型先背核心句，再用真题索引里的材料抓手改写。不要整段照抄，考场要把材料事实塞进规则句。", fill=LIGHT_BLUE)
    for idx, (title, trigger, skeleton, core, trap) in enumerate(MODEL_BANK, 1):
        add_heading(doc, f"模型{idx:02d}  {title}", 2)
        add_label(doc, "触发词", trigger)
        add_label(doc, "答题骨架", skeleton)
        add_label(doc, "核心句", core, red=True, fill="FFF7F7")
        add_label(doc, "别这样写", trap)


def add_question_index(doc: Document, infos: list[dict], overrides: dict[str, str]) -> None:
    doc.add_page_break()
    add_heading(doc, "五、真题索引：按年份回练，按细则改写", 1)
    add_box(
        doc,
        "阅读规则",
        "这里不再铺满原材料。每题只保留入口、材料抓手、设问和核心给分句。需要完整题面和细则时，回到全量汇编查原文。",
        fill=LIGHT_GOLD,
        color=GOLD,
    )
    current_year = None
    for idx, info in enumerate(infos, 1):
        if info["year"] != current_year:
            current_year = info["year"]
            add_heading(doc, current_year, 2)
        add_heading(doc, f"{idx:02d}. {info['year']} · {info['district']} · {info['paper']} · 第{info['qno']}题", 3)
        add_label(doc, "入口", f"{info['a_axis']} ｜ {info['b_axis']}", red=True)
        add_label(doc, "材料抓手", info["material_hook"])
        for sub_idx, entry in enumerate(info["entries"], 1):
            prompt = truncate(old.display_prompt(entry), 78)
            rubric = old.clean_student_rubric(entry, overrides)
            atoms = "；".join(score_atoms(rubric, 1))
            label = f"第{sub_idx}问" if len(info["entries"]) > 1 else "设问"
            add_label(doc, label, prompt)
            add_label(doc, "给分句", truncate(atoms, 115), red=True, fill="FFF7F7")


def add_appendix(doc: Document, stats, coverage_rows: list[dict], packets: list[dict]) -> None:
    doc.add_page_break()
    add_heading(doc, "附录：覆盖范围与暂不入正文题卡", 1)
    add_box(
        doc,
        "边界",
        "本附录只说明覆盖口径。退回稿中的完整材料/细则堆叠不再进入学生主阅读线，避免把宝典做成证据库。",
        fill=LIGHT_GRAY,
        color=GRAY,
    )
    for label, value in [
        ("源文件", f"扫描登记 {stats.source_total} 个；纳入候选 {stats.included_sources} 个。"),
        ("套卷", f"{stats.suites_total} 套；has_xuanbier {stats.has_xuanbier} 套，no_xuanbier {stats.no_xuanbier} 套。"),
        ("题量", f"{stats.big_questions_total} 道大题，{stats.subquestions_total} 个分问。"),
        ("边界备注", f"综合法治或材料待补 {stats.pending_total} 个；正式细则待补/不完整 {stats.formal_incomplete_total} 个。"),
    ]:
        add_label(doc, label, value)

    excluded = [r for r in packets if r["entry_id"] in old.BODY_EXCLUDE_ENTRY_IDS]
    if excluded:
        add_heading(doc, "正文暂不展示题卡", 2)
        for r in excluded:
            add_label(doc, r["title"], old.BODY_EXCLUDE_ENTRY_IDS[r["entry_id"]])

    zero_rows = [r for r in coverage_rows if int(r.get("legal_subjective_count") or 0) == 0]
    if zero_rows:
        add_heading(doc, "暂列无题套卷", 2)
        for r in zero_rows:
            add_label(doc, f"{r['year']} {r['district_or_exam']} {r['paper_type']}", r["status"])


def write_qa(stats, infos: list[dict], packets: list[dict]) -> None:
    lines = [
        "# STUDENT_BAODIAN_POLISHED_LAYOUT_CHECK_20260604",
        "",
        f"- Output: `{OUT_DOCX}`",
        f"- Desktop copy: `{DESKTOP_COPY}`",
        "- Supersedes: `选必二法律与生活_主观题细则宝典_学生版_20260604.docx` for student delivery.",
        "- Design: no main-body tables; paragraph handbook style modeled on 选必一/必修四 final handbooks.",
        "- Main reading line: 总路线 -> A线 -> B线 -> 高频模型 -> 真题索引.",
        "- Full material/rubric dump removed from main body; complete evidence remains in full compilation and appendix pointers.",
        f"- Question groups indexed: {len(infos)}",
        f"- Source stats: {stats.big_questions_total} big questions / {stats.subquestions_total} subquestions.",
        f"- Raw packet count: {len(packets)}",
        "- Claude review route: must use Claude desktop app cowork, not web Claude.",
        "",
    ]
    QA_MD.write_text("\n".join(lines), encoding="utf-8")


def build_doc() -> None:
    packets = old.BASE.read_jsonl(old.BASE.PACKETS)
    coverage_rows = old.BASE.read_csv(old.BASE.COVERAGE)
    source_rows = old.BASE.read_csv(old.BASE.SOURCE_LEDGER)
    overrides = old.BASE.load_rubric_overrides()
    overrides.update(old.BASE.EXTRA_RUBRIC_OVERRIDES)
    stats = old.BASE.source_status_stats(source_rows, coverage_rows, packets)
    groups = old.make_body_groups(packets)
    infos = [group_info(group, overrides) for group in groups]

    doc = Document()
    style_document(doc)
    add_cover(doc, stats, infos)
    add_route(doc)
    add_axis_sections(doc, infos)
    add_models(doc, infos)
    add_question_index(doc, infos, overrides)
    add_appendix(doc, stats, coverage_rows, packets)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    shutil.copy2(OUT_DOCX, DESKTOP_COPY)
    write_qa(stats, infos, packets)
    print(OUT_DOCX)
    print(DESKTOP_COPY)
    print(QA_MD)


if __name__ == "__main__":
    build_doc()
