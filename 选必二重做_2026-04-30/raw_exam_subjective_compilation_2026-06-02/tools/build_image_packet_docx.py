#!/usr/bin/env python3
from __future__ import annotations

import csv
import json
import re
import shutil
import textwrap
import zipfile
from collections import defaultdict
from pathlib import Path

import fitz
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


RUN_DIR = Path(__file__).resolve().parents[1]
PACKETS = RUN_DIR / "03_source_packets" / "source_packets_final.jsonl"
SOURCE_LEDGER = RUN_DIR / "00_control" / "SOURCE_LEDGER.csv"
OUT_DIR = RUN_DIR / "05_output"
ASSET_DIR = OUT_DIR / "image_packet_assets"
OUT_DOCX = OUT_DIR / "选必二法律与生活_习题汇编_2024-2026_图片版.docx"
REPORT_CSV = ASSET_DIR / "image_packet_report.csv"

FONT_CANDIDATES = [
    "/System/Library/Fonts/PingFang.ttc",
    "/System/Library/Fonts/STHeiti Light.ttc",
    "/System/Library/Fonts/Supplemental/Songti.ttc",
    "/Library/Fonts/Arial Unicode.ttf",
]


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    for path in FONT_CANDIDATES:
        if Path(path).exists():
            try:
                return ImageFont.truetype(path, size=size, index=0)
            except Exception:
                continue
    return ImageFont.load_default()


FONT_BODY = load_font(34)
FONT_TITLE = load_font(42, bold=True)
FONT_SMALL = load_font(24)

# Fractional crop boxes are relative to the rendered PDF page image:
# (left, top, right, bottom). They are used only for scanned/no-coordinate
# rubric PDFs after visual verification against the original rubric page.
MANUAL_RUBRIC_CROPS: dict[str, list[dict]] = {
    "E001": [{"page": 1, "box": (0.51, 0.30, 0.98, 0.415)}],
    "E002": [{"page": 1, "box": (0.50, 0.235, 0.985, 0.455)}],
    "E003": [{"page": 1, "box": (0.50, 0.455, 0.985, 0.498)}],
    "E041": [
        {"page": 7, "box": (0.00, 0.07, 0.98, 0.87)},
        {"page": 8, "box": (0.00, 0.08, 0.98, 0.90)},
    ],
    "E047": [{"page": 2, "box": (0.03, 0.02, 0.95, 0.45)}],
    "E051": [
        {"page": 54, "box": (0.00, 0.12, 1.00, 0.68)},
        {"page": 55, "box": (0.00, 0.09, 1.00, 0.95)},
    ],
    "E055": [{"page": 6, "box": (0.00, 0.04, 1.00, 0.92)}],
    "E056": [{"page": 7, "box": (0.00, 0.07, 1.00, 0.93)}],
    "E057": [{"page": 9, "box": (0.00, 0.04, 1.00, 0.95)}],
}

RUBRIC_TEXT_OVERRIDES: dict[str, dict[str, str]] = {
    "E001": {
        "source_id": "SRC_1f17f2c91d_细则",
        "note": "formal scoring rubric from raw PPT; reference-answer PDF removed",
        "text": """
法律规定诉讼时效有助于督促当事人积极主张权利/行使权利/积极维权/强化当事人依法维权/提高法律意识/积极解决纠纷/节约司法资源（2分）
法律对赡养不适用诉讼时效的规定因为：
成年人对父母有赡养的义务（2分）。
这样做有利于保护老年人的合法权益/保护老年人的基本生存权利/构建和谐家庭（弘扬家庭美德）/彰显公序良俗（社会主义和谐价值观、社会秩序）（2分）。
""".strip(),
    },
    "E002": {
        "source_id": "SRC_9080f96f97_19_1",
        "note": "formal per-question marking summary from raw DOCX; reference-answer PDF removed",
        "text": """
高三政治二模阅卷总结             2024.5.09
题号与设问：19（1）请你选择其中一方进行代理，并撰写起诉状（7分）
阅卷细则
原告被告，不给分；
示例1:原告:甲公司;被告:乙公司
诉讼请求：判令被告承担全部仓储费用/继续履行，被告承担本案诉讼费用（2分）
事实与理由:
事实: 我公司于2024年3月10日自乙公司购买一批电视,因收货不便,该批货物由快递代为储存，且约定3月20日及以前的仓储费用由卖方承担。3月19日,我公司与被告协商再延长10天,获得肯定答复。3月30日,我公司收货并垫付全部仓储费用,但乙公司拒绝承担3月21日及以后产生的仓储费用。（踩意2分）
理由:①我公司提出延长仓储时间，乙公司回复“好的”，际应按原合同内容全面履行，乙公司的行为违背诚信原则。
②乙公司拒绝承担3月21日以后产生的仓储费用，并未与我公司协商一致，属于单方变更合同行为。
其行为构成违约，应当承担违约责任
（1个理由2分，2个理由3分）

示例2:原告:乙公司;被告:甲公司
诉讼请求：判令被告支付尾款/继续履行，被告承担本案诉讼费用。（2分）
事实与理由:
事实：甲公司于2024年3月10日自我公司购买一批电视,因收货不便,该批货物由快递代为储存，且约定3月20日及以前的仓储费用由我公司承担。3月19日,被告与我司协商再延长10天,我公司同意。3月30日,甲公司收货后，我公司不同意支付3月21日及以后产生的仓储费用，甲公司拒不支付合同尾款。 （踩意2分）
理由: ① 3月19日提出新合同，合同只对延期进行约定，并未协商延长仓储时间的费用承担问题，但由于费用产生由甲公司原因导致，该笔费用理应由甲公司承担，若逾期产生的费用仍由我方承担，有违公平原则。（只谈要约承诺构成新合同给1分 ，有公平原则给2分）
②买卖合同有效，甲公司不支付尾款违反全面履行原则。
其行为构成违约，应承担违约责任。
（1个理由2分，2个理由3分）
""".strip(),
    },
    "E003": {
        "source_id": "SRC_c492176058_19_2",
        "note": "formal per-question marking summary from raw DOCX; reference-answer PDF removed",
        "text": """
高三政治二模阅卷总结             2024.5.09
题号与设问：19（2）为避免类似问题产生，我们应该        。（2分）
阅卷细则
指出“合同内容应清晰明确”即可得2分。相关替代有：明确双方的权利与义务/对合同内容
要完善、细化、具体等/意思表示要明确/避免模糊不清的意思表示等。
只提出要协商一致，得1分。
以下内容不给分：
意思表示一致；诚信履行；全面履行
""".strip(),
    },
    "E006": {
        "source_id": "SRC_24a14c0c23_细则",
        "note": "formal scoring table from raw DOCX appended to rubric image",
        "text": """
19（8分）民事主体的著作权受法律保护，甲公司创作虚拟数字人Ada，对其享有著作权。乙公司未经许可使用甲公司创作的相关视频，并在展示过程中故意替换有关标识、加注注册商标，构成著作权侵权。乙公司进行虚假和引人误解的商业宣传，属于不正当竞争。法院依法判决乙公司赔偿甲公司损失，有利于激发创新活力，维护市场秩序。
【细则】
 | 知识 1分 | 材料 1分 | 备注
层次1 | 著作权   例：甲公司享有著作权 | 甲公司创作虚拟数字人Ada/甲公司制作并发布介绍虚拟数字人和真人与虚拟数字人动作捕捉的视频。 | 1.未准确答出“著作权”，写成“知识产权”“作品”“著作”“视频内容”不得分，后边的材料也不得分。   2.专利权不得分。  3.写了“甲公司创作虚拟数字人Ada正面形象”的材料但没有著作权不得分。  4.“著作权”“Ada”写错别字不得分。   5.没有写甲公司，可以和材料结合得2分，例：对虚拟数字人Ada享有著作权。
层次2 | 乙公司构成侵权/侵犯了甲公司的著作权 | 乙公司未经许可使用甲公司创作的相关视频，并在展示过程中故意替换有关标识、加注注册商标 | 1.写出“侵权” 并加材料，得2分。  2.没写“侵权”，只写“甲公司创作的作品/视频受著作权保护，乙公司未经许可使用……”，得2分。  3.写“混淆行为”“侵权行为”“侵犯了知识产权”并加材料，得2分。  4.写“搭便车”并加材料，得2分。  5.写“侵犯商标权”不得分。
层次3 | 不正当竞争/进行虚假或引人误解的商业宣传 | 乙公司擅自使用/替换视频中的标识、添加注册商标。 | 1.如果与上述侵权结合，需要有新的材料支撑，没有新的材料支撑不重复给材料分。  2.只写“乙公司未依法诚信经营”，有材料，也得2分。
层次4 | 1.激发创新/保护创作/尊重创意。  2.市场秩序/公平竞争/合法竞争/营造良好市场环境  3.保护创作者合法权益 | 乙公司消除影响、赔偿经济损失 | 1.答出任意一点知识，加材料，得2分。  2.只答判决本身的意义不得分。  3.只写赔偿损失的意义，不写材料，给1分。  4.层次4和层次2结合时，写“保护创作者合法权益”，算一个角度。
水平4  | 5-6分 | 紧扣判决内容展开分析，观点明确，知识运用准确、全面，论述逻辑严密，条理清晰。
水平3 | 3-4分 | 能扣判决内容展开分析，观点比较明确，知识运用比较准确、全面，论述符合逻辑，有条理。
水平2 | 1-2分 | 能扣判决内容展开分析，观点不明确，知识运用不准确、不全面，论述缺乏逻辑，条理性差。
水平1 | 0分 | 应答与试题无关，或重复试题内容，或没有应答。
""".strip(),
    },
    "E018": {
        "source_id": "SRC_066dbcf5b7_细则",
        "note": "formal scoring standard from raw DOCX; reference answer removed",
        "text": """
19题 评分标准说明：
企业层面：本案判决维护当事人合法权益，保护原告方（当事人）的知识产权，1分；严惩或者严厉打击侵权行为 1分
对科技创新的意义：规范市场竞争秩序、营造公平的市场环境、鼓励创新，2分
对司法领域的价值：对同类案件提供范例、借鉴、参考，提高效率，增强公信力2分。
变通的地方：彰显法律权威或者司法权威作为1分的替换答案
对法治国家、法治社会而言，法治信仰（意识方面面），依法层面（行动方面）1点1分
出现的问题：维护社会公平正义不得分；没有铺垫，直接表达法治中国建设的不得分。
""".strip(),
    },
    "E033": {
        "source_id": "SRC_700c0fe3eb_细则",
        "note": "formal scoring table from raw DOC; title-only rubric removed",
        "text": """
20.（8分）任务类型：解释与论证（法理依据+事实依据+对发展新质生产力的保障作用）
分项
标准
评分细则（0-8分）
法理依据
①根据民法典规定，当事人不履行合同义务或者履行合同义务不符合约定的，守约方可以要求违约方承担违约责任。
采点赋分。
对应案件一，学生能够从民法典角度说明，准确分析其中一方面得2分；准确分析两方面得4分。
结合案件二，根据反不正当竞争法，被告构成不正当竞争，应当承担停止侵害、赔偿损失的侵权责任。准确回答一点得2分；准确回答两点得4分。
答非所问，或没有应答，得0分。
事实依据
①被告开发时间超过履行期限，阶段性产品未达合同约定标准，构成违约，应当承担违约责任。
保障作用
②依法打击不正当竞争行为，维护消费者知情权和选择权，保护知识产权，保护经营者的合法权益，维护公平竞争的市场秩序，营造良好营商环境，为数字经济发展提供法治保障。
""".strip(),
    },
    "E034": {
        "source_id": "SRC_61b68f1421_细则",
        "note": "same-question formal scoring rubric found in raw term-exam PPT; answer-only PDF removed",
        "text": """
根据反不正当竞争法，陈某在明知他人已经使用“小爱同学”唤醒词并有一定影响的情况下，抢先注册，并发布引人误解的商业宣传信息，构成混淆及虚假宣传，属于不正当竞争，违背了诚实信用原则。该案判决合理确定各方利益，规制了滥用权利的行为，保护了科技创新型企业的品牌商誉，有利于促进新业态的发展，为数字经济发展营造良好的法治化营商环境。
细则：结合材料描述侵权事实 1分；引人误解/混淆/虚假宣传  1分；属于不正当竞争/违反反不正当竞争法  1分； 违背了诚实信用原则 /商业道德1分；
该案判决保护了~或规范了~  2分
有利于新业态/新经济/数字经济发展/创新/1分，维护市场秩序/营商环境 1分
""".strip(),
    },
}

RUBRIC_MEDIA_OVERRIDES: dict[str, dict] = {
    "E007": {
        "source_id": "SRC_e6f22d1048_细则",
        "media": ["word/media/image3.png"],
        "note": "formal scoring rubric image embedded in raw DOCX",
    },
    "E028": {
        "source_id": "SRC_cda046c2d3_细则",
        "media": ["word/media/image1.png"],
        "note": "formal scoring rubric image embedded in raw DOCX",
    },
    "E031": {
        "source_id": "SRC_a043a3fcb2_细则",
        "media": ["word/media/image12.png"],
        "note": "formal scoring rubric image embedded in raw DOCX",
    },
}

RUBRIC_UNRESOLVED: dict[str, str] = {
    "E009": "原始石景山一模细则 PPT 中未定位到本题带分值分布的正式阅卷细则，已移除参考答案式细则。",
    "E043": "原始丰台一模细则 PPT 中仅定位到答案表述，未找到第20题带分值分布的正式阅卷细则，已移除参考答案式细则。",
}


def set_east_asian(run, font: str = "宋体") -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def paragraph_spacing(paragraph, before: int = 0, after: int = 6, line: float = 1.15) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 12, "1F4D78"),
        ("Heading 3", 11, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def read_packets() -> list[dict]:
    with PACKETS.open(encoding="utf-8") as f:
        return [json.loads(line) for line in f if line.strip()]


def read_source_ledger() -> dict[str, dict]:
    with SOURCE_LEDGER.open(encoding="utf-8-sig", newline="") as f:
        return {row["source_id"]: row for row in csv.DictReader(f)}


def normalize(text: str) -> str:
    return re.sub(r"[\s　，。；;：:、（）()\[\]【】《》“”\"'·.．,\-—_]+", "", text or "")


def split_units(text: str) -> list[tuple[int, str]]:
    parts: list[tuple[int, str]] = []
    current = 1
    buf: list[str] = []
    for line in text.splitlines():
        m = re.match(r"^===== .* page (\d+) =====$", line.strip())
        if not m:
            m = re.match(r"^\[PDF_PAGE (\d+)\]$", line.strip())
        if not m:
            m = re.match(r"^\[PPT_SLIDE (\d+)\]$", line.strip())
        if m:
            if buf:
                parts.append((current, "\n".join(buf)))
                buf = []
            current = int(m.group(1))
        else:
            buf.append(line)
    if buf:
        parts.append((current, "\n".join(buf)))
    return parts


def snippets_for(text: str) -> list[str]:
    n = normalize(text)
    if not n:
        return []
    starts = [0, max(0, len(n) // 4), max(0, len(n) // 2), max(0, len(n) * 3 // 4)]
    out: list[str] = []
    for start in starts:
        s = n[start : start + 42]
        if len(s) >= 10 and s not in out:
            out.append(s)
    return out


def rolling_chunks(text: str, size: int = 7, step: int = 3) -> set[str]:
    n = normalize(text)
    if len(n) < size:
        return {n} if n else set()
    return {n[i : i + size] for i in range(0, len(n) - size + 1, step)}


def strip_repeated_headers(text: str) -> str:
    kept: list[str] = []
    for line in (text or "").splitlines():
        stripped = line.strip()
        compact = normalize(stripped)
        if not compact:
            continue
        if (
            "答案及评分参考" in stripped
            or "参考答案第" in stripped
            or re.search(r"第\s*\d+\s*页", stripped)
            or re.search(r"共\s*\d+\s*页", stripped)
        ) and len(compact) >= 12:
            continue
        kept.append(stripped)
    return "\n".join(kept)


def is_cover_like_page(text: str) -> bool:
    compact = normalize(text)
    if not compact:
        return True
    if len(compact) <= 80 and re.search(r"主观题.*细则|综合练习", text or "") and not re.search(r"\d+\s*分|评分|答案|原答案|替[:：]", text or ""):
        return True
    if compact in {"谢谢"}:
        return True
    return False


def score_text_match(text: str, target_text: str, fallback_no: str | None = None) -> int:
    norm = normalize(text)
    target_norm = normalize(target_text)
    if not norm or not target_norm:
        return 0

    score = 0
    for snip in snippets_for(target_text):
        if snip in norm:
            score += 180
        else:
            chunks = [snip[i : i + 8] for i in range(0, max(0, len(snip) - 7), 8)]
            score += sum(12 for c in chunks if c and c in norm)

    target_chunks = rolling_chunks(target_norm, size=8, step=4)
    if target_chunks:
        block_chunks = rolling_chunks(norm, size=8, step=4)
        score += min(120, len(target_chunks & block_chunks) * 5)

    if fallback_no and re.search(rf"(^|\D){re.escape(fallback_no)}[.．、（(]", text):
        score += 30
    if re.search(r"\d+\s*分|评分|标准|变通|参考答案|原答案|答案", text):
        score += 8
    return score


def locate_units(text_path: str | None, target_text: str, fallback_no: str | None = None, max_units: int = 2) -> list[int]:
    if not text_path or not Path(text_path).exists():
        return []
    source = Path(text_path).read_text(encoding="utf-8", errors="ignore")
    units = split_units(source)
    if not units:
        return []
    snippets = snippets_for(target_text)
    scores: list[tuple[int, int]] = []
    for unit_no, unit_text in units:
        norm = normalize(unit_text)
        score = 0
        for snip in snippets:
            if snip in norm:
                score += 120
            else:
                chunks = [snip[i : i + 8] for i in range(0, max(0, len(snip) - 7), 8)]
                score += sum(8 for c in chunks if c and c in norm)
        if fallback_no and re.search(rf"(^|\D){re.escape(fallback_no)}[.．、]", unit_text):
            score += 25
        if score:
            scores.append((score, unit_no))
    if not scores:
        return []
    scores.sort(reverse=True)
    chosen = sorted(unit_no for score, unit_no in scores[:max_units] if score >= max(16, scores[0][0] * 0.35))
    return chosen or [scores[0][1]]


def question_number(entry: dict) -> str | None:
    m = re.search(r"第(\d+)题", entry.get("title", ""))
    return m.group(1) if m else None


def render_pdf_pages(pdf_path: Path, page_nums: list[int], prefix: Path) -> list[Path]:
    if not pdf_path.exists() or not page_nums:
        return []
    out: list[Path] = []
    try:
        pdf = fitz.open(pdf_path)
        for page_no in page_nums:
            if not (1 <= page_no <= len(pdf)):
                continue
            dest = prefix.parent / f"{prefix.name}_p{page_no}.jpg"
            if not dest.exists():
                pix = pdf[page_no - 1].get_pixmap(matrix=fitz.Matrix(1.45, 1.45), alpha=False)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                img.save(dest, "JPEG", quality=86, optimize=True)
            out.append(dest)
        pdf.close()
    except Exception:
        return out
    return out


def candidate_pdf_pages(pdf: fitz.Document, target_text: str, fallback_no: str | None, max_pages: int = 3) -> list[int]:
    scores: list[tuple[int, int]] = []
    for idx in range(len(pdf)):
        text = pdf[idx].get_text()
        if is_cover_like_page(text):
            continue
        score = score_text_match(text, target_text, fallback_no)
        if score:
            scores.append((score, idx + 1))
    if not scores:
        return []
    scores.sort(reverse=True)
    best = scores[0][0]
    threshold = max(28, int(best * 0.32))
    chosen = {page_no for score, page_no in scores[:max_pages] if score >= threshold}
    primary_chosen = set(chosen)
    # Some rubrics continue with one or two scoring lines at the top of the next page.
    # Keep adjacent low-score pages when they still match the locked rubric text.
    target_tail = normalize(strip_repeated_headers(target_text))[-260:]
    tail_chunks = rolling_chunks(target_tail, size=10, step=5)
    for score, page_no in scores:
        if score < 20 or not any(page_no == picked + 1 for picked in primary_chosen):
            continue
        page_start = normalize(strip_repeated_headers(pdf[page_no - 1].get_text()[:800]))
        if tail_chunks and any(chunk and chunk in page_start for chunk in tail_chunks):
            chosen.add(page_no)
    return sorted(chosen)


def union_rects(rects: list[fitz.Rect]) -> fitz.Rect | None:
    if not rects:
        return None
    out = fitz.Rect(rects[0])
    for rect in rects[1:]:
        out.include_rect(rect)
    return out


def expand_rect(rect: fitz.Rect, page_rect: fitz.Rect, margin: float = 18) -> fitz.Rect:
    return fitz.Rect(
        max(page_rect.x0, rect.x0 - margin),
        max(page_rect.y0, rect.y0 - margin),
        min(page_rect.x1, rect.x1 + margin),
        min(page_rect.y1, rect.y1 + margin),
    )


def is_footer_or_page_number(text: str, rect: fitz.Rect, page_rect: fitz.Rect) -> bool:
    stripped = re.sub(r"\s+", "", text or "")
    if stripped in {"", "|"}:
        return True
    if re.fullmatch(r"\d{1,3}", stripped) and (rect.y1 > page_rect.height * 0.9 or rect.y0 < page_rect.height * 0.08):
        return True
    if "参考答案第" in stripped and rect.y0 > page_rect.height * 0.9:
        return True
    return False


def header_question_no(text: str) -> str | None:
    first_line = next((line.strip() for line in (text or "").splitlines() if line.strip()), "")
    m = re.match(r"^(\d{1,2})\s*[.．、]", first_line)
    if m:
        return m.group(1)
    m = re.match(r"^(\d{1,2})\s*[（(]", first_line)
    if m:
        return m.group(1)
    return None


def section_rect_from_blocks(
    blocks: list[tuple[fitz.Rect, str]],
    scored_indices: list[int],
    page_rect: fitz.Rect,
    fallback_no: str | None,
) -> fitz.Rect | None:
    if not scored_indices:
        return None

    start = min(scored_indices)
    end = max(scored_indices)

    for idx in range(start, -1, -1):
        qno = header_question_no(blocks[idx][1])
        if qno and (fallback_no is None or qno == fallback_no):
            start = idx
            break
        if idx < start and header_question_no(blocks[idx][1]) and blocks[idx][0].y1 < blocks[start][0].y0:
            break

    for idx in range(start + 1, len(blocks)):
        qno = header_question_no(blocks[idx][1])
        if qno and (fallback_no is None or qno != fallback_no) and idx > end:
            end = idx - 1
            break
    else:
        end = len(blocks) - 1

    rects = [blocks[idx][0] for idx in range(start, end + 1)]
    return expand_rect(union_rects(rects), page_rect, margin=8) if rects else None


def crop_pdf_rubric_pages(pdf_path: Path, page_nums: list[int], target_text: str, fallback_no: str | None, prefix: Path) -> tuple[list[Path], str]:
    if not pdf_path.exists() or not page_nums:
        return [], "PDF page location failed; used locked rubric text rendered as image"

    out: list[Path] = []
    notes: list[str] = []
    try:
        pdf = fitz.open(pdf_path)
        for page_no in page_nums:
            if not (1 <= page_no <= len(pdf)):
                continue
            page = pdf[page_no - 1]
            if is_cover_like_page(page.get_text()):
                notes.append(f"p{page_no}: cover/closing page skipped")
                continue
            page_rect = page.rect
            raw_blocks = [b for b in page.get_text("blocks") if len(b) >= 5 and (b[4] or "").strip()]
            blocks: list[tuple[fitz.Rect, str]] = []
            for block in sorted(raw_blocks, key=lambda b: (b[1], b[0])):
                rect = fitz.Rect(block[:4])
                text = block[4]
                if is_footer_or_page_number(text, rect, page_rect):
                    continue
                blocks.append((rect, text))

            scored: list[tuple[int, int, fitz.Rect, str]] = []
            for idx, (rect, text) in enumerate(blocks):
                score = score_text_match(text, target_text, fallback_no)
                if score:
                    scored.append((score, idx, rect, text))
            if not scored:
                notes.append(f"p{page_no}: no confident text block")
                continue

            best = max(score for score, _, _, _ in scored)
            target_tail = normalize(strip_repeated_headers(target_text))[-260:]
            tail_chunks = rolling_chunks(target_tail, size=10, step=5)
            continuation_indices = [
                idx
                for idx, (_, text) in enumerate(blocks[:4])
                if tail_chunks and any(chunk and chunk in normalize(strip_repeated_headers(text)) for chunk in tail_chunks)
            ]
            if best < 20 and continuation_indices:
                scored_indices = continuation_indices[:1]
            else:
                threshold = max(12, int(best * 0.18))
                scored_indices = [idx for score, idx, _, _ in scored if score >= threshold]
            clip = section_rect_from_blocks(blocks, scored_indices, page_rect, fallback_no)
            if clip is None:
                notes.append(f"p{page_no}: empty crop")
                continue
            page_area = page_rect.width * page_rect.height
            clip_area = clip.width * clip.height
            if clip.height > page_rect.height * 0.94 or clip_area > page_area * 0.90:
                notes.append(f"p{page_no}: crop too close to full page")
                continue

            dest = prefix.parent / f"{prefix.name}_p{page_no}_crop.jpg"
            pix = page.get_pixmap(matrix=fitz.Matrix(2.2, 2.2), clip=clip, alpha=False)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            img.save(dest, "JPEG", quality=88, optimize=True)
            out.append(dest)
        pdf.close()
    except Exception as exc:
        return out, f"PDF crop failed ({exc}); used locked rubric text rendered as image"

    if out:
        if notes:
            return out, "PDF rubric cropped to local text block(s); " + "; ".join(notes)
        return out, "PDF rubric cropped to local text block(s)"
    return [], "PDF has no usable local text block; used locked rubric text rendered as image"


def render_manual_pdf_rubric(entry: dict, pdf_path: Path, prefix: Path) -> tuple[list[Path], str]:
    specs = MANUAL_RUBRIC_CROPS.get(entry["entry_id"])
    if not specs or not pdf_path.exists():
        return [], ""
    out: list[Path] = []
    try:
        pdf = fitz.open(pdf_path)
        for idx, spec in enumerate(specs, start=1):
            page_no = int(spec["page"])
            if not (1 <= page_no <= len(pdf)):
                continue
            page = pdf[page_no - 1]
            pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            l, t, r, b = spec["box"]
            crop_box = (
                max(0, int(img.width * l)),
                max(0, int(img.height * t)),
                min(img.width, int(img.width * r)),
                min(img.height, int(img.height * b)),
            )
            if crop_box[2] <= crop_box[0] or crop_box[3] <= crop_box[1]:
                continue
            dest = prefix.parent / f"{prefix.name}_manual_p{page_no}_{idx}.jpg"
            img.crop(crop_box).save(dest, "JPEG", quality=88, optimize=True)
            out.append(dest)
        pdf.close()
    except Exception as exc:
        return out, f"manual PDF rubric crop failed ({exc}); used locked rubric text rendered as image"
    if out:
        return out, "manual original PDF rubric crop"
    return [], "manual PDF rubric crop failed; used locked rubric text rendered as image"


def render_pdf_rubric(entry: dict, pdf_path: Path, source: dict, prefix: Path) -> tuple[list[Path], str]:
    manual_assets, manual_note = render_manual_pdf_rubric(entry, pdf_path, prefix)
    if manual_assets:
        return manual_assets, manual_note

    qno = question_number(entry)
    rubric_text = entry.get("rubric", "")
    page_nums: list[int] = []

    try:
        pdf = fitz.open(pdf_path)
        page_nums = candidate_pdf_pages(pdf, rubric_text, qno, max_pages=3)
        pdf.close()
    except Exception:
        page_nums = []

    if not page_nums:
        page_nums = locate_units(source.get("text_path"), rubric_text, qno, max_units=3)

    assets, note = crop_pdf_rubric_pages(pdf_path, page_nums, rubric_text, qno, prefix)
    if assets:
        return assets, note

    return text_images("细则图片", entry["title"], rubric_text, prefix), note


def wrap_line(line: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    if not line:
        return [""]
    pieces: list[str] = []
    current = ""
    for ch in line:
        trial = current + ch
        if font.getlength(trial) <= max_width:
            current = trial
        else:
            if current:
                pieces.append(current)
            current = ch
    if current:
        pieces.append(current)
    return pieces


def text_images(label: str, title: str, text: str, prefix: Path) -> list[Path]:
    width = 1500
    margin = 70
    max_height = 2150
    line_gap = 14
    lines: list[str] = []
    for raw in (text or "").splitlines():
        lines.extend(wrap_line(raw.rstrip(), FONT_BODY, width - 2 * margin))
    if not lines:
        lines = ["（空）"]

    pages: list[list[str]] = []
    page: list[str] = []
    height = margin + 120
    for line in lines:
        line_h = int(FONT_BODY.getbbox(line or " ")[3] - FONT_BODY.getbbox(line or " ")[1] + line_gap)
        if page and height + line_h + margin > max_height:
            pages.append(page)
            page = []
            height = margin + 120
        page.append(line)
        height += line_h
    if page:
        pages.append(page)

    out: list[Path] = []
    for idx, page_lines in enumerate(pages, start=1):
        dest = prefix.parent / f"{prefix.name}_{label}_{idx}.jpg"
        content_height = margin + 120 + len(page_lines) * 52 + margin
        img = Image.new("RGB", (width, min(max_height, max(content_height, 900))), "white")
        draw = ImageDraw.Draw(img)
        draw.rectangle((0, 0, width - 1, img.height - 1), outline=(190, 198, 208), width=2)
        draw.text((margin, margin), f"{label}｜{title}", fill=(31, 77, 120), font=FONT_TITLE)
        y = margin + 96
        for line in page_lines:
            draw.text((margin, y), line, fill=(20, 20, 20), font=FONT_BODY)
            y += 52
        img.save(dest, "JPEG", quality=88, optimize=True)
        out.append(dest)
    return out


def extract_docx_media(docx_path: Path, dest_dir: Path, limit: int = 6) -> list[Path]:
    if not docx_path.exists() or docx_path.suffix.lower() != ".docx":
        return []
    dest_dir.mkdir(parents=True, exist_ok=True)
    out: list[Path] = []
    try:
        with zipfile.ZipFile(docx_path) as z:
            names = [n for n in z.namelist() if n.startswith("word/media/")]
            for name in names[:limit]:
                suffix = Path(name).suffix.lower()
                if suffix not in {".png", ".jpg", ".jpeg"}:
                    continue
                dest = dest_dir / f"{docx_path.stem}_{Path(name).name}"
                if not dest.exists():
                    with z.open(name) as src, dest.open("wb") as dst:
                        shutil.copyfileobj(src, dst)
                out.append(dest)
    except Exception:
        return []
    return out


def extract_named_docx_media(docx_path: Path, media_names: list[str], dest_dir: Path, prefix: str) -> list[Path]:
    if not docx_path.exists() or docx_path.suffix.lower() != ".docx":
        return []
    dest_dir.mkdir(parents=True, exist_ok=True)
    out: list[Path] = []
    wanted = set(media_names)
    try:
        with zipfile.ZipFile(docx_path) as z:
            for name in z.namelist():
                if name not in wanted:
                    continue
                suffix = Path(name).suffix.lower()
                if suffix not in {".png", ".jpg", ".jpeg"}:
                    continue
                dest = dest_dir / f"{prefix}_{Path(name).name}"
                with z.open(name) as src, dest.open("wb") as dst:
                    shutil.copyfileobj(src, dst)
                out.append(dest)
    except Exception:
        return []
    return out


def rubric_source_for(entry: dict, source_map: dict[str, dict]) -> tuple[dict, Path, str, str]:
    eid = entry["entry_id"]
    source_id = entry.get("rubric_source_id", "")
    rubric_text = entry.get("rubric", "")
    note_prefix = ""

    if eid in RUBRIC_TEXT_OVERRIDES:
        override = RUBRIC_TEXT_OVERRIDES[eid]
        source_id = override["source_id"]
        rubric_text = override["text"]
        note_prefix = override["note"]
    elif eid in RUBRIC_UNRESOLVED:
        rubric_text = (
            "【未找到正式评分细则】\n"
            f"{entry['title']}\n"
            f"{RUBRIC_UNRESOLVED[eid]}\n"
            "本题不再以参考答案充当细则；需补充正式阅卷细则后替换。"
        )
        note_prefix = "formal scoring rubric not found in raw source; reference answer removed"
    elif eid in RUBRIC_MEDIA_OVERRIDES:
        override = RUBRIC_MEDIA_OVERRIDES[eid]
        source_id = override["source_id"]
        note_prefix = override["note"]

    source = source_map.get(source_id, source_map.get(entry.get("rubric_source_id", ""), {}))
    path = Path(source.get("path") or entry.get("rubric_path", ""))
    if eid in RUBRIC_UNRESOLVED:
        path = Path(entry.get("rubric_path", ""))
    return source, path, rubric_text, note_prefix


def build_assets(entries: list[dict], source_map: dict[str, dict]) -> tuple[dict[str, list[Path]], dict[str, list[Path]], list[dict]]:
    if ASSET_DIR.exists():
        shutil.rmtree(ASSET_DIR)
    (ASSET_DIR / "question").mkdir(parents=True)
    (ASSET_DIR / "rubric").mkdir(parents=True)
    question_assets: dict[str, list[Path]] = defaultdict(list)
    rubric_assets: dict[str, list[Path]] = defaultdict(list)
    rows: list[dict] = []

    for entry in entries:
        eid = entry["entry_id"]
        q_source = source_map.get(entry["question_source_id"], {})
        q_path = Path(entry["question_path"])
        r_source, r_path, rubric_text, override_note = rubric_source_for(entry, source_map)
        q_note = ""
        r_note = ""

        q_ext = q_path.suffix.lower()
        if q_ext == ".pdf":
            q_text = f"{entry.get('material','')}\n{entry.get('prompt','')}"
            q_pages = locate_units(q_source.get("text_path"), q_text, question_number(entry), max_units=2)
            q_assets = render_pdf_pages(q_path, q_pages, ASSET_DIR / "question" / f"{eid}_question")
            if not q_assets:
                q_assets = text_images("试题图片", entry["title"], q_text, ASSET_DIR / "question" / f"{eid}_question")
                q_note = "PDF page location/render failed; used locked text rendered as image"
            question_assets[eid].extend(q_assets)
        else:
            q_text = f"{entry.get('material','')}\n\n{entry.get('prompt','')}"
            question_assets[eid].extend(text_images("试题图片", entry["title"], q_text, ASSET_DIR / "question" / f"{eid}_question"))
            q_note = f"{q_ext} source rendered from locked text because direct page rasterization is unavailable"

        if eid in RUBRIC_UNRESOLVED:
            rubric_assets[eid].extend(text_images("细则图片", entry["title"], rubric_text, ASSET_DIR / "rubric" / f"{eid}_rubric"))
            r_note = override_note
        elif eid in RUBRIC_MEDIA_OVERRIDES:
            override = RUBRIC_MEDIA_OVERRIDES[eid]
            r_assets = extract_named_docx_media(
                r_path,
                override["media"],
                ASSET_DIR / "rubric",
                f"{eid}_rubric",
            )
            rubric_assets[eid].extend(r_assets)
            r_note = override_note
            if not r_assets:
                failure_text = (
                    "【正式细则图片提取失败】\n"
                    f"{entry['title']}\n"
                    "本题的正式细则应来自原始 DOCX 内嵌图片，但自动提取失败；"
                    "已避免回退为参考答案。请人工检查原始细则文件后替换。"
                )
                rubric_assets[eid].extend(text_images("细则图片", entry["title"], failure_text, ASSET_DIR / "rubric" / f"{eid}_rubric"))
                r_note = f"{override_note}; embedded image extraction failed; reference answer not used"
        else:
            effective_entry = dict(entry)
            effective_entry["rubric"] = rubric_text
            r_ext = r_path.suffix.lower()
            if r_ext == ".pdf":
                r_assets, r_note = render_pdf_rubric(effective_entry, r_path, r_source, ASSET_DIR / "rubric" / f"{eid}_rubric")
                rubric_assets[eid].extend(r_assets)
            else:
                rubric_assets[eid].extend(text_images("细则图片", entry["title"], rubric_text, ASSET_DIR / "rubric" / f"{eid}_rubric"))
                r_note = f"{r_ext} source rendered from locked rubric text"
        if override_note and override_note not in r_note:
            r_note = f"{override_note}; {r_note}" if r_note else override_note

        rows.append(
            {
                "entry_id": eid,
                "title": entry["title"],
                "question_source": str(q_path),
                "question_asset_count": str(len(question_assets[eid])),
                "question_assets": ";".join(str(p) for p in question_assets[eid]),
                "question_note": q_note,
                "rubric_source": str(r_path),
                "rubric_asset_count": str(len(rubric_assets[eid])),
                "rubric_assets": ";".join(str(p) for p in rubric_assets[eid]),
                "rubric_note": r_note,
            }
        )

    with REPORT_CSV.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=[
                "entry_id",
                "title",
                "question_source",
                "question_asset_count",
                "question_assets",
                "question_note",
                "rubric_source",
                "rubric_asset_count",
                "rubric_assets",
                "rubric_note",
            ],
        )
        writer.writeheader()
        writer.writerows(rows)
    return question_assets, rubric_assets, rows


def add_picture(doc: Document, image_path: Path, caption: str) -> None:
    with Image.open(image_path) as im:
        w, h = im.size
    max_width = 6.65
    max_height = 8.55
    scaled_height = max_width * (h / w) if w else max_height
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if scaled_height > max_height:
        p.add_run().add_picture(str(image_path), height=Inches(max_height))
    else:
        p.add_run().add_picture(str(image_path), width=Inches(max_width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_spacing(cap, after=8, line=1.0)
    for run in cap.runs:
        set_east_asian(run)
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string("666666")


def build_doc() -> None:
    entries = read_packets()
    source_map = read_source_ledger()
    question_assets, rubric_assets, _ = build_assets(entries, source_map)

    doc = Document()
    style_document(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("选必二《法律与生活》习题汇编（2024-2026）图片版")
    set_east_asian(r)
    r.font.size = Pt(20)
    r.font.bold = True
    paragraph_spacing(title, after=4)
    sub = doc.add_paragraph("每道题按“试题图片 + 细则图片”合并排列；不收录材料/设问/细则文字版。")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_spacing(sub, after=10)
    for run in sub.runs:
        set_east_asian(run)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string("555555")

    current_year = None
    for idx, entry in enumerate(entries, start=1):
        if idx > 1:
            doc.add_page_break()
        if str(entry["year"]) != current_year:
            current_year = str(entry["year"])
            doc.add_heading(current_year, level=1)
        doc.add_heading(entry["title"], level=2)
        doc.add_heading("试题图片", level=3)
        for n, image in enumerate(question_assets.get(entry["entry_id"], []), start=1):
            add_picture(doc, image, f"试题图片 {n}")
        doc.add_heading("细则图片", level=3)
        for n, image in enumerate(rubric_assets.get(entry["entry_id"], []), start=1):
            add_picture(doc, image, f"细则图片 {n}")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)
    print(REPORT_CSV)


if __name__ == "__main__":
    build_doc()
