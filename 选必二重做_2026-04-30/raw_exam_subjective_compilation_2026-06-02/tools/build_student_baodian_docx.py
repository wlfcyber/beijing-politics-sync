#!/usr/bin/env python3
from __future__ import annotations

import importlib.util
import re
import shutil
import sys
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


RUN_DIR = Path(__file__).resolve().parents[1]
OUT_DIR = RUN_DIR / "05_output"
OUT_DOCX = OUT_DIR / "选必二法律与生活_主观题细则宝典_学生版_20260604.docx"
DESKTOP_COPY = Path.home() / "Desktop" / OUT_DOCX.name
QA_MD = OUT_DIR / "STUDENT_BAODIAN_LAYOUT_CHECK_20260604.md"

FONT_EA_BODY = "宋体"
FONT_EA_HEAD = "微软雅黑"
FONT_WEST = "Calibri"
BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F8FC"
GRAY = "666666"
RED = "C00000"
GOLD = "7A5A00"
BORDER = "B8C2CC"


def load_base_builder():
    path = RUN_DIR / "tools" / "build_full_text_subjective_rubric_docx.py"
    spec = importlib.util.spec_from_file_location("full_text_builder", path)
    if not spec or not spec.loader:
        raise RuntimeError(f"cannot load {path}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


BASE = load_base_builder()

BODY_EXCLUDE_ENTRY_IDS = {
    "E009": "模块边界和正式细则完整性待教师复核，正文暂不作为学生刷题题卡。",
    "E005": "题面为多栏 OCR 拼接，材料与资料包行间错位，暂入附录待重排。",
    "E006": "虚拟数字人题面与资料卡在当前文本层串行错位，暂入附录待重排。",
    "E015": "表格多栏材料在当前文本层行列错位，暂入附录待重排。",
    "E016": "表格多栏材料在当前文本层行列错位，暂入附录待重排。",
    "E022": "三列表格材料在当前文本层拆分错位，暂入附录待重排。",
    "E023": "三列表格材料在当前文本层拆分错位，暂入附录待重排。",
    "E024": "示例表与案情段落行列错位，暂入附录待重排。",
    "E035": "案情与侧栏法条在当前文本层交错，暂入附录待重排。",
    "E036": "案件事实与裁判要点两栏在当前文本层交错，暂入附录待重排。",
    "E053": "校园欺凌与治安管理处罚法跨模块色彩强，暂入附录由教师取舍。",
    "E074": "题面材料表在当前文本层缺失，暂入附录，待回源补齐后再进正文。",
}

DISTRICT_ORDER = {
    name: idx
    for idx, name in enumerate(
        [
            "东城",
            "西城",
            "朝阳",
            "海淀",
            "丰台",
            "石景山",
            "通州",
            "顺义",
            "顺义思政",
            "房山",
            "大兴",
            "昌平",
            "门头沟",
            "平谷",
            "延庆",
            "密云",
            "怀柔",
        ]
    )
}
PAPER_ORDER = {
    "期中": 5,
    "期末": 10,
    "一模": 20,
    "二模": 30,
    "三模": 40,
    "模拟": 50,
}

ENTRY_AXIS_OVERRIDES = {
    "E004": ("A6 侵权责任", "B5 意义价值"),
    "E007": ("A7 婚姻家庭继承", "B2 判责/理由"),
    "E010": ("A7 婚姻家庭继承", "B4 评析认识"),
    "E011": ("A7 婚姻家庭继承", "B4 评析认识"),
    "E012": ("A7 婚姻家庭继承", "B4 评析认识"),
    "E014": ("A3 物权/相邻", "B7 问题识别/短答"),
    "E018": ("A5 知识产权/竞争", "B5 意义价值"),
    "E026": ("A2 人格权/隐私 + A3 物权/相邻 + A6 侵权责任", "B2 判责/理由"),
    "E027": ("A2 人格权/隐私 + A3 物权/相邻 + A6 侵权责任", "B6 维权路径"),
    "E029": ("A6 侵权责任", "B1 表格/补链"),
    "E021": ("A5 知识产权/竞争", "B5 意义价值"),
    "E030": ("A8 劳动关系", "B5 意义价值"),
    "E033": ("A4 合同 + A5 知识产权/竞争", "B1 表格/补链"),
    "E037": ("A3 物权/相邻 + A5 知产/竞争", "B2 判责/理由"),
    "E038": ("A3 物权/相邻 + A5 知产/竞争", "B2 判责/理由"),
    "E042": ("A5 知识产权/竞争", "B4 评析认识"),
    "E045": ("A1 民事主体/法律关系", "B4 评析认识"),
    "E046": ("A1 民事主体/法律关系", "B4 评析认识"),
    "E048": ("A5 知识产权/竞争", "B4 评析认识"),
    "E052": ("A10 纠纷解决程序", "B2 判责/理由"),
    "E058": ("A3 物权/相邻", "B2 判责/理由"),
    "E060": ("A10 纠纷解决程序", "B2 判责/理由"),
    "E062": ("A8 劳动关系", "B5 意义价值"),
    "E064": ("A10 纠纷解决程序", "B5 意义价值"),
    "E069": ("A5 知识产权/竞争", "B4 评析认识"),
    "E070": ("A5 知识产权/竞争", "B4 评析认识"),
    "E074": ("A3 物权/相邻", "B5 意义价值"),
}

STUDENT_MATERIAL_OVERRIDES = {
    "E014": (
        "某校开展“我为社区献一策”的社会实践活动。调查发现：某社区有 446 辆电动自行车，却只有 60 个充电接口，"
        "新增充电设施引发绿地占用、集中充电、消防安全、通行采光等争议。学生建议在公共区域设置“车电分离”的电池充电柜。\n"
        "资料卡：《高层民用建筑消防安全管理规定》第三十七条要求电动自行车存放、充电场所配备必要消防器材，充电设施具备充满自动断电功能；"
        "《中华人民共和国民法典》第二百七十八条规定，改变共有部分用途或者利用共有部分从事经营活动，由业主共同决定。"
    ),
    "E059": (
        "20.（8分）阅读材料，完成下表。\n"
        "基本案情：李某系社区广场舞义务组织者，长期无偿提供音响、安排场地、组织活动，并时常提醒参与者量力而行、注意安全。"
        "张某有轻微心脏不适史但未告知李某，自愿参加活动。某日张某跳舞时突然倒地，李某立即拨打120并实施心肺复苏，"
        "张某最终因“心源性猝死”抢救无效死亡。张某家属以李某未尽安全保障义务为由诉至法院，要求赔偿60万元。\n"
        "裁判结果：人民法院驳回张某家属诉讼请求，判定李某无需承担责任。\n"
        "法律依据：《中华人民共和国民法典》第一千一百七十六条、第一千一百九十八条相关规定。"
    )
}

STUDENT_PROMPT_OVERRIDES = {
    "E014": "（1）遇到的法律问题：________、________。",
}

STUDENT_RUBRIC_OVERRIDES = {
    "E017": (
        "参考答案：\n"
        "根据《中华人民共和国劳动法》和《中华人民共和国劳动合同法》规定，劳动者应当完成劳动任务，遵守劳动纪律和职业道德。（1分）\n"
        "尹某没有履行劳动者义务，违反相关法律规定，不符合敬业等社会主义核心价值观。（2分）\n"
        "购物中心依据管理规章制度解除劳动合同，符合《中华人民共和国劳动合同法》第三十九条规定。（1分）\n"
        "法院驳回尹某的诉讼请求，保护了用人单位合理的用工自主权，引导劳动者坚持权利义务相统一，构建健康劳动关系。（3分）"
    ),
    "E030": (
        "21.（8分）\n"
        "规范“竞业限制”：有利于保护企业商业秘密、保护知识产权、推动社会创新。（任意两点2分）\n"
        "限制“竞业限制”：防止过度限制劳动者基本劳动权利，有利于人才资源合理配置。（任意一点2分）\n"
        "综合：平衡劳动者合法权益与企业商业秘密保护，维护市场公平竞争，体现公平原则和权利义务相统一。（2分）\n"
        "案例一：李某作为负有保密义务的劳动者，违反竞业限制约定，应承担违约责任。（1分）\n"
        "案例二：刘某作为普通员工，并非负有保密义务的人员，餐饮公司与其签订的竞业限制协议违反法律规定，协议无效。（1分）"
    ),
    "E041": (
        "18.（8分）\n"
        "案例一：人民法院通过诉讼调解，推动纠纷实质性化解，节约司法资源，实现多方共赢。（2分）\n"
        "案例二：对知识产权侵权适用惩罚性赔偿，严惩主观恶性强、社会危害大的侵权行为，提高侵权成本，"
        "保护权利人合法权益，维护市场秩序和公平竞争。（2分）\n"
        "案例三：驳回并谴责恶意诉讼，惩治滥用诉讼权利的行为，贯彻诚信原则，明确权利行使界限，"
        "维护诉讼秩序、节约司法资源。（2分）\n"
        "总说：人民法院坚持公正司法，以事实为依据、以法律为准绳，发挥司法裁判导向作用；"
        "保护知识产权就是保护创新，有利于提升创新活力。（2分）"
    ),
    "E058": (
        "（1）（8分）\n"
        "法院判决以事实为根据，以法律为准绳。老旧小区加装电梯属民生工程，已获法定比例业主同意，程序合法。"
        "电梯采用玻璃幕墙设计，在提升本楼业主出行便利的同时，能尽量减少对相邻楼业主的影响。"
        "依据《中华人民共和国民法典》规定，不动产的相邻权利人应当按照有利生产、方便生活、团结互助、公平合理的原则，"
        "正确处理相邻关系。本案判决维护了相邻业主合法权益，有助于促进邻里和谐，践行友善的社会主义核心价值观。"
    ),
    "E059": (
        "20.（8分）\n"
        "理由：根据《中华人民共和国民法典》规定，参与者不存在故意或重大过失，组织者尽到安全保障义务，则无需承担责任。（2分）"
        "本案中张某自愿参加活动，因自身原因突发意外，并非他人行为造成。（1分）"
        "李某虽为广场舞义务组织者，但已尽到合理的安全保障义务，事前提醒参与者注意安全，事后及时救助，"
        "且不存在故意或重大过失，与张某死亡不存在因果关系，不承担侵权责任。（2分）\n"
        "意义：引导参与者增强风险意识，对自己的行为负责；合理界定义务组织者的责任范围，体现公平原则；"
        "鼓励社会成员在互帮互助中传递社会主义核心价值观，维护和谐有序的社区生活秩序。（3分）"
    ),
    "E060": (
        "18.（1）（8分）\n"
        "从诉讼主体看：公益社会组织为维护生态环境公共利益，对涉案企业提起环境民事公益诉讼，推动损害风险进入司法程序。\n"
        "从纠纷解决方式看：本案综合运用和解、司法调解、诉讼等方式，发挥多元纠纷解决机制作用。\n"
        "从法律效力看：法院对首份和解协议依法审查，对不足以保护公共利益的协议不予确认；经调解达成协议后出具民事调解书，赋予其强制执行效力。\n"
        "从履行监督看：人民法院持续督导履行，推动封场、环境修复和验收，维护生态环境公共利益。"
    ),
}

RUBRIC_TRUNCATE_MARKERS = [
    "【附中版补充】",
    "（二）学生表现",
    "学生表现和问题",
    "海淀期末21题讲评",
    "一、试题分析",
    "出现的问题",
    "水平4",
    "部分同学",
    "建议最后冲刺",
    "首先，案例分析时",
    "第一，案例分析时",
    "第二，学生",
    "能力要求\n分析角度",
    "能力要求",
    "学生学习",
    "教师教学",
    "太阳风带电粒子",
    "中国科学院大学星际航行学院",
    "独特优势：",
    "全球治理倡议",
    "友好关系；主持上合组织",
    "论坛部长级会议",
]

BAD_RUBRIC_LINE_CONTAINS = [
    "阅卷总结",
    "学生问题",
    "学生表现",
    "存在问题",
    "出现的问题",
    "试题分析",
    "讲评",
    "评分标准说明",
    "阅卷前制定的参考答案",
    "【附中版补充】",
    "部分同学",
    "高三政治综合练习",
    "统一练习非选择题评分细则",
    "主观题细则",
    "2026.5",
    "PAGEPAGE",
    "病象",
    "病状",
    "病因",
    "药方",
    "教师教学",
    "学生学习",
    "教学启示",
    "改进措施",
    "谢谢",
]


def set_fonts(run, ea: str = FONT_EA_BODY, west: str = FONT_WEST) -> None:
    run.font.name = west
    run._element.rPr.rFonts.set(qn("w:eastAsia"), ea)


def para_spacing(p, before: int = 0, after: int = 6, line: float = 1.25) -> None:
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=90, start=140, bottom=90, end=140) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120, border_color: str = BORDER) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), border_color)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def style_cell_text(cell, size: float = 9.5, bold_first: bool = False) -> None:
    for p in cell.paragraphs:
        para_spacing(p, after=2, line=1.18)
        for idx, run in enumerate(p.runs):
            set_fonts(run)
            run.font.size = Pt(size)
            if bold_first and idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(BLUE)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, head in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = head
        set_shading(cell, LIGHT_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    set_table_geometry(table, widths)
    for ridx, row in enumerate(table.rows):
        for cell in row.cells:
            style_cell_text(cell, size=9 if ridx else 9.5, bold_first=ridx == 0)


def add_callout(doc: Document, title: str, text: str, fill: str = PALE_BLUE, title_color: str = BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_shading(cell, fill)
    p = cell.paragraphs[0]
    para_spacing(p, after=3, line=1.2)
    r = p.add_run(title)
    set_fonts(r, FONT_EA_HEAD)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(title_color)
    if text:
        p.add_run().add_break()
        r = p.add_run(text.strip())
        set_fonts(r)
        r.font.size = Pt(10)
    set_table_geometry(table, [9360], border_color="D5DFEA")


def add_plain_paragraph(doc: Document, text: str, size: float = 10.5, color: str | None = None, bold: bool = False) -> None:
    p = doc.add_paragraph()
    para_spacing(p, after=5, line=1.24)
    for idx, part in enumerate((text or "").split("\n")):
        if idx:
            p.add_run().add_break()
        r = p.add_run(part)
        set_fonts(r)
        r.font.size = Pt(size)
        r.bold = bold
        if color:
            r.font.color.rgb = RGBColor.from_string(color)


def add_small_source(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    para_spacing(p, before=1, after=4, line=1.12)
    r = p.add_run(text)
    set_fonts(r)
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(GRAY)


def add_label_text(doc: Document, label: str, text: str, fill: str = "FFFFFF") -> None:
    p = doc.add_paragraph()
    para_spacing(p, before=4, after=3, line=1.2)
    r = p.add_run(f"【{label}】")
    set_fonts(r, FONT_EA_HEAD)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    body = (text or "【缺失】").strip()
    if body:
        p.add_run().add_break()
        for idx, part in enumerate(body.split("\n")):
            if idx:
                p.add_run().add_break()
            r = p.add_run(part)
            set_fonts(r)
            r.font.size = Pt(9.6)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_fonts(run, FONT_EA_HEAD)


def add_field(run, field: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_WEST
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EA_BODY)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, "0B2545", 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = FONT_WEST
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EA_HEAD)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    if "FG CardTitle" not in doc.styles:
        style = doc.styles.add_style("FG CardTitle", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT_WEST
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EA_HEAD)
        style.font.size = Pt(12)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = header.add_run("飞哥正志讲堂｜选必二《法律与生活》主观题细则宝典")
    set_fonts(r, FONT_EA_HEAD)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("第 ")
    set_fonts(r)
    r.font.size = Pt(9)
    add_field(footer.add_run(), "PAGE")
    r = footer.add_run(" 页")
    set_fonts(r)
    r.font.size = Pt(9)


def add_cover(doc: Document, stats, body_groups: list[list[dict]]) -> None:
    p = doc.add_paragraph()
    para_spacing(p, after=0, line=1.1)
    r = p.add_run("选择性必修二")
    set_fonts(r, FONT_EA_HEAD)
    r.font.size = Pt(14)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    para_spacing(p, after=4, line=1.1)
    r = p.add_run("《法律与生活》主观题细则宝典")
    set_fonts(r, FONT_EA_HEAD)
    r.font.size = Pt(24)
    r.bold = True

    p = doc.add_paragraph()
    para_spacing(p, after=18, line=1.1)
    r = p.add_run("学生版｜题源材料 + 设问拆解 + 评分细则")
    set_fonts(r, FONT_EA_HEAD)
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(GRAY)

    add_callout(
        doc,
        "本册怎么用",
        "先看前两页速查卡，弄清题目争点和法律入口；再按年份刷题卡；每题先读“案件摘要”和“细则原文”，最后回到原材料改写成自己的答案。题源索引放在附录，不打断学生阅读。",
    )
    add_table(
        doc,
        ["阅读顺序", "学生动作", "拿分结果"],
        [
            ["1. 案件摘要", "先圈主体、利益冲突、题干动词", "知道这题问什么"],
            ["2. 入口", "用 A 轴定法律关系，用 B 轴定答案形状", "不跑模块、不空喊价值"],
            ["3. 细则", "先看给分句，再回扣完整材料", "答题句子有踩分点"],
            ["4. 仿写", "用材料词改写规则句", "形成可迁移答案"],
        ],
        [1500, 3900, 3960],
    )
    body_subquestions = sum(len(group) for group in body_groups)
    add_plain_paragraph(
        doc,
        f"当前正文入册：{len(body_groups)} 张题卡，{body_subquestions} 个分问。原始题包共 {stats.big_questions_total} 道大题、{stats.subquestions_total} 个分问；材料缺失或边界待复核题卡移至附录。",
        size=9.5,
        color=GRAY,
    )
    doc.add_page_break()


def add_front_navigation(doc: Document) -> None:
    add_heading(doc, "第一部分  考场速查：一眼看出这题怎么答", 1)
    add_callout(
        doc,
        "母公式",
        "生活事实 -> 争点 -> 法律翻译 -> 法律结果 -> 必要价值。不要先堆法条，也不要先喊法治；每句话都要把材料事实压进规则。",
        fill="FFF7E6",
        title_color=GOLD,
    )
    add_table(
        doc,
        ["先看什么", "学生要问自己", "落笔形状"],
        [
            ["主体", "谁和谁发生冲突？谁受损？", "先写法律关系，不先写口号。"],
            ["争点", "到底争的是钱、合同、人格、劳动、消费，还是程序？", "入口错，全题会散。"],
            ["题干动词", "问理由、诉求、评析、意义、路径，还是填表？", "B 轴决定答案长相。"],
            ["材料事实", "哪些事实能对应构成要件、义务、责任或程序？", "一句规则后面必须跟一句事实。"],
        ],
        [1500, 3900, 3960],
    )

    add_heading(doc, "A 轴：法律入口速查", 2)
    add_table(
        doc,
        ["入口", "看到什么先想到", "容易误判"],
        [
            ["A1 民事行为效力", "未成年人、追认、意思表示、代理、时效", "不要把所有民法题都塞进 A1"],
            ["A2 人格权", "名誉、隐私、个人信息、肖像、生命健康", "有消费背景也可能先是人格权"],
            ["A3 物权/相邻", "共有部分、通行、采光、业主共同决定", "审批事实不等于程序题"],
            ["A4 合同", "要约承诺、履行、违约、格式条款、解除", "有损害再补侵权"],
            ["A5 知产/竞争", "商业秘密、商标、著作权、混淆、诋毁", "创新材料必须落到具体权利"],
            ["A6 侵权责任", "过错、损害、因果、安全保障义务", "不要泛泛写保护权利"],
            ["A7 婚姻家庭继承", "夫妻财产、赡养、继承、遗赠扶养", "亲情价值后面要有具体义务"],
            ["A8 劳动关系", "招聘、合同、解除、竞业限制、社保", "劳动法既保护劳动者也维护有序用工"],
            ["A9 消费者权益", "欺诈、知情选择、公平交易、惩罚性赔偿", "先看有没有经营者-消费者关系"],
            ["A10 纠纷解决", "证据、调解、仲裁、诉讼、维权路径", "只有问路径时才主导"],
        ],
        [1450, 4300, 3610],
    )

    add_heading(doc, "B 轴：题干动作速查", 2)
    add_table(
        doc,
        ["动作", "题干常见词", "一句话形状"],
        [
            ["B1 表格/补链", "完成下表、参考示例、补裁判理由", "事实 -> 规则 -> 裁判要点"],
            ["B2 判责/理由", "为什么这样判、谁担责", "法律关系 -> 规则 -> 事实 -> 责任"],
            ["B3 诉求支持", "诉求能否支持、请求是否成立", "先结论，再逐项匹配权利基础"],
            ["B4 评析认识", "评析、认识、谈谈理解", "合理处 -> 片面处 -> 正确法律立场"],
            ["B5 意义价值", "意义、价值、启示、原因", "保护对象 -> 约束行为 -> 秩序价值"],
            ["B6 维权路径", "如何维权、准备什么证据、走什么程序", "证据 -> 路径 -> 请求"],
            ["B7 问题识别", "指出问题、填空、短答", "短语或短句，不写长判决"],
        ],
        [1450, 3900, 4010],
    )
    doc.add_page_break()


def infer_a_axis(text: str) -> str:
    rules = [
        ("A8 劳动关系", ["劳动", "用工", "竞业", "招聘", "社保", "工资", "劳动者"]),
        ("A9 消费者权益", ["消费者", "经营者", "欺诈", "惩罚性", "退费", "直播", "知情权"]),
        ("A5 知识产权/竞争", ["商业秘密", "著作权", "商标", "专利", "混淆", "不正当竞争", "技术秘密"]),
        ("A4 合同", ["合同", "要约", "承诺", "违约", "履行", "格式条款", "解除"]),
        ("A6 侵权责任", ["侵权", "损害", "过错", "安全保障", "赔偿", "因果关系"]),
        ("A2 人格权", ["隐私", "个人信息", "名誉", "肖像", "姓名", "生命健康"]),
        ("A3 物权/相邻", ["所有权", "共有", "相邻", "通行", "采光", "业主", "物业"]),
        ("A7 婚姻家庭继承", ["夫妻", "赡养", "继承", "遗嘱", "遗赠", "抚养"]),
        ("A10 纠纷解决程序", ["调解", "仲裁", "诉讼", "举证", "证据", "维权"]),
        ("A1 民事行为效力", ["未成年人", "追认", "行为能力", "意思表示", "代理", "诉讼时效"]),
    ]
    for label, keys in rules:
        if any(k in text for k in keys):
            return label
    return "A 入口待由材料确定"


def infer_b_axis(prompt: str) -> str:
    checks = [
        ("B1 表格/补链", ["表", "补全", "示例", "填"]),
        ("B3 诉求支持", ["诉求", "请求", "支持", "能否"]),
        ("B4 评析认识", ["评析", "认识", "理解", "谈谈", "观点"]),
        ("B5 意义价值", ["意义", "价值", "原因", "启示", "作用"]),
        ("B6 维权路径", ["维权", "路径", "建议", "怎么办", "如何"]),
        ("B2 判责/理由", ["理由", "责任", "承担", "判", "法院", "仲裁"]),
    ]
    for label, keys in checks:
        if any(k in prompt for k in keys):
            return label
    return "B7 问题识别/短答"


def first_sentence(text: str, max_len: int = 110) -> str:
    clean = re.sub(r"\s+", " ", (text or "").strip())
    if not clean:
        return "先看细则原文。"
    parts = re.split(r"(?<=[。！？])", clean)
    sent = parts[0].strip() if parts else clean
    if len(sent) > max_len:
        sent = sent[:max_len].rstrip() + "..."
    return sent


def q_number(entry: dict) -> int:
    qno = BASE.q_base(entry)
    m = re.search(r"\d+", str(qno))
    return int(m.group(0)) if m else 999


def student_group_key(entry: dict) -> tuple:
    return (
        str(entry["year"]),
        entry["district_or_exam"],
        entry["paper_type"],
        BASE.q_base(entry),
    )


def group_sort_key(group: list[dict]) -> tuple:
    first = group[0]
    year = int(first["year"]) if str(first["year"]).isdigit() else 9999
    district = DISTRICT_ORDER.get(first["district_or_exam"], 999)
    paper = PAPER_ORDER.get(first["paper_type"], 999)
    return (year, district, paper, q_number(first), min(e["entry_id"] for e in group))


def make_body_groups(packets: list[dict]) -> list[list[dict]]:
    grouped: dict[tuple, list[dict]] = defaultdict(list)
    for entry in packets:
        if entry["entry_id"] in BODY_EXCLUDE_ENTRY_IDS:
            continue
        grouped[student_group_key(entry)].append(entry)
    return sorted((sorted(group, key=lambda e: (q_number(e), e["entry_id"])) for group in grouped.values()), key=group_sort_key)


def clean_common_student_text(text: str) -> str:
    lines: list[str] = []
    for raw in (text or "").splitlines():
        line = re.sub(r"\s+", " ", raw).strip()
        if not line:
            if lines and lines[-1]:
                lines.append("")
            continue
        compact = re.sub(r"\s+", "", line)
        if re.search(r"高三年级[（(]?思想政治[）)]?第\d+页[（(]?共\d+页[）)]?", compact):
            continue
        if re.search(r"高三政治第\d+页[（(]?共\d+页[）)]?", compact):
            continue
        if re.fullmatch(r"第\d+页[（(]?共\d+页[）)]?", compact):
            continue
        if re.fullmatch(r"[—\\-－_ ]*\\d+[—\\-－_ ]*", line):
            continue
        lines.append(line)
    return "\n".join(lines).strip()


def display_material(entry: dict) -> str:
    return clean_common_student_text(STUDENT_MATERIAL_OVERRIDES.get(entry["entry_id"], (entry.get("material") or "").strip()))


def display_prompt(entry: dict) -> str:
    return clean_common_student_text(STUDENT_PROMPT_OVERRIDES.get(entry["entry_id"], (entry.get("prompt") or "").strip()))


def clean_student_rubric(entry: dict, overrides: dict[str, str]) -> str:
    entry_id = entry["entry_id"]
    text = STUDENT_RUBRIC_OVERRIDES.get(entry_id) or BASE.rubric_text(entry, overrides)
    text = clean_common_student_text(text)
    text = text.replace("\t", " ")
    text = re.sub(r"\s+\|\s+", "\n", text)
    for marker in RUBRIC_TRUNCATE_MARKERS:
        pos = text.find(marker)
        if pos > 0:
            text = text[:pos].rstrip()
    lines: list[str] = []
    for line in text.splitlines():
        clean = line.strip()
        if not clean:
            if lines and lines[-1]:
                lines.append("")
            continue
        if any(marker in clean for marker in BAD_RUBRIC_LINE_CONTAINS):
            continue
        lines.append(clean)
    cleaned = "\n".join(lines).strip()
    return cleaned or "细则原文待补。"


def group_axis(group: list[dict], material: str, prompt: str, rubric: str) -> tuple[str, str]:
    for entry in group:
        override = ENTRY_AXIS_OVERRIDES.get(entry["entry_id"])
        if override:
            return override
    return infer_a_axis(material + " " + prompt + " " + rubric), infer_b_axis(prompt)


def add_question_card(doc: Document, card_no: int, group: list[dict], overrides: dict[str, str]) -> None:
    first = group[0]
    year = str(first["year"])
    district = first["district_or_exam"]
    paper = first["paper_type"]
    qno = BASE.q_base(first)
    material = display_material(first)
    combined_prompt = " ".join(display_prompt(e) for e in group)
    rubric_for_infer = " ".join(clean_student_rubric(e, overrides) for e in group)
    a_axis, b_axis = group_axis(group, material, combined_prompt, rubric_for_infer)

    p = doc.add_paragraph(style="FG CardTitle")
    r = p.add_run(f"{card_no:02d}. {year} · {district} · {paper} · 第{qno}题")
    set_fonts(r, FONT_EA_HEAD)

    add_table(
        doc,
        ["案件摘要", "优先入口", "答题动作"],
        [[first_sentence(material, 90), a_axis, b_axis]],
        [3900, 2600, 2860],
    )
    add_label_text(doc, "材料", material, fill="FFFFFF")
    for idx, entry in enumerate(group, start=1):
        prompt = display_prompt(entry)
        rubric = clean_student_rubric(entry, overrides)
        prefix = f"第{idx}问" if len(group) > 1 else "设问"
        add_label_text(doc, prefix, prompt, fill="F7FBFF")
        add_label_text(doc, "细则原文", rubric, fill="FFFFFF")


def add_entries(doc: Document, body_groups: list[list[dict]], overrides: dict[str, str]) -> None:
    add_heading(doc, "第二部分  真题题卡：按年份刷，按细则改写", 1)
    current_year = None
    card_no = 1
    for group in body_groups:
        year = str(group[0]["year"])
        if year != current_year:
            current_year = year
            add_heading(doc, year, 2)
        add_question_card(doc, card_no, group, overrides)
        card_no += 1


def add_appendix(doc: Document, stats, coverage_rows: list[dict], packets: list[dict]) -> None:
    doc.add_page_break()
    add_heading(doc, "附录  题源索引与补证备注", 1)
    add_callout(
        doc,
        "阅读边界",
        "本附录用于教师查看题源、细则和后续补证状态；学生正常刷题可先不看。当前非期中套卷已按新口径回源补入，期中无选必二仍单列。",
        fill="F4F6F9",
    )
    add_table(
        doc,
        ["项目", "结果"],
        [
            ["源文件", f"扫描登记 {stats.source_total} 个；本轮纳入候选 {stats.included_sources} 个。"],
            ["套卷", f"{stats.suites_total} 套；当前矩阵识别 has_xuanbier {stats.has_xuanbier} 套，no_xuanbier {stats.no_xuanbier} 套。"],
            ["题量", f"{stats.big_questions_total} 道大题，{stats.subquestions_total} 个分问。"],
            ["边界备注", f"综合法治或材料待补 {stats.pending_total} 个；正式细则待补/不完整 {stats.formal_incomplete_total} 个。"],
        ],
        [1800, 7560],
    )
    excluded = [r for r in packets if r["entry_id"] in BODY_EXCLUDE_ENTRY_IDS]
    add_heading(doc, "正文暂不展示题卡", 2)
    add_table(
        doc,
        ["出处", "处理原因"],
        [[r["title"], BODY_EXCLUDE_ENTRY_IDS[r["entry_id"]]] for r in excluded],
        [3900, 5460],
    )
    zero_rows = [r for r in coverage_rows if int(r.get("legal_subjective_count") or 0) == 0]
    add_heading(doc, "暂列无题套卷：以期中为主", 2)
    add_table(
        doc,
        ["年份", "套卷", "卷型", "当前状态"],
        [[r["year"], r["district_or_exam"], r["paper_type"], r["status"]] for r in zero_rows],
        [900, 3900, 1200, 3360],
    )
    pending = [r for r in packets if r.get("pending_reason")]
    add_heading(doc, "边界备注题卡", 2)
    add_table(
        doc,
        ["出处", "说明"],
        [[r["title"], r["pending_reason"]] for r in pending],
        [3900, 5460],
    )


def write_qa(stats, packets: list[dict], body_groups: list[list[dict]]) -> None:
    body_ids = {entry["entry_id"] for group in body_groups for entry in group}
    excluded_ids = sorted(set(BODY_EXCLUDE_ENTRY_IDS) & {entry["entry_id"] for entry in packets})
    rows = [
        "# STUDENT_BAODIAN_LAYOUT_CHECK_20260604",
        "",
        f"- Output: `{OUT_DOCX}`",
        f"- Desktop copy: `{DESKTOP_COPY}`",
        "- Design preset: compact_reference_guide with Chinese manual overrides.",
        "- Front matter: cover, student use guide, A/B quick cards before evidence appendix.",
        "- Body: grouped question cards with 案件摘要 / 优先入口 / 答题动作 / 材料 / 设问 / 细则原文.",
        "- Appendix: 题源索引、暂不展示题卡、边界备注 moved to the end.",
        f"- Body cards shown to students: {len(body_groups)}",
        f"- Body subquestions shown to students: {len(body_ids)}",
        f"- Body-excluded entry ids: {', '.join(excluded_ids) if excluded_ids else 'none'}",
        f"- Big questions in current packets: {stats.big_questions_total}",
        f"- Subquestions in current packets: {stats.subquestions_total}",
        f"- Boundary-note subquestions: {stats.pending_total}",
        "",
    ]
    QA_MD.write_text("\n".join(rows), encoding="utf-8")


def build_doc() -> None:
    packets = BASE.read_jsonl(BASE.PACKETS)
    coverage_rows = BASE.read_csv(BASE.COVERAGE)
    source_rows = BASE.read_csv(BASE.SOURCE_LEDGER)
    overrides = BASE.load_rubric_overrides()
    overrides.update(BASE.EXTRA_RUBRIC_OVERRIDES)
    stats = BASE.source_status_stats(source_rows, coverage_rows, packets)
    body_groups = make_body_groups(packets)

    doc = Document()
    style_document(doc)
    add_cover(doc, stats, body_groups)
    add_front_navigation(doc)
    add_entries(doc, body_groups, overrides)
    add_appendix(doc, stats, coverage_rows, packets)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    shutil.copy2(OUT_DOCX, DESKTOP_COPY)
    write_qa(stats, packets, body_groups)
    print(OUT_DOCX)
    print(DESKTOP_COPY)
    print(QA_MD)


if __name__ == "__main__":
    build_doc()
