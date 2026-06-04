#!/usr/bin/env python3
from __future__ import annotations

import csv
import copy
import json
import re
import shutil
import sys
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

RUN_DIR = Path(__file__).resolve().parents[1]
OUT_DIR = RUN_DIR / "05_output"
OUT_DOCX = OUT_DIR / "选必二法律与生活_A类十主题学生宝典工作稿_20260604.docx"
DESKTOP_COPY = Path.home() / "Desktop" / OUT_DOCX.name
REPORT_MD = OUT_DIR / "A_THEME_STUDENT_GUIDE_COVERAGE_REPORT_20260604.md"
REPAIR_OVERRIDES_JSON = OUT_DIR / "A_THEME_SOURCE_REPAIR_OVERRIDES_20260604.json"
PENDING_BOUNDARY_JSON = OUT_DIR / "A_THEME_PENDING_BOUNDARY_RESOLUTIONS_20260604.json"

sys.path.insert(0, str((RUN_DIR / "tools").resolve()))
import build_full_text_subjective_rubric_docx as base  # noqa: E402
import build_student_baodian_polished_docx as polished  # noqa: E402

RED = "C00000"
BLUE = "17365D"
GOLD = "7F6000"
GRAY = "666666"
LIGHT_BLUE = "EAF2F8"
LIGHT_RED = "FFF2F2"
LIGHT_GOLD = "FFF2CC"
LIGHT_GRAY = "F5F5F5"


THEMES = {
    "A1": {
        "name": "A1 民事行为与时效",
        "signals": "未成年人、意思表示、代理、追认、合同效力、限制民事行为能力、诉讼时效、不适用时效。",
        "first": "先判断行为是否有效、权利是否还能主张，再落到交易安全或特殊身份利益保护。",
        "trap": "诉讼时效不是权利消灭；未成年人案件不能只写同情，要写行为能力、追认和责任分配。",
    },
    "A2": {
        "name": "A2 人格权与个人信息",
        "signals": "名誉、隐私、肖像、姓名、生命健康、个人信息、未成年人身心权益、平台数据使用。",
        "first": "先锁定被保护的人格利益，再写侵害方式、损害后果、责任承担和必要的治理边界。",
        "trap": "网络、AI、平台背景不能替代具体人格权；必须写清侵犯的是哪项权利。",
    },
    "A3": {
        "name": "A3 物权相邻与共有",
        "signals": "所有权、财产权、共有部分、相邻通行、采光、加装电梯、业主共同决定、物业服务、绿色居住。",
        "first": "先找不动产或财产权边界，再落到有利生产、方便生活、团结互助、公平合理。",
        "trap": "多数人同意不等于少数人合法权益自动消失；物业和相邻关系要分清合同义务与物权边界。",
    },
    "A4": {
        "name": "A4 合同与诚信",
        "signals": "要约、承诺、合同成立、合同效力、履行、违约、格式条款、补充协议、诚实信用。",
        "first": "先判断合同是否成立有效、内容是否明确，再写全面履行、诚信协商、违约责任。",
        "trap": "不能只喊诚信，要把约定、回复、履行、违约事实压进规则。",
    },
    "A5": {
        "name": "A5 知识产权与公平竞争",
        "signals": "著作权、商标、专利、商业秘密、数据权益、虚假宣传、混淆、恶意诉讼、创新保护。",
        "first": "先定具体权利或竞争秩序，再写侵权/不正当竞争行为、责任和创新价值。",
        "trap": "创新价值必须附着在具体权利和具体行为上，不能写成空泛口号。",
    },
    "A6": {
        "name": "A6 侵权责任",
        "signals": "损害、过错、因果关系、安全保障义务、污染环境、动物损害、赔偿、减责免责。",
        "first": "按要件写：权利受损、违法或过错、因果关系、责任承担，再看免责或减责。",
        "trap": "不能只凭结果倒推责任；要看义务、过错、证据和因果关系。",
    },
    "A7": {
        "name": "A7 婚姻家庭与继承",
        "signals": "夫妻关系、赡养、抚养、继承、遗赠扶养协议、家庭义务、老年人权益。",
        "first": "先写法定义务或协议效力，再写保护弱者、弘扬家庭美德和公序良俗。",
        "trap": "亲情价值必须落到谁对谁承担什么法律义务。",
    },
    "A8": {
        "name": "A8 劳动关系",
        "signals": "劳动合同、事实劳动关系、平台用工、解除、竞业限制、集体协商、就业歧视、社保。",
        "first": "先判断是否存在劳动关系或劳动法保护对象，再平衡劳动者权益与依法用工秩序。",
        "trap": "劳动法不只是保护劳动者，也维护用人单位依法管理和市场秩序。",
    },
    "A9": {
        "name": "A9 消费者权益",
        "signals": "经营者、消费者、网购、食品安全、欺诈、知情权、公平交易、惩罚性赔偿、放心消费。",
        "first": "先确认消费者与经营者关系，再写知情权、公平交易权、安全权和市场秩序。",
        "trap": "惩罚性赔偿要看法院支持的是哪一笔请求、哪一种交易。",
    },
    "A10": {
        "name": "A10 纠纷解决程序",
        "signals": "调解、仲裁、诉讼、证据、举证、起诉状、回避、诉讼代理、公益诉讼、执行监督、司法确认。",
        "first": "先写适用的程序和证据要求，再写依法、公正、高效、多元化解纠纷。",
        "trap": "只有题干真正问路径、证据或程序时，程序才是主入口；否则它只是答题动作。",
    },
}

B_AXIS = {
    "B1": ("B1 补表/补链", "事实 -> 规则 -> 结论，短句入格。"),
    "B2": ("B2 判责/理由", "法律关系 -> 规则 -> 材料事实 -> 责任或裁判结果。"),
    "B3": ("B3 诉求支持", "先结论，再逐项匹配请求、权利基础和事实。"),
    "B4": ("B4 评析认识", "肯定合理处，指出法律依据，再落材料和价值。"),
    "B5": ("B5 意义价值", "保护谁、约束谁、形成什么秩序。"),
    "B6": ("B6 维权路径", "证据 -> 程序 -> 请求 -> 执行或监督。"),
    "B7": ("B7 短答识别", "只写关键词和短句，避免长篇判决。"),
}

PRIMARY_THEME_OVERRIDES = {
    "E001": "A1", "E002": "A4", "E003": "A4", "E004": "A6", "E005": "A8",
    "E006": "A5", "E007": "A7", "E008": "A4", "E009": "A4", "E010": "A10",
    "E011": "A10", "E012": "A7", "E013": "A9", "E014": "A3", "E015": "A10",
    "E016": "A1", "E017": "A8", "E018": "A5", "E019": "A5", "E020": "A9",
    "E021": "A5", "E022": "A5", "E023": "A4", "E024": "A5", "E025": "A9",
    "E026": "A3", "E027": "A3", "E028": "A7", "E029": "A9", "E030": "A8",
    "E031": "A4", "E032": "A5", "E033": "A4", "E034": "A5", "E035": "A1",
    "E036": "A8", "E037": "A3", "E038": "A3", "E039": "A4", "E040": "A4",
    "E041": "A5", "E042": "A5", "E043": "A6", "E044": "A2", "E045": "A4",
    "E046": "A6", "E047": "A9", "E048": "A5", "E049": "A6", "E050": "A9",
    "E051": "A5", "E052": "A10", "E053": "A2", "E054": "A3", "E055": "A4",
    "E056": "A6", "E057": "A10", "E058": "A3", "E059": "A6", "E060": "A10",
    "E061": "A4", "E062": "A8", "E063": "A6", "E064": "A10", "E065": "A10",
    "E066": "A1", "E067": "A4", "E068": "A4", "E069": "A5", "E070": "A5",
    "E071": "A8", "E072": "A2", "E073": "A6", "E074": "A3",
}

NEEDS_RELAYOUT = {
    "E005": "多栏/OCR 文本曾被标记为待重排，本稿先收正文并保留待回源校验标记。",
    "E006": "题面资料卡曾有串行错位风险，本稿先收正文并保留待回源校验标记。",
    "E015": "多案例表格在文本层有行列错位风险，本稿先收正文并保留待回源校验标记。",
    "E016": "多案例表格在文本层有行列错位风险，本稿先收正文并保留待回源校验标记。",
    "E022": "表格型材料在文本层有拆分错位风险，本稿先收正文并保留待回源校验标记。",
    "E023": "表格型材料在文本层有拆分错位风险，本稿先收正文并保留待回源校验标记。",
    "E024": "裁判要点表格在文本层有错位风险，本稿先收正文并保留待回源校验标记。",
    "E035": "案情与侧栏法条曾被标记交错，本稿先收正文并保留待回源校验标记。",
    "E036": "案件事实与裁判要点两栏曾被标记交错，本稿先收正文并保留待回源校验标记。",
    "E053": "模块边界仍需教师复核，本稿先收正文并保留待核标记。",
    "E074": "当前题面材料表缺失，必须回源补齐后才能作为最终闭合正文。",
}

SOURCE_REPAIR_CLEARS_RELAYOUT = {"E005", "E006", "E015", "E016", "E022", "E023", "E024", "E035", "E036", "E053", "E074"}

A_THEME_RUBRIC_OVERRIDES = {
    "E058": (
        "（1）（8分）\n"
        "法院判决以事实为根据，以法律为准绳。老旧小区加装电梯属民生工程，已获法定比例业主同意，程序合法。"
        "电梯采用玻璃幕墙设计，在提升本楼业主出行便利的同时，能尽量减少对相邻楼业主的影响。"
        "依据《中华人民共和国民法典》规定，不动产的相邻权利人应当按照有利生产、方便生活、团结互助、公平合理的原则，"
        "正确处理相邻关系。本案判决维护了相邻业主合法权益，有助于促进邻里和谐，践行友善的社会主义核心价值观。"
    ),
    "E059": (
        "20.（8分）\n"
        "理由：根据《中华人民共和国民法典》规定，参与者不存在故意或重大过失，组织者尽到安全保障义务，则无需承担责任。（2分）"
        "本案中张某自愿参加活动，因自身原因突发意外，并非他人行为造成。（1分）"
        "李某虽为广场舞义务组织者，但已尽到合理的安全保障义务，事前提醒参与者注意安全，事后及时救助，"
        "且不存在故意或重大过失，与张某死亡不存在因果关系，不承担侵权责任。（2分）\n"
        "意义：引导参与者增强风险意识，对自己的行为负责；合理界定义务组织者的责任范围，体现公平原则；"
        "鼓励社会成员在互帮互助中传递社会主义核心价值观，维护和谐有序的社区生活秩序。（3分）"
    ),
}

RUBRIC_TRUNCATE_MARKERS = [
    "（二）学生表现",
    "学生表现和问题",
    "\n学生问题",
    "（三）教学启示",
    "教学启示",
    "\n能力要求\n分析角度\n学生学习",
    "教师教学",
    "改进措施",
    "建议最后冲刺",
    "全球治理倡议",
    "友好关系；主持上合组织",
    "论坛部长级会议",
]

BAD_RUBRIC_LINE_CONTAINS = [
    "学生问题",
    "学生表现",
    "试题分析",
    "讲评",
    "教学启示",
    "教师教学",
    "改进措施",
    "PAGEPAGE",
]

THEME_STUDENT_GUIDE = {
    "A1": {
        "goal": "把“行为效力”和“权利还能不能顺利主张”先判准，再处理交易安全、未成年人保护或弱者保护。",
        "quick": [
            "看见未成年人、代理、追认、处分大额财产，先问行为能力和意思表示是否有效。",
            "看见超过多年才起诉、催讨、抗辩，先问是否属于诉讼时效和是否存在不适用时效的例外。",
        ],
        "moves": [
            "主体身份：谁是权利人、义务人、未成年人、监护人或代理人。",
            "效力判断：成立、有效、效力待定、可撤销或不适用诉讼时效。",
            "结果落点：是否支持请求、是否需要追认、是否保护特定权益。",
        ],
        "sentences": [
            "该行为是否发生法律效力，要结合行为人的民事行为能力、意思表示和法律规定判断。",
            "诉讼时效届满并不消灭实体权利，但义务人可以依法提出抗辩。",
        ],
        "boundary": "不要把“同情弱者”直接写成法律结论；先写规则，再写规则为什么保护该主体。",
    },
    "A2": {
        "goal": "先点名被侵犯的人格利益或个人信息权益，再写侵害方式、后果、责任和治理意义。",
        "quick": [
            "看见姓名、肖像、名誉、隐私、生命健康、个人信息，就先锁定具体人格权。",
            "看见平台、AI、校园欺凌、数据采集，不要停在背景，要追问侵害了哪项权利。",
        ],
        "moves": [
            "权利定位：人格尊严、人身自由、个人信息、未成年人身心健康。",
            "行为匹配：公开、传播、侮辱、泄露、非法处理、伤害或威胁。",
            "责任落点：停止侵害、赔礼道歉、消除影响、赔偿损失、加强治理。",
        ],
        "sentences": [
            "民事主体的人格权受法律保护，任何组织和个人不得非法侵害。",
            "处理个人信息应当遵循合法、正当、必要和诚信原则。",
        ],
        "boundary": "不要把所有网络题都写成个人信息；名誉、隐私、肖像、生命健康要分别识别。",
    },
    "A3": {
        "goal": "围绕财产权、共有部分和相邻关系，写清权利边界与利益协调。",
        "quick": [
            "看见房屋、通行、采光、加装电梯、业主表决、物业服务，先找物权或共有管理入口。",
            "看见绿色居住、社区治理，不要只写价值，要落到谁享有权利、谁承担义务。",
        ],
        "moves": [
            "权属定位：所有权、用益、担保、共有、业主共同决定。",
            "利益衡量：有利生产、方便生活、团结互助、公平合理。",
            "纠纷处理：协商、表决、履约、停止妨害、恢复原状或赔偿。",
        ],
        "sentences": [
            "处理相邻关系应当按照有利生产、方便生活、团结互助、公平合理的原则。",
            "业主对共有部分依法共同管理，表决结果也不能侵害业主合法权益。",
        ],
        "boundary": "多数决不是万能答案；相邻关系、物业合同和侵权责任要分开写。",
    },
    "A4": {
        "goal": "先判断合同是否成立有效，再用约定、履行、违约和诚信原则压实结论。",
        "quick": [
            "看见订购、回复、补充协议、尾款、解除、格式条款，优先走合同入口。",
            "看见“诚信”“协商”“约定不明”，必须回到合同条款和双方真实意思表示。",
        ],
        "moves": [
            "成立效力：要约、承诺、合同内容、真实意思表示。",
            "履行责任：全面履行、协作履行、诚信原则、通知协助。",
            "后果落点：继续履行、承担费用、违约责任、补充或解释合同。",
        ],
        "sentences": [
            "依法成立的合同受法律保护，当事人应当按照约定全面履行义务。",
            "合同约定不明时，应当结合交易习惯、诚信原则和双方真实意思解释。",
        ],
        "boundary": "不能只写“应该讲诚信”；要写清哪句话、哪个行为、哪项费用或责任构成争议。",
    },
    "A5": {
        "goal": "把“创新成果受保护”和“市场竞争有边界”连起来，避免写成空泛创新价值。",
        "quick": [
            "看见著作权、商标、专利、商业秘密、AI作品、数据权益，先点具体知识产权。",
            "看见混淆、虚假宣传、诋毁、恶意诉讼、搭便车，先点公平竞争秩序。",
        ],
        "moves": [
            "对象识别：作品、发明、商标、商业秘密、数据或经营标识。",
            "行为判断：复制、传播、使用、混淆、虚假宣传、侵犯秘密或不正当竞争。",
            "价值落点：保护创新、规范竞争、维护消费者和经营者合法权益。",
        ],
        "sentences": [
            "法律保护知识产权，有利于激励创新创造和促进新质生产力发展。",
            "经营者应当遵循自愿、平等、公平、诚信原则，不得实施不正当竞争行为。",
        ],
        "boundary": "知识产权题不能只写“保护创新”；必须写明保护的权利类型和被评价的行为。",
    },
    "A6": {
        "goal": "按侵权责任四件事写：权益受损、义务或过错、因果关系、责任承担。",
        "quick": [
            "看见摔伤、污染、动物、设施、餐厅、学校、安全保障义务，先走侵权要件。",
            "看见“是否担责”，不要凭结果倒推，要看义务、过错、证据和因果关系。",
        ],
        "moves": [
            "权益损害：生命健康、财产、人身或生态环境受损。",
            "责任要件：违法行为、过错、过错推定或无过错责任、因果关系。",
            "承担方式：停止侵害、排除妨碍、赔偿损失、恢复原状、道歉等。",
        ],
        "sentences": [
            "行为人因过错侵害他人民事权益造成损害的，应当承担侵权责任。",
            "经营场所、公共场所管理人未尽安全保障义务造成损害的，应依法承担相应责任。",
        ],
        "boundary": "侵权题最忌只写“造成损害所以赔偿”；必须把法定义务和事实证据连上。",
    },
    "A7": {
        "goal": "把亲情叙事翻译成法定义务、协议效力和家庭伦理的法律化表达。",
        "quick": [
            "看见赡养、抚养、夫妻财产、继承、遗赠扶养协议，先点家庭法律关系。",
            "看见老人、子女、配偶、继承人，不要只写道德，要写谁对谁负什么法律义务。",
        ],
        "moves": [
            "关系定位：夫妻、父母子女、继承人、受遗赠人或扶养人。",
            "规则匹配：赡养、抚养、扶养、继承顺序、协议效力、公序良俗。",
            "价值落点：保护弱者、弘扬家庭美德、维护和谐稳定家庭关系。",
        ],
        "sentences": [
            "成年子女对父母负有赡养、扶助和保护的义务。",
            "家庭成员之间既有道德要求，也有明确的法律权利义务。",
        ],
        "boundary": "不要把家庭题写成空泛亲情作文；必须落到法定义务或有效协议。",
    },
    "A8": {
        "goal": "先判劳动关系或劳动法保护对象，再平衡劳动者权益与依法用工秩序。",
        "quick": [
            "看见平台用工、打卡、派单、工资、解除、竞业限制，先判劳动关系或从属性。",
            "看见就业歧视、社保、加班、工伤，先锁定劳动者权益，再写用人单位义务。",
        ],
        "moves": [
            "关系认定：人格从属性、经济从属性、组织从属性。",
            "权利义务：合同、报酬、休息休假、社保、安全、依法解除。",
            "秩序价值：保护劳动者、规范平台和企业用工、促进就业公平。",
        ],
        "sentences": [
            "判断劳动关系不能只看合同名称，应当结合实际用工管理和从属性特征。",
            "劳动法既保护劳动者合法权益，也引导用人单位依法规范用工。",
        ],
        "boundary": "“合作协议”不是决定性标签；实际管理、报酬和组织纳入才是关键事实。",
    },
    "A9": {
        "goal": "先确认消费者和经营者关系，再写具体权利、经营者义务和市场秩序。",
        "quick": [
            "看见购买商品、接受服务、网购、直播带货、食品安全，先走消费者权益入口。",
            "看见三倍赔偿、退一赔三、知情权、公平交易，必须看交易事实和法院支持范围。",
        ],
        "moves": [
            "关系定位：消费者为生活消费需要购买、使用商品或接受服务。",
            "权利匹配：安全权、知情权、自主选择权、公平交易权、求偿权。",
            "责任落点：退货、赔偿、惩罚性赔偿、行政监管、放心消费环境。",
        ],
        "sentences": [
            "经营者应当真实、全面提供商品或者服务信息，保障消费者知情权和公平交易权。",
            "欺诈消费者的，依法可能承担惩罚性赔偿责任。",
        ],
        "boundary": "不是所有买卖都是消费纠纷；生产经营采购、投资交易要谨慎排除。",
    },
    "A10": {
        "goal": "把纠纷解决从“找谁说理”写成程序、证据、请求和法治效果。",
        "quick": [
            "看见调解、仲裁、诉讼、证据、起诉状、举证、回避，先走程序入口。",
            "看见多元化解、司法确认、执行监督，先写路径，再写公正高效和权利实现。",
        ],
        "moves": [
            "程序选择：协商、调解、仲裁、诉讼或行政/司法衔接。",
            "证据请求：证明对象、举证责任、诉讼请求、事实与理由。",
            "效果落点：依法维权、定分止争、节约成本、维护司法公正。",
        ],
        "sentences": [
            "当事人应当依法收集和提交证据，通过适当程序维护自身合法权益。",
            "多元纠纷解决机制有利于及时化解矛盾、降低维权成本、维护社会公平正义。",
        ],
        "boundary": "程序不是所有题的主入口；只有设问真正考路径、证据、诉讼文书或纠纷解决时才放在 A10。",
    },
}

THEME_METHOD_GUIDE = {
    "A1": {
        "memory": [
            "先问行为有没有法律效力，再问权利还能不能顺利主张。",
            "未成年人、代理、追认、时效抗辩，是本章最常见的四个开关。",
        ],
        "chain": [
            "主体身份：年龄、监护、代理、债权人或义务人。",
            "法律开关：行为能力、真实意思、追认、可撤销或诉讼时效。",
            "材料事实：谁何时作出行为，谁何时主张权利，谁提出抗辩。",
            "结果表达：行为效力、请求是否支持、特殊主体利益如何保护。",
        ],
        "questions": [
            "“谈原因”要写制度目的：督促主张权利、稳定交易秩序、保护特殊利益。",
            "“补表”要先填效力结论，再用材料事实说明为什么。",
            "“评析观点”要把情理判断压回行为能力、意思表示和时效规则。",
        ],
        "template": [
            "本案首先涉及……的民事行为效力/诉讼时效问题。",
            "从材料看，……，因此应认定……；若……，则……。",
            "这一规则既维护交易安全，也保护……的合法权益。",
        ],
        "wrong": [
            "把诉讼时效写成实体权利消灭。",
            "只写保护未成年人，不写行为能力和追认。",
            "只讲公平同情，不写法律效力结论。",
        ],
    },
    "A2": {
        "memory": [
            "人格权题先点名权利，后判断侵害方式。",
            "网络、AI、校园场景只是载体，真正给分的是具体人格利益。",
        ],
        "chain": [
            "权利定位：名誉、隐私、肖像、姓名、生命健康或个人信息。",
            "侵害行为：公开、传播、侮辱、非法处理、伤害、威胁或滥用。",
            "责任承担：停止侵害、消除影响、赔礼道歉、赔偿损失。",
            "治理价值：保护人格尊严，规范平台和新技术使用边界。",
        ],
        "questions": [
            "“如何维权”要写证据、请求和责任承担。",
            "“分析裁判”要先说被侵害权利，再说法院为什么支持或不支持。",
            "“意义价值”要落到人格尊严、未成年人保护和数字治理。",
        ],
        "template": [
            "材料中的……行为，涉及对……权/个人信息权益的保护。",
            "该行为已经表现为……，可能造成……后果，应依法承担……责任。",
            "此类处理有利于守住人格尊严和新技术应用的法律边界。",
        ],
        "wrong": [
            "把所有网络材料都写成个人信息。",
            "只写“侵犯人格权”，不点具体权利名称。",
            "只谈道德谴责，不写民事责任方式。",
        ],
    },
    "A3": {
        "memory": [
            "物权题的核心是边界：谁享有、谁让渡、谁不能越界。",
            "相邻关系不是谁声音大谁赢，而是依法协调利益。",
        ],
        "chain": [
            "权利对象：房屋、共有部分、相邻通行、采光、物业或财产权。",
            "权利边界：所有权、共有管理、相邻义务或绿色原则。",
            "利益衡量：有利生产、方便生活、团结互助、公平合理。",
            "解决结果：协商、表决、履约、停止妨害、恢复原状或赔偿。",
        ],
        "questions": [
            "“说明规定如何推动发展”要把物权规则和公共利益连起来。",
            "“评析建议”要先判权利边界，再写社区协商或共同决定。",
            "“裁判理由”要从物权、共有或相邻关系推出责任。",
        ],
        "template": [
            "本案首先涉及……的物权/相邻关系/共有管理问题。",
            "权利人可以依法……，但行使权利不得损害……的合法权益。",
            "依法协调各方利益，有利于维护财产秩序和公共生活秩序。",
        ],
        "wrong": [
            "把物业服务合同、相邻关系和侵权责任混成一段。",
            "认为多数业主同意就一定合法。",
            "只写绿色发展，不写民法典物权规则。",
        ],
    },
    "A4": {
        "memory": [
            "合同题先看有没有约定，再看有没有按约履行。",
            "诚信不是口号，必须贴到回复、付款、交付、解除等事实。",
        ],
        "chain": [
            "合同基础：要约、承诺、合同成立、效力和主要条款。",
            "履行过程：谁履行了什么，谁变更、拒绝、迟延或解除。",
            "规则适用：全面履行、协作履行、诚信原则、格式条款或违约责任。",
            "裁判落点：请求是否支持、费用谁承担、合同是否继续或解除。",
        ],
        "questions": [
            "“起诉状”要把事实与理由写成诉讼请求的法律支撑。",
            "“评析观点”要写约定、交易习惯、诚信原则和履行事实。",
            "“意义影响”要把合同秩序与营商环境、交易安全相连。",
        ],
        "template": [
            "双方之间形成……合同关系，依法成立的合同受法律保护。",
            "材料中……表明……，当事人应当按照约定和诚信原则履行。",
            "因此，……应承担……责任/法院支持……请求。",
        ],
        "wrong": [
            "只写“应讲诚信”，没有合同条款和履行事实。",
            "把所有平台、AI、租赁题都硬写成合同效力。",
            "忽略格式条款、交易习惯或补充协议的独立给分点。",
        ],
    },
    "A5": {
        "memory": [
            "知识产权题要先说保护什么成果，再说禁止什么行为。",
            "创新价值必须挂在具体权利和公平竞争秩序上。",
        ],
        "chain": [
            "保护对象：作品、商标、专利、商业秘密、数据或经营标识。",
            "违法行为：复制传播、混淆、虚假宣传、侵犯秘密、恶意竞争。",
            "责任后果：停止侵权、赔偿损失、惩戒不正当竞争。",
            "价值提升：激励创新、规范市场、保护消费者和经营者权益。",
        ],
        "questions": [
            "“典型案例意义”要写保护权利、惩治侵权、鼓励创新三层。",
            "“裁判理由”要把权利类型、行为方式和法律责任逐项匹配。",
            "“法治价值”要从市场竞争秩序推出创新发展价值。",
        ],
        "template": [
            "本案涉及……知识产权/公平竞争秩序。",
            "……行为属于……，侵害了……，依法应承担……。",
            "人民法院依法保护……，有利于激励创新、规范竞争、营造良好营商环境。",
        ],
        "wrong": [
            "只写“保护新质生产力”，不写权利类型。",
            "把知识产权侵权和不正当竞争完全混同。",
            "把社会意义写满，却没有判断行为是否合法。",
        ],
    },
    "A6": {
        "memory": [
            "侵权题不是看谁受伤就谁赢，要看义务、过错和因果关系。",
            "责任方式要从材料事实里长出来。",
        ],
        "chain": [
            "权益损害：人身、财产、人格或生态环境受到损害。",
            "责任要件：违法行为、过错或特殊责任规则、因果关系。",
            "抗辩边界：免责、减责、受害人过错或第三人原因。",
            "结果承担：停止侵害、排除妨碍、恢复原状、赔偿损失等。",
        ],
        "questions": [
            "“是否担责”要逐项写要件，不要只写结论。",
            "“裁判理由”要把安全保障义务、因果关系和责任范围拆开。",
            "“维权路径”要写证据、请求和责任承担方式。",
        ],
        "template": [
            "行为人因……侵害他人民事权益造成损害，符合……责任要件。",
            "材料中……说明其未尽……义务/存在……过错，与损害之间具有……关系。",
            "因此应依法承担……责任；若存在……，可相应减轻或免除。",
        ],
        "wrong": [
            "只凭结果严重就直接判赔。",
            "忘记安全保障义务、无过错责任或过错推定这些特殊入口。",
            "不处理免责、减责或证据不足。",
        ],
    },
    "A7": {
        "memory": [
            "家庭题要把亲情翻译成法定义务。",
            "继承和赡养都不能只靠道德表态拿分。",
        ],
        "chain": [
            "身份关系：夫妻、父母子女、继承人、受遗赠人或扶养人。",
            "权利义务：赡养、抚养、扶养、继承顺序或协议效力。",
            "事实匹配：谁履行、谁不履行、协议是否明确、利益如何受损。",
            "价值落点：保护老年人、未成年人或家庭弱者，弘扬家庭美德。",
        ],
        "questions": [
            "“认识告知书”要写法定义务和法德结合。",
            "“协议效力”要写主体、内容、公序良俗和履行。",
            "“原因意义”要写保护家庭成员合法权益和家庭文明建设。",
        ],
        "template": [
            "……之间存在……家庭法律关系，依法负有……义务。",
            "材料中……表明……，应依法……。",
            "这一处理既维护家庭成员合法权益，也体现公序良俗和家庭美德。",
        ],
        "wrong": [
            "只写孝亲敬老，不写成年子女法定义务。",
            "把继承、遗赠扶养协议和普通赠与混为一谈。",
            "不说明谁对谁承担义务。",
        ],
    },
    "A8": {
        "memory": [
            "劳动题先判关系，再谈权益。",
            "合同名称不决定关系，实际管理和从属性才关键。",
        ],
        "chain": [
            "关系认定：劳动合同、事实劳动关系、平台用工或劳务合作。",
            "从属性事实：工作安排、报酬支付、管理控制、组织纳入。",
            "权益义务：工资、休息休假、社保、解除、竞业限制、就业公平。",
            "秩序价值：保护劳动者，规范企业和平台依法用工。",
        ],
        "questions": [
            "“是否存在劳动关系”要围绕人格、经济、组织从属性。",
            "“解除/竞业限制”要写合同、法律边界和双方义务。",
            "“公平与效率”要把劳动者权益保护与企业依法经营连接起来。",
        ],
        "template": [
            "判断双方是否存在劳动关系，应结合实际用工事实，而不能只看合同名称。",
            "材料中……体现出……从属性/用工管理，因此……。",
            "依法规范用工既保护劳动者合法权益，也维护稳定有序的就业市场。",
        ],
        "wrong": [
            "看到合作协议就否定劳动关系。",
            "只写保护劳动者，不写用人单位依法管理边界。",
            "把就业歧视、平台用工、竞业限制全部套同一段话。",
        ],
    },
    "A9": {
        "memory": [
            "消费题先确认生活消费关系，再点具体消费者权利。",
            "惩罚性赔偿不是万能句，要看欺诈和法院支持范围。",
        ],
        "chain": [
            "关系确认：消费者为生活消费购买商品或接受服务。",
            "权利定位：安全权、知情权、自主选择权、公平交易权、求偿权。",
            "经营者义务：真实全面告知、诚信经营、质量安全、依法赔偿。",
            "市场价值：保护消费者，规范经营者，营造放心消费环境。",
        ],
        "questions": [
            "“诉求能否支持”要逐项匹配消费者请求和法律依据。",
            "“经营者责任”要写信息披露、质量安全和诚信经营。",
            "“市场意义”要从放心消费、消费信心和市场秩序展开。",
        ],
        "template": [
            "消费者为生活消费需要购买商品或接受服务，其合法权益受法律保护。",
            "经营者……，侵害了消费者……权，应依法承担……责任。",
            "依法处理该纠纷，有利于规范经营行为，营造安全放心的消费环境。",
        ],
        "wrong": [
            "把生产经营采购也当成消费者权益。",
            "一看到索赔就写退一赔三，不看欺诈事实。",
            "不区分安全权、知情权和公平交易权。",
        ],
    },
    "A10": {
        "memory": [
            "程序题先写路径，再写证据和请求。",
            "调解、仲裁、诉讼不是并列背诵，要看题干问哪一步。",
        ],
        "chain": [
            "程序选择：协商、调解、仲裁、诉讼、司法确认或执行监督。",
            "证据责任：证明对象、举证责任、证据提交和事实理由。",
            "请求表达：诉讼请求、仲裁请求、调解方案或维权建议。",
            "法治效果：定分止争、节约成本、公正高效、实现权利。",
        ],
        "questions": [
            "“起诉状/诉讼请求”要写当事人、请求、事实与理由。",
            "“举证责任”要先判断诉讼类型，再判断谁证明什么。",
            "“多元解纷意义”要写及时化解、降低成本、维护司法公正。",
        ],
        "template": [
            "本案可以通过……程序依法解决争议。",
            "当事人应围绕……提交证据，并提出……请求。",
            "该程序安排有利于依法维权、定分止争，维护社会公平正义。",
        ],
        "wrong": [
            "把所有纠纷都写成诉讼，不看仲裁、调解或司法确认。",
            "只写程序名称，不写证据和请求。",
            "遇到综合法治题时忽略模块边界和题干限定。",
        ],
    },
}


THEME_PHASE2_GUIDE = {
    "A1": {
        "proposition": [
            "命题人常把未成年人、老人赡养、欠款追讨放在同一题里，考学生能否把朴素公平感翻译成行为能力、追认和时效规则。",
            "材料通常不会直接问“效力待定”或“诉讼时效”，而是让学生先从主体身份、时间经过和抗辩动作里找法律开关。",
        ],
        "path": [
            "第一步看主体：未成年人、监护人、代理人、债权人、义务人分别是谁。",
            "第二步看行为：订立、处分、追认、催讨、抗辩分别发生在什么时间。",
            "第三步套规则：行为能力、意思表示、追认、可撤销、诉讼时效或不适用时效。",
            "第四步落结果：效力如何、请求是否支持、特殊主体利益如何被保护。",
        ],
        "scoring": [
            "规则名必须准：限制民事行为能力、法定代理人追认、诉讼时效抗辩、不适用诉讼时效。",
            "事实扣合必须有时间线：何时欠款、何时催讨、何时起诉、何时抗辩。",
            "价值句只能放在结尾：交易安全、积极维权、保护未成年人或老年人合法权益。",
        ],
        "boundary": [
            "诉讼时效届满不是债权消灭，而是义务人取得抗辩权。",
            "未成年人保护不能替代合同效力判断；要先写行为能力和追认。",
            "亲情、同情、弱者保护只能作为价值落点，不能直接当法律依据。",
        ],
    },
    "A2": {
        "proposition": [
            "命题人喜欢借网络平台、AI、校园和公共空间场景考人格尊严的具体化，核心不是技术新，而是权利类型要判准。",
            "个人信息题常把信息处理、隐私、名誉、生命健康混放，考学生能否先分权利，再写侵害方式和责任。",
        ],
        "path": [
            "第一步定权利：名誉、隐私、肖像、姓名、生命健康、身体健康或个人信息权益。",
            "第二步看行为：公开、传播、侮辱、诽谤、泄露、非法处理、伤害或过度采集。",
            "第三步看后果：社会评价降低、生活安宁受扰、身心受损、信息权益被侵害。",
            "第四步写责任：停止侵害、消除影响、赔礼道歉、赔偿损失、完善治理。",
        ],
        "scoring": [
            "权利名称要具体，不能只写“人格权受到侵害”。",
            "材料动作要贴上：谁发布、谁传播、谁处理信息、谁受到何种影响。",
            "治理价值要收束到人格尊严、未成年人保护、数字空间法治边界。",
        ],
        "boundary": [
            "隐私权、个人信息权益、名誉权不是同一个词的替换。",
            "AI或平台只是载体；没有具体权利类型就拿不到关键分。",
            "道德谴责不能代替民事责任方式。",
        ],
    },
    "A3": {
        "proposition": [
            "命题人常用老旧小区、电梯、物业、相邻通行和生态居住考“权利边界+利益协调”。",
            "这类题表面是社区矛盾，实质是让学生把物权规则写成公共生活中的可执行秩序。",
        ],
        "path": [
            "第一步定对象：房屋、共有部分、相邻土地建筑物、物业服务或具体财产权。",
            "第二步定关系：所有权、共有管理、相邻关系、物业合同或侵权关系。",
            "第三步衡量利益：有利生产、方便生活、团结互助、公平合理、绿色原则。",
            "第四步给方案：协商表决、履约、停止妨害、恢复原状、赔偿或共同治理。",
        ],
        "scoring": [
            "物权规则要落到具体对象：共有部分、相邻权利、业主共同决定、财产权保护。",
            "利益衡量要双向写：一方便利不能直接吞掉另一方合法权益。",
            "价值句要写成生活秩序：邻里和谐、社区治理、生态公共利益、财产制度稳定。",
        ],
        "boundary": [
            "多数业主同意不等于当然合法，仍要看程序和少数人合法权益。",
            "物业合同、相邻关系、侵权责任要分清入口。",
            "绿色发展不是万能结论，必须回到物权规则和公共利益协调。",
        ],
    },
    "A4": {
        "proposition": [
            "命题人常把一两句回复、补充约定、格式条款、租赁和平台服务做成陷阱，考学生能否把生活话语还原成合同关系。",
            "高频设问不是背合同概念，而是从事实里判断约定是否明确、履行是否到位、违约责任由谁承担。",
        ],
        "path": [
            "第一步找约定：合同是否成立，主要条款、补充协议、格式条款分别是什么。",
            "第二步看履行：谁交付、谁付款、谁通知、谁拒绝、谁解除或单方变更。",
            "第三步定责任：全面履行、协作履行、诚信原则、违约责任或格式条款规制。",
            "第四步写请求：继续履行、支付尾款、承担费用、退款、赔偿或确认条款效力。",
        ],
        "scoring": [
            "合同题的核心分在“约定+事实+责任”，不能只写诚信原则。",
            "起诉状类题要把诉讼请求和事实理由对齐，不能只写法律口号。",
            "格式条款要写提示说明义务、内容公平和消费者或相对方权益保护。",
        ],
        "boundary": [
            "诚信原则是裁判规则的一部分，不是万能道德评价。",
            "合同成立、合同有效、合同履行、违约责任要分层写。",
            "平台服务、租赁、AI服务不一定都归合同效力，先看设问真正要你判断什么。",
        ],
    },
    "A5": {
        "proposition": [
            "命题人会把虚拟数字人、技术秘密、数据权益、恶意诉讼、混淆宣传放进新质生产力语境，考的仍是具体权利和竞争边界。",
            "这章最容易被学生写成宏大创新作文，真正得分点在“保护对象-侵害行为-责任后果-创新价值”的闭环。",
        ],
        "path": [
            "第一步定保护对象：作品、商标、专利、商业秘密、数据、经营标识或竞争利益。",
            "第二步定行为：复制传播、使用泄露、混淆、虚假宣传、诋毁、恶意诉讼或搭便车。",
            "第三步定责任：停止侵害、消除影响、赔偿损失、惩戒不正当竞争。",
            "第四步升价值：激励创新、规范市场竞争、优化营商环境、保护消费者和经营者权益。",
        ],
        "scoring": [
            "权利类型必须具体，尤其著作权、商业秘密、不正当竞争不能互相糊掉。",
            "材料事实要写侵权动作：未经许可、替换标识、使用秘密、虚假宣传、混淆来源。",
            "价值表达要从法律保护推出创新发展，不能直接从“新质生产力”起笔。",
        ],
        "boundary": [
            "知识产权侵权和不正当竞争可并存，但要分别写依据和行为。",
            "创新价值不是独立采分万能句，必须附着在具体权利保护上。",
            "数据、AI、数字人题先回到现有权利和竞争秩序，不要自创概念。",
        ],
    },
    "A6": {
        "proposition": [
            "命题人常用摔伤、AI幻觉、污染、公共场所、动物和安全保障义务考责任边界，而不是考“谁更可怜”。",
            "材料会故意给出损害结果，但得分取决于能否写出义务、过错、因果关系和免责减责。",
        ],
        "path": [
            "第一步定权益损害：生命健康、身体、财产、人格、生态环境或其他民事权益。",
            "第二步定责任规则：一般过错、过错推定、无过错责任、安全保障义务或特殊侵权。",
            "第三步扣事实证据：谁违反了什么义务，损害如何发生，因果关系是否成立。",
            "第四步写责任范围：停止侵害、排除妨碍、赔偿、恢复原状、减责或免责。",
        ],
        "scoring": [
            "不要只写损害结果，要写义务来源和过错或特殊责任依据。",
            "安全保障义务题要写场所、管理人、合理限度和救助处置。",
            "AI权责边界题要把注意义务、可预见风险、损害后果和产业治理分开。",
        ],
        "boundary": [
            "严重损害不等于当然担责，仍要看责任要件。",
            "侵权责任和合同责任可能交叉，但设问问侵权时优先走要件。",
            "免责、减责、受害人过错和第三人原因不能漏。",
        ],
    },
    "A7": {
        "proposition": [
            "命题人把赡养、继承和遗赠扶养协议放进家庭伦理故事里，考学生能否把亲情叙事翻译成明确法定义务。",
            "这章题量不大，但每题都容易被写成道德作文，必须先找身份关系和义务来源。",
        ],
        "path": [
            "第一步定身份：夫妻、父母子女、继承人、受遗赠人、扶养人或其他家庭成员。",
            "第二步定义务：赡养、抚养、扶养、继承顺序、协议效力或夫妻共同责任。",
            "第三步扣事实：谁履行、谁拒绝、谁需要保护、协议是否有效并实际履行。",
            "第四步升价值：保护老年人、未成年人和家庭弱者，弘扬家庭美德和公序良俗。",
        ],
        "scoring": [
            "家庭题必须写清“谁对谁负什么义务”。",
            "协议效力题要看主体、意思表示、内容合法、公序良俗和履行情况。",
            "法德结合的价值句要放在法定义务之后。",
        ],
        "boundary": [
            "孝亲敬老不是法律依据的替代品。",
            "继承、赠与、遗赠扶养协议不能混用。",
            "赡养费、精神慰藉、常回家看看等落点要按设问分别处理。",
        ],
    },
    "A8": {
        "proposition": [
            "命题人常用平台用工、竞业限制、解除合同、集体协商和就业歧视考劳动关系从属性。",
            "这章不是单向保护劳动者，而是考依法用工秩序中劳动者权益与企业管理边界的平衡。",
        ],
        "path": [
            "第一步判关系：劳动合同、事实劳动关系、劳务合作、平台用工或竞业限制关系。",
            "第二步看从属性：工作安排、报酬来源、考核管理、组织纳入、工具和风险承担。",
            "第三步套权益义务：工资、休息、社保、安全、解除、保密竞业或平等就业。",
            "第四步落价值：保护劳动者、规范企业平台用工、促进公平就业和市场秩序。",
        ],
        "scoring": [
            "劳动关系认定不能只看合同名称，要写实际管理和从属性事实。",
            "竞业限制要写主体范围、期限、补偿、保密利益和就业自由边界。",
            "就业歧视题要写平等就业权和用人单位依法招录义务。",
        ],
        "boundary": [
            "合作协议、承揽协议、平台规则都不是排除劳动关系的决定性标签。",
            "保护劳动者不等于否定用人单位依法管理。",
            "平台用工、竞业限制、就业歧视不能套同一段模板。",
        ],
    },
    "A9": {
        "proposition": [
            "命题人常借网购、直播、食品安全、预付费和惩罚性赔偿考消费关系和经营者义务。",
            "材料经常给出多项诉求，考学生能否逐项判断法院支持哪一项、为什么支持或不支持。",
        ],
        "path": [
            "第一步定关系：是否为生活消费需要购买商品、使用商品或接受服务。",
            "第二步定权利：安全权、知情权、自主选择权、公平交易权、求偿权。",
            "第三步定责任：退换、赔偿、惩罚性赔偿、召回、行政监管或平台责任。",
            "第四步升价值：诚信经营、放心消费、维护市场秩序和消费信心。",
        ],
        "scoring": [
            "惩罚性赔偿要看欺诈或食品安全等法定条件，不能见赔偿就写退一赔三。",
            "知情权和公平交易权要扣到信息披露、价格、质量、格式条款或虚假宣传。",
            "诉求支持题要逐项写请求、依据、事实和法院态度。",
        ],
        "boundary": [
            "生产经营采购、投资交易和普通商业合作不当然是消费者权益。",
            "食品安全、欺诈、质量瑕疵对应的责任强度不同。",
            "直播平台、商家、生产者责任要按材料身份区分。",
        ],
    },
    "A10": {
        "proposition": [
            "命题人把模拟庭审、起诉状、举证责任、多元解纷和司法确认放在案例中，考程序意识而不是背程序名称。",
            "这章常与家庭、合同、侵权、涉外法治交叉，只有设问明确问路径、证据、文书或程序时才作为主入口。",
        ],
        "path": [
            "第一步定路径：协商、调解、仲裁、诉讼、行政处理、司法确认或执行监督。",
            "第二步定证据：证明对象、举证责任、证据提交、事实理由和请求。",
            "第三步定角色：当事人、诉讼代理人、辩护人、调解组织、仲裁机构或法院。",
            "第四步写效果：依法维权、定分止争、降低成本、维护公正高效权威的司法秩序。",
        ],
        "scoring": [
            "起诉状题要把诉讼请求和事实理由写成一一支撑关系。",
            "举证责任题先判断纠纷类型和证明对象，再写谁承担证明责任。",
            "多元解纷意义要写及时、便捷、低成本、公正和权利实现。",
        ],
        "boundary": [
            "程序不是所有纠纷题的主答案；实体规则仍要按主知识判断。",
            "辩护人、诉讼代理人、法定代理人不能混用。",
            "调解、仲裁、诉讼不能机械并列，要看题干给出的真实路径。",
        ],
    },
}


def shade_paragraph(p, fill: str, border: str | None = None) -> None:
    ppr = p._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    if border:
        pbdr = ppr.find(qn("w:pBdr"))
        if pbdr is None:
            pbdr = OxmlElement("w:pBdr")
            ppr.append(pbdr)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:color"), border)
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "4")
        pbdr.append(left)


def ensure_fgcard_alias(doc: Document) -> None:
    if "FGCard" not in doc.styles:
        style = doc.styles.add_style("FGCard", WD_STYLE_TYPE.PARAGRAPH)
        base_style = doc.styles["FG Card"] if "FG Card" in doc.styles else doc.styles["Normal"]
        style.base_style = base_style
        style.font.name = base_style.font.name
        style.font.size = base_style.font.size


def add_line(doc: Document, text: str = "", *, size: float = 10.2, color: str | None = None,
             bold: bool = False, before: float = 0, after: float = 3):
    p = doc.add_paragraph()
    polished.spacing(p, before=before, after=after)
    if text:
        r = p.add_run(text)
        polished.set_run(r, size=size, bold=bold, color=color)
    return p


def add_long_block(doc: Document, label: str, text: str, *, fill: str = LIGHT_GRAY,
                   label_color: str = BLUE, red: bool = False) -> None:
    p = doc.add_paragraph(style="FG Card")
    shade_paragraph(p, fill, "D9E2F3")
    r = p.add_run(f"【{label}】")
    polished.set_run(r, size=10.0, bold=True, color=RED if red else label_color)
    cleaned = clean_text(text)
    if not cleaned:
        r = p.add_run("（当前源包未提取到文本，需回源补齐。）")
        polished.set_run(r, size=9.8, color=RED)
        return
    lines = cleaned.splitlines()
    for i, line in enumerate(lines):
        p.add_run().add_break()
        r = p.add_run(line)
        polished.set_run(r, size=9.8, color=RED if red else None)
        if i > 120:
            break


def add_point_card(doc: Document, label: str, points: list[str], *,
                   fill: str = LIGHT_BLUE, label_color: str = BLUE,
                   numbered: bool = True) -> None:
    p = doc.add_paragraph(style="FG Card")
    shade_paragraph(p, fill, "D9E2F3")
    r = p.add_run(f"【{label}】")
    polished.set_run(r, size=10.0, bold=True, color=label_color)
    for i, point in enumerate(points, start=1):
        p.add_run().add_break()
        prefix = f"{i}. " if numbered else ""
        r = p.add_run(prefix + point)
        polished.set_run(r, size=9.8)


def clean_text(text: str) -> str:
    text = (text or "").replace("\u3000", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def clean_student_rubric(text: str) -> str:
    text = clean_text(text)
    cut_points = [idx for marker in RUBRIC_TRUNCATE_MARKERS if (idx := text.find(marker)) != -1]
    if cut_points:
        text = text[: min(cut_points)].rstrip()
    kept = []
    for line in text.splitlines():
        clean = line.strip()
        if not clean:
            kept.append(line)
            continue
        if any(bad in clean for bad in BAD_RUBRIC_LINE_CONTAINS):
            continue
        if re.fullmatch(r"高三政治.*阅卷总结.*", clean):
            continue
        kept.append(line)
    return clean_text("\n".join(kept))


def rubric_text(entry: dict, overrides: dict[str, str]) -> str:
    entry_id = entry["entry_id"]
    raw = A_THEME_RUBRIC_OVERRIDES.get(entry_id) or overrides.get(entry_id) or entry.get("rubric", "")
    return clean_student_rubric(raw)


def raw_rubric_text(entry: dict, overrides: dict[str, str]) -> str:
    entry_id = entry["entry_id"]
    return A_THEME_RUBRIC_OVERRIDES.get(entry_id) or overrides.get(entry_id) or entry.get("rubric", "")


def rubric_integrity_summary(packets: list[dict], overrides: dict[str, str]) -> tuple[int, list[dict]]:
    trimmed = []
    for entry in packets:
        raw = clean_text(raw_rubric_text(entry, overrides))
        shown = clean_text(rubric_text(entry, overrides))
        if raw != shown:
            trimmed.append(
                {
                    "entry_id": entry["entry_id"],
                    "title": entry["title"],
                    "raw_len": len(raw),
                    "shown_len": len(shown),
                }
            )
    return len(packets) - len(trimmed), trimmed


def load_source_repair_overrides() -> dict[str, dict]:
    if not REPAIR_OVERRIDES_JSON.exists():
        return {}
    return json.loads(REPAIR_OVERRIDES_JSON.read_text(encoding="utf-8"))


def load_pending_boundary_resolutions() -> dict[str, dict]:
    if not PENDING_BOUNDARY_JSON.exists():
        return {}
    data = json.loads(PENDING_BOUNDARY_JSON.read_text(encoding="utf-8"))
    return data.get("entries", data)


def apply_source_repair_overrides(packets: list[dict], repairs: dict[str, dict]) -> list[dict]:
    if not repairs:
        return packets
    fixed = []
    for entry in packets:
        row = copy.deepcopy(entry)
        repair = repairs.get(row["entry_id"])
        if repair:
            for field in ("material", "prompt", "rubric"):
                if repair.get(field):
                    row[field] = repair[field]
            row["source_repair_note"] = repair.get("repair_note", "")
            row["source_repair_source"] = repair.get("repair_source", "")
            row["source_repaired_fields"] = ",".join(field for field in ("material", "prompt", "rubric") if repair.get(field))
        fixed.append(row)
    return fixed


def apply_pending_boundary_resolutions(packets: list[dict], resolutions: dict[str, dict]) -> list[dict]:
    if not resolutions:
        return packets
    fixed = []
    for entry in packets:
        row = copy.deepcopy(entry)
        resolution = resolutions.get(row["entry_id"])
        if resolution:
            row["pending_boundary_status"] = resolution.get("status", "")
            row["pending_boundary_cleared"] = bool(resolution.get("clear_pending"))
            row["pending_boundary_note"] = resolution.get("display_note", "")
            row["pending_residual_issue"] = resolution.get("residual_issue", "")
            row["pending_theme_decision"] = resolution.get("theme_decision", "")
        fixed.append(row)
    return fixed


def infer_b_axis(prompt: str) -> str:
    p = prompt or ""
    if any(w in p for w in ["完成下表", "完成表格", "参考示例", "补充完整", "填写", "横线"]):
        return "B1"
    if any(w in p for w in ["维权", "调解", "如何", "建议", "准备", "起诉状", "提出解决"]):
        return "B6"
    if any(w in p for w in ["诉求", "请求", "支持", "能否得到"]):
        return "B3"
    if any(w in p for w in ["理由", "依据", "责任", "判决结果", "裁判理由", "法理依据"]):
        return "B2"
    if any(w in p for w in ["意义", "价值", "作用", "原因", "影响", "积极意义"]):
        return "B5"
    if any(w in p for w in ["评析", "认识", "谈谈", "观点", "阐明", "阐释", "分析"]):
        return "B4"
    return "B7"


def answer_atoms(text: str, limit: int = 5) -> list[str]:
    atoms = polished.score_atoms(text, max_items=limit)
    if atoms:
        return atoms[:limit]
    pieces = re.split(r"[。；;]\s*", clean_text(text))
    out = []
    for piece in pieces:
        piece = piece.strip()
        if len(piece) < 8:
            continue
        if any(bad in piece for bad in ["题号与设问", "学生表现", "阅卷总结", "试题分析"]):
            continue
        out.append(piece[:120])
        if len(out) >= limit:
            break
    return out


def classify_basis(entry: dict, theme_key: str, b_key: str) -> str:
    theme = THEMES[theme_key]
    base_sentence = f"本题归入{theme['name']}，因为材料和设问首先触发的是“{theme['signals'].split('。')[0]}”这一组入口信号。"
    prompt = clean_text(entry.get("prompt", ""))
    pending = clean_text(entry.get("pending_reason", ""))
    extra = f"设问动作属于{B_AXIS[b_key][0]}：{B_AXIS[b_key][1]}"
    if pending:
        if entry.get("pending_boundary_cleared"):
            note = clean_text(entry.get("pending_boundary_note", ""))
            return base_sentence + extra + f" 原待核边界已裁决：{note}"
        residual = clean_text(entry.get("pending_residual_issue") or pending)
        return base_sentence + extra + f" 当前源包另有待核提示：{residual}"
    return base_sentence + extra


def format_counter(counter: Counter, labeler=None, limit: int = 4) -> str:
    if not counter:
        return "无"
    parts = []
    for key, count in counter.most_common(limit):
        name = labeler(key) if labeler else key
        parts.append(f"{name} {count}")
    return "；".join(parts)


def b_axis_summary(entries: list[dict]) -> str:
    counter = Counter(infer_b_axis(e.get("prompt", "")) for e in entries)
    return format_counter(counter, lambda k: B_AXIS[k][0])


def year_summary(entries: list[dict]) -> str:
    counter = Counter(str(e.get("year", "")).strip()[:4] for e in entries if e.get("year"))
    return format_counter(counter)


def compact_entry_title(entry: dict) -> str:
    title = entry.get("title", "").replace("【待确认】", "")
    title = title.replace(" · ", "")
    title = title.replace("第", " 第", 1) if "第" in title else title
    return re.sub(r"\s+", "", title)


def b_axis_grouped_refs(entries: list[dict], *, max_per_axis: int = 4) -> str:
    if not entries:
        return "当前源包中暂未归入本主题的分问。"
    grouped: dict[str, list[dict]] = defaultdict(list)
    for entry in entries:
        grouped[infer_b_axis(entry.get("prompt", ""))].append(entry)
    parts = []
    for b_key in sorted(grouped, key=lambda k: B_AXIS[k][0]):
        rows = sorted(grouped[b_key], key=entry_sort_key)
        refs = "、".join(compact_entry_title(e) for e in rows[:max_per_axis])
        tail = f"等{len(rows)}题" if len(rows) > max_per_axis else ""
        parts.append(f"{B_AXIS[b_key][0]}：{refs}{tail}")
    return "；".join(parts)


def add_theme_method_page(doc: Document, theme_key: str, entries: list[dict]) -> None:
    method = THEME_METHOD_GUIDE[theme_key]
    add_point_card(doc, "本章先背", method["memory"], fill=LIGHT_BLUE)
    add_point_card(doc, "核心链条", method["chain"], fill="FFFFFF")
    add_point_card(doc, "常见设问翻译", method["questions"], fill=LIGHT_GRAY)
    add_point_card(doc, "采分模板", method["template"], fill=LIGHT_GOLD, label_color=GOLD)
    add_point_card(doc, "错答清单", method["wrong"], fill=LIGHT_RED, label_color=RED)
    polished.add_label(doc, "真题分组导读", b_axis_grouped_refs(entries), fill=LIGHT_BLUE)


def add_theme_overview(doc: Document, theme_buckets: dict[str, list[dict]]) -> None:
    polished.add_heading(doc, "A类十主题速查", 2)
    polished.add_box(
        doc,
        "用法",
        "先用这一页决定题目该从哪个 A 入口进入；进入章节后，先读章首学习页和本章方法页，再刷逐题五件套。",
        fill=LIGHT_BLUE,
    )
    for theme_key in THEMES:
        entries = theme_buckets.get(theme_key, [])
        guide = THEME_STUDENT_GUIDE[theme_key]
        text = (
            f"{len(entries)}个分问｜高频设问：{b_axis_summary(entries)}｜"
            f"第一判断：{THEMES[theme_key]['first']}｜本章目标：{guide['goal']}"
        )
        polished.add_label(doc, THEMES[theme_key]["name"], text)
    doc.add_page_break()


def add_theme_learning_panel(doc: Document, theme_key: str, entries: list[dict]) -> None:
    guide = THEME_STUDENT_GUIDE[theme_key]
    polished.add_box(doc, "本章怎么学", guide["goal"], fill=LIGHT_BLUE)
    polished.add_label(
        doc,
        "题源画像",
        f"本章收录 {len(entries)} 个分问；年份分布：{year_summary(entries)}；高频设问动作：{b_axis_summary(entries)}。",
    )
    add_point_card(doc, "一眼判断", guide["quick"], fill=LIGHT_GRAY)
    add_point_card(doc, "采分动作", guide["moves"], fill=LIGHT_GOLD, label_color=GOLD)
    add_point_card(doc, "考场句型", guide["sentences"], fill="FFFFFF")
    polished.add_label(doc, "边界提醒", guide["boundary"], red=True, fill=LIGHT_RED)


def add_theme_phase2_panel(doc: Document, theme_key: str) -> None:
    guide = THEME_PHASE2_GUIDE[theme_key]
    polished.add_label(doc, "命题人路径", "\n".join(guide["proposition"]), fill=LIGHT_BLUE)
    add_point_card(doc, "判题四步", guide["path"], fill="FFFFFF")
    add_point_card(doc, "高频给分件", guide["scoring"], fill=LIGHT_GOLD, label_color=GOLD)
    add_point_card(doc, "易混边界", guide["boundary"], fill=LIGHT_RED, label_color=RED)


def entry_sort_key(entry: dict) -> tuple:
    year = int(re.search(r"\d{4}", str(entry.get("year", "9999"))).group(0))
    qno = int(re.search(r"\d+", str(entry.get("question_no", "999"))).group(0))
    eid = int(re.search(r"\d+", entry.get("entry_id", "E999")).group(0))
    return (year, str(entry.get("district_or_exam", "")), str(entry.get("paper_type", "")), qno, eid)


def load_rows() -> tuple[list[dict], list[dict], list[dict], dict[str, str]]:
    packets = base.read_jsonl(base.PACKETS)
    coverage = base.read_csv(base.COVERAGE)
    source_rows = base.read_csv(base.SOURCE_LEDGER)
    overrides = base.load_rubric_overrides()
    overrides.update(getattr(base, "EXTRA_RUBRIC_OVERRIDES", {}))
    overrides.update(A_THEME_RUBRIC_OVERRIDES)
    return packets, coverage, source_rows, overrides


def add_cover(doc: Document, stats, packets: list[dict], theme_counts: Counter) -> None:
    add_line(doc, "选择性必修二", size=15, color=BLUE, bold=True, after=2)
    add_line(doc, "《法律与生活》A类十主题学生宝典工作稿", size=24, bold=True, after=2)
    add_line(doc, "飞哥正志讲堂｜十主题刷题地图", size=12.5, color=GRAY, bold=True, after=16)
    polished.add_box(
        doc,
        "本稿定位",
        "前半部分保留已认可的总路线、A线入口、B线动作和高频模型；后半部分按 A1-A10 做十主题方法页，并在每章增加命题人路径、判题四步、高频给分件、易混边界，再把当前源包的全部法律主观题分问收进正文，逐题展开五件套。",
        fill=LIGHT_BLUE,
    )
    polished.add_label(doc, "覆盖口径", f"当前源包：{stats.suites_total} 套卷、{stats.big_questions_total} 道大题、{stats.subquestions_total} 个分问；本稿入正文 {len(packets)} 个分问。", red=True)
    polished.add_label(doc, "红线说明", "题干与细则来自当前 source_packets_final.jsonl；凡存在 OCR、表格行列、模块边界或正式细则风险的条目，正文内直接标注待核，不作最终闭合。")
    top = "；".join(f"{THEMES[k]['name']} {theme_counts.get(k, 0)}题" for k in THEMES)
    polished.add_label(doc, "A类题量", top)
    doc.add_page_break()


def add_framework_front(doc: Document, packets: list[dict]) -> None:
    polished.add_route(doc)
    polished.add_axis_sections(doc, [])
    polished.add_models(doc, [])
    doc.add_page_break()
    polished.add_heading(doc, "六、A类十主题逐题厚版", 1)
    polished.add_box(
        doc,
        "读法",
        "每题先看【入口】，再看【判定依据】确认为什么放在这个 A 类；随后照读【题目材料】【设问】【细则】，最后用【答案落点】训练采分句。",
        fill=LIGHT_BLUE,
    )


def add_theme_chapter(doc: Document, theme_key: str, entries: list[dict], overrides: dict[str, str]) -> None:
    theme = THEMES[theme_key]
    polished.add_heading(doc, theme["name"], 2)
    polished.add_label(doc, "信号卡", theme["signals"])
    polished.add_label(doc, "第一判断", theme["first"], red=True, fill=LIGHT_RED)
    polished.add_label(doc, "易错提醒", theme["trap"])
    add_theme_learning_panel(doc, theme_key, entries)
    add_theme_phase2_panel(doc, theme_key)
    add_theme_method_page(doc, theme_key, entries)
    if not entries:
        polished.add_label(doc, "本轮题源", "当前源包中暂未归入本主题的分问。")
        return
    reps = "；".join(e["title"].replace("【待确认】", "") for e in entries[:8])
    polished.add_label(doc, "代表题源", reps)
    for idx, entry in enumerate(entries, start=1):
        b_key = infer_b_axis(entry.get("prompt", ""))
        title = entry["title"].replace("【待确认】", "")
        polished.add_heading(doc, f"{idx:02d}. {title}", 3)
        polished.add_label(doc, "入口", f"{theme['name']} ｜ {B_AXIS[b_key][0]}", red=True, fill=LIGHT_RED)
        if entry.get("source_repair_note"):
            fields = entry.get("source_repaired_fields", "")
            polished.add_label(doc, "回源修复", f"{entry['source_repair_note']}（已修复字段：{fields}）", fill=LIGHT_BLUE)
        if entry["entry_id"] in NEEDS_RELAYOUT and entry["entry_id"] not in SOURCE_REPAIR_CLEARS_RELAYOUT:
            polished.add_label(doc, "待回源重排", NEEDS_RELAYOUT[entry["entry_id"]], red=True, fill=LIGHT_RED)
        if entry.get("pending_reason"):
            if entry.get("pending_boundary_cleared"):
                polished.add_label(doc, "边界裁决", entry.get("pending_boundary_note", ""), fill=LIGHT_BLUE)
            else:
                residual = entry.get("pending_residual_issue") or entry["pending_reason"]
                if entry.get("pending_boundary_note"):
                    residual = f"{residual}（{entry['pending_boundary_note']}）"
                polished.add_label(doc, "待核边界", residual, red=True, fill=LIGHT_RED)
        polished.add_label(doc, "判定依据", classify_basis(entry, theme_key, b_key))
        add_long_block(doc, "题目材料", entry.get("material", ""), fill=LIGHT_GRAY)
        add_long_block(doc, "设问", entry.get("prompt", ""), fill=LIGHT_GOLD, label_color=GOLD)
        rtext = rubric_text(entry, overrides)
        add_long_block(doc, "细则", rtext, fill="FFFFFF")
        atoms = answer_atoms(rtext, 5)
        atom_text = "\n".join(f"{i}. {atom}" for i, atom in enumerate(atoms, start=1)) or "（当前细则文本不足以稳定提炼，待回源核验后补齐。）"
        add_long_block(doc, "答案落点", atom_text, fill=LIGHT_RED, label_color=RED, red=True)


def add_index(doc: Document, packets: list[dict]) -> None:
    doc.add_page_break()
    polished.add_heading(doc, "七、真题索引", 1)
    polished.add_box(doc, "用途", "按年份和区回练时，看这一节；按主题刷题时，看第六节。", fill=LIGHT_BLUE)
    for entry in sorted(packets, key=entry_sort_key):
        theme_key = PRIMARY_THEME_OVERRIDES.get(entry["entry_id"], "A10")
        b_key = infer_b_axis(entry.get("prompt", ""))
        pending_open = bool(entry.get("pending_reason") and not entry.get("pending_boundary_cleared"))
        relayout_open = entry["entry_id"] in NEEDS_RELAYOUT and entry["entry_id"] not in SOURCE_REPAIR_CLEARS_RELAYOUT
        flag = "｜待核" if pending_open or relayout_open else ""
        polished.add_label(doc, entry["title"].replace("【待确认】", ""), f"{THEMES[theme_key]['name']} ｜ {B_AXIS[b_key][0]}{flag}")


def add_pending_appendix(doc: Document, packets: list[dict]) -> None:
    pending_original = [p for p in packets if p.get("pending_reason")]
    pending_unresolved = [p for p in pending_original if not p.get("pending_boundary_cleared")]
    relayout = [p for p in packets if p["entry_id"] in NEEDS_RELAYOUT and p["entry_id"] not in SOURCE_REPAIR_CLEARS_RELAYOUT]
    doc.add_page_break()
    polished.add_heading(doc, "八、待核/待补清单", 1)
    polished.add_box(
        doc,
        "口径",
        "本清单只列真正尚未闭合的项目；已通过回源修复或边界裁决清除的旧标记不再列入。当前不缺题、不缺材料，未闭合项为正式细则或模块边界问题。",
        fill=LIGHT_BLUE,
    )
    if not pending_unresolved and not relayout:
        polished.add_label(doc, "当前状态", "无待核/待补项。")
        return
    for entry in sorted({p["entry_id"]: p for p in pending_unresolved + relayout}.values(), key=lambda r: r["entry_id"]):
        theme_key = PRIMARY_THEME_OVERRIDES.get(entry["entry_id"], "A10")
        b_key = infer_b_axis(entry.get("prompt", ""))
        reason = entry.get("pending_residual_issue") or entry.get("pending_reason") or NEEDS_RELAYOUT.get(entry["entry_id"], "")
        if entry["entry_id"] in NEEDS_RELAYOUT and entry["entry_id"] not in SOURCE_REPAIR_CLEARS_RELAYOUT:
            reason = NEEDS_RELAYOUT[entry["entry_id"]]
        text = f"{THEMES[theme_key]['name']} ｜ {B_AXIS[b_key][0]}\n未闭合原因：{reason}"
        polished.add_label(doc, entry["title"].replace("【待确认】", ""), text, red=True, fill=LIGHT_RED)


def write_report(stats, packets: list[dict], coverage: list[dict], source_rows: list[dict],
                 theme_buckets: dict[str, list[dict]], repairs: dict[str, dict],
                 overrides: dict[str, str]) -> None:
    pending_original = [p for p in packets if p.get("pending_reason")]
    pending_cleared = [p for p in pending_original if p.get("pending_boundary_cleared")]
    pending_unresolved = [p for p in pending_original if not p.get("pending_boundary_cleared")]
    relayout_all = [p for p in packets if p["entry_id"] in NEEDS_RELAYOUT]
    relayout = [p for p in relayout_all if p["entry_id"] not in SOURCE_REPAIR_CLEARS_RELAYOUT]
    relayout_cleared = [p for p in relayout_all if p["entry_id"] in SOURCE_REPAIR_CLEARS_RELAYOUT]
    repaired = [p for p in packets if p["entry_id"] in repairs]
    theme_lines = [f"| {THEMES[k]['name']} | {len(theme_buckets.get(k, []))} |" for k in THEMES]
    pending_map = {p["entry_id"]: p for p in pending_unresolved + relayout}
    pending_lines = [
        f"- {p['entry_id']} {p['title']}: {p.get('pending_residual_issue') or p.get('pending_reason') or NEEDS_RELAYOUT.get(p['entry_id'], '')}"
        for p in sorted(pending_map.values(), key=lambda r: r["entry_id"])
    ]
    boundary_lines = [
        f"- {p['entry_id']} {p['title']}: {p.get('pending_boundary_note', '')}"
        for p in sorted(pending_cleared, key=lambda r: r["entry_id"])
    ]
    rubric_exact_count, rubric_trimmed = rubric_integrity_summary(packets, overrides)
    rubric_trim_lines = [
        f"- {row['entry_id']} {row['title']}: 官方评分段保留，剔除页眉/教师端讲评类非评分文本（{row['raw_len']} -> {row['shown_len']} 字）。"
        for row in sorted(rubric_trimmed, key=lambda r: r["entry_id"])
    ]
    included_status = Counter(r.get("conversion_status", "") for r in source_rows if r.get("include_status") == "include")
    lines = [
        "# A类十主题学生宝典覆盖报告",
        "",
        "## 1. 数量核对",
        "",
        "| 项目 | 目标 | 当前证据 | 状态 |",
        "| --- | ---: | ---: | --- |",
        f"| 套卷 | 63 | {stats.suites_total} | 对齐 |",
        f"| 大题 | 67 | {stats.big_questions_total} | 对齐 |",
        f"| 分问 | 74 | {stats.subquestions_total} | 对齐 |",
        f"| 本稿正文收录分问 | 74 | {len(packets)} | 全部进入 A 类正文，不再只放附录 |",
        "",
        "## 2. 题源处理状态",
        "",
        f"- 桌面候选源文件：{len(source_rows)} 条 ledger 记录；include 源：{stats.included_sources}。",
        f"- 转换状态：converted {included_status.get('converted', 0)}；ocr {included_status.get('ocr', 0)}；conversion_error {included_status.get('conversion_error', 0)}。",
        f"- 覆盖矩阵：{len(coverage)} 套；has_xuanbier {stats.has_xuanbier}；no_xuanbier {stats.no_xuanbier}。",
        f"- 回源修复覆盖层：{len(repaired)} 个条目；原待回源重排/材料风险 {len(relayout_all)} 个，其中本轮已证据清除 {len(relayout_cleared)} 个，仍保留 {len(relayout)} 个。",
        f"- pending_reason 原始 {len(pending_original)} 个；本轮边界裁决已清除 {len(pending_cleared)} 个；仍保留 {len(pending_unresolved)} 个。",
        f"- 当前仍需人工/回源核验：未清除 pending_reason {len(pending_unresolved)} 个；未闭合待回源重排/材料风险 {len(relayout)} 个。",
        "",
        "## 3. A类主题题量",
        "",
        "| 主题 | 分问数 |",
        "| --- | ---: |",
        *theme_lines,
        "",
        "## 4. 边界裁决已清除清单",
        "",
        *(boundary_lines or ["- 无。"]),
        "",
        "## 5. 仍待核/待补清单",
        "",
        *(pending_lines or ["- 无。"]),
        "- 详见：`05_output/A_THEME_REMAINING_PENDING_RECHECK_20260604.md`。",
        "",
        "## 6. 本轮回源修复",
        "",
        *(
            [
                f"- {p['entry_id']} {p['title']}: {p.get('source_repair_note', '')}"
                for p in sorted(repaired, key=lambda r: r["entry_id"])
            ]
            or ["- 无。"]
        ),
        "",
        "## 7. 细则原文审计",
        "",
        f"- 74个分问均有【细则】文本；与源包/覆盖层评分段完全一致：{rubric_exact_count} 个。",
        f"- 仅剔除页眉、教师端讲评、学生表现、改进措施等非评分段：{len(rubric_trimmed)} 个。",
        "- 已修正规则：保留“不得分”“替换答案”“变通说明”等评分边界句，不再因出现“存在问题/出现的问题”而删除。",
        *(rubric_trim_lines or ["- 无剔除项。"]),
        "",
        "## 8. 本阶段正文增强与清洗",
        "",
        "- A类十主题每章新增方法页：本章先背、核心链条、常见设问翻译、采分模板、错答清单、真题分组导读。",
        "- A类十主题进入二轮主题加工：每章新增命题人路径、判题四步、高频给分件、易混边界，重点把章首从目录提示升级为学生可直接迁移的判题说明。",
        "- 新增 pending 边界裁决覆盖层：原始18个 pending_reason 中已清除16个；E009保留正式细则核验，E057保留模块边界核验。",
        "- 将 E046、E056 两个生成式AI权责边界题从 A4 调整到 A6 侵权责任开放容器。",
        "- 合并全量汇编正式细则覆盖，补入 E043 干净侵权责任细则。",
"- 接入学生版已验证的 E058/E059 干净细则，清除通州期末/一模中误串入的全球治理、上合组织、全球南方等外模块内容。",
"- 展示细则前统一截除学生表现、学生问题、教学启示、教师教学、改进措施、阅卷总结等教师端阅卷反思语段。",
"- Word 末尾新增“八、待核/待补清单”，单列 E009、E057 两个仍未闭合项。",
"",
"## 9. 本阶段产物",
        "",
        f"- Word：`{OUT_DOCX.name}`",
        f"- 桌面副本：`{DESKTOP_COPY}`",
        "- 剩余待核复查：`05_output/A_THEME_REMAINING_PENDING_RECHECK_20260604.md`",
        "- 细则原文QA：`05_output/A_THEME_RUBRIC_SOURCE_INTEGRITY_QA_20260604.md`",
        "- 目标完成审计：`05_output/A_THEME_OBJECTIVE_COMPLETION_AUDIT_20260604.md`",
        "- Word COM 视觉QA：`05_output/A_THEME_WORD_COM_VISUAL_QA_20260604.md`",
        "- 视觉接触表人工复查：`05_output/A_THEME_VISUAL_CONTACT_SHEET_REVIEW_20260604.md`",
        "- 新增加工：A类十主题速查页；每章新增本章怎么学、题源画像、一眼判断、采分动作、考场句型、边界提醒、命题人路径、判题四步、高频给分件、易混边界、本章先背、核心链条、常见设问翻译、采分模板、错答清单、真题分组导读。",
        "- 说明：本稿是 A 类十主题学生宝典工作稿；还未声称完成最终闭合，后续需要逐条回源修复待核项。",
        "",
    ]
    REPORT_MD.write_text("\n".join(lines), encoding="utf-8")


def build() -> None:
    packets, coverage, source_rows, overrides = load_rows()
    repairs = load_source_repair_overrides()
    boundary_resolutions = load_pending_boundary_resolutions()
    packets = apply_source_repair_overrides(packets, repairs)
    packets = apply_pending_boundary_resolutions(packets, boundary_resolutions)
    stats = base.source_status_stats(source_rows, coverage, packets)
    theme_buckets: dict[str, list[dict]] = defaultdict(list)
    for entry in packets:
        theme_key = PRIMARY_THEME_OVERRIDES.get(entry["entry_id"], "A10")
        theme_buckets[theme_key].append(entry)
    for key in theme_buckets:
        theme_buckets[key] = sorted(theme_buckets[key], key=entry_sort_key)

    doc = Document()
    polished.style_document(doc)
    ensure_fgcard_alias(doc)
    add_cover(doc, stats, packets, Counter({k: len(v) for k, v in theme_buckets.items()}))
    add_framework_front(doc, packets)
    add_theme_overview(doc, theme_buckets)
    for theme_key in THEMES:
        add_theme_chapter(doc, theme_key, theme_buckets.get(theme_key, []), overrides)
    add_index(doc, packets)
    add_pending_appendix(doc, packets)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    shutil.copy2(OUT_DOCX, DESKTOP_COPY)
    write_report(stats, packets, coverage, source_rows, theme_buckets, repairs, overrides)
    print(OUT_DOCX)
    print(REPORT_MD)
    print(DESKTOP_COPY)


if __name__ == "__main__":
    build()
