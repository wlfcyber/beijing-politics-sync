import csv, os, re, shutil
from pathlib import Path
from datetime import datetime

SCRIPT_DIR = Path(__file__).resolve().parent
LOCAL_RUN = SCRIPT_DIR.parent
DEFAULT_RUN = Path(r'C:\bp_sync_visible\reports\full_all_suites_independent_rerun_2026-04-29')
DEFAULT_DESKTOP = (
    Path.home() / 'Desktop' / '4.29凌晨跑完的结果v6'
    if os.name != 'nt'
    else Path(r'C:\Users\Administrator\Desktop\4.29凌晨跑完的结果v6')
)

RUN = Path(os.environ.get('V6_RUN_DIR') or (LOCAL_RUN if LOCAL_RUN.exists() else DEFAULT_RUN))
DESKTOP = Path(os.environ.get('V6_OUTPUT_DIR') or DEFAULT_DESKTOP)
STUDENT = DESKTOP / '01_学生版Word'
AUDIT = DESKTOP / '02_审计证据'
CSVOUT = DESKTOP / '03_结构化CSV'
LOGS = DESKTOP / '04_过程日志'
for p in [STUDENT, AUDIT, CSVOUT, LOGS]:
    p.mkdir(parents=True, exist_ok=True)

HEADER = f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"

def read_text(path):
    if path.exists():
        return path.read_text(encoding='utf-8', errors='replace').strip()
    return ''

def strip_title(text):
    lines = text.splitlines()
    if lines and lines[0].startswith('# '):
        return '\n'.join(lines[1:]).strip()
    return text.strip()

def sanitize_student_text(text):
    """Keep final student-facing documents free of audit/runtime vocabulary."""
    replacements = {
        'module-boundary/excluded': '模块边界：不入库',
        'rendered-ocr-needed': '待人工核答案',
        'NEED_EVIDENCE': '待补证据',
        'included': '已入库',
        'Blocked Objective Notes': '客观题待补说明',
        'OCR记录': '技术记录',
        'OCR': '文本识别',
        'source_path': '证据路径',
        '不进入学生版': '不展开必修四框架',
        '进入学生版': '纳入对应框架',
        '入学生版': '纳入对应框架',
        'suite | q | wrong option | wrong-option type | why wrong': '套卷｜题号｜错项｜错项类型｜为什么错',
        'suite | q | answer | module | correct-option chain': '套卷｜题号｜答案｜模块｜正确项链条',
        'bundle text': '资料包文本',
        '文本识别记录': '技术记录',
        '文本识别 记录': '技术记录',
        'source id': '证据编号',
        'correct_and_wrong_option_chain': '正确项/错肢链',
        'correct_option_chain': '正确项链',
        'choice_chain': '选择题链',
        'correct_chain': '正确项链',
        'wrong_option': '错肢链',
        'chain_confirmed': '链条已确认',
        'PASS': '已纳入',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text.replace('`', '')
    text = text.replace('这里不放路径、行号、文件号、文本识别记录；证据路径放在审计样稿和 CSV 里。',
                        '这里不放证据路径和技术记录；证据细节放在审计样稿和 CSV 里。')
    text = text.replace('这里不放路径、行号、文件号、技术记录；证据路径放在审计样稿和 CSV 里。',
                        '这里不放证据路径和技术记录；证据细节放在审计样稿和 CSV 里。')
    text = text.replace('这里不放路径、行号、技术记录、证据编号。证据边界见审计 CSV。',
                        '这里不放证据路径和技术记录；证据边界见审计 CSV。')
    text = text.replace('这里不放路径、行号、文本识别、证据编号。证据边界见审计 CSV。',
                        '这里不放证据路径和技术记录；证据边界见审计 CSV。')
    text = text.replace('说明：本文件不放路径、行号、文本识别、证据编号。证据边界见审计 CSV。',
                        '说明：本文件不放证据路径和技术记录；证据边界见审计 CSV。')
    return text

sample_student = read_text(RUN / 'sample_outputs' / 'S040_2026东城一模_student_facing_sample.md')
student_parts = []
if sample_student:
    student_parts.append('## 样例已确认：2026东城一模\n\n' + sanitize_student_text(strip_title(sample_student)))
for f in sorted((RUN / 'worker_outputs').glob('*_v6_student_entries.md')):
    student_parts.append(f'## {f.stem}\n\n' + sanitize_student_text(strip_title(read_text(f))))

philosophy_md = '# 必修四哲学材料-知识触发框架 v6\n\n' + HEADER + '\n\n'.join(student_parts) + '\n'
(STUDENT / '必修四哲学材料-知识触发框架_v6.md').write_text(philosophy_md, encoding='utf-8')

choice_parts=[]
choice_sample = read_text(RUN / 'sample_outputs' / 'S040_2026东城一模_choice_review_full.md')
if choice_sample:
    choice_parts.append('## 样例已确认：2026东城一模\n\n' + sanitize_student_text(strip_title(choice_sample)))
for f in sorted((RUN / 'worker_outputs').glob('*_v6_choice_review.md')):
    choice_parts.append(f'## {f.stem}\n\n' + sanitize_student_text(strip_title(read_text(f))))
late_choice = read_text(RUN / 'late_objective_closure' / 'late_objective_choice_review_supplement.md')
if late_choice:
    choice_parts.append('## 迟到客观题补证新增链条\n\n' + sanitize_student_text(strip_title(late_choice)))
choice_md = '# 北京高考政治选择题错肢总结 v6\n\n' + HEADER + '\n\n'.join(choice_parts) + '\n'
(STUDENT / '北京高考政治选择题错肢总结_v6.md').write_text(choice_md, encoding='utf-8')

# Culture document is filtered from student entries where headings or content mark culture/文化.
culture_chunks=[]
for block in student_parts:
    # coarse block split by second-level headings, preserve blocks mentioning culture.
    splits = re.split(r'(?=^## )', block, flags=re.M)
    for s in splits:
        if '文化' in s or '中华民族精神' in s or '核心价值观' in s:
            culture_chunks.append(s.strip())
culture_md = '# 必修四文化材料-知识触发框架 v6\n\n' + HEADER + '\n\n'.join(culture_chunks) + '\n'
(STUDENT / '必修四文化材料-知识触发框架_v6.md').write_text(culture_md, encoding='utf-8')

# Merge audit CSVs.
def merge_csv(pattern, out_name, include_sample=None, extra_files=None):
    files=list((RUN/'worker_outputs').glob(pattern))
    rows=[]; fields=None
    if include_sample and include_sample.exists():
        files=[include_sample]+files
    if extra_files:
        files=files+[f for f in extra_files if f.exists()]
    for f in files:
        with f.open('r', encoding='utf-8-sig', newline='') as fh:
            reader=csv.DictReader(fh)
            if fields is None:
                fields=reader.fieldnames or []
            for r in reader:
                r['_source_file']=str(f)
                rows.append(r)
    if not fields:
        fields=['_source_file']
    if '_source_file' not in fields:
        fields=fields+['_source_file']
    out=CSVOUT/out_name
    with out.open('w', encoding='utf-8-sig', newline='') as fh:
        w=csv.DictWriter(fh, fieldnames=fields, extrasaction='ignore')
        w.writeheader(); w.writerows(rows)
    return len(rows)

entries_count = merge_csv(
    '*_v6_audit_entries.csv',
    'v6_entries_merged.csv',
    RUN/'sample_outputs'/'S040_2026东城一模_entries.csv',
    [
        RUN/'late_objective_closure'/'late_objective_entries_normalized.csv',
        RUN/'late_objective_closure'/'objective_gap_closure_matrix.csv',
    ],
)
# Choice review is markdown, keep copied separately; if workers also output CSV later, merge can be extended.

# Audit evidence bundle.
audit_parts=[
    '# 必修四 v6 审计证据\n\n' + HEADER,
    '## Governor 初始规则\n\n' + read_text(RUN/'governor_reports'/'evidence_governor.md'),
    '## 补丁者预审\n\n' + read_text(RUN/'patcher_reports'/'multipoint_patch_review.md'),
]
for f in sorted((RUN/'worker_reports').glob('*_v6.md')):
    audit_parts.append(f'## {f.stem}\n\n' + read_text(f))
for f in sorted((RUN/'late_objective_closure').glob('*.md')):
    audit_parts.append(f'## late_objective_closure/{f.stem}\n\n' + read_text(f))
(AUDIT/'必修四_v6_审计证据.md').write_text('\n\n'.join(audit_parts), encoding='utf-8')
for src_name, dst_dir in [
    ('objective_gap_closure_report.md', AUDIT),
    ('objective_gap_closure_matrix.csv', CSVOUT),
    ('late_objective_entries_normalized.csv', CSVOUT),
]:
    src = RUN / 'late_objective_closure' / src_name
    if src.exists():
        shutil.copy2(src, dst_dir / src_name)

# Copy control files/logs.
for name in ['TASK_BRIEF.md','ROLE_CONTRACTS.md','STUDENT_OUTPUT_TEMPLATE.md','PROGRESS.md','DECISION_LOG.md','THREAD_REGISTRY.md','SUITE_ROSTER.csv','COVERAGE_MATRIX.csv']:
    src=RUN/name
    if src.exists():
        shutil.copy2(src, LOGS/name)

# Lightweight DOCX generation.
def md_to_docx(md_path, docx_path):
    from docx import Document
    from docx.shared import Pt
    doc=Document()
    style=doc.styles['Normal']
    style.font.name='Microsoft YaHei'
    style.font.size=Pt(10.5)
    def pipe_row_to_sentence(line):
        if not line.lstrip().startswith('|'):
            return None
        cols=[c.strip() for c in line.strip().strip('|').split('|')]
        if not any(cols):
            return ''
        joined=''.join(cols)
        if set(joined) <= set('-: '):
            return ''
        if cols[:6] == ['题号','答案','模块','正确项为什么对','错项怎么错','入库处理']:
            return '题号｜答案｜模块｜正确项为什么对｜错项怎么错｜入库处理'
        if cols[:6] == ['题号','答案','模块','正确项/边界处理','错项处理','入库处理']:
            return '题号｜答案｜模块｜正确项/边界处理｜错项处理｜入库处理'
        if len(cols) >= 6 and re.fullmatch(r'\d+', cols[0] or ''):
            labels=['题','答案','模块','正确项/边界处理','错项处理','入库处理']
            return f"第{cols[0]}题｜答案：{cols[1]}｜模块：{cols[2]}。{labels[3]}：{cols[3]}。{labels[4]}：{cols[4]}。{labels[5]}：{cols[5]}。"
        return '｜'.join(cols)

    for line in sanitize_student_text(md_path.read_text(encoding='utf-8', errors='replace')).splitlines():
        if not line.strip():
            continue
        pipe_sentence=pipe_row_to_sentence(line)
        if pipe_sentence is not None:
            if pipe_sentence:
                doc.add_paragraph(pipe_sentence)
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        else:
            # Remove bold markdown markers for Word readability.
            clean=re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            doc.add_paragraph(clean)
    doc.save(docx_path)

for md_name, docx_name in [
    ('必修四哲学材料-知识触发框架_v6.md','必修四哲学材料-知识触发框架_v6.docx'),
    ('北京高考政治选择题错肢总结_v6.md','北京高考政治选择题错肢总结_v6.docx'),
    ('必修四文化材料-知识触发框架_v6.md','必修四文化材料-知识触发框架_v6.docx'),
]:
    md_to_docx(STUDENT/md_name, STUDENT/docx_name)

report = f'''# v6 构建报告

生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

- 学生版片段来源数：{len(student_parts)}
- 选择题 review 来源数：{len(choice_parts)}
- 合并审计 CSV 行数：{entries_count}
- Word 已生成：3
'''
(DESKTOP/'BUILD_REPORT.md').write_text(report, encoding='utf-8')
print(report)
