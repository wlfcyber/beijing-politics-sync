import csv
import json
import re
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


RUN_DIR = Path(__file__).resolve().parents[1]
FUSION_MD = RUN_DIR / "fusion" / "framework_first_fusion" / "FRAMEWORK_FIRST_FUSION_P2_PATCHED.md"
P2_DECISIONS = RUN_DIR / "claudecode_lane" / "p2_recheck" / "P2_RECHECK_DECISIONS.csv"
DELIVERY = RUN_DIR / "delivery"

OUT_MD = DELIVERY / "选必三_逻辑与思维_学生版_框架闭合稿.md"
OUT_DOCX = DELIVERY / "选必三_逻辑与思维_学生版_框架闭合稿.docx"
QA_JSON = DELIVERY / "STUDENT_DELIVERY_QA.json"
QA_MD = DELIVERY / "STUDENT_DELIVERY_QA.md"

FORBIDDEN = [
    "固定分析流程",
    "P0证据",
    "P1证据",
    "P2证据",
    "A-formal",
    "A-support",
    "B-choice-signal",
    "source_insufficient",
    "block_from_student_body",
    "can_enter_fusion",
    "FRAMEWORK_FIRST",
    "framework_first",
    "RECHECK_MANIFEST",
    ".txt",
    ".pdf",
    ".pptx",
    ".docx",
    "C:\\",
    "/Users/",
    "FINAL_PASS",
    "终稿已通过",
    "需 Codex 回源复核",
]


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def source_label(row: dict[str, str]) -> str:
    qid = row["question_id"]
    parent_qid = row["parent_question_id"]
    if qid.startswith("Q-2024朝阳期中-18-"):
        return qid
    raw = parent_qid.removeprefix("Q-")
    parts = raw.split("-")
    if len(parts) >= 3 and parts[-2].isdigit() and parts[-1].isdigit():
        suite = "-".join(parts[:-2])
        q = f"{parts[-2]}({parts[-1]})"
    else:
        suite = "-".join(parts[:-1])
        q = parts[-1]
    return f"S-{suite} Q{q}"


def excluded_aliases() -> set[str]:
    aliases: set[str] = set()
    for row in read_csv(P2_DECISIONS):
        if row.get("can_enter_fusion") == "no" or row.get("decision") in {"source_insufficient", "block_from_student_body"}:
            aliases.add(row["question_id"])
            aliases.add(row["parent_question_id"])
            aliases.add(source_label(row))
    return aliases


def parse_fusion() -> list[dict[str, object]]:
    lines = FUSION_MD.read_text(encoding="utf-8", errors="replace").splitlines()
    entries: list[dict[str, object]] = []
    top = ""
    node = ""
    current = None
    heading_re = re.compile(r"^\*\*(.+?)\*\*（([^）]+)）$")
    for line in lines:
        if line.startswith("## "):
            if current:
                entries.append(current)
                current = None
            top = line[3:].strip()
            continue
        if line.startswith("### "):
            if current:
                entries.append(current)
                current = None
            node = line[4:].strip()
            continue
        match = heading_re.match(line)
        if match:
            if current:
                entries.append(current)
            current = {
                "top": top,
                "node": node,
                "label": match.group(1).strip(),
                "meta": match.group(2).strip(),
                "bullets": [],
            }
            continue
        if current is not None and line.startswith("- "):
            current["bullets"].append(line[2:])
    if current:
        entries.append(current)
    return entries


def display_label(label: str) -> str:
    if label.startswith("S-"):
        text = label[2:].replace(" Q", " 第") + "题"
        return text
    if label.startswith("Q-"):
        raw = label[2:]
        parts = raw.split("-")
        if len(parts) >= 3:
            suite = "-".join(parts[:-1])
            return f"{suite} 第{parts[-1]}题"
        return raw
    return label


def split_source_tail(text: str) -> str:
    cut_markers = [
        "源已验证：",
        "源端：",
        "源端答案表",
        "源040",
        "FRAMEWORK_FIRST",
        "framework_first",
        "RECHECK_MANIFEST",
        "fusion/",
        "can_enter_fusion",
        "covered_by_74_review_body",
        "choice trap结构成立",
        "paper.txt",
        "*paper.txt",
        " 行",
        "::",
    ]
    for marker in cut_markers:
        idx = text.find(marker)
        if idx >= 0:
            text = text[:idx]
    return text


def scrub(text: str) -> str:
    text = split_source_tail(text)
    text = text.replace("本次P2源级复核完成后应在后续fusion patch中去除该回源复核标记，", "")
    text = text.replace("本次P2源级复核完成后应在后续fusion patch中去除该回源复核标记", "")
    text = text.replace("正文锁定答案", "答案")
    text = re.sub(r"答案([A-D])由[^，。；]*?(锁定|确认|明示)", r"答案\1", text)
    text = re.sub(r"由20\d{2}[^，。；]*?(锁定|确认|明示)", "", text)
    text = re.sub(r"evidence_level维持[^，。；]*", "", text)
    text = re.sub(r"旧稿待核[^，。；。]*", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    text = text.strip("；，。 ")
    return text


def student_bullets(entry: dict[str, object]) -> list[tuple[str, str]]:
    out: list[tuple[str, str]] = []
    for bullet in entry["bullets"]:
        if bullet.startswith(("P0证据：", "P1证据：", "P2证据：")):
            continue
        if "source_insufficient" in bullet or "block_from_student_body" in bullet:
            continue
        mapping = [
            ("材料怎么看：", "材料触发点"),
            ("为什么想到：", "为什么能想到"),
            ("卷面句：", "答案落点"),
        ]
        for prefix, title in mapping:
            if bullet.startswith(prefix):
                value = scrub(bullet[len(prefix) :])
                if value:
                    out.append((title, value))
                break
    return out


def build_markdown(entries: list[dict[str, object]]) -> tuple[str, dict[str, object]]:
    skip_aliases = excluded_aliases()
    grouped: dict[str, dict[str, list[dict[str, object]]]] = defaultdict(lambda: defaultdict(list))
    skipped = []
    kept = []
    for entry in entries:
        label = str(entry["label"])
        meta = str(entry["meta"])
        body_text = "\n".join(entry["bullets"])
        if label in skip_aliases or "source_insufficient" in body_text or "block_from_student_body" in body_text:
            skipped.append(label)
            continue
        bullets = student_bullets(entry)
        if len(bullets) < 2:
            skipped.append(label)
            continue
        item = {**entry, "student_bullets": bullets}
        grouped[str(entry["top"])][str(entry["node"])].append(item)
        kept.append(label)

    lines = [
        "# 选必三《逻辑与思维》学生版框架闭合稿",
        "",
        "这份稿子按“框架 -> 题目 -> 材料触发 -> 为什么能想到 -> 答案落点”的顺序组织，供高三学生复习选必三《逻辑与思维》使用。",
        "",
        "## 使用方法",
        "",
        "- 先看材料中出现的动作、关系、变化、抽象链或推理结构。",
        "- 再判断它属于科学思维、辩证思维、创新思维，还是形式逻辑/推理题型。",
        "- 最后把“方法术语 + 本题材料事实 + 作用或结论”写成答案落点。",
        "",
    ]
    for top in grouped:
        if top == "Status: `FUSION_DRAFT_NOT_FINAL`":
            continue
        lines.extend([f"## {top}", ""])
        for node, node_entries in grouped[top].items():
            lines.extend([f"### {node}", ""])
            for entry in node_entries:
                type_hint = "选择题" if "choice_trap" in str(entry["meta"]) else "主观题"
                lines.extend([f"#### {type_hint}：{display_label(str(entry['label']))}", ""])
                for title, value in entry["student_bullets"]:
                    lines.append(f"- **{title}：**{value}")
                lines.append("")

    text = "\n".join(lines).rstrip() + "\n"
    qa = {
        "kept_entries": len(kept),
        "skipped_entries": len(skipped),
        "top_nodes": {top: sum(len(v) for v in nodes.values()) for top, nodes in grouped.items()},
        "node_count": sum(len(nodes) for nodes in grouped.values()),
        "forbidden_hits": {token: text.count(token) for token in FORBIDDEN},
        "kept_sample": kept[:10],
        "skipped_sample": skipped[:10],
    }
    qa["verdict"] = "STUDENT_MARKDOWN_OK" if not any(qa["forbidden_hits"].values()) and len(kept) >= 90 else "STUDENT_MARKDOWN_FAIL"
    return text, qa


def set_cell_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_markdown_to_docx(md_text: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size, color in [
        ("Title", 22, RGBColor(31, 78, 121)),
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(63, 63, 63)),
        ("Heading 3", 11, RGBColor(31, 78, 121)),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color

    for line in md_text.splitlines():
        if line.startswith("# "):
            p = doc.add_paragraph(line[2:], style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
        elif line.startswith("#### "):
            p = doc.add_paragraph(line[5:], style="Heading 3")
            set_cell_shading(p, "EAF2F8")
        elif line.startswith("- **"):
            match = re.match(r"- \*\*(.+?)：\*\*(.*)", line)
            if match:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(match.group(1) + "：")
                run.bold = True
                run.font.name = "Microsoft YaHei"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                p.add_run(match.group(2))
            else:
                doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line)
        else:
            continue

    doc.core_properties.title = "选必三《逻辑与思维》学生版框架闭合稿"
    doc.core_properties.author = "飞哥政治庄园"
    doc.save(OUT_DOCX)


def main() -> None:
    DELIVERY.mkdir(exist_ok=True)
    entries = parse_fusion()
    md_text, qa = build_markdown(entries)
    OUT_MD.write_text(md_text, encoding="utf-8")
    add_markdown_to_docx(md_text)

    qa["docx"] = str(OUT_DOCX)
    qa["markdown"] = str(OUT_MD)
    QA_JSON.write_text(json.dumps(qa, ensure_ascii=False, indent=2), encoding="utf-8")
    QA_MD.write_text(
        "\n".join(
            [
                "# Student Delivery QA",
                "",
                f"Verdict: `{qa['verdict']}`",
                "",
                f"- kept entries: `{qa['kept_entries']}`",
                f"- skipped entries: `{qa['skipped_entries']}`",
                f"- node count: `{qa['node_count']}`",
                f"- top nodes: `{qa['top_nodes']}`",
                f"- forbidden hits: `{qa['forbidden_hits']}`",
                f"- markdown: `{OUT_MD}`",
                f"- docx: `{OUT_DOCX}`",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    print(json.dumps(qa, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
