#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/wanglifei/Desktop/北京高考政治/选必二法律主观题框架_从题源生长")
IN_MD = ROOT / "12_final_baodian" / "选必二法律主观题满分训练宝典_v5_9_27核心38题保分索引_双审门禁修补版_20260521.md"
OUT_DOCX = ROOT / "12_final_baodian" / "选必二法律主观题满分训练宝典_v5_9_27核心38题保分索引_双审门禁修补版_20260521.docx"


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para_spacing(paragraph, before: int = 0, after: int = 4, line: float = 1.08) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def add_bottom_border(paragraph, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_cell_text(cell, text: str, bold: bool = False, fill: str | None = None) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)
    p = cell.paragraphs[0]
    set_para_spacing(p, after=1, line=1.0)
    r = p.add_run(text)
    set_run_font(r, 8.5, bold)


def parse_cells(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_separator(row: list[str]) -> bool:
    return all(c and set(c) <= {"-", ":", " "} for c in row)


def add_table(doc: Document, table_lines: list[str]) -> None:
    rows = [parse_cells(line) for line in table_lines]
    rows = [row for row in rows if not is_separator(row)]
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            set_cell_text(
                table.cell(r_idx, c_idx),
                row[c_idx] if c_idx < len(row) else "",
                bold=(r_idx == 0),
                fill=("EAF2F8" if r_idx == 0 else None),
            )
    doc.add_paragraph()


def clean_inline_markdown(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("**", "")
    return text


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(clean_inline_markdown(text), level=level)
    set_para_spacing(p, before=6 if level else 0, after=5)
    for run in p.runs:
        if level == 0:
            set_run_font(run, 18, True, "7A1F1F")
        elif level == 1:
            set_run_font(run, 15, True, "1F4E79")
        else:
            set_run_font(run, 12, True, "333333")


def add_warning(doc: Document, text: str, low_freq: bool = False) -> None:
    p = doc.add_paragraph()
    set_para_spacing(p, before=3, after=5, line=1.05)
    shade_paragraph(p, "FFF2CC" if low_freq else "F4CCCC")
    add_bottom_border(p, "D6B656" if low_freq else "CC0000")
    r = p.add_run(clean_inline_markdown(text))
    set_run_font(r, 9.5, True, "7F6000" if low_freq else "990000")


def add_body_paragraph(doc: Document, line: str) -> None:
    p = doc.add_paragraph()
    set_para_spacing(p)
    r = p.add_run(clean_inline_markdown(line))
    set_run_font(r, 10.2)


def add_bullet(doc: Document, line: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    set_para_spacing(p, after=2, line=1.04)
    r = p.add_run(clean_inline_markdown(line))
    set_run_font(r, 9.8)


def add_numbered(doc: Document, line: str) -> None:
    p = doc.add_paragraph(style="List Number")
    set_para_spacing(p, after=2, line=1.04)
    r = p.add_run(clean_inline_markdown(line))
    set_run_font(r, 9.8)


def export() -> None:
    lines = IN_MD.read_text(encoding="utf-8").splitlines()
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(title, before=0, after=8)
    r = title.add_run("选必二《法律与生活》主观题满分训练宝典")
    set_run_font(r, 20, True, "7A1F1F")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(sub, after=10)
    r = sub.add_run("V5.9 双审门禁修补版：27 核心满分训练 + 38 保分/边界/回源索引")
    set_run_font(r, 11, True, "666666")

    guard = doc.add_paragraph()
    set_para_spacing(guard, before=4, after=8)
    shade_paragraph(guard, "FCE4D6")
    add_bottom_border(guard, "C00000")
    r = guard.add_run("守门口径：本稿不是 65 题全部核心满分闭环。27 道核心题用于满分训练；其余 38 道只作为保分、边界、回源或参考索引。")
    set_run_font(r, 10.5, True, "9C0006")

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].rstrip().startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            add_table(doc, table_lines)
            continue
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), 0)
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), 1)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), 2)
        elif line.startswith("> "):
            text = line[2:].strip()
            add_warning(doc, text, low_freq=("LOW-FREQUENCY" in text or "低频" in text))
        elif line.startswith("- "):
            add_bullet(doc, line[2:].strip())
        elif re.match(r"^\\d+\\.\\s+", line):
            add_numbered(doc, re.sub(r"^\\d+\\.\\s+", "", line).strip())
        elif line == "---":
            p = doc.add_paragraph("—" * 24)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.runs[0], 9, False, "999999")
        else:
            add_body_paragraph(doc, line)
        i += 1

    props = doc.core_properties
    props.title = "选必二法律主观题满分训练宝典 V5.9"
    props.subject = "北京高考政治选必二法律与生活主观题"
    props.author = "Codex"
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    export()
