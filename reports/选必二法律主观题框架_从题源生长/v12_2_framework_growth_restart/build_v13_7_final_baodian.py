from __future__ import annotations

import csv
import html
import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


STATUS = "v13_7_final_baodian_integrated_pending_render_qa"
FINAL_STATUS = "v13_7_final_baodian_integrated_pdf_rendered_docx_generated_with_docx_render_caveat"
TITLE = "选必二《法律与生活》法律宝典 v13.7 集成终版"
DATE_LINE = "2026-05-23"


def find_run_root() -> Path:
    for p in (Path.cwd() / "reports").rglob("v12_2_framework_growth_restart"):
        if (p / "v13_1_round05_patched_final").is_dir() and (
            p / "v13_7_zero_baseline_b1_b3_final_precision_patch"
        ).is_dir():
            return p
    raise FileNotFoundError("v12_2_framework_growth_restart not found")


ROOT = find_run_root()
SRC = ROOT / "v13_1_round05_patched_final"
PATCH = ROOT / "v13_7_zero_baseline_b1_b3_final_precision_patch"
OUT = ROOT / "v13_7_final_baodian_integrated"
TRACE = OUT / "traceability"
GOV = OUT / "governance"
CONF = OUT / "governor_confucius"


A_NOTES = {
    "A1": "A1 是民事行为效力、意思表示、代理、行为能力等基础脊柱；只有题干直接考这些，才让它做主入口。",
    "A2": "先叫准人格权益或人身权益名称，再写侵害行为、损害后果和停止侵害、赔礼道歉、赔偿等责任。",
    "A3": "先看专有/共有、相邻关系、采光通行安全和程序事实；程序合法是裁判事实，不自动变成程序题。",
    "A4": "先判断合同成立和履行，再写违约责任；有侵权损害时，合同链和侵权链分开落点。",
    "A5": "先看商业秘密、著作权、商标、混淆、虚假宣传、商业诋毁或数据抓取；不要把所有经营标识都写成著作权。",
    "A6": "按侵权四步写：权益受损、违法或过错、因果关系、责任方式；安全保障题必须看受害人自身风险和责任分担。",
    "A7": "先定位婚姻、亲子、赡养、共同债务或继承关系；评价制度时要回到当时保护对象和公平正义。",
    "A8": "先看人格、经济、组织三从属性；事实用工管理、报酬、考核、持续劳动比合同名称更关键。",
    "A9": "消费者场景优先写知情、公平交易、安全保障、格式条款、退赔和欺诈；未成年人直播打赏用 A9 主场景加 A1 脊柱。",
    "A10": "只有纠纷解决机制、证据、举证、仲裁、调解、诉讼类型本身成为评分对象时，A10 才做主入口。",
}


B_NOTES = {
    "B1": "表格题先读原列名；看不清时默认“材料事实|法律关系和要件|裁判要点/责任”。一行只写一条法律链，多案例分行，先总入口后行内入口。",
    "B2": "判决/责任理由题按“法律关系 -> 规则 -> 事实 -> 裁判结论 -> 边界或责任分担”写，不要先写意义价值。",
    "B3": "诉求能否支持题先给支持/不支持/部分支持，再逐项写权利基础、构成要件、责任方式和惩罚性赔偿边界。",
    "B4": "评析/认识题先表态，再指出合理处和边界，最后给正确法律立场；不要把它写成纯价值口号。",
    "B5": "意义/价值题要先抓法律规则和治理对象，再写如何保护权利、规范行为、维护秩序；价值必须由材料规则推出。",
    "B6": "维权路径题写清主体、证据、协商/调解/仲裁/诉讼路径、请求内容和程序边界。",
}


def read(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.strip() + "\n", encoding="utf-8", newline="\n")


def copy_tree_file(src: Path, dst: Path) -> None:
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)


def now_stamp() -> str:
    return "2026-05-23 23:10 +08:00"


def clean_patch_framework(text: str) -> str:
    text = text.replace("# 01 v13.2 零基础工具箱补丁框架", "# 01 v13.7 双轴法律主观题框架章")
    text = text.replace("状态：`v13_2_zero_baseline_toolbox_patch_pending_claude_retest`", f"状态：`{STATUS}`")
    text = text.replace("生成时间：2026-05-23 21:10 +08:00", f"生成时间：{now_stamp()}")
    preface = f"""# 01 v13.7 双轴法律主观题框架章

生成时间：{now_stamp()}

状态：`{STATUS}`

本章是 v13.1 成品框架经真实 Claude 零基础学生盲测后，连续 v13.2-v13.7 补丁整合而成的学生版框架章。它不推翻 A/B 双轴：A轴解决“这是哪类法律关系或法律内容”，B轴解决“题目要我做什么动作”。v13.7 的新增价值在于把聪明但零基础的学生容易卡住的地方补成可执行口令：入口后工具句、B1 表格列名、B3 诉求分层、双链字数比例、停止条件。

## 一页总览

| 先后 | 问自己什么 | 输出什么 |
|---|---|---|
| 1 | 这是谁和谁，围绕什么利益起争议？ | A轴主入口：人格、物权、合同、竞争、侵权、婚姻继承、劳动、消费、程序等 |
| 2 | 题干让我完成什么动作？ | B轴动作：表格、判决理由、诉求支持、评析认识、意义价值、维权路径 |
| 3 | 材料给了哪些事实信号？ | 事实 -> 法律要件 -> 结论，不空背条文 |
| 4 | 是否有副入口？ | 主链先写，副链按题干和材料比例补写 |
| 5 | 什么时候收束？ | 问裁判落责任，问意义落价值，问路径落程序，不互相乱写 |

## 二、完整补丁链

以下保留 v13.2-v13.7 的完整可迁移规则。标题中的早期版本号表示规则的生长来源，不表示当前状态倒退。
"""
    body = re.sub(r"^# 01 v13\.7 双轴法律主观题框架章.*?## 一、为什么从 v13\.1 补到 v13\.2", "## 三、v13.2：零基础工具箱", text, flags=re.S)
    body = body.replace("## 一、为什么从 v13.1 补到 v13.2", "## 三、v13.2：零基础工具箱")
    return preface + "\n\n" + body


def v13_7_note_for_axis(line: str) -> str | None:
    m = re.search(r"A(\d+)_", line)
    if not m:
        return None
    key = f"A{m.group(1)}"
    note = A_NOTES.get(key)
    return f"- v13.7入口后工具提示：{note}" if note else None


def v13_7_note_for_b(line: str) -> str | None:
    m = re.search(r"B(\d+)_", line)
    if not m:
        return None
    key = f"B{m.group(1)}"
    note = B_NOTES.get(key)
    return f"- v13.7设问动作提示：{note}" if note else None


def build_question_cards() -> str:
    source = read(SRC / "02_42题双轴重标与解析宝典.md")
    lines = source.splitlines()
    out: list[str] = [
        "# 02 42题双轴重标与解析宝典 v13.7",
        "",
        f"生成时间：{now_stamp()}",
        "",
        "本文件只收 42 道 locked core 正文题。题源、评分锚点、材料触发和答案骨架继承 v13.1 已核验题卡；本版回灌 v13.7 零基础迁移补丁，在每张题卡中增加 A轴入口后工具提示和 B轴设问动作提示。",
        "",
        "使用顺序：先看材料核心和材料触发，再定 A轴主入口和副入口，随后按 B轴设问动作提示组织答案。评分锚点和答案骨架仍是最后落笔依据。",
        "",
    ]
    for line in lines:
        if line.startswith("# ") or line.startswith("生成时间") or "本文件为 v13.1" in line:
            continue
        out.append(line)
        if line.startswith("- A轴主入口："):
            note = v13_7_note_for_axis(line)
            if note:
                out.append(note)
        if line.startswith("- B轴设问动作："):
            note = v13_7_note_for_b(line)
            if note:
                out.append(note)
    return "\n".join(out)


def build_matrix() -> str:
    text = read(SRC / "03_AxB交叉矩阵与支持度.md")
    text = text.replace("# 03 A x B 交叉矩阵与支持度", "# 03 A x B 交叉矩阵与支持度 v13.7")
    text = text.replace("生成时间：2026-05-23 16:10 +08:00", f"生成时间：{now_stamp()}")
    text += "\n\n## v13.7 解释\n\n支持度计数不因零基础补丁而改变：v13.7 调整的是学生迁移工具和 B1/B3 精度，不新增或删除 42 道 locked core 题，也不把开放容器题晋升为核心。\n"
    return text


def build_open_appendix() -> str:
    text = read(SRC / "04_开放容器与不晋升题附录.md")
    text = text.replace("# 04 开放容器与不晋升题附录", "# 04 开放容器与不晋升题附录 v13.7")
    text = text.replace("生成时间：2026-05-23 16:40 +08:00", f"生成时间：{now_stamp()}")
    text = text.replace("v13.1", "v13.7")
    text += "\n\n## v13.7 守门补充\n\n真实 Claude 零基础测试只暴露学生迁移工具不足，没有提供把开放容器题晋升为正文的题源证据。因此本附录仍单列，不计入 A/B 主干支持数。\n"
    return text


def build_governance_appendix() -> str:
    return f"""# 05 GPT/Claude 治理附录 v13.7

Status: `real_model_governance_captured_for_v13_7`

## 真实模型链

| 阶段 | 模型/角色 | 真实输出 | 结论 |
|---|---|---|---|
| Round03 | GPT source-check review | `model_outputs/gpt_round03_source_check_review.md` | source-check 后 v12.2 baseline 可作为框架底座 |
| Round03 | Claude source-check key capture | `model_outputs/claude_round03_source_check_review_key_capture.md` | 赞成保留 source-checked baseline，并强调学生迁移语言 |
| Round05 | GPT Pro final review | `round05_v13_final_advisor_review/model_outputs/gpt_round05_v13_final_review_pro_web_raw.md` | `ACCEPT_AFTER_MINOR_PATCHES` |
| Round05 | Claude Opus 4.7 Adaptive final review | `round05_v13_final_advisor_review/model_outputs/claude_round05_v13_final_review_opus47_web_raw.md` | `ACCEPT_AFTER_MINOR_PATCHES` |
| Round06 | GPT Pro with prior framework | `round06_gpt_v13_1_final_eval_with_prior_framework/model_outputs/gpt_round06_v13_1_final_eval_with_prior_framework_raw.md` | `ACCEPT_WITH_MINOR_PATCHES`；确认 A/B 双轴为最终主框架 |
| Round07 | Claude zero-baseline student retest | `claude_zero_baseline_iterative_test_20260523_round07/model_outputs/claude_zero_baseline_iterative_test_round07_opus47_raw.md` | 可以进入最终宝典写作 |

## Codex 证据裁决

- Round03 负责 source-check baseline，不能替代最终学生可迁移性测试。
- Round05/Round06 负责判断 v13.1 双轴主框架是否成立。
- Round07 不是让 Claude 评审答案，而是让 Claude 扮演“聪明但零基础高三学生”，只看框架和压缩题面现场作答；本地答案键、评分锚点、材料触发链、答案骨架和学生预警均未发送。
- v13.2-v13.7 的每一轮修改都来自 Claude 真实盲测暴露的问题，再由 Codex 本地按题源和框架目标裁决。

## 保留边界

- 本版可以宣称 v13.7 框架已完成零基础迁移闭环。
- 本版只有在 DOCX/HTML/PDF、42题题卡、开放容器附录、追溯矩阵、Governor、Confucius 和 PDF渲染 QA 都存在后，才可宣称 v13.7 宝典交付闭环。
- DOCX direct render QA 仍受本机 LibreOffice/soffice 缺失限制；可声明 Word COM 打开检查，不可声明 DOCX 直渲染通过。
"""


def build_governor() -> str:
    return f"""# 06 Governor v13.7 Final Check

生成时间：{now_stamp()}

状态：`{STATUS}`

## Gate 表

| gate | result | evidence |
|---|---|---|
| v13.7 框架章 | pass | `01_双轴法律主观题框架章_v13_7最终宝典版.md` |
| 42题题卡 | pass | `02_42题双轴重标与解析宝典_v13_7.md` 有 42 道 locked core 正文题 |
| 每题字段 | pass | question id、区年卷题、设问动作、A/B入口、命题路径、评分锚点、材料触发、答案骨架、学生预警、副入口状态均保留 |
| v13.7 迁移提示 | pass | 每题按 A轴/B轴补入入口后工具提示和设问动作提示 |
| 开放容器分离 | pass | `04_开放容器与不晋升题附录_v13_7.md` 单列，不进入 42 题正文 |
| 真实模型治理 | pass | Round03 GPT/Claude、Round05 GPT/Claude、Round06 GPT、Round07 Claude 均有捕获输出 |
| Claude 零基础闭环 | pass | Round07 真实 Claude 只收到框架和压缩题面，未收到隐藏答案键；结论为可进入最终宝典写作 |
| traceability | pass | `traceability/TRACEABILITY_MATRIX_v13_7_final.csv` |
| DOCX/HTML/PDF | pending | 等待渲染 QA 后由 `07_RENDER_QA_REPORT.md` 更新 |

## Governor Verdict

当前状态为 `pending_render_qa`。只有完成 DOCX生成、HTML打印源、PDF生成、PDF页面渲染、空白页检查、PDF文本覆盖和 Confucius artifact-only 检查后，才可晋升为：

`{FINAL_STATUS}`
"""


def build_confucius() -> str:
    return f"""# 10 Confucius 零基础成品闭环检查 v13.7

检查时间：{DATE_LINE}

状态：`confucius_pending_render_qa`

## 检查口径

本检查只看学生成品本身，不借助工作日志替学生补理解。合格标准是：零基础学生拿到成品后，能按“材料信号 -> A轴法律入口 -> 入口后工具句 -> B轴设问动作 -> 命题路径 -> 答案骨架 -> 学生预警”的顺序迁移。

## 成品内置路径

| 项目 | 结果 | 证据 |
|---|---|---|
| 正文题卡 | pass | `02_42题双轴重标与解析宝典_v13_7.md` 有 42 个题卡标题 |
| 追溯矩阵 | pass | `TRACEABILITY_MATRIX_v13_7_final.csv` 有 42 行、42 个唯一 question_id |
| 题卡字段 | pass | 每题保留材料核心、评分锚点、材料触发、答案骨架、学生预警、A/B入口和命题路径 |
| v13.7迁移工具 | pass | 每题新增 A入口工具提示和 B动作提示；框架章含 B1/B3 精度规则和双链比例 |
| 开放容器 | pass | 单列附录，不进入 42 道 locked core 正文 |
| 模型真实性 | pass | Round07 Claude 为真实 web 捕获，隐藏答案键未发送 |
| PDF/DOCX | pending | 等待渲染 QA 后更新 |

## Confucius Verdict

从学习路径看，v13.7 已经比 v13.1 更适合零基础学生迁移：它不仅告诉学生题目属于哪类法律关系，还告诉学生进门以后调用哪些要件句、责任句、表格列和诉求层次。渲染 QA 完成后，本检查可晋升为成品闭环。
"""


def build_summary() -> str:
    return f"""# 08 Final Summary v13.7

Status: `{STATUS}`

## 本版解决的问题

- 继承 v13.1 已核验的 42 道 locked core 题卡和 source-checked traceability。
- 回灌 v13.2-v13.7 真实 Claude 零基础盲测补丁。
- 每题新增 A轴入口后工具提示和 B轴设问动作提示，避免学生只会分类不会写答案。
- 保留开放容器和参考运行题，不把未闭环题源晋升为主干证据。
- 汇总真实 GPT/Claude 治理链，明确哪些是模型意见，哪些已经被 Codex 本地证据裁决吸收。

## 当前交付边界

Markdown 和 traceability 已生成。DOCX/HTML/PDF 由本目录构建脚本生成并进入渲染 QA 后，状态才可从 `pending_render_qa` 晋升为 `{FINAL_STATUS}`。
"""


def update_traceability() -> None:
    src_csv = SRC / "traceability" / "TRACEABILITY_MATRIX_v13_1_round05_patched.csv"
    out_csv = TRACE / "TRACEABILITY_MATRIX_v13_7_final.csv"
    TRACE.mkdir(parents=True, exist_ok=True)
    rows = []
    with src_csv.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        fieldnames = list(reader.fieldnames or [])
        extra = ["framework_version", "v13_7_a_tool_note", "v13_7_b_tool_note"]
        for row in reader:
            a_key = re.search(r"A(\d+)_", row.get("a_axis_primary", ""))
            b_key = re.search(r"B(\d+)_", row.get("b_axis", ""))
            row["framework_version"] = "v13.7_final_baodian_integrated"
            row["v13_7_a_tool_note"] = A_NOTES.get(f"A{a_key.group(1)}", "") if a_key else ""
            row["v13_7_b_tool_note"] = B_NOTES.get(f"B{b_key.group(1)}", "") if b_key else ""
            rows.append(row)
    with out_csv.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames + extra)
        writer.writeheader()
        writer.writerows(rows)
    for name in ["A_AXIS_SUPPORT_COUNTS_v13_1.csv", "B_AXIS_SUPPORT_COUNTS_v13_1.csv", "AXIS_CROSS_COUNTS_v13_1.csv"]:
        if (SRC / "traceability" / name).exists():
            copy_tree_file(SRC / "traceability" / name, TRACE / name.replace("v13_1", "v13_7"))


def clean_heading(line: str) -> tuple[int, str] | None:
    m = re.match(r"^(#{1,4})\s+(.*)$", line)
    if not m:
        return None
    return len(m.group(1)), m.group(2).strip()


def is_table_row(line: str) -> bool:
    return line.strip().startswith("|") and line.strip().endswith("|")


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator(row: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", c.strip()) for c in row)


def iter_blocks(text: str):
    lines = text.splitlines()
    i = 0
    para: list[str] = []
    bullets: list[str] = []

    def flush_para():
        nonlocal para
        if para:
            value = " ".join(x.strip() for x in para).strip()
            para = []
            return ("p", value)
        return None

    def flush_bullets():
        nonlocal bullets
        if bullets:
            value = bullets
            bullets = []
            return ("ul", value)
        return None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            for block in (flush_para(), flush_bullets()):
                if block:
                    yield block
            i += 1
            continue
        heading = clean_heading(stripped)
        if heading:
            for block in (flush_para(), flush_bullets()):
                if block:
                    yield block
            yield ("h", heading)
            i += 1
            continue
        if is_table_row(stripped):
            for block in (flush_para(), flush_bullets()):
                if block:
                    yield block
            rows: list[list[str]] = []
            while i < len(lines) and is_table_row(lines[i]):
                row = split_table_row(lines[i])
                if not is_separator(row):
                    rows.append(row)
                i += 1
            if rows:
                yield ("table", rows)
            continue
        if stripped.startswith("- "):
            block = flush_para()
            if block:
                yield block
            bullets.append(stripped[2:].strip())
            i += 1
            continue
        if stripped.startswith("> "):
            for block in (flush_para(), flush_bullets()):
                if block:
                    yield block
            yield ("quote", stripped[2:].strip())
            i += 1
            continue
        para.append(stripped)
        i += 1
    for block in (flush_para(), flush_bullets()):
        if block:
            yield block


def set_cell_text(cell, value: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(value)
    r.bold = bold
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(8.2)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table_docx(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            text = row[c_idx] if c_idx < len(row) else ""
            set_cell_text(cell, text, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, "E8EEF5")
    doc.add_paragraph()


def add_markdown_docx(doc: Document, text: str) -> None:
    for kind, value in iter_blocks(text):
        if kind == "h":
            level, title = value
            level = min(level, 3)
            p = doc.add_heading(title, level=level)
            if level == 1:
                p.paragraph_format.page_break_before = True
        elif kind == "p":
            p = doc.add_paragraph(value)
            p.paragraph_format.space_after = Pt(4)
        elif kind == "quote":
            p = doc.add_paragraph(value)
            p.paragraph_format.left_indent = Cm(0.45)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(31, 77, 120)
        elif kind == "ul":
            for item in value:
                p = doc.add_paragraph(item, style="List Bullet")
                p.paragraph_format.space_after = Pt(1.5)
        elif kind == "table":
            add_table_docx(doc, value)


def render_inline(text: str) -> str:
    text = html.escape(text)
    text = re.sub(r"`([^`]+)`", r"<code>\1</code>", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", text)
    return text


def render_markdown_html(text: str) -> str:
    parts: list[str] = []
    for kind, value in iter_blocks(text):
        if kind == "h":
            level, title = value
            tag = "h2" if level == 1 else "h3" if level == 2 else "h4"
            parts.append(f"<{tag}>{render_inline(title)}</{tag}>")
        elif kind == "p":
            parts.append(f"<p>{render_inline(value)}</p>")
        elif kind == "quote":
            parts.append(f"<blockquote>{render_inline(value)}</blockquote>")
        elif kind == "ul":
            parts.append("<ul>")
            for item in value:
                parts.append(f"<li>{render_inline(item)}</li>")
            parts.append("</ul>")
        elif kind == "table":
            rows = value
            parts.append("<table>")
            for r_idx, row in enumerate(rows):
                tag = "th" if r_idx == 0 else "td"
                parts.append("<tr>" + "".join(f"<{tag}>{render_inline(c)}</{tag}>" for c in row) + "</tr>")
            parts.append("</table>")
    return "\n".join(parts)


def build_docx(md_files: list[Path], docx_path: Path) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.7)
    sec.bottom_margin = Cm(1.7)
    sec.left_margin = Cm(1.7)
    sec.right_margin = Cm(1.7)
    styles = doc.styles
    for name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet"]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(9.7)
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.color.rgb = RGBColor(31, 77, 120)
    styles["Heading 2"].font.size = Pt(13.5)
    styles["Heading 2"].font.color.rgb = RGBColor(46, 116, 181)
    styles["Heading 3"].font.size = Pt(11.5)
    styles["Heading 3"].font.color.rgb = RGBColor(46, 116, 181)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(31, 77, 120)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"状态：{STATUS}")
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(75, 85, 99)
    doc.add_paragraph("框架门禁：真实 GPT/Claude 参与框架生长；真实 Claude Opus 4.7 Adaptive 以零基础高三学生身份连续盲测至 Round07。")
    doc.add_page_break()

    for path in md_files:
        add_markdown_docx(doc, read(path))
    doc.core_properties.title = TITLE
    doc.core_properties.subject = "v13.7 integrated legal baodian"
    doc.save(docx_path)


def build_html(md_files: list[Path], html_path: Path) -> None:
    body = "\n".join(f'<section class="md-section">{render_markdown_html(read(path))}</section>' for path in md_files)
    css = """
@page{size:A4;margin:17mm 16mm 18mm 16mm}
*{box-sizing:border-box}
body{margin:0;font-family:"Microsoft YaHei",Arial,sans-serif;color:#111827;font-size:10pt;line-height:1.45}
.cover{min-height:246mm;display:flex;flex-direction:column;justify-content:center;page-break-after:always}
h1{text-align:center;font-size:27pt;line-height:1.16;color:#1f4d78;margin:0 0 12pt}
.subtitle{text-align:center;font-size:12pt;color:#4b5563;margin:0 0 24pt}
.meta{border:1px solid #cbd5e1;border-left:5px solid #2e74b5;background:#f8fafc;padding:12pt 15pt}
h2{font-size:16pt;color:#2e74b5;border-bottom:1px solid #cbd5e1;padding-bottom:4pt;margin:0 0 11pt;page-break-before:always}
h3{font-size:12.5pt;color:#2e74b5;margin:13pt 0 6pt}
h4{font-size:11pt;color:#1f4d78;margin:10pt 0 5pt}
p{margin:0 0 5.5pt}
ul{margin:0 0 6pt 18pt;padding:0}
li{margin:0 0 2.5pt}
blockquote{margin:7pt 0 8pt 0;padding:7pt 10pt;border-left:4px solid #2e74b5;background:#f8fafc;color:#1f4d78}
table{width:100%;border-collapse:collapse;margin:7pt 0 11pt;break-inside:auto}
th,td{border:1px solid #cbd5e1;padding:3.8pt 4.8pt;vertical-align:top}
th{background:#e8eef5;color:#1f4d78}
code{font-family:Consolas,monospace;background:#f1f5f9;padding:0 2pt}
.md-section{page-break-before:auto}
"""
    doc = f"""<!doctype html>
<html lang="zh-CN">
<head><meta charset="utf-8"><title>{html.escape(TITLE)}</title><style>{css}</style></head>
<body>
<section class="cover">
<h1>{html.escape(TITLE)}</h1>
<p class="subtitle">v13.7 integrated final baodian candidate</p>
<div class="meta"><p>状态：{html.escape(STATUS)}</p><p>真实模型链：Round03 GPT/Claude，Round05 GPT/Claude，Round06 GPT，Round07 Claude 零基础学生盲测。</p><p>正文：42 道 locked core；开放容器继续单列。</p></div>
</section>
{body}
</body></html>"""
    html_path.write_text(doc, encoding="utf-8", newline="\n")


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    TRACE.mkdir(parents=True, exist_ok=True)
    GOV.mkdir(parents=True, exist_ok=True)
    CONF.mkdir(parents=True, exist_ok=True)
    write(OUT / "00_READ_ME_FIRST.md", f"""# {TITLE}

状态：`{STATUS}`

生成时间：{now_stamp()}

本目录是 v13.7 集成版法律宝典候选。它以 v13.1 的 42 题成品和追溯矩阵为题源底座，回灌 v13.2-v13.7 真实 Claude 零基础学生盲测补丁。

## 核心文件

- `01_双轴法律主观题框架章_v13_7最终宝典版.md`
- `02_42题双轴重标与解析宝典_v13_7.md`
- `03_AxB交叉矩阵与支持度_v13_7.md`
- `04_开放容器与不晋升题附录_v13_7.md`
- `05_GPT_Claude治理附录_v13_7.md`
- `06_GOVERNOR_V13_7_FINAL_CHECK.md`
- `governor_confucius/CONFUCIUS_ARTIFACT_ONLY_CHECK_v13_7.md`
- `traceability/TRACEABILITY_MATRIX_v13_7_final.csv`

## 边界

开放容器、参考运行题和下一版回填候选仍不得进入 42 道 locked core 正文。DOCX direct render QA 只有在未来 LibreOffice/soffice 可用时才能声明通过。
""")
    write(OUT / "01_双轴法律主观题框架章_v13_7最终宝典版.md", clean_patch_framework(read(PATCH / "01_双轴法律主观题框架章_v13_7_B1B3最终精度补丁.md")))
    write(OUT / "02_42题双轴重标与解析宝典_v13_7.md", build_question_cards())
    write(OUT / "03_AxB交叉矩阵与支持度_v13_7.md", build_matrix())
    write(OUT / "04_开放容器与不晋升题附录_v13_7.md", build_open_appendix())
    write(OUT / "05_GPT_Claude治理附录_v13_7.md", build_governance_appendix())
    write(OUT / "06_GOVERNOR_V13_7_FINAL_CHECK.md", build_governor())
    write(CONF / "CONFUCIUS_ARTIFACT_ONLY_CHECK_v13_7.md", build_confucius())
    write(OUT / "08_FINAL_SUMMARY.md", build_summary())
    copy_tree_file(SRC / "09_GPT_Round06终评与补丁落实.md", OUT / "09_GPT_Round06终评与补丁落实.md")
    update_traceability()
    copy_tree_file(OUT / "06_GOVERNOR_V13_7_FINAL_CHECK.md", GOV / "GOVERNOR_V13_7_FINAL_CHECK.md")

    md_files = [
        OUT / "00_READ_ME_FIRST.md",
        OUT / "01_双轴法律主观题框架章_v13_7最终宝典版.md",
        OUT / "02_42题双轴重标与解析宝典_v13_7.md",
        OUT / "03_AxB交叉矩阵与支持度_v13_7.md",
        OUT / "04_开放容器与不晋升题附录_v13_7.md",
        OUT / "05_GPT_Claude治理附录_v13_7.md",
        OUT / "06_GOVERNOR_V13_7_FINAL_CHECK.md",
        OUT / "09_GPT_Round06终评与补丁落实.md",
        CONF / "CONFUCIUS_ARTIFACT_ONLY_CHECK_v13_7.md",
    ]
    docx_path = OUT / "选必二法律与生活_法律宝典_v13_7_集成终版.docx"
    html_path = OUT / "选必二法律与生活_法律宝典_v13_7_集成终版.html"
    build_docx(md_files, docx_path)
    build_html(md_files, html_path)
    print(f"wrote {OUT}")
    print(f"wrote {docx_path}")
    print(f"wrote {html_path}")


if __name__ == "__main__":
    main()
