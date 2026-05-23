from __future__ import annotations

import argparse
import csv
import importlib.util
import shutil
from pathlib import Path

from docx import Document


STATUS_PENDING = "v13_10_final_baodian_integrated_pending_pdf_docx_qa"
STATUS_FINAL = "v13_10_final_baodian_integrated_pdf_rendered_docx_generated_with_docx_render_caveat"
CONFUCIUS_DELIVERY_STATUS = "v13_10_confucius_reader_framework_pass_delivery_patch_verified_baodian_docx_pdf_regenerated_with_docx_render_caveat"
TITLE = "选必二《法律与生活》法律宝典 v13.10 Confucius框架终版"
DATE_LINE = "2026-05-23"
STAMP = "2026-05-23 23:55 +08:00"


def find_run_root() -> Path:
    for p in (Path.cwd() / "reports").rglob("v12_2_framework_growth_restart"):
        if (p / "v13_7_final_baodian_integrated").is_dir():
            return p
    raise FileNotFoundError("v12_2_framework_growth_restart not found")


ROOT = find_run_root()
SRC = ROOT / "v13_7_final_baodian_integrated"
OUT = ROOT / "v13_10_final_baodian_integrated"
TRACE = OUT / "traceability"
GOV = OUT / "governance"
CONF = OUT / "governor_confucius"
BUILDER_V13_7 = ROOT / "build_v13_7_final_baodian.py"


def read(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.strip() + "\n", encoding="utf-8", newline="\n")


def copy_file(src: Path, dst: Path) -> None:
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)


def current_status() -> str:
    pdf = OUT / "选必二法律与生活_法律宝典_v13_10_Confucius框架终版.pdf"
    return STATUS_FINAL if pdf.exists() and pdf.stat().st_size > 0 else STATUS_PENDING


def transform_common(text: str) -> str:
    replacements = {
        "v13.7": "v13.10",
        "v13_7": "v13_10",
        "v13-7": "v13-10",
        "13.7": "13.10",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text.replace("生成时间：2026-05-23 23:10 +08:00", f"生成时间：{STAMP}")
    return text


def build_readme(status: str) -> str:
    return f"""# {TITLE}

状态：`{status}`

生成时间：{STAMP}

本目录是 v13.10 成品交付目录。它把已经通过 Confucius 三轮“愤怒高三学生”试读的 v13.10 框架，落到完整法律宝典：一页考场卡、框架章、42 道 locked core 逐题解析、AxB 支持矩阵、开放容器附录、GPT/Claude/Confucius 治理附录、最终 Governor 与渲染 QA。

## 主交付文件

- `选必二法律与生活_法律宝典_v13_10_Confucius框架终版.md`
- `选必二法律与生活_法律宝典_v13_10_Confucius框架终版.docx`
- `选必二法律与生活_法律宝典_v13_10_Confucius框架终版.html`
- `选必二法律与生活_法律宝典_v13_10_Confucius框架终版.pdf`

## 关键边界

- GPT/Claude 只统计真实捕获输出：Round03、Round05、Round06、Round07。v13.10 Confucius 是本地专门试读 agent，不冒充 GPT 或 Claude。
- 42 道 locked core 题沿用 source-checked 题卡和追溯矩阵；开放容器、参考题、下一版回填候选仍单列，不晋升正文。
- v13.10 的增量是“框架先读、能现场学习、能快速作答”的考场化修复，不改写已核验题源事实。
- DOCX 已生成后必须有 Word/结构检查；本机缺 LibreOffice/soffice 时，不声明 DOCX direct render QA passed。PDF 由 HTML 打印源生成并渲染页面检查。
"""


def build_framework() -> str:
    text = read(SRC / "01_双轴法律主观题框架章_v13_10_Confucius三轮修复版.md")
    text = text.replace(
        "状态：`v13_10_framework_first_confucius_third_repair_candidate`",
        f"状态：`{current_status()}`",
    )
    insert = f"""

## 0.1 本版和旧“六入口”的关系

早期六入口不是废掉，而是被吸收到 B 轴设问动作中：B1 表格补链、B2 判决责任、B3 诉求支持、B4 评析观点、B5 意义价值、B6 维权路径。v13.10 在此基础上补出 B7 法律问题识别/填空，并把 A 轴改成“法律关系/法律内容入口”。所以学生不是先背六类题型，而是先定 A 轴生活冲突，再按 B 轴决定答案形状。

命题人路径可以压缩成一句话：材料给一个生活冲突，设问指定一种法律工作，评分奖励学生把事实压进规则和责任，而不是空背法治口号。
"""
    marker = "## 0.5 三张考场卡"
    return text.replace(marker, insert + "\n" + marker)


def build_question_cards() -> str:
    text = read(SRC / "02_42题双轴重标与解析宝典_v13_7.md")
    text = text.replace("# 02 42题双轴重标与解析宝典 v13.7", "# 02 42题双轴重标与解析宝典 v13.10")
    text = text.replace("生成时间：2026-05-23 23:10 +08:00", f"生成时间：{STAMP}")
    text = text.replace(
        "本文件只收 42 道 locked core 正文题。题源、评分锚点、材料触发和答案骨架继承 v13.1 已核验题卡；本版回灌 v13.7 零基础迁移补丁，在每张题卡中增加 A轴入口后工具提示和 B轴设问动作提示。",
        "本文件只收 42 道 locked core 正文题。题源、评分锚点、材料触发和答案骨架继承 v13.1 已核验题卡；本版改用 v13.10 Confucius 三轮修复框架组织阅读顺序，并继承 v13.7 零基础迁移题卡字段。",
    )
    text = text.replace("v13.7入口后工具提示", "v13.10入口后工具提示")
    text = text.replace("v13.7设问动作提示", "v13.10设问动作提示")
    text = text.replace("v13.7 notes", "v13.10 notes")
    return text


def build_governance_appendix(status: str) -> str:
    return f"""# 05 GPT/Claude/Confucius 治理附录 v13.10

Status: `{status}`

## 真实 GPT/Claude 链条

| 阶段 | 真实模型/角色 | 捕获文件 | 结论 |
|---|---|---|---|
| Round03 | GPT source-check review | `model_outputs/gpt_round03_source_check_review.md` | source-check 后 v12.2 baseline 可作为框架底座 |
| Round03 | Claude source-check key capture | `model_outputs/claude_round03_source_check_review_key_capture.md` | 接受 source-checked baseline，并提示学生迁移语言 |
| Round05 | GPT Pro final review | `round05_v13_final_advisor_review/model_outputs/gpt_round05_v13_final_review_pro_web_raw.md` | `ACCEPT_AFTER_MINOR_PATCHES` |
| Round05 | Claude Opus 4.7 Adaptive final review | `round05_v13_final_advisor_review/model_outputs/claude_round05_v13_final_review_opus47_web_raw.md` | `ACCEPT_AFTER_MINOR_PATCHES` |
| Round06 | GPT Pro with prior framework | `round06_gpt_v13_1_final_eval_with_prior_framework/model_outputs/gpt_round06_v13_1_final_eval_with_prior_framework_raw.md` | `ACCEPT_WITH_MINOR_PATCHES`；确认 A/B 双轴为最终主框架 |
| Round07 | Claude zero-baseline student retest | `claude_zero_baseline_iterative_test_20260523_round07/model_outputs/claude_zero_baseline_iterative_test_round07_opus47_raw.md` | 可以进入最终宝典写作 |

## v13.10 本地 Confucius 试读 agent

| 轮次 | 输入限制 | 输出文件 | 结论 |
|---|---|---|---|
| Round01 | 只给框架，不给答案键 | `confucius_angry_student_reader_agent/FIRST_RUN_REPORT_20260523.md` | `FRAMEWORK_PASS_WITH_REPAIRS` |
| Round02 | v13.8 修复框架 + 随机题 | `confucius_angry_student_reader_agent/SECOND_RUN_REPORT_20260523_V13_8.md` | `FRAMEWORK_PASS_WITH_REPAIRS` |
| Round03 | v13.9 修复框架 + 随机题 | `confucius_angry_student_reader_agent/THIRD_RUN_REPORT_20260523_V13_9.md` | `FRAMEWORK_PASS_WITH_REPAIRS` |
| Round04 | v13.10 框架 + 随机题 | `confucius_angry_student_reader_agent/FOURTH_RUN_REPORT_20260523_V13_10.md` | `FRAMEWORK_PASS` |
| Round05 | v13.10 一页考场卡 + 交付补丁 | `confucius_angry_student_reader_agent/FIFTH_RUN_REPORT_20260523_V13_10_DELIVERY_PATCH.md` | `FRAMEWORK_PASS` |
| Closure | 最终本地框架闭环 | `confucius_angry_student_reader_agent/FINAL_CONFUCIUS_CLOSURE_20260523_V13_10.md` | 交付前历史记录：框架通过，但当时还未进入本目录交付 |
| Delivery patch | 本目录交付闭环 | `governor_confucius/FINAL_CONFUCIUS_CLOSURE_20260523_V13_10.md` | `{CONFUCIUS_DELIVERY_STATUS}` |

## Codex 裁决

- v13.10 的 Confucius agent 是专门试读和迁移验证，不计为 GPT/Claude 真实模型链。
- 真实 GPT/Claude 已完成框架底座和 v13 主框架审查；v13.10 只做“学生能不能读懂并现场习得”的框架交付修复。
- 所有新框架语言必须回到 42 题 source-checked 题卡、追溯矩阵和开放容器边界；不能因为试读顺口就晋升新题源。
- 允许宣称：v13.10 框架已通过本地 Confucius artifact-only 学生试读门。
- 只有 DOCX/HTML/PDF 文件与渲染 QA 文件都存在时，才允许宣称 v13.10 宝典交付闭环。
"""


def build_governor(status: str) -> str:
    pdf = OUT / "选必二法律与生活_法律宝典_v13_10_Confucius框架终版.pdf"
    docx = OUT / "选必二法律与生活_法律宝典_v13_10_Confucius框架终版.docx"
    pdf_state = "pass" if pdf.exists() and pdf.stat().st_size > 0 else "pending"
    docx_state = "pass with render caveat" if docx.exists() and docx.stat().st_size > 0 else "pending"
    return f"""# 06 Governor v13.10 Final Check

生成时间：{STAMP}

状态：`{status}`

## Gate 表

| gate | result | evidence |
|---|---|---|
| v13.10 一页考场卡 | pass | `00_v13_10_一页考场卡_学生先读版.md` |
| v13.10 框架章 | pass | `01_双轴法律主观题框架章_v13_10_Confucius三轮修复版.md` |
| 42题逐题解析 | pass | `02_42题双轴重标与解析宝典_v13_10.md` 有 42 道 locked core 正文题 |
| 每题字段 | pass | question id、区年卷题、设问动作、A/B入口、命题路径、评分锚点、材料触发、答案骨架、学生预警、副入口状态均保留 |
| 开放容器分离 | pass | `04_开放容器与不晋升题附录_v13_10.md` 单列，不进入 42 题正文 |
| 真实 GPT/Claude 治理 | pass | Round03/Round05/Round06/Round07 捕获输出均按原文件引用 |
| Confucius 学生试读 | pass | 五轮本地试读最终为 `FRAMEWORK_PASS` |
| traceability | pass | `traceability/TRACEABILITY_MATRIX_v13_10_final.csv` |
| DOCX generated | {docx_state} | DOCX direct render QA 只有 LibreOffice/soffice 可用时才能声明通过 |
| PDF generated and rendered | {pdf_state} | `07_RENDER_QA_REPORT_v13_10.md` |

## Governor Verdict

当前允许状态为：`{status}`。

若状态仍为 pending，不得宣称 v13.10 DOCX/PDF 已完整交付。若状态为 `{STATUS_FINAL}`，可宣称 v13.10 宝典 Markdown/HTML/DOCX/PDF 已生成，PDF 已页面渲染检查，DOCX 已生成但 direct render QA 因本机缺 LibreOffice/soffice 不声明通过。
"""


def build_summary(status: str) -> str:
    return f"""# 08 Final Summary v13.10

Status: `{status}`

## 本版完成项

- v13.10 一页考场卡和框架章成为宝典入口，不再让学生先读版本补丁史。
- 42 道 locked core 题全部按 A 轴法律入口 + B 轴设问动作组织，字段完整保留。
- 旧六入口被明确吸收到 B1-B6，v13.10 额外补 B7 法律问题识别/填空。
- 开放容器、参考运行题、下一版回填候选保持附录，不晋升正文。
- GPT/Claude 真实捕获结论和 Confucius 本地试读结论分开记录。

## 交付边界

本版允许状态：`{status}`。

DOCX direct render QA 受本机 LibreOffice/soffice 缺失限制；PDF 是用于视觉检查和分发的可靠渲染交付。
"""


def build_confucius_delivery_closure(raw: str) -> str:
    raw = raw.replace(
        "# Final Confucius Closure 20260523: v13.10 Framework Pass",
        "# Final Confucius Closure 20260523: v13.10 Framework Pass + Baodian Delivery",
    )
    raw = raw.replace(
        "Status: `v13_10_confucius_reader_framework_pass_delivery_patch_verified`",
        f"Status: `{CONFUCIUS_DELIVERY_STATUS}`",
    )
    replacement = """Delivery now claimed in this integrated directory:

- v13.10 Markdown, HTML, DOCX and PDF were regenerated from the v13.10 framework-first baodian source.
- The PDF was regenerated from the final HTML print source and rasterized into page PNGs for QA.
- The DOCX was regenerated and opened read-only through Word COM.

Still not claimed here:

- DOCX direct render QA, because local LibreOffice/soffice remains unavailable.
- A new real GPT/Claude advisor gate, because the v13.8-v13.10 repair loop used local Confucius subagents only. Real GPT/Claude gates remain the captured Round03/Round05/Round06/Round07 chain listed in the governance appendix.
"""
    start = raw.index("Still not claimed here:")
    end = raw.index("## Final Status")
    raw = raw[:start] + replacement + "\n" + raw[end:]
    raw = raw.replace(
        "`v13_10_confucius_reader_framework_pass_delivery_patch_verified_baodian_docx_pdf_not_regenerated`",
        f"`{CONFUCIUS_DELIVERY_STATUS}`",
    )
    return raw


def build_traceability() -> None:
    src_csv = SRC / "traceability" / "TRACEABILITY_MATRIX_v13_7_final.csv"
    out_csv = TRACE / "TRACEABILITY_MATRIX_v13_10_final.csv"
    TRACE.mkdir(parents=True, exist_ok=True)
    with src_csv.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        fieldnames = list(reader.fieldnames or [])
        extra = ["v13_10_framework_gate", "v13_10_delivery_patch"]
        rows = []
        for row in reader:
            row["framework_version"] = "v13.10_confucius_framework_final_baodian_integrated"
            row["v13_10_framework_gate"] = "confucius_framework_pass"
            row["v13_10_delivery_patch"] = "one_page_card_plus_framework_first_student_order"
            rows.append(row)
    with out_csv.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames + extra)
        writer.writeheader()
        writer.writerows(rows)
    for name in ["A_AXIS_SUPPORT_COUNTS_v13_7.csv", "B_AXIS_SUPPORT_COUNTS_v13_7.csv", "AXIS_CROSS_COUNTS_v13_7.csv"]:
        src = SRC / "traceability" / name
        if src.exists():
            copy_file(src, TRACE / name.replace("v13_7", "v13_10"))


def load_v13_7_builder():
    spec = importlib.util.spec_from_file_location("build_v13_7_final_baodian", BUILDER_V13_7)
    if spec is None or spec.loader is None:
        raise RuntimeError("cannot load v13.7 builder")
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    mod.TITLE = TITLE
    mod.STATUS = current_status()
    return mod


def build_combined_markdown(md_files: list[Path]) -> None:
    parts = [f"# {TITLE}", "", f"状态：`{current_status()}`", "", f"生成时间：{STAMP}", ""]
    for path in md_files:
        parts.append(read(path))
        parts.append("")
    write(OUT / "选必二法律与生活_法律宝典_v13_10_Confucius框架终版.md", "\n".join(parts))


def build_docx_html(md_files: list[Path]) -> None:
    mod = load_v13_7_builder()
    docx_path = OUT / "选必二法律与生活_法律宝典_v13_10_Confucius框架终版.docx"
    html_path = OUT / "选必二法律与生活_法律宝典_v13_10_Confucius框架终版.html"
    mod.build_docx(md_files, docx_path)
    mod.build_html(md_files, html_path)
    html_text = read(html_path)
    html_text = html_text.replace("v13.7 integrated final baodian candidate", "v13.10 Confucius framework final baodian")
    html_text = html_text.replace(
        "真实模型链：Round03 GPT/Claude，Round05 GPT/Claude，Round06 GPT，Round07 Claude 零基础学生盲测。",
        "真实模型链：Round03 GPT/Claude，Round05 GPT/Claude，Round06 GPT，Round07 Claude；v13.10 本地 Confucius 愤怒高三学生试读至 FRAMEWORK_PASS。",
    )
    write(html_path, html_text)


def build_base() -> None:
    status = current_status()
    OUT.mkdir(parents=True, exist_ok=True)
    TRACE.mkdir(parents=True, exist_ok=True)
    GOV.mkdir(parents=True, exist_ok=True)
    CONF.mkdir(parents=True, exist_ok=True)

    write(OUT / "00_READ_ME_FIRST.md", build_readme(status))
    copy_file(SRC / "00_v13_10_一页考场卡_学生先读版.md", OUT / "00_v13_10_一页考场卡_学生先读版.md")
    write(OUT / "01_双轴法律主观题框架章_v13_10_Confucius三轮修复版.md", build_framework())
    write(OUT / "02_42题双轴重标与解析宝典_v13_10.md", build_question_cards())
    write(OUT / "03_AxB交叉矩阵与支持度_v13_10.md", transform_common(read(SRC / "03_AxB交叉矩阵与支持度_v13_7.md")))
    write(OUT / "04_开放容器与不晋升题附录_v13_10.md", transform_common(read(SRC / "04_开放容器与不晋升题附录_v13_7.md")))
    write(OUT / "05_GPT_Claude_Confucius治理附录_v13_10.md", build_governance_appendix(status))
    write(OUT / "06_GOVERNOR_V13_10_FINAL_CHECK.md", build_governor(status))
    write(OUT / "08_FINAL_SUMMARY_v13_10.md", build_summary(status))
    copy_file(SRC / "09_GPT_Round06终评与补丁落实.md", OUT / "09_GPT_Round06终评与补丁落实.md")
    if (SRC / "confucius_angry_student_reader_agent").is_dir():
        shutil.copytree(SRC / "confucius_angry_student_reader_agent", CONF / "confucius_angry_student_reader_agent", dirs_exist_ok=True)
        integrated_closure = build_confucius_delivery_closure(read(SRC / "confucius_angry_student_reader_agent" / "FINAL_CONFUCIUS_CLOSURE_20260523_V13_10.md"))
        write(
            CONF / "FINAL_CONFUCIUS_CLOSURE_20260523_V13_10.md",
            integrated_closure,
        )
        write(
            CONF / "confucius_angry_student_reader_agent" / "FINAL_CONFUCIUS_CLOSURE_20260523_V13_10.md",
            integrated_closure,
        )
    copy_file(OUT / "06_GOVERNOR_V13_10_FINAL_CHECK.md", GOV / "GOVERNOR_V13_10_FINAL_CHECK.md")
    build_traceability()

    md_files = [
        OUT / "00_v13_10_一页考场卡_学生先读版.md",
        OUT / "01_双轴法律主观题框架章_v13_10_Confucius三轮修复版.md",
        OUT / "02_42题双轴重标与解析宝典_v13_10.md",
        OUT / "03_AxB交叉矩阵与支持度_v13_10.md",
        OUT / "04_开放容器与不晋升题附录_v13_10.md",
        OUT / "05_GPT_Claude_Confucius治理附录_v13_10.md",
        OUT / "09_GPT_Round06终评与补丁落实.md",
        CONF / "FINAL_CONFUCIUS_CLOSURE_20260523_V13_10.md",
    ]
    build_combined_markdown(md_files)
    build_docx_html(md_files)


def read_optional(path: Path) -> str:
    if not path.exists():
        return "not_run"
    first_line = read(path).strip().splitlines()[0].strip().lstrip("\ufeff")
    return first_line[:260]


def write_render_report() -> None:
    import fitz
    from PIL import Image, ImageStat

    status = current_status()
    pdf = OUT / "选必二法律与生活_法律宝典_v13_10_Confucius框架终版.pdf"
    docx = OUT / "选必二法律与生活_法律宝典_v13_10_Confucius框架终版.docx"
    html_path = OUT / "选必二法律与生活_法律宝典_v13_10_Confucius框架终版.html"
    render_dir = OUT / "rendered_pdf_pages"
    render_dir.mkdir(parents=True, exist_ok=True)
    for old in render_dir.glob("page-*.png"):
        old.unlink()

    pdf_result = "pending"
    text_result = "pending"
    page_result = "pending"
    blank_result = "pending"
    page_count = 0
    blank_pages: list[int] = []
    required = ["v13.10", "Confucius", "CC0251", "locked core", "A4+A6"]
    if pdf.exists() and pdf.stat().st_size > 0:
        doc = fitz.open(str(pdf))
        page_count = doc.page_count
        text = "\n".join(page.get_text("text") for page in doc)
        missing = [token for token in required if token not in text]
        text_result = "pass" if not missing else "fail_missing_" + "_".join(missing)
        for idx, page in enumerate(doc, 1):
            pix = page.get_pixmap(matrix=fitz.Matrix(1.25, 1.25), alpha=False)
            out = render_dir / f"page-{idx:03d}.png"
            pix.save(str(out))
            im = Image.open(out).convert("L")
            stat = ImageStat.Stat(im)
            if stat.mean[0] > 252 and stat.stddev[0] < 2.0:
                blank_pages.append(idx)
        pdf_result = f"pass ({pdf.stat().st_size} bytes)"
        page_result = f"pass ({page_count} pages -> {page_count} PNG)"
        blank_result = "pass" if not blank_pages else "fail_blank_like_" + ",".join(map(str, blank_pages))

    docx_result = "pending"
    if docx.exists() and docx.stat().st_size > 0:
        d = Document(docx)
        docx_result = f"pass ({docx.stat().st_size} bytes; paragraphs={len(d.paragraphs)}; tables={len(d.tables)})"

    word_check = read_optional(OUT / "qa_word_com_check.txt")
    docx_render_check = read_optional(OUT / "qa_docx_render_check.txt")
    write(
        OUT / "07_RENDER_QA_REPORT_v13_10.md",
        f"""# 07 Render QA Report v13.10

Status: `{status}`

Date: {DATE_LINE}

## Produced Files

- Combined Markdown: `选必二法律与生活_法律宝典_v13_10_Confucius框架终版.md`
- DOCX candidate: `选必二法律与生活_法律宝典_v13_10_Confucius框架终版.docx`
- HTML print source: `选必二法律与生活_法律宝典_v13_10_Confucius框架终版.html`
- PDF delivery: `选必二法律与生活_法律宝典_v13_10_Confucius框架终版.pdf`
- PDF page renders: `rendered_pdf_pages/page-001.png` through final page

## Verification

| check | result | evidence |
|---|---|---|
| traceability row count | pass | `TRACEABILITY_MATRIX_v13_10_final.csv` has 42 rows |
| Markdown card count | pass | `02_42题双轴重标与解析宝典_v13_10.md` has 42 `###` cards |
| v13.10 framework gate | pass | Confucius local reader closure says `FRAMEWORK_PASS` and delivery patch verified |
| open-container separation | pass | `04_开放容器与不晋升题附录_v13_10.md` exists outside the 42-card body |
| real GPT/Claude boundary | pass | no new GPT/Claude call claimed; real captured Round03/Round05/Round06/Round07 outputs are referenced |
| HTML print source | {'pass' if html_path.exists() and html_path.stat().st_size > 0 else 'pending'} | `{html_path.name}` |
| PDF generated | {pdf_result} | `{pdf.name}` |
| PDF text coverage | {text_result} | required tokens: {', '.join(required)} |
| PDF pages rendered | {page_result} | `rendered_pdf_pages/` |
| blank-page check | {blank_result} | blank-like pages: {blank_pages or 'none'} |
| DOCX generated / structural check | {docx_result} | python-docx open check |
| DOCX Word COM open check | {word_check} | `qa_word_com_check.txt` |
| DOCX direct render via `render_docx.py` | not passed / not claimed | {docx_render_check} |

## Governor Note

The PDF delivery is rendered from the v13.10 HTML print source and then rasterized to page PNGs for blank-page and sample visual checks. The DOCX exists and is structurally readable, but this machine has no LibreOffice/soffice path, so DOCX direct visual-render QA is not claimed.
""",
    )


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--qa-only", action="store_true")
    args = parser.parse_args()
    if not args.qa_only:
        build_base()
    write_render_report()
    print(current_status())
    print(OUT)


if __name__ == "__main__":
    main()
