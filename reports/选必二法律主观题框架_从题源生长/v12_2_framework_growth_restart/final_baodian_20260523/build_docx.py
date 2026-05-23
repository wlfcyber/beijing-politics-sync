from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import generate_baodian as data


OUT = Path(__file__).resolve().parent
DOCX_PATH = OUT / "选必二法律与生活_法律宝典_v12_2.docx"
GENERATED_AT = "2026-05-23 02:10 +08:00"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(17, 24, 39)
MUTED = RGBColor(75, 85, 99)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


def set_east_asia_font(style_or_run, font_name: str = "Microsoft YaHei") -> None:
    rpr = style_or_run._element.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        style_or_run._element.insert(0, rpr)
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    set_east_asia_font(normal, "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        set_east_asia_font(style, "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        set_east_asia_font(style, "Microsoft YaHei")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("飞哥政治庄园｜选必二《法律与生活》v12.2｜Markdown+DOCX 交付")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    set_east_asia_font(run, "Microsoft YaHei")

    props = doc.core_properties
    props.title = "选必二法律与生活法律宝典 v12.2"
    props.subject = "法律与生活主观题框架与42题解析"
    props.author = "Codex, with real GPT and Claude framework council captures"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        # OOXML order is tblPr, tblGrid, then rows. Putting tblGrid before
        # tblPr makes Word's conversion engine crawl or hang on some builds.
        tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_label_paragraph(doc: Document, label: str, text: str, *, compact: bool = False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3 if compact else 5)
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    set_east_asia_font(r, "Microsoft YaHei")
    body = p.add_run(text)
    set_east_asia_font(body, "Microsoft YaHei")
    return p


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("选必二《法律与生活》法律宝典")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = DARK_BLUE
    set_east_asia_font(r, "Microsoft YaHei")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(18)
    r = sub.add_run("v12.2｜从42道题源生长出的主观题框架与逐题解析")
    r.font.size = Pt(13)
    r.font.color.rgb = MUTED
    set_east_asia_font(r, "Microsoft YaHei")

    rows = [
        ("状态", "DOCX baodian rendered candidate; framework and Markdown gates already passed"),
        ("生成时间", GENERATED_AT),
        ("证据底座", "42 locked core 主观题；参考/开放/下一版回填题不晋升正文"),
        ("模型链条", "真实 GPT Round 03 + Claude Round 03 审查；Codex 源检查裁决"),
        ("使用方式", "先用设问动作定入口，再用材料事实压实法律关系，最后回到评分锚点"),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [1800, 7560])
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        row = table.rows[0] if i == 0 else table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        shade_cell(row.cells[0], LIGHT_BLUE)
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    set_east_asia_font(run, "Microsoft YaHei")
            set_cell_margins(cell)
    doc.add_page_break()


def add_counts_table(doc: Document, grouped) -> None:
    doc.add_heading("正文覆盖", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_width(table, [1800, 1100, 6460])
    headers = ["入口", "题数", "适用场景"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        shade_cell(cell, LIGHT_BLUE)
    for code in data.ENTRY_ORDER:
        spec = data.ENTRY_DEFS[code]
        row = table.add_row()
        row.cells[0].text = code
        row.cells[1].text = str(len(grouped[code]))
        row.cells[2].text = spec["cue"]
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    set_east_asia_font(run, "Microsoft YaHei")


def add_framework(doc: Document, grouped) -> None:
    doc.add_heading("第一部分：六入口框架", level=1)
    add_label_paragraph(doc, "总框架：", "选必二主观题不是按教材页码出题，而是按设问动作出题。先看题目让你做什么，再判断命题人要学生完成哪一种法律工作：补链、判责、判诉求、评观点、写意义、走路径。")
    add_label_paragraph(doc, "命题人路径：", "生活化法律场景 -> 材料事实触发点 -> 设问动作 -> 评分锚点。学生反向作答：设问动作定入口，材料触发定知识，评分锚点定落点。")
    add_counts_table(doc, grouped)
    for code in data.ENTRY_ORDER:
        spec = data.ENTRY_DEFS[code]
        rows = grouped[code]
        doc.add_heading(f"{spec['title']}（{len(rows)}题）", level=2)
        add_label_paragraph(doc, "识别口令：", spec["cue"], compact=True)
        add_label_paragraph(doc, "命题路径：", spec["path"], compact=True)
        add_label_paragraph(doc, "作答骨架：", spec["skeleton"], compact=True)
        add_label_paragraph(doc, "学生预警：", spec["warning"], compact=True)
        add_label_paragraph(doc, "本版题源：", "、".join(r["question_id"] for r in rows), compact=True)


def add_question_cards(doc: Document, grouped) -> None:
    doc.add_page_break()
    doc.add_heading("第二部分：42题按框架解析", level=1)
    n = 1
    for code in data.ENTRY_ORDER:
        spec = data.ENTRY_DEFS[code]
        doc.add_heading(f"{spec['title']}（{len(grouped[code])}题）", level=2)
        for row in grouped[code]:
            qid = row["question_id"]
            trace = data.TRACE.get(qid, {})
            doc.add_heading(f"{n}. {qid}", level=3)
            add_label_paragraph(doc, "区年卷题：", f"{data.cell(row, 'year')}年 {data.cell(row, 'district')} {data.cell(row, 'exam_stage')} 第{data.cell(row, 'question_no')}题", compact=True)
            add_label_paragraph(doc, "设问动作：", data.cell(row, "真实设问"), compact=True)
            add_label_paragraph(doc, "材料核心：", data.cell(row, "真实材料核心"), compact=True)
            add_label_paragraph(doc, "框架入口：", f"{spec['title']}；原题源主入口：{data.cell(row, '主入口')}", compact=True)
            add_label_paragraph(doc, "副入口：", data.secondary_entry(row), compact=True)
            add_label_paragraph(doc, "命题路径：", spec["path"], compact=True)
            add_label_paragraph(doc, "细则/评分锚点：", data.scoring_anchor(row), compact=True)
            for idx, trig in enumerate(data.trigger_lines(row), 1):
                add_label_paragraph(doc, f"材料触发{idx}：", trig, compact=True)
            add_label_paragraph(doc, "答案骨架：", data.cell(row, "答案骨架"), compact=True)
            add_label_paragraph(doc, "学生预警：", data.cell(row, "禁止命中") or spec["warning"], compact=True)
            add_label_paragraph(doc, "讲解口径：", data.cell(row, "飞哥想说"), compact=True)
            add_label_paragraph(doc, "证据与来源：", f"{data.cell(row, 'evidence_status')}；{data.cell(row, 'chain_state')}；source_check_state={data.cell(trace, 'source_check_state', 'not_recorded')}", compact=True)
            n += 1


def add_appendices(doc: Document, grouped) -> None:
    doc.add_page_break()
    doc.add_heading("第三部分：开放容器与治理附录", level=1)
    doc.add_heading("A. 不晋升规则", level=2)
    add_label_paragraph(doc, "规则：", "参考题、开放容器、别名风险行、下一版回填候选不能直接提升为第七入口。只有当新证据证明 locked core 正文题无法被六入口覆盖，才允许重开框架结构。")

    doc.add_heading("B. Round 03 明确不晋升行", level=2)
    non_promoted = [
        r for r in data.PENDING_CHECKS
        if not data.cell(r, "decision").startswith("KEEP_CORE")
    ]
    for r in non_promoted:
        add_label_paragraph(
            doc,
            f"{data.cell(r, 'question_id')}：",
            f"{data.cell(r, 'source_status')}；{data.cell(r, 'entrance_after_check')}；{data.cell(r, 'decision')}；{data.cell(r, 'guardrail')}",
            compact=True,
        )

    doc.add_heading("C. GPT / Claude 真实结论", level=2)
    add_label_paragraph(doc, "GPT Round 03：", "ChatGPT web clean conversation，visible mode 为“进阶专业”，保留模型标签 caution；verdict 为 accept_source_checked_candidate_no_structural_change。")
    add_label_paragraph(doc, "Claude Round 03：", "Claude web 可见模型 Opus 4.7 Adaptive；verdict 为 accept_source_checked_candidate_no_structural_change。")
    add_label_paragraph(doc, "Codex 裁决：", "本地源证据优先。双模型一致意见只确认六入口结构可作为 v12.2 source-checked baseline；下一版回填仍不自动进入正文。")

    doc.add_heading("D. Governor 结论", level=2)
    add_label_paragraph(doc, "Markdown gate：", "已通过：框架章、42题解析、参考/开放附录、模型治理附录均已存在。")
    add_label_paragraph(doc, "DOCX/PDF gate：", "本文件为 DOCX 渲染候选；需通过 render_docx.py 生成 PNG/PDF 并完成视觉检查后，才能把 DOCX/PDF delivery 标成 pass。")


def build() -> None:
    grouped = {code: [] for code in data.ENTRY_ORDER}
    for row in data.CORE:
        grouped[data.entry_code(row)].append(row)

    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_framework(doc, grouped)
    add_question_cards(doc, grouped)
    add_appendices(doc, grouped)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
