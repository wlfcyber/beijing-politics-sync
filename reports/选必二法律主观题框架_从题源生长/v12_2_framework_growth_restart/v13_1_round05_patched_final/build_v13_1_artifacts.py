from __future__ import annotations

import html
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def find_run_dir() -> Path:
    for p in (Path.cwd() / "reports").rglob("v12_2_framework_growth_restart"):
        candidate = p / "v13_1_round05_patched_final"
        if candidate.is_dir():
            return candidate
    raise FileNotFoundError("v13_1_round05_patched_final not found")


OUT = find_run_dir()
MD_FILES = [next(OUT.glob(f"{prefix}_*.md")) for prefix in ["01", "02", "03", "04", "05", "06"]]
if list(OUT.glob("09_*.md")):
    MD_FILES.append(next(OUT.glob("09_*.md")))
CONFUCIUS = OUT / "governor_confucius" / "CONFUCIUS_ARTIFACT_ONLY_CHECK_v13_1.md"
if CONFUCIUS.exists():
    MD_FILES.append(CONFUCIUS)
README = OUT / "00_READ_ME_FIRST.md"
DOCX_PATH = next(OUT.glob("*.docx"))
HTML_PATH = next(OUT.glob("*.html"))


def read(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def first_status_line() -> str:
    for line in read(README).splitlines():
        if line.startswith("状态"):
            return line.replace("`", "")
    return ""


def clean_heading(line: str) -> tuple[int, str] | None:
    m = re.match(r"^(#{1,4})\s+(.*)$", line)
    if not m:
        return None
    return len(m.group(1)), m.group(2).strip()


def is_table_row(line: str) -> bool:
    return line.strip().startswith("|") and line.strip().endswith("|")


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator(row: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", c.strip()) for c in row)


def iter_blocks(text: str):
    lines = text.splitlines()
    i = 0
    para: list[str] = []
    bullets: list[str] = []

    def flush_para():
        nonlocal para
        if para:
            value = " ".join(x.strip() for x in para).strip()
            para = []
            return ("p", value)
        return None

    def flush_bullets():
        nonlocal bullets
        if bullets:
            value = bullets
            bullets = []
            return ("ul", value)
        return None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            b = flush_para()
            if b:
                yield b
            b = flush_bullets()
            if b:
                yield b
            i += 1
            continue
        heading = clean_heading(stripped)
        if heading:
            b = flush_para()
            if b:
                yield b
            b = flush_bullets()
            if b:
                yield b
            yield ("h", heading)
            i += 1
            continue
        if is_table_row(stripped):
            b = flush_para()
            if b:
                yield b
            b = flush_bullets()
            if b:
                yield b
            rows: list[list[str]] = []
            while i < len(lines) and is_table_row(lines[i]):
                row = split_table_row(lines[i])
                if not is_separator(row):
                    rows.append(row)
                i += 1
            if rows:
                yield ("table", rows)
            continue
        if stripped.startswith("- "):
            b = flush_para()
            if b:
                yield b
            bullets.append(stripped[2:].strip())
            i += 1
            continue
        if stripped.startswith("> "):
            b = flush_para()
            if b:
                yield b
            b = flush_bullets()
            if b:
                yield b
            yield ("quote", stripped[2:].strip())
            i += 1
            continue
        para.append(stripped)
        i += 1

    b = flush_para()
    if b:
        yield b
    b = flush_bullets()
    if b:
        yield b


def set_cell_text(cell, value: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(value)
    r.bold = bold
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table_docx(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            text = row[c_idx] if c_idx < len(row) else ""
            set_cell_text(cell, text, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, "E8EEF5")
    doc.add_paragraph()


def add_markdown_docx(doc: Document, text: str):
    for kind, value in iter_blocks(text):
        if kind == "h":
            level, title = value
            level = min(level, 3)
            p = doc.add_heading(title, level=level)
            if level == 1:
                p.paragraph_format.page_break_before = True
        elif kind == "p":
            p = doc.add_paragraph(value)
            p.paragraph_format.space_after = Pt(5)
        elif kind == "quote":
            p = doc.add_paragraph(value)
            p.paragraph_format.left_indent = Cm(0.45)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(31, 77, 120)
        elif kind == "ul":
            for item in value:
                p = doc.add_paragraph(item, style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
        elif kind == "table":
            add_table_docx(doc, value)


def build_docx():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)

    styles = doc.styles
    for name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet"]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.color.rgb = RGBColor(31, 77, 120)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 2"].font.color.rgb = RGBColor(46, 116, 181)
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.color.rgb = RGBColor(46, 116, 181)

    title_line = read(README).splitlines()[0].lstrip("# ").strip()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title_line)
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(31, 77, 120)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(first_status_line())
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(75, 85, 99)
    doc.add_paragraph("模型终审：GPT Pro 与 Claude Opus 4.7 Adaptive 均为 ACCEPT_AFTER_MINOR_PATCHES；本版已落实 Codex 证据裁决补丁。")
    doc.add_page_break()

    for path in MD_FILES:
        add_markdown_docx(doc, read(path))

    doc.core_properties.title = title_line
    doc.core_properties.subject = "v13.1 Round05 patched legal baodian"
    doc.save(DOCX_PATH)


def render_inline(text: str) -> str:
    text = html.escape(text)
    text = re.sub(r"`([^`]+)`", r"<code>\1</code>", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", text)
    return text


def render_markdown_html(text: str) -> str:
    parts: list[str] = []
    for kind, value in iter_blocks(text):
        if kind == "h":
            level, title = value
            tag = "h2" if level == 1 else "h3" if level == 2 else "h4"
            parts.append(f"<{tag}>{render_inline(title)}</{tag}>")
        elif kind == "p":
            parts.append(f"<p>{render_inline(value)}</p>")
        elif kind == "quote":
            parts.append(f"<blockquote>{render_inline(value)}</blockquote>")
        elif kind == "ul":
            parts.append("<ul>")
            for item in value:
                parts.append(f"<li>{render_inline(item)}</li>")
            parts.append("</ul>")
        elif kind == "table":
            rows = value
            parts.append("<table>")
            for r_idx, row in enumerate(rows):
                tag = "th" if r_idx == 0 else "td"
                parts.append("<tr>" + "".join(f"<{tag}>{render_inline(c)}</{tag}>" for c in row) + "</tr>")
            parts.append("</table>")
    return "\n".join(parts)


def build_html():
    title_line = read(README).splitlines()[0].lstrip("# ").strip()
    body_parts = []
    for path in MD_FILES:
        body_parts.append(f'<section class="md-section">{render_markdown_html(read(path))}</section>')
    status = render_inline(first_status_line())
    body = "\n".join(body_parts)
    css = """
@page{size:A4;margin:18mm 17mm 19mm 17mm}
*{box-sizing:border-box}
body{margin:0;font-family:"Microsoft YaHei",Arial,sans-serif;color:#111827;font-size:10.4pt;line-height:1.48}
.cover{min-height:245mm;display:flex;flex-direction:column;justify-content:center;page-break-after:always}
h1{text-align:center;font-size:28pt;line-height:1.16;color:#1f4d78;margin:0 0 12pt}
.subtitle{text-align:center;font-size:12.5pt;color:#4b5563;margin:0 0 24pt}
.meta{border:1px solid #cbd5e1;border-left:5px solid #2e74b5;background:#f8fafc;padding:12pt 15pt}
h2{font-size:17pt;color:#2e74b5;border-bottom:1px solid #cbd5e1;padding-bottom:4pt;margin:0 0 12pt;page-break-before:always}
h3{font-size:13pt;color:#2e74b5;margin:14pt 0 7pt}
h4{font-size:11.5pt;color:#1f4d78;margin:11pt 0 5pt}
p{margin:0 0 6pt}
ul{margin:0 0 7pt 18pt;padding:0}
li{margin:0 0 3pt}
blockquote{margin:7pt 0 9pt 0;padding:7pt 10pt;border-left:4px solid #2e74b5;background:#f8fafc;color:#1f4d78}
table{width:100%;border-collapse:collapse;margin:8pt 0 12pt;break-inside:auto}
th,td{border:1px solid #cbd5e1;padding:4pt 5pt;vertical-align:top}
th{background:#e8eef5;color:#1f4d78}
code{font-family:Consolas,monospace;background:#f1f5f9;padding:0 2pt}
.md-section{page-break-before:auto}
"""
    doc = f"""<!doctype html>
<html lang="zh-CN">
<head><meta charset="utf-8"><title>{html.escape(title_line)}</title><style>{css}</style></head>
<body>
<section class="cover">
<h1>{html.escape(title_line)}</h1>
<p class="subtitle">Round05 patched final candidate</p>
<div class="meta"><p>{status}</p><p>Model gate: GPT Pro + Claude Opus 4.7 Adaptive, both real web captures.</p><p>Core: 42 locked questions; open-container rows remain outside the framework support count.</p></div>
</section>
{body}
</body></html>"""
    HTML_PATH.write_text(doc, encoding="utf-8", newline="\n")


if __name__ == "__main__":
    build_docx()
    build_html()
    print(f"wrote {DOCX_PATH}")
    print(f"wrote {HTML_PATH}")
