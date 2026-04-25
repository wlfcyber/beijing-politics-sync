from __future__ import annotations

import re
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\Administrator\Desktop\02_代码项目与工具\mac-thread-restore\beijing-politics-sync-visible")
ARTIFACTS = ROOT / "artifacts"
REPORTS = ROOT / "reports" / "overnight_2026-04-25"
DESKTOP = Path(r"C:\Users\Administrator\Desktop")

PHILOSOPHY_MD = ARTIFACTS / "必修四哲学材料-知识触发总框架_持续更新版_v2.md"
CULTURE_MD = ARTIFACTS / "必修四文化材料-知识触发总框架_持续更新版.md"
WRONG_MD = ARTIFACTS / "北京高考政治错肢库_持续更新版.md"

OUT_PHILOSOPHY = DESKTOP / "必修四哲学材料-知识触发总框架_穷尽修订版.docx"
OUT_WRONG = DESKTOP / "北京高考政治选择题错肢总结_穷尽版.docx"
OUT_CULTURE = DESKTOP / "必修四文化材料-知识触发总框架_穷尽修订版.docx"

FONT_BODY = "Microsoft YaHei"
FONT_SONG = "SimSun"
ACCENT = RGBColor(31, 78, 121)
LIGHT = "D9EAF7"
LIGHT_2 = "EEF5FB"
LANDSCAPE_TEXT_WIDTH = 10.79


def set_east_asia_font(run, font_name: str = FONT_BODY) -> None:
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def style_run(run, size: float | None = None, bold: bool | None = None, color=None, font: str = FONT_BODY) -> None:
    set_east_asia_font(run, font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_inches: float) -> None:
    width = Inches(width_inches)
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])


def generic_widths(cols: int) -> list[float]:
    if cols <= 1:
        return [LANDSCAPE_TEXT_WIDTH]
    if cols == 2:
        return [2.05, LANDSCAPE_TEXT_WIDTH - 2.05]
    if cols == 3:
        return [2.2, 5.8, LANDSCAPE_TEXT_WIDTH - 8.0]
    if cols == 4:
        return [1.75, 3.0, 2.5, LANDSCAPE_TEXT_WIDTH - 7.25]
    if cols == 5:
        return [1.35, 2.65, 2.45, 2.0, LANDSCAPE_TEXT_WIDTH - 8.45]
    unit = LANDSCAPE_TEXT_WIDTH / cols
    return [unit] * cols


def set_cell_text(cell, text: str, size: float = 8.2, bold: bool = False, center: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(clean_inline(text))
    style_run(run, size=size, bold=bold, font=FONT_BODY)


def configure_document(doc: Document, landscape: bool = False) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.left_margin = Inches(0.45)
        section.right_margin = Inches(0.45)
        section.top_margin = Inches(0.42)
        section.bottom_margin = Inches(0.42)
    else:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.left_margin = Inches(0.62)
        section.right_margin = Inches(0.62)
        section.top_margin = Inches(0.58)
        section.bottom_margin = Inches(0.58)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.font.size = Pt(10.3)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.12

    for name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
        st = styles[name]
        st.font.name = FONT_BODY
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = ACCENT
        st.paragraph_format.space_before = Pt(8)
        st.paragraph_format.space_after = Pt(4)


def add_cover(doc: Document, title: str, subtitle: str, notes: list[str], landscape: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title)
    style_run(run, size=22 if not landscape else 20, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(subtitle)
    style_run(run, size=11.5, bold=False)

    for idx, note in enumerate(notes, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.55)
        p.paragraph_format.right_indent = Inches(0.55)
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(f"{idx}. {note}")
        style_run(run, size=10.0, bold=(idx == 1))
    doc.add_paragraph()


def clean_inline(text: str) -> str:
    text = text.replace("<br>", "\n").replace("<br/>", "\n")
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("**", "").replace("__", "")
    return text.strip()


def split_md_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [clean_inline(cell) for cell in line.split("|")]


def is_separator_row(line: str) -> bool:
    cells = split_md_row(line)
    return bool(cells) and all(re.fullmatch(r"\s*:?-{3,}:?\s*", c) for c in cells)


def add_markdown_table(doc: Document, rows: list[list[str]], small: bool = False) -> None:
    rows = [r for r in rows if r]
    if not rows:
        return
    headers = rows[0]
    data = rows[1:] if len(rows) > 1 else []
    if not data:
        add_plain_para(doc, "；".join(headers))
        return
    for idx, row in enumerate(data, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.08
        title = row[0] if row else ""
        run = p.add_run(f"{idx}. {title}")
        style_run(run, size=9.0 if small else 9.6, bold=True, color=ACCENT)
        for h, cell in zip(headers[1:], row[1:]):
            if not cell:
                continue
            run = p.add_run(f"\n{h}：{cell}")
            style_run(run, size=8.5 if small else 9.2)


def add_plain_para(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(clean_inline(text))
    style_run(run, size=10.0 if style is None else None)


def markdown_to_docx(md_path: Path, out_path: Path, title: str, subtitle: str, cover_notes: list[str]) -> None:
    text = md_path.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc, landscape=True)
    add_cover(doc, title, subtitle, cover_notes, landscape=True)

    lines = text.splitlines()
    i = 0
    in_code = False
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        if line.strip().startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            if line.strip():
                add_plain_para(doc, line)
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        if line.lstrip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                if not is_separator_row(lines[i]):
                    table_lines.append(split_md_row(lines[i]))
                i += 1
            add_markdown_table(doc, table_lines, small=True)
            continue
        if line.startswith("# "):
            doc.add_heading(clean_inline(line[2:]), level=1)
        elif line.startswith("## "):
            doc.add_heading(clean_inline(line[3:]), level=1)
        elif line.startswith("### "):
            doc.add_heading(clean_inline(line[4:]), level=2)
        elif line.startswith("#### "):
            doc.add_heading(clean_inline(line[5:]), level=3)
        elif line.startswith("- "):
            add_plain_para(doc, line[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            add_plain_para(doc, re.sub(r"^\d+\.\s+", "", line), style="List Number")
        else:
            add_plain_para(doc, line)
        i += 1
    doc.save(out_path)


def parse_wrong_rows(md_path: Path) -> tuple[list[dict], Counter]:
    rows: list[dict] = []
    section_counts: Counter = Counter()
    current_section = "未分组"
    for line in md_path.read_text(encoding="utf-8").splitlines():
        if line.startswith("## "):
            current_section = clean_inline(line[3:])
            continue
        if not line.startswith("|") or is_separator_row(line):
            continue
        cells = split_md_row(line)
        if len(cells) != 9 or cells[0] == "所属模块":
            continue
        row = {
            "section": current_section,
            "module": cells[0],
            "wrong": cells[1],
            "type": cells[2],
            "why": cells[3],
            "correct": cells[4],
            "pure": cells[5],
            "standalone": cells[6],
            "similar": cells[7],
            "source": cells[8],
        }
        rows.append(row)
        section_counts[current_section] += 1
    return rows, section_counts


def first_sample(rows: list[dict], key: str, value: str) -> dict:
    for row in rows:
        if row[key] == value:
            return row
    return rows[0]


def compact(text: str, n: int = 70) -> str:
    text = re.sub(r"\s+", "", text)
    return text if len(text) <= n else text[: n - 1] + "…"


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    for k, v in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{k}：")
        style_run(run, size=10.0, bold=True, color=ACCENT)
        run = p.add_run(v)
        style_run(run, size=10.0)


def add_table_from_records(
    doc: Document,
    headers: list[str],
    records: list[list[str]],
    size: float = 8.2,
    widths: list[float] | None = None,
) -> None:
    for idx, record in enumerate(records, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.08
        lead = record[0] if record else ""
        run = p.add_run(f"{idx}. {lead}")
        style_run(run, size=max(size, 8.6), bold=True, color=ACCENT)
        for h, value in zip(headers[1:], record[1:]):
            if value == "":
                continue
            run = p.add_run(f"\n{h}：{value}")
            style_run(run, size=max(size - 0.1, 7.8))


def build_wrong_summary(out_path: Path) -> None:
    rows, section_counts = parse_wrong_rows(WRONG_MD)
    type_counter = Counter(r["type"] for r in rows)
    module_counter = Counter(r["module"] for r in rows)
    similar_counter = Counter(r["similar"] for r in rows if r["similar"] and r["similar"] != "-")
    pure_counter = Counter(r["pure"] for r in rows)
    standalone_counter = Counter(r["standalone"] for r in rows)

    doc = Document()
    configure_document(doc, landscape=True)
    add_cover(
        doc,
        "北京高考政治选择题错肢总结",
        "2024-2026 套卷错肢规律、同类合并与客观答案源闭环",
        [
            f"正文实测错肢 {len(rows)} 条，覆盖 {len(section_counts)} 个批次；以错肢库正文为最终统计口径。",
            "本总结优先提炼可跨题复用的错法：主体错位、作用夸大、条件倒置、材料外推、概念范围错误、逻辑规则误用。",
            "2026丰台一模、2026房山一模、2026丰台期末、2024门头沟一模已补证第1-15题客观答案源；不凭答案字母反推错肢。",
        ],
        landscape=True,
    )

    doc.add_heading("一、覆盖口径", level=1)
    add_kv_table(
        doc,
        [
            ("错肢正文行数", str(len(rows))),
            ("已解析批次", str(len(section_counts))),
            ("纯知识性错肢", f"是 {pure_counter.get('是', 0)}；否 {pure_counter.get('否', 0)}"),
            ("可脱离材料直接判错", f"是 {standalone_counter.get('是', 0)}；结合材料更稳 {standalone_counter.get('结合材料更稳', 0)}"),
            ("监管提示", "ledger 合计 1445 与正文 1429 有历史差异；本 Word 以正文实测为准，保留台账待整理。"),
        ],
    )

    doc.add_heading("二、高频错误类型", level=1)
    type_records = []
    for err_type, count in type_counter.most_common(28):
        sample = first_sample(rows, "type", err_type)
        type_records.append([err_type, str(count), compact(sample["wrong"], 58), compact(sample["correct"], 70), sample["source"]])
    add_table_from_records(doc, ["错误类型", "次数", "典型错肢", "规范说法", "来源"], type_records, size=7.5, widths=[1.55, 0.55, 3.15, 3.55, 1.99])

    doc.add_heading("三、模块高频错法", level=1)
    module_records = []
    for module, count in module_counter.most_common(32):
        sample = first_sample(rows, "module", module)
        module_records.append([module, str(count), sample["type"], compact(sample["wrong"], 62), sample["source"]])
    add_table_from_records(doc, ["模块", "条数", "典型错误类型", "代表错肢", "来源"], module_records, size=7.5, widths=[1.9, 0.55, 1.55, 4.65, 2.14])

    doc.add_heading("四、同类合并优先背诵", level=1)
    similar_records = []
    for similar, count in similar_counter.most_common(55):
        if count < 2:
            continue
        sample = first_sample(rows, "similar", similar)
        similar_records.append([similar, str(count), sample["type"], compact(sample["why"], 76), sample["source"]])
    add_table_from_records(doc, ["同类合并", "次数", "错法", "判错理由", "代表来源"], similar_records, size=7.4, widths=[2.25, 0.55, 1.35, 4.55, 2.09])

    doc.add_heading("五、跨题通用判错清单", level=1)
    rules = [
        ("主体错位", "先问动作主体是谁。政协、人大、政府、法院、检察院、基层群众自治组织、社会组织不能互相替代。"),
        ("作用夸大", "看到“决定、确保、根本保证、直接实现、完全避免、充分条件”等强词，优先核查材料和教材是否支撑。"),
        ("条件倒置", "必要条件、充分条件、逆命题、否定后件是选必三高频陷阱；不要把“只有…才…”倒成“有…就…”。"),
        ("材料外推", "材料只给现象、局部效果或一个政策工具时，不能推出制度性、全局性、最终性结论。"),
        ("概念范围错误", "把马克思主义品质泛化为全部哲学、把文化载体写成文化本身、把行政法规写成地方性法规，均属范围错。"),
        ("哲学关系误写", "普遍性与特殊性、同一性与斗争性、量变质变、辩证否定、自在联系与人为联系要按关系本身作答。"),
        ("文化题泛化", "文化传播、文化交流、文化自信、文化功能、双创要按材料触发，不能见传统文化就固定套文化自信。"),
        ("经济题方向错", "规模经济、宏观调控、社会保障水平、消费升级、贸易政策要核查作用方向和适用条件。"),
        ("法律程序错", "先定法律关系与诉讼类型，再判断举证、调解、继承、夫妻共同财产、劳动保护等规则。"),
        ("国际政治绝对化", "联合国、全球南方、国际组织、国际合作通常是重要作用或推动作用，少写决定作用。"),
    ]
    add_table_from_records(doc, ["通用错法", "判错动作"], rules, size=8.2, widths=[1.7, LANDSCAPE_TEXT_WIDTH - 1.7])

    doc.add_heading("六、批次覆盖表", level=1)
    batch_records = [[sec, str(cnt)] for sec, cnt in section_counts.items()]
    add_table_from_records(doc, ["批次", "错肢条数"], batch_records, size=7.8, widths=[8.7, 2.09])

    doc.add_heading("七、客观答案源补证", level=1)
    key_records = [
        ["2026北京丰台高三一模", "1.B 2.A 3.D 4.A 5.A 6.D 7.B 8.C 9.D 10.C 11.D 12.B 13.A 14.A 15.C", "北京高考在线答案版 PDF；本地已保存并核验"],
        ["2026北京房山高三一模", "1.C 2.D 3.B 4.A 5.C 6.D 7.B 8.A 9.D 10.B 11.D 12.C 13.B 14.C 15.A", "北京高考在线答案版 PDF；渲染页图已核验"],
        ["2026北京丰台高三（上）期末", "1.B 2.C 3.A 4.B 5.C 6.A 7.B 8.D 9.C 10.D 11.D 12.A 13.C 14.B 15.D", "北京高考在线答案版 PDF；渲染页图已核验"],
        ["2024北京门头沟高三一模", "1.B 2.C 3.A 4.C 5.B 6.C 7.B 8.D 9.A 10.D 11.C 12.C 13.B 14.C 15.A", "北京高考在线答案版 PDF；旧本地缺独立套卷目录，本轮已补证"],
    ]
    add_table_from_records(doc, ["套卷", "第1-15题答案", "证据状态"], key_records, size=7.8, widths=[2.1, 6.2, 2.49])

    doc.save(out_path)


def main() -> None:
    markdown_to_docx(
        PHILOSOPHY_MD,
        OUT_PHILOSOPHY,
        "必修四哲学材料-知识触发总框架",
        "2024-2026 套卷穷尽修订版",
        [
            "大题权重高于选择题；优先海淀、西城、东城、朝阳，再处理郊区。",
            "新增 2026-04-25 夜间监管补丁：2024顺义二模第16（1）题三条哲学链已并入。",
            "所有条目保留来源套卷、题号、材料触发和答题逻辑；不把普通参考答案当评分细则。",
        ],
    )
    build_wrong_summary(OUT_WRONG)
    markdown_to_docx(
        CULTURE_MD,
        OUT_CULTURE,
        "必修四文化材料-知识触发总框架",
        "2024-2026 套卷穷尽修订版",
        [
            "按 0-7 文化框架归位：载体、特点、作用、横向、纵向、文化强国/文化自信、民族精神、习近平文化思想。",
            "新增 2026-04-25 夜间监管补丁：东城一模第16题及朝阳、海淀、西城、石景山可证实文化链已并入。",
            "题号待核、串页、旧题回顾材料未直接升级为答题点。",
        ],
    )
    print(f"Wrote: {OUT_PHILOSOPHY}")
    print(f"Wrote: {OUT_WRONG}")
    print(f"Wrote: {OUT_CULTURE}")


if __name__ == "__main__":
    main()
