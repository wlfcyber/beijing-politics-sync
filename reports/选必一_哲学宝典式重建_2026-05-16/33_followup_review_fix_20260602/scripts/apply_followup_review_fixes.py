from __future__ import annotations

import argparse
import json
import re
from copy import deepcopy
from pathlib import Path

from docx import Document


ENTRY_RE = re.compile(r"^\d+\.\s+")
FIELD_RE = re.compile(r"^【([^】]+)】\s*(.*)$")


def style_name(paragraph) -> str:
    return paragraph.style.name if paragraph.style else ""


def set_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def insert_clone_before(marker, paragraph) -> None:
    marker._element.addprevious(deepcopy(paragraph._element))


def insert_para_before(marker, text: str, style: str = "Normal"):
    return marker.insert_paragraph_before(text, style=style)


def parse_entries(doc: Document) -> list[dict]:
    entries: list[dict] = []
    current_h1 = current_h2 = current_h3 = ""
    starts: list[tuple[int, str, str, str]] = []
    for idx, par in enumerate(doc.paragraphs):
        text = par.text.strip()
        st = style_name(par)
        if st == "Heading 1":
            current_h1 = text
            current_h2 = ""
            current_h3 = ""
        elif st == "Heading 2":
            current_h2 = text
            current_h3 = ""
        elif st == "Heading 3":
            current_h3 = text
        elif ENTRY_RE.match(text):
            starts.append((idx, current_h1, current_h2, current_h3))

    for pos, (start, h1, h2, h3) in enumerate(starts):
        end = len(doc.paragraphs)
        if pos + 1 < len(starts):
            end = starts[pos + 1][0]
        for idx in range(start + 1, end):
            if style_name(doc.paragraphs[idx]) in {"Heading 1", "Heading 2", "Heading 3"}:
                end = idx
                break
        fields = {}
        field_paras = {}
        for idx in range(start + 1, end):
            m = FIELD_RE.match(doc.paragraphs[idx].text.strip())
            if m:
                fields[m.group(1)] = m.group(2)
                field_paras[m.group(1)] = doc.paragraphs[idx]
        entries.append(
            {
                "start": start,
                "end": end,
                "title": doc.paragraphs[start].text.strip(),
                "h1": h1,
                "h2": h2,
                "h3": h3,
                "fields": fields,
                "field_paras": field_paras,
                "paragraphs": [doc.paragraphs[i] for i in range(start, end)],
            }
        )
    return entries


def find_section_end_marker(doc: Document, h3_text: str):
    in_section = False
    for par in doc.paragraphs:
        st = style_name(par)
        text = par.text.strip()
        if st == "Heading 3" and text == h3_text:
            in_section = True
            continue
        if in_section and st in {"Heading 1", "Heading 2", "Heading 3"}:
            return par
    if in_section:
        # Use the last paragraph as a stable append marker fallback.
        return doc.paragraphs[-1]
    raise RuntimeError(f"找不到目标核心点：{h3_text}")


def move_entries_to_h3(doc: Document, entries: list[dict], predicates: list, target_h3: str, changes: list[str]) -> None:
    selected = [entry for entry in entries if any(pred(entry) for pred in predicates)]
    if not selected:
        return
    marker = find_section_end_marker(doc, target_h3)
    for entry in selected:
        insert_para_before(marker, "", "Normal")
        for par in entry["paragraphs"]:
            insert_clone_before(marker, par)
        changes.append(f"moved: {entry['title']} -> {target_h3}")


def delete_entries(entries: list[dict], predicates: list, changes: list[str]) -> list:
    selected = [entry for entry in entries if any(pred(entry) for pred in predicates)]
    refs = []
    for entry in selected:
        refs.extend(entry["paragraphs"])
        changes.append(f"deleted: {entry['title']} from {entry['h3']}")
    return refs


def update_field(entry: dict, label: str, value: str, changes: list[str]) -> None:
    par = entry["field_paras"].get(label)
    if par is None:
        return
    set_text(par, f"【{label}】 {value}")
    changes.append(f"updated {label}: {entry['title']} / {entry['h3']}")


def entry_matches(entry: dict, title_part: str, h3_part: str | None = None) -> bool:
    if title_part not in entry["title"]:
        return False
    if h3_part is not None and h3_part not in entry["h3"]:
        return False
    return True


def add_fengtai_finance_entry(doc: Document, changes: list[str]) -> None:
    target_h3 = "核心答题点：维护全球产业链供应链稳定畅通"
    marker = find_section_end_marker(doc, target_h3)
    paragraphs = [
        "2024丰台一模Q20（供应链金融创新）",
        "【材料触发点】 中国发起设立国际金融机构及投资基金，加快供应链金融创新，助推全球供应链向广大发展中国家延伸。",
        "【设问】 结合材料，运用《当代国际政治与经济》知识，说明我国的实践是如何推动供应链成为国际合作“共赢链”的。",
        "【为什么能想到】 供应链要成为“共赢链”，不仅要有物流通道和合作平台，还要解决链条延伸中的资金约束。材料把国际金融机构、投资基金和供应链金融创新放在一起，说明资金支持是全球供应链稳定延伸的重要机制。",
        "【答案落点】 中国加快供应链金融创新，缓解供应链发展中的资金短缺问题，助推全球供应链向广大发展中国家延伸，增强全球供应链运行韧性和稳定性。",
        "【同题组】 经济全球化：利用两个市场两种资源优化全球资源配置；推进贸易和投资自由化便利化；以产业链供应链合作推动互利共赢",
    ]
    insert_para_before(marker, "", "Normal")
    for text in paragraphs:
        insert_para_before(marker, text, "Normal")
    changes.append("added: 2024丰台一模Q20（供应链金融创新） -> 维护全球产业链供应链稳定畅通")


def remove_backend_bullet_residue(doc: Document, changes: list[str]) -> list:
    prefixes = (
        "· 【同题内分层",
        "· 【什么时候写】",
        "· 材料信号：",
        "· 设问意图：",
        "· 答题动作：",
        "· 【卷面句】",
    )
    refs = []
    for par in doc.paragraphs:
        text = par.text.strip()
        if text.startswith(prefixes):
            refs.append(par)
    if refs:
        changes.append(f"removed backend bullet residue paragraphs: {len(refs)}")
    return refs


def refine_formulaic_why(entries: list[dict], changes: list[str]) -> None:
    """Trim review-flagged formula tails in why fields without changing evidence."""
    replacements = [
        (
            "材料把一国行动与世界共同挑战、共同发展连接起来，已经进入命运相连的层面；因此要用推动构建人类命运共同体统摄合作方向和最终价值。",
            "材料不是在讲单方受益，而是在讲中国行动如何回应共同挑战、扩展共同发展空间；这正是把题面事实上升到人类命运共同体的关键。",
        ),
        (
            "看到这类关系，应联想到",
            "这一关系提示答案应落到",
        ),
        (
            "据此可联想到",
            "这些事实共同指向",
        ),
        (
            "由此可判断采分核心是",
            "这些事实共同指向",
        ),
        (
            "材料事实服务于这一判断",
            "作答时应让材料回扣这一核心",
        ),
        (
            "材料事实应围绕该核心展开",
            "材料事实要围绕这一核心展开",
        ),
    ]
    regex_replacements = [
        (re.compile(r"^设问追问的是(?:原因和逻辑|做法和机制|作用和意义|作用和价值)。"), ""),
        (
            re.compile(r"这道题可从多角度作答，本层归入“([^”]+)”，用来解释材料事实的特定作用。"),
            r"同一材料可服务多个答点；本处只取“\1”这一层，说明材料怎样支撑该术语。",
        ),
        (
            re.compile(r"同题组存在多条线索；这里对应“([^”]+)”，用于说明材料事实的这一层功能。"),
            r"同题组存在多条线索；这里取“\1”这一层，说明材料怎样支撑该术语。",
        ),
        (
            re.compile(r"同题组线索较多，本处选择“([^”]+)”这一核心，说明材料在该处的采分功能。"),
            r"同题组线索较多；本处取“\1”这一层，说明材料怎样支撑该术语。",
        ),
        (
            re.compile(r"同一题组还有其他角度，本处抓“([^”]+)”，突出材料事实在该核心下的作用。"),
            r"同一题组还有其他角度；本处取“\1”这一层，突出材料事实在该核心下的作用。",
        ),
        (
            re.compile(r"同一材料可服务多个答点；这里归到“([^”]+)”，强调这一核心的材料支撑。"),
            r"同一材料可服务多个答点；本处取“\1”这一层，强调这一核心的材料支撑。",
        ),
        (
            re.compile(r"，?看到这类关系，应联想到“([^”]+)”，再用材料事实支撑这一判断。"),
            r"。这一关系提示答案应落到“\1”，再用题面事实支撑。",
        ),
        (
            re.compile(r"，?据此可联想到“([^”]+)”，用它说明上述事实的答题依据。"),
            r"。这些事实共同指向“\1”，作答时要说明它们为何支撑这一核心。",
        ),
        (
            re.compile(r"，?由此可判断采分核心是“([^”]+)”，材料事实服务于这一判断。"),
            r"。这些事实共同指向“\1”，作答时应让材料回扣这一核心。",
        ),
        (
            re.compile(r"这[一类]?关系的答题支点是“([^”]+)”，材料事实应围绕该核心展开。"),
            r"这一组关系提示答案支点在“\1”，材料事实要围绕这一核心展开。",
        ),
        (
            re.compile(r"这一关系的答题支点是“([^”]+)”，材料事实要围绕这一核心展开。"),
            r"这一组关系提示答案支点在“\1”，材料事实要围绕这一核心展开。",
        ),
    ]
    changed = 0
    for entry in entries:
        why = entry["fields"].get("为什么能想到", "")
        if not why:
            continue
        new_why = why
        for old, new in replacements:
            new_why = new_why.replace(old, new)
        for pattern, repl in regex_replacements:
            new_why = pattern.sub(repl, new_why)
        new_why = re.sub(r"。{2,}", "。", new_why).strip()
        if new_why and new_why != why:
            update_field(entry, "为什么能想到", new_why, changes)
            changed += 1
    changes.append(f"refined formulaic why fields: {changed}")


def renumber_entries_by_h3(doc: Document) -> int:
    current_h3 = ""
    counter = 0
    changed = 0
    for par in doc.paragraphs:
        st = style_name(par)
        text = par.text.strip()
        if st == "Heading 3":
            current_h3 = text
            counter = 0
            continue
        if current_h3 and ENTRY_RE.match(text):
            counter += 1
            new_text = ENTRY_RE.sub(f"{counter}. ", text, count=1)
            if new_text != text:
                set_text(par, new_text)
                changed += 1
        elif current_h3 and not re.match(r"^\d+\.\s+", text) and re.match(r"^\d{4}", text):
            counter += 1
            set_text(par, f"{counter}. {text}")
            changed += 1
    return changed


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--summary", type=Path, required=True)
    args = parser.parse_args()

    doc = Document(args.input)
    changes: list[str] = []
    entries = parse_entries(doc)

    trade_h3 = "核心答题点：推进贸易和投资自由化便利化"
    regional_h3 = "核心答题点：推动区域经济融合与区域经济一体化"
    def regional_wrong_entry(entry: dict) -> bool:
        if regional_h3 not in entry["h3"]:
            return False
        title = entry["title"]
        return any(
            key in title
            for key in [
                "2024丰台一模Q20",
                "2025朝阳一模Q20",
                "2025海淀一模Q21(2)",
                "2026东城二模Q20(3)",
                "2026丰台二模Q20",
                "2026西城一模Q20(2)",
            ]
        ) or title.startswith("29. 2025朝阳二模Q21")

    move_entries_to_h3(
        doc,
        entries,
        [regional_wrong_entry],
        trade_h3,
        changes,
    )

    # Move the 朝阳一模 supply-chain autonomy item out of the 周边工作四定位 bucket.
    move_entries_to_h3(
        doc,
        entries,
        [
            lambda e: entry_matches(
                e,
                "2025朝阳一模Q20",
                "核心答题点：周边工作新局面四定位",
            )
        ],
        "核心答题点：坚持独立自主、自力更生，增强发展主动权",
        changes,
    )

    delete_refs = []
    delete_refs.extend(
        delete_entries(
            entries,
            [regional_wrong_entry],
            changes,
        )
    )
    delete_refs.extend(
        delete_entries(
            entries,
            [
                lambda e: entry_matches(
                    e,
                    "2025朝阳一模Q20",
                    "核心答题点：周边工作新局面四定位",
                )
            ],
            changes,
        )
    )
    delete_refs.extend(
        delete_entries(
            entries,
            [
                lambda e: entry_matches(
                    e,
                    "2026东城一模Q19(3)",
                    "核心答题点：维护全球产业链供应链稳定畅通",
                )
            ],
            changes,
        )
    )
    delete_refs.extend(
        delete_entries(
            entries,
            [
                lambda e: entry_matches(
                    e,
                    "2026朝阳二模Q20(2)",
                    "核心答题点：共商共建共享的全球治理观",
                )
            ],
            changes,
        )
    )
    delete_refs.extend(
        delete_entries(
            entries,
            [
                lambda e: entry_matches(
                    e,
                    "2025西城一模Q21",
                    "核心答题点：四大全球倡议相互促进",
                )
            ],
            changes,
        )
    )

    for entry in entries:
        if entry_matches(entry, "2026顺义一模Q20", "核心答题点：坚持正确义利观"):
            answer = entry["fields"].get("答案落点", "")
            if "无偿" in answer:
                update_field(
                    entry,
                    "答案落点",
                    "中国坚持正确义利观，义利兼顾、以义为先，与南方国家共享农业技术、培养本土人才、增强内生发展动力，惠及当地民生，成为南南合作典范。",
                    changes,
                )

        if entry_matches(entry, "2026东城期末Q20", "核心答题点：坚持正确义利观，在发展合作中共享机遇"):
            update_field(
                entry,
                "材料触发点",
                "题面把全球发展倡议与共建“一带一路”等发展合作相连；“以发展促繁荣”要求把发展合作接到互利共赢、共享机遇和共同繁荣。",
                changes,
            )
            update_field(
                entry,
                "为什么能想到",
                "四大全球倡议题要说明各倡议如何系统支撑人类命运共同体。发展倡议的材料证据集中在共建“一带一路”等发展合作，应从互利共赢、共享发展机遇写到共同繁荣的物质根基，不能混入其他题目的发展合作细节。",
                changes,
            )
            update_field(
                entry,
                "答案落点",
                "全球发展倡议以发展促繁荣，坚持互利共赢，在共建“一带一路”等发展合作中共享机遇，为构建人类命运共同体奠定共同繁荣的物质根基。",
                changes,
            )

        if entry_matches(entry, "2026通州期末Q21", "核心答题点：贡献中国智慧、中国方案、中国力量"):
            update_field(
                entry,
                "为什么能想到",
                "本题不是问中国对世界作出何种贡献，而是综合阐释中国式现代化建设如何展开。题面列举多场开放平台和国际合作实践，说明高水平对外开放已经成为推进中国式现代化的重要支撑；因此应把“中国智慧、中国方案、中国力量”落回中国式现代化建设的实践路径。",
                changes,
            )
            update_field(
                entry,
                "答案落点",
                "我国依托中国—中亚峰会、“一带一路”国际合作高峰论坛、中非合作论坛、进博会、服贸会、消博会等开放平台，展现中国式现代化的开放道路和实践能力，为推进中国式现代化贡献中国智慧、中国方案、中国力量。",
                changes,
            )

        if entry_matches(entry, "2026通州一模Q19", "核心答题点：中国特色大国外交与人类命运共同体"):
            update_field(
                entry,
                "答案落点",
                "中国元首外交展现中国特色大国外交的主动作为，坚持对话合作、维护以联合国为核心的国际体系，推动构建人类命运共同体，为动荡变革的世界注入稳定性与正能量。",
                changes,
            )

        if entry_matches(entry, "2026西城二模Q19(1)", "核心答题点：推动构建人类命运共同体"):
            update_field(
                entry,
                "为什么能想到",
                "数据跨境流动和数字经济发展形成全球性议题，中国参与全球数据治理要先说明全球化背景和共同价值基础。题目问参与逻辑，关键在于说明中国为什么既维护自身数据安全和发展利益，又主张各国共同应对数字治理问题；这一路径体现命运相连、共同治理的共同体思维。",
                changes,
            )

    # Improve the highest-risk copied HMCT tails when they explain another core.
    for entry in entries:
        if entry["h3"] == "核心答题点：推动国际秩序和全球治理体系更加公正合理" and "推动构建人类命运共同体统摄" in entry["fields"].get("为什么能想到", ""):
            update_field(
                entry,
                "为什么能想到",
                "设问问“四大全球倡议如何系统推动构建人类命运共同体”，其中全球治理倡议对应的是制度保障。材料指向全球治理难题和国际调解等制度性实践，说明作答应突出治理规则如何更公平、更有效，而不是只重复共同体总纲；因此要落到推动国际秩序和全球治理体系更加公正合理。",
                changes,
            )
        if entry["h3"] == "核心答题点：践行真正的多边主义" and "推动构建人类命运共同体统摄" in entry["fields"].get("为什么能想到", ""):
            update_field(
                entry,
                "为什么能想到",
                "材料把个别国家鼓噪“脱钩断链”和多国企业到链博会寻求合作放在对照中，焦点是用多边合作反对单边割裂。设问要求回应外部质疑，不能只写共同体愿景，而要先说明中国为何坚持通过多边平台凝聚产供链合作共识。",
                changes,
            )
        if entry["h3"] == "核心答题点：维护《联合国宪章》宗旨和原则" and "推动构建人类命运共同体统摄" in entry["fields"].get("为什么能想到", ""):
            update_field(
                entry,
                "为什么能想到",
                "开源战略题强调技术开放合作如何形成全球合作新模式。若落在联合国宪章层，关键不是泛写共同体愿景，而是说明开放合作应在平等、合作、共同发展的国际法和多边治理框架中展开。",
                changes,
            )

    add_fengtai_finance_entry(doc, changes)
    delete_refs.extend(remove_backend_bullet_residue(doc, changes))
    refine_formulaic_why(entries, changes)

    # Delete originals and wrong entries after all insertions/updates.
    seen = set()
    for par in delete_refs:
        ident = id(par._element)
        if ident in seen:
            continue
        seen.add(ident)
        delete_paragraph(par)

    renumbered = renumber_entries_by_h3(doc)
    changes.append(f"renumbered entry headings by core point: {renumbered}")

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    args.summary.write_text(json.dumps({"changes": changes}, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"change_count": len(changes), "output": str(args.output)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
