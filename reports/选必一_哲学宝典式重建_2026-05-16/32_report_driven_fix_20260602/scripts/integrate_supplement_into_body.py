from __future__ import annotations

import argparse
import json
import re
from collections import defaultdict
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ENTRY_RE = re.compile(r"^\d+\.\s+")
APPENDIX_H1 = "三维报告补入题（按六大要素归入）"
SPEED_H1 = "六大要素术语极简速记版"


NEW_CORE_H2 = {
    ("理论", "核心答题点：共同利益是合作的基础"): "合作",
    ("经济全球化", "核心答题点：中国市场红利与全球经济包容性增长"): "兼：国内国际两种资源两个市场",
    ("经济全球化", "核心答题点：中国市场红利推动互利共赢"): "兼：国内国际两种资源两个市场",
    ("经济全球化", "核心答题点：以产业链供应链合作推动互利共赢"): "产业链供应链与开放安全",
    ("经济全球化", "核心答题点：促进国内国际双循环，让世界共享中国发展机遇"): "兼：国内国际两种资源两个市场",
    ("经济全球化", "核心答题点：发挥我国超大规模市场优势"): "兼：国内国际两种资源两个市场",
    ("经济全球化", "核心答题点：发挥比较优势并打造国际竞争新优势"): "国际分工、企业出海与国际竞争新优势",
    ("经济全球化", "核心答题点：坚持多边贸易、多边主义与全球经济治理体系改革"): "全球经济治理与规则权益",
    ("经济全球化", "核心答题点：推动全球要素向中国市场集聚"): "兼：国内国际两种资源两个市场",
    ("经济全球化", "核心答题点：推动区域经济融合与区域经济一体化"): "贸资：贸易投资自由化便利化",
    ("经济全球化", "核心答题点：提升发展中国家的代表性和发言权，打破规则制定垄断"): "全球经济治理与规则权益",
    ("经济全球化", "核心答题点：数据要素在全球范围内流动"): "数字经济、绿色转型与新兴领域",
    ("经济全球化", "核心答题点：统筹发展和安全，形成多元稳定的经贸关系"): "产业链供应链与开放安全",
    ("经济全球化", "核心答题点：维护以世界贸易组织为核心的多边贸易体制"): "贸易保护主义与多边贸易体制",
    ("经济全球化", "核心答题点：通过区域经济合作维护自由贸易和以世界贸易组织为核心的多边贸易体制"): "贸易保护主义与多边贸易体制",
    ("政治多极化", "核心答题点：国际关系民主化与国际政治经济新秩序"): "民主：国际关系民主化",
    ("政治多极化", "核心答题点：维护地区和平与稳定"): "新国关：新型国际关系",
    ("中国", "核心答题点：中国特色大国外交与人类命运共同体"): "政策",
    ("中国", "核心答题点：促进技术共享和民生改善，为全球可持续发展贡献力量"): "责任",
    ("中国", "核心答题点：周边工作新局面四定位：发展繁荣的重要基础、维护国家安全的重点、运筹外交全局的首要、推动构建人类命运共同体的关键"): "政策",
    ("联合国", "核心答题点：遵循联合国宪章宗旨和原则"): "联合国宪章宗旨和原则",
}


def style_name(paragraph) -> str:
    return paragraph.style.name if paragraph.style else ""


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def insert_clone_before(marker, paragraph) -> None:
    marker._element.addprevious(deepcopy(paragraph._element))


def insert_text_before(marker, text: str, style: str) -> None:
    marker.insert_paragraph_before(text, style=style)


def find_appendix_bounds(doc: Document) -> tuple[int, int]:
    start = end = None
    for idx, par in enumerate(doc.paragraphs):
        text = par.text.strip()
        st = style_name(par)
        if st == "Heading 1" and text == APPENDIX_H1:
            start = idx
        elif start is not None and st == "Heading 1" and text == SPEED_H1:
            end = idx
            break
    if start is None or end is None:
        raise RuntimeError("找不到补入区或速记版边界。")
    return start, end


def extract_entries(doc: Document, start: int, end: int) -> list[dict]:
    entries = []
    current_h1 = current_h3 = ""
    current = None

    def finish() -> None:
        nonlocal current
        if current:
            entries.append(current)
            current = None

    for par in doc.paragraphs[start + 1 : end]:
        text = par.text.strip()
        st = style_name(par)
        if st == "Heading 2":
            finish()
            current_h1 = text
            current_h3 = ""
            continue
        if st == "Heading 3":
            finish()
            current_h3 = text
            continue
        if ENTRY_RE.match(text):
            finish()
            current = {
                "h1": current_h1,
                "h3": current_h3,
                "title": text,
                "paragraphs": [par],
            }
            continue
        if current is not None:
            current["paragraphs"].append(par)
    finish()
    return entries


def build_section_indexes(doc: Document, appendix_start: int):
    h3_markers: dict[tuple[str, str], object] = {}
    h2_markers: dict[tuple[str, str], object] = {}
    h1_markers: dict[str, object] = {}
    current_h1 = current_h2 = current_h3 = ""
    active_h3_key = None
    active_h2_key = None
    active_h1 = None

    for par in doc.paragraphs[:appendix_start]:
        st = style_name(par)
        text = par.text.strip()
        if st in {"Heading 1", "Heading 2", "Heading 3"}:
            if active_h3_key is not None:
                h3_markers.setdefault(active_h3_key, par)
                active_h3_key = None
            if st in {"Heading 1", "Heading 2"} and active_h2_key is not None:
                h2_markers.setdefault(active_h2_key, par)
                active_h2_key = None
            if st == "Heading 1" and active_h1 is not None:
                h1_markers.setdefault(active_h1, par)
                active_h1 = None

        if st == "Heading 1":
            current_h1 = text
            current_h2 = ""
            current_h3 = ""
            active_h1 = current_h1
        elif st == "Heading 2":
            current_h2 = text
            current_h3 = ""
            active_h2_key = (current_h1, current_h2)
        elif st == "Heading 3":
            current_h3 = text
            active_h3_key = (current_h1, current_h3)

    appendix_marker = doc.paragraphs[appendix_start]
    if active_h3_key is not None:
        h3_markers.setdefault(active_h3_key, appendix_marker)
    if active_h2_key is not None:
        h2_markers.setdefault(active_h2_key, appendix_marker)
    if active_h1 is not None:
        h1_markers.setdefault(active_h1, appendix_marker)
    return h3_markers, h2_markers, h1_markers


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--summary", type=Path, required=True)
    args = parser.parse_args()

    doc = Document(args.input)
    appendix_start, appendix_end = find_appendix_bounds(doc)
    entries = extract_entries(doc, appendix_start, appendix_end)
    h3_markers, h2_markers, h1_markers = build_section_indexes(doc, appendix_start)

    grouped: dict[tuple[str, str], list[dict]] = defaultdict(list)
    for entry in entries:
        grouped[(entry["h1"], entry["h3"])].append(entry)

    placed = []
    created_cores = []
    missing = []

    for key, group in grouped.items():
        h1, h3 = key
        if key in h3_markers:
            marker = h3_markers[key]
            target_mode = "existing_h3"
        else:
            h2 = NEW_CORE_H2.get(key)
            if h2 and (h1, h2) in h2_markers:
                marker = h2_markers[(h1, h2)]
                insert_text_before(marker, "", "Normal")
                insert_text_before(marker, h3, "Heading 3")
                marker = h2_markers[(h1, h2)]
                target_mode = f"new_h3_under_h2:{h2}"
                created_cores.append({"h1": h1, "h2": h2, "h3": h3, "entries": len(group)})
            elif h1 in h1_markers:
                marker = h1_markers[h1]
                insert_text_before(marker, "", "Normal")
                insert_text_before(marker, h3, "Heading 3")
                marker = h1_markers[h1]
                target_mode = "new_h3_under_h1_fallback"
                created_cores.append({"h1": h1, "h2": "", "h3": h3, "entries": len(group)})
            else:
                missing.append({"h1": h1, "h3": h3, "entries": len(group)})
                continue

        for entry in group:
            for par in entry["paragraphs"]:
                insert_clone_before(marker, par)
            insert_text_before(marker, "", "Normal")
            placed.append({"h1": h1, "h3": h3, "title": entry["title"], "target_mode": target_mode})

    # Remove the whole appendix block after insertion. Recompute bounds because
    # insertions before the appendix change paragraph indexes.
    appendix_start, appendix_end = find_appendix_bounds(doc)
    for par in list(doc.paragraphs[appendix_start:appendix_end]):
        remove_paragraph(par)

    # Remove appendix line from manual TOC if it exists.
    for par in list(doc.paragraphs[:90]):
        if par.text.strip().startswith(APPENDIX_H1 + "\t"):
            remove_paragraph(par)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    summary = {
        "input": str(args.input),
        "output": str(args.output),
        "entries_extracted": len(entries),
        "entries_placed": len(placed),
        "groups": len(grouped),
        "created_cores": created_cores,
        "missing": missing,
        "sample_placed": placed[:20],
    }
    args.summary.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({k: v for k, v in summary.items() if k not in {"sample_placed"}}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
