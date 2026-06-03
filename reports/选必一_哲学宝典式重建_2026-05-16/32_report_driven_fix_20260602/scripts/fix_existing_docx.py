from __future__ import annotations

import argparse
import re
from copy import deepcopy
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document

from parse_baodian_docx import iter_entries, normalize


TARGET_Q17_TITLE = "3. 2026朝阳期中Q17"
TARGET_Q17_HEADING = "核心答题点：科教兴国战略与人才强国战略（人才与国际竞争力）"


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def qtype(question: str) -> str:
    if "如何" in question or "怎样" in question or "怎么" in question:
        return "设问追问的是做法和机制"
    if "为什么" in question or "为何" in question or "理论逻辑" in question or "深层逻辑" in question:
        return "设问追问的是原因和逻辑"
    if "意义" in question or "贡献" in question or "价值" in question:
        return "设问追问的是作用和价值"
    if "理解" in question or "阐释" in question or "说明" in question or "分析" in question:
        return "设问要求把材料事实上升为教材判断"
    return "设问要求把材料关系转化为教材术语"


def cue_from_trigger(trigger: str) -> str:
    text = re.sub(r"[。；;].*$", "", trigger or "").strip()
    text = re.sub(r"^(题目|题面|设问|材料|资料包)", lambda m: m.group(0), text)
    if len(text) > 72:
        text = text[:72].rstrip("，、；。") + "……"
    return text or "题干给出的材料关系"


def landing_from_answer(answer: str, heading: str) -> str:
    h = heading.replace("核心答题点：", "")
    for key in [
        "正确义利观",
        "贸易和投资自由化便利化",
        "经济全球化",
        "两个市场两种资源",
        "自力更生",
        "总体国家安全观",
        "国际关系民主化",
        "共商共建共享",
        "人类命运共同体",
        "真正的多边主义",
        "《联合国宪章》",
        "中国智慧",
        "中国方案",
        "大国责任",
        "大国担当",
        "共同利益",
        "国家利益",
        "综合国力",
    ]:
        if key in answer or key in h:
            return key
    return h[:28]


def synthesize_why(entry: dict) -> str:
    fields = entry["fields"]
    question = fields.get("设问", "")
    trigger = fields.get("材料触发点", "")
    answer = fields.get("答案落点", "")
    heading = entry.get("h3", "")
    cue = cue_from_trigger(trigger)
    landing = landing_from_answer(answer, heading)
    qt = qtype(question)

    if "正确义利观" in landing or "义利观" in heading:
        return (
            f"{qt}。{cue}，关键不只是写中国做了什么，还要解释中国怎样处理自身发展、他国关切和共同利益的关系；"
            f"因此要从义利相济、以义为先的正确义利观切入，再把材料中的合作或治理行动接到大国担当。"
        )
    if "自力更生" in landing or "自主可控" in answer:
        return (
            f"{qt}。{cue}，一边指向核心能力不能受制于人，一边又保留开放合作空间；"
            f"这组张力不是普通技术措施，而是在考处理好自力更生与对外开放的关系。"
        )
    if "国家安全" in landing or "安全" in heading:
        return (
            f"{qt}。{cue}，材料把发展收益和风险约束放在同一问题里，说明答案不能只写增长或开放；"
            f"需要用统筹发展和安全把风险防控、科技安全和高质量发展连成一条线。"
        )
    if "贸易和投资自由化便利化" in landing or "贸易投资" in heading:
        return (
            f"{qt}。{cue}，材料呈现的是跨境交易流程、成本或市场准入的变化；"
            f"这些事实的教材落点不是一般开放口号，而是贸易和投资自由化便利化对开放型经济的推动。"
        )
    if "两个市场两种资源" in landing or "资源配置" in heading:
        return (
            f"{qt}。{cue}，题干把国内发展与国际资源、国际市场联系起来，说明作答要解释资源如何跨国配置；"
            f"因此要从用好两个市场两种资源切入，再写对双循环和开放型经济的带动。"
        )
    if "经济全球化" in landing or "全球化" in heading:
        return (
            f"{qt}。{cue}，材料事实已经超出单个国家内部发展，指向生产、贸易、投资或规则在全球范围内联动；"
            f"因此要把它提升为经济全球化的方向、机制或治理问题。"
        )
    if "国际关系民主化" in landing:
        return (
            f"{qt}。{cue}，题干强调的是各国能否平等参与、规则能否公正形成；"
            f"这类线索指向国际关系运行方式，所以要落到推动国际关系民主化。"
        )
    if "共商共建共享" in landing:
        return (
            f"{qt}。{cue}，材料中的合作不是单边安排，而是多方共同参与、共同治理、共同受益；"
            f"这种治理关系对应共商共建共享的全球治理观。"
        )
    if "人类命运共同体" in landing:
        return (
            f"{qt}。{cue}，材料把一国行动与世界共同挑战、共同发展连接起来，已经进入命运相连的层面；"
            f"因此要用推动构建人类命运共同体统摄合作方向和最终价值。"
        )
    if "真正的多边主义" in landing or "多边主义" in heading:
        return (
            f"{qt}。{cue}，题干强调通过多边机制处理共同问题，而不是少数国家单边决定；"
            f"这正对应践行真正的多边主义及维护以联合国为核心的国际体系。"
        )
    if "联合国" in heading or "《联合国宪章》" in landing:
        return (
            f"{qt}。{cue}，材料把中国行动放在联合国或国际法框架内，说明作答要先给出制度合法性；"
            f"因此要联系《联合国宪章》宗旨原则和联合国在全球治理中的作用。"
        )
    if "共同利益" in landing or "共同利益" in heading:
        return (
            f"{qt}。{cue}，材料不是单向帮助，而是显示相关国家存在共同需求和共同收益；"
            f"回答合作为什么成立时，应先用国家间共同利益是合作基础来解释。"
        )
    if "国家利益" in landing or "国家利益" in heading:
        return (
            f"{qt}。{cue}，题干要求说明国家对外行动的根本依据；"
            f"因此要从维护国家利益这一出发点和落脚点解释政策选择。"
        )
    if "综合国力" in landing or "竞争" in heading:
        return (
            f"{qt}。{cue}，材料把科技、产业、人才或创新能力放进国际比较中，说明问题已经进入综合国力竞争层面；"
            f"因此要联系国际竞争实质，再说明相关能力如何支撑中国掌握主动。"
        )
    return (
        f"{qt}。{cue}，材料提供的是能触发“{landing}”的关系线索，而不是孤立事实；"
        f"作答要把这组关系转化为教材术语，再接上本题具体事实和作用结果。"
    )


def similarity(a: str, b: str) -> float:
    a2, b2 = normalize(a), normalize(b)
    if not a2 or not b2:
        return 0.0
    return SequenceMatcher(None, a2, b2).ratio()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--threshold", type=float, default=0.6)
    args = parser.parse_args()

    entries, _ = iter_entries(args.source)
    doc = Document(args.source)

    remove_para_idxs: set[int] = set()
    rewritten_why = 0
    forced_why = 0

    for entry in entries:
        title = entry["entry_title"]
        heading = entry["h3"]
        fields = entry["fields"]

        if title == TARGET_Q17_TITLE and heading == TARGET_Q17_HEADING:
            start = entry["entry_para_idx"]
            field_idxs = list(entry["field_para_idx"].values())
            end = max(field_idxs) if field_idxs else start
            for idx in range(start, end + 1):
                remove_para_idxs.add(idx)
            continue

        why_idx = entry["field_para_idx"].get("为什么能想到")
        trigger = fields.get("材料触发点", "")
        why = fields.get("为什么能想到", "")
        should_rewrite = similarity(trigger, why) >= args.threshold
        if "2026朝阳期中Q17" in title and should_rewrite is False:
            should_rewrite = True
            forced_why += 1
        if why_idx is not None and should_rewrite:
            set_paragraph_text(doc.paragraphs[why_idx], "【为什么能想到】 " + synthesize_why(entry))
            rewritten_why += 1

        answer_idx = entry["field_para_idx"].get("答案落点")
        if title == "3. 2025海淀二模Q21" and "坚持正确义利观" in heading and answer_idx is not None:
            set_paragraph_text(
                doc.paragraphs[answer_idx],
                "【答案落点】 中国担任安理会轮值主席、推动安理会讨论聚焦实质性问题，并同全球南方国家发起“和平之友”小组，体现义利相济、以义为先的正确义利观，为完善全球治理、构建新型国际关系和构建人类命运共同体贡献中国智慧和中国方案。",
            )

    single_window_hits = 0
    for para in doc.paragraphs:
        if para._element.getparent() is None:
            continue
        text = para.text
        if "单一窗口" in text:
            single_window_hits += text.count("单一窗口")
            text = text.replace('题干中"分钟通关""单一窗口"等改革', '题干中"分钟通关"等通关创新')
            text = text.replace('材料把"分钟通关""单一窗口"等通关改革具体细节摆出来', '材料把"分钟通关"等通关改革细节摆出来')
            text = text.replace('先把"分钟通关""单一窗口"等改革提炼', '先把"分钟通关"等通关创新提炼')
            text = text.replace('"分钟通关""单一窗口"等改革', '"分钟通关"等通关创新')
            text = text.replace('"分钟通关""单一窗口"等通关改革', '"分钟通关"等通关改革')
            set_paragraph_text(para, text)

    for idx in sorted(remove_para_idxs, reverse=True):
        delete_paragraph(doc.paragraphs[idx])

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    print(
        {
            "removed_wrong_q17_entry_paragraphs": len(remove_para_idxs),
            "rewritten_why": rewritten_why,
            "forced_q17_why": forced_why,
            "single_window_hits_removed": single_window_hits,
            "output": str(args.output),
        }
    )


if __name__ == "__main__":
    main()
