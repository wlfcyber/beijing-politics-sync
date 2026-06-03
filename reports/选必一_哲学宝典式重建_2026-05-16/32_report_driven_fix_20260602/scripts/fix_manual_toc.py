from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


TOC_TAIL = [
    "三维报告补入题（按六大要素归入）\t317",
    "六大要素术语极简速记版\t364",
    "时代背景\t364",
    "理论\t365",
    "经济全球化\t365",
    "政治多极化\t367",
    "中国\t368",
    "联合国\t370",
]


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()

    doc = Document(args.input)
    start = None
    for idx, para in enumerate(doc.paragraphs[:90]):
        if para.text.strip().startswith("六大要素术语极简速记版\t"):
            start = idx
            break
    if start is None:
        raise RuntimeError("找不到目录中的“六大要素术语极简速记版”行。")

    for offset, text in enumerate(TOC_TAIL):
        doc.paragraphs[start + offset].text = text

    heading_idx = None
    for idx, para in enumerate(doc.paragraphs):
        if para.style.name == "Heading 1" and para.text.strip() == "时代背景":
            heading_idx = idx
            break
    if heading_idx is not None and doc.paragraphs[heading_idx - 1].text.strip():
        doc.paragraphs[heading_idx].insert_paragraph_before("")

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    print(args.output)


if __name__ == "__main__":
    main()
