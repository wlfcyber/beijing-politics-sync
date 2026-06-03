from __future__ import annotations

import argparse
import json
import re
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


H1_ORDER = ["时代背景", "理论", "经济全球化", "政治多极化", "中国", "联合国"]
TARGET_ORDER = [
    "2024丰台一模Q20",
    "2024丰台一模Q21",
    "2024石景山一模Q19",
    "2025朝阳一模Q20",
    "2025朝阳二模Q21",
    "2025海淀一模Q21",
    "2025石景山一模Q17",
    "2025门头沟一模Q19",
    "2026东城二模Q20",
    "2026丰台二模Q20",
    "2026丰台期末Q21",
    "2026房山二模Q20",
    "2026朝阳二模Q20",
    "2026海淀期中Q22",
    "2026西城一模Q20",
    "2026通州一模Q19",
]


def norm(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def clean_title(title: str) -> str:
    title = re.sub(r"^\s*\d+\.\s*", "", title or "").strip()
    return title


def clean_core(core: str) -> str:
    return re.sub(r"^核心答题点：", "", core or "").strip()


def split_group_lines(text: str) -> list[str]:
    lines = []
    for raw in (text or "").splitlines():
        s = raw.strip()
        if not s:
            continue
        s = re.sub(r"^[-•]\s*", "", s)
        lines.append(s)
    return lines


def supplemental_manual_entries() -> list[dict]:
    q21_question = (
        "结合材料，运用《当代国际政治与经济》知识，"
        "以“四大全球倡议为构建人类命运共同体凝聚磅礴合力”为主题，撰写一篇短评。"
    )
    q21_trigger = (
        "题目要求围绕“四大全球倡议为构建人类命运共同体凝聚磅礴合力”写短评，"
        "题干呈现四大全球倡议回应时代课题、凝聚国际共识与共同未来之间的关系。"
    )
    q21_common = (
        "短评题不能只列四个倡议名称，要把“四大全球倡议”整体解释为凝聚国际共识、"
        "推动构建人类命运共同体的中国方案。"
    )
    tq_question = (
        "结合材料，运用《当代国际政治与经济》知识，"
        "分析中国元首外交如何为世界注入稳定性与正能量。（8分）"
    )
    return [
        {
            "target": "2026丰台期末Q21",
            "h1": "时代背景",
            "core": "和平与发展仍是时代主题",
            "title": "2026丰台期末Q21",
            "fields": {
                "材料触发点": q21_trigger
                + "回答这一关系时，应先说明和平与发展仍是时代主题，各国人民求和平、谋发展、促合作、图共赢的愿望构成合作基础。",
                "设问": q21_question,
                "为什么能想到": q21_common
                + "“凝聚磅礴合力”指向倡议为什么能被各方接受，先用和平与发展仍是时代主题说明合作有时代基础，再把共同愿望接到人类命运共同体。",
                "答案落点": "和平与发展仍是时代主题，四大全球倡议顺应各国人民求和平、谋发展、促合作、图共赢的共同愿望，为构建人类命运共同体凝聚合作力量。",
            },
        },
        {
            "target": "2026丰台期末Q21",
            "h1": "政治多极化",
            "core": "共商共建共享的全球治理观",
            "title": "2026丰台期末Q21",
            "fields": {
                "材料触发点": q21_trigger
                + "回答这一关系时，应说明全球治理事务不能由少数国家垄断，而要由各方共同参与、共同建设、共同受益。",
                "设问": q21_question,
                "为什么能想到": q21_common
                + "“凝聚合力”本身就是共同治理的表达，写法上应落到共商共建共享，而不是把中国倡议写成单向输出。",
                "答案落点": "四大全球倡议坚持共商共建共享，推动各方在共同参与、共同建设、共同受益中凝聚构建人类命运共同体的合力。",
            },
        },
        {
            "target": "2026丰台期末Q21",
            "h1": "政治多极化",
            "core": "相互尊重、公平正义、合作共赢的新型国际关系",
            "title": "2026丰台期末Q21",
            "fields": {
                "材料触发点": q21_trigger
                + "回答这一关系时，应说明相关合作同时增进各方利益，不能只解释为单方面受益或单向援助。",
                "设问": q21_question,
                "为什么能想到": q21_common
                + "材料要点集中在倡议如何把不同国家带到共同合作方向上，答题应从国际关系形态落到相互尊重、公平正义、合作共赢。",
                "答案落点": "四大全球倡议为构建人类命运共同体凝聚合作力量，有利于推动构建相互尊重、公平正义、合作共赢的新型国际关系。",
            },
        },
        {
            "target": "2026丰台期末Q21",
            "h1": "政治多极化",
            "core": "践行真正的多边主义",
            "title": "2026丰台期末Q21",
            "fields": {
                "材料触发点": q21_trigger
                + "回答这一关系时，应说明共同问题要在多边框架中平等协商解决，反对单边主义和阵营对抗。",
                "设问": q21_question,
                "为什么能想到": q21_common
                + "题眼是“为构建人类命运共同体凝聚合力”，不是单个国家单独行动；因此要写各国平等协商、共同应对共同问题的真正多边主义。",
                "答案落点": "四大全球倡议为构建人类命运共同体凝聚合作力量，体现我国践行真正的多边主义，主张各国在平等协商中合作解决共同问题。",
            },
        },
        {
            "target": "2026通州一模Q19",
            "h1": "时代背景",
            "core": "和平与发展仍是时代主题",
            "title": "2026通州一模Q19",
            "fields": {
                "材料触发点": "设问要求分析中国元首外交如何为世界注入稳定性与正能量；材料呈现中俄、中美、欧洲、周边国家和全球南方合作，说明答题不能只写外交活动本身，还要说明这些活动顺应和平、发展、合作、共赢的时代潮流。",
                "设问": tq_question,
                "为什么能想到": "题干直接把元首外交的作用概括为“稳定性与正能量”，这是时代主题题的信号。写法上先判断世界处在动荡变革中，再说明中国以对话合作回应共同愿望，为和平发展提供稳定支撑。",
                "答案落点": "中国元首外交顺应和平与发展的时代主题，通过沟通对话、互信合作和共同发展，为动荡变革的世界注入稳定性与正能量。",
            },
        },
        {
            "target": "2026通州一模Q19",
            "h1": "理论",
            "core": "共同利益是合作的基础",
            "title": "2026通州一模Q19",
            "fields": {
                "材料触发点": "材料既有中俄关系稳定定位，也有中美关系穿越风浪、周边睦邻友好和全球南方合作，说明不同国家之间存在可以通过合作实现的共同利益。",
                "设问": tq_question,
                "为什么能想到": "题目问“如何注入稳定性与正能量”，不能只写中国单方面努力；多组双边和多边镜头共同出现，提示答案要抓国家间共同利益，通过扩大合作基础来解释稳定来源。",
                "答案落点": "中国元首外交立足各方共同利益，推动大国关系平稳前行、周边关系睦邻友好、全球南方团结合作，从而为国际关系稳定提供合作基础。",
            },
        },
        {
            "target": "2026通州一模Q19",
            "h1": "政治多极化",
            "core": "相互尊重、公平正义、合作共赢的新型国际关系",
            "title": "2026通州一模Q19",
            "fields": {
                "材料触发点": "材料中习近平主席强调推动中美关系迈向相互尊重、合作共赢，并提出大国要带头讲平等、讲法治、讲合作、讲诚信。",
                "设问": tq_question,
                "为什么能想到": "材料直接出现“相互尊重、合作共赢”和“平等、法治、合作、诚信”，这些不是普通外交修辞，而是处理国际关系新思路的关键词，应上升到新型国际关系。",
                "答案落点": "中国元首外交推动构建相互尊重、公平正义、合作共赢的新型国际关系，引导大国关系和国际交往走向平等、法治、合作、诚信。",
            },
        },
        {
            "target": "2026通州一模Q19",
            "h1": "政治多极化",
            "core": "国际关系民主化与国际政治经济新秩序",
            "title": "2026通州一模Q19",
            "fields": {
                "材料触发点": "材料强调大国要带头讲平等、讲法治，并坚定维护以联合国为核心的国际体系和以国际法为基础的国际秩序。",
                "设问": tq_question,
                "为什么能想到": "“讲平等、讲法治”和“国际法为基础的国际秩序”指向国际事务不能由强权支配，而应由各国在规则基础上平等协商，答题应落到国际关系民主化和更加公正合理的国际秩序。",
                "答案落点": "中国元首外交倡导国际关系民主化，推动国际秩序和全球治理体系朝着更加公正合理的方向发展。",
            },
        },
        {
            "target": "2026通州一模Q19",
            "h1": "中国",
            "core": "中国特色大国外交与人类命运共同体",
            "title": "2026通州一模Q19",
            "fields": {
                "材料触发点": "设问聚焦中国元首外交；材料呈现中俄、中美、欧洲、周边、上合组织和中拉论坛等多层外交活动，最后落到让世界看到真实立体的中国并赢得更深信任。",
                "设问": tq_question,
                "为什么能想到": "“元首外交”是中国特色大国外交的直接题眼，材料又把多层外交活动统一到世界稳定和中国形象上，说明答案应从中国方案、中国担当和人类命运共同体方向组织。",
                "答案落点": "中国元首外交展现中国特色大国外交的主动作为，推动构建人类命运共同体，向世界展示负责任大国形象，为全球治理贡献中国智慧和中国方案。",
            },
        },
        {
            "target": "2026通州一模Q19",
            "h1": "联合国",
            "core": "遵循联合国宪章宗旨和原则",
            "title": "2026通州一模Q19",
            "fields": {
                "材料触发点": "材料明确写到坚定维护以联合国为核心的国际体系和以国际法为基础的国际秩序。",
                "设问": tq_question,
                "为什么能想到": "出现“以联合国为核心的国际体系”和“国际法为基础的国际秩序”，就不能只写一般外交合作，而要点出维护联合国权威、遵循联合国宪章宗旨和原则。",
                "答案落点": "中国元首外交坚定维护以联合国为核心的国际体系和以国际法为基础的国际秩序，维护《联合国宪章》宗旨和原则。",
            },
        },
    ]


def load_old_candidates(path: Path) -> list[dict]:
    raw = json.loads(path.read_text(encoding="utf-8"))
    entries = []
    for item in raw:
        fields = item.get("current_fields") or {}
        entries.append(
            {
                "target": item.get("target", ""),
                "h1": item.get("h1", ""),
                "core": clean_core(item.get("core", "")),
                "title": clean_title(item.get("title", "")),
                "fields": {
                    "材料触发点": fields.get("材料触发点", "").strip(),
                    "设问": fields.get("设问", "").strip(),
                    "为什么能想到": fields.get("为什么能想到", "").strip(),
                    "答案落点": fields.get("答案落点", "").strip(),
                    "同题组": fields.get("同题组", "").strip(),
                },
            }
        )
    return entries


def target_rank(target: str) -> int:
    for i, t in enumerate(TARGET_ORDER):
        if target == t:
            return i
    return len(TARGET_ORDER)


def h1_rank(h1: str) -> int:
    return H1_ORDER.index(h1) if h1 in H1_ORDER else len(H1_ORDER)


def add_para_before(marker, text: str = "", style: str | None = None):
    return marker.insert_paragraph_before(text, style=style)


def write_entry(marker, number: int, entry: dict) -> int:
    title = entry.get("title") or entry.get("target") or "补入题"
    add_para_before(marker, f"{number}. {title}", "Normal")
    fields = entry.get("fields", {})
    for label in ["材料触发点", "设问", "为什么能想到", "答案落点"]:
        value = fields.get(label, "").strip()
        if value:
            add_para_before(marker, f"【{label}】{value}", "Normal")
    group = fields.get("同题组", "").strip()
    if group:
        lines = split_group_lines(group)
        if lines:
            add_para_before(marker, "【同题组】" + lines[0], "Normal")
            for line in lines[1:]:
                add_para_before(marker, f"· {line}", "Normal")
    add_para_before(marker, "", "Normal")
    return number + 1


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--candidates", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--summary", type=Path, required=True)
    args = parser.parse_args()

    doc = Document(args.input)
    markers = [p for p in doc.paragraphs if p.style.name == "Heading 1" and p.text.strip() == "六大要素术语极简速记版"]
    if not markers:
        raise RuntimeError("找不到“六大要素术语极简速记版”标题，无法定位补入位置。")
    marker = markers[0]

    entries = load_old_candidates(args.candidates) + supplemental_manual_entries()
    entries.sort(key=lambda e: (h1_rank(e["h1"]), e["core"], target_rank(e["target"]), e["title"]))

    pagebreak = add_para_before(marker, "", "Normal")
    pagebreak.add_run().add_break(WD_BREAK.PAGE)
    add_para_before(marker, "三维报告补入题（按六大要素归入）", "Heading 1")
    add_para_before(
        marker,
        "说明：本区只补入三维审核报告点名缺漏的题目，按六大要素归类；为保留正文排版，未重排原主体条目。",
        "Normal",
    )

    grouped: dict[str, dict[str, list[dict]]] = defaultdict(lambda: defaultdict(list))
    for entry in entries:
        grouped[entry["h1"]][entry["core"]].append(entry)

    total_written = 0
    per_target: dict[str, int] = defaultdict(int)
    number = 1
    for h1 in H1_ORDER:
        if h1 not in grouped:
            continue
        add_para_before(marker, h1, "Heading 2")
        for core in sorted(grouped[h1].keys()):
            add_para_before(marker, f"核心答题点：{core}", "Heading 3")
            for entry in grouped[h1][core]:
                number = write_entry(marker, number, entry)
                total_written += 1
                per_target[entry["target"]] += 1

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    summary = {
        "input": str(args.input),
        "output": str(args.output),
        "total_added_entries": total_written,
        "per_target": dict(sorted(per_target.items(), key=lambda kv: target_rank(kv[0]))),
    }
    args.summary.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
