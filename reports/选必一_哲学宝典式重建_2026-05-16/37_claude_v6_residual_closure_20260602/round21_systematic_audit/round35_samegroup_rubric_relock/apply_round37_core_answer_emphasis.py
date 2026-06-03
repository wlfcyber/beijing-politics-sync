#!/usr/bin/env python3
from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor


INPUT = Path("/Users/wanglifei/Desktop/选必一6.1终极版_同题组细则再锁_学生版无工程痕迹_带水印_20260603.docx")
OUTPUT = Path("/Users/wanglifei/Desktop/选必一6.1终极版_同题组细则再锁_核心答题点红色强化_学生版_带水印_20260603.docx")
AUDIT = Path(__file__).with_name("round37_core_answer_emphasis_audit.json")

DEEP_RED = RGBColor(192, 0, 0)


def iter_paragraphs(container):
    for paragraph in container.paragraphs:
        yield paragraph
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def is_core_answer_paragraph(text: str) -> bool:
    stripped = text.strip()
    return stripped.startswith("核心答题点：") or stripped.startswith("【核心答题点】")


def emphasize_core_answer_points(doc: Document) -> dict:
    touched = 0
    bad_after = []

    for paragraph in iter_paragraphs(doc):
        if not is_core_answer_paragraph(paragraph.text):
            continue
        touched += 1
        for run in paragraph.runs:
            if not run.text:
                continue
            run.bold = True
            run.font.color.rgb = DEEP_RED
            run.font.size = Pt(12)

    for idx, paragraph in enumerate(iter_paragraphs(doc)):
        if not is_core_answer_paragraph(paragraph.text):
            continue
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            color = run.font.color.rgb if run.font.color is not None else None
            size = run.font.size.pt if run.font.size is not None else None
            if run.bold is not True or color != DEEP_RED or abs(size - 12) > 0.01:
                bad_after.append(
                    {
                        "paragraph_index": idx,
                        "text": paragraph.text[:120],
                        "run": run.text[:80],
                        "bold": run.bold,
                        "color": str(color),
                        "size": size,
                    }
                )

    return {
        "input": str(INPUT),
        "output": str(OUTPUT),
        "core_answer_points_touched": touched,
        "bad_after": bad_after,
    }


def main() -> None:
    if not INPUT.exists():
        raise FileNotFoundError(INPUT)
    shutil.copyfile(INPUT, OUTPUT)
    doc = Document(OUTPUT)
    audit = emphasize_core_answer_points(doc)
    doc.save(OUTPUT)
    AUDIT.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(audit, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
