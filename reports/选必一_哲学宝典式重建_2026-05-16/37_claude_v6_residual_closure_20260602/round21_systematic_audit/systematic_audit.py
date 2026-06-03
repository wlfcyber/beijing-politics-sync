from __future__ import annotations

import csv
import json
import re
import unicodedata
from collections import defaultdict
from pathlib import Path

from docx import Document


DOCX = Path("/Users/wanglifei/Desktop/选必一6.1最终版_第20轮同题组纠错_带水印_20260603.docx")
RAW_DIR = Path("/Users/wanglifei/GaokaoPolitics/beijing-politics-sync/reports/选必一_哲学宝典式重建_2026-05-16/27_from_zero_source_locked_rebuild_20260531/02_source_cards/raw_cards")
OUT_DIR = Path("/Users/wanglifei/GaokaoPolitics/beijing-politics-sync/reports/选必一_哲学宝典式重建_2026-05-16/37_claude_v6_residual_closure_20260602/round21_systematic_audit")


def norm(s: str) -> str:
    s = unicodedata.normalize("NFKC", s)
    s = s.replace("«", "《").replace("»", "》")
    s = s.replace("ꎬ", "，").replace("ꎮ", "。").replace("􀆰", ".")
    s = s.replace("（", "(").replace("）", ")")
    s = s.replace("“", "\"").replace("”", "\"").replace("‘", "'").replace("’", "'")
    s = s.replace("；", ";").replace("，", ",").replace("。", ".").replace("：", ":")
    s = s.replace("—", "-").replace("－", "-")
    return re.sub(r"\s+", "", s)


def compact(s: str) -> str:
    return re.sub(r"\s+", "", s or "")


def clean_key_title(title: str) -> str | None:
    m = re.search(r"(20\d{2}.+?Q\d+(?:\(\d+\))?)", title.replace(" ", ""))
    if not m:
        return None
    key = m.group(1)
    key = key.replace("细则复练", "")
    return key


def key_to_card_id(key: str) -> str:
    key = key.replace("思政", "")
    m = re.match(r"(20\d{2})(.+?)Q(\d+)(?:\((\d+)\))?$", key)
    if not m:
        return ""
    year, suite, q, sub = m.groups()
    # Longest exam-name first.
    exam_names = ["期末", "期中", "一模", "二模"]
    exam = ""
    district = suite
    for name in exam_names:
        if suite.endswith(name):
            exam = name
            district = suite[: -len(name)]
            break
    parts = [year, district, exam, f"Q{q}"]
    if sub:
        parts.append(sub)
    return "_".join([p for p in parts if p])


def extract_source_sections(card_text: str) -> tuple[str, str]:
    paper = ""
    rubric = ""
    if "## 原卷片段" in card_text:
        paper = card_text.split("## 原卷片段", 1)[1].split("## 细则片段", 1)[0]
    if "## 细则片段" in card_text:
        rubric = card_text.split("## 细则片段", 1)[1]
    return paper, rubric


def parse_docx(path: Path):
    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs]
    styles = [p.style.name if p.style else "" for p in doc.paragraphs]
    entries = []
    cur_h1 = cur_h2 = cur_h3 = ""
    for idx, (para, text) in enumerate(zip(doc.paragraphs, paras)):
        style = para.style.name if para.style else ""
        if style.startswith("Heading 1"):
            cur_h1 = text
        elif style.startswith("Heading 2"):
            cur_h2 = text
        elif style.startswith("Heading 3") or text.startswith("核心答题点："):
            cur_h3 = text.replace("核心答题点：", "").strip()
        if re.match(r"^\d+\.\s*20\d{2}.+?Q\d+", text):
            end = idx + 1
            while end < len(paras):
                next_text = paras[end]
                next_style = styles[end]
                if end > idx and (
                    re.match(r"^\d+\.\s*20\d{2}.+?Q\d+", next_text)
                    or next_style.startswith("Heading 1")
                    or next_style.startswith("Heading 2")
                    or next_style.startswith("Heading 3")
                    or next_text.startswith("核心答题点：")
                ):
                    break
                end += 1
            block = paras[idx:end]
            fields = {}
            tong_lines = []
            in_tong = False
            for line in block:
                for label in ["材料触发点", "设问", "为什么能想到", "答案落点"]:
                    if line.startswith(f"【{label}】"):
                        fields[label] = line.split("】", 1)[1].strip()
                if line.startswith("【同题组】"):
                    in_tong = True
                    tong_lines.append(line)
                    continue
                if in_tong and line.startswith("·"):
                    tong_lines.append(line)
                    continue
                if in_tong and line.startswith("【"):
                    in_tong = False
            key = clean_key_title(text) or text
            entries.append(
                {
                    "idx": idx,
                    "end": end,
                    "title": text,
                    "key": key,
                    "card_id": key_to_card_id(key),
                    "h1": cur_h1,
                    "h2": cur_h2,
                    "h3": cur_h3,
                    "fields": fields,
                    "tong_lines": tong_lines,
                    "block": block,
                }
            )
    return doc, paras, entries


def parse_tong_terms(lines: list[str]) -> set[tuple[str, str]]:
    terms = set()
    for line in lines:
        if not line.startswith("·"):
            continue
        body = line.lstrip("·").strip()
        if "：" not in body:
            continue
        bucket, rest = body.split("：", 1)
        bucket = bucket.strip()
        for term in re.split(r"[；;]", rest):
            term = term.strip()
            if term:
                terms.add((bucket, term))
    return terms


def issue(kind: str, severity: str, key: str, title: str, detail: str, suggestion: str = ""):
    return {
        "kind": kind,
        "severity": severity,
        "key": key,
        "title": title,
        "detail": detail,
        "suggestion": suggestion,
    }


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc, paras, entries = parse_docx(DOCX)
    by_key = defaultdict(list)
    for e in entries:
        by_key[e["key"]].append(e)

    issues = []

    # Required fields and same-key prompt consistency.
    for e in entries:
        for label in ["材料触发点", "设问", "为什么能想到", "答案落点"]:
            if not e["fields"].get(label):
                issues.append(issue("missing_field", "high", e["key"], e["title"], f"缺少【{label}】"))

    for key, group in by_key.items():
        prompts = {compact(e["fields"].get("设问", "")) for e in group if e["fields"].get("设问")}
        if len(prompts) > 1:
            issues.append(issue("same_question_prompt_mismatch", "high", key, group[0]["title"], f"同题号出现 {len(prompts)} 种设问文本"))

    # 同题组 consistency: existing lines may omit the current entry, but must not omit others or include nonexistent terms.
    for key, group in by_key.items():
        if len(group) <= 1:
            continue
        expected = {(e["h1"], e["h3"]) for e in group if e["h1"] and e["h3"]}
        for e in group:
            if not e["tong_lines"]:
                issues.append(issue("missing_tongtizu", "high", key, e["title"], "同一题号有多个条目，但本条没有【同题组】"))
                continue
            actual = parse_tong_terms(e["tong_lines"])
            actual_plus_current = set(actual)
            actual_plus_current.add((e["h1"], e["h3"]))
            missing = expected - actual_plus_current
            extra = actual - expected
            if missing:
                issues.append(
                    issue(
                        "tongtizu_missing_peer",
                        "high",
                        key,
                        e["title"],
                        "同题组漏列真实同题核心：" + "；".join([f"{a}:{b}" for a, b in sorted(missing)]),
                    )
                )
            if extra:
                issues.append(
                    issue(
                        "tongtizu_extra_peer",
                        "high",
                        key,
                        e["title"],
                        "同题组列了不存在的同题核心：" + "；".join([f"{a}:{b}" for a, b in sorted(extra)]),
                    )
                )

    # Prompt source-card rough matching.
    raw_cards = {p.stem: p for p in RAW_DIR.glob("*.md")}
    for key, group in by_key.items():
        card_id = group[0]["card_id"]
        card = raw_cards.get(card_id)
        if not card:
            issues.append(issue("source_card_missing", "medium", key, group[0]["title"], f"未找到源卡 {card_id}"))
            continue
        text = card.read_text(encoding="utf-8", errors="ignore")
        paper, rubric = extract_source_sections(text)
        paper_norm = norm(paper)
        for prompt in sorted({e["fields"].get("设问", "") for e in group if e["fields"].get("设问")}):
            prompt_norm = norm(prompt)
            if prompt_norm and prompt_norm not in paper_norm:
                # Give a softer severity because OCR/source cards may contain selected prompts split by artifacts.
                issues.append(
                    issue(
                        "prompt_not_exact_in_source_card",
                        "medium",
                        key,
                        group[0]["title"],
                        f"当前设问未能在源卡原卷片段中精确匹配：{prompt[:120]}",
                        "需要人工回源核对，不能据此单独判错。",
                    )
                )

    # Known old-error relapse scans.
    full_text = "\n".join(paras)
    known_patterns = [
        ("old_A_tongzhou_qizhong", "2026通州期中Q20|2026通州期中Q21", "通州期中/期末题号旧错回流"),
        ("old_A_fengtai_qizhong", "2026丰台期中Q20", "丰台期中/期末题号旧错回流"),
        ("old_B_haidian_tieqie", "知识运用准确、贴切", "海淀二模设问旧错回流"),
        ("old_B_yanqing_rewrite", "阐释中国对维护全球产供链韧性和稳定的政策主张", "延庆一模设问改写旧错回流"),
        ("old_B_fengtai_q20_rewrite", "说明.?五大工程.?如何推动中拉双方在各自现代化征程上并肩前行", "丰台期末Q20设问改写旧错回流"),
        ("old_B_shijingshan_q18_meaning", "中国与东盟经贸合作不断提质升级的原因和意义", "石景山二模Q18设问虚增意义旧错回流"),
        ("old_B_dongcheng_q20_title", "关于.?同球共济.?精神和中国行动的小论文", "东城二模Q20题面限定虚增旧错回流"),
        ("old_B_xicheng_q19_prefix", "面对新一轮前沿科技发展浪潮", "西城一模Q19(6)前缀虚增旧错回流"),
        ("old_B_freetrade_missing", "自由贸易试验区/自由贸易港", "丰台二模Q20自由贸易试验港漏字旧错回流"),
        ("old_B_tongzhou_punct", "正逢其时、指引方向、彰显担当", "通州期末Q20顿号旧错回流"),
        ("old_C_yugong_fabrication", "中国愿做愚公移山", "海淀期末Q22材料编造旧错回流"),
        ("old_C_deepseek_140_more", "140多个国家", "房山一模Q18(2)数字虚增旧错回流"),
        ("backend_terms", "细则依据|证据层级|v7漏了|要落到|采分点", "后台/复查话术残留"),
    ]
    for kind, pat, desc in known_patterns:
        hits = re.findall(pat, full_text)
        if hits:
            issues.append(issue(kind, "high", "GLOBAL", "", f"{desc}：命中 {len(hits)} 次，pattern={pat}"))

    # Produce tables.
    issue_path = OUT_DIR / "systematic_audit_issues.json"
    issue_path.write_text(json.dumps(issues, ensure_ascii=False, indent=2), encoding="utf-8")

    csv_path = OUT_DIR / "entries_snapshot.csv"
    with csv_path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=["idx", "key", "card_id", "h1", "h2", "h3", "title", "question", "tong_lines"])
        writer.writeheader()
        for e in entries:
            writer.writerow(
                {
                    "idx": e["idx"],
                    "key": e["key"],
                    "card_id": e["card_id"],
                    "h1": e["h1"],
                    "h2": e["h2"],
                    "h3": e["h3"],
                    "title": e["title"],
                    "question": e["fields"].get("设问", ""),
                    "tong_lines": " | ".join(e["tong_lines"]),
                }
            )

    # Human-readable summary.
    by_kind = defaultdict(int)
    by_sev = defaultdict(int)
    for it in issues:
        by_kind[it["kind"]] += 1
        by_sev[it["severity"]] += 1
    lines = [
        "# 第21轮系统审计 初扫报告",
        "",
        f"- docx: `{DOCX}`",
        f"- entries: {len(entries)}",
        f"- unique question keys: {len(by_key)}",
        f"- repeated question groups: {sum(1 for g in by_key.values() if len(g) > 1)}",
        f"- issues: {len(issues)}",
        "",
        "## Issue Counts By Severity",
        "",
    ]
    for k in sorted(by_sev):
        lines.append(f"- {k}: {by_sev[k]}")
    lines += ["", "## Issue Counts By Kind", ""]
    for k, v in sorted(by_kind.items()):
        lines.append(f"- {k}: {v}")
    lines += ["", "## High Issues", ""]
    for it in issues:
        if it["severity"] == "high":
            lines.append(f"- [{it['kind']}] {it['key']} {it['title']} :: {it['detail']}")
    lines += ["", "## Medium Issues Needing Manual Review", ""]
    for it in issues:
        if it["severity"] == "medium":
            lines.append(f"- [{it['kind']}] {it['key']} {it['title']} :: {it['detail']}")
    (OUT_DIR / "systematic_audit_summary.md").write_text("\n".join(lines), encoding="utf-8")
    print(json.dumps({"entries": len(entries), "unique_keys": len(by_key), "issues": len(issues), "by_kind": dict(by_kind), "by_severity": dict(by_sev)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
