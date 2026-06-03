from __future__ import annotations

import json
import re
import shutil
from collections import defaultdict
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


SRC = Path("/Users/wanglifei/Desktop/选必一6.1最终版_第20轮同题组纠错_带水印_20260603.docx")
OUT = Path("/Users/wanglifei/Desktop/选必一6.1最终版_第21轮系统同题组与设问校正_带水印_20260603.docx")
RAW_DIR = Path("/Users/wanglifei/GaokaoPolitics/beijing-politics-sync/reports/选必一_哲学宝典式重建_2026-05-16/27_from_zero_source_locked_rebuild_20260531/02_source_cards/raw_cards")
OUT_DIR = Path("/Users/wanglifei/GaokaoPolitics/beijing-politics-sync/reports/选必一_哲学宝典式重建_2026-05-16/37_claude_v6_residual_closure_20260602/round21_systematic_audit")


def norm(s: str) -> str:
    s = s.replace("（", "(").replace("）", ")")
    s = s.replace("“", "\"").replace("”", "\"").replace("‘", "'").replace("’", "'")
    s = s.replace("；", ";").replace("，", ",").replace("。", ".").replace("：", ":")
    s = s.replace("？", "?").replace("！", "!")
    return re.sub(r"\s+", "", s)


def clean_key_title(title: str) -> str | None:
    m = re.search(r"(20\d{2}.+?Q\d+(?:\(\d+\))?)", title.replace(" ", ""))
    if not m:
        return None
    return m.group(1).replace("细则复练", "")


def key_to_card_id(key: str) -> str:
    key = key.replace("思政", "")
    m = re.match(r"(20\d{2})(.+?)Q(\d+)(?:\((\d+)\))?$", key)
    if not m:
        return ""
    year, suite, q, sub = m.groups()
    exam = ""
    district = suite
    for name in ["期末", "期中", "一模", "二模"]:
        if suite.endswith(name):
            exam = name
            district = suite[: -len(name)]
            break
    parts = [year, district, exam, f"Q{q}"]
    if sub:
        parts.append(sub)
    return "_".join([p for p in parts if p])


def source_paper_norm(card_id: str) -> str:
    path = RAW_DIR / f"{card_id}.md"
    if not path.exists():
        return ""
    text = path.read_text(encoding="utf-8", errors="ignore")
    if "## 原卷片段" in text:
        text = text.split("## 原卷片段", 1)[1].split("## 细则片段", 1)[0]
    return norm(text)


def set_para_text(p: Paragraph, text: str):
    runs = list(p.runs)
    if not runs:
        p.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def insert_after_clone(anchor: Paragraph, template: Paragraph, text: str) -> Paragraph:
    new_p = deepcopy(template._p)
    anchor._p.addnext(new_p)
    new_para = Paragraph(new_p, anchor._parent)
    set_para_text(new_para, text)
    return new_para


def delete_para(p: Paragraph):
    el = p._element
    el.getparent().remove(el)


def parse(doc: Document):
    paras = [p.text.strip() for p in doc.paragraphs]
    styles = [p.style.name if p.style else "" for p in doc.paragraphs]
    entries = []
    cur_h1 = cur_h2 = cur_h3 = ""
    for idx, (para, text) in enumerate(zip(doc.paragraphs, paras)):
        style = para.style.name if para.style else ""
        if style.startswith("Heading 1"):
            cur_h1 = text
        elif style.startswith("Heading 2"):
            cur_h2 = text
        elif style.startswith("Heading 3") or text.startswith("核心答题点："):
            cur_h3 = text.replace("核心答题点：", "").strip()
        if re.match(r"^\d+\.\s*20\d{2}.+?Q\d+", text):
            end = idx + 1
            while end < len(paras):
                nt = paras[end]
                ns = styles[end]
                if end > idx and (
                    re.match(r"^\d+\.\s*20\d{2}.+?Q\d+", nt)
                    or ns.startswith("Heading 1")
                    or ns.startswith("Heading 2")
                    or ns.startswith("Heading 3")
                    or nt.startswith("核心答题点：")
                ):
                    break
                end += 1
            block = paras[idx:end]
            fields = {}
            field_indices = {}
            tong_start = None
            tong_end = None
            for i in range(idx, end):
                line = paras[i]
                for label in ["材料触发点", "设问", "为什么能想到", "答案落点"]:
                    if line.startswith(f"【{label}】"):
                        fields[label] = line.split("】", 1)[1].strip()
                        field_indices[label] = i
                if line.startswith("【同题组】") and tong_start is None:
                    tong_start = i
                    j = i + 1
                    while j < end and paras[j].startswith("·"):
                        j += 1
                    tong_end = j
            key = clean_key_title(text) or text
            entries.append(
                {
                    "idx": idx,
                    "end": end,
                    "title": text,
                    "key": key,
                    "card_id": key_to_card_id(key),
                    "h1": cur_h1,
                    "h2": cur_h2,
                    "h3": cur_h3,
                    "fields": fields,
                    "field_indices": field_indices,
                    "tong_start": tong_start,
                    "tong_end": tong_end,
                }
            )
    return paras, entries


def canonical_prompt(key: str, group: list[dict]) -> str:
    prompts = sorted({e["fields"].get("设问", "") for e in group if e["fields"].get("设问")})
    if not prompts:
        return ""
    paper = source_paper_norm(group[0]["card_id"])

    def score(prompt: str) -> tuple[int, int, int, int, str]:
        pnorm = norm(prompt)
        in_source = 1 if paper and pnorm in paper else 0
        chinese_quote = prompt.count("“") + prompt.count("”")
        has_score = 1 if re.search(r"（\d+分）|\(\d+分\)", prompt) else 0
        return (in_source, chinese_quote, has_score, len(pnorm), prompt)

    # If no source card confirms any variant, choose the most common variant, then the most natural Chinese-quote version.
    counts = defaultdict(int)
    for e in group:
        if e["fields"].get("设问"):
            counts[e["fields"]["设问"]] += 1
    best = max(prompts, key=lambda p: (score(p), counts[p]))
    return best


def group_lines(group: list[dict]) -> list[str]:
    seen_buckets = []
    for e in group:
        if e["h1"] and e["h1"] not in seen_buckets:
            seen_buckets.append(e["h1"])
    lines = []
    for bucket in seen_buckets:
        terms = []
        seen_terms = set()
        for e in group:
            if e["h1"] == bucket and e["h3"] and e["h3"] not in seen_terms:
                seen_terms.add(e["h3"])
                terms.append(e["h3"])
        if terms:
            lines.append(f"· {bucket}：" + "；".join(terms))
    return lines


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)
    header_tpl = next((p for p in doc.paragraphs if p.text.strip().startswith("【同题组】")), None)
    bullet_tpl = next((p for p in doc.paragraphs if p.text.strip().startswith("· ")), None)
    if header_tpl is None or bullet_tpl is None:
        raise RuntimeError("同题组模板段落不存在")

    paras, entries = parse(doc)
    by_key = defaultdict(list)
    for e in entries:
        by_key[e["key"]].append(e)

    prompt_changes = []
    tong_rebuilds = []
    canonical_by_key = {key: canonical_prompt(key, group) for key, group in by_key.items()}
    group_lines_by_key = {key: group_lines(group) for key, group in by_key.items() if len(group) > 1}
    header_by_key = {}
    for key, group in by_key.items():
        score_header = ""
        for e in group:
            if e["tong_start"] is not None:
                old = paras[e["tong_start"]]
                if "分" in old:
                    score_header = old
                    break
        header_by_key[key] = score_header or "【同题组】"

    # Change field text in place first. This does not affect paragraph indices.
    for e in entries:
        canon = canonical_by_key.get(e["key"], "")
        old = e["fields"].get("设问", "")
        idx = e["field_indices"].get("设问")
        if canon and old != canon and idx is not None:
            set_para_text(doc.paragraphs[idx], f"【设问】 {canon}")
            prompt_changes.append({"key": e["key"], "title": e["title"], "old": old, "new": canon})

    # Rebuild all 同题组 blocks from actual entries, bottom-up to keep indices stable.
    for e in sorted(entries, key=lambda x: x["idx"], reverse=True):
        if len(by_key[e["key"]]) <= 1:
            continue
        new_lines = [header_by_key[e["key"]]] + group_lines_by_key[e["key"]]
        start = e["tong_start"]
        end = e["tong_end"]
        if start is None:
            # Insert after answer point if missing.
            anchor_idx = e["field_indices"].get("答案落点", e["end"] - 1)
            anchor = doc.paragraphs[anchor_idx]
            for n, line in enumerate(new_lines):
                anchor = insert_after_clone(anchor, header_tpl if n == 0 else bullet_tpl, line)
            tong_rebuilds.append({"key": e["key"], "title": e["title"], "action": "inserted", "lines": new_lines})
            continue
        old_lines = paras[start:end]
        if old_lines == new_lines:
            continue
        # Set first old paragraphs, add missing, remove extras.
        available = list(range(start, end))
        for n, idx in enumerate(available[: len(new_lines)]):
            set_para_text(doc.paragraphs[idx], new_lines[n])
        anchor = doc.paragraphs[available[min(len(available), len(new_lines)) - 1]]
        for n in range(len(available), len(new_lines)):
            anchor = insert_after_clone(anchor, bullet_tpl, new_lines[n])
        for idx in reversed(available[len(new_lines) :]):
            delete_para(doc.paragraphs[idx])
        tong_rebuilds.append({"key": e["key"], "title": e["title"], "action": "rebuilt", "old": old_lines, "new": new_lines})

    doc.save(OUT)
    report = {
        "source": str(SRC),
        "output": str(OUT),
        "entries": len(entries),
        "prompt_changes": prompt_changes,
        "tong_rebuild_count": len(tong_rebuilds),
        "tong_rebuilds_sample": tong_rebuilds[:50],
    }
    (OUT_DIR / "round21_apply_report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"output": str(OUT), "prompt_changes": len(prompt_changes), "tong_rebuilds": len(tong_rebuilds)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
