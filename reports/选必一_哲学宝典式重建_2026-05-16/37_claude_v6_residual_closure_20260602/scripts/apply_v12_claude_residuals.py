from __future__ import annotations

import argparse
import glob
import json
from pathlib import Path

from docx import Document


def set_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def update_para_snapshot(entry: dict, para_idx: int, text: str) -> None:
    for para in entry.get("paras", []):
        if para.get("idx") == para_idx:
            para["text"] = text


def apply_field(entry: dict, label: str, value: str, changes: dict[int, str]) -> None:
    para_idx = entry["field_para_idx"][label]
    text = f"【{label}】 {value}"
    entry["fields"][label] = value
    changes[para_idx] = text
    update_para_snapshot(entry, para_idx, text)


def load_s5_outputs(outputs_dir: Path) -> dict[int, str]:
    result: dict[int, str] = {}
    for path_str in sorted(glob.glob(str(outputs_dir / "s5_out_*.json"))):
        path = Path(path_str)
        data = json.loads(path.read_text(encoding="utf-8"))
        for key, value in data.items():
            idx = int(key)
            if idx in result and result[idx] != value:
                raise RuntimeError(f"duplicate S5 key with different value: {idx}")
            result[idx] = value
    return result


def load_s5_targets(outputs_dir: Path) -> dict[int, dict]:
    result: dict[int, dict] = {}
    for path_str in sorted(glob.glob(str(outputs_dir / "s5_batch_*.json"))):
        data = json.loads(Path(path_str).read_text(encoding="utf-8"))
        for question, rec in data.items():
            for item in rec.get("items", []):
                idx = int(item["wh_idx"])
                result[idx] = {
                    "question": question,
                    "core": item.get("核心点", ""),
                    "answer": item.get("答案落点", ""),
                }
    return result


def norm(text: str) -> str:
    return (
        text.replace("核心答题点：", "")
        .replace("“", '"')
        .replace("”", '"')
        .replace("（", "(")
        .replace("）", ")")
        .replace(" ", "")
        .strip()
    )


def find_entry_for_s5(entries: list[dict], target: dict, fallback_para_idx: int | None = None) -> dict:
    question = target["question"]
    core = norm(target["core"])
    answer = norm(target["answer"])
    candidates = [entry for entry in entries if question in entry.get("entry_title", "")]
    if not candidates:
        candidates = entries

    scored: list[tuple[int, dict]] = []
    for entry in candidates:
        score = 0
        h3 = norm(entry.get("h3", ""))
        entry_answer = norm(entry.get("fields", {}).get("答案落点", ""))
        if core and (core in h3 or h3 in core):
            score += 5
        if answer and (answer in entry_answer or entry_answer in answer):
            score += 6
        if target["question"] in entry.get("entry_title", ""):
            score += 3
        if score:
            scored.append((score, entry))

    if fallback_para_idx is not None:
        for entry in entries:
            if entry.get("field_para_idx", {}).get("为什么能想到") == fallback_para_idx:
                scored.append((20, entry))

    if not scored:
        raise RuntimeError(f"S5 target entry not found: {target}")

    scored.sort(key=lambda item: item[0], reverse=True)
    best_score = scored[0][0]
    best = [entry for score, entry in scored if score == best_score]
    if len(best) > 1:
        # If the same core appears twice for one question, answer text should normally break the tie.
        names = [(entry["entry_title"], entry["h3"], entry.get("fields", {}).get("答案落点", "")[:80]) for entry in best]
        raise RuntimeError(f"ambiguous S5 target {target}: {names}")
    return best[0]


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input-docx", type=Path, required=True)
    parser.add_argument("--input-entries", type=Path, required=True)
    parser.add_argument("--s5-dir", type=Path, required=True)
    parser.add_argument("--output-docx", type=Path, required=True)
    parser.add_argument("--output-entries", type=Path, required=True)
    parser.add_argument("--log-json", type=Path, required=True)
    args = parser.parse_args()

    entries = json.loads(args.input_entries.read_text(encoding="utf-8"))
    changes: dict[int, str] = {}
    log: list[dict] = []

    s5 = load_s5_outputs(args.s5_dir)
    s5_targets = load_s5_targets(args.s5_dir)
    para_to_entry: dict[int, dict] = {}
    for entry in entries:
        why_idx = entry.get("field_para_idx", {}).get("为什么能想到")
        if why_idx is not None:
            para_to_entry[int(why_idx)] = entry

    for para_idx, why in sorted(s5.items()):
        target = s5_targets.get(para_idx)
        if target is None:
            entry = para_to_entry.get(para_idx)
            if entry is None:
                raise RuntimeError(f"S5 target not found in batch or entries: {para_idx}")
        else:
            entry = find_entry_for_s5(entries, target, fallback_para_idx=para_idx if para_idx in para_to_entry else None)
        apply_field(entry, "为什么能想到", why, changes)
        log.append(
            {
                "kind": "s5_rewrite",
                "s5_idx": para_idx,
                "para_idx": entry["field_para_idx"]["为什么能想到"],
                "title": entry["entry_title"],
                "h1": entry["h1"],
                "h2": entry["h2"],
                "h3": entry["h3"],
            }
        )

    targeted = [
        {
            "title": "2026石景山二模Q18",
            "h3_contains": "利用两个市场两种资源优化全球资源配置",
            "fields": {
                "为什么能想到": "材料真正的线索是互为重要市场、自贸区升级、数字基础设施和绿色发展合作。双方通过市场、技术、经验和规则联通扩大资源配置范围，因此会想到两个市场两种资源和资源配置效率。",
            },
        },
        {
            "title": "2025石景山一模Q17(2)",
            "h3_contains": "推进普惠包容的经济全球化",
            "fields": {
                "材料触发点": "原卷材料二写多个“北京宣言”，其中《和平共处五项原则发表70周年纪念大会北京宣言》提出“做大并分好全球经济‘蛋糕’”，让发展成果更多更公平地惠及各国，推动全球化朝着普惠包容方向发展。",
                "为什么能想到": "“做大并分好全球经济蛋糕”同时包含发展和共享两层意思：既要扩大世界经济发展空间，也要让成果更公平惠及各国。概述完善全球治理的中国主张时，这句话应转化为推进普惠包容的经济全球化。",
            },
        },
        {
            "title": "2024丰台一模Q20",
            "h3_contains": "推进贸易和投资自由化便利化",
            "fields": {
                "为什么能想到": "515家中外企业和机构参展、签署200多项协议，说明链博会把供应链上下游企业放到同一合作平台中对接。参展和签约数据体现跨境合作成本下降、合作渠道拓宽，因此会想到贸易投资自由化便利化。",
            },
        },
        {
            "title": "2024丰台一模Q20供应链如何成为",
            "h3_contains": "利用两个市场两种资源优化全球资源配置",
            "fields": {
                "为什么能想到": "高速铁路网、高速公路网和世界级港口群为跨国货物、服务、生产要素流动提供通道。供应链成为共赢链，靠的是国内外市场和资源被更顺畅地连接起来，提升全球配置效率。",
            },
        },
        {
            "title": "2024丰台一模Q20",
            "h3_contains": "维护全球产业链供应链稳定畅通",
            "field_contains": {"材料触发点": "数字化技术"},
            "fields": {
                "为什么能想到": "数字化技术、绿色低碳技术对应供应链成本高、信息不对称、环节不透明、流程不标准等痛点。技术赋能能降低成本、提高效率、改善绿色水平，因此指向供应链稳定畅通和高质量发展。",
            },
        },
        {
            "title": "2024丰台一模Q20",
            "h3_contains": "以产业链供应链合作推动互利共赢",
            "fields": {
                "为什么能想到": "基础设施联通、供应链金融创新、链博会平台、数字绿色技术分别打通物流、资金、对接渠道和技术供给。多维合作共同服务供应链畅通，最终指向各国在产业链供应链合作中互利共赢。",
            },
        },
        {
            "title": "2026顺义二模Q20",
            "h3_contains": "互利共赢的开放战略",
            "fields": {
                "为什么能想到": "微电网、打井、物流、希望农场等项目把中国开放同各国民生改善连在一起。材料说让经济全球化更有活力，说明开放带来的不是单方收益，而是发展成果惠及各国，正对应互利共赢的开放战略。",
            },
        },
        {
            "title": "2025海淀二模Q21",
            "h3_contains": "维护《联合国宪章》宗旨和原则",
            "fields": {
                "材料触发点": "中国担任安理会轮值主席、推动安理会讨论聚焦实质性问题，体现中国以实际行动维护联合国权威和宪章宗旨原则。",
                "为什么能想到": "安理会是联合国维护国际和平与安全的重要机构，中国担任轮值主席并推动讨论聚焦实质性问题，说明中国不是旁观者，而是在联合国框架内维护宪章宗旨和原则。",
            },
        },
    ]

    for rule in targeted:
        hits = []
        for entry in entries:
            if rule["title"] not in entry["entry_title"]:
                continue
            if rule.get("h3_contains") and rule["h3_contains"] not in entry["h3"]:
                continue
            ok = True
            for label, needle in rule.get("field_contains", {}).items():
                if needle not in entry.get("fields", {}).get(label, ""):
                    ok = False
                    break
            if ok:
                hits.append(entry)
        if not hits:
            raise RuntimeError(f"targeted rule missed: {rule}")
        for entry in hits:
            for label, value in rule["fields"].items():
                apply_field(entry, label, value, changes)
            log.append(
                {
                    "kind": "targeted",
                    "title": entry["entry_title"],
                    "h1": entry["h1"],
                    "h2": entry["h2"],
                    "h3": entry["h3"],
                    "fields": sorted(rule["fields"]),
                }
            )

    cleanup_replacements = {
        "设问要求": "题目要求",
        "设问要": "题目指向",
        "所以会想到": "因此联到",
    }
    for entry in entries:
        for label, value in list(entry.get("fields", {}).items()):
            new_value = value
            for old, new in cleanup_replacements.items():
                new_value = new_value.replace(old, new)
            if new_value != value:
                apply_field(entry, label, new_value, changes)
                log.append(
                    {
                        "kind": "cleanup",
                        "title": entry["entry_title"],
                        "h1": entry["h1"],
                        "h2": entry["h2"],
                        "h3": entry["h3"],
                        "fields": [label],
                    }
                )

    # Keep the cached paragraph snapshots in sync with the cleaned fields.
    for entry in entries:
        for label, value in entry.get("fields", {}).items():
            para_idx = entry.get("field_para_idx", {}).get(label)
            if para_idx is not None:
                update_para_snapshot(entry, int(para_idx), f"【{label}】 {value}")

    doc = Document(args.input_docx)
    if changes and max(changes) >= len(doc.paragraphs):
        raise RuntimeError(f"paragraph index {max(changes)} >= paragraph count {len(doc.paragraphs)}")
    for para_idx, text in changes.items():
        set_text(doc.paragraphs[para_idx], text)
    doc.save(args.output_docx)

    args.output_entries.write_text(json.dumps(entries, ensure_ascii=False, indent=2), encoding="utf-8")
    args.log_json.write_text(json.dumps(log, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"s5={sum(1 for item in log if item['kind'] == 's5_rewrite')}")
    print(f"targeted={sum(1 for item in log if item['kind'] == 'targeted')}")
    print(f"paragraph_changes={len(changes)}")


if __name__ == "__main__":
    main()
