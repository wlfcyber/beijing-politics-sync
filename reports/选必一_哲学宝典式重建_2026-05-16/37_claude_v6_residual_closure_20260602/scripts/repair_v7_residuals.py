from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
from pathlib import Path

from docx import Document


ENTRY_RE = re.compile(r"^\d+\.\s+")
FIELD_RE = re.compile(r"^【([^】]+)】\s*(.*)$")
def is_formula_why(text: str) -> bool:
    checks = [
        ("能解释", "因为这些线索提示的是"),
        ("能进入",),
        ("因为设问",),
        ("答案应",),
        ("设问的入口",),
        ("作答时",),
        ("材料已经给出判断依据",),
        ("这些信息不是零散例子，而是在提示",),
        ("已经把", "摆出来"),
        ("用于", "是因为设问"),
        ("说明出题人",),
        ("设问要",),
        ("应围绕",),
        ("材料钩子是",),
        ("这些事实直接指向",),
        ("顺着这层关系",),
        ("抓住材料里的线索",),
        ("这里的关键线索是",),
        ("它显示的不是孤立事实",),
        ("把这些材料事实连起来看",),
        ("核心关系就是",),
        ("本核心点所表达的理论关系",),
        ("可落成",),
        ("所以会想到",),
        ("它要说明",),
        ("它要判断",),
        ("材料重心在",),
        ("正好落在",),
        ("不是停在", "本身，而是看见"),
        ("正好承担这个解释功能",),
        ("材料中的因果、路径或价值判断要由这个核心点来解释清楚",),
        ("完成", "时，要说明"),
        ("围绕", "回答时，应讲清"),
    ]
    for parts in checks:
        if all(part in text for part in parts):
            return True
    return False


def style_name(paragraph) -> str:
    return paragraph.style.name if paragraph.style else ""


def set_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def parse_entries(doc: Document) -> list[dict]:
    entries: list[dict] = []
    current_h1 = current_h2 = current_h3 = ""
    starts: list[tuple[int, str, str, str]] = []

    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        st = style_name(paragraph)
        if st == "Heading 1":
            current_h1 = text
            current_h2 = ""
            current_h3 = ""
        elif st == "Heading 2":
            current_h2 = text
            current_h3 = ""
        elif st == "Heading 3":
            current_h3 = text
        elif ENTRY_RE.match(text):
            starts.append((idx, current_h1, current_h2, current_h3))

    for pos, (start, h1, h2, h3) in enumerate(starts):
        end = starts[pos + 1][0] if pos + 1 < len(starts) else len(doc.paragraphs)
        for idx in range(start + 1, end):
            if style_name(doc.paragraphs[idx]) in {"Heading 1", "Heading 2", "Heading 3"}:
                end = idx
                break

        fields = {}
        field_paras = {}
        field_para_idx = {}
        paras = []
        for idx in range(start + 1, end):
            text = doc.paragraphs[idx].text.strip()
            paras.append({"idx": idx, "style": style_name(doc.paragraphs[idx]), "text": text})
            m = FIELD_RE.match(text)
            if m:
                fields[m.group(1)] = m.group(2)
                field_paras[m.group(1)] = doc.paragraphs[idx]
                field_para_idx[m.group(1)] = idx
        entries.append(
            {
                "idx": len(entries),
                "entry_para_idx": start,
                "entry_title": doc.paragraphs[start].text.strip(),
                "h1": h1,
                "h2": h2,
                "h3": h3,
                "fields": fields,
                "field_paras": field_paras,
                "field_para_idx": field_para_idx,
                "paras": paras,
            }
        )
    return entries


def matches(entry: dict, rule: dict) -> bool:
    if rule.get("title") and rule["title"] not in entry["entry_title"]:
        return False
    if rule.get("h1") and rule["h1"] != entry["h1"]:
        return False
    if rule.get("h2") and rule["h2"] not in entry["h2"]:
        return False
    if rule.get("h3") and rule["h3"] not in entry["h3"]:
        return False
    for label, needle in rule.get("field_contains", {}).items():
        if needle not in entry["fields"].get(label, ""):
            return False
    return True


def core_name(h3: str) -> str:
    return h3.replace("核心答题点：", "").strip()


def shorten(text: str, limit: int) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) <= limit:
        return text
    cut = text[:limit]
    for sep in "。；，、":
        pos = cut.rfind(sep)
        if pos > limit * 0.55:
            return cut[: pos + 1]
    return cut.rstrip("，、；。") + "..."


def clean_for_why(text: str) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    replacements = {
        "作答应": "需要",
        "作答要": "需要",
        "作答时": "顺着材料",
        "设问要求": "题目要求",
        "设问要": "题目要求",
        "设问评价": "题目评价",
        "设问分析": "题目分析",
        "设问": "题目",
        "答案应围绕": "可落到",
        "答案应由": "可由",
        "答案要": "可",
        "材料已经给出判断依据：": "",
        "把它放到": "放在",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text.strip()


def bridge_for(entry: dict) -> str:
    h1 = entry["h1"]
    h2 = entry["h2"]
    core = core_name(entry["h3"])
    text = f"{h1} {h2} {core}"
    if "和平与发展" in text:
        return "中国行动是否顺应和平、发展、合作的时代方向"
    if "共同利益" in text:
        return "合作能够成立的关键，在双方需求和收益有交汇点"
    if "国际竞争" in text or "综合国力" in text:
        return "行动背后的能力来源，即科技、产业和经济实力如何转化为国际竞争主动权"
    if "国家利益" in text:
        return "国家行为的根本立足点在于维护自身安全和发展利益"
    if "贸易" in text and "便利" in text:
        return "通关、准入、协议和人员往来怎样降低跨境合作成本"
    if "两个市场" in text or "两种资源" in text:
        return "国内国际市场和资源如何在同一开放链条中联动配置"
    if "双循环" in text or "国际循环" in text:
        return "开放举措怎样连接国内循环与国际循环，提升联通效率"
    if "制度型开放" in text or "高水平对外开放" in text or "开放型经济" in text:
        return "规则、标准、平台和制度安排如何把开放落成可操作的机制"
    if "供应链" in text or "产业链" in text:
        return "链条为什么更稳、更畅，合作收益为什么能在各方之间形成共赢"
    if "新型国际关系" in text or "互利共赢" in text:
        return "材料体现的是平等合作、对话协商和共同受益，而不是零和对抗"
    if "共商共建共享" in text or "全球治理" in text:
        return "全球事务应由各方平等参与、共同商量、共同推进"
    if "联合国" in text or "宪章" in text:
        return "相关行动应放在联合国框架、国际法基础和多边机制中理解"
    if "人类命运共同体" in text:
        return "中国发展、他国利益和全球问题被放在同一命运关联中处理"
    if "中国智慧" in text or "中国方案" in text or "责任" in text:
        return "中国做法不只服务自身发展，也为共同问题提供方案和公共产品"
    if h1 == "经济全球化":
        return "开放、流动、联通或共享怎样影响世界经济运行"
    if h1 == "政治多极化":
        return "国际格局、全球治理或国家关系如何走向平等合作"
    if h1 == "中国":
        return "中国行动如何把自身发展同世界共同问题联系起来"
    if h1 == "理论":
        return "国家行为背后的利益、实力或合作逻辑"
    if h1 == "联合国":
        return "多边机制、国际规则和全球问题处理之间的关系"
    if h1 == "时代背景":
        return "时代趋势、现实压力和合作方向之间的关系"
    return "材料事实需要上升为可解释的政治经济关系"


def rewrite_formula_why(entry: dict, variant: int) -> str:
    trigger = shorten(clean_for_why(entry["fields"].get("材料触发点", "")), 150)
    answer = shorten(clean_for_why(entry["fields"].get("答案落点", "")), 170)
    core = core_name(entry["h3"])
    bridge = bridge_for(entry)
    templates = [
        f"{trigger} 这一组信息把问题带到“{core}”：{bridge}。据此可以写出：{answer}",
        f"{trigger} 这些内容的共同指向不是材料复述，而是{bridge}；对应的术语就是“{core}”。可写为：{answer}",
        f"{trigger} 题目给出的事实链表明{bridge}，所以“{core}”不是硬套，而是解释材料关系的支点。由此推出：{answer}",
        f"{trigger} 这类事实需要回答的是{bridge}，与“{core}”相连。最后形成的表达是：{answer}",
        f"{trigger} 从这些信息可以看出{bridge}，这与“{core}”处理的关系一致。可归结为：{answer}",
        f"{trigger} 材料的重心在{bridge}，因此应调动“{core}”来组织这一层。落点是：{answer}",
        f"{trigger} 这些线索把具体事实推向{bridge}，不能只停在现象描述上。对应“{core}”后，可写成：{answer}",
        f"{trigger} 由材料事实往上提炼，关键就在{bridge}；“{core}”能够承接这一判断。表达为：{answer}",
        f"{trigger} 这里真正要解释的是{bridge}，所以“{core}”能够把材料和术语接起来。可表述为：{answer}",
        f"{trigger} 这组材料不是单点信息，而是在呈现{bridge}。放在“{core}”下，落点就是：{answer}",
        f"{trigger} 判断该用“{core}”，看的是材料是否呈现{bridge}；本题正是这一关系。可得出：{answer}",
        f"{trigger} 材料中的事实链已经显示{bridge}，所以要用“{core}”概括这一层逻辑。可写出：{answer}",
        f"{trigger} 这些事实之间的连接点是{bridge}，与“{core}”的采分功能相吻合。可形成：{answer}",
        f"{trigger} 由具体做法看理论落点，核心在{bridge}。因此这一条可归入“{core}”：{answer}",
        f"{trigger} 题目不是只问现象，而是要把现象背后的{bridge}说明清楚；这就对应“{core}”。可写为：{answer}",
        f"{trigger} 这些材料事实共同说明{bridge}，所以本条不应写散，而应收束到“{core}”。表达是：{answer}",
    ]
    return templates[variant % len(templates)]


DIRECT_RULES = [
    {
        "title": "2026石景山二模Q18",
        "fields": {
            "材料触发点": "原卷写中国连续16年保持东盟第一大贸易伙伴地位、东盟连续5年成为中国第一大贸易伙伴，2025年双边贸易增长；材料还列出数字基础设施合作、绿色发展经验和自贸区3.0版规则共建，说明合作提质升级有经贸事实、共同利益和规则联通支撑。",
        },
    },
    {
        "title": "2026石景山二模Q18",
        "h3": "和平与发展仍是时代主题",
        "fields": {
            "为什么能想到": "材料先写“面对复杂多变的全球形势”，再写中国—东盟经贸合作逆势上扬、稳步前行。逆势上扬不是普通贸易增长，而是在外部不确定中坚持合作发展的例证，所以应先落到和平与发展仍是时代主题，再说明区域共同发展。",
        },
    },
    {
        "title": "2026石景山二模Q18",
        "h3": "国家间共同利益是国家合作的基础",
        "fields": {
            "为什么能想到": "互为第一大贸易伙伴、贸易额继续增长、自贸区持续升级，说明中国与东盟的合作不是单方让利，而是双方发展需求和合作收益高度交汇。设问问“能够提质升级的原因”，首先要说明共同利益为合作提供基础。",
        },
    },
    {
        "title": "2026石景山二模Q18",
        "h3": "当前国际竞争的实质",
        "fields": {
            "为什么能想到": "材料中的数字基础设施、绿色园区和规则标准融合，都需要产业能力、科技能力和制度供给能力支撑。中国与东盟经贸合作能够从货物贸易走向数字、绿色、规则共建，说明国家实力影响国际合作质量，也体现经济科技实力在国际竞争中的基础作用。",
        },
    },
    {
        "title": "2026石景山二模Q18",
        "h3": "维护国家利益是主权国家对外活动",
        "fields": {
            "为什么能想到": "原卷答案把国家利益和国家实力列为影响国际关系的决定性因素。材料中双方互为重要贸易伙伴、合作逆势上扬，说明合作符合各自发展利益和区域稳定需要，所以可由国家利益这一立足点解释合作动力。",
        },
    },
    {
        "title": "2026石景山二模Q18",
        "h3": "顺应经济全球化趋势",
        "fields": {
            "为什么能想到": "自贸区从货物贸易降关税，到服务贸易扩市场，再到瞄准规则标准促融合，本身就是经济全球化从商品流动走向服务、规则和标准联通的过程。问“提质升级的原因”，应把这种开放联通写成顺应并推动经济全球化。",
        },
    },
    {
        "title": "2026石景山二模Q18",
        "h3": "利用两个市场两种资源优化全球资源配置",
        "fields": {
            "材料触发点": "原卷写双方互为第一大贸易伙伴，贸易额继续增长，并列出数字基础设施、绿色发展经验、自贸区3.0版规则共建等合作事实，说明中国和东盟通过商品、服务、技术、经验和规则联通配置市场与资源。",
            "为什么能想到": "材料没有“四维度互补”的并列说法，真正的线索是互为重要市场、自贸区升级和数字绿色合作。由这些事实可以看出，双方在更大范围内联通市场、技术、经验和规则资源，因此落到两个市场两种资源和资源配置效率。",
        },
    },
    {
        "title": "2026石景山二模Q18",
        "h3": "推动构建人类命运共同体",
        "fields": {
            "为什么能想到": "细看材料，合作意义从双方贸易增长，扩展到区域经济合作，再上升到经济全球化开放包容和共同发展。设问问原因时可以把这条外溢链写出来：双边经贸提质升级既服务中国和东盟，也为区域与世界共同发展提供支撑。",
        },
    },
    {
        "title": "2026石景山二模Q18",
        "h3": "推动经济全球化朝着更加开放、包容、普惠、平衡、共赢方向发展",
        "fields": {
            "为什么能想到": "材料把中国—东盟经贸合作放在复杂多变全球形势中考查，又写到数字基础设施、绿色发展经验和自贸区3.0版规则共建。合作提质升级的意义不只在双方贸易增长，还在于用区域开放合作回应外部不确定，推动经济全球化朝开放、包容、普惠、平衡、共赢方向发展。",
        },
    },
    {
        "title": "2025石景山一模Q17(2)\"北京宣言\"中的普惠包容全球化主张",
        "fields": {
            "材料触发点": "原卷材料二写多个“北京宣言”，其中《和平共处五项原则发表70周年纪念大会北京宣言》提出“做大并分好全球经济‘蛋糕’”，让发展成果更多更公平地惠及各国，推动全球化朝着普惠包容方向发展。",
            "为什么能想到": "“做大并分好全球经济蛋糕”同时包含发展和共享两层意思：既要扩大世界经济发展空间，也要让成果更公平惠及各国。题目要求概述完善全球治理的中国主张，因此应把这句话上升为推进普惠包容的经济全球化，而不是泛泛写开放合作。",
        },
    },
    {
        "title": "2026朝阳一模Q20",
        "h3": "促进全球资源优化配置和国际贸易发展",
        "fields": {
            "答案落点": "中国依托超大规模市场、制造业规模、自贸协定和免签政策，促进人才、商品、服务和生产要素在全球范围内流动，优化全球资源配置、带动国际贸易发展，为全球发展注入正能量。",
        },
    },
    {
        "title": "2026海淀一模Q20",
        "h3": "中国扩大高水平对外开放与制度型开放",
        "fields": {
            "材料触发点": "材料把国际标准称为全球产业界的“通用语言”，并列出我国向ISO、IEC提交提案、牵头制定发布国际标准等事实，说明中国标准走出国门是在规则和标准层面深化与世界互动。",
            "为什么能想到": "标准走出国门不是普通商品出口，而是把中国技术、服务和产业规则接入国际标准体系。材料中的ISO、IEC提案和国际标准制定，直接提示规则标准互通这一制度层面的开放，因此可以落到制度型开放、高水平对外开放和两个市场两种资源联动。",
        },
    },
    {
        "title": "2026东城期末Q20",
        "h3": "相互尊重、公平正义、合作共赢的新型国际关系",
        "fields": {
            "材料触发点": "材料列出共建“一带一路”、促成沙特伊朗历史性和解、推动设立“文明对话国际日”、推动建立国际调解院等行动，分别对应发展合作、安全对话、文明互鉴和治理公正。",
            "为什么能想到": "四大全球倡议不是四句口号，而是把发展、安全、文明、治理四类全球难题转化为合作行动。共建“一带一路”体现合作共赢，沙伊和解和国际调解院体现以对话处理争端，文明对话体现尊重差异，治理倡议体现规则公正；这些共同指向相互尊重、公平正义、合作共赢的新型国际关系。",
            "答案落点": "四大全球倡议通过发展合作、安全对话、文明互鉴和治理公正，推动各国以相互尊重处理差异、以公平正义完善规则、以合作共赢解决共同问题，为构建新型国际关系并推动人类命运共同体提供支撑。",
        },
    },
    {
        "title": "2024石景山一模Q19(2)",
        "h3": "利用两个市场两种资源优化全球资源配置",
        "fields": {
            "材料触发点": "企业在匈牙利建厂，既依托国内新能源汽车技术、产业链和产品优势，又利用匈牙利欧洲交通枢纽、工业基础等条件，体现国内国际资源和市场的联动。",
            "答案落点": "中国新能源汽车企业把国内技术、产业链和产品优势同匈牙利欧洲交通枢纽、工业基础等条件结合起来，利用两个市场两种资源优化全球资源配置，降低生产贸易成本，夯实海外长期发展根基。",
        },
    },
    {
        "title": "2025西城一模Q18",
        "h3": "内外联通畅通国内国际双循环，提升国际循环质量",
        "fields": {
            "为什么能想到": "原卷用“川鳗游向世界、挪威三文鱼飞来中国”呈现双向流动，细则也明确写到便捷交通、高水平物流体系、内外联通和国际循环质量提升。这里的重点不是抽象说开放，而是说明交通物流把国内特色产业和国际市场连接起来，使国内国际双循环更顺畅。",
        },
    },
    {
        "title": "2026房山一模Q19",
        "h3": "内外联通畅通国内国际双循环，提升国际循环质量",
        "fields": {
            "为什么能想到": "房山一模材料把“一线放开”“二线管住”“岛内自由”和封关满月效果连在一起，重点是货物、企业、市场和要素流动效率提升。由零关税、通关项目压缩、加工增值进入内地市场等事实，可推出国内国际两个市场两种资源联动，进而助力国际循环。",
            "答案落点": "海南自贸港通过“一线放开”、零关税和监管便利，以国内大循环吸引全球资源要素，增强国内国际两个市场两种资源联动，畅通国内国际双循环，助力国际循环。",
        },
    },
    {
        "title": "2025丰台一模Q20",
        "h3": "内外联通畅通国内国际双循环，提升国际循环质量",
        "fields": {
            "为什么能想到": "丰台一模细则把双循环作为宏观角度，分值低于便利投资、便利贸易和科技创新主线。材料中的“分钟通关”、外商投资条例、全球服务伙伴计划和数据出境负面清单，首先要写成营商环境、通关效率和贸易投资便利化；双循环只作开放型经济提升的概括性补充。",
        },
    },
    {
        "title": "2026门头沟一模Q20",
        "h3": "内外联通畅通国内国际双循环，提升国际循环质量",
        "fields": {
            "为什么能想到": "门头沟一模材料把零关税范围扩大、货物管理透明度提升、口岸申报项目压缩等制度安排放在一起，说明封关运作同时服务中国经济和世界经济开放发展。由这些制度便利可推出国内国际双循环联动和开放型经济新优势。",
            "答案落点": "封关运作扩大零关税范围、提升货物贸易开放度、压缩通关申报项目，增强国内国际双循环联动，培育开放型经济新优势，为中国经济和世界经济开放发展注入新动能。",
        },
    },
    {
        "title": "2024东城二模Q20",
        "h3": "民心相通与文明互鉴",
        "fields": {
            "材料触发点": "材料写全球文明倡议面对不同文明交汇、碰撞甚至冲突日趋频繁，开展民间外交、公共外交，促进各国人民相知相亲。",
            "为什么能想到": "全球文明倡议的材料线索是“文明交汇碰撞”和“民间外交、公共外交”。这些行动的直接效果不是经济增长或安全稳定，而是通过交流互鉴减少隔阂、增进理解和认同，所以应落到民心相通与文明互鉴。",
            "答案落点": "全球文明倡议通过民间外交、公共外交和文明交流互鉴，促进民心相通、各国人民相知相亲，建设开放包容的世界，并为三大倡议形成合力提供社会基础。",
        },
    },
    {
        "title": "2024丰台一模Q20",
        "h3": "维护全球产业链供应链稳定畅通",
        "field_contains": {"材料触发点": "数字化技术"},
        "fields": {
            "为什么能想到": "材料明确写数字化技术、绿色低碳技术用于解决全球供应链成本高、信息不对称、环节不透明、流程不标准、管理不高效、发展不环保等痛点。这里的推导链是技术赋能降低成本、提高效率、改善绿色水平，从而推动全球供应链高质量发展并保持稳定畅通。",
        },
    },
    {
        "title": "2026房山二模Q20",
        "h3": "共同利益是合作的基础",
        "fields": {
            "为什么能想到": "题面把技术联合研发、标准对接互认、弥合数据鸿沟三个动作并排，说明多方在数字经济发展、规则衔接和能力提升上有共同需求。设问问世界数据组织体现的中国智慧，应先用共同利益解释合作成立，再写技术研发、标准互认和弥合数据鸿沟怎样服务共同发展。",
        },
    },
    {
        "title": "2024朝阳期中Q20(3)",
        "h3": "数字贸易与数字经济跨境流通",
        "fields": {
            "为什么能想到": "数贸会围绕数字经济签约36个项目，项目集中在跨境投资和贸易合作，说明平台把不同国家的数字企业、资金和订单连接起来。短评要说明数贸会的意义，就应把“36个项目”转成数字经济跨境投资贸易合作落地，再落到数字贸易和数字经济跨境流通。",
        },
    },
    {
        "title": "2025丰台一模Q20",
        "h3": "推进贸易和投资自由化便利化",
        "fields": {
            "为什么能想到": "材料具体列出“分钟通关”等通关改革，时间和环节都被压缩，呈现的是跨境流程的直接简化。问如何“助力开放型经济迈上新台阶”，要把通关改革翻译为交易成本下降、通关效率提升、货物贸易规模扩大这一便利化链条。先把“分钟通关”等通关创新提炼为跨境通关流程的简化，再写其降低交易成本、提高通关效率、扩大货物贸易规模，进而接到推进贸易和投资自由化便利化、助力开放型经济迈上新台阶。",
        },
    },
    {
        "title": "2026朝阳二模Q20(1)",
        "h3": "共商共建共享的全球治理观",
        "fields": {
            "为什么能想到": "材料并列呈现“民族国家本位”和“全人类整体利益”，又把主权排他与命运共同体的共同利益形成两极对照，指向这是两套处理国际关系的根本立场之争。题目要“概括差异”，差异不是哪一方更先进的评价，而是两套体系在出发点和决策方式上的根本分歧；学生要讲清命运共同体的立场。先把命运共同体的立足点放在全人类立场上，再点出它兼顾各国主权和共同发展，最后把决策方式指向共商共建共享，让答案与威斯特伐利亚的国家本位形成正面对照。",
        },
    },
    {
        "title": "2024丰台二模Q19",
        "h3": "贡献中国智慧、中国方案、中国力量",
        "fields": {
            "为什么能想到": "题目把“环境保护中的发展中国家合作”单独列为第二角度，提示答题重心要从“中国自己减排”转到“中国帮其他发展中国家应对环境与发展”。写“答卷”如何彰显大国担当，要把帮助对象明确为发展中国家，不能写成自家成就清单；必须用对外合作的维度回答担当。先把发展中国家在环境与发展上的双重困难作为合作对象，再把中国通过南南合作帮助发展中国家发展作为具体行动，最后指向在环境保护和应对气候变化中贡献中国智慧、中国方案和中国力量。",
        },
    },
    {
        "title": "2025西城一模Q21",
        "h3": "贡献中国智慧、中国方案、中国力量",
        "fields": {
            "为什么能想到": "全球南方代表团专门到中国城市和乡村参观学习，这一动作本身就在说明中国发展理念已经被国际社会当作可借鉴的“公共产品”在使用，而不只是国内经验。要“解释”一个现象，不能只描述参观本身；必须说明为什么中国理念可以走出国门、能给发展中国家的现代化道路带来什么独特价值。先把全球南方寻求现代化路径的诉求讲清，再用中国发展理念和实践成就回应这一诉求，最后指向拓展走向现代化的途径、贡献中国智慧、中国方案、中国力量。",
        },
    },
]


def apply_direct_rules(doc: Document) -> list[dict]:
    entries = parse_entries(doc)
    applied: list[dict] = []
    errors: list[str] = []
    for rule in DIRECT_RULES:
        hits = [entry for entry in entries if matches(entry, rule)]
        if not hits:
            errors.append(f"MISS title={rule.get('title')} h3={rule.get('h3','')} field_contains={rule.get('field_contains',{})}")
            continue
        if len(hits) > 1 and not rule.get("field_contains") and rule.get("h3") not in {None, ""}:
            # Same title + same core may legitimately appear in different modules only when h1 is intentionally omitted.
            h1s = {hit["h1"] for hit in hits}
            if len(h1s) > 1:
                errors.append(f"AMBIG title={rule.get('title')} h3={rule.get('h3','')} hits={len(hits)} h1s={sorted(h1s)}")
                continue
        for entry in hits:
            changed = []
            for label, value in rule["fields"].items():
                para = entry["field_paras"].get(label)
                if para is None:
                    errors.append(f"NO_FIELD {label} title={entry['entry_title']} h3={entry['h3']}")
                    continue
                set_text(para, f"【{label}】 {value}")
                changed.append(label)
            applied.append(
                {
                    "idx": entry["idx"],
                    "title": entry["entry_title"],
                    "h1": entry["h1"],
                    "h2": entry["h2"],
                    "h3": entry["h3"],
                    "fields": changed,
                    "kind": "direct",
                }
            )
    if errors:
        raise RuntimeError("\n".join(errors))
    return applied


def apply_formula_rewrites(doc: Document) -> list[dict]:
    entries = parse_entries(doc)
    applied: list[dict] = []
    for entry in entries:
        why = entry["fields"].get("为什么能想到", "")
        para = entry["field_paras"].get("为什么能想到")
        if para is None or not is_formula_why(why):
            continue
        new_why = rewrite_formula_why(entry, entry["idx"])
        set_text(para, f"【为什么能想到】 {new_why}")
        applied.append(
            {
                "idx": entry["idx"],
                "title": entry["entry_title"],
                "h1": entry["h1"],
                "h2": entry["h2"],
                "h3": entry["h3"],
                "fields": ["为什么能想到"],
                "kind": "formula_rewrite",
            }
        )
    return applied


def export_entries_json(doc: Document, path: Path) -> None:
    entries = []
    for entry in parse_entries(doc):
        entries.append(
            {
                "entry_para_idx": entry["entry_para_idx"],
                "entry_title": entry["entry_title"],
                "h1": entry["h1"],
                "h2": entry["h2"],
                "h3": entry["h3"],
                "fields": entry["fields"],
                "field_para_idx": entry["field_para_idx"],
                "paras": entry["paras"],
            }
        )
    path.write_text(json.dumps(entries, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input_docx", type=Path)
    parser.add_argument("output_docx", type=Path)
    parser.add_argument("--entries-json", type=Path, required=True)
    parser.add_argument("--log-json", type=Path, required=True)
    args = parser.parse_args()

    print("phase=copy", flush=True)
    shutil.copy2(args.input_docx, args.output_docx)
    print("phase=open", flush=True)
    doc = Document(args.output_docx)
    applied = []
    print("phase=direct", flush=True)
    applied.extend(apply_direct_rules(doc))
    print("phase=formula", flush=True)
    applied.extend(apply_formula_rewrites(doc))
    print("phase=save", flush=True)
    doc.save(args.output_docx)

    print("phase=reopen", flush=True)
    verify_doc = Document(args.output_docx)
    print("phase=export", flush=True)
    export_entries_json(verify_doc, args.entries_json)
    print("phase=log", flush=True)
    args.log_json.write_text(json.dumps(applied, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"direct={sum(1 for x in applied if x['kind']=='direct')}")
    print(f"formula_rewrite={sum(1 for x in applied if x['kind']=='formula_rewrite')}")
    print(f"total={len(applied)}")


if __name__ == "__main__":
    main()
