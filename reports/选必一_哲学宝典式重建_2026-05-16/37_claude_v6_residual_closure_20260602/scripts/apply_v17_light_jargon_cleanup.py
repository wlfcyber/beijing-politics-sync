#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document


REPLACEMENTS = [
    (
        '题面提到"营造"良好外部政治环境，与上一点的"观念层"相对，这里要解决的是国际秩序方向问题——国际关系是否民主化、联合国地位是否稳固、治理体系是否改革、新秩序是否公正合理。"良好外部环境"在秩序层不能用一两个词收尾，要把四个相互衔接的主张串联起来给出完整方向。先把推动国际关系民主化作为方向定位，再把维护以联合国为核心的国际体系、以国际法为基础的国际秩序列为体系支柱，接着用推动全球治理体系变革完善作为更新路径，最后指向推动建立更加公正合理的国际政治经济新秩序。',
        '题面提到"营造"良好外部政治环境，材料集中呈现单边霸凌、动荡冲突和意识形态分歧等外部变量，说明本题要回答中国如何推动国际秩序朝更公正合理方向发展。国际关系民主化、维护以联合国为核心的国际体系、完善全球治理体系和建设更加公正合理的新秩序，是营造良好外部环境的连续主张。先把推动国际关系民主化作为方向定位，再把维护以联合国为核心的国际体系、以国际法为基础的国际秩序列为体系支柱，接着用推动全球治理体系变革完善作为更新路径，最后指向推动建立更加公正合理的国际政治经济新秩序。',
    ),
    (
        '中国在巴黎AI行动峰会上推动联合国大会通过人工智能能力建设国际合作决议；这把开源战略与全球AI治理串联，需要用"构建人类命运共同体"承担全球合作新模式的价值高点。',
        '中国在巴黎AI行动峰会上推动联合国大会通过人工智能能力建设国际合作决议；开源战略由此同全球AI治理相连，需要用"构建人类命运共同体"说明全球合作新模式的价值方向。',
    ),
    (
        '材料把中国开源战略与"巴黎AI行动峰会""联合国大会通过人工智能能力建设国际合作决议"串联，指向开源已超出技术层面，进入全球AI治理的合作架构。要"理解"开源战略是"全球合作新模式"，必须把开源放到国际合作的价值层面，而不是停留在技术开放本身。先把中国推动联合国通过AI能力建设国际合作决议作为关键事实点出，再写中国构建人类命运共同体的价值定位，进而接到开源战略让各国共享AI红利、形成全球合作新模式。',
        '材料把中国开源战略同"巴黎AI行动峰会""联合国大会通过人工智能能力建设国际合作决议"相连，指向开源已超出技术层面，进入全球AI治理的合作架构。要"理解"开源战略是"全球合作新模式"，必须把开源放到国际合作的价值层面，而不是停留在技术开放本身。先把中国推动联合国通过AI能力建设国际合作决议作为关键事实点出，再写中国构建人类命运共同体的价值定位，进而接到开源战略让各国共享AI红利、形成全球合作新模式。',
    ),
    (
        '题面指向"新时代中国特色大国外交"，外交理念在新阶段被进一步丰富；这一过渡说明指导思想本身在随时代演进，而非外交动作的零散变化。分析"变"的思想层面，要解释指导思想如何与时俱进、为何在新时代形成新的理论形态——不能停留在外交行动层。先把新时代中国特色大国外交作为思想载体接进来，再说明实践需要催生新的理论概括，最后指向习近平外交思想为中国特色大国外交提供根本遵循和行动指南。',
        '题面指向"新时代中国特色大国外交"，外交理念在新阶段被进一步丰富；这一过渡说明外交指导思想随时代演进，而不是外交动作的零散变化。分析新中国外交之"变"，要解释指导思想如何与时俱进、为何在新时代形成新的理论形态。先把新时代中国特色大国外交作为思想载体接进来，再说明实践需要催生新的理论概括，最后指向习近平外交思想为中国特色大国外交提供根本遵循和行动指南。',
    ),
    (
        '设问把中国特色大国外交的主动作为同中国发展大局相连，材料呈现“践行外交为民宗旨—构建海外安全保障体系—让中国公民、法人在海外更安全、更踏实、更放心”的关系链，触发维护人民利益与服务中国式现代化的目标层表达。',
        '设问把中国特色大国外交的主动作为同中国发展大局相连，材料呈现“践行外交为民宗旨—构建海外安全保障体系—让中国公民、法人在海外更安全、更踏实、更放心”的关系链，触发维护人民利益与服务中国式现代化的目标表达。',
    ),
    (
        '材料给出"外交为民宗旨—海外安全保障体系—中国公民、法人在海外更安全、更踏实、更放心"完整链条，三环都指向人民利益，不是抽象的国家利益。问外交为什么要"更有作为"，必须从目标层回答——外交为民、服务中国式现代化是更高层次的"作为方向"，不能仅写做了什么事。先把外交为民宗旨作为根本立场，再用海外安全保障体系建设作为践行宗旨的具体动作，最后指向维护人民利益、服务中国式现代化。',
        '材料给出"外交为民宗旨—海外安全保障体系—中国公民、法人在海外更安全、更踏实、更放心"完整链条，三环都指向人民利益，不是抽象的国家利益。问外交为什么要"更有作为"，要回答外交为民、服务中国式现代化这一作为方向，不能仅写做了什么事。先把外交为民宗旨作为根本立场，再用海外安全保障体系建设作为践行宗旨的具体动作，最后指向维护人民利益、服务中国式现代化。',
    ),
]

DELETE_PARAGRAPH_GROUPS = [
    [
        '· 【设问】结合材料，运用《当代国际政治与经济》知识，说明我国的实践是如何推动供应链成为国际合作“共赢链”的。',
        '· 【为什么能想到】',
        '· 【同题组】',
        '· 经济全球化：利用两个市场两种资源优化全球资源配置；推进贸易和投资自由化便利化；维护全球产业链供应链稳定畅通；以产业链供应链合作推动互利共赢',
    ],
    [
        '· 【设问】结合材料，运用《当代国际政治与经济》的相关知识，分析中国做“赋能型大国”的世界意义？',
        '· 【为什么能想到】',
        '· 【同题组】',
        '· 经济全球化：推动经济全球化朝着更加开放、包容、普惠、平衡、共赢方向发展；中国市场红利与全球经济包容性增长',
        '· 政治多极化：相互尊重、公平正义、合作共赢的新型国际关系',
        '· 中国：贡献中国智慧、中国方案、中国力量；中国推动构建人类命运共同体；促进技术共享和民生改善，为全球可持续发展贡献力量；倡导文明平等互鉴，弘扬全人类共同价值，尊重世界文明多样性',
    ],
]


def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    if old not in paragraph.text:
        return False
    full = paragraph.text.replace(old, new)
    if paragraph.runs:
        paragraph.runs[0].text = full
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(full)
    return True


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)


def delete_exact_groups(doc: Document) -> list[dict[str, object]]:
    deleted: list[dict[str, object]] = []
    paragraphs = list(doc.paragraphs)
    texts = [p.text for p in paragraphs]
    delete_ranges: list[tuple[int, int, list[str]]] = []
    for group in DELETE_PARAGRAPH_GROUPS:
        matches = [idx for idx in range(0, len(texts) - len(group) + 1) if texts[idx : idx + len(group)] == group]
        if len(matches) != 1:
            raise SystemExit(f"delete group match count {len(matches)} for: {group[0][:60]}")
        start = matches[0]
        delete_ranges.append((start, start + len(group), group))

    for start, end, group in sorted(delete_ranges, reverse=True):
        for paragraph in paragraphs[start:end]:
            delete_paragraph(paragraph)
        deleted.append({"start_paragraph": start, "paragraphs": len(group), "texts": group})
    deleted.sort(key=lambda item: item["start_paragraph"])
    return deleted


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input-docx", required=True)
    parser.add_argument("--output-docx", required=True)
    parser.add_argument("--log-json", required=True)
    args = parser.parse_args()

    doc = Document(args.input_docx)
    hits = []
    for idx, paragraph in enumerate(doc.paragraphs):
        for old, new in REPLACEMENTS:
            if replace_in_paragraph(paragraph, old, new):
                hits.append({"paragraph": idx, "old": old, "new": new})

    missing = [old for old, _ in REPLACEMENTS if not any(hit["old"] == old for hit in hits)]
    if missing:
        raise SystemExit(f"missing replacements: {len(missing)}")

    deleted = delete_exact_groups(doc)
    if len(deleted) != len(DELETE_PARAGRAPH_GROUPS):
        raise SystemExit(f"deleted groups mismatch: {len(deleted)} / {len(DELETE_PARAGRAPH_GROUPS)}")

    out = Path(args.output_docx)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    Path(args.log_json).write_text(
        json.dumps(
            {"changes": len(hits), "deleted_groups": len(deleted), "missing": 0, "hits": hits, "deleted": deleted},
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
