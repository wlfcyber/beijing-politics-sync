from __future__ import annotations

import argparse
import importlib.util
import json
import shutil
import sys
from pathlib import Path

from docx import Document


def load_module(path: Path, name: str):
    spec = importlib.util.spec_from_file_location(name, path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Cannot load module: {path}")
    mod = importlib.util.module_from_spec(spec)
    sys.modules[name] = mod
    spec.loader.exec_module(mod)
    return mod


def set_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input_docx", type=Path)
    parser.add_argument("input_entries_json", type=Path)
    parser.add_argument("output_docx", type=Path)
    parser.add_argument("--entries-json", type=Path, required=True)
    parser.add_argument("--log-json", type=Path, required=True)
    parser.add_argument("--helper", type=Path, required=True)
    parser.add_argument("--xml-helper", type=Path, required=True)
    args = parser.parse_args()

    helper = load_module(args.helper, "repair_v7_helper")
    xml_helper = load_module(args.xml_helper, "repair_v7_xml_helper")
    entries = json.loads(args.input_entries_json.read_text(encoding="utf-8"))
    changes, log = xml_helper.build_changes(entries, helper)

    shutil.copy2(args.input_docx, args.output_docx)
    doc = Document(args.output_docx)
    if max(changes) >= len(doc.paragraphs):
        raise RuntimeError(f"paragraph index {max(changes)} >= paragraph count {len(doc.paragraphs)}")
    for idx, text in changes.items():
        set_text(doc.paragraphs[idx], text)
    doc.save(args.output_docx)

    args.entries_json.write_text(json.dumps(entries, ensure_ascii=False, indent=2), encoding="utf-8")
    args.log_json.write_text(json.dumps(log, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"changes={len(changes)}")
    print(f"direct={sum(1 for x in log if x['kind']=='direct')}")
    print(f"formula_rewrite={sum(1 for x in log if x['kind']=='formula_rewrite')}")


if __name__ == "__main__":
    main()
