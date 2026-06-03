from __future__ import annotations

import csv
import json
import re
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


RUN_DIR = Path(__file__).resolve().parents[1]
BASE_DIR = RUN_DIR.parent
SOURCE_MD = (
    BASE_DIR
    / "21_claude_review_patch_20260528"
    / "选必一_当代国际政治与经济_主观题术语宝典_237条逐条核实二次复核最终稿_20260528.md"
)

OUT_MD = RUN_DIR / "选必一_当代国际政治与经济_主观题术语宝典_237条_二级结构版_20260528.md"
OUT_DOCX = RUN_DIR / "选必一_当代国际政治与经济_主观题术语宝典_237条_二级结构版_20260528.docx"
OUT_MAP = RUN_DIR / "SECOND_LEVEL_CORE_MAPPING.csv"
OUT_SUMMARY = RUN_DIR / "SECOND_LEVEL_REORG_QA.md"
OUT_JSON = RUN_DIR / "SECOND_LEVEL_REORG_QA.json"

BUCKETS = ["时代背景", "理论", "经济全球化", "政治多极化", "中国", "联合国"]

SECOND_LEVEL_ORDER = {
    "时代背景": [
        "时代主题与世界大势",
        "百年变局与外部风险",
        "全球治理议题与共同挑战",
    ],
    "理论": [
        "国家利益与国家安全",
        "共同利益与合作基础",
        "国际竞争与综合国力",
    ],
    "经济全球化": [
        "经济全球化趋势与开放型世界经济",
        "两个市场两种资源与国际分工",
        "贸易投资自由化便利化与制度型开放",
        "全球经济治理、规则制定与权益维护",
        "产业链供应链与开放安全",
        "数字经济、绿色转型与新兴领域",
        "贸易保护主义、多边贸易体制与开放环境",
    ],
    "政治多极化": [
        "全球治理观与治理体系变革",
        "国际关系民主化与新型国际关系",
        "人类命运共同体与共同现代化",
        "多边主义、国际秩序与国际法治",
        "世界多极化、全球南方与发展中国家",
        "气候治理、安全治理与文明互鉴",
    ],
    "中国": ["政策", "智慧", "责任", "地位"],
    "联合国": [
        "《联合国宪章》与国际体系",
        "联合国地位作用与多边舞台",
        "中国与联合国相互需要",
        "联合国发展议程与全球治理贡献",
    ],
}

BUCKET_GUIDES = {
    "时代背景": "先判断材料是否在说明中国主张、合作倡议或外交行动为什么顺势、应势、必要；这里负责回答“世界处在什么背景下”。",
    "理论": "先判断设问是否在追问合作、竞争、国家行为为什么成立；这里负责回答“国际关系运行的基本原理”。",
    "经济全球化": "先判断材料是否围绕开放、市场、贸易、产业链、规则治理、数字绿色转型和共享发展成果；这里负责回答“国际经济如何联通和运行”。",
    "政治多极化": "先判断材料是否围绕国际关系、国际秩序、全球治理方向和共同应对全球问题；这里负责回答“国际政治秩序怎样变、怎样治”。",
    "中国": "先判断材料是否强调中国主张、中国方案、中国行动、中国外交和中国贡献；本桶按政策、智慧、责任、地位检索。",
    "联合国": "先判断材料是否直接涉及《联合国宪章》、联合国地位作用、以联合国为核心的国际体系、联合国与中国关系或联合国发展议程。",
}


@dataclass
class CoreBlock:
    bucket: str
    title: str
    count: int | None
    lines: list[str]


def has_core_before_next_section(lines: list[str], idx: int) -> bool:
    for line in lines[idx + 1 : min(len(lines), idx + 20)]:
        if line.startswith("## "):
            return False
        if line.startswith("### 核心答题点："):
            return True
    return False


def find_actual_bucket_starts(lines: list[str]) -> dict[str, int]:
    starts: dict[str, int] = {}
    cursor = 0
    for bucket in BUCKETS:
        needle = f"## {bucket}"
        for i in range(cursor, len(lines)):
            if lines[i].strip() == needle and has_core_before_next_section(lines, i):
                starts[bucket] = i
                cursor = i + 1
                break
        else:
            raise RuntimeError(f"Cannot find actual bucket start: {bucket}")
    return starts


def find_speed_start(lines: list[str], after: int) -> int:
    for i in range(after, len(lines)):
        if lines[i].strip() == "## 六大要素术语极简速记版":
            return i
    return len(lines)


def parse_core_heading(line: str) -> tuple[str, int | None]:
    m = re.match(r"### 核心答题点：(.+?)(?:（出现(\d+)次）)?$", line.strip())
    if not m:
        raise RuntimeError(f"Bad core heading: {line}")
    title = m.group(1)
    count = int(m.group(2)) if m.group(2) else None
    return title, count


def parse_source(lines: list[str]) -> tuple[list[CoreBlock], list[str]]:
    starts = find_actual_bucket_starts(lines)
    speed_start = find_speed_start(lines, starts["联合国"] + 1)
    blocks: list[CoreBlock] = []

    for bi, bucket in enumerate(BUCKETS):
        start = starts[bucket]
        end = starts[BUCKETS[bi + 1]] if bi + 1 < len(BUCKETS) else speed_start
        section = lines[start + 1 : end]
        positions = [i for i, line in enumerate(section) if line.startswith("### 核心答题点：")]
        for pi, pos in enumerate(positions):
            block_end = positions[pi + 1] if pi + 1 < len(positions) else len(section)
            block_lines = section[pos:block_end]
            title, count = parse_core_heading(block_lines[0])
            blocks.append(CoreBlock(bucket=bucket, title=title, count=count, lines=block_lines))

    speed_lines = lines[speed_start:] if speed_start < len(lines) else []
    return blocks, speed_lines


def contains_any(text: str, words: list[str]) -> bool:
    return any(word in text for word in words)


def classify(block: CoreBlock) -> tuple[str, str]:
    title = block.title
    t = title
    bucket = block.bucket

    if bucket == "时代背景":
        if contains_any(t, ["和平与发展", "时代主题"]):
            return "时代主题与世界大势", "和平与发展、时代主题属于宏观世界大势"
        if contains_any(t, ["百年变局", "外部风险", "风险挑战"]):
            return "百年变局与外部风险", "材料背景侧重变局、风险和外部挑战"
        return "全球治理议题与共同挑战", "气候变化等全球性问题作为共同挑战进入时代背景"

    if bucket == "理论":
        if contains_any(t, ["国家利益", "国家安全"]):
            return "国家利益与国家安全", "解释主权国家行为的出发点和安全底线"
        if contains_any(t, ["共同利益", "合作的基础"]):
            return "共同利益与合作基础", "解释国家间合作为什么能够成立"
        return "国际竞争与综合国力", "解释竞争实质、科技经济实力和综合国力"

    if bucket == "经济全球化":
        if contains_any(t, ["贸易保护主义", "多边贸易体制", "开放合作的国际经济环境", "反对贸易保护主义"]):
            return "贸易保护主义、多边贸易体制与开放环境", "回应贸易壁垒、单边措施和多边贸易秩序"
        if contains_any(t, ["产业链", "供应链", "统筹开放与安全", "安全稳定"]):
            return "产业链供应链与开放安全", "围绕产业链供应链稳定和开放安全"
        if contains_any(t, ["数字", "数据", "科学技术", "技术", "新能源", "绿色转型", "未来产业", "创新型"]):
            return "数字经济、绿色转型与新兴领域", "围绕数字经济、科技创新、绿色转型和未来产业"
        if contains_any(t, ["规则", "标准", "全球经济治理", "国际组织赋予", "权益", "话语权", "经贸规则"]):
            return "全球经济治理、规则制定与权益维护", "围绕经济治理、规则制定、标准和权益维护"
        if contains_any(t, ["贸易和投资", "贸易投资", "制度型开放", "高水平对外开放", "开放型经济水平", "更高层次开放型经济", "自由化便利化", "零关税"]):
            return "贸易投资自由化便利化与制度型开放", "围绕市场开放、制度型开放和贸易投资便利化"
        if contains_any(t, ["两个市场", "两种资源", "资源配置", "国际分工", "国内国际双循环", "国际循环", "国际市场", "市场空间", "市场红利"]):
            return "两个市场两种资源与国际分工", "围绕市场资源联动、国际分工和国内国际循环"
        return "经济全球化趋势与开放型世界经济", "围绕经济全球化方向、开放型世界经济和共享成果"

    if bucket == "政治多极化":
        if contains_any(t, ["国际关系民主化", "新型国际关系", "平等互利", "互利共赢", "和平解决国际争端"]):
            return "国际关系民主化与新型国际关系", "围绕国家间关系形态和民主化方向"
        if contains_any(t, ["人类命运共同体", "共同现代化", "共同发展", "持久和平"]):
            return "人类命运共同体与共同现代化", "围绕命运共同体和共同现代化的国际秩序取向"
        if contains_any(t, ["多边主义", "国际秩序", "国际法治", "新秩序", "主权平等"]):
            return "多边主义、国际秩序与国际法治", "围绕多边主义、法治原则和公正合理国际秩序"
        if contains_any(t, ["世界多极化", "全球南方", "发展中国家", "国际组织"]):
            return "世界多极化、全球南方与发展中国家", "围绕多极化格局和发展中国家力量变化"
        if contains_any(t, ["气候", "安全", "文明", "绿色低碳", "跨国公司", "集团政治", "实力至上"]):
            return "气候治理、安全治理与文明互鉴", "围绕具体全球治理议题和治理逻辑转换"
        return "全球治理观与治理体系变革", "围绕共商共建共享和全球治理体系改革"

    if bucket == "中国":
        if contains_any(t, ["综合国力", "国际地位", "话语权", "创始会员国", "安理会", "重要力量", "参与者、贡献者和改革者"]):
            return "地位", "中国在国际体系中的身份、实力和影响力"
        if contains_any(t, ["总体国家安全观", "外交政策", "和平共处五项原则", "独立自主", "国家性质", "党对外交", "习近平外交思想", "自力更生"]):
            return "政策", "中国外交政策、国家安全和对外关系准则"
        if contains_any(t, ["智慧", "方案", "力量", "全球发展倡议", "全球安全倡议", "全球文明倡议", "全球治理倡议", "四大全球倡议", "人类命运共同体", "中拉命运共同体", "国际共识", "合作理念", "普遍价值", "文明平等互鉴", "全人类共同价值"]):
            return "智慧", "中国主张、中国方案、全球倡议和理念供给"
        return "责任", "中国行动、发展合作、南南合作、气候治理和负责任大国担当"

    if bucket == "联合国":
        if contains_any(t, ["联合国对中国", "中国与联合国"]):
            return "中国与联合国相互需要", "围绕中国和联合国双向关系"
        if contains_any(t, ["宪章", "以联合国为核心", "国际法"]):
            return "《联合国宪章》与国际体系", "围绕宪章原则、联合国核心体系和国际法秩序"
        if contains_any(t, ["联合国是", "最具普遍性", "代表性", "权威性", "主要舞台"]):
            return "联合国地位作用与多边舞台", "围绕联合国自身地位作用和多边平台功能"
        if contains_any(t, ["对中国", "中国作为联合国", "中国与联合国"]):
            return "中国与联合国相互需要", "围绕中国和联合国双向关系"
        return "联合国发展议程与全球治理贡献", "围绕联合国发展议程和全球性问题治理"

    raise RuntimeError(f"Unknown bucket: {bucket}")


def demote_core_block(block: CoreBlock) -> list[str]:
    out = []
    for i, line in enumerate(block.lines):
        if i == 0:
            out.append(line.replace("### 核心答题点：", "#### 核心答题点：", 1))
        else:
            out.append(line)
    return out


def demote_speed_lines(speed_lines: list[str]) -> list[str]:
    out: list[str] = []
    if not speed_lines:
        return out
    for i, line in enumerate(speed_lines):
        if i == 0 and line.startswith("## 六大要素术语极简速记版"):
            out.append(line)
        elif line in {f"## {b}" for b in BUCKETS}:
            out.append(line.replace("## ", "### ", 1))
        else:
            out.append(line)
    return out


def build_markdown(blocks: list[CoreBlock], speed_lines: list[str], assignments: dict[str, tuple[str, str]]) -> str:
    by_bucket_level: dict[str, dict[str, list[CoreBlock]]] = {
        bucket: {level: [] for level in SECOND_LEVEL_ORDER[bucket]} for bucket in BUCKETS
    }
    for block in blocks:
        level, _reason = assignments[block.title]
        by_bucket_level[block.bucket][level].append(block)

    lines: list[str] = []
    lines.extend(
        [
            "# 选择性必修一《当代国际政治与经济》主观题术语宝典",
            "学生厚版 · 237条逐条核实二级结构版（2026.5.28 · 共104术语/237条目）",
            "飞哥正志讲堂",
            "",
            "## 二级结构说明",
            "本版按“六大要素 → 术语/框架容器 → 核心答题点 → 真题条目”组织。六大要素用于找方向，二级容器用于判断材料关系，核心答题点保留原细则高信息量表述，题例用于训练迁移。",
            "",
            "“出现N次”仅表示本宝典收录样本中的命中次数，用于判断复习优先级，不等同于考试预测，也不代表该点在所有设问中都可直接套用。",
            "",
            "## 目录",
            "",
        ]
    )
    for bucket in BUCKETS:
        present = [level for level in SECOND_LEVEL_ORDER[bucket] if by_bucket_level[bucket][level]]
        lines.append(f"- {bucket}")
        for level in present:
            lines.append(f"  - {level}（{len(by_bucket_level[bucket][level])}个核心点）")
    lines.append("- 六大要素术语极简速记版")
    lines.append("")

    for bucket in BUCKETS:
        lines.append(f"## {bucket}")
        lines.append("")
        lines.append(f"**本桶怎么用：**{BUCKET_GUIDES[bucket]}")
        if bucket == "政治多极化":
            lines.append("")
            lines.append("**边界提醒：**本桶侧重国际关系、国际秩序、共同利益、合作共赢和全球治理方向；如果材料重心是中国主张、中国方案和中国行动，优先回到“中国”桶。")
        if bucket == "中国":
            lines.append("")
            lines.append("**边界提醒：**本桶按政策、智慧、责任、地位检索；同一个“人类命运共同体”点，若强调中国提出、推动、落实和贡献，归入中国；若强调国际秩序和全球治理方向，归入政治多极化。")
        if bucket == "时代背景":
            lines.append("")
            lines.append("**边界提醒：**当经济全球化只是说明中国倡议为何正当、必要时，可作为时代背景；当材料重点是开放、市场、贸易、产业链和规则治理时，应转入经济全球化。")
        lines.append("")

        for level in SECOND_LEVEL_ORDER[bucket]:
            level_blocks = by_bucket_level[bucket][level]
            if not level_blocks:
                continue
            lines.append(f"### {level}")
            lines.append("")
            for block in level_blocks:
                lines.extend(demote_core_block(block))
                lines.append("")

    lines.extend(demote_speed_lines(speed_lines))
    text = "\n".join(lines).rstrip() + "\n"
    return text


def set_font(run, name="PingFang SC", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_label_paragraph(doc: Document, line: str):
    m = re.match(r"^(【[^】]+】)(.*)$", line)
    p = doc.add_paragraph(style="BodyLabel")
    if not m:
        p.add_run(line)
        return
    label, rest = m.groups()
    r1 = p.add_run(label)
    set_font(r1, bold=True, color=(31, 78, 121))
    r2 = p.add_run(rest)
    set_font(r2)


def build_docx(md_text: str, out_path: Path):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.7)
    sec.bottom_margin = Cm(1.7)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "PingFang SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(3)

    for name, base, size, color, before, after in [
        ("DocTitle", "Title", 20, (31, 78, 121), 0, 8),
        ("DocSubtitle", "Subtitle", 11, (89, 89, 89), 0, 10),
        ("H1Custom", "Heading 1", 16, (31, 78, 121), 12, 6),
        ("H2Custom", "Heading 2", 13, (47, 84, 150), 9, 4),
        ("H3Custom", "Heading 3", 11.5, (84, 130, 53), 7, 3),
    ]:
        style = styles.add_style(name, 1)
        style.base_style = styles[base]
        style.font.name = "PingFang SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(*color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = name in {"H1Custom", "H2Custom", "H3Custom"}

    for name, left_cm, first_cm, after in [
        ("QuestionItem", 0.3, -0.3, 2),
        ("BodyLabel", 0, 0, 2),
        ("BodyIndented", 0.35, 0, 2),
    ]:
        style = styles.add_style(name, 1)
        style.base_style = normal
        style.font.name = "PingFang SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(10)
        style.paragraph_format.left_indent = Cm(left_cm)
        style.paragraph_format.first_line_indent = Cm(first_cm)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.06

    def add_text_para(text: str, style_name: str = "Normal", bold_prefix=False):
        p = doc.add_paragraph(style=style_name)
        if bold_prefix and "：" in text:
            prefix, rest = text.split("：", 1)
            r1 = p.add_run(prefix + "：")
            set_font(r1, bold=True, color=(31, 78, 121))
            r2 = p.add_run(rest)
            set_font(r2)
        else:
            r = p.add_run(text)
            set_font(r)
        return p

    first_title = True
    for raw in md_text.splitlines():
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="DocTitle")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[2:])
            set_font(r, size=20, bold=True, color=(31, 78, 121))
            first_title = False
        elif first_title:
            add_text_para(line, "DocSubtitle")
        elif line.startswith("## "):
            doc.add_paragraph(line[3:], style="H1Custom")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:], style="H2Custom")
        elif line.startswith("#### "):
            doc.add_paragraph(line[5:], style="H3Custom")
        elif line.startswith("**本桶怎么用：**"):
            add_text_para(line.replace("**", ""), "BodyIndented", bold_prefix=True)
        elif line.startswith("**边界提醒：**"):
            add_text_para(line.replace("**", ""), "BodyIndented", bold_prefix=True)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(line[2:])
            set_font(r, size=10)
        elif line.startswith("  - "):
            p = doc.add_paragraph(style="List Bullet 2")
            r = p.add_run(line[4:])
            set_font(r, size=10)
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="QuestionItem")
            r = p.add_run(line)
            set_font(r, bold=True, color=(68, 68, 68))
        elif line.startswith("【"):
            add_label_paragraph(doc, line)
        else:
            add_text_para(line, "Normal")

    core_props = doc.core_properties
    core_props.title = "选必一《当代国际政治与经济》主观题术语宝典 237条 二级结构版"
    core_props.subject = "北京高考政治 选择性必修一 主观题术语宝典"
    core_props.author = "飞哥正志讲堂"
    doc.save(out_path)


def count_entries(block: CoreBlock) -> int:
    return sum(1 for line in block.lines if re.match(r"^\d+\. ", line))


def main():
    source_text = SOURCE_MD.read_text(encoding="utf-8")
    lines = source_text.splitlines()
    blocks, speed_lines = parse_source(lines)

    assignments: dict[str, tuple[str, str]] = {}
    rows = []
    for block in blocks:
        level, reason = classify(block)
        assignments[block.title] = (level, reason)
        rows.append(
            {
                "bucket": block.bucket,
                "second_level": level,
                "core_point": block.title,
                "frequency": block.count if block.count is not None else "",
                "entry_count": count_entries(block),
                "reason": reason,
            }
        )

    md_text = build_markdown(blocks, speed_lines, assignments)
    OUT_MD.write_text(md_text, encoding="utf-8")

    with OUT_MAP.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=["bucket", "second_level", "core_point", "frequency", "entry_count", "reason"],
        )
        writer.writeheader()
        writer.writerows(rows)

    build_docx(md_text, OUT_DOCX)

    source_counts = {
        "source_core_count": len(blocks),
        "source_entry_count": sum(count_entries(b) for b in blocks),
        "source_when_count": source_text.count("【什么时候写】"),
        "source_question_count": source_text.count("【设问】"),
        "source_answer_count": source_text.count("【卷面句】"),
    }
    output_counts = {
        "output_core_count": md_text.count("#### 核心答题点："),
        "output_entry_count": sum(1 for line in md_text.splitlines() if re.match(r"^\d+\. ", line)),
        "output_when_count": md_text.count("【什么时候写】"),
        "output_question_count": md_text.count("【设问】"),
        "output_answer_count": md_text.count("【卷面句】"),
    }

    level_counter = Counter((row["bucket"], row["second_level"]) for row in rows)
    forbidden_terms = ["Codex", "Claude", "GPT", "prompt", "rubric_angle_id", ".txt"]
    forbidden_hits = {term: md_text.count(term) for term in forbidden_terms if md_text.count(term)}

    expected = {
        "core_count_ok": output_counts["output_core_count"] == 104,
        "entry_count_ok": output_counts["output_entry_count"] == 237,
        "when_count_ok": output_counts["output_when_count"] == source_counts["source_when_count"],
        "question_count_ok": output_counts["output_question_count"] == source_counts["source_question_count"],
        "answer_count_ok": output_counts["output_answer_count"] == source_counts["source_answer_count"],
        "forbidden_terms_ok": not forbidden_hits,
    }

    summary = {
        "source_md": str(SOURCE_MD),
        "out_md": str(OUT_MD),
        "out_docx": str(OUT_DOCX),
        "out_mapping": str(OUT_MAP),
        "source_counts": source_counts,
        "output_counts": output_counts,
        "checks": expected,
        "forbidden_hits": forbidden_hits,
        "second_level_counts": {
            f"{bucket} / {level}": level_counter[(bucket, level)]
            for bucket in BUCKETS
            for level in SECOND_LEVEL_ORDER[bucket]
            if level_counter[(bucket, level)]
        },
    }
    OUT_JSON.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")

    qa_lines = [
        "# SECOND_LEVEL_REORG_QA",
        "",
        "## 核验结论",
        "",
        f"- 核心答题点：{output_counts['output_core_count']}（目标 104）",
        f"- 条目：{output_counts['output_entry_count']}（目标 237）",
        f"- 【什么时候写】：{output_counts['output_when_count']}（与基稿一致）",
        f"- 【设问】：{output_counts['output_question_count']}（与基稿一致）",
        f"- 【卷面句】：{output_counts['output_answer_count']}（与基稿一致）",
        f"- 后台词扫描：{'PASS' if not forbidden_hits else 'FAIL ' + str(forbidden_hits)}",
        "",
        "## 二级结构核心点分布",
        "",
        "| 一级要素 | 二级容器 | 核心点数 |",
        "|---|---:|---:|",
    ]
    for bucket in BUCKETS:
        for level in SECOND_LEVEL_ORDER[bucket]:
            n = level_counter[(bucket, level)]
            if n:
                qa_lines.append(f"| {bucket} | {level} | {n} |")
    qa_lines.extend(
        [
            "",
            "## 输出文件",
            "",
            f"- Markdown：`{OUT_MD}`",
            f"- Word：`{OUT_DOCX}`",
            f"- 映射表：`{OUT_MAP}`",
        ]
    )
    OUT_SUMMARY.write_text("\n".join(qa_lines) + "\n", encoding="utf-8")

    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
