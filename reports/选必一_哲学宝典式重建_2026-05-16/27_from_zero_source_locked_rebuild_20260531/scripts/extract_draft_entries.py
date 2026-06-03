#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import csv
import argparse
import re
from pathlib import Path

from docx import Document


RUN = Path("/Users/wanglifei/GaokaoPolitics/beijing-politics-sync/reports/选必一_哲学宝典式重建_2026-05-16/27_from_zero_source_locked_rebuild_20260531")
DEFAULT_DRAFT_DOCX = Path("/Users/wanglifei/Library/Application Support/Claude/local-agent-mode-sessions/2c0a1a14-78ce-4ddc-8f72-9d4f74dbff6e/703c2d6a-ade6-4ee8-94a4-92b2abb47bf1/local_ea19f1a6-c672-4010-9aa6-235ad1631aa4/outputs/选必一_当代国际政治与经济_主观题术语宝典_哲学对齐版_20260531.docx")
FRAMEWORK_DOCX = Path("/Users/wanglifei/Desktop/给Claude上传_最新_选必一讲义二级框架_Claude排版.docx")
INDEX = RUN / "02_source_cards" / "RAW_CARD_INDEX.csv"
OUT = RUN / "03_structured"

MODULES = {"时代背景", "理论", "经济全球化", "政治多极化", "中国", "联合国"}
FIELD_RE = re.compile(r"^【([^】]+)】\s*(.*)$")
ENTRY_RE = re.compile(r"^(\d+)\.\s+(.+?Q\d+(?:\(\d+\))?)(.*)$")
LABEL_RE = re.compile(r"^(20\d{2})(.+?)(一模|二模|期中|期末)Q(\d+)(?:\((\d+)\))?")


def docx_text(path: Path):
    doc = Document(str(path))
    lines = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            vals = [c.text.strip() for c in row.cells if c.text.strip()]
            if vals:
                lines.append(" | ".join(vals))
    return lines


def make_card_id(label: str):
    label = label.replace("思政", "")
    m = LABEL_RE.search(label)
    if not m:
        return ""
    year, district, exam_type, q_no, q_sub = m.groups()
    sub = f"_{q_sub}" if q_sub else ""
    return f"{year}_{district}_{exam_type}_Q{q_no}{sub}"


def load_index():
    rows = list(csv.DictReader(INDEX.open(encoding="utf-8-sig")))
    return {r["card_id"].lstrip("\ufeff"): r for r in rows}


def parse_entries(lines):
    current_module = ""
    current_subsection = ""
    current_core = ""
    entries = []
    cur = None

    def finish():
        nonlocal cur
        if not cur:
            return
        cur["body"] = "\n".join(cur.pop("body_lines")).strip()
        entries.append(cur)
        cur = None

    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line in MODULES:
            finish()
            current_module = line
            current_subsection = ""
            current_core = ""
            continue
        if line.startswith("核心答题点："):
            finish()
            current_core = line.removeprefix("核心答题点：").strip()
            continue
        m_entry = ENTRY_RE.match(line)
        if m_entry:
            finish()
            ordinal, source_label, title_tail = m_entry.groups()
            cur = {
                "module": current_module,
                "subsection": current_subsection,
                "core_point": current_core,
                "ordinal": ordinal,
                "source_label": source_label.strip(),
                "title_tail": title_tail.strip(),
                "card_id": make_card_id(source_label.strip()),
                "body_lines": [],
            }
            continue
        m_field = FIELD_RE.match(line)
        if cur and m_field:
            key, value = m_field.groups()
            cur[key] = value.strip()
            cur["body_lines"].append(line)
            continue
        if cur:
            cur["body_lines"].append(line)
        elif current_module and not line.startswith("本桶怎么用") and not line.startswith("边界提醒") and not line.startswith("目录"):
            if not line.startswith("学生厚版") and not line.startswith("飞哥正志讲堂") and not line.startswith("二级结构说明"):
                current_subsection = line
    finish()
    return entries


def parse_args():
    parser = argparse.ArgumentParser()
    parser.add_argument("--draft-docx", type=Path, default=DEFAULT_DRAFT_DOCX)
    parser.add_argument("--prefix", default="DRAFT")
    return parser.parse_args()


def main():
    args = parse_args()
    prefix = args.prefix.upper()
    OUT.mkdir(parents=True, exist_ok=True)
    lines = docx_text(args.draft_docx)
    (RUN / "01_inputs" / f"{prefix}_CLAUDE_DRAFT_TEXT.txt").write_text("\n".join(lines), encoding="utf-8")
    framework_lines = docx_text(FRAMEWORK_DOCX)
    (RUN / "01_inputs" / f"{prefix}_FRAMEWORK_TEXT.txt").write_text("\n".join(framework_lines), encoding="utf-8")

    cards = load_index()
    entries = parse_entries(lines)
    for e in entries:
        card = cards.get(e["card_id"], {})
        e["evidence_status"] = card.get("status", "NO_CARD_MATCH")
        e["paper_found"] = card.get("paper_found", "")
        e["rubric_found"] = card.get("rubric_found", "")
        e["card_path"] = card.get("card_path", "")

    fieldnames = [
        "module", "subsection", "core_point", "ordinal", "source_label", "title_tail",
        "card_id", "evidence_status", "paper_found", "rubric_found", "card_path",
        "什么时候写", "设问", "得分层次", "卷面句", "同题组",
    ]
    with (OUT / f"{prefix}_ENTRY_AUDIT.csv").open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fieldnames, extrasaction="ignore")
        w.writeheader()
        w.writerows(entries)

    problem_rows = [e for e in entries if e["evidence_status"] != "RAW_CARD_READY"]
    with (OUT / f"{prefix}_ENTRY_PROBLEMS.csv").open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fieldnames, extrasaction="ignore")
        w.writeheader()
        w.writerows(problem_rows)

    print("draft_entries", len(entries))
    print("ready_entries", sum(1 for e in entries if e["evidence_status"] == "RAW_CARD_READY"))
    print("problem_entries", len(problem_rows))
    print("unique_card_ids", len({e["card_id"] for e in entries if e["card_id"]}))


if __name__ == "__main__":
    main()
