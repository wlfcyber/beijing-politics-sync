#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import csv
import hashlib
import subprocess
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from pypdf import PdfReader


RUN = Path("/Users/wanglifei/GaokaoPolitics/beijing-politics-sync/reports/选必一_哲学宝典式重建_2026-05-16/27_from_zero_source_locked_rebuild_20260531")
MATCH = RUN / "01_inputs" / "QUESTION_SOURCE_MATCH.csv"
TEXT_DIR = RUN / "01_inputs" / "source_texts"
MANIFEST = RUN / "01_inputs" / "EXTRACTED_SOURCE_MANIFEST.csv"
OCR_NEEDED = RUN / "01_inputs" / "OCR_NEEDED.csv"


def sha(path: Path) -> str:
    h = hashlib.sha1()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()[:12]


def split_paths(value):
    return [p.strip() for p in (value or "").split(" || ") if p.strip()]


def extract_pdf(path: Path):
    reader = PdfReader(str(path))
    pages = []
    for idx, page in enumerate(reader.pages, 1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            text = f"[PAGE_EXTRACT_ERROR {exc}]"
        pages.append(f"--- page {idx} ---\n{text}")
    return "\n\n".join(pages), len(reader.pages)


def extract_docx(path: Path):
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [" ".join(c.text.split()) for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts), 0


def extract_pptx(path: Path):
    texts = []
    ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    with zipfile.ZipFile(path) as z:
        slide_names = sorted(n for n in z.namelist() if n.startswith("ppt/slides/slide") and n.endswith(".xml"))
        for slide_no, name in enumerate(slide_names, 1):
            root = ET.fromstring(z.read(name))
            vals = [node.text for node in root.findall(".//a:t", ns) if node.text]
            texts.append(f"--- slide {slide_no} ---\n" + "\n".join(vals))
    return "\n\n".join(texts), len(texts)


def extract_doc_or_ppt(path: Path):
    proc = subprocess.run(["textutil", "-convert", "txt", "-stdout", str(path)], capture_output=True, text=True)
    if proc.returncode != 0:
        return f"[TEXTUTIL_ERROR]\n{proc.stderr}", 0
    return proc.stdout, 0


def extract(path: Path):
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".pptx":
        return extract_pptx(path)
    if suffix in {".doc", ".ppt"}:
        return extract_doc_or_ppt(path)
    return "", 0


def main():
    TEXT_DIR.mkdir(parents=True, exist_ok=True)
    rows = list(csv.DictReader(MATCH.open(encoding="utf-8-sig")))
    paths = []
    for row in rows:
        for field in ["matched_papers", "matched_rubrics", "matched_reviews", "matched_answers"]:
            paths.extend(split_paths(row.get(field, "")))
    unique = []
    seen = set()
    for p in paths:
        if p not in seen and Path(p).exists():
            seen.add(p)
            unique.append(Path(p))

    out_rows = []
    ocr_rows = []
    for path in unique:
        sid = sha(path)
        out = TEXT_DIR / f"{sid}_{path.stem}.txt"
        try:
            text, pages = extract(path)
            status = "OK"
        except Exception as exc:
            text, pages = f"[EXTRACT_ERROR]\n{exc}", 0
            status = "ERROR"
        out.write_text(f"SOURCE: {path}\n\n{text}", encoding="utf-8")
        char_count = len(text.strip())
        needs_ocr = path.suffix.lower() == ".pdf" and char_count < 500
        row = {
            "source_id": sid,
            "path": str(path),
            "suffix": path.suffix.lower(),
            "pages_or_slides": pages,
            "status": status,
            "char_count": char_count,
            "needs_ocr": "YES" if needs_ocr else "NO",
            "text_path": str(out),
        }
        out_rows.append(row)
        if needs_ocr:
            ocr_rows.append(row)

    with MANIFEST.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=list(out_rows[0].keys()))
        w.writeheader()
        w.writerows(out_rows)
    with OCR_NEEDED.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=list(out_rows[0].keys()))
        w.writeheader()
        w.writerows(ocr_rows)
    print("extracted", len(out_rows), "ocr_needed", len(ocr_rows))


if __name__ == "__main__":
    main()
