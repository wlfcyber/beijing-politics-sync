#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document


RUN = Path("/Users/wanglifei/GaokaoPolitics/beijing-politics-sync/reports/选必一_哲学宝典式重建_2026-05-16/27_from_zero_source_locked_rebuild_20260531")
DRAFT_DOCX = Path("/Users/wanglifei/Library/Application Support/Claude/local-agent-mode-sessions/2c0a1a14-78ce-4ddc-8f72-9d4f74dbff6e/703c2d6a-ade6-4ee8-94a4-92b2abb47bf1/local_ea19f1a6-c672-4010-9aa6-235ad1631aa4/outputs/选必一_当代国际政治与经济_主观题术语宝典_哲学对齐版_20260531.docx")
OUT_DIR = RUN / "04_student_draft"
OUT_DOCX = OUT_DIR / "选必一_当代国际政治与经济_主观题术语宝典_确认版_学生版_20260531.docx"
OUT_MD = OUT_DIR / "选必一_当代国际政治与经济_主观题术语宝典_确认版_学生版_20260531.md"
DESKTOP_COPY = Path("/Users/wanglifei/Desktop/选必一_当代国际政治与经济_主观题术语宝典_确认版_学生版_20260531.docx")
REPORT = RUN / "05_qa" / "FINAL_CLEANING_REPORT.md"

ENTRY_RE = re.compile(r"^(\d+)\.\s+(.+?Q\d+(?:\(\d+\))?)(.*)$")
LABEL_RE = re.compile(r"^(20\d{2})(.+?)(一模|二模|期中|期末)Q(\d+)(?:\((\d+)\))?")
MODULES = {"时代背景", "理论", "经济全球化", "政治多极化", "中国", "联合国"}
BANNED_CARD_IDS = {"2026_海淀_期中_Q22_1", "2026_西城_期末_Q20"}
PROMPT_FIXES = {
    "2025_西城_一模_Q18": "【设问】 从“两条鱼循环”看增长新空间。产自中国的鳗鱼“游”向世界，挪威三文鱼“飞”来中国，增长新空间在不断拓展。结合材料，运用《经济与社会》所学，谈谈中国经济增长新空间是如何拓展的。",
    "2026_顺义_一模_Q21": "【设问】 结合材料，综合运用所学，谈谈以中国式现代化全面推进中华民族伟大复兴，应如何运用“审势、乘势、驭势”的智慧。",
}


def make_card_id(label: str):
    label = label.replace("思政", "")
    m = LABEL_RE.search(label)
    if not m:
        return ""
    year, district, exam_type, q_no, q_sub = m.groups()
    sub = f"_{q_sub}" if q_sub else ""
    return f"{year}_{district}_{exam_type}_Q{q_no}{sub}"


def remove_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    parent.remove(p)


def set_paragraph_text(paragraph, text):
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    first = paragraph.runs[0]
    first.text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def paragraph_texts(doc):
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    (RUN / "05_qa").mkdir(parents=True, exist_ok=True)
    doc = Document(str(DRAFT_DOCX))
    current_card_id = ""
    skip = False
    removed_entries = []
    removed_score_lines = 0
    prompt_fixes = 0

    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            if skip:
                remove_paragraph(paragraph)
            continue

        if "235 条" in text:
            text = text.replace("235 条", "229 条")
            text = text.replace("235 条目", "229 条目")
            if "学生最终版" in text:
                text = text.replace("学生最终版", "学生确认版")
            set_paragraph_text(paragraph, text)

        entry = ENTRY_RE.match(text)
        if entry:
            current_card_id = make_card_id(entry.group(2).strip())
            skip = current_card_id in BANNED_CARD_IDS
            if skip:
                removed_entries.append(text)
                remove_paragraph(paragraph)
            continue

        if text in MODULES or text.startswith("核心答题点："):
            skip = False
            current_card_id = ""

        if skip:
            remove_paragraph(paragraph)
            continue

        if text.startswith("【得分层次】"):
            removed_score_lines += 1
            remove_paragraph(paragraph)
            continue

        if current_card_id in PROMPT_FIXES and text.startswith("【设问】"):
            set_paragraph_text(paragraph, PROMPT_FIXES[current_card_id])
            prompt_fixes += 1

    doc.save(str(OUT_DOCX))
    shutil.copy2(OUT_DOCX, DESKTOP_COPY)

    cleaned = Document(str(OUT_DOCX))
    OUT_MD.write_text("\n".join(paragraph_texts(cleaned)) + "\n", encoding="utf-8")
    REPORT.write_text(
        "\n".join([
            "# FINAL_CLEANING_REPORT",
            "",
            "- 学生版规则：不显示细则依据、证据层级、源文件路径；后台证据保留在 `02_source_cards/` 和 `05_qa/`。",
            f"- 删除未能按“原卷+评分来源”确认的条目数：{len(removed_entries)}。",
            f"- 删除学生版 `【得分层次】` 行数：{removed_score_lines}。",
            f"- 改回原卷设问行数：{prompt_fixes}。",
            "",
            "## Deleted Entries",
            *[f"- {x}" for x in removed_entries],
            "",
            "## Prompt Fixes",
            "- 2025西城一模Q18：改回原卷“两条鱼循环”设问。",
            "- 2026顺义一模Q21：改回原卷“以中国式现代化全面推进中华民族伟大复兴，应如何运用‘审势、乘势、驭势’的智慧”。",
            "",
            f"## Output",
            f"- {OUT_DOCX}",
            f"- {DESKTOP_COPY}",
        ]) + "\n",
        encoding="utf-8",
    )
    print("output", OUT_DOCX)
    print("desktop", DESKTOP_COPY)
    print("removed_entries", len(removed_entries))
    print("removed_score_lines", removed_score_lines)
    print("prompt_fixes", prompt_fixes)


if __name__ == "__main__":
    main()
