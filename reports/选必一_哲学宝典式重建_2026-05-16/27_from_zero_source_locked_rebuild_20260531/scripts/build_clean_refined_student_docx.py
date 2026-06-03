#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import re
import shutil
from pathlib import Path

from docx import Document


RUN = Path("/Users/wanglifei/GaokaoPolitics/beijing-politics-sync/reports/选必一_哲学宝典式重建_2026-05-16/27_from_zero_source_locked_rebuild_20260531")
DRAFT_DOCX = Path("/Users/wanglifei/Library/Application Support/Claude/local-agent-mode-sessions/2c0a1a14-78ce-4ddc-8f72-9d4f74dbff6e/703c2d6a-ade6-4ee8-94a4-92b2abb47bf1/local_ea19f1a6-c672-4010-9aa6-235ad1631aa4/outputs/选必一_当代国际政治与经济_主观题术语宝典_哲学精炼版_20260531.docx")
OUT_DIR = RUN / "04_student_draft"
OUT_DOCX = OUT_DIR / "选必一_当代国际政治与经济_主观题术语宝典_哲学精炼审核版_学生版_20260531.docx"
OUT_MD = OUT_DIR / "选必一_当代国际政治与经济_主观题术语宝典_哲学精炼审核版_学生版_20260531.md"
DESKTOP_COPY = Path("/Users/wanglifei/Desktop/选必一_当代国际政治与经济_主观题术语宝典_哲学精炼审核版_学生版_20260531.docx")
REPORT = RUN / "05_qa" / "REFINED_CLEANING_REPORT.md"

ENTRY_RE = re.compile(r"^(\d+)\.\s+(.+?Q\d+(?:\(\d+\))?)(.*)$")
LABEL_RE = re.compile(r"^(20\d{2})(.+?)(一模|二模|期中|期末)Q(\d+)(?:\((\d+)\))?")
MODULES = {"时代背景", "理论", "经济全球化", "政治多极化", "中国", "联合国"}
BANNED_CARD_IDS = {"2026_海淀_期中_Q22_1"}
BANNED_PREFIXES = (
    "【得分层次】",
    "【判别标注】",
    "【术语家族】",
    "【使用标注】",
    "【双要素判别规则】",
)
PROMPT_FIXES = {
    "2025_西城_一模_Q18": "【设问】 从“两条鱼循环”看增长新空间。产自中国的鳗鱼“游”向世界，挪威三文鱼“飞”来中国，增长新空间在不断拓展。结合材料，运用《经济与社会》所学，谈谈中国经济增长新空间是如何拓展的。",
    "2026_西城_期末_Q20": "【设问】 结合材料，运用《国际政治与经济》知识，阐释参与全球气候治理的中国实践。",
    "2024_海淀_二模_Q18_1": "【设问】 任选一个议题，参考资料包中的内容，运用《当代国际政治与经济》知识，围绕所选的议题写一篇时政述评。（8分）要求：观点明确；知识运用准确；论述合乎逻辑，条理清晰。",
    "2025_海淀_二模_Q21": "【设问】 结合材料，运用《当代国际政治与经济》知识，阐释“中国需要联合国，联合国也需要中国”。",
    "2025_昌平_二模_Q21": "【设问】 结合材料，说明上述举措对我国建设更高水平的开放型经济具有什么重要意义。",
    "2026_延庆_一模_Q19_2": "【设问】 结合材料二，运用《当代国际政治与经济》知识，说明中国“推动重塑全球能源治理格局”的理论逻辑和价值意蕴。",
    "2026_石景山_二模_Q18": "【设问】 结合材料，运用《当代国际政治与经济》知识，分析中国与东盟经贸合作能够提质升级的原因。",
    "2024_朝阳_二模_Q20": "【设问】 结合研讨背景，运用《当代国际政治与经济》知识，完成下表。",
}
TEXT_REPLACEMENTS = {
    "学生厚版 · 235 条逐条核实讲义二级框架版 · Claude 排版优化（2026.5.28 · 共 104 术语 / 235 条目 · 2026.5.31 学生审核版·回原卷细则逐条核校）": "终极审核版 · 461 条逐条回源核准 · 六大要素二级框架版（2026.5.31）",
    "学生厚版 · 235 条逐条核实讲义二级框架版 · Claude 排版优化（2026.5.28 · 共 104 术语 / 235 条目 · 2026.5.31 学生审核版·逐条核校）": "终极审核版 · 461 条逐条回源核准 · 六大要素二级框架版（2026.5.31）",
    "学生最终版": "学生审核版",
    "回原卷细则逐条核校": "逐条核校",
    "核心答题点保留原细则高信息量表述": "核心答题点保留高信息量表述",
    "（细则可替代为「和平与发展」）": "，也可从和平与发展的时代主题角度展开",
    "（细则用语：人才与国际竞争力的关系）": "（人才与国际竞争力）",
    "细则用语：": "",
    "细则明确": "题目评价重点提示",
    "细则集中呈现": "题目评价重点集中在",
    "回到细则关键表述": "紧扣关键表述",
    "本题是《经济与社会》＋《当代国际政治与经济》跨模块题（细则：必修二4分＋选必一2分、企业4分＋政府2分）；选必一2分对应细则“贸易摩擦/规则环境”角度。": "本题是《经济与社会》＋《当代国际政治与经济》跨模块题；选必一角度集中在“贸易摩擦/规则环境”。",
    "细则强调": "题目评价重点强调",
    "细则把": "题目评价重点把",
    "细则": "评分重点",
    "设问要求": "题目指向",
    "题目要求": "题目指向",
    "要落到": "应指向",
    "落到": "指向",
    "先写": "先说明",
    "要写": "需要说明",
    "答题时": "作答时",
    "后台映射表": "术语映射表",
    "后台": "内部",
    "和平与发展仍是时代主题，中国倡议顺应这一潮流，致力于亚太和平合作与共同发展。": "中国倡议顺应和平、发展、合作、共赢的潮流，致力于亚太和平合作与共同发展。",
    "当前国际竞争实质是综合国力较量，教育培养人才、夯实科技实力，是赢得竞争、支撑强国复兴之基。": "教育具有战略支撑地位，能够促进人才培养、提升劳动者技术水平和创新能力，增强我国综合国力和国际竞争力，支撑强国建设和民族复兴。",
    "中国与发展中国家共享发展新机遇，科技小院培育人才、惠及民生、增进福祉，实现了双方互利共赢。": "科技小院面向全球南方开展农业合作，培养人才、增强内生发展动力，推动共享发展成果，实现互利共赢。",
    "材料中": "题干中",
    "采分点": "作答方向",
    "得分关键词": "关键表述",
    "采分关键词": "关键表述",
    "细则得分点": "评分方向",
}


def make_card_id(label: str):
    label = label.replace("思政", "")
    m = LABEL_RE.search(label)
    if not m:
        return ""
    year, district, exam_type, q_no, q_sub = m.groups()
    sub = f"_{q_sub}" if q_sub else ""
    return f"{year}_{district}_{exam_type}_Q{q_no}{sub}"


def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)


def set_paragraph_text(paragraph, text):
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def paragraph_texts(doc):
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def strip_feige_intro(text: str):
    if "【飞哥说】" not in text:
        return text
    prefix, suffix = text.split("【飞哥说】", 1)
    marker = "题面"
    if marker in suffix:
        return prefix + marker + suffix.split(marker, 1)[1]
    return prefix.strip()


def cleanup_text(text: str):
    text = strip_feige_intro(text)
    for old, new in TEXT_REPLACEMENTS.items():
        text = text.replace(old, new)
    return text


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    REPORT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(DRAFT_DOCX))
    current_card_id = ""
    skip = False
    removed_entries = []
    removed_prefix_counts = {p: 0 for p in BANNED_PREFIXES}
    prompt_fixes = {}
    text_cleanups = 0

    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            if skip:
                remove_paragraph(paragraph)
            continue

        if "学生最终版" in text:
            text = text.replace("学生最终版", "学生审核版")
            set_paragraph_text(paragraph, text)

        entry = ENTRY_RE.match(text)
        if entry:
            current_card_id = make_card_id(entry.group(2).strip())
            skip = current_card_id in BANNED_CARD_IDS
            if skip:
                removed_entries.append(text)
                remove_paragraph(paragraph)
            continue

        if text in MODULES or text.startswith("核心答题点："):
            skip = False
            current_card_id = ""

        if skip:
            remove_paragraph(paragraph)
            continue

        if text.startswith(BANNED_PREFIXES):
            for prefix in BANNED_PREFIXES:
                if text.startswith(prefix):
                    removed_prefix_counts[prefix] += 1
                    break
            remove_paragraph(paragraph)
            continue

        if current_card_id in PROMPT_FIXES and text.startswith("【设问】"):
            set_paragraph_text(paragraph, PROMPT_FIXES[current_card_id])
            prompt_fixes[current_card_id] = prompt_fixes.get(current_card_id, 0) + 1
            continue

        cleaned = cleanup_text(text)
        if cleaned != text:
            set_paragraph_text(paragraph, cleaned)
            text_cleanups += 1

    doc.save(str(OUT_DOCX))
    shutil.copy2(OUT_DOCX, DESKTOP_COPY)

    cleaned_doc = Document(str(OUT_DOCX))
    lines = paragraph_texts(cleaned_doc)
    OUT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")
    report_lines = [
        "# REFINED_CLEANING_REPORT",
        "",
        "- 输入：Claude 哲学精炼版 20260531。",
        "- 学生版规则：删除后台证据展示，不显示细则依据、细则位置、证据层级、源文件路径和得分层次。",
        f"- 删除无法闭合正式细则的条目数：{len(removed_entries)}。",
        f"- 改回原卷设问的题号数：{len(prompt_fixes)}。",
        f"- 清理后台表述段落数：{text_cleanups}。",
        "",
        "## Removed Entries",
        *[f"- {x}" for x in removed_entries],
        "",
        "## Removed Prefixes",
        *[f"- {k}: {v}" for k, v in removed_prefix_counts.items()],
        "",
        "## Prompt Fixes",
        *[f"- {k}: {v}处" for k, v in prompt_fixes.items()],
        "",
        "## Output",
        f"- {OUT_DOCX}",
        f"- {DESKTOP_COPY}",
    ]
    REPORT.write_text("\n".join(report_lines) + "\n", encoding="utf-8")
    print("output", OUT_DOCX)
    print("desktop", DESKTOP_COPY)
    print("removed_entries", len(removed_entries))
    print("removed_prefix_counts", removed_prefix_counts)
    print("prompt_fixes", prompt_fixes)
    print("text_cleanups", text_cleanups)


if __name__ == "__main__":
    main()
