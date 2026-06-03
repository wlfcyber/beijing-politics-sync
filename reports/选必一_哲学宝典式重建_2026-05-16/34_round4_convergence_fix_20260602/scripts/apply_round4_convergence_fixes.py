from __future__ import annotations

import argparse
import json
import re
from copy import deepcopy
from pathlib import Path

from docx import Document


ENTRY_RE = re.compile(r"^\d+\.\s+")
FIELD_RE = re.compile(r"^【([^】]+)】\s*(.*)$")


def style_name(paragraph) -> str:
    return paragraph.style.name if paragraph.style else ""


def set_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def insert_clone_before(marker, paragraph) -> None:
    marker._element.addprevious(deepcopy(paragraph._element))


def insert_para_before(marker, text: str = "", style: str = "Normal"):
    return marker.insert_paragraph_before(text, style=style)


def parse_entries(doc: Document) -> list[dict]:
    entries: list[dict] = []
    current_h1 = current_h2 = current_h3 = ""
    starts: list[tuple[int, str, str, str]] = []

    for idx, par in enumerate(doc.paragraphs):
        text = par.text.strip()
        st = style_name(par)
        if st == "Heading 1":
            current_h1 = text
            current_h2 = ""
            current_h3 = ""
        elif st == "Heading 2":
            current_h2 = text
            current_h3 = ""
        elif st == "Heading 3":
            current_h3 = text
        elif ENTRY_RE.match(text):
            starts.append((idx, current_h1, current_h2, current_h3))

    for pos, (start, h1, h2, h3) in enumerate(starts):
        end = starts[pos + 1][0] if pos + 1 < len(starts) else len(doc.paragraphs)
        for idx in range(start + 1, end):
            if style_name(doc.paragraphs[idx]) in {"Heading 1", "Heading 2", "Heading 3"}:
                end = idx
                break
        fields = {}
        field_paras = {}
        for idx in range(start + 1, end):
            m = FIELD_RE.match(doc.paragraphs[idx].text.strip())
            if m:
                fields[m.group(1)] = m.group(2)
                field_paras[m.group(1)] = doc.paragraphs[idx]
        entries.append(
            {
                "start": start,
                "end": end,
                "title": doc.paragraphs[start].text.strip(),
                "h1": h1,
                "h2": h2,
                "h3": h3,
                "fields": fields,
                "field_paras": field_paras,
                "paragraphs": [doc.paragraphs[i] for i in range(start, end)],
            }
        )
    return entries


def update_field(entry: dict, label: str, value: str, changes: list[str]) -> None:
    paragraph = entry["field_paras"].get(label)
    if paragraph is None:
        return
    set_text(paragraph, f"【{label}】 {value}")
    entry["fields"][label] = value
    changes.append(f"updated {label}: {entry['title']} / {entry['h3']}")


def find_section_end_marker(doc: Document, h3_text: str):
    in_section = False
    for par in doc.paragraphs:
        st = style_name(par)
        text = par.text.strip()
        if st == "Heading 3" and text == h3_text:
            in_section = True
            continue
        if in_section and st in {"Heading 1", "Heading 2", "Heading 3"}:
            return par
    if in_section:
        return doc.paragraphs[-1]
    raise RuntimeError(f"找不到目标核心点：{h3_text}")


def clone_entry_to_h3(doc: Document, entry: dict, target_h3: str, changes: list[str]) -> None:
    marker = find_section_end_marker(doc, target_h3)
    insert_para_before(marker, "", "Normal")
    for paragraph in entry["paragraphs"]:
        insert_clone_before(marker, paragraph)
    changes.append(f"moved: {entry['title']} -> {target_h3}")


def delete_refs_for_entry(entry: dict, changes: list[str], reason: str) -> list:
    changes.append(f"deleted: {entry['title']} from {entry['h3']} ({reason})")
    return list(entry["paragraphs"])


def matches(entry: dict, title_part: str, h3_part: str | None = None) -> bool:
    if title_part not in entry["title"]:
        return False
    if h3_part is not None and h3_part not in entry["h3"]:
        return False
    return True


def repair_chaoyang_yimo_q20(entry: dict, changes: list[str]) -> None:
    if "2026朝阳一模Q20" not in entry["title"]:
        return

    h3 = entry["h3"]
    if "当前国际竞争的实质" in h3:
        update_field(
            entry,
            "材料触发点",
            "设问问“为什么能”，要回答中国向世界释放稳定性和正能量的能力来源。材料镜头一列出人形机器人、AI眼镜、开源大模型下载量全球第一、创新药76款，说明中国创新势能持续迸发。",
            changes,
        )
        update_field(
            entry,
            "为什么能想到",
            "人形机器人、AI眼镜、开源大模型和创新药集中呈现中国科技实力与创新能力增强。问“中国为什么能”为全球发展注入稳定性和正能量，首先要解释能力来源；这一层应从国际竞争实质和创新驱动战略切入，再说明中国科技实力增强如何为全球产业升级、民生改善和产业链供应链稳定提供支撑。",
            changes,
        )
        update_field(
            entry,
            "答案落点",
            "当前国际竞争的实质是以经济和科技实力为基础的综合国力较量。中国坚持创新驱动战略，科技实力日益增强，人形机器人、AI眼镜、开源大模型和创新药等创新成果为全球产业升级、民生改善和产业链供应链稳定提供支撑。",
            changes,
        )
        return

    if "维护全球产业链供应链稳定畅通" in h3:
        update_field(
            entry,
            "材料触发点",
            "中国AI眼镜、开源大模型、创新药等创新成果说明中国科技实力增强，答案要把“创新驱动—全球创新策源地—产业链供应链稳定”写成能力链条。",
            changes,
        )
        update_field(
            entry,
            "为什么能想到",
            "材料列出AI眼镜、开源大模型、创新药等具体创新成果，说明中国不是只表达合作意愿，而是具备稳定全球产业链供应链的科技和产业能力。问“中国为什么能”注入稳定性，要把这些创新成果转化为全球创新策源和产业链供应链韧性的支撑。",
            changes,
        )
        return

    if "共享发展新机遇" in h3 or "改善发展中国家民生" in h3:
        update_field(
            entry,
            "材料触发点",
            "鲁班工坊、菌草种植、减贫和生态治理经验直接服务发展中国家生产生活改善，说明中国把自身发展经验转化为发展中国家的可持续发展能力。",
            changes,
        )
        update_field(
            entry,
            "为什么能想到",
            "本题的“正能量”不只来自中国科技创新，也来自中国对发展中国家的赋能。鲁班工坊、菌草技术、减贫和生态治理经验都指向民生改善和内生发展能力提升，应归入与发展中国家共享发展新机遇、惠民生增福祉、实现互利共赢这一层。",
            changes,
        )
        return

    if "正确义利观" in h3:
        update_field(
            entry,
            "材料触发点",
            "鲁班工坊、菌草技术、减贫和生态治理经验表明中国在合作中兼顾自身发展与发展中国家民生改善，不是单方获利，而是义利兼顾、互利共赢。",
            changes,
        )
        update_field(
            entry,
            "为什么能想到",
            "设问问中国为什么能为全球发展注入稳定性和正能量，关键不只是写中国做了什么，还要解释中国怎样处理自身发展、他国关切和共同利益的关系。材料中的经验分享和技术合作体现义利相济、以义为先，应从正确义利观切入。",
            changes,
        )
        return

    if "大国责任" in h3:
        update_field(
            entry,
            "材料触发点",
            "中国通过创新成果、开放政策和发展合作把自身发展同世界发展连接起来，特别是以经验分享和技术合作帮助发展中国家改善民生。",
            changes,
        )
        update_field(
            entry,
            "为什么能想到",
            "本题要求说明中国能够为全球发展注入稳定性和正能量，材料中的开放合作、经验分享和对发展中国家的赋能体现中国不是旁观者，而是全球发展的贡献者。这里应从大国责任担当说明中国行动的价值。",
            changes,
        )
        return

    if "贡献中国智慧" in h3:
        update_field(
            entry,
            "材料触发点",
            "人形机器人、AI眼镜、开源大模型、创新药以及鲁班工坊、菌草技术等材料，说明中国既有创新能力，也能把发展经验和技术方案转化为全球发展资源。",
            changes,
        )
        update_field(
            entry,
            "为什么能想到",
            "题目问中国为什么能注入稳定性和正能量，既要写能力来源，也要写方案供给。中国创新成果和发展合作实践共同说明，中国以可复制、可共享的技术和发展方案回应全球发展问题。",
            changes,
        )
        update_field(
            entry,
            "答案落点",
            "中国坚持创新驱动战略，推动开源大模型、创新药等成果发展，并通过鲁班工坊、菌草技术等合作提供发展方案，为全球发展贡献中国智慧、中国方案、中国力量。",
            changes,
        )
        return

    if "促进全球资源优化配置" in h3:
        update_field(
            entry,
            "材料触发点",
            "中国消费市场、制造业规模、货物贸易、自贸协定和免签政策共同说明中国市场与全球资源要素双向联动。",
            changes,
        )
        update_field(
            entry,
            "为什么能想到",
            "设问问中国为什么能注入稳定性和正能量，材料中的超大规模市场、制造业规模、货物贸易、自贸协定和免签政策说明中国能够吸引和配置全球商品、服务、人才等要素。这里应从全球资源优化配置和国际贸易发展说明中国开放对世界的支撑。",
            changes,
        )
        return

    if "开放、包容、普惠、平衡、共赢" in h3:
        update_field(
            entry,
            "材料触发点",
            "中国扩大自贸朋友圈、实施单方面免签，并通过鲁班工坊、菌草技术、减贫和生态治理经验分享让更多国家共享发展机遇。",
            changes,
        )
        update_field(
            entry,
            "为什么能想到",
            "本题材料把中国开放市场和帮助发展中国家提升能力放在一起，说明中国推动的全球化不是少数国家获利，而是让市场、规则、技术和发展经验更广泛流动。这里应落到经济全球化更加开放、包容、普惠、平衡、共赢的方向。",
            changes,
        )
        return

    if "和平与发展仍是时代主题" in h3:
        update_field(
            entry,
            "材料触发点",
            "题目问中国为什么能为全球发展注入稳定性和正能量，材料从创新、开放和发展合作三个镜头说明中国行动顺应和平与发展的时代需要。",
            changes,
        )
        update_field(
            entry,
            "为什么能想到",
            "全球发展面临不稳定不确定因素时，题目把“中国能注入稳定性和正能量”作为核心追问。中国的创新成果、开放政策和发展合作共同服务于世界发展，应先用和平与发展仍是时代主题统摄这一价值方向。",
            changes,
        )


def refine_formulaic_why(entries: list[dict], changes: list[str]) -> None:
    changed = 0
    for entry in entries:
        why = entry["fields"].get("为什么能想到", "")
        if not why:
            continue
        new = why
        new = re.sub(
            r"这说明本层应抓“([^”]+)”，从该角度解释材料事实。",
            r"这里不能只复述题面事实，而要把事实归入“\1”这一作答角度。",
            new,
        )
        new = re.sub(
            r"本题这一层对应“([^”]+)”，需要用材料事实说明其成立。",
            r"本题这一层要落在“\1”，并把材料事实写成因果链条。",
            new,
        )
        new = re.sub(
            r"(同题组存在多条线索；这里取|同题组线索较多；本处取|同一材料可服务多个答点；本处只取)“([^”]+)”这一层，说明材料怎样支撑该术语。",
            r"\1“\2”这一层，把材料事实和该核心连成因果链。",
            new,
        )
        new = re.sub(
            r"这些事实共同指向“([^”]+)”，用它说明上述事实的答题依据。",
            r"这些事实的共同功能是支撑“\1”，作答时要把材料事实和该核心连成因果链。",
            new,
        )
        new = re.sub(
            r"这些事实共同指向“([^”]+)”，作答时应让材料回扣这一核心。",
            r"这些事实的共同功能是支撑“\1”，作答时要把材料事实和该核心连成因果链。",
            new,
        )
        new = re.sub(
            r"抓住“([^”]+)”，才能把这些事实解释为可采分的判断。",
            r"这些事实只有归入“\1”，才不会停留在材料罗列。",
            new,
        )
        new = re.sub(
            r"这里对应“([^”]+)”，它把题面事实和答案落点连接起来。",
            r"这里应落在“\1”，并把题面事实写成清楚的因果链。",
            new,
        )
        new = new.replace(
            "题干要求说明国家对外行动的根本依据；因此要从维护国家利益这一出发点和落脚点解释政策选择。",
            "这里要把具体做法追溯到国家利益这一根本立足点，说明主权、安全、发展利益如何决定国家对外选择。",
        )
        new = new.replace("说明材料怎样支撑该术语", "把材料事实和该核心连成因果链")
        new = new.replace("支撑该术语", "支撑该核心")
        new = new.replace("而，这里", "这里")
        new = new.replace("而，这说明", "这说明")
        new = re.sub(r"。{2,}", "。", new).strip()
        if new != why:
            update_field(entry, "为什么能想到", new, changes)
            changed += 1
    changes.append(f"refined formulaic why fields: {changed}")


def renumber_entries_by_h3(doc: Document) -> int:
    current_h3 = ""
    counter = 0
    changed = 0
    for par in doc.paragraphs:
        st = style_name(par)
        text = par.text.strip()
        if st == "Heading 3":
            current_h3 = text
            counter = 0
            continue
        if current_h3 and ENTRY_RE.match(text):
            counter += 1
            new_text = ENTRY_RE.sub(f"{counter}. ", text, count=1)
            if new_text != text:
                set_text(par, new_text)
                changed += 1
        elif current_h3 and not ENTRY_RE.match(text) and re.match(r"^\d{4}", text):
            counter += 1
            set_text(par, f"{counter}. {text}")
            changed += 1
    return changed


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--summary", type=Path, required=True)
    args = parser.parse_args()

    doc = Document(args.input)
    entries = parse_entries(doc)
    changes: list[str] = []
    delete_refs = []

    for entry in entries:
        repair_chaoyang_yimo_q20(entry, changes)

        if matches(entry, "2026通州期末Q21", "贡献中国智慧、中国方案、中国力量"):
            update_field(
                entry,
                "为什么能想到",
                "本题不是要求评价中国向外输出什么，而是说明“十四五”时期中国式现代化建设为什么能取得进展。题面列出中国—中亚峰会、“一带一路”国际合作高峰论坛、中非合作论坛、进博会、服贸会、消博会等开放平台，说明高水平对外开放是推进中国式现代化的实践路径；这里的“中国智慧、中国方案、中国力量”应落回中国自身现代化建设。",
                changes,
            )
            update_field(
                entry,
                "答案落点",
                "我国依托中国—中亚峰会、“一带一路”国际合作高峰论坛、中非合作论坛、进博会、服贸会、消博会等开放平台推进高水平对外开放，为推进中国式现代化贡献中国智慧、中国方案、中国力量。",
                changes,
            )

        if matches(entry, "2026通州期末Q21", "坚持独立自主、自力更生"):
            update_field(
                entry,
                "为什么能想到",
                "“十四五”圆满收官题既看中国式现代化建设成就，也看这些成就背后的自主能力。作答时不能把开放平台简单写成对外输出，而要把科技自立自强、发展主动权和高水平开放区分开；本层只回答中国自身发展主动权的来源。",
                changes,
            )

        if matches(entry, "2025东城一模Q20", "中国扩大高水平对外开放与制度型开放"):
            delete_refs.extend(delete_refs_for_entry(entry, changes, "东城一模Q20不支持制度型开放"))

        if matches(entry, "2025东城一模Q20", "推动全球治理体系变革完善，建立更加公正合理的国际政治经济新秩序"):
            update_field(
                entry,
                "为什么能想到",
                "此次出访既有APEC、G20等多边活动，也有回应全球南方求平等、促变革的立场表达。评价元首外交的“精彩”之处，不能写成泛泛外交依据，而要说明中国如何代表发展中国家诉求，推动全球治理体系朝更加公正合理方向变革。",
                changes,
            )

        if matches(entry, "2025东城一模Q20", "与发展中国家共享发展新机遇，惠民生增福祉，实现互利共赢"):
            update_field(
                entry,
                "为什么能想到",
                "中国同秘鲁签署自贸协定升级议定书和共建“一带一路”合作规划，并在11天内达成60多项合作文件，呈现的是同发展中国家深化合作、共享发展机遇的关系。这里应把材料写成互利共赢和惠民生增福祉，而不是转向泛泛外交依据。",
                changes,
            )

        if matches(entry, "2024西城一模Q19(6)", "维护国家利益是主权国家对外活动的出发点和落脚点"):
            delete_refs.extend(delete_refs_for_entry(entry, changes, "本问限定经济角度，细则无维护国家利益"))

        if matches(entry, "2026丰台一模Q19", "展现大国责任担当"):
            delete_refs.extend(delete_refs_for_entry(entry, changes, "与负责任大国条目重复"))

        if matches(entry, "2026海淀二模Q20(2)", "全球发展倡议：以发展促繁荣"):
            delete_refs.extend(delete_refs_for_entry(entry, changes, "海淀二模Q20(2)细则仅列四角度，无该独立角度"))

        if matches(entry, "2026海淀二模Q20(2)", "从中国主张到国际共识，从合作理念到共同行动"):
            delete_refs.extend(delete_refs_for_entry(entry, changes, "海淀二模Q20(2)题面话不作为独立细则角度"))

        if matches(entry, "2026朝阳二模Q20(2)", "和平与发展仍是时代主题"):
            update_field(
                entry,
                "材料触发点",
                "材料说人类命运共同体理念在民族国家之上强调全人类立场，推动各国在兼顾全人类共同利益的前提下维护自身主权和发展。",
                changes,
            )
            update_field(
                entry,
                "为什么能想到",
                "设问问人类命运共同体理念为什么能维护世界和平与发展，首要逻辑是它超越民族国家本位，维护全人类共同利益，并为协调国家利益与全球共同利益提供价值指引；这一层应归入共同利益，而不是停留在“和平与发展”题眼本身。",
                changes,
            )
            update_field(
                entry,
                "答案落点",
                "人类命运共同体理念超越民族国家本位，在尊重国家主权的前提下强调全人类立场，维护全人类共同利益，为协调国家利益与全球共同利益提供价值指引。",
                changes,
            )
            clone_entry_to_h3(doc, entry, "核心答题点：国家间共同利益是国家合作的基础", changes)
            delete_refs.extend(delete_refs_for_entry(entry, changes, "移入理论/合作共同利益"))

        if matches(entry, "2026朝阳二模Q20(2)", "相互尊重、公平正义、合作共赢的新型国际关系"):
            update_field(
                entry,
                "材料触发点",
                "材料把人类命运共同体理念与集团对抗、零和博弈相对照，说明这一理念倡导合作共赢，推动各国在发展自身利益时兼顾他国合理关切。",
                changes,
            )
            update_field(
                entry,
                "为什么能想到",
                "本题第二层不是概括Q20(1)的“差异”，而是解释理念为什么能减少冲突、维护和平。材料中的“走出集团对抗、零和博弈”和“兼顾他国合理关切”共同指向合作共赢这一作答角度。",
                changes,
            )
            update_field(
                entry,
                "答案落点",
                "人类命运共同体理念摒弃零和博弈、倡导合作共赢，推动各国在发展自身利益时兼顾他国合理关切，从而减少冲突，为世界和平奠定理念基础。",
                changes,
            )
            clone_entry_to_h3(doc, entry, "核心答题点：在平等互利基础上开展合作，实现互利共赢", changes)
            delete_refs.extend(delete_refs_for_entry(entry, changes, "移入合作共赢核心"))

        if matches(entry, "2026朝阳二模Q20(2)", "贡献中国智慧、中国方案、中国力量"):
            update_field(
                entry,
                "为什么能想到",
                "材料明确说人类命运共同体理念是对“世界怎么了、我们怎么办、人类向何处去”的中国回答。设问要求解释其维护和平与发展的原因，应把这一理念作为中国提供的全球问题解决方案来写。",
                changes,
            )
            update_field(
                entry,
                "答案落点",
                "人类命运共同体理念从全人类福祉出发，为解决全球性问题提供中国方案，推动构建更加公正合理的国际治理体系，促进各国在共同发展中实现持久和平与繁荣。",
                changes,
            )

    refine_formulaic_why(entries, changes)

    seen = set()
    for paragraph in delete_refs:
        ident = id(paragraph._element)
        if ident in seen:
            continue
        seen.add(ident)
        delete_paragraph(paragraph)

    renumbered = renumber_entries_by_h3(doc)
    changes.append(f"renumbered entry headings by core point: {renumbered}")

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    args.summary.parent.mkdir(parents=True, exist_ok=True)
    args.summary.write_text(json.dumps({"changes": changes}, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"change_count": len(changes), "output": str(args.output)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
