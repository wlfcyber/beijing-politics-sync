from pathlib import Path

from docx import Document


DOCX = Path("/Users/wanglifei/Desktop/选必一_当代国际政治与经济_主观题术语宝典_终审修订版_20260601.docx")

REPLACEMENTS = {
    "终极审核版 · 461 条逐条回源核准 · 六大要素二级框架版（2026.5.31）":
    "主观题术语宝典 · 六大要素二级框架版（2026.6.1）",
    "终审修订版 · 450 条逐条回源核准 · 六大要素二级框架版（2026.6.1）":
    "主观题术语宝典 · 六大要素二级框架版（2026.6.1）",
}


def replace_in_runs(paragraph, old: str, new: str) -> int:
    count = 0
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            count += 1
    return count


def main() -> None:
    doc = Document(DOCX)
    counts = {old: 0 for old in REPLACEMENTS}

    for paragraph in doc.paragraphs:
        for old, new in REPLACEMENTS.items():
            counts[old] += replace_in_runs(paragraph, old, new)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old, new in REPLACEMENTS.items():
                        counts[old] += replace_in_runs(paragraph, old, new)

    doc.save(DOCX)

    report = Path(__file__).resolve().parents[1] / "03_outputs" / "FIX_REVISED_METADATA_REPORT.md"
    lines = [
        "# FIX_REVISED_METADATA_REPORT",
        "",
        f"- DOCX: `{DOCX}`",
    ]
    for old, count in counts.items():
        lines.append(f"- Replaced `{old}`: {count}")
    report.write_text("\n".join(lines) + "\n", encoding="utf-8")


if __name__ == "__main__":
    main()
