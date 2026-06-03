from pathlib import Path

from docx import Document


DOCX = Path("/Users/wanglifei/Desktop/选必一_当代国际政治与经济_主观题术语宝典_终审修订版_20260601.docx")

# Paragraph indexes are stable for the manually-built table of contents at the
# beginning of this document. Values were recalculated from the Word-exported PDF
# after removing unclosed entries, blank tail pages, and direct body font overrides.
TOC_PAGES = {
    10: 4,
    11: 4,
    12: 4,
    13: 16,
    14: 18,
    15: 18,
    16: 27,
    17: 34,
    18: 42,
    19: 42,
    20: 44,
    21: 45,
    22: 48,
    23: 57,
    24: 79,
    25: 97,
    26: 111,
    27: 116,
    28: 124,
    29: 130,
    30: 136,
    31: 140,
    32: 140,
    33: 154,
    34: 181,
    35: 197,
    36: 205,
    37: 213,
    38: 221,
    39: 222,
    40: 224,
    41: 224,
    42: 236,
    43: 272,
    44: 306,
    45: 311,
    46: 314,
    47: 314,
    48: 315,
    49: 323,
    50: 325,
    51: 326,
    52: 326,
    53: 326,
    54: 327,
    55: 329,
    56: 330,
    57: 332,
}


def set_paragraph_text_preserving_para_format(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def main() -> None:
    doc = Document(DOCX)
    updates = []
    for idx, page in TOC_PAGES.items():
        paragraph = doc.paragraphs[idx]
        label = paragraph.text.split("\t", 1)[0].rstrip()
        new_text = f"{label}\t{page}"
        old_text = paragraph.text
        if old_text != new_text:
            set_paragraph_text_preserving_para_format(paragraph, new_text)
            updates.append((idx, old_text, new_text))

    doc.save(DOCX)

    report = Path(__file__).resolve().parents[1] / "03_outputs" / "UPDATE_MANUAL_TOC_PAGES_REPORT.md"
    lines = [
        "# UPDATE_MANUAL_TOC_PAGES_REPORT",
        "",
        f"- DOCX: `{DOCX}`",
        f"- TOC entries updated: {len(updates)}",
        "",
        "## Updates",
    ]
    for idx, old, new in updates:
        lines.append(f"- paragraph {idx}: `{old}` -> `{new}`")
    report.write_text("\n".join(lines) + "\n", encoding="utf-8")


if __name__ == "__main__":
    main()
