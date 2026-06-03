#!/usr/bin/env python3
from __future__ import annotations

import re
from pathlib import Path

from docx import Document


DOCX = Path("/Users/wanglifei/Desktop/选必一_当代国际政治与经济_主观题术语宝典_终审修订版_20260601.docx")
REPORT = Path(
    "/Users/wanglifei/GaokaoPolitics/beijing-politics-sync/reports/"
    "选必一_哲学宝典式重建_2026-05-16/28_final_review_20260601/03_outputs/"
    "FIX_HARD_FAILURES_REPORT.md"
)

ENTRY_RE = re.compile(r"^\d+\.\s+.+?Q\d+(?:\(\d+\))?")


def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def is_entry(text: str) -> bool:
    return bool(ENTRY_RE.match(text.strip()))


def main() -> None:
    doc = Document(DOCX)
    fixed_prompt = 0
    for p in doc.paragraphs:
        text = p.text
        if (
            text.startswith("【设问】 任选一个议题")
            and "2024海淀二模" not in text
            and "知识运用准确；论述合乎逻辑" in text
        ):
            p.text = text.replace("知识运用准确；论述合乎逻辑", "知识运用准确、贴切；论述合乎逻辑")
            fixed_prompt += 1

    paragraphs = list(doc.paragraphs)
    delete_indexes: set[int] = set()
    removed_heads = []
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if is_entry(text) and "2026丰台期末Q20" in text:
            removed_heads.append(text)
            delete_indexes.add(i)
            j = i + 1
            while j < len(paragraphs):
                t = paragraphs[j].text.strip()
                if paragraphs[j].style.name.startswith("Heading") or is_entry(t):
                    break
                delete_indexes.add(j)
                j += 1

    for i in sorted(delete_indexes, reverse=True):
        delete_paragraph(paragraphs[i])

    doc.save(DOCX)
    REPORT.parent.mkdir(parents=True, exist_ok=True)
    REPORT.write_text(
        "\n".join(
            [
                "# FIX_HARD_FAILURES_REPORT",
                "",
                f"- DOCX: `{DOCX}`",
                f"- Fixed `2024海淀二模Q18(1)` prompt occurrences: {fixed_prompt}",
                f"- Removed unclosed `2026丰台期末Q20` entries: {len(removed_heads)}",
                "",
                "## Removed Entries",
                "",
                *[f"- {x}" for x in removed_heads],
                "",
                "Reason: OCR of the desktop original paper shows Q21 is the `四大全球倡议` short-review question; the `五大工程/中拉利益汇合点` text only appears in the rubric PDF page 64 as Q20. The original paper and rubric therefore do not close for these entries.",
                "",
            ]
        ),
        encoding="utf-8",
    )
    print(REPORT)


if __name__ == "__main__":
    main()
