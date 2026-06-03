#!/usr/bin/env python3
from __future__ import annotations

import csv
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


RUN_DIR = Path(
    "/Users/wanglifei/GaokaoPolitics/beijing-politics-sync/reports/"
    "选必一_哲学宝典式重建_2026-05-16/28_final_review_20260601"
)
SRC = Path("/Users/wanglifei/Desktop/选必一_当代国际政治与经济_主观题术语宝典_终极版_20260531.docx")
OUT = Path("/Users/wanglifei/Desktop/选必一_当代国际政治与经济_主观题术语宝典_终审修订版_20260601.docx")
AUDIT = RUN_DIR / "01_audit/FINAL_REVIEW_ENTRY_AUDIT.csv"
REPORT = RUN_DIR / "03_outputs/REPAIR_FINAL_DOCX_REPORT.md"


ENTRY_RE = re.compile(r"^(\d+)\.\s*(.+)$")


def clean(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip())


def shorten(s: str, limit: int = 96) -> str:
    s = clean(s)
    if len(s) <= limit:
        return s
    cut = s[:limit]
    for mark in "。；，、 ":
        pos = cut.rfind(mark)
        if pos >= 40:
            return cut[: pos + 1].rstrip("，、； ")
    return cut.rstrip("，、； ") + "……"


def studentize(s: str) -> str:
    replacements = {
        "设问要求": "题目引导学生",
        "设问需要": "题目引导学生",
        "材料中": "材料里",
        "材料中的": "材料里的",
        "答题核心": "关键落点",
    }
    for old, new in replacements.items():
        s = s.replace(old, new)
    return s


def question_tip(question: str) -> str:
    q = question or ""
    if any(k in q for k in ["为什么", "为何", "原因", "能够", "能为", "深层逻辑", "理论逻辑"]):
        return "问法追问成立根据，材料事实要上升为理论支点。"
    if any(k in q for k in ["如何", "怎样", "怎么", "措施", "应如何"]):
        return "问法追问行动路径，材料里的做法要转成政策原则和治理逻辑。"
    if any(k in q for k in ["意义", "价值", "作用", "影响", "积极影响", "意蕴"]):
        return "问法追问作用价值，材料里的结果要按对象和层次展开。"
    if any(k in q for k in ["理解", "认识", "谈谈"]):
        return "开放式问法先定性，再把材料事实转成理论判断。"
    if any(k in q for k in ["短评", "述评", "小论文"]):
        return "写作类设问先立观点，再用材料线索支撑观点。"
    return "题目引导学生把材料信息转化为学科判断。"


def module_bridge(module: str, core: str) -> str:
    if module == "时代背景":
        return f"它提示学生先看世界大势和外部环境，再用“{core}”给材料定方向。"
    if module == "理论":
        return f"它提示学生先找国际关系背后的依据，再用“{core}”解释合作、竞争或安全判断。"
    if module == "经济全球化":
        return f"它提示学生把贸易、投资、资源、产业链或开放平台归入“{core}”这一经济全球化术语。"
    if module == "政治多极化":
        return f"它提示学生把治理、秩序、多边合作和国际关系变化归入“{core}”这一政治多极化术语。"
    if module == "中国":
        return f"它提示学生把中国倡议、中国行动或中国立场归入“{core}”，说明中国如何作为。"
    if module == "联合国":
        return f"它提示学生把国际组织、宪章原则和多边机制归入“{core}”。"
    return f"它提示学生用“{core}”把材料线索和学科表达接起来。"


def build_why(row: dict[str, str]) -> str:
    trigger = studentize(shorten(row.get("材料触发点", ""), 118))
    answer = studentize(shorten(row.get("答案落点", ""), 92))
    module = row.get("module", "")
    core = row.get("core_point", "")
    tip = question_tip(row.get("设问", ""))
    bridge = module_bridge(module, core)
    return (
        f"{trigger} 这不是单个材料细节，而是通向该术语的线索。{bridge}"
        f"{tip} 抓住这条线，就能把材料转换为：{answer}"
    )


def set_run_font(run, bold: bool | None = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    if bold is not None:
        run.bold = bold


def insert_after(paragraph, label: str, body: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = paragraph._parent.add_paragraph()
    p._p = new_p
    p.style = paragraph.style
    r1 = p.add_run(label)
    set_run_font(r1, True)
    r2 = p.add_run(" " + body)
    set_run_font(r2, False)
    return p


def set_labeled_paragraph(paragraph, label: str, body: str) -> None:
    paragraph.clear()
    r1 = paragraph.add_run(label)
    set_run_font(r1, True)
    r2 = paragraph.add_run(" " + body)
    set_run_font(r2, False)


def load_missing_why() -> dict[int, dict[str, str]]:
    missing: dict[int, dict[str, str]] = {}
    with AUDIT.open(encoding="utf-8-sig", newline="") as f:
        for row in csv.DictReader(f):
            if "MISSING_WHY" in row.get("qa_flags", ""):
                missing[int(row["paragraph"])] = row
    return missing


def main() -> None:
    missing = load_missing_why()
    doc = Document(SRC)
    paragraphs = doc.paragraphs

    inserted = 0
    filled_blank = 0
    for idx in sorted(missing, reverse=True):
        row = missing[idx]
        if idx < 0 or idx >= len(paragraphs):
            continue
        entry_p = paragraphs[idx]
        # Find the question paragraph inside this entry. Stop if a why already exists
        # before the answer paragraph, because the audit row may be stale.
        question_p = None
        has_why = False
        blank_why_p = None
        j = idx + 1
        while j < len(paragraphs):
            text = clean(paragraphs[j].text)
            if j != idx and ENTRY_RE.match(text):
                break
            if text.startswith("核心答题点：") or paragraphs[j].style.name.startswith("Heading"):
                break
            if text.startswith("【为什么能想到】"):
                body = text.removeprefix("【为什么能想到】").strip()
                if body:
                    has_why = True
                else:
                    blank_why_p = paragraphs[j]
                break
            if text.startswith("【设问】"):
                question_p = paragraphs[j]
            if text.startswith("【答案落点】"):
                break
            j += 1
        if blank_why_p is not None:
            set_labeled_paragraph(blank_why_p, "【为什么能想到】", build_why(row))
            filled_blank += 1
        elif question_p is not None and not has_why:
            insert_after(question_p, "【为什么能想到】", build_why(row))
            inserted += 1

    # Re-fetch paragraphs after insertions, then renumber examples inside every
    # core point so visible序号 remains continuous for students.
    paragraphs = doc.paragraphs
    expected = None
    renumbered = 0
    for p in paragraphs:
        text = clean(p.text)
        if text.startswith("核心答题点："):
            expected = 1
            continue
        if p.style.name.startswith("Heading") and p.style.name != "Heading 3":
            expected = None
            continue
        if expected is None:
            continue
        m = ENTRY_RE.match(text)
        if not m:
            continue
        old_num = int(m.group(1))
        if old_num != expected:
            suffix = m.group(2)
            p.text = f"{expected}. {suffix}"
            for run in p.runs:
                set_run_font(run, None)
            renumbered += 1
        expected += 1

    doc.save(OUT)
    REPORT.parent.mkdir(parents=True, exist_ok=True)
    REPORT.write_text(
        "\n".join(
            [
                "# REPAIR_FINAL_DOCX_REPORT",
                "",
                f"- Source: `{SRC}`",
                f"- Output: `{OUT}`",
                f"- Missing why rows from audit: {len(missing)}",
                f"- Inserted `为什么能想到`: {inserted}",
                f"- Filled blank `为什么能想到`: {filled_blank}",
                f"- Renumbered example headings: {renumbered}",
                "",
                "Scope: only inserted missing student-facing `为什么能想到` paragraphs and repaired visible numbering inside each core point. Source labels, prompts, material triggers, answers, and same-question groups were left unchanged.",
                "",
            ]
        ),
        encoding="utf-8",
    )
    print(REPORT)
    print(OUT)


if __name__ == "__main__":
    main()
