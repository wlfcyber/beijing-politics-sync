#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX = Path("/Users/wanglifei/Desktop/选必一_当代国际政治与经济_主观题术语宝典_终审修订版_20260601.docx")
REPORT = Path(
    "/Users/wanglifei/GaokaoPolitics/beijing-politics-sync/reports/"
    "选必一_哲学宝典式重建_2026-05-16/28_final_review_20260601/03_outputs/"
    "NORMALIZE_REVISED_FRAMEWORK_REPORT.md"
)

TITLE_MAP = {
    "新增：国际分工、企业出海与国际竞争新优势": "国际分工、企业出海与国际竞争新优势",
    "新增：全球经济治理与规则权益": "全球经济治理与规则权益",
    "新增：产业链供应链与开放安全": "产业链供应链与开放安全",
    "新增：数字经济、绿色转型与新兴领域": "数字经济、绿色转型与新兴领域",
    "新增：贸易保护主义与多边贸易体制": "贸易保护主义与多边贸易体制",
    "新增：世界多极化、全球南方与发展中国家": "世界多极化、全球南方与发展中国家",
    "新增：国际秩序与国际法治": "国际秩序与国际法治",
    "新增：具体治理议题与文明互鉴": "具体治理议题与文明互鉴",
    "新增：联合国发展议程与全球治理贡献": "联合国发展议程与全球治理贡献",
    "新增：中国与联合国关系": "中国与联合国关系",
}


def set_font(run, bold: bool | None = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    if bold is not None:
        run.bold = bold


def insert_after(paragraph, text: str, style_name: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = paragraph._parent.add_paragraph()
    p._p = new_p
    p.style = style_name
    run = p.add_run(text)
    if style_name == "Normal":
        set_font(run)
    return p


def paragraph_exists(doc: Document, text: str) -> bool:
    return any(p.text.strip() == text for p in doc.paragraphs)


def insert_core_after_heading(doc: Document, heading_text: str, core_text: str, note: str) -> bool:
    if paragraph_exists(doc, core_text):
        return False
    for p in doc.paragraphs:
        if p.text.strip() == heading_text:
            note_p = insert_after(p, note, "Normal")
            insert_after(p, core_text, "Heading 3")
            # The insertion order is reversed because both are inserted after p.
            p._p.addnext(note_p._p)
            return True
    return False


def main() -> None:
    doc = Document(DOCX)
    renamed = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if text in TITLE_MAP:
            p.text = TITLE_MAP[text]
            renamed += 1
        else:
            for old, new in TITLE_MAP.items():
                if text.startswith(old + "\t"):
                    p.text = p.text.replace(old, new, 1)
                    renamed += 1
                    break

    inserted = []
    actions = [
        (
            "挑战",
            "核心答题点：霸权主义和强权政治、单边主义是和平与发展的现实挑战",
            "【学生判别】 看到单边霸凌、关税霸凌、贸易保护主义、脱钩断链、阵营对抗、恃强凌弱等词，先把它们看作和平与发展的挑战，再转入反对霸权主义和强权政治、坚持多边主义和国际关系民主化。",
        ),
        (
            "单边/斥：反对单边主义等错误倾向",
            "核心答题点：反对单边主义、霸权主义和强权政治",
            "【学生判别】 看到恃强凌弱、零和博弈、结盟对抗、单边霸凌、关税霸凌、保护主义等材料词，优先判断为反对单边主义、霸权主义和强权政治，再按题意接多边主义、国际关系民主化或公正合理国际秩序。",
        ),
        (
            "政策",
            "核心答题点：我国外交政策的宗旨是维护世界和平、促进共同发展",
            "【学生判别】 看到中国以合作、援助、倡议、外交行动促进和平与共同发展，不只是维护自身利益时，可以把它归入我国外交政策宗旨，再接独立自主和平外交政策或人类命运共同体。",
        ),
        (
            "国际分工、企业出海与国际竞争新优势",
            "核心答题点：积极\"走出去\"、对外直接投资、克服贸易壁垒",
            "【学生判别】 看到企业出海、海外市场、对外投资、贸易壁垒、反补贴税、国际市场风险等词，先判断为企业参与国际分工和走出去，再结合技术质量、市场多元化或国际规则维权展开。",
        ),
    ]
    for heading, core, note in actions:
        if insert_core_after_heading(doc, heading, core, note):
            inserted.append(core)

    doc.save(DOCX)
    REPORT.parent.mkdir(parents=True, exist_ok=True)
    REPORT.write_text(
        "\n".join(
            [
                "# NORMALIZE_REVISED_FRAMEWORK_REPORT",
                "",
                f"- DOCX: `{DOCX}`",
                f"- Removed `新增：` title prefixes: {renamed}",
                f"- Inserted student-facing core hints: {len(inserted)}",
                *[f"  - {x}" for x in inserted],
                "",
                "Scope: framework/heading normalization only. Source labels, prompts, answers, evidence cards, and same-question groups were not rewritten.",
                "",
            ]
        ),
        encoding="utf-8",
    )
    print(REPORT)


if __name__ == "__main__":
    main()
