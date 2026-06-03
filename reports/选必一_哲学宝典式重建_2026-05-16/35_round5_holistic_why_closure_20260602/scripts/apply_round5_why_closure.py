from __future__ import annotations

import argparse
import csv
import json
import re
from collections import defaultdict
from pathlib import Path

from docx import Document


ENTRY_RE = re.compile(r"^\d+\.\s+")
FIELD_RE = re.compile(r"^【([^】]+)】\s*(.*)$")

BAD_PATTERNS = [
    "对应",
    "支撑该核心",
    "支撑该术语",
    "把材料事实和该核心连成因果链",
    "这里应落在",
    "不能只复述题面事实",
    "题面事实和答案落点",
    "同一材料可服务多个答点",
    "同题组线索较多",
    "同题组存在多条线索",
    "抓住",
    "可采分",
    "用它说明",
    "这一层",
    "材料事实",
    "作答时",
    "答案支点",
    "落到",
    "要写",
    "设问要求",
    "材料关系与理论核心",
    "关键关系是",
    "这一结论",
    "这些线索要说明合作条件、治理路径或中国行动意义怎样成立",
    "从零散信息整理成",
    "贴标签",
    "理论入口",
]

FORBIDDEN_AFTER = [
    "支撑该术语",
    "支撑该核心",
    "把材料事实和该核心连成因果链",
    "这里应落在",
    "这说明本层应抓",
    "从该角度解释材料事实",
    "本题这一层对应",
    "说明材料怎样支撑该术语",
    "用它说明上述事实的答题依据",
    "可采分的判断",
    "题干要求说明国家对外行动的根本依据",
    "作答时",
    "答案支点",
    "设问要求",
    "材料关系与理论核心",
    "关键关系是",
    "这一结论",
    "这些线索要说明合作条件、治理路径或中国行动意义怎样成立",
    "从零散信息整理成",
    "贴标签",
    "理论入口",
]


def style_name(paragraph) -> str:
    return paragraph.style.name if paragraph.style else ""


def set_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def compact(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def parse_entries(doc: Document) -> list[dict]:
    entries: list[dict] = []
    current_h1 = current_h2 = current_h3 = ""
    starts: list[tuple[int, str, str, str]] = []

    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        st = style_name(paragraph)
        if st == "Heading 1":
            current_h1 = text
            current_h2 = ""
            current_h3 = ""
        elif st == "Heading 2":
            current_h2 = text
            current_h3 = ""
        elif st == "Heading 3":
            current_h3 = text
        elif ENTRY_RE.match(text):
            starts.append((idx, current_h1, current_h2, current_h3))

    for pos, (start, h1, h2, h3) in enumerate(starts):
        end = starts[pos + 1][0] if pos + 1 < len(starts) else len(doc.paragraphs)
        for idx in range(start + 1, end):
            if style_name(doc.paragraphs[idx]) in {"Heading 1", "Heading 2", "Heading 3"}:
                end = idx
                break
        fields = {}
        field_paras = {}
        for idx in range(start + 1, end):
            text = doc.paragraphs[idx].text.strip()
            m = FIELD_RE.match(text)
            if m:
                fields[m.group(1)] = m.group(2)
                field_paras[m.group(1)] = doc.paragraphs[idx]
        entries.append(
            {
                "idx": len(entries),
                "title": doc.paragraphs[start].text.strip(),
                "h1": h1,
                "h2": h2,
                "h3": h3,
                "fields": fields,
                "field_paras": field_paras,
            }
        )
    return entries


def clean_title(title: str) -> str:
    return ENTRY_RE.sub("", title).strip()


def core_name(h3: str) -> str:
    return h3.replace("核心答题点：", "").strip()


def prompt_intent(question: str) -> str:
    q = question or ""
    if "为什么" in q or "为何" in q or "原因" in q or "深层逻辑" in q or "理论逻辑" in q:
        return "解释原因和成立逻辑"
    if "如何" in q or "怎么" in q or "措施" in q or "应如何" in q:
        return "说明做法如何产生作用"
    if "意义" in q or "价值" in q or "作用" in q or "助力" in q or "注入" in q:
        return "说明作用和价值"
    if "理解" in q or "谈谈你对" in q or "阐释" in q or "阐述" in q:
        return "把材料判断上升为理论解释"
    if "短评" in q or "小论文" in q or "述评" in q:
        return "围绕主题组织论证"
    if "评价" in q or "精彩" in q or "担当" in q:
        return "评价材料行动的理论含义"
    if "概括" in q or "差异" in q:
        return "概括材料之间的关系差异"
    return "说明题面事实背后的理论关系"


def relation_label(entry: dict) -> str:
    h1, h2, h3 = entry["h1"], entry["h2"], entry["h3"]
    text = f"{h1} {h2} {h3}"
    if "共同利益" in text:
        return "合作成立的利益基础"
    if "国际竞争" in text or "综合国力" in text:
        return "能力来源与国际竞争"
    if "国家利益" in text or "国家安全" in text or "自力更生" in text:
        return "利益安全与发展主动权"
    if h1 == "时代背景":
        return "时代条件与外部环境"
    if "贸易" in text or "投资" in text or "自由化便利化" in text:
        return "跨境交易和要素流动条件"
    if "资源" in text or "两个市场" in text:
        return "国内国际市场资源联动"
    if "开放" in text or "开放型经济" in text or "制度型" in text:
        return "开放平台、规则和市场联通"
    if "产业链" in text or "供应链" in text:
        return "产业链供应链运行机制"
    if "共享经济全球化成果" in text or "共同繁荣" in text:
        return "经济全球化成果共享"
    if "全球经济治理" in text or "规则制定" in text:
        return "全球经济治理参与"
    if "国际标准" in text or "规则话语权" in text or "国际规则" in text:
        return "国际规则与标准参与"
    if "文明互鉴" in text or "共同价值" in text or "文明多样性" in text:
        return "文明交流互鉴"
    if "南南合作" in text or "全球南方" in text:
        return "全球南方合作"
    if "民心相通" in text:
        return "人文交流和民心基础"
    if "多边主义" in text:
        return "多边平台和共同规则"
    if "国际关系民主化" in text:
        return "国际事务平等参与"
    if "新型国际关系" in text or "互利共赢" in text:
        return "国际关系方向"
    if "全球治理" in text or "国际秩序" in text:
        return "全球治理结构"
    if "人类命运共同体" in text:
        return "共同体价值统摄"
    if "创始会员国" in text or "安理会" in text or "国际政治经济格局" in text:
        return "中国国际地位和治理分量"
    if "联合国" in text:
        return "联合国框架和多边机制"
    if "习近平外交思想" in text or "中国特色大国外交" in text or "中国式现代化" in text:
        return "外交指导思想和现代化服务方向"
    if h1 == "中国" and "政策" in h2:
        return "中国外交政策选择"
    if h1 == "中国" and "智慧" in h2:
        return "中国方案和中国主张"
    if h1 == "中国" and "责任" in h2:
        return "大国责任与发展合作"
    if h1 == "联合国":
        return "联合国框架和多边机制"
    return "材料关系与理论核心"


ANCHOR_PATTERNS = [
    r"全球治理倡议",
    r"全球发展倡议",
    r"人类命运共同体",
    r"中国—中亚峰会",
    r"一带一路",
    r"中非合作论坛",
    r"进博会",
    r"服贸会",
    r"消博会",
    r"链博会",
    r"鲁班工坊",
    r"菌草",
    r"AI眼镜",
    r"开源大模型",
    r"创新药",
    r"人形机器人",
    r"自贸协定",
    r"单方面免签",
    r"零关税",
    r"二线管住",
    r"数字化技术",
    r"绿色低碳技术",
    r"供应链金融",
    r"全球南方",
    r"联合国",
    r"安理会",
    r"《联合国宪章》",
    r"减贫",
    r"教育",
    r"卫生健康",
    r"气候治理",
    r"脱钩断链",
    r"贸易保护主义",
    r"单边主义",
    r"霸权主义",
    r"强权政治",
    r"\d+个国家",
    r"\d+多个国家",
    r"\d+项",
    r"\d+多项",
    r"\d+家",
]


def anchors_from(entry: dict) -> str:
    blob = "；".join(
        [
            entry["fields"].get("材料触发点", ""),
            entry["fields"].get("答案落点", ""),
            entry["fields"].get("设问", ""),
        ]
    )
    anchors: list[str] = []
    for pattern in ANCHOR_PATTERNS:
        for m in re.findall(pattern, blob):
            if m not in anchors:
                anchors.append(m)
            if len(anchors) >= 4:
                break
        if len(anchors) >= 4:
            break
    for quoted in re.findall(r"[“\"]([^”\"]{2,18})[”\"]", blob):
        if quoted not in anchors and not quoted.startswith("为什么能想到"):
            anchors.append(quoted)
        if len(anchors) >= 4:
            break
    if anchors:
        return "、".join(anchors[:4])
    answer = entry["fields"].get("答案落点", "")
    answer = re.sub(r"[。；，,].*$", "", answer).strip()
    return answer[:32] or "题面给出的关键事实"


def bridge_text(entry: dict) -> str:
    core = core_name(entry["h3"])
    text = f"{entry['h1']} {entry['h2']} {entry['h3']} {entry['fields'].get('答案落点','')}"
    if "共同利益" in core:
        return "这些主体能合作，是因为需求和收益有交汇点"
    if "国际竞争" in core or "综合国力" in core:
        return "中国能够提供稳定性，必须先说明科技和经济实力这类能力基础"
    if "维护国家利益" in core:
        return "国家行为要回到主权、安全和发展利益这一根本立足点"
    if "和平与发展" in core:
        return "中国行动的正当性要放在和平、发展、合作的时代方向中判断"
    if "贸易" in core or "投资" in core:
        return "这些安排改变的是跨境合作成本、市场准入和交易便利程度"
    if "资源" in core or "两个市场" in core:
        return "材料中的市场、资源和要素不是孤立出现，而是在国内国际之间重新配置"
    if "共享经济全球化成果" in core or "共同繁荣" in core:
        return "材料强调发展机会由更多国家共享，不能只停留在中国自身增长"
    if "顺应经济全球化" in core or "推动经济全球化" in core or "普惠包容" in core or "开放型世界经济" in core:
        return "材料中的开放举措要指向更广范围的参与、更低的合作成本和更公平的发展机会"
    if "开放" in core or "开放型" in core:
        return "开放不是口号，而是通过平台、规则、通关和市场条件变成实际联通"
    if "产业链" in core or "供应链" in core:
        return "题面要解释链条为什么更稳、更畅或更能共享收益"
    if "技术创新" in core or "产品质量" in core or "竞争力" in core:
        return "材料中的技术和质量线索要转化为企业或国家竞争能力的来源"
    if "新型国际关系" in core or "互利共赢" in core:
        return "材料强调的不是单方获利，而是以平等合作替代对抗和零和"
    if "共商共建共享" in core:
        return "治理行动由多方参与、共同建设并共享结果，而不是单边安排"
    if "国际关系民主化" in core:
        return "材料中的平等参与和发展中国家诉求，指向国际事务不能由少数国家垄断"
    if "多边主义" in core:
        return "问题的解决依靠多边平台和共同规则，而不是少数国家单边决定"
    if "人类命运共同体" in core:
        return "材料把中国发展、他国利益和全球问题放在同一命运关联中"
    if "中国智慧" in core or "中国方案" in core:
        return "中国行动提供的是可供共同问题使用的理念、机制或实践方案"
    if "正确义利观" in core:
        return "中国处理合作时不是单方取利，而是兼顾自身发展、他国关切和共同收益"
    if "改善发展中国家民生" in core or "小而美" in core:
        return "材料中的项目要指向民生改善和可持续发展能力提升"
    if "民心相通" in core or "文明互鉴" in core or "文明平等" in core or "共同价值" in core:
        return "人文和文明线索说明合作需要价值认同、相互理解和社会基础"
    if "南南合作" in core or "全球南方" in core:
        return "材料要说明中国与发展中国家在共同发展诉求中形成合作合力"
    if "绿色低碳" in core or "气候治理" in core:
        return "气候和绿色线索要说明中国把转型能力转化为全球治理贡献"
    if "全球治理体系" in core or "国际政治经济新秩序" in core:
        return "材料关注的是治理规则和秩序方向是否更加公正合理"
    if "责任" in core or "担当" in core:
        return "中国把自身能力转化为公共贡献，承担的是发展合作和全球治理责任"
    if "习近平外交思想" in core or "中国特色大国外交" in core or "中国式现代化" in core:
        return "外交行动既要服务国家现代化，也要营造和平稳定的外部环境"
    if "联合国" in text or "宪章" in core:
        return "这类行动要放回联合国框架、国际法基础和多边秩序中理解"
    return "材料中的因果、路径或价值判断要由这个核心点来解释清楚"


def answer_focus(entry: dict) -> str:
    answer = entry["fields"].get("答案落点", "")
    answer = re.sub(r"\s+", "", answer)
    if not answer:
        return "该核心结论"
    if len(answer) <= 140:
        return answer
    # Prefer the segment containing the core keyword when possible.
    core = core_name(entry["h3"])
    parts = re.split(r"[。；;]", answer)
    for part in parts:
        if part and any(token and token in part for token in re.split(r"[，、和与及（）()]", core)[:4]):
            return clip_text(part)
    return clip_text(answer)


def clip_text(text: str, limit: int = 140) -> str:
    text = text.strip()
    if len(text) <= limit:
        return text
    chunks = re.split(r"([，,；;。])", text)
    current = ""
    for i in range(0, len(chunks), 2):
        part = chunks[i]
        punct = chunks[i + 1] if i + 1 < len(chunks) else ""
        candidate = current + part + punct
        if len(candidate) > limit:
            break
        current = candidate
    if len(current) >= 24:
        return current.rstrip("，,；;。")
    return text[: limit - 1].rstrip("，,；;。") + "…"


def make_why(entry: dict) -> str:
    title = clean_title(entry["title"])
    core = core_name(entry["h3"])
    intent = prompt_intent(entry["fields"].get("设问", ""))
    relation = relation_label(entry)
    anchors = anchors_from(entry)
    bridge = bridge_text(entry)
    focus = answer_focus(entry)
    if len(core) > 38:
        why = (
            f"在{title}里，“{core}”用于解释{anchors}中的{relation}。"
            f"设问要{intent}，重点是{bridge}；答案应写向“{focus}”。"
        )
        return why.replace("。。", "。")
    variant = entry["idx"] % 4
    if variant == 0:
        why = (
            f"“{core}”用于{title}，是因为设问要{intent}，{anchors}已经把{relation}摆出来。"
            f"{bridge}，所以答案应围绕“{focus}”展开。"
        )
    elif variant == 1:
        why = (
            f"在{title}里想到“{core}”，不是停在{anchors}本身，而是看见它们共同指向{relation}。"
            f"完成{intent}时，要说明{bridge}，答案就落在“{focus}”。"
        )
    elif variant == 2:
        why = (
            f"“{core}”能解释{title}中的{anchors}，因为这些线索提示的是{relation}。"
            f"围绕{intent}回答时，应讲清{bridge}，再写出“{focus}”。"
        )
    else:
        why = (
            f"“{core}”解释{title}，因为{anchors}不只是材料细节，而是在呈现{relation}。"
            f"这个核心能说明{bridge}，因此答案写“{focus}”。"
        )
    why = why.replace("。。", "。")
    return why


def should_rewrite(entry: dict, duplicate_prefixes: set[str]) -> bool:
    why = entry["fields"].get("为什么能想到", "")
    trigger = entry["fields"].get("材料触发点", "")
    w = compact(why)
    t = compact(trigger)
    if w[:40] in duplicate_prefixes:
        return True
    if t and (w.startswith(t[:30]) or t[:45] in w[:100]):
        return True
    if any(pat in why for pat in BAD_PATTERNS):
        return True
    return False


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--summary", type=Path, required=True)
    args = parser.parse_args()

    doc = Document(args.input)
    entries = parse_entries(doc)
    prefix_map: dict[str, list[dict]] = defaultdict(list)
    for entry in entries:
        prefix_map[compact(entry["fields"].get("为什么能想到", ""))[:40]].append(entry)
    duplicate_prefixes = {k for k, v in prefix_map.items() if k and len(v) > 1}

    changes = []
    for entry in entries:
        if not should_rewrite(entry, duplicate_prefixes):
            continue
        paragraph = entry["field_paras"].get("为什么能想到")
        if paragraph is None:
            continue
        old = entry["fields"].get("为什么能想到", "")
        new = make_why(entry)
        if any(bad in new for bad in FORBIDDEN_AFTER):
            raise RuntimeError(f"Generated forbidden phrase for {entry['idx']}: {new}")
        set_text(paragraph, f"【为什么能想到】 {new}")
        changes.append(
            {
                "idx": entry["idx"],
                "title": entry["title"],
                "h1": entry["h1"],
                "h2": entry["h2"],
                "h3": entry["h3"],
                "old_prefix": old[:80],
                "new": new,
            }
        )

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    args.summary.parent.mkdir(parents=True, exist_ok=True)
    args.summary.write_text(
        json.dumps(
            {
                "entry_count": len(entries),
                "rewritten_count": len(changes),
                "changes": changes,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    print(json.dumps({"entry_count": len(entries), "rewritten_count": len(changes)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
