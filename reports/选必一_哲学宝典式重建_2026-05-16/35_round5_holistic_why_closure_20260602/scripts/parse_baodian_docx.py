from __future__ import annotations

import argparse
import csv
import json
import re
from pathlib import Path

from docx import Document


FIELD_PAT = re.compile(r"^【([^】]+)】\s*(.*)$")
ENTRY_PAT = re.compile(r"^\d+\.\s+")


def normalize(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def iter_entries(docx_path: Path):
    doc = Document(docx_path)
    current = {
        "h1": "",
        "h2": "",
        "h3": "",
        "entry": None,
    }
    entries = []

    def finish():
        entry = current.get("entry")
        if entry:
            entries.append(entry)
            current["entry"] = None

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name if para.style else ""
        if not text:
            continue

        if style == "Heading 1":
            finish()
            current["h1"] = text
            current["h2"] = ""
            current["h3"] = ""
            continue
        if style == "Heading 2":
            finish()
            current["h2"] = text
            current["h3"] = ""
            continue
        if style == "Heading 3":
            finish()
            current["h3"] = text
            continue

        if ENTRY_PAT.match(text):
            finish()
            current["entry"] = {
                "entry_para_idx": idx,
                "entry_title": text,
                "h1": current["h1"],
                "h2": current["h2"],
                "h3": current["h3"],
                "fields": {},
                "field_para_idx": {},
                "paras": [],
            }
            continue

        entry = current.get("entry")
        if entry is None:
            continue

        entry["paras"].append({"idx": idx, "style": style, "text": text})
        m = FIELD_PAT.match(text)
        if m:
            label, value = m.group(1), m.group(2)
            entry["fields"][label] = value
            entry["field_para_idx"][label] = idx
        elif entry["paras"]:
            # Continuation lines are mainly for 同题组. Keep them as raw paras.
            pass

    finish()
    return entries, len(doc.paragraphs)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--json", type=Path)
    parser.add_argument("--csv", type=Path)
    args = parser.parse_args()

    entries, paragraph_count = iter_entries(args.docx)
    identical = []
    for i, entry in enumerate(entries):
        trigger = entry["fields"].get("材料触发点", "")
        why = entry["fields"].get("为什么能想到", "")
        if trigger and why and normalize(trigger) == normalize(why):
            identical.append(i)

    summary = {
        "paragraph_count": paragraph_count,
        "entry_count": len(entries),
        "question_fields": sum(1 for e in entries if "设问" in e["fields"]),
        "trigger_fields": sum(1 for e in entries if "材料触发点" in e["fields"]),
        "why_fields": sum(1 for e in entries if "为什么能想到" in e["fields"]),
        "answer_fields": sum(1 for e in entries if "答案落点" in e["fields"]),
        "identical_trigger_why": len(identical),
    }
    print(json.dumps(summary, ensure_ascii=False, indent=2))

    if args.json:
        args.json.write_text(json.dumps(entries, ensure_ascii=False, indent=2), encoding="utf-8")

    if args.csv:
        with args.csv.open("w", encoding="utf-8-sig", newline="") as f:
            writer = csv.DictWriter(
                f,
                fieldnames=[
                    "idx",
                    "entry_para_idx",
                    "h1",
                    "h2",
                    "h3",
                    "entry_title",
                    "材料触发点",
                    "设问",
                    "为什么能想到",
                    "答案落点",
                    "同题组",
                    "identical_trigger_why",
                ],
                extrasaction="ignore",
            )
            writer.writeheader()
            for i, entry in enumerate(entries):
                row = {
                    "idx": i,
                    "entry_para_idx": entry["entry_para_idx"],
                    "h1": entry["h1"],
                    "h2": entry["h2"],
                    "h3": entry["h3"],
                    "entry_title": entry["entry_title"],
                    "identical_trigger_why": normalize(entry["fields"].get("材料触发点", ""))
                    == normalize(entry["fields"].get("为什么能想到", "")),
                }
                row.update(entry["fields"])
                writer.writerow(row)


if __name__ == "__main__":
    main()
