from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

import fitz
from docx import Document


FOOTER_RE = re.compile(r"—\s*\d+\s*—")


def set_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def extract_manual_toc(doc: Document) -> list[tuple[int, str, str]]:
    labels: list[tuple[int, str, str]] = []
    for idx, paragraph in enumerate(doc.paragraphs[:80]):
        text = paragraph.text.strip()
        if idx >= 9 and "\t" in text:
            label, page = text.rsplit("\t", 1)
            if page.strip().isdigit():
                labels.append((idx, label.strip(), page.strip()))
    return labels


def clean_page_lines(page) -> list[str]:
    lines = []
    for line in page.get_text("text").splitlines():
        line = line.strip()
        if not line:
            continue
        if FOOTER_RE.fullmatch(line):
            continue
        if "...." in line:
            continue
        lines.append(line)
    return lines


def find_actual_pages(pdf_path: Path, labels: list[tuple[int, str, str]]) -> list[tuple[int, str, str, int]]:
    pdf = fitz.open(str(pdf_path))
    page_lines = [clean_page_lines(page) for page in pdf]
    results: list[tuple[int, str, str, int]] = []
    start_page = 0
    for idx, label, old_page in labels:
        found = None
        for pno in range(start_page, len(page_lines)):
            if any(line == label for line in page_lines[pno]):
                found = pno + 1
                break
        if found is None:
            for pno in range(start_page, len(page_lines)):
                if label in "\n".join(page_lines[pno]):
                    found = pno + 1
                    break
        if found is None:
            raise RuntimeError(f"Cannot locate TOC label in PDF: {label}")
        results.append((idx, label, old_page, found))
        start_page = max(0, found - 1)
    return results


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", type=Path, required=True)
    parser.add_argument("--pdf", type=Path, required=True)
    parser.add_argument("--out", type=Path, required=True)
    parser.add_argument("--summary", type=Path, required=True)
    args = parser.parse_args()

    doc = Document(args.docx)
    labels = extract_manual_toc(doc)
    actual = find_actual_pages(args.pdf, labels)
    changes = []
    for idx, label, old_page, actual_page in actual:
        paragraph = doc.paragraphs[idx]
        new_text = f"{label}\t{actual_page}"
        if paragraph.text.strip() != new_text:
            set_text(paragraph, new_text)
            changes.append({"label": label, "old": old_page, "new": actual_page})

    doc.save(args.out)
    args.summary.write_text(
        json.dumps({"toc_items": len(actual), "changes": changes}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(json.dumps({"toc_items": len(actual), "changed": len(changes)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
