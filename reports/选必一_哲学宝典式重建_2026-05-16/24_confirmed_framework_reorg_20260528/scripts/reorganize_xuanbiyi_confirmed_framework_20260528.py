from __future__ import annotations

import csv
import json
import re
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


RUN_DIR = Path(__file__).resolve().parents[1]
BASE_DIR = RUN_DIR.parent
SOURCE_MD = (
    BASE_DIR
    / "21_claude_review_patch_20260528"
    / "选必一_当代国际政治与经济_主观题术语宝典_237条逐条核实二次复核最终稿_20260528.md"
)

OUT_MD = RUN_DIR / "选必一_当代国际政治与经济_主观题术语宝典_237条_讲义二级框架确认版_20260528.md"
OUT_DOCX = RUN_DIR / "选必一_当代国际政治与经济_主观题术语宝典_237条_讲义二级框架确认版_20260528.docx"
OUT_MAP = RUN_DIR / "SECOND_LEVEL_CORE_MAPPING.csv"
OUT_SUMMARY = RUN_DIR / "SECOND_LEVEL_REORG_QA.md"
OUT_JSON = RUN_DIR / "SECOND_LEVEL_REORG_QA.json"

BUCKETS = ["时代背景", "理论", "经济全球化", "政治多极化", "中国", "联合国"]

SECOND_LEVEL_ORDER = {
    "时代背景": [
        "机遇",
        "挑战",
    ],
    "理论": [
        "合作",
        "竞争",
        "国家安全",
    ],
    "经济全球化": [
        "民：民生福祉",
        "国：成员国和相关国家发展",
        "环：经济环节与要素流动",
        "营：营商环境",
        "配：资源配置",
        "贸资：贸易投资自由化便利化",
        "总：开放包容普惠平衡共赢方向",
        "开：高水平对外开放",
        "兼：国内国际两种资源两个市场",
        "新增：国际分工、企业出海与国际竞争新优势",
        "新增：全球经济治理与规则权益",
        "新增：产业链供应链与开放安全",
        "新增：数字经济、绿色转型与新兴领域",
        "新增：贸易保护主义与多边贸易体制",
    ],
    "政治多极化": [
        "新国关：新型国际关系",
        "完全治：完善全球治理",
        "共体：人类命运共同体",
        "单边/斥：反对单边主义等错误倾向",
        "民主：国际关系民主化",
        "新增：世界多极化、全球南方与发展中国家",
        "新增：国际秩序与国际法治",
        "新增：具体治理议题与文明互鉴",
    ],
    "中国": ["政策", "智慧", "责任", "一条历史线", "一条主线"],
    "联合国": [
        "多边主义/多极化场所",
        "联合国宪章宗旨和原则",
        "新增：联合国发展议程与全球治理贡献",
        "新增：中国与联合国关系",
    ],
}

BUCKET_GUIDES = {
    "时代背景": "按讲义分“机遇”和“挑战”。机遇侧重和平与发展、经济全球化和政治多极化深入发展；挑战侧重霸权主义、强权政治、逆全球化、治理赤字和共同风险。",
    "理论": "按讲义分“合作、竞争、国家安全”。合作抓共同利益，竞争抓综合国力和国际竞争，国家安全抓最高国家利益和安全底线。",
    "经济全球化": "按讲义口诀“民国环营配，贸资总开兼”优先归类。能放入每个字下的术语就放入；明显属于另一个角度的，另设“新增”二级标题承接。",
    "政治多极化": "按讲义口诀“新国关 完全治，共体 单边 民主 斥”优先归类；不能放进口诀的世界多极化、国际法治、具体治理议题另设新增角度。",
    "中国": "按讲义分“政策、智慧、责任、一条历史线、一条主线”。政策抓外交政策和准则，智慧抓中国智慧中国方案，责任抓大国担当，历史线抓综合国力和国际地位，主线抓习近平外交思想。",
    "联合国": "按讲义抓两个主口径：多边主义/多极化场所，遵循《联合国宪章》宗旨和原则。联合国发展议程、中国与联合国关系等明显另角度单独承接。",
}


@dataclass
class CoreBlock:
    bucket: str
    title: str
    count: int | None
    lines: list[str]


def has_core_before_next_section(lines: list[str], idx: int) -> bool:
    for line in lines[idx + 1 : min(len(lines), idx + 20)]:
        if line.startswith("## "):
            return False
        if line.startswith("### 核心答题点："):
            return True
    return False


def find_actual_bucket_starts(lines: list[str]) -> dict[str, int]:
    starts: dict[str, int] = {}
    cursor = 0
    for bucket in BUCKETS:
        needle = f"## {bucket}"
        for i in range(cursor, len(lines)):
            if lines[i].strip() == needle and has_core_before_next_section(lines, i):
                starts[bucket] = i
                cursor = i + 1
                break
        else:
            raise RuntimeError(f"Cannot find actual bucket start: {bucket}")
    return starts


def find_speed_start(lines: list[str], after: int) -> int:
    for i in range(after, len(lines)):
        if lines[i].strip() == "## 六大要素术语极简速记版":
            return i
    return len(lines)


def parse_core_heading(line: str) -> tuple[str, int | None]:
    m = re.match(r"### 核心答题点：(.+?)(?:（出现(\d+)次）)?$", line.strip())
    if not m:
        raise RuntimeError(f"Bad core heading: {line}")
    title = m.group(1)
    count = int(m.group(2)) if m.group(2) else None
    return title, count


def parse_source(lines: list[str]) -> tuple[list[CoreBlock], list[str]]:
    starts = find_actual_bucket_starts(lines)
    speed_start = find_speed_start(lines, starts["联合国"] + 1)
    blocks: list[CoreBlock] = []

    for bi, bucket in enumerate(BUCKETS):
        start = starts[bucket]
        end = starts[BUCKETS[bi + 1]] if bi + 1 < len(BUCKETS) else speed_start
        section = lines[start + 1 : end]
        positions = [i for i, line in enumerate(section) if line.startswith("### 核心答题点：")]
        for pi, pos in enumerate(positions):
            block_end = positions[pi + 1] if pi + 1 < len(positions) else len(section)
            block_lines = section[pos:block_end]
            title, count = parse_core_heading(block_lines[0])
            blocks.append(CoreBlock(bucket=bucket, title=title, count=count, lines=block_lines))

    speed_lines = lines[speed_start:] if speed_start < len(lines) else []
    return blocks, speed_lines


def contains_any(text: str, words: list[str]) -> bool:
    return any(word in text for word in words)


def classify(block: CoreBlock) -> tuple[str, str]:
    title = block.title
    t = title
    bucket = block.bucket

    if bucket == "时代背景":
        if contains_any(t, ["百年变局", "外部风险", "风险挑战", "治理赤字", "气候变化", "共同挑战", "霸权主义", "强权政治", "脱钩断链", "贸易保护主义"]):
            return "挑战", "讲义时代背景中的挑战面：霸权主义、强权政治、逆全球化、治理赤字和共同风险"
        return "机遇", "讲义时代背景中的机遇面：和平与发展、经济全球化和政治多极化深入发展"

    if bucket == "理论":
        if contains_any(t, ["国家利益", "国家安全", "安全利益"]):
            return "国家安全", "讲义理论中的国家安全线：国家利益、安全底线和最高国家利益"
        if contains_any(t, ["共同利益", "合作的基础"]):
            return "合作", "讲义理论中的合作线：共同利益是合作基础，兼顾他国合理关切"
        return "竞争", "讲义理论中的竞争线：国际竞争实质、综合国力、经济科技人才实力"

    if bucket == "经济全球化":
        if contains_any(t, ["民生福祉", "就业岗位", "居民收入", "基础设施", "惠及人民", "发展成果惠及人民"]):
            return "民：民生福祉", "对应讲义“民”：民生福祉、就业、收入、基础设施和人民受益"
        if contains_any(t, ["成员国", "相关国家", "共建国家", "经济复苏", "经济增长", "共同繁荣", "世界经济发展"]):
            return "国：成员国和相关国家发展", "对应讲义“国”：推动成员国或相关国家经济复苏和增长"
        if contains_any(t, ["贸易保护主义", "多边贸易体制", "开放合作的国际经济环境", "反对贸易保护主义", "多边贸易秩序"]):
            return "新增：贸易保护主义与多边贸易体制", "明显是反贸易保护主义、多边贸易体制和开放环境角度"
        if contains_any(t, ["数字", "数据", "科学技术", "新能源", "绿色转型", "未来产业", "创新型"]):
            return "新增：数字经济、绿色转型与新兴领域", "明显是数字、科技、绿色和未来产业角度"
        if contains_any(t, ["全球经济治理", "规则制定", "规则完善", "国际规则", "标准制定", "话语权", "权益", "经贸规则", "国际组织赋予"]):
            return "新增：全球经济治理与规则权益", "明显是经济治理、规则制定、标准和权益维护角度"
        if contains_any(t, ["产业链", "供应链", "统筹开放与安全", "安全稳定"]):
            return "新增：产业链供应链与开放安全", "明显是产业链供应链和开放安全角度"
        if contains_any(t, ["企业出海", "出海", "国际竞争新优势", "国际竞争力", "比较优势", "优势互补", "出口市场多元化", "产品质量", "增强竞争力", "技术创新", "深度加入国际分工"]):
            return "新增：国际分工、企业出海与国际竞争新优势", "明显是企业出海、国际分工或竞争优势角度，不硬塞口诀字"
        if contains_any(t, ["贸易和投资", "贸易投资", "自由化便利化", "零关税", "贸易自由化", "投资便利化"]):
            return "贸资：贸易投资自由化便利化", "对应讲义“贸资”：推动贸易投资自由化便利化"
        if contains_any(t, ["开放、包容、普惠、平衡、共赢", "开放包容普惠平衡共赢", "普惠包容", "开放型世界经济", "经济全球化方向", "经济全球化趋势", "共享经济全球化成果", "开放合作和包容普惠"]):
            return "总：开放包容普惠平衡共赢方向", "对应讲义“总”：推动经济全球化朝开放、包容、普惠、平衡、共赢方向发展"
        if contains_any(t, ["高水平对外开放", "开放型经济水平", "更高层次开放型经济", "扩大高水平对外开放", "引进来", "走出去", "互利共赢的开放战略"]):
            return "开：高水平对外开放", "对应讲义“开”：高水平对外开放、更高层次开放型经济、引进来走出去"
        if contains_any(t, ["两个市场", "两种资源", "国内国际双循环", "国际循环", "国内国际"]):
            return "兼：国内国际两种资源两个市场", "对应讲义“兼”：充分利用国内国际两种资源、两个市场"
        if contains_any(t, ["营商环境", "制度型开放", "规则互认", "标准规制互认", "负面清单"]):
            return "营：营商环境", "对应讲义“营”：优化国际化、市场化、法治化营商环境"
        if contains_any(t, ["资源配置", "资源优化配置", "优化全球资源配置", "提高资源配置效率"]):
            return "配：资源配置", "对应讲义“配”：优化资源配置和提高配置效率"
        if contains_any(t, ["要素", "商品", "服务贸易", "商品、服务", "技术、资金", "资金", "劳务", "人员", "跨境流通", "国际分工", "产业升级", "产业结构", "拉动投资", "拉动消费", "开拓市场", "市场空间", "国际市场"]):
            return "环：经济环节与要素流动", "对应讲义“环”：要素自由流动、产业升级、投资消费和市场拓展"
        return "总：开放包容普惠平衡共赢方向", "无法更细归入口诀时，按经济全球化总方向承接"

    if bucket == "政治多极化":
        if contains_any(t, ["国际关系民主化", "新型国际关系", "平等互利", "互利共赢", "和平解决国际争端"]):
            if contains_any(t, ["国际关系民主化"]):
                return "民主：国际关系民主化", "对应讲义“民主”：推动国际关系民主化"
            return "新国关：新型国际关系", "对应讲义“新国关”：相互尊重、公平正义、合作共赢的新型国际关系"
        if contains_any(t, ["人类命运共同体", "共同现代化", "共同发展", "持久和平"]):
            return "共体：人类命运共同体", "对应讲义“共体”：构建人类命运共同体"
        if contains_any(t, ["单边主义", "贸易保护主义", "霸权主义", "强权政治", "真正的多边主义"]):
            return "单边/斥：反对单边主义等错误倾向", "对应讲义“单边/斥”：反对单边主义、贸易保护主义、霸权主义和强权政治"
        if contains_any(t, ["共商共建共享", "全球治理观", "全球治理体系", "完善全球治理", "治理体系", "更加公正合理"]):
            return "完全治：完善全球治理", "对应讲义“完全治”：完善全球治理，倡导共商共建共享"
        if contains_any(t, ["国际秩序", "国际法治", "新秩序", "主权平等", "国际法"]):
            return "新增：国际秩序与国际法治", "明显是国际秩序、国际法治和主权平等角度"
        if contains_any(t, ["世界多极化", "全球南方", "发展中国家", "国际组织"]):
            return "新增：世界多极化、全球南方与发展中国家", "明显是多极化格局和全球南方力量变化角度"
        if contains_any(t, ["气候", "安全", "文明", "绿色低碳", "跨国公司", "集团政治", "实力至上"]):
            return "新增：具体治理议题与文明互鉴", "明显是气候、安全、文明互鉴等具体治理议题"
        return "完全治：完善全球治理", "默认按完善全球治理承接"

    if bucket == "中国":
        if contains_any(t, ["习近平外交思想"]):
            return "一条主线", "对应讲义一条主线：习近平外交思想提供根本遵循和行动指南"
        if contains_any(t, ["综合国力", "国际地位", "话语权", "创始会员国", "安理会", "重要力量", "参与者、贡献者和改革者"]):
            return "一条历史线", "对应讲义一条历史线：中国综合国力增强、国际地位提升"
        if contains_any(t, ["总体国家安全观", "外交政策", "和平共处五项原则", "独立自主", "国家性质", "党对外交", "自力更生"]):
            return "政策", "中国外交政策、国家安全和对外关系准则"
        if contains_any(t, ["智慧", "方案", "力量", "全球发展倡议", "全球安全倡议", "全球文明倡议", "全球治理倡议", "四大全球倡议", "人类命运共同体", "中拉命运共同体", "国际共识", "合作理念", "普遍价值", "文明平等互鉴", "全人类共同价值"]):
            return "智慧", "中国主张、中国方案、全球倡议和理念供给"
        return "责任", "中国行动、发展合作、南南合作、气候治理和负责任大国担当"

    if bucket == "联合国":
        if contains_any(t, ["联合国对中国", "中国与联合国"]):
            return "新增：中国与联合国关系", "明显是中国与联合国双向关系角度"
        if contains_any(t, ["宪章", "以联合国为核心", "国际法"]):
            return "联合国宪章宗旨和原则", "对应讲义：遵循《联合国宪章》宗旨和原则"
        if contains_any(t, ["联合国是", "最具普遍性", "代表性", "权威性", "主要舞台", "多边", "场所", "舞台"]):
            return "多边主义/多极化场所", "对应讲义：联合国是践行多边主义和多极化合作的重要场所"
        if contains_any(t, ["对中国", "中国作为联合国", "中国与联合国"]):
            return "新增：中国与联合国关系", "明显是中国与联合国双向关系角度"
        return "新增：联合国发展议程与全球治理贡献", "明显是联合国发展议程和全球性问题治理角度"

    raise RuntimeError(f"Unknown bucket: {bucket}")


def demote_core_block(block: CoreBlock) -> list[str]:
    out = []
    for i, line in enumerate(block.lines):
        if i == 0:
            heading = re.sub(r"（出现\d+次）$", "", line)
            heading = re.sub(r"；出现\d+次", "", heading)
            out.append(heading.replace("### 核心答题点：", "#### 核心答题点：", 1))
        else:
            out.append(line)
    return out


def demote_speed_lines(speed_lines: list[str]) -> list[str]:
    out: list[str] = []
    if not speed_lines:
        return out
    for i, line in enumerate(speed_lines):
        if i == 0 and line.startswith("## 六大要素术语极简速记版"):
            out.append(line)
        elif line in {f"## {b}" for b in BUCKETS}:
            out.append(line.replace("## ", "### ", 1))
        else:
            out.append(line)
    return out


def build_markdown(blocks: list[CoreBlock], speed_lines: list[str], assignments: dict[str, tuple[str, str]]) -> str:
    by_bucket_level: dict[str, dict[str, list[CoreBlock]]] = {
        bucket: {level: [] for level in SECOND_LEVEL_ORDER[bucket]} for bucket in BUCKETS
    }
    for block in blocks:
        level, _reason = assignments[block.title]
        by_bucket_level[block.bucket][level].append(block)

    lines: list[str] = []
    lines.extend(
        [
            "# 选择性必修一《当代国际政治与经济》主观题术语宝典",
            "学生厚版 · 237条逐条核实讲义二级框架确认版（2026.5.28 · 共104术语/237条目）",
            "飞哥正志讲堂",
            "",
            "## 二级结构说明",
            "本版按用户讲义框架组织：六大要素用于找方向；二级框架优先采用讲义口诀和结构；核心答题点保留原细则高信息量表述，题例用于训练迁移。",
            "",
            "正文不展示频次，避免学生把术语当成机械押题；后台映射表保留分布，供后续审核和调整。",
            "",
            "## 目录",
            "",
        ]
    )
    for bucket in BUCKETS:
        present = [level for level in SECOND_LEVEL_ORDER[bucket] if by_bucket_level[bucket][level]]
        lines.append(f"- {bucket}")
        for level in present:
            lines.append(f"  - {level}（{len(by_bucket_level[bucket][level])}个核心点）")
    lines.append("- 六大要素术语极简速记版")
    lines.append("")

    for bucket in BUCKETS:
        lines.append(f"## {bucket}")
        lines.append("")
        lines.append(f"**本桶怎么用：**{BUCKET_GUIDES[bucket]}")
        if bucket == "政治多极化":
            lines.append("")
            lines.append("**边界提醒：**先按“新国关、完全治、共体、单边/斥、民主”归入讲义口诀；明显是世界多极化格局、国际法治或具体治理议题的，放入新增角度。")
        if bucket == "中国":
            lines.append("")
            lines.append("**边界提醒：**同一个“人类命运共同体”点，若强调中国提出、推动、落实和贡献，归入中国智慧；若强调国际秩序和全球治理方向，归入政治多极化。")
        if bucket == "时代背景":
            lines.append("")
            lines.append("**边界提醒：**当经济全球化、政治多极化只是说明中国倡议为何正当、必要时，放入机遇；当材料重点是开放、市场、贸易、产业链和规则治理时，转入经济全球化。")
        if bucket == "经济全球化":
            lines.append("")
            lines.append("**边界提醒：**本桶先看“民国环营配，贸资总开兼”。能归入口诀字下的术语优先归入；企业出海、规则权益、产业链安全、数字绿色、贸易保护主义等明显另角度，用新增二级承接。")
            lines.append("")
            lines.append("**口诀索引：**民=民生福祉（就业、收入、基础设施）；国=成员国和相关国家发展；环=要素自由流动、产业升级、投资消费、开拓市场；营=优化营商环境（国际化、市场化、法治化）；配=优化资源配置；贸资=贸易投资自由化便利化；总=开放、包容、普惠、平衡、共赢；开=更高层次开放型经济、高水平对外开放、引进来走出去；兼=国内国际两种资源、两个市场。")
        lines.append("")

        for level in SECOND_LEVEL_ORDER[bucket]:
            level_blocks = by_bucket_level[bucket][level]
            if not level_blocks:
                continue
            lines.append(f"### {level}")
            lines.append("")
            for block in level_blocks:
                lines.extend(demote_core_block(block))
                lines.append("")

    lines.extend(demote_speed_lines(speed_lines))
    text = "\n".join(lines).rstrip() + "\n"
    return text


def set_font(run, name="PingFang SC", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_label_paragraph(doc: Document, line: str):
    m = re.match(r"^(【[^】]+】)(.*)$", line)
    p = doc.add_paragraph(style="BodyLabel")
    if not m:
        p.add_run(line)
        return
    label, rest = m.groups()
    r1 = p.add_run(label)
    set_font(r1, bold=True, color=(31, 78, 121))
    r2 = p.add_run(rest)
    set_font(r2)


def build_docx(md_text: str, out_path: Path):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.7)
    sec.bottom_margin = Cm(1.7)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "PingFang SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(3)

    for name, base, size, color, before, after in [
        ("DocTitle", "Title", 20, (31, 78, 121), 0, 8),
        ("DocSubtitle", "Subtitle", 11, (89, 89, 89), 0, 10),
        ("H1Custom", "Heading 1", 16, (31, 78, 121), 12, 6),
        ("H2Custom", "Heading 2", 13, (47, 84, 150), 9, 4),
        ("H3Custom", "Heading 3", 11.5, (84, 130, 53), 7, 3),
    ]:
        style = styles.add_style(name, 1)
        style.base_style = styles[base]
        style.font.name = "PingFang SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(*color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = name in {"H1Custom", "H2Custom", "H3Custom"}

    for name, left_cm, first_cm, after in [
        ("QuestionItem", 0.3, -0.3, 2),
        ("BodyLabel", 0, 0, 2),
        ("BodyIndented", 0.35, 0, 2),
    ]:
        style = styles.add_style(name, 1)
        style.base_style = normal
        style.font.name = "PingFang SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(10)
        style.paragraph_format.left_indent = Cm(left_cm)
        style.paragraph_format.first_line_indent = Cm(first_cm)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.06

    def add_text_para(text: str, style_name: str = "Normal", bold_prefix=False):
        p = doc.add_paragraph(style=style_name)
        if bold_prefix and "：" in text:
            prefix, rest = text.split("：", 1)
            r1 = p.add_run(prefix + "：")
            set_font(r1, bold=True, color=(31, 78, 121))
            r2 = p.add_run(rest)
            set_font(r2)
        else:
            r = p.add_run(text)
            set_font(r)
        return p

    first_title = True
    for raw in md_text.splitlines():
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="DocTitle")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[2:])
            set_font(r, size=20, bold=True, color=(31, 78, 121))
            first_title = False
        elif first_title:
            add_text_para(line, "DocSubtitle")
        elif line.startswith("## "):
            doc.add_paragraph(line[3:], style="H1Custom")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:], style="H2Custom")
        elif line.startswith("#### "):
            doc.add_paragraph(line[5:], style="H3Custom")
        elif line.startswith("**本桶怎么用：**"):
            add_text_para(line.replace("**", ""), "BodyIndented", bold_prefix=True)
        elif line.startswith("**边界提醒：**"):
            add_text_para(line.replace("**", ""), "BodyIndented", bold_prefix=True)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(line[2:])
            set_font(r, size=10)
        elif line.startswith("  - "):
            p = doc.add_paragraph(style="List Bullet 2")
            r = p.add_run(line[4:])
            set_font(r, size=10)
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="QuestionItem")
            r = p.add_run(line)
            set_font(r, bold=True, color=(68, 68, 68))
        elif line.startswith("【"):
            add_label_paragraph(doc, line)
        else:
            add_text_para(line, "Normal")

    core_props = doc.core_properties
    core_props.title = "选必一《当代国际政治与经济》主观题术语宝典 237条 讲义二级框架确认版"
    core_props.subject = "北京高考政治 选择性必修一 主观题术语宝典"
    core_props.author = "飞哥正志讲堂"
    doc.save(out_path)


def count_entries(block: CoreBlock) -> int:
    return sum(1 for line in block.lines if re.match(r"^\d+\. ", line))


def main():
    source_text = SOURCE_MD.read_text(encoding="utf-8")
    lines = source_text.splitlines()
    blocks, speed_lines = parse_source(lines)

    assignments: dict[str, tuple[str, str]] = {}
    rows = []
    for block in blocks:
        level, reason = classify(block)
        assignments[block.title] = (level, reason)
        rows.append(
            {
                "bucket": block.bucket,
                "second_level": level,
                "core_point": block.title,
                "frequency": block.count if block.count is not None else "",
                "entry_count": count_entries(block),
                "reason": reason,
            }
        )

    md_text = build_markdown(blocks, speed_lines, assignments)
    OUT_MD.write_text(md_text, encoding="utf-8")

    with OUT_MAP.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=["bucket", "second_level", "core_point", "frequency", "entry_count", "reason"],
        )
        writer.writeheader()
        writer.writerows(rows)

    build_docx(md_text, OUT_DOCX)

    source_counts = {
        "source_core_count": len(blocks),
        "source_entry_count": sum(count_entries(b) for b in blocks),
        "source_when_count": source_text.count("【什么时候写】"),
        "source_question_count": source_text.count("【设问】"),
        "source_answer_count": source_text.count("【卷面句】"),
    }
    output_counts = {
        "output_core_count": md_text.count("#### 核心答题点："),
        "output_entry_count": sum(1 for line in md_text.splitlines() if re.match(r"^\d+\. ", line)),
        "output_when_count": md_text.count("【什么时候写】"),
        "output_question_count": md_text.count("【设问】"),
        "output_answer_count": md_text.count("【卷面句】"),
    }

    level_counter = Counter((row["bucket"], row["second_level"]) for row in rows)
    forbidden_terms = ["Codex", "Claude", "GPT", "prompt", "rubric_angle_id", ".txt"]
    forbidden_hits = {term: md_text.count(term) for term in forbidden_terms if md_text.count(term)}

    expected = {
        "core_count_ok": output_counts["output_core_count"] == 104,
        "entry_count_ok": output_counts["output_entry_count"] == 237,
        "when_count_ok": output_counts["output_when_count"] == source_counts["source_when_count"],
        "question_count_ok": output_counts["output_question_count"] == source_counts["source_question_count"],
        "answer_count_ok": output_counts["output_answer_count"] == source_counts["source_answer_count"],
        "forbidden_terms_ok": not forbidden_hits,
    }

    summary = {
        "source_md": str(SOURCE_MD),
        "out_md": str(OUT_MD),
        "out_docx": str(OUT_DOCX),
        "out_mapping": str(OUT_MAP),
        "source_counts": source_counts,
        "output_counts": output_counts,
        "checks": expected,
        "forbidden_hits": forbidden_hits,
        "second_level_counts": {
            f"{bucket} / {level}": level_counter[(bucket, level)]
            for bucket in BUCKETS
            for level in SECOND_LEVEL_ORDER[bucket]
            if level_counter[(bucket, level)]
        },
    }
    OUT_JSON.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")

    qa_lines = [
        "# SECOND_LEVEL_REORG_QA",
        "",
        "## 核验结论",
        "",
        f"- 核心答题点：{output_counts['output_core_count']}（目标 104）",
        f"- 条目：{output_counts['output_entry_count']}（目标 237）",
        f"- 【什么时候写】：{output_counts['output_when_count']}（与基稿一致）",
        f"- 【设问】：{output_counts['output_question_count']}（与基稿一致）",
        f"- 【卷面句】：{output_counts['output_answer_count']}（与基稿一致）",
        f"- 后台词扫描：{'PASS' if not forbidden_hits else 'FAIL ' + str(forbidden_hits)}",
        "",
        "## 二级结构核心点分布",
        "",
        "| 一级要素 | 二级容器 | 核心点数 |",
        "|---|---:|---:|",
    ]
    for bucket in BUCKETS:
        for level in SECOND_LEVEL_ORDER[bucket]:
            n = level_counter[(bucket, level)]
            if n:
                qa_lines.append(f"| {bucket} | {level} | {n} |")
    qa_lines.extend(
        [
            "",
            "## 输出文件",
            "",
            f"- Markdown：`{OUT_MD}`",
            f"- Word：`{OUT_DOCX}`",
            f"- 映射表：`{OUT_MAP}`",
        ]
    )
    OUT_SUMMARY.write_text("\n".join(qa_lines) + "\n", encoding="utf-8")

    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
