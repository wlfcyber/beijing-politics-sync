#!/usr/bin/env python3
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document


RUN_DIR = Path(
    "/Users/wanglifei/GaokaoPolitics/beijing-politics-sync/reports/"
    "选必一_哲学宝典式重建_2026-05-16/30_make_final_20260601"
)
WORK_DOCX = RUN_DIR / "01_work/选必一_当代国际政治与经济_主观题术语宝典_最终版_工作稿_20260601.docx"
OUT_DOCX = RUN_DIR / "03_outputs/选必一_当代国际政治与经济_主观题术语宝典_最终版_20260601.docx"
DESKTOP_DOCX = Path("/Users/wanglifei/Desktop/选必一_当代国际政治与经济_主观题术语宝典_最终版_20260601.docx")


OLD_DONGCHENG = {
    "【材料触发点】 题面写龙头企业联合上下游和科研机构共建研发中心、强链补链，说明产业链韧性反过来支撑开放，而不是开放的附属结果。":
    "【材料触发点】 题面以“两区”建设为抓手，推动制度型开放，支持企业参与国际标准制定，并通过世界级产业大会、国际科创枢纽联络处开展国际科创交流，说明高水平开放本身为产业链韧性提升提供制度和资源支撑。",
    "【为什么能想到】 材料写\"龙头企业联合上下游和科研机构共建研发中心\"\"强链补链\"——这是补内功的产业链动作，而非直接的开放举措。问\"如何统筹\"——本条需要说明的是\"韧性反向支撑开放\"这一方向，避免把所有材料都套到\"开放促进韧性\"一头。先把龙头企业、上下游、科研机构共建研发中心作为强链补链抓手点出，再写产业链稳定性与竞争力提升，最后指向为高水平对外开放筑牢安全屏障。设问问“如何统筹高水平对外开放与产业链韧性”，本条挂在“扩大高水平对外开放与制度型开放”，就要落到开放举措本身：把坚持引进来与走出去相结合、对接国际高标准经贸规则、以“两区”建设为抓手推进制度型开放写出，而不是只写强链补链的产业链动作。":
    "【为什么能想到】 材料先写“两区”建设、国际标准制定、世界级产业大会和国际科创枢纽联络处，这些是开放举措层面的证据。问“如何统筹高水平对外开放与产业链韧性”，本条归入“中国扩大高水平对外开放与制度型开放”，应抓住开放促进韧性的方向：先写引进来与走出去相结合，再写对接国际高标准经贸规则、稳步扩大制度型开放和开展国际科创交流，最后指向为产业链韧性提升提供制度支撑和全球创新资源。",
    "【答案落点】 北京坚持引进来与走出去相结合，对接国际高标准经贸规则，以“两区”建设为抓手推进制度型开放，支持企业参与国际标准制定，扩大高水平对外开放。":
    "【答案落点】 北京以“两区”建设为抓手，坚持引进来与走出去相结合，对接国际高标准经贸规则，支持企业参与国际标准制定，开展国际科创交流合作，稳步扩大制度型开放，为产业链韧性提升提供制度支撑和全球创新资源。",
}


def replace_in_runs(paragraph, old: str, new: str) -> int:
    count = 0
    for run in paragraph.runs:
        if old in run.text:
            count += run.text.count(old)
            run.text = run.text.replace(old, new)
    return count


def main() -> int:
    doc = Document(str(WORK_DOCX))

    fengtai_replacements = 0
    dongcheng_replacements = 0
    for paragraph in doc.paragraphs:
        fengtai_replacements += replace_in_runs(
            paragraph,
            "2026丰台期末Q20",
            "2026丰台期末细则复练Q20",
        )
        text = paragraph.text
        if text in OLD_DONGCHENG:
            style = paragraph.style
            paragraph.clear()
            paragraph.style = style
            paragraph.add_run(OLD_DONGCHENG[text])
            dongcheng_replacements += 1

    if fengtai_replacements < 11:
        raise RuntimeError(f"expected at least 11 Fengtai replacements, got {fengtai_replacements}")
    if dongcheng_replacements != 3:
        raise RuntimeError(f"expected 3 Dongcheng replacements, got {dongcheng_replacements}")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    shutil.copy2(OUT_DOCX, DESKTOP_DOCX)
    print(f"saved={OUT_DOCX}")
    print(f"desktop={DESKTOP_DOCX}")
    print(f"fengtai_replacements={fengtai_replacements}")
    print(f"dongcheng_replacements={dongcheng_replacements}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
