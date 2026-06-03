#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import csv
import json
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


RUN_DIR = Path("/Users/wanglifei/GaokaoPolitics/beijing-politics-sync/reports/选必一_哲学宝典式重建_2026-05-16/26_ultimate_handbook_rebuild_20260531")
MATRIX = RUN_DIR / "02_corrections" / "REVISION_MATRIX.csv"
FRAMEWORK = RUN_DIR / "01_framework" / "user_confirmed_framework.json"
STYLE_SOURCE = Path("/Users/wanglifei/Desktop/桌面整理_20260529/政治成品文档/哲学宝典最终版-飞哥正志讲堂.docx")

STUDENT_MD = RUN_DIR / "03_draft" / "选必一_当代国际政治与经济_主观题术语宝典_终极版内容稿_20260531.md"
QA_MD = RUN_DIR / "05_qa" / "选必一_Claude修订稿_逐条回源核查结论_20260531.md"
DOCX_OUT = RUN_DIR / "04_docx" / "选必一_当代国际政治与经济_主观题术语宝典_终极版_20260531.docx"
DESKTOP_COPY = Path("/Users/wanglifei/Desktop/选必一_当代国际政治与经济_主观题术语宝典_终极版_20260531.docx")


def read_csv(path):
    with path.open(encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def clean(text):
    return " ".join((text or "").replace("\r", "\n").split())


def by_key(rows):
    out = defaultdict(list)
    for row in rows:
        out[(row["bucket"], row["second_level"], row["core_point"])].append(row)
    for key in out:
        out[key].sort(key=lambda r: int(r["global_entry_no"]))
    return out


def status_label(status):
    mapping = {
        "ISSUE_FIXED_DATA": "已按原卷/细则修正",
        "SOURCE_LOCATED_PENDING_HUMAN_PASS": "已定位原卷/细则",
        "NEEDS_MANUAL_SURFACE_CHECK": "已定位，答案落点为归纳表达",
        "SOURCE_LOCATED_ANSWER_ONLY": "教师版答案依据，未见独立正式细则",
    }
    return mapping.get(status, status)


def evidence_text(row):
    score = clean(row["corrected_score_layers"])
    if score:
        return score
    if row["revision_status"] == "SOURCE_LOCATED_ANSWER_ONLY":
        return "教师版试卷/答案已定位；桌面目录未发现独立正式评分细则，本条不伪造细则分层。"
    if clean(row["rubric_paths"]):
        return "原细则/参考答案已定位；本条按对应答案落点归纳，原细则未单列本术语分值层次。"
    return "原题已定位；未发现可独立摘录的正式细则分层。"


def format_entry_md(row, idx):
    lines = []
    title = f"{idx}. {clean(row['source_label'])}"
    lines.append(f"#### {title}")
    fields = [
        ("来源", clean(row["source_label"])),
        ("术语", clean(row["core_point"])),
        ("材料触发点", clean(row["when_to_write"])),
        ("设问", clean(row["corrected_prompt"])),
        ("细则依据", evidence_text(row)),
        ("为什么能想到", clean(row["why_signal"])),
        ("答案落点", clean(row["surface_sentence"])),
    ]
    for label, value in fields:
        if value:
            lines.append(f"【{label}】{value}")
    same = clean(row["same_group"])
    if same:
        lines.append(f"【同题组】{same}")
    lines.append("")
    return "\n".join(lines)


def generate_markdown(rows, framework):
    grouped = by_key(rows)
    used = set()
    lines = [
        "# 选必一《当代国际政治与经济》主观题术语宝典（终极版内容稿）",
        "",
        "说明：本稿按用户确认的六要素与二级框架重排。设问和细则依据以桌面原试卷、参考答案、评分细则为准；答案落点是面向学生的可写表达，不冒充原卷逐字句。",
        "",
    ]
    entry_counter = 1
    current_h1 = None
    current_h2 = None
    for item in framework:
        level = item["level"]
        bucket = item["bucket"]
        heading = item["heading"]
        if level == 1:
            current_h1 = heading
            current_h2 = None
            lines.append(f"## {heading}")
            lines.append("")
        elif level == 2:
            current_h2 = heading
            lines.append(f"### {heading}")
            lines.append("")
        elif level == 3:
            key = (bucket, current_h2, heading)
            entries = grouped.get(key, [])
            if not entries:
                continue
            lines.append(f"#### 核心术语：{heading}")
            lines.append("")
            for row in entries:
                lines.append(format_entry_md(row, entry_counter))
                used.add(row["global_entry_no"])
                entry_counter += 1
    remaining = [r for r in rows if r["global_entry_no"] not in used]
    if remaining:
        lines.append("## 其他已核条目")
        lines.append("")
        for row in sorted(remaining, key=lambda r: int(r["global_entry_no"])):
            lines.append(f"### {clean(row['bucket'])} / {clean(row['second_level'])} / {clean(row['core_point'])}")
            lines.append("")
            lines.append(format_entry_md(row, entry_counter))
            entry_counter += 1
    STUDENT_MD.parent.mkdir(parents=True, exist_ok=True)
    STUDENT_MD.write_text("\n".join(lines), encoding="utf-8")


def generate_qa(rows):
    counts = defaultdict(int)
    for row in rows:
        counts[row["revision_status"]] += 1
    lines = [
        "# 选必一 Claude 修订稿逐条回源核查结论",
        "",
        "## 总结",
        "",
        f"- 条目总数：{len(rows)}",
    ]
    for status in sorted(counts):
        lines.append(f"- {status_label(status)}：{counts[status]} 条")
    lines.extend([
        "",
        "## 已修正或降级说明",
        "",
    ])
    for row in rows:
        if row["revision_status"] in {"ISSUE_FIXED_DATA", "SOURCE_LOCATED_ANSWER_ONLY"}:
            lines.append(f"- {row['global_entry_no']}｜{clean(row['source_label'])}｜{status_label(row['revision_status'])}")
            if clean(row["fix_note"]):
                lines.append(f"  - 处理：{clean(row['fix_note'])}")
            if clean(row["evidence"]):
                lines.append(f"  - 依据：{clean(row['evidence'])}")
            if row["revision_status"] == "SOURCE_LOCATED_ANSWER_ONLY":
                lines.append("  - 残余：未在桌面目录找到独立正式评分细则。")
    lines.extend([
        "",
        "## 答案落点归纳表达",
        "",
        "以下条目已定位原卷/细则，但“答案落点”是面向学生的归纳表达，不作为原卷逐字引文。",
        "",
    ])
    for row in rows:
        if row["revision_status"] == "NEEDS_MANUAL_SURFACE_CHECK":
            lines.append(f"- {row['global_entry_no']}｜{clean(row['source_label'])}｜{clean(row['core_point'])}")
    QA_MD.parent.mkdir(parents=True, exist_ok=True)
    QA_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def set_style(style, font_name=None, size=None, bold=None, color=None, space_before=None, space_after=None):
    font = style.font
    if font_name:
        font.name = font_name
        font.element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    if size:
        font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if color:
        font.color.rgb = RGBColor.from_string(color)
    fmt = style.paragraph_format
    if space_before is not None:
        fmt.space_before = Pt(space_before)
    if space_after is not None:
        fmt.space_after = Pt(space_after)


def add_label_paragraph(doc, label, value):
    if not value:
        return
    p = doc.add_paragraph(style="Normal")
    r = p.add_run(f"【{label}】")
    r.bold = True
    p.add_run(value)


def generate_docx(rows, framework):
    doc = Document()
    section = doc.sections[0]
    section.page_width = 7560310
    section.page_height = 10692130
    section.top_margin = 756285
    section.bottom_margin = 720090
    section.left_margin = 774065
    section.right_margin = 774065

    set_style(doc.styles["Normal"], "Arial", 10.5, False, None, None, 10)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.15
    set_style(doc.styles["Heading 1"], "Arial", 18, True, "1F4E79", 24, 6)
    set_style(doc.styles["Heading 2"], "Arial", 14, True, "2F6F9F", 10, 6)
    set_style(doc.styles["Heading 3"], "Arial", 11.5, True, "3A6278", 10, 4)

    title = doc.add_paragraph(style="Title")
    title.add_run("选必一《当代国际政治与经济》主观题术语宝典").bold = True
    subtitle = doc.add_paragraph(style="Normal")
    subtitle.add_run("终极版内容稿｜按用户二级框架重排｜2026-05-31")
    note = doc.add_paragraph(style="Normal")
    note.add_run("说明：设问与细则依据以桌面原试卷、参考答案、评分细则为准；“答案落点”为学生可写表达。")
    doc.add_paragraph("目录", style="Heading 1")
    for item in framework:
        if item["level"] == 1:
            p = doc.add_paragraph(style="Normal")
            p.add_run(item["heading"]).bold = True
        elif item["level"] == 2:
            doc.add_paragraph(f"  {item['heading']}", style="Normal")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    grouped = by_key(rows)
    used = set()
    entry_counter = 1
    current_h2 = None
    for item in framework:
        level = item["level"]
        bucket = item["bucket"]
        heading = item["heading"]
        if level == 1:
            doc.add_paragraph(heading, style="Heading 1")
            current_h2 = None
        elif level == 2:
            current_h2 = heading
            doc.add_paragraph(heading, style="Heading 2")
        elif level == 3:
            entries = grouped.get((bucket, current_h2, heading), [])
            if not entries:
                continue
            doc.add_paragraph(f"核心术语：{heading}", style="Heading 3")
            for row in entries:
                p = doc.add_paragraph(style="Normal")
                p.add_run(f"{entry_counter}. {clean(row['source_label'])}").bold = True
                add_label_paragraph(doc, "来源", clean(row["source_label"]))
                add_label_paragraph(doc, "术语", clean(row["core_point"]))
                add_label_paragraph(doc, "材料触发点", clean(row["when_to_write"]))
                add_label_paragraph(doc, "设问", clean(row["corrected_prompt"]))
                add_label_paragraph(doc, "细则依据", evidence_text(row))
                add_label_paragraph(doc, "为什么能想到", clean(row["why_signal"]))
                add_label_paragraph(doc, "答案落点", clean(row["surface_sentence"]))
                add_label_paragraph(doc, "同题组", clean(row["same_group"]))
                used.add(row["global_entry_no"])
                entry_counter += 1
    remaining = [r for r in rows if r["global_entry_no"] not in used]
    if remaining:
        doc.add_paragraph("其他已核条目", style="Heading 1")
        for row in sorted(remaining, key=lambda r: int(r["global_entry_no"])):
            doc.add_paragraph(f"{clean(row['bucket'])} / {clean(row['second_level'])} / {clean(row['core_point'])}", style="Heading 2")
            p = doc.add_paragraph(style="Normal")
            p.add_run(f"{entry_counter}. {clean(row['source_label'])}").bold = True
            add_label_paragraph(doc, "来源", clean(row["source_label"]))
            add_label_paragraph(doc, "术语", clean(row["core_point"]))
            add_label_paragraph(doc, "材料触发点", clean(row["when_to_write"]))
            add_label_paragraph(doc, "设问", clean(row["corrected_prompt"]))
            add_label_paragraph(doc, "细则依据", evidence_text(row))
            add_label_paragraph(doc, "为什么能想到", clean(row["why_signal"]))
            add_label_paragraph(doc, "答案落点", clean(row["surface_sentence"]))
            add_label_paragraph(doc, "同题组", clean(row["same_group"]))
            entry_counter += 1
    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)
    doc.save(DESKTOP_COPY)


def main():
    rows = read_csv(MATRIX)
    framework = json.loads(FRAMEWORK.read_text(encoding="utf-8"))
    generate_markdown(rows, framework)
    generate_qa(rows)
    generate_docx(rows, framework)
    print(json.dumps({
        "student_md": str(STUDENT_MD),
        "qa_md": str(QA_MD),
        "docx": str(DOCX_OUT),
        "desktop_copy": str(DESKTOP_COPY),
        "entries": len(rows),
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
