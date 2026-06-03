#!/usr/bin/env python3
import csv
import hashlib
import json
import os
import re
import subprocess
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from pypdf import PdfReader


ROOTS = [
    Path("/Users/wanglifei/Desktop/2024模拟题"),
    Path("/Users/wanglifei/Desktop/2025模拟题"),
    Path("/Users/wanglifei/Desktop/2026模拟题"),
]
DISTRICTS = ["朝阳", "海淀", "丰台", "西城", "东城", "顺义", "石景山", "延庆", "门头沟", "房山", "通州", "昌平"]
EXAM_TYPES = ["期中", "期末", "一模", "二模", "三模"]
VALID_EXTS = {".docx", ".doc", ".pdf", ".pptx", ".rtf"}


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def stable_id(path: Path) -> str:
    return hashlib.sha1(str(path).encode("utf-8")).hexdigest()[:16]


def classify(path: Path):
    parts = path.parts
    s = str(path)
    year = ""
    for y in ["2024", "2025", "2026"]:
        if y in s:
            year = y
            break
    district = ""
    exam_type = ""
    # Prefer the nearest directory name like "2026通州期末" over parent
    # grouping names such as "2026各区期末和期中", which contain two exam types.
    for part in reversed(parts):
        part_district = next((d for d in DISTRICTS if d in part), "")
        part_exam_type = next((t for t in EXAM_TYPES if t in part), "")
        if part_district and part_exam_type:
            district = part_district
            exam_type = part_exam_type
            break
    if not district:
        district = next((d for d in DISTRICTS if d in s), "")
    if not exam_type:
        exam_type = next((t for t in EXAM_TYPES if t in s), "")
    if "/试卷/" in s:
        kind = "试卷"
    elif "/细则/" in s:
        kind = "细则"
    elif "/答案/" in s:
        kind = "答案"
    else:
        kind = "其他"
    return year, district, exam_type, kind


def extract_docx(path: Path) -> str:
    doc = Document(path)
    chunks = []
    for para in doc.paragraphs:
        t = clean(para.text)
        if t:
            chunks.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [clean(cell.text) for cell in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                chunks.append(line)
    return "\n".join(chunks)


def extract_office_xml(path: Path) -> str:
    chunks = []
    with zipfile.ZipFile(path) as zf:
        names = [n for n in zf.namelist() if n.startswith(("ppt/slides/", "ppt/notesSlides/")) and n.endswith(".xml")]
        names.sort()
        for name in names:
            try:
                root = ET.fromstring(zf.read(name))
            except Exception:
                continue
            texts = []
            for elem in root.iter():
                if elem.tag.endswith("}t") and elem.text:
                    texts.append(elem.text)
            line = clean(" ".join(texts))
            if line:
                chunks.append(f"[{name}] {line}")
    return "\n".join(chunks)


def extract_doc_or_rtf(path: Path) -> str:
    proc = subprocess.run(
        ["/usr/bin/textutil", "-convert", "txt", "-stdout", str(path)],
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        errors="replace",
        timeout=60,
    )
    if proc.returncode != 0:
        raise RuntimeError(proc.stderr.strip()[:500])
    return "\n".join(clean(line) for line in proc.stdout.splitlines() if clean(line))


def extract_pdf(path: Path):
    reader = PdfReader(str(path))
    chunks = []
    pages = 0
    for i, page in enumerate(reader.pages, 1):
        pages += 1
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        text = clean(text)
        if text:
            chunks.append(f"[page {i}] {text}")
    return "\n".join(chunks), pages


def extract_text(path: Path):
    suffix = path.suffix.lower()
    pages = ""
    if suffix == ".docx":
        return extract_docx(path), pages
    if suffix == ".pptx":
        return extract_office_xml(path), pages
    if suffix in {".doc", ".rtf"}:
        return extract_doc_or_rtf(path), pages
    if suffix == ".pdf":
        text, pages = extract_pdf(path)
        return text, pages
    return "", pages


def main() -> int:
    if len(sys.argv) != 2:
        print("usage: build_source_text_cache.py OUTDIR", file=sys.stderr)
        return 2
    out_dir = Path(sys.argv[1])
    text_dir = out_dir / "texts"
    out_dir.mkdir(parents=True, exist_ok=True)
    text_dir.mkdir(parents=True, exist_ok=True)

    files = []
    for root in ROOTS:
        for path in root.rglob("*"):
            if not path.is_file():
                continue
            if path.name.startswith("~$") or path.name == ".DS_Store":
                continue
            if path.suffix.lower() not in VALID_EXTS:
                continue
            files.append(path)
    files.sort(key=lambda p: str(p))

    rows = []
    for path in files:
        fid = stable_id(path)
        text_path = text_dir / f"{fid}.txt"
        year, district, exam_type, kind = classify(path)
        status = "OK"
        error = ""
        text = ""
        pages = ""
        try:
            text, pages = extract_text(path)
        except Exception as exc:
            status = "EXTRACT_ERROR"
            error = f"{exc.__class__.__name__}: {str(exc)[:400]}"
        text_len = len(text)
        if status == "OK" and path.suffix.lower() == ".pdf" and text_len < 80:
            status = "NEEDS_OCR"
        if status == "OK" and text_len == 0:
            status = "EMPTY_TEXT"
        text_path.write_text(text, encoding="utf-8")
        rows.append(
            {
                "file_id": fid,
                "path": str(path),
                "year": year,
                "district": district,
                "exam_type": exam_type,
                "kind": kind,
                "suffix": path.suffix.lower(),
                "pages": pages,
                "text_len": text_len,
                "status": status,
                "error": error,
                "text_path": str(text_path),
            }
        )

    fieldnames = ["file_id", "path", "year", "district", "exam_type", "kind", "suffix", "pages", "text_len", "status", "error", "text_path"]
    with (out_dir / "SOURCE_FILE_INVENTORY.csv").open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)
    summary = {
        "files": len(rows),
        "by_status": {},
        "by_suffix": {},
        "by_kind": {},
    }
    for row in rows:
        for key, target in [("status", "by_status"), ("suffix", "by_suffix"), ("kind", "by_kind")]:
            summary[target][row[key]] = summary[target].get(row[key], 0) + 1
    (out_dir / "source_extract_summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
