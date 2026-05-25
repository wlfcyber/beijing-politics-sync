from __future__ import annotations

import csv
import json
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

import apply_p0_thickness_batch01_20260525 as helper
from docx import Document


ROOT = Path(__file__).resolve().parent

TARGETS = [
    {
        "queue_id": "T0182",
        "heading": "16. 2026房山一模 第18题第（1）问（主观题）",
        "node": "系统观念 / 系统优化",
        "old_answer_excerpt": "北京通过科学治理实现“常态蓝天”，首先在于坚持系统治理和系统优化，统筹四大结构调整",
        "evidence_note": "材料写“系统治理”，并列统筹产业、能源、交通、用地四大结构调整和多污染物协同控制。",
        "new_why": "材料直接使用“系统治理”，并把产业、能源、交通、用地四大结构调整同工业减排、机动车清洁化、扬尘精细化管控、散煤清零、多污染物协同控制放在一起。空气质量改善不是某一个污染源单独变化的结果，而是多个要素相互联系、协同优化形成整体效果，所以应落到系统观念和系统优化。",
        "new_answer": "①系统具有整体性、有序性和内部结构优化趋向，科学治理要用系统观念统筹各要素。②北京实现“常态蓝天”，不是只治理某一个污染源，而是统筹产业、能源、交通、用地四大结构调整，推进工业减排、机动车清洁化、扬尘精细化管控、散煤清零和多污染物协同控制。③把各治理环节组织成相互配合的系统，能够形成整体合力，提升空气治理的整体效能。",
    },
    {
        "queue_id": "T0374",
        "heading": "16. 2026海淀期末 第17题（主观题）",
        "node": "矛盾的普遍性和特殊性",
        "old_answer_excerpt": "红色文化把共同精神价值与具体历史场景、时代传播方式结合起来，在普遍性与特殊性的统一中增强年轻人的认同",
        "evidence_note": "正式评分标准page_002在Q17“红色文化的内涵”维度列出“矛盾的普遍性和特殊性的辩证关系”。",
        "new_why": "红色文化能“圈粉”年轻人，既因为它承载理想信念、爱国主义、奋斗精神等共同价值，又因为这些共同价值总是在井冈山、遵义、延安、红色场馆、沉浸体验和青年传播方式中具体呈现。细则明确列普遍性和特殊性的辩证关系，说明要写共同精神和具体表达的统一。",
        "new_answer": "①矛盾的普遍性和特殊性相互联结，普遍性寓于特殊性之中，并通过特殊性表现出来。②红色文化承载革命理想、爱国情怀、奋斗精神等共同精神价值，这是吸引青年的普遍内核；同时，它又通过具体历史场景、人物故事、场馆体验、文创表达和网络传播等特殊形式呈现。③共同价值借助鲜活具体的表达进入青年生活，才能增强代入感和认同感，使红色文化真正“圈粉”年轻人。",
    },
    {
        "queue_id": "T0094",
        "heading": "16. 2026通州期末 第16题（主观题）",
        "node": "规律的客观性",
        "old_answer_excerpt": "治水必须尊重水流、地势和生态规律，在规律允许的范围内进行工程设计和维护",
        "evidence_note": "正式细则第16题写明“要尊重客观规律”，并把顺势而为、因地制宜归入该角度。",
        "new_why": "都江堰的智慧不在于人任意支配自然，而在于顺应水流、地势和生态条件来设计工程。细则把“顺势而为、因地制宜”直接归入尊重客观规律，说明卷面要写人必须按水利运行规律办事，在规律允许范围内发挥工程设计和维护作用。",
        "new_answer": "①规律具有客观性，人们不能违背规律，只能在认识和把握规律的基础上利用规律。②都江堰治水不是凭主观愿望强行改造水势，而是顺应地势高低、水流走向、泥沙运动和灌溉防洪需要，因地制宜设计分水、泄洪、灌溉等工程环节。③正因为工程建设和历代维护都尊重客观规律，都江堰才能跨越千年持续发挥防洪、灌溉和生态效益。",
    },
    {
        "queue_id": "T0165",
        "heading": "17. 2026房山一模 第16题第（2）问（主观题）",
        "node": "整体与部分",
        "old_answer_excerpt": "卷面可写：要立足整体、统筹全局，把地方产业作为国家能源转型和区域协同中的部分来定位",
        "evidence_note": "评标细则高档要求写“融入大局中找定位/系统/整体”，并在细则提示中写“在大局中找准定位”。",
        "new_why": "题目谈因地制宜，但细则的高档要求不是只看本地资源，而是写“融入大局中找定位”“在大局中找准定位”。这说明地方发展要放进国家能源转型、区域协同和产业体系整体中理解：部分要服从整体，同时也要通过发挥局部优势促进整体功能提升。",
        "new_answer": "①整体居于主导地位，部分服从和服务于整体；部分的功能及其变化也会影响整体功能。②地方因地制宜发展产业，不能只算本地小账，而要把当地资源禀赋、产业基础和区位条件放到国家能源转型、区域协同和现代化建设大局中定位。③在整体布局中发挥清洁能源、口岸、农牧产品和文化资源等局部优势，既能推动本地高质量发展，也能服务国家和区域整体发展目标。",
    },
    {
        "queue_id": "T0493",
        "heading": "17. 2026西城期末 第16题第（2）问（主观题）",
        "node": "社会存在与社会意识",
        "old_answer_excerpt": "真实内容之所以能打动人，是因为它如实反映社会存在、社会矛盾和人民需求",
        "evidence_note": "正式评分细则page_002酌情原则列明“社会存在社会意识”，唯物史观角度要求如实反映社会矛盾和人民需求。",
        "new_why": "“真实”直抵人心，不能只解释为表达技巧好。细则把社会存在与社会意识列为可得分角度，并要求写如实反映社会矛盾和人民需求，说明真实内容之所以有力量，是因为它来源于真实社会生活，能够把人民关切、现实矛盾和时代问题反映出来。",
        "new_answer": "①社会存在决定社会意识，社会意识是对社会存在的反映。②真实内容能够打动人，是因为它不是脱离现实的虚构说教，而是如实反映社会生活、社会矛盾和人民群众的真实需求，使人们在其中看见现实处境和共同关切。③这样的社会意识能够回应社会存在中的问题，增强认同和共鸣，并对人们认识现实、参与实践产生积极影响。",
    },
    {
        "queue_id": "T0095",
        "heading": "17. 2026通州期末 第21题（主观题）",
        "node": "规律的客观性",
        "old_answer_excerpt": "谋划和推进中国式现代化要把握经济社会发展规律，根据实际形势制定和实施规划",
        "evidence_note": "正式细则第21题思维方法角度列明“尊重客观规律”。",
        "new_why": "题目问我国如何为中国式现代化建设谱写宏伟篇章，细则在思维方法角度列出“尊重客观规律”。这说明不能只写政策成果，还要写现代化建设必须根据我国发展阶段、国情实际和经济社会发展规律来谋划，按规律确定目标、路径和节奏。",
        "new_answer": "①规律具有客观性，推进中国式现代化必须尊重经济社会发展规律和现代化建设规律。②我国不是凭主观愿望制定发展目标，而是立足国情和发展阶段，分析实际形势、主要任务和历史方位，科学谋划战略安排和实施步骤。③在尊重规律基础上制定并落实规划，才能把宏伟目标转化为有序实践，推动中国式现代化建设不断取得新进展。",
    },
    {
        "queue_id": "T0438",
        "heading": "18. 2025丰台期末 第16题（主观题）",
        "node": "实践是认识的基础",
        "old_answer_excerpt": "要做到胸中有“数”，必须深入基层、调查研究，在实践中获取真实“数”，并以此指导工作和国家治理",
        "evidence_note": "材料强调“数”从哪里来，做计划、定方案、看大数都离不开关注实际、调查研究。",
        "new_why": "材料反复追问“数从哪里来、数当如何定、数为何所用”，并强调做计划、定方案、国家治理都要关注实际和调查研究。这里的“数”不是主观想象出来的，而是人在实践中接触实际、处理问题、总结经验形成的认识，因此应落到实践是认识的来源和基础。",
        "new_answer": "①实践是认识的来源和基础，正确认识必须从实践和实际中来。②胸中有“数”，不是坐在办公室凭感觉估算，而是要深入基层、调查研究，在生产生活和国家治理实践中掌握真实情况、数量关系和发展趋势。③只有从实践中取得真实可靠的“数”，再把这些认识用于做计划、定方案、看大数、看长远，才能提高决策和治理的科学性，避免胸中无数导致错误。",
    },
    {
        "queue_id": "T0166",
        "heading": "18. 2025朝阳一模 第16题（主观题）",
        "node": "整体与部分",
        "old_answer_excerpt": "部分影响整体，关键部分的功能及其变化甚至会对整体功能起重要作用。国产动画创作团队通过关键戏份打磨",
        "evidence_note": "细则第16题写创作者团队“以局部推动整体”，推动中华优秀传统文化以新形式展现、绽放新光彩。",
        "new_why": "国产动画让传统文化绽放新光彩，不是抽象地靠“文化自信”四个字，而是通过创作团队、关键戏份、人物形象、技术表达等具体部分带动整部作品的质量提升，并进一步影响中华优秀传统文化的整体传播效果。细则写“以局部推动整体”，正对应部分对整体的作用。",
        "new_answer": "①整体由部分构成，部分影响整体，关键部分的功能及其变化甚至会对整体功能起决定性作用。②国产动画创作者通过打磨关键戏份、塑造人物形象、改进视觉技术和创新叙事表达，使一部作品中的具体部分更好承载中华优秀传统文化。③这些局部创新提升了作品整体感染力和传播力，并带动传统文化以新的形式被看见、被理解、被喜爱，从而让中华优秀传统文化在新时代绽放新光彩。",
    },
]


def now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M +08")


def write_csv(path: Path, rows: list[dict[str, object]], fields: list[str]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fields})


def write_draft_files() -> None:
    rows = []
    for target in TARGETS:
        rows.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "evidence_note": target["evidence_note"],
                "new_why_chars": len(target["new_why"]),
                "new_answer_chars": len(target["new_answer"]),
                "new_why": target["new_why"],
                "new_answer": target["new_answer"],
            }
        )
    (ROOT / "P0_THICKNESS_REPAIR_BATCH04_DRAFT_20260525.json").write_text(
        json.dumps({"updated": now(), "targets": rows}, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / "P0_THICKNESS_REPAIR_BATCH04_DRAFT_20260525.csv",
        rows,
        ["queue_id", "heading", "node", "evidence_note", "new_why_chars", "new_answer_chars", "new_why", "new_answer"],
    )
    lines = [
        "# P0 Thickness Repair Batch04 Draft 20260525",
        "",
        f"Updated: {now()}",
        "",
        "Status: `DRAFT_READY_FOR_DOCX_APPLICATION`",
        "",
        "- Scope: next 8 P0 subjective triple-thin rows after Batch03 refresh.",
        "- Repair method: thicken only existing why/answer paragraphs inside existing rubric-supported nodes.",
        "- No new principle is introduced; evidence notes are copied from current node-level support.",
        "",
        "| queue_id | node | heading | new why chars | new answer chars | evidence note |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in rows:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | {row['new_why_chars']} | {row['new_answer_chars']} | {row['evidence_note']} |"
        )
    (ROOT / "P0_THICKNESS_REPAIR_BATCH04_DRAFT_20260525.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )


def main() -> int:
    write_draft_files()
    docx = helper.current_docx()
    backup = docx.with_name(f"{docx.stem}_backup_before_p0_thickness_batch04_{datetime.now():%Y%m%d_%H%M%S}.docx")
    shutil.copy2(docx, backup)

    doc = Document(str(docx))
    blocks = helper.entry_blocks(doc)
    applied: list[dict[str, object]] = []
    for target in TARGETS:
        block = helper.matching_block(blocks, target)
        fields = block["fields"]
        field_paras = block["field_paras"]
        before_why = str(fields.get("why", ""))
        before_answer = str(fields.get("answer", ""))
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["why"])], "why", target["new_why"])
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["answer"])], "answer", target["new_answer"])
        applied.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "heading_para": block["heading_para"],
                "why_para": field_paras["why"],
                "answer_para": field_paras["answer"],
                "before_why_chars": len(before_why),
                "after_why_chars": len(target["new_why"]),
                "before_answer_chars": len(before_answer),
                "after_answer_chars": len(target["new_answer"]),
                "evidence_note": target["evidence_note"],
                "old_answer_excerpt": target["old_answer_excerpt"],
            }
        )

    doc.save(str(docx))
    with zipfile.ZipFile(docx) as z:
        zip_ok = "[Content_Types].xml" in z.namelist() and "word/document.xml" in z.namelist()

    result = {
        "updated": now(),
        "status": "P0_BATCH04_APPLIED_REQUIRES_RENDER_AND_EXTERNAL_RECHECK",
        "docx": str(docx),
        "backup_docx": str(backup),
        "zip_ok": zip_ok,
        "targets": len(TARGETS),
        "applied": applied,
        "boundary": [
            "This is a local thickness repair batch, not final acceptance.",
            "PDF/render outputs must be regenerated after this DOCX edit.",
            "External review remains pending for the current post-repair artifact.",
        ],
    }
    (ROOT / "P0_THICKNESS_REPAIR_BATCH04_APPLY_20260525.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / "P0_THICKNESS_REPAIR_BATCH04_APPLY_20260525.csv",
        applied,
        [
            "queue_id",
            "heading",
            "node",
            "heading_para",
            "why_para",
            "answer_para",
            "before_why_chars",
            "after_why_chars",
            "before_answer_chars",
            "after_answer_chars",
            "evidence_note",
            "old_answer_excerpt",
        ],
    )
    lines = [
        "# P0 Thickness Repair Batch04 Apply 20260525",
        "",
        f"Updated: {result['updated']}",
        "",
        f"Status: `{result['status']}`",
        "",
        f"- DOCX: `{docx.name}`.",
        f"- Backup: `{backup.name}`.",
        f"- Targets applied: `{len(applied)}` / `{len(TARGETS)}`.",
        f"- DOCX zip structural check: `{str(zip_ok).lower()}`.",
        "",
        "## Applied Rows",
        "",
        "| queue_id | node | heading | why chars | answer chars | paragraph refs |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in applied:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | "
            f"{row['before_why_chars']} -> {row['after_why_chars']} | "
            f"{row['before_answer_chars']} -> {row['after_answer_chars']} | "
            f"h={row['heading_para']}; why={row['why_para']}; answer={row['answer_para']} |"
        )
    lines.extend(
        [
            "",
            "## Boundary",
            "",
            "- This batch repairs only 8 additional P0 rows.",
            "- Remaining thickness queue, render QA, and model gates stay open until rechecked.",
        ]
    )
    (ROOT / "P0_THICKNESS_REPAIR_BATCH04_APPLY_20260525.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
