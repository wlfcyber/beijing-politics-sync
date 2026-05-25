from __future__ import annotations

import csv
import json
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

import apply_p0_thickness_batch01_20260525 as helper
from docx import Document


ROOT = Path(__file__).resolve().parent

TARGETS = [
    {
        "queue_id": "T0256",
        "heading": "11. 2026房山一模 第18题第（1）问（主观题）",
        "node": "量变与质变 / 适度原则",
        "old_answer_excerpt": "北京实现“常态蓝天”，在于坚持长期治理、持续治理",
        "evidence_note": "材料和评标细则强调久久为功、日积月累、持续治理，支持量变积累引起空气质量质变。",
        "new_why": "题目问北京如何实现“常态蓝天”，材料不是只写一次性治理措施，而是反复强调“久久为功”“一茬接着一茬干”“日积月累的治理行动”。这些表述说明空气质量改善依赖长期、连续、渐进的量的积累；从雾霾困扰到优良天数接近95%，则体现量变达到一定程度后引起质的变化。",
        "new_answer": "①量变是质变的必要准备，生态环境治理不能指望一次行动立刻发生根本变化，而要靠长期持续积累。②北京坚持一茬接着一茬干，把减排、监测、执法、产业调整等治理行动长期推进，使空气质量改善的量逐步积累。③当这种积累达到一定程度，就推动城市空气质量从“雾霾困扰”转向“常态蓝天”，体现了以持续量变促成治理成效质变。",
    },
    {
        "queue_id": "T0402",
        "heading": "11. 2026房山一模 第18题第（1）问（主观题）",
        "node": "两点论与重点论",
        "old_answer_excerpt": "卷面可写：治理常态蓝天既要全面分析污染治理的多方面矛盾",
        "evidence_note": "评标细则明列矛盾分析法、具体问题具体分析、两点论与重点论统一，并写到聚焦重点领域、分区分类、一企一策。",
        "new_why": "材料和细则同时出现两组信号：一是污染治理涉及能源、产业、交通、区域协同等多方面因素，要求全面分析；二是“聚焦重点领域”“分区分类”“一企一策”强调抓住重点环节和主要矛盾。看到这种“全面治理+重点突破”的结构，就应落到两点论和重点论相统一。",
        "new_answer": "①治理常态蓝天要坚持两点论，全面分析污染成因、产业结构、交通排放、区域传输和群众生活等多方面矛盾，不能只盯单一因素。②又要坚持重点论，抓住重点区域、重点行业、重点时段和关键排放源，实行分区分类、一企一策、精准施策。③把全面统筹与重点突破结合起来，才能以关键环节带动整体治理提升，推动空气质量稳定改善。",
    },
    {
        "queue_id": "T0471",
        "heading": "11. 2026朝阳期中 第19题（主观题）",
        "node": "认识对实践的反作用",
        "old_answer_excerpt": "要形成对抗战精神的正确认识，并用这种认识指导新时代实践",
        "evidence_note": "正式评分细则把正确认识对实践的反作用列为弘扬抗战精神的哲学角度。",
        "new_why": "卷首语任务不是只复述抗战史实，而是说明怎样继承和弘扬伟大抗战精神。细则明确给出“正确认识对实践的反作用”，说明学生要把对抗战精神内涵的认识，转化为新时代坚定理想信念、担当强国建设和民族复兴任务的实践方向。",
        "new_answer": "①认识对实践具有反作用，正确认识能够指导实践、推动实践发展。②正确理解抗战精神中的爱国情怀、民族气节、英雄气概和必胜信念，能帮助青年把历史记忆转化为价值追求和行动方向。③继承弘扬抗战精神，就要用这种正确认识引导新时代实践，把精神力量落实为强国建设、民族复兴和担当奉献的实际行动。",
    },
    {
        "queue_id": "T0016",
        "heading": "11. 2026通州期末 第16题（主观题）",
        "node": "一切从实际出发 / 实事求是 / 主观与客观具体的历史的统一",
        "old_answer_excerpt": "都江堰的治水智慧在于李冰从蜀地西北高、东南低的地形和水势实际出发",
        "evidence_note": "正式细则第16题哲学角度写明坚持一切从实际出发、实事求是。",
        "new_why": "材料抓手是“当地西北高、东南低”的地形和水势，以及李冰因地制宜、顺势而为的工程设计。它说明治水方案不是主观想象出来的，而是把人的目的、技术和方案同客观自然条件结合起来，使主观认识和客观实际达到具体的历史的统一。",
        "new_answer": "①物质决定意识，办事情要坚持一切从实际出发、实事求是。②李冰修建都江堰时不是凭空设计工程，而是从蜀地西北高、东南低的地形、水势和灌溉防洪需要出发，因地制宜、顺势而为，使治水方案符合客观自然条件。③正因为主观设计同客观实际相符合，都江堰才能长期发挥防洪、灌溉和生态效益，体现跨越千年的治水智慧。",
    },
    {
        "queue_id": "T0096",
        "heading": "12. 2026东城期末 第21题（主观题）",
        "node": "规律的客观性",
        "old_answer_excerpt": "推进区域协调发展要尊重规律的客观性",
        "evidence_note": "细则列明尊重客观规律等角度，材料强调不同区域自然禀赋与发展基础存在差异。",
        "new_why": "区域发展落差不是靠主观愿望就能消除，也不能用同一种模式强行处理。材料要求把各区域自然禀赋、发展基础、比较优势和协同关系放在一起分析，说明现代化建设必须尊重区域分工、优势互补和协调发展的客观规律，才能把落差转化为发展动能。",
        "new_answer": "①规律具有客观性，推进区域协调发展必须尊重不同地区自然禀赋、产业基础、区位条件和发展阶段的差异。②我国不是简单平均用力，而是引导京津冀、长三角、粤港澳以及东北、中部、西部等区域立足比较优势，在分工协作中优势互补。③这样才能按区域发展和现代化建设规律办事，把发展落差中的势能转化为协同推进中国式现代化的动能。",
    },
    {
        "queue_id": "T0168",
        "heading": "12. 2026丰台期末 第16题（主观题）",
        "node": "整体与部分",
        "old_answer_excerpt": "留白体现立足整体、正确处理整体与部分关系",
        "evidence_note": "材料把留白放到山林溪谷、高楼绿地、方寸绿意与周边环境关系中理解，支持整体与部分、联系观点。",
        "new_why": "留白的价值不在于一块空地本身，而在于它同山林、溪谷、高楼、绿地、人的感受和城市空间共同构成整体意境。材料说留白让周边元素相互滋养、形成和谐有序之美，正说明部分只有放进整体联系中才能显现功能，整体也通过各部分协调配合形成审美效果。",
        "new_answer": "①整体由部分构成，部分离不开整体，部分的功能及其变化会影响整体功能。②留白作为空间中的一个部分，只有同山林溪谷、高楼绿地、城市喧嚣和人的审美感受相互联系、相互配合，才能形成疏密有致、和谐有序的整体意境。③理解留白不能孤立看“空”，而要立足整体空间关系，把留白作为优化整体功能和审美效果的重要部分来把握。",
    },
    {
        "queue_id": "T0472",
        "heading": "12. 2026海淀期末 第17题（主观题）",
        "node": "认识对实践的反作用",
        "old_answer_excerpt": "正确认识对实践具有反作用。青年理解红色文化的精神内涵和时代价值后",
        "evidence_note": "正式评分标准写明正确的认识有助于推动年轻人积极投身于实践。",
        "new_why": "材料中的红色文化“圈粉”年轻人，不只是因为形式新颖，更因为青年在沉浸体验、讲解传播和实践活动中形成了对红色文化精神内涵和时代价值的正确认识。细则明确说这种正确认识会推动年轻人积极投身实践，因此应从认识反作用于实践来答。",
        "new_answer": "①正确认识对实践具有促进作用，能够增强实践的自觉性和方向性。②青年通过红色场馆体验、主题活动、志愿讲解和网络传播等方式理解红色文化，形成对理想信念、家国情怀和奋斗精神的正确认识。③这种认识会进一步激发他们主动参与传承传播、社会服务和现实实践，使红色文化从“被看见”转化为青年愿意践行的精神力量。",
    },
    {
        "queue_id": "T0169",
        "heading": "13. 2025丰台期末 第16题（主观题）",
        "node": "整体与部分",
        "old_answer_excerpt": "用系统观念分析“数”，从整体把握部分之间及部分与整体的关联",
        "evidence_note": "题目把国家治理中的看大数、看长远、做计划、定方案作为作答抓手，支持系统观念和整体部分关系。",
        "new_why": "“胸中有数”在国家治理中不是掌握零散数字，而是把人口、资源、产业、民生、区域、时间节奏等多方面数量关系放到整体中统筹。材料中的“看大数、看长远、做计划、定方案”说明数字要服务整体谋划，部分数据只有同整体目标和相互关系结合起来才有治理意义。",
        "new_answer": "①整体居于主导地位，部分服从和服务于整体；系统观念要求把各要素联系起来统筹把握。②国家治理中的“数”不能停留在孤立数字，而要把人口、资源、产业、区域、民生等部分数据放进发展全局中分析，看到部分之间以及部分与整体之间的联系。③这样才能看大数、看长远，科学制定计划和方案，使局部安排服务国家治理整体目标。",
    },
]


def now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M +08")


def write_csv(path: Path, rows: list[dict[str, object]], fields: list[str]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fields})


def write_draft_files() -> None:
    rows = []
    for target in TARGETS:
        rows.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "evidence_note": target["evidence_note"],
                "new_why_chars": len(target["new_why"]),
                "new_answer_chars": len(target["new_answer"]),
                "new_why": target["new_why"],
                "new_answer": target["new_answer"],
            }
        )
    (ROOT / "P0_THICKNESS_REPAIR_BATCH02_DRAFT_20260525.json").write_text(
        json.dumps({"updated": now(), "targets": rows}, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / "P0_THICKNESS_REPAIR_BATCH02_DRAFT_20260525.csv",
        rows,
        ["queue_id", "heading", "node", "evidence_note", "new_why_chars", "new_answer_chars", "new_why", "new_answer"],
    )
    lines = [
        "# P0 Thickness Repair Batch02 Draft 20260525",
        "",
        f"Updated: {now()}",
        "",
        "Status: `DRAFT_READY_FOR_DOCX_APPLICATION`",
        "",
        "- Scope: next 8 P0 subjective triple-thin rows after Batch01 refresh.",
        "- Repair method: thicken only existing why/answer paragraphs inside existing nodes.",
        "",
        "| queue_id | node | heading | new why chars | new answer chars | evidence note |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in rows:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | {row['new_why_chars']} | {row['new_answer_chars']} | {row['evidence_note']} |"
        )
    (ROOT / "P0_THICKNESS_REPAIR_BATCH02_DRAFT_20260525.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )


def main() -> int:
    write_draft_files()
    docx = helper.current_docx()
    backup = docx.with_name(f"{docx.stem}_backup_before_p0_thickness_batch02_{datetime.now():%Y%m%d_%H%M%S}.docx")
    shutil.copy2(docx, backup)

    doc = Document(str(docx))
    blocks = helper.entry_blocks(doc)
    applied: list[dict[str, object]] = []
    for target in TARGETS:
        block = helper.matching_block(blocks, target)
        fields = block["fields"]
        field_paras = block["field_paras"]
        before_why = str(fields.get("why", ""))
        before_answer = str(fields.get("answer", ""))
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["why"])], "why", target["new_why"])
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["answer"])], "answer", target["new_answer"])
        applied.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "heading_para": block["heading_para"],
                "why_para": field_paras["why"],
                "answer_para": field_paras["answer"],
                "before_why_chars": len(before_why),
                "after_why_chars": len(target["new_why"]),
                "before_answer_chars": len(before_answer),
                "after_answer_chars": len(target["new_answer"]),
                "evidence_note": target["evidence_note"],
                "old_answer_excerpt": target["old_answer_excerpt"],
            }
        )

    doc.save(str(docx))
    with zipfile.ZipFile(docx) as z:
        zip_ok = "[Content_Types].xml" in z.namelist() and "word/document.xml" in z.namelist()

    result = {
        "updated": now(),
        "status": "P0_BATCH02_APPLIED_REQUIRES_RENDER_AND_EXTERNAL_RECHECK",
        "docx": str(docx),
        "backup_docx": str(backup),
        "zip_ok": zip_ok,
        "targets": len(TARGETS),
        "applied": applied,
        "boundary": [
            "This is a local thickness repair batch, not final acceptance.",
            "PDF/render outputs must be regenerated after this DOCX edit.",
            "External review remains pending for the current post-repair artifact.",
        ],
    }
    (ROOT / "P0_THICKNESS_REPAIR_BATCH02_APPLY_20260525.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / "P0_THICKNESS_REPAIR_BATCH02_APPLY_20260525.csv",
        applied,
        [
            "queue_id",
            "heading",
            "node",
            "heading_para",
            "why_para",
            "answer_para",
            "before_why_chars",
            "after_why_chars",
            "before_answer_chars",
            "after_answer_chars",
            "evidence_note",
            "old_answer_excerpt",
        ],
    )
    lines = [
        "# P0 Thickness Repair Batch02 Apply 20260525",
        "",
        f"Updated: {result['updated']}",
        "",
        f"Status: `{result['status']}`",
        "",
        f"- DOCX: `{docx.name}`.",
        f"- Backup: `{backup.name}`.",
        f"- Targets applied: `{len(applied)}` / `{len(TARGETS)}`.",
        f"- DOCX zip structural check: `{str(zip_ok).lower()}`.",
        "",
        "## Applied Rows",
        "",
        "| queue_id | node | heading | why chars | answer chars | paragraph refs |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in applied:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | "
            f"{row['before_why_chars']} -> {row['after_why_chars']} | "
            f"{row['before_answer_chars']} -> {row['after_answer_chars']} | "
            f"h={row['heading_para']}; why={row['why_para']}; answer={row['answer_para']} |"
        )
    lines.extend(
        [
            "",
            "## Boundary",
            "",
            "- This batch repairs only 8 additional P0 rows.",
            "- Remaining thickness queue, render QA, and model gates stay open until rechecked.",
        ]
    )
    (ROOT / "P0_THICKNESS_REPAIR_BATCH02_APPLY_20260525.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
