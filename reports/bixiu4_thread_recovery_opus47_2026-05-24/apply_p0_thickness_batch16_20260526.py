from __future__ import annotations

import csv
import json
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

import apply_p0_thickness_batch01_20260525 as helper
from docx import Document


ROOT = Path(__file__).resolve().parent
BATCH = "16"
RUN_DATE = "20260526"

TARGET_TEXTS = {
    "T0386": {
        "node": "认识对实践的反作用",
        "evidence_note": "矩阵 M1047：PPT评分细则第337-340行明列“以新的理论指导新的实践”。",
        "new_why": "谱写马克思主义中国化时代化新篇章，不是把理论创新停留在概念表述中，而是要求新的理论认识继续指导党治国理政、中国式现代化和认识世界改造世界的实践。矩阵 M1047 记录PPT评分细则第337-340行明列“以新的理论指导新的实践”，所以本条应写出正确理论形成后如何反作用于实践、转化为现实力量。",
        "new_answer": "①认识对实践具有反作用，科学理论能够指导实践、推动实践发展。②党的创新理论是在实践基础上形成和发展的新的理论认识，不能只停留在文本和口号中。③谱写马克思主义中国化时代化新篇章，就是要把习近平新时代中国特色社会主义思想的世界观和方法论运用于治国理政、改革发展、现代化建设等实践，把理论优势转化为分析问题、解决问题、改造世界的实践力量，在新的实践中继续推进理论创新。",
    },
    "T0142": {
        "node": "系统观念 / 系统优化",
        "evidence_note": "矩阵 M0051/M0130：丰台一模主观题正式阅卷细则Q16第（2）点明确“坚持联系的观点，立足整体，统筹部分”，要求统筹科技、伦理与社会公平。",
        "new_why": "推动人工智能健康发展涉及科技创新、文化传承、伦理治理、社会公平和民生福祉等多个部分，任何一个环节单独推进都可能造成整体失衡。矩阵 M0051/M0130 记录丰台一模正式阅卷细则要求坚持联系观点、立足整体、统筹部分，并具体指向科技、伦理与社会公平，因此本条要把AI治理写成系统优化问题。",
        "new_answer": "①系统观念要求立足整体，把各部分、各要素联系起来统筹考虑，优化结构以实现整体功能。②人工智能发展不是单纯技术效率问题，它同时牵动文化保护、伦理安全、就业公平、公共服务和人的全面发展。③推动人工智能健康发展，要统筹技术创新与规则治理、产业效率与社会公平、文化传承与数字创造，使各环节相互协调、相互促进。④只有把部分放到整体文明进步中安排，技术进步才不会偏离人的价值和社会整体利益。",
    },
    "T0425": {
        "node": "社会发展的两大基本规律和基本矛盾",
        "evidence_note": "矩阵 M1212：正式细则第444行列“上层建筑”；材料《国家发展规划法》提供法律支撑。",
        "new_why": "国家发展规划法和规划制度属于制度、法律、治理体系层面的安排，落在唯物史观中就是上层建筑调整。矩阵 M1212 记录正式细则第444行列“上层建筑”，材料又突出用法律支撑规划实施、用制度保障现代化建设，因此不能只写一般法治意义，而要写上层建筑如何适应并服务经济基础和社会发展需要。",
        "new_answer": "①生产力和生产关系、经济基础和上层建筑的矛盾运动，是推动社会发展的基本矛盾和基本规律。②国家发展规划法、规划制度和法律支撑属于上层建筑的重要内容，它们要同我国经济社会发展阶段、现代化建设任务和治理需要相适应。③通过把发展规划制度化、法治化，可以规范战略目标、职责分工和实施机制，使上层建筑更好服务经济社会发展、国家治理和中国式现代化实践。④这体现了通过调整完善上层建筑来促进社会发展的要求。",
    },
    "T0405": {
        "node": "真理观",
        "evidence_note": "矩阵 M1400：正式评分细则page_002认识论角度写明追求真理是人类认识活动根本任务，并在酌情原则列出真理最基本属性。",
        "new_why": "本题问“真实”何以直抵人心，不能只把真实理解成情感真诚。矩阵 M1400 记录正式评分细则page_002把追求真实同追求真理相连，并列出真理的最基本属性；这说明答案应从真理观说明：真实认识要经实践检验，符合客观事实，因而具有跨越偏见、形成共识的力量。",
        "new_answer": "①真理是标志主观同客观相符合的哲学范畴，真理最基本的属性是客观性。②追求真实，就是反对虚假、偏见和主观臆断，使认识经得起实践检验、符合客观事实。③当新闻、历史、司法、科学、纪录片等领域都以真实为追求时，人们获得的不是任意表达，而是更接近客观事实和真理的认识。④正因真实认识具有客观根据，能够突破立场隔阂和情绪遮蔽，促成人们形成共同判断，所以具有直抵人心的力量。",
    },
    "T0005": {
        "node": "物质决定意识",
        "evidence_note": "矩阵 M1438：正式细则第16题唯物论角度写明“体现了物质决定意识”。",
        "new_why": "都江堰治水方案不是先有主观设想再强加给自然，而是由岷江水势、成都平原地势、水沙关系、灌溉防洪需求等客观条件所制约。矩阵 M1438 记录正式细则第16题唯物论角度明确写“体现了物质决定意识”，所以本条应突出客观自然条件决定治水认识和工程设计，不能写成单纯技术赞美。",
        "new_answer": "①物质决定意识，要求人的认识和方案必须以客观存在的实际情况为依据。②都江堰的治水方案不是主观臆造出来的，而是从岷江水势、地形高差、水沙运行、灌溉防洪和区域生产生活等客观实际出发形成的。③鱼嘴分水、飞沙堰排沙、宝瓶口控水等工程安排，正是对当地自然条件和水利实际的反映与利用。④因此，都江堰的成功说明治水意识和工程方案必须符合客观实际，主观设计只有以物质条件为基础才能发挥作用。",
    },
    "T0012": {
        "node": "一切从实际出发 / 实事求是 / 主观与客观具体的历史的统一",
        "evidence_note": "矩阵 M0995：2025丰台期末正式PPT细则第205-263行支持实践、实际出发、联系观、量变质变、矛盾观等；本条落到从实际出发。",
        "new_why": "毛泽东强调胸中有“数”，不是要求机械崇拜数字，而是要求决策者从调查统计、客观数量关系、事物发展界限和具体条件出发。矩阵 M0995 记录正式PPT细则支持实际出发等哲学角度；本条落在“一切从实际出发”节点，应写清“数”如何使主观计划、方案和治理判断同客观实际相符合。",
        "new_answer": "①物质决定意识，要求坚持一切从实际出发、实事求是，使主观认识和行动符合客观实际。②“数”首先来自调查统计和客观情况，不能凭主观感觉、经验印象或愿望编造。③“数”还要结合事物质量、数量界限和具体条件来确定，脱离实际条件的数字会误导判断。④胸中有“数”的目的，是把客观数量事实转化为计划制定、政策安排和国家治理的依据，使主观方案同真实情况相符合，避免盲目决策和过犹不及。",
    },
    "T0218": {
        "node": "辩证否定 / 守正创新",
        "evidence_note": "矩阵 M0049/M0052：丰台一模正式阅卷细则Q16要求全面辩证看AI积极作用和风险挑战，并以正确价值观引导其健康发展；本条落到扬弃式创新。",
        "new_why": "面对人工智能，材料和细则都不是要求全盘肯定或全盘否定，而是要求看到AI推动文物修复、数字创作、公共服务的积极作用，同时克服信息失真、文化窄化、智能鸿沟等风险。正式阅卷细则要求全面辩证看AI积极作用和风险挑战，并以正确价值观引导其健康发展，本条据此落到辩证否定的“扬弃”逻辑。",
        "new_answer": "①辩证否定是事物自身的否定，实质是扬弃，既不是简单肯定一切，也不是简单否定一切。②面对人工智能，应保留和发展其在文物修复、数字创作、公共服务、文化传播和文明进步中的积极作用，使技术创新服务人的发展。③同时要否定和克服其可能带来的信息失真、文化窄化、价值偏差和智能鸿沟等风险，通过规则治理、伦理约束和价值引导加以规范。④这样才能在守住人的主体地位和文明价值的基础上实现技术创新，让AI健康发展。",
    },
    "T0335": {
        "node": "实践与认识（总）",
        "evidence_note": "矩阵 M1258：正式评分细则第100-102行在“如何弘扬”中列“立足实践、知行合一”。",
        "new_why": "抗战精神来源于全民族抗战的历史实践，新时代弘扬也不能停留在口号、纪念和情绪感动上，而要落实到强国建设、民族复兴和现实奋斗中。矩阵 M1258 记录正式评分细则第100-102行在“如何弘扬”中列“立足实践、知行合一”，所以本条要写出认识与实践统一：理解精神内涵后必须回到现实行动。",
        "new_answer": "①实践是认识的基础，也是认识的目的和归宿，正确认识必须回到实践中接受检验并指导行动。②抗战精神是在全民族抗战的伟大实践中形成的，对其内涵的理解不能停留在书本记忆和情感表达上。③新时代弘扬抗战精神，要立足强国建设和民族复兴的现实实践，把爱国情怀、民族气节、英雄气概和必胜信念转化为岗位奋斗、攻坚克难、服务人民的实际行动。④做到知行合一，才能使精神力量真正转化为推动现实发展的实践力量。",
    },
}


def now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M +08")


def load_queue_targets() -> list[dict[str, str]]:
    queue_path = ROOT / "THICKNESS_REPAIR_PRIORITY_QUEUE_20260525.csv"
    wanted = set(TARGET_TEXTS)
    found: dict[str, dict[str, str]] = {}
    with queue_path.open(encoding="utf-8-sig", newline="") as f:
        for row in csv.DictReader(f):
            qid = row.get("queue_id", "")
            if qid in wanted and qid not in found:
                found[qid] = row
    missing = wanted - set(found)
    if missing:
        raise RuntimeError(f"Missing target queue rows: {sorted(missing)}")
    targets: list[dict[str, str]] = []
    for qid in TARGET_TEXTS:
        row = found[qid]
        text = TARGET_TEXTS[qid]
        targets.append(
            {
                "queue_id": qid,
                "heading": row["heading"],
                "node": text["node"],
                "old_answer_excerpt": row["answer_excerpt"],
                "evidence_note": text["evidence_note"],
                "new_why": text["new_why"],
                "new_answer": text["new_answer"],
            }
        )
    return targets


def write_csv(path: Path, rows: list[dict[str, object]], fields: list[str]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fields})


def write_draft_files(targets: list[dict[str, str]]) -> None:
    rows = []
    for target in targets:
        rows.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "evidence_note": target["evidence_note"],
                "new_why_chars": len(target["new_why"]),
                "new_answer_chars": len(target["new_answer"]),
                "new_why": target["new_why"],
                "new_answer": target["new_answer"],
            }
        )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_{RUN_DATE}.json").write_text(
        json.dumps({"updated": now(), "targets": rows}, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_{RUN_DATE}.csv",
        rows,
        ["queue_id", "heading", "node", "evidence_note", "new_why_chars", "new_answer_chars", "new_why", "new_answer"],
    )
    lines = [
        f"# P0 Thickness Repair Batch{BATCH} Draft {RUN_DATE}",
        "",
        f"Updated: {now()}",
        "",
        "Status: `DRAFT_READY_FOR_DOCX_APPLICATION`",
        "",
        "- Scope: next 8 P0 subjective triple-thin rows after Batch15 refresh.",
        "- Repair method: thicken only existing why/answer paragraphs inside existing evidence-supported nodes.",
        "- Evidence boundary: all evidence notes cite current matrix rows; ordinary reference answers are not upgraded into scoring rubrics.",
        "",
        "| queue_id | node | heading | new why chars | new answer chars | evidence note |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in rows:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | {row['new_why_chars']} | {row['new_answer_chars']} | {row['evidence_note']} |"
        )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_{RUN_DATE}.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )


def main() -> int:
    targets = load_queue_targets()
    write_draft_files(targets)
    docx = helper.current_docx()
    backup = docx.with_name(f"{docx.stem}_backup_before_p0_thickness_batch{BATCH}_{datetime.now():%Y%m%d_%H%M%S}.docx")
    shutil.copy2(docx, backup)

    doc = Document(str(docx))
    blocks = helper.entry_blocks(doc)
    applied: list[dict[str, object]] = []
    for target in targets:
        block = helper.matching_block(blocks, target)
        fields = block["fields"]
        field_paras = block["field_paras"]
        before_why = str(fields.get("why", ""))
        before_answer = str(fields.get("answer", ""))
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["why"])], "why", target["new_why"])
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["answer"])], "answer", target["new_answer"])
        applied.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "heading_para": block["heading_para"],
                "why_para": field_paras["why"],
                "answer_para": field_paras["answer"],
                "before_why_chars": len(before_why),
                "after_why_chars": len(target["new_why"]),
                "before_answer_chars": len(before_answer),
                "after_answer_chars": len(target["new_answer"]),
                "evidence_note": target["evidence_note"],
                "old_answer_excerpt": target["old_answer_excerpt"],
            }
        )

    doc.save(str(docx))
    with zipfile.ZipFile(docx) as z:
        zip_ok = "[Content_Types].xml" in z.namelist() and "word/document.xml" in z.namelist()

    result = {
        "updated": now(),
        "status": f"P0_BATCH{BATCH}_APPLIED_REQUIRES_RENDER_AND_EXTERNAL_RECHECK",
        "docx": str(docx),
        "backup_docx": str(backup),
        "zip_ok": zip_ok,
        "targets": len(targets),
        "applied": applied,
        "boundary": [
            "This is a local thickness repair batch, not a terminal acceptance claim.",
            "PDF/render outputs must be regenerated after this DOCX edit.",
            "External review remains pending for the current post-repair artifact.",
        ],
    }
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_{RUN_DATE}.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_{RUN_DATE}.csv",
        applied,
        [
            "queue_id",
            "heading",
            "node",
            "heading_para",
            "why_para",
            "answer_para",
            "before_why_chars",
            "after_why_chars",
            "before_answer_chars",
            "after_answer_chars",
            "evidence_note",
            "old_answer_excerpt",
        ],
    )
    lines = [
        f"# P0 Thickness Repair Batch{BATCH} Apply {RUN_DATE}",
        "",
        f"Updated: {result['updated']}",
        "",
        f"Status: `{result['status']}`",
        "",
        f"- DOCX: `{docx.name}`.",
        f"- Backup: `{backup.name}`.",
        f"- Targets applied: `{len(applied)}` / `{len(targets)}`.",
        f"- DOCX zip structural check: `{str(zip_ok).lower()}`.",
        "",
        "## Applied Rows",
        "",
        "| queue_id | node | heading | why chars | answer chars | paragraph refs |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in applied:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | "
            f"{row['before_why_chars']} -> {row['after_why_chars']} | "
            f"{row['before_answer_chars']} -> {row['after_answer_chars']} | "
            f"h={row['heading_para']}; why={row['why_para']}; answer={row['answer_para']} |"
        )
    lines.extend(
        [
            "",
            "## Boundary",
            "",
            "- This batch repairs only 8 additional P0 rows.",
            "- Remaining thickness queue, render QA, and model gates stay open until rechecked.",
        ]
    )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_{RUN_DATE}.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
