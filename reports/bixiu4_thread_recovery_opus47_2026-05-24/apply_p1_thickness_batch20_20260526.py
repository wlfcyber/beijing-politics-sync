from __future__ import annotations

import csv
import json
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

import apply_p0_thickness_batch01_20260525 as helper
from docx import Document


ROOT = Path(__file__).resolve().parent
BATCH = "20"
RUN_DATE = "20260526"

TARGET_TEXTS = [
    {
        "queue_id": "T0365",
        "heading": "1. 2025海淀期末 第17题第（2）问（主观题）",
        "node": "认识发展原理",
        "evidence_note": "矩阵 M1107：2025海淀期末 Q17(2) PPT评标细则第69-70、91-94行明列“认识有反复性或无限性/追求真理是一个过程”。",
        "new_why": "永动机材料的关键词不是“失败无意义”，而是“多次尝试、不断否定、由浅入深、最终发现能量守恒”。这些信息说明人的认识不是一次完成的，而是在实践探索、失败反思和理论概括中不断发展。正式评标细则把本题列入认识反复性、无限性和追求真理过程，所以答案要写出“失败探索推动认识深化”的链条，不能只停留在纪念失败者的情感表达。",
        "new_answer": "①认识具有反复性和无限性，追求真理是一个永无止境的过程。永动机探索者一次次尝试又一次次失败，说明人们对能量转化规律的认识不可能一蹴而就。②失败并没有切断认识发展，反而暴露旧设想的问题，促使后来科学家从经验事实、实验结果和理论分析中不断修正认识，逐渐形成能量守恒思想。③纪念失败者，就是承认真理发现常常经历由浅入深、由含糊到清晰的过程，失败实践同样是人类认识接近真理的重要环节。",
    },
    {
        "queue_id": "T0348",
        "heading": "1. 2025西城一模 第16题（主观题）",
        "node": "认识对实践的反作用",
        "evidence_note": "矩阵 M0180/M0214：2025西城一模 Q16 正式评分细则给哲学3分，明确“以知导行、亲身实践、在实践基础上不断发展知促进行”。",
        "new_why": "研学活动不是先实践再随意总结，而是从行前准备就形成目标、问题意识和方法安排，再把这种认识带入实地调查和后续行动。材料中的诵读《诗经》、观看纪录片、确立研学目标，是“知”；寻找植物、比较水文、设计课程和线路，是“行”。正式细则写明以知导行，所以本条要突出正确认识对实践具有指导作用。",
        "new_answer": "①认识对实践具有能动反作用，正确认识能够指导实践有目的、有计划地开展。行前准备阶段，学生通过诵读《诗经》、观看纪录片，形成对文化溯源、学科融合和实践创新的认识，明确研学方向。②这种认识不是停留在书本上，而是指导后续实地研学：学生带着问题去寻植物、比较水文、设计课程和旅游线路，使实践活动避免盲目。③从行前准备到持续行动，体现了在正确认识引导下开展实践，并在实践中继续发展认识、改进活动方案。",
    },
    {
        "queue_id": "T0285",
        "heading": "1. 2026东城一模 第16题（主观题）",
        "node": "矛盾的主要方面和次要方面 / 两点论与重点论",
        "evidence_note": "矩阵 M0219：2026东城一模 Q16 正式PPT评分细则明确哲学核心考点含“矛盾的主次方面（情绪需求是次要方面，思想性、艺术性是主要方面）”。",
        "new_why": "题目观点把“满足大众情绪需求”当作判断优秀文艺的根本标准，正好触碰细则中的主次方面区分。材料承认文艺可以回应情绪需求，但优秀文艺更要有思想高度、艺术品质、家国情怀和时代精神。正式细则点明思想性、艺术性是主要方面，情绪需求是次要方面，因此答案要同时写出“两点”和“重点”。",
        "new_answer": "①看待优秀文艺要坚持两点论与重点论相统一。优秀文艺可以回应大众情绪需求，给人审美愉悦和情感共鸣，这是需要看到的一面。②但在判定“优秀”时，思想性、艺术性、时代精神和家国情怀才是主要方面；情绪满足只是次要方面，不能替代作品的价值高度和艺术质量。③题中观点把次要方面当成主要方面，用情绪需求遮蔽思想性和艺术性，所以是片面的；真正优秀的文艺应在满足人民精神需要的同时突出正确价值和高水平艺术表达。",
    },
    {
        "queue_id": "T0214",
        "heading": "1. 2026海淀期末 第16题（主观题）",
        "node": "矛盾就是对立统一",
        "evidence_note": "矩阵 M1353：2026海淀期末 Q16 正式评分标准列明“矛盾的观点”，细则写“一分为二、两重统一”。",
        "new_why": "数字原住民与技术进步的关系同时包含便利与风险：互联互通、智能便捷是积极面，感官刺激陷阱、思维惰化、表达弱化是消极面。正式评分标准要求用矛盾观点、一分为二和两重统一分析，说明本题不能单纯赞美技术，也不能简单否定技术，而要把同一技术环境中的两个方面说清楚。",
        "new_answer": "①矛盾就是对立统一，同一事物内部包含既相互区别又相互联系的两个方面。技术进步给数字原住民带来互联互通、海量信息、智能便捷等红利，也带来感官刺激陷阱、思维惰化、表达和深度思考能力弱化等风险。②这两方面共同存在于数字生活之中，不能只看便利而忽视危害，也不能因风险而全盘否定技术。③数字原住民要坚持一分为二地看待技术，在享受技术红利的同时增强主体意识和深度思考能力，防止被技术牵着走。",
    },
    {
        "queue_id": "T0289",
        "heading": "1. 2026海淀期末 第16题（主观题）",
        "node": "两点论与重点论",
        "evidence_note": "矩阵 M1346：2026海淀期末 Q16 正式评分标准写“一分为二、两重统一”，并提示“重点论”出不来只能1分。",
        "new_why": "本条的旧答案已有重点论方向，但还需要把细则要求的“两重统一”讲清。数字技术既有帮助成长的积极作用，也有弱化深度思考和主体性的消极影响；学生不能平均罗列两边，而要抓住人的主体成长这个重点。评分标准提示重点论不出只能低分，所以答案要明确“全面看两点，并把保持主体性放在主要位置”。",
        "new_answer": "①坚持两点论与重点论相统一，既要全面看待矛盾双方，又要抓住主要方面。数字原住民应看到技术一方面提供学习资源、沟通渠道和创造工具，另一方面也可能诱发碎片化、复制粘贴和沉溺感官刺激。②在两点之中，人的主体性、独立思考和深度学习应当放在主要位置，技术只能作为辅助工具，不能替代人的判断和创造。③因此，正确使用技术不是各打五十大板，而是在承认技术价值的同时重点防范其支配人、削弱人的风险，让技术服务人的成长。",
    },
    {
        "queue_id": "T0320",
        "heading": "1. 2026海淀期末 第16题（主观题）",
        "node": "实践是认识的基础",
        "evidence_note": "矩阵 M1352：2026海淀期末 Q16 正式评分标准列明“实践的观点”，题干要求结合材料或个人体验说明数字原住民如何在技术使用实践中成长。",
        "new_why": "题目问数字原住民如何成长，不能只写观念表态。材料和设问都要求回到技术使用实践：选择信息、阅读思考、规范表达、真实交往，都是在实践中检验“技术如何为我所用”的过程。正式评分标准把实践观点列为可用角度，因此本条要写出实践是认识来源、检验和发展的基础。",
        "new_answer": "①实践是认识的基础，认识技术、驾驭技术，最终要回到具体使用技术的实践中。数字原住民不能只在观念上说要正确使用技术，而要在真实学习和生活中主动选择高质量信息、静心阅读、独立书写、规范表达。②这些实践会不断检验和深化他们对技术利弊的认识，使其明白哪些使用方式促进成长，哪些使用方式导致依赖和惰化。③只有在实践中反复调整技术使用习惯，才能让技术真正服务学习、创造和社会参与，而不是削弱人的主体能力。",
    },
    {
        "queue_id": "T0099",
        "heading": "1. 2026海淀期末 第17题（主观题）",
        "node": "根据固有联系建立新的具体联系",
        "evidence_note": "矩阵 M1358：2026海淀期末 Q17 正式评分标准“红色文化的发展”维度列出联系观点，材料通过体验、数字技术和青年讲述建立红色文化与青年生活的新具体联系。",
        "new_why": "红色文化能圈粉年轻人，不是因为历史本身自动变得年轻，而是因为材料中的挑粮体验、数字艺术馆、青年讲解视频把红色文化内核与当代青年的感受方式重新连接起来。正式评分标准支持联系观点，本题更具体的落点是根据固有联系建立新的具体联系：尊重红色文化与青年成长、理想信念之间的固有联系，再用新场景激活。",
        "new_answer": "①人们可以根据事物固有联系，改变事物状态、调整原有联系，建立新的具体联系。红色文化本身与青年理想信念、精神成长和责任担当存在内在联系，但这种联系如果只停留在课本叙述中，年轻人未必容易感受。②井冈山挑粮体验、长征数字艺术馆光影沉浸、延安青年讲解视频，把红色历史同体验、视觉、讲述和传播方式结合起来，建立了红色文化与青年生活的新联系。③因此，红色文化能够从“遥远历史”转化为青年可参与、可感知、可传播的现实情境。",
    },
    {
        "queue_id": "T0005",
        "heading": "1. 2026西城期末 第21题（主观题）",
        "node": "一切从实际出发 / 实事求是 / 主观与客观具体的历史的统一",
        "evidence_note": "矩阵 M1388：2026西城期末 Q21 正式评分细则 page_005 角度3列出“实事求是”。",
        "new_why": "五年规划之所以科学，不能只写“规划很多、坚持很久”，而要说明它为什么能与中国发展阶段相适应。材料从早期经济发展到后来覆盖民生、科技、生态、文化等，体现规划随客观实际变化而调整。正式评分细则明确列“实事求是”，因此答案要把“从实际出发、随实际调整、形成科学规划”写完整。",
        "new_answer": "①一切从实际出发、实事求是，要求使主观认识和客观实际相符合。五年规划不是凭空设定目标，而是立足我国不同历史阶段的发展水平、主要任务和现实条件来制定。②从早期侧重经济建设，到逐步覆盖社会民生、科技创新、生态环境、文化建设等，规划内容随着国家发展实际和人民需要变化而不断拓展。③正是因为五年规划坚持从实际出发，在实践中根据新问题、新趋势、新挑战调整路径，才保持科学性和连续性，成为治国理政的重要经验。",
    },
    {
        "queue_id": "T0304",
        "heading": "10. 2025朝阳一模 第16题（主观题）",
        "node": "实践是认识的基础",
        "evidence_note": "矩阵 M0177/M0209/M0424：2025朝阳一模 Q16 正式评分细则原图支持借助现代科技推动文化创新、立足时代之基和创作实践。",
        "new_why": "国产动画电影成功不是靠抽象热爱传统文化，而是靠长期制作实践把传统文化、现代技术和时代审美结合。材料中的九个月特效打磨、现代科技、观众审美和传统文化转化，都说明实践提出问题、检验方案、推动认识深化。正式细则支持现代科技推动文化创新和立足时代之基，所以本条要落到实践推动认识和文化表达发展。",
        "new_answer": "①实践是认识的来源和发展动力，文化创新也必须立足具体创作实践。国产动画电影对中华优秀传统文化的理解，不是停留在概念层面，而是在剧本、人物、镜头、视效和传播效果的实践打磨中不断形成。②创作团队长期投入，关键戏份历经数月特效制作，把传统文化形象与现代电影技术、当代观众审美结合，在试错和修改中深化对传统文化现代表达的认识。③正是这种实践基础上的创新，使传统文化由静态符号转化为可观看、可共情、可传播的新时代艺术表达。",
    },
    {
        "queue_id": "T0422",
        "heading": "10. 2026东城期末 第16题（主观题）",
        "node": "价值观的导向作用",
        "evidence_note": "矩阵 M1171：2026东城期末 Q16 正式细则第156行列“价值观”，第174行示例明确“价值观具有导向作用”。",
        "new_why": "“白不止于白”的关键在于它从颜色表象上升为文化价值符号。材料中的清白、清正廉洁、刚正不阿、责任担当，都指向价值观对人认识和行为的引导。正式细则直接列价值观并示例“价值观具有导向作用”，因此答案要把“白”的审美意象如何转化为价值准则、如何引导人格和实践写清。",
        "new_answer": "①价值观对人们认识世界、改造世界的活动具有重要导向作用，也会影响人的价值判断和行为选择。“清白”作为文化符号，不只是颜色上的白，更承载清正廉洁、刚正不阿、心怀家国百姓等价值追求。②这种价值追求引导人们把外在审美转化为内在道德准则，形成崇清尚廉、坚守正直、担当责任的价值判断。③所以“白”不止于白，是因为它在文明传承中沉淀为引导人格塑造和实践选择的价值符号，能够持续涵养人的精神世界。",
    },
    {
        "queue_id": "T0011",
        "heading": "10. 2026房山一模 第16题第（2）问（主观题）",
        "node": "一切从实际出发 / 实事求是 / 主观与客观具体的历史的统一",
        "evidence_note": "矩阵 M0071/M0133：2026房山一模 Q16(2) 评标/细则 paras13-21 明示“一切从实际出发”，高档要求融入大局中找定位。",
        "new_why": "“因地制宜，本质就是实事求是”的设问已经给出明确方向。材料列举房山区能源资源、农业资源、乡村文化资源，以及宜林、宜草、宜沙、宜荒等安排，都是从客观条件出发，而不是用一个模板改造所有地区。正式评标细则明示一切从实际出发，所以答案要把“客观条件、具体定位、具体路径”结合起来。",
        "new_answer": "①一切从实际出发、实事求是，要求尊重客观实际，使主观认识和实践方案符合具体地区的条件。“因地制宜”首先要看各地真实存在的资源禀赋、产业基础、自然条件和发展定位。②房山区依托能源资源建设战略保障基地，依托农产品禀赋发展特色农业，依托乡村文化资源开展文化活动；“宜林则林、宜草则草、宜沙则沙、宜荒则荒”也体现按不同条件作不同安排。③这种做法避免脱离实际和照搬模式，正是把主观规划建立在客观实际基础上的实事求是。",
    },
    {
        "queue_id": "T0423",
        "heading": "11. 2025东城一模 第16题（主观题）",
        "node": "价值观的导向作用",
        "evidence_note": "矩阵 M0081：2025东城一模 Q16 正式阅卷报告明确价值观导向作用、价值判断和价值选择可用于价值/文化价值角度。",
        "new_why": "中国年能够生生不息，不只是因为节日形式热闹，而是因为春节承载共同价值追求。材料中的阖家团圆、和睦喜庆、人与自然和谐、追求美好生活，都是中华民族共同价值取向。正式阅卷报告明确可用价值观导向作用，因此答案要把春节价值观如何凝聚认同、引导行动、推动传承写出来。",
        "new_answer": "①价值观对人们认识世界和改造世界具有导向作用，正确价值观能够凝聚共识、引导实践。春节承载阖家团圆、和睦喜庆、敬祖爱亲、人与自然和谐、追求美好生活等共同价值取向。②这种价值取向让不同地区、不同代际的中国人在节日中形成情感认同和文化认同，增强民族自豪感与凝聚力。③春节形式可以与时俱进，但其共同价值始终引导人们珍视家庭、延续文化、追求美好生活，所以中国年能够年复一年焕发新的生命力。",
    },
    {
        "queue_id": "T0397",
        "heading": "12. 2025朝阳期末 第16题（主观题）",
        "node": "人民群众",
        "evidence_note": "矩阵 M1032：2025朝阳期末 Q16 PPT评分细则第137行明列“尊重人民群众主体地位/群众观点群众路线”。",
        "new_why": "用文化同世界对话，不能只从输出者视角安排内容。材料中汉代文物出境展出之所以受欢迎，是因为策展关注当地普通观众的生活方式、互动体验和理解方式，把观众作为文化交流主体。PPT评分细则明列尊重人民群众主体地位和群众路线，因此答案要强调面向人民、贴近群众、依靠群众理解来实现对话。",
        "new_answer": "①人民群众是社会历史的主体，也是文化交流和精神财富创造、接受、传播的重要主体。用文化同世界对话，不能只考虑我们想展示什么，还要尊重不同国家普通观众的理解方式、生活经验和文化需求。②汉代文物出境展出注重趣味互动，关注生活方式、文化交流等贴近普通观众的视角，寻找其与现代生活的契合点，拉近观众与文物的距离。③这体现了群众观点和群众路线：把人民作为文化对话的主体，文化交流才能真正被理解、接受并产生共鸣。",
    },
    {
        "queue_id": "T0077",
        "heading": "12. 2026丰台一模 第16题（主观题）",
        "node": "系统观念 / 系统优化",
        "evidence_note": "矩阵 M0051/M0130：2026丰台一模 Q16 正式阅卷细则明确“坚持联系的观点，立足整体，统筹部分”，要求统筹科技、伦理与社会公平。",
        "new_why": "AI 发展问题牵涉科技创新、伦理风险、社会公平、文化传承和不同群体现实需要，不能孤立看单一技术指标。正式阅卷细则要求坚持联系观点、立足整体、统筹部分，所以本条要从系统观念写出：把 AI 的发展目标、治理边界、群体差异和公共利益放进同一个整体中协调。",
        "new_answer": "①系统观念要求立足整体、统筹部分，注重各要素之间的相互联系和结构优化。人工智能的发展不是孤立的技术问题，而同社会生活、文化传承、民生福祉、科技伦理和社会公平紧密相关。②妥善处理这一时代命题，需要把技术效率、伦理安全、公共利益、弱势群体需求和文化多样性放在整体中统筹，而不能只追求技术扩张。③只有立足整体、协调各部分关系，才能使 AI 发展朝着有益、安全、公平的方向推进，实现社会整体效益最优。",
    },
    {
        "queue_id": "T0133",
        "heading": "12. 2026房山一模 第20题（主观题）",
        "node": "系统观念 / 系统优化",
        "evidence_note": "矩阵 M0569：2026房山一模 Q20 评标/细则 paras69-81 明示矛盾、联系/对立统一、认识与实践、系统优化等。",
        "new_why": "依法治国和依规治党有机统一，不是两个口号并列，而是把国家法律体系、党内法规体系、党的领导和国家治理放进一个整体中理解。房山一模 Q20 细则明示系统优化等角度，因此答案要写出系统内部各要素相互联系、目标一致、功能协同，避免只解释单个概念。",
        "new_answer": "①系统观念要求着眼整体，统筹部分，使系统内部各要素相互配合、协同运行。依法治国和依规治党有机统一，说明宪法法律、党内法规、党的领导和国家治理不是彼此割裂的，而是中国特色社会主义法治体系中的相互联系要素。②坚持依法治国，能够用国家法律规范社会治理；坚持依规治党，能够用党内法规管党治党、保障党的领导更加制度化规范化。③把二者统一起来，才能在价值目标、制度规则和运行机制上形成协同，提升治党治国的整体效能。",
    },
    {
        "queue_id": "T0199",
        "heading": "12. 2026门头沟一模 第16题（主观题）",
        "node": "辩证否定 / 守正创新",
        "evidence_note": "矩阵 M0076/M0170/M0232：2026门头沟一模 Q16 细则明确哲学角度包括辩证否定观/扬弃，答出一点2分，两点及以上3分，结合材料1分。",
        "new_why": "永定河古渠的活态传承不是把旧形态原封不动保存，也不是拆掉历史另起炉灶。材料中的灌溉、防洪功能是要保留的合理内核，生态补水、文化教育、文旅融合等新功能则是面向新时代的生发。正式细则明列辩证否定观，所以答案要把“保留、克服、发展”的扬弃逻辑写完整。",
        "new_answer": "①辩证否定是事物自身的否定，是联系和发展的环节，实质是扬弃。永定河古渠的保护传承首先要保留其灌溉、防洪、水利智慧和历史文化价值，延续古渠同区域生活和城市记忆的联系。②同时，不能让古渠停留在沉睡文物状态，而要克服旧有功能单一、展示方式封闭的局限，发展生态补水、文化教育、文旅融合等新功能。③这种守正创新使古渠在保留历史根脉中获得新的时代功能，成为奔流不息的活态遗产。",
    },
]


def now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M +08")


def json_path(path: Path) -> str:
    return str(path).replace("\\", "/")


def load_queue_targets() -> list[dict[str, str]]:
    queue_path = ROOT / "THICKNESS_REPAIR_PRIORITY_QUEUE_20260525.csv"
    found: dict[tuple[str, str], dict[str, str]] = {}
    with queue_path.open(encoding="utf-8-sig", newline="") as f:
        for row in csv.DictReader(f):
            found[(row.get("queue_id", ""), row.get("heading", ""))] = row
    targets: list[dict[str, str]] = []
    for target in TARGET_TEXTS:
        key = (target["queue_id"], target["heading"])
        row = found.get(key)
        if row is None:
            raise RuntimeError(f"Missing target queue row: {key}")
        targets.append({**target, "old_answer_excerpt": row["answer_excerpt"]})
    return targets


def write_csv(path: Path, rows: list[dict[str, object]], fields: list[str]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fields})


def write_draft_files(targets: list[dict[str, str]]) -> None:
    rows = []
    for target in targets:
        rows.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "evidence_note": target["evidence_note"],
                "new_why_chars": len(target["new_why"]),
                "new_answer_chars": len(target["new_answer"]),
                "new_why": target["new_why"],
                "new_answer": target["new_answer"],
            }
        )
    (ROOT / f"P1_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_{RUN_DATE}.json").write_text(
        json.dumps({"updated": now(), "targets": rows}, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / f"P1_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_{RUN_DATE}.csv",
        rows,
        ["queue_id", "heading", "node", "evidence_note", "new_why_chars", "new_answer_chars", "new_why", "new_answer"],
    )
    lines = [
        f"# P1 Thickness Repair Batch{BATCH} Draft {RUN_DATE}",
        "",
        f"Updated: {now()}",
        "",
        "Status: `DRAFT_READY_FOR_DOCX_APPLICATION`",
        "",
        "- Scope: 16 P1 subjective rows after P0 was cleared.",
        "- Repair method: thicken only existing why/answer paragraphs inside existing framework nodes.",
        "- Evidence boundary: all selected rows cite same-question scoring/rubric/marking-rule matrix support.",
        "- Excluded from this batch: T0257 and T0055 because current matrix support is not same-question philosophy scoring support for the candidate answer.",
        "",
        "| queue_id | node | heading | new why chars | new answer chars | evidence note |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in rows:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | {row['new_why_chars']} | {row['new_answer_chars']} | {row['evidence_note']} |"
        )
    (ROOT / f"P1_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_{RUN_DATE}.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )


def main() -> int:
    targets = load_queue_targets()
    write_draft_files(targets)
    docx = helper.current_docx()
    backup = docx.with_name(f"{docx.stem}_backup_before_p1_thickness_batch{BATCH}_{datetime.now():%Y%m%d_%H%M%S}.docx")
    shutil.copy2(docx, backup)

    doc = Document(str(docx))
    blocks = helper.entry_blocks(doc)
    applied: list[dict[str, object]] = []
    for target in targets:
        block = helper.matching_block(blocks, target)
        fields = block["fields"]
        field_paras = block["field_paras"]
        before_why = str(fields.get("why", ""))
        before_answer = str(fields.get("answer", ""))
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["why"])], "why", target["new_why"])
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["answer"])], "answer", target["new_answer"])
        applied.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "heading_para": block["heading_para"],
                "why_para": field_paras["why"],
                "answer_para": field_paras["answer"],
                "before_why_chars": len(before_why),
                "after_why_chars": len(target["new_why"]),
                "before_answer_chars": len(before_answer),
                "after_answer_chars": len(target["new_answer"]),
                "evidence_note": target["evidence_note"],
                "old_answer_excerpt": target["old_answer_excerpt"],
            }
        )

    doc.save(str(docx))
    with zipfile.ZipFile(docx) as z:
        zip_ok = "[Content_Types].xml" in z.namelist() and "word/document.xml" in z.namelist()

    result = {
        "updated": now(),
        "status": f"P1_BATCH{BATCH}_APPLIED_REQUIRES_RENDER_AND_EXTERNAL_RECHECK",
        "docx": json_path(docx),
        "backup_docx": json_path(backup),
        "zip_ok": zip_ok,
        "targets": len(targets),
        "applied": applied,
        "boundary": [
            "This is a local P1 thickness repair batch, not a terminal acceptance claim.",
            "PDF/render outputs must be regenerated after this DOCX edit.",
            "External current-version GPTPro/Claude review remains pending for the current post-repair artifact.",
        ],
    }
    (ROOT / f"P1_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_{RUN_DATE}.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / f"P1_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_{RUN_DATE}.csv",
        applied,
        [
            "queue_id",
            "heading",
            "node",
            "heading_para",
            "why_para",
            "answer_para",
            "before_why_chars",
            "after_why_chars",
            "before_answer_chars",
            "after_answer_chars",
            "evidence_note",
            "old_answer_excerpt",
        ],
    )
    lines = [
        f"# P1 Thickness Repair Batch{BATCH} Apply {RUN_DATE}",
        "",
        f"Updated: {result['updated']}",
        "",
        f"Status: `{result['status']}`",
        "",
        f"- DOCX: `{docx.name}`.",
        f"- Backup: `{backup.name}`.",
        f"- Targets applied: `{len(applied)}` / `{len(targets)}`.",
        f"- DOCX zip structural check: `{str(zip_ok).lower()}`.",
        "",
        "## Applied Rows",
        "",
        "| queue_id | node | heading | why chars | answer chars | paragraph refs |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in applied:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | "
            f"{row['before_why_chars']} -> {row['after_why_chars']} | "
            f"{row['before_answer_chars']} -> {row['after_answer_chars']} | "
            f"h={row['heading_para']}; why={row['why_para']}; answer={row['answer_para']} |"
        )
    lines.extend(
        [
            "",
            "## Open Gates",
            "",
            "- Re-run density queue after application.",
            "- Re-export PDF and render all pages.",
            "- Refresh Governor, Confucius, model ledger, format QA, and deferred upload scope.",
            "- Current-version GPTPro/Claude external review remains `real_call_pending` until a real new review is completed.",
        ]
    )
    (ROOT / f"P1_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_{RUN_DATE}.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
