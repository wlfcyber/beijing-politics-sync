from __future__ import annotations

import csv
import json
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

import apply_p0_thickness_batch01_20260525 as helper
from docx import Document


ROOT = Path(__file__).resolve().parent
BATCH = "17"
RUN_DATE = "20260526"

TARGET_TEXTS = {
    "T0207": {
        "node": "事物发展是前进性与曲折性的统一",
        "evidence_note": "矩阵 M1316：2026海淀期中 Q22(2) 正式阅卷细则 page_094 常见知识列明“发展观（前途光明、道路曲折）”。",
        "new_why": "“势不可挡”不是只写乐观信心，而是把民族复兴放在百年变局、风险挑战和长期奋斗中理解。正式细则把本问常见知识列到“发展观（前途光明、道路曲折）”，说明答题必须同时扣住两层：中华民族伟大复兴符合历史大势，前途光明；但推进复兴不可能一帆风顺，要在风险挑战、矛盾斗争和曲折考验中持续前进。只写“有信心”会变成口号，必须写出光明前途与曲折道路的统一。",
        "new_answer": "①事物发展的方向是前进的、上升的，道路是曲折的、迂回的。中华民族伟大复兴顺应中国社会发展和人类历史进步的大势，前途是光明的，所以“势不可挡”。②同时，复兴进程处在复杂国际环境和国内改革发展任务之中，会遇到风险挑战、矛盾斗争和阶段性困难，不能把前途光明理解为过程轻松。③正确做法是在坚定历史信心的同时增强忧患意识，准备经受风高浪急甚至惊涛骇浪的考验，在克服曲折中不断把民族复兴推向前进。",
    },
    "T0418": {
        "node": "社会发展的两大基本规律和基本矛盾",
        "evidence_note": "矩阵 M1318：2026海淀期中 Q22(2) 正式阅卷细则 page_094 常见知识列明“人类社会发展规律”，属于社会历史发展规律的正式宽口径支持。",
        "new_why": "本条回答“复兴不可挡”的深层根据，不能停留在情绪判断，而要说明民族复兴为什么合乎历史规律。正式细则列出“人类社会发展规律”，意味着可以从历史唯物主义角度解释：中国式现代化不是偶然选择，而是在生产力发展、制度完善、人民实践和社会进步的规律性运动中形成的方向。把复兴写成符合社会发展规律，才能说明“势”不是主观愿望，而是客观历史趋势。",
        "new_answer": "①社会历史发展有其客观规律，生产力与生产关系、经济基础与上层建筑的矛盾运动推动社会向前发展。②中华民族伟大复兴建立在中国社会生产力发展、制度优势发挥、改革开放深化和人民群众实践创造的基础上，符合中国社会发展的方向，也顺应人类社会追求现代化和文明进步的趋势。③因此，推进民族复兴要尊重和把握社会发展规律，把制度完善、生产力发展和人民实践统一起来，在顺应历史规律中增强历史主动，把光明前景转化为现实进程。",
    },
    "T0166": {
        "node": "发展的观点 / 发展的普遍性",
        "evidence_note": "矩阵 M0163/M0593：2026海淀一模 Q16 正式评分标准第1页写明可从“发展的观点、社会存在与社会意识、价值观、文化的功能”等角度作答。",
        "new_why": "人工智能时代的人文学科价值不能按静态有用无用来判断。材料把人文学科放在春秋战国、文艺复兴和人工智能时代等历史变革中比较，说明社会技术条件变化会不断提出新的思想、伦理和价值问题，人文学科也会在回应这些问题中推陈出新。正式评分标准明列“发展的观点”，所以本条应写出人文学科在发展中更新形态、拓展任务，而不是被人工智能简单替代。",
        "new_answer": "①发展具有普遍性，任何事物都处在变化发展之中，人文学科也会随着社会实践和技术条件变化而不断更新。②人工智能能够处理信息和生成文本，但它带来的价值冲突、伦理风险、意义追问和文明方向问题，反而要求人文学科在新的时代条件下发挥作用。③因此，人工智能时代不是人文学科终结的时代，而是人文学科回应新实践、提出新问题、形成新表达的时代；应在发展中理解其价值，使其继续为人的全面发展和文明进步提供思想支撑。",
    },
    "T0399": {
        "node": "社会存在与社会意识",
        "evidence_note": "矩阵 M0163/M0593：2026海淀一模 Q16 正式评分标准第1页写明可从“社会存在与社会意识”等角度作答。",
        "new_why": "材料中的人工智能技术进步、社会生产生活方式变化，属于社会存在层面的新变化；人文学科关于价值、伦理、历史、文化和人的意义的思考，属于社会意识。正式评分标准列明“社会存在与社会意识”，所以本条要写清双向关系：技术时代的新社会存在提出新的人文问题，同时先进的人文社会意识又能反过来规范、引导人工智能发展，给技术“赋魂”。",
        "new_answer": "①社会存在决定社会意识，社会意识是对社会存在的反映。人工智能改变生产生活方式和交往方式，催生新的伦理、价值、文化和意义问题，从而提出人文学科继续发展的现实需要。②社会意识又具有相对独立性并反作用于社会存在，先进的人文思想能够引导技术发展方向。③人文学科通过价值辨析、伦理规范、文化理解和历史反思，帮助人工智能服务人的全面发展和文明进步，避免技术脱离人的价值目标，这正是其不可替代的价值。",
    },
    "T0452": {
        "node": "价值观的导向作用",
        "evidence_note": "矩阵 M0163/M0593：2026海淀一模 Q16 正式评分标准第1页写明可从“价值观”等角度作答。",
        "new_why": "人工智能可以提高信息处理和内容生成效率，但它不能自行决定技术应服务什么价值、守住什么伦理底线、面向什么样的人类未来。题目材料强调“辨析价值冲突”“伦理意识”“文化理解”，正式评分标准又列出“价值观”角度，说明应从价值观导向作用落笔：人文学科的核心价值不在于替代机器功能，而在于为技术发展提供正确方向和评价尺度。",
        "new_answer": "①价值观对人们认识世界和改造世界具有重要导向作用，影响人们判断什么值得追求、技术应服务谁、发展应走向哪里。②人工智能时代会出现效率与公平、便利与隐私、创新与伦理、工具理性与人的主体性等价值冲突，这些问题不能靠算法自动解决。③人文学科能够提供价值辨析、伦理规范和文化理解，引导人工智能坚持以人为本、服务文明进步、防止技术异化，因此其价值引领功能不能被写作或信息处理功能替代。",
    },
    "T0193": {
        "node": "量变与质变 / 适度原则",
        "evidence_note": "矩阵 M0564：2026延庆一模 Q20 正式评分细则支持发展观、规律、矛盾观、发挥主观能动性等必修四角度，并在节点中登记量变质变链条。",
        "new_why": "材料中的“阶梯式递进”“承前启后”和“五年规划”都在说明现代化不是一次完成的跳跃，而是在连续阶段、连续任务和连续积累中走向整体跃升。每个五年规划都承担特定阶段的量的积累，到关键节点推动发展质量、结构和水平发生新的变化。细则支持发展观及相关规律角度，本条落在量变质变节点，应把“十五五”写成现代化进程中承接前期积累、推动新的质变的重要阶段。",
        "new_answer": "①量变是质变的必要准备，质变是量变的必然结果，事物发展表现为量变和质变的统一。②我国现代化建设通过一个个五年规划接续推进，每个阶段都在目标、任务、产业、民生、治理等方面积累条件和力量；这些连续积累不是零散工作，而是通向整体跃升的必要准备。③“十五五”承前启后，既要巩固前期发展成果，又要为基本实现社会主义现代化打下更坚实基础，在阶段性量的积累中推动中国式现代化实现新的质的提升。",
    },
    "T0142": {
        "node": "系统观念 / 系统优化",
        "evidence_note": "矩阵 M0068/M0164：2026石景山一模 Q17(1) 正式评分细则示例明确列出系统优化方法，并把规律、联系、发展、矛盾、文化传承创新作为可答角度。",
        "new_why": "中医药材料不是孤立介绍一个药方或一种技术，而是把人体整体、个体差异、辨证施治、医德伦理和文化传承放在同一体系中理解。正式细则示例列出系统优化方法，说明可以从系统观念回答：人体脏腑经络、病因病机、治疗方法和文化价值相互联系，只有立足整体、统筹各要素，才能说明中医药为什么既有医学价值又有文化价值。",
        "new_answer": "①系统观念要求立足整体，把各部分、各要素联系起来考察，并通过优化结构实现整体功能。②中医药强调人体是有机整体，脏腑经络、阴阳平衡、个体体质和病症变化相互联系，辨证论治正是根据整体关系统筹治疗，而不是头痛医头、脚痛医脚。③同时，中医药还把医学技术、生命观念、医德伦理和中华优秀传统文化连接起来，形成治病救人、传承文化、坚定文化自信的整体价值。这样理解，才能把中医药的哲学智慧和文化价值写完整。",
    },
    "T0282": {
        "node": "矛盾的特殊性 / 具体问题具体分析",
        "evidence_note": "矩阵 M0085-M0088/M0184：2025延庆一模 Q16 正式细则要求从实际出发了解老年人精神文化需要，并明示可从矛盾、价值观、文化发展基本途径等角度作答。",
        "new_why": "网络文艺适老化不能套用年轻化流量逻辑，因为中老年群体在媒介使用、审美经验、情感需要、操作能力和精神文化需求上有自身特点。正式细则要求从实际出发了解老年人精神文化需要，并允许从矛盾角度作答；这说明本条应落到矛盾特殊性和具体问题具体分析：不同受众构成不同具体矛盾，创作和传播方式必须随对象变化而变化。",
        "new_answer": "①矛盾具有特殊性，不同事物、不同群体、不同阶段有不同矛盾，必须坚持具体问题具体分析。②中老年题材网络文艺面对的是中老年人的精神文化需要，其审美经验、情感诉求、操作习惯和内容偏好不同于年轻群体，不能简单复制年轻化、流量化表达。③创作者应从中老年人的真实生活和精神需求出发，优化题材选择、人物塑造、叙事节奏、传播入口和互动方式，使作品既方便观看、愿意观看，又能表现银发力量、满足多层次精神文化需要。",
    },
}


def now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M +08")


def load_queue_targets() -> list[dict[str, str]]:
    queue_path = ROOT / "THICKNESS_REPAIR_PRIORITY_QUEUE_20260525.csv"
    wanted = set(TARGET_TEXTS)
    found: dict[str, dict[str, str]] = {}
    with queue_path.open(encoding="utf-8-sig", newline="") as f:
        for row in csv.DictReader(f):
            qid = row.get("queue_id", "")
            if qid in wanted and qid not in found:
                found[qid] = row
    missing = wanted - set(found)
    if missing:
        raise RuntimeError(f"Missing target queue rows: {sorted(missing)}")
    targets: list[dict[str, str]] = []
    for qid in TARGET_TEXTS:
        row = found[qid]
        text = TARGET_TEXTS[qid]
        targets.append(
            {
                "queue_id": qid,
                "heading": row["heading"],
                "node": text["node"],
                "old_answer_excerpt": row["answer_excerpt"],
                "evidence_note": text["evidence_note"],
                "new_why": text["new_why"],
                "new_answer": text["new_answer"],
            }
        )
    return targets


def write_csv(path: Path, rows: list[dict[str, object]], fields: list[str]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fields})


def write_draft_files(targets: list[dict[str, str]]) -> None:
    rows = []
    for target in targets:
        rows.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "evidence_note": target["evidence_note"],
                "new_why_chars": len(target["new_why"]),
                "new_answer_chars": len(target["new_answer"]),
                "new_why": target["new_why"],
                "new_answer": target["new_answer"],
            }
        )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_{RUN_DATE}.json").write_text(
        json.dumps({"updated": now(), "targets": rows}, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_{RUN_DATE}.csv",
        rows,
        ["queue_id", "heading", "node", "evidence_note", "new_why_chars", "new_answer_chars", "new_why", "new_answer"],
    )
    lines = [
        f"# P0 Thickness Repair Batch{BATCH} Draft {RUN_DATE}",
        "",
        f"Updated: {now()}",
        "",
        "Status: `DRAFT_READY_FOR_DOCX_APPLICATION`",
        "",
        "- Scope: 8 P0 subjective triple-thin rows with row-level matrix support from formal scoring/rubric sources.",
        "- Repair method: thicken only existing why/answer paragraphs inside existing framework nodes.",
        "- Evidence boundary: rows with only suite-level summary support are intentionally left for later source recheck.",
        "",
        "| queue_id | node | heading | new why chars | new answer chars | evidence note |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in rows:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | {row['new_why_chars']} | {row['new_answer_chars']} | {row['evidence_note']} |"
        )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_{RUN_DATE}.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )


def main() -> int:
    targets = load_queue_targets()
    write_draft_files(targets)
    docx = helper.current_docx()
    backup = docx.with_name(f"{docx.stem}_backup_before_p0_thickness_batch{BATCH}_{datetime.now():%Y%m%d_%H%M%S}.docx")
    shutil.copy2(docx, backup)

    doc = Document(str(docx))
    blocks = helper.entry_blocks(doc)
    applied: list[dict[str, object]] = []
    for target in targets:
        block = helper.matching_block(blocks, target)
        fields = block["fields"]
        field_paras = block["field_paras"]
        before_why = str(fields.get("why", ""))
        before_answer = str(fields.get("answer", ""))
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["why"])], "why", target["new_why"])
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["answer"])], "answer", target["new_answer"])
        applied.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "heading_para": block["heading_para"],
                "why_para": field_paras["why"],
                "answer_para": field_paras["answer"],
                "before_why_chars": len(before_why),
                "after_why_chars": len(target["new_why"]),
                "before_answer_chars": len(before_answer),
                "after_answer_chars": len(target["new_answer"]),
                "evidence_note": target["evidence_note"],
                "old_answer_excerpt": target["old_answer_excerpt"],
            }
        )

    doc.save(str(docx))
    with zipfile.ZipFile(docx) as z:
        zip_ok = "[Content_Types].xml" in z.namelist() and "word/document.xml" in z.namelist()

    result = {
        "updated": now(),
        "status": f"P0_BATCH{BATCH}_APPLIED_REQUIRES_RENDER_AND_EXTERNAL_RECHECK",
        "docx": str(docx),
        "backup_docx": str(backup),
        "zip_ok": zip_ok,
        "targets": len(targets),
        "applied": applied,
        "boundary": [
            "This is a local thickness repair batch, not a terminal acceptance claim.",
            "PDF/render outputs must be regenerated after this DOCX edit.",
            "External review remains pending for the current post-repair artifact.",
        ],
    }
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_{RUN_DATE}.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_{RUN_DATE}.csv",
        applied,
        [
            "queue_id",
            "heading",
            "node",
            "heading_para",
            "why_para",
            "answer_para",
            "before_why_chars",
            "after_why_chars",
            "before_answer_chars",
            "after_answer_chars",
            "evidence_note",
            "old_answer_excerpt",
        ],
    )
    lines = [
        f"# P0 Thickness Repair Batch{BATCH} Apply {RUN_DATE}",
        "",
        f"Updated: {result['updated']}",
        "",
        f"Status: `{result['status']}`",
        "",
        f"- DOCX: `{docx.name}`.",
        f"- Backup: `{backup.name}`.",
        f"- Targets applied: `{len(applied)}` / `{len(targets)}`.",
        f"- DOCX zip structural check: `{str(zip_ok).lower()}`.",
        "",
        "## Applied Rows",
        "",
        "| queue_id | node | heading | why chars | answer chars | paragraph refs |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in applied:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | "
            f"{row['before_why_chars']} -> {row['after_why_chars']} | "
            f"{row['before_answer_chars']} -> {row['after_answer_chars']} | "
            f"h={row['heading_para']}; why={row['why_para']}; answer={row['answer_para']} |"
        )
    lines.extend(
        [
            "",
            "## Boundary",
            "",
            "- This batch repairs only 8 additional P0 rows with row-level evidence support.",
            "- P0 rows with only suite-level summary support remain queued for later source recheck.",
            "- Remaining thickness queue, render QA, and model gates stay open until rechecked.",
        ]
    )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_{RUN_DATE}.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
