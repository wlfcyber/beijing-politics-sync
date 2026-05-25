from __future__ import annotations

import csv
import json
import shutil
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document


RECOVERY = Path(__file__).resolve().parent
RUN = RECOVERY.parent / "bixiu4_baodian_52_base_insert_second_mock_first_mock_audit_2026-05-24"
DELIVERY = RUN / "05_delivery"
MATRIX = RECOVERY / "FULL_QUESTION_COVERAGE_AND_PLACEMENT_MATRIX.csv"
BUNDLE = (
    Path.home()
    / "Desktop"
    / "beijing_politics_research"
    / "data"
    / "preprocessed_corpus"
    / "gpt_suite_bundles"
    / "2025各区模拟题__2025各区一模__2025延庆一模.md"
)
REPORT = RECOVERY / "YANQING_2025_YIMO_CANDIDATE_QUEUE_REPAIR_20260525.md"
MANIFEST = RECOVERY / "YANQING_2025_YIMO_CANDIDATE_QUEUE_REPAIR_20260525.json"


def current_docx() -> Path:
    files = [
        p
        for p in DELIVERY.glob("*.docx")
        if "_backup_" not in p.stem and not p.name.startswith("~$")
    ]
    if len(files) != 1:
        raise RuntimeError(f"Expected one current DOCX, found {files}")
    return files[0]


def paragraph_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs).strip()


def is_heading(paragraph) -> bool:
    name = getattr(paragraph.style, "name", "") or ""
    return name.startswith("Heading")


def find_heading_blocks(doc: Document, source: str, question: str) -> list[dict[str, object]]:
    paragraphs = list(doc.paragraphs)
    blocks: list[dict[str, object]] = []
    for idx, paragraph in enumerate(paragraphs):
        text = paragraph_text(paragraph)
        if not is_heading(paragraph) or source not in text or question not in text:
            continue
        end = len(paragraphs)
        for j in range(idx + 1, len(paragraphs)):
            if is_heading(paragraphs[j]):
                end = j
                break
        blocks.append(
            {
                "start": idx,
                "end_exclusive": end,
                "heading": text,
                "paragraphs": [paragraph_text(p) for p in paragraphs[idx:end]],
            }
        )
    return blocks


def remove_blocks(doc: Document, blocks: list[dict[str, object]]) -> int:
    paragraphs = list(doc.paragraphs)
    removed = 0
    ranges = sorted(
        [(int(block["start"]), int(block["end_exclusive"])) for block in blocks],
        reverse=True,
    )
    for start, end in ranges:
        for idx in range(end - 1, start - 1, -1):
            element = paragraphs[idx]._element
            element.getparent().remove(element)
            removed += 1
    return removed


def rewrite_matrix() -> dict[str, object]:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = MATRIX.with_name(
        f"FULL_QUESTION_COVERAGE_AND_PLACEMENT_MATRIX_backup_before_2025_yanqing_yimo_candidate_repair_{timestamp}.csv"
    )
    shutil.copy2(MATRIX, backup)
    with MATRIX.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        fieldnames = reader.fieldnames or []
        rows = list(reader)

    before = deepcopy(rows)
    changed = 0

    def set_row(row: dict[str, str], **updates: str) -> None:
        nonlocal changed
        local_changed = False
        for key, value in updates.items():
            if row.get(key, "") != value:
                row[key] = value
                local_changed = True
        if local_changed:
            changed += 1

    source_artifact = (
        "data/preprocessed_corpus/gpt_suite_bundles/2025各区模拟题__2025各区一模__2025延庆一模.md:"
    )
    q16_support = (
        "延庆一模Q16正式细则：网络文艺作品创作要从实际出发了解老年人精神文化需要；坚持正确价值导向、"
        "创作表现银发力量故事；满足多层次精神文化需要，并明示可从一切从实际出发、矛盾、价值观、文化发展基本途径等角度作答。"
    )
    q18_support = (
        "延庆一模Q18题干明确要求运用《逻辑与思维》知识、辩证思维应对低空经济挑战，属选必三边界，不进入必修四正文。"
    )
    q21_support = (
        "延庆一模Q21细则含人民群众、社会主义核心价值观、民族精神或文化自信等必修四相关落点；党的领导、一国两制等非必修四点只作边界说明。"
    )

    for row in rows:
        if row.get("题源") != "2025延庆一模":
            continue
        q = row.get("题号", "")
        if q in {f"Q{i}" for i in range(1, 16)}:
            set_row(
                row,
                **{
                    "是否进宝典": "否：选择题候选词命中，非主观题正文",
                    "宝典节点": "选择题客观题边界",
                    "细则支持原理": "延庆一模官方答案键Q1-Q15仅支持客观题正误判定；本行是早期术语命中，不是主观题评分细则落点。",
                    "证据等级": "官方答案键+题面；非主观题评分细则",
                    "是否误放": "否：未进正文",
                    "是否需补厚": "否",
                    "当前处理": "CHOICE_TERM_CANDIDATE_REVIEWED_NON_BODY_YANQING_2025",
                    "备注": "按选择题边界裁决：不把客观题答案键冒充主观题细则。",
                    "source_artifact": source_artifact + "21-43",
                },
            )
        elif q.startswith("Q16") or q == "Qunknown":
            set_row(
                row,
                **{
                    "是否进宝典": "是：当前DOCX已有Q16覆盖，不重复新增",
                    "宝典节点": "一切从实际出发/矛盾/价值观/文化发展基本途径",
                    "细则支持原理": q16_support,
                    "证据等级": "正式细则",
                    "是否误放": "否",
                    "是否需补厚": "否",
                    "当前处理": "COVERED_BY_CURRENT_DOCX_YANQING_2025_Q16",
                    "备注": "早期候选已并入Q16正式细则覆盖记录；不再保留待核/补厚状态。",
                    "source_artifact": source_artifact + "59-61",
                },
            )
        elif q == "Q17":
            set_row(
                row,
                **{
                    "是否进宝典": "否：政治与法治/民营经济立法边界",
                    "宝典节点": "-",
                    "细则支持原理": "延庆一模Q17围绕民营经济促进法、协商民主、科学立法、民主立法、依法立法，不作为必修四哲学与文化正文条目。",
                    "证据等级": "正式细则+题干模块边界",
                    "是否误放": "否：未进正文",
                    "是否需补厚": "否",
                    "当前处理": "MODULE_BOUNDARY_EXCLUDED_YANQING_2025_Q17_POLITICS_LAW",
                    "备注": "候选词命中来自跨模块细则片段，已按模块边界排除。",
                    "source_artifact": source_artifact + "63-69",
                },
            )
        elif q == "Q18":
            set_row(
                row,
                **{
                    "是否进宝典": "否：选必三《逻辑与思维》边界，已从当前DOCX撤下",
                    "宝典节点": "-",
                    "细则支持原理": q18_support,
                    "证据等级": "正式细则/题干模块边界",
                    "是否误放": "否：旧正文误放已撤下",
                    "是否需补厚": "否",
                    "当前处理": "REMOVED_FROM_CURRENT_DOCX_YANQING_2025_Q18_XUANBISAN_BOUNDARY",
                    "备注": "当前DOCX原有两个Q18正文块，已按选必三边界撤下；保留为排除证据记录。",
                    "source_artifact": source_artifact + "72-88; CURRENT_DOCX_2025_YANQING_Q18_CONTEXT_20260525.md",
                },
            )
        elif q == "Q19":
            set_row(
                row,
                **{
                    "是否进宝典": "否：法律与生活/纠纷调解边界",
                    "宝典节点": "-",
                    "细则支持原理": "延庆一模Q19围绕网络购物纠纷、调解、诚信原则、合同解除和证据意识，属法律模块，不作为必修四正文条目。",
                    "证据等级": "正式细则+题干模块边界",
                    "是否误放": "否：未进正文",
                    "是否需补厚": "否",
                    "当前处理": "MODULE_BOUNDARY_EXCLUDED_YANQING_2025_Q19_LAW",
                    "备注": "候选词命中来自法律题细则，已按模块边界排除。",
                    "source_artifact": source_artifact + "90-94",
                },
            )
        elif q == "Q20":
            set_row(
                row,
                **{
                    "是否进宝典": "否：经济/当代国际政治经济边界",
                    "宝典节点": "-",
                    "细则支持原理": "延庆一模Q20第(1)问为产业链供应链与经济高质量发展，第(2)问为全球产供链、经济全球化、多极化和人类命运共同体，非必修四正文条目。",
                    "证据等级": "正式细则+题干模块边界",
                    "是否误放": "否：未进正文",
                    "是否需补厚": "否",
                    "当前处理": "MODULE_BOUNDARY_EXCLUDED_YANQING_2025_Q20_ECON_IR",
                    "备注": "候选词命中来自跨模块细则，已按模块边界排除。",
                    "source_artifact": source_artifact + "97-106",
                },
            )
        elif q == "Q21":
            set_row(
                row,
                **{
                    "是否进宝典": "是：当前DOCX已有Q21必修四相关覆盖",
                    "宝典节点": "人民群众/社会主义核心价值观/民族精神或文化自信",
                    "细则支持原理": q21_support,
                    "证据等级": "正式细则",
                    "是否误放": "否",
                    "是否需补厚": "否",
                    "当前处理": "COVERED_BY_CURRENT_DOCX_YANQING_2025_Q21_BIXIU4_RELEVANT_POINTS",
                    "备注": "只保留必修四相关评分点；党的领导、一国两制等点不作为必修四节点。",
                    "source_artifact": source_artifact + "108-110; CURRENT_DOCX_2025_YANQING_Q21_CONTEXT_20260525.md",
                },
            )
        elif q == "SUITE_LEVEL":
            set_row(
                row,
                **{
                    "是否进宝典": "套卷级索引：已拆为逐题处置",
                    "宝典节点": "SUITE_LEVEL_SUMMARY",
                    "细则支持原理": "2025延庆一模已按Q1-Q21逐题裁决；套卷级行只作索引，不能替代逐题证据。",
                    "证据等级": "套卷索引+逐题裁决已完成",
                    "是否误放": "不适用",
                    "是否需补厚": "否",
                    "当前处理": "SUITE_LEVEL_INDEX_SUPERSEDED_BY_ROW_REVIEW_YANQING_2025",
                    "备注": "风险队列中的套卷级补厚标记已由逐题回源处置替代。",
                    "source_artifact": source_artifact + "7-110",
                },
            )

    with MATRIX.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)

    remaining_yanqing_risk_flags = []
    for row in rows:
        if row.get("题源") != "2025延庆一模":
            continue
        text = " ".join(row.values())
        if any(token in text for token in ["待核", "需逐题", "term_hit_needs_review", "needs_review"]):
            remaining_yanqing_risk_flags.append(
                {
                    "题号": row.get("题号", ""),
                    "当前处理": row.get("当前处理", ""),
                    "是否进宝典": row.get("是否进宝典", ""),
                    "证据等级": row.get("证据等级", ""),
                }
            )

    return {
        "matrix_backup": str(backup),
        "rows_changed": changed,
        "rows_total": len(rows),
        "rows_before": len(before),
        "remaining_yanqing_risk_flags": remaining_yanqing_risk_flags,
    }


def write_q21_context(docx: Path, block: dict[str, object] | None) -> Path:
    out = RECOVERY / "CURRENT_DOCX_2025_YANQING_Q21_CONTEXT_20260525.md"
    lines = [
        "# Current DOCX Context - 2025 Yanqing Q21",
        "",
        f"- DOCX: `{docx}`",
        f"- Target headings: `{0 if block is None else 1}`",
        "",
    ]
    if block is not None:
        lines.append("## Q21 Block")
        lines.append("")
        lines.append(f"- heading: `{block['heading']}`")
        for idx, text in enumerate(block["paragraphs"]):
            lines.append(f"- p{idx}: {text}")
    out.write_text("\n".join(lines) + "\n", encoding="utf-8", newline="\n")
    return out


def main() -> None:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = current_docx()
    doc = Document(str(docx_path))
    q18_blocks_before = find_heading_blocks(doc, "2025延庆一模", "第18题")
    q21_blocks_before = find_heading_blocks(doc, "2025延庆一模", "第21题")
    q21_context = write_q21_context(docx_path, q21_blocks_before[0] if q21_blocks_before else None)

    docx_backup = None
    removed_paragraphs = 0
    if q18_blocks_before:
        docx_backup = docx_path.with_name(
            f"{docx_path.stem}_backup_before_2025_yanqing_q18_xuanbisan_removal_{timestamp}.docx"
        )
        shutil.copy2(docx_path, docx_backup)
        removed_paragraphs = remove_blocks(doc, q18_blocks_before)
        doc.save(str(docx_path))

    post_doc = Document(str(docx_path))
    q18_blocks_after = find_heading_blocks(post_doc, "2025延庆一模", "第18题")
    q21_blocks_after = find_heading_blocks(post_doc, "2025延庆一模", "第21题")
    matrix_info = rewrite_matrix()

    manifest = {
        "status": "YANQING_2025_YIMO_CANDIDATE_QUEUE_REPAIRED",
        "timestamp": timestamp,
        "docx": str(docx_path),
        "docx_backup": str(docx_backup) if docx_backup else None,
        "q18_blocks_before": len(q18_blocks_before),
        "q18_headings_before": [b["heading"] for b in q18_blocks_before],
        "q18_removed_paragraphs": removed_paragraphs,
        "q18_blocks_after": len(q18_blocks_after),
        "q21_blocks_after": len(q21_blocks_after),
        "q21_context": str(q21_context),
        "bundle": str(BUNDLE),
        **matrix_info,
    }
    MANIFEST.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    lines = [
        "# 2025延庆一模候选队列修复记录",
        "",
        "- Status: `YANQING_2025_YIMO_CANDIDATE_QUEUE_REPAIRED`.",
        f"- Current DOCX: `{docx_path}`.",
        f"- DOCX backup before removal: `{docx_backup}`.",
        f"- Matrix backup before rewrite: `{matrix_info['matrix_backup']}`.",
        f"- Source bundle: `{BUNDLE}`.",
        "",
        "## DOCX 正文边界修复",
        "",
        f"- Q18 headings before repair: `{len(q18_blocks_before)}`.",
        f"- Q18 removed paragraph count: `{removed_paragraphs}`.",
        f"- Q18 headings after repair: `{len(q18_blocks_after)}`.",
        "- Adjudication: Q18 is excluded because the question explicitly asks for 《逻辑与思维》/辩证思维; it is an 选必三 boundary item, not a 必修四 philosophy/culture body item.",
        f"- Q21 current context record: `{q21_context}`.",
        "",
        "## Matrix 修复范围",
        "",
        f"- Rows changed: `{matrix_info['rows_changed']}`.",
        "- Q1-Q15: closed as choice-question answer-key/term-hit boundary, non-body.",
        "- Q16: retained as current-DOCX coverage supported by formal Q16 scoring rules, with no duplicate insertion.",
        "- Q17: excluded as politics/law-making module boundary.",
        "- Q18: excluded and removed from current DOCX as 选必三 boundary.",
        "- Q19: excluded as law/dispute module boundary.",
        "- Q20: excluded as economics/international module boundary.",
        "- Q21: retained only for Bixiu4-relevant scoring points: 人民群众、社会主义核心价值观、民族精神/文化自信.",
        "- SUITE_LEVEL rows: downgraded to index-only after row-level review.",
        "",
        "## Remaining Boundaries",
        "",
        "- This is a local source-and-DOCX repair, not an external-model acceptance.",
        "- GPTPro web external review remains `real_call_pending`.",
        "- Claude Opus web/app external review remains `real_call_pending`; direct `https://claude.ai` auto-login path is the required retry route.",
        "- ClaudeCode post-repair evidence remains subject to `BLOCKED_MODEL_CONFIRMATION_REQUIRED` until model confirmation is clean.",
        "- No GitHub push was attempted; final upload is deferred under ORDER_063 after all active Beijing politics lines reach terminal state.",
    ]
    if matrix_info["remaining_yanqing_risk_flags"]:
        lines.extend(["", "## Remaining Yanqing Flags", ""])
        for item in matrix_info["remaining_yanqing_risk_flags"]:
            lines.append(
                f"- `{item['题号']}` / `{item['当前处理']}` / `{item['是否进宝典']}` / `{item['证据等级']}`"
            )
    REPORT.write_text("\n".join(lines) + "\n", encoding="utf-8", newline="\n")
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
