from __future__ import annotations

import csv
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document


RECOVERY = Path(__file__).resolve().parent
RUN = RECOVERY.parent / "bixiu4_baodian_52_base_insert_second_mock_first_mock_audit_2026-05-24"
DELIVERY = RUN / "05_delivery"
MATRIX = RECOVERY / "FULL_QUESTION_COVERAGE_AND_PLACEMENT_MATRIX.csv"

SHIJINGSHAN_REPORT = RECOVERY / "SHIJINGSHAN_2024_YIMO_Q16_SOURCE_EXHAUSTION_AND_DOCX_REMOVAL_20260525.md"
XICHENG_REPORT = RECOVERY / "XICHENG_2026_ERMO_Q20_FORMAL_PINGBIAO_REPAIR_20260525.md"

XICHENG_RUBRIC_PAGE = (
    RUN
    / "99_logs"
    / "weak_gate_sources"
    / "renders"
    / "xicheng_rubric"
    / "page_014.png"
)

SHIJINGSHAN_ROW_IDS = {"M0144", "M0195", "M0201", "M0315"}
XICHENG_ROW_ID = "M0771"
SHIJINGSHAN_HEADING_MARKER = "2024石景山一模 第16题"


def current_docx() -> Path:
    candidates = [
        p
        for p in DELIVERY.glob("*.docx")
        if "_backup_" not in p.stem and not p.name.startswith("~$")
    ]
    if len(candidates) != 1:
        raise RuntimeError(f"Expected one current DOCX, found {candidates}")
    return candidates[0]


def remove_shijingshan_q16_entries(timestamp: str) -> dict[str, object]:
    docx = current_docx()
    backup = docx.with_name(f"{docx.stem}_backup_before_shijingshan_q16_non_rubric_removal_{timestamp}.docx")
    shutil.copy2(docx, backup)

    doc = Document(str(docx))
    paragraphs = list(doc.paragraphs)
    targets: list[tuple[int, int, str, list[str]]] = []
    for idx, para in enumerate(paragraphs):
        text = para.text.strip()
        if para.style.name == "Heading 3" and SHIJINGSHAN_HEADING_MARKER in text:
            end = len(paragraphs)
            for next_idx in range(idx + 1, len(paragraphs)):
                if paragraphs[next_idx].style.name.startswith("Heading "):
                    end = next_idx
                    break
            targets.append((idx, end, text, [p.text.strip() for p in paragraphs[idx:end]]))

    for start, end, _heading, _block in reversed(targets):
        for para in paragraphs[start:end]:
            parent = para._element.getparent()
            if parent is not None:
                parent.remove(para._element)

    doc.save(str(docx))

    verify = Document(str(docx))
    remaining = [
        p.text.strip()
        for p in verify.paragraphs
        if p.style.name == "Heading 3" and SHIJINGSHAN_HEADING_MARKER in p.text
    ]
    return {
        "docx": str(docx),
        "backup": str(backup),
        "removed_count": len(targets),
        "removed_headings": [item[2] for item in targets],
        "remaining_headings": remaining,
    }


def update_matrix(timestamp: str) -> dict[str, object]:
    backup = MATRIX.with_name(
        f"FULL_QUESTION_COVERAGE_AND_PLACEMENT_MATRIX_backup_before_remaining_non_rubric_repair_{timestamp}.csv"
    )
    shutil.copy2(MATRIX, backup)
    with MATRIX.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        fieldnames = reader.fieldnames or []
        rows = list(reader)

    changed: list[str] = []
    for row in rows:
        row_id = row.get("matrix_row_id", "")
        if row_id in SHIJINGSHAN_ROW_IDS:
            row["题型或模块判断"] = "main_question_non_rubric_boundary_removed_from_current_docx"
            row["是否进宝典"] = "否：2024石景山一模Q16无正式细则，当前DOCX两处残留已撤下"
            row["宝典节点"] = "非正文边界：教师版宽角度不得作为细则节点"
            row["细则支持原理"] = (
                "本地原始目录只见教师版带答案DOCX与讲评PPT，细则目录为空；"
                "缓存清单同样只列上述两源。Q16仅有教师版宽角度：文化继承与发展、文化自信、认识的作用、联系、发展等，不能冒充评分细则。"
            )
            row["证据等级"] = "无正式细则；教师版宽角度仅作非正文边界"
            row["是否误放"] = "否：已从当前DOCX撤下"
            row["是否需补厚"] = "否：无细则，不补正文"
            row["当前处理"] = "REMOVED_FROM_CURRENT_DOCX_SHIJINGSHAN_2024_Q16_NON_RUBRIC_BOUNDARY"
            row["备注"] = (
                "M0144/M0195/M0201/M0315为同一Q16多生产线镜像；已按硬规则撤出正文。"
                "如后续用户提供正式细则或评标文件，可单独重开。"
            )
            row["source_artifact"] = (
                "SHIJINGSHAN_2024_YIMO_Q16_SOURCE_EXHAUSTION_AND_DOCX_REMOVAL_20260525.md; "
                "BATCH18_2024_SHIJINGSHAN_YIMO_SOURCE_TRANSCRIPTION_20260525.md"
            )
            changed.append(row_id)
        elif row_id == XICHENG_ROW_ID:
            row["题型或模块判断"] = "main_question_formal_pingbiao_broad_scoring_supported"
            row["是否进宝典"] = "是：本批新增进入当前DOCX/PDF正文"
            row["宝典节点"] = "一切从实际出发 / 实事求是；人民群众；发展的观点 / 发展的普遍性"
            row["细则支持原理"] = (
                "2026西城二模评标PDF渲染页第14页Q20细则："
                "①正确政绩观说明可围绕人民立场、人民至上、以人民为中心、党的性质宗旨、执政理念、群众观点等；"
                "②重要性分析要求能结合哲学观点，从对党、对人民、对国家事业等角度展开，含哲学观点2分+具体分析2分；"
                "③实践要求写明科学决策：解放思想实事求是、尊重规律、发展的观点、立足长远、与时俱进等；"
                "实干担当：价值观引领、群众路线、求真务实、实践第一、先锋模范作用。"
            )
            row["证据等级"] = "正式评标PDF渲染证据-哲学观点明示+材料支撑+当前DOCX正文覆盖"
            row["是否误放"] = "否"
            row["是否需补厚"] = "否"
            row["当前处理"] = "REPAIRED_WITH_FORMAL_PINGBIAO_RENDERED_Q20_EVIDENCE"
            row["备注"] = (
                "该评标为宽角度/模块组合评分，不是逐点唯一答案；当前正文落点限于评标明示的群众观点、实事求是、发展观点等。"
            )
            row["source_artifact"] = (
                "XICHENG_2026_ERMO_Q20_FORMAL_PINGBIAO_REPAIR_20260525.md; "
                "reports/bixiu4_baodian_52_base_insert_second_mock_first_mock_audit_2026-05-24/"
                "99_logs/weak_gate_sources/renders/xicheng_rubric/page_014.png"
            )
            changed.append(row_id)

    with MATRIX.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)

    return {"backup": str(backup), "changed": changed}


def write_reports(timestamp: str, removal: dict[str, object], matrix_update: dict[str, object]) -> None:
    SHIJINGSHAN_REPORT.write_text(
        "\n".join(
            [
                "# 2024石景山一模 Q16 Source Exhaustion And DOCX Removal",
                "",
                f"timestamp: `{timestamp}`",
                "status: `REMOVED_FROM_CURRENT_DOCX_NON_RUBRIC_BOUNDARY`",
                "",
                "## Source Exhaustion",
                "",
                "- Raw directory checked: `C:\\Users\\Administrator\\Desktop\\2024各区模拟题\\2024各区一模\\2024石景山一模`.",
                "- Files present: `其他材料\\2024年石景山一模.pptx`; `试卷\\2024北京石景山高三一模政治（教师版带答案）.docx`.",
                "- `细则` folder is present but contains no source file in the local listing.",
                "- Preprocessed manifest contains only the PPT and teacher-version DOCX for this suite.",
                "- Q16 only has teacher-version broad angle wording; no formal scoring rubric, marking report, or user-confirmed scoring file was found.",
                "",
                "## Decision",
                "",
                "- The two current DOCX entries headed `2024石景山一模 第16题（主观题）` were not kept as final正文证据.",
                f"- Removed headings: `{removal['removed_count']}`.",
                f"- Remaining matching headings after removal: `{len(removal['remaining_headings'])}`.",
                f"- DOCX backup: `{removal['backup']}`.",
                "- Matrix rows M0144/M0195/M0201/M0315 are now marked as non-body boundary rows.",
                "",
                "## Boundary",
                "",
                "- This is not a content rejection of the teacher-version answer itself; it is a scoring-evidence boundary.",
                "- If a formal 石景山 Q16 scoring file is later supplied, this question can be reopened and reinserted with proper evidence.",
            ]
        )
        + "\n",
        encoding="utf-8",
        newline="\n",
    )

    XICHENG_REPORT.write_text(
        "\n".join(
            [
                "# 2026西城二模 Q20 Formal Pingbiao Repair",
                "",
                f"timestamp: `{timestamp}`",
                "status: `REPAIRED_WITH_FORMAL_PINGBIAO_RENDERED_Q20_EVIDENCE`",
                "",
                "## Source",
                "",
                "- Original file: `C:\\Users\\Administrator\\Desktop\\2026各区模拟题\\2026各区二模\\2026西城二模\\细则\\西城二模评标.pdf`.",
                "- Text layer in source bundle is empty, so the accepted evidence is the rendered page image.",
                f"- Rendered evidence page: `{XICHENG_RUBRIC_PAGE}`.",
                "",
                "## Q20 Evidence Transcription",
                "",
                "- Q20 asks for understanding of牢固树立和践行正确政绩观, 10分.",
                "- Part 1: correct政绩观 explanation, 2分. The page lists人民立场、人民至上、以人民为中心、党的性质宗旨、执政理念、群众观点等, 任意一个2分.",
                "- Part 2: importance analysis, 4分. The page requires结合哲学观点, 从对党、对人民、对国家事业等角度展开分析, 哲学观点2分+具体分析2分.",
                "- Part 3: practice requirement, 4分. It lists科学决策: 解放思想实事求是、尊重规律、发展的观点、立足长远、与时俱进等; and实干担当: 价值观引领、群众路线、求真务实、实践第一、先锋模范作用.",
                "",
                "## Matrix Impact",
                "",
                "- M0771 is upgraded from teacher-version broad answer boundary to formal rendered评标 evidence.",
                "- The evidence remains broad-angle scoring support, not a claim of point-by-point unique answer.",
                f"- Matrix backup: `{matrix_update['backup']}`.",
            ]
        )
        + "\n",
        encoding="utf-8",
        newline="\n",
    )


def main() -> None:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    removal = remove_shijingshan_q16_entries(timestamp)
    matrix_update = update_matrix(timestamp)
    write_reports(timestamp, removal, matrix_update)
    print(f"removed_shijingshan_q16={removal['removed_count']}")
    print(f"remaining_shijingshan_q16={len(removal['remaining_headings'])}")
    print(f"matrix_changed={','.join(matrix_update['changed'])}")
    print(f"docx={removal['docx']}")
    print(f"matrix_backup={matrix_update['backup']}")


if __name__ == "__main__":
    main()
