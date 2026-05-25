from __future__ import annotations

import csv
import json
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

import apply_p0_thickness_batch01_20260525 as helper
from docx import Document


ROOT = Path(__file__).resolve().parent
BATCH = "10"

TARGETS = [
    {
        "queue_id": "T0380",
        "heading": "33. 2026朝阳期中 第18题（主观题）",
        "node": "实践与认识（总）",
        "old_answer_excerpt": "AI可以提供辅助性情绪疏导，但认识和能力的形成要回到实践。人应在真实互动中理解他人、处理冲突、提升情绪调节能力，不能用虚拟回应替代现实实践。",
        "evidence_note": "正式评分细则第60-62、78-81行在反对理由中列实践等，并指向现实交往能力；当前矩阵 M1246 以正式细则宽角度/术语支持该节点。",
        "new_why": "材料讨论 AI 情绪陪伴的边界，关键不是否认技术能提供即时回应，而是指出情绪理解、关系处理和能力成长不能脱离真实交往。正式细则在反对理由中列“实践”等，并指向过度依赖会削弱现实交往能力。看到“真实互动”“处理冲突”“提升能力”，就应从实践与认识、能力在实践中形成来组织答案。",
        "new_answer": "①实践是认识和能力形成发展的基础，人的情绪理解、沟通判断和关系处理能力必须在真实生活中训练。②AI 可以作为辅助性情绪疏导工具，提供表达空间和初步回应，但不能替代人与人之间的现实互动、责任承担和冲突解决。③正确使用 AI，要把技术辅助同真实实践结合起来，在现实交往中提升情绪调节和解决问题的能力。",
    },
    {
        "queue_id": "T0178",
        "heading": "33. 2026朝阳期末 第21题（主观题）",
        "node": "系统观念 / 系统优化",
        "old_answer_excerpt": "认识四大优势要坚持系统优化方法，把制度优势、超大规模市场优势、完整产业体系优势和人才优势作为相互依存的系统来统筹，推动各要素协同发力。",
        "evidence_note": "正式阅卷细则第152-153行明列系统优化，题面第469-470行写四大优势相互依存、协同赋能；当前矩阵 M1290 以正式细则支持该节点。",
        "new_why": "题面把制度优势、超大规模市场优势、完整产业体系优势和人才优势写成相互依存、协同赋能的有机整体，细则第152-153行也明列系统优化。由此可知，不能把四大优势拆成彼此孤立的条件罗列，而要看它们如何在整体结构中相互支撑、互相转化，共同形成推进中国式现代化的系统合力。",
        "new_answer": "①系统优化方法要求立足整体，把各要素放在相互联系、相互作用的结构中统筹。②四大优势不是孤立存在的：制度优势提供组织保障，超大规模市场提供需求和空间，完整产业体系提供供给基础，人才优势提供创新动力。③把这些优势作为系统协同起来，才能实现要素互补、协同赋能，形成推进中国式现代化的整体效能。",
    },
    {
        "queue_id": "T0048",
        "heading": "33. 2026海淀期末 第17题（主观题）",
        "node": "主观能动性 / 意识的能动作用",
        "old_answer_excerpt": "红色文化能够通过精神感召和意义阐释发挥意识的能动作用，激励青年理解历史、坚定理想并主动参与实践。",
        "evidence_note": "正式评分标准 page_002 Q17 哲学知识栏列出“意识的能动作用/认识的作用/价值观的导向作用”；当前矩阵 M1354 以正式评分标准支持该节点。",
        "new_why": "红色文化滋养青年，并不只是提供历史知识，更通过精神感召、意义阐释和榜样力量影响青年怎样认识历史、理解责任、选择行动。正式评分标准 page_002 Q17 哲学知识栏列出“意识的能动作用/认识的作用”。因此应说明红色文化作为正确意识和精神力量，如何促进青年坚定理想并投身实践。",
        "new_answer": "①意识对改造世界具有能动作用，正确意识能够指导实践、激励行动。②红色文化通过讲述革命历史、阐释精神意义、呈现先进榜样，帮助青年形成正确历史认识和责任意识，坚定理想信念。③这种精神力量不是停留在感动上，而是引导青年把价值认同转化为学习奋斗、服务社会和参与实践的主动行动。",
    },
    {
        "queue_id": "T0491",
        "heading": "33. 2026海淀期末 第21题（主观题）",
        "node": "人民群众",
        "old_answer_excerpt": "人民群众是历史的主体。中国式现代化要为了人民、依靠人民、成果由人民共享，在群众观点和群众路线中实现活力有序释放。",
        "evidence_note": "正式评分标准 page_006 Q21“以人民为中心”角度写明人民群众是历史的主体，坚持群众观点、群众路线；当前矩阵 M1362 以正式评分标准支持该节点。",
        "new_why": "Q21 围绕活力与秩序展开，但评分标准 page_006 同时给出“以人民为中心”角度，明写人民群众是历史的主体、坚持群众观点和群众路线。材料若只谈治理技术或制度安排，会漏掉活力从哪里来、秩序为谁服务。应把中国式现代化的动力来源和价值归宿落到人民主体地位上。",
        "new_answer": "①人民群众是社会历史的主体，是推动中国式现代化的根本力量。②实现活力有序释放，必须坚持群众观点和群众路线，尊重人民创造，听取人民意见，把人民的积极性、主动性、创造性组织到现代化建设中。③中国式现代化既要依靠人民推进，也要让发展成果由人民共享，使治理秩序服务人民利益和人民生活。",
    },
    {
        "queue_id": "T0207",
        "heading": "34. 2025海淀期末 第16题（主观题）",
        "node": "发展的观点 / 发展的普遍性",
        "old_answer_excerpt": "打造中医药新时尚，要坚持发展的观点，在传承药食同源等中医药文化根脉的基础上，根据当代生活方式和健康需求推动其发展，使中药元素进入曲奇、面包、吐司等现代食品形态，让传统中医药文化在创新发展中焕发新活力。",
        "evidence_note": "PPT细则第37-44行明列“坚持发展的观点”，当前矩阵 M1094 以强细则/PPT评标细则支持该节点。",
        "new_why": "“中医药新时尚”不是简单把旧元素搬进新包装，而是在保留药食同源等文化根脉的同时，根据当代生活方式、健康需求和消费场景进行转化。PPT细则明列“坚持发展的观点”，说明要从传统文化不断适应时代、获得新形态和新生命力来作答。看到“曲奇、面包、吐司”等现代形态，就应落到发展。",
        "new_answer": "①事物是变化发展的，传统中医药文化也要在传承中发展、在发展中创新。②打造中医药新时尚，要守住药食同源等文化根脉，同时根据当代健康需求和生活方式，把中药元素融入曲奇、面包、吐司等现代食品形态。③这样既不是割裂传统，也不是原样照搬，而是在创新发展中让传统中医药文化焕发新的活力。",
    },
    {
        "queue_id": "T0526",
        "heading": "34. 2025西城期末 第18题（主观题）",
        "node": "价值观的导向作用",
        "old_answer_excerpt": "价值观对人们认识世界和改造世界具有导向作用。看待技术恐惧，既要尊重技术发展规律，也要坚持科技以人为本，把人民利益作为技术治理的重要尺度，关注受技术变革冲击的人群，让人工智能发展服务人的全面发展和社会公共利益。",
        "evidence_note": "正式细则第116-118行明列“价值观具有导向作用”，科技要以人为本；当前矩阵 M1141 以强细则支持该节点。",
        "new_why": "技术恐惧表面上是对人工智能等新技术的担忧，本质上涉及用什么价值尺度评价技术、治理技术、安排技术发展方向。正式细则第116-118行明列“价值观具有导向作用”，并要求科技以人为本。因此作答不能只写技术会进步，还要说明正确价值观如何引导我们把人民利益、人的全面发展和公共利益放在技术治理中心。",
        "new_answer": "①价值观对人们认识世界和改造世界具有导向作用，会影响人们对技术的评价、选择和治理。②面对技术恐惧，既要尊重技术发展规律，又要坚持科技以人为本，把人民利益和社会公共利益作为重要尺度。③这样才能关注受技术变革冲击的人群，防止技术发展偏离人的需要，使人工智能等新技术服务人的全面发展和社会进步。",
    },
    {
        "queue_id": "T0107",
        "heading": "34. 2026丰台期末 第22题（主观题）",
        "node": "联系的普遍性 / 联系的观点（总）",
        "old_answer_excerpt": "理解五年规划要坚持联系观点。它把长期目标与阶段任务、中央统筹与地方落实、部门协同与社会参与联系起来，形成持续推进现代化建设的治理链条。",
        "evidence_note": "正式细则第452行明列联系的观点，材料涉及长期战略、中期灵活、中央地方部门协同；当前矩阵 M1215 以正式细则支持该节点。",
        "new_why": "五年规划的意义不在于孤立列出某一时期指标，而在于把长期现代化目标、阶段性任务、中央统筹、地方落实、部门协同和社会参与连接起来。正式细则第452行明列联系的观点，材料也突出长期战略与中期安排相衔接。由此可知，答题要说明规划怎样通过多方面联系形成连续推进的治理链条。",
        "new_answer": "①联系具有普遍性，要用联系的观点看问题。②五年规划把长期目标与阶段任务联系起来，把中央统筹与地方落实、部门协同与社会参与联系起来，使战略目标、政策安排和具体行动相互衔接。③这种联系能够避免短期化和碎片化治理，形成持续推进现代化建设的治理链条和工作合力。",
    },
    {
        "queue_id": "T0492",
        "heading": "34. 2026西城期末 第16题第（1）问（主观题）",
        "node": "人民群众",
        "old_answer_excerpt": "纪录片创作要坚持人民主体，贴近人民精神生活，展现人民奋发有为的精神风貌，因为人民群众是社会精神财富的创造者。",
        "evidence_note": "正式评分细则 page_001 Q16(1) 角度1变通写明“人民群众是精神财富的创造者，要坚持人民主体”；当前矩阵 M1392 以正式评分标准宽角度/术语支持该节点。",
        "new_why": "纪录片创作不是只追求镜头技巧或国际传播效果，关键在于谁是精神财富的创造主体、作品应贴近谁的生活。正式评分细则 page_001 Q16(1) 变通写明“人民群众是精神财富的创造者，要坚持人民主体”。因此应把纪录片的真实感、感染力和精神风貌，落到人民群众创造历史和精神财富的主体地位上。",
        "new_answer": "①人民群众是社会历史的主体，也是社会精神财富的创造者。②纪录片创作要坚持人民主体，贴近人民生活和精神世界，从普通人的奋斗、创造和情感中发现时代气象。③只有把镜头对准人民、把表达扎根人民，才能真实展现人民奋发有为的精神风貌，使作品具有打动人心的力量和持久的社会价值。",
    },
]


def now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M +08")


def write_csv(path: Path, rows: list[dict[str, object]], fields: list[str]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fields})


def write_draft_files() -> None:
    rows = []
    for target in TARGETS:
        rows.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "evidence_note": target["evidence_note"],
                "new_why_chars": len(target["new_why"]),
                "new_answer_chars": len(target["new_answer"]),
                "new_why": target["new_why"],
                "new_answer": target["new_answer"],
            }
        )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_20260525.json").write_text(
        json.dumps({"updated": now(), "targets": rows}, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_20260525.csv",
        rows,
        ["queue_id", "heading", "node", "evidence_note", "new_why_chars", "new_answer_chars", "new_why", "new_answer"],
    )
    lines = [
        f"# P0 Thickness Repair Batch{BATCH} Draft 20260525",
        "",
        f"Updated: {now()}",
        "",
        "Status: `DRAFT_READY_FOR_DOCX_APPLICATION`",
        "",
        "- Scope: next 8 P0 subjective triple-thin rows after Batch09 refresh.",
        "- Repair method: thicken only existing why/answer paragraphs inside existing nodes.",
        "- Evidence boundary: evidence notes are carried from the current coverage matrix; no ordinary reference answer is upgraded into a scoring rubric.",
        "",
        "| queue_id | node | heading | new why chars | new answer chars | evidence note |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in rows:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | {row['new_why_chars']} | {row['new_answer_chars']} | {row['evidence_note']} |"
        )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_20260525.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )


def main() -> int:
    write_draft_files()
    docx = helper.current_docx()
    backup = docx.with_name(f"{docx.stem}_backup_before_p0_thickness_batch{BATCH}_{datetime.now():%Y%m%d_%H%M%S}.docx")
    shutil.copy2(docx, backup)

    doc = Document(str(docx))
    blocks = helper.entry_blocks(doc)
    applied: list[dict[str, object]] = []
    for target in TARGETS:
        block = helper.matching_block(blocks, target)
        fields = block["fields"]
        field_paras = block["field_paras"]
        before_why = str(fields.get("why", ""))
        before_answer = str(fields.get("answer", ""))
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["why"])], "why", target["new_why"])
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["answer"])], "answer", target["new_answer"])
        applied.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "heading_para": block["heading_para"],
                "why_para": field_paras["why"],
                "answer_para": field_paras["answer"],
                "before_why_chars": len(before_why),
                "after_why_chars": len(target["new_why"]),
                "before_answer_chars": len(before_answer),
                "after_answer_chars": len(target["new_answer"]),
                "evidence_note": target["evidence_note"],
                "old_answer_excerpt": target["old_answer_excerpt"],
            }
        )

    doc.save(str(docx))
    with zipfile.ZipFile(docx) as z:
        zip_ok = "[Content_Types].xml" in z.namelist() and "word/document.xml" in z.namelist()

    result = {
        "updated": now(),
        "status": f"P0_BATCH{BATCH}_APPLIED_REQUIRES_RENDER_AND_EXTERNAL_RECHECK",
        "docx": str(docx),
        "backup_docx": str(backup),
        "zip_ok": zip_ok,
        "targets": len(TARGETS),
        "applied": applied,
        "boundary": [
            "This is a local thickness repair batch, not final acceptance.",
            "PDF/render outputs must be regenerated after this DOCX edit.",
            "External review remains pending for the current post-repair artifact.",
        ],
    }
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_20260525.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_20260525.csv",
        applied,
        [
            "queue_id",
            "heading",
            "node",
            "heading_para",
            "why_para",
            "answer_para",
            "before_why_chars",
            "after_why_chars",
            "before_answer_chars",
            "after_answer_chars",
            "evidence_note",
            "old_answer_excerpt",
        ],
    )
    lines = [
        f"# P0 Thickness Repair Batch{BATCH} Apply 20260525",
        "",
        f"Updated: {result['updated']}",
        "",
        f"Status: `{result['status']}`",
        "",
        f"- DOCX: `{docx.name}`.",
        f"- Backup: `{backup.name}`.",
        f"- Targets applied: `{len(applied)}` / `{len(TARGETS)}`.",
        f"- DOCX zip structural check: `{str(zip_ok).lower()}`.",
        "",
        "## Applied Rows",
        "",
        "| queue_id | node | heading | why chars | answer chars | paragraph refs |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in applied:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | "
            f"{row['before_why_chars']} -> {row['after_why_chars']} | "
            f"{row['before_answer_chars']} -> {row['after_answer_chars']} | "
            f"h={row['heading_para']}; why={row['why_para']}; answer={row['answer_para']} |"
        )
    lines.extend(
        [
            "",
            "## Boundary",
            "",
            "- This batch repairs only 8 additional P0 rows.",
            "- Remaining thickness queue, render QA, and model gates stay open until rechecked.",
        ]
    )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_20260525.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
