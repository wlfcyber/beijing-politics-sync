from __future__ import annotations

import csv
import json
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

import apply_p0_thickness_batch01_20260525 as helper
from docx import Document


ROOT = Path(__file__).resolve().parent
BATCH = "08"

TARGETS = [
    {
        "queue_id": "T0500",
        "heading": "30. 2026朝阳期末 第21题（主观题）",
        "node": "群众观点 / 群众路线 / 人民主体地位",
        "old_answer_excerpt": "发挥制度优势要坚持群众观点和群众路线，坚持以人民为中心，激发人民群众投身中国式现代化建设的积极性、主动性、创造性。",
        "evidence_note": "正式阅卷细则直接列坚持以人民为中心、群众观，并要求激发人民群众投身现代化建设的积极性、主动性、创造性。",
        "new_why": "题目要求说明人民至上与现代化建设的关系，材料把人民至上作为制度优势之一，细则又明确要求从群众观说明人民群众在现代化建设中的主体作用。这里不能只写一般价值口号，而要看到中国式现代化依靠人民推进、为了人民展开、由人民共享，人民群众的积极性主动性创造性是制度优势转化为治理效能和现代化成果的关键。",
        "new_answer": "①人民群众是社会历史的主体，是推进中国式现代化的实践力量。②发挥制度优势，必须坚持群众观点和群众路线，坚持以人民为中心，把人民对美好生活的需要作为现代化建设的重要方向。③通过制度安排激发人民群众投身现代化建设的积极性、主动性、创造性，才能把人民主体作用转化为推进中国式现代化的强大合力。",
    },
    {
        "queue_id": "T0538",
        "heading": "30. 2026西城二模 第16题（主观题）",
        "node": "价值观的导向作用",
        "old_answer_excerpt": "中式生活方式承载的健康、和谐、顺应自然等价值追求契合当代青年对美好生活的期待。卷面可写：价值观具有导向作用，共通的价值追求引导海外青年理解、认同并模仿中式生活方式。",
        "evidence_note": "评标渲染页明确列价值观的导向作用；材料写全球青年从中式生活方式中感受和谐、顺应自然、身心兼顾等智慧。",
        "new_why": "材料中的喝热水、八段锦、节气起居、和睦共居之所以能跨越文化差异，并不只是因为形式新鲜，而是因为背后承载健康生活、身心协调、顺应自然、人与人和谐相处等价值取向。评标渲染页明确列价值观的导向作用，因此应解释这些共同价值如何影响海外青年对中式生活方式的认识、评价和行为选择。",
        "new_answer": "①价值观对人们认识世界和改造世界具有重要导向作用，会影响人们对生活方式的理解、评价和选择。②中式生活方式蕴含健康、和谐、顺应自然、身心兼顾等价值追求，契合全球青年对美好生活、身心平衡和文明互鉴的期待。③正是这些共通价值引导海外青年理解、认同并愿意体验中式生活方式，使其能够跨越文化差异、形成广泛共鸣。",
    },
    {
        "queue_id": "T0212",
        "heading": "30. 2026西城二模 第20题（主观题）",
        "node": "发展的观点 / 当前与长远",
        "old_answer_excerpt": "树立正确政绩观，要用发展的观点处理当前和长远的关系，既办群众当下看得见的实事，也做打基础、利长远的好事，以持续奋斗推动事业长远发展。",
        "evidence_note": "材料强调眼前实事、为后人铺垫、打基础、利长远、久久为功和行稳致远，支持发展的观点落点。",
        "new_why": "正确政绩观不仅要看当下办了什么，还要看这些工作能否为后人铺垫、打基础、利长远。材料中的谷文昌“久久为功”和一代代共产党人推动事业行稳致远，说明政绩评价不能停留在短期显绩和即时效果，而要把当前工作放在长期发展过程中衡量。看到“利长远、打基础、久久为功”，应落到发展的观点。",
        "new_answer": "①事物是变化发展的，要用发展的观点看问题，正确处理当前和长远的关系。②树立正确政绩观，既要解决群众眼前急难愁盼，办好当下看得见、摸得着的实事，也要做打基础、利长远、惠后人的工作。③像谷文昌那样久久为功，才能避免急功近利和短期工程，把一时成绩转化为推动党和国家事业行稳致远的持续成果。",
    },
    {
        "queue_id": "T0213",
        "heading": "31. 2025丰台期末 第16题（主观题）",
        "node": "发展的观点 / 用发展眼光看数据和方案",
        "old_answer_excerpt": "作答可用发展的观点说明：胸中有“数”要把对象看成变化发展的过程，既分析当下数量状态，也判断发展趋势和长远影响，使计划方案适应事物发展的方向和阶段性要求。",
        "evidence_note": "细则多处把发展列为可作答角度；材料要求在工作生活中做计划、定方案，在国家治理中看大数、看长远。",
        "new_why": "“胸中有数”不是对某个静态数字的记忆，而是要通过数据把握对象所处阶段、变化趋势和长远影响。材料写做计划、定方案、国家治理中看大数、看长远，细则又把发展列为可作答角度，说明学生应把数据理解为发展过程中的信息：既看当前状态，也看未来走向和阶段性要求。",
        "new_answer": "①事物是变化发展的，要用发展的观点看问题。②胸中有“数”，既要掌握当前数量状态和现实条件，也要分析数据背后的变化趋势、阶段特征和长远影响。③只有把“数”放到发展过程中理解，计划方案才能适应事物发展的方向和阶段性要求，避免只看眼前、静止用数或凭短期指标作决策。",
    },
    {
        "queue_id": "T0539",
        "heading": "31. 2025朝阳一模 第21题（主观题）",
        "node": "价值观的导向作用 / 正确价值判断和价值选择",
        "old_answer_excerpt": "价值观具有导向作用。党员干部树立以人民为中心的政绩观，才能把工作导向人民利益和民生难题，防止脱离群众、只看表面政绩。",
        "evidence_note": "细则写重视价值观导向作用、自觉站在最广大人民立场上；材料以造福人民的政绩观评价干部业绩。",
        "new_why": "“大事小事观”“为民办事观”“造福人民的政绩观”本质上是在回答干部如何判断事情价值、为谁办事、用什么标准评价政绩。细则写重视价值观导向作用，并要求站在最广大人民立场上；材料也把造福人民作为评价干部业绩的标准。因此应落到价值观对认识、选择和行动的导向作用。",
        "new_answer": "①价值观对人们认识世界和改造世界具有导向作用，影响人们作出价值判断和价值选择。②党员干部只有树立以人民为中心的正确政绩观，才能把“大事小事”的判断标准放在人民利益上，把工作重点指向群众最关心、最直接、最现实的问题。③这样才能防止脱离群众、追求表面政绩或个人声望，使为民办事真正服务造福人民。",
    },
    {
        "queue_id": "T0181",
        "heading": "31. 2025朝阳期末 第22题（主观题）",
        "node": "系统观念 / 整体性推进理论创新和事业发展",
        "old_answer_excerpt": "不断谱写马克思主义中国化时代化新篇章，要坚持系统观念，把理论创新放在党和国家事业整体推进中谋划，统筹执政规律、社会主义建设规律、人类社会发展规律以及不同领域重大时代课题，使新的理论成果成为整体性推进事业发展的科学思想方法。",
        "evidence_note": "PPT评分细则在辩证法部分明列坚持系统观念，为整体性推进党和国家各项事业提供科学思想方法。",
        "new_why": "马克思主义中国化时代化不是某个局部观点的零散更新，而是围绕党和国家事业发展，系统回答治国理政、社会主义建设和人类文明发展中的重大时代课题。细则明确列坚持系统观念，并指出其为整体性推进党和国家各项事业提供科学思想方法，因此应把理论创新放在整体布局、规律统筹和多领域协同中理解。",
        "new_answer": "①系统观念要求立足整体，统筹各方面关系，优化整体推进。②谱写马克思主义中国化时代化新篇章，要把理论创新放在党和国家事业整体发展中谋划，统筹共产党执政规律、社会主义建设规律、人类社会发展规律以及不同领域重大时代课题。③这样形成的理论成果才不是零散观点，而能成为整体性推进中国式现代化和党治国理政实践的科学思想方法。",
    },
    {
        "queue_id": "T0048",
        "heading": "31. 2026朝阳期中 第19题（主观题）",
        "node": "正确意识的能动作用 / 精神力量转化为行动",
        "old_answer_excerpt": "弘扬抗战精神要发挥正确意识的能动作用。以爱国情怀、民族气节、英雄气概和必胜信念激励青年主动担当，把精神力量转化为砥砺前行的行动自觉。",
        "evidence_note": "正式评分细则把正确意识的能动作用列为弘扬抗战精神的哲学角度。",
        "new_why": "素材引用伟大抗战精神的丰富内涵，并要求新时代继承和弘扬。细则明确把正确意识的能动作用列为哲学角度，说明抗战精神不能只作为历史记忆停留在口号中，而应作为正确意识和精神力量，激励青年认识责任、坚定信念，并把爱国情怀、民族气节和必胜信念转化为新时代奋斗行动。",
        "new_answer": "①正确意识对实践具有促进作用，能够指导人们改造世界、推动行动。②伟大抗战精神蕴含爱国情怀、民族气节、英雄气概和必胜信念，是激励新时代青年奋进的正确精神力量。③继承和弘扬抗战精神，就是要把这种精神转化为主动担当、攻坚克难、服务国家和民族复兴的行动自觉，而不是停留在纪念和口号上。",
    },
    {
        "queue_id": "T0540",
        "heading": "32. 2025丰台期末 第16题（主观题）",
        "node": "价值观的导向作用 / 数据使用的价值方向",
        "old_answer_excerpt": "胸中有“数”不能只是技术性掌握数据，还要在正确价值观导向下使用数据。只有把“数”用于科学决策、公共治理和人民利益，才能使数据分析转化为正确行动，避免片面追求指标或凭偏好用数。",
        "evidence_note": "细则把“数为何所用”明确列为可涉及正确价值观；材料把“数”用于计划方案和国家治理。",
        "new_why": "“数为何所用”直接指向数据服务什么目标。材料把“数”用于计划方案和国家治理，细则又明确列正确价值观，说明胸中有数不能只是技术层面的统计和计算，还要回答用数为了谁、服务什么公共目标。若缺少正确价值导向，数据可能被片面指标、个人偏好或短期利益牵引，反而造成错误决策。",
        "new_answer": "①价值观对认识活动和实践活动具有导向作用。②胸中有“数”，不仅要掌握真实数据，还要在正确价值观指导下决定数据如何使用，把“数”用于科学决策、公共治理、人民利益和长远发展。③这样才能让数据分析转化为正确行动，避免片面追求指标、为数字而数字，或按照个人偏好随意用数。",
    },
]


def now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M +08")


def write_csv(path: Path, rows: list[dict[str, object]], fields: list[str]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fields})


def write_draft_files() -> None:
    rows = []
    for target in TARGETS:
        rows.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "evidence_note": target["evidence_note"],
                "new_why_chars": len(target["new_why"]),
                "new_answer_chars": len(target["new_answer"]),
                "new_why": target["new_why"],
                "new_answer": target["new_answer"],
            }
        )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_20260525.json").write_text(
        json.dumps({"updated": now(), "targets": rows}, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_20260525.csv",
        rows,
        ["queue_id", "heading", "node", "evidence_note", "new_why_chars", "new_answer_chars", "new_why", "new_answer"],
    )
    lines = [
        f"# P0 Thickness Repair Batch{BATCH} Draft 20260525",
        "",
        f"Updated: {now()}",
        "",
        "Status: `DRAFT_READY_FOR_DOCX_APPLICATION`",
        "",
        "- Scope: next 8 P0 subjective triple-thin rows after Batch07 refresh.",
        "- Repair method: thicken only existing why/answer paragraphs inside existing nodes.",
        "- Evidence boundary: formal-rubric support is named only where the existing row already carries it; material-trigger rows are not upgraded to formal rubric evidence.",
        "",
        "| queue_id | node | heading | new why chars | new answer chars | evidence note |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in rows:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | {row['new_why_chars']} | {row['new_answer_chars']} | {row['evidence_note']} |"
        )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_20260525.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )


def main() -> int:
    write_draft_files()
    docx = helper.current_docx()
    backup = docx.with_name(f"{docx.stem}_backup_before_p0_thickness_batch{BATCH}_{datetime.now():%Y%m%d_%H%M%S}.docx")
    shutil.copy2(docx, backup)

    doc = Document(str(docx))
    blocks = helper.entry_blocks(doc)
    applied: list[dict[str, object]] = []
    for target in TARGETS:
        block = helper.matching_block(blocks, target)
        fields = block["fields"]
        field_paras = block["field_paras"]
        before_why = str(fields.get("why", ""))
        before_answer = str(fields.get("answer", ""))
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["why"])], "why", target["new_why"])
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["answer"])], "answer", target["new_answer"])
        applied.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "heading_para": block["heading_para"],
                "why_para": field_paras["why"],
                "answer_para": field_paras["answer"],
                "before_why_chars": len(before_why),
                "after_why_chars": len(target["new_why"]),
                "before_answer_chars": len(before_answer),
                "after_answer_chars": len(target["new_answer"]),
                "evidence_note": target["evidence_note"],
                "old_answer_excerpt": target["old_answer_excerpt"],
            }
        )

    doc.save(str(docx))
    with zipfile.ZipFile(docx) as z:
        zip_ok = "[Content_Types].xml" in z.namelist() and "word/document.xml" in z.namelist()

    result = {
        "updated": now(),
        "status": f"P0_BATCH{BATCH}_APPLIED_REQUIRES_RENDER_AND_EXTERNAL_RECHECK",
        "docx": str(docx),
        "backup_docx": str(backup),
        "zip_ok": zip_ok,
        "targets": len(TARGETS),
        "applied": applied,
        "boundary": [
            "This is a local thickness repair batch, not final acceptance.",
            "PDF/render outputs must be regenerated after this DOCX edit.",
            "External review remains pending for the current post-repair artifact.",
        ],
    }
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_20260525.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_20260525.csv",
        applied,
        [
            "queue_id",
            "heading",
            "node",
            "heading_para",
            "why_para",
            "answer_para",
            "before_why_chars",
            "after_why_chars",
            "before_answer_chars",
            "after_answer_chars",
            "evidence_note",
            "old_answer_excerpt",
        ],
    )
    lines = [
        f"# P0 Thickness Repair Batch{BATCH} Apply 20260525",
        "",
        f"Updated: {result['updated']}",
        "",
        f"Status: `{result['status']}`",
        "",
        f"- DOCX: `{docx.name}`.",
        f"- Backup: `{backup.name}`.",
        f"- Targets applied: `{len(applied)}` / `{len(TARGETS)}`.",
        f"- DOCX zip structural check: `{str(zip_ok).lower()}`.",
        "",
        "## Applied Rows",
        "",
        "| queue_id | node | heading | why chars | answer chars | paragraph refs |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in applied:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | "
            f"{row['before_why_chars']} -> {row['after_why_chars']} | "
            f"{row['before_answer_chars']} -> {row['after_answer_chars']} | "
            f"h={row['heading_para']}; why={row['why_para']}; answer={row['answer_para']} |"
        )
    lines.extend(
        [
            "",
            "## Boundary",
            "",
            "- This batch repairs only 8 additional P0 rows.",
            "- Remaining thickness queue, render QA, and model gates stay open until rechecked.",
        ]
    )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_20260525.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
