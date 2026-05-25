from __future__ import annotations

import csv
import json
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

import apply_p0_thickness_batch01_20260525 as helper
from docx import Document


ROOT = Path(__file__).resolve().parent
BATCH = "14"

TARGETS = [
    {
        "queue_id": "T0279",
        "heading": "43. 2026朝阳期末 第16题（主观题）",
        "node": "矛盾就是对立统一",
        "old_answer_excerpt": "传统戏曲焕发生命力要坚持矛盾观点，把守住文化之根和紧跟时代之变统一起来，在保留精髓与突破旧形式中实现新的发展。",
        "evidence_note": "矩阵 M1284：正式阅卷细则第21行补充“矛盾”；第15-18行围绕剔除与保留、传统与时代统一。",
        "new_why": "传统戏曲创新不是在传统和现代之间二选一，而是要处理保留与剔除、守根与求变、传统程式与时代审美的对立统一关系。正式阅卷细则第21行补充“矛盾”，且第15-18行围绕精髓保留、新内涵注入和时代转化展开，说明答案要写出矛盾双方怎样在同一文化传承实践中实现统一。",
        "new_answer": "①矛盾双方既相互区别又相互联系，传统戏曲中的守住文化之根与回应时代之变就是一对对立统一关系。②焕发生命力，既不能丢掉唱腔、程式、人物精神和审美精髓，也不能固守旧形式、拒绝现代舞台表达和青年审美。③在保留精髓中剔除不合时宜的表达，在紧跟时代中守住文化根脉，才能把传统与现代统一起来，实现传统戏曲新的发展。",
    },
    {
        "queue_id": "T0529",
        "heading": "43. 2026西城期末 第16题第（2）问（主观题）",
        "node": "价值判断与价值选择",
        "old_answer_excerpt": "追求真实体现正确价值判断与价值选择，引导人们尊重事实、规律和人民利益，因而能够跨领域达成价值共识。",
        "evidence_note": "矩阵 M1396：正式评分细则 page_002 价值观角度写明引导人们作出正确价值判断与选择。",
        "new_why": "设问问“真实”为何直抵人心，不能只写真实让人感动。正式评分细则 page_002 在价值观角度写明，对真实的追求能引导人们作出正确价值判断与选择。新闻、历史、司法、科学、纪录片等领域共同追求真实，说明真实提供了尊重事实、尊重规律、维护人民利益的价值尺度。",
        "new_answer": "①价值判断和价值选择影响人们怎样评价事实、采取行动，正确价值选择要符合社会发展要求并站在人民立场上。②追求真实，就是尊重客观事实、遵循规律、反对虚假和偏见，这能引导人们作出正确价值判断与选择。③正因为真实以事实和公道为基础，能够跨越身份、立场和情绪的隔阂，促成人们在新闻、历史、司法、科学等领域形成价值共识，所以具有直抵人心的力量。",
    },
    {
        "queue_id": "T0280",
        "heading": "44. 2026海淀期末 第21题（主观题）",
        "node": "矛盾就是对立统一",
        "old_answer_excerpt": "活力与秩序是对立统一的关系。中国式现代化既要释放发展活力，也要维护社会秩序，在动态平衡中推进现代化。",
        "evidence_note": "矩阵 M1363：正式评分标准 page_006 Q21要求运用对立统一观点分析“活力与秩序”的关系。",
        "new_why": "设问核心就是“活而不乱、活跃有序”的动态平衡，正式评分标准 page_006 明确要求运用对立统一观点分析活力与秩序的关系。活力强调创新、流动和创造，秩序强调规则、稳定和安全，二者有张力但又相互需要，因此不能只强调一端。",
        "new_answer": "①矛盾双方既对立又统一，活力与秩序也是中国式现代化进程中的对立统一关系。②没有活力，社会发展就会缺少创新、创造和流动；没有秩序，发展活力又可能滑向无序、风险和不公平。③中国式现代化要在法治、制度和治理能力保障下释放市场活力、社会活力和创新活力，同时以良好秩序维护安全稳定，形成活而不乱、活跃有序的动态平衡。",
    },
    {
        "queue_id": "T0530",
        "heading": "44. 2026通州期末 第16题（主观题）",
        "node": "价值判断与价值选择",
        "old_answer_excerpt": "治水利民体现正确价值判断和价值选择，把人民利益作为治理、维护和创新的根本尺度。",
        "evidence_note": "矩阵 M1434：正式细则第16题价值观部分写明“自觉站在最广大人民群众的立场上”。",
        "new_why": "都江堰跨越千年的治水智慧，不只是工程技术智慧，也包含为什么治水、为了谁治水的价值尺度。正式细则第16题价值观部分写明要自觉站在最广大人民群众的立场上，因此应把防洪、灌溉、民生安定和持续维护落到正确价值判断与价值选择上。",
        "new_answer": "①正确价值判断和价值选择要自觉站在最广大人民群众的立场上，把人民利益作为根本尺度。②都江堰治水智慧服务防洪、灌溉、生产生活和区域发展，体现了治水利民、造福人民的价值取向。③今天保护、维护和创新利用都江堰，也应坚持人民立场，把工程安全、生态福祉、文化传承和民生需要统一起来，使千年水利工程继续服务人民生活。",
    },
    {
        "queue_id": "T0249",
        "heading": "5. 2025西城期末 第18题（主观题）",
        "node": "矛盾就是对立统一",
        "old_answer_excerpt": "要承认矛盾、坚持一分为二看待技术恐惧，既看到它抑制创新的一面，也看到它倒逼安全监管和政策完善的一面；对待新技术既要扶持其成长，又要正视恐惧背后的合理诉求，把矛盾转化为推动新技术健康发展的动力。",
        "evidence_note": "矩阵 M1143：正式细则第111-113行明列“矛盾就是对立统一，要一分为二辩证看问题”。",
        "new_why": "技术恐惧既表现为排斥、拒绝和抵制新技术，也可能提示安全、伦理、就业和制度风险，倒逼监管与规则完善。正式细则明列矛盾就是对立统一，要一分为二辩证看问题。材料列举印刷机、机器生产、火车、核技术、计算机等历史经验，正是要求学生看到技术恐惧的消极面和合理警示面。",
        "new_answer": "①矛盾就是对立统一，要坚持一分为二、全面地看问题。②技术恐惧一方面可能放大风险、阻碍创新、导致排斥新技术；另一方面也反映人们对安全、伦理、秩序和利益调整的合理担忧，能够倒逼制度法规、风险评估和安全治理完善。③对待新技术，既要支持其健康成长，又要正视恐惧背后的现实诉求，把矛盾转化为改进技术治理、推动新技术造福社会的动力。",
    },
    {
        "queue_id": "T0110",
        "heading": "5. 2026丰台期末 第16题（主观题）",
        "node": "联系的客观性",
        "old_answer_excerpt": "留白要尊重联系的客观性。只有把留白放回具体空间和周围事物的真实联系中，才能营造和谐有序的整体意境；随意空置或过度留白反而会破坏效果。",
        "evidence_note": "矩阵 M1206：正式细则第52-53行写明“联系具有普遍性、客观性”。",
        "new_why": "材料把留白放到山林、溪谷、高楼旁绿地、旷野沉思等具体场景中，说明留白效果取决于空间结构、周围景物、人的活动和整体意境之间的真实联系。正式细则第52-53行写明联系具有普遍性、客观性，所以留白不是主观随意空置，而要尊重具体环境中的客观联系。",
        "new_answer": "①联系具有客观性，事物之间的联系不以人的主观愿望为转移。②留白的美感来自它同山林、溪谷、建筑、绿地、视线、活动空间和整体意境之间的真实联系，不能脱离具体场景随意安排。③只有尊重这些客观联系，把留白放在适合的空间结构和环境关系中，才能形成含蓄、舒展、和谐的审美效果；若主观任意空置或过度留白，反而会破坏整体意境。",
    },
    {
        "queue_id": "T0285",
        "heading": "5. 2026房山二模 第21题（主观题·矛盾普遍性角度）",
        "node": "矛盾的普遍性",
        "old_answer_excerpt": "可从矛盾普遍性角度作答：新的长征路上必然会遇到发展不平衡、改革攻坚、科技创新、生态文明、国际环境等矛盾和挑战。要坚持问题导向，承认并分析矛盾，通过伟大斗争和系统实践解决矛盾，推动中国式现代化不断前进。",
        "evidence_note": "矩阵 M0697：第21题正式细则在“走好新的长征路”角度列出“矛盾普遍性，承认分析解决矛盾/伟大斗争”。",
        "new_why": "“新征程是新的长征”强调现代化道路不会一帆风顺。正式细则在走好新的长征路角度列出矛盾普遍性、承认分析解决矛盾和伟大斗争，说明作答不能停留在口号式奋斗，而要写出新征程中风险挑战和复杂矛盾客观存在，必须坚持问题导向去解决。",
        "new_answer": "①矛盾具有普遍性，事事有矛盾、时时有矛盾，推进新的长征路必然会遇到风险挑战和复杂问题。②中国式现代化进程中，会面对发展不平衡、改革攻坚、科技创新、生态文明、民生保障和国际环境变化等矛盾。③走好新的长征路，必须坚持问题导向，敢于承认矛盾、深入分析矛盾、积极解决矛盾，通过伟大斗争和系统实践不断破解难题，推动现代化建设向前发展。",
    },
    {
        "queue_id": "T0423",
        "heading": "5. 2026朝阳期中 第19题（主观题）",
        "node": "社会存在与社会意识",
        "old_answer_excerpt": "要弘扬抗战精神，因为它作为先进社会意识对社会存在具有能动反作用，能凝聚民族奋进合力，推动新时代强国建设和民族复兴的实践。",
        "evidence_note": "矩阵 M1260：正式评分细则第98-99行在“为什么弘扬”中列正确社会意识的推动作用。",
        "new_why": "抗战精神在全民族抗战的社会实践中孕育形成，同时又能作为先进社会意识影响今天的民族复兴实践。正式评分细则第98-99行把正确社会意识的推动作用列入为什么弘扬的角度，说明答案要写出社会存在决定社会意识、先进社会意识反作用于社会存在这条双向链。",
        "new_answer": "①社会存在决定社会意识，伟大抗战精神是在全民族抗战的历史实践中孕育形成的。②先进的社会意识对社会存在具有能动反作用，能够凝聚共识、鼓舞斗志、引导实践。③新时代弘扬抗战精神，就是把爱国情怀、民族气节、英雄气概和必胜信念转化为强国建设、民族复兴的精神力量，推动人民在现实实践中接续奋斗。",
    },
]


def now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M +08")


def write_csv(path: Path, rows: list[dict[str, object]], fields: list[str]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fields})


def write_draft_files() -> None:
    rows = []
    for target in TARGETS:
        rows.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "evidence_note": target["evidence_note"],
                "new_why_chars": len(target["new_why"]),
                "new_answer_chars": len(target["new_answer"]),
                "new_why": target["new_why"],
                "new_answer": target["new_answer"],
            }
        )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_20260525.json").write_text(
        json.dumps({"updated": now(), "targets": rows}, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_20260525.csv",
        rows,
        ["queue_id", "heading", "node", "evidence_note", "new_why_chars", "new_answer_chars", "new_why", "new_answer"],
    )
    lines = [
        f"# P0 Thickness Repair Batch{BATCH} Draft 20260525",
        "",
        f"Updated: {now()}",
        "",
        "Status: `DRAFT_READY_FOR_DOCX_APPLICATION`",
        "",
        "- Scope: next 8 P0 subjective triple-thin rows after Batch13 refresh.",
        "- Repair method: thicken only existing why/answer paragraphs inside existing evidence-supported nodes.",
        "- Evidence boundary: all evidence notes cite current matrix rows; ordinary reference answers are not upgraded into scoring rubrics.",
        "",
        "| queue_id | node | heading | new why chars | new answer chars | evidence note |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in rows:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | {row['new_why_chars']} | {row['new_answer_chars']} | {row['evidence_note']} |"
        )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_20260525.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )


def main() -> int:
    write_draft_files()
    docx = helper.current_docx()
    backup = docx.with_name(f"{docx.stem}_backup_before_p0_thickness_batch{BATCH}_{datetime.now():%Y%m%d_%H%M%S}.docx")
    shutil.copy2(docx, backup)

    doc = Document(str(docx))
    blocks = helper.entry_blocks(doc)
    applied: list[dict[str, object]] = []
    for target in TARGETS:
        block = helper.matching_block(blocks, target)
        fields = block["fields"]
        field_paras = block["field_paras"]
        before_why = str(fields.get("why", ""))
        before_answer = str(fields.get("answer", ""))
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["why"])], "why", target["new_why"])
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["answer"])], "answer", target["new_answer"])
        applied.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "heading_para": block["heading_para"],
                "why_para": field_paras["why"],
                "answer_para": field_paras["answer"],
                "before_why_chars": len(before_why),
                "after_why_chars": len(target["new_why"]),
                "before_answer_chars": len(before_answer),
                "after_answer_chars": len(target["new_answer"]),
                "evidence_note": target["evidence_note"],
                "old_answer_excerpt": target["old_answer_excerpt"],
            }
        )

    doc.save(str(docx))
    with zipfile.ZipFile(docx) as z:
        zip_ok = "[Content_Types].xml" in z.namelist() and "word/document.xml" in z.namelist()

    result = {
        "updated": now(),
        "status": f"P0_BATCH{BATCH}_APPLIED_REQUIRES_RENDER_AND_EXTERNAL_RECHECK",
        "docx": str(docx),
        "backup_docx": str(backup),
        "zip_ok": zip_ok,
        "targets": len(TARGETS),
        "applied": applied,
        "boundary": [
            "This is a local thickness repair batch, not final acceptance.",
            "PDF/render outputs must be regenerated after this DOCX edit.",
            "External review remains pending for the current post-repair artifact.",
        ],
    }
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_20260525.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_20260525.csv",
        applied,
        [
            "queue_id",
            "heading",
            "node",
            "heading_para",
            "why_para",
            "answer_para",
            "before_why_chars",
            "after_why_chars",
            "before_answer_chars",
            "after_answer_chars",
            "evidence_note",
            "old_answer_excerpt",
        ],
    )
    lines = [
        f"# P0 Thickness Repair Batch{BATCH} Apply 20260525",
        "",
        f"Updated: {result['updated']}",
        "",
        f"Status: `{result['status']}`",
        "",
        f"- DOCX: `{docx.name}`.",
        f"- Backup: `{backup.name}`.",
        f"- Targets applied: `{len(applied)}` / `{len(TARGETS)}`.",
        f"- DOCX zip structural check: `{str(zip_ok).lower()}`.",
        "",
        "## Applied Rows",
        "",
        "| queue_id | node | heading | why chars | answer chars | paragraph refs |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in applied:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | "
            f"{row['before_why_chars']} -> {row['after_why_chars']} | "
            f"{row['before_answer_chars']} -> {row['after_answer_chars']} | "
            f"h={row['heading_para']}; why={row['why_para']}; answer={row['answer_para']} |"
        )
    lines.extend(
        [
            "",
            "## Boundary",
            "",
            "- This batch repairs only 8 additional P0 rows.",
            "- Remaining thickness queue, render QA, and model gates stay open until rechecked.",
        ]
    )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_20260525.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
