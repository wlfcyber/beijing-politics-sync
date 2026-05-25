from __future__ import annotations

import csv
import json
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

import apply_p0_thickness_batch01_20260525 as helper
from docx import Document


ROOT = Path(__file__).resolve().parent
BATCH = "11"

TARGETS = [
    {
        "queue_id": "T0048",
        "heading": "34. 2026西城期末 第16题第（2）问（主观题）",
        "node": "主观能动性 / 意识的能动作用",
        "old_answer_excerpt": "符合客观真实的正确意识能够有效发挥能动作用，正确指导实践、打破偏见、凝聚共识，所以真实具有直抵人心的力量。",
        "evidence_note": "矩阵 M1394：正式评分细则 page_002 Q16(2) 唯物论角度写明正确意识有效发挥能动作用、正确指导实践。",
        "new_why": "设问追问“真实”为何能直抵人心，关键不是把真实写成一般情感感染力，而是从主观认识是否符合客观事实来说明。正式评分细则 page_002 写明，只有符合客观的正确意识才能有效发挥能动作用、正确指导实践。看到新闻、历史、司法、科学、纪录片共同追求真实，就应落到正确意识的能动作用。",
        "new_answer": "①真实意味着主观认识同客观事实相符合，能够帮助人们摆脱虚假、偏见和隔阂，形成可靠认识。②正确意识对改造世界具有能动作用，能够指导实践、凝聚共识、推动人们作出正确判断和行动。③新闻、历史、司法、科学和纪录片中的真实之所以直抵人心，就在于它以客观事实支撑正确认识，使人们在事实面前产生理解、信任和行动力量。",
    },
    {
        "queue_id": "T0519",
        "heading": "35. 2026朝阳期中 第18题（主观题）",
        "node": "价值观的导向作用",
        "old_answer_excerpt": "使用AI提供情绪价值要发挥正确价值观的导向作用，坚持科技服务人的健康成长，让AI成为辅助疏导工具，而不是引导人逃避现实关系、沉溺虚拟安慰。",
        "evidence_note": "矩阵 M1245：正式评分细则第63-66、72-85行在“如何做”中列正确价值观的导向作用。",
        "new_why": "材料要求讨论是否使用 AI 提供情绪价值，真正的落点不是技术能不能回应，而是用什么价值方向约束技术使用。正式评分细则在“如何做”中列正确价值观的导向作用，说明答题要把科技向善、人的健康成长、真实交往和主体能力放在评价尺度中，而不是让便利性压倒人的发展。",
        "new_answer": "①价值观对人们认识世界和改造世界具有导向作用，会影响我们怎样评价和使用 AI。②使用 AI 提供情绪价值，应坚持科技服务人的健康成长，把人的真实需要、现实交往能力和身心发展作为价值尺度。③AI 可以作为表达、倾诉和初步疏导的辅助工具，但不能替代真实关系和现实问题解决；要防止沉溺虚拟安慰、逃避现实冲突，使技术使用服从人的全面发展。",
    },
    {
        "queue_id": "T0106",
        "heading": "35. 2026朝阳期中 第19题（主观题）",
        "node": "联系的普遍性 / 联系的观点（总）",
        "old_answer_excerpt": "弘扬抗战精神要坚持联系观点，把中华民族家国传统、全民族抗战实践和新时代强国复兴使命联系起来理解，使伟大精神在现实奋斗中延续。",
        "evidence_note": "矩阵 M1261：正式评分细则第100-102行在如何弘扬抗战精神中列联系的观点，证据等级为正式细则宽角度/术语支持。",
        "new_why": "抗战精神不是孤立的历史口号。材料把古代家国情怀、全民族抗战实践、习近平总书记对抗战精神的阐释以及新时代弘扬要求放在同一链条中，细则也把联系的观点列入如何弘扬的哲学角度。因此应说明抗战精神如何在历史传统、民族实践和现实使命的联系中被理解和传承。",
        "new_answer": "①联系具有普遍性，要把事物放在相互联系中理解。②弘扬抗战精神，要看到中华民族家国传统、全民族抗战实践、伟大民族精神和新时代强国复兴使命之间的内在联系。③把这些联系讲清楚，才能避免把抗战精神变成空泛口号，而是把历史记忆、民族情感和现实奋斗贯通起来，使伟大精神在新时代继续转化为爱国担当和强国行动。",
    },
    {
        "queue_id": "T0204",
        "heading": "35. 2026朝阳期中 第19题（主观题）",
        "node": "发展的观点 / 发展的普遍性",
        "old_answer_excerpt": "弘扬抗战精神要坚持发展的观点。既看到它是中华民族精神在抗战实践中的发展升华，也要在新时代强国建设中不断继承、弘扬和发展其精神价值。",
        "evidence_note": "矩阵 M1257：正式评分细则第100-102行在如何弘扬抗战精神中列发展的观点，证据等级为正式细则宽角度/术语支持。",
        "new_why": "材料强调从中华民族精神的发展和全民族抗战实践中深刻领会抗战精神内涵，说明抗战精神不是静止不变的文本，而是在特定历史实践中形成、丰富和升华的精神形态。细则把发展的观点列为如何弘扬的哲学角度，因此要写出其从历史形成到新时代继承发展的过程。",
        "new_answer": "①事物是变化发展的，民族精神也在历史实践中不断丰富和发展。②伟大抗战精神是中华民族精神在全民族抗战实践中的集中体现和发展升华，承载着爱国情怀、民族气节、英雄气概和必胜信念。③新时代弘扬抗战精神，既要继承其历史内涵，又要结合强国建设和民族复兴实践不断赋予现实表达，使其成为激励青年奋斗、担当时代使命的精神力量。",
    },
    {
        "queue_id": "T0377",
        "heading": "35. 2026海淀期末 第16题（主观题）",
        "node": "实践与认识（总）",
        "old_answer_excerpt": "对技术的认识和使用能力要在实践中形成并检验，学生应在学习、表达和问题解决实践中训练独立探索，而不是停留在复制粘贴和即时答案上。",
        "evidence_note": "矩阵 M1351：正式评分标准 page_001 Q16 列明“实践的观点”，且设问要求结合材料或个人体验作答。",
        "new_why": "设问要求谈数字原住民如何与技术进步相伴成长，不能只写“会用技术”，而要说明认识和能力从哪里来、怎样检验。正式评分标准把“实践的观点”列入可用角度，并要求结合材料或个人体验作答。材料中的学习、表达、探索和解决问题，都指向在真实实践中形成正确技术认识和行动能力。",
        "new_answer": "①实践是认识的基础，技术认识和技术使用能力要在真实学习、表达、探索和问题解决中形成、发展和检验。②数字原住民不能停留在复制粘贴、等待即时答案或被算法牵着走，而要主动在实践中提出问题、辨别信息、独立思考。③把技术使用放回真实学习和社会生活场景中，才能提升判断力、创造力和解决问题能力，与技术进步相伴成长。",
    },
    {
        "queue_id": "T0486",
        "heading": "35. 2026西城期末 第21题（主观题）",
        "node": "人民群众",
        "old_answer_excerpt": "科学制定五年规划要坚持人民群众观点，广泛听取人民意见、凝聚共识，使规划反映人民需要并服务人民利益。",
        "evidence_note": "矩阵 M1389：正式评分细则 page_005 Q21 角度2变通写明“人民至上/人民群众观点等”，证据等级为正式评分标准宽角度/术语支持。",
        "new_why": "题目问科学制定和接续实施五年规划的政治优势，材料中的开门问策、集思广益、全过程人民民主，指向规划不是少数人闭门制定，而是把人民意见纳入国家治理过程。正式评分细则 page_005 变通写明人民至上、人民群众观点等，因此应落到人民群众主体地位和群众路线。",
        "new_answer": "①人民群众是社会历史的主体，要坚持群众观点和群众路线。②科学制定五年规划，必须通过开门问策、集思广益、全过程人民民主等方式广泛听取人民意见，把人民的实践经验、现实需求和共同愿望转化为规划内容。③这样制定和接续实施规划，才能凝聚社会共识，保证规划反映人民需要、服务人民利益，并把人民群众的积极性创造性转化为推进现代化建设的力量。",
    },
    {
        "queue_id": "T0049",
        "heading": "35. 2026通州期末 第16题（主观题）",
        "node": "主观能动性 / 意识的能动作用",
        "old_answer_excerpt": "人在尊重规律基础上能动设计、修建、维护并创新表达治水智慧，推动工程持续发挥作用。",
        "evidence_note": "矩阵 M1431：正式细则第16题写明“同时发挥人的主观能动性，实现人与自然的和谐共生”。",
        "new_why": "都江堰跨越千年的治水智慧，既来自顺应地势水势，也来自人的主动设计、修建、维护和持续创新。正式细则明写“同时发挥人的主观能动性，实现人与自然的和谐共生”。因此作答不能只赞美工程古老，而要说明人在尊重客观条件基础上如何能动改造自然、保护工程并延续智慧。",
        "new_answer": "①人能够在尊重客观规律的基础上发挥主观能动性，认识条件、制定方案并改造世界。②都江堰的治水智慧体现为人们依据水势、地形和生产生活需要进行工程设计、修建、维护和治理，使水利工程长期发挥防洪、灌溉和生态调节作用。③今天继续保护利用都江堰，也要发挥人的能动性，在传承、修缮、治理和创新表达中实现人与自然和谐共生。",
    },
    {
        "queue_id": "T0107",
        "heading": "36. 2026朝阳期末 第16题（主观题）",
        "node": "联系的普遍性 / 联系的观点（总）",
        "old_answer_excerpt": "理解传统戏曲焕发生命力要坚持联系观点，把过去的戏曲传统、当代艺术表达和未来文化发展联系起来，在传统与现代元素融合中实现传承创新。",
        "evidence_note": "矩阵 M1285：正式阅卷细则第19-21行写传统戏曲连接过去、现在与未来，并补充“联系”，证据等级为正式细则宽角度/术语支持。",
        "new_why": "传统戏曲焕发生命力，不能理解为单纯保留旧形式，也不能理解为脱离传统的现代包装。正式阅卷细则把传统戏曲说成连接过去、现在与未来的文化纽带，并补充“联系”。材料中的传统精髓、当代表达、现代元素、观众需求和未来发展，要求用联系观点说明生命力从何而来。",
        "new_answer": "①联系具有普遍性，要把传统戏曲放在历史传承、现实表达和未来发展的联系中理解。②传统戏曲一方面承载过去的艺术积淀和文化记忆，另一方面又要同当代审美、现代技术、青年观众和现实生活发生联系。③只有在传统与现代、内容与形式、舞台表达与时代需求的融合中传承创新，传统戏曲才能不断焕发新的生命力。",
    },
]


def now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M +08")


def write_csv(path: Path, rows: list[dict[str, object]], fields: list[str]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fields})


def write_draft_files() -> None:
    rows = []
    for target in TARGETS:
        rows.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "evidence_note": target["evidence_note"],
                "new_why_chars": len(target["new_why"]),
                "new_answer_chars": len(target["new_answer"]),
                "new_why": target["new_why"],
                "new_answer": target["new_answer"],
            }
        )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_20260525.json").write_text(
        json.dumps({"updated": now(), "targets": rows}, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_20260525.csv",
        rows,
        ["queue_id", "heading", "node", "evidence_note", "new_why_chars", "new_answer_chars", "new_why", "new_answer"],
    )
    lines = [
        f"# P0 Thickness Repair Batch{BATCH} Draft 20260525",
        "",
        f"Updated: {now()}",
        "",
        "Status: `DRAFT_READY_FOR_DOCX_APPLICATION`",
        "",
        "- Scope: next 8 P0 subjective triple-thin rows after Batch10 refresh.",
        "- Repair method: thicken only existing why/answer paragraphs inside existing evidence-supported nodes.",
        "- Evidence boundary: all evidence notes cite current matrix rows; ordinary reference answers are not upgraded into scoring rubrics.",
        "",
        "| queue_id | node | heading | new why chars | new answer chars | evidence note |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in rows:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | {row['new_why_chars']} | {row['new_answer_chars']} | {row['evidence_note']} |"
        )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_20260525.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )


def main() -> int:
    write_draft_files()
    docx = helper.current_docx()
    backup = docx.with_name(f"{docx.stem}_backup_before_p0_thickness_batch{BATCH}_{datetime.now():%Y%m%d_%H%M%S}.docx")
    shutil.copy2(docx, backup)

    doc = Document(str(docx))
    blocks = helper.entry_blocks(doc)
    applied: list[dict[str, object]] = []
    for target in TARGETS:
        block = helper.matching_block(blocks, target)
        fields = block["fields"]
        field_paras = block["field_paras"]
        before_why = str(fields.get("why", ""))
        before_answer = str(fields.get("answer", ""))
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["why"])], "why", target["new_why"])
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["answer"])], "answer", target["new_answer"])
        applied.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "heading_para": block["heading_para"],
                "why_para": field_paras["why"],
                "answer_para": field_paras["answer"],
                "before_why_chars": len(before_why),
                "after_why_chars": len(target["new_why"]),
                "before_answer_chars": len(before_answer),
                "after_answer_chars": len(target["new_answer"]),
                "evidence_note": target["evidence_note"],
                "old_answer_excerpt": target["old_answer_excerpt"],
            }
        )

    doc.save(str(docx))
    with zipfile.ZipFile(docx) as z:
        zip_ok = "[Content_Types].xml" in z.namelist() and "word/document.xml" in z.namelist()

    result = {
        "updated": now(),
        "status": f"P0_BATCH{BATCH}_APPLIED_REQUIRES_RENDER_AND_EXTERNAL_RECHECK",
        "docx": str(docx),
        "backup_docx": str(backup),
        "zip_ok": zip_ok,
        "targets": len(TARGETS),
        "applied": applied,
        "boundary": [
            "This is a local thickness repair batch, not final acceptance.",
            "PDF/render outputs must be regenerated after this DOCX edit.",
            "External review remains pending for the current post-repair artifact.",
        ],
    }
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_20260525.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_20260525.csv",
        applied,
        [
            "queue_id",
            "heading",
            "node",
            "heading_para",
            "why_para",
            "answer_para",
            "before_why_chars",
            "after_why_chars",
            "before_answer_chars",
            "after_answer_chars",
            "evidence_note",
            "old_answer_excerpt",
        ],
    )
    lines = [
        f"# P0 Thickness Repair Batch{BATCH} Apply 20260525",
        "",
        f"Updated: {result['updated']}",
        "",
        f"Status: `{result['status']}`",
        "",
        f"- DOCX: `{docx.name}`.",
        f"- Backup: `{backup.name}`.",
        f"- Targets applied: `{len(applied)}` / `{len(TARGETS)}`.",
        f"- DOCX zip structural check: `{str(zip_ok).lower()}`.",
        "",
        "## Applied Rows",
        "",
        "| queue_id | node | heading | why chars | answer chars | paragraph refs |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in applied:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | "
            f"{row['before_why_chars']} -> {row['after_why_chars']} | "
            f"{row['before_answer_chars']} -> {row['after_answer_chars']} | "
            f"h={row['heading_para']}; why={row['why_para']}; answer={row['answer_para']} |"
        )
    lines.extend(
        [
            "",
            "## Boundary",
            "",
            "- This batch repairs only 8 additional P0 rows.",
            "- Remaining thickness queue, render QA, and model gates stay open until rechecked.",
        ]
    )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_20260525.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
