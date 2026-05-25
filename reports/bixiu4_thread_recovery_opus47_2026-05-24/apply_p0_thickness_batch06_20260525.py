from __future__ import annotations

import csv
import json
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

import apply_p0_thickness_batch01_20260525 as helper
from docx import Document


ROOT = Path(__file__).resolve().parent
BATCH = "06"

TARGETS = [
    {
        "queue_id": "T0025",
        "heading": "22. 2025朝阳一模 第21题（主观题）",
        "node": "一切从实际出发 / 实事求是 / 主观与客观具体的历史的统一",
        "old_answer_excerpt": "坚持一切从实际出发、实事求是。党员干部要从人民群众最关心、最直接、最现实的利益问题出发，找准来回跑、看病贵、落户难等实际难题，求真务实解决问题。",
        "evidence_note": "细则写坚持实事求是、求真务实；材料列举具体民生难题，支持一切从实际出发和实事求是落点。",
        "new_why": "材料把“大事小事观”“为民办事观”“造福人民的政绩观”落到“来回跑、看病贵、落户难”等具体难题，细则又写坚持实事求是、求真务实。这个题的触发点不是抽象表态，而是党员干部判断事情大小、确定工作重点时，必须从群众真实处境和现实利益诉求出发，把主观愿望同客观民情统一起来。",
        "new_answer": "①物质决定意识，要求一切从实际出发、实事求是。②中国共产党人的“大事小事观”要以人民群众真实需要为尺度，把“来回跑、看病贵、落户难”等具体问题作为工作对象，而不是追求表面声势或个人政绩。③只有求真务实、把难事办妥、把实事办实，才能使政绩经得起群众和实践检验，体现为民造福的正确政绩观。",
    },
    {
        "queue_id": "T0273",
        "heading": "23. 2025朝阳一模 第16题（主观题）",
        "node": "辩证否定 / 守正创新 / 创造性转化创新性发展",
        "old_answer_excerpt": "坚持辩证否定和守正创新：国产动画保留中华优秀传统文化的精神内核和故事资源，同时用现代技术、叙事和传播方式加以创新表达，使传统文化在新时代获得新的生命力。",
        "evidence_note": "细则写文化传承、现代科技、文化创新和传统文化与现代传播方式结合，支持守正创新与创造性转化落点。",
        "new_why": "材料中的国产动画不是把传统文化当作静态符号展示，也不是脱离传统另起炉灶，而是在人物、故事、审美意象和价值内核上延续中华优秀传统文化，同时借助现代视效技术、电影叙事和传播渠道让其进入当代观众经验。细则把文化传承、现代科技和文化创新连在一起，因此应落到辩证否定、守正创新和创造性转化。",
        "new_answer": "①辩证否定是联系和发展的环节，要求在继承中发展、在守正中创新。②国产动画让中华优秀传统文化绽放新光彩，首先在故事资源、价值精神和审美气质上守住文化根脉，避免把传统文化空心化。③同时，它又运用现代科技、电影语言和传播方式进行创新表达，使古老文化被当代观众理解、接受和喜爱，从而实现创造性转化、创新性发展。",
    },
    {
        "queue_id": "T0026",
        "heading": "23. 2026丰台期末 第22题（主观题）",
        "node": "一切从实际出发 / 实事求是 / 主观与客观具体的历史的统一",
        "old_answer_excerpt": "科学制定五年规划要坚持从中国具体实际出发。我国根据不同历史阶段的发展任务和现实条件接续编制规划，把顶层设计与实际国情结合起来，使战略安排符合现代化建设需要。",
        "evidence_note": "细则把“与中国具体实际相结合”列为可用知识；题目围绕五年规划的科学制定和接续实施。",
        "new_why": "五年规划之所以能成为治国理政经验，关键在于它不是照搬外部模式，也不是凭主观愿望画蓝图，而是依据我国不同历史阶段的发展水平、主要任务、资源条件和人民需要来确定目标与路径。细则明确列“与中国具体实际相结合”，说明答题应把规划的科学性落到从实际出发、使战略安排符合中国式现代化实践条件上。",
        "new_answer": "①物质决定意识，制定政策和规划必须坚持一切从实际出发，使主观方案符合客观国情。②我国接续编制和实施五年规划，是在不同历史阶段把现代化目标、发展任务、现实条件和群众意见结合起来，形成分阶段、可执行的战略安排。③这样既发挥顶层设计的引领作用，又避免脱离国情的空想，使国家治理能够沿着符合中国具体实际的方向持续推进。",
    },
    {
        "queue_id": "T0275",
        "heading": "25. 2025朝阳期末 第22题（主观题）",
        "node": "辩证否定 / 守正创新 / 创造性转化创新性发展",
        "old_answer_excerpt": "谱写马克思主义中国化时代化新篇章，要坚持守正创新。守正，就是坚持马克思主义基本原理和贯穿党的创新理论的立场观点方法；创新，就是立足新的实践要求和重大时代课题，不断提出新理念新思路新办法，使马克思主义在中国实践中持续焕发生命力。",
        "evidence_note": "PPT评分细则在辩证法部分列守正创新和用发展的观点看待理论创新，支持该原理落点。",
        "new_why": "题目问如何谱写马克思主义中国化时代化新篇章，材料同时出现“追求真理、揭示真理、笃行真理”和“回答重大时代课题”。这表明理论发展既不能丢掉马克思主义基本原理和立场观点方法，也不能停留在既有结论上，而要在新时代中国实践中推进理论创新。细则直接列守正创新和发展的观点，因此该节点需要讲清守正与创新的统一。",
        "new_answer": "①谱写马克思主义中国化时代化新篇章，要坚持守正创新，把坚持马克思主义基本原理同回答中国具体实践问题统一起来。②守正，就是坚持马克思主义的立场观点方法，守住人民立场、实践观点和科学真理，不把理论创新变成脱离根本方向的随意翻新。③创新，就是立足新时代国内外形势变化和重大实践课题，形成新的理念、思路和办法，使马克思主义在中国实践中不断丰富发展、焕发生命力。",
    },
    {
        "queue_id": "T0028",
        "heading": "25. 2026朝阳期末 第16题（主观题）",
        "node": "一切从实际出发 / 实事求是 / 主观与客观具体的历史的统一",
        "old_answer_excerpt": "传统戏曲焕发新生命力要坚持一切从实际出发，立足当代文化生活和戏曲艺术特点，在保留传统精髓的基础上赋予新的时代内涵，而不是简单复制粘贴。",
        "evidence_note": "正式阅卷细则补充列“从实际出发”；题面强调以当代之思叩问传统、保留戏曲精髓并赋予时代内涵。",
        "new_why": "题目讨论传统戏曲“焕发新的生命力”，材料不是要求把传统形式原封不动搬上舞台，也不是迎合流行而丢掉戏曲本体，而是要求从当代文化生活、观众审美、戏曲艺术规律和传统精髓之间的真实关系出发。细则补充列“从实际出发”，说明这一题应强调主观创作必须符合戏曲传承创新的客观条件。",
        "new_answer": "①物质决定意识，文化创新也要坚持一切从实际出发，不能脱离传统戏曲自身特点和当代文化生活需要。②传统戏曲要焕发新生命力，必须立足戏曲程式、唱念做打、审美精神等传统根基，同时把当代人的情感、问题意识和表达方式融入其中。③这样才能避免简单复制粘贴或表面包装，在保留传统精髓的基础上赋予时代内涵，使传统戏曲真正被今天的观众理解和喜爱。",
    },
    {
        "queue_id": "T0511",
        "heading": "25. 2026西城二模 第20题（主观题）",
        "node": "人民群众是社会历史的主体 / 群众观点与群众路线",
        "old_answer_excerpt": "树立正确政绩观，要坚持人民主体地位，把为民造福作为出发点和落脚点，让工作成效由人民需要和人民评价来检验，而不是追求个人声望或表面工程。",
        "evidence_note": "材料反复写最广大人民群众的最大利益、为民初心、为民造福和人民检验，支持人民主体与群众观点落点。",
        "new_why": "正确政绩观首先回答政绩为谁而树、由谁来评。材料反复把“最广大人民群众的最大利益”“为民初心”“为民造福”和“经得起人民检验”放在核心位置，说明政绩不能以个人升迁、短期声势或表面工程为标准，而要以人民群众的实践需要和真实评价为依据。因此应落到人民群众是社会历史的主体以及群众观点。",
        "new_answer": "①人民群众是社会历史的主体，是社会实践和社会评价的根本力量，党员干部必须坚持群众观点和群众路线。②树立正确政绩观，就要把为民造福作为出发点和落脚点，围绕人民群众最关心的现实问题开展工作，而不是追求个人声望、短期显绩或形象工程。③工作成效最终要由人民需要、人民获得感和人民评价来检验，只有让发展成果更多更公平惠及人民，政绩才真正站得住。",
    },
    {
        "queue_id": "T0029",
        "heading": "26. 2026西城期末 第16题第（2）问（主观题）",
        "node": "一切从实际出发 / 实事求是 / 主观与客观具体的历史的统一",
        "old_answer_excerpt": "追求真实就是坚持从客观事实出发，使主观认识符合客观实际，反对主观臆造和偏见，因此能形成说服力和共鸣。",
        "evidence_note": "正式评分细则写“真实是指客观事物及其规律的本来面目，追求真实就是实现主观与客观事实相符合”。",
        "new_why": "细则直接把“真实”界定为客观事物及其规律的本来面目，并把追求真实说成主观与客观事实相符合。这里的哲学抓手不是泛泛说情感真诚，而是认识和表达必须以客观事实、客观规律为依据。作品或判断之所以直抵人心，正是因为它摆脱主观臆造和偏见，能够让人看到真实生活和真实规律。",
        "new_answer": "①物质决定意识，正确意识必须如实反映客观存在；坚持真实，就是坚持从客观事实出发，使主观认识符合客观实际。②“真实”具有直抵人心的力量，是因为它呈现客观事物及其规律的本来面目，不靠夸饰、虚构和偏见制造感动。③当表达建立在真实生活、真实矛盾和真实情感之上时，受众更容易产生理解、信服和共鸣，也更能接受其中蕴含的价值判断。",
    },
    {
        "queue_id": "T0513",
        "heading": "27. 2025海淀期末 第22题（主观题）",
        "node": "人民群众是社会历史的主体 / 群众观点与群众路线",
        "old_answer_excerpt": "围绕愚公移山写作，可以落到人民群众是社会历史的主体。真正能搬走压在人民头上的“大山”的，不是神秘力量，而是全中国人民大众的团结奋斗。要依靠人民、团结人民，发挥人民群众的主体作用，才能汇聚推进民族复兴和现代化建设的磅礴力量。",
        "evidence_note": "PPT细则将人民主体地位列为必修四可选知识；材料把“上帝”解释为全中国人民大众共同挖山。",
        "new_why": "材料把愚公移山从个人神话转化为人民大众共同奋斗的历史叙事，毛泽东明确说“这个上帝不是别人，就是全中国的人民大众”。这说明推动历史前进、搬走压在人民头上的“大山”的根本力量不是神秘力量，也不是少数人的孤立意志，而是人民群众的团结实践。细则列人民主体地位，因此应把作文立意落到人民群众历史主体作用。",
        "new_answer": "①人民群众是社会历史的主体，是历史的创造者，是推动社会变革和民族复兴的根本力量。②围绕愚公移山写作时，要把“搬山”的力量理解为全中国人民大众的团结奋斗：只有人民一起行动、持续实践，才能搬走阻碍民族独立、人民解放和国家发展的“大山”。③因此卷面可写依靠人民、团结人民、服务人民，发挥人民群众主体作用，汇聚推进现代化建设和民族复兴的磅礴力量。",
    },
]


def now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M +08")


def write_csv(path: Path, rows: list[dict[str, object]], fields: list[str]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fields})


def write_draft_files() -> None:
    rows = []
    for target in TARGETS:
        rows.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "evidence_note": target["evidence_note"],
                "new_why_chars": len(target["new_why"]),
                "new_answer_chars": len(target["new_answer"]),
                "new_why": target["new_why"],
                "new_answer": target["new_answer"],
            }
        )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_20260525.json").write_text(
        json.dumps({"updated": now(), "targets": rows}, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_20260525.csv",
        rows,
        ["queue_id", "heading", "node", "evidence_note", "new_why_chars", "new_answer_chars", "new_why", "new_answer"],
    )
    lines = [
        f"# P0 Thickness Repair Batch{BATCH} Draft 20260525",
        "",
        f"Updated: {now()}",
        "",
        "Status: `DRAFT_READY_FOR_DOCX_APPLICATION`",
        "",
        "- Scope: next 8 P0 subjective triple-thin rows after Batch05 refresh.",
        "- Repair method: thicken only existing why/answer paragraphs inside existing nodes.",
        "- Evidence boundary: formal-rubric support is named only where the existing row already carries it; material-trigger rows are not upgraded to formal rubric evidence.",
        "",
        "| queue_id | node | heading | new why chars | new answer chars | evidence note |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in rows:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | {row['new_why_chars']} | {row['new_answer_chars']} | {row['evidence_note']} |"
        )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_20260525.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )


def main() -> int:
    write_draft_files()
    docx = helper.current_docx()
    backup = docx.with_name(f"{docx.stem}_backup_before_p0_thickness_batch{BATCH}_{datetime.now():%Y%m%d_%H%M%S}.docx")
    shutil.copy2(docx, backup)

    doc = Document(str(docx))
    blocks = helper.entry_blocks(doc)
    applied: list[dict[str, object]] = []
    for target in TARGETS:
        block = helper.matching_block(blocks, target)
        fields = block["fields"]
        field_paras = block["field_paras"]
        before_why = str(fields.get("why", ""))
        before_answer = str(fields.get("answer", ""))
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["why"])], "why", target["new_why"])
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["answer"])], "answer", target["new_answer"])
        applied.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "heading_para": block["heading_para"],
                "why_para": field_paras["why"],
                "answer_para": field_paras["answer"],
                "before_why_chars": len(before_why),
                "after_why_chars": len(target["new_why"]),
                "before_answer_chars": len(before_answer),
                "after_answer_chars": len(target["new_answer"]),
                "evidence_note": target["evidence_note"],
                "old_answer_excerpt": target["old_answer_excerpt"],
            }
        )

    doc.save(str(docx))
    with zipfile.ZipFile(docx) as z:
        zip_ok = "[Content_Types].xml" in z.namelist() and "word/document.xml" in z.namelist()

    result = {
        "updated": now(),
        "status": f"P0_BATCH{BATCH}_APPLIED_REQUIRES_RENDER_AND_EXTERNAL_RECHECK",
        "docx": str(docx),
        "backup_docx": str(backup),
        "zip_ok": zip_ok,
        "targets": len(TARGETS),
        "applied": applied,
        "boundary": [
            "This is a local thickness repair batch, not final acceptance.",
            "PDF/render outputs must be regenerated after this DOCX edit.",
            "External review remains pending for the current post-repair artifact.",
        ],
    }
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_20260525.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_20260525.csv",
        applied,
        [
            "queue_id",
            "heading",
            "node",
            "heading_para",
            "why_para",
            "answer_para",
            "before_why_chars",
            "after_why_chars",
            "before_answer_chars",
            "after_answer_chars",
            "evidence_note",
            "old_answer_excerpt",
        ],
    )
    lines = [
        f"# P0 Thickness Repair Batch{BATCH} Apply 20260525",
        "",
        f"Updated: {result['updated']}",
        "",
        f"Status: `{result['status']}`",
        "",
        f"- DOCX: `{docx.name}`.",
        f"- Backup: `{backup.name}`.",
        f"- Targets applied: `{len(applied)}` / `{len(TARGETS)}`.",
        f"- DOCX zip structural check: `{str(zip_ok).lower()}`.",
        "",
        "## Applied Rows",
        "",
        "| queue_id | node | heading | why chars | answer chars | paragraph refs |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in applied:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | "
            f"{row['before_why_chars']} -> {row['after_why_chars']} | "
            f"{row['before_answer_chars']} -> {row['after_answer_chars']} | "
            f"h={row['heading_para']}; why={row['why_para']}; answer={row['answer_para']} |"
        )
    lines.extend(
        [
            "",
            "## Boundary",
            "",
            "- This batch repairs only 8 additional P0 rows.",
            "- Remaining thickness queue, render QA, and model gates stay open until rechecked.",
        ]
    )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_20260525.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
