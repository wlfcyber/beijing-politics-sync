from __future__ import annotations

import csv
import json
import re
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent
RUN = ROOT.parent / "bixiu4_baodian_52_base_insert_second_mock_first_mock_audit_2026-05-24"
DELIVERY = RUN / "05_delivery"

LBRACKET = chr(0x3010)
RBRACKET = chr(0x3011)
LABELS = {
    "material": "".join(chr(c) for c in [0x6750, 0x6599, 0x89E6, 0x53D1, 0x70B9]),
    "question": "".join(chr(c) for c in [0x8BBE, 0x95EE]),
    "why": "".join(chr(c) for c in [0x4E3A, 0x4EC0, 0x4E48, 0x80FD, 0x60F3, 0x5230]),
    "answer": "".join(chr(c) for c in [0x7B54, 0x6848, 0x843D, 0x70B9]),
}
QUESTION_RE = re.compile(r"^\s*\d+\.\s*20(?:24|25|26).+" + chr(0x7B2C) + r".+" + chr(0x9898))


TARGETS = [
    {
        "queue_id": "T0159",
        "heading": "1. 2025海淀一模 第22题（主观题）",
        "node": "整体与部分",
        "old_answer_excerpt": "以“全面深化改革”或“中国式现代化”等任一论断为例,从联系的整体性出发说明该领域各项举措要相互配合、统筹兼顾、整体推进",
        "evidence_note": "Q22正式评分标准要求阐释系统观念及其重要性，源包支持全局局部、当前长远、全面系统等系统观念链条。",
        "new_why": "材料中的“弹钢琴”不是一般比喻，而是把十个指头、节奏、配合和成调作为一个整体系统来说明治国理政方法；五个“系统工程”论断又反复出现“统筹兼顾、全局性谋划、整体性推进”。这些关键词共同指向整体与部分的辩证关系：部分相互联系、相互作用，整体处于统帅地位，局部必须服从和服务于整体功能，所以应当用系统观念而不是孤立观点作答。",
        "new_answer": "①以全面深化改革或中国式现代化为例，各领域举措不是彼此孤立的单项任务，而是由经济、政治、文化、社会、生态等部分相互联系、相互作用构成的有机整体。②坚持系统观念，就是从整体出发统筹全局与局部、当前与长远，使各项举措相互配合、协调推进，形成整体功能优于部分功能简单相加的效果。③如果只抓局部或单打一，就会顾此失彼；因此系统观念能够把重大实践中的多方面关系组织起来，是具有基础性的思想和工作方法。",
    },
    {
        "queue_id": "T0267",
        "heading": "1. 2025海淀期末 第17题第（2）问（主观题）",
        "node": "事物发展是前进性与曲折性的统一",
        "old_answer_excerpt": "纪念失败者，是因为科学发展是前进性与曲折性的统一",
        "evidence_note": "PPT细则明列“事物发展是前进性和曲折性的统一”。",
        "new_why": "材料把两层意思放在一起：一方面，永动机制造尝试无一例外失败，说明科学探索的道路充满反复、挫折和曲折；另一方面，这些失败又推动人们转向能量转化规律，最终促成能量守恒定律发现，说明科学发展的总趋势是前进和上升。看到“失败并没有白费，而是推动重大发现”这一链条，就应落到前进性与曲折性的统一。",
        "new_answer": "①科学发展的前途是光明的，人类认识能够在实践探索中不断接近和发现真理，能量守恒定律的形成就是科学认识前进的结果。②科学发展的道路又是曲折的，永动机探索中的失败、否定和反复并不等于没有价值，而是暴露问题、积累经验、促使人们寻找更深层原因。③纪念失败者，就是承认科学进步往往是在曲折中前进，失败探索同样构成通向真理的重要环节。",
    },
    {
        "queue_id": "T0515",
        "heading": "1. 2025西城期末 第21题（主观题）",
        "node": "社会发展的两大基本规律和基本矛盾",
        "old_answer_excerpt": "改革是社会主义制度的自我完善和自我发展，是解决我国社会基本矛盾、推动社会主义社会前进的直接动力",
        "evidence_note": "正式细则明列改革是社会主义制度的自我完善和自我发展、是解决社会基本矛盾和推动社会主义社会发展的直接动力。",
        "new_why": "题目问“改革和法治相统一”，材料抓手是“改革和法治相互促动、完善制度法规、推动制度更加成熟定型”。这不是单纯政治表述，而是把改革放到社会主义社会基本矛盾的解决过程里：生产关系、上层建筑需要随着社会发展不断调整完善，改革提供动力，法治把改革成果制度化、规范化。",
        "new_answer": "①改革是社会主义制度的自我完善和自我发展，是解决社会主义社会基本矛盾、推动社会前进的直接动力。②法治能够以制度和规则确认改革成果、规范改革过程，使改革在稳定秩序中持续推进。③坚持改革和法治相统一，就是用改革破解制约发展的体制机制问题，又用法治巩固制度成果、保障改革行稳致远，从而推动中国特色社会主义制度更加成熟更加定型。",
    },
    {
        "queue_id": "T0348",
        "heading": "1. 2026朝阳期中 第18题（主观题）",
        "node": "矛盾的普遍性",
        "old_answer_excerpt": "要全面、一分为二地看待AI提供情绪价值",
        "evidence_note": "正式评分细则支持从支持理由、反对理由和如何做三个层面展开，对AI情绪价值作一分为二分析。",
        "new_why": "材料同时呈现支持者和反对者的理由：AI能实时回应、提供表达空间，也存在共情缺陷、依赖风险和削弱真实交往的问题。同一事物内部既有积极方面又有消极方面，且二者共存于AI情绪陪伴这一对象之中，这正是矛盾普遍存在、要用一分为二观点分析问题的典型设问。",
        "new_answer": "①矛盾具有普遍性，任何事物都包含对立统一的两个方面，AI提供情绪价值也不例外。②应肯定其在即时回应、陪伴表达、缓解部分情绪需求上的积极作用，同时看到其缺少真实生命体验、深度共情和现实责任关系，可能带来依赖和交往弱化。③正确态度不是简单赞成或反对，而是在合理使用中发挥技术长处，并以真实人际交往、主体能力提升和价值判断来限制其风险。",
    },
    {
        "queue_id": "T0095",
        "heading": "10. 2025丰台期末 第16题（主观题）",
        "node": "规律的客观性",
        "old_answer_excerpt": "做到胸中有“数”，首先要承认事物发展有客观规律和客观数量界限",
        "evidence_note": "细则把“数从哪里来”对应到实践、实际、规律和发挥主观能动性。",
        "new_why": "材料中的“数”不是随手估计的数字，而是从实践调查、实际情况和事物发展规律中来；“数量界限”又说明事物发展有客观的度，超过或不到一定界限都会影响质量变化。看到“数从哪里来、数当如何定、不注意数量界限就犯错误”，就应想到规律具有客观性，必须尊重规律、按规律办事。",
        "new_answer": "①胸中有“数”，首先要从实践和实际出发，通过调查研究把握客观情况，而不能凭主观愿望定数。②事物发展有客观规律，也有数量界限和度的规定，决策用数必须符合规律、符合实际条件。③在尊重客观规律的基础上发挥主观能动性，才能把数字、比例、范围和节奏转化为科学决策依据，避免盲目冒进、主观臆断或胸中无数造成失误。",
    },
    {
        "queue_id": "T0476",
        "heading": "10. 2026丰台期末 第22题（主观题）",
        "node": "正确认识指导实践",
        "old_answer_excerpt": "正确认识对实践具有指导作用。五年规划把对发展阶段、任务和趋势的科学认识转化为战略安排和行动方案",
        "evidence_note": "细则列“正确认识”，当前框架对应“正确认识指导实践”。",
        "new_why": "材料强调五年规划科学制定、接续实施和“一张蓝图绘到底”，说明规划不是空泛口号，而是对发展阶段、发展目标、战略任务和实践路径的认识成果。认识形成后又转化为指标、任务、政策和行动安排，反过来指导各地区各部门的现代化建设实践，因此应落到正确认识对实践的指导作用。",
        "new_answer": "①实践决定认识，五年规划必须建立在对我国发展阶段、现实条件、主要任务和未来趋势的科学把握之上。②正确认识对实践具有指导作用，科学规划能够把发展认识转化为目标体系、战略安排和行动方案，引导各方面有方向、有步骤地推进现代化建设。③接续实施规划并保持“一张蓝图绘到底”，体现用正确认识统领长期实践，使实践少走弯路、形成连续推进的合力。",
    },
    {
        "queue_id": "T0643",
        "heading": "10. 2026海淀期末 第16题（主观题）",
        "node": "实现人生价值",
        "old_answer_excerpt": "数字原住民应在正确使用技术、发展思考能力和参与真实实践中促进自身成长",
        "evidence_note": "正式评分标准在行动角度列出实践、主观能动性和实现人生价值。",
        "new_why": "设问要求谈“数字原住民”如何与技术进步相伴成长，核心不是评价技术本身，而是回答人在技术环境中怎样成长成才、创造和实现价值。细则把实践、主观能动性、实现人生价值列为行动角度，说明学生要把技术使用放到现实实践、能力发展、社会责任和人生价值实现的链条中理解。",
        "new_answer": "①数字原住民要正确发挥主观能动性，把技术作为学习、创造和服务社会的工具，而不是被技术支配或沉迷于便利。②要在真实实践中发展判断力、思考力和解决问题能力，把数字技能转化为参与社会生活、回应现实需要的本领。③人生价值在劳动和奉献中创造和实现，青年应把个人成长同社会需要结合起来，让技术进步成为提升自我、贡献社会和实现人生价值的条件。",
    },
    {
        "queue_id": "T0276",
        "heading": "10. 2026通州期末 第16题（主观题）",
        "node": "事物发展是前进性与曲折性的统一",
        "old_answer_excerpt": "都江堰跨越千年说明发展方向具有前进性",
        "evidence_note": "正式细则第16题发展观括注“前途光明道路曲折”。",
        "new_why": "都江堰跨越千年的治水智慧本身体现了发展不是一次完成的静态成果，而是在长期维护、修缮、治理和创新中延续。细则明确给出“前途光明道路曲折”，材料也提示艰难曲折不能阻挡历史前进，因此应把都江堰的延续理解为在曲折实践中不断向前发展的过程。",
        "new_answer": "①事物发展的总趋势是前进的、上升的，都江堰能够跨越千年，说明科学治水理念和人与自然和谐共生的方向具有生命力。②事物发展的道路又是曲折的，水利工程的维护、灾害应对、技术更新和保护利用都会遇到困难与反复。③阐释其治水智慧时，应看到人们在长期实践中正视曲折、解决问题、不断创新，使都江堰在历史发展中持续发挥作用。",
    },
]


def now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M +08")


def current_docx() -> Path:
    candidates = [
        p
        for p in DELIVERY.glob("*.docx")
        if "backup" not in p.name.lower() and not p.name.startswith("~$")
    ]
    if len(candidates) != 1:
        raise RuntimeError(f"Expected one current docx, found {len(candidates)}: {[p.name for p in candidates]}")
    return candidates[0]


def label_prefix(role: str) -> str:
    return f"{LBRACKET}{LABELS[role]}{RBRACKET}"


def strip_label(text: str, role: str) -> str:
    prefix = label_prefix(role)
    if text.startswith(prefix):
        return text[len(prefix) :].strip()
    return ""


def set_labeled_paragraph(paragraph, role: str, value: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    label = paragraph.add_run(f"{label_prefix(role)} ")
    label.bold = True
    paragraph.add_run(value)


def entry_blocks(doc: Document) -> list[dict[str, object]]:
    blocks: list[dict[str, object]] = []
    current: dict[str, object] | None = None
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if QUESTION_RE.match(text):
            if current is not None:
                current["end_para"] = idx
                blocks.append(current)
            current = {
                "heading": text,
                "heading_para": idx,
                "paragraphs": [],
                "fields": {},
                "field_paras": {},
            }
            continue
        if current is None:
            continue
        current["paragraphs"].append(idx)
        for role in LABELS:
            value = strip_label(text, role)
            if value and role not in current["fields"]:
                current["fields"][role] = value
                current["field_paras"][role] = idx
    if current is not None:
        current["end_para"] = len(doc.paragraphs)
        blocks.append(current)
    return blocks


def matching_block(blocks: list[dict[str, object]], target: dict[str, str]) -> dict[str, object]:
    matches = [block for block in blocks if block["heading"] == target["heading"]]
    answer_matches = [
        block
        for block in matches
        if target["old_answer_excerpt"] in str(block.get("fields", {}).get("answer", ""))
    ]
    if len(answer_matches) != 1:
        raise RuntimeError(
            f"{target['queue_id']} expected one answer-excerpt match; "
            f"heading matches={len(matches)}, answer matches={len(answer_matches)}"
        )
    return answer_matches[0]


def write_csv(path: Path, rows: list[dict[str, object]], fields: list[str]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fields})


def write_draft_files() -> None:
    draft_rows = []
    for target in TARGETS:
        draft_rows.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "evidence_note": target["evidence_note"],
                "new_why_chars": len(target["new_why"]),
                "new_answer_chars": len(target["new_answer"]),
                "new_why": target["new_why"],
                "new_answer": target["new_answer"],
            }
        )
    (ROOT / "P0_THICKNESS_REPAIR_BATCH01_DRAFT_20260525.json").write_text(
        json.dumps({"updated": now(), "targets": draft_rows}, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / "P0_THICKNESS_REPAIR_BATCH01_DRAFT_20260525.csv",
        draft_rows,
        ["queue_id", "heading", "node", "evidence_note", "new_why_chars", "new_answer_chars", "new_why", "new_answer"],
    )
    lines = [
        "# P0 Thickness Repair Batch01 Draft 20260525",
        "",
        f"Updated: {now()}",
        "",
        "Status: `DRAFT_READY_FOR_DOCX_APPLICATION`",
        "",
        "- Scope: first 8 P0 subjective triple-thin rows.",
        "- Repair method: thicken only the existing node-supported why/answer paragraphs; no new unsupported principle is introduced.",
        "- Matching rule for application: exact heading plus old answer excerpt, so duplicate same-question headings in other nodes are not overwritten.",
        "",
        "| queue_id | node | heading | new why chars | new answer chars | evidence note |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in draft_rows:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | {row['new_why_chars']} | {row['new_answer_chars']} | {row['evidence_note']} |"
        )
    (ROOT / "P0_THICKNESS_REPAIR_BATCH01_DRAFT_20260525.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )


def main() -> int:
    write_draft_files()
    docx = current_docx()
    backup = docx.with_name(f"{docx.stem}_backup_before_p0_thickness_batch01_{datetime.now():%Y%m%d_%H%M%S}.docx")
    shutil.copy2(docx, backup)

    doc = Document(str(docx))
    blocks = entry_blocks(doc)
    applied: list[dict[str, object]] = []
    for target in TARGETS:
        block = matching_block(blocks, target)
        fields = block["fields"]
        field_paras = block["field_paras"]
        before_why = str(fields.get("why", ""))
        before_answer = str(fields.get("answer", ""))
        why_para = doc.paragraphs[int(field_paras["why"])]
        answer_para = doc.paragraphs[int(field_paras["answer"])]
        set_labeled_paragraph(why_para, "why", target["new_why"])
        set_labeled_paragraph(answer_para, "answer", target["new_answer"])
        applied.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "heading_para": block["heading_para"],
                "why_para": field_paras["why"],
                "answer_para": field_paras["answer"],
                "before_why_chars": len(before_why),
                "after_why_chars": len(target["new_why"]),
                "before_answer_chars": len(before_answer),
                "after_answer_chars": len(target["new_answer"]),
                "evidence_note": target["evidence_note"],
                "old_answer_excerpt": target["old_answer_excerpt"],
            }
        )

    doc.save(str(docx))
    with zipfile.ZipFile(docx) as z:
        zip_ok = "[Content_Types].xml" in z.namelist() and "word/document.xml" in z.namelist()

    result = {
        "updated": now(),
        "status": "P0_BATCH01_APPLIED_REQUIRES_RENDER_AND_EXTERNAL_RECHECK",
        "docx": str(docx),
        "backup_docx": str(backup),
        "zip_ok": zip_ok,
        "targets": len(TARGETS),
        "applied": applied,
        "boundary": [
            "This is a local thickness repair batch, not final acceptance.",
            "PDF/render outputs must be regenerated after this DOCX edit.",
            "GPTPro and Claude Opus review must be retried after all planned thickness repairs, not just this batch.",
        ],
    }
    (ROOT / "P0_THICKNESS_REPAIR_BATCH01_APPLY_20260525.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / "P0_THICKNESS_REPAIR_BATCH01_APPLY_20260525.csv",
        applied,
        [
            "queue_id",
            "heading",
            "node",
            "heading_para",
            "why_para",
            "answer_para",
            "before_why_chars",
            "after_why_chars",
            "before_answer_chars",
            "after_answer_chars",
            "evidence_note",
            "old_answer_excerpt",
        ],
    )
    lines = [
        "# P0 Thickness Repair Batch01 Apply 20260525",
        "",
        f"Updated: {result['updated']}",
        "",
        f"Status: `{result['status']}`",
        "",
        f"- DOCX: `{docx.name}`.",
        f"- Backup: `{backup.name}`.",
        f"- Targets applied: `{len(applied)}` / `{len(TARGETS)}`.",
        f"- DOCX zip structural check: `{str(zip_ok).lower()}`.",
        "",
        "## Applied Rows",
        "",
        "| queue_id | node | heading | why chars | answer chars | paragraph refs |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in applied:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | "
            f"{row['before_why_chars']} -> {row['after_why_chars']} | "
            f"{row['before_answer_chars']} -> {row['after_answer_chars']} | "
            f"h={row['heading_para']}; why={row['why_para']}; answer={row['answer_para']} |"
        )
    lines.extend(
        [
            "",
            "## Boundary",
            "",
            "- This batch repairs only 8 P0 rows.",
            "- The full thickness queue remains open until the refreshed audit shows no unresolved density candidates or the user approves a narrower boundary.",
            "- Render/PDF QA and later real-model review remain required.",
        ]
    )
    (ROOT / "P0_THICKNESS_REPAIR_BATCH01_APPLY_20260525.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
