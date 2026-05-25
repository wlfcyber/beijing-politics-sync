from __future__ import annotations

import csv
import json
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

import apply_p0_thickness_batch01_20260525 as helper
from docx import Document


ROOT = Path(__file__).resolve().parent
BATCH = "09"

TARGETS = [
    {
        "queue_id": "T0210",
        "heading": "32. 2025朝阳期末 第16题（主观题）",
        "node": "发展的观点 / 发展的普遍性",
        "old_answer_excerpt": "用文化同世界对话，要坚持发展的观点，把中华优秀传统文化放在过去、现在、未来的连续发展中理解。既要看到《论语》、汉代文物、春节等传统文化资源的历史根脉，也要通过当代阐释、互动展览和国际传播使其获得新的生命力，推动不同文化共同发展。",
        "evidence_note": "PPT评分细则第134行明列“发展”，当前矩阵 M1034 以强细则支持该节点。",
        "new_why": "材料中的《论语》、汉代文物、春节等不是被静态展示为旧物，而是通过当代阐释、互动展览、国际传播进入新的交流场景。PPT评分细则明列“发展”，说明作答不能只写传统文化值得传播，还要看到传统文化在历史延续中不断获得新的表现形式和生命力。看到“过去、现在、未来连续起来”“传统资源进入当代对话”，就应落到发展的观点。",
        "new_answer": "①事物是变化发展的，中华优秀传统文化也要放在历史延续和时代创新的过程中理解。②用文化同世界对话，既要讲清《论语》、汉代文物、春节等传统文化资源的历史根脉，也要通过当代阐释、互动展览和国际传播让其进入现代生活和世界交流。③这样才能使传统文化在创新发展中获得新的生命力，推动不同文化在交流互鉴中共同发展。",
    },
    {
        "queue_id": "T0417",
        "heading": "32. 2025朝阳期末 第22题（主观题）",
        "node": "实践是认识的基础",
        "old_answer_excerpt": "不断谱写马克思主义中国化时代化新篇章，要坚持实践是认识的基础。要从国内外形势新变化和中国具体实践新要求出发，总结实践经验、回答实践课题，在实践基础上形成新的理论认识。",
        "evidence_note": "PPT评分细则第337-340行明列实践观点和实践基础上的理论创新，当前矩阵 M1044 以强细则支持该节点。",
        "new_why": "设问要求说明如何不断谱写马克思主义中国化时代化新篇章，材料抓手是国内外形势新变化、中国具体实践新要求和实践经验总结。PPT评分细则明确把“实践观点/实践基础上的理论创新”列为得分角度，因此这里不是抽象谈理论自我更新，而是要说明新的理论认识来自实践、服务实践，并在回答现实课题中发展。",
        "new_answer": "①实践是认识的来源和发展动力，理论创新必须建立在中国具体实践基础上。②面对国内外形势新变化和中国实践新要求，要深入总结实践经验、研究实践中的新问题，把人民创造和时代课题转化为理论概括。③只有坚持从实践中来、到实践中去，才能形成新的理论认识，持续谱写马克思主义中国化时代化新篇章。",
    },
    {
        "queue_id": "T0180",
        "heading": "32. 2026丰台期末 第22题（主观题）",
        "node": "系统观念 / 系统优化",
        "old_answer_excerpt": "五年规划体现系统优化方法。通过中央统筹、地方分解、部门协同，把目标、任务、资源和政策组织为统一系统，推动国家治理形成“一盘棋”效能。",
        "evidence_note": "正式细则第452行明列“系统优化”，当前矩阵 M1214 以正式细则支持该节点。",
        "new_why": "五年规划不是单个部门的孤立计划，而是把发展目标、阶段任务、资源配置、政策工具和执行主体组织成一个治理系统。正式细则第452行明列“系统优化”，材料也写到中央统筹、地方分解、部门协同和全国一盘棋。由这些词可以判断，答题重点应放在立足整体、统筹要素、优化结构，而不是只写一般计划管理。",
        "new_answer": "①系统优化方法要求立足整体，把各要素、各层级、各环节统筹起来，形成整体功能。②五年规划通过中央统筹确定总目标，地方分解落实任务，部门协同配置政策和资源，把长期战略、阶段安排和具体行动连接为统一系统。③这种“一盘棋”治理能够避免各自为政和局部失衡，使国家治理在整体协同中提高效能。",
    },
    {
        "queue_id": "T0048",
        "heading": "32. 2026海淀期中 第22题第（2）问（主观题）",
        "node": "主观能动性 / 意识的能动作用",
        "old_answer_excerpt": "中华民族伟大复兴要在尊重历史大势和发展规律的基础上发挥主观能动性，把民族复兴的目标、信心和行动统一起来，通过主动奋斗把有利态势转化为现实成果。",
        "evidence_note": "正式阅卷细则 page_094 Q22(2) 常见知识列明“发挥主观能动性”，当前矩阵 M1315 以正式阅卷细则宽角度支持该节点。",
        "new_why": "材料围绕中华民族伟大复兴的“势”展开，既强调历史大势和发展规律，也强调信心、目标和主动奋斗。阅卷细则 page_094 明列“发挥主观能动性”，并要求分析扣在“势”上，说明不能把有利态势当作自动实现的结果。学生需要写清：尊重规律是前提，主动奋斗和正确意识把可能性转化为现实成果。",
        "new_answer": "①发挥主观能动性必须以尊重客观规律和历史大势为基础，不能脱离现实条件空喊目标。②中华民族伟大复兴有历史趋势和发展基础，但这种“势”还需要通过坚定信心、明确目标、主动作为来把握和推动。③把民族复兴的目标、信念和行动统一起来，才能把有利态势转化为现实成果，在主动奋斗中推进复兴进程。",
    },
    {
        "queue_id": "T0497",
        "heading": "32. 2026海淀期中 第22题第（2）问（主观题）",
        "node": "人民群众",
        "old_answer_excerpt": "人民群众是历史的创造者，是推进中华民族伟大复兴的力量之源。坚持人民主体地位，汇聚全体人民智慧和力量，才能使复兴大势转化为现实力量。",
        "evidence_note": "正式阅卷细则 page_094 Q22(2) 常见知识列明“人民群众主体地位”，当前矩阵 M1317 以正式阅卷细则支持该节点。",
        "new_why": "同一道题在“势”的分析中还要求看到主体力量。阅卷细则 page_094 明列“人民群众主体地位”，示例也写到人民群众是推动中华民族伟大复兴的力量之源。材料若只写历史大势，会漏掉谁来顺势而为、谁来把趋势变成现实的问题；因此应把复兴大势同人民主体地位、群众智慧和群众力量联系起来。",
        "new_answer": "①人民群众是社会历史的主体，是推动中华民族伟大复兴的根本力量。②复兴大势不是少数人单独推动的结果，必须坚持人民主体地位，尊重人民创造，凝聚全体人民的智慧、力量和奋斗热情。③只有把历史趋势同人民群众的实践力量结合起来，才能把复兴的有利态势转化为现实力量和持续成果。",
    },
    {
        "queue_id": "T0211",
        "heading": "33. 2025朝阳期末 第22题（主观题）",
        "node": "发展的观点 / 发展的普遍性",
        "old_answer_excerpt": "谱写马克思主义中国化时代化新篇章，要坚持发展的观点看待理论创新。随着国内外形势新变化和实践新要求不断出现，马克思主义中国化时代化也要在长期坚持中不断丰富发展，用发展着的理论回答发展着的时代课题。",
        "evidence_note": "PPT评分细则第332-336行明列“发展的观点”，当前矩阵 M1042 以强细则支持该节点。",
        "new_why": "材料不是要求背诵既有理论结论，而是要求说明在新形势、新实践、新课题中怎样继续推进理论创新。PPT评分细则明列“发展的观点”，说明马克思主义中国化时代化要被理解为随实践和时代课题不断丰富发展的过程。看到“新变化”“新要求”“新篇章”，就应从发展的普遍性和理论与时代共同前进来组织答案。",
        "new_answer": "①事物是变化发展的，理论创新也必须随着实践和时代的发展而不断前进。②国内外形势新变化和中国具体实践新要求不断出现，要求马克思主义中国化时代化在坚持基本立场观点方法的基础上不断丰富和发展。③用发展着的理论回答发展着的时代课题，才能使理论保持生命力，持续指导新的伟大实践。",
    },
    {
        "queue_id": "T0533",
        "heading": "33. 2025海淀期末 第22题（主观题）",
        "node": "价值观的导向作用",
        "old_answer_excerpt": "围绕愚公移山写作，可以落到价值观的导向作用。愚公移山精神引导我们把国家发展、民族复兴和人民幸福作为奋斗目标，面对困难不退缩、面对任务不懈怠，以正确价值观凝聚坚持不懈、团结奋斗的行动力量。",
        "evidence_note": "PPT细则第520-521行将“价值观”列为必修四可选知识，当前矩阵 M1111 以强细则/PPT评标细则支持该节点。",
        "new_why": "愚公移山精神之所以能用于写作，不只是因为故事感人，而是因为它提供了面对困难时怎样判断目标价值、怎样选择行动方向的价值引领。PPT细则把“价值观”列为可选知识，说明可以从价值观导向作用切入。材料中的国家发展、民族复兴、人民幸福和坚持不懈奋斗，都是价值目标对认识和行动的引导。",
        "new_answer": "①价值观对人们认识世界和改造世界具有导向作用，影响人们的价值判断、目标选择和行动方式。②愚公移山精神引导我们把国家发展、民族复兴和人民幸福作为值得追求的目标，在困难面前不退缩、不懈怠。③用正确价值观凝聚坚持不懈、团结奋斗的行动力量，才能把精神认同转化为推进事业发展的现实行动。",
    },
    {
        "queue_id": "T0108",
        "heading": "33. 2026丰台期末 第16题（主观题）",
        "node": "联系的普遍性 / 联系的观点（总）",
        "old_answer_excerpt": "分析留白要坚持联系的观点。留白与周围景物、空间节奏和生活情境相互作用，共同构成有张弛、有层次的审美整体，不能把留白孤立理解为空缺。",
        "evidence_note": "正式细则第51-53行明列留白与周围事物相互联系、相互影响以及坚持联系的观点，当前矩阵 M1207 以正式细则支持该节点。",
        "new_why": "“留白”容易被学生误解为空缺或少画一点，但细则明确写到留白与周围事物相互联系、相互影响，并要求坚持联系的观点。材料中的留白、景物、空间节奏和生活情境共同形成审美效果，说明它的意义来自关系结构，而不是孤立存在。因此答题要抓住联系的普遍性和整体审美效果。",
        "new_answer": "①联系具有普遍性，事物总是在同其他事物的相互联系中存在和发挥作用。②留白不是孤立的空缺，而是与周围景物、画面空间、节奏层次和生活情境相互作用，共同构成有张弛、有意境的审美整体。③分析留白要把它放在整体关系中理解，看到它如何通过与周边元素的联系增强画面表达和文化意味。",
    },
]


def now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M +08")


def write_csv(path: Path, rows: list[dict[str, object]], fields: list[str]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fields})


def write_draft_files() -> None:
    rows = []
    for target in TARGETS:
        rows.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "evidence_note": target["evidence_note"],
                "new_why_chars": len(target["new_why"]),
                "new_answer_chars": len(target["new_answer"]),
                "new_why": target["new_why"],
                "new_answer": target["new_answer"],
            }
        )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_20260525.json").write_text(
        json.dumps({"updated": now(), "targets": rows}, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_20260525.csv",
        rows,
        ["queue_id", "heading", "node", "evidence_note", "new_why_chars", "new_answer_chars", "new_why", "new_answer"],
    )
    lines = [
        f"# P0 Thickness Repair Batch{BATCH} Draft 20260525",
        "",
        f"Updated: {now()}",
        "",
        "Status: `DRAFT_READY_FOR_DOCX_APPLICATION`",
        "",
        "- Scope: next 8 P0 subjective triple-thin rows after Batch08 refresh.",
        "- Repair method: thicken only existing why/answer paragraphs inside existing nodes.",
        "- Evidence boundary: evidence notes are carried from the current coverage matrix; no ordinary reference answer is upgraded into a scoring rubric.",
        "",
        "| queue_id | node | heading | new why chars | new answer chars | evidence note |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in rows:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | {row['new_why_chars']} | {row['new_answer_chars']} | {row['evidence_note']} |"
        )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_20260525.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )


def main() -> int:
    write_draft_files()
    docx = helper.current_docx()
    backup = docx.with_name(f"{docx.stem}_backup_before_p0_thickness_batch{BATCH}_{datetime.now():%Y%m%d_%H%M%S}.docx")
    shutil.copy2(docx, backup)

    doc = Document(str(docx))
    blocks = helper.entry_blocks(doc)
    applied: list[dict[str, object]] = []
    for target in TARGETS:
        block = helper.matching_block(blocks, target)
        fields = block["fields"]
        field_paras = block["field_paras"]
        before_why = str(fields.get("why", ""))
        before_answer = str(fields.get("answer", ""))
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["why"])], "why", target["new_why"])
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["answer"])], "answer", target["new_answer"])
        applied.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "heading_para": block["heading_para"],
                "why_para": field_paras["why"],
                "answer_para": field_paras["answer"],
                "before_why_chars": len(before_why),
                "after_why_chars": len(target["new_why"]),
                "before_answer_chars": len(before_answer),
                "after_answer_chars": len(target["new_answer"]),
                "evidence_note": target["evidence_note"],
                "old_answer_excerpt": target["old_answer_excerpt"],
            }
        )

    doc.save(str(docx))
    with zipfile.ZipFile(docx) as z:
        zip_ok = "[Content_Types].xml" in z.namelist() and "word/document.xml" in z.namelist()

    result = {
        "updated": now(),
        "status": f"P0_BATCH{BATCH}_APPLIED_REQUIRES_RENDER_AND_EXTERNAL_RECHECK",
        "docx": str(docx),
        "backup_docx": str(backup),
        "zip_ok": zip_ok,
        "targets": len(TARGETS),
        "applied": applied,
        "boundary": [
            "This is a local thickness repair batch, not final acceptance.",
            "PDF/render outputs must be regenerated after this DOCX edit.",
            "External review remains pending for the current post-repair artifact.",
        ],
    }
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_20260525.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_20260525.csv",
        applied,
        [
            "queue_id",
            "heading",
            "node",
            "heading_para",
            "why_para",
            "answer_para",
            "before_why_chars",
            "after_why_chars",
            "before_answer_chars",
            "after_answer_chars",
            "evidence_note",
            "old_answer_excerpt",
        ],
    )
    lines = [
        f"# P0 Thickness Repair Batch{BATCH} Apply 20260525",
        "",
        f"Updated: {result['updated']}",
        "",
        f"Status: `{result['status']}`",
        "",
        f"- DOCX: `{docx.name}`.",
        f"- Backup: `{backup.name}`.",
        f"- Targets applied: `{len(applied)}` / `{len(TARGETS)}`.",
        f"- DOCX zip structural check: `{str(zip_ok).lower()}`.",
        "",
        "## Applied Rows",
        "",
        "| queue_id | node | heading | why chars | answer chars | paragraph refs |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in applied:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | "
            f"{row['before_why_chars']} -> {row['after_why_chars']} | "
            f"{row['before_answer_chars']} -> {row['after_answer_chars']} | "
            f"h={row['heading_para']}; why={row['why_para']}; answer={row['answer_para']} |"
        )
    lines.extend(
        [
            "",
            "## Boundary",
            "",
            "- This batch repairs only 8 additional P0 rows.",
            "- Remaining thickness queue, render QA, and model gates stay open until rechecked.",
        ]
    )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_20260525.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
