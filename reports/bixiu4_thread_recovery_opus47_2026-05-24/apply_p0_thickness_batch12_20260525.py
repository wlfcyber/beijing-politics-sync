from __future__ import annotations

import csv
import json
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

import apply_p0_thickness_batch01_20260525 as helper
from docx import Document


ROOT = Path(__file__).resolve().parent
BATCH = "12"

TARGETS = [
    {
        "queue_id": "T0200",
        "heading": "36. 2026朝阳期末 第16题（主观题）",
        "node": "发展的观点 / 发展的普遍性",
        "old_answer_excerpt": "传统戏曲焕发生命力要用发展的观点看问题，在继承精髓的基础上创造性转化、创新性发展，使传统艺术随着时代和实践变化不断获得新的表达。",
        "evidence_note": "矩阵 M1283：正式阅卷细则第18、21行写明创造性转化、创新性发展，并补充“发展”。",
        "new_why": "题面围绕传统戏曲在当代艺术创作中的转化和突破，不能只写“保护传统”，也不能把创新理解成脱离传统的包装。正式阅卷细则把创造性转化、创新性发展和“发展”放在一起，说明落点应是用发展的观点看传统戏曲怎样在继承精髓的基础上适应时代、回应观众并获得新的表现形态。",
        "new_answer": "①事物是变化发展的，传统戏曲也要在时代实践中不断获得新的生命力。②传统戏曲焕发生命力，必须在继承唱腔、程式、人物精神和审美精髓的基础上，进行创造性转化、创新性发展。③把传统艺术同当代舞台、数字技术、青年审美和现实生活结合起来，既守住文化根脉，又形成新的表达方式，才能使传统戏曲随着时代变化不断发展。",
    },
    {
        "queue_id": "T0174",
        "heading": "36. 2026通州期末 第16题（主观题）",
        "node": "系统观念 / 系统优化",
        "old_answer_excerpt": "都江堰的持续运行依赖系统优化，把工程保护、水资源治理、污染防治和文化传播统筹起来。",
        "evidence_note": "矩阵 M1439：正式细则第16题联系观括注“系统观念/整体部分”。",
        "new_why": "都江堰不是一个孤立工程，而是由分水、排沙、灌溉、生态调节、遗产保护和文化传承共同构成的运行系统。正式细则在联系观下括注系统观念、整体部分，说明答题要把各环节作为相互联系的部分来统筹，而不能只写某一项工程技术或单点维护。",
        "new_answer": "①系统优化要求立足整体，把各个部分、要素和环节统筹起来，使整体功能得到更好发挥。②都江堰的持续运行，需要把工程修缮保护、水资源调配、污染防治、生态涵养、文化遗产保护和传播利用统一考虑。③只有让各部分相互配合、协同发力，才能既维护水利工程功能，又延续治水智慧和文化价值，实现人与自然和谐共生。",
    },
    {
        "queue_id": "T0104",
        "heading": "37. 2026朝阳期末 第21题（主观题）",
        "node": "联系的普遍性 / 联系的观点（总）",
        "old_answer_excerpt": "认识四大优势要坚持联系观点，看到四种优势之间相互依存、相互作用、协同赋能，共同为全面建成社会主义现代化强国提供支撑。",
        "evidence_note": "矩阵 M1291：正式阅卷细则第152-153行列“联系的观点看问题”；题面第469-470行写相互依存、协同赋能、有机整体。",
        "new_why": "设问要求认识四大优势，不能把制度、市场、产业体系和人才资源写成并列清单。题面明确使用相互依存、协同赋能、有机整体等表述，正式阅卷细则也列出联系的观点看问题。看到这些提示，就应把四种优势之间的相互支撑关系写出来。",
        "new_answer": "①联系具有普遍性，要用联系的观点看问题，把四大优势作为相互依存、相互作用的有机整体来理解。②党的领导和社会主义制度优势能够为市场、产业和人才发展提供方向和制度保障，超大规模市场、完整产业体系和丰富人才资源又为制度优势转化为治理效能与发展效能提供现实基础。③四种优势协同赋能，才能共同支撑全面建成社会主义现代化强国。",
    },
    {
        "queue_id": "T0201",
        "heading": "37. 2026海淀期末 第17题（主观题）",
        "node": "发展的观点 / 发展的普遍性",
        "old_answer_excerpt": "红色文化要在时代发展中创新表达，通过数字技术、沉浸体验和青年化叙事焕发生机。",
        "evidence_note": "矩阵 M1357：正式评分标准 page_002 Q17“红色文化的发展”维度列出“发展的观点/联系的观点”。",
        "new_why": "设问围绕红色文化如何吸引青年，材料中的数字展陈、沉浸体验、短视频或青年化表达，说明红色文化不能停留在静态陈列中。正式评分标准在“红色文化的发展”维度列出发展的观点，因此答案要写清精神内核如何在新时代实践和传播方式变化中获得新表达。",
        "new_answer": "①事物是变化发展的，红色文化也要在新时代实践中不断创新表达方式。②红色文化的精神内核不能丢，但传播形式、叙事语言和体验方式要随着青年生活方式、数字技术和社会实践变化而发展。③通过数字展馆、沉浸式体验、青年化叙事和现实奋斗场景，把红色资源转化为青年可感、可知、可参与的文化体验，才能增强理解认同并焕发生机。",
    },
    {
        "queue_id": "T0373",
        "heading": "37. 2026西城期末 第21题（主观题）",
        "node": "实践与认识（总）",
        "old_answer_excerpt": "五年规划在实践中形成、实施并根据新情况优化，体现了实践与认识的互动统一，使规划不断贴合国家发展实际。",
        "evidence_note": "矩阵 M1391：正式评分细则 page_005 Q21角度3列出“实践与认识”。",
        "new_why": "五年规划不是凭空提出的设想，而是在国家发展实践、调查研究、人民意见和形势判断中形成，并在实施过程中接受实践检验。正式评分细则把实践与认识列为角度，说明要写出规划从实践中来、到实践中去，并随着新情况新问题不断完善的过程。",
        "new_answer": "①实践是认识的基础，认识又反作用于实践。②科学制定五年规划，要从经济社会发展实践、人民群众实践经验和现实问题出发，通过调查研究、开门问策和科学论证形成对发展目标与路径的认识。③规划实施后还要接受实践检验，并根据国内外形势、发展阶段和新问题不断调整完善，使认识在实践中深化，再以更科学的规划指导现代化建设。",
    },
    {
        "queue_id": "T0481",
        "heading": "37. 2026通州期末 第21题（主观题）",
        "node": "人民群众",
        "old_answer_excerpt": "中国式现代化建设要坚持人民主体地位，使发展成果更多更公平惠及人民并激发人民参与。",
        "evidence_note": "矩阵 M1446：正式细则第21题写明“坚守人民立场”“以人民为中心，增进民生福祉”。",
        "new_why": "设问讨论中国式现代化建设的推进，材料和细则都把坚守人民立场、以人民为中心、增进民生福祉作为关键落点。看到这些表述，就不能只写宏观发展成就，而要写人民群众既是现代化建设主体，也是发展成果的享有者。",
        "new_answer": "①人民群众是社会历史的主体，是推动中国式现代化的根本力量。②推进中国式现代化必须坚守人民立场，坚持以人民为中心，把增进民生福祉、满足人民对美好生活的需要作为重要目标。③让发展成果更多更公平惠及人民，同时尊重人民首创精神、激发人民参与热情，才能把人民群众的积极性创造性转化为现代化建设的持久动力。",
    },
    {
        "queue_id": "T0541",
        "heading": "39. 2026东城期末 第16题（主观题）",
        "node": "价值判断与价值选择",
        "old_answer_excerpt": "分析“清白”时要落到正确价值判断与价值选择。清正廉洁、刚正不阿体现了符合社会发展要求和人民根本利益的价值取向；把“白”理解为人格操守和家国责任，能引导人们在实际生活中作出正确价值选择。",
        "evidence_note": "矩阵 M1170：细则第156行列“价值判断和价值选择”；第174行以“清白”道德准则和责任担当作示例。",
        "new_why": "题目借“清白”讨论做人做事的价值尺度，不是单纯解释颜色或词义。细则明确列出价值判断和价值选择，并把清正廉洁、刚正不阿、道德准则和责任担当作为示例。这说明答案要把“清白”落实为符合社会发展要求和人民利益的价值标准，再说明它怎样引导现实行动。",
        "new_answer": "①价值判断和价值选择具有社会历史性，也必须自觉遵循社会发展规律并站在最广大人民立场上。②“清白”所体现的清正廉洁、刚正不阿、守住底线和勇担责任，是符合社会发展要求和人民根本利益的正确价值取向。③把“清白”作为做人做事的价值尺度，能够引导人们在学习、工作和公共生活中抵制私利诱惑，坚持事实、公道和责任，作出正确价值选择。",
    },
    {
        "queue_id": "T0284",
        "heading": "39. 2026房山一模 第20题（主观题）",
        "node": "矛盾就是对立统一",
        "old_answer_excerpt": "卷面可写：矛盾双方既有区别又有联系。依法治国和依规治党各有侧重，又统一于坚持党的领导、人民当家作主和依法治国有机统一的实践中，必须协同推进。",
        "evidence_note": "矩阵 M0569：房山一模评标/细则 paras69-81 明示矛盾、联系/对立统一、认识与实践、系统优化等。",
        "new_why": "设问讨论依法治国和依规治党的关系，材料并不是要求把二者简单相加。评标/细则明示矛盾、联系/对立统一等哲学角度，并在高分层级要求写出二者既有区别又相互联系、统一于中国式现代化建设和国家治理实践中。因此要用对立统一说明协同推进的必要性。",
        "new_answer": "①矛盾就是对立统一，矛盾双方既相互区别、各有特点，又相互联系、相互依存。②依法治国侧重以宪法法律规范国家和社会治理，依规治党侧重用党内法规管党治党、保持党的先进性和纯洁性，二者作用对象和规范体系不同。③但二者又统一于坚持党的领导、人民当家作主和依法治国有机统一的治理实践中，必须协同推进，才能把党的领导优势、法治优势和制度优势转化为现代化治理效能。",
    },
]


def now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M +08")


def write_csv(path: Path, rows: list[dict[str, object]], fields: list[str]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fields})


def write_draft_files() -> None:
    rows = []
    for target in TARGETS:
        rows.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "evidence_note": target["evidence_note"],
                "new_why_chars": len(target["new_why"]),
                "new_answer_chars": len(target["new_answer"]),
                "new_why": target["new_why"],
                "new_answer": target["new_answer"],
            }
        )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_20260525.json").write_text(
        json.dumps({"updated": now(), "targets": rows}, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_20260525.csv",
        rows,
        ["queue_id", "heading", "node", "evidence_note", "new_why_chars", "new_answer_chars", "new_why", "new_answer"],
    )
    lines = [
        f"# P0 Thickness Repair Batch{BATCH} Draft 20260525",
        "",
        f"Updated: {now()}",
        "",
        "Status: `DRAFT_READY_FOR_DOCX_APPLICATION`",
        "",
        "- Scope: next 8 P0 subjective triple-thin rows after Batch11 refresh.",
        "- Repair method: thicken only existing why/answer paragraphs inside existing evidence-supported nodes.",
        "- Evidence boundary: all evidence notes cite current matrix rows; ordinary reference answers are not upgraded into scoring rubrics.",
        "",
        "| queue_id | node | heading | new why chars | new answer chars | evidence note |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in rows:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | {row['new_why_chars']} | {row['new_answer_chars']} | {row['evidence_note']} |"
        )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_DRAFT_20260525.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )


def main() -> int:
    write_draft_files()
    docx = helper.current_docx()
    backup = docx.with_name(f"{docx.stem}_backup_before_p0_thickness_batch{BATCH}_{datetime.now():%Y%m%d_%H%M%S}.docx")
    shutil.copy2(docx, backup)

    doc = Document(str(docx))
    blocks = helper.entry_blocks(doc)
    applied: list[dict[str, object]] = []
    for target in TARGETS:
        block = helper.matching_block(blocks, target)
        fields = block["fields"]
        field_paras = block["field_paras"]
        before_why = str(fields.get("why", ""))
        before_answer = str(fields.get("answer", ""))
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["why"])], "why", target["new_why"])
        helper.set_labeled_paragraph(doc.paragraphs[int(field_paras["answer"])], "answer", target["new_answer"])
        applied.append(
            {
                "queue_id": target["queue_id"],
                "heading": target["heading"],
                "node": target["node"],
                "heading_para": block["heading_para"],
                "why_para": field_paras["why"],
                "answer_para": field_paras["answer"],
                "before_why_chars": len(before_why),
                "after_why_chars": len(target["new_why"]),
                "before_answer_chars": len(before_answer),
                "after_answer_chars": len(target["new_answer"]),
                "evidence_note": target["evidence_note"],
                "old_answer_excerpt": target["old_answer_excerpt"],
            }
        )

    doc.save(str(docx))
    with zipfile.ZipFile(docx) as z:
        zip_ok = "[Content_Types].xml" in z.namelist() and "word/document.xml" in z.namelist()

    result = {
        "updated": now(),
        "status": f"P0_BATCH{BATCH}_APPLIED_REQUIRES_RENDER_AND_EXTERNAL_RECHECK",
        "docx": str(docx),
        "backup_docx": str(backup),
        "zip_ok": zip_ok,
        "targets": len(TARGETS),
        "applied": applied,
        "boundary": [
            "This is a local thickness repair batch, not final acceptance.",
            "PDF/render outputs must be regenerated after this DOCX edit.",
            "External review remains pending for the current post-repair artifact.",
        ],
    }
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_20260525.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    write_csv(
        ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_20260525.csv",
        applied,
        [
            "queue_id",
            "heading",
            "node",
            "heading_para",
            "why_para",
            "answer_para",
            "before_why_chars",
            "after_why_chars",
            "before_answer_chars",
            "after_answer_chars",
            "evidence_note",
            "old_answer_excerpt",
        ],
    )
    lines = [
        f"# P0 Thickness Repair Batch{BATCH} Apply 20260525",
        "",
        f"Updated: {result['updated']}",
        "",
        f"Status: `{result['status']}`",
        "",
        f"- DOCX: `{docx.name}`.",
        f"- Backup: `{backup.name}`.",
        f"- Targets applied: `{len(applied)}` / `{len(TARGETS)}`.",
        f"- DOCX zip structural check: `{str(zip_ok).lower()}`.",
        "",
        "## Applied Rows",
        "",
        "| queue_id | node | heading | why chars | answer chars | paragraph refs |",
        "|---|---|---|---:|---:|---|",
    ]
    for row in applied:
        lines.append(
            f"| {row['queue_id']} | {row['node']} | {row['heading']} | "
            f"{row['before_why_chars']} -> {row['after_why_chars']} | "
            f"{row['before_answer_chars']} -> {row['after_answer_chars']} | "
            f"h={row['heading_para']}; why={row['why_para']}; answer={row['answer_para']} |"
        )
    lines.extend(
        [
            "",
            "## Boundary",
            "",
            "- This batch repairs only 8 additional P0 rows.",
            "- Remaining thickness queue, render QA, and model gates stay open until rechecked.",
        ]
    )
    (ROOT / f"P0_THICKNESS_REPAIR_BATCH{BATCH}_APPLY_20260525.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
