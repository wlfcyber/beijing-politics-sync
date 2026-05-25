from __future__ import annotations

import csv
import json
import re
import shutil
import tempfile
import zipfile
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from lxml import etree


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_NS = "http://www.w3.org/XML/1998/namespace"
W = f"{{{W_NS}}}"
XML = f"{{{XML_NS}}}"
NS = {"w": W_NS}

ROOT = Path(__file__).resolve().parents[2]
RECOVERY = ROOT / "reports" / "bixiu4_thread_recovery_opus47_2026-05-24"
RUN = ROOT / "reports" / "bixiu4_baodian_52_base_insert_second_mock_first_mock_audit_2026-05-24"
DELIVERY = RUN / "05_delivery"
ACCEPTED = RUN / "04_fusion_audit" / "student_patch_entries.accepted.jsonl"
LEDGER = DELIVERY / "docx_insert_ledger.csv"
MATRIX = RECOVERY / "FULL_QUESTION_COVERAGE_AND_PLACEMENT_MATRIX.csv"
BATCH_REPORT = RECOVERY / "COVERAGE_FUSION_BATCH08_2025_DONGCHENG_YIMO_CODEX_20260525.md"
CARTOON = RECOVERY / "_source_image_extracts" / "2025_dongcheng_yimo_q5_cartoon.png"
SUITE = "2025东城一模"

QUESTION_1 = (
    "2024年是中华人民共和国成立75周年。75年披荆斩棘、砥砺前进，我们党团结和带领人民绘就了一幅波澜壮阔、气势恢弘的历史画卷，"
    "谱写了一曲感天动地、气壮山河的奋斗赞歌。总结历史经验，为我们新征程上带来的深刻启迪有："
    "①新时代必须坚持和加强中国共产党的全面领导；②人民群众是决定党和国家前途命运的根本力量；"
    "③中国梦是实现中华民族伟大复兴的必由之路；④改革开放是基本国策，只有进行时没有完成时。"
    "A.①② B.①③ C.②④ D.③④。官方答案：A。"
)

QUESTION_4 = (
    "东城区以历史文化遗产保护带动城市更新、治理创新和发展焕新。社稷坛、太庙等重点文物腾退完成，“左祖右社”风采再现。"
    "钟鼓楼周边、前门大街等重要节点进行环境整治，壮美中轴尽显古都风韵。“中轴线上”“故宫以东”古建艺术季等活动精彩纷呈，"
    "让更多市民沉浸式感受古都文化魅力。这体现了：①坚持人民至上，使人民共享文化保护与发展成果；"
    "②坚持突出重点，以遗产利用为首位促进城市更新；③把握适度原则，通过恢复重要遗产展现历史文脉；"
    "④建立人为联系，发挥文化资源优势赋能城市发展。A.①② B.①④ C.②③ D.③④。官方答案：B。"
)

QUESTION_5 = (
    "下列古诗文反映的哲理和右侧漫画寓意相同的是："
    "A.人事有代谢，往来成古今；B.近水知鱼性，近山识鸟音；C.蝉噪林逾静，鸟鸣山更幽；D.海日生残夜，江春入旧年。"
    "漫画文字：已知的越多，未知的更多。官方答案：C。"
)

QUESTION_16 = (
    "结合材料，运用《哲学与文化》知识，分析中国年何以“生生不息”。"
)

ENTRIES = [
    {
        "canonical_node": "人民群众",
        "question_no": "Q1",
        "heading_suffix": "2025东城一模 第1题（选择题）",
        "material_trigger": "题干回顾新中国75年奋斗历程，正确项②明确“人民群众是决定党和国家前途命运的根本力量”。",
        "question_prompt": QUESTION_1,
        "why_trigger": "看到“党团结和带领人民绘就历史画卷”“人民群众是决定党和国家前途命运的根本力量”，应落到历史唯物主义的人民群众观点。历史不是少数人单独创造的，国家发展和民族复兴的根本力量来自广大人民群众的实践。",
        "answer_landing": "本题应选A。本节点处理②：总结75年历史经验，必须看到人民群众是社会历史的主体，是推动国家发展和民族复兴的根本力量；新征程上仍要坚持人民主体地位，依靠人民创造新的历史伟业。",
        "evidence_level": "选择题官方答案键+题干正确项链条",
        "source_lines": "01_source_inventory/suite_source_bundles/2025东城一模.md:264-271;528-565",
    },
    {
        "canonical_node": "根据固有联系建立新的具体联系",
        "question_no": "Q4",
        "heading_suffix": "2025东城一模 第4题（选择题）",
        "material_trigger": "东城区把历史文化遗产保护同城市更新、治理创新、发展焕新结合起来；正确项④明确“建立人为联系，发挥文化资源优势赋能城市发展”。",
        "question_prompt": QUESTION_4,
        "why_trigger": "看到“以历史文化遗产保护带动城市更新”“文化资源优势赋能城市发展”，应想到人可以根据事物固有联系建立新的具体联系。这里不是任意拼接文化和城市，而是在尊重古都历史文脉、空间资源和市民文化需求的基础上，把文物保护、环境治理、文旅活动和城市发展联系起来。",
        "answer_landing": "本题应选B。本节点处理④：人们可以根据事物固有联系，改变事物状态、调整原有联系，建立新的具体联系。东城区以中轴线、故宫以东等文化资源为基础，把遗产保护、城市更新和市民文化体验联系起来，使文化资源转化为城市发展的现实优势。",
        "evidence_level": "选择题官方答案键+题干正确项链条",
        "source_lines": "01_source_inventory/suite_source_bundles/2025东城一模.md:287-298;528-565",
    },
    {
        "canonical_node": "矛盾就是对立统一",
        "question_no": "Q5",
        "heading_suffix": "2025东城一模 第5题（选择题）",
        "material_trigger": "漫画文字为“已知的越多，未知的更多”；正确项C“蝉噪林逾静，鸟鸣山更幽”用动与静、鸣与幽的相互映衬说明对立面相互依存、相反相成。",
        "question_prompt": QUESTION_5,
        "why_trigger": "漫画把“已知”和“未知”放在同一知识探索过程中：知道得越多，越会发现未知更多。正确诗句也把“蝉噪”和“林静”、“鸟鸣”和“山幽”放在一起，说明矛盾双方既相互区别又相互依存、相互映衬。",
        "answer_landing": "本题应选C。已知与未知、噪与静、鸣与幽不是彼此孤立的两个方面，而是在同一情境中相互规定、相互映衬；这体现矛盾双方的对立统一。学习中认识越深入，越能发现新的未知，也正是在这种已知与未知的矛盾运动中推动认识继续发展。",
        "evidence_level": "选择题官方答案键+题干正确项链条+源页漫画图",
        "source_lines": "01_source_inventory/suite_source_bundles/2025东城一模.md:299-303;528-565;_source_image_extracts/2025_dongcheng_yimo_q5_cartoon.png",
        "cartoon": str(CARTOON),
    },
    {
        "canonical_node": "发展的观点 / 发展的普遍性",
        "question_no": "Q16",
        "heading_suffix": "2025东城一模 第16题（主观题）",
        "material_trigger": "“现代的中国年”中，机器人小狮子与醒狮共舞，现代科技与传统习俗完美结合；运动过大年成为时尚。阅卷报告在“文化创新”角度下明确可用哲学知识“发展”。",
        "question_prompt": QUESTION_16,
        "why_trigger": "看到春节从写春联、贴年画、守岁、拜年、观灯，发展到机器人醒狮、滑雪过年、联合国非遗和世界共鸣，说明中国年不是静止保存旧形式，而是在继承中不断加入新的时代内容和新的庆祝方式。",
        "answer_landing": "答案落点：事物是变化发展的，要用发展的观点看问题。中国年之所以生生不息，是因为它在保持团圆、和睦、祈福等传统内核的同时，不断吸收现代科技、运动休闲、国际传播等新形式，实现内涵和庆祝方式的与时俱进。",
        "evidence_level": "正式阅卷报告评分点+哲学可用知识明示",
        "source_lines": "01_source_inventory/suite_source_bundles/2025东城一模.md:24-30;390-402;569-573",
    },
]

SECTION_NEXT = {
    "人民群众": "价值观的导向作用",
    "根据固有联系建立新的具体联系": "联系的多样性",
    "矛盾就是对立统一": "矛盾的普遍性",
    "发展的观点 / 发展的普遍性": "量变与质变 / 适度原则",
}

LABELS = [
    ("【材料触发点】", "material_trigger"),
    ("【设问】", "question_prompt"),
    ("【为什么能想到】", "why_trigger"),
    ("【答案落点】", "answer_landing"),
]

KEYS = {
    "id": "matrix_row_id",
    "source": "题源",
    "year": "年份",
    "stage": "阶段",
    "question": "题号",
    "type": "题型或模块判断",
    "in_body": "是否进宝典",
    "node": "宝典节点",
    "principle": "细则支持原理",
    "evidence": "证据等级",
    "misplaced": "是否误放",
    "needs": "是否需补厚",
    "action": "当前处理",
    "note": "备注",
    "artifact": "source_artifact",
}


def current_docx() -> Path:
    docs = [p for p in DELIVERY.glob("*.docx") if "_backup_" not in p.stem and not p.name.startswith("~$")]
    if len(docs) != 1:
        raise RuntimeError(f"Expected one current DOCX, found {docs}")
    return docs[0]


def para_text(p: etree._Element) -> str:
    return "".join(t.text or "" for t in p.xpath(".//w:t", namespaces=NS))


def set_plain_text(p: etree._Element, text: str) -> None:
    texts = p.xpath(".//w:t", namespaces=NS)
    if not texts:
        r = etree.SubElement(p, W + "r")
        t = etree.SubElement(r, W + "t")
        t.text = text
        return
    texts[0].text = text
    for node in texts[1:]:
        node.text = ""


def make_run(text: str, *, label: bool = False) -> etree._Element:
    r = etree.Element(W + "r")
    if label:
        rpr = etree.SubElement(r, W + "rPr")
        etree.SubElement(rpr, W + "b")
        color = etree.SubElement(rpr, W + "color")
        color.set(W + "val", "21574C")
    t = etree.SubElement(r, W + "t")
    if text.startswith(" ") or text.endswith(" "):
        t.set(XML + "space", "preserve")
    t.text = text
    return r


def set_label_text(p: etree._Element, label: str, rest: str) -> None:
    for child in list(p):
        if child.tag == W + "r":
            p.remove(child)
    p.append(make_run(label, label=True))
    p.append(make_run(" " + rest))
    ppr = p.find("./w:pPr", namespaces=NS)
    if ppr is None:
        ppr = etree.Element(W + "pPr", nsmap=p.nsmap)
        p.insert(0, ppr)
    spacing = ppr.find("./w:spacing", namespaces=NS)
    if spacing is None:
        spacing = etree.SubElement(ppr, W + "spacing")
    spacing.set(W + "after", "80")


def find_section(paras: list[etree._Element], heading: str) -> tuple[int, int]:
    next_heading = SECTION_NEXT[heading]
    start = next((i for i, p in enumerate(paras) if para_text(p).strip() == heading), None)
    if start is None:
        raise RuntimeError(f"section heading not found: {heading}")
    end = next((i for i in range(start + 1, len(paras)) if para_text(paras[i]).strip() == next_heading), None)
    if end is None:
        raise RuntimeError(f"next heading not found after {heading}: {next_heading}")
    return start, end


def insert_entry(root: etree._Element, entry: dict[str, str]) -> str:
    body = root.find("w:body", namespaces=NS)
    if body is None:
        raise RuntimeError("word/document.xml has no body")
    paras = [p for p in body if p.tag == W + "p"]
    start, end = find_section(paras, entry["canonical_node"])
    existing = next((para_text(paras[i]).strip() for i in range(start + 1, end) if para_text(paras[i]).strip().endswith(entry["heading_suffix"])), None)
    if existing:
        return existing

    numbered = []
    for idx in range(start + 1, end):
        text = para_text(paras[idx]).strip()
        m = re.match(r"^(\d+)\.\s", text)
        if m:
            numbered.append((idx, int(m.group(1))))
    if not numbered:
        raise RuntimeError(f"no numbered item under {entry['canonical_node']}")
    last_idx, last_no = numbered[-1]
    block_end = end
    for idx in range(last_idx + 1, end):
        if re.match(r"^\d+\.\s", para_text(paras[idx]).strip()):
            block_end = idx
            break
    template = [deepcopy(p) for p in paras[last_idx:block_end]]
    if len(template) < 5:
        raise RuntimeError(f"template too short under {entry['canonical_node']}: {len(template)}")

    heading = f"{last_no + 1}. {entry['heading_suffix']}"
    set_plain_text(template[0], heading)
    for p, (label, key) in zip(template[1:5], LABELS):
        set_label_text(p, label, entry[key])
    if len(template) > 5:
        set_plain_text(template[5], "")
    ref = paras[end]
    for p in template:
        body.insert(body.index(ref), p)
    return heading


def add_q5_cartoon(docx: Path) -> None:
    if not CARTOON.exists():
        raise RuntimeError(f"Q5 cartoon image is missing: {CARTOON}")
    doc = Document(str(docx))
    target_idx = None
    for i, para in enumerate(doc.paragraphs):
        if SUITE in para.text and "第5题" in para.text:
            target_idx = i
            break
    if target_idx is None:
        raise RuntimeError("Q5 heading not found after DOCX insertion")
    window = doc.paragraphs[target_idx + 1 : target_idx + 7]
    if any("w:drawing" in p._p.xml for p in window):
        return
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_p.add_run()
    run.add_picture(str(CARTOON), width=Inches(2.65))
    doc.paragraphs[target_idx + 1]._p.addnext(pic_p._p)
    doc.save(str(docx))


def update_docx(timestamp: str) -> list[str]:
    docx = current_docx()
    backup = docx.with_name(f"{docx.stem}_backup_before_2025_dongcheng_yimo_batch08_{timestamp}.docx")
    shutil.copy2(docx, backup)
    with zipfile.ZipFile(docx, "r") as zin:
        infos = zin.infolist()
        data = {info.filename: zin.read(info.filename) for info in infos}
    root = etree.fromstring(data["word/document.xml"])
    headings = [insert_entry(root, entry) for entry in ENTRIES]
    data["word/document.xml"] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)
    try:
        with zipfile.ZipFile(tmp_path, "w") as zout:
            for info in infos:
                payload = data[info.filename]
                zi = zipfile.ZipInfo(info.filename, date_time=info.date_time)
                zi.compress_type = info.compress_type
                zi.external_attr = info.external_attr
                zout.writestr(zi, payload)
        shutil.move(tmp_path, docx)
    finally:
        if tmp_path.exists():
            tmp_path.unlink()
    add_q5_cartoon(docx)
    return headings


def update_ledger(timestamp: str, headings: list[str]) -> None:
    backup = LEDGER.with_name(f"{LEDGER.stem}_backup_before_batch08_dongcheng_yimo_{timestamp}{LEDGER.suffix}")
    shutil.copy2(LEDGER, backup)
    with LEDGER.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        rows = list(reader)
        fieldnames = reader.fieldnames or []
    desired = {
        (entry["canonical_node"], SUITE, entry["question_no"]): {
            "canonical_node": entry["canonical_node"],
            "source_suite": SUITE,
            "question_no": entry["question_no"],
            "inserted_heading": heading,
        }
        for entry, heading in zip(ENTRIES, headings)
    }
    cleaned: list[dict[str, str]] = []
    seen = set()
    for row in rows:
        key = (row["canonical_node"], row["source_suite"], row["question_no"])
        if key in desired:
            if key not in seen:
                cleaned.append(desired[key])
                seen.add(key)
            continue
        cleaned.append(row)
    rows = cleaned
    for key, row in desired.items():
        if key not in seen:
            rows.append(row)
            seen.add(key)
    with LEDGER.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def accepted_record(entry: dict[str, str], heading: str) -> dict[str, str]:
    record = {
        "source_suite": SUITE,
        "question_no": entry["question_no"],
        "canonical_node": entry["canonical_node"],
        "student_facing_text": (
            f"{heading}\n"
            f"【材料触发点】 {entry['material_trigger']}\n"
            f"【设问】 {entry['question_prompt']}\n"
            f"【为什么能想到】 {entry['why_trigger']}\n"
            f"【答案落点】 {entry['answer_landing']}"
        ),
        "evidence_level": entry["evidence_level"],
        "source_lines": entry["source_lines"],
        "batch": "batch08_2025_dongcheng_yimo",
    }
    if "cartoon" in entry:
        record["cartoon_image"] = entry["cartoon"]
    return record


def update_accepted(timestamp: str, headings: list[str]) -> None:
    backup = ACCEPTED.with_name(f"{ACCEPTED.stem}_backup_before_batch08_dongcheng_yimo_{timestamp}{ACCEPTED.suffix}")
    shutil.copy2(ACCEPTED, backup)
    records = [json.loads(line) for line in ACCEPTED.read_text(encoding="utf-8").splitlines() if line.strip()]
    desired = {
        (SUITE, entry["question_no"], entry["canonical_node"]): accepted_record(entry, heading)
        for entry, heading in zip(ENTRIES, headings)
    }
    cleaned: list[dict[str, str]] = []
    seen = set()
    for record in records:
        key = (record.get("source_suite"), record.get("question_no"), record.get("canonical_node"))
        if key in desired:
            if key not in seen:
                cleaned.append(desired[key])
                seen.add(key)
            continue
        cleaned.append(record)
    records = cleaned
    for key, record in desired.items():
        if key not in seen:
            records.append(record)
            seen.add(key)
    ACCEPTED.write_text("\n".join(json.dumps(r, ensure_ascii=False) for r in records) + "\n", encoding="utf-8")


def patch_row(row: dict[str, str], **values: str) -> None:
    for key, value in values.items():
        row[KEYS[key]] = value


def row_values(row_id: str, q: str, type_: str, in_body: str, node: str, principle: str, evidence: str, misplaced: str, needs: str, action: str, note: str, artifact: str) -> dict[str, str]:
    return {
        KEYS["id"]: row_id,
        "row_source": "batch08_2025_dongcheng_yimo_recovery",
        KEYS["source"]: SUITE,
        KEYS["year"]: "2025",
        KEYS["stage"]: "一模",
        KEYS["question"]: q,
        KEYS["type"]: type_,
        KEYS["in_body"]: in_body,
        KEYS["node"]: node,
        KEYS["principle"]: principle,
        KEYS["evidence"]: evidence,
        KEYS["misplaced"]: misplaced,
        KEYS["needs"]: needs,
        KEYS["action"]: action,
        KEYS["note"]: note,
        KEYS["artifact"]: artifact,
    }


def update_matrix(timestamp: str, headings: list[str]) -> None:
    backup = MATRIX.with_name(f"{MATRIX.stem}_backup_before_batch08_2025_dongcheng_yimo_{timestamp}{MATRIX.suffix}")
    shutil.copy2(MATRIX, backup)
    with MATRIX.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        rows = list(reader)
        fieldnames = reader.fieldnames or []

    by_id = {r[KEYS["id"]]: r for r in rows}
    common = "01_source_inventory/suite_source_bundles/2025东城一模.md"
    q1_heading, q4_heading, q5_heading, q16_dev_heading = headings
    decisions = {
        "M0081": dict(in_body="是：已由当前DOCX覆盖", node="价值判断与价值选择；价值观的导向作用", principle="16题阅卷报告明确价值观导向作用、价值判断和价值选择可用于价值/文化价值角度；当前DOCX已有价值观导向和价值判断两处Q16条目。", evidence="正式阅卷报告-哲学可用知识明示+现有DOCX正文覆盖", misplaced="否", needs="否", action="COVERED_BY_EXISTING_DOCX_BATCH08", note="保留旧条目并校正为两个价值观节点覆盖。", artifact=f"{common}:21-30;390-402;569-573"),
        "M0082": dict(in_body="是：已由当前DOCX覆盖", node="辩证否定 / 守正创新", principle="16题阅卷报告在文化创新角度明确哲学可用知识包括辩证否定观。", evidence="正式阅卷报告-哲学可用知识明示+现有DOCX正文覆盖", misplaced="否", needs="否", action="COVERED_BY_EXISTING_DOCX_BATCH08", note="现有DOCX有Q16辩证否定条目。", artifact=f"{common}:24-30;390-402;569-573"),
        "M0083": dict(in_body="是：本批新增进入当前DOCX/PDF正文", node="发展的观点 / 发展的普遍性", principle="16题阅卷报告在文化创新角度明确哲学可用知识包括发展；材料写春节形式与时俱进、传统与现代结合。", evidence="正式阅卷报告评分点+哲学可用知识明示", misplaced="否", needs="否", action="INSERTED_BATCH08_2025_DONGCHENG_Q16_DEVELOPMENT", note=f"新增条目：{q16_dev_heading}", artifact=f"{common}:24-30;390-402;569-573"),
        "M0084": dict(in_body="是：已由当前DOCX覆盖", node="矛盾的普遍性和特殊性", principle="16题阅卷报告在共鸣角度明确哲学可用知识包括矛盾普遍性与特殊性的辩证关系。", evidence="正式阅卷报告-哲学可用知识明示+现有DOCX正文覆盖", misplaced="否", needs="否", action="COVERED_BY_EXISTING_DOCX_BATCH08", note="现有DOCX有Q16矛盾普遍性与特殊性条目。", artifact=f"{common}:27-30;390-402;569-573"),
        "M0172": dict(in_body="是：已由当前DOCX覆盖并本批补足发展节点", node="价值观导向；价值判断与价值选择；辩证否定；发展；矛盾普遍性与特殊性；社会存在与社会意识", principle="Q16正式阅卷报告四角度已核对；发展节点本批补入，其余节点现有DOCX覆盖。", evidence="正式阅卷报告-哲学可用知识明示+现有DOCX正文覆盖", misplaced="否", needs="否", action="COVERED_AND_DEVELOPMENT_INSERTED_BATCH08", note="Q16不再挂起。", artifact=f"{common}:16-30;390-402;569-573"),
        "M0173": dict(in_body="是：已由当前DOCX覆盖", node="联系的普遍性/联系观点；发展的观点；主要矛盾；两点论与重点论", principle="18(1)虽以辩证思维设问，但阅卷报告明确整体性可替换为联系观点/系统优化，动态性可替换为发展/量变质变，抓主要矛盾可替换为两点论重点论。", evidence="正式阅卷报告-可替换哲学原理明示+现有DOCX正文覆盖", misplaced="否", needs="否", action="COVERED_BY_EXISTING_DOCX_BATCH08", note="不按纯选必三排除；按正式细则允许的必修四哲学替代点保留。", artifact=f"{common}:76-87;418-424;578-581"),
        "M0203": dict(in_body="是：已由当前DOCX覆盖并本批补足发展节点", node="价值观/辩证否定/发展/矛盾普遍特殊", principle="Q16母版覆盖已核；发展节点本批补入。", evidence="正式阅卷报告-哲学可用知识明示+现有DOCX正文覆盖", misplaced="否", needs="否", action="COVERED_AND_DEVELOPMENT_INSERTED_BATCH08", note="关闭母版核定挂起。", artifact=f"{common}:16-30;390-402;569-573"),
        "M0334": dict(in_body="是：本批新增进入当前DOCX/PDF正文", node="人民群众", principle="Q1正确项②明确人民群众是决定党和国家前途命运的根本力量。", evidence="选择题官方答案键+题干正确项链条", misplaced="否", needs="否", action="INSERTED_BATCH08_2025_DONGCHENG_Q1_PEOPLE", note=f"新增条目：{q1_heading}", artifact=f"{common}:264-271;528-565"),
        "M0335": dict(in_body="否：中特/文化精神边界题", node="不进入当前哲学宝典正文", principle="Q2官方答案B，正确项为行动指南和弘扬长征精神，属于中特/文化精神边界，不拆入哲学主框架。", evidence="官方答案键+模块边界", misplaced="否", needs="否", action="MODULE_BOUNDARY_EXCLUDED_BATCH08", note="若文化精神专题另线处理。", artifact=f"{common}:272-279;528-565"),
        "M0336": dict(in_body="否：政治制度边界题", node="不进入当前哲学宝典正文", principle="Q3官方答案A，澳门一国两制实践证明融入国家发展大局，属于政治制度边界。", evidence="官方答案键+模块边界", misplaced="否", needs="否", action="MODULE_BOUNDARY_EXCLUDED_BATCH08", note="不纳入必修四哲学正文。", artifact=f"{common}:280-286;528-565"),
        "M0337": dict(in_body="是：本批新增进入当前DOCX/PDF正文", node="根据固有联系建立新的具体联系", principle="Q4正确项④明确建立人为联系，发挥文化资源优势赋能城市发展。", evidence="选择题官方答案键+题干正确项链条", misplaced="否", needs="否", action="INSERTED_BATCH08_2025_DONGCHENG_Q4_NEW_CONNECTION", note=f"新增条目：{q4_heading}", artifact=f"{common}:287-298;528-565"),
        "M0338": dict(in_body="是：本批新增进入当前DOCX/PDF正文", node="矛盾就是对立统一", principle="Q5官方答案C，诗句与漫画共同体现对立面相互依存、相反相成；本批嵌入源页漫画图。", evidence="选择题官方答案键+题干正确项链条+源页漫画图", misplaced="否", needs="否", action="INSERTED_BATCH08_2025_DONGCHENG_Q5_CARTOON_CONTRADICTION", note=f"新增条目：{q5_heading}; 漫画图已嵌入。", artifact=f"{common}:299-303;528-565;reports/bixiu4_thread_recovery_opus47_2026-05-24/_source_image_extracts/2025_dongcheng_yimo_q5_cartoon.png"),
        "M0339": dict(in_body="是：已由当前DOCX覆盖", node="认识发展原理", principle="Q6官方答案C，现有DOCX已有Q6选择题条目。", evidence="官方答案键+现有DOCX正文覆盖", misplaced="否", needs="否", action="COVERED_BY_EXISTING_DOCX_BATCH08", note="不新增重复条目。", artifact=f"{common}:304-310;528-565"),
        "M0340": dict(in_body="否：政治与法治边界题", node="不进入当前哲学宝典正文", principle="Q7官方答案A，人民政协反映社情民意信息，属于政治与法治。", evidence="官方答案键+模块边界", misplaced="否", needs="否", action="MODULE_BOUNDARY_EXCLUDED_BATCH08", note="政治模块排除。", artifact=f"{common}:311-317;528-565"),
        "M0341": dict(in_body="否：政治与法治/民族法治边界题", node="不进入当前哲学宝典正文", principle="Q8官方答案C，围绕铸牢中华民族共同体意识的法治保障。", evidence="官方答案键+模块边界", misplaced="否", needs="否", action="MODULE_BOUNDARY_EXCLUDED_BATCH08", note="民族法治边界排除。", artifact=f"{common}:318-324;528-565"),
        "M0342": dict(in_body="否：经济与法治边界题", node="不进入当前哲学宝典正文", principle="Q10官方答案B，增值税法实施体现税收法定和减负政策，属于经济/法治。", evidence="官方答案键+模块边界", misplaced="否", needs="否", action="MODULE_BOUNDARY_EXCLUDED_BATCH08", note="Q9缺失独立行已另补。", artifact=f"{common}:335-345;528-565"),
        "M0343": dict(in_body="否：法律与生活边界题", node="不进入当前哲学宝典正文", principle="Q11官方答案A，AI作品著作权问题属于法律与生活。", evidence="官方答案键+模块边界", misplaced="否", needs="否", action="MODULE_BOUNDARY_EXCLUDED_BATCH08", note="法律模块排除。", artifact=f"{common}:346-352;528-565"),
        "M0344": dict(in_body="否：经济与社会边界题", node="不进入当前哲学宝典正文", principle="Q12官方答案D，民营经济促进法和高质量发展，属于经济与社会。", evidence="官方答案键+模块边界", misplaced="否", needs="否", action="MODULE_BOUNDARY_EXCLUDED_BATCH08", note="经济模块排除。", artifact=f"{common}:353-358;528-565"),
        "M0345": dict(in_body="否：选必三逻辑与思维边界题", node="不进入当前哲学宝典正文", principle="Q13官方答案D，考查必要条件判断，属于逻辑与思维。", evidence="官方答案键+模块边界", misplaced="否", needs="否", action="MODULE_BOUNDARY_EXCLUDED_BATCH08", note="逻辑模块排除。", artifact=f"{common}:359-365;528-565"),
        "M0346": dict(in_body="否：选必三/文化线边界题", node="不进入当前哲学宝典正文", principle="Q15官方答案D，低脂饮食实验考查归纳推理，属于逻辑与思维。", evidence="官方答案键+模块边界", misplaced="否", needs="否", action="MODULE_BOUNDARY_EXCLUDED_BATCH08", note="Q14缺失独立文化边界行已另补；本行对应Q15。", artifact=f"{common}:379-386;528-565"),
        "M0347": dict(in_body="是：已由当前DOCX覆盖并本批补足发展节点", node="价值观导向；价值判断与价值选择；社会存在与社会意识；辩证否定；发展；矛盾普遍性与特殊性", principle="Q16正式阅卷报告哲学可用知识已逐项核定；发展节点本批补入。", evidence="正式阅卷报告-哲学可用知识明示+现有DOCX正文覆盖", misplaced="否", needs="否", action="COVERED_AND_DEVELOPMENT_INSERTED_BATCH08", note="社会存在条目由丰富内涵评分点与参考答案支撑，发展/价值/矛盾由哲学可用知识明示支撑。", artifact=f"{common}:16-30;390-402;569-573"),
        "M0348": dict(in_body="否：法律与生活/法治边界题", node="不进入当前哲学宝典正文", principle="Q17正式阅卷报告明确运用法治知识，评分点为立法、执法、司法。", evidence="正式阅卷报告-模块边界", misplaced="否", needs="否", action="MODULE_BOUNDARY_EXCLUDED_BATCH08", note="法治模块排除。", artifact=f"{common}:50-75;404-416;574-577"),
        "M0349": dict(in_body="是：已由当前DOCX覆盖", node="联系观点/发展观点/主要矛盾/两点论重点论", principle="Q18(1)正式阅卷报告明确联系、系统优化、发展、量变质变、主要矛盾、两点论重点论等可替换哲学原理。", evidence="正式阅卷报告-可替换哲学原理明示+现有DOCX正文覆盖", misplaced="否", needs="否", action="COVERED_BY_EXISTING_DOCX_BATCH08", note="Q18(2)经济逻辑另按经济模块排除。", artifact=f"{common}:76-126;418-429;578-586"),
        "M0350": dict(in_body="否：法律与生活边界题", node="不进入当前哲学宝典正文", principle="Q19正式阅卷报告为诉讼调解、劳动权利义务、遗产管理制度、民法典与核心价值观，属于法律与生活。", evidence="正式阅卷报告-模块边界", misplaced="否", needs="否", action="MODULE_BOUNDARY_EXCLUDED_BATCH08", note="法律模块排除。", artifact=f"{common}:129-184;431-478;587-592"),
        "M0351": dict(in_body="否：当代国际政治与经济边界题", node="不进入当前哲学宝典正文", principle="Q20正式阅卷报告与答案角度为共同利益、高水平开放、全球治理、人类命运共同体，属于选必一。", evidence="正式阅卷报告/答案-模块边界", misplaced="否", needs="否", action="MODULE_BOUNDARY_EXCLUDED_BATCH08", note="国际政治经济边界排除。", artifact=f"{common}:186-210;479-488;593-611"),
        "M0352": dict(in_body="是：已由当前DOCX覆盖", node="社会发展的两大基本规律和基本矛盾", principle="Q21正式阅卷报告明确教育作为上层建筑反作用于经济基础；当前DOCX已有Q21条目。", evidence="正式阅卷报告-角度提示+等级评分+现有DOCX正文覆盖", misplaced="否", needs="否", action="COVERED_BY_EXISTING_DOCX_BATCH08", note="不把等级评分角度升级为逐点强细则；保留上层建筑反作用经济基础落点。", artifact=f"{common}:212-230;490-523;612-626"),
        "M0353": dict(in_body="否：抽取残留/模块边界", node="不进入当前哲学宝典正文", principle="该行由Q20/Q21等级评分或跨模块答案抽取残留形成；没有独立哲学题号。", evidence="抽取残留清理+源包核对", misplaced="否", needs="否", action="EXTRACTION_RESIDUE_CLOSED_BATCH08", note="关闭Qunknown残留。", artifact=f"{common}:593-626"),
        "M0785": dict(in_body="套卷逐题已由Batch08闭合", node="SUITE_LEVEL_SUMMARY", principle="逐题回源已闭合：Q1/Q4/Q5/Q16新增，Q6/Q16/Q18/Q21现有覆盖，其余模块边界排除。", evidence="Batch08逐题矩阵闭合", misplaced="否", needs="否", action="SUITE_CLOSED_BY_BATCH08", note="套卷级记录不替代逐题记录。", artifact=f"{common}:1-626"),
        "M0830": dict(in_body="套卷逐题已由Batch08闭合", node="SUITE_LEVEL_SUMMARY", principle="逐题回源已闭合：Q1/Q4/Q5/Q16新增，Q6/Q16/Q18/Q21现有覆盖，其余模块边界排除。", evidence="Batch08逐题矩阵闭合", misplaced="否", needs="否", action="SUITE_CLOSED_BY_BATCH08", note="套卷级挂起项关闭。", artifact=f"{common}:1-626"),
    }

    for row_id, values in decisions.items():
        if row_id not in by_id:
            raise RuntimeError(f"matrix row not found: {row_id}")
        patch_row(by_id[row_id], **values)

    max_id = max(int(r[KEYS["id"]][1:]) for r in rows if re.match(r"^M\d+$", r[KEYS["id"]]))
    additions = [
        row_values(f"M{max_id + 1:04d}", "Q9", "政治与法治边界题", "否：政治与法治边界题", "不进入当前哲学宝典正文", "Q9官方答案C，行政执法监督典型案例保护行政相对人合法权益，属于法治政府/行政执法。", "官方答案键+模块边界", "否", "否", "MISSING_ROW_ADDED_AND_EXCLUDED_BATCH08", "原矩阵缺独立Q9行，本批补充边界闭合。", f"{common}:325-334;528-565"),
        row_values(f"M{max_id + 2:04d}", "Q14", "必修四文化线边界题", "否：当前哲学宝典不补；若扩展文化宝典需另审", "不进入当前哲学宝典正文", "Q14官方答案B，考查文化越是民族的越是世界的、传统文化创新表达，属于文化线。", "官方答案键+文化模块边界", "否", "否", "MISSING_ROW_ADDED_AND_EXCLUDED_BATCH08", "原矩阵缺独立Q14行，本批补充文化边界闭合。", f"{common}:366-378;528-565"),
    ]
    existing_keys = {(r[KEYS["source"]], r[KEYS["question"]], r[KEYS["action"]]) for r in rows}
    for row in additions:
        key = (row[KEYS["source"]], row[KEYS["question"]], row[KEYS["action"]])
        if key not in existing_keys:
            rows.append(row)
            existing_keys.add(key)

    with MATRIX.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def write_report(timestamp: str, headings: list[str]) -> None:
    text = f"""# Coverage Fusion Batch08 - 2025东城一模

timestamp: {timestamp}
operator: Codex recovery thread
suite: 2025东城一模
status: CODEX_BATCH08_SOURCE_COVERAGE_APPLIED

## Source Basis

- source bundle: `01_source_inventory/suite_source_bundles/2025东城一模.md`
- formal scoring/rubric lines checked: Q16 `16-30`, Q17 `50-75`, Q18(1) `76-87`, Q19 `129-184`, Q20 `186-210`, Q21 `212-230`
- paper/answer lines checked: Q1-Q15 `264-386`, answers `528-565`
- Q5 cartoon image extracted from the source PDF page and saved as `_source_image_extracts/2025_dongcheng_yimo_q5_cartoon.png`

## Inserted Into DOCX

- {headings[0]} -> 人民群众
- {headings[1]} -> 根据固有联系建立新的具体联系
- {headings[2]} -> 矛盾就是对立统一, with source cartoon image embedded
- {headings[3]} -> 发展的观点 / 发展的普遍性

## Question Disposition

- Q1: inserted under 人民群众, supported by official choice key and correct option ②.
- Q2-Q3: excluded as 中国特色/政治制度 boundary.
- Q4: inserted under 根据固有联系建立新的具体联系.
- Q5: inserted under 矛盾就是对立统一 and the source cartoon image was embedded.
- Q6: existing DOCX coverage under 认识发展原理 retained.
- Q7-Q13: excluded by political, legal, economic, or logic/thinking module boundary; Q9 missing row added.
- Q14: missing row added and excluded as culture-line boundary for the current philosophy handbook.
- Q15: excluded as logic/thinking.
- Q16: existing coverage retained; development node added because the formal marking report explicitly lists 发展 as usable philosophy knowledge.
- Q17: excluded by formal 法治 rubric.
- Q18(1): existing DOCX coverage retained because the formal marking report explicitly allows philosophy substitutions including 联系、系统优化、发展、量变质变、主要矛盾、两点论重点论.
- Q18(2): excluded as economics.
- Q19: excluded by formal 法律与生活 rubric.
- Q20: excluded as 当代国际政治与经济.
- Q21: existing DOCX coverage retained for 上层建筑反作用经济基础; evidence remains angle/level-scoring support, not per-point detailed rubric.
- Qunknown: closed as extraction residue.

## Controls

- Sonnet evidence not used.
- GPTPro web / external Claude Opus full-artifact review remains `real_call_pending`.
- ClaudeCode Batch08 Opus 4.7 recheck still must be called after this Codex production patch.
"""
    BATCH_REPORT.write_text(text, encoding="utf-8")


def main() -> None:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    headings = update_docx(timestamp)
    update_ledger(timestamp, headings)
    update_accepted(timestamp, headings)
    update_matrix(timestamp, headings)
    write_report(timestamp, headings)
    print(f"Batch08 applied at {timestamp}")
    for heading in headings:
        print(heading)


if __name__ == "__main__":
    main()
