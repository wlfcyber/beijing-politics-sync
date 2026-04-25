"""Render the philosophy framework markdown into a Word .docx.

Style choices:
- Use Heading 1/2/3 styles for ## / ### / #### headings.
- Convert numbered list items "1." into a numbered paragraph style.
- Convert bullet items "- " into bullet paragraphs.
- Preserve **bold** runs.
- For the framework summary table at the start, render as a small native table.
- For all other "tables" we render as labeled paragraphs (avoids the
  narrow-vertical-table failure noted in thread-defect-patches.md).
"""
from __future__ import annotations

import os
import re
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor

SRC = (
    r"C:\Users\Administrator\Desktop\飞哥的政治庄园\reports"
    r"\philosophy_rerun_claudecode_2026-04-25_2230\philosophy_framework_body.md"
)
OUT = r"C:\Users\Administrator\Desktop\哲学大框架_ClaudeCode重跑版.docx"


def add_run_with_bold(paragraph, text: str, base_size: int = 12,
                      base_font: str = "Microsoft YaHei") -> None:
    """Parse **bold** segments and append runs."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        run.font.size = Pt(base_size)
        run.font.name = base_font
        # Set East-Asian font on the run rPr too
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts"
        )
        if rFonts is None:
            from docx.oxml.ns import qn
            rFonts = rPr.makeelement(qn("w:rFonts"), {})
            rPr.append(rFonts)
        from docx.oxml.ns import qn
        rFonts.set(qn("w:eastAsia"), base_font)


def set_doc_default_font(doc: Document, font_name: str = "Microsoft YaHei",
                         size_pt: int = 12) -> None:
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    from docx.oxml.ns import qn
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def render() -> None:
    with open(SRC, encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()
    set_doc_default_font(doc)

    # Cover
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_run = cover.add_run("必修四《哲学与文化》大框架")
    cover_run.font.size = Pt(28)
    cover_run.font.bold = True
    from docx.oxml.ns import qn
    cover_run._element.get_or_add_rPr().append(
        cover_run._element.makeelement(qn("w:rFonts"), {qn("w:eastAsia"): "Microsoft YaHei"})
    )

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("Claude Code 重跑版 · 2026-04-25")
    sub_run.font.size = Pt(14)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run(
        "适配北京高考 · 来源覆盖 2024、2025、2026 各区一模 / 二模 / 期中 / 期末"
    )
    note_run.font.size = Pt(11)
    note_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph("")

    in_summary_table = False
    pending_table_rows: list[list[str]] = []
    pending_after_summary = False

    def flush_summary_table():
        nonlocal pending_table_rows, in_summary_table
        if not pending_table_rows:
            return
        header = pending_table_rows[0]
        body = pending_table_rows[1:]
        table = doc.add_table(rows=1 + len(body), cols=len(header))
        table.style = "Light Grid Accent 1"
        # set fixed widths to avoid the narrow-vertical-table failure
        col_widths = [Cm(2.5), Cm(4.5), Cm(9.0)]
        for i, w in enumerate(col_widths[:len(header)]):
            for cell in table.columns[i].cells:
                cell.width = w
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(header):
            p = hdr_cells[i].paragraphs[0]
            r = p.add_run(h)
            r.bold = True
            r.font.size = Pt(11)
        for row_idx, row in enumerate(body, start=1):
            cells = table.rows[row_idx].cells
            for i, cell_text in enumerate(row):
                p = cells[i].paragraphs[0]
                p.style = doc.styles["Normal"]
                add_run_with_bold(p, cell_text, base_size=10)
        pending_table_rows = []
        in_summary_table = False

    for raw in lines:
        line = raw.rstrip("\n")
        if not line.strip():
            if in_summary_table:
                flush_summary_table()
            doc.add_paragraph("")
            continue

        # Markdown table rows
        if line.startswith("|"):
            # The framework body uses one summary table (after "## 二、").
            # We capture only the first encountered table; others are
            # rendered as labeled paragraphs to avoid wide-table layout
            # failures.
            if not in_summary_table and pending_after_summary is False:
                # parse table cells
                parts = [c.strip() for c in line.strip("|").split("|")]
                # skip the separator row that looks like ---|---|---
                if all(re.match(r"^[-: ]+$", c) for c in parts):
                    continue
                pending_table_rows.append(parts)
                in_summary_table = True
                continue
            else:
                if in_summary_table:
                    parts = [c.strip() for c in line.strip("|").split("|")]
                    if all(re.match(r"^[-: ]+$", c) for c in parts):
                        continue
                    pending_table_rows.append(parts)
                    continue
                else:
                    # convert table row into a labeled paragraph
                    parts = [c.strip() for c in line.strip("|").split("|")]
                    if all(re.match(r"^[-: ]+$", c) for c in parts):
                        continue
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.5)
                    add_run_with_bold(p, " | ".join(parts), base_size=10.5)
                    continue
        else:
            if in_summary_table:
                flush_summary_table()
                pending_after_summary = True

        # Headings
        if line.startswith("# "):
            h = doc.add_heading(level=0)
            r = h.add_run(line[2:].strip())
            r.font.size = Pt(22)
            continue
        if line.startswith("## "):
            h = doc.add_heading(level=1)
            add_run_with_bold(h, line[3:].strip(), base_size=18)
            continue
        if line.startswith("### "):
            h = doc.add_heading(level=2)
            add_run_with_bold(h, line[4:].strip(), base_size=15)
            continue
        if line.startswith("#### "):
            h = doc.add_heading(level=3)
            add_run_with_bold(h, line[5:].strip(), base_size=13)
            continue
        if line.startswith("##### "):
            h = doc.add_heading(level=4)
            add_run_with_bold(h, line[6:].strip(), base_size=12)
            continue

        # Blockquote
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            r = p.add_run(line[2:].strip())
            r.italic = True
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue

        # Bullet — render as Normal paragraph with explicit "•" prefix, so
        # numbering counters never auto-advance and the layout stays
        # predictable across sections.
        m = re.match(r"^(\s*)- (.*)$", line)
        if m:
            indent_spaces = len(m.group(1))
            level = indent_spaces // 2
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5 + 0.6 * level)
            add_run_with_bold(p, "• " + m.group(2), base_size=11.5)
            continue

        # Numbered — keep the literal number from the source (per-section),
        # rendered as Normal paragraph so we don't get cross-section
        # auto-numbering drift.
        m = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m:
            indent_spaces = len(m.group(1))
            level = indent_spaces // 2
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5 + 0.6 * level)
            add_run_with_bold(
                p, f"{m.group(2)}. " + m.group(3), base_size=11.5
            )
            continue

        # Plain paragraph
        if line.strip().startswith("---"):
            doc.add_paragraph("―" * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        p = doc.add_paragraph()
        add_run_with_bold(p, line, base_size=12)

    if in_summary_table:
        flush_summary_table()

    doc.save(OUT)
    print(f"wrote: {OUT}")
    print(f"  size: {os.path.getsize(OUT)} bytes")


if __name__ == "__main__":
    render()
