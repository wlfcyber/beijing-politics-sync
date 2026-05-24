from __future__ import annotations

import csv
import json
import re
import shutil
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor


DESKTOP = Path.home() / "Desktop"
ROOT = DESKTOP / "02_代码项目与工具" / "mac-thread-restore" / "beijing-politics-sync-visible"
RUN = ROOT / "reports" / "bixiu4_philosophy_all_questions_strict_dual_audit_2026-05-23"
STRICT = ROOT / "reports" / "bixiu4_philosophy_strict_v8_2026-05-23"
BASE_MD = DESKTOP / "5.23哲学宝典_认可版仅插新增卷子v8" / "01_学生版Word" / "必修四哲学材料-知识触发框架_v8_认可版仅插新增卷子.md"
BASE_DOCX = DESKTOP / "5.23哲学宝典_认可版仅插新增卷子v8" / "01_学生版Word" / "必修四哲学材料-知识触发框架_v8_认可版仅插新增卷子.docx"
SUBJ_CSV = RUN / "02_codex_lane" / "unique_subjective_prompt_packets.csv"
CHOICE_CSV = STRICT / "remaining_old_choice_presence_gaps_after_v7.csv"
ROSTER_CSV = STRICT / "current_65_suite_roster.csv"
SECOND_MOCK_JSON = ROOT / "reports" / "bixiu4_philosophy_accepted_base_second_mock_patch_2026-05-23" / "04_fusion" / "fused_entries_2026_second_mock.json"
OUT_DIR = DESKTOP / "5.24二模同流程修订v10"

FORBIDDEN = [
    "补漏",
    "补入",
    "补丁",
    "审计",
    "证据",
    "CSV",
    "评标",
    "评分细则",
    "评分参考",
    "评分",
    "阅卷细则",
    "阅卷总结",
    "讲评",
    "参考答案",
    "可从",
    "给分点",
    "pdf",
    "docx",
    "source_path",
    "gpt_markdown_path",
    "sha256",
    "render_dir",
    "text-extracted",
    "docx-xml",
    "pypdf",
    "C:\\Users",
    "gpt_sources",
    "pptx",
    "OCR",
]

HIGH_RISK = ["主要矛盾", "矛盾的主要方面", "两点论", "重点论", "主流", "支流", "辩证否定", "量变", "质变", "价值观"]


SECOND_MOCK_THICK_CARDS: list[dict[str, str]] = [
    {
        "title": "历史传承与现代创新：先守住根，再让传统进入今天",
        "source": "2026朝阳二模 第16题",
        "question": "结合材料，从哲学角度，说明在城市文化建设中怎样实现历史传承与现代创新的有机结合。",
        "material": "材料把北京中轴线、平遥古城等历史文化古迹的保护，同四合院更新、市民文化中心、线上文化产品、群众点单和政府买单等现代公共文化供给放在一起，设问直接问“历史传承”和“现代创新”怎样有机结合。",
        "principles": "对立统一和联系观点、辩证否定和守正创新、具体问题具体分析、价值观导向。",
        "logic": "这道题不能只写“保护传统”，也不能只写“创新形式”。传承和创新是一对关系：传承解决城市文化从哪里来，创新解决城市文化怎样活在今天。材料中的古迹保护对应“守住城市文脉”，四合院焕新和公共文化服务对应“让传统空间承担现代生活功能”，群众需求和核心价值观对应“创新不能跑偏，要服务人民精神生活和城市认同”。",
        "landing": "答案可以落成四层：第一，承认传承与创新相互依存、相互促进，城市文化建设要把保护文脉和激活生活统一起来。第二，坚持辩证否定，在保留优秀历史文化根脉的基础上推动传统空间、传统资源创造性转化。第三，从不同城市、不同空间、不同群众需求出发，具体设计公共文化服务和文化产品。第四，用社会主义核心价值观引领城市文化软实力建设，让文化空间凝聚认同感、归属感，而不是只做景观改造。",
        "reminder": "课堂讲这题时，要把“传承”和“创新”当成一条链讲：保护根脉是前提，现代转化是路径，群众需求是落点，价值引领是方向。",
    },
    {
        "title": "系统思维与战略定力：四个中国不是四件散活，而是一张蓝图",
        "source": "2026朝阳二模 第21题",
        "question": "“十五五”规划纲要将“四个中国”确立为推进中国式现代化的战略支柱，深刻体现了中国共产党治国理政的系统思维和一张蓝图绘到底的战略定力。结合材料，综合运用所学，谈谈对系统思维和战略定力的认识。",
        "material": "材料把数字中国、健康中国、平安中国、美丽中国共同作为中国式现代化战略支柱，又强调规划接续、方向稳定、长期推进。",
        "principles": "系统观念、联系观点、整体与部分、量变质变、坚持党的领导和长期主义。",
        "logic": "“四个中国”不能拆成四句口号。系统思维看的是整体性、关联性和协同性：数字中国能赋能健康、生态和治理，平安中国提供稳定环境，美丽中国和健康中国又共同指向人民生活质量。战略定力看的是方向不变、目标执着、行动持久：不是今天换一个主题、明天换一个工程，而是把长期目标分解到阶段规划中持续落实。",
        "landing": "答案可以先写系统思维：把“四个中国”作为中国式现代化的有机整体统筹推进，避免单兵突进和顾此失彼，使各战略支柱相互支撑、协同发力。再写战略定力：在党的领导下坚持既定方向，把宏伟蓝图转化为阶段任务，通过久久为功的量的积累推动现代化目标逐步实现。最后点二者统一：系统思维解决“怎么布局”，战略定力解决“怎么坚持到底”，共同保证中国式现代化行稳致远。",
        "reminder": "不要把这题写成一般“联系观”。关键词是“整体布局”和“持续落地”，一定要把四个中国之间的互相支撑说出来。",
    },
    {
        "title": "京彩课堂：实践场景让思政课从书本走向真实理解",
        "source": "2026东城二模 第16题",
        "question": "结合材料，运用《哲学与文化》知识，谈谈“京彩课堂”的精彩所在。",
        "material": "材料写北京学校打破传统课堂局限，依托中轴线、故宫、科创基地、临空经济等真实场景，构建实景、任务、探究结合的育人模式，并借助数字化方式创新课堂。",
        "principles": "实践是认识的基础、守正创新、理论与实践统一、文化创新和文化自信。",
        "logic": "“精彩”不是因为活动热闹，而是因为它改变了思政课的理解方式。学生不是只在教室里背概念，而是在真实文化场景和社会实践中理解理论、检验认识、形成认同。它也不是抛开思政育人目标去玩形式，而是在守住铸魂育人方向的基础上，用探究式、项目式、数字化方式激活课堂生命力。",
        "landing": "答案可以这样落：从哲学上看，实践是认识的来源、动力和检验标准，京彩课堂把抽象理论放入真实场景，让学生在看、做、问、研中深化认识，并把认识转化为文化自信和理想信念。从文化上看，京彩课堂坚持守正创新，既坚守思政育人和中华优秀传统文化根基，又创新内容形式和传播手段，使思政教育更可感、更入脑入心。",
        "reminder": "讲这题要抓“真实场景”。没有真实实践，就解释不了为什么课堂会更精彩；没有守正，就解释不了为什么它仍然是思政课。",
    },
    {
        "title": "有字之书和无字之书：书本知识必须回到生活和社会实践",
        "source": "2026海淀二模 第16题",
        "question": "从哲学角度，谈谈为什么要把读“有字之书”和读“无字之书”结合起来。",
        "material": "“有字之书”承载经典学说、专业知识和文明智慧；“无字之书”镌刻于广袤大地，蕴藏于社会万象之中。设问问二者为什么要结合。",
        "principles": "联系观点、实践与认识、认识发展、理论联系实际。",
        "logic": "这道题的核心不是劝学生多读书，而是说明书本知识和社会实践不能割裂。“有字之书”提供系统认识，“无字之书”提供现实生活、社会实践和问题情境。只读有字之书，容易停留在纸面；只读无字之书，又可能只有零散经验。二者结合，才能让认识从书本走向现实，再在现实中得到检验和发展。",
        "landing": "答案可以写：世界是普遍联系的，青年成长要把书本知识同社会生活、时代实践和现实问题联系起来。实践是认识的基础，读无字之书能让人在社会实践中理解、检验和发展已有认识；读有字之书又能为实践提供理论滋养和方法指导。把二者结合起来，才能把知识变成本领，把见闻提升为真正的认识。",
        "reminder": "不要把“无字之书”写成单纯旅游见闻，它指的是社会实践、生活经验和真实问题。",
    },
    {
        "title": "中式生活方式跨文化共鸣：独特中国表达承载人类共同追求",
        "source": "2026西城二模 第16题",
        "question": "结合材料，运用《哲学与文化》知识，分析中式生活方式为什么能跨越文化差异、引发全球青年广泛共鸣。",
        "material": "材料写全球青年体验喝热水、练八段锦、节气起居、和睦共居，并在烟火气中感受中华文明崇尚和谐、顺应自然、身心兼顾的智慧。",
        "principles": "矛盾普遍性和特殊性、实践是认识的基础、中华优秀传统文化的当代价值、文化交流。",
        "logic": "“跨越差异”说明它不是取消文化差异，而是在特殊的中国生活方式里包含了普遍的人类追求。喝热水、八段锦、节气起居是中国表达，这是特殊性；追求健康、和谐、身心平衡、人与自然协调，是不同文化中的青年都能理解的普遍性。材料还强调外国青年亲身体验和旅行感知，说明文化理解不是靠抽象宣传，而是在实践体验中形成。",
        "landing": "答案可以落成三层：第一，中式生活方式以独特形式承载健康、和谐、顺应自然等共同价值，实现特殊性和普遍性的统一。第二，全球青年通过亲身体验和日常模仿，在实践中理解中华文化智慧，形成真实共鸣。第三，中华优秀传统文化具有当代价值，通过可感可用的生活方式进入世界青年的日常生活，促进跨文化理解与交流。",
        "reminder": "这题不要只写“中华文化博大精深”。要说明为什么外国青年也能共鸣：因为中国方式里有人的共同生活追求。",
    },
    {
        "title": "正确政绩观：为人民、从实际、看长远",
        "source": "2026西城二模 第20题",
        "question": "结合材料，综合运用所学，谈谈你对“要牢固树立和践行正确政绩观”的理解。",
        "material": "材料围绕党员干部“为谁干事、干什么事、怎么干事”展开，强调必须以最广大人民群众的最大利益为最高标准，既做群众得实惠的实事，也做为后人打基础、利长远的好事，并坚持从实际出发、按规律办事、科学决策、苦干实干。",
        "principles": "人民群众主体地位、价值判断和价值选择、一切从实际出发、尊重规律、发展观点、实践检验。",
        "logic": "正确政绩观不是单纯评价干部做了多少项目，而是判断政绩方向是否正确、方法是否科学、结果是否经得起检验。材料里的“为谁干事”指向人民立场，“怎么干事”指向实事求是和按规律办事，“干什么事”指向既顾眼前民生又顾长远发展。政绩如果脱离人民、脱离实际、只追求短期显绩，就不是正确政绩。",
        "landing": "答案可以这样写：第一，正确政绩观必须站稳人民立场，把人民群众根本利益作为最高标准，创造人民拥护的实绩。第二，坚持一切从实际出发和按规律办事，科学决策、真抓实干，反对形象工程和违背规律的短期行为。第三，坚持发展观点，把眼前利益和长远利益统一起来，多做惠民生、打基础、利长远的事。第四，政绩最终要接受实践、人民和历史检验，不能只看一时数据和表面热闹。",
        "reminder": "讲这题要围绕三问展开：为谁干事是人民立场，怎么干事是实事求是，干什么事是眼前和长远统一。",
    },
    {
        "title": "湿地保护：守护生态之美，也守护文化根脉",
        "source": "2026丰台二模 第16题",
        "question": "结合材料，运用《哲学与文化》知识，谈谈你对“守护湿地，就是守护生态之美与文化根脉”的认识。",
        "material": "材料把湿地放在生态系统、流域治理、传统生态智慧、诗意审美和文化记忆中理解，要求回应“生态之美”和“文化根脉”两个层次。",
        "principles": "尊重客观规律和发挥主观能动性、系统观念、具体问题具体分析、两点论和重点论、生态价值观、中华优秀传统文化。",
        "logic": "这道题最容易写薄，因为学生会只写“湿地很美”或“保护环境”。真正的逻辑是：湿地首先是有自身规律的生态系统，必须顺应水网、流域、动植物和人的活动关系；湿地又承载天人合一、和合共生等传统生态智慧，是文化根脉的一部分。守护湿地就要把生态治理和文化传承一起看，用系统治理和具体施策让生态价值、文化价值、人民长远利益统一起来。",
        "landing": "答案可以这样组织：第一，尊重湿地生态规律，在保护修复中发挥人的主动性，不能凭主观愿望任意改造湿地。第二，用系统观念统筹水系、流域、植被、村落、农田和人的生产生活，解决湿地问题要看到整体联系。第三，具体问题具体分析，因水制宜、因地制宜，不同湿地采用不同保护路径。第四，坚持正确生态价值观，继承天人合一、和合共生的传统生态智慧，使湿地保护既延续生态之美，也延续文化根脉。",
        "reminder": "这题一定要回扣“守护湿地”。单写中华文化源远流长、博大精深，或者通篇罗列原理，都撑不起高质量答案。",
    },
    {
        "title": "科技自立自强为什么是关键：抓主要环节，还要系统推进",
        "source": "2026丰台二模 第22题",
        "question": "综合运用所学，谈谈对“建设社会主义现代化强国，关键在科技自立自强”的理解。",
        "material": "材料强调建设社会主义现代化强国关键在科技自立自强，同时写到关键核心技术买不来、要不来，要完善新型举国体制，统筹教育、科技、人才和产业创新资源，提升创新体系整体效能。",
        "principles": "主要矛盾和关键环节、两点论和重点论、系统观念、联系观点、实践观点、新质生产力。",
        "logic": "“关键在科技自立自强”直接提示要抓关键环节：现代化强国建设有很多任务，但在新一轮科技革命和产业变革中，科技能力决定发展主动权、安全主动权和新质生产力水平。同时，抓关键不是只盯一个实验室或一项技术，还要把教育、科技、人才、产业、制度优势组织成整体合力。",
        "landing": "答案可以分两层写：第一，从重点看，科技自立自强是现代化强国建设的关键环节，关键核心技术不能依赖别人，只有掌握科技主动权，才能支撑新质生产力和国家安全。第二，从全局看，科技自立自强要坚持系统观念，完善新型举国体制，统筹教育、科技、人才和产业创新资源，让制度优势、市场活力和社会协同形成创新合力。最后回扣：抓住科技这个关键，不等于忽视其他方面，而是用关键突破带动现代化建设全局。",
        "reminder": "这题要防止两种偏差：只写“科技重要”会空，只写“教育科技人才一起抓”又抓不住“关键”。要把重点论和系统观念合起来。",
    },
    {
        "title": "中华工业文化的浪漫：把天工、匠心和现代制造连成一条长链",
        "source": "2026房山二模 第16题",
        "question": "结合材料，运用《哲学与文化》知识，谈谈如何从中华工业文化读懂中华民族最感动人的浪漫。",
        "material": "材料从良渚玉器、西周青铜、《考工记》中的“天有时、地有气、材有美、工有巧”，延伸到现代工业体系和北斗原子钟等当代制造成就。",
        "principles": "系统观念、量变质变、尊重客观规律和发挥主观能动性、价值观导向、中华民族精神、中华优秀传统文化的连续性和创新性。",
        "logic": "“浪漫”不是空泛抒情，而是中华民族长期把自然条件、材料特性、工艺智慧、劳动创造和国家发展连接起来。古代工匠尊重天时、地气、材料之美，并发挥人的巧思；现代工业把这种精益求精、协同攻关、追求极致的精神延续到复杂系统制造中。长期积累形成质的跃升，传统工艺智慧在现代工业中实现新的发展。",
        "landing": "答案可以写：中华工业文化体现尊重规律和发挥主观能动性的统一，能工巧匠在把握材料、工艺和自然条件中创造器物之美；体现系统观念，工业成就来自劳动者、技术、标准、协作和国家工业体系的整体支撑；体现量变到质变，代代积累的工艺追求和创新实践最终转化为现代制造能力；体现中华文明连续性和创新性，传统工匠精神在新时代工业实践中创造性发展，成为民族精神和文化自信的重要来源。",
        "reminder": "讲这题不要只剩“工匠精神”四个字，要把古代工艺、现代制造、长期积累、系统协作连起来。",
    },
    {
        "title": "OPC与辩证否定：新形态不是全盘抛弃，而是扬弃旧形态",
        "source": "2026房山二模 第18（2）题",
        "question": "结合材料一和材料二，运用《逻辑与思维》辩证否定观知识，分析OPC的出现和发展。",
        "material": "材料把传统一人公司同“一个创意大脑+数字员工”的OPC新创业形态对照，同时列出数字员工可能带来的商业诋毁、商业秘密泄露、著作权风险等问题。",
        "principles": "辩证否定、扬弃、联系的环节、发展的环节、新事物成长。",
        "logic": "题干直接要求用辩证否定观。OPC不是把一人公司完全推倒，也不是原封不动保留旧模式，而是在旧的一人公司基础上引入数字员工，形成新的创业范式。新事物发展也不是没有问题，数字员工的功能要肯定，法律风险和治理短板要改造，这正是扬弃：既保留合理因素，又克服消极因素。",
        "landing": "答案可以写：OPC的出现是对传统一人公司这一旧矛盾统一体的辩证否定，它以数字员工扩展个人创业能力，形成“单人成军”的新形态，体现否定是联系和发展的环节。OPC的发展要坚持扬弃，肯定并保留数字员工在效率、创意和协作上的积极功能，同时针对商业诋毁、泄密、侵权等风险进行规范和改造，把风险治理纳入新事物发展过程。",
        "reminder": "这题不要写成“新的一定战胜旧的”。高质量答案要讲清楚：从旧形态中生长出来，保留优势，改造问题，这才叫辩证否定。",
    },
    {
        "title": "新征程就是新的长征：历史精神要落到现代化实践",
        "source": "2026房山二模 第21题",
        "question": "综合运用所学，谈谈对“新征程就是新的长征”的理解。",
        "material": "材料把长征精神同以中国式现代化全面推进中华民族伟大复兴联系起来，强调新征程仍要坚持党的领导、坚持真理、弘扬伟大精神、面对风险挑战继续前进。",
        "principles": "联系观点、发展观点、矛盾普遍性、人民群众、价值观导向、伟大长征精神。",
        "logic": "这道题不是让学生复述长征历史，而是把“历史长征”和“现代化新征程”建立联系。新的长征说明任务变了、时代条件变了，但奋斗精神、人民主体、党的领导、攻坚克难没有变。新征程上的现代化建设会遇到新矛盾新问题，所以要在发展中解决矛盾，在实践中继续推进民族复兴。",
        "landing": "答案可以这样写：新征程与长征路具有历史联系，长征精神为现代化建设提供精神动力和价值引领。中国式现代化是不断发展的历史进程，必须在党的领导下坚持正确方向，依靠人民群众，弘扬艰苦奋斗和攻坚克难精神。面对新征程中的风险挑战，要承认矛盾、分析矛盾、解决矛盾，把历史精神转化为推进强国建设、民族复兴的现实行动。",
        "reminder": "不要把这题写成纯文化传承。它的哲学骨架是联系历史与现实、在发展中解决矛盾、把精神力量转化为实践力量。",
    },
    {
        "title": "良法与善治：制度规则要在治理实践中实现价值",
        "source": "2026石景山二模 第17（3）题",
        "question": "从哲学角度，分析“良法”和“善治”的关联。",
        "material": "材料讲养老服务条例把制度设计落实到老年人的一餐一饭、一医一药，使制度责任转化为可及感、踏实感、获得感，并指出良法只是起点，善治才是目标。",
        "principles": "联系观点、实践观点、矛盾对立统一、价值实现。",
        "logic": "良法和善治不是两个孤立概念。良法提供治理依据、规则边界和价值方向；善治检验法律是否真正落到人民生活中。没有良法，治理容易无章可循；只有良法而没有执行、协同和服务，法律也停留在纸面。二者既有区别，又相互依赖、相互促进。",
        "landing": "答案可以写：良法和善治是紧密联系的统一体，良法为善治提供制度依据和规范保障，能够最大程度凝聚认同、发挥法律效力；善治是良法的实践目标和现实展开，只有把法律规则落实到权力运行、公共服务和权益保障中，才能推动治理现代化。材料中的养老服务把制度条文转化为人人可感的照护，说明良法要通过治理实践实现其价值。",
        "reminder": "不要只背法治术语。设问要“从哲学角度”，核心是关系：起点与目标、规则与实践、制度设计与人民获得感。",
    },
    {
        "title": "北京在中国式现代化大局中的作用：示范不是单点优秀，而是系统带动",
        "source": "2026石景山二模 第20题",
        "question": "结合材料，综合运用所学，分析北京如何在中国式现代化建设大局中发挥作用。",
        "material": "材料围绕北京“十五五”发展，强调北京在经济高质量发展、文化强国建设、制度优势、城市治理等方面的代表性、指向性和示范引领作用。",
        "principles": "系统观念、整体与部分、联系观点、发展观点。",
        "logic": "北京不是孤立建设一座城市，而是在国家现代化全局中承担示范、引领、服务和支撑功能。作为部分，北京的发展要服从和服务中国式现代化整体；作为首都，北京又能通过制度探索、科技创新、文化中心建设、治理经验输出带动整体发展。",
        "landing": "答案可以这样落：从整体与部分看，北京要把自身发展放进中国式现代化建设大局中谋划，服务国家战略而不是只追求地方指标。从系统观念看，北京要统筹经济、科技、文化、生态、治理等方面，形成相互支撑的现代化实践。从发展观点看，北京通过高质量发展、全国文化中心建设和治理创新形成示范效应，为全国现代化建设提供经验和支撑。",
        "reminder": "这题不要只罗列北京做了什么，要写清“北京这个部分怎样服务全国这个整体”。",
    },
    {
        "title": "新大众文艺的三组平衡：多样不能失向，人民性不能粗糙化，流量不能压倒价值",
        "source": "2026顺义二模 第16题",
        "question": "结合材料，运用《哲学与文化》知识，谈谈对新大众文艺面对“尊重多样与弘扬主流、人民性与艺术性、社会效益与经济效益”之间寻找平衡点这一时代课题的思考。",
        "material": "材料把新大众文艺放在多元内容、网络传播、人民创作和市场流量中考察，提出三组关系：尊重多样与弘扬主流，人民性与艺术性，社会效益与经济效益。",
        "principles": "两点论和重点论、矛盾普遍性和特殊性、辩证否定和守正创新、人民群众主体地位、价值观导向、社会存在与社会意识。",
        "logic": "这道题最重要的是“平衡点”。尊重多样不是放任无序，弘扬主流也不是取消多样；人民性不是粗糙化、低俗化，艺术性也不能脱离人民；经济效益可以支撑传播和创作，但社会效益必须放在首位。材料每组关系都要求既看两方面，又抓住重点。",
        "landing": "答案可以分三组写：第一，尊重文艺样态多样和群众表达差异，同时以主流价值、社会主义核心价值观和先进文化方向引领创作，做到多样而不跑偏。第二，坚持人民群众是文化创造主体，扎根人民生活、反映人民心声，同时提升艺术水准和思想内涵，做到通俗不低俗。第三，处理好社会效益和经济效益，坚持社会效益优先，以正能量和高质量内容赢得市场，再用市场传播扩大优秀作品影响力。",
        "reminder": "这题必须讲“既要又要，但重点在哪里”。如果只平均写三组关系，就缺少两点论和重点论的方法味。",
    },
    {
        "title": "先见与远虑：预判规律，谋划长远，落到持续实践",
        "source": "2026顺义二模 第21题",
        "question": "综合运用所学，阐述“先见”与“远虑”的传统理政智慧为何能成为支撑中国式现代化稳步前行、护航民族伟大复兴行稳致远的深厚精神底气与重要实践指引。",
        "material": "材料围绕“先见”和“远虑”展开，强调在复杂环境中预判趋势、谋定后动、守住底线、长期推进，并把这种传统理政智慧用于中国式现代化实践。",
        "principles": "尊重客观规律和发挥主观能动性、实践与认识、发展观点、量变质变、价值观导向。",
        "logic": "“先见”不是算命式预言，而是在实践和调查研究中把握规律、看清趋势；“远虑”不是空喊长远，而是把长远目标转化为阶段安排和持续行动。中国式现代化面对复杂环境，既需要对风险和趋势有科学认识，也需要在正确价值引领下保持战略定力。",
        "landing": "答案可以写：先见体现尊重客观规律和认识指导实践，只有把握发展趋势和风险变化，才能作出科学决策。远虑体现发挥主观能动性和发展观点，把民族复兴的长远目标落实为连续不断的实践安排。二者统一起来，就是在正确价值观引领下谋全局、谋长远，通过持续行动和量的积累推动中国式现代化稳步前行。",
        "reminder": "不要把“先见远虑”写成一般传统文化赞美。它要落到现代化治理：看清规律、提前布局、长期坚持。",
    },
]


EXTRA_SECOND_MOCK_NODE_ROWS: list[dict[str, str]] = [
    {
        "section": "人民群众",
        "source": "2026石景山二模 第20题",
        "kind": "主观题",
        "trigger": "材料把北京放在中国式现代化大局中分析，并强调中国式现代化是全体人民共同富裕的现代化，北京发展要回应人民生活、文化服务和城市治理需要。",
        "question": "结合材料，综合运用所学，分析北京如何在中国式现代化建设大局中发挥作用。",
        "why": "中国式现代化不是只看城市指标和产业规模，而是要落到人民共同富裕、公共文化服务和治理效能提升上。北京发挥示范作用，最终要通过高质量发展和公共服务改善人民生活。",
        "answer": "北京在中国式现代化大局中发挥作用，要坚持以人民为中心，把高质量发展、文化强国建设和城市治理创新转化为人民群众可感可及的获得感，在服务人民生活中体现首都示范引领作用。",
    }
]


ROW_REMINDER_OVERRIDES: dict[tuple[str, str], str] = {
    ("2026石景山二模 第20题", "人民群众"): "讲这题写人民立场时，落点不是泛泛说人民重要，而是北京高质量发展、文化建设和治理创新最终要服务共同富裕和人民可感可及的获得感。",
    ("2026朝阳二模 第21题", "量变与质变 / 适度原则"): "写战略定力时，久久为功不是慢慢做就行，要把宏伟蓝图分解为可执行的阶段任务，并通过持续落实形成累积效果。",
    ("2026房山二模 第16题", "价值观的导向作用"): "写价值观角度时，重点抓“极致”这一价值追求怎样引导劳动创造，不要只把工匠精神写成态度好。",
    ("2026房山二模 第16题", "系统观念 / 系统优化"): "写系统观念角度时，重点抓材料、工艺、劳动者、标准、协作和国家工业体系怎样形成整体创造力。",
    ("2026房山二模 第16题", "量变与质变 / 适度原则"): "写量变质变角度时，重点抓长期精度积累、技术积累和代际传承怎样转化为现代制造能力的跃升。",
    ("2026房山二模 第16题", "尊重客观规律与发挥主观能动性相结合"): "写规律和能动性角度时，重点抓“天时、地气、材美、工巧”：先尊重客观条件，再发挥人的工巧创造。",
    ("2026石景山二模 第17（3）题", "矛盾就是对立统一"): "写矛盾角度时，突出良法和善治的区别与统一：良法是制度起点，善治是治理目标，二者不能割裂。",
    ("2026石景山二模 第17（3）题", "联系的普遍性 / 联系的观点（总）"): "写联系角度时，突出良法和善治的相互支撑：法律规则要通过治理实践转化为实际效能。",
    ("2026石景山二模 第17（3）题", "实践与认识（总）"): "写实践角度时，突出法律来自治理问题，又要回到一餐一饭、一医一药等具体服务中接受检验。",
}


ROW_ANSWER_APPEND: dict[tuple[str, str], str] = {
    ("2026石景山二模 第2题", "矛盾就是对立统一"): "注意错肢①：不能把“保留传统剪纸合理因素”说成被辩证否定；辩证否定是否定消极因素、保留合理因素，题中正确链条落在艺术性与思想性的对立统一。",
}


CURATED: dict[tuple[str, str], dict[str, str]] = {
    ("2024东城一模", "16"): {
        "prompt": "结合材料，综合运用所学，说明中华文明新形态为什么既体现中华文化特色，又能够为世界文明进步贡献中国智慧。",
    },
    ("2024东城一模", "18"): {
        "prompt": "结合材料，运用所学，说明北京科技体制改革如何推动新质生产力发展。",
    },
    ("2024东城一模", "21"): {
        "prompt": "以京津冀协同发展为例，综合运用所学，分析战略性有利条件在推进现代化建设中是如何发挥作用的。",
    },
    ("2024东城二模", "18"): {
        "prompt": "围绕新就业形态劳动者权益保障，说明应如何通过制度完善和多主体协同回应现实问题。",
    },
    ("2024东城二模", "21"): {
        "prompt": "实现中华民族伟大复兴进入了不可逆转的历史进程。结合材料，综合运用所学，分析战略性有利条件在这一历史进程中是如何发挥作用的。",
    },
    ("2025丰台期末", "16"): {
        "node": "实践与认识 / 一切从实际出发",
        "prompt": "围绕“胸中有数”，说明在工作生活和国家治理中为什么要重视真实数据、基本统计和主要百分比。",
        "answer": "“数”不是主观想象，而来自实践、调查和对客观情况的把握。作答时应说明：做计划、定方案、治理国家，都要从真实情况和数据出发，在实践中取得可靠认识，再用这些认识指导行动。",
        "note": "本题不再把十几个原理分散挂载。主链保留实践与认识、一切从实际出发；涉及“主要矛盾的主要方面”的内容只作为判断事物性质时的提醒，不作为本题唯一主线。",
    },
    ("2024东城二模", "16"): {
        "node": "尊重规律与发挥主观能动性 / 联系系统 / 人民群众与价值选择",
        "prompt": "结合材料，运用《哲学与文化》知识，分析为何二千多岁的桑基鱼塘仍未老。",
        "answer": "桑基鱼塘形成于人民农业实践，顺应水、桑、蚕、鱼、塘之间的生态规律；它不是孤立保留旧形式，而是在新时代把生态保护、产业融合、文化研学和乡村振兴联系起来，继续服务人民生活。作答可按“形成有规律、发展成系统、价值为人民”三层展开。",
        "note": "不再单独强挂辩证否定；若课堂使用守正创新，只作为文化转化表达，不写成唯一哲学主线。",
    },
    ("2024丰台二模", "20"): {
        "node": "矛盾分析法 / 联系观点 / 实践推动认识发展",
        "prompt": "运用哲学知识，谈谈对“统筹发展和安全”的认识。",
        "answer": "发展与安全不是二选一。安全为发展提供条件，发展又为安全提供基础；在不同时期，我们党在实践中不断深化对二者关系的认识。作答应围绕二者相互联系、对立统一和在实践中不断深化认识展开。",
        "note": "本题作为综合等级题处理，不再拆成三个彼此重复的独立条目。",
    },
    ("2026海淀期中", "22"): {
        "node": "社会历史发展规律 / 人民群众 / 社会意识反作用 / 系统观念",
        "prompt": "围绕中华民族伟大复兴势不可挡，说明这一判断的历史基础、实践主体和现实支撑。",
        "answer": "复兴势不可挡，不是口号，而是社会历史发展趋势、长期实践积累和现实条件共同作用的结果。中国共产党领导人民长期奋斗，道路、理论、制度、文化形成支撑力量；作答要把历史趋势、人民主体、制度文化支撑和国内国际条件组织成一条逻辑链。",
        "note": "作答时把四层合成一条逻辑链：历史趋势、人民主体、制度文化支撑、国内国际条件。",
    },
    ("2024朝阳二模", "19"): {
        "node": "发展观点 / 价值观导向 / 文化创造性转化",
        "prompt": "说明中华优秀传统文化的特质如何赋予中国式现代化深厚底蕴。",
        "answer": "“生生之宇宙观”提示不断变化发展，“一体之天人观”提示人与自然、人与他人的共同体关系，“法天地之人生观”提示自强不息、厚德载物的价值追求。作答应把传统文化特质转化为现代化的精神底蕴和实践方向。",
        "note": "主条更适合放文化宝典；哲学宝典只保留发展观点和价值观导向的轻量触发。",
    },
    ("2024丰台一模", "21"): {
        "prompt": "结合材料，综合运用所学，谈谈对全人类共同价值的理解。",
    },
    ("2024朝阳一模", "16"): {
        "prompt": "结合材料，综合运用所学，分析“接受人民监督”与“勇于自我革命”的关系。",
    },
    ("2024西城二模", "18"): {
        "prompt": "结合材料，运用《哲学与文化》知识，说明应如何推动乡村文化振兴。",
    },
    ("2025丰台期末", "17"): {
        "prompt": "一座座城市的成长，体现着“人民城市人民建，人民城市为人民”理念。结合材料，运用所学，谈谈对这一理念的认识。",
    },
    ("2025朝阳期末", "22"): {
        "prompt": "推进马克思主义中国化时代化是一个追求真理、揭示真理、笃行真理的过程。结合材料，综合运用所学，说明如何不断谱写马克思主义中国化时代化新篇章。",
    },
    ("2025石景山一模", "21"): {
        "prompt": "结合材料，综合运用所学，阐述我们党在进一步全面深化改革中统筹破立关系，对全面建成社会主义现代化强国的重要意义。",
    },
    ("2025门头沟一模", "21"): {
        "prompt": "结合材料，综合运用所学，说明门头沟区转型实践对探索中国式现代化区域样本的启示。",
    },
    ("2026石景山期末", "20"): {
        "prompt": "结合材料，运用哲学知识，说明“脉冲”试验法如何体现按规律办事与发挥主观能动性。",
    },
}


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def clean(text: str, limit: int = 520) -> str:
    text = re.sub(r"\s+", " ", text or "").strip()
    text = strip_metadata(text)
    for old, new in [
        ("评分细则明确", "题目要求"),
        ("评分参考明确", "题目要求"),
        ("评分细则", "作答要求"),
        ("评分参考", "作答要求"),
        ("评分", "作答"),
        ("阅卷细则", "作答要求"),
        ("阅卷总结", "作答要求"),
        ("评标 PPT", "作答要求"),
        ("评标", "作答要求"),
        ("答案线索", "题目要求"),
        ("可从", "可以围绕"),
        ("补入", "放入"),
        ("证据", "来源"),
        ("审计", "检查"),
        ("给分点", "答题点"),
        ("参考答案", "作答示例"),
        ("细则", "要求"),
    ]:
        text = text.replace(old, new)
    text = text.replace("CSV", "表格")
    return text if len(text) <= limit else text[: limit - 1] + "…"


def strip_metadata(text: str) -> str:
    text = text or ""
    text = re.sub(r"## [^`]*? - source_path: `[^`]*`", " ", text)
    text = re.sub(r"- (?:source_path|status|method|sha256|gpt_markdown_path|render_dir): `[^`]*`", " ", text)
    text = re.sub(r"`[A-Z]:\\[^`]+`", " ", text)
    text = re.sub(r"[A-Z]:\\Users\\Administrator\\[^ ]+", " ", text)
    text = re.sub(r"PAGE \\?\* MERGEFORMAT \d+", " ", text)
    text = re.sub(r"第\d+页/共\d+页", " ", text)
    text = re.sub(r"---\s*## .*", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def clean_prompt(text: str, fallback: str, limit: int = 520) -> str:
    text = strip_metadata(text)
    for marker in [
        "参考答案",
        "第一部分",
        "一、选择题",
        "二、非选择题",
        "阅卷细则",
        "阅卷总结",
        "评分细则",
        "等级赋分",
        "试题分析",
        "讲评",
    ]:
        pos = text.find(marker)
        if pos > 60:
            text = text[:pos]
    text = re.sub(r"\(考生务必.*?无效\)", "", text)
    text = re.sub(r"高三政治[一二]模.*?(?=\d+[\.．、（])", "", text)
    text = text.strip(" ：；，。")
    if not text or text.startswith("##") or "source_path" in text or len(text) < 12:
        text = fallback
    return clean(text, limit)


def source_text(row: dict[str, str]) -> str:
    return f"{row['suite_name']} 第{row['question_no']}题"


def build_second_mock_thick_cards() -> str:
    lines = [
        "",
        "# 2026年二模主观题加厚主讲卡",
        "",
        "这一部分把同一道题的多个哲学点合并成一张主讲卡。课堂上先按这里讲完整逻辑，再回到后面的原理节点查对应知识。",
        "",
    ]
    for card in SECOND_MOCK_THICK_CARDS:
        lines.extend(
            [
                f"## {card['title']}",
                "",
                f"**来源题目**：{card['source']}",
                "",
                f"**完整设问**：{card['question']}",
                "",
                f"**材料触发点**：{card['material']}",
                "",
                f"**能用哪些原理**：{card['principles']}",
                "",
                f"**为什么这样对应**：{card['logic']}",
                "",
                f"**答案落点**：{card['landing']}",
                "",
                f"**使用提醒**：{card['reminder']}",
                "",
            ]
        )
    return "\n".join(lines)


def build_second_mock_thick_md() -> Path:
    out_md = OUT_DIR / "04_二模加厚卡" / "2026二模哲学主观题加厚主讲卡_v10.md"
    out_md.parent.mkdir(parents=True, exist_ok=True)
    out_md.write_text(build_second_mock_thick_cards().lstrip(), encoding="utf-8")
    return out_md


def second_mock_contexts() -> dict[str, dict[str, str]]:
    contexts: dict[str, dict[str, str]] = {}
    for card in SECOND_MOCK_THICK_CARDS:
        contexts[card["source"]] = card
    return contexts


def drop_existing_second_mock_blocks(text: str) -> str:
    lines = text.splitlines()
    kept: list[str] = []
    i = 0
    while i < len(lines):
        if lines[i].startswith("## "):
            j = i + 1
            while j < len(lines) and not lines[j].startswith("## "):
                j += 1
            block = "\n".join(lines[i:j])
            if re.search(r"\*\*来源题目\*\*：2026\S+二模", block):
                i = j
                continue
        kept.append(lines[i])
        i += 1
    return "\n".join(kept).rstrip() + "\n"


def read_second_mock_rows() -> list[dict[str, str]]:
    rows = json.loads(SECOND_MOCK_JSON.read_text(encoding="utf-8"))
    return rows + EXTRA_SECOND_MOCK_NODE_ROWS


def theme_from_source(source: str) -> str:
    themes = {
        "2026朝阳二模 第16题": "城市文化传承创新",
        "2026朝阳二模 第21题": "四个中国与战略定力",
        "2026东城二模 第16题": "京彩课堂",
        "2026海淀二模 第16题": "有字之书和无字之书",
        "2026西城二模 第16题": "中式生活方式跨文化共鸣",
        "2026西城二模 第20题": "正确政绩观",
        "2026丰台二模 第16题": "湿地保护",
        "2026丰台二模 第22题": "科技自立自强",
        "2026房山二模 第16题": "中华工业文化",
        "2026房山二模 第18（2）题": "OPC新创业形态",
        "2026房山二模 第21题": "新征程就是新的长征",
        "2026石景山二模 第17（3）题": "良法与善治",
        "2026石景山二模 第20题": "北京服务现代化大局",
        "2026顺义二模 第16题": "新大众文艺",
        "2026顺义二模 第21题": "先见与远虑",
        "2026朝阳二模 第3题": "词元概念普及",
        "2026朝阳二模 第4题": "智慧物业治理",
        "2026东城二模 第2题": "德字内涵演变",
        "2026东城二模 第4题": "灯诱法与智能识虫",
        "2026海淀二模 第4题": "深圳湾护鸟照明",
        "2026西城二模 第3题": "现代城市规划",
        "2026西城二模 第4题": "外卖骑手治理闭环",
        "2026石景山二模 第1题": "十五五规划",
        "2026石景山二模 第2题": "立体抗战剪纸",
        "2026顺义二模 第3题": "沈括隙积术",
        "2026房山二模 第2题": "新大众文艺选择题",
        "2026房山二模 第5题": "工位情绪充电站",
    }
    return themes.get(source, source.replace("2026", "").replace("二模", "").strip())


def node_short(section: str) -> str:
    return (
        section.replace("（总）", "")
        .replace(" / ", "、")
        .replace("/", "、")
        .replace("相结合", "统一")
    )


def build_second_mock_node_entries() -> str:
    rows = read_second_mock_rows()
    contexts = second_mock_contexts()
    grouped: dict[str, list[dict[str, str]]] = defaultdict(list)
    section_order: list[str] = []
    for row in rows:
        if row["section"] not in grouped:
            section_order.append(row["section"])
        grouped[row["section"]].append(row)

    lines = [
        "",
        "# 2026年二模题例（按原理节点）",
        "",
        "以下按原理节点整理二模题。同一张卷子的多个原理分别进入对应节点，但每条都保留完整设问、材料触发、对应逻辑和答案落点。选择题只写正确项或错肢的判断链，不混成主观题答题链。",
        "",
    ]

    for section in section_order:
        lines.extend([f"## {section}", ""])
        for row in grouped[section]:
            source = row["source"]
            ctx = contexts.get(source, {})
            kind = row.get("kind", "")
            title_prefix = theme_from_source(source)
            title = f"{title_prefix}中的{node_short(section)}"
            if kind == "选择题":
                title = f"{title_prefix}选择题中的{node_short(section)}"
            material = row.get("trigger", "").strip()
            if ctx:
                material = f"{material} 放到整题里看：{ctx.get('material', '')}"
            why = row.get("why", "").strip()
            if ctx:
                why = f"{why} 整题逻辑是：{ctx.get('logic', '')}"
            answer = row.get("answer", "").strip()
            if ctx and kind != "选择题":
                answer = f"{answer} 作答时还要回到本题主线：{ctx.get('landing', '')}"
            if kind == "选择题":
                answer = f"{answer} 选择题判断时只把它当作正确项或错肢识别链，不写成主观题万能模板。"
            if (source, section) in ROW_ANSWER_APPEND:
                answer = f"{answer} {ROW_ANSWER_APPEND[(source, section)]}"
            question_label = "完整题干" if kind == "选择题" else "完整设问"
            lines.extend(
                [
                    f"### {title}",
                    "",
                    f"**来源题目**：{source}",
                    "",
                    f"**{question_label}**：{row.get('question', '').strip()}",
                    "",
                    f"**材料触发点**：{clean(material, 760)}",
                    "",
                    f"**为什么能想到这个原理**：{clean(why, 900)}",
                    "",
                    f"**答案落点**：{clean(answer, 900)}",
                    "",
                ]
            )
            reminder = ROW_REMINDER_OVERRIDES.get((source, section), ctx.get("reminder", ""))
            if reminder and kind != "选择题":
                lines.extend([f"**使用提醒**：{clean(reminder, 420)}", ""])
    return "\n".join(lines)


def build_subjective_cards() -> str:
    rows = read_csv(SUBJ_CSV)
    lines = [
        "",
        "# 2024-2026 主观题重点深化卡",
        "",
        "以下题例按“材料触发点—为什么想到这个原理—答案落点”整理。一个题目如果同时考多个哲学点，就在同一张卡里分层讲清楚。",
        "",
    ]
    for r in rows:
        key = (r["suite_name"], r["question_no"])
        curated = CURATED.get(key, {})
        triggers = clean(curated.get("node") or r["triggers"], 180)
        prompt = clean_prompt(
            curated.get("prompt") or r["prompt_snippet"],
            f"{r['suite_name']} 第{r['question_no']}题围绕材料设问。",
            520,
        )
        material = clean(r["materials"], 520)
        answer = clean(curated.get("answer") or r["logic_summary"], 620)
        note = clean(curated.get("note", ""), 360)
        lines.extend(
            [
                f"## {triggers}",
                "",
                f"**来源题目**：{source_text(r)}",
                "",
                f"**设问**：{prompt}",
                "",
                f"**材料触发点**：{material}",
                "",
                f"**为什么能想到这个原理**：{answer}",
                "",
                f"**答案落点**：{answer}",
                "",
            ]
        )
        if note:
            lines.extend([f"**使用提醒**：{note}", ""])
    return "\n".join(lines)


def build_main_md() -> Path:
    out_md = OUT_DIR / "01_学生版Word" / "必修四哲学宝典_v10_二模同流程修订学生版.md"
    base = BASE_MD.read_text(encoding="utf-8")
    text = drop_existing_second_mock_blocks(base.rstrip())
    node_entries = build_second_mock_node_entries()
    match = re.search(r"(生成时间：[^\n]+\n)", text)
    if match:
        text = text[: match.end()] + node_entries + "\n" + text[match.end() :]
    else:
        text = text + "\n" + node_entries
    text = text + "\n" + build_subjective_cards() + "\n"
    text = text.replace("# 必修四哲学材料-知识触发框架 v6", "# 必修四哲学宝典（2024-2026北京区卷学生版）")
    text = re.sub(r"生成时间：2026-04-29 01:36:22", "生成时间：2026-05-24", text)
    text = text.replace("## 2024_v6_student_entries", "## 2024年题例")
    text = text.replace("## 2025A_v6_student_entries", "## 2025年一模与期末题例")
    text = text.replace("## 2025B_v6_student_entries", "## 2025年二模与补充题例")
    text = text.replace("## 2026A_v6_student_entries", "## 2026年一模、期中与期末题例")
    text = text.replace("## 2026B_v6_student_entries", "## 2026年期中期末补充题例")
    text = re.sub(r"说明：本文件只展示最终学生版应采用的写法。.*?\n\n", "", text, count=1)
    # Student-facing file cannot contain process words.
    for token in FORBIDDEN:
        text = text.replace(token, "")
    out_md.parent.mkdir(parents=True, exist_ok=True)
    out_md.write_text(text, encoding="utf-8")
    return out_md


def build_choice_md() -> Path:
    rows = read_csv(CHOICE_CSV)
    out_md = OUT_DIR / "02_选择题专册" / "必修四哲学选择题触发与错肢专册_v10.md"
    lines = [
        "# 必修四哲学选择题触发与错肢专册（2024-2026北京区卷）",
        "",
        "说明：选择题按“题干或选项触发—对应知识—判断逻辑—风险等级”整理，不混入主观题给分链。",
        "",
    ]
    grouped: dict[str, list[dict[str, str]]] = defaultdict(list)
    for r in rows:
        grouped[clean(r.get("trigger") or r.get("section") or "未分类", 80)].append(r)
    idx = 1
    for trigger in sorted(grouped):
        lines.extend([f"## {trigger}", ""])
        for r in grouped[trigger]:
            blob = " ".join(r.values())
            risk = "待回原卷确认" if any(t in blob for t in HIGH_RISK) else "可作选择题链"
            lines.extend(
                [
                    f"### {idx}. {r.get('suite_name')} 第{r.get('question_no_norm')}题",
                    "",
                    f"**题干或选项触发**：{clean(r.get('material', ''), 520)}",
                    "",
                    f"**对应知识**：{clean(r.get('trigger', ''), 220)}",
                    "",
                    f"**判断逻辑**：{clean(r.get('logic', ''), 620)}",
                    "",
                    f"**风险等级**：{risk}",
                    "",
                ]
            )
            idx += 1
    text = "\n".join(lines)
    for token in ["评分细则", "评分参考", "评分", "阅卷细则", "评标", "补漏", "补丁", "审计", "CSV", "证据"]:
        text = text.replace(token, "")
    out_md.parent.mkdir(parents=True, exist_ok=True)
    out_md.write_text(text, encoding="utf-8")
    return out_md


def markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("**") and "**：" in line:
            label, rest = line.split("**：", 1)
            p = doc.add_paragraph()
            r = p.add_run(label.strip("*") + "：")
            r.bold = True
            p.add_run(rest)
        else:
            doc.add_paragraph(line)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def build_report() -> Path:
    roster = read_csv(ROSTER_CSV)
    subj = read_csv(SUBJ_CSV)
    choice = read_csv(CHOICE_CSV)
    out = OUT_DIR / "03_严审报告" / "65套覆盖与模型审核状态.md"
    lines = [
        "# 65套覆盖与模型审核状态",
        "",
        "## 当前结论",
        "",
        "- 严格题源基数：65套。",
        "- 主宝典：以认可 v6 为底稿，先删除旧底稿中的薄二模块，再按原宝典节点写法重写 2026 二模题例。",
        "- 二模题：已按融合清单处理 52 条二模哲学条目，并按复核意见为石景山二模第20题增补 1 条人民立场节点；主观题按材料链和答案落点写入对应原理节点，选择题按正确项或错肢判断链处理。",
        "- 二模主观题：另保留一份同题合并的加厚主讲卡，方便课堂先讲整题逻辑。",
        "- 选择题：174 条旧漏项已单独成册，不再混入主观题给分链。",
        "- ClaudeCode：已完成独立严审，结论是 v7 不能作为终稿，v8 必须按本轮严格口径处理。",
        "- GPT Pro 网页版：审核包已生成，但 Chrome 扩展当前断连，不能写成已完成。",
        "",
        "## 数字",
        "",
        f"- 65套总清单：{len(roster)}",
        f"- 旧主观质量失败唯一题组：{len(subj)}",
        f"- 旧选择题待闭环：{len(choice)}",
        "",
        "## 交付边界",
        "",
        "- 本夜间版优先保证不再把过程性内容塞进学生正文。",
        "- 高风险选择题标为“待回原卷确认”，避免再次发生过度归因。",
        "- GPT Pro 网页端恢复后，应按 `04_review_packages/GPTPRO_WEB_BATCH_*.md` 分批继续审核。",
        "",
    ]
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text("\n".join(lines), encoding="utf-8")
    return out


def scan_forbidden(paths: list[Path]) -> str:
    lines = ["# 学生版禁词扫描", ""]
    for path in paths:
        text = path.read_text(encoding="utf-8", errors="replace")
        lines.append(f"## {path.name}")
        hit = False
        for token in FORBIDDEN:
            count = text.count(token)
            if count:
                lines.append(f"- {token}: {count}")
                hit = True
        if not hit:
            lines.append("- PASS: 0")
        lines.append("")
    out = OUT_DIR / "03_严审报告" / "student_forbidden_token_scan.md"
    out.write_text("\n".join(lines), encoding="utf-8")
    return "\n".join(lines)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    main_md = build_main_md()
    thick_md = build_second_mock_thick_md()
    choice_md = build_choice_md()
    report = build_report()
    main_docx = main_md.with_suffix(".docx")
    thick_docx = thick_md.with_suffix(".docx")
    choice_docx = choice_md.with_suffix(".docx")
    report_docx = report.with_suffix(".docx")
    markdown_to_docx(main_md, main_docx)
    markdown_to_docx(thick_md, thick_docx)
    markdown_to_docx(choice_md, choice_docx)
    markdown_to_docx(report, report_docx)
    scan = scan_forbidden([main_md, choice_md, thick_md])
    print(main_docx)
    print(thick_docx)
    print(choice_docx)
    print(report_docx)
    print(scan)


if __name__ == "__main__":
    main()
