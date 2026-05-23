from __future__ import annotations

import csv
import re
from collections import Counter
from pathlib import Path

from docx import Document


DESKTOP = Path.home() / "Desktop"
REPO = DESKTOP / "02_代码项目与工具" / "mac-thread-restore" / "beijing-politics-sync-visible"
RUN = REPO / "reports" / "bixiu4_philosophy_full_coverage_double_lane_2026-05-23"
DOCX = DESKTOP / "哲学宝典最终版-飞哥正志讲堂_主次矛盾全覆盖补丁v7_2026-05-23.docx"
V3 = RUN / "01_codex_lane" / "v3_inventory_vs_latest_docx.csv"
OUT = RUN / "05_governor_confucius" / "v7_coverage_verification.md"


DISTRICTS = "海淀|西城|东城|朝阳|丰台|石景山|门头沟|房山|通州|顺义|昌平|延庆|大兴|平谷|怀柔|密云"
KEY_RE = re.compile(rf"(20\d{{2}}).*?({DISTRICTS}).*?(一模|二模|期中|期末).*?第[（(]?(\d+)")


def main() -> None:
    doc = Document(str(DOCX))
    h2: list[str] = []
    h3: list[str] = []
    keys: set[str] = set()

    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        text = p.text.strip()
        if style == "Heading 2":
            h2.append(text)
        if style == "Heading 3":
            h3.append(text)
            m = KEY_RE.search(text)
            if m:
                year, district, stage, question = m.groups()
                keys.add(f"{year}|{district}|{stage}|{question}")

    rows = list(csv.DictReader(V3.open("r", encoding="utf-8-sig")))
    candidates = [
        r
        for r in rows
        if (r["question_nature"] == "subjective" and r["evidence_grade_initial"] in {"A", "B"})
        or (r["question_nature"] == "choice" and r["evidence_grade_initial"] == "C")
    ]

    missing: list[dict] = []
    for r in candidates:
        key = "|".join([r["norm_year"], r["norm_district"], r["norm_stage"], r["norm_question"]])
        if key not in keys:
            missing.append(r | {"verify_key": key})

    all_text = "\n".join(p.text for p in doc.paragraphs)
    terms = [
        "主要矛盾和次要矛盾",
        "矛盾的主要方面和次要方面",
        "主流和支流",
        "两点论与重点论",
        "主要矛盾",
        "次要矛盾",
        "矛盾的主要方面",
        "矛盾的次要方面",
        "主流",
        "支流",
    ]
    counts = Counter({term: all_text.count(term) for term in terms})

    lines = [
        "# v7 覆盖验证",
        "",
        f"- 文档：{DOCX}",
        f"- Heading 2 数量：{len(h2)}",
        f"- Heading 3 数量：{len(h3)}",
        f"- V3 候选总数（主观 A/B + 选择 C）：{len(candidates)}",
        f"- 仍未被套卷+题号覆盖：{len(missing)}",
        "",
        "## 新增关键节点存在性",
    ]
    for term in ["主要矛盾和次要矛盾", "矛盾的主要方面和次要方面", "主流和支流", "两点论与重点论"]:
        lines.append(f"- {term}: {'YES' if term in h2 else 'NO'}")
    lines.extend(["", "## 关键词频次"])
    for term in terms:
        lines.append(f"- {term}: {counts[term]}")
    if missing:
        lines.extend(["", "## 仍未覆盖明细"])
        for r in missing:
            lines.append(
                f"- {r['verify_key']} | {r['source']} | {r['question_nature']} | {r['evidence_grade_initial']} | {r['trigger']}"
            )

    OUT.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print("\n".join(lines[:24]))


if __name__ == "__main__":
    main()
