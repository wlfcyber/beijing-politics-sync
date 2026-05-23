from __future__ import annotations

import csv
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt


DESKTOP = Path.home() / "Desktop"
REPO = DESKTOP / "02_代码项目与工具" / "mac-thread-restore" / "beijing-politics-sync-visible"
RUN = REPO / "reports" / "bixiu4_philosophy_full_coverage_double_lane_2026-05-23"
V3_CSV = RUN / "01_codex_lane" / "v3_inventory_vs_latest_docx.csv"
OUT_DIR = RUN / "04_delivery"
OUT_DOCX = DESKTOP / "哲学宝典最终版-飞哥正志讲堂_主次矛盾全覆盖补丁v7_2026-05-23.docx"
OUT_REPORT = RUN / "03_fusion" / "FUSION_PATCH_REPORT_v7.md"


def latest_base_docx() -> Path:
    candidates = sorted(
        DESKTOP.glob("*哲学宝典*二模补题版*.docx"),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    if not candidates:
        raise FileNotFoundError("未找到桌面上的 2026二模补题版 哲学宝典 docx")
    return candidates[0]


def clean_text(text: str, limit: int = 420) -> str:
    text = re.sub(r"\s+", " ", (text or "").strip())
    return text if len(text) <= limit else text[: limit - 1] + "…"


def load_gap_rows():
    rows = list(csv.DictReader(V3_CSV.open("r", encoding="utf-8-sig")))
    subjectives = [
        r
        for r in rows
        if r.get("covered_by_latest_docx") == "False"
        and r.get("question_nature") == "subjective"
        and r.get("evidence_grade_initial") in {"A", "B"}
    ]
    choices = [
        r
        for r in rows
        if r.get("covered_by_latest_docx") == "False"
        and r.get("question_nature") == "choice"
        and r.get("evidence_grade_initial") == "C"
    ]
    return subjectives, choices


def add_para_before(target, text: str = "", style: str | None = None):
    p = target.insert_paragraph_before(text)
    if style:
        p.style = style
    return p


def add_body_before(target, label: str, text: str):
    p = add_para_before(target)
    r1 = p.add_run(label)
    r1.bold = True
    p.add_run(clean_text(text))
    return p


def add_entry_before(target, title: str, material: str, trigger: str, landing: str, note: str | None = None):
    add_para_before(target, title, "Heading 3")
    add_body_before(target, "材料触发点：", material)
    add_body_before(target, "知识触发：", trigger)
    add_body_before(target, "答案落点：", landing)
    if note:
        add_body_before(target, "使用提醒：", note)


def find_heading(doc: Document, text: str, style_name: str = "Heading 2"):
    for p in doc.paragraphs:
        if p.style and p.style.name == style_name and p.text.strip() == text:
            return p
    raise ValueError(f"找不到标题：{style_name} {text}")


def first_heading_after(doc: Document, heading_para, style_name: str = "Heading 3"):
    seen = False
    for p in doc.paragraphs:
        if p._p is heading_para._p:
            seen = True
            continue
        if seen and p.style and p.style.name == style_name:
            return p
    raise ValueError(f"找不到 {heading_para.text} 之后的 {style_name}")


def source_title(row: dict) -> str:
    year = row.get("norm_year", "")
    district = row.get("norm_district", "")
    stage = row.get("norm_stage", "")
    question = row.get("norm_question", "")
    qtype = "主观题" if row.get("question_nature") == "subjective" else "选择题"
    return f"{year}{district}{stage} 第{question}题（{qtype}）"


def append_gap_index(doc: Document, subjectives: list[dict], choices: list[dict]):
    doc.add_page_break()
    doc.add_heading("六、本次补入题目索引", level=1)
    p = doc.add_paragraph(
        "这一章只做补漏索引：正文仍按前面的哲学框架复习；这里把扫描中没有被同套卷同题号覆盖的题目放进来，避免备课时只反复看到少数题。"
    )
    p.runs[0].bold = True

    doc.add_heading("主观题补漏索引", level=2)
    for i, r in enumerate(subjectives, 1):
        doc.add_heading(f"补主观{i}. {source_title(r)}", level=3)
        add_labeled(doc, "材料触发点：", r.get("material", ""))
        add_labeled(doc, "知识触发：", r.get("trigger", ""))
        add_labeled(doc, "答案落点：", r.get("logic", ""))
        if "文化" in r.get("trigger", "") and "哲学" not in r.get("trigger", ""):
            add_labeled(doc, "使用提醒：", "这类题更适合在文化宝典做主条，哲学宝典中只保留轻量触发，避免把文化题硬塞进哲学。")

    doc.add_heading("哲学选择题补漏索引", level=2)
    doc.add_paragraph(
        "选择题按“正确项对应/错误项辨析”方式处理；这里先保证每个哲学触发点进入宝典索引，完整选项仍以原卷为准。"
    )
    for i, r in enumerate(choices, 1):
        doc.add_heading(f"补选择{i}. {source_title(r)}", level=3)
        add_labeled(doc, "题干或选项触发：", r.get("material", ""))
        add_labeled(doc, "对应知识：", r.get("trigger", ""))
        add_labeled(doc, "判断逻辑：", r.get("logic", ""))


def add_labeled(doc: Document, label: str, text: str):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    p.add_run(clean_text(text))
    return p


def patch_docx():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    OUT_REPORT.parent.mkdir(parents=True, exist_ok=True)

    base = latest_base_docx()
    shutil.copy2(base, OUT_DOCX)
    doc = Document(str(OUT_DOCX))

    # Insert the missing contradiction nodes at their natural place in dialectics:
    # after "矛盾的普遍性和特殊性" and before the existing "两点论与重点论" section.
    two_points_heading = find_heading(doc, "两点论与重点论", "Heading 2")

    add_para_before(two_points_heading, "主要矛盾和次要矛盾", "Heading 2")
    add_body_before(
        two_points_heading,
        "原理方法论：",
        "主要矛盾在事物发展过程中处于支配地位，对事物发展起决定作用；次要矛盾处于从属地位。办事情要善于抓重点，集中力量解决主要矛盾，同时统筹兼顾，恰当处理次要矛盾。",
    )
    add_body_before(
        two_points_heading,
        "快速判别：",
        "题目问“推进工作、解决问题、抓核心领域、战略全局中的重点”时，优先想到主要矛盾；但题目问“评价事物性质、看利弊主流”时，不要误写主要矛盾，要转到矛盾主要方面。",
    )
    add_entry_before(
        two_points_heading,
        "1. 2025东城一模 第18题第（1）问（主观题）",
        "“两重”建设既强化硬投资，又抓软建设，材料把重点落在新质生产力发展等核心领域。",
        "抓主要矛盾；也可替换为坚持两点论和重点论相统一。",
        "答题要写出：在复杂任务中抓住影响发展全局的核心领域，集中力量推进重大战略实施和重点领域安全能力建设，同时不把其他建设割裂开来。",
        "这题不能只写“矛盾分析法”，细则提醒只笼统写矛盾分析法且材料分析准确，只能拿本层次低分。",
    )
    add_entry_before(
        two_points_heading,
        "2. 2024东城一模 第21题（主观题）",
        "京津冀功能圈、产业圈、交通圈协同发展，功能定位和区域协调需要分清重点与整体。",
        "主次矛盾；两点论；整体与部分。",
        "功能圈答题可从主次矛盾切入：协同发展不是平均用力，而是抓住功能定位、区域协调中的重点问题，再带动产业圈、交通圈等整体联动。",
    )
    add_entry_before(
        two_points_heading,
        "3. 2024海淀一模 第17题第（2）问（主观题）",
        "京津冀三地从协同发展的战略大局出发，分别分析各自功能定位和特色优势。",
        "抓主要矛盾；分析与综合；系统观念。",
        "答题要落到：三地不是各讲各的优势，而是在战略大局中抓住制约协同发展的关键矛盾，科学定位各自功能，再整体把握区域差异，推动优势互补。",
    )
    add_entry_before(
        two_points_heading,
        "4. 2026丰台二模 第16题（主观题）",
        "评标材料在哲学角度中明列“主要矛盾、系统观念、联系、创新、生产力”。",
        "主要矛盾；系统观念；联系观点。",
        "答题时把主要矛盾写成“在国内国际复杂条件中抓住最能牵动发展全局的关键问题”，再用系统观念说明多领域协同推进。",
    )
    add_entry_before(
        two_points_heading,
        "5. 2026顺义一模 第21题（主观题）",
        "材料要求分析中国共产党洞察时与势、辨析危与机，统筹中华民族伟大复兴战略全局和世界百年未有之大变局。",
        "主次矛盾和矛盾主次方面；两点论和重点论相结合。",
        "答题要写出：面对“危”与“机”、国内全局与国际变局，既全面看问题，又抓住推动复兴进程的关键矛盾和主要方面，沉着应对风险挑战。",
    )
    add_entry_before(
        two_points_heading,
        "6. 选择题易错：把“主要矛盾”当万能钥匙",
        "常见错项会把课间休息、争论、文化交融、生态保护等局部议题硬说成主要矛盾。",
        "主要矛盾误用；矛盾主次方面偷换。",
        "判断选择题时先问：材料有没有呈现一组矛盾体系，并要求抓决定全局的关键问题？没有，就不要套“主要矛盾”。若题干是在评价事物性质，则应考虑“主要矛盾的主要方面”。",
        "典型错项包括“重点抓住和解决次要矛盾”“事物性质由主要矛盾决定”“生态保护成为社会主要矛盾的主要方面”等。",
    )

    add_para_before(two_points_heading, "矛盾的主要方面和次要方面", "Heading 2")
    add_body_before(
        two_points_heading,
        "原理方法论：",
        "事物的性质主要是由主要矛盾的主要方面决定的；矛盾的主、次方面在一定条件下可以相互转化。看问题要分清主流和支流，既不能忽视支流，又要着重把握主流。",
    )
    add_body_before(
        two_points_heading,
        "快速判别：",
        "题目让评价一个事物“本质上怎样、总体上该怎么看、利弊哪个是主流”时，优先想到矛盾主要方面；不要把“关键环节/重点工作”误写成矛盾主要方面。",
    )
    add_entry_before(
        two_points_heading,
        "1. 2026东城一模 第16题（主观题）",
        "文艺作品既要回应观众情绪需求，又不能失去思想性、艺术性。",
        "矛盾的主次方面辩证关系；抓主流。",
        "答题要写出：文艺作品的思想性、艺术性是矛盾的主要方面，情绪需求是次要方面；可以回应情绪，但不能让次要方面替代作品本质。",
    )
    add_entry_before(
        two_points_heading,
        "2. 2025西城一模 第17题（主观题）",
        "“投资于人”和“投资于物”存在对立，资源有限条件下必须判断现代化发展的方向。",
        "矛盾的主要方面；以人民为中心。",
        "答题要写出：“投资于人”体现现代化的方向和人的发展目的，是这组对立中应当把握的主要方面；因此更多资金资源投向人，才是以人民为中心的生动体现。",
    )
    add_entry_before(
        two_points_heading,
        "3. 2025丰台期末 第16题（主观题）",
        "材料强调懂得注意基本统计、主要百分比，做到胸中有“数”。",
        "主要矛盾的主要方面；度；数量界限。",
        "答题要写出：事物的性质由主要矛盾的主要方面决定，把握主要百分比才能看清性质和趋势，避免胸中无“数”。",
        "细则限制很重要：在论证“数为何所用”时，如果从“看大数、看长远”硬引出主要矛盾或矛盾主次方面，但没有回应设问维度，不给分。",
    )
    add_entry_before(
        two_points_heading,
        "4. 2026丰台一模 第21题（主观题）",
        "人工智能既为文明进步注入动能，也带来风险与隐忧。",
        "矛盾的主要方面和次要方面；主流和支流；两点论和重点论。",
        "答题要写出：AI 的积极作用是主流，风险与隐忧是支流；要抓住主流，用好 AI 促进发展，同时不能忽视支流，要加强规范引导。",
    )
    add_entry_before(
        two_points_heading,
        "5. 选择题易错：把“关键环节”错当“矛盾主要方面”",
        "选择题常把“转体”“独特创意”“智能养老普及”等材料词包装成矛盾主要方面。",
        "矛盾主要方面误用；性质判断与发展关键的区别。",
        "判断时先问：材料是否要求判断事物性质，且明确哪一方面起主导作用？如果只是关键工序、创新条件、普及现象，一般不能直接说成矛盾主要方面。",
    )

    add_para_before(two_points_heading, "主流和支流", "Heading 2")
    add_body_before(
        two_points_heading,
        "原理方法论：",
        "主流、支流本质上是矛盾主要方面和次要方面的表达。看问题要抓主流以把握事物性质，同时看到支流，防止小问题扩大化。",
    )
    add_entry_before(
        two_points_heading,
        "1. 2026丰台一模 第21题（主观题）",
        "人工智能的积极作用与风险隐忧并存。",
        "主流和支流；两点论与重点论。",
        "积极作用是主流，风险隐忧是支流。答案不能只写风险治理，也不能只写技术乐观，而要写“抓主流、看支流、用规范保障发展”。",
    )
    add_entry_before(
        two_points_heading,
        "2. 2026顺义二模 第16题（主观题）",
        "新大众文艺内容多元、形式多样，但不能多元无序、价值失向。",
        "主流价值；两点论与重点论；矛盾普遍性和特殊性。",
        "答题要写出：尊重多样性符合矛盾特殊性要求，但必须塑造主流舆论新格局、弘扬主旋律、传播正能量，不能让流量化、低俗化成为方向。",
    )
    add_entry_before(
        two_points_heading,
        "3. 选择题提醒：题干里的“主流”不一定是哲学主流",
        "有些题干出现“主流媒体”“主流消费模式”“主流商城”等词。",
        "语词主流与哲学主流的区别。",
        "如果题目只是日常语义的“主流”，没有呈现矛盾双方及事物性质判断，就不要硬套矛盾主要方面；只有当正确项明确要求抓主流、分清利弊主次时，才进入本节点。",
    )

    # Add stronger entries into the existing "两点论与重点论" section.
    two_points_first_entry = first_heading_after(doc, two_points_heading, "Heading 3")
    add_body_before(
        two_points_first_entry,
        "本节补强：",
        "两点论不是“两个方面都写一点”这么简单；重点论也不是“只写重点”。北京题常把它落实为：既看到矛盾双方，又抓住主要矛盾或主要方面。",
    )
    add_entry_before(
        two_points_first_entry,
        "补强1. 2024朝阳一模 第18题第（2）问（主观题）",
        "科学普及与科技创新各有侧重又具有同一性，现实中存在“重科研、轻科普”的短板。",
        "两点论和重点论相统一；系统优化；具体问题具体分析。",
        "答题要写出：既要看到科普和科技创新两翼都重要，又要针对“重科研、轻科普”的短板，把人才队伍和基层科普服务作为重点补强，推动两翼齐飞。",
    )
    add_entry_before(
        two_points_first_entry,
        "补强2. 2026丰台一模 第21题（主观题）",
        "AI 发展带来积极作用，也带来风险与隐忧。",
        "两点论与重点论；主流和支流。",
        "答案要既看到积极作用和风险挑战，又抓住积极作用这个主流，把规范治理作为保障发展而不是否定发展。",
    )
    add_entry_before(
        two_points_first_entry,
        "补强3. 2025东城一模 第18题第（1）问（主观题）",
        "“两重”建设要聚焦新质生产力发展等核心领域，同时推进重大战略实施和重点领域安全能力建设。",
        "抓主要矛盾；两点论和重点论相统一。",
        "答案要写出：复杂建设任务不能平均用力，要抓核心领域和关键问题；同时不能把其他配套建设丢掉。",
    )
    add_entry_before(
        two_points_first_entry,
        "补强4. 2026顺义二模 第16题（主观题）",
        "新大众文艺多元发展，同时需要主流价值引领。",
        "矛盾普遍性和特殊性；两点论和重点论；价值观导向。",
        "答案要写出：多样性不等于无序，多元表达必须由主流价值引领，做到尊重差异和弘扬主旋律统一。",
    )
    add_entry_before(
        two_points_first_entry,
        "补强5. 2026房山一模 第16题第（2）问（主观题）",
        "治理路径材料允许从矛盾分析法、具体问题具体分析、两点论与重点论统一等角度作答。",
        "具体问题具体分析；两点论和重点论。",
        "答案要落到治理措施：既全面分析治理对象的多方面矛盾，又抓住最影响治理成效的关键环节，分类施策、重点突破。",
        "只有写出材料中的治理对象和措施，才算真正落到两点论重点论；空写原理不能替代答案。",
    )

    subjectives, choices = load_gap_rows()
    append_gap_index(doc, subjectives, choices)

    # Apply a light touch to newly added text through existing Word styles. Do not redesign the accepted handbook.
    for p in doc.paragraphs:
        if p.style and p.style.name == "Normal":
            p.paragraph_format.space_after = Pt(3)

    doc.save(str(OUT_DOCX))

    report = [
        "# v7 融合补丁报告",
        "",
        f"- 底稿：{base}",
        f"- 输出：{OUT_DOCX}",
        "- 处理方式：保留认可版宝典主体，只在辩证法原位置新增“主要矛盾和次要矛盾 / 矛盾的主要方面和次要方面 / 主流和支流”，并补强原有“两点论与重点论”节点。",
        f"- 主观题补漏索引：{len(subjectives)} 条",
        f"- 哲学选择题补漏索引：{len(choices)} 条",
        "- 证据口径：正文只写细则/教师版/评标材料已经出现或允许的角度；文化-only 条目在索引中提示轻量挂载或转文化线。",
    ]
    OUT_REPORT.write_text("\n".join(report) + "\n", encoding="utf-8")
    print(OUT_DOCX)
    print(OUT_REPORT)


if __name__ == "__main__":
    patch_docx()
