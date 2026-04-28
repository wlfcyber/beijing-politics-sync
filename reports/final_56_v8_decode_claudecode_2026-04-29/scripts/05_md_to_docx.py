"""V8 decode-version Markdown -> DOCX converter.

Translates the student-facing MDs into Word using a deliberately minimal
mapping: # / ## / ### / #### are mapped to Heading 1..4; --- creates a page
break; bold runs come from **...**; everything else is body text. The first
H1 is rendered as a centered cover title; the next "## 序" stays on its own
page.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt

sys.stdout.reconfigure(encoding="utf-8")

ROOT = Path(r"C:/bp_sync_visible/reports/final_56_v8_decode_claudecode_2026-04-29")
PAIRS = [
    ("2026北京高考政治哲学宝典---三年模拟全触发全链条_v8_decode版_学生版.md",
     "2026北京高考政治哲学宝典---三年模拟全触发全链条_v8_decode版_学生版.docx"),
    ("北京高考政治选择题错肢总结_v8_decode版.md",
     "北京高考政治选择题错肢总结_v8_decode版.docx"),
]

H_RE = re.compile(r"^(#{1,6})\s+(.*?)\s*$")
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def add_paragraph(doc: Document, line: str, bold: bool = False, size: int | None = None, align: int | None = None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    cursor = 0
    for m in BOLD_RE.finditer(line):
        if m.start() > cursor:
            run = p.add_run(line[cursor:m.start()])
            if bold:
                run.bold = True
            if size:
                run.font.size = Pt(size)
        run = p.add_run(m.group(1))
        run.bold = True
        if size:
            run.font.size = Pt(size)
        cursor = m.end()
    if cursor < len(line):
        run = p.add_run(line[cursor:])
        if bold:
            run.bold = True
        if size:
            run.font.size = Pt(size)


def convert(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    seen_first_h1 = False
    seen_xu = False

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph()
            continue
        if line.strip() == "---":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            continue
        m = H_RE.match(line)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2)
            heading_text = re.sub(r"\*\*", "", heading_text)
            if level == 1 and not seen_first_h1:
                seen_first_h1 = True
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(heading_text)
                run.bold = True
                run.font.size = Pt(28)
                continue
            if level == 2 and "序" in heading_text and not seen_xu:
                seen_xu = True
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            try:
                doc.add_heading(heading_text, level=min(level, 4))
            except Exception:
                add_paragraph(doc, heading_text, bold=True, size=14)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            cursor = 0
            text_after = line[2:]
            for m in BOLD_RE.finditer(text_after):
                if m.start() > cursor:
                    p.add_run(text_after[cursor:m.start()])
                run = p.add_run(m.group(1))
                run.bold = True
                cursor = m.end()
            if cursor < len(text_after):
                p.add_run(text_after[cursor:])
            continue
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(re.match(r"^-{2,}$", c.replace(":", "")) for c in cells):
                continue
            p = doc.add_paragraph()
            p.add_run(" | ".join(cells))
            continue
        add_paragraph(doc, line)

    doc.save(docx_path)
    print(f"wrote {docx_path}; size={docx_path.stat().st_size}")


def main() -> None:
    for md_name, docx_name in PAIRS:
        md_path = ROOT / "outputs" / md_name
        docx_path = ROOT / "outputs" / docx_name
        if not md_path.exists():
            print(f"SKIP missing {md_path}")
            continue
        convert(md_path, docx_path)


if __name__ == "__main__":
    main()
