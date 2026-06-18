#!/usr/bin/env python3
from __future__ import annotations

import csv
import json
import re
import subprocess
import unicodedata
import zipfile
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET

import docx
import pdfplumber


RUN_DIR = Path(__file__).resolve().parents[1]
CACHE_ROOT = Path("/Users/wanglifei/GaokaoPolitics/beijing-politics-sync/data/preprocessed_corpus")
INVENTORY = RUN_DIR / "01_source_inventory" / "source_inventory.csv"

WINDOWS_CACHE_ROOT = r"C:\Users\Administrator\Desktop\beijing_politics_research\data\preprocessed_corpus"
OCR_CACHE_ROOTS = [
    Path("/Users/wanglifei/GaokaoPolitics/beijing-politics-sync"),
    Path("/Users/wanglifei/Desktop/北京高考政治"),
]


TARGET_KEYWORDS = {
    "B2_ECONOMICS": [
        "经济与社会",
        "基本经济制度",
        "公有制",
        "非公有制",
        "国有经济",
        "民营经济",
        "民营企业",
        "市场经济",
        "市场机制",
        "有效市场",
        "有为政府",
        "宏观调控",
        "经济职能",
        "新发展理念",
        "高质量发展",
        "现代化经济体系",
        "实体经济",
        "数字经济",
        "新质生产力",
        "供给侧",
        "乡村振兴",
        "共同富裕",
        "收入分配",
        "按劳分配",
        "社会保障",
        "财政",
        "税收",
        "消费",
        "产业链",
        "供应链",
        "扩大内需",
        "营商环境",
        "银发经济",
        "医保",
        "养老服务",
        "低保",
        "社会救助",
        "就业",
        "劳动者",
        "人力资源",
        "市场主体",
        "生产要素",
        "数据要素",
        "居民收入",
        "民生服务",
        "产业发展",
        "区域协调",
        "基础设施",
        "经济增长",
        "经济总量",
        "GDP",
    ],
    "B3_POLITICS_RULE_OF_LAW": [
        "政治与法治",
        "中国共产党",
        "党的领导",
        "全面从严治党",
        "人民当家作主",
        "人民代表大会",
        "人大代表",
        "政协",
        "协商民主",
        "基层群众自治",
        "村民自治",
        "居民自治",
        "全过程人民民主",
        "民族区域自治",
        "宗教政策",
        "法治政府",
        "依法行政",
        "全面依法治国",
        "法治国家",
        "法治社会",
        "严格执法",
        "公正司法",
        "全民守法",
        "宪法",
        "国家治理",
        "行政机关",
        "政府",
        "检察",
        "法院",
        "司法机关",
        "民主监督",
        "民主决策",
        "民主协商",
        "全国人大",
        "全国人大常委会",
        "人大常委会",
        "立法权",
        "科学立法",
        "行政许可",
        "行政处罚",
        "党组织",
        "基层治理",
        "社区治理",
        "社会治理",
        "公共服务",
        "人民政协",
        "模拟政协",
    ],
    "B4_CULTURE": [
        "文化",
        "文化功能",
        "文化载体",
        "中华优秀传统文化",
        "优秀传统文化",
        "传统文化",
        "创造性转化",
        "创新性发展",
        "文化传承",
        "文化创新",
        "革命文化",
        "社会主义先进文化",
        "文化自信",
        "文化强国",
        "民族精神",
        "爱国主义",
        "社会主义核心价值观",
        "文明交流互鉴",
        "文化遗产",
        "非遗",
        "博物馆",
        "大运河",
        "中轴线",
        "文艺",
        "文旅",
        "文化产业",
        "文化事业",
        "精神家园",
        "文化基因",
        "精神谱系",
        "红色基因",
        "长征精神",
        "革命精神",
        "时代精神",
        "传统色",
        "纹样",
        "中华文明",
        "文明传承",
        "文物",
        "古籍",
        "文创",
        "文化软实力",
        "中式生活",
    ],
}

BOUNDARY_KEYWORDS = {
    "B4_PHILOSOPHY_EXCLUDED": [
        "实践",
        "认识",
        "真理",
        "联系的观点",
        "联系的普遍性",
        "整体与部分",
        "系统优化",
        "发展的观点",
        "事物发展",
        "矛盾",
        "辩证否定",
        "量变",
        "质变",
        "社会存在",
        "社会意识",
        "人民群众",
        "价值观的导向",
        "价值判断",
        "价值选择",
        "人生价值",
        "主观能动性",
        "物质决定意识",
        "意识的能动作用",
        "尊重客观规律",
        "历史唯物主义",
        "意识活动",
        "自觉选择性",
        "能动创造性",
        "量的渐进性",
        "质的飞跃",
        "质的连续性",
        "适度原则",
    ],
    "XB1_EXCLUDED": [
        "当代国际政治与经济",
        "国际关系",
        "国家利益",
        "中国外交",
        "联合国",
        "全球治理",
        "世界多极化",
        "经济全球化",
        "和平与发展",
        "人类命运共同体",
        "国际组织",
        "主权国家",
        "外交",
        "元首外交",
        "多边外交",
        "APEC",
        "G20",
        "金砖",
        "上合组织",
        "南南合作",
        "对外开放",
        "国际贸易",
        "国际投资",
        "全球价值链",
        "贸易强国",
    ],
    "XB2_EXCLUDED": [
        "法律与生活",
        "民法典",
        "合同",
        "侵权",
        "知识产权",
        "劳动合同",
        "劳动关系",
        "继承",
        "诉讼",
        "调解",
        "仲裁",
        "消费者权益",
        "格式条款",
        "相邻关系",
        "人格权",
        "物权",
        "司法确认",
        "举证",
        "民事",
        "民事主体",
        "民事行为能力",
        "民事诉讼",
        "侵权责任",
        "赔偿",
        "违约",
        "专利",
        "专利法",
        "所有权",
        "健康权",
        "监护人",
        "法定监护人",
        "辩护人",
        "上诉",
    ],
    "XB3_EXCLUDED": [
        "逻辑与思维",
        "科学思维",
        "辩证思维",
        "创新思维",
        "超前思维",
        "发散思维",
        "聚合思维",
        "逆向思维",
        "联想思维",
        "迁移",
        "想象",
        "分析与综合",
        "推理",
        "三段论",
        "概念外延",
        "正确运用判断",
        "假言判断",
        "性质判断",
        "关系判断",
        "周延",
        "属种关系",
        "类比推理",
        "演绎推理",
        "归纳推理",
        "推论",
        "非对称性关系判断",
        "充分必要条件",
        "必要条件",
        "充分条件",
        "逻辑规则",
    ],
    "B1_EXCLUDED": [
        "中国特色社会主义",
        "科学社会主义",
        "改革开放",
        "中国梦",
        "社会主义现代化强国",
    ],
}

BOOST_KEYWORDS = {
    "B2_ECONOMICS": [
        "经济与社会",
        "基本经济制度",
        "公有制",
        "非公有制",
        "国有经济",
        "民营经济",
        "市场经济",
        "宏观调控",
        "经济职能",
        "新发展理念",
        "高质量发展",
        "现代化经济体系",
        "实体经济",
        "数字经济",
        "新质生产力",
        "乡村振兴",
        "共同富裕",
        "收入分配",
        "社会保障",
        "财政",
        "税收",
        "产业链",
        "供应链",
        "扩大内需",
        "营商环境",
        "银发经济",
        "医保",
        "养老服务",
        "低保",
        "社会救助",
        "就业",
        "生产要素",
        "数据要素",
        "经济增长",
        "经济总量",
        "GDP",
    ],
    "B3_POLITICS_RULE_OF_LAW": [
        "政治与法治",
        "中国共产党",
        "党的领导",
        "人民代表大会",
        "人大代表",
        "政协",
        "协商民主",
        "基层群众自治",
        "全过程人民民主",
        "民族区域自治",
        "法治政府",
        "依法行政",
        "全面依法治国",
        "法治国家",
        "法治社会",
        "严格执法",
        "公正司法",
        "全民守法",
        "宪法",
        "国家治理",
        "全国人大",
        "人大常委会",
        "科学立法",
        "党组织",
        "基层治理",
        "人民政协",
        "模拟政协",
    ],
    "B4_CULTURE": [
        "中华优秀传统文化",
        "优秀传统文化",
        "传统文化",
        "创造性转化",
        "创新性发展",
        "文化传承",
        "文化创新",
        "革命文化",
        "社会主义先进文化",
        "文化自信",
        "文化强国",
        "民族精神",
        "爱国主义",
        "社会主义核心价值观",
        "文明交流互鉴",
        "文化遗产",
        "非遗",
        "博物馆",
        "大运河",
        "中轴线",
        "文旅",
        "精神家园",
        "文化基因",
        "精神谱系",
        "红色基因",
        "长征精神",
        "革命精神",
        "时代精神",
        "传统色",
        "纹样",
        "中华文明",
        "文明传承",
        "文物",
        "古籍",
        "文创",
        "文化软实力",
        "中式生活",
    ],
    "B4_PHILOSOPHY_EXCLUDED": [
        "真理",
        "联系的观点",
        "联系的普遍性",
        "整体与部分",
        "系统优化",
        "发展的观点",
        "矛盾",
        "辩证否定",
        "量变",
        "质变",
        "社会存在",
        "社会意识",
        "人民群众",
        "价值观的导向",
        "价值判断",
        "价值选择",
        "人生价值",
        "物质决定意识",
        "意识的能动作用",
        "尊重客观规律",
        "历史唯物主义",
        "意识活动",
        "自觉选择性",
        "能动创造性",
        "量的渐进性",
        "质的飞跃",
        "质的连续性",
        "适度原则",
    ],
    "XB1_EXCLUDED": [
        "当代国际政治与经济",
        "国际关系",
        "国家利益",
        "中国外交",
        "联合国",
        "全球治理",
        "世界多极化",
        "经济全球化",
        "和平与发展",
        "人类命运共同体",
        "国际组织",
        "主权国家",
        "外交",
        "元首外交",
        "多边外交",
        "APEC",
        "G20",
        "金砖",
        "上合组织",
        "南南合作",
        "国际贸易",
        "国际投资",
        "全球价值链",
        "贸易强国",
    ],
    "XB2_EXCLUDED": [
        "法律与生活",
        "民法典",
        "合同",
        "侵权",
        "知识产权",
        "劳动合同",
        "劳动关系",
        "继承",
        "诉讼",
        "调解",
        "仲裁",
        "消费者权益",
        "格式条款",
        "相邻关系",
        "人格权",
        "物权",
        "司法确认",
        "举证",
        "民事",
        "民事主体",
        "民事行为能力",
        "民事诉讼",
        "侵权责任",
        "违约",
        "专利",
        "专利法",
        "所有权",
        "健康权",
        "监护人",
        "法定监护人",
        "辩护人",
        "上诉",
    ],
    "XB3_EXCLUDED": [
        "逻辑与思维",
        "科学思维",
        "辩证思维",
        "创新思维",
        "超前思维",
        "发散思维",
        "聚合思维",
        "逆向思维",
        "联想思维",
        "迁移",
        "想象",
        "分析与综合",
        "推理",
        "三段论",
        "概念外延",
        "正确运用判断",
        "假言判断",
        "性质判断",
        "关系判断",
        "周延",
        "属种关系",
        "类比推理",
        "演绎推理",
        "归纳推理",
        "推论",
        "非对称性关系判断",
        "充分必要条件",
        "必要条件",
        "充分条件",
        "逻辑规则",
    ],
}

EXPLICIT_TARGET_TERMS = {
    "B2_ECONOMICS": ["经济与社会"],
    "B3_POLITICS_RULE_OF_LAW": ["政治与法治"],
}

EXPLICIT_BOUNDARY_TERMS = {
    "XB1_EXCLUDED": [
        "当代国际政治与经济",
        "经济全球化",
        "世界多极化",
        "国际关系",
        "国家利益",
        "联合国",
        "全球治理",
        "人类命运共同体",
        "国际组织",
    ],
    "XB2_EXCLUDED": [
        "法律与生活",
        "民法典",
        "合同",
        "侵权",
        "知识产权",
        "诉讼",
        "民事",
        "专利",
        "所有权",
        "人格权",
        "物权",
        "劳动合同",
    ],
    "XB3_EXCLUDED": [
        "逻辑与思维",
        "科学思维",
        "辩证思维",
        "创新思维",
        "超前思维",
        "三段论",
        "推理",
        "概念外延",
        "假言",
        "充分条件",
        "必要条件",
        "迁移",
        "想象",
        "分析与综合",
    ],
    "B4_PHILOSOPHY_EXCLUDED": [
        "历史唯物主义",
        "社会存在",
        "社会意识",
        "价值判断",
        "价值选择",
        "整体与部分",
        "系统优化",
        "辩证否定",
        "意识活动",
        "自觉选择性",
        "能动创造性",
        "物质决定意识",
    ],
}

COVERAGE_SOURCE_TYPES = {"paper", "reference-answer"}
DISTRICTS = [
    "海淀",
    "西城",
    "东城",
    "朝阳",
    "丰台",
    "石景山",
    "顺义",
    "房山",
    "延庆",
    "门头沟",
    "昌平",
    "通州",
    "大兴",
    "怀柔",
    "密云",
    "平谷",
    "燕山",
]
STAGES = [("一模", "一模"), ("二模", "二模"), ("期末", "期末"), ("期中", "期中")]


def localize_cache_path(value: str) -> Path | None:
    if not value:
        return None
    value = value.strip().lstrip("\ufeff")
    rel = None
    if value.startswith(WINDOWS_CACHE_ROOT):
        rel = value[len(WINDOWS_CACHE_ROOT) :].lstrip("\\/")
    elif "beijing_politics_research\\data\\preprocessed_corpus\\" in value:
        rel = value.split("beijing_politics_research\\data\\preprocessed_corpus\\", 1)[1]
    elif "preprocessed_corpus/" in value:
        rel = value.split("preprocessed_corpus/", 1)[1]
    if rel is None:
        return Path(value)
    return CACHE_ROOT / rel.replace("\\", "/")


def read_cache_index() -> dict[str, dict[str, str]]:
    out: dict[str, dict[str, str]] = {}
    manifest = CACHE_ROOT / "manifest.csv"
    if not manifest.exists():
        return out
    with manifest.open("r", encoding="utf-8-sig", newline="") as f:
        for row in csv.DictReader(f):
            sha = row.get("sha256", "")
            if sha:
                out[sha] = row
    return out


def infer_suite_id_from_text(text: str, path: Path) -> str:
    # Parent folders often contain run dates such as 2026-05-25; do not let
    # those dates contaminate the exam year inference.
    parent_hints = []
    for part in path.parts[-4:-1]:
        clean = part.replace(" ", "")
        has_year = bool(re.search(r"202[456]", clean))
        has_district = any(district in clean for district in DISTRICTS)
        has_stage = any(needle in clean for needle, _stage in STAGES)
        if has_year and has_district and has_stage:
            parent_hints.append(clean)
    probe = (path.name + "\n" + "\n".join(parent_hints) + "\n" + text[:8000]).replace(" ", "")
    year_match = re.search(r"(202[456])", probe)
    year = year_match.group(1) if year_match else ""
    district_hits = [(probe.find(district), district) for district in DISTRICTS if district in probe]
    stage_hits = [(probe.find(needle), stage) for needle, stage in STAGES if needle in probe]
    district = min(district_hits)[1] if district_hits else ""
    stage = min(stage_hits)[1] if stage_hits else ""
    if not (year and district and stage):
        return ""
    return f"{year}_{district}_{stage}"


def ocr_cache_score(text: str, path: Path) -> int:
    score = len(text)
    name = path.name
    if "第一部分" in text or "本部分共15" in text:
        score += 5000
    if "试卷" in name:
        score += 3000
    if "merged" in name:
        score += 1500
    if any(k in name for k in ["细则", "评标", "阅卷", "答案"]):
        score -= 3000
    return score


def read_existing_ocr_cache() -> dict[str, dict[str, str]]:
    best: dict[str, dict[str, str]] = {}
    for root in OCR_CACHE_ROOTS:
        if not root.exists():
            continue
        for path in root.glob("**/ocr_cache/**/*.ocr.txt"):
            try:
                text = path.read_text(encoding="utf-8", errors="ignore")
            except OSError:
                continue
            if len(text) < 500:
                continue
            suite_id = infer_suite_id_from_text(text, path)
            if not suite_id:
                continue
            score = ocr_cache_score(text, path)
            old = best.get(suite_id)
            if not old or score > int(old["score"]):
                best[suite_id] = {"path": str(path), "text": text, "score": str(score)}
    return best


def read_docx(path: Path) -> str:
    document = docx.Document(path)
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_pdf(path: Path) -> str:
    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for idx, page in enumerate(pdf.pages, start=1):
            text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
            if text.strip():
                parts.append(f"\n[page {idx}]\n{text}")
    return "\n".join(parts)


def read_pptx(path: Path) -> str:
    parts: list[str] = []
    with zipfile.ZipFile(path) as zf:
        slide_names = sorted(
            name for name in zf.namelist() if name.startswith("ppt/slides/slide") and name.endswith(".xml")
        )
        for slide_name in slide_names:
            raw = zf.read(slide_name)
            root = ET.fromstring(raw)
            texts = [node.text for node in root.iter() if node.tag.endswith("}t") and node.text]
            if texts:
                parts.append(f"\n[{slide_name}]\n" + "\n".join(texts))
    return "\n".join(parts)


def read_doc_legacy(path: Path) -> str:
    proc = subprocess.run(
        ["textutil", "-convert", "txt", "-stdout", str(path)],
        check=False,
        capture_output=True,
        text=True,
        timeout=30,
    )
    if proc.returncode != 0:
        raise RuntimeError(proc.stderr.strip() or f"textutil return code {proc.returncode}")
    return proc.stdout


def read_raw(path: Path, file_type: str) -> str:
    if file_type == "pdf":
        return read_pdf(path)
    if file_type == "docx":
        return read_docx(path)
    if file_type == "pptx":
        return read_pptx(path)
    if file_type == "doc":
        return read_doc_legacy(path)
    if file_type in {"txt", "md", "csv"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    return ""


def normalize_text(text: str) -> str:
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\U001001b0", ".")
    text = text.replace("\ua3ac", "，").replace("\ua3ae", "。").replace("\ua3bb", "；")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u3000", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


QUESTION_NO = r"([1-9]|1[0-9]|2[01])"
LINE_START_QUESTION = re.compile(rf"(?m)^\s*(?:第\s*)?{QUESTION_NO}\s*(?:题)?[\.．、]\s*")
INLINE_QUESTION = re.compile(rf"(?<!\d)(?:第\s*)?{QUESTION_NO}\s*(?:题)?[\.．、]\s*")


def split_questions(text: str) -> list[dict[str, str]]:
    text = normalize_text(text)
    if not text:
        return []

    matches = list(LINE_START_QUESTION.finditer(text))
    if len(matches) < 4:
        marked = INLINE_QUESTION.sub(lambda m: f"\n@@Q{m.group(1)}@@ ", text)
        matches = list(re.finditer(r"@@Q([1-9]|1[0-9]|2[01])@@\s*", marked))
        working = marked
    else:
        working = text

    segments = []
    for idx, match in enumerate(matches):
        q = match.group(1)
        start = match.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(working)
        body = working[start:end].replace(f"@@Q{q}@@", "").strip()
        body = re.sub(r"@@Q([1-9]|1[0-9]|2[01])@@", r"\n\1. ", body)
        if len(body) < 8:
            continue
        if re.fullmatch(r"[ABCD答案\s]+", body[:80]):
            continue
        segments.append({"question": q, "text": body})

    return segments


def compact_snippet(text: str, limit: int = 1200) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    return text[:limit]


def score_keywords(text: str, keywords: Iterable[str]) -> tuple[int, list[str]]:
    matched = []
    score = 0
    for kw in keywords:
        count = text.count(kw)
        if count:
            score += count
            matched.append(kw)
    return score, matched


def classify_text(text: str, file_path: str, source_type: str) -> dict[str, str]:
    scores: dict[str, int] = {}
    matches: dict[str, list[str]] = {}

    for module, keywords in TARGET_KEYWORDS.items():
        score, matched = score_keywords(text, keywords)
        scores[module] = score
        matches[module] = matched

    for module, keywords in BOUNDARY_KEYWORDS.items():
        score, matched = score_keywords(text, keywords)
        scores[module] = score
        matches[module] = matched

    for module, keywords in BOOST_KEYWORDS.items():
        for kw in keywords:
            if kw in text:
                scores[module] = scores.get(module, 0) + 1
                matches.setdefault(module, []).append(f"boost:{kw}")

    target_modules = list(TARGET_KEYWORDS)
    boundary_modules = list(BOUNDARY_KEYWORDS)
    target = max(target_modules, key=lambda k: scores[k])
    boundary = max(boundary_modules, key=lambda k: scores[k])
    target_score = scores[target]
    boundary_score = scores[boundary]

    all_matches = {
        key: ",".join(dict.fromkeys(value[:12]))
        for key, value in matches.items()
        if value
    }

    explicit_targets = {
        module: [term for term in terms if term in text]
        for module, terms in EXPLICIT_TARGET_TERMS.items()
    }
    explicit_targets = {module: terms for module, terms in explicit_targets.items() if terms}
    explicit_boundaries = {
        module: [term for term in terms if term in text]
        for module, terms in EXPLICIT_BOUNDARY_TERMS.items()
    }
    explicit_boundaries = {module: terms for module, terms in explicit_boundaries.items() if terms}

    if explicit_targets and explicit_boundaries:
        return {
            "book_module": "UNKNOWN_OR_MIXED",
            "status": "blocked",
            "decision_reason": json.dumps(
                {
                    "target": target,
                    "target_score": target_score,
                    "boundary": boundary,
                    "boundary_score": boundary_score,
                    "matches": all_matches,
                    "explicit_target_modules": explicit_targets,
                    "explicit_boundary_modules": explicit_boundaries,
                    "override": "manual_subquestion_split_needed",
                },
                ensure_ascii=False,
            ),
            "next_action": "manual_subquestion_split_needed",
        }
    if explicit_boundaries:
        explicit_boundary = max(explicit_boundaries, key=lambda module: (scores.get(module, 0), len(explicit_boundaries[module])))
        return {
            "book_module": explicit_boundary,
            "status": "module-boundary-excluded",
            "decision_reason": json.dumps(
                {
                    "target": target,
                    "target_score": target_score,
                    "boundary": boundary,
                    "boundary_score": boundary_score,
                    "matches": all_matches,
                    "explicit_boundary_modules": explicit_boundaries,
                    "override": "explicit_boundary_term",
                },
                ensure_ascii=False,
            ),
            "next_action": "exclude_from_three_target_lines",
        }

    if target_score == 0 and boundary_score == 0:
        module = "UNKNOWN_OR_MIXED"
        status = "blocked"
        next_action = "manual_review_needed"
    elif boundary_score >= target_score + 2 and boundary_score >= 2:
        module = boundary
        status = "module-boundary-excluded"
        next_action = "exclude_from_three_target_lines"
    elif target_score >= 2 and target_score >= boundary_score + 1:
        module = target
        status = "included"
        next_action = "future_baodian_intake"
    elif source_type == "module-classification" and target_score >= 1:
        module = target
        status = "reference-only"
        next_action = "use_as_helper_not_coverage_evidence"
    else:
        module = "UNKNOWN_OR_MIXED"
        status = "blocked"
        next_action = "manual_review_needed"

    return {
        "book_module": module,
        "status": status,
        "decision_reason": json.dumps(
            {
                "target": target,
                "target_score": target_score,
                "boundary": boundary,
                "boundary_score": boundary_score,
                "matches": all_matches,
            },
            ensure_ascii=False,
        ),
        "next_action": next_action,
    }


def question_type(question: str) -> str:
    if question.isdigit():
        n = int(question)
        if n <= 15:
            return "objective"
        return "subjective"
    return "unknown"


def main() -> None:
    cache_index = read_cache_index()
    ocr_cache_index = read_existing_ocr_cache()
    text_dir = RUN_DIR / "02_text_cache" / "texts"
    text_dir.mkdir(parents=True, exist_ok=True)
    ocr_absorb_dir = RUN_DIR / "02_text_cache" / "ocr_absorbed"
    ocr_absorb_dir.mkdir(parents=True, exist_ok=True)
    for old_path in ocr_absorb_dir.glob("*.txt"):
        old_path.unlink()

    with INVENTORY.open("r", encoding="utf-8", newline="") as f:
        inventory_rows = list(csv.DictReader(f))

    canonical_rows = [
        row
        for row in inventory_rows
        if row["status"] != "duplicate-or-drift"
    ]

    extraction_rows = []
    candidate_rows = []

    for row in canonical_rows:
        sha = row["sha256"]
        text_path = text_dir / f"{sha[:16]}.txt" if sha else text_dir / f"nohash_{len(extraction_rows)}.txt"
        extraction_status = "not-attempted"
        cache_hit = "no"
        source_text_path = ""
        error = ""
        text = ""

        if row["status"] == "module-boundary-excluded":
            extraction_status = "skipped-excluded"
        else:
            cache_row = cache_index.get(sha, {})
            for cache_field in ("text_path", "gpt_markdown_path"):
                candidate = localize_cache_path(cache_row.get(cache_field, ""))
                if candidate and candidate.exists():
                    try:
                        text = candidate.read_text(encoding="utf-8", errors="ignore")
                        extraction_status = "cache-hit"
                        cache_hit = "yes"
                        source_text_path = str(candidate)
                        break
                    except OSError as exc:
                        error = f"cache_read_error={exc}"
            if not text:
                try:
                    text = read_raw(Path(row["file_path"]), row["file_type"])
                    extraction_status = "raw-extracted" if text.strip() else "empty-or-unsupported"
                    source_text_path = row["file_path"]
                except Exception as exc:  # noqa: BLE001
                    extraction_status = "failed"
                    error = str(exc)

        text = normalize_text(text)
        if text:
            text_path.write_text(text, encoding="utf-8")
        else:
            text_path = Path("")

        extraction_rows.append(
            {
                "suite_id": row["suite_id"],
                "year": row["year"],
                "district": row["district"],
                "stage": row["stage"],
                "source_type": row["source_type"],
                "file_path": row["file_path"],
                "sha256": sha,
                "cache_hit": cache_hit,
                "extraction_status": extraction_status,
                "char_count": len(text),
                "run_text_path": str(text_path),
                "source_text_path": source_text_path,
                "error": error,
            }
        )

        segments = split_questions(text)
        if row["source_type"] == "paper" and not segments and len(text) > 1500:
            segments = [{"question": "unsegmented", "text": text}]

        for seg in segments:
            candidate_rows.append(
                {
                    "suite_id": row["suite_id"],
                    "year": row["year"],
                    "district": row["district"],
                    "stage": row["stage"],
                    "question": seg["question"],
                    "question_type": question_type(seg["question"]),
                    "source_type": row["source_type"],
                    "file_path": row["file_path"],
                    "sha256": sha,
                    "run_text_path": str(text_path),
                    "snippet": compact_snippet(seg["text"]),
                    "segment_char_count": len(seg["text"]),
                }
            )

    cache_manifest_path = RUN_DIR / "02_text_cache" / "cache_manifest.csv"
    cache_fields = [
        "suite_id",
        "year",
        "district",
        "stage",
        "source_type",
        "file_path",
        "sha256",
        "cache_hit",
        "extraction_status",
        "char_count",
        "run_text_path",
        "source_text_path",
        "error",
    ]
    with cache_manifest_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=cache_fields)
        writer.writeheader()
        writer.writerows(extraction_rows)

    candidates_path = RUN_DIR / "03_question_index" / "question_candidates.csv"
    candidate_fields = [
        "suite_id",
        "year",
        "district",
        "stage",
        "question",
        "question_type",
        "source_type",
        "file_path",
        "sha256",
        "run_text_path",
        "snippet",
        "segment_char_count",
    ]
    with candidates_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=candidate_fields)
        writer.writeheader()
        writer.writerows(candidate_rows)

    grouped: dict[tuple[str, str], list[dict[str, str]]] = defaultdict(list)
    for row in candidate_rows:
        if "unknown_district" in row["suite_id"] or row["question"] == "unsegmented":
            continue
        if row["source_type"] not in COVERAGE_SOURCE_TYPES:
            continue
        grouped[(row["suite_id"], row["question"])].append(row)

    matrix_rows = []
    for (suite_id, question), rows in sorted(grouped.items()):
        paper_rows = [row for row in rows if row["source_type"] == "paper"]
        classification_rows = paper_rows or rows
        combined = "\n".join(row["snippet"] for row in classification_rows)
        primary = paper_rows[0] if paper_rows else rows[0]
        classification = classify_text(combined, primary["file_path"], primary["source_type"])
        source_types = ",".join(sorted({r["source_type"] for r in rows}))
        if classification["status"] == "included" and "paper" not in {r["source_type"] for r in rows}:
            classification["status"] = "reference-only"
            classification["next_action"] = "locate_paper_for_confirmation"
        evidence_paths = " | ".join(dict.fromkeys(r["file_path"] for r in rows[:4]))
        matrix_rows.append(
            {
                "suite_id": suite_id,
                "year": primary["year"],
                "district": primary["district"],
                "stage": primary["stage"],
                "question": question,
                "book_module": classification["book_module"],
                "question_type": question_type(question),
                "evidence_source": evidence_paths,
                "source_types": source_types,
                "status": classification["status"],
                "artifact_location": str(candidates_path),
                "decision_reason": classification["decision_reason"],
                "next_action": classification["next_action"],
            }
        )

    inventory_by_suite: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in inventory_rows:
        inventory_by_suite[row["suite_id"]].append(row)

    suites_with_rows = {row["suite_id"] for row in matrix_rows}
    for suite_id in sorted(inventory_by_suite):
        if suite_id in suites_with_rows:
            continue
        inv_rows = inventory_by_suite[suite_id]
        first = inv_rows[0]
        source_types = sorted({row["source_type"] for row in inv_rows})
        evidence_paths = " | ".join(row["file_path"] for row in inv_rows[:4])
        ocr_hit = ocr_cache_index.get(suite_id)
        hard_excluded_suite = first["year"] == "2026" and first["district"] == "石景山" and first["stage"] == "期末"
        if ocr_hit and not hard_excluded_suite:
            ocr_text = normalize_text(ocr_hit["text"])
            segments = split_questions(ocr_text)
            if len(segments) >= 5:
                ocr_text_path = ocr_absorb_dir / f"{suite_id}.txt"
                ocr_text_path.write_text(ocr_text, encoding="utf-8")
                segments_by_question: dict[str, list[str]] = defaultdict(list)
                for seg in segments:
                    segments_by_question[seg["question"]].append(seg["text"])
                for question, texts in sorted(segments_by_question.items(), key=lambda item: (int(item[0]) if item[0].isdigit() else 999, item[0])):
                    combined_text = "\n".join(texts)
                    candidate_rows.append(
                        {
                            "suite_id": suite_id,
                            "year": first["year"],
                            "district": first["district"],
                            "stage": first["stage"],
                            "question": question,
                            "question_type": question_type(question),
                            "source_type": "ocr-cache",
                            "file_path": ocr_hit["path"],
                            "sha256": "",
                            "run_text_path": str(ocr_text_path),
                            "snippet": compact_snippet(combined_text),
                            "segment_char_count": len(combined_text),
                        }
                    )
                    classification = classify_text(combined_text, ocr_hit["path"], "ocr-cache")
                    matrix_rows.append(
                        {
                            "suite_id": suite_id,
                            "year": first["year"],
                            "district": first["district"],
                            "stage": first["stage"],
                            "question": question,
                            "book_module": classification["book_module"],
                            "question_type": question_type(question),
                            "evidence_source": ocr_hit["path"],
                            "source_types": "ocr-cache",
                            "status": classification["status"],
                            "artifact_location": str(ocr_text_path),
                            "decision_reason": classification["decision_reason"],
                            "next_action": classification["next_action"],
                        }
                    )
                continue
        if hard_excluded_suite:
            module = "EXCLUDED_BY_HARD_RULE"
            status = "module-boundary-excluded"
            next_action = "do_not_process_without_new_user_provided_rubric"
            reason = "2026石景山期末 hard-rule exclusion"
        elif "unknown_district" in suite_id:
            module = "REFERENCE_HELPER_ONLY"
            status = "reference-only"
            next_action = "use_as_helper_not_coverage_evidence"
            reason = "no stable district/stage identity; helper compilation only"
        elif "paper" in source_types:
            module = "UNKNOWN_OR_MIXED"
            status = "ocr-needed"
            next_action = "ocr_or_manual_split_needed"
            reason = "paper exists but no reliable question-level split was produced"
        elif "module-classification" in source_types:
            module = "REFERENCE_HELPER_ONLY"
            status = "reference-only"
            next_action = "use_as_helper_not_coverage_evidence"
            reason = "module classification helper has no single suite question ids"
        else:
            module = "UNKNOWN_OR_MIXED"
            status = "blocked"
            next_action = "manual_source_review_needed"
            reason = "no paper-based question split was produced"
        matrix_rows.append(
            {
                "suite_id": suite_id,
                "year": first["year"],
                "district": first["district"],
                "stage": first["stage"],
                "question": "SUITE_BLOCKER",
                "book_module": module,
                "question_type": "unknown",
                "evidence_source": evidence_paths,
                "source_types": ",".join(source_types),
                "status": status,
                "artifact_location": str(INVENTORY),
                "decision_reason": json.dumps({"reason": reason, "source_types": source_types}, ensure_ascii=False),
                "next_action": next_action,
            }
        )

    matrix_rows.sort(key=lambda r: (r["suite_id"], r["question"]))

    with candidates_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=candidate_fields)
        writer.writeheader()
        writer.writerows(candidate_rows)

    matrix_path = RUN_DIR / "04_module_classification" / "module_classification_matrix.csv"
    matrix_fields = [
        "suite_id",
        "year",
        "district",
        "stage",
        "question",
        "book_module",
        "question_type",
        "evidence_source",
        "source_types",
        "status",
        "artifact_location",
        "decision_reason",
        "next_action",
    ]
    with matrix_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=matrix_fields)
        writer.writeheader()
        writer.writerows(matrix_rows)

    coverage_path = RUN_DIR / "00_control" / "COVERAGE_MATRIX.csv"
    coverage_fields = [
        "suite_id",
        "question",
        "book_module",
        "question_type",
        "evidence_source",
        "status",
        "artifact_location",
        "decision_reason",
        "next_action",
    ]
    with coverage_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=coverage_fields)
        writer.writeheader()
        for row in matrix_rows:
            writer.writerow({field: row[field] for field in coverage_fields})

    status_counts = Counter(row["status"] for row in matrix_rows)
    module_counts = Counter(row["book_module"] for row in matrix_rows)
    type_counts = Counter(row["question_type"] for row in matrix_rows)
    extraction_counts = Counter(row["extraction_status"] for row in extraction_rows)
    cache_counts = Counter(row["cache_hit"] for row in extraction_rows)
    suite_counts = Counter(row["suite_id"] for row in matrix_rows)

    matrix_by_suite: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in matrix_rows:
        matrix_by_suite[row["suite_id"]].append(row)

    suite_overview_rows = []
    for suite_id in sorted(inventory_by_suite):
        inv_rows = inventory_by_suite[suite_id]
        mat_rows = matrix_by_suite.get(suite_id, [])
        status_counter = Counter(row["status"] for row in mat_rows)
        module_counter = Counter(row["book_module"] for row in mat_rows if row["status"] in {"included", "reference-only"})
        source_types = {row["source_type"] for row in inv_rows}
        first = inv_rows[0]
        suite_overview_rows.append(
            {
                "suite_id": suite_id,
                "year": first["year"],
                "district": first["district"],
                "stage": first["stage"],
                "source_file_rows": len(inv_rows),
                "has_paper": "yes" if "paper" in source_types else "no",
                "has_reference_answer": "yes" if "reference-answer" in source_types else "no",
                "has_rubric": "yes" if "rubric" in source_types else "no",
                "has_marking_report": "yes" if "marking-report" in source_types else "no",
                "classified_rows": len(mat_rows),
                "included_B2": sum(1 for row in mat_rows if row["book_module"] == "B2_ECONOMICS" and row["status"] == "included"),
                "included_B3": sum(1 for row in mat_rows if row["book_module"] == "B3_POLITICS_RULE_OF_LAW" and row["status"] == "included"),
                "included_B4_CULTURE": sum(1 for row in mat_rows if row["book_module"] == "B4_CULTURE" and row["status"] == "included"),
                "reference_only_rows": status_counter["reference-only"],
                "boundary_excluded_rows": status_counter["module-boundary-excluded"],
                "blocked_rows": status_counter["blocked"],
                "included_or_reference_modules": ",".join(sorted(module_counter)),
            }
        )

    suite_overview_path = RUN_DIR / "05_reports" / "suite_readiness_overview.csv"
    suite_overview_fields = [
        "suite_id",
        "year",
        "district",
        "stage",
        "source_file_rows",
        "has_paper",
        "has_reference_answer",
        "has_rubric",
        "has_marking_report",
        "classified_rows",
        "included_B2",
        "included_B3",
        "included_B4_CULTURE",
        "reference_only_rows",
        "boundary_excluded_rows",
        "blocked_rows",
        "included_or_reference_modules",
    ]
    with suite_overview_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=suite_overview_fields)
        writer.writeheader()
        writer.writerows(suite_overview_rows)

    suites_without_classified_rows = [row for row in suite_overview_rows if int(row["classified_rows"]) == 0]

    report_path = RUN_DIR / "05_reports" / "classification_readiness_report.md"
    with report_path.open("w", encoding="utf-8") as f:
        f.write("# Classification Readiness Report\n\n")
        f.write(f"- generated_at: {datetime.now().isoformat(timespec='seconds')}\n")
        f.write(f"- source_inventory_rows: {len(inventory_rows)}\n")
        f.write(f"- canonical_source_rows: {len(canonical_rows)}\n")
        f.write(f"- extracted_or_checked_rows: {len(extraction_rows)}\n")
        f.write(f"- question_candidate_rows: {len(candidate_rows)}\n")
        f.write(f"- classified_suite_question_rows: {len(matrix_rows)}\n")
        f.write(f"- suites_with_classified_rows: {len(suite_counts)}\n\n")
        f.write(f"- suite_overview_rows: {len(suite_overview_rows)}\n")
        f.write(f"- suites_without_classified_rows: {len(suites_without_classified_rows)}\n\n")

        f.write("## Cache And Extraction\n\n")
        for key, value in sorted(cache_counts.items()):
            f.write(f"- cache_hit_{key}: {value}\n")
        f.write("\n")
        for key, value in sorted(extraction_counts.items()):
            f.write(f"- {key}: {value}\n")

        f.write("\n## Classification Status Counts\n\n")
        for key, value in sorted(status_counts.items()):
            f.write(f"- {key}: {value}\n")

        f.write("\n## Module Counts\n\n")
        for key, value in sorted(module_counts.items()):
            f.write(f"- {key}: {value}\n")

        f.write("\n## Question Type Counts\n\n")
        for key, value in sorted(type_counts.items()):
            f.write(f"- {key}: {value}\n")

        f.write("\n## Included Target Counts\n\n")
        for module in TARGET_KEYWORDS:
            value = sum(1 for row in matrix_rows if row["book_module"] == module and row["status"] == "included")
            f.write(f"- {module}: {value}\n")

        f.write("\n## Boundary And Blocker Notes\n\n")
        f.write("- Rows marked `module-boundary-excluded` should not enter the three future宝典 lines.\n")
        f.write("- Rows marked `blocked` require manual review before any future coverage claim.\n")
        f.write("- This is a source-forward preparation matrix, not a student-facing artifact.\n")

        f.write("\n## Top Suites By Row Count\n\n")
        for suite, value in suite_counts.most_common(20):
            f.write(f"- {suite}: {value}\n")

        f.write("\n## Suites Without Classified Rows\n\n")
        for row in suites_without_classified_rows[:40]:
            f.write(
                f"- {row['suite_id']}: files={row['source_file_rows']}, "
                f"paper={row['has_paper']}, reference={row['has_reference_answer']}, rubric={row['has_rubric']}\n"
            )

    print(f"wrote {cache_manifest_path}")
    print(f"wrote {candidates_path}")
    print(f"wrote {matrix_path}")
    print(f"wrote {coverage_path}")
    print(f"wrote {suite_overview_path}")
    print(f"wrote {report_path}")
    print(f"candidates={len(candidate_rows)} classified_rows={len(matrix_rows)}")


if __name__ == "__main__":
    main()
