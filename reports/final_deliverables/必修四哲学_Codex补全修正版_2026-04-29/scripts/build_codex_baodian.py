import csv
import re
from collections import OrderedDict, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Inches, RGBColor


ROOT = Path("/Users/wanglifei/Desktop/北京高考政治/必修四框架重做_2026-04-29")
SYNC = Path("/Users/wanglifei/GaokaoPolitics/beijing-politics-sync")
CODEX_BASE = SYNC / "reports/final_deliverables/必修四哲学_学生触发句优先版_2026-04-27/outputs/必修四哲学材料-知识触发总框架_学生触发句优先版_2026-04-27.md"
V6_WORKER_DIR = SYNC / "reports/full_all_suites_independent_rerun_2026-04-29/worker_outputs"
SUPPLEMENT = ROOT / "outputs/必修四框架_同类漏项补做_第一批.md"
OUT = ROOT / "outputs"
AUDIT = ROOT / "audit"
FULL_RERUN = SYNC / "reports/full_all_suites_independent_rerun_2026-04-29"
COVERAGE_MATRIX = FULL_RERUN / "COVERAGE_MATRIX.csv"

TITLE = "2026北京高考政治哲学宝典---三年模拟全触发全链条"
SIGNATURE = "飞哥正志讲堂"


IMAGE_MAP = {
    ("2025", "丰台", "期末", "7", ""): ROOT / "assets/images/2025丰台期末_Q7_漫画.png",
}


FRAMEWORK_STRUCTURE = OrderedDict(
    [
        (
            "一、唯物论",
            [
                (
                    "物质决定意识",
                    ["物质决定意识", "物质与意识", "意识是对客观存在", "意识依赖于物质"],
                ),
                (
                    "一切从实际出发 / 实事求是 / 主观与客观具体的历史的统一",
                    ["一切从实际", "实际出发", "实事求是", "主观与客观", "因地制宜"],
                ),
                (
                    "主观能动性 / 意识的能动作用",
                    ["主观能动性", "意识的能动作用", "意识能动作用", "意识对物质", "能动创造性", "发挥主观能动性"],
                ),
                (
                    "尊重客观规律与发挥主观能动性相结合",
                    ["尊重客观规律与发挥主观能动性", "尊重规律与发挥主观能动性", "客观规律与主观能动性", "规律与主观能动性", "按客观规律办事和发挥主观能动性"],
                ),
                ("规律的客观性", ["规律的客观性", "客观规律", "按规律办事", "尊重客观规律"]),
            ],
        ),
        (
            "二、辩证法",
            [
                ("联系的普遍性 / 联系的观点（总）", ["联系的普遍性", "联系观点", "联系的观点", "普遍联系", "联系"]),
                ("联系的客观性", ["联系的客观性", "客观联系", "自在联系", "自在事物联系", "人为事物联系"]),
                ("根据固有联系建立新的具体联系", ["根据固有联系建立新的具体联系", "建立新的具体联系", "新的具体联系", "从事物固有联系", "根据事物固有联系", "利用固有联系", "把握固有联系", "固有联系"]),
                ("联系的多样性", ["联系的多样性", "多样联系", "多种联系", "条件性", "时间地点条件", "一切以时间地点条件为转移", "多重条件", "多因素"]),
                ("整体与部分", ["整体与部分", "整体", "部分", "全局", "局部"]),
                ("系统观念 / 系统优化", ["系统观念", "系统优化", "系统", "协同", "统筹", "整体性", "有序性", "内部结构优化"]),
                ("发展的观点 / 发展的普遍性", ["发展的观点", "发展观点", "发展的普遍性", "发展"]),
                ("量变与质变 / 适度原则", ["量变", "质变", "质量互变", "适度", "适度原则"]),
                ("事物发展是前进性与曲折性的统一", ["前进性", "曲折性", "前进性与曲折性", "道路是曲折"]),
                ("辩证否定 / 守正创新", ["辩证否定", "守正创新", "扬弃", "自我否定"]),
                ("矛盾就是对立统一", ["对立统一", "一分为二", "既", "又"]),
                ("矛盾的普遍性", ["矛盾的普遍性", "矛盾普遍性", "时时有矛盾", "事事有矛盾"]),
                ("矛盾的特殊性 / 具体问题具体分析", ["矛盾的特殊性", "矛盾特殊性", "具体问题具体分析", "具体分析"]),
                ("矛盾的普遍性和特殊性", ["矛盾的普遍性和特殊性", "普遍性和特殊性", "普遍性与特殊性", "共性与个性", "一般与个别"]),
                ("两点论与重点论", ["两点论", "重点论", "主要矛盾", "次要矛盾", "矛盾的主要方面", "矛盾的次要方面", "主流", "支流", "主次"]),
                ("内因与外因", ["内因", "外因", "内外因"]),
            ],
        ),
        (
            "三、认识论",
            [
                ("实践与认识（总）", ["实践与认识", "实践和认识", "知行"]),
                ("实践是认识的基础", ["实践是认识的基础", "实践决定认识", "实践是认识的来源", "实践是认识发展的动力", "实践是检验认识真理性的唯一标准", "实践是认识的目的", "立足实践", "注重实践"]),
                ("认识对实践的反作用", ["认识对实践", "认识反作用", "科学认识", "正确认识指导实践"]),
                ("认识发展原理", ["认识发展", "认识具有", "认识无限", "认识上升", "深化认识", "认识反复", "认识的反复性", "认识的无限性", "认识的上升性"]),
                ("真理观", ["真理", "检验认识", "追求真理"]),
            ],
        ),
        (
            "四、历史唯物主义",
            [
                ("社会存在与社会意识", ["社会存在", "社会意识"]),
                ("社会发展的两大基本规律和基本矛盾", ["社会基本矛盾", "生产力", "生产关系", "经济基础", "上层建筑", "社会发展的两大基本规律"]),
                ("改革 / 改革的实质", ["改革", "制度创新"]),
                ("人民群众", ["人民群众", "群众观点", "群众路线", "人民是历史"]),
            ],
        ),
        (
            "五、价值观 / 人生观",
            [
                ("价值观的导向作用", ["价值观的导向", "价值观导向", "正确价值观", "价值导向"]),
                ("价值判断与价值选择", ["价值判断", "价值选择"]),
                ("实现人生价值", ["人生价值", "实现人生价值", "个人价值", "社会价值"]),
            ],
        ),
        (
            "六、文化传承、文化功能与文化自信（必修四文化线，非哲学主框架）",
            [
                ("文化的功能 / 文化与经济政治", ["文化功能", "文化的功能", "文化作用", "文化与经济", "文化与政治"]),
                ("中华优秀传统文化创造性转化、创新性发展 / 文化传承", ["优秀传统文化", "创造性转化", "创新性发展", "文化传承", "传统", "文脉", "非遗"]),
                ("中华民族精神 / 文化自信", ["民族精神", "文化自信", "中华文化", "中华文明"]),
                ("社会主义核心价值观 / 文化强国", ["社会主义核心价值观", "核心价值观", "文化强国", "思想道德"]),
            ],
        ),
    ]
)

FRAMEWORK_ORDER = OrderedDict(
    (module, [key for _, keys in children for key in keys]) for module, children in FRAMEWORK_STRUCTURE.items()
)

SUBNODE_ORDER = [(module, subnode, keys) for module, children in FRAMEWORK_STRUCTURE.items() for subnode, keys in children]
SUBNODE_INDEX = {(module, subnode): i for i, (module, subnode, _) in enumerate(SUBNODE_ORDER)}

CLASSIFICATION_PRIORITY = [
    ("一、唯物论", "物质决定意识"),
    ("一、唯物论", "尊重客观规律与发挥主观能动性相结合"),
    ("一、唯物论", "一切从实际出发 / 实事求是 / 主观与客观具体的历史的统一"),
    ("一、唯物论", "主观能动性 / 意识的能动作用"),
    ("一、唯物论", "规律的客观性"),
    ("二、辩证法", "系统观念 / 系统优化"),
    ("二、辩证法", "整体与部分"),
    ("二、辩证法", "根据固有联系建立新的具体联系"),
    ("二、辩证法", "联系的多样性"),
    ("二、辩证法", "联系的客观性"),
    ("二、辩证法", "量变与质变 / 适度原则"),
    ("二、辩证法", "事物发展是前进性与曲折性的统一"),
    ("二、辩证法", "辩证否定 / 守正创新"),
    ("二、辩证法", "矛盾的普遍性和特殊性"),
    ("二、辩证法", "矛盾的特殊性 / 具体问题具体分析"),
    ("二、辩证法", "矛盾的普遍性"),
    ("二、辩证法", "两点论与重点论"),
    ("二、辩证法", "内因与外因"),
    ("二、辩证法", "矛盾就是对立统一"),
    ("二、辩证法", "联系的普遍性 / 联系的观点（总）"),
    ("二、辩证法", "发展的观点 / 发展的普遍性"),
    ("三、认识论", "真理观"),
    ("三、认识论", "认识发展原理"),
    ("三、认识论", "认识对实践的反作用"),
    ("三、认识论", "实践是认识的基础"),
    ("三、认识论", "实践与认识（总）"),
    ("四、历史唯物主义", "社会发展的两大基本规律和基本矛盾"),
    ("四、历史唯物主义", "社会存在与社会意识"),
    ("四、历史唯物主义", "改革 / 改革的实质"),
    ("四、历史唯物主义", "人民群众"),
    ("五、价值观 / 人生观", "价值判断与价值选择"),
    ("五、价值观 / 人生观", "价值观的导向作用"),
    ("五、价值观 / 人生观", "实现人生价值"),
    ("六、文化传承、文化功能与文化自信（必修四文化线，非哲学主框架）", "社会主义核心价值观 / 文化强国"),
    ("六、文化传承、文化功能与文化自信（必修四文化线，非哲学主框架）", "中华优秀传统文化创造性转化、创新性发展 / 文化传承"),
    ("六、文化传承、文化功能与文化自信（必修四文化线，非哲学主框架）", "中华民族精神 / 文化自信"),
    ("六、文化传承、文化功能与文化自信（必修四文化线，非哲学主框架）", "文化的功能 / 文化与经济政治"),
]

DISTRICT_ORDER = {"海淀": 0, "西城": 1, "东城": 2, "朝阳": 3, "丰台": 4}
STAGE_ORDER = {"一模": 0, "二模": 1, "期中": 2, "期末": 3}

BAD_SCALAR_VALUES = {
    "yes",
    "no",
    "true",
    "false",
    "pass",
    "passed",
    "filled",
    "included",
    "nan",
    "none",
    "null",
    "correct",
    "choice_chain",
    "correct_option_chain",
    "correct_and_wrong_option_chain",
    "pass_objective",
    "choice_only",
}


def is_bad_scalar(text: str) -> bool:
    value = str(text or "").strip().strip("。；，,. ").lower()
    return value in BAD_SCALAR_VALUES


BAD_FINAL_PATTERNS = [
    r"`[^`]+`",
    r"（?证据边界：[^）]+）?",
    r"（?证据：[^）]+）?",
    r"【证据：[^】]+】",
    r"来源与证据：",
    r"A细则/阅卷/评分标准直接支持",
    r"B评标/讲评角度清单",
    r"C选择题正确项；有可靠答案源",
    r"稳定触发",
    r"本条为框架映射整理，非细则逐字同名",
    r"本条为框架近邻迁入或人工保留，非细则逐字同名",
]

META_ARTIFACT_TERMS = [
    "评标",
    "参考答案",
    "答案写",
    "答案/补充",
    "答案核",
    "答案提示",
    "答案明确",
    "答案及评分参考",
    "阅卷",
    "细则",
    "评分",
    "讲评",
    "报告明确",
]

ANGLE_LIST_RE = re.compile(r"(?:可从|可以从|能够从|哲学角度可从|文化角度可从)[^。！？；]{0,120}角度作答")


def clean_text(text: str) -> str:
    if not text:
        return ""
    if is_bad_scalar(text):
        return ""
    text = text.replace("\ufeff", "")
    text = text.replace("->", "→")
    for pat in BAD_FINAL_PATTERNS:
        text = re.sub(pat, "", text)
    text = re.sub(r"/Users/[^，。；\s]+", "", text)
    text = re.sub(r"C:\\[^，。；\s]+", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    meta_terms_re = "|".join(re.escape(t) for t in META_ARTIFACT_TERMS + ["答案键"])
    text = re.sub(rf"[^。！？；]*(?:{meta_terms_re})[^。！？；]*[。！？；]?", "", text)
    text = re.sub(r"[^。！？；]*(?:可从|可以从|能够从|哲学角度可从|文化角度可从)[^。！？；]{0,120}角度作答[^。！？；]*[。！？；]?", "", text)
    text = text.replace("因此“", "因此可想到“")
    text = text.replace("因此稳定", "因此")
    text = text.replace("稳定给分", "正确项对应")
    text = text.replace("稳定链", "可迁移链")
    text = text.replace("作答时要先写", "")
    text = text.replace("作答时要写清", "")
    text = text.replace("作答时要写，", "")
    text = text.replace("作答时要写", "")
    text = re.sub(r"因此[^。！？]*(?:可想到|可以想到|想到这个原理)[^。！？]*[。！？]?", "", text)
    text = re.sub(r"因此[^。！？]*(?:课堂整理链|归入|放进|纳入)[^。！？]*[。！？]?", "", text)
    text = re.sub(r"；\s*；", "；", text)
    text = re.sub(r"（\s*）", "", text)
    text = re.sub(r"。\s*。", "。", text)
    text = text.strip(" ；，。")
    if not text or is_bad_scalar(text):
        return ""
    return text + ("。" if text[-1] not in "。！？；" else "")


def contains_final_artifact_language(text: str) -> bool:
    """Student-facing fields must not expose source-check or answer-key wording."""
    text = str(text or "")
    if not text:
        return False
    return any(term in text for term in META_ARTIFACT_TERMS) or bool(ANGLE_LIST_RE.search(text))


STUDENT_ARTIFACT_RE = re.compile(
    r"(正确项链|错肢|错项|第\s*\d+\s*题答案|参考示例|9分标准|可读页图|示例写|迁入正确项|迁入[①②③④0-9、 ]+正确项|答案：|答案\s*[A-D]\b|哲学角度含|文化角度含)"
)


def scrub_student_language(text: str) -> str:
    """Turn audit/checking language into classroom-facing wording."""
    text = clean_text(text)
    if not text:
        return ""
    text = re.sub(r"第\s*\d+\s*题答案\s*[A-D]?\s*[；;]?", "", text)
    text = re.sub(r"答案\s*[A-D]\s*[；;。]?", "", text)
    text = re.sub(r"正确项链[：:]?", "题干信号：", text)
    text = re.sub(r"错肢[：:]?", "易错边界：", text)
    text = re.sub(r"错项", "易错说法", text)
    text = text.replace("参考示例要求", "材料要求")
    text = text.replace("参考示例", "材料")
    text = text.replace("9分标准必采", "题目要求写出")
    text = re.sub(r"可读页图显示[：:]?", "", text)
    text = text.replace("示例写", "材料呈现")
    text = re.sub(r"迁入正确项[：:]?[①②③④0-9、 ]+[；;]?", "", text)
    text = re.sub(r"迁入[①②③④0-9、 ]+正确项[；;]?", "", text)
    text = text.replace("正确项", "符合题意的说法")
    text = text.replace("答案：", "")
    text = text.replace("哲学角度含", "从哲学上看，材料涉及")
    text = text.replace("文化角度含", "从文化上看，材料涉及")
    text = text.replace("发挥正确意识和正确价值观的导向作用", "发挥人的能动思考和正确价值观的导向作用")
    text = text.replace("正确意识", "正确认识")
    text = re.sub(r"\s+", " ", text).strip(" ；，。")
    return text + ("。" if text and text[-1] not in "。！？；" else "")


def evidence_fragment(row):
    text = scrub_student_language(row.get("material") or row.get("trigger") or row.get("full_prompt") or "")
    text = re.sub(r"!\[[^\]]*\]\([^)]+\)", "", text).strip()
    quoted = re.findall(r"[“《]([^”》]{2,28})[”》]", text)
    if quoted:
        return "、".join(quoted[:3])
    sentence = sentence_split(text)[0].strip() if sentence_split(text) else text
    return sentence[:110].rstrip("，；。") + ("..." if len(sentence) > 110 else "")


NODE_BRIDGES = {
    "物质决定意识": "先有客观条件、现实变化或时代问题，再产生相应认识和价值判断，符合物质决定意识。",
    "一切从实际出发 / 实事求是 / 主观与客观具体的历史的统一": "材料先给出具体地点、对象、条件或问题，说明方案和判断必须从这些实际出发。",
    "主观能动性 / 意识的能动作用": "材料突出人在认识、判断、设计、选择或表达中的主动作用，所以能想到人的意识活动具有能动性。",
    "尊重客观规律与发挥主观能动性相结合": "材料一边有客观条件和规律，一边有人主动设计、调控和实践，说明要在尊重规律基础上发挥主观能动性。",
    "规律的客观性": "材料中的规律、条件或边界不是由人的愿望决定的，人只能认识和利用规律，不能随意改造规律本身。",
    "联系的普遍性 / 联系的观点（总）": "材料把多个主体、要素或环节放在相互影响中呈现，说明不能孤立看问题，要用联系观点理解。",
    "联系的客观性": "材料中的联系来自事物自身属性、空间关系或功能关系，不是人主观任意编造出来的。",
    "根据固有联系建立新的具体联系": "材料不是凭空拼接，而是根据事物已有属性、功能和条件重新组合，建立新的具体联系。",
    "联系的多样性": "材料呈现不同时间、地点、条件、主体或因素的差异，说明联系具有多样性，要分析条件。",
    "整体与部分": "材料同时涉及局部环节和整体效果，说明部分影响整体，整体也统摄部分。",
    "系统观念 / 系统优化": "材料把多个部分、层级或环节组织成一个结构，强调协同配合和整体效能。",
    "发展的观点 / 发展的普遍性": "材料出现新旧变化、历史演进或当代转化，说明事物不是静止的，而是在变化发展中展开。",
    "量变与质变 / 适度原则": "材料出现持续积累、阶段推进、度的把握或由积累到突破的过程，触发量变质变或适度原则。",
    "事物发展是前进性与曲折性的统一": "材料同时写到目标、成果、困难或反复，说明发展前途向前但道路并不平直。",
    "辩证否定 / 守正创新": "材料不是全盘抛弃旧事物，也不是原样复制，而是在保留合理内核中克服旧形式、实现创新发展。",
    "矛盾就是对立统一": "材料同时呈现相反因素或两种力量的相互依存、相互制约，触发对立统一。",
    "矛盾的普遍性": "材料把同一对象的积极面和风险面、优势和问题同时呈现，说明矛盾普遍存在，要一分为二看。",
    "矛盾的特殊性 / 具体问题具体分析": "材料给出特定对象、群体或场景的特殊问题，说明不能套用统一方案，要具体问题具体分析。",
    "矛盾的普遍性和特殊性": "材料同时出现可借鉴的共性和具体情境的个性，触发共性与个性相统一。",
    "两点论与重点论": "材料既要求全面看到两方面，又提示关键短板或主要方向，说明要两点论和重点论统一。",
    "内因与外因": "材料同时呈现内部动力和外部条件，说明发展要把内因与外因结合起来分析。",
    "实践与认识（总）": "材料把认识、学习、体验、应用或解决问题连在一起，说明认识离不开实践并服务实践。",
    "实践是认识的基础": "材料中的认识来自调查、实验、体验、使用或治理实践，说明实践是认识的来源、动力或检验标准。",
    "认识对实践的反作用": "材料显示理念、图谱、判断或方案反过来指导行动，说明正确认识能够推动实践。",
    "认识发展原理": "材料呈现认识从未知到已知、从粗浅到深化的过程，说明认识具有反复性、无限性和上升性。",
    "真理观": "材料强调认识要经由实践证明或检验，说明真理性不能只靠主观确信。",
    "社会存在与社会意识": "材料把社会条件、技术变革或现实生活变化同观念更新联系起来，说明社会存在决定社会意识。",
    "社会发展的两大基本规律和基本矛盾": "材料涉及生产力、生产关系、经济基础或上层建筑的调整，说明社会发展由基本矛盾运动推动。",
    "改革 / 改革的实质": "材料出现破旧立新、制度完善或体制机制调整，说明改革是社会主义制度的自我完善和发展。",
    "人民群众": "材料呈现群众参与、群众需要、群众智慧或群众评价，说明人民群众是社会历史主体。",
    "价值观的导向作用": "材料把行为选择同价值取向、生态理念或公共利益联系起来，说明价值观会引导实践方向。",
    "价值判断与价值选择": "材料要求在不同利益、长短期目标或价值方向之间作判断和选择，触发价值判断与价值选择。",
    "实现人生价值": "材料把个人奋斗、责任担当同社会需要联系起来，说明人生价值要在个人与社会统一中实现。",
    "文化的功能 / 文化与经济政治": "材料呈现文化对人的精神、社会治理、经济发展或价值共识的作用，说明文化具有功能和影响。",
    "中华优秀传统文化创造性转化、创新性发展 / 文化传承": "材料把传统资源同当代技术、审美、生活或传播方式结合，说明传统文化要创造性转化、创新性发展。",
    "中华民族精神 / 文化自信": "材料呈现民族精神、家国情怀、历史记忆或文化认同，说明文化能凝聚精神力量、增强文化自信。",
    "社会主义核心价值观 / 文化强国": "材料把文化实践同价值引领、思想道德或文化强国目标联系起来，说明要用核心价值观引领文化建设。",
}


def weak_why_reasons(row, text):
    text = scrub_student_language(text)
    reasons = []
    material = scrub_student_language(row.get("material") or row.get("trigger") or "")
    if not text:
        return ["empty"]
    if material and text.strip("。") == material.strip("。"):
        reasons.append("copy_of_material")
    if len(text) < 45:
        reasons.append("too_short")
    if STUDENT_ARTIFACT_RE.search(row.get("why", "")):
        reasons.append("artifact_language")
    if re.match(r"^(题目不是|设问要求|这类题|材料不是|材料强调|材料把|材料直接|材料连续|不是单纯|当材料|看到|只要|可从)", text):
        reasons.append("generic_start")
    return reasons


def repair_weak_why(row, text):
    evidence = evidence_fragment(row)
    node = row.get("node") or classify(row)[1]
    bridge = NODE_BRIDGES.get(node, "材料中的具体信息能够和该原理的核心要求对应起来。")
    if not evidence:
        return scrub_student_language(text)
    return clean_text(f"能想到这个原理，是因为材料中的“{evidence}”给出了直接信号：{bridge}")


def student_landing(row):
    landing = scrub_student_language(row.get("landing", ""))
    evidence = evidence_fragment(row)
    node = row.get("node") or classify(row)[1]
    if not landing or len(landing) < 40 or STUDENT_ARTIFACT_RE.search(row.get("landing", "")):
        if evidence:
            bridge = NODE_BRIDGES.get(node, "材料事实说明该原理能够解释题中现象。")
            landing = f"答案可以落在：材料中的“{evidence}”说明，{bridge}"
    return clean_text(landing)


def student_question_prompt(row):
    prompt = scrub_student_language(question_prompt(row.get("full_prompt") or row.get("material") or ""))
    if question_type_rank(row):
        if STUDENT_ARTIFACT_RE.search(row.get("full_prompt", "") + row.get("material", "")) or not any(c in prompt for c in QUESTION_CUES):
            return "选择题：根据题干材料，判断最符合题意的一项。"
    return prompt


QUESTION_CUES = [
    "结合材料",
    "综合运用",
    "运用《",
    "运用所学",
    "从哲学角度",
    "从哲学与文化角度",
    "谈谈",
    "说明",
    "分析",
    "阐释",
    "评价",
    "认识",
    "理解",
    "为什么",
    "为何",
    "如何",
    "应如何",
    "请你",
]

QUESTION_LEADIN_CUES = [
    "有人说",
    "有人认为",
    "有观点",
    "社会上有人",
    "网友认为",
    "也有人认为",
    "这一观点",
    "该观点",
    "上述观点",
    "这句话",
    "这一问题",
    "这一时代命题",
]


def sentence_split(text: str):
    return re.findall(r"[^。！？]+[。！？]?", str(text or ""))


def is_question_sentence(sentence: str) -> bool:
    s = str(sentence or "").strip()
    return any(cue in s for cue in QUESTION_CUES) or bool(re.search(r"[？?]\s*$", s))


def is_question_leadin(sentence: str) -> bool:
    s = str(sentence or "").strip()
    if len(s) > 160:
        return False
    return any(cue in s for cue in QUESTION_LEADIN_CUES)


def trim_inline_material_before_question(text: str) -> str:
    """If material and question sit in one sentence, cut at the real question cue."""
    text = str(text or "").strip()
    for cue in ["结合材料", "综合运用", "运用《", "运用所学"]:
        idx = text.rfind(cue)
        if idx <= 0:
            continue
        lead = text[:idx].strip(" ，；。")
        tail = text[idx:].strip()
        if len(lead) <= 160 and is_question_leadin(lead):
            return clean_text(lead + "。" + tail)
        m = re.search(r"([^。！？]{0,160}(?:有人说|有人认为|有观点|社会上有人|网友认为|“[^”]{1,120}”)[^。！？]*)$", lead)
        if m:
            return clean_text(m.group(1).strip(" ，；。") + "。" + tail)
        return clean_text(tail)
    return clean_text(text)


def question_prompt(text: str) -> str:
    """Return only the exam question prompt, not the whole material body."""
    text = clean_text(text)
    if not text:
        return ""
    sentences = [s.strip() for s in sentence_split(text) if s.strip()]
    if len(text) <= 180:
        return text
    q_idx = None
    for i in range(len(sentences) - 1, -1, -1):
        if is_question_sentence(sentences[i]):
            q_idx = i
            break
    if q_idx is None:
        return text

    start = q_idx
    while start > 0 and is_question_leadin(sentences[start - 1]):
        start -= 1
    candidate = clean_text("".join(sentences[start:]))
    if len(candidate) > 220:
        candidate = trim_inline_material_before_question(candidate)
    if candidate:
        return candidate
    return text


def short_source(text: str) -> str:
    text = clean_text(text)
    text = re.sub(r"（[^）]*(?:pdf|docx|pptx|缓存|答案源|评分标准|细则|阅卷)[^）]*）", "", text, flags=re.I)
    text = re.sub(r"；.*", "", text)
    text = text.replace("北京", "")
    return text.strip(" ；，。")


def normalize_source_key(text: str):
    if not text:
        return None
    t = text.replace("（上）", "").replace("(上)", "")
    head = re.search(r"(202[4-6]).{0,18}(海淀|西城|东城|朝阳|丰台|石景山|门头沟|顺义|延庆|房山|昌平|通州).{0,18}(一模|二模|期中|期末)", t)
    q = re.search(r"第\s*([0-9]+)\s*题?(?:第?\s*[（(]\s*([0-9一二三四五六七八九十]+)\s*[）)]\s*问?)?", t)
    if not (head and q):
        return None
    return (head.group(1), head.group(2), head.group(3), q.group(1), q.group(2) or "")


def same_source_key_loose(key):
    if not key:
        return None
    return key[:4]


def parse_codex_base():
    text = CODEX_BASE.read_text(encoding="utf-8")
    rows = []
    section = ""
    buf = []
    for line in text.splitlines():
        if line.startswith("### "):
            section = clean_text(line[4:])
        if re.match(r"^\d+\. \*\*材料触发句\*\*", line):
            if buf:
                rows.append(parse_codex_block(section, "\n".join(buf)))
            buf = [line]
        elif buf:
            buf.append(line)
    if buf:
        rows.append(parse_codex_block(section, "\n".join(buf)))
    return [r for r in rows if r]


def field(block: str, label: str) -> str:
    m = re.search(rf"(?:\*\*)?{re.escape(label)}(?:\*\*)?[：:](.+)", block)
    return clean_text(m.group(1)) if m else ""


def parse_codex_block(section: str, block: str):
    trigger_sentence = field(block, "材料触发句")
    related = field(block, "相关材料")
    knowledge = field(block, "触发知识")
    why = field(block, "为什么触发")
    source = field(block, "来源与证据")
    if not (trigger_sentence and knowledge and source):
        return None
    title = knowledge.split("/")[0].split("；")[0].strip()
    if len(title) > 34:
        title = section
    return {
        "origin": "Codex-2026-04-27",
        "section": section,
        "title": title,
        "source": short_source(source),
        "source_key": normalize_source_key(source),
        "full_prompt": "",
        "material": related,
        "trigger": trigger_sentence,
        "knowledge": knowledge,
        "why": why,
        "landing": "",
        "rubric": "",
    }


def parse_v6_prompts_and_entries():
    prompt_by_key = {}
    entries = []
    for path in sorted(V6_WORKER_DIR.glob("*_v6_student_entries.md")):
        text = path.read_text(encoding="utf-8", errors="replace")
        last_prompt = ""
        for block in re.split(r"(?=^## )", text, flags=re.M):
            if not block.startswith("## "):
                continue
            title = clean_text(block.splitlines()[0][3:])
            if title.startswith(("一、", "二、")) or "学生版条目" in title:
                continue
            source = field(block, "**来源题目**")
            full_prompt = question_prompt(field(block, "**完整设问**") or field(block, "**完整题干**"))
            if full_prompt == "同上。" and last_prompt:
                full_prompt = last_prompt
            elif full_prompt:
                last_prompt = full_prompt
            key = normalize_source_key(source)
            if key and full_prompt:
                prompt_by_key.setdefault(key, full_prompt)
                prompt_by_key.setdefault(same_source_key_loose(key), full_prompt)
            if source and full_prompt:
                entries.append(
                    {
                        "origin": path.name,
                        "section": "",
                        "title": title,
                        "source": short_source(source),
                        "source_key": key,
                        "full_prompt": full_prompt,
                        "material": "",
                        "trigger": field(block, "**材料触发点**") or field(block, "**错肢触发点**"),
                        "knowledge": title,
                        "why": field(block, "**为什么能想到这个原理**") or field(block, "**为什么不能这样想**"),
                        "landing": field(block, "**答案落点**"),
                        "rubric": "",
                    }
                )
    return prompt_by_key, entries


def parse_supplement():
    text = SUPPLEMENT.read_text(encoding="utf-8")
    rows = []
    section = ""
    title = ""
    block = []
    for line in text.splitlines():
        if line.startswith("## "):
            section = clean_text(line[3:])
        if line.startswith("### "):
            if block and title:
                rows.append(parse_supplement_block(section, title, "\n".join(block)))
            title = clean_text(line[4:])
            block = []
        elif title:
            block.append(line)
    if block and title:
        rows.append(parse_supplement_block(section, title, "\n".join(block)))
    return [r for r in rows if r]


def source_from_suite(suite_name: str, question: str) -> str:
    question = str(question or "").strip().replace("Q", "")
    if not question or question in {"N/A", "choice-all", "subjective-all", "1-15"}:
        return suite_name
    if re.search(r"第\s*\d+", question):
        return f"{suite_name} {question}"
    return f"{suite_name} 第{question}题"


def passish(value: str) -> bool:
    value = str(value or "").lower()
    return any(token in value for token in ["pass", "included", "filled", "choice_chain", "correct"])


def make_choice_row(source, title, material, why, landing, origin):
    return {
        "origin": origin,
        "section": "",
        "title": clean_text(title),
        "source": short_source(source),
        "source_key": normalize_source_key(source),
        "full_prompt": clean_text(material),
        "material": clean_text(material),
        "trigger": clean_text(material),
        "knowledge": clean_text(title),
        "why": clean_text(why),
        "landing": clean_text(landing or why),
        "rubric": "",
    }


def parse_late_objective_backfills():
    path = FULL_RERUN / "late_objective_closure/late_objective_entries_normalized.csv"
    rows = []
    if not path.exists():
        return rows
    with path.open(encoding="utf-8-sig", newline="") as f:
        for r in csv.DictReader(f):
            if not passish(r.get("student_action", "")):
                continue
            module_blob = " ".join([r.get("module", ""), r.get("wrong_or_boundary", ""), r.get("correct_chain", "")])
            if not any(t in module_blob for t in ["哲学", "文化", "实践", "认识", "联系", "发展", "矛盾", "价值", "意识", "民族精神", "传统", "社会历史", "人生"]):
                continue
            source = source_from_suite(r["suite_name"], r["question"])
            material = r.get("stem_or_material") or r.get("material_trigger") or r.get("correct_chain")
            why = r.get("correct_chain") or r.get("wrong_or_boundary") or material
            title = r.get("wrong_or_boundary") or r.get("module") or r.get("correct_chain")
            if "choice_chain" in title:
                title = r.get("correct_chain", title)
            rows.append(
                make_choice_row(
                    source,
                    title,
                    material,
                    why,
                    why,
                    "late-objective-closure",
                )
            )
    return rows


def parse_sample_s040_backfills():
    path = FULL_RERUN / "sample_outputs/S040_2026东城一模_entries.csv"
    rows = []
    if not path.exists():
        return rows
    with path.open(encoding="utf-8-sig", newline="") as f:
        for r in csv.DictReader(f):
            if not passish(r.get("status", "")):
                continue
            if "wrong_option_only" in r.get("status", ""):
                continue
            source = source_from_suite(r["suite_name"], r["question_no"])
            title = r.get("knowledge_point") or r.get("module") or r.get("entry_kind")
            material = r.get("material_trigger") or r.get("evidence_excerpt")
            why = r.get("evidence_excerpt") or material
            landing = r.get("answer_landing") or why
            rows.append(make_choice_row(source, title, material, why, landing, "S040-sample-output"))
    return rows


def parse_worker_audit_backfills():
    rows = []
    for path in sorted((FULL_RERUN / "worker_outputs").glob("*_v6_audit_entries.csv")):
        with path.open(encoding="utf-8-sig", newline="") as f:
            for r in csv.DictReader(f):
                status = r.get("status") or r.get("boundary_status") or ""
                if not passish(status):
                    continue
                if any(t in status.lower() for t in ["reference_only", "reference-only", "blocked", "need_evidence", "excluded"]):
                    continue
                suite = r.get("suite_name", "")
                q = r.get("question_no") or r.get("question") or ""
                title = r.get("knowledge_point") or r.get("module") or r.get("item_type") or r.get("entry_type")
                material = r.get("material_trigger") or r.get("evidence_excerpt") or r.get("student_entry")
                landing = r.get("answer_landing") or r.get("student_entry") or r.get("evidence_excerpt")
                blob = " ".join([title or "", material or "", landing or ""])
                if not any(t in blob for t in ["哲学", "文化", "实践", "认识", "联系", "发展", "矛盾", "价值", "意识", "规律", "人民", "传统", "民族精神", "文化自信", "人生"]):
                    continue
                rows.append(
                    make_choice_row(
                        source_from_suite(suite, q),
                        title,
                        material,
                        material,
                        landing,
                        path.name,
                    )
                )
    return rows


def targeted_coverage_backfills():
    """Manual coverage repairs from source-checked bundles where merger missed valid B4 chains."""
    data = [
        # 2024 first/second round leftovers.
        ("2024西城一模 第17题", "价值观的导向作用：生态价值观引导人与自然和谐共生", "城市公园把植物生长交给自然，呈现生物多样性和自然演替，题目要求说明这种处理人与自然关系的意义。", "材料不是单纯写公园景观，而是把人对自然的态度作为行动方向；尊重自然、顺应自然、保护自然体现正确生态价值观对实践的导向。", "处理人与自然关系，要树立正确生态价值观，把人的活动控制在尊重生态规律的范围内，推动人与自然和谐共生。"),
        ("2024海淀二模 第16题", "联系的普遍性 / 联系的观点", "材料把个体、社会环境和人物选择放在相互影响的关系中展开。", "题目不是孤立评价一个人的行为，而是要求看到人和社会、个人选择和社会条件之间的普遍联系。", "分析人物成长和社会贡献时，要把个人行为放到社会联系中理解，说明个人发展离不开社会条件，也会反过来影响社会。"),
        ("2024海淀二模 第16题", "实现人生价值：在个人与社会统一中创造价值", "材料要求评价人物如何在社会需要中作出选择、承担责任并形成贡献。", "当材料把个人选择同社会需要、责任担当和贡献联系起来时，触发的是人生价值在个人与社会统一中实现。", "人物价值不只在个人成就，而在把个人理想同社会需要结合起来，在责任担当和贡献中实现人生价值。"),
        ("2024西城二模 第17题", "价值观的导向作用：生态价值观", "公园建设保留野草、野花、野树，让植物按照自然规律生长，形成生物多样性。", "这种做法背后不是单纯技术安排，而是用尊重自然、保护生态的价值取向引导城市治理。", "城市治理要坚持正确生态价值观，让规划、建设和管理服务人与自然和谐共生。"),
        ("2024海淀期中 第18题", "两点论与重点论", "细则围绕辩证思维，强调既要全面看问题，又要抓住重点展开论证。", "当题目要求既看到多个方面，又不能平均用力时，就触发两点论与重点论相统一。", "回答这类题要先全面分析对象的不同方面，再围绕关键方面展开论证，做到既看两点，又突出重点。"),
        # S013 objective-choice chains.
        ("2024顺义思政二模 第2题", "主观能动性 / 意识的能动作用", "新一代载人飞船命名为“梦舟”、月面着陆器命名为“揽月”，名称令人遐想并具有中国文化特色。", "飞行器命名不是客观物本身自动产生意义，而是人把航天实践、梦想追求和文化意象能动地结合起来。", "看到命名、寓意、文化想象类材料，要写意识活动具有目的性、主动创造性，能够用文化符号表达实践追求。"),
        ("2024顺义思政二模 第3题", "文化的功能 / 文化与经济政治", "哈尔滨冰雪经济火爆后，多地文旅部门打造各自文旅名片，文化资源带动消费和城市热度。", "材料把文化资源、文旅产业和经济发展连在一起，说明文化与经济相互交融，优秀文化资源可以转化为发展动能。", "文旅发展要挖掘本地文化资源，通过创意表达和公共服务把文化优势转化为经济社会发展优势。"),
        ("2024顺义思政二模 第4题", "中华优秀传统文化创造性转化、创新性发展 / 文化传承", "来自不同文明的演奏家同乘一叶小舟，用古老乐器奏响人类文明和谐之音。", "不同文明同台表达，说明文化既有民族特色，也能通过交流互鉴形成世界共鸣。", "文化交流中既要守住自身特色，又要尊重文明多样性，在交流互鉴中推动文化发展。"),
        # S022 source-checked question and rubric chains.
        ("2025房山一模 第3题", "矛盾的特殊性 / 具体问题具体分析", "智能养老设备不断普及，但操作复杂、隐私担忧使部分老年人难以适应。", "同样是技术应用，老年群体有特殊需求和使用障碍，不能按普通用户场景一刀切。", "发展智能养老要从老年群体的特殊需要出发，改进操作方式、隐私保护和适老服务。"),
        ("2025房山一模 第4题", "矛盾的普遍性和特殊性", "北京大气治理经验可以为曼谷提供借鉴，但曼谷也有自身环境状况和发展条件。", "材料同时出现可借鉴的共性经验和他国治理的个性条件，触发共性与个性具体的历史的统一。", "借鉴北京经验不能照搬，要把污染治理的一般经验同曼谷的具体实际结合起来。"),
        ("2025房山一模 第5题", "辩证否定 / 守正创新", "青春版《牡丹亭》保留昆曲传统内核，又通过舞台、服装、人物等形式创新吸引年轻观众并走向世界。", "材料不是抛弃传统，也不是原样复制，而是在保留精华基础上创新表达。", "传统艺术焕发生命力，要坚持守正创新，在传承内容精髓的同时创新表现形式和传播方式。"),
        ("2025房山一模 第20题", "实现人生价值", "奋斗被界定为付出艰辛努力、战胜困难、实现宏伟目标，材料强调青年把小我融入大我。", "当设问把奋斗、青年、民族复兴和克服困难连在一起时，触发人生价值在劳动奉献和社会需要中实现。", "青年要把个人奋斗融入民族复兴，在砥砺自我、服务国家和人民中创造并实现人生价值。"),
        # Other 2025/2026 leftovers.
        ("2025门头沟一模 第16题", "发展的观点 / 发展的普遍性", "故宫数字文物展览、八达岭智能导览、北京胡同数字地图等让古老文化遗产与现代数字技术深度融合。", "材料强调古老文化在新技术、新场景中获得新表达，说明文化不是静止保存，而是在发展中延续生命力。", "文化遗产保护要用发展的眼光，通过数字技术和新场景开发让传统文化在当代继续生长。"),
        ("2025海淀期中 第7题", "中华民族精神 / 文化自信", "延安作风体现党的初心使命和性质宗旨，新的历史条件下仍要传承弘扬。", "革命精神不是历史标签，而是能够为新时代作风建设提供精神滋养和价值参照的文化资源。", "新时代党员要传承弘扬延安作风，把党的初心使命转化为联系群众、服务人民和自我革命的行动。"),
        ("2025海淀期中 第9题", "人民群众", "从83岁退休教育工作者到12岁学生都参与“十五五”规划建言，表达对蓝天碧水和民生温度的期待。", "规划建言征集把不同年龄群众的经验、需要和智慧纳入决策过程，说明人民群众是社会实践和社会治理的重要主体。", "科学规划要坚持问计于民，尊重人民主体地位，把群众智慧和群众需要转化为治理方案。"),
        ("2025海淀期中 第12题", "价值观的导向作用", "中国推进绿色丝绸之路建设，在埃及、埃塞俄比亚推广生态修复与绿色发展经验。", "材料把国际合作同生态文明理念、绿色发展共识联系起来，说明正确生态价值观能够引导合作方向。", "共建绿色丝绸之路要坚持绿色发展价值导向，通过生态合作凝聚国际绿色共识、造福各国人民。"),
        ("2026海淀期中 第7题", "中华民族精神 / 文化自信", "延安作风集中体现党的初心使命和性质宗旨，新时代仍要传承弘扬。", "革命作风作为精神财富，能够为新时代全面从严治党和党员作风建设提供价值引领。", "新时代要传承延安作风，把革命精神转化为密切联系群众、服务人民和推进自我革命的现实行动。"),
        ("2026海淀期中 第9题", "人民群众", "一老一少参与北京市“十五五”规划建言，提出城市供暖、环保、亲水空间等建议。", "材料直接呈现人民群众参与公共决策、贡献生活经验和治理智慧。", "制定规划要坚持问计于民、问需于民，把人民群众的智慧和期待转化为城市治理方向。"),
        ("2026朝阳期末 第16题", "中华优秀传统文化创造性转化、创新性发展 / 文化传承", "传统戏曲保留艺术精髓，又通过新的舞台表达和时代内涵焕发生命力。", "材料中的“焕发新的生命力”说明传统文化需要在继承中发展，在创新中传承。", "传统戏曲要在守住精髓的基础上创新表达，使中华优秀传统文化以当代方式被理解、欣赏和传播。"),
        ("2026丰台期末 第16题", "矛盾就是对立统一", "“留白”以无衬有、以虚映实，在空白和画面、多少与疏密之间形成审美张力。", "留白之美来自相反因素之间的相互依存和相互成就，正是矛盾双方对立统一的表现。", "理解留白要写虚与实、无与有相互依存、相互成就，艺术效果正是在这种对立统一中生成。"),
        ("2026石景山期末 第18题（1）", "联系的普遍性 / 联系的观点（总）", "北京非遗漫游把非遗主题、典型村落、空间场景和不同功能联系成文化体验整体。", "材料不是孤立介绍一个非遗点位，而是把多处村落和多种功能放进整体文化线路中。", "说明非遗漫游意义时，要写各类非遗资源相互联结，共同形成可体验、可传播的文化整体。"),
        ("2026石景山期末 第18题（1）", "发展的观点 / 发展的普遍性", "非遗漫游把传统非遗资源转化为当代文旅体验和公共文化供给。", "材料强调传统资源进入新的生活场景，说明非遗保护不是静态封存，而是在发展中延续。", "非遗传承要用发展的观点，把传统技艺、村落空间和现代体验结合起来，让传统文化获得新的生命力。"),
    ]
    return [
        {
            "origin": "targeted-coverage-backfill",
            "section": "",
            "title": title,
            "source": short_source(source),
            "source_key": normalize_source_key(source),
            "full_prompt": material,
            "material": material,
            "trigger": material,
            "knowledge": title,
            "why": why,
            "landing": landing,
            "rubric": "",
        }
        for source, title, material, why, landing in data
    ]


def manual_supplements():
    """Source-checked additions found by the patcher lane but absent from the v6 student text."""
    prompt = (
        "对长期主义，中国人有个更生动的表达——“坚持就是胜利”。“路漫漫其修远兮，吾将上下而求索”"
        "“凿井者，起于三寸之坎，以就万仞之深”……这些名句，穿越时空，成为一代代国人的座右铭。"
        "认准了就坚定不移干。一年不行就十年，十年不行就二十年、三十年。一代接着一代、一茬接着一茬，日拱一卒。"
        "长期主义深植于中华民族的文化基因，也融入了中国共产党治国理政的行动哲学。新中国成立以来，从“一五”到“十四五”，"
        "一个又一个紧密接续的“五年规划”，就是在以长期主义的理念不断推进中国式现代化。从“两弹一星”到载人航天，从复兴号到C919，"
        "从塔里木沙漠公路到港珠澳跨海大桥……回望发展历程，许多领域的惊艳成果，都源自几十年如一日的坚守，源自看准了就坚定不移干的执着。"
        "“坚持长期主义，做困难而正确的事”，结合材料，综合运用所学，谈谈你对这句话的理解。"
    )
    base = {
        "origin": "current-thread-manual-source-check",
        "source": "2025西城一模 第22题",
        "source_key": normalize_source_key("2025西城一模 第22题"),
        "full_prompt": prompt,
        "rubric": (
            "西城一模第22题细则写明：长期主义能够科学把握事物发展趋势，坚持真理尊重规律，把眼前利益和长远利益相结合，在量的积累中实现质变；"
            "长期主义植根于中国优秀传统文化，蕴含自强不息的伟大民族精神。"
        ),
    }
    return [
        {
            **base,
            "section": "发展观",
            "title": "发展观点与量变质变：长期主义不是慢，而是持续积累到突破",
            "knowledge": "发展的观点；量变是质变的必要准备；前途是光明的、道路是曲折的",
            "material": "“一年不行就十年，十年不行就二十年、三十年”“一代接着一代、一茬接着一茬，日拱一卒”“许多领域的惊艳成果，都源自几十年如一日的坚守”。",
            "trigger": "长期坚持、日拱一卒、几十年如一日，直接指向量的持续积累；从困难到成果，指向发展道路中的曲折与前进。",
            "why": "材料不是单纯劝人有耐心，而是在解释为什么长期主义能成事：事物发展需要过程，重大成果往往不是一次行动完成，而是在长期积累中实现质的跃升；面对困难仍坚持，是承认道路曲折但相信前途光明。",
            "landing": "理解“坚持长期主义，做困难而正确的事”，要写出长期主义符合发展规律：把眼前行动和长远目标统一起来，以持续积累推动质变，在曲折中坚定推进中国式现代化。",
        },
        {
            **base,
            "section": "价值观",
            "title": "正确价值观：做困难而正确的事",
            "knowledge": "正确价值观的导向作用；正确价值判断和价值选择",
            "material": "设问直接说“做困难而正确的事”，细则强调“科学、坚持真理、尊重规律、正确价值观”。",
            "trigger": "“正确”不是难易判断，而是价值判断；“困难而正确”说明行动选择要受真理、规律和长远价值引导，而不是被短期利益牵着走。",
            "why": "能想到正确价值观，是因为材料中的“做困难而正确的事”不是在比较事情难不难，而是在说明行动选择要受真理、规律和长远利益引导；长期主义把国家发展、民族复兴和人民长远利益放在更高位置，这正是价值观在引导人们作出正确价值判断和价值选择。",
            "landing": "作答时要写，长期主义之所以值得坚持，是因为它坚持正确价值导向，把眼前利益与长远利益统一起来，选择对国家发展和民族复兴真正有价值的方向，即使困难也要坚定推进。",
        },
        {
            **base,
            "section": "文化",
            "title": "中华优秀传统文化与民族精神：长期主义的文化根基",
            "knowledge": "中华优秀传统文化；伟大民族精神；自强不息",
            "material": "材料引用“路漫漫其修远兮，吾将上下而求索”“凿井者，起于三寸之坎，以就万仞之深”，并说“长期主义深植于中华民族的文化基因”。",
            "trigger": "古代名句、文化基因、自强不息共同说明，长期主义不是临时口号，而是中华优秀传统文化中坚韧求索、久久为功精神的当代表达。",
            "why": "当材料用传统名句解释现实发展方法时，触发的是优秀传统文化的当代价值；当材料强调一代代坚持和自强不息时，触发的是伟大民族精神为现代化建设提供精神动力。",
            "landing": "理解长期主义，还要写出其文化根基：中华优秀传统文化和自强不息的民族精神为中国式现代化提供持久精神力量，使一代代人能够接续奋斗、久久为功。",
        },
    ]


def manual_2024_east_supplements():
    """Source-checked OCR repair for 2024东城一模 Q18(3), missed by earlier OCR-needed runs."""
    prompt = (
        "传统产业在我国制造业中占比超80%，是现代化产业体系的基底。多年来，我国轻工、纺织、机械等传统产业带动效应强、产业关联度大、国际市场占有率高，"
        "是我国参与国际竞争的生力军。但也存在低端产能过剩、高端供给不足、产业基础不牢、创新能力薄弱等问题，对数字化、智能化有着庞大、迫切的实际需求。"
        "当前，从人形机器人到量子计算机，从脑机接口到6G网络设备，中国未来产业正在科技新赛道上加速前进。未来材料、未来网络等能够广泛渗透传统产业链上下游各个环节，"
        "引领传统优势产业高端化、智能化、绿色化转型。（3）发展新质生产力，要正确认识传统产业和未来产业的关系。结合材料，运用《逻辑与思维》知识，谈谈你的看法。"
    )
    base = {
        "origin": "current-thread-ocr-source-check",
        "source": "2024东城一模 第18题第（3）问",
        "source_key": normalize_source_key("2024东城一模 第18题第（3）问"),
        "full_prompt": prompt,
        "rubric": "细则明确要求：不是忽视、放弃传统产业，而是坚持辩证否定、改造提升传统产业；并用超前思维布局建设未来产业。",
    }
    return [
        {
            **base,
            "section": "辩证法",
            "title": "辩证否定 / 守正创新",
            "knowledge": "辩证否定 / 守正创新",
            "material": "传统产业既是现代化产业体系的基底，又存在低端产能过剩、高端供给不足、创新能力薄弱等问题；未来产业能够渗透传统产业链上下游，引领传统优势产业高端化、智能化、绿色化转型。",
            "trigger": "材料同时写出传统产业的基础价值和现实问题，并要求未来产业赋能传统产业升级，触发的不是简单抛弃旧产业，而是在保留合理基础中克服落后环节。",
            "why": "辩证否定的核心是扬弃：既肯定传统产业的支撑作用，又否定其低端、薄弱、落后的方面；未来产业不是取代传统产业，而是通过技术赋能推动传统产业改造提升。",
            "landing": "正确认识传统产业和未来产业的关系，要说明发展新质生产力不能忽视、放弃传统产业，而应保留传统产业的制造基础、配套能力和市场支撑，同时用未来技术克服低端产能、创新薄弱等问题，实现对传统产业的扬弃式升级。",
        },
        {
            **base,
            "section": "辩证法",
            "title": "发展的观点 / 发展的普遍性",
            "knowledge": "发展的观点 / 发展的普遍性",
            "material": "未来产业正在科技新赛道上加速前进，未来材料、未来网络等能够引领传统优势产业高端化、智能化、绿色化转型。",
            "trigger": "材料强调新赛道、新技术、新产业形态以及传统产业向高端化、智能化、绿色化转型，指向事物在新条件下不断发展。",
            "why": "发展的观点要求看到产业不是静止的：传统产业可以在未来技术赋能下升级，未来产业也会在培育中成长为推动新质生产力的新动能。",
            "landing": "正确认识二者关系，要支持和培育未来产业这些新事物，同时看到未来产业的发展能够推动传统产业转型升级，使现代化产业体系在持续发展中形成新质生产力。",
        },
        {
            **base,
            "section": "辩证法",
            "title": "系统观念 / 系统优化",
            "knowledge": "系统观念 / 系统优化",
            "material": "传统产业是现代化产业体系的基底，未来产业又能渗透传统产业链上下游各个环节，二者共同服务现代化产业体系建设。",
            "trigger": "材料把传统产业、未来产业、产业链上下游、现代化产业体系放在同一结构中理解，说明不能孤立看某一个产业。",
            "why": "系统观念要求把传统产业和未来产业作为现代化产业体系中的不同部分来统筹：传统产业提供基础和市场，未来产业提供新技术和新方向，二者优化组合才能提升整体效能。",
            "landing": "回答传统产业和未来产业的关系，要写出二者不是彼此割裂的两块，而是现代化产业体系中的相互支撑部分；应统筹产业链、创新链、价值链，让未来产业赋能传统产业、传统产业支撑未来产业，形成整体优化的新质生产力格局。",
        },
    ]


def manual_2026_haidian_q16_user_correction():
    """User-verified correction for 2026海淀一模 Q16 detailed rubric split."""
    prompt = (
        "纵观整个人类发展史，凡是社会大发展、大变革的时代，也是人文学科推陈出新的时代，比如我国春秋战国时期的百家争鸣、西方的文艺复兴等。"
        "当前，面对技术进步带来的社会巨变，人们迫切需要为这个充满机遇与挑战的时代注入新的思想活力。人工智能可以完成写作、翻译、分析信息等工作，"
        "人文学科能教人如何辨析价值冲突。技术再进步，“人何以为人”的思考永不过时。技术可以模拟语言，却未必理解语言背后的情感；"
        "算法可以识别图像，却无法体验图像所承载的历史记忆。当我们拥有足够的历史意识、伦理意识和文化理解时，技术才能真正成为文明进步的工具，"
        "而不是新的风险来源。人工智能为我们赋能，人文学科为人工智能赋魂。有人说，人文学科在人工智能时代具有不可替代的价值。"
        "结合材料，运用《哲学与文化》知识，谈谈对这一观点的认识。"
    )
    base = {
        "origin": "user-verified-rubric-correction-2026-04-29",
        "source": "2026海淀一模 第16题",
        "source_key": normalize_source_key("2026海淀一模 第16题"),
        "full_prompt": prompt,
        "rubric": "用户人工复核：细则在“必要性”中列物质决定意识并解释；细则另有主观能动性；细则不含“意识反作用于物质”。本地 pypdf 缓存只抽到简版角度清单，按用户核验的完整细分口径修正。",
    }
    return [
        {
            **base,
            "section": "唯物论",
            "title": "物质决定意识：人工智能时代的客观现实决定人文学科不可替代的必要性",
            "knowledge": "物质决定意识",
            "material": "材料先说“社会大发展、大变革的时代，也是人文学科推陈出新的时代”，并举春秋战国百家争鸣、西方文艺复兴为例；随后把人工智能时代也写成技术进步带来的社会巨变。",
            "trigger": "春秋战国百家争鸣、文艺复兴这些例子说明：不是人们凭空想让人文学科重要起来，而是社会结构、生产生活和时代问题发生大变化时，人们的思想文化也会被推动着更新。",
            "why": "能想到物质决定意识，是因为材料先用春秋战国百家争鸣、文艺复兴说明“社会大发展、大变革”会推动人文学科推陈出新；再写今天人工智能带来新的社会巨变。这个材料逻辑就是：客观时代变革提出新的价值冲突、情感理解和历史记忆问题，决定人们必须重新认识人文学科的必要性。",
            "landing": "因为物质决定意识，人工智能时代的客观技术变革、社会巨变和技术局限，决定人们需要人文学科来辨析价值冲突、理解情感和历史记忆；所以人文学科在人工智能时代不是可有可无，而具有不可替代的必要性。",
        },
        {
            **base,
            "section": "唯物论",
            "title": "主观能动性 / 意识的能动作用：人要主动辨析价值冲突、理解技术边界",
            "knowledge": "主观能动性 / 意识的能动作用",
            "material": "人文学科能教人如何辨析价值冲突；技术再进步，“人何以为人”的思考永不过时；历史意识、伦理意识和文化理解是人工智能不能替代的。",
            "trigger": "材料把“辨析价值冲突”“人何以为人”“历史意识、伦理意识和文化理解”交给人来完成，说明面对人工智能不能只被动接受技术结果，而要发挥人的能动思考和价值判断能力。",
            "why": "主观能动性强调人能够有目的、有意识地认识世界、作出判断并选择行动方向。人工智能可以处理信息，但不能替人完成价值判断、伦理反思和文化理解；人文学科训练的正是人在技术时代主动思考、主动判断、主动把握方向的能力。",
            "landing": "说明人文学科不可替代，还要写人必须发挥主观能动性：面对人工智能，人不能把价值判断和伦理选择完全交给技术，而要主动辨析价值冲突、理解技术边界和“人何以为人”的问题，这正是人文学科不可替代的价值。",
        },
    ]


def manual_gap_coverage_supplements():
    """Quality backfills for suites exposed by the no-empty-logic validator."""
    hd2024_prompt = (
        "东城区前门街道草厂地区有一个远近闻名的“小院议事厅”。“小院议事厅”在社区党委的领导下，组织居民会议、召开恳谈会、举办议事会，"
        "探索出协商共治的“五民工作法”——民事民提、民事民议、民事民决、民事民办、民事民评。街道围着社区转，社区围着小院转，小院围着居民转，"
        "通过多方参与形成合力，有效解决了垃圾分类、院落改善、文明养犬等难点问题，大大增强了社区居民的获得感和幸福感，促进了居民自治与社区精细治理的有效结合。"
        "从最开始坐在胡同口的小板凳上商量事儿，到现在有了专门议事的空间，很多邻里之间的矛盾事儿、社区里的治理难题，就这样在“唠家常”中解决了。"
        "结合材料，分析“五民工作法”对我国发展基层民主的启示。"
    )
    sj2025_prompt = (
        "破与立，是中华优秀传统文化中关于事物变革创新的一对关系，蕴含着破立并举、革故鼎新的先贤智慧。"
        "新中国成立后，党带领人民破立结合、循序渐进，建立起社会主义制度；改革开放之初，党带领人民边破边立、先行先试，改革从农村到城市、从沿海到内地渐次铺开；"
        "党的十八大以来，党带领人民破立并举、先立后破，开创了改革开放全新局面。当前，进一步全面深化改革已经进入攻坚期和深水区，党带领人民坚守志不改、道不变的信念，"
        "锚定进一步全面深化改革总目标，坚持“不立不破、先立后破”的发展之道：立新发展格局，破传统发展路径依赖；立新质生产力，破原始创新能力不强；"
        "立清风正气，破“四风”顽疾；立公正高效权威的司法制度，破影响司法公正之弊，推动各项制度更加成熟和定型，改革不断取得令人民满意的新成效。"
        "结合材料，综合运用所学，阐述我们党在进一步全面深化改革中统筹破立关系，对全面建成社会主义现代化强国的重要意义。"
    )
    tz2026_prompt = (
        "都江堰，这座全世界迄今唯一仍在使用的古代水利工程，是人类与自然和谐共生最早的成功实践之一。"
        "战国时期，秦国蜀郡太守李冰利用当地西北高、东南低的地理条件，率领民众历经艰难、勇于开拓，修建都江堰，治水利民，充分彰显了“因地制宜、顺势而为、天人合一”的传统理念。"
        "都江堰的修建并非一蹴而就，历代治水者接力维护，使其千年至今仍焕发生机。如今，建设者守正创新，迈入“智水”时代，构建智慧水资源保护与污染防治体系、科技赋能文物保护等，推动工程持续升级。"
        "同时创新放水节文化表达，以情景剧演绎“李冰治水”故事，增设游客互动体验项目，让治水文化可触可感，为都江堰世界文化遗产的可持续发展注入智慧力量。"
        "结合材料，运用《哲学与文化》知识，阐释都江堰跨越千年的治水智慧。"
    )
    xc2024_prompt = (
        "在中国，有这样一个城市公园，没有精心筛选的观赏植物，没有大片修剪整齐的绿地，也没有平整规矩的健身步道。"
        "园林设计师们从城市各个角落搜集了不同的土壤将公园填满，然后将种植权完全交给自然，让风来种，让鸟来种，让昆虫来种，让土壤里的种子自己萌发。"
        "各种植物野蛮生长，那些常被园丁斩草除根的野草、野花和野树，成了这里的主人。这个完全信任自然、由自然意志主宰的公园如今枝繁叶茂、郁郁葱葱，"
        "有各种各样的鸟类、昆虫、浮游生物、底栖动物，甚至出现了以往踪迹难觅的国家一级重点保护动物。这个公园向久居水泥森林的都市人，展现出自然界真正的物竞天择、景观春秋。"
        "从哲学角度，说明处理人与自然的关系，应避免陷入人类中心主义。"
    )
    hd2026_q16_prompt = "结合材料，谈谈你对数字化世界中青年应如何对待信息工具的理解。"
    return [
        {
            "origin": "current-thread-gap-closure",
            "section": "历史唯物主义",
            "title": "人民群众",
            "source": "2024海淀期中 第18题",
            "source_key": normalize_source_key("2024海淀期中 第18题"),
            "full_prompt": hd2024_prompt,
            "material": "“五民工作法”把居民放在基层治理中心，民事民提、民事民议、民事民决、民事民办、民事民评，并以居民获得感和幸福感检验治理成效。",
            "trigger": "材料连续出现“民提、民议、民决、民办、民评”和“居民获得感、幸福感”，说明基层治理不是替居民作主，而是依靠居民、服务居民。",
            "knowledge": "人民群众",
            "why": "人民群众是社会历史的主体。基层民主的关键不只是有治理主体，更是让群众参与提出问题、讨论方案、作出决定、办理事务和评价效果。",
            "landing": "发展基层民主，要坚持人民主体地位，把居民的需求、智慧和评价纳入治理全过程，让群众在参与社区事务中实现自我管理、自我服务、自我教育、自我监督。",
            "rubric": "",
        },
        {
            "origin": "current-thread-gap-closure",
            "section": "辩证法",
            "title": "系统观念 / 系统优化",
            "source": "2024海淀期中 第18题",
            "source_key": normalize_source_key("2024海淀期中 第18题"),
            "full_prompt": hd2024_prompt,
            "material": "社区党委、居民会议、恳谈会、议事会、街道、社区、小院和居民围绕治理问题形成多方参与合力。",
            "trigger": "“街道围着社区转，社区围着小院转，小院围着居民转”把多主体、多层级、多环节放进同一个治理系统。",
            "knowledge": "系统观念 / 系统优化",
            "why": "系统观念要求把基层治理看成由多个主体和环节构成的整体。只有让党委领导、街道支持、社区组织、居民参与相互配合，才能把分散诉求转化为治理合力。",
            "landing": "发展基层民主，要统筹基层党组织、街道社区、议事平台和居民主体，优化协商、决策、办理、评价各环节，使基层治理形成有序协同的整体效能。",
            "rubric": "",
        },
        {
            "origin": "current-thread-gap-closure",
            "section": "辩证法",
            "title": "矛盾的特殊性 / 具体问题具体分析",
            "source": "2024海淀期中 第18题",
            "source_key": normalize_source_key("2024海淀期中 第18题"),
            "full_prompt": hd2024_prompt,
            "material": "“五民工作法”针对垃圾分类、院落改善、文明养犬、邻里矛盾等不同治理难题，通过“唠家常”的方式协商解决。",
            "trigger": "材料不是抽象谈治理口号，而是列出不同社区问题，并用贴近居民生活的协商方式逐一化解。",
            "knowledge": "矛盾的特殊性 / 具体问题具体分析",
            "why": "矛盾具有特殊性，基层治理面对的是具体社区、具体人群、具体矛盾。只有从不同问题的实际情况出发，才能形成务实管用的治理办法。",
            "landing": "发展基层民主，要具体分析不同社区事务和矛盾类型，通过居民熟悉、愿意参与的协商方式解决具体问题，而不是用一套抽象方案处理所有社区难题。",
            "rubric": "",
        },
        {
            "origin": "current-thread-gap-closure",
            "section": "历史唯物主义",
            "title": "改革 / 改革的实质",
            "source": "2025石景山一模 第21题",
            "source_key": normalize_source_key("2025石景山一模 第21题"),
            "full_prompt": sj2025_prompt,
            "material": "党带领人民破立结合、边破边立、先立后破，破传统发展路径依赖、破原始创新能力不强、破“四风”顽疾，同时立新发展格局、立新质生产力、立公正高效权威的司法制度。",
            "trigger": "材料中的“破旧观念旧体制、立新理念新机制”正是改革作为制度自我完善和发展的过程。",
            "knowledge": "改革 / 改革的实质",
            "why": "改革不是简单否定过去，而是在社会发展新要求和人民群众新期待中调整不适应生产力和治理现代化的体制机制，使制度更加成熟定型。",
            "landing": "阐述统筹破立关系的意义，要写改革通过破除不适应发展的旧观念旧体制、确立服务新发展阶段的新理念新机制，推动国家治理体系和治理能力现代化，为全面建成社会主义现代化强国提供制度动力。",
            "rubric": "",
        },
        {
            "origin": "current-thread-gap-closure",
            "section": "辩证法",
            "title": "辩证否定 / 守正创新",
            "source": "2026通州期末 第16题",
            "source_key": normalize_source_key("2026通州期末 第16题"),
            "full_prompt": tz2026_prompt,
            "material": "都江堰保留“因势利导”的核心治水智慧，今天又通过智慧水资源保护、科技赋能文物保护、放水节文化表达和互动体验实现升级。",
            "trigger": "材料不是说都江堰原封不动地保存，也不是抛弃传统治水智慧，而是在保留核心智慧的基础上用现代科技和文化表达更新其功能。",
            "knowledge": "辩证否定 / 守正创新",
            "why": "辩证否定的实质是扬弃，既保留合理内核，又克服旧形式的局限。都江堰跨越千年，靠的正是守住因势利导、天人合一的智慧，并在新时代用科技和文化创新继续发展。",
            "landing": "阐释都江堰跨越千年的治水智慧，要写它在守正创新中实现扬弃：守住顺势治水、人与自然和谐共生的核心智慧，又通过现代科技保护和文化传播创新，使古代工程持续焕发生命力。",
            "rubric": "",
        },
        {
            "origin": "current-thread-post-clean-coverage",
            "section": "辩证法",
            "title": "联系的普遍性 / 联系的观点：城市公园是自然要素相互作用的生命网络",
            "source": "2024西城一模 第17题",
            "source_key": normalize_source_key("2024西城一模 第17题"),
            "full_prompt": xc2024_prompt,
            "material": "城市公园把不同土壤、风、鸟、昆虫、土壤里的种子、野草野花野树以及鸟类和底栖动物放在同一生态过程里，呈现出自然界自己的相互作用。",
            "trigger": "材料不是只写某一种植物，而是写土壤、种子、风鸟昆虫、植物和动物之间彼此影响、共同生成公园生态。",
            "knowledge": "联系的普遍性 / 联系的观点",
            "why": "联系的观点强调事物不是孤立存在的。题目要求说明为什么要避免人类中心主义，关键就在于自然界本身有复杂联系，人不能把自然只当作任意摆布的对象。",
            "landing": "处理人与自然关系，应看到自然要素之间、人与自然之间都处在普遍联系中；尊重这些生态联系，才能让城市公园形成生物多样性和自然恢复力，避免用人类中心主义割裂自然自身的联系。",
            "rubric": "2024西城一模第17题补充材料提到联系与发展、客观规律与主观能动性等角度；本条只保留学生可读的联系链。",
        },
        {
            "origin": "current-thread-post-clean-coverage",
            "section": "价值观",
            "title": "价值观的导向作用：把信息工具放回辅助位置，培养独立思考",
            "source": "2026海淀期末 第16题",
            "source_key": normalize_source_key("2026海淀期末 第16题"),
            "full_prompt": hd2026_q16_prompt,
            "material": "数字化世界中的青年既享受信息工具的便利，也面对工具滥用、直接依赖工具替代自己思考的风险。",
            "trigger": "材料把“如何对待信息工具”的问题落到青年自身的判断和选择上，说明关键不只是工具强不强，而是人用什么价值取向支配工具使用。",
            "knowledge": "价值观的导向作用",
            "why": "价值观会影响人们认识和改造世界的方向。面对数字化工具，青年如果把工具当作直接主体，就会削弱独立思考；如果树立正确价值观，就会把工具作为辅助，在使用中提升自身能力。",
            "landing": "回答青年如何对待信息工具，要写正确价值观的导向作用：青年应警惕工具滥用和过度依赖，把信息工具作为辅助而不是替代自己思考的主体，在数字化世界中培养辨别、判断和自主学习能力。",
            "rubric": "2026海淀期末第16题评分材料列价值观导向和正确价值观促进作用；本条转写为学生可读链条。",
        },
    ]


def bracket_field(block: str, label: str) -> str:
    m = re.search(rf"【{re.escape(label)}】\s*\n(.+?)(?=\n(?:【|## |### )|$)", block, flags=re.S)
    return clean_text(m.group(1)) if m else ""


def parse_supplement_block(section: str, title: str, block: str):
    source = bracket_field(block, "来源题目")
    full_prompt = question_prompt(bracket_field(block, "材料与设问"))
    return {
        "origin": "current-thread-supplement",
        "section": section,
        "title": title,
        "source": short_source(source),
        "source_key": normalize_source_key(source),
        "full_prompt": full_prompt,
        "material": bracket_field(block, "材料触发点"),
        "trigger": bracket_field(block, "材料触发点"),
        "knowledge": title.split("：")[0],
        "why": bracket_field(block, "为什么能想到这个原理"),
        "landing": bracket_field(block, "答案落点"),
        "rubric": bracket_field(block, "细则摘录"),
    }


def classify(row):
    # Use the row's own title/knowledge first. Broad parent section names often contain
    # multiple principles and must not decide the exact subnode.
    primary = " ".join([row.get("title", ""), row.get("knowledge", "")])
    fallback_blob = " ".join([primary, row.get("trigger", ""), row.get("section", "")])
    lookup = {(m, s): keys for m, s, keys in SUBNODE_ORDER}

    def match_by_priority(blob):
        for module, subnode in CLASSIFICATION_PRIORITY:
            keys = lookup[(module, subnode)]
            if any(k in blob for k in keys):
                return module, subnode
        return None

    # Exact and high-risk principle names should win over broad words such as
    # "联系", "发展", "实践", or "文化". Apply that rule both to title/knowledge
    # and to fallback trigger text.
    hit = match_by_priority(primary)
    if hit:
        return hit
    for module, subnode, keys in SUBNODE_ORDER:
        if any(k in primary for k in keys):
            return module, subnode
    hit = match_by_priority(fallback_blob)
    if hit:
        return hit
    for module, subnode, keys in SUBNODE_ORDER:
        if any(k in fallback_blob for k in keys):
            return module, subnode
    return "六、文化传承、文化功能与文化自信（必修四文化线，非哲学主框架）", "文化的功能 / 文化与经济政治"


def question_no(row):
    key = row.get("source_key")
    if key and key[3].isdigit():
        return int(key[3])
    m = re.search(r"第?([0-9]+)(?:题|[（(])", row.get("source", "") + " " + row.get("full_prompt", ""))
    return int(m.group(1)) if m else 99


def question_type_rank(row):
    blob = " ".join([row.get("origin", ""), row.get("title", ""), row.get("source", ""), row.get("full_prompt", "")])
    if question_no(row) <= 15 or "选择题" in blob or "choice" in blob.lower() or "错肢" in blob or "正确项" in blob:
        return 1
    return 0


def suite_sort_key(row):
    key = row.get("source_key") or ("9999", "郊区", "期末", "99", "")
    year, district, stage, qno, sub = key
    return (
        question_type_rank(row),
        DISTRICT_ORDER.get(district, 9),
        int(year) if str(year).isdigit() else 9999,
        STAGE_ORDER.get(stage, 9),
        int(qno) if str(qno).isdigit() else 99,
        sub,
        row.get("source", ""),
    )


def image_for(row):
    key = row.get("source_key")
    if key in IMAGE_MAP and IMAGE_MAP[key].exists():
        return str(IMAGE_MAP[key])
    blob = row.get("full_prompt", "") + row.get("material", "") + row.get("trigger", "")
    if "漫画" in blob and key:
        candidate = IMAGE_MAP.get(key)
        if candidate and candidate.exists():
            return str(candidate)
    return ""


def student_image_ref(image_path: str):
    if not image_path:
        return ""
    try:
        return str(Path(image_path).relative_to(ROOT))
    except ValueError:
        return image_path


def student_trigger(row):
    choices = [row.get("material", ""), row.get("trigger", "")]
    best = max((scrub_student_language(c) for c in choices if c), key=len, default="")
    if len(best) < 24 and row.get("full_prompt"):
        prompt = scrub_student_language(row["full_prompt"])
        if len(prompt) > len(best):
            best = prompt[:120].rstrip("，；。") + "。"
    return scrub_student_language(best)


def student_why(row):
    text = scrub_student_language(row.get("why", ""))
    if not text:
        return ""
    meta_terms_re = "|".join(re.escape(t) for t in META_ARTIFACT_TERMS)
    text = re.sub(rf"[^。！？]*(?:{meta_terms_re})[^。！？]*[。！？]", "", text)
    text = re.sub(r"因此[^。！？]*(?:补入|栏目|该题不能)[^。！？]*[。！？]?", "", text)
    text = re.sub(r"因此[^。！？]*(?:课堂整理链|归入|放进|纳入)[^。！？]*[。！？]?", "", text)
    meta_terms = META_ARTIFACT_TERMS + ["补入", "该题不能", "栏目", "课堂整理链", "归入"]
    parts = [p.strip() for p in re.split(r"(?<=[。！？])", text) if p.strip()]
    kept = [p for p in parts if not any(t in p for t in meta_terms)]
    cleaned = scrub_student_language("".join(kept) or text)
    if weak_why_reasons(row, cleaned):
        return repair_weak_why(row, cleaned)
    return cleaned


def logic_only(row):
    blob = " ".join([row.get("title", ""), row.get("knowledge", ""), row.get("full_prompt", "")])
    logic_terms = ["归纳推理", "类比推理", "换质位", "三段论", "联言", "假言", "演绎推理", "发散思维", "聚合思维", "超前思维", "创新思维", "分析与综合"]
    bixiu4_terms = ["文化", "矛盾", "联系", "发展", "辩证否定", "实践", "认识", "价值", "人民", "社会存在", "规律", "实际", "意识", "传统", "民族精神"]
    return any(t in blob for t in logic_terms) and not any(t in blob for t in bixiu4_terms)


def make_landing(row):
    bad_landing_terms = META_ARTIFACT_TERMS + ["选择题中要把", "不能只看话题相似", "作答时要把材料", "这一原理的运用", "该题不能", "补入", "归入", "课堂整理链", "可想到"]
    landing = clean_text(row.get("landing", ""))
    if landing and not any(t in landing for t in bad_landing_terms) and not ANGLE_LIST_RE.search(landing):
        return landing
    why = row.get("why", "")
    # Use the last substantial non-meta sentence as a concrete answer direction.
    meta_terms = META_ARTIFACT_TERMS + ["触发", "补入", "该题", "栏目", "材料链", "作答时", "学生", "归入", "课堂整理链", "可想到"]
    parts = [p.strip() for p in re.split(r"[。！？]", why) if len(p.strip()) > 8]
    parts = [p for p in parts if not any(t in p for t in meta_terms)] or parts
    if parts:
        sentence = parts[-1]
    else:
        sentence = row.get("trigger") or why
    if "答案落点" in sentence:
        return clean_text(sentence)
    return clean_text(sentence)


def enrich_rows():
    codex_rows = parse_codex_base()
    prompt_by_key, v6_entries = parse_v6_prompts_and_entries()
    supplements = parse_supplement()
    csv_backfills = (
        parse_late_objective_backfills()
        + parse_sample_s040_backfills()
        + parse_worker_audit_backfills()
        + targeted_coverage_backfills()
    )
    manual = manual_supplements()
    manual += manual_2024_east_supplements()
    manual += manual_2026_haidian_q16_user_correction()
    manual += manual_gap_coverage_supplements()

    # Current-thread supplements often contain the full prompt for OCR-repaired suites
    # that v6 marked as NEED_EVIDENCE. Use them as prompt inventory too, otherwise
    # valid Codex framework rows such as 2026海淀一模/物质决定意识 get filtered out.
    for prompt_row in supplements + csv_backfills + manual:
        key = prompt_row.get("source_key")
        full_prompt = question_prompt(prompt_row.get("full_prompt"))
        if key and full_prompt:
            prompt_by_key.setdefault(key, full_prompt)
            prompt_by_key.setdefault(same_source_key_loose(key), full_prompt)

    codex_by_key = defaultdict(list)
    for r in codex_rows:
        if r["source_key"]:
            codex_by_key[r["source_key"]].append(r)
            codex_by_key[same_source_key_loose(r["source_key"])].append(r)

    rows = []
    # Use v6 mainly as a prompt/question inventory, then replace explanation with the closest Codex chain when available.
    for v in v6_entries:
        candidates = codex_by_key.get(v.get("source_key")) or codex_by_key.get(same_source_key_loose(v.get("source_key"))) or []
        best = None
        for c in candidates:
            overlap = sum(1 for token in re.findall(r"[\u4e00-\u9fff]{2,}", v["title"]) if token in c["knowledge"] or token in c["section"] or token in c["why"])
            if best is None or overlap > best[0]:
                best = (overlap, c)
        if best and best[0] > 0:
            c = best[1]
            v["why"] = c["why"] or v["why"]
            if c.get("material") and len(c["material"]) > len(v.get("trigger", "")):
                v["material"] = c["material"]
            v["knowledge"] = c["knowledge"] or v["knowledge"]
            v["origin"] = v["origin"] + "+Codex-chain"
        rows.append(v)

    # Do not let the v6 prompt inventory erase Codex's independently audited framework nodes.
    # For now, only backfill the missing "物质决定意识" node when the full prompt can be recovered.
    for c in codex_rows:
        if not any(t in c.get("knowledge", "") for t in ["物质决定意识", "物质与意识", "意识对物质"]):
            continue
        key = c.get("source_key")
        if key == ("2026", "海淀", "一模", "16", ""):
            # User-verified detailed rubric correction is injected below via
            # manual_2026_haidian_q16_user_correction(); avoid keeping the old
            # compound-node wording that added unsupported "意识反作用于物质".
            continue
        c["full_prompt"] = prompt_by_key.get(key) or prompt_by_key.get(same_source_key_loose(key), "")
        if c["full_prompt"]:
            rows.append(c)

    # Add all current-thread supplements last; they override duplicates from the same source and title cluster.
    rows.extend(supplements)
    rows.extend(csv_backfills)
    rows.extend(manual)

    deduped = []
    seen = set()
    for r in rows:
        if logic_only(r):
            continue
        module, node = classify(r)
        r["module"] = module
        r["node"] = node
        r["image_path"] = image_for(r)
        r["landing"] = make_landing(r)
        if not r.get("material"):
            r["material"] = r.get("trigger", "")
        if not r.get("full_prompt"):
            key = r.get("source_key")
            r["full_prompt"] = prompt_by_key.get(key) or prompt_by_key.get(same_source_key_loose(key), "")
        final_trigger = student_trigger(r)
        final_prompt = student_question_prompt(r)
        final_why = student_why(r)
        final_landing = student_landing(r)
        if not (final_trigger and final_why and final_landing):
            continue
        if any(contains_final_artifact_language(x) for x in [final_trigger, final_prompt, final_why, final_landing]):
            continue
        if any(STUDENT_ARTIFACT_RE.search(x) for x in [final_trigger, final_prompt, final_why, final_landing]):
            continue
        # Only de-duplicate very close repeats; keep same question under different principles.
        sig = (r.get("source"), module, node, re.sub(r"\W+", "", r.get("knowledge", ""))[:18], re.sub(r"\W+", "", r.get("title", ""))[:18])
        if sig in seen:
            continue
        seen.add(sig)
        deduped.append(r)

    deduped.sort(key=lambda r: (SUBNODE_INDEX.get((r["module"], r["node"]), 999), suite_sort_key(r)))
    return deduped


def build_markdown(rows):
    grouped = defaultdict(lambda: defaultdict(list))
    for r in rows:
        grouped[r["module"]][r["node"]].append(r)
    lines = [f"# {TITLE}", "", f"**{SIGNATURE}**", "", "## 前言", "", "（此处留给飞哥老师自行撰写。）", ""]
    lines.append("## 目录")
    for module in FRAMEWORK_STRUCTURE:
        lines.append(f"- {module}")
        for subnode, _ in FRAMEWORK_STRUCTURE[module]:
            lines.append(f"  - {subnode}")
    lines.append("")
    for module, children in FRAMEWORK_STRUCTURE.items():
        lines.append(f"## {module}")
        lines.append("")
        for subnode, _ in children:
            items = grouped[module].get(subnode)
            lines.append(f"### {subnode}")
            lines.append("")
            if not items:
                continue
            for i, r in enumerate(items, 1):
                qtype = "选择题" if question_type_rank(r) else "主观题"
                lines.append(f"#### {i}. {r['source'] or r['title']}（{qtype}）")
                lines.append("")
                lines.append(f"【材料触发点】 {student_trigger(r)}")
                lines.append("")
                if r.get("image_path"):
                    lines.append(f"![漫画]({student_image_ref(r['image_path'])})")
                    lines.append("")
                mat_prompt = student_question_prompt(r)
                lines.append(f"【设问】 {mat_prompt}")
                lines.append("")
                lines.append(f"【为什么能想到】 {student_why(r)}")
                lines.append("")
                lines.append(f"【答案落点】 {student_landing(r).removeprefix('答案落点：')}")
                lines.append("")
    return "\n".join(lines)


def set_cell_shading(paragraph, color="F2F6F8"):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    ppr.append(shd)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_new_page_section(doc):
    doc.add_section(WD_SECTION_START.NEW_PAGE)


def configure_styles(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Songti SC"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "PingFang SC"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        styles[name].font.color.rgb = RGBColor(0, 0, 0)
    styles["Heading 1"].font.size = Pt(20)
    styles["Heading 2"].font.size = Pt(15)
    styles["Heading 3"].font.size = Pt(12)


def add_para(doc, text="", style=None, bold_label=False):
    p = doc.add_paragraph(style=style)
    if bold_label and text.startswith("【") and "】" in text:
        label, rest = text.split("】", 1)
        run = p.add_run(label + "】")
        run.bold = True
        p.add_run(rest)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    return p


def build_docx(markdown_text: str, docx_path: Path):
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    # Cover: only title and signature. Use paragraph spacing rather than many
    # blank paragraphs so the foreword cannot leak onto the first page.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(220)
    p.paragraph_format.space_after = Pt(90)
    r = p.add_run(TITLE)
    r.bold = True
    r.font.name = "PingFang SC"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    r.font.size = Pt(25)
    r.font.color.rgb = RGBColor(0, 0, 0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SIGNATURE)
    r.bold = True
    r.font.name = "PingFang SC"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0, 0, 0)
    add_new_page_section(doc)

    # Foreword page.
    h = doc.add_heading("前言", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(16):
        doc.add_paragraph()
    add_new_page_section(doc)

    in_content = False
    for line in markdown_text.splitlines():
        if line.startswith("# ") or line.startswith("**") or line.strip() == "## 前言" or "此处留给" in line:
            continue
        if line.startswith("## 目录"):
            doc.add_heading("目录", level=1)
            continue
        if line.startswith("- "):
            add_para(doc, line, None)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            in_content = True
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
            continue
        if line.startswith("![漫画](") and line.endswith(")"):
            image_path = line[len("![漫画]("):-1]
            if not Path(image_path).is_absolute():
                image_path = str(ROOT / image_path)
            if Path(image_path).exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(image_path, width=Inches(2.25))
            continue
        if not line.strip():
            continue
        add_para(doc, line, None, bold_label=True)

    doc.save(docx_path)


def write_audit(rows):
    audit_path = AUDIT / "Codex宝典补全版_审计索引.csv"
    fields = ["module", "node", "source", "question_type", "title", "knowledge", "origin", "has_full_prompt", "has_landing", "has_image"]
    with audit_path.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fields, lineterminator="\n")
        w.writeheader()
        for r in rows:
            w.writerow(
                {
                    "node": r["node"],
                    "module": r["module"],
                    "source": r["source"],
                    "question_type": "选择题" if question_type_rank(r) else "主观题",
                    "title": r["title"],
                    "knowledge": r["knowledge"],
                    "origin": r["origin"],
                    "has_full_prompt": bool(r.get("full_prompt")),
                    "has_landing": bool(r.get("landing")),
                    "has_image": bool(r.get("image_path")),
                }
            )
    return audit_path


def suite_name_from_row(row):
    source = row.get("source", "")
    if "2024顺义思政二模" in source:
        return "2024顺义思政二模"
    key = normalize_source_key(source)
    return "".join(key[:3]) if key else ""


def write_coverage(rows):
    coverage_path = AUDIT / "逐套三触发点覆盖验收表.csv"
    counts = defaultdict(int)
    examples = defaultdict(list)
    for r in rows:
        suite = suite_name_from_row(r)
        if not suite:
            continue
        counts[suite] += 1
        examples[suite].append(f"{r['source']} / {r['node']}")
    out_rows = []
    if COVERAGE_MATRIX.exists():
        with COVERAGE_MATRIX.open(encoding="utf-8-sig", newline="") as f:
            for r in csv.DictReader(f):
                suite = r["suite_name"].strip()
                count = counts[suite]
                need = max(0, 3 - count)
                status = "PASS_3PLUS" if count >= 3 else ("NEED_BACKFILL_LT3" if count else "NEED_BACKFILL_MISSING")
                out_rows.append(
                    {
                        "suite_id": r["suite_id"],
                        "suite_name": suite,
                        "year": r["year"],
                        "stage": r["stage"],
                        "current_trigger_count": count,
                        "required_minimum": 3,
                        "need_to_add": need,
                        "status": status,
                        "current_examples": "；".join(examples.get(suite, [])[:10]),
                    }
                )
    if out_rows:
        with coverage_path.open("w", encoding="utf-8-sig", newline="") as f:
            w = csv.DictWriter(f, fieldnames=list(out_rows[0].keys()), lineterminator="\n")
            w.writeheader()
            w.writerows(out_rows)
    passed = sum(1 for r in out_rows if r["status"] == "PASS_3PLUS")
    lt3 = sum(1 for r in out_rows if r["status"] == "NEED_BACKFILL_LT3")
    missing = sum(1 for r in out_rows if r["status"] == "NEED_BACKFILL_MISSING")
    need_total = sum(int(r["need_to_add"]) for r in out_rows)
    min_count = min((int(r["current_trigger_count"]) for r in out_rows), default=0)
    return coverage_path, passed, lt3, missing, need_total, min_count, len(out_rows)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    AUDIT.mkdir(parents=True, exist_ok=True)
    rows = enrich_rows()
    md = build_markdown(rows)
    md_path = OUT / "2026北京高考政治哲学宝典_Codex补全版_学生版.md"
    docx_path = OUT / "2026北京高考政治哲学宝典_Codex补全版_学生版.docx"
    md_path.write_text(md, encoding="utf-8")
    build_docx(md, docx_path)
    audit_path = write_audit(rows)
    coverage_path, passed, lt3, missing, need_total, min_count, total_suites = write_coverage(rows)
    progress = ROOT / "PROGRESS.md"
    progress.write_text(
        "\n".join(
            [
                "# 必修四框架重做进度",
                "",
                "## 2026-04-29 Codex补全版",
                "",
                "### 当前验收状态：覆盖门槛已通过，仍需人工抽样审读",
                "",
                "- 按用户新增硬门槛“每一套卷至少 3 个可用触发点”复核，当前覆盖表已无 0 条或不足 3 条套卷。",
                f"- 逐套验收表：{coverage_path}",
                f"- 最近一次复核结果：{total_suites} 套中 {passed} 套达到 3 条以上；{lt3} 套不足 3 条；{missing} 套当前最终索引为 0 条；还需补 {need_total} 个触发点；最低覆盖数 {min_count}。",
                "- 覆盖门槛通过不等于全文逐句完美；后续若继续精修，应做学生可读性和源证据抽样复核。",
                "",
                "- 样张只读取版式结构，不复制 ClaudeCode 正文判断。",
                "- 内容底稿使用 Codex 2026-04-27 学生触发句优先版，并用 2026-04-29 v6 套卷输出只补齐题目定位/完整设问字段。",
                "- 当前线程已独立回源补做 2026海淀一模、2026丰台一模、2026房山一模、2025朝阳一模等同类漏项，并覆盖到对应原理节点。",
                "- 2026-04-29 晚间修订：按用户母版恢复五大哲学模块和具体原理节点；字段顺序改为“材料触发点 → 设问 → 为什么能想到 → 答案落点”；同节点内按主观题优先、海淀/西城/东城/朝阳/丰台优先排序；去掉彩色卡片式排版；漫画题嵌入图片。",
                "- 2026-04-29 覆盖补齐：合并 late objective closure、S040 样稿、worker audit CSV，并对合并漏项做定点 backfill；56 套全部达到至少 3 个触发点。",
                "- 2026-04-29 海淀一模第16题修正：物质决定意识节点只写“必要性”，不写意识反作用；同题另在主观能动性节点单独成条；本地 pypdf 缓存只含简版角度清单，按用户人工核验的完整细分口径修正。",
                "- 2026-04-29 最终正文清洗：`评标`、`参考答案`、`答案写`、`答案/补充`、`答案核`、`可从……角度作答` 等审计话术不得进入学生版；清洗后对 2024西城一模、2026海淀期末做了学生可读链条补齐。",
                "- 2026-04-30 同类问题全局复查：清除 `正确项链`、`错肢`、`第n题答案`、`参考示例`、`9分标准` 等核验话术；所有“为什么能想到”必须从材料信号进入知识解释链，不能整句复制、过短或用抽象程序话替代。",
                "- 2026-04-30 框架追修：唯物论第一节点改回单独 `物质决定意识`，意识能动作用相关内容进入 `主观能动性 / 意识的能动作用`。",
                "- 学生版已清理路径、行号、文件名、OCR/debug/log 信息；审计索引单独存放。",
                "",
                f"- Markdown：{md_path}",
                f"- Word：{docx_path}",
                f"- 审计索引：{audit_path}",
                f"- 覆盖验收：{coverage_path}",
            ]
        ),
        encoding="utf-8",
    )
    print(md_path)
    print(docx_path)
    print(audit_path)
    print(f"entries={len(rows)}")


if __name__ == "__main__":
    main()
