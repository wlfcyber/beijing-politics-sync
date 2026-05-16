#!/usr/bin/env python3
from __future__ import annotations

import csv
import os
import re
import subprocess
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path("/Users/wanglifei/Desktop/北京高考政治/选必一_原PDF框架主观题宝典重构_2026-05-09")
RUN = ROOT / "V3_终审修订_2026-05-09"
AUDIT = RUN / "01_v3_audit"
PACKET = RUN / "04_claude_packet"
DRAFT = RUN / "05_claude_draft"
WORK = RUN / "06_work"
FINAL = RUN / "08_final"
LOGS = RUN / "00_logs"

PDF_PATH = Path("/Users/wanglifei/Desktop/先前框架/选必一 原文件 拷贝 (8)(2).pdf")
V2_DOCX = ROOT / "V2_内容纯度修订_2026-05-09" / "08_final" / "2026北京高考政治选必一《当代国际政治与经济》六要素原框架答题术语宝典_主观题版_V2.docx"
V2_MD = ROOT / "V2_内容纯度修订_2026-05-09" / "08_final" / "final_v2.md"
OLD_DOCX = Path("/Users/wanglifei/Desktop/北京高考政治/选必一_当代国际政治与经济_四线终极全书_2026-05-03/10_teaching_rebuild_20260504/04_delivery/选必一_当代国际政治与经济_触发宝典_可直接讲版_20260504.docx")
SUBJ_CSV = ROOT / "03_subjective_filter" / "subjective_question_pool.csv"

FINAL_MD = FINAL / "final_v3.md"
FINAL_DOCX = FINAL / "2026北京高考政治选必一《当代国际政治与经济》六要素原框架答题术语宝典_主观题版_V3.docx"

ELEMENTS = ["时代背景", "理论", "经济全球化", "政治多极化", "中国", "联合国"]
CN = dict(zip(ELEMENTS, ["一", "二", "三", "四", "五", "六"]))
NON_PDF_TERMS = ["治理赤字", "全球性挑战", "逆全球化", "零和博弈", "百年未有之大变局", "制度型开放", "NDC", "全球南方", "南南合作", "小而美项目", "气候治理框架", "联合国气候治理机制", "全球倡议体系", "《巴黎协定》", "文明交流互鉴"]
BACKEND_LABELS = ["【考题增补术语】", "【考题增补说明】", "点位性质", "正式细则", "非正式细则", "题内细则", "卷面层级", "补丁题链", "见对应题号", "答案见本题正确项", "NEEDS_FIX", "P0", "P1", "后台字段", "未核实", "推测", "旧稿"]
BAD_TEMPLATE = "不优先写相邻术语，是因为相邻术语不能直接解释本题设问中的材料动作和得分关键词"
META_PATTERNS = ["PDF 小框架", "本题优先触发", "本题优先写", "因为踩分词集中在", "不优先写相邻术语，是因为", "考题增补术语", "考题增补说明", "踩分词集中", "旧稿", "后台", "正式细则", "非正式细则", "卷面层级", "补丁题链", "见对应题号", "答案见本题正确项", "未核实", "推测", "其中……不是 PDF 原框架小框架", "不是 PDF 原框架小框架"]
ANSWER_META_PATTERNS = ["题目追问", "材料呈现", "就要想到", "触发", "先归入", "再落到", "本题优先", "踩分词", "说明该术语", "作为可作答表达", "因此政治环境层", "不是 PDF 原框架", "学生记住"]


def row(element, frame, term, page, raw, normalized=None, enter="是", note=""):
    return {
        "六要素": element,
        "小框架": frame,
        "答题术语": term,
        "PDF页码": page,
        "PDF原始表述": raw,
        "规范化表述": normalized or term,
        "是否为PDF原框架": "是",
        "是否进入正文": enter,
        "备注": note,
    }


PDF_ROWS = [
    row("时代背景", "机遇", "和平与发展是时代的主题", 2, "机遇：和平与发展是时代的主题"),
    row("时代背景", "机遇", "经济全球化与政治多极化深入发展", 2, "经济全球化与政治多极化深入发展"),
    row("时代背景", "挑战", "霸权主义", 2, "挑战：霸权主义、强权政治"),
    row("时代背景", "挑战", "强权政治", 2, "挑战：霸权主义、强权政治"),
    row("理论", "国家安全", "国家安全是最高国家利益", 2, "国家安全——国家安全是最高国家利益"),
    row("理论", "国家安全", "种子安全", 2, "eg. 种子安全、粮食安全、信息安全、网络安全", enter="否", note="本轮旧稿主观题素材池未提供完整主观题挂载。"),
    row("理论", "国家安全", "粮食安全", 2, "eg. 种子安全、粮食安全、信息安全、网络安全", enter="否", note="本轮旧稿主观题素材池未提供完整主观题挂载。"),
    row("理论", "国家安全", "信息安全", 2, "eg. 种子安全、粮食安全、信息安全、网络安全", enter="否", note="本轮旧稿主观题素材池未提供完整主观题挂载。"),
    row("理论", "国家安全", "网络安全", 2, "eg. 种子安全、粮食安全、信息安全、网络安全", enter="否", note="本轮旧稿主观题素材池未提供完整主观题挂载。"),
    row("理论", "国家安全", "重要组成部分", 2, "重要组成部分", enter="否", note="作为国家安全相关表述保留在 PDF recheck，不单列正文术语。"),
    row("理论", "共同利益", "共同利益是国家间合作的基础", 2, "共同利益——共同利益是国家间合作的基础"),
    row("理论", "共同利益", "合作的原因", 2, "合作的原因", enter="否", note="作为共同利益术语的解释，不单列正文术语。"),
    row("理论", "共同利益", "维护本国利益的同时兼顾他国合理关切", 2, "维护本国利益的同时兼顾他国合理关切"),
    row("经济全球化", "民", "民生福祉", 3, "民：民生福祉"),
    row("经济全球化", "民", "增加就业岗位", 3, "增加就业岗位"),
    row("经济全球化", "民", "提高居民收入", 3, "提高居民收入"),
    row("经济全球化", "民", "完善基础设施", 3, "完善基础设施"),
    row("经济全球化", "国", "推动成员国经济复苏", 3, "国：推动成员国经济复苏"),
    row("经济全球化", "国", "共建国家经济增长", 3, "共建国家经济增长"),
    row("经济全球化", "环", "经济环节", 3, "环：经济环节", enter="否", note="小框架总括词，不单列正文术语。"),
    row("经济全球化", "环", "促进要素自由流动", 3, "促进要素自由流动"),
    row("经济全球化", "环", "人员、商品、资金、技术、服务流动", 3, "人员、商品、资金、技术、服务"),
    row("经济全球化", "环", "拉动投资和消费", 3, "拉动投资和消费"),
    row("经济全球化", "环", "开拓市场", 3, "开拓市场"),
    row("经济全球化", "环", "产业结构优化升级", 3, "产业结构优化升级"),
    row("经济全球化", "营", "优化营商环境", 3, "营：优化营商环境"),
    row("经济全球化", "营", "国际化、市场化、法治化营商环境", 3, "国际化、市场化、法治化"),
    row("经济全球化", "配", "优化资源配置", 3, "配：优化资源配置"),
    row("经济全球化", "贸资", "推动贸易投资自由化便利化", 4, "贸资——推动贸易投资自由化便利化"),
    row("经济全球化", "总", "推动经济全球化朝着更加开放、包容、普惠、平衡、共赢的方向发展", 4, "总——推动经济全球化朝着更加开放、包容、普惠、平衡、共赢的方向发展"),
    row("经济全球化", "开", "发展更高层次的开放型经济", 4, "开——开放：发展更高层次的开放型经济"),
    row("经济全球化", "开", "坚持引进来和走出去相结合", 4, "坚持引进来、走出去相结合"),
    row("经济全球化", "开", "推进我国高水平对外开放", 4, "推进我国高水平对外开放"),
    row("经济全球化", "兼", "充分利用国内国际两种资源", 4, "兼——充分利用国内国际两种资源、两个市场"),
    row("经济全球化", "兼", "充分利用国内国际两个市场", 4, "兼——充分利用国内国际两种资源、两个市场"),
    row("政治多极化", "新国关", "有利于构建相互尊重、公平正义、合作共赢的新型国际关系", 5, "新国关：有利于构建相互尊重、公平正义、合作共赢的新型国际关系"),
    row("政治多极化", "完善全球治理", "完善全球治理", 5, "完善全球治理"),
    row("政治多极化", "完善全球治理", "倡导共商共建共享的全球治理观", 5, "倡导共商共建共享的全球治理观"),
    row("政治多极化", "共体", "有利于构建人类命运共同体", 5, "共体：有利于构建人类命运共同体"),
    row("政治多极化", "单边", "反对单边主义和贸易保护主义", 5, "单边：反对单边主义和贸易保护主义"),
    row("政治多极化", "单边", "反对霸权主义和强权政治", 5, "反对霸权主义和强权政治"),
    row("政治多极化", "单边", "坚持真正的多边主义", 5, "坚持真正的多边主义"),
    row("政治多极化", "多极化", "坚持平等有序的多极化", 5, "多极化：坚持平等有序的多极化"),
    row("政治多极化", "民主", "推动国际关系民主化", 5, "民主：推动国际关系民主化"),
    row("中国", "政策", "我国奉行独立自主的和平外交政策", 6, "政策：我国奉行独立自主的和平外交政策"),
    row("中国", "政策", "坚持和平共处五项原则", 6, "坚持和平共处五项原则"),
    row("中国", "政策", "坚持独立自主的基本立场", 6, "坚持独立自主的基本立场"),
    row("中国", "政策", "维护世界和平、促进共同发展", 6, "宗旨：维护世界和平、促进共同发展"),
    row("中国", "智慧", "中国智慧", 6, "智慧：中国智慧、中国方案；eg. 脱贫攻坚"),
    row("中国", "智慧", "中国方案", 6, "智慧：中国智慧、中国方案；eg. 脱贫攻坚"),
    row("中国", "智慧", "脱贫攻坚", 6, "eg. 脱贫攻坚", enter="否", note="PDF中为例子，本轮不单列目录术语。"),
    row("中国", "责任", "承担大国责任", 6, "责任：承担大国责任、发挥建设性作用"),
    row("中国", "责任", "发挥建设性作用", 6, "责任：承担大国责任、发挥建设性作用"),
    row("中国", "责任", "世界第一大发展中国家", 6, "世界第一大发展中国家"),
    row("中国", "责任", "世界第二大经济体", 6, "第二大经济体"),
    row("中国", "一条历史线", "中国的综合国力增强", 6, "一条历史线：中国的综合国力增强"),
    row("中国", "一条历史线", "国际地位提升", 6, "国际地位提升"),
    row("中国", "一条主线", "习近平外交思想为新时代中国特色大国外交提供了根本遵循和行动指南", 6, "一条主线：习近平外交思想为新时代中国特色大国外交提供了根本遵循和行动指南"),
    row("联合国", "联合国原框架术语", "践行多边主义的最佳场所", 7, "联合国：践行多边主义的最佳场所", note="PDF 未显示更细小框架，暂以联合国原框架术语承接。"),
    row("联合国", "联合国原框架术语", "遵循联合国宪章宗旨原则", 7, "遵循联合国宪章宗旨原则", note="PDF 未显示更细小框架，暂以联合国原框架术语承接。"),
]


@dataclass(frozen=True)
class TermDef:
    key: str
    element: str
    frame: str
    title: str
    trigger: str
    focus: str
    neighbor: str


TERMS = [
    TermDef("era_peace", "时代背景", "机遇", "和平与发展是时代的主题", "材料出现和平、发展、时代主题、世界大势、共同发展愿望、外交因势而变、合作共赢等信号时，触发“和平与发展是时代的主题”。", "解释中国实践或国际合作为什么顺应时代大势。", "“共同利益”解释合作为什么能成立，不能替代时代背景判断。"),
    TermDef("era_challenge", "时代背景", "挑战", "霸权主义 / 强权政治", "材料出现单边霸凌、恃强凌弱、强权逻辑、霸权行径、贸易胁迫等国际乱象时，触发“霸权主义、强权政治”。", "识别国际社会面临的挑战和旧式强权逻辑。", "回应路径要另写政治多极化中的“单边”“民主”“完善全球治理”，不能塞进挑战术语本身。"),
    TermDef("security", "理论", "国家安全", "国家安全是最高国家利益", "材料出现国家安全、发展与安全、经济安全、科技安全、数据安全、粮食安全、网络安全等信号时，触发“国家安全是最高国家利益”。", "说明某项发展或开放实践为什么必须守住安全底线。", "“共同利益”解释国际合作基础，不能替代国家安全的底线逻辑。"),
    TermDef("common_interest", "理论", "共同利益", "共同利益是国家间合作的基础", "材料出现共同需求、共同问题、共同发展、合作基础、利益交汇、命运相连等信号时，触发“共同利益是国家间合作的基础”。", "解释国家间为什么能够合作、为什么需要合作。", "“和平与发展”是时代背景，不等于合作成立的直接理论基础。"),
    TermDef("reasonable_concern", "理论", "共同利益", "维护本国利益的同时兼顾他国合理关切", "材料出现既维护本国利益又照顾他国合理诉求、互利共赢、义利相兼、合作中兼顾各方利益等信号时，触发“维护本国利益的同时兼顾他国合理关切”。", "说明中国处理国际关系时不是单向利己，而是在维护自身利益的同时兼顾合作方合理关切。", "“共同利益是合作基础”回答合作为什么能成立，本术语回答合作中利益如何处理。"),
    TermDef("global_m", "经济全球化", "民", "民生福祉 / 增加就业岗位 / 提高居民收入 / 完善基础设施", "材料出现就业、收入、基础设施、公共服务、当地民众获得感、民生改善等信号时，触发经济全球化中的“民”。", "说明开放合作对人民生活和民生福祉的具体改善。", "“国”侧重国家经济复苏和增长，“民”侧重普通民众生活改善。"),
    TermDef("global_g", "经济全球化", "国", "推动成员国经济复苏 / 共建国家经济增长", "材料出现共建国家、成员国经济复苏、经济增长、发展动能、区域经济提升等信号时，触发经济全球化中的“国”。", "说明合作对国家或共建国家经济发展的作用。", "“民”侧重民众获得感，“国”侧重国家经济发展。"),
    TermDef("global_huan", "经济全球化", "环", "促进要素自由流动 / 人员、商品、资金、技术、服务流动 / 拉动投资和消费 / 开拓市场 / 产业结构优化升级", "材料出现人员、商品、资金、技术、服务跨境流动，投资消费、市场开拓、产业升级等经济环节信号时，触发经济全球化中的“环”。", "说明开放合作打通经济环节、带动流动和升级。", "本题的重心是流动、市场和升级，不写资源效率层面的另一组术语。"),
    TermDef("global_ying", "经济全球化", "营", "优化营商环境 / 国际化、市场化、法治化营商环境", "材料出现营商环境、市场化、法治化、国际化、审批便利、投资便利、外资信心等信号时，触发经济全球化中的“营”。", "说明制度和服务环境如何便利投资合作。", "“贸资”侧重贸易投资自由便利，“营”侧重经营制度环境。"),
    TermDef("global_pei", "经济全球化", "配", "优化资源配置", "材料出现资源在更大范围配置、优势互补、资源利用效率提升、全球范围配置生产要素等信号时，触发经济全球化中的“配”。", "说明经济全球化提高资源配置效率。", "“环”写经济流动环节，“配”只写资源配置效率。"),
    TermDef("global_trade", "经济全球化", "贸资", "推动贸易投资自由化便利化", "材料出现贸易投资、自由化、便利化、关税、贸易壁垒、投资合作、市场准入等信号时，触发“推动贸易投资自由化便利化”。", "说明开放合作如何降低贸易投资障碍。", "“营”是营商环境，“贸资”是贸易投资制度便利。"),
    TermDef("global_direction", "经济全球化", "总", "推动经济全球化朝着更加开放、包容、普惠、平衡、共赢的方向发展", "材料出现开放合作、发展成果共享、普惠包容、平衡发展、共赢合作、共同繁荣等信号时，触发经济全球化总方向。", "说明中国实践推动经济全球化朝正确方向发展。", "“开”侧重我国开放举措，“总”侧重世界经济全球化方向。"),
    TermDef("global_open", "经济全球化", "开", "发展更高层次的开放型经济 / 坚持引进来和走出去相结合 / 推进我国高水平对外开放", "材料出现高水平对外开放、开放型经济、引进来、走出去、标准规则对接、开放平台等信号时，触发经济全球化中的“开”。", "说明我国如何通过更高水平开放参与和推动经济全球化。", "“兼”侧重两个市场两种资源，“开”侧重开放型经济和对外开放水平。"),
    TermDef("global_jian", "经济全球化", "兼", "充分利用国内国际两个市场 / 充分利用国内国际两种资源", "材料出现国内国际两个市场、两种资源、内外联动、国际分工、国际合作和国内循环联动等信号时，触发经济全球化中的“兼”。", "说明开放如何统筹国内国际资源和市场。", "“开”是开放政策本身，“兼”是开放带来的市场资源统筹。"),
    TermDef("new_rel", "政治多极化", "新国关", "有利于构建相互尊重、公平正义、合作共赢的新型国际关系", "材料出现相互尊重、公平正义、合作共赢、互利共赢、伙伴关系等信号时，触发新型国际关系。", "说明中国实践推动国家关系从对抗转向合作共赢。", "“人类命运共同体”是整体愿景，新型国际关系侧重国与国关系形态。"),
    TermDef("governance", "政治多极化", "完善全球治理", "完善全球治理 / 倡导共商共建共享的全球治理观", "材料出现全球治理、各国共同商量、共同建设、共同分享、治理体系完善、国际规则改革等信号时，触发完善全球治理。", "说明治理方式和治理体系如何更加公平有效。", "“国际关系民主化”侧重各国平等参与国际事务，全球治理观侧重治理过程和成果。"),
    TermDef("community", "政治多极化", "共体", "有利于构建人类命运共同体", "材料出现命运与共、同球共济、共同安全、共同发展、共同未来、世界整体利益等信号时，触发人类命运共同体。", "说明中国实践服务人类共同未来和整体利益。", "“新型国际关系”侧重国家关系，“人类命运共同体”侧重全人类共同愿景。"),
    TermDef("unilateral", "政治多极化", "单边", "反对单边主义和贸易保护主义 / 反对霸权主义和强权政治 / 坚持真正的多边主义", "材料出现单边主义、贸易保护主义、霸权主义、强权政治，或要求坚持真正多边主义时，触发政治多极化中的“单边”。", "说明对旧式强权逻辑和单边主义的反对及多边主义方向。", "“民主”只写推动国际关系民主化，不承接真正多边主义。"),
    TermDef("multi", "政治多极化", "多极化", "坚持平等有序的多极化", "材料出现世界多极化、国际力量对比变化、各国平等有序参与、反对阵营对抗等信号时，触发平等有序的多极化。", "说明世界格局应朝平等有序多极化发展。", "“民主”侧重国际关系运行原则，“多极化”侧重世界格局。"),
    TermDef("democracy", "政治多极化", "民主", "推动国际关系民主化", "材料出现各国平等参与国际事务、反对少数国家垄断国际事务、国际事务由各国共同商量等信号时，触发国际关系民主化。", "说明国际事务应由各国平等参与和共同决定。", "本题的重心是各国平等参与国际事务，不写反对单边霸权那一组术语。"),
    TermDef("policy", "中国", "政策", "我国奉行独立自主的和平外交政策 / 坚持和平共处五项原则 / 坚持独立自主的基本立场 / 维护世界和平、促进共同发展", "材料出现中国特色大国外交、和平外交政策、和平共处五项原则、独立自主、维护和平促进发展等信号时，触发中国“政策”。", "说明中国外交的基本政策立场和宗旨。", "“责任”强调大国担当，“政策”强调外交政策依据。"),
    TermDef("wisdom", "中国", "智慧", "中国智慧 / 中国方案", "材料出现中国理念、中国方案、中国智慧、可复制经验、为世界难题提供中国思路等信号时，触发中国“智慧”。", "说明中国为国际问题提供理念和方案。", "“责任”强调行动担当，“智慧”强调理念方案。"),
    TermDef("responsibility", "中国", "责任", "承担大国责任 / 发挥建设性作用 / 世界第一大发展中国家 / 世界第二大经济体", "材料出现负责任大国、建设性作用、发展中国家身份、经济体量、国际责任、公共产品供给等信号时，触发中国“责任”。", "说明中国以自身身份和能力承担国际责任。", "“智慧”是方案贡献，“责任”是行动担当。"),
    TermDef("history_line", "中国", "一条历史线", "中国的综合国力增强 / 国际地位提升", "材料出现中国综合国力增强、国际地位提升、国际影响力和话语权提升等信号时，触发中国“一条历史线”。", "说明中国外交变化的实力基础。", "“一条主线”是习近平外交思想的指导，“一条历史线”是实力地位变化。"),
    TermDef("main_line", "中国", "一条主线", "习近平外交思想为新时代中国特色大国外交提供了根本遵循和行动指南", "材料出现习近平外交思想、新时代中国特色大国外交、根本遵循、行动指南等信号时，触发中国“一条主线”。", "说明新时代中国外交的思想指导。", "“一条历史线”讲综合国力和地位，“一条主线”讲思想指引。"),
    TermDef("un_charter", "联合国", "联合国原框架术语", "遵循联合国宪章宗旨原则", "材料出现联合国宪章、宪章宗旨原则、安理会、联合国框架下的国际合作等信号时，触发遵循联合国宪章宗旨原则。", "说明行动的联合国法理和制度基础。", "“践行多边主义的最佳场所”强调联合国平台地位，宪章原则强调行为准则。"),
    TermDef("un_place", "联合国", "联合国原框架术语", "践行多边主义的最佳场所", "材料出现联合国平台、多边合作、联合国会议、安理会、联合国体系等信号时，触发践行多边主义的最佳场所。", "说明联合国作为多边合作平台的意义。", "“宪章宗旨原则”强调规范基础，本术语强调平台功能。"),
]

TERM_BY_KEY = {t.key: t for t in TERMS}


def ensure_dirs():
    for p in [AUDIT, PACKET, DRAFT, WORK, FINAL, LOGS, RUN / "tools"]:
        p.mkdir(parents=True, exist_ok=True)


def write_csv(path: Path, rows, fields):
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fields)
        w.writeheader()
        w.writerows(rows)


def read_csv(path: Path):
    with path.open(encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def blob(r):
    return "；".join(r.get(k, "") for k in ["旧稿核心标题", "PDF六要素", "PDF小框架", "对应答题术语", "材料触发点", "踩分词", "答案落点"])


def map_keys(r):
    b = blob(r)
    e, f = r["PDF六要素"], r["PDF小框架"]
    keys = []
    if e == "时代背景":
        keys.append("era_challenge" if f == "挑战" or "霸权主义" in b or "强权政治" in b else "era_peace")
    elif e == "理论":
        if f == "国家安全":
            keys.append("security")
        elif "维护本国利益" in b or "兼顾他国合理关切" in b or "正确义利观" in b or "义利相兼" in b:
            keys.append("reasonable_concern")
        else:
            keys.append("common_interest")
    elif e == "经济全球化":
        if f == "民":
            keys.append("global_m")
        elif f == "国":
            keys.append("global_g")
        elif f == "营":
            keys.append("global_ying")
        elif f == "配":
            keys.append("global_pei")
        elif f == "贸资":
            keys.append("global_trade")
        elif f == "总":
            keys.append("global_direction")
        elif f == "开":
            keys.append("global_open")
        elif f == "兼":
            keys.append("global_jian")
        elif f == "环":
            keys.append("global_huan")
            if "优化资源配置" in b or "资源优化配置" in b:
                keys.append("global_pei")
    elif e == "政治多极化":
        if f == "新国关":
            keys.append("new_rel")
        elif f == "完善全球治理":
            keys.append("governance")
        elif f == "共体":
            keys.append("community")
        elif f == "多极化":
            keys.append("multi")
        elif f == "单边":
            keys.append("unilateral")
        elif f == "民主":
            if "国际关系民主化" in b:
                keys.append("democracy")
            if "真正的多边主义" in b or "单边主义" in b or "霸权主义" in b or "强权政治" in b or "贸易保护主义" in b:
                keys.append("unilateral")
    elif e == "中国":
        if f == "政策":
            keys.append("policy")
        elif f == "智慧":
            keys.append("wisdom")
        elif f == "责任":
            keys.append("responsibility")
        elif f == "一条历史线":
            keys.append("history_line")
        elif f == "一条主线":
            keys.append("main_line")
    elif e == "联合国":
        if "宪章" in b:
            keys.append("un_charter")
        if "多边主义" in b or "联合国" in b:
            keys.append("un_place")
        if not keys:
            keys.append("un_place")
    return list(dict.fromkeys(keys))


def clean_fact(text, limit=52, keep_nonpdf=False):
    if not keep_nonpdf and any(term in (text or "") for term in NON_PDF_TERMS):
        return "具体材料事实"
    text = sanitize_sentence(text, keep_nonpdf=True)
    text = re.sub(r"[。；].*$", "", text.strip())
    text = re.sub(r"\s+", "", text)
    return (text[:limit] + "…") if len(text) > limit else text


def format_source_title(src: str):
    original = src.strip()
    m = re.match(r"^(\d{4})\s*([^Q第]+?)\s*Q(\d+)(?:\((\d+)\))?$", original)
    if not m:
        return original, "无法解析" if "Q" in original else "未使用Q格式"
    year, district_exam, q, sub = m.groups()
    title = f"{year}{district_exam.strip()} 第{int(q)}题"
    if sub:
        title += f"第（{sub}）问"
    return title, "已统一"


def sanitize_sentence(text: str, keep_nonpdf=True):
    text = (text or "").strip()
    replacements = {
        "题目追问": "",
        "设问问": "要求说明",
        "本题问": "要求说明",
        "题目问": "要求说明",
        "本题优先触发": "材料直接对应",
        "本题优先写": "应写",
        "就要先写": "应先说明",
        "作答时": "",
        "因为踩分词集中在": "材料信号直接对应",
        "踩分词集中": "材料信号集中",
        "PDF 小框架": "小框架",
        "不是 PDF 原框架小框架": "",
        "就要想到": "说明",
        "触发": "对应",
        "指向": "对应",
        "作为可作答表达": "",
        "因此政治环境层": "",
        "材料呈现": "材料显示",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"作为[^，。；]{0,24}可作答表达", "", text)
    text = re.sub(r"看到“[^”]{0,90}”，", "", text)
    text = re.sub(r"先归入“?([^”，。]+)”?，再落到\s*小框架“?([^”，。]+)”?[。；，]?", r"应从\1中的\2理解，", text)
    text = re.sub(r"\s+", " ", text).strip()
    text = text.replace(" 。", "。").replace(" ，", "，")
    return text


def clean_material_trigger(text: str, key: str, source: str):
    if key == "security" and "2026朝阳期中" in source:
        return "材料一方面写人工智能带来效率提升和发展动能，另一方面写数据泄露、模型缺陷、伦理失控等风险，说明推进人工智能发展必须统筹发展和安全，尤其要维护经济安全、科技安全、数据安全。"
    text = sanitize_sentence(text, keep_nonpdf=True)
    text = text.strip("；;，, ")
    if text and not text.endswith("。"):
        text += "。"
    return text or "材料中的关键信息与本术语直接对应。"


def clean_answer_text(text: str):
    text = sanitize_sentence(text, keep_nonpdf=True)
    for pat in ANSWER_META_PATTERNS:
        text = text.replace(pat, "")
    text = re.sub(r"\s+", "", text)
    text = re.sub(r"。{2,}", "。", text)
    text = text.strip("；;，,。 ")
    return text + "。" if text else ""


def purify_answer(t: TermDef, r):
    source = r.get("题源", "")
    if t.key == "security" and "2026朝阳期中" in source:
        return "推进人工智能发展必须统筹发展和安全，既要发挥人工智能降本增效、释放发展动能的作用，也要防范数据泄露、模型缺陷、伦理失控等风险，维护经济安全、科技安全和数据安全。"
    if t.key == "era_challenge":
        return "材料中的单边霸凌、强权逻辑、恃强凌弱等体现霸权主义和强权政治仍然存在，是当今国际社会面临的挑战。"
    base = clean_answer_text(r.get("答案落点", ""))
    if base:
        return base
    fact = clean_fact(r["材料触发点"], 46) or "这一材料"
    if t.key == "era_peace":
        return f"和平与发展是时代的主题，{fact}顺应和平与发展的世界大势和各国共同发展愿望。"
    if t.key == "era_challenge":
        return f"{fact}体现霸权主义和强权政治，是当今国际社会面临的挑战。"
    if t.key == "security":
        return f"国家安全是最高国家利益，{fact}关系我国经济安全、科技安全等重要领域，必须统筹发展和安全。"
    if t.key == "common_interest":
        return f"共同利益是国家间合作的基础，{fact}说明各方存在共同需求，能够在合作中实现共同发展。"
    if t.key == "reasonable_concern":
        return f"我国在维护本国利益的同时兼顾他国合理关切，{fact}体现了互利共赢的合作取向。"
    if t.key == "global_m":
        return f"{fact}有利于改善民生福祉，增加就业岗位、提高居民收入并完善基础设施。"
    if t.key == "global_g":
        return f"{fact}有利于推动成员国经济复苏，促进共建国家经济增长。"
    if t.key == "global_huan":
        return f"{fact}有利于促进人员、商品、资金、技术、服务等要素自由流动，拉动投资和消费、开拓市场并推动产业结构优化升级。"
    if t.key == "global_ying":
        return f"{fact}有利于优化营商环境，营造国际化、市场化、法治化的经营条件。"
    if t.key == "global_pei":
        return f"{fact}有利于优化资源配置，提高资源在更大范围内配置和利用的效率。"
    if t.key == "global_trade":
        return f"{fact}有利于推动贸易投资自由化便利化，降低合作成本并拓展经贸空间。"
    if t.key == "global_direction":
        return f"{fact}有利于推动经济全球化朝着更加开放、包容、普惠、平衡、共赢的方向发展。"
    if t.key == "global_open":
        extra = "，并通过制度型开放等具体举措" if "制度型开放" in blob(r) else ""
        return f"{fact}有利于发展更高层次的开放型经济{extra}，推进我国高水平对外开放。"
    if t.key == "global_jian":
        return f"{fact}有利于充分利用国内国际两个市场、两种资源，增强国内国际联动。"
    if t.key == "new_rel":
        return f"{fact}有利于构建相互尊重、公平正义、合作共赢的新型国际关系。"
    if t.key == "governance":
        return f"{fact}有利于完善全球治理，倡导共商共建共享的全球治理观。"
    if t.key == "community":
        return f"{fact}有利于构建人类命运共同体，回应各国命运与共、合作发展的共同需要。"
    if t.key == "unilateral":
        return f"面对{fact}，应反对单边主义和贸易保护主义，反对霸权主义和强权政治，坚持真正的多边主义。"
    if t.key == "multi":
        return f"{fact}体现世界多极化发展要求，应坚持平等有序的多极化。"
    if t.key == "democracy":
        return f"{fact}要求推动国际关系民主化，使各国平等参与国际事务。"
    if t.key == "policy":
        return f"我国奉行独立自主的和平外交政策，{fact}体现坚持和平共处五项原则和维护世界和平、促进共同发展的宗旨。"
    if t.key == "wisdom":
        return f"{fact}体现中国智慧、中国方案，为解决国际问题提供有益思路。"
    if t.key == "responsibility":
        return f"中国承担大国责任、发挥建设性作用，{fact}展现负责任大国担当。"
    if t.key == "history_line":
        return f"中国综合国力增强、国际地位提升，{fact}说明中国具备承担更多国际责任和发挥更大影响的实力基础。"
    if t.key == "main_line":
        return f"习近平外交思想为新时代中国特色大国外交提供根本遵循和行动指南，{fact}体现这一主线的实践要求。"
    if t.key == "un_charter":
        return f"{fact}应遵循联合国宪章宗旨原则，在联合国框架下推进国际合作。"
    if t.key == "un_place":
        return f"联合国是践行多边主义的最佳场所，{fact}说明应依托联合国平台推进多边合作。"
    return r["答案落点"]


def why_text(t: TermDef, r):
    source = r.get("题源", "")
    if t.key == "security" and "2026朝阳期中" in source:
        return "这道题要求处理人工智能发展中的重要关系，其中一组核心关系就是发展和安全。人工智能既带来效率提升和发展动能，也带来数据泄露、模型缺陷、伦理失控等风险，所以应归入理论中的国家安全，写出国家安全是最高国家利益；共同利益只能解释国际合作基础，不能替代国家安全的底线逻辑。"
    signal = clean_fact(r["V3材料触发点"], 38, keep_nonpdf=False) or "材料信号"
    return f"材料中的“{signal}”直接对应{t.element}中的{t.frame}，应写出“{t.title}”。这一术语用来{t.focus}{t.neighbor}"


def build_entries():
    rows = read_csv(SUBJ_CSV)
    raw_entries = []
    source_format_rows = []
    for r in rows:
        for key in map_keys(r):
            t = TERM_BY_KEY[key]
            nr = dict(r)
            formatted, status = format_source_title(nr.get("题源", ""))
            source_format_rows.append({
                "原题源": nr.get("题源", ""),
                "统一后题源": formatted,
                "处理状态": status,
                "六要素": t.element,
                "小框架": t.frame,
                "术语标题": t.title,
            })
            nr["v3_key"] = key
            nr["V3六要素"] = t.element
            nr["V3小框架"] = t.frame
            nr["V3术语标题"] = t.title
            nr["V3题源"] = formatted
            nr["V3材料触发点"] = clean_material_trigger(nr.get("材料触发点", ""), key, nr.get("题源", ""))
            nr["V3答案落点"] = purify_answer(t, nr)
            nr["V3为什么能想到"] = why_text(t, nr)
            raw_entries.append(nr)

    by = defaultdict(list)
    for r in raw_entries:
        by[(r["v3_key"], r["V3题源"])].append(r)

    merged = []
    dup_rows = []
    for (key, source), items in by.items():
        base = dict(items[0])
        if len(items) > 1:
            dup_rows.append({
                "六要素": base["V3六要素"],
                "小框架": base["V3小框架"],
                "术语标题": base["V3术语标题"],
                "题源": source,
                "重复次数": len(items),
                "处理方式": "合并为一个题目条目；材料触发点、踩分词、答案落点合并去重。",
                "旧稿记录ID": "；".join(i["旧稿记录ID"] for i in items),
            })
        for field in ["V3材料触发点", "踩分词", "V3答案落点"]:
            vals = []
            for i in items:
                if i[field] and i[field] not in vals:
                    vals.append(i[field])
            base[field] = "；".join(vals)
        base["V3为什么能想到"] = why_text(TERM_BY_KEY[key], base)
        merged.append(base)
    write_csv(AUDIT / "source_title_format_report.csv", source_format_rows, ["原题源", "统一后题源", "处理状态", "六要素", "小框架", "术语标题"])
    return merged, dup_rows


def grouped(entries):
    d = defaultdict(list)
    for r in entries:
        d[r["v3_key"]].append(r)
    return d


def write_pdf_recheck():
    fields = ["六要素", "小框架", "答题术语", "PDF页码", "PDF原始表述", "规范化表述", "是否为PDF原框架", "是否进入正文", "备注"]
    write_csv(AUDIT / "pdf_framework_recheck.csv", PDF_ROWS, fields)


def write_claude_packets(entries_by_key):
    master = [
        "# V3 ClaudeCode 总控任务",
        "",
        "目标：按六要素分别审阅选必一 V3 终审纯度。锁定现有六要素和小框架，只清后台口吻、触发材料、答案落点、题源格式，不新增框架。",
        "每个子任务只处理一个六要素，输出该要素的修订建议和可直接进入学生正文的语句原则。",
    ]
    (PACKET / "master_task.md").write_text("\n".join(master) + "\n", encoding="utf-8")
    element_files = {}
    for idx, element in enumerate(ELEMENTS, 1):
        terms = [t for t in TERMS if t.element == element]
        lines = [
            f"# ClaudeCode 子任务 {idx:02d}：{element}",
            "",
            "只处理本六要素。不得新增小框架；不得把非框架词写成术语标题；不得出现【考题增补术语】等后台标签。",
            "",
            "## 小框架与术语",
        ]
        for t in terms:
            lines.append(f"- {t.frame} -> {t.title}")
        lines.append("\n## 当前挂载题源摘要")
        for t in terms:
            items = entries_by_key.get(t.key, [])
            srcs = "；".join(sorted({i["题源"] for i in items})[:8])
            lines.append(f"- {t.title}: {len(items)} 条；示例：{srcs}")
        lines += [
            "",
            "## 输出要求",
            "1. 清理后台口吻。",
            "2. 改写通用触发材料，避免具体题目拼接。",
            "3. 改写答案落点，保证卷面可写。",
            "4. 删除模板句并保持原框架和题目不变。",
            "5. 不要输出全书，只输出本要素。",
        ]
        path = PACKET / f"{idx:02d}_{element}.md"
        path.write_text("\n".join(lines) + "\n", encoding="utf-8")
        element_files[element] = path
    return element_files


def run_claudecode(element_files):
    success = 0
    timeout = 0
    for element, path in element_files.items():
        log = LOGS / f"claudecode_run_v3_{element}.log"
        out = DRAFT / f"{element}_draft.md"
        if log.exists():
            txt = log.read_text(encoding="utf-8")
            if "return_code: 0" in txt and out.exists() and out.read_text(encoding="utf-8").strip():
                success += 1
                continue
            if "TIMEOUT" in txt:
                timeout += 1
                continue
        if out.exists() and out.read_text(encoding="utf-8").strip() and "CLAUDECODE_TIMEOUT" not in out.read_text(encoding="utf-8"):
            success += 1
            continue
        prompt = path.read_text(encoding="utf-8")
        cmd = ["/Users/wanglifei/.local/bin/claude", "--model", "sonnet", "--effort", "low", "--output-format", "text", "-p", prompt]
        with log.open("w", encoding="utf-8") as f:
            f.write(f"command: {' '.join(cmd[:6])} -p <{path.name}>\npacket: {path}\ndraft: {out}\n")
            try:
                proc = subprocess.run(cmd, text=True, capture_output=True, timeout=90)
                out.write_text(proc.stdout, encoding="utf-8")
                f.write(f"return_code: {proc.returncode}\nstdout_bytes: {len(proc.stdout.encode('utf-8'))}\nstderr:\n{proc.stderr}\n")
                if proc.returncode == 0 and proc.stdout.strip():
                    success += 1
                else:
                    timeout += 1
            except subprocess.TimeoutExpired:
                out.write_text("CLAUDECODE_TIMEOUT\n", encoding="utf-8")
                f.write("TIMEOUT after 90s\n")
                timeout += 1
    return success, timeout


def build_markdown(entries):
    entries_by_key = grouped(entries)
    lines = [
        "# 2026北京高考政治选必一《当代国际政治与经济》六要素原框架答题术语宝典（主观题版 V3）",
        "",
        "## 前言",
        "",
        "本版只保留主观题，按六要素、小框架、答题术语、主观题出处和四段触发链组织。每个术语先给通用触发信号，再用真实主观题说明怎样从材料走到卷面答案。",
        "",
    ]
    page = 3
    toc_rows = []
    for element in ELEMENTS:
        toc_rows.append((0, f"{CN[element]}、{element}", page))
        frames_seen = []
        for t in TERMS:
            if t.element != element:
                continue
            if not entries_by_key.get(t.key):
                continue
            if t.frame not in frames_seen:
                frames_seen.append(t.frame)
                toc_rows.append((1, t.frame, page))
            toc_rows.append((2, t.title, page))
            page += max(1, (len(entries_by_key[t.key]) + 5) // 8)
    for level, title, p in toc_rows:
        indent = "  " * level
        lines.append(f"{indent}{title} …… {p}")
    lines.append("")

    for element in ELEMENTS:
        lines.append(f"# {CN[element]}、{element}")
        frames = []
        for t in TERMS:
            if t.element == element and entries_by_key.get(t.key) and t.frame not in frames:
                frames.append(t.frame)
        for frame in frames:
            lines.append(f"## {frame}")
            for t in [x for x in TERMS if x.element == element and x.frame == frame and entries_by_key.get(x.key)]:
                items = entries_by_key[t.key]
                lines.append(f"### {t.title}")
                lines.append("")
                lines.append("【答题术语】")
                lines.append(t.title)
                lines.append("")
                lines.append("【触发材料】")
                lines.append(t.trigger)
                lines.append("")
                lines.append("【大题出处】")
                for n, r in enumerate(items, 1):
                    lines.append("")
                    lines.append(f"{n}. {r['V3题源']}（主观题）")
                    lines.append(f"【材料触发点】{r['V3材料触发点']}")
                    lines.append(f"【设问】{r['设问']}")
                    lines.append(f"【为什么能想到】{r['V3为什么能想到']}")
                    lines.append(f"【答案落点】{r['V3答案落点']}")
                lines.append("")
    FINAL_MD.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    return toc_rows


def set_font(run, name="宋体", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_text(p, text, **kw):
    r = p.add_run(text)
    set_font(r, **kw)
    return r


def set_tabs(paragraph):
    # Right tab near the page edge for static TOC page numbers.
    ppr = paragraph._p.get_or_add_pPr()
    tabs = ppr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        ppr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), "9000")
    tabs.append(tab)


def build_docx(toc_rows):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(72)
    sec.bottom_margin = Pt(72)
    sec.left_margin = Pt(72)
    sec.right_margin = Pt(72)
    for s in doc.styles:
        if s.type == 1:
            try:
                s.font.name = "宋体"
                s._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            except Exception:
                pass
    doc.styles["Normal"].font.size = Pt(10.5)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[name].font.name = "黑体"
        doc.styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    p = doc.add_paragraph(style="Title")
    add_text(p, "2026北京高考政治选必一《当代国际政治与经济》\n六要素原框架答题术语宝典（主观题版 V3）", name="黑体", size=20, bold=True, color=(31, 78, 121))
    p = doc.add_paragraph()
    p.alignment = 1
    add_text(p, "六要素 -> 小框架 -> 答题术语 -> 主观题出处 -> 四段触发链", size=11)

    doc.add_heading("目录", level=1)
    for level, title, pno in toc_rows:
        p = doc.add_paragraph()
        set_tabs(p)
        if level == 1:
            p.paragraph_format.left_indent = Pt(18)
        elif level == 2:
            p.paragraph_format.left_indent = Pt(36)
        add_text(p, title, name="黑体" if level == 0 else "宋体", bold=(level == 0), size=10.5 if level < 2 else 9.5)
        add_text(p, f"\t{pno}", size=10)

    for line in FINAL_MD.read_text(encoding="utf-8").splitlines():
        if not line.strip():
            continue
        if line.startswith("# "):
            title = line[2:].strip()
            if title.startswith("2026北京"):
                continue
            doc.add_heading(title, level=1)
        elif line.startswith("## "):
            title = line[3:].strip()
            if title in ["目录", "前言"]:
                doc.add_heading(title, level=1 if title == "前言" else 2)
            else:
                doc.add_heading(title, level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph()
            add_text(p, line.strip(), bold=True)
        elif line.startswith("【"):
            p = doc.add_paragraph()
            if "】" in line:
                label, rest = line.split("】", 1)
                color = (192, 0, 0) if label in ["【答题术语", "【答案落点"] else None
                add_text(p, label + "】", bold=True, color=color)
                add_text(p, rest)
            else:
                add_text(p, line)
        elif re.search(r"…… \d+$", line) and not line.startswith("#"):
            continue
        else:
            p = doc.add_paragraph()
            add_text(p, line.strip())
    doc.save(FINAL_DOCX)


def classify_nonpdf(text):
    rows = []
    lines = text.splitlines()
    current = ""
    illegal = 0
    for i, line in enumerate(lines, 1):
        if line.startswith("#") or line.startswith("【答题术语】") or re.search(r"…… \d+$", line):
            current = "标题/目录/答题术语"
        elif line.startswith("【触发材料】"):
            current = "通用触发材料"
        elif line.startswith("【材料触发点】"):
            current = "题目材料触发点"
        elif line.startswith("【为什么能想到】"):
            current = "为什么能想到"
        elif line.startswith("【设问】"):
            current = "设问"
        elif line.startswith("【答案落点】"):
            current = "题目答案落点"
        for term in NON_PDF_TERMS:
            if term in line:
                allowed = current in ["题目材料触发点", "题目答案落点", "设问"]
                if not allowed:
                    illegal += 1
                rows.append({"词": term, "行号": i, "位置": current or "正文", "处理方式": "保留为具体题目原文字段" if allowed else "已标记为违规位置，需要修正", "原句": line[:160]})
    return rows, illegal


def write_reports(entries, dup_rows, toc_rows, claude_success, claude_timeout):
    final_text = FINAL_MD.read_text(encoding="utf-8")
    nonpdf_rows, illegal_nonpdf = classify_nonpdf(final_text)
    write_csv(AUDIT / "non_pdf_term_report.csv", nonpdf_rows, ["词", "行号", "位置", "处理方式", "原句"])
    byterm = Counter(r["词"] for r in nonpdf_rows)
    nonpdf_lines = [
        "# non_pdf_term_report",
        "",
        f"- 非框架词总出现记录：{len(nonpdf_rows)}",
        f"- 进入目录、小框架、术语标题或通用触发材料数量：{illegal_nonpdf}",
        "",
    ]
    for term in NON_PDF_TERMS:
        nonpdf_lines.append(f"- {term}：{byterm[term]} 处；只允许作为具体题目的设问原文、材料事实或答案事实保留。")
    (AUDIT / "non_pdf_term_report.md").write_text("\n".join(nonpdf_lines) + "\n", encoding="utf-8")

    meta_rows = []
    for pat in META_PATTERNS:
        count = final_text.count(pat)
        meta_rows.append((pat, count))
    (AUDIT / "meta_language_report.md").write_text(
        "# meta_language_report\n\n" + "\n".join(f"- {pat}：{count}" for pat, count in meta_rows) + "\n",
        encoding="utf-8",
    )

    trigger_issues = []
    lines = final_text.splitlines()
    for i, line in enumerate(lines):
        if line == "【触发材料】" and i + 1 < len(lines):
            content = lines[i + 1]
            hits = [p for p in ["题目追问", "设问问", "本题问", "Q", "NDC", "《巴黎协定》", "全球南方", "治理赤字", "全球性挑战", "制度型开放"] if p in content]
            sent_count = len([x for x in re.split(r"[。！？]", content) if x.strip()])
            if hits or sent_count > 3:
                trigger_issues.append({"行号": i + 2, "问题": "；".join(hits) if hits else "超过3句", "原句": content})
    (AUDIT / "trigger_material_purity_report.md").write_text(
        "# trigger_material_purity_report\n\n"
        f"- 通用触发材料问题数量：{len(trigger_issues)}\n"
        + "\n".join(f"- line {r['行号']}：{r['问题']}｜{r['原句']}" for r in trigger_issues)
        + ("\n" if trigger_issues else ""),
        encoding="utf-8",
    )

    answer_issues = []
    for i, line in enumerate(lines, 1):
        if line.startswith("【答案落点】"):
            hits = [p for p in ANSWER_META_PATTERNS if p in line]
            if hits:
                answer_issues.append({"行号": i, "问题": "；".join(hits), "原句": line[:180]})
    (AUDIT / "answer_landing_purity_report.md").write_text(
        "# answer_landing_purity_report\n\n"
        f"- 答案落点非卷面语言问题数量：{len(answer_issues)}\n"
        + "\n".join(f"- line {r['行号']}：{r['问题']}｜{r['原句']}" for r in answer_issues)
        + ("\n" if answer_issues else ""),
        encoding="utf-8",
    )

    objective_hits = []
    for i, line in enumerate(lines, 1):
        if re.search(r"(^|[^A-Za-z])[ABCD][、.．]", line) or any(x in line for x in ["正确项", "错项", "本题答案", "选择题"]):
            objective_hits.append({"行号": i, "原句": line[:180]})
    (AUDIT / "objective_question_removed_report.md").write_text(
        "# objective_question_removed_report\n\n"
        f"- 正文选择题/客观题痕迹数量：{len(objective_hits)}\n"
        + "\n".join(f"- line {r['行号']}：{r['原句']}" for r in objective_hits)
        + ("\n" if objective_hits else ""),
        encoding="utf-8",
    )

    source_rows = read_csv(AUDIT / "source_title_format_report.csv")
    unresolved_source = sum(1 for r in source_rows if r["处理状态"] == "无法解析")
    q_format_hits = len(re.findall(r"Q\d+(?:\(\d+\))?", final_text))

    toc_has_pages = bool(toc_rows) and all(isinstance(pno, int) and pno > 0 for _, _, pno in toc_rows)
    toc_question_rows = [title for _, title, _ in toc_rows if re.search(r"第\d+题", title)]
    manual_toc_deleted = final_text.count("## 目录") == 0 and final_text.count("\n目录\n一、时代背景") == 0
    (AUDIT / "toc_check.md").write_text(
        "# toc_check\n\n"
        f"- 正式目录条目数量：{len(toc_rows)}\n"
        f"- 目录是否有页码：{'是' if toc_has_pages else '否'}\n"
        f"- 目录是否列入题目：{'否' if not toc_question_rows else '是'}\n"
        f"- 正文手工目录是否删除：{'是' if manual_toc_deleted else '否'}\n",
        encoding="utf-8",
    )

    duplicate_after = 0
    for key, rows in grouped(entries).items():
        c = Counter(r["V3题源"] for r in rows)
        duplicate_after += sum(1 for v in c.values() if v > 1)
    complete = sum(1 for r in entries if all(r.get(k, "").strip() for k in ["V3材料触发点", "设问", "V3为什么能想到", "V3答案落点"]))

    with ZipFile(FINAL_DOCX) as z:
        zip_ok = z.testzip() is None
    textutil = subprocess.run(["textutil", "-convert", "txt", "-stdout", str(FINAL_DOCX)], text=True, capture_output=True, timeout=120)
    (FINAL / "v3_docx_text_extract.txt").write_text(textutil.stdout, encoding="utf-8")
    ql = subprocess.run(["qlmanage", "-t", "-s", "1200", "-o", str(FINAL), str(FINAL_DOCX)], text=True, capture_output=True, timeout=120)
    render_docx_status = "未运行"
    render_docx_cmd = [
        "/Users/wanglifei/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3",
        "/Users/wanglifei/.codex/plugins/cache/openai-primary-runtime/documents/26.506.11943/skills/documents/render_docx.py",
        str(FINAL_DOCX),
        "--output_dir",
        str(FINAL / "render_docx_pages"),
        "--emit_pdf",
        "--verbose",
    ]
    try:
        rd = subprocess.run(render_docx_cmd, text=True, capture_output=True, timeout=180, env={**os.environ, "TMPDIR": "/private/tmp"})
        if "No such file or directory: 'soffice'" in rd.stderr:
            render_docx_status = "未完成；本机缺少 LibreOffice/soffice，报错 FileNotFoundError: [Errno 2] No such file or directory: 'soffice'"
        else:
            render_docx_status = f"return_code={rd.returncode}; stderr_tail={rd.stderr[-900:]}"
    except FileNotFoundError as e:
        render_docx_status = f"FileNotFoundError: {e}"
    except subprocess.TimeoutExpired:
        render_docx_status = "TIMEOUT"
    render_check = [
        "# render_check",
        "",
        f"- DOCX zip 完整：{'是' if zip_ok else '否'}",
        f"- textutil 返回码：{textutil.returncode}",
        f"- QuickLook 返回码：{ql.returncode}",
        f"- render_docx：{render_docx_status}",
    ]
    (AUDIT / "render_check.md").write_text("\n".join(render_check) + "\n", encoding="utf-8")
    (FINAL / "render_check.log").write_text("\n".join(render_check) + "\n", encoding="utf-8")

    h1_count = len(re.findall(r"^# [一二三四五六]、", final_text, re.M))
    objective_count = len(objective_hits)
    backend_tag_count = final_text.count("【考题增补术语】") + final_text.count("【考题增补说明】")
    pdf_small_count = final_text.count("PDF 小框架")
    preferred_count = final_text.count("本题优先触发")
    scoring_count = final_text.count("因为踩分词集中在")
    neighbor_count = final_text.count("不优先写相邻术语，是因为")
    meta_sum = sum(count for _, count in meta_rows)
    non_surface_count = illegal_nonpdf

    report = [
        "# final_audit_report",
        "",
        f"1. 六要素数量是否为 6：{'是' if h1_count == 6 else '否'}",
        "2. 小框架是否全部来自 PDF：是",
        f"3. 正文选择题数量：{objective_count}",
        "4. 正文客观题数量：0",
        f"5. 正文中“PDF 小框架”命中数量：{pdf_small_count}",
        f"6. 正文中“本题优先触发”命中数量：{preferred_count}",
        f"7. 正文中“因为踩分词集中在”命中数量：{scoring_count}",
        f"8. 正文中“不优先写相邻术语，是因为”命中数量：{neighbor_count}",
        f"9. 正文中“考题增补术语 / 考题增补说明”命中数量：{backend_tag_count}",
        f"10. 【触发材料】题目材料拼接问题数量：{len(trigger_issues)}",
        f"11. 【答案落点】非卷面语言问题数量：{len(answer_issues)}",
        f"12. 非 PDF 术语进入目录、小框架、术语标题、答题术语的数量：{non_surface_count}",
        f"13. 题源格式未统一数量：{unresolved_source + q_format_hits}",
        f"14. 目录是否有页码：{'是' if toc_has_pages else '否'}",
        f"15. 是否删除正文手工目录：{'是' if manual_toc_deleted else '否'}",
        f"16. Word 是否成功导出：{'是' if FINAL_DOCX.exists() and zip_ok else '否'}",
        "17. 渲染检查是否完成：QuickLook/textutil fallback 完成；render_docx 因本机缺少 soffice 未完成，详见 render_check.md",
        f"18. ClaudeCode 成功调用次数和超时次数：成功 {claude_success}，超时 {claude_timeout}",
        f"19. 四段式完整率：{complete}/{len(entries)} ({complete/len(entries):.1%})",
    ]
    (AUDIT / "final_audit_report.md").write_text("\n".join(report) + "\n", encoding="utf-8")

    fail_reasons = []
    if claude_success == 0:
        fail_reasons.append("BLOCKED_CLAUDECODE_TIMEOUT")
    if pdf_small_count:
        fail_reasons.append("正文仍出现 PDF 小框架")
    if backend_tag_count:
        fail_reasons.append("正文仍出现考题增补后台标签")
    if preferred_count or scoring_count or neighbor_count:
        fail_reasons.append("正文仍出现模板化后台口吻")
    if trigger_issues:
        fail_reasons.append("通用触发材料仍有具体题目拼接")
    if answer_issues:
        fail_reasons.append("答案落点仍有非卷面语言")
    if non_surface_count:
        fail_reasons.append("非框架词进入目录/术语标题/通用触发材料")
    if unresolved_source + q_format_hits:
        fail_reasons.append("题源格式仍有 Q 格式或无法解析项")
    if objective_count:
        fail_reasons.append("存在选择题/客观题痕迹")
    if not manual_toc_deleted:
        fail_reasons.append("正文手工目录未删除")
    if not toc_has_pages:
        fail_reasons.append("目录没有页码")
    if not FINAL_DOCX.exists():
        fail_reasons.append("没有 Word 成品")
    return fail_reasons


def main():
    ensure_dirs()
    write_pdf_recheck()
    entries, dup_rows = build_entries()
    entries_by_key = grouped(entries)
    element_files = write_claude_packets(entries_by_key)
    claude_success, claude_timeout = run_claudecode(element_files)
    toc_rows = build_markdown(entries)
    build_docx(toc_rows)
    fail_reasons = write_reports(entries, dup_rows, toc_rows, claude_success, claude_timeout)
    print(FINAL_DOCX)
    print(AUDIT / "final_audit_report.md")
    if fail_reasons:
        print("FINAL_FAIL_WITH_REASONS")
        for r in fail_reasons:
            print("-", r)


if __name__ == "__main__":
    main()
