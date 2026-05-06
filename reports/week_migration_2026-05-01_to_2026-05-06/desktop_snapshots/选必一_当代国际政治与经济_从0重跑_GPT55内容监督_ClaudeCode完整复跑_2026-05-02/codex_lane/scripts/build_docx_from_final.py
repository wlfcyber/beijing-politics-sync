#!/usr/bin/env python3
"""Create a polished DOCX from the final student Markdown."""

from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


RUN = Path(__file__).resolve().parents[2]
MD = RUN / "09_delivery" / "飞哥政治庄园_选必一_当代国际政治与经济_细则术语成品版_2026-05-03.md"
OUT = RUN / "09_delivery" / "飞哥政治庄园_选必一_当代国际政治与经济_细则术语成品版_2026-05-03.docx"


def set_east_asia_font(style, font_name: str) -> None:
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_bottom_border(paragraph, color="D8DEE9", size="6") -> None:
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def add_label_paragraph(doc: Document, label: str, body: str) -> None:
    p = doc.add_paragraph(style="Body")
    p.paragraph_format.keep_together = True
    r = p.add_run(label + "：")
    r.bold = True
    r.font.color.rgb = RGBColor(55, 65, 81)
    p.add_run(body)


def build() -> None:
    text = MD.read_text(encoding="utf-8")
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    normal = styles["Normal"]
    set_east_asia_font(normal, "PingFang SC")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(31, 41, 55)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(4)

    for name, size, color in [
        ("Title", 20, "111827"),
        ("Heading 1", 15, "0F766E"),
        ("Heading 2", 12.5, "1F2937"),
        ("Heading 3", 11, "374151"),
    ]:
        style = styles[name]
        set_east_asia_font(style, "PingFang SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(4)

    body = styles.add_style("Body", 1)
    set_east_asia_font(body, "PingFang SC")
    body.font.size = Pt(10.2)
    body.font.color.rgb = RGBColor(31, 41, 55)
    body.paragraph_format.line_spacing = 1.12
    body.paragraph_format.space_after = Pt(4)

    note = styles.add_style("DocNote", 1)
    set_east_asia_font(note, "PingFang SC")
    note.font.size = Pt(9.5)
    note.font.color.rgb = RGBColor(75, 85, 99)
    note.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.text = "飞哥政治庄园 · 选必一《当代国际政治与经济》"
    header.style = styles["DocNote"]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bottom_border(header)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.style = styles["DocNote"]
    footer.add_run("第 ")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    footer._p.append(fld_begin)
    footer._p.append(instr)
    footer._p.append(fld_end)
    footer.add_run(" 页")

    for idx, line in enumerate(text.splitlines()):
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(line[2:], style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_bottom_border(p, color="0F766E", size="10")
            continue
        if line.startswith("## "):
            p = doc.add_paragraph(line[3:], style="Heading 1")
            add_bottom_border(p, color="99F6E4", size="8")
            continue
        if line.startswith("### 术语："):
            p = doc.add_paragraph(line[4:], style="Heading 2")
            p.paragraph_format.keep_with_next = True
            continue
        if line.startswith("#### "):
            p = doc.add_paragraph(line[5:], style="Heading 3")
            p.paragraph_format.keep_with_next = True
            continue
        label_match = re.match(r"^(框架位置|完整设问|细则位置|来源|材料触发|答案句)：(.+)$", line)
        if label_match:
            add_label_paragraph(doc, label_match.group(1), label_match.group(2))
        else:
            doc.add_paragraph(line, style="DocNote")

    doc.core_properties.title = "飞哥政治庄园 选必一细则术语成品版"
    doc.core_properties.subject = "选择性必修一《当代国际政治与经济》主观题评分术语"
    doc.core_properties.author = "Codex"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
