#!/usr/bin/env python3
"""Build a student-facing Phase11E candidate Markdown/DOCX from the 29-entry source lock.

This is a candidate builder, not a final-pass signer. It preserves the four
student-facing fields while adding Word structure, page numbers, TOC fields,
and source-derived images for visual questions.
"""

from __future__ import annotations

import re
import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap, qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "09_student_draft/phase11D_combined_source_verified_29_REVIEW_ONLY.md"
OUT_MD = ROOT / "outputs/2026北京高考政治选必三逻辑与思维宝典---思维与推理全触发全链条_学生版_CANDIDATE_PENDING_GPT.md"
OUT_DOCX = ROOT / "outputs/2026北京高考政治选必三逻辑与思维宝典---思维与推理全触发全链条_学生版_CANDIDATE_PENDING_GPT.docx"
QA_REPORT = ROOT / "08_review/phase11E_candidate_docx_build_report.md"
ASSET_DIR = ROOT / "outputs/assets_phase11E_candidate"

nsmap.setdefault("v", "urn:schemas-microsoft-com:vml")

IMAGE_MAP = {
    "2025 海淀期末第17题第(1)问": ROOT / "07_student_doc/extracted_images/image2.png",
    "2025 海淀期末第18题": ROOT / "07_student_doc/extracted_images/image3.png",
    "2025 丰台期末第7题": ROOT / "07_student_doc/2025丰台期末_Q7_漫画_crop.png",
}

SECTION_ORDER = [
    "一、思维方法：从材料动作锁定方法",
    "二、推理题型：先定形式再套规则",
    "三、交叉综合：同一道题拆两层",
    "四、边界陷阱：看见术语先判模块",
]

SECTION_BY_TITLE = {
    "2025 东城期末第18题第(2)问": SECTION_ORDER[0],
    "2024 海淀二模第17题第(1)问": SECTION_ORDER[0],
    "2024 海淀二模第17题第(2)问": SECTION_ORDER[0],
    "2026 顺义一模第19题第(2)问": SECTION_ORDER[2],
    "2024 朝阳一模第7题": SECTION_ORDER[0],
    "2025 海淀期末第18题": SECTION_ORDER[0],
    "2025 丰台期末第8题": SECTION_ORDER[0],
    "2026 通州期末第11题": SECTION_ORDER[0],
    "2026 朝阳期中第20题": SECTION_ORDER[0],
    "2026 东城一模第19题第(4)问": SECTION_ORDER[0],
    "2025 海淀期末第17题第(1)问": SECTION_ORDER[0],
    "2024 朝阳一模第20题第(1)问": SECTION_ORDER[1],
    "2025 顺义一模第7题": SECTION_ORDER[1],
    "2026 朝阳期中第21题第(2)问": SECTION_ORDER[0],
    "2025 西城二模第16题第(3)问": SECTION_ORDER[0],
    "2024 朝阳一模第9题": SECTION_ORDER[0],
    "2024 朝阳一模第20题第(2)问": SECTION_ORDER[1],
    "2025 东城期末第13题": SECTION_ORDER[1],
    "2024 朝阳期中第7题": SECTION_ORDER[1],
    "2024 西城一模第19题第(2)问": SECTION_ORDER[1],
    "2024 西城一模第19题第(3)问": SECTION_ORDER[1],
    "2024 西城一模第19题第(5)问": SECTION_ORDER[1],
    "2025 西城二模第16题第(2)问": SECTION_ORDER[1],
    "2026 通州期末第19题第(2)问": SECTION_ORDER[1],
    "2024 朝阳二模第19题第(1)问": SECTION_ORDER[2],
    "2024 朝阳二模第19题第(2)问": SECTION_ORDER[2],
    "2024 朝阳期中第9题": SECTION_ORDER[2],
    "2026 顺义一模第19题第(1)问": SECTION_ORDER[2],
    "2025 丰台期末第7题": SECTION_ORDER[3],
}


@dataclass
class Entry:
    order: int
    title: str
    fields: dict[str, str]


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace('"', "“").replace("'", "”")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def normalize_question(text: str) -> str:
    text = normalize_text(text)
    # Split statement units like ①...②..., but keep answer-combo units such as A.①② together.
    text = re.sub(r"(?<!\n)(?<![ABCD][.．])(?<![-—–])([①②③④])(?=[^①②③④—–-])", r"\n\1", text)
    text = re.sub(r"(?<!\n)([ABCD][.．])", r"\n\1", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def parse_entries() -> list[Entry]:
    raw = INPUT.read_text(encoding="utf-8")
    chunks = re.split(r"\n(?=## )", raw)
    entries: list[Entry] = []
    for chunk in chunks:
        if not chunk.startswith("## "):
            continue
        title_line, _, body = chunk.partition("\n")
        title = title_line[3:].strip()
        fields: dict[str, str] = {}
        for match in re.finditer(r"【([^】]+)】[ \t]*(?:\n\n|\n)?(.*?)(?=\n\n?【[^】]+】|\n---|\Z)", body, re.S):
            label = match.group(1).strip()
            value = match.group(2).strip()
            if label == "设问":
                value = normalize_question(value)
            else:
                value = normalize_text(value)
            fields[label] = value
        if {"材料触发点", "设问", "为什么能想到", "答案落点"} <= set(fields):
            entries.append(Entry(len(entries) + 1, title, fields))
    return entries


def grouped(entries: list[Entry]) -> dict[str, list[Entry]]:
    groups = {section: [] for section in SECTION_ORDER}
    for entry in entries:
        groups[SECTION_BY_TITLE.get(entry.title, SECTION_ORDER[2])].append(entry)
    return groups


def build_markdown(entries: list[Entry]) -> str:
    groups = grouped(entries)
    lines: list[str] = [
        "# 2026北京高考政治选必三《逻辑与思维》宝典",
        "",
        "## 使用口令",
        "",
        "- 思维题：先看材料动作，再锁定思维方法，最后写成“方法 + 材料 + 作用”的卷面句。",
        "- 推理题：先定形式，再套规则，再把规则落回本题材料。",
        "- 选择题：四个选项都要看，先找正确项的材料依据，再找错项的偷换、绝对化或模块错位。",
        "- 边界题：看见“未来”“创新”“系统”等词，不要急着套选必三，先判断设问真正考哪本书。",
        "",
        "---",
    ]
    global_no = 1
    for section in SECTION_ORDER:
        lines.extend(["", f"## {section}", ""])
        for entry in groups[section]:
            lines.extend(["", f"### {global_no}. {entry.title}", ""])
            for label in ("材料触发点", "设问", "为什么能想到", "答案落点"):
                lines.extend([f"**{label}**", ""])
                text = entry.fields[label]
                lines.extend(text.splitlines())
                lines.append("")
                if label == "材料触发点" and entry.title in IMAGE_MAP and IMAGE_MAP[entry.title].exists():
                    image_path = ASSET_DIR / IMAGE_MAP[entry.title].name
                    lines.append(f"![{entry.title}](assets_phase11E_candidate/{image_path.name})")
                    lines.append("")
            global_no += 1
    return "\n".join(lines).rstrip() + "\n"


def set_font(run, east_asia: str = "宋体", latin: str = "Aptos", size: float | None = None):
    if size is not None:
        run.font.size = Pt(size)
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def style_document(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, color in [
        ("Title", 22, "1F4E79"),
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 13.5, "2F75B5"),
        ("Heading 3", 12, "1F4E79"),
    ]:
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10 if style_name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)


def add_field(paragraph, instr: str, display: str = "") -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr_text, fld_sep, text, fld_end])


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, size=9)
    add_field(paragraph, "PAGE", "1")
    run = paragraph.add_run(" 页")
    set_font(run, size=9)


def add_watermark(section) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pict_run = paragraph.add_run()
    pict = OxmlElement("w:pict")
    shape = OxmlElement("v:shape")
    shape.set("id", "FeigeWatermark")
    shape.set("type", "#_x0000_t136")
    shape.set(
        "style",
        "position:absolute;left:0;text-align:left;margin-left:0;"
        "margin-top:240pt;width:500pt;height:90pt;rotation:315;"
        "z-index:-251654144;mso-position-horizontal:center;"
        "mso-position-horizontal-relative:margin;"
        "mso-position-vertical:center;"
        "mso-position-vertical-relative:margin",
    )
    shape.set("fillcolor", "#C9DDF3")
    shape.set("stroked", "f")
    fill = OxmlElement("v:fill")
    fill.set("opacity", "0.20")
    textpath = OxmlElement("v:textpath")
    textpath.set("style", "font-family:'Microsoft YaHei';font-size:30pt;font-weight:bold")
    textpath.set("string", "飞哥正志讲堂")
    shape.append(fill)
    shape.append(textpath)
    pict.append(shape)
    pict_run._r.append(pict)


def enable_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_body_paragraph(doc: Document, text: str, option: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = None if option else Cm(0.74)
    paragraph.paragraph_format.left_indent = Cm(0.74 if option else 0)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_font(run, size=10.5)


def build_docx(entries: list[Entry]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.05)
    section.right_margin = Cm(2.05)
    style_document(doc)
    add_page_number(section)
    add_watermark(section)
    enable_update_fields(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("2026北京高考政治选必三《逻辑与思维》宝典")
    set_font(run, east_asia="微软雅黑", latin="Aptos Display", size=22)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("思维方法 + 推理题型 全触发全链条 学生版")
    set_font(run, east_asia="微软雅黑", size=12)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()
    doc.add_heading("目录", level=1)
    toc = doc.add_paragraph()
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u', "右键更新目录")
    doc.add_page_break()

    doc.add_heading("使用口令", level=1)
    for bullet in [
        "思维题：先看材料动作，再锁定思维方法，最后写成“方法 + 材料 + 作用”的卷面句。",
        "推理题：先定形式，再套规则，再把规则落回本题材料。",
        "选择题：四个选项都要看，先找正确项的材料依据，再找错项的偷换、绝对化或模块错位。",
        "边界题：看见“未来”“创新”“系统”等词，不要急着套选必三，先判断设问真正考哪本书。",
    ]:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["List Bullet"]
        run = p.add_run(bullet)
        set_font(run, size=10.5)

    groups = grouped(entries)
    global_no = 1
    for section_title in SECTION_ORDER:
        doc.add_page_break()
        doc.add_heading(section_title, level=1)
        for entry in groups[section_title]:
            doc.add_heading(f"{global_no}. {entry.title}", level=2)
            for label in ("材料触发点", "设问", "为什么能想到", "答案落点"):
                label_p = doc.add_paragraph()
                label_p.paragraph_format.keep_with_next = True
                run = label_p.add_run(f"【{label}】")
                set_font(run, east_asia="微软雅黑", size=10.5)
                run.bold = True
                run.font.color.rgb = RGBColor(31, 78, 121)
                text = entry.fields[label]
                for line in text.splitlines():
                    stripped = line.strip()
                    if not stripped:
                        continue
                    add_body_paragraph(
                        doc,
                        stripped,
                        option=bool(re.match(r"^([①②③④]|[ABCD][.．])", stripped)),
                    )
                if label == "材料触发点" and entry.title in IMAGE_MAP and IMAGE_MAP[entry.title].exists():
                    img_p = doc.add_paragraph()
                    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = img_p.add_run()
                    width = Inches(2.9 if entry.title != "2025 海淀期末第18题" else 2.25)
                    run.add_picture(str(IMAGE_MAP[entry.title]), width=width)
            global_no += 1

    doc.save(OUT_DOCX)


def main() -> int:
    OUT_MD.parent.mkdir(parents=True, exist_ok=True)
    QA_REPORT.parent.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    for path in IMAGE_MAP.values():
        if path.exists():
            shutil.copy2(path, ASSET_DIR / path.name)
    entries = parse_entries()
    if len(entries) != 29:
        raise RuntimeError(f"expected 29 entries, got {len(entries)}")
    md = build_markdown(entries)
    OUT_MD.write_text(md, encoding="utf-8")
    build_docx(entries)
    missing_images = [title for title, path in IMAGE_MAP.items() if not path.exists()]
    report = [
        "# Phase11E Candidate DOCX Build Report",
        "",
        f"- entries: {len(entries)}",
        f"- markdown: `{OUT_MD}`",
        f"- docx: `{OUT_DOCX}`",
        f"- image_slots: {len(IMAGE_MAP)}",
        f"- missing_images: {len(missing_images)}",
        f"- missing_image_titles: {', '.join(missing_images) if missing_images else 'none'}",
        "- status: `CANDIDATE_BUILT_PENDING_OPUS_CAPTURE_GPT_REVIEW_GOVERNOR_CONFUCIUS`",
    ]
    QA_REPORT.write_text("\n".join(report) + "\n", encoding="utf-8")
    print(f"entries={len(entries)}")
    print(f"wrote={OUT_MD}")
    print(f"wrote={OUT_DOCX}")
    print(f"wrote={QA_REPORT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
