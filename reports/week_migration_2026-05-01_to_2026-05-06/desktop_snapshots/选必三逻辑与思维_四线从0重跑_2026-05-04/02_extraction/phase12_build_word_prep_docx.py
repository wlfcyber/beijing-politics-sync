#!/usr/bin/env python3
"""Build the Phase12 student-clean Word-prep DOCX.

This builder consumes the already-gated student-clean Markdown and clean dual
indexes. It does not change teaching content. It only packages the content into
Word structure: title page, TOC field, headings, page numbers, watermark, and
source images for visual questions.
"""

from __future__ import annotations

import re
import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap, qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BODY_MD = ROOT / "09_student_draft/phase12_student_clean_candidate.md"
REASONING_INDEX = ROOT / "09_student_draft/phase12_reasoning_typology_index_STUDENT_CLEAN_CANDIDATE.md"
THINKING_INDEX = ROOT / "09_student_draft/phase12_thinking_method_index_STUDENT_CLEAN_CANDIDATE.md"
OUT_MD = ROOT / "outputs/2026北京高考政治选必三逻辑与思维宝典---思维与推理全触发全链条_学生版_WORD_PREP.md"
OUT_DOCX = ROOT / "outputs/2026北京高考政治选必三逻辑与思维宝典---思维与推理全触发全链条_学生版_WORD_PREP.docx"
QA_REPORT = ROOT / "08_review/phase12_word_prep_docx_build_report.md"
ASSET_DIR = ROOT / "outputs/assets_phase12_word_prep"

nsmap.setdefault("v", "urn:schemas-microsoft-com:vml")

IMAGE_MAP = {
    "2025 海淀期末第17题第(1)问": ROOT / "07_student_doc/extracted_images/image2.png",
    "2025 海淀期末第18题": ROOT / "07_student_doc/extracted_images/image3.png",
    "2025 丰台期末第7题": ROOT / "07_student_doc/2025丰台期末_Q7_漫画_crop.png",
}

INTERNAL_MARKERS = [
    "REVIEW_ONLY",
    "Status:",
    "<!-- question_id",
    "phase12_decision",
    "source_pool",
    "manual_lock",
    "NEEDS_TYPE",
    "NEEDS_METHOD",
    "交叉题次挂载",
]


@dataclass
class RenderState:
    current_title: str | None = None
    image_inserted_for_title: set[str] | None = None

    def __post_init__(self) -> None:
        if self.image_inserted_for_title is None:
            self.image_inserted_for_title = set()


def set_run_font(run, east_asia: str = "宋体", latin: str = "Aptos", size: float | None = None) -> None:
    if size is not None:
        run.font.size = Pt(size)
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def style_document(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)

    style_specs = {
        "Title": (23, "1F4E79", "微软雅黑"),
        "Subtitle": (12, "595959", "微软雅黑"),
        "Heading 1": (16, "1F4E79", "微软雅黑"),
        "Heading 2": (13.2, "2F75B5", "微软雅黑"),
        "Heading 3": (11.5, "1F4E79", "微软雅黑"),
    }
    for name, (size, color, font) in style_specs.items():
        style = styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10 if name not in {"Title", "Subtitle"} else 0)
        style.paragraph_format.space_after = Pt(6 if name != "Subtitle" else 12)
        style.paragraph_format.keep_with_next = name.startswith("Heading")

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(3)


def add_field(paragraph, instr: str, display: str = "") -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr_text, fld_sep, text, fld_end])


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9)
    add_field(paragraph, "PAGE", "1")
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9)


def add_running_header(section) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("飞哥正志讲堂｜选必三《逻辑与思维》")
    set_run_font(run, east_asia="微软雅黑", size=8.5)
    run.font.color.rgb = RGBColor(127, 127, 127)


def add_watermark(section) -> None:
    header = section.header
    paragraph = header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pict_run = paragraph.add_run()
    pict = OxmlElement("w:pict")
    shape = OxmlElement("v:shape")
    shape.set("id", "FeigeWatermark")
    shape.set("type", "#_x0000_t136")
    shape.set(
        "style",
        "position:absolute;left:0;text-align:left;margin-left:0;"
        "margin-top:240pt;width:500pt;height:90pt;rotation:315;"
        "z-index:-251654144;mso-position-horizontal:center;"
        "mso-position-horizontal-relative:margin;"
        "mso-position-vertical:center;"
        "mso-position-vertical-relative:margin",
    )
    shape.set("fillcolor", "#8EAADB")
    shape.set("stroked", "f")
    fill = OxmlElement("v:fill")
    fill.set("opacity", "0.28")
    textpath = OxmlElement("v:textpath")
    textpath.set("style", "font-family:'Microsoft YaHei';font-size:30pt;font-weight:bold")
    textpath.set("string", "飞哥正志讲堂")
    shape.append(fill)
    shape.append(textpath)
    pict.append(shape)
    pict_run._r.append(pict)


def enable_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def normalize_quotes(text: str) -> str:
    # Preserve apostrophes inside English identifiers; normalize Chinese quote style.
    text = text.replace("“", "“").replace("”", "”")
    text = re.sub(r"(?<![A-Za-z0-9])'([^']+)'(?![A-Za-z0-9])", r"“\1”", text)
    return text


def combined_markdown() -> str:
    body = BODY_MD.read_text(encoding="utf-8").strip()
    reasoning = REASONING_INDEX.read_text(encoding="utf-8").strip()
    thinking = THINKING_INDEX.read_text(encoding="utf-8").strip()

    body_without_title = re.sub(r"^# .+?\n+", "", body, count=1)
    intro, sep, rest = body_without_title.partition("\n## 主观题")
    lines = [
        "# 2026北京高考政治选必三《逻辑与思维》宝典：思维与推理全触发全链条",
        "",
        intro.strip(),
        "",
        "## 思维方法索引",
        "",
        re.sub(r"^# .+?\n+", "", thinking, count=1).strip(),
        "",
        "## 推理题型索引",
        "",
        re.sub(r"^# .+?\n+", "", reasoning, count=1).strip(),
        "",
        "## 主观题",
        rest.strip(),
    ]
    return normalize_quotes("\n".join(lines).rstrip() + "\n")


def add_label_paragraph(doc: Document, line: str) -> None:
    match = re.match(r"^(【[^】]+】)(.*)$", line)
    if not match:
        add_body_paragraph(doc, line)
        return
    label, rest = match.groups()
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True if not rest.strip() else False
    run = paragraph.add_run(label)
    set_run_font(run, east_asia="微软雅黑", size=10.5)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)
    if rest.strip():
        run = paragraph.add_run(rest.strip())
        set_run_font(run, size=10.5)


def add_body_paragraph(doc: Document, line: str) -> None:
    option_like = bool(re.match(r"^([①②③④]|[ABCD][.．]|[0-9]+[.、])", line))
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.7 if option_like else 0)
    paragraph.paragraph_format.first_line_indent = None if option_like else Cm(0.72)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.13
    run = paragraph.add_run(line)
    set_run_font(run, size=10.5)


def add_bullet(doc: Document, line: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Cm(0.75)
    paragraph.paragraph_format.first_line_indent = Cm(-0.25)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(line)
    set_run_font(run, size=9.8 if len(line) > 80 else 10.2)


def add_image_if_needed(doc: Document, state: RenderState, just_rendered_line: str) -> None:
    title = state.current_title
    if not title or title not in IMAGE_MAP or title in state.image_inserted_for_title:
        return
    if not re.match(r"^(【材料触发点】|【题干信号】|- 材料信号：)", just_rendered_line):
        return
    image = IMAGE_MAP[title]
    if not image.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run()
    width = Inches(2.25 if title == "2025 海淀期末第18题" else 3.05)
    if title == "2025 丰台期末第7题":
        width = Inches(2.8)
    run.add_picture(str(image), width=width)
    state.image_inserted_for_title.add(title)


def render_markdown_to_docx(doc: Document, md: str) -> None:
    state = RenderState()
    for raw_line in md.splitlines():
        line = raw_line.strip()
        if not line or line == "---":
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            heading = line[3:].strip()
            doc.add_heading(heading, level=1)
            continue
        if line.startswith("### "):
            title = line[4:].strip()
            state.current_title = title
            doc.add_heading(title, level=2)
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            continue
        if line.startswith("- "):
            add_bullet(doc, line[2:].strip())
            add_image_if_needed(doc, state, line)
            continue
        if line.startswith("【"):
            add_label_paragraph(doc, line)
            add_image_if_needed(doc, state, line)
            continue
        add_body_paragraph(doc, line)
        add_image_if_needed(doc, state, line)


def build_docx(md: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.75)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    style_document(doc)
    add_running_header(section)
    add_page_number(section)
    add_watermark(section)
    enable_update_fields(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("2026北京高考政治选必三《逻辑与思维》宝典")
    set_run_font(run, east_asia="微软雅黑", latin="Aptos Display", size=23)

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("思维与推理全触发全链条｜学生版")
    set_run_font(run, east_asia="微软雅黑", size=12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("主观题在前，选择题在后；按区县与时间顺序整理；含思维方法索引与推理题型索引")
    set_run_font(run, east_asia="微软雅黑", size=10)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()
    doc.add_heading("目录", level=1)
    toc = doc.add_paragraph()
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u', "打开后请更新目录")
    doc.add_page_break()

    render_markdown_to_docx(doc, md)
    doc.save(OUT_DOCX)


def audit_outputs(md: str) -> list[str]:
    marker_hits = {m: md.count(m) for m in INTERNAL_MARKERS if md.count(m)}
    body = BODY_MD.read_text(encoding="utf-8")
    image_missing = [title for title, path in IMAGE_MAP.items() if title in body and not path.exists()]
    return [
        "# Phase12 Word Prep DOCX Build Report",
        "",
        "- status: `WORD_PREP_DOCX_BUILT_PENDING_MICROSOFT_WORD_VALIDATION`",
        f"- markdown: `{OUT_MD}`",
        f"- docx: `{OUT_DOCX}`",
        f"- body_entries: {body.count(chr(10) + '### ')}",
        f"- complete_option_blocks: {body.count('【完整选项】')}",
        f"- correct_item_blocks: {body.count('【正确项】')}",
        f"- wrong_trap_blocks: {body.count('【错项陷阱】')}",
        f"- images_expected: {len(IMAGE_MAP)}",
        f"- images_missing: {len(image_missing)}",
        f"- missing_image_titles: {', '.join(image_missing) if image_missing else 'none'}",
        f"- internal_marker_hits: {marker_hits if marker_hits else 0}",
        "- final_permission: `no`",
        "- next: open in Microsoft Word, update fields/TOC, save, export/render-check, then final gate.",
    ]


def main() -> int:
    OUT_MD.parent.mkdir(parents=True, exist_ok=True)
    QA_REPORT.parent.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    for path in IMAGE_MAP.values():
        if path.exists():
            shutil.copy2(path, ASSET_DIR / path.name)

    md = combined_markdown()
    OUT_MD.write_text(md, encoding="utf-8")
    build_docx(md)
    QA_REPORT.write_text("\n".join(audit_outputs(md)) + "\n", encoding="utf-8")

    print(f"wrote={OUT_MD}")
    print(f"wrote={OUT_DOCX}")
    print(f"wrote={QA_REPORT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
