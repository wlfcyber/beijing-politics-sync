from __future__ import annotations

import html
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase import pdfmetrics
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
DELIVERY = ROOT / "09_delivery"
MD_PATH = DELIVERY / "选必一_当代国际政治与经济_完整学生讲义_最终闭环版_20260504.md"
DOCX_PATH = DELIVERY / "选必一_当代国际政治与经济_完整学生讲义_最终闭环版_20260504.docx"
PDF_PATH = DELIVERY / "选必一_当代国际政治与经济_完整学生讲义_最终闭环版_20260504.pdf"
QA_PATH = DELIVERY / "document_generation_qa_最终闭环版_20260504.md"


TITLE = "选必一《当代国际政治与经济》主观题完整学生讲义"
FONT_EAST_ASIA = "Songti SC"
FONT_LATIN = "Arial"
ACCENT = RGBColor(25, 84, 125)
SCORING_RED = RGBColor(192, 0, 0)
SCORING_RED_HEX = "#C00000"
RED_SPAN_RE = re.compile(r'(<span style="color:#c00000">.*?</span>)')


def clean_inline(text: str) -> str:
    text = text.replace("`", "")
    text = text.replace("->", "→")
    return text.strip()


def add_field(paragraph, instr: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def list_number_abstract_id(doc: Document) -> str:
    numbering = doc.part.numbering_part.element
    for abstract in numbering.findall(qn("w:abstractNum")):
        for style in abstract.findall(".//" + qn("w:pStyle")):
            if style.get(qn("w:val")) == "ListNumber":
                return abstract.get(qn("w:abstractNumId"))
    for abstract in numbering.findall(qn("w:abstractNum")):
        for fmt in abstract.findall(".//" + qn("w:numFmt")):
            if fmt.get(qn("w:val")) == "decimal":
                return abstract.get(qn("w:abstractNumId"))
    return "0"


def new_numbering_id(doc: Document) -> str:
    numbering = doc.part.numbering_part.element
    used = [
        int(num.get(qn("w:numId")))
        for num in numbering.findall(qn("w:num"))
        if num.get(qn("w:numId")) and num.get(qn("w:numId")).isdigit()
    ]
    num_id = str(max(used or [0]) + 1)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), list_number_abstract_id(doc))
    num.append(abstract)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num = num_pr.find(qn("w:numId"))
    if num is None:
        num = OxmlElement("w:numId")
        num_pr.append(num)
    num.set(qn("w:val"), num_id)


def parse_md_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_md_table_separator(line: str) -> bool:
    cells = parse_md_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def is_md_table_start(lines: list[str], idx: int) -> bool:
    if idx + 1 >= len(lines):
        return False
    return lines[idx].strip().startswith("|") and is_md_table_separator(lines[idx + 1])


def collect_md_table(lines: list[str], idx: int) -> tuple[list[list[str]], int]:
    header = parse_md_table_row(lines[idx])
    rows = [header]
    idx += 2
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        row = parse_md_table_row(lines[idx])
        if len(row) == len(header):
            rows.append(row)
        idx += 1
    return rows, idx


def set_run_font(run, size: float | None = None, bold: bool | None = None, color=None) -> None:
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def inline_segments(text: str, bold: bool = False, red: bool = False):
    pattern = re.compile(r'(\*\*[^*]+\*\*|<span style="color:#c00000">.*?</span>)')
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            yield text[pos:match.start()], bold, red
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            yield from inline_segments(token[2:-2], bold=True, red=red)
        elif token.startswith('<span style="color:#c00000">') and token.endswith("</span>"):
            yield from inline_segments(token[len('<span style="color:#c00000">'):-len("</span>")], bold=bold, red=True)
        pos = match.end()
    if pos < len(text):
        yield text[pos:], bold, red


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, margin_twips: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ["top", "left", "bottom", "right"]:
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def add_markdown_runs(paragraph, text: str, size: float | None = None) -> None:
    text = clean_inline(text)
    for content, is_bold, is_red in inline_segments(text):
        if not content:
            continue
        run = paragraph.add_run(content)
        set_run_font(run, size=size, bold=is_bold, color=SCORING_RED if is_red else None)


def table_col_widths(cols: int) -> list[float]:
    if cols == 3:
        return [3.6, 6.2, 7.4]
    if cols == 2:
        return [5.2, 12.0]
    return [17.2 / cols] * cols


def add_docx_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False
    widths = table_col_widths(len(rows[0]))
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.width = Cm(widths[c_idx])
            set_cell_margins(cell, 100)
            if r_idx == 0:
                set_cell_shading(cell, "EAF2F8")
            paragraph = cell.paragraphs[0]
            paragraph.text = ""
            add_markdown_runs(paragraph, value, size=9.2)
            for run in paragraph.runs:
                run.font.bold = r_idx == 0 or run.font.bold
    doc.add_paragraph()


def normalize_styles(doc: Document) -> None:
    styles = doc.styles
    for name in ["Normal", "Body Text"]:
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        style.font.size = Pt(10.5)
    for name, size in [
        ("Title", 20),
        ("Heading 1", 16),
        ("Heading 2", 14),
        ("Heading 3", 12),
        ("Heading 4", 10.5),
    ]:
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT
    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        style.font.size = Pt(10.5)


def add_page_furniture(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = ""
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_p.add_run("飞哥政治庄园｜选必一完整学生讲义")
    set_run_font(run, size=9, bold=False, color=RGBColor(100, 100, 100))

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("第 ")
    set_run_font(run, size=9, color=RGBColor(120, 120, 120))
    add_field(footer_p, "PAGE")
    run = footer_p.add_run(" 页 / 共 ")
    set_run_font(run, size=9, color=RGBColor(120, 120, 120))
    add_field(footer_p, "NUMPAGES")
    run = footer_p.add_run(" 页")
    set_run_font(run, size=9, color=RGBColor(120, 120, 120))


def add_docx_paragraph(doc: Document, line: str, state: dict) -> None:
    stripped = line.strip()
    if not stripped:
        state["ordered_num_id"] = None
        return

    if stripped.startswith("# "):
        state["ordered_num_id"] = None
        p = doc.add_paragraph(style="Title")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_markdown_runs(p, stripped[2:], size=20)
        doc.add_paragraph()
        return
    if stripped.startswith("## "):
        state["ordered_num_id"] = None
        p = doc.add_paragraph(style="Heading 1")
        add_markdown_runs(p, stripped[3:], size=16)
        state["in_questions"] = "按题训练闭环" in stripped
        state["in_review"] = "六桶复盘" in stripped or "慎用" in stripped
        return
    if stripped.startswith("### "):
        state["ordered_num_id"] = None
        if state.get("in_questions") and re.match(r"### \d+\.", stripped):
            doc.add_paragraph()
        p = doc.add_paragraph(style="Heading 2")
        add_markdown_runs(p, stripped[4:], size=14)
        return
    if stripped.startswith("#### "):
        state["ordered_num_id"] = None
        p = doc.add_paragraph(style="Heading 3")
        add_markdown_runs(p, stripped[5:], size=12)
        return
    if re.match(r"^\d+\.\s+\*\*", stripped):
        state["ordered_num_id"] = None
        p = doc.add_paragraph(style="Heading 4")
        add_markdown_runs(p, stripped, size=10.5)
        return
    if stripped.startswith("- "):
        state["ordered_num_id"] = None
        p = doc.add_paragraph(style="List Bullet")
        add_markdown_runs(p, stripped[2:], size=10.5)
        return
    numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
    if numbered:
        if state.get("ordered_num_id") is None or numbered.group(1) == "1":
            state["ordered_num_id"] = new_numbering_id(doc)
        p = doc.add_paragraph(style="List Number")
        apply_numbering(p, state["ordered_num_id"])
        add_markdown_runs(p, numbered.group(2), size=10.5)
        return

    state["ordered_num_id"] = None
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08
    add_markdown_runs(p, stripped, size=10.5)


def build_docx(text: str) -> None:
    doc = Document()
    normalize_styles(doc)
    add_page_furniture(doc)
    state = {"in_questions": False, "in_review": False}
    lines = text.splitlines()
    idx = 0
    while idx < len(lines):
        if is_md_table_start(lines, idx):
            rows, idx = collect_md_table(lines, idx)
            add_docx_table(doc, rows)
            continue
        add_docx_paragraph(doc, lines[idx], state)
        idx += 1
    doc.core_properties.title = TITLE
    doc.core_properties.subject = "选必一 当代国际政治与经济 主观题学生讲义"
    doc.core_properties.author = "飞哥政治庄园"
    doc.save(DOCX_PATH)


def pdf_styles():
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    base = dict(fontName="STSong-Light", wordWrap="CJK")
    return {
        "title": ParagraphStyle(
            "title",
            parent=styles["Title"],
            fontName="STSong-Light",
            fontSize=20,
            leading=27,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#1a5276"),
            spaceAfter=16,
        ),
        "h1": ParagraphStyle(
            "h1",
            parent=styles["Heading1"],
            **base,
            fontSize=15,
            leading=21,
            textColor=colors.HexColor("#1a5276"),
            spaceBefore=12,
            spaceAfter=8,
        ),
        "h2": ParagraphStyle(
            "h2",
            parent=styles["Heading2"],
            **base,
            fontSize=13,
            leading=18,
            textColor=colors.HexColor("#1a5276"),
            spaceBefore=9,
            spaceAfter=6,
        ),
        "h3": ParagraphStyle(
            "h3",
            parent=styles["Heading3"],
            **base,
            fontSize=11.5,
            leading=16,
            textColor=colors.HexColor("#1a5276"),
            spaceBefore=7,
            spaceAfter=4,
        ),
        "body": ParagraphStyle(
            "body",
            parent=styles["BodyText"],
            **base,
            fontSize=10.2,
            leading=15,
            alignment=TA_LEFT,
            spaceAfter=4,
        ),
        "bullet": ParagraphStyle(
            "bullet",
            parent=styles["BodyText"],
            **base,
            fontSize=10.0,
            leading=14.5,
            leftIndent=14,
            firstLineIndent=-9,
            spaceAfter=3,
        ),
        "table_header": ParagraphStyle(
            "table_header",
            parent=styles["BodyText"],
            **base,
            fontSize=8.8,
            leading=12,
            textColor=colors.HexColor("#1a5276"),
        ),
        "table_cell": ParagraphStyle(
            "table_cell",
            parent=styles["BodyText"],
            **base,
            fontSize=8.6,
            leading=12,
            textColor=colors.HexColor("#111111"),
        ),
    }


def md_to_pdf_para(text: str) -> str:
    text = clean_inline(text)
    pieces: list[str] = []
    for content, is_bold, is_red in inline_segments(text):
        value = html.escape(content, quote=False)
        if is_bold:
            value = f"<b>{value}</b>"
        if is_red:
            value = f'<font color="{SCORING_RED_HEX}">{value}</font>'
        pieces.append(value)
    return "".join(pieces)


def pdf_table_col_widths(cols: int) -> list[float]:
    total = A4[0] - 3.1 * cm
    if cols == 3:
        weights = [0.19, 0.36, 0.45]
    elif cols == 2:
        weights = [0.30, 0.70]
    else:
        weights = [1 / cols] * cols
    return [total * weight for weight in weights]


def make_pdf_table(rows: list[list[str]], styles) -> Table:
    table_data = [
        [
            Paragraph(md_to_pdf_para(cell), styles["table_header" if r_idx == 0 else "table_cell"])
            for cell in row
        ]
        for r_idx, row in enumerate(rows)
    ]
    table = Table(table_data, colWidths=pdf_table_col_widths(len(rows[0])), repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EAF2F8")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1a5276")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B9C7D3")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def draw_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("STSong-Light", 8)
    canvas.setFillColor(colors.HexColor("#777777"))
    canvas.drawCentredString(A4[0] / 2, 0.9 * cm, f"飞哥政治庄园｜选必一完整学生讲义｜第 {doc.page} 页")
    canvas.restoreState()


def build_pdf(text: str) -> None:
    styles = pdf_styles()
    story = []
    in_questions = False
    lines = text.splitlines()
    idx = 0
    while idx < len(lines):
        if is_md_table_start(lines, idx):
            rows, idx = collect_md_table(lines, idx)
            story.append(make_pdf_table(rows, styles))
            story.append(Spacer(1, 8))
            continue
        raw = lines[idx]
        idx += 1
        line = raw.strip()
        if not line:
            story.append(Spacer(1, 2))
            continue
        if line.startswith("# "):
            story.append(Paragraph(md_to_pdf_para(line[2:]), styles["title"]))
            continue
        if line.startswith("## "):
            in_questions = "按题训练闭环" in line
            story.append(Paragraph(md_to_pdf_para(line[3:]), styles["h1"]))
            continue
        if line.startswith("### "):
            if in_questions and re.match(r"### \d+\.", line) and story:
                story.append(PageBreak())
            story.append(Paragraph(md_to_pdf_para(line[4:]), styles["h2"]))
            continue
        if line.startswith("#### "):
            story.append(Paragraph(md_to_pdf_para(line[5:]), styles["h3"]))
            continue
        if re.match(r"^\d+\.\s+\*\*", line):
            story.append(Paragraph(md_to_pdf_para(line), styles["h3"]))
            continue
        if line.startswith("- "):
            story.append(Paragraph("· " + md_to_pdf_para(line[2:]), styles["bullet"]))
            continue
        if re.match(r"^\d+\.\s+", line):
            story.append(Paragraph(md_to_pdf_para(line), styles["bullet"]))
            continue
        story.append(Paragraph(md_to_pdf_para(line), styles["body"]))

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=1.55 * cm,
        rightMargin=1.55 * cm,
        topMargin=1.55 * cm,
        bottomMargin=1.55 * cm,
        title=TITLE,
        author="飞哥政治庄园",
    )
    doc.build(story, onFirstPage=draw_footer, onLaterPages=draw_footer)


def structural_qa(text: str) -> list[str]:
    checks = []
    main_text = text.split("## 慎用与跨模块表达积累")[0]
    main_questions = len(re.findall(r"^### \d+\. ", main_text, re.M))
    counts = {
        "主线按题训练题": main_questions,
        "完整设问": main_text.count("**完整设问**"),
        "设问触发（题型通用）": main_text.count("**设问触发（题型通用）**"),
        "设问触发（本题独有）": main_text.count("**设问触发（本题独有）**"),
        "本题踩分点汇总": main_text.count("**本题踩分点汇总**"),
        "整题汇总卷面答案": main_text.count("**整题汇总卷面答案**"),
        "条目拆解": main_text.count("**条目拆解**"),
        "材料触发": main_text.count("材料触发："),
        "框架落点": main_text.count("框架落点："),
        "踩分词": main_text.count("踩分词："),
        "红色踩分词标记": main_text.count('<span style="color:#c00000">'),
        "答题点自身积累（可替换表达）": main_text.count("答题点自身积累（可替换表达，不必全背）："),
        "卷面答案句": main_text.count("卷面答案句"),
    }
    checks.append("## Structure Counts")
    for key, value in counts.items():
        checks.append(f"- {key}: {value}")

    forbidden = [
        "发展与和平",
        "经济全球化正确方向",
        "本题实际命中这些框架",
        "2026石景山期末",
        "/Users",
        "Desktop",
        "08_review",
        "07_student_doc",
        "SOURCE_LEDGER",
        "COVERAGE_MATRIX",
        "Codex",
        "Claude",
        "GPT",
        "Opus",
        "candidate",
        "guarded",
        "证据层级",
        "评标",
        "参考答案",
        "评分细则",
        "采分点",
        "要落到",
        "设问要求",
        "非累计",
        "替代表述",
        "同一评分位置",
        "同槽",
        "不并入",
        "fallback",
        " 的NDC",
        " 的公共产品",
        "，的",
        "：的",
    ]
    hits = {term: text.count(term) for term in forbidden if text.count(term)}
    checks.append("\n## Forbidden-Term Scan")
    if hits:
        for term, count in hits.items():
            checks.append(f"- HIT {term}: {count}")
    else:
        checks.append("- PASS: no hits")

    checks.append("\n## Output Files")
    for path in [MD_PATH, DOCX_PATH, PDF_PATH]:
        checks.append(f"- {path.name}: {path.stat().st_size} bytes")
    reader = PdfReader(str(PDF_PATH))
    checks.append(f"- PDF pages: {len(reader.pages)}")
    return checks


def main() -> None:
    text = MD_PATH.read_text(encoding="utf-8")
    build_docx(text)
    build_pdf(text)
    QA_PATH.write_text("\n".join(structural_qa(text)) + "\n", encoding="utf-8")
    print(DOCX_PATH)
    print(PDF_PATH)
    print(QA_PATH)


if __name__ == "__main__":
    main()
