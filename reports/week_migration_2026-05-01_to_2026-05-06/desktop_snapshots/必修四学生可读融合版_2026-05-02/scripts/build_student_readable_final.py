#!/usr/bin/env python3
from __future__ import annotations

import importlib.util
import json
import os
import re
import shutil
from collections import Counter
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path("/Users/wanglifei/Desktop/北京高考政治")
WORK = ROOT / "必修四学生可读融合版_2026-05-02"
AUDIT = WORK / "audit"
REPORTS = WORK / "reports"
OUTPUTS = WORK / "outputs"

PREPARE_SCRIPT = WORK / "scripts" / "prepare_student_readability_inputs.py"
FUSION_SCRIPT = ROOT / "必修四终极融合版_2026-05-02" / "scripts" / "build_ultimate_fusion.py"
REBUILD_PATH = ROOT / "哲学宝典5.2原地修订" / "scripts" / "rebuild_in_place_baodian.py"

COMMON_PAIRS = AUDIT / "common_entry_pairs_for_student_agent.jsonl"
CODEX_ONLY = AUDIT / "codex_only_entries.json"
CLAUDE_ONLY = AUDIT / "claude_only_entries.json"
STUDENT_DECISIONS = AUDIT / "student_agent_decisions.jsonl"

FINAL_STEM = "2026北京高考政治哲学宝典---三年模拟全触发全链条_学生可读终极版"
FINAL_MD = OUTPUTS / f"{FINAL_STEM}.md"
FINAL_DOCX = OUTPUTS / f"{FINAL_STEM}.docx"
DESKTOP_DOCX = Path("/Users/wanglifei/Desktop/哲学宝典最终版 5.2学生可读终极版.docx")
TITLE = "2026北京高考政治哲学宝典---三年模拟全触发全链条"
DOC_TITLE = "2026北京高考政治哲学宝典\n三年模拟全触发全链条"
SIGNATURE = "飞哥正志讲堂"
STUDENT_KEYS = ["材料触发点", "设问", "为什么能想到", "答案落点"]

FORBIDDEN_RE = re.compile(
    r"细则|评标|参考答案|答案写|答案核|答案/补充|可从.*角度|/Users/|"
    r"\.pdf|\.docx|\.pptx|OCR|debug|slide|line id|file id|PASS|"
    r"correct_option_chain|filled",
    re.IGNORECASE,
)
AWARENESS_MATERIAL_RE = re.compile(
    r"意识对物质|意识反作用于物质|意识[^。；;，,]*对物质[^。；;，,]*反作用|"
    r"意识[^。；;，,]*能动[^。；;，,]*反作用于物质"
)

spec = importlib.util.spec_from_file_location("prepare_inputs", PREPARE_SCRIPT)
prepare = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(prepare)

fusion_spec = importlib.util.spec_from_file_location("fusion", FUSION_SCRIPT)
fusion = importlib.util.module_from_spec(fusion_spec)
assert fusion_spec.loader is not None
fusion_spec.loader.exec_module(fusion)

rebuild_spec = importlib.util.spec_from_file_location("rebuild", REBUILD_PATH)
rebuild = importlib.util.module_from_spec(rebuild_spec)
assert rebuild_spec.loader is not None
rebuild_spec.loader.exec_module(rebuild)


def clean_text(text: str) -> str:
    return rebuild.clean_text(text)


def key_tuple(key: dict) -> tuple[str, str, str, str, str]:
    return (key["module"], key["node"], key["suite"], key["q"], key["qtype"])


def entry_key(e: dict) -> tuple[str, str, str, str, str]:
    return (
        e["module"],
        e["node"],
        rebuild.normalize_suite(e["suite"]),
        rebuild.normalize_q(e["q"]),
        e["qtype"],
    )


def heading_from_key(key: dict) -> str:
    return rebuild.title_from_parts(key["suite"], key["q"], key["qtype"])


def load_jsonl(path: Path) -> list[dict]:
    rows = []
    for line in path.read_text(encoding="utf-8").splitlines():
        if line.strip():
            rows.append(json.loads(line))
    return rows


def load_decisions() -> dict[str, dict]:
    decisions = {}
    if not STUDENT_DECISIONS.exists():
        return decisions
    for row in load_jsonl(STUDENT_DECISIONS):
        pair_id = row.get("pair_id")
        if pair_id:
            decisions[pair_id] = row
    return decisions


def fallback_choice(pair: dict) -> dict:
    codex_fields = pair["codex"]["fields"]
    claude_fields = pair["claude"]["fields"]
    choices = {}
    chosen = {}
    for key in STUDENT_KEYS:
        c = codex_fields.get(key, "")
        d = claude_fields.get(key, "")
        if key == "设问":
            pick = "claude" if len(d) >= len(c) and not d.startswith("选择题：根据题干材料") else "codex"
        elif key == "材料触发点":
            pick = "claude" if len(d) > len(c) + 20 and len(d) < 520 else "codex"
        else:
            pick = "claude" if len(d) > len(c) + 40 and len(d) < 900 else "codex"
        choices[key] = pick
        chosen[key] = clean_text(d if pick == "claude" else c)
    preferred = Counter(choices.values()).most_common(1)[0][0]
    return {
        "pair_id": pair["pair_id"],
        "preferred_source": preferred,
        "field_choices": choices,
        "chosen_fields": chosen,
        "readable_reason": "fallback: choose more concrete material and fuller answer while avoiding overlong text",
        "reject_reason": "",
    }


def fields_from_decision(pair: dict, decision: dict | None) -> tuple[dict[str, str], str, str]:
    if not decision:
        decision = fallback_choice(pair)
    fields = {}
    chosen_fields = decision.get("chosen_fields") or {}
    field_choices = decision.get("field_choices") or {}
    for key in STUDENT_KEYS:
        choice = field_choices.get(key) or decision.get("preferred_source") or "codex"
        if key in chosen_fields and chosen_fields[key]:
            value = chosen_fields[key]
        elif choice == "claude":
            value = pair["claude"]["fields"].get(key, "")
        elif choice == "mix":
            value = chosen_fields.get(key) or pair["claude"]["fields"].get(key) or pair["codex"]["fields"].get(key, "")
        else:
            value = pair["codex"]["fields"].get(key, "")
        fields[key] = clean_text(value)
    source = decision.get("preferred_source", "fallback")
    return fields, source, decision.get("readable_reason", "")


def hard_rule_reason(node: str, blob: str) -> str:
    if not AWARENESS_MATERIAL_RE.search(blob):
        return ""
    if node == "物质决定意识":
        return "material-node-mixes-awareness-reaction"
    if node != "主观能动性 / 意识的能动作用":
        return "awareness-material-reaction-outside-initiative-node"
    if "物质决定意识" in blob:
        return "initiative-node-combines-material-and-awareness-reaction"
    return ""


def apply_local_hard_rules(e: dict) -> dict:
    e = json.loads(json.dumps(e, ensure_ascii=False))
    q_norm = rebuild.normalize_q(e.get("q", ""))
    if (
        e.get("suite") == "2024朝阳二模"
        and q_norm in {"16(2)", "16((2))"}
        and e.get("node") == "主观能动性 / 意识的能动作用"
    ):
        e["fields"]["为什么能想到"] = (
            "材料写科学家受神经系统研究启发开发人工神经网络，推动深度学习和类脑智能，"
            "并且设问追问未来如何处理人与人工智能的关系。这里学生要抓住的是：人不是被技术牵着走，"
            "而是能够在认识技术条件和发展趋势的基础上，有目的、有计划地设计、使用、规范和引导人工智能。"
            "因此这一栏只落到主观能动性 / 意识的能动作用。"
        )
        e["fields"]["答案落点"] = (
            "从主观能动性角度看，未来不能把人与人工智能的关系理解成技术单方面塑造人。"
            "人应当主动认识人工智能的发展条件、应用场景和潜在风险，通过科研创新、制度规范、伦理引导和价值选择，"
            "能动地塑造人工智能的发展方向，使人工智能更好服务人的发展。"
        )
        e["student_merge_action"] = e.get("student_merge_action", "") + "+hard-rule-overridden"
    return e


def build_entries() -> tuple[list[dict], dict]:
    common_pairs = load_jsonl(COMMON_PAIRS)
    decisions = load_decisions()
    codex_only = json.loads(CODEX_ONLY.read_text(encoding="utf-8"))
    claude_only = json.loads(CLAUDE_ONLY.read_text(encoding="utf-8"))
    entries = []
    logs = {"common": [], "codex_only": [], "claude_only": [], "missing_decisions": []}

    for pair in common_pairs:
        decision = decisions.get(pair["pair_id"])
        if not decision:
            logs["missing_decisions"].append(pair["pair_id"])
        fields, source, reason = fields_from_decision(pair, decision)
        key = pair["key"]
        e = {
            "module": key["module"],
            "node": key["node"],
            "suite": key["suite"],
            "q": key["q"],
            "qtype": key["qtype"],
            "heading_base": heading_from_key(key),
            "fields": fields,
            "origin": "student-readable-common",
            "image": "",
            "student_merge_action": f"common-{source}",
            "student_pair_id": pair["pair_id"],
            "student_reason": reason,
            "codex_origin": pair["codex"].get("origin", ""),
            "claude_source_entry_id": pair["claude"].get("source_entry_id", ""),
        }
        e = apply_local_hard_rules(e)
        entries.append(e)
        logs["common"].append({"pair_id": pair["pair_id"], "source": source, "field_choices": (decision or {}).get("field_choices", {})})

    for row in codex_only:
        key = row["key"]
        e = {
            "module": key["module"],
            "node": key["node"],
            "suite": key["suite"],
            "q": key["q"],
            "qtype": key["qtype"],
            "heading_base": row.get("heading_base") or heading_from_key(key),
            "fields": {k: clean_text(row.get("fields", {}).get(k, "")) for k in STUDENT_KEYS},
            "origin": row.get("origin", "codex-only"),
            "image": "",
            "student_merge_action": "codex-only-kept",
        }
        entries.append(apply_local_hard_rules(e))
        logs["codex_only"].append(entry_key(e))

    for row in claude_only:
        key = row["key"]
        e = {
            "module": key["module"],
            "node": key["node"],
            "suite": key["suite"],
            "q": key["q"],
            "qtype": key["qtype"],
            "heading_base": row.get("heading_base") or heading_from_key(key),
            "fields": {k: clean_text(row.get("fields", {}).get(k, "")) for k in STUDENT_KEYS},
            "origin": "claude-only-readable",
            "image": "",
            "student_merge_action": "claude-only-kept",
            "claude_source_entry_id": row.get("source_entry_id", ""),
        }
        entries.append(apply_local_hard_rules(e))
        logs["claude_only"].append(entry_key(e))

    # De-duplicate in case q normalization makes a pair collide after loading.
    by_key = {}
    removed = []
    for e in entries:
        key = entry_key(e)
        if key in by_key:
            previous = by_key[key]
            previous_len = sum(len(previous["fields"].get(k, "")) for k in STUDENT_KEYS)
            current_len = sum(len(e["fields"].get(k, "")) for k in STUDENT_KEYS)
            if current_len > previous_len:
                removed.append(previous)
                by_key[key] = e
            else:
                removed.append(e)
        else:
            by_key[key] = e
    logs["dedup_removed"] = removed
    entries = sorted(by_key.values(), key=rebuild.sort_key)
    return entries, logs


def write_markdown(entries: list[dict], path: Path) -> None:
    lines = [f"# {TITLE}", "", f"**{SIGNATURE}**", "", "## 前言", "", "（此处留给飞哥老师自行撰写。）", ""]
    current_module = current_node = None
    for e in entries:
        if e["module"] != current_module:
            current_module = e["module"]
            current_node = None
            lines.extend([f"## {current_module}", ""])
        if e["node"] != current_node:
            current_node = e["node"]
            lines.extend([f"### {current_node}", ""])
        lines.extend([f"#### {e['heading_base']}", ""])
        for key in STUDENT_KEYS:
            lines.extend([f"【{key}】 {e['fields'][key]}", ""])
    path.write_text("\n".join(lines), encoding="utf-8")


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def polish_docx_cover(path: Path) -> None:
    doc = Document(path)
    title_index = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.startswith("2026北京高考政治哲学宝典"):
            title_index = idx
            break
    if title_index is not None:
        for paragraph in list(doc.paragraphs[:title_index]):
            if not paragraph.text.strip():
                remove_paragraph(paragraph)
        title = doc.paragraphs[0]
        title.paragraph_format.space_before = Pt(215)
        title.paragraph_format.space_after = Pt(28)
    doc.save(path)


def write_docx(entries: list[dict], path: Path) -> None:
    tmp = path.with_suffix(".tmp.docx")
    rebuild.TITLE = DOC_TITLE
    rebuild.build_doc(entries, tmp)
    polish_docx_cover(tmp)
    os.replace(tmp, path)


def qa(entries: list[dict], logs: dict) -> dict:
    body = "\n".join("\n".join(e["fields"].values()) for e in entries)
    keys = Counter(entry_key(e) for e in entries)
    hard_hits = []
    for e in entries:
        blob = "\n".join(e["fields"].values())
        reason = hard_rule_reason(e["node"], blob)
        if reason:
            hard_hits.append({"reason": reason, "heading": e["heading_base"], "node": e["node"], "action": e.get("student_merge_action", "")})
    return {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "final_entries": len(entries),
        "action_counts": dict(Counter(e.get("student_merge_action", "") for e in entries)),
        "common_entries": len(logs["common"]),
        "codex_only": len(logs["codex_only"]),
        "claude_only": len(logs["claude_only"]),
        "missing_student_decisions": logs["missing_decisions"],
        "dedup_removed": len(logs.get("dedup_removed", [])),
        "forbidden_hits": sorted(set(m.group(0) for m in FORBIDDEN_RE.finditer(body))),
        "missing_student_fields": [e["heading_base"] for e in entries if any(not e["fields"].get(k) for k in STUDENT_KEYS)],
        "duplicate_keys": [str(k) for k, count in keys.items() if count > 1],
        "hard_rule_hits": hard_hits,
        "hard_samples": {
            "2026海淀一模_Q16": [(e["module"], e["node"], e.get("student_merge_action", "")) for e in entries if e["suite"] == "2026海淀一模" and rebuild.normalize_q(e["q"]) == "16"],
            "2024朝阳二模_Q16(2)": [(e["module"], e["node"], e.get("student_merge_action", "")) for e in entries if e["suite"] == "2024朝阳二模" and rebuild.normalize_q(e["q"]) in {"16(2)", "16((2))"}],
            "2026石景山期末": [e["heading_base"] for e in entries if e["suite"] == "2026石景山期末"],
        },
    }


def write_reports(qa_data: dict, logs: dict) -> None:
    lines = [
        "# 必修四学生可读终极版融合报告 2026-05-02",
        "",
        "## 条目有无",
        "",
        f"- 两边都有：{qa_data['common_entries']}",
        f"- Codex 独有保留：{qa_data['codex_only']}",
        f"- Claude 独有保留：{qa_data['claude_only']}",
        f"- 最终条目：{qa_data['final_entries']}",
        "",
        "## 学生可读融合",
        "",
        "- 对两边都有的 136 条，按学生智能体输出逐字段选择。",
        "- 若学生智能体未覆盖某条，使用本地保守 fallback：设问优先完整，材料和答案优先更具体但不过长。",
        "- 独有条目不互删：先保留，再由硬规则和禁词 QA 兜底。",
        "",
        "## QA",
        "",
        f"- 学生字段禁词：{qa_data['forbidden_hits']}",
        f"- 四字段缺失：{len(qa_data['missing_student_fields'])}",
        f"- 重复键：{len(qa_data['duplicate_keys'])}",
        f"- 小本本硬规则违规：{len(qa_data['hard_rule_hits'])}",
        f"- 缺学生决策 pair：{qa_data['missing_student_decisions']}",
        "",
        "## 输出",
        "",
        f"- Markdown：{FINAL_MD}",
        f"- Word：{FINAL_DOCX}",
        f"- 桌面 Word：{DESKTOP_DOCX}",
    ]
    (REPORTS / "必修四学生可读终极版融合报告_2026-05-02.md").write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    OUTPUTS.mkdir(parents=True, exist_ok=True)
    AUDIT.mkdir(parents=True, exist_ok=True)
    REPORTS.mkdir(parents=True, exist_ok=True)
    entries, logs = build_entries()
    qa_data = qa(entries, logs)
    if qa_data["forbidden_hits"] or qa_data["missing_student_fields"] or qa_data["duplicate_keys"] or qa_data["hard_rule_hits"]:
        raise SystemExit(json.dumps(qa_data, ensure_ascii=False, indent=2))
    (AUDIT / "student_readable_final_entries.json").write_text(json.dumps(entries, ensure_ascii=False, indent=2), encoding="utf-8")
    (AUDIT / "student_readable_merge_log.json").write_text(json.dumps(logs, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
    (AUDIT / "student_readable_final_qa.json").write_text(json.dumps(qa_data, ensure_ascii=False, indent=2), encoding="utf-8")
    write_markdown(entries, FINAL_MD)
    write_docx(entries, FINAL_DOCX)
    shutil.copy2(FINAL_DOCX, DESKTOP_DOCX)
    write_reports(qa_data, logs)
    print(json.dumps(qa_data, ensure_ascii=False, indent=2))
    print(f"FINAL_DOCX={FINAL_DOCX}")
    print(f"DESKTOP_DOCX={DESKTOP_DOCX}")


if __name__ == "__main__":
    main()
