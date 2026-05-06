#!/usr/bin/env python3
"""把学生版 Markdown 转为 Word（教学排版，干净版面，无彩色块）。

封面：仅标题 + 飞哥正志讲堂 大署名。
第二页：前言占位（飞哥稍后填写）。
其后：按节点正文。
"""
from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("/Users/wanglifei/Desktop/北京高考政治/必修四从0重跑_ClaudeCode_2026-05-02")
SCRIPTS_DIR = ROOT / "scripts"
ENTRIES_DIR = ROOT / "audit" / "entries"
TITLE = "2026 北京高考政治哲学宝典——三年模拟全触发全链条"
BLUE_TITLE = RGBColor(0x18, 0x45, 0x6B)
BLUE_H1 = RGBColor(0x1F, 0x4E, 0x79)
BLUE_H2 = RGBColor(0x2F, 0x6F, 0x9F)
BLUE_H3 = RGBColor(0x3A, 0x6B, 0x88)

# 与 synthesize_student.py 共用的框架与匹配
import sys
sys.path.insert(0, str(SCRIPTS_DIR))
from synthesize_student import (  # noqa: E402
    FRAMEWORK, DISTRICT_ORDER, STAGE_ORDER, FORBIDDEN_TOKENS,
    normalize_node_path, match_framework_node, district_rank, stage_rank,
    load_all_entries, sanitize, render_question_label,
)


def _set_zh_font(
    run,
    size_pt: float,
    bold: bool = False,
    name_zh: str = "宋体",
    name_ascii: str = "Times New Roman",
    color: Optional[RGBColor] = None,
):
    run.bold = bold
    run.font.name = name_ascii
    run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name_zh)
    rFonts.set(qn("w:ascii"), name_ascii)
    rFonts.set(qn("w:hAnsi"), name_ascii)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def build_cover(doc: Document):
    # 上方留白：插入若干空行
    for _ in range(5):
        doc.add_paragraph("")
    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(TITLE)
    _set_zh_font(run, 30, bold=True, name_zh="黑体", name_ascii="Times New Roman", color=BLUE_TITLE)
    # 中部留白
    for _ in range(8):
        doc.add_paragraph("")
    # 大署名
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("飞哥正志讲堂")
    _set_zh_font(run, 36, bold=True, name_zh="楷体", name_ascii="Times New Roman")


def build_foreword(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("前  言")
    _set_zh_font(run, 22, bold=True, name_zh="黑体")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run("（此页留位，由飞哥本人填写。）")
    _set_zh_font(run, 12, bold=False, name_zh="楷体")


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_zh_font(run, 22, bold=True, name_zh="黑体", color=BLUE_H1)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _set_zh_font(run, 16, bold=True, name_zh="黑体", color=BLUE_H2)


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _set_zh_font(run, 13, bold=True, name_zh="楷体", color=BLUE_H3)


def add_field(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.5)
    run_l = p.add_run(f"{label}：")
    _set_zh_font(run_l, 12, bold=True, name_zh="黑体")
    run_v = p.add_run(value)
    _set_zh_font(run_v, 12, bold=False, name_zh="宋体")


def add_blank(doc, pt: float = 4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)


def add_empty_node_note(doc):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("（本轮三年模拟未放入稳定案例。）")
    _set_zh_font(run, 11, bold=False, name_zh="楷体")


def main():
    entries = load_all_entries()
    buckets: dict = {}
    for e in entries:
        cat, node = normalize_node_path(e.get("target_node_path", ""))
        m = match_framework_node(cat, node)
        if not m or not m[1]:
            continue
        buckets.setdefault(m, []).append(e)

    doc = Document()
    # 全局默认字号
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")

    # 1) 封面
    build_cover(doc)
    add_page_break(doc)
    # 2) 前言
    build_foreword(doc)
    add_page_break(doc)

    # 3) 正文：按框架节点
    for fw_cat, nodes in FRAMEWORK:
        add_h1(doc, fw_cat)
        for fw_node in nodes:
            add_h2(doc, fw_node)
            ents = buckets.get((fw_cat, fw_node), [])
            if not ents:
                add_empty_node_note(doc)
                continue
            ents.sort(key=lambda e: (
                0 if e.get("question_type") == "subjective" or e.get("sort_priority") == "main_question" else 1,
                district_rank(e.get("district", "")),
                -int(e.get("year", 0) or 0),
                stage_rank(e.get("stage", "")),
                str(e.get("question_no", "")),
            ))
            for e in ents:
                label = render_question_label(e)
                add_h3(doc, label)
                add_field(doc, "材料触发点", sanitize(e.get("材料触发点", ""), "材料触发点"))
                add_field(doc, "设问", sanitize(e.get("设问", ""), "设问"))
                add_field(doc, "为什么能想到", sanitize(e.get("为什么能想到", ""), "为什么能想到"))
                add_field(doc, "答案落点", sanitize(e.get("答案落点", ""), "答案落点"))
                add_blank(doc, 4)

    out = ROOT / "outputs" / "2026北京高考政治哲学宝典---三年模拟全触发全链条_学生版.docx"
    doc.save(str(out))
    print(f"OK 学生版 Word -> {out}")


if __name__ == "__main__":
    main()
