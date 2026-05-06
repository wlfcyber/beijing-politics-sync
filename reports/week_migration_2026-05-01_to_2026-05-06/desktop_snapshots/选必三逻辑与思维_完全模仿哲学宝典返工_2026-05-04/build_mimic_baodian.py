from __future__ import annotations

import csv
import re
import zipfile
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/wanglifei/Desktop/北京高考政治")
RUN = ROOT / "选必三逻辑与思维_完全模仿哲学宝典返工_2026-05-04"
OLD_RUN = ROOT / "选必三逻辑与思维双线四线终极跑_2026-05-04"

THINKING_SOURCE = ROOT / "选必三_逻辑与思维_思维部分_逐题材料-思维路径积累_穷尽版.md"
REASONING_MATRIX = OLD_RUN / "claudecode_lane/reasoning_type_matrix.csv"
REASONING_REPORT = OLD_RUN / "codex_lane/agents/reasoning_worker_report.md"
OUT_MD = RUN / "04_delivery/选必三逻辑与思维_完全模仿哲学宝典_学生版_2026-05-04.md"
OUT_DOCX = RUN / "04_delivery/选必三逻辑与思维_完全模仿哲学宝典_学生版_2026-05-04.docx"
LEDGER_MD = RUN / "02_ledgers/coverage_counts.md"


BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(35, 35, 35)
MUTED = RGBColor(100, 100, 100)


def clean_text(s: str) -> str:
    if s is None:
        return ""
    s = s.replace("`", "")
    s = s.replace("->", "→")
    s = s.replace("A-formal", "")
    s = s.replace("B-choice-signal", "")
    s = s.replace("candidate", "")
    s = s.replace("cache", "")
    s = s.replace("待核验", "")
    s = s.replace("回源核验", "")
    s = s.replace("候选", "")
    s = re.sub(r"\s+", " ", s).strip()
    s = s.replace("细则实际", "题目落点")
    s = s.replace("细则", "题目落点")
    s = s.replace("题面写", "题目表述为")
    s = s.replace("总帽子", "总方法")
    s = s.replace("。。", "。")
    return s


def parse_thinking_entries() -> list[dict]:
    text = THINKING_SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    entries: list[dict] = []
    section = ""
    current: dict | None = None

    def flush() -> None:
        if current:
            for key in ["材料信息", "触发思维", "路径积累"]:
                current[key] = clean_text(current.get(key, ""))
            if current.get("材料信息") or current.get("路径积累"):
                entries.append(current.copy())

    for line in lines:
        if line.startswith("## 一、A类"):
            section = "A主观题"
            continue
        if line.startswith("## 二、B类"):
            flush()
            current = None
            section = "B选择题"
            continue
        if line.startswith("## 三、C类"):
            flush()
            current = None
            section = "C边界"
            continue
        m = re.match(r"^### `([^`]+)`", line)
        if m and section in {"A主观题", "B选择题"}:
            flush()
            current = {"label": clean_text(m.group(1)), "kind": section}
            continue
        if current and line.startswith("- "):
            m2 = re.match(r"- ([^：:]+)[：:](.*)", line)
            if m2:
                current[m2.group(1).strip()] = clean_text(m2.group(2))
    flush()
    return entries


NODE_RULES = [
    ("科学思维", "科学思维三特征：客观性、预见性、可检验性", ["科学思维", "客观性", "预见性", "可检验性", "真实", "实验", "测试", "验证", "调研", "调查"]),
    ("科学思维", "调查研究与研究设计题", ["调查研究", "问卷", "访谈", "实地调研", "研究方法", "调研活动", "项目研究"]),
    ("辩证思维", "分析与综合 / 整体性", ["分析与综合", "整体性", "整体", "综合", "多主体", "协同", "一盘棋", "多功能", "统筹"]),
    ("辩证思维", "动态性 / 质量互变", ["动态性", "质量互变", "量变", "质变", "渐进", "持续", "日积月累", "由小到大", "阶梯"]),
    ("辩证思维", "辩证否定 / 适度原则 / 矛盾分析", ["辩证否定", "适度", "矛盾", "挑战", "一分为二", "扬弃", "守正创新", "旧资源"]),
    ("辩证思维", "思维抽象与思维具体", ["思维抽象", "思维具体", "感性具体", "杂多", "河—合—和"]),
    ("创新思维", "联想思维 / 迁移和想象", ["联想", "迁移", "想象", "文创", "意象", "跨界", "嫁接", "类比"]),
    ("创新思维", "发散思维与聚合思维", ["发散", "聚合", "多方案", "多方向", "多元", "收束"]),
    ("创新思维", "逆向思维", ["逆向", "反向", "倒过来", "从人找书到书找人", "被动到主动", "静态资源"]),
    ("创新思维", "超前思维", ["超前", "未来", "趋势", "提前", "前瞻", "预判", "布局", "中长期"]),
    ("创新思维", "创新思维综合题", ["创新思维", "创新", "新场景", "新业态", "新产品", "首发经济"]),
]


def node_for_entry(entry: dict) -> list[tuple[str, str]]:
    blob = " ".join([entry.get("材料信息", ""), entry.get("触发思维", ""), entry.get("路径积累", "")])
    hits: list[tuple[str, str]] = []
    for major, node, terms in NODE_RULES:
        if any(t in blob for t in terms):
            hits.append((major, node))
    if entry["kind"] == "B选择题" and not hits:
        hits.append(("选择题陷阱库", "综合选择题信号与错项排除"))
    if not hits:
        hits.append(("综合复合题", "多方法综合迁移题"))
    dedup = []
    seen = set()
    for x in hits:
        if x not in seen:
            dedup.append(x)
            seen.add(x)
    return dedup[:5]


def make_thinking_entry(entry: dict, idx: int) -> list[str]:
    kind = "选择题" if entry["kind"] == "B选择题" else "主观题"
    material = entry.get("材料信息", "")
    trigger = entry.get("触发思维", "")
    path = entry.get("路径积累", "")
    if not trigger:
        trigger = "本题考查材料信号与思维方法、错项之间的匹配关系。"
    if entry["kind"] == "B选择题":
        ask = "本题要求依据材料判断正确选项，并排除把思维方法乱套、夸大或错位的选项。"
        answer = f"{path} 因此做选择题时先圈材料动作，再看选项是否把方法、主体、关系或作用说准。"
    else:
        ask = "本题要求结合材料说明其体现的思维方法，或运用相关思维方法分析、评价、提出做法。"
        answer = f"卷面要把材料中的具体动作写进方法里：{path}"
    return [
        f"### {idx}. {entry['label']}（{kind}）",
        f"【材料触发点】 {material}",
        f"【设问】 {ask}",
        f"【为什么能想到】 {trigger}。{path}".replace("。。", "。"),
        f"【答案落点】 {answer}".replace("。。", "。"),
        "",
    ]


def parse_reasoning_matrix() -> list[dict]:
    rows = []
    with REASONING_MATRIX.open(encoding="utf-8") as f:
        for row in csv.DictReader(f):
            rows.append({k: clean_text(v) for k, v in row.items()})
    return rows


def parse_reasoning_report_extras(existing_labels: set[str]) -> list[dict]:
    text = REASONING_REPORT.read_text(encoding="utf-8").splitlines()
    family = ""
    in_high = False
    extras = []
    for line in text:
        m = re.match(r"^### \d+\. (.+)$", line)
        if m:
            family = clean_text(m.group(1))
            in_high = False
            continue
        if line.strip() == "**高价值候选**":
            in_high = True
            continue
        if in_high and line.startswith("**"):
            in_high = False
        if in_high and line.startswith("- `"):
            m2 = re.match(r"- `([^`]+)`[：:](.*)", line)
            if not m2:
                continue
            label = clean_text(m2.group(1))
            if label in existing_labels:
                continue
            desc = clean_text(m2.group(2))
            extras.append(
                {
                    "suite": label,
                    "question": "",
                    "reasoning_family": family,
                    "rule_or_error": desc,
                    "tested_move": desc,
                    "common_trap": "同类题容易把判断类型、推理规则或有效式混用。",
                    "solution_protocol": "先翻译逻辑结构，再套本题型最小规则，最后判断能否必然推出。",
                    "notes": "题型报告中的同类题。",
                }
            )
    return extras


def family_key(row: dict) -> str:
    fam = row.get("reasoning_family") or row.get("family") or "综合逻辑题"
    if "三段论" in fam:
        return "三段论"
    if "假言" in fam:
        return "假言推理"
    if "选言" in fam:
        return "选言判断与选言推理"
    if "联言" in fam:
        return "联言判断与联言推理"
    if "归纳" in fam or "求异" in fam:
        return "归纳推理与探求因果联系"
    if "类比" in fam:
        return "类比推理"
    if "周延" in fam:
        return "周延判断"
    if "换质" in fam or "换位" in fam:
        return "换质换位"
    if "三律" in fam or "概念" in fam:
        return "逻辑三律与概念规则"
    return fam


FAMILY_INTROS = {
    "三段论": "先找大项、小项、中项，再查四条规则：三项是否唯一，中项是否至少周延一次，大小项是否不当扩大，是否两个否定前提。",
    "假言推理": "先把自然语言翻译成充分、必要或充要条件，再只套对应有效式，不能凭生活常识倒推。",
    "选言判断与选言推理": "先判相容还是不相容。相容选言不能肯定一支就否定另一支，不相容选言才可以肯定否定。",
    "联言判断与联言推理": "联言判断各支都真才真；分解式可以从整体推出各支，合成式要各支都真才能合成整体。",
    "归纳推理与探求因果联系": "归纳从个别到一般，不完全归纳是或然的；求同、求异、共变、剩余法是探求因果联系的方法。",
    "类比推理": "类比从一个对象迁移到另一个对象，相似根据越多、越本质，结论越可靠。",
    "周延判断": "主项看量项，谓项看联项。全称和单称主项周延，否定判断谓项周延，肯定判断谓项不周延。",
    "换质换位": "换质两变两不变；换位主谓互换但不能扩大不周延项。遇到换质位要一步一步做。",
    "逻辑三律与概念规则": "同一律防偷换，矛盾律防同真，排中律防两不可；概念题还要查定义和划分是否合规。",
}


def make_reasoning_entry(row: dict, idx: int) -> list[str]:
    suite = row.get("suite", "")
    q = row.get("question", "")
    label = f"{suite} {q}".strip()
    if not label:
        label = f"同类题 {idx}"
    fam = family_key(row)
    material = row.get("tested_move") or row.get("rule_or_error") or fam
    if "/Users/" in material or material.startswith("/"):
        material = row.get("rule_or_error") or FAMILY_INTROS.get(fam, fam)
    ask = f"本题要求识别或运用“{fam}”规则，判断推理是否成立、错在何处，或补全合乎规则的推理。"
    why = row.get("rule_or_error") or FAMILY_INTROS.get(fam, "先识别逻辑结构，再套对应规则。")
    trap = row.get("common_trap", "")
    protocol = row.get("solution_protocol") or FAMILY_INTROS.get(fam, "")
    if "/Users/" in protocol or protocol.startswith("/") or "Desktop/" in protocol:
        protocol = FAMILY_INTROS.get(fam, "先翻译逻辑结构，再套本题型最小规则。")
    return [
        f"### {idx}. {label}（推理题）",
        f"【材料触发点】 {material}",
        f"【设问】 {ask}",
        f"【为什么能想到】 {why}。高频陷阱：{trap}",
        f"【答案落点】 {protocol}",
        "",
    ]


def build_markdown() -> str:
    entries = parse_thinking_entries()
    node_map: dict[tuple[str, str], list[dict]] = defaultdict(list)
    for e in entries:
        for node in node_for_entry(e):
            node_map[node].append(e)

    reasoning_rows = parse_reasoning_matrix()
    existing = {clean_text(f"{r.get('suite','')} {r.get('question','')}").strip() for r in reasoning_rows}
    reasoning_rows.extend(parse_reasoning_report_extras(existing))
    family_map: dict[str, list[dict]] = defaultdict(list)
    for r in reasoning_rows:
        family_map[family_key(r)].append(r)

    lines: list[str] = []
    lines.extend(
        [
            "# 选必三《逻辑与思维》思维部分与推理部分全题触发宝典（学生版）",
            "",
            "本版按哲学宝典的写法重做：不写概念提纲，而把每一道题放到可迁移的方法节点下。每条都按“材料触发点、设问、为什么能想到、答案落点”展开。",
            "",
            "目录",
            "",
            "- 第一部分：思维部分，完全按哲学题一样处理",
            "- 第二部分：思维选择题陷阱库",
            "- 第三部分：推理部分，按题型族群处理",
            "- 第四部分：按年份题目总索引",
            "",
            "# 第一部分：思维部分，完全按哲学题一样处理",
            "",
            "思维题的本质不是背定义，而是看材料动作：调查、预测、统筹、变化、扬弃、联想、发散、反转、抽象。材料动作一旦稳定，就能触发对应思维方法。",
            "",
        ]
    )

    major_order = ["科学思维", "辩证思维", "创新思维", "综合复合题"]
    idx = 1
    for major in major_order:
        nodes = [
            k
            for k in node_map
            if k[0] == major and any(e["kind"] == "A主观题" for e in node_map[k])
        ]
        if not nodes:
            continue
        lines.append(f"# {major}")
        lines.append("")
        for _, node in sorted(nodes, key=lambda x: x[1]):
            node_entries = [e for e in node_map[(major, node)] if e["kind"] == "A主观题"]
            if not node_entries:
                continue
            lines.append(f"## {node}")
            lines.append("")
            for e in node_entries:
                lines.extend(make_thinking_entry(e, idx))
                idx += 1

    lines.append("# 第二部分：思维选择题陷阱库")
    lines.append("")
    choice_entries = [e for e in entries if e["kind"] == "B选择题"]
    for e in choice_entries:
        lines.extend(make_thinking_entry(e, idx))
        idx += 1

    lines.append("# 第三部分：推理部分，按题型族群处理")
    lines.append("")
    lines.append("推理题先不看材料话题，先翻译逻辑结构。所有题都放在对应题型下面，学生复习时按题型刷。")
    lines.append("")
    ridx = 1
    family_order = [
        "三段论",
        "假言推理",
        "选言判断与选言推理",
        "联言判断与联言推理",
        "归纳推理与探求因果联系",
        "类比推理",
        "周延判断",
        "换质换位",
        "逻辑三律与概念规则",
    ]
    for fam in family_order:
        rows = family_map.get(fam, [])
        if not rows:
            continue
        lines.append(f"## {fam}")
        lines.append("")
        lines.append(FAMILY_INTROS.get(fam, "先识别题型，再套最小规则。"))
        lines.append("")
        for row in rows:
            lines.extend(make_reasoning_entry(row, ridx))
            ridx += 1

    lines.append("# 第四部分：按年份题目总索引")
    lines.append("")
    by_year = defaultdict(list)
    for e in entries:
        m = re.match(r"(202[456])", e["label"])
        by_year[m.group(1) if m else "其他"].append(e)
    for year in ["2024", "2025", "2026", "其他"]:
        if year not in by_year:
            continue
        lines.append(f"## {year}")
        lines.append("")
        for e in by_year[year]:
            lines.append(f"- {e['label']}：{e.get('触发思维') or '选择题信号'}")
        lines.append("")

    LEDGER_MD.write_text(
        "\n".join(
            [
                "# Coverage Counts",
                "",
                f"- 思维主观题原始条目：{sum(1 for e in entries if e['kind']=='A主观题')}",
                f"- 思维选择题原始条目：{sum(1 for e in entries if e['kind']=='B选择题')}",
                f"- 思维部分按节点挂载条目：{idx - 1}",
                f"- 推理题型矩阵与题型报告条目：{ridx - 1}",
                f"- 最终 Markdown：{OUT_MD}",
                f"- 最终 Word：{OUT_DOCX}",
            ]
        ),
        encoding="utf-8",
    )
    return "\n".join(lines).replace("“三新”", "“多向性、跨越性、独特性”")


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def set_font(run, size: int | None = None, bold: bool | None = None, color=None) -> None:
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    styles = doc.styles
    for name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        st = styles[name]
        st.font.name = "微软雅黑"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.size = Pt(22)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = BLUE
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = BLUE
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(48, 84, 150)
    styles["Heading 3"].font.size = Pt(11)
    styles["Heading 3"].font.bold = True
    styles["Heading 3"].font.color.rgb = DARK
    for sname in ["toc 1", "toc 2", "toc 3", "Hyperlink"]:
        if sname in styles:
            styles[sname].font.color.rgb = DARK
            styles[sname].font.name = "微软雅黑"
            styles[sname]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def build_docx(md: str) -> None:
    doc = Document()
    style_doc(doc)
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings.append(update)
    title = md.splitlines()[0].lstrip("# ").strip()
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_font(r, 22, True, BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("完全模仿哲学宝典条目法：材料触发点 / 设问 / 为什么能想到 / 答案落点")
    set_font(r, 11, False, MUTED)

    toc_title = doc.add_paragraph(style="Heading 1")
    toc_title.add_run("目录")
    p = doc.add_paragraph()
    add_field(p, r'TOC \o "1-3" \h \z \u')
    doc.add_section(WD_SECTION_START.NEW_PAGE)

    skip_first = True
    for raw in md.splitlines()[1:]:
        line = raw.strip()
        if skip_first and line in {"目录", "- 第一部分：思维部分，完全按哲学题一样处理", "- 第二部分：思维选择题陷阱库", "- 第三部分：推理部分，按题型族群处理", "- 第四部分：按年份题目总索引"}:
            continue
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(line[2:], style="Heading 1")
        elif line.startswith("## "):
            p = doc.add_paragraph(line[3:], style="Heading 2")
        elif line.startswith("### "):
            p = doc.add_paragraph(line[4:], style="Heading 3")
        elif line.startswith("- "):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Cm(0.55)
            p.add_run("• ")
            p.add_run(line[2:])
        else:
            p = doc.add_paragraph(style="Normal")
            if line.startswith("【材料触发点】"):
                label, body = "【材料触发点】", line[len("【材料触发点】") :].strip()
            elif line.startswith("【设问】"):
                label, body = "【设问】", line[len("【设问】") :].strip()
            elif line.startswith("【为什么能想到】"):
                label, body = "【为什么能想到】", line[len("【为什么能想到】") :].strip()
            elif line.startswith("【答案落点】"):
                label, body = "【答案落点】", line[len("【答案落点】") :].strip()
            else:
                label, body = "", line
            if label:
                r = p.add_run(label + " ")
                set_font(r, 10.5, True, BLUE)
            r = p.add_run(body)
            set_font(r, 10.5, False, DARK)
    for i, sec in enumerate(doc.sections):
        header_p = sec.header.paragraphs[0]
        header_p.text = "飞哥正志讲堂 · 选必三《逻辑与思维》"
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i == 0:
            footer_p = sec.footer.paragraphs[0]
            footer_p.text = ""
            footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_p.add_run("第 ")
            add_field(footer_p, "PAGE")
            footer_p.add_run(" 页")
    doc.save(OUT_DOCX)


def main() -> None:
    RUN.joinpath("04_delivery").mkdir(parents=True, exist_ok=True)
    RUN.joinpath("02_ledgers").mkdir(parents=True, exist_ok=True)
    md = build_markdown()
    OUT_MD.write_text(md, encoding="utf-8")
    build_docx(md)
    with zipfile.ZipFile(OUT_DOCX) as z:
        media_count = len([n for n in z.namelist() if n.startswith("word/media/")])
    print(OUT_MD)
    print(OUT_DOCX)
    print(f"media_count={media_count}")


if __name__ == "__main__":
    main()
