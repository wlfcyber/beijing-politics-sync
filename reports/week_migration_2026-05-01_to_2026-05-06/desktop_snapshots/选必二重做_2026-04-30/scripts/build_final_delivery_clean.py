#!/usr/bin/env python3
from __future__ import annotations

import csv
import html
import re
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


BASE = Path("/Users/wanglifei/Desktop/北京高考政治/选必二重做_2026-04-30")
FINAL = BASE / "final_legal_outputs"
DELIVERY = BASE / "delivery"
TEXT_CACHE = BASE / "preprocess_v2_2026-05-03" / "text_cache"
LEDGER = FINAL / "legal_question_chain_ledger_2026-05-03.csv"

COLLECTION_MD = DELIVERY / "选必二法律题处理合集_最终交付版_2026-05-04.md"
FRAMEWORK_MD = DELIVERY / "选必二法律与生活最终进化框架_最终交付版_2026-05-04.md"
COLLECTION_DOCX = DELIVERY / "选必二法律题处理合集_最终交付版_2026-05-04.docx"
FRAMEWORK_DOCX = DELIVERY / "选必二法律与生活最终进化框架_最终交付版_2026-05-04.docx"
COLLECTION_HTML = DELIVERY / "选必二法律题处理合集_最终交付版_2026-05-04.html"
FRAMEWORK_HTML = DELIVERY / "选必二法律与生活最终进化框架_最终交付版_2026-05-04.html"


FRONT_DOMAIN_ORDER = [
    "交易关系与消费者保护",
    "劳动用工与职业边界",
    "财产、相邻、继承与家庭",
    "人格权与侵权责任",
    "创新竞争与公共救济",
]

FRAMEWORK_DOMAIN_ORDER = [
    "交易关系：合同规则与消费者保护",
    "劳动就业与职业边界",
    "财产权、物权与相邻关系",
    "婚姻家庭与继承扶养",
    "人格权与侵权责任",
    "知识产权与竞争秩序",
    "纠纷解决与程序救济",
    "生态公益、新业态与法治价值",
]

FRONT_DOMAIN_METHOD = {
    "交易关系与消费者保护": "先识别买卖、服务、租赁、平台交易等生活交易关系；消费者保护不是单独口号，而是合同交易中信息不对称、经营者优势地位和格式条款滥用的纠偏规则。",
    "劳动用工与职业边界": "先识别劳动关系、用工管理或职业边界，再看从属性事实、劳动者义务、解除补偿、竞业限制和劳动权保护。",
    "财产、相邻、继承与家庭": "先识别财产归属、身份关系、扶养监护或继承安排，再落到效力、归属、相邻利益平衡和家庭伦理。",
    "人格权与侵权责任": "先识别人格利益和侵害事实，再看过错、损害、因果关系、免责或责任方式。",
    "创新竞争与公共救济": "先判断核心是创新成果、竞争秩序、程序救济还是公共利益保护；知识产权和不正当竞争保护创新与市场秩序，程序和公益诉讼把权利救济、公共利益保护落地。",
}

FRONT_DOMAIN_SHORT = {
    "交易关系与消费者保护": "交易保护域",
    "劳动用工与职业边界": "劳动用工域",
    "财产、相邻、继承与家庭": "财产家庭域",
    "人格权与侵权责任": "人格侵权域",
    "创新竞争与公共救济": "秩序救济域",
}

DOMAIN_CORE = {
    "交易关系：合同规则与消费者保护": "合同成立、合同效力、履行与违约、格式条款、消费者知情权、安全权、经营者责任和惩罚性赔偿。",
    "劳动就业与职业边界": "劳动关系认定、从属性事实、劳动合同、试用期、解除补偿、新就业形态权益保护。",
    "财产权、物权与相邻关系": "所有权、他物权、登记与交付、相邻关系、物权保护与财产归属。",
    "婚姻家庭与继承扶养": "婚姻家庭身份、监护扶养、遗嘱继承、法定继承、遗赠扶养协议和家庭美德。",
    "人格权与侵权责任": "姓名、肖像、名誉、隐私、个人信息、一般侵权责任和特殊侵权责任。",
    "知识产权与竞争秩序": "著作权、商标、专利、商业秘密、许可使用、不正当竞争和创新秩序。",
    "纠纷解决与程序救济": "调解、仲裁、诉讼、举证、管辖、二审、执行、公益诉讼和司法建议。",
    "生态公益、新业态与法治价值": "生态公益保护、平台经济、数字技术、新业态治理以及法治价值收束。",
}

FORBIDDEN_PATTERNS = [
    "Codex",
    "Claude",
    "GPT",
    "Governor",
    "Confucius",
    "ClaudeCode",
    "official_answer_key",
    "formal_or_scoring_source",
    "objective_answer_locked",
    "formal_scoring_matched",
    "rubric_match_pending",
    "answer_pending_official",
    "参考答案",
    "答案及评分参考",
    "评分参考",
    "评标",
    "pipeline",
    "debug",
    "gpt55",
    "claudecode",
    "OCR",
    "ocr",
    "扫描全能王",
    "source_pdf",
    "vision",
    "[slide",
    "/Users/",
]


FRONT_DOMAIN_RENAME = {
    "纠纷解决生态公益与新业态": "创新竞争与公共救济",
}

FRAMEWORK_DOMAIN_RENAME = {
    "合同交易与消费者保护": "交易关系：合同规则与消费者保护",
}

LEGACY_FRONT_RENAME = {
    "物权相邻继承家庭": "财产、相邻、继承与家庭",
    "人格权侵权": "人格权与侵权责任",
    "知识产权不正当竞争": "创新竞争与公共救济",
    "救济与公共利益": "创新竞争与公共救济",
}

LABOR_FRAMEWORK_DOMAIN = "劳动就业与职业边界"


def normalize_record_domains(row: dict) -> None:
    if row.get("framework_domain") in FRAMEWORK_DOMAIN_RENAME:
        row["framework_domain"] = FRAMEWORK_DOMAIN_RENAME[row["framework_domain"]]
    front = row.get("front_domain", "")
    if front == "合同消费者劳动":
        if row.get("framework_domain") == LABOR_FRAMEWORK_DOMAIN:
            row["front_domain"] = "劳动用工与职业边界"
        else:
            row["front_domain"] = "交易关系与消费者保护"
    elif front in FRONT_DOMAIN_RENAME:
        row["front_domain"] = FRONT_DOMAIN_RENAME[front]
    elif front in LEGACY_FRONT_RENAME:
        row["front_domain"] = LEGACY_FRONT_RENAME[front]


DOMAIN_OVERRIDES = {
    "2024_海淀_二模_010_Q11": {
        "front_domain": "物权相邻继承家庭",
        "framework_domain": "财产权、物权与相邻关系",
        "front_domain_reason": "共享单车占用小区出入口的主争点是建筑区划共有部分、通行利益和排除妨碍，不是人格权侵权。",
        "framework_domain_reason": "先判断业主对共有部分的共有和共同管理权，再落到排除妨碍、恢复正常通行。",
    },
    "2024_石景山_一模_011_Q11": {
        "front_domain": "合同消费者劳动",
        "framework_domain": "合同交易与消费者保护",
        "front_domain_reason": "牙膏虚假宣传题的主争点是消费者权益、经营者欺诈和惩罚性赔偿，不是人格权侵权。",
        "framework_domain_reason": "先看经营者宣传行为是否构成欺诈，再看消费者三倍赔偿与医疗损失因果证明。",
    },
    "2024_海淀_一模_005_Q11": {
        "front_domain": "人格权侵权",
        "framework_domain": "人格权与侵权责任",
        "front_domain_reason": "球场运动伤害题的主争点是一般侵权责任与公平原则，调解、诉讼只是题肢中的程序干扰。",
        "framework_domain_reason": "先看侵害事实、过错或公平分担，再排除程序节点错项。",
    },
    "2024_西城_一模_012_Q18": {
        "front_domain": "物权相邻继承家庭",
        "framework_domain": "婚姻家庭与继承扶养",
        "front_domain_reason": "抽取的小问主争点是成年子女赡养义务、精神慰藉与法治德治统一，不是单纯程序题。",
        "framework_domain_reason": "先落家庭扶养义务，再说明法院文书对权利保护和家庭美德的引导。",
    },
    "2024_顺义_二模_014_Q11": {
        "front_domain": "人格权侵权",
        "framework_domain": "人格权与侵权责任",
        "front_domain_reason": "匿名评价和新闻报道题的主争点是名誉侵权边界与公共利益抗辩，程序内容只是错项干扰。",
        "framework_domain_reason": "先判断是否构成人格权侵权，再处理举证责任、上诉等程序错项。",
    },
    "2025_东城_二模_025_Q10": {
        "front_domain": "人格权侵权",
        "framework_domain": "人格权与侵权责任",
        "front_domain_reason": "景区游客受伤题的主争点是安全保障义务与侵权责任，调解只是材料场景。",
        "framework_domain_reason": "先看注意义务、安全保障义务和损害后果，再排除举证倒置等错项。",
    },
    "2025_海淀_期末_035_Q14": {
        "front_domain": "知识产权不正当竞争",
        "framework_domain": "知识产权与竞争秩序",
        "front_domain_reason": "擅自使用画作图案的核心是著作权保护，法院调解只是题肢中的程序信息。",
        "framework_domain_reason": "先识别作品、权利保护与使用边界，再处理举证和调解题肢。",
    },
    "2025_海淀_期末_035_Q21": {
        "front_domain": "合同消费者劳动",
        "framework_domain": "劳动就业与职业边界",
        "front_domain_reason": "竞业限制题的主争点是劳动者择业自由、企业商业秘密保护和用人单位与劳动者的利益平衡。",
        "framework_domain_reason": "先看劳动者是否负有保密义务，再看竞业限制协议的效力与边界。",
        "rubric_override": "规范竞业限制，有利于保护企业商业秘密、保护知识产权、推动社会创新；对竞业限制加以限制，有利于保护劳动者基本劳动权利，促进人才资源合理配置。案例一中，李某作为负有保密义务的劳动者，违反竞业限制或保密约定，应承担违约责任。案例二中，刘某作为普通员工，并非负有保密义务人员，公司与其签订的竞业限制协议违反规定，应认定无效。总的看，法律规定竞业限制并设置边界，是为了平衡劳动者合法权益与企业商业秘密保护。",
    },
    "2025_西城_二模_030_Q18": {
        "front_domain": "合同消费者劳动",
        "framework_domain": "合同交易与消费者保护",
        "front_domain_reason": "未成年人直播打赏题的主争点是民事法律行为效力、平台责任和监护责任，新业态只是场景。",
        "framework_domain_reason": "先判断限制民事行为能力人的大额交易效力，再综合各方过错落责任。",
    },
    "2026_东城_一模_037_Q18": {
        "front_domain": "知识产权不正当竞争",
        "framework_domain": "知识产权与竞争秩序",
        "front_domain_reason": "题目要求说明司法如何保护创新，核心仍是专利、植物新品种权和恶意诉讼中的知识产权保护。",
        "framework_domain_reason": "先抓知识产权权利保护和竞争秩序，再把调解、惩罚性赔偿、驳回恶意诉讼作为救济路径。",
        "rubric_override": "案例一：通过诉讼调解或和解，实质化解专利权归属纠纷，减少诉讼对创新创业的干扰，提高效率、节约司法资源。案例二：对故意知识产权侵权适用惩罚性赔偿，提高侵权成本，保护权利人合法权益，维护公平竞争和市场秩序。案例三：驳回并谴责恶意诉讼、滥用诉讼权利行为，明确权利行使边界，维护诉讼秩序和诚信原则。总说：人民法院坚持以事实为根据、以法律为准绳，发挥司法裁判导向作用，保护知识产权就是保护创新，有利于激发创新活力。",
    },
    "2026_朝阳_一模_041_Q18": {
        "front_domain": "知识产权不正当竞争",
        "framework_domain": "知识产权与竞争秩序",
        "front_domain_reason": "题眼是人民法院依法保护知识产权以护航新质生产力，程序措施服务于知识产权保护。",
        "framework_domain_reason": "先归知识产权，再按调、惩、辨三条司法措施落到创新秩序。",
    },
    "2026_房山_一模_040_Q11": {
        "front_domain": "救济与公共利益",
        "framework_domain": "纠纷解决与程序救济",
        "front_domain_reason": "个人债务集中清偿题的主争点是执行、信用修复和债权债务利益平衡，不是一般合同违约责任。",
        "framework_domain_reason": "先看法院执行与信用修复机制，再判断其如何平衡债权人和债务人利益。",
    },
    "2026_通州_期末_054_Q13": {
        "front_domain": "知识产权不正当竞争",
        "framework_domain": "知识产权与竞争秩序",
        "front_domain_reason": "未成年人摄影作品被图书封面使用的核心是著作权保护与侵权救济，不是一般侵权题。",
        "framework_domain_reason": "先识别作品与著作权归属，再判断使用、署名、报酬、举证和救济路径。",
    },
}


EXCLUDED_QUESTION_IDS = {
    "2025_石景山_一模_021_Q8": "入党誓词题属于党史/政治与法治，不是《法律与生活》法律关系题。",
    "2025_海淀_二模_029_Q7": "票根经济题属于消费场景与经济联动，不以法律关系、权利义务或责任效力为设问核心。",
}


EXCLUDED_QUESTION_LABELS = {
    "2025_石景山_一模_021_Q8": "2025 石景山一模第8题",
    "2025_海淀_二模_029_Q7": "2025 海淀二模第7题",
}


SUBJECTIVE_CORE_BY_FRONT = {
    "交易关系与消费者保护": "主观题敢考且常考：合同成立与效力、格式条款、经营者责任、消费者维权、未成年人交易效力。",
    "劳动用工与职业边界": "主观题敢考且常考：劳动关系认定、新就业形态从属性、劳动合同解除补偿、竞业限制边界、劳动权与商业秘密平衡。",
    "财产、相邻、继承与家庭": "主观题敢考：相邻关系与共有管理、绿色物权、遗赠扶养协议、赡养与精神慰藉、继承安排效力。",
    "人格权与侵权责任": "主观题敢考：一般侵权责任、安全保障义务合理限度、见义勇为补偿规则、人格权益保护、未成年人权益、个人信息权益和网络侵权边界。",
    "创新竞争与公共救济": "主观题敢考：著作权、作品独创性、AI作品和数字人场景中的权利保护、商标/商业标识、不正当竞争、商业秘密、调解诉讼、举证责任、公益诉讼和生态修复。",
}


CHOICE_POOL_BY_FRONT = {
    "交易关系与消费者保护": "只在选择题里重点识别：要约承诺细节、促销/价格规则、先用后付、生鲜灯、直播带货、民事行为能力具体年龄分段等。它们可做题肢排错，不直接推成主观题大模板。",
    "劳动用工与职业边界": "只在选择题里重点识别：录用通知、女职工保护、一函两书、劳动争议调解机构、劳动者义务细节等。",
    "财产、相邻、继承与家庭": "只在选择题里重点识别：质权、抵押、担保物权、登记交付、一房二卖、土地经营权流转、宅基地、法定继承顺位和丧失继承权等细节。",
    "人格权与侵权责任": "只在选择题里重点识别：运动伤害公平分担、公共利益新闻报道抗辩、动物饲养、高空抛物、产品责任、个人信息单独同意、举证倒置具体情形。公共场所安全保障义务和见义勇为补偿规则不降级，主观题中要会写。",
    "创新竞争与公共救济": "只在选择题里重点识别：著作权登记误区、发表权保护期、法定许可、专利许可、数字人模型、数据权益边界、诉讼代理人与辩护人、二审终审、司法确认、行政裁决、诉讼时效等。",
}


DUAL_STATE_BY_FRONT = {
    "交易关系与消费者保护": "有时作为主观题一句采分点：限制民事行为能力人的交易效力、格式条款提示说明义务。",
    "劳动用工与职业边界": "有时作为主观题一句采分点：劳动权与商业秘密保护的平衡、职业边界和诚信义务。",
    "财产、相邻、继承与家庭": "有时作为主观题一句采分点：相邻关系、共有部分管理、排除妨碍、继承安排效力。",
    "人格权与侵权责任": "有时作为主观题一句采分点：举证责任分配、安全保障义务合理限度、个人信息作为人格权益及处理者义务。",
    "创新竞争与公共救济": "有时作为主观题一句采分点：AI作品独创性、数据竞争利益、未经许可使用作品或商业标识、司法确认强制执行力、谁主张谁举证、公益诉讼与生态修复责任。",
}


DOUBLE_TAG_NOTES = [
    "竞业限制题：前台先按劳动就业处理，若材料同时涉及商业秘密或知识产权保护，再把创新秩序作为价值收束。",
    "AI、数字人、直播、平台题：先看核心争议是作品保护、消费者权益、人格权益还是平台责任，不靠技术词直接归域。",
    "调解、诉讼、司法确认出现在合同、侵权、家庭或知识产权题中时，程序通常是救济路径；只有设问主问程序效果时，才进创新竞争与公共救济域。",
    "生态公益、新业态材料出现时，先找实体法律关系；公共利益多数是意义收束，不自动单独成域。",
]

NOISE_HINTS = [
    "政治与法治",
    "经济与社会",
    "哲学",
    "文化",
    "逻辑与思维",
    "当代国际政治与经济",
    "中国共产党",
    "国际关系",
    "新质生产力",
]

LEGAL_KEYWORDS = [
    "民法典", "民事", "合同", "要约", "承诺", "违约", "不可抗力", "物业", "消费者", "经营者", "仅退款",
    "劳动", "劳动关系", "劳动合同", "劳动仲裁", "用人单位", "经济补偿", "竞业限制", "工伤",
    "物权", "所有权", "不动产", "相邻关系", "通行", "排除妨碍", "财产权", "用益物权",
    "继承", "遗嘱", "遗产", "遗赠扶养", "赡养", "扶养", "监护", "夫妻",
    "人格权", "名誉权", "肖像权", "隐私", "个人信息", "侵权", "过错", "无过错", "赔偿", "损害",
    "著作权", "商标", "专利", "商业秘密", "不正当竞争", "虚拟数字人", "AI", "人工智能",
    "调解", "仲裁", "诉讼", "举证", "司法确认", "公益诉讼", "二审", "上诉", "执行", "司法建议",
    "生态环境", "碳排放", "平台", "新业态",
]


def clean_text(value: str | None, limit: int | None = None) -> str:
    text = (value or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    text = re.sub(r"/Users/[^\n，。；;）)】]+", "", text)
    replacements = {
        "`": "",
        "参考答案": "答案要点",
        "答案及评分参考": "答案与给分要点",
        "评分参考": "给分要点",
        "评分细则": "给分依据",
        "阅卷细则": "给分依据",
        "评标": "给分依据",
        "合同交易与消费者保护": "交易关系：合同规则与消费者保护",
        "合同消费者劳动": "交易关系与消费者保护/劳动用工",
        "rubric": "给分依据",
        "formal": "正式",
        "official": "正式",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"(?m)^.*答案与给分要点.*$", "", text)
    text = re.sub(r"(?m)^\s*\[slide\s+\d+\]\s*$", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\[ocr page \d+\]", "", text, flags=re.IGNORECASE)
    text = re.sub(r"(?m)^.*扫描全能王.*$", "", text)
    text = re.sub(r"(?m)^source_pdf:.*$", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    if limit and len(text) > limit:
        return text[:limit].rstrip() + "……"
    return text


def parse_cache_file(path: Path) -> tuple[str, str]:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    lines = raw.splitlines()
    source = ""
    if lines and lines[0].startswith("source: "):
        source = lines[0].replace("source: ", "", 1).strip()
    text = raw.split("\n\n", 1)[1] if "\n\n" in raw else raw
    return source, text


def load_source_cache() -> dict[str, str]:
    cache: dict[str, str] = {}
    for path in sorted(TEXT_CACHE.glob("*.txt")):
        source, text = parse_cache_file(path)
        if source:
            cache[source] = text
    return cache


SOURCE_TEXT_CACHE = load_source_cache()


def load_records() -> list[dict]:
    with LEDGER.open(encoding="utf-8-sig") as f:
        rows = list(csv.DictReader(f))
    for row in rows:
        qid = row.get("question_id", "")
        if qid in DOMAIN_OVERRIDES:
            row.update(DOMAIN_OVERRIDES[qid])
        normalize_record_domains(row)
    records = [r for r in rows if r.get("include_status", "").startswith("included")]
    # This suite item is a cross-module family/culture prompt without a locked
    # legal scoring block in the local source; keeping it would force a wrong
    # rubric match into the legal framework.
    records = [
        r for r in records
        if r.get("label") != "2025 顺义 一模 2025顺义一模 第21题"
    ]
    records = [
        r for r in records
        if r.get("question_id") not in EXCLUDED_QUESTION_IDS
    ]
    records.sort(key=lambda r: (r["year"], r["district"], r["stage"], int(r["question_no"])))
    return records


def source_rubric_text(record: dict) -> str:
    source = record.get("rubric_source", "")
    return SOURCE_TEXT_CACHE.get(source, record.get("rubric_text", ""))


def compact_for_split(text: str) -> str:
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def segment_source_text(text: str, qno: str) -> list[str]:
    text = compact_for_split(text)
    if not text.strip():
        return []
    # First split obvious PPT slide and question-answer boundaries. Then keep
    # medium-sized windows so a rubric line and its follow-up scoring notes stay together.
    boundaries = [0, len(text)]
    patterns = [
        r"\[slide\s+\d+\]",
        r"(?m)^\s*(?:1[6-9]|2[0-1])\s*[.．、]\s*(?:（?\d+\s*分）?)?",
        r"(?m)^\s*(?:1[6-9]|2[0-1])\s*题",
        r"(?m)^\s*第\s*(?:1[6-9]|2[0-1])\s*题",
        r"(?m)^\s*[（(]\d+[）)]\s*(?:（?\d+\s*分）?)?",
        r"(?m)^\s*(?:答案|评分细则|细则|阅卷细则)[:：]",
    ]
    for pattern in patterns:
        for match in re.finditer(pattern, text):
            boundaries.append(match.start())
    boundaries = sorted(set(boundaries))
    rough = []
    for start, end in zip(boundaries, boundaries[1:]):
        seg = text[start:end].strip()
        if len(seg) >= 20:
            rough.append(seg)
    if not rough:
        rough = [text]
    windows = []
    for i, seg in enumerate(rough):
        candidates = [seg]
        if i + 1 < len(rough):
            candidates.append(seg + "\n" + rough[i + 1])
        if i + 2 < len(rough) and len(seg) < 500:
            candidates.append(seg + "\n" + rough[i + 1] + "\n" + rough[i + 2])
        windows.extend(candidates)
    # Include direct qno windows in case the source is a continuous DOCX answer sheet.
    for pattern in [
        rf"(?<!\d){re.escape(qno)}\s*[.．、]\s*",
        rf"(?<!\d){re.escape(qno)}\s*题",
        rf"第\s*{re.escape(qno)}\s*题",
    ]:
        for match in re.finditer(pattern, text):
            start = max(0, match.start() - 250)
            end = min(len(text), match.start() + 1800)
            windows.append(text[start:end].strip())
    # De-duplicate while preserving order.
    seen = set()
    unique = []
    for seg in windows:
        seg = clean_text(seg, 2200)
        key = seg[:180]
        if key and key not in seen:
            seen.add(key)
            unique.append(seg)
    return unique


def score_segment(record: dict, segment: str) -> float:
    qno = record.get("question_no", "")
    prompt = record.get("complete_prompt", "")
    qtext = record.get("question_text", "")
    triggers = [x for x in re.split(r"[、,，\s]+", record.get("material_triggers", "")) if len(x) >= 2]
    parties = re.findall(r"[甲乙丙丁]公司|[A-Z]公司|[\u4e00-\u9fa5]{0,3}某[\u4e00-\u9fa5]{0,3}|小[\u4e00-\u9fa5]", qtext + prompt)
    keywords = list(dict.fromkeys(triggers + parties + [k for k in LEGAL_KEYWORDS if k in (qtext + prompt)]))
    score = 0.0
    if re.search(rf"(?<!\d){re.escape(qno)}\s*[.．、题]", segment):
        score += 35
    for sub in re.findall(r"[（(]\s*\d+\s*[）)]", prompt):
        if sub in segment or sub.replace("(", "（").replace(")", "）") in segment:
            score += 18
    if "评分" in segment or "细则" in segment or "分）" in segment or "分)" in segment:
        score += 10
    if "答案" in segment or "应当" in segment or "属于" in segment or "有利于" in segment:
        score += 8
    for kw in keywords:
        if kw and kw in segment:
            score += 9 if kw in parties or kw in triggers else 5
    for kw in LEGAL_KEYWORDS:
        if kw in segment:
            score += 1.5
    noise_hits = sum(1 for hint in NOISE_HINTS if hint in segment)
    legal_hits = sum(1 for kw in LEGAL_KEYWORDS if kw in segment)
    if noise_hits and legal_hits < 4:
        score -= 22 * noise_hits
    if len(segment) > 1800:
        score -= 8
    return score


def strong_anchors(record: dict) -> list[str]:
    prompt = record.get("complete_prompt", "")
    qtext = record.get("question_text", "")
    triggers = [x for x in re.split(r"[、,，\s]+", record.get("material_triggers", "")) if len(x) >= 2]
    parties = re.findall(r"[甲乙丙丁]公司|[A-Z]\s*公司|[A-Z]\s*集团|[\u4e00-\u9fa5]{0,3}某[\u4e00-\u9fa5]{0,3}|小[\u4e00-\u9fa5]", qtext + prompt)
    legal = [k for k in LEGAL_KEYWORDS if k in (qtext + prompt)]
    return list(dict.fromkeys(parties + triggers + legal))


def crop_relevant_segment(record: dict, segment: str) -> str:
    if not segment:
        return segment
    anchors = strong_anchors(record)
    generic = {"合同", "承诺", "法院", "人民法院", "判决", "裁判", "赔偿", "调解", "诉讼", "侵权", "民法典", "法律", "法治"}
    party_like = [
        a for a in anchors
        if re.search(r"公司|集团|某|小[\u4e00-\u9fa5]", a) and a not in generic
    ]
    search_anchors = party_like or [a for a in anchors if a not in generic] or anchors
    party_positions = [segment.find(anchor) for anchor in party_like if anchor and segment.find(anchor) >= 0]
    positions = [segment.find(anchor) for anchor in search_anchors if anchor and segment.find(anchor) >= 0]
    qno = record.get("question_no", "")
    qno_positions: list[int] = []
    for pattern in [
        rf"(?<!\d){re.escape(qno)}\s*[.．、]\s*",
        rf"(?<!\d){re.escape(qno)}\s*[（(]\s*\d+\s*[）)]",
        rf"(?<!\d){re.escape(qno)}\s*题",
        rf"第\s*{re.escape(qno)}\s*题",
    ]:
        match = re.search(pattern, segment)
        if match:
            qno_positions.append(match.start())
    sub_positions: list[int] = []
    for sub in re.findall(r"[（(]\s*\d+\s*[）)]", record.get("complete_prompt", "")):
        pos = segment.find(sub)
        if pos >= 0:
            sub_positions.append(pos)
        alt = sub.replace("(", "（").replace(")", "）")
        pos = segment.find(alt)
        if pos >= 0:
            sub_positions.append(pos)
    prefer_qno = False
    if qno_positions:
        wanted_sub = re.findall(r"[（(]\s*(\d+)\s*[）)]", record.get("complete_prompt", ""))
        if wanted_sub:
            sub_no = wanted_sub[-1]
            close_qno = []
            for pos in qno_positions:
                tail = segment[pos: pos + 80]
                if re.search(rf"[（(]\s*{sub_no}\s*[）)]", tail):
                    close_qno.append(pos)
            if close_qno:
                positions = close_qno
                prefer_qno = True
            elif sub_positions and min(sub_positions) < min(qno_positions):
                positions = sub_positions + party_positions
            else:
                positions = qno_positions
                prefer_qno = True
        else:
            positions = qno_positions
            prefer_qno = True
    elif sub_positions:
        positions = sub_positions + party_positions
    elif party_positions:
        positions = party_positions
    if not positions:
        return segment
    anchor = min(positions)
    start = anchor if prefer_qno else max(0, anchor - 220)
    # Snap to a nearby slide/question/scoring boundary before the anchor.
    prefix = segment[: anchor + 1]
    boundary_patterns = [
        r"\[slide\s+\d+\]",
        r"(?m)^\s*(?:1[6-9]|2[0-1])\s*[.．、]\s*",
        r"(?m)^\s*(?:1[6-9]|2[0-1])\s*题",
        r"(?m)^\s*[（(]\s*\d+\s*[）)]\s*(?:[（(]\s*\d+\s*分\s*[）)])?",
        r"(?m)^\s*[（(]?\s*\d+\s*分\s*[）)]?",
        r"(?m)^\s*(?:答案|答案要点|评分细则|细则|阅卷细则)[:：]?",
    ]
    if not prefer_qno:
        for pattern in boundary_patterns:
            matches = list(re.finditer(pattern, prefix))
            if matches:
                start = max(start, matches[-1].start())
    cropped = segment[start:]
    # If this is an explicit subquestion, stop before the next subquestion.
    sub_nums = re.findall(r"[（(]\s*(\d+)\s*[）)]", record.get("complete_prompt", ""))
    if sub_nums:
        current = int(sub_nums[-1])
        next_num = current + 1
        next_match = re.search(rf"\n\s*(?:{re.escape(qno)}\s*)?[（(]\s*{next_num}\s*[）)]", cropped)
        if next_match and next_match.start() > 160:
            cropped = cropped[: next_match.start()]
    # Stop before the next full subjective question when it appears after the selected answer.
    try:
        next_qno = str(int(qno) + 1)
    except ValueError:
        next_qno = ""
    if next_qno:
        next_match = re.search(rf"\n\s*(?:【[^】]{{0,10}}】\s*)?{re.escape(next_qno)}\s*[.．、题(（]", cropped)
        if next_match and next_match.start() > 260:
            cropped = cropped[: next_match.start()]
    for marker in ["反馈试题", "下一题", "学生答题情况", "复习建议", "典型问题", "教学建议", "政治学科的法律模块"]:
        pos = cropped.find(marker)
        if pos > 220:
            cropped = cropped[:pos]
            break
    for marker in ["等级水平", "水平 4", "水平4"]:
        pos = cropped.find(marker)
        if pos > 500:
            cropped = cropped[:pos]
            break
    # PPT: if the selected slide already contains a compact full scoring answer, cut at the next slide.
    slide_match = re.match(r"\[slide\s+\d+\]", cropped)
    if slide_match:
        next_slide = re.search(r"\n\s*\[slide\s+\d+\]", cropped[slide_match.end():])
        first_slide = cropped[: slide_match.end() + next_slide.start()] if next_slide else cropped
        if next_slide and score_segment(record, first_slide) >= 45 and len(first_slide) >= 120:
            cropped = first_slide
    return cropped


def refined_rubric(record: dict) -> str:
    if record.get("rubric_override"):
        return clean_text(record.get("rubric_override"), 1400)
    if record["question_type"] == "choice":
        return ""
    full_text = source_rubric_text(record)
    segments = segment_source_text(full_text or record.get("rubric_text", ""), record.get("question_no", ""))
    if not segments:
        return clean_text(record.get("answer_points") or record.get("rubric_text"), 1200)
    ranked = sorted(((score_segment(record, seg), seg) for seg in segments), key=lambda x: x[0], reverse=True)
    best_score, best = ranked[0]
    # If the best segment is still weak, prefer the already narrowed ledger text but mark no uncertainty in the artifact.
    if best_score < 22:
        return clean_text(record.get("answer_points") or record.get("rubric_text"), 1200)
    refined = clean_text(crop_relevant_segment(record, best), 1400)
    try:
        next_qno = str(int(record.get("question_no", "")) + 1)
    except ValueError:
        next_qno = ""
    if next_qno:
        match = re.search(rf"(?:\n|\s【[^】]{{0,10}}】\s*){re.escape(next_qno)}\s*[.．、题(（]", refined)
        if match and match.start() > 100:
            refined = refined[: match.start()].rstrip()
    for marker in ["反馈试题", "下一题", "学生答题情况", "复习建议", "典型问题", "教学建议", "政治学科的法律模块"]:
        pos = refined.find(marker)
        if pos > 220:
            refined = refined[:pos].rstrip()
            break
    return refined


def qtype_name(record: dict) -> str:
    return "主观题" if record["question_type"] == "subjective" else "选择题"


def review_note(record: dict) -> str:
    if record["include_status"] == "included_needs_review":
        return "复合题：已抽取其中法律小问入框，频次统计按法律答题动作归位。"
    return "标准入框题。"


def answer_block(record: dict) -> list[str]:
    if record["question_type"] == "choice":
        lines = [
            f"答案：{clean_text(record['answer'])}。",
            "答案落点：" + clean_text(record["answer_points"]),
        ]
        wrong = clean_text(record.get("wrong_option_sentences"), 900)
        if wrong:
            lines.append("错项处理：" + wrong)
        return lines
    rubric = refined_rubric(record)
    points = clean_text(record.get("answer_points"), 1200)
    # The ledger's heuristic answer_points can be polluted when a source is a whole PPT/DOCX.
    # Use the refined source segment as the authoritative classroom-facing scoring excerpt.
    if rubric:
        return ["给分要点：" + rubric]
    return ["给分要点：" + points]


def chain_text(record: dict) -> str:
    trigger = clean_text(record["material_triggers"])
    if record.get("question_id") in DOMAIN_OVERRIDES:
        front_reason = clean_text(record.get("front_domain_reason"), 700)
        framework_reason = clean_text(record.get("framework_domain_reason"), 700)
        return (
            f"材料触发为“{trigger}”。本题先确定{record['framework_domain']}下的主体关系，"
            f"再把材料事实转成规则条件，最后按设问落到答案结论。{front_reason} "
            f"{framework_reason} 这类题不能只凭法院、调解、诉讼、平台或新业态等表层词归位，"
            "必须按命题人的主争点来定框架位置。"
        )
    why = clean_text(record["why_think"], 800)
    path = clean_text(record["proposition_path"], 800)
    return (
        f"材料触发为“{trigger}”。先据此确定{record['framework_domain']}下的主体关系，"
        f"再把材料事实转成规则条件，最后按设问落到答案结论。{why} {path}"
    )


def answer_summary(record: dict, limit: int = 700) -> str:
    if record["question_type"] == "choice":
        return clean_text(record["answer_points"], limit)
    return refined_rubric(record) or clean_text(record["answer_points"], limit)


def write_collection(records: list[dict]) -> None:
    counts = Counter(r["question_type"] for r in records)
    front_counts = Counter(r["front_domain"] for r in records)
    front_subj_counts = Counter(r["front_domain"] for r in records if r["question_type"] == "subjective")
    front_choice_counts = Counter(r["front_domain"] for r in records if r["question_type"] == "choice")
    domain_counts = Counter(r["framework_domain"] for r in records)
    lines: list[str] = [
        "# 选必二《法律与生活》法律题处理合集",
        "",
        f"本合集收录已入框法律题 {len(records)} 道，其中主观题 {counts['subjective']} 道、选择题 {counts['choice']} 道。主观题均配有给分要点，选择题配有答案落点或题肢判断说明；每道题均放入“设问、材料触发、答案落点、命题逻辑、完整解释链”的同一结构中。",
        "",
        "## 总览",
        "",
        f"- 主观题：{counts['subjective']} 道",
        f"- 选择题：{counts['choice']} 道",
        "- 主观题处理链：设问 -> 材料触发 -> 给分要点 -> 命题逻辑 -> 完整解释链。",
        "- 选择题处理链：题干事实 -> 题肢关键词 -> 正确答案 -> 错项排除 -> 命题逻辑。",
        "",
        "### 前台五域频次",
        "",
    ]
    for domain in FRONT_DOMAIN_ORDER:
        lines.append(f"- {domain}：{front_counts[domain]} 道（主观 {front_subj_counts[domain]}，选择 {front_choice_counts[domain]}）")
    lines.extend(["", "### 后台八域频次", ""])
    for domain in FRAMEWORK_DOMAIN_ORDER:
        lines.append(f"- {domain}：{domain_counts[domain]} 道")
    lines.append("")
    for idx, record in enumerate(records, 1):
        lines.extend(
            [
                f"## {idx}. {clean_text(record['label'])}",
                "",
                f"- 题型：{qtype_name(record)}",
                f"- 前台五域：{record['front_domain']}",
                f"- 后台归位：{record['framework_domain']}",
                f"- 入框说明：{review_note(record)}",
                "",
                "### 设问",
                "",
                clean_text(record["complete_prompt"], 1400),
                "",
                "### 材料触发",
                "",
                clean_text(record["material_triggers"]),
                "",
                "### 答案落点",
                "",
            ]
        )
        for line in answer_block(record):
            lines.append(clean_text(line))
            lines.append("")
        lines.extend(
            [
                "### 命题逻辑",
                "",
                clean_text(record["proposition_path"], 1000),
                "",
                "### 完整解释链",
                "",
                chain_text(record),
                "",
            ]
        )
    COLLECTION_MD.write_text("\n".join(lines), encoding="utf-8")


def write_framework(records: list[dict]) -> None:
    front_counts = Counter(r["front_domain"] for r in records)
    front_subj_counts = Counter(r["front_domain"] for r in records if r["question_type"] == "subjective")
    front_choice_counts = Counter(r["front_domain"] for r in records if r["question_type"] == "choice")
    domain_counts = Counter(r["framework_domain"] for r in records)
    domain_subj_counts = Counter(r["framework_domain"] for r in records if r["question_type"] == "subjective")
    domain_choice_counts = Counter(r["framework_domain"] for r in records if r["question_type"] == "choice")
    by_domain: dict[str, list[dict]] = defaultdict(list)
    for record in records:
        by_domain[record["framework_domain"]].append(record)
    lines: list[str] = [
        "# 选必二《法律与生活》最终进化框架",
        "",
        "本框架的原则是：主干必须来自最高频、最可迁移的答题动作；覆盖必须足够宽，能让新题无论落在合同、劳动、物权、家庭、人格、知识产权、程序还是新业态，都能自然归位。",
        "",
        "## 一核",
        "",
        "以事实为根据、以法律为准绳，在具体法律关系中适用规则，解决“谁有什么权利义务、谁承担什么责任或产生什么法律效果”的问题，最终实现定分止争。",
        "",
        "## 二线",
        "",
        "1. 权利保护与权利边界：先确认合法权益，再判断权利行使有没有事实条件、诚信边界、程序边界和他人权益边界。",
        "2. 法治规则与德治价值：这条线在答题中要落实为“规则适用与价值校准”。先用规则处理本案，再把处理结果上升到公平、诚信、友善、家庭伦理、劳动保护、创新秩序等价值；价值必须后置，不能替代本案法律分析。",
        "",
        "## 三问",
        "",
        "1. 判什么/怎么处理：本案要确认关系、效力、责任、归属、救济路径还是价值意义。",
        "2. 凭什么：材料中的事实怎样对应教材规则、法律条件、制度功能或程序要求。",
        "3. 有什么意义：本案处理怎样保护权利、划清边界、维护秩序、弘扬价值。",
        "",
        "## 四步",
        "",
        "定主体关系 -> 找争点事实 -> 对规则条件 -> 落答案结果。",
        "",
        "- 定主体关系：谁和谁之间是什么法律关系，先把题放进正确领域。",
        "- 找争点事实：材料中哪些事实决定权利义务、效力、责任、归属或程序。",
        "- 对规则条件：把事实对应教材规则、制度要求、程序要求，而不是空背法律名词。",
        "- 落答案结果：按题型落到责任、效力、归属、救济路径、裁判理由或价值意义。",
        "",
        "## 五域",
        "",
        "五域是学生前台的第一层索引，划分依据不是法考目录，而是“学生读题时最容易识别的生活关系 + 高频题量 + 答题动作差异”。",
        "",
        "本框架的主干主要服务主观题答题。选择题中的细碎制度、程序节点、概念边界和错项陷阱，主要用于选择题排错，不得反推为主观题大模板。判断一个节点能不能成为“写答案要展开的部分”，看它是否在主观题中反复要求学生写出“关系、规则、结论、意义”。",
        "",
        "如果某个细节在主观题中只作为一句采分点出现，可以写一句，但不展开成整段。也就是说，学生要分清三层：主观题要展开的部分；只在选择题里识别错项的部分；偶尔在主观题中补一句的过渡点。",
        "",
        "第五域的边界要特别守住：看到法院、调解、诉讼、平台、生态、新业态，不自动归第五域。若设问核心仍是侵权、合同、家庭、知识产权等实体法律关系，就先归实体域；程序只是救济路径，公共利益只是价值收束。",
        "",
        "### 主观题占比提醒",
        "",
        "备考优先级看主观题占比，不只看总题量。主观题要练完整解释链；选择题高频的细节主要练排错，不必全部背成主观题段落。",
        "",
        "### 双标签题读法",
        "",
    ]
    for note in DOUBLE_TAG_NOTES:
        lines.append(f"- {note}")
    lines.append("")
    lines.extend(
        [
            "### 五域明细",
            "",
        ]
    )
    for domain in FRONT_DOMAIN_ORDER:
        subjective_ratio = front_subj_counts[domain] / front_counts[domain] if front_counts[domain] else 0
        lines.extend(
            [
                f"### {domain}（课堂短名：{FRONT_DOMAIN_SHORT[domain]}；{front_counts[domain]} 道：主观 {front_subj_counts[domain]}，选择 {front_choice_counts[domain]}，主观占比 {subjective_ratio:.0%}）",
                "",
                FRONT_DOMAIN_METHOD[domain],
                "",
                "写答案要展开的部分：" + SUBJECTIVE_CORE_BY_FRONT[domain],
                "",
                "只在选择题里识别错项的部分：" + CHOICE_POOL_BY_FRONT[domain],
                "",
                DUAL_STATE_BY_FRONT[domain],
                "",
            ]
        )
    lines.extend(
        [
            "### 五域分布",
            "",
        ]
    )
    for domain in FRONT_DOMAIN_ORDER:
        lines.append(f"- {domain}（{FRONT_DOMAIN_SHORT[domain]}）：{front_counts[domain]} 道，主观 {front_subj_counts[domain]}，选择 {front_choice_counts[domain]}")
    lines.extend(
        [
            "",
        ]
    )
    lines.extend(
        [
            "## 八域",
            "",
            "八域是教师后台的精细索引，用来把每道题准确放回框架。学生最终仍按五域记忆，遇到具体题再由五域下钻到八域。",
            "",
        ]
    )
    for domain in FRAMEWORK_DOMAIN_ORDER:
        lines.append(
            f"- {domain}：{domain_counts[domain]} 道（主观 {domain_subj_counts[domain]}，选择 {domain_choice_counts[domain]}）。核心范围：{DOMAIN_CORE[domain]}"
        )
    lines.extend(
        [
            "",
            "## 主观题与选择题分层",
            "",
            "写主观题时，只展开能形成“关系句、规则句、结论句、价值句”的内容。选择题可以考得更细、更碎、更边界化，例如期限、程序身份、登记误区、举证细节、单个概念陷阱；这些内容适合做排错训练，不适合背成主观题模板。",
            "",
            "第三状态要记住：有些细节平时放在选择题排错里，但遇到主观题明确问到或材料需要时，可以作为一句采分点写入，不能扩写成整段。",
            "",
            "### 主观题优先训练",
            "",
            "- 合同、消费、劳动：合同成立/效力、格式条款、消费者维权、未成年人交易、劳动关系与竞业限制。",
            "- 物权、家庭、继承：共有与相邻关系、绿色物权、遗赠扶养、赡养与家庭伦理。",
            "- 人格与侵权：人格权益、一般侵权、安全保障义务合理限度、见义勇为补偿、网络/个人信息侵权。",
            "- 知识产权与竞争：著作权、作品独创性、AI/数字人场景下的权利边界、商标和商业标识、不正当竞争、商业秘密。",
            "- 创新竞争与公共救济：知识产权、不正当竞争、调解、仲裁、诉讼、公益诉讼、举证责任、司法确认与强制执行效力、司法措施。",
            "",
            "### 选择题优先识别",
            "",
            "- 物权细节：质押、抵押、担保物权、登记交付、一房二卖、宅基地、土地经营权。",
            "- 程序细节：辩护人/诉讼代理人、自诉/民事诉讼、二审终审、审判监督、司法确认具体程序、行政裁决、诉讼时效具体期间。",
            "- 知识产权细节：著作权登记误区、发表权保护期、法定许可、专利许可、数字人模型、AI作品登记或主体细节。",
            "- 消费与劳动细节：生鲜灯、先用后付、价格规则、直播带货、录用通知、一函两书、女职工保护、民事行为能力具体年龄分段。",
            "- 侵权细节：运动伤害公平分担、新闻报道公共利益抗辩、个人信息单独同意、动物饲养、高空抛物、产品责任等特殊侵权排错。",
            "",
            "### 答题决策树",
            "",
            "看到一道题，先判题型。选择题用细节池排错；主观题用关系、规则、结论、意义展开；综合民法题或多案例并列题逐项处理，再统一收束。",
        ]
    )
    lines.extend(
        [
            "",
            "## 高频主干",
            "",
            "### 主干一：法律关系判断",
            "",
            "先确定主体之间是什么关系。没有关系判断，权利、义务、责任、程序和价值都会漂。",
            "",
            "### 主干二：事实触发规则",
            "",
            "材料细节不是故事背景，而是给分条件。平台用工题看从属性事实，继承题看遗嘱或协议效力，侵权题看损害和责任，程序题看路径和节点。",
            "",
            "### 主干三：答案结果落点",
            "",
            "主观题不能只写规则名称，必须落到本案结论：是否成立、是否有效、归谁所有、谁承担责任、走什么程序、为什么有意义。程序救济有双身份：在第五域里是独立主线，在其他四域里是落结果时的救济路径。",
            "",
            "### 主干四：价值后置",
            "",
            "凡意义题，先写本案中如何保护权利、划清边界、落实处理结果，再上升到制度、秩序和价值。价值不能替代法律分析。",
            "",
            "## 七类命题路径",
            "",
            "1. 裁判依据题：案例事实 -> 双方争议 -> 规则适用 -> 裁判结果。",
            "2. 程序救济题：当事人处境 -> 程序选择 -> 节点要求 -> 法律效果。",
            "3. 司法措施题：法律关系 -> 司法动作 -> 规则效果 -> 秩序价值。",
            "4. 表格补全题：看示例 -> 找同类事实 -> 补规则理由 -> 补结果意义。",
            "5. 观点评析题：拆观点 -> 判合理与片面 -> 给法律依据 -> 回到本案结论。",
            "6. 意义认识题：本案结果 -> 权利保护 -> 秩序价值，必须先案后升。",
            "7. 选择题：题干事实 -> 题肢关键词 -> 排主体错位、关系错位、责任错位、价值偷换。",
            "",
            "## 3必+3选解释链",
            "",
            "- 3 必：关系句、规则句、结论句。每道主观题至少要说清谁和谁、依据什么规则、最后怎么处理。",
            "- 3 选：争点句、适用句、价值句。争点和适用按分值展开；价值句只在意义、认识、作用、启示类设问展开。",
            "- 意义题双段：先写本案处理，再写制度价值。选择题则把这条链压缩成“定关系、看规则、排错项”。",
            "",
            "## 学生答题黑名单",
            "",
            "以下表达只能在本案法律分析之后使用，不能当作答案主体：全面依法治国、社会公平正义、新质生产力、营商环境、司法温度、弘扬文明新风、法治与德治相结合。出现这些词前，必须先写清主体关系、规则依据和处理结果。",
            "",
            "## 框架自我进化流程",
            "",
            "1. 新题进入后，先锁定设问、材料触发、给分要点和答案落点。",
            "2. 再判断它是否能被现有一核、二线、三问、四步、五域自然容纳。",
            "3. 如果不能容纳，先作为候选节点，不立刻改主干。",
            "4. 候选节点必须至少解释多道题，并能转化成学生可执行动作，才允许升级。",
            "5. 升级前做两道反向测试：是否把高中政治做成专业法考；是否把法律题写成空泛法治口号。",
            "6. 最后把每道题重新回填到框架中，确认没有题被硬塞、漏放或只靠热点词归位。",
            "",
            "## 每题回填",
            "",
        ]
    )
    for domain in FRAMEWORK_DOMAIN_ORDER:
        domain_records = by_domain.get(domain, [])
        if not domain_records:
            continue
        lines.extend(
            [
                f"## {domain}",
                "",
                f"本节点回填 {len(domain_records)} 道题。核心范围：{DOMAIN_CORE[domain]}",
                "",
            ]
        )
        for record in domain_records:
            lines.extend(
                [
                    f"### {clean_text(record['label'])}",
                    "",
                    f"- 题型：{qtype_name(record)}",
                    f"- 前台五域：{record['front_domain']}",
                    f"- 材料触发：{clean_text(record['material_triggers'])}",
                    f"- 设问：{clean_text(record['complete_prompt'], 700)}",
                    f"- 答案落点：{answer_summary(record, 700)}",
                    f"- 完整解释链：{chain_text(record)}",
                    "",
                ]
            )
    FRAMEWORK_MD.write_text("\n".join(lines), encoding="utf-8")


def set_run_font(run, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = 2
    run = p.add_run("第 ")
    set_run_font(run, 9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run2 = p.add_run(" 页")
    set_run_font(run2, 9)


def configure_doc(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(60)
    section.right_margin = Pt(60)
    header = section.header.paragraphs[0]
    header.text = title
    header.alignment = 1
    for run in header.runs:
        set_run_font(run, 9, False)
    add_page_number(section)
    styles = doc.styles
    for name, east, size, bold in [
        ("Normal", "宋体", 10.5, None),
        ("Title", "黑体", 22, True),
        ("Heading 1", "黑体", 16, True),
        ("Heading 2", "黑体", 13, True),
        ("Heading 3", "黑体", 11, True),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east)
        style.font.size = Pt(size)
        if bold is not None:
            style.font.bold = bold


def add_para(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 10.5, text.startswith("【") and text.endswith("】"))
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08


def md_to_docx(md_path: Path, docx_path: Path, title: str) -> None:
    doc = Document()
    configure_doc(doc, title)
    first_title = False
    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            if first_title:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                add_para(doc, line[2:].strip(), "Heading 1")
            else:
                add_para(doc, line[2:].strip(), "Title")
                first_title = True
        elif line.startswith("## "):
            add_para(doc, line[3:].strip(), "Heading 1")
        elif line.startswith("### "):
            add_para(doc, line[4:].strip(), "Heading 2")
        elif line.startswith("- "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.first_line_indent = Pt(-9)
            run = p.add_run("• " + line[2:].strip())
            set_run_font(run, 10)
            p.paragraph_format.space_after = Pt(3)
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.first_line_indent = Pt(-9)
            run = p.add_run(line)
            set_run_font(run, 10)
            p.paragraph_format.space_after = Pt(3)
        else:
            add_para(doc, line)
    doc.save(docx_path)


def md_to_html(md_path: Path, html_path: Path, title: str) -> None:
    body: list[str] = []
    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if not line:
            continue
        escaped = html.escape(line)
        if line.startswith("# "):
            body.append(f"<h1>{html.escape(line[2:].strip())}</h1>")
        elif line.startswith("## "):
            body.append(f"<h2>{html.escape(line[3:].strip())}</h2>")
        elif line.startswith("### "):
            body.append(f"<h3>{html.escape(line[4:].strip())}</h3>")
        elif line.startswith("- "):
            body.append(f"<p class='bullet'>• {html.escape(line[2:].strip())}</p>")
        elif re.match(r"^\d+\. ", line):
            body.append(f"<p class='bullet'>{escaped}</p>")
        else:
            body.append(f"<p>{escaped}</p>")
    html_path.write_text(
        """<!doctype html>
<html><head><meta charset="utf-8"><title>"""
        + html.escape(title)
        + """</title><style>
@page { margin: 22mm 20mm; }
body { font-family: "Songti SC", "STSong", "STHeiti", serif; font-size: 11pt; line-height: 1.5; color: #111; }
h1 { font-family: "Heiti SC", "STHeiti", sans-serif; font-size: 22pt; margin: 0 0 12pt; page-break-after: avoid; }
h2 { font-family: "Heiti SC", "STHeiti", sans-serif; font-size: 16pt; margin: 18pt 0 8pt; page-break-after: avoid; border-bottom: 1px solid #ddd; padding-bottom: 3pt; }
h3 { font-family: "Heiti SC", "STHeiti", sans-serif; font-size: 13pt; margin: 12pt 0 6pt; page-break-after: avoid; }
p { margin: 0 0 5pt; }
.bullet { margin-left: 16pt; text-indent: -9pt; }
</style></head><body>
"""
        + "\n".join(body)
        + "\n</body></html>\n",
        encoding="utf-8",
    )


def check_forbidden(path: Path) -> list[str]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return [p for p in FORBIDDEN_PATTERNS if p in text]


def main() -> int:
    DELIVERY.mkdir(parents=True, exist_ok=True)
    records = load_records()
    write_collection(records)
    write_framework(records)
    md_to_docx(COLLECTION_MD, COLLECTION_DOCX, "选必二《法律与生活》法律题处理合集")
    md_to_docx(FRAMEWORK_MD, FRAMEWORK_DOCX, "选必二《法律与生活》最终进化框架")
    md_to_html(COLLECTION_MD, COLLECTION_HTML, "选必二《法律与生活》法律题处理合集")
    md_to_html(FRAMEWORK_MD, FRAMEWORK_HTML, "选必二《法律与生活》最终进化框架")
    for path in [COLLECTION_MD, FRAMEWORK_MD]:
        bad = check_forbidden(path)
        if bad:
            raise SystemExit(f"forbidden tokens in {path}: {bad}")
    print(COLLECTION_MD)
    print(FRAMEWORK_MD)
    print(COLLECTION_DOCX)
    print(FRAMEWORK_DOCX)
    print(COLLECTION_HTML)
    print(FRAMEWORK_HTML)
    print(f"records={len(records)} subjective={sum(1 for r in records if r['question_type']=='subjective')} choice={sum(1 for r in records if r['question_type']=='choice')}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
