#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


BASE = Path("/Users/wanglifei/Desktop/北京高考政治/选必二重做_2026-04-30")
OUT = BASE / "final_legal_outputs"


FILES = [
    (
        OUT / "选必二法律题全量处理合集_2026-05-03.md",
        OUT / "选必二法律题全量处理合集_2026-05-03.docx",
        "选必二《法律与生活》法律题全量处理合集",
    ),
    (
        OUT / "选必二法律与生活最终进化框架_2026-05-03.md",
        OUT / "选必二法律与生活最终进化框架_2026-05-03.docx",
        "选必二《法律与生活》最终进化框架",
    ),
]


def set_run_font(run, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_style_font(style, east_asia="宋体", latin="Times New Roman", size=10.5, bold=None) -> None:
    style.font.name = latin
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = 2
    run = p.add_run("第 ")
    set_run_font(run, 9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run2 = p.add_run(" 页")
    set_run_font(run2, 9)


def configure_doc(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(60)
    section.right_margin = Pt(60)
    header = section.header.paragraphs[0]
    header.text = title
    header.alignment = 1
    for run in header.runs:
        set_run_font(run, 9, False)
    add_page_number(section)
    styles = doc.styles
    set_style_font(styles["Normal"], size=10.5)
    set_style_font(styles["Title"], east_asia="黑体", size=22, bold=True)
    set_style_font(styles["Heading 1"], east_asia="黑体", size=16, bold=True)
    set_style_font(styles["Heading 2"], east_asia="黑体", size=13, bold=True)
    set_style_font(styles["Heading 3"], east_asia="黑体", size=11, bold=True)


def add_para(doc: Document, text: str, style: str | None = None) -> None:
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    run = p.add_run(text)
    if text.startswith("【") and text.endswith("】"):
        set_run_font(run, 10.5, True)
    else:
        set_run_font(run, 10.5)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08


def md_to_docx(md_path: Path, docx_path: Path, title: str) -> None:
    doc = Document()
    configure_doc(doc, title)
    lines = md_path.read_text(encoding="utf-8").splitlines()
    first_title_written = False
    for raw in lines:
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            if first_title_written:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                add_para(doc, line[2:].strip(), "Heading 1")
            else:
                add_para(doc, line[2:].strip(), "Title")
                first_title_written = True
            continue
        if line.startswith("## "):
            add_para(doc, line[3:].strip(), "Heading 1")
            continue
        if line.startswith("### "):
            add_para(doc, line[4:].strip(), "Heading 2")
            continue
        if line.startswith("#### "):
            add_para(doc, line[5:].strip(), "Heading 3")
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style=None)
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.first_line_indent = Pt(-9)
            run = p.add_run("• " + line[2:].strip())
            set_run_font(run, 10)
            p.paragraph_format.space_after = Pt(3)
            continue
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            run = p.add_run(line[2:].strip())
            set_run_font(run, 10, False)
            run.italic = True
            continue
        add_para(doc, line)
    doc.save(docx_path)


def main() -> int:
    for md_path, docx_path, title in FILES:
        md_to_docx(md_path, docx_path, title)
        print(docx_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
