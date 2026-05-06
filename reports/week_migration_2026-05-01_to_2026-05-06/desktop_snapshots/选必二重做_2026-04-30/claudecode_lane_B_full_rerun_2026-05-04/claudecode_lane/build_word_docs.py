#!/usr/bin/env python3
"""
Round 7: 生成两份 Word 交付
A. 框架版（按主观题情境主干 + 选择题章节双线）
B. 情境版（按情境穷尽 + 主观题与选择题分开）

设计原则：
- 顶层服从韩金阳总纲（知识链段+功能落点段；制度内容细节；不法考化/不必修三化）
- 学生文档绝不出现：Codex/Claude/GPT/参考答案/评标/PASS/FAIL/路径/pipeline 等后台话
- 情境-争点-知识链-功能落点 三层为主观题主干
- 自动字段尽量填，未填的明示"[待手写细化]"，不糊弄学生
"""
import json, re, os
from pathlib import Path
from collections import defaultdict
import re as _re_mod  # ensure re is available globally
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

PROJECT = Path("/Users/wanglifei/Desktop/北京高考政治/选必二重做_2026-04-30/claudecode_lane_B_full_rerun_2026-05-04")
CTRL = PROJECT / "00_control"
DELIVERY = PROJECT / "delivery"
DELIVERY.mkdir(exist_ok=True, parents=True)

# 域顺序（学习顺序，从最高频起）
DOMAIN_ORDER = ["交易交换", "身份财产", "人身权益", "创新竞争", "救济公共"]
DOMAIN_BLURB = {
    "交易交换": "学生最常遇到的法律情境——网购、平台、合同、消费者权益、劳动用工。共同原理：承诺—履行—信息对称—权益保护。",
    "身份财产": "婚姻家庭、继承、物权、宅基地、相邻关系。共同原理：身份/产权法定确认+流转规则。",
    "人身权益": "肖像权、声音权、个人信息、名誉、隐私、一般与特殊侵权。共同原理：人格独立 + 不当行为责任。",
    "创新竞争": "知识产权、不正当竞争、商业秘密、新业态。共同原理：私权排他性 vs. 市场公平秩序的张力。",
    "救济公共": "调解、仲裁、诉讼、公益诉讼、举证、执行。共同原理：争议解决程序 + 公益边界。",
}

# 选择题按制度细节章节组织
CHOICE_CHAPTER_ORDER = [
    ("第一章 民法基础与合同", ["交易交换"]),
    ("第二章 物权、婚姻、继承", ["身份财产"]),
    ("第三章 人格权与侵权责任", ["人身权益"]),
    ("第四章 创新竞争与市场秩序", ["创新竞争"]),
    ("第五章 争议解决与公益保护", ["救济公共"]),
    ("第六章 其他情境", ["未分类"]),
]

# ========== 完整手写题卡（46 道全覆盖）==========
import sys
sys.path.insert(0, str(Path(__file__).parent))
from hand_crafted_full import HAND_CRAFTED_SUBJ as HAND_CRAFTED, CHOICE_TYPE_PATTERNS

_OLD_INLINE = {
    # 2025 朝阳一模 19 网络平台搭售欺诈（最完整 L1 阅卷总结）
    "2025/2025朝阳一模::19": {
        "命题路径": "命题人想培养'识别经营者欺诈+消费者特殊保护'的法律应用能力。载体选 网络平台搭售红包 这一典型新业态情境。设问路径：事实切片→欺诈四要件→可撤销+三倍赔偿→数字消费秩序意义。细则奖励四件事：合同成立的要约/承诺识别、搭售三个事实条件全写、消费者权益保护法精确援引、回扣本案后再上升数字经济功能。",
        "知识链段": "1) 法律关系：王某与A公司是网络服务合同关系。2) 要件：合同成立(支付+出票=承诺)；欺诈四要件(故意+搭售红包行为+材料3条件全写+违背真意)。3) 规则：民法典关于可撤销合同；消费者权益保护法。4) 效力：合同可撤销+经营者欺诈成立。5) 救济：撤销合同+退还票款310元+三倍赔偿。6) 程序：王某诉至人民法院。",
        "功能落点段": "维护消费者合法权益（东向）；规范在线文旅平台经济企业经营、引导电商主体规范经营（南向）；持续优化网络消费环境、明确权利行使边界和责任范围、促进数字经济持续健康发展（北向）。",
        "踩分硬卡口": "▌'消费者权益保护法'必须出现'消费者法律'字样，写'保护法'不给分。 ▌搭售情境必须把材料 3 个条件全写：擅自搭售红包+界面不能清楚知悉费用+无法拒绝支付额外费用，缺一不得分。 ▌'欺诈'必须扣到主观故意+客观行为+因果+违背真意四要件。",
        "等价表达": "诱骗购买 ≈ 欺诈行为；违背真实意思表示 ≈ 违背真意；网络消费环境 ≈ 数字消费环境；规范企业经营 ≈ 引导电商主体合规。",
        "必修三化风险": "✗ 写'弘扬法治精神''全面依法治国'不给分；✗ 用'保护法'代替'消费者权益保护法'不给分；✗ 把意义只写到'保护消费者权益'就停（缺数字经济落点）少 1-2 分。",
        "元层评注": "本题是网络消费法律保护的标杆题，命题人在阅卷桌上画了三道硬线：术语精确（消费者权益保护法）、事实完整（材料 3 条件）、意义具体（数字经济）。学生背口诀'撤销+三倍+数字'三步走。",
    },
    # 2026 东城一模 18 知识产权三案例
    "2026/2026东城一模::18": {
        "命题路径": "命题人想培养'多案例并列时识别不同处理方式'的高级法律思维。载体选 三个并列司法案例（调解+严判+驳回）。设问路径：每个案例分别识别处理方式→对应法律依据→分别落点→总说收束。细则奖励：四个独立得分点+总说独立可得+不同案例间的相同价值话术不重复给。",
        "知识链段": "案例一：诉讼调解→推动纠纷实质性化解→以和为贵+节约司法资源。案例二：知产恶意侵权（主观恶性+社会危害大）→惩罚性赔偿→警示规范+保护当事人权益+维护市场秩序+诚信原则+公平竞争。案例三：恶意诉讼/滥用诉讼权利→驳回并谴责→诚信原则+权利义务统一+维护诉讼/司法秩序。总说：法院公正司法/裁判导向作用 + 保护知识产权就是保护创新。",
        "功能落点段": "公正司法/以事实为依据以法律为准绳/司法裁判导向作用（西向）；保护知识产权就是保护创新、提升创新活力（南向）；维护市场公平竞争秩序（北向）；和谐价值观、法德结合、诚信原则（东向）。",
        "踩分硬卡口": "▌案例独立性：三个案例各自得分，不能合并写。 ▌'节约司法资源'在案例一给过，案例三再写不重复给。 ▌'诚信原则'在案例二给过，案例三再写不重复给。 ▌案例三必须写'诉讼/司法秩序'，写'社会秩序'不给。 ▌案例三必须辨清是滥用诉讼权利，不是商业诋毁/不正当竞争（错位扣分）。 ▌案例二必须辨清是惩罚性赔偿场景，不是简单的混淆行为'搭便车'。",
        "等价表达": "诉讼调解 ≈ 调解；惩罚性赔偿 ≈ 严惩主观恶性强、社会危害大的知产侵权；驳回滥用诉讼权利 ≈ 谴责恶意诉讼。",
        "必修三化风险": "✗ 把案例三写成'维护社会秩序'丢 1 分（必须司法/诉讼秩序）；✗ 总说写'弘扬法治精神'不给（必须公正司法+保护创新双视角）；✗ 案例间共享价值话术（'诚信''节约司法资源'）超过给分次数浪费篇幅。",
        "元层评注": "本题是 2026 高频新题型——多案例聚合。学生最容易踩的坑是把所有案例的价值话术堆到一起写（看似全面实则因不重复给规则只能拿 1 次分），正确做法是每个案例独立给一个法律依据+一个责任落点+一个价值收束。",
    },
    # 2024 西城二模 16 食品惩罚性赔偿+消费者保护边界
    "2024/西城二模::16": {
        "命题路径": "命题人想培养'消费者保护权利的边界识别'。载体选 知假买假大量加购 这一争议情境。设问路径：经营者违法→首单支持→明知加购不支持→维护合理消费秩序的双向价值。细则奖励：同时讲清'支持'和'不支持'两侧法律依据。",
        "知识链段": "1) 法律关系：甲与经营者是食品买卖合同关系。2) 经营者违法：销售不符合食品安全标准的食品。3) 食品安全法第148条：消费者可主张十倍惩罚性赔偿。4) 边界：明知问题后大量加购已偏离生活消费目的。5) 处理：首单合理消费部分十倍赔偿支持；明知加购部分不支持。",
        "功能落点段": "打击违法经营、保护消费者权益（南向）；引导经营者诚信经营（南向）；防止消费者保护被滥用（西向）；维护食品安全制度的真正功能性（北向）；维护市场公平交易秩序（北向）。",
        "踩分硬卡口": "▌必须同时写'首单支持'和'明知加购不支持'两侧，缺一只能拿一半分。 ▌必须明确指出'生活消费'的边界——明知加购不属生活消费。 ▌惩罚性赔偿术语必须精确，不能简化为'赔偿'。",
        "等价表达": "知假买假 ≈ 明知问题后大量加购；偏离生活消费目的 ≈ 不再是为生活需要购买。",
        "必修三化风险": "✗ 只写'保护消费者权益'不写'防止滥用'丢一半分；✗ 写'弘扬诚信价值'代替具体经营者诚信义务空洞；✗ 用'消费者权益保护法'统一术语而不区分食品安全法十倍赔偿场景。",
        "元层评注": "本题是消费者保护场景的'反例平衡题'——命题人借此教学生：法律保护是有边界的，权利不能滥用。这是韩金阳总纲'政治性↔学理性'轴的典型体现：既要保护权益（政治性），又要守住学理边界（学理性）。",
    },
}


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h

def add_para(doc, text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p

def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def split_long_text(s, max_chars=280):
    """把长文本按句子折成短段"""
    # 这里只用于自动展示，避免一行 800 字
    return s

def sanitize(text):
    """学生文档禁词清洗：移除/替换后台话。"""
    if not text:
        return text
    # 长串先替换，避免子串被错误覆盖
    repl = [
        ("阅卷前制定的", "本题"),
        ("阅卷桌", "命题人"),
        ("阅卷反馈", "答题反馈"),
        ("阅卷总结", "答题总结"),
        ("阅卷人", "命题人"),
        ("阅卷组", "命题组"),
        ("阅卷前", "本题"),
        ("阅卷", "命题"),                  # 兜底
        ("评标", "给分依据"),
        ("参考答案", "答题要点"),
        ("讲评", "答题指引"),
        ("勿传", ""),
        ("Round N+1", "下一轮"),
        ("Round N", "下一轮"),
        ("===== Slide", "（页"),           # 把幻灯片标记降级
    ]
    for old, new in repl:
        text = text.replace(old, new)
    return text

def render_subjective_question(doc, q, hand=None):
    """把一道主观题渲染到 doc。hand: 手写字段 dict，无则用占位。"""
    suite_label = q["suite"].replace("/", " · ")
    title = f"【{suite_label}  第{q['qnum']}题】"
    add_heading(doc, title, level=4, color=RGBColor(0x1F, 0x4E, 0x79))

    settings = q.get("setting_extracted") or "（未提取到设问，请回卷复核）"
    add_para(doc, f"设问： {settings}", size=10.5)

    mat = q.get("material_compressed") or ""
    add_para(doc, f"材料一句话（起因—经过—结果）： {mat}", size=10)

    issues = q.get("issues") or []
    if issues:
        add_para(doc, "争点：", bold=True, size=10)
        for i, iss in enumerate(issues, 1):
            add_bullet(doc, f"{i}. {iss}", size=10)

    if hand:
        add_para(doc, "命题路径：", bold=True, size=10)
        add_para(doc, hand.get("命题路径", ""), size=10)
        add_para(doc, "知识链段（要件→规则→效力→责任→救济）：", bold=True, size=10)
        add_para(doc, hand.get("知识链段", ""), size=10)
        add_para(doc, "功能落点段（论证法律推动法治建设/社会经济发展）：", bold=True, size=10)
        add_para(doc, hand.get("功能落点段", ""), size=10)
        p = doc.add_paragraph()
        r = p.add_run("踩分硬卡口：")
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        add_para(doc, hand.get("踩分硬卡口", ""), size=10)
        add_para(doc, "等价表达：", bold=True, size=10)
        add_para(doc, hand.get("等价表达", ""), size=10)
        add_para(doc, "必修三化风险：", bold=True, size=10)
        add_para(doc, hand.get("必修三化风险", ""), size=10)
        add_para(doc, "元层评注：", bold=True, italic=True, size=10)
        add_para(doc, hand.get("元层评注", ""), italic=True, size=10)
    else:
        # 未手写题：贴细则段 + 标"待细化"
        rubric = q.get("rubric_segments") or []
        if rubric:
            add_para(doc, "细则核心要点（命题人给分依据）：", bold=True, size=10)
            seg = sanitize(rubric[0]["snippet"][:1500])
            add_para(doc, seg, size=9.5)
        else:
            add_para(doc, "（本题未自动匹配到细则段，需回原卷查阅。）", italic=True, size=10)
        add_para(doc, "命题路径 / 知识链段 / 功能落点 / 踩分硬卡口 / 等价表达 / 必修三化风险 / 元层评注：[本题示例暂缺，建议对照同情境已写题作答]", italic=True, size=10)

    add_para(doc, "─" * 30, size=8)


def render_choice_question(doc, q):
    suite_label = q["suite"].replace("/", " · ")
    title = f"【{suite_label}  第{q['qnum']}题】"
    add_heading(doc, title, level=4, color=RGBColor(0x1F, 0x4E, 0x79))
    add_para(doc, f"情境分类： {q.get('situation_l1','')} > {q.get('situation_l2','')}", size=9.5, italic=True)

    qtext_short = sanitize(q["qtext"][:600].replace("\n", " ").strip())
    add_para(doc, f"题面： {qtext_short}", size=10)

    options = q.get("options_abcd") or []
    if options:
        for letter, content in options[:4]:
            add_bullet(doc, f"{letter}．{content.strip()[:200]}", size=9.5)

    # 提取已知答案
    ans_match = re.search(r'【答案】\s*([A-D])|本题选\s*([A-D])', q["qtext"])
    if ans_match:
        ans = ans_match.group(1) or ans_match.group(2)
        add_para(doc, f"答案：{ans}", bold=True, size=10)

    # 引用情境陷阱模式
    l2 = q.get('situation_l2','')
    pattern = CHOICE_TYPE_PATTERNS.get(l2)
    if pattern:
        add_para(doc, f"本题考点：{pattern['考点核心']}", size=9.5, italic=True)
        add_para(doc, f"判断方向：{pattern['正确判断方向']}", size=9.5, italic=True)
    add_para(doc, "─" * 20, size=8)


def write_framework_doc(subj, choice):
    """框架版 Word"""
    doc = Document()

    # 顶层
    h = doc.add_heading("选必二《法律与生活》框架版", level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    add_heading(doc, "一、命题总纲（先看这页，每次都看）", level=1)
    add_para(doc, "选必二的题永远在这两条主轴上：", bold=True)
    add_bullet(doc, "政治性 ↔ 学理性：法律的精确知识与价值取向同时考。")
    add_bullet(doc, "价值性 ↔ 知识性：法律知识是答案主体，价值是法律功能性的落点。")

    add_para(doc, "主观题答题结构（永远两段）：", bold=True)
    add_bullet(doc, "知识链段：法律关系判断 → 法定要件套用 → 规则援引 → 效力/责任判定 → 救济选择 → 程序落点。")
    add_bullet(doc, "功能落点段：论证这条法律规则如何推动 法治建设 / 社会经济发展 / 市场秩序 / 新业态规范。回扣本案再上升，不许写抽象口号。")

    add_para(doc, "选择题考查重点：", bold=True)
    add_bullet(doc, "考制度本身的内容乃至细节（具体规则、条文、要件、程序），不只考价值取向。")

    add_para(doc, "两条永不可越的边界：", bold=True)
    add_bullet(doc, "不法考化：不超纲，不出过多主体/过多法律关系/疑难问题。")
    add_bullet(doc, "不必修三化：法律知识是答案主体，不能让宏观法治口号占主导。")

    # 主观题主干
    add_heading(doc, "二、主观题主干（按情境分组）", level=1)

    by_l1_l2 = defaultdict(lambda: defaultdict(list))
    for e in subj:
        by_l1_l2[e.get("situation_l1", "未分类")][e.get("situation_l2", "未分类")].append(e)

    for l1 in DOMAIN_ORDER:
        if l1 not in by_l1_l2:
            continue
        add_heading(doc, l1 + "域", level=2)
        add_para(doc, DOMAIN_BLURB.get(l1, ""), italic=True, size=10)
        for l2, qs in by_l1_l2[l1].items():
            add_heading(doc, f"  {l2}（{len(qs)} 道）", level=3)
            for q in qs:
                key = f"{q['suite']}::{q['qnum']}"
                hand = HAND_CRAFTED.get(key)
                render_subjective_question(doc, q, hand=hand)

    # 横切层
    add_heading(doc, "三、横切工具层（六张表）", level=1)

    add_heading(doc, "表 1：要件库（高频法律概念的要件组合）", level=2)
    rule_kit = [
        ("合同成立", "要约+承诺，承诺到达发生效力"),
        ("合同欺诈（可撤销）", "故意 + 欺诈行为 + 因果关系 + 违背真实意思表示"),
        ("消费者欺诈（三倍赔偿）", "经营者欺诈故意 + 误导消费 + 影响真实选择 + 属生活消费"),
        ("食品安全（十倍赔偿）", "食品不符合食品安全标准 + 消费者购买"),
        ("夫妻共同债务", "为夫妻共同生活 / 共同经营 / 共同意思表示之一"),
        ("一般侵权", "行为 + 损害 + 因果关系 + 过错"),
        ("肖像权侵权", "未经许可 + 使用可识别肖像 + 无法定免责事由"),
        ("惩罚性赔偿（知产）", "主观恶性强 + 社会危害大 + 已有赔偿不足以惩罚"),
        ("不正当竞争混淆", "有影响的标识近似 + 容易导致消费者误认 + 攀附商誉"),
        ("竞业限制", "高管/高级技术/保密义务人员 + 期限不超 2 年 + 范围合理 + 用人单位给经济补偿"),
        ("公益诉讼（民事）", "损害公共利益 + 起诉主体适格（检察机关/适格社会组织）"),
        ("格式条款无效", "免除提供方主要责任 / 加重对方责任 / 排除对方主要权利 / 未尽提示说明义务"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "法律概念"
    hdr[1].text = "要件组合（缺一不可）"
    for name, kit in rule_kit:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = kit

    add_heading(doc, "表 2：救济库（按救济类型）", level=2)
    relief_kit = [
        ("合同救济", "继续履行 / 解除合同 / 赔偿损失 / 定金罚则 / 违约金"),
        ("消费者救济", "退还价款 / 三倍赔偿（消费欺诈）/ 十倍赔偿（食品）"),
        ("物权救济", "返还原物 / 排除妨害 / 消除危险 / 恢复原状"),
        ("人格权救济", "停止侵害 / 排除妨碍 / 消除危险 / 赔礼道歉 / 恢复名誉 / 消除影响 / 赔偿损失"),
        ("知产救济", "停止侵害 / 赔偿损失 / 惩罚性赔偿"),
        ("程序救济", "协商 / 调解（人民调解/诉讼调解）/ 仲裁 / 诉讼 / 公益诉讼 / 司法确认"),
    ]
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Table Grid"
    h2 = t2.rows[0].cells
    h2[0].text = "救济类型"
    h2[1].text = "具体救济方式"
    for k, v in relief_kit:
        row = t2.add_row().cells
        row[0].text = k
        row[1].text = v

    add_heading(doc, "表 3：法律功能落点表（四向指南针）", level=2)
    add_para(doc, "意义部分必须用具体功能话语，禁止抽象口号。", bold=True)
    fn_kit = [
        ("法治建设方向", "公正司法 / 以事实为依据以法律为准绳 / 司法裁判导向作用 / 明确权利行使界限"),
        ("社会经济发展方向", "市场秩序 / 营商环境 / 数字经济 / 新业态规范 / 行业规范化经营"),
        ("核心价值方向", "诚信原则 / 公平竞争 / 平等保护 / 弱势保护 / 和谐价值"),
        ("制度引导方向", "规范企业经营 / 警示侵权 / 激励创新 / 保护消费 / 维护行业秩序"),
    ]
    t3 = doc.add_table(rows=1, cols=2)
    t3.style = "Table Grid"
    h3 = t3.rows[0].cells
    h3[0].text = "方向"
    h3[1].text = "具体话语（可直接上卷面）"
    for k, v in fn_kit:
        row = t3.add_row().cells
        row[0].text = k
        row[1].text = v

    add_heading(doc, "表 4：法律精确性卡口表", level=2)
    add_para(doc, "学理性硬要求——术语不精确就丢分。", bold=True)
    prec_kit = [
        ("A1 法律名精确", "✓ 消费者权益保护法 / 食品安全法 / 民法典 ／  ✗ 保护法 / 民法 / 经济法"),
        ("A2 法律概念精确", "✓ 行使损害赔偿请求权 / 构成不正当竞争中的混淆行为 ／ ✗ 维权 / 违法 / 侵权（笼统）"),
        ("A3 责任/秩序定位精确", "✓ 违约责任 vs 侵权责任 / 司法秩序 vs 诉讼秩序 / 市场秩序 ／ ✗ 违法责任 / 社会秩序"),
    ]
    t4 = doc.add_table(rows=1, cols=2)
    t4.style = "Table Grid"
    h4 = t4.rows[0].cells
    h4[0].text = "精确性方向"
    h4[1].text = "对照"
    for k, v in prec_kit:
        row = t4.add_row().cells
        row[0].text = k
        row[1].text = v

    add_heading(doc, "表 5：答案内容主体卡口表（不必修三化）", level=2)
    add_para(doc, "韩金阳总纲：法律知识不能只是背景，答案内容均来自法治模块。", bold=True)
    body_kit = [
        ("答案主体", "✓ 民法典 X 条 / 法定要件 / 责任落点 / 救济方式 ／ ✗ 全面依法治国 / 法治精神 / 科学立法严格执法（不能为答案主体）"),
        ("功能话语", "✓ 维护消费者合法权益 / 规范企业经营 / 优化网络消费环境 / 促进数字经济发展 / 公正司法 / 保护创新激励 ／ ✗ 弘扬法治精神 / 推进依法治国 / 维护社会主义核心价值观（空洞收束）"),
    ]
    t5 = doc.add_table(rows=1, cols=2)
    t5.style = "Table Grid"
    h5 = t5.rows[0].cells
    h5[0].text = "层面"
    h5[1].text = "对照"
    for k, v in body_kit:
        row = t5.add_row().cells
        row[0].text = k
        row[1].text = v

    add_heading(doc, "表 6：踩分边界禁区表（命题人在阅卷桌的硬规则）", level=2)
    boundaries = [
        ("不重复给分", "多案例题中相同价值话术只能给一次（如'诚信原则''节约司法资源'）"),
        ("材料触发完整性", "材料给的多个事实条件必须全部写出（如搭售欺诈三条件全写）"),
        ("案例独立性", "多案例题每个案例必须独立产出法律依据+责任落点，不能合并"),
        ("法律名硬卡口", "'消费者权益保护法'必须写完整名称，'保护法'不给分"),
        ("秩序定位硬卡口", "诉讼/司法秩序场景不写'社会秩序'"),
        ("反错位提醒", "案例三若为滥用诉讼权利，不能写成'商业诋毁'或'不正当竞争'"),
    ]
    t6 = doc.add_table(rows=1, cols=2)
    t6.style = "Table Grid"
    h6 = t6.rows[0].cells
    h6[0].text = "边界规则"
    h6[1].text = "说明"
    for k, v in boundaries:
        row = t6.add_row().cells
        row[0].text = k
        row[1].text = v

    # 选择题主干（独立另一支）
    add_heading(doc, "四、选择题主干（独立于主观题）", level=1)
    add_para(doc, "选择题考'制度本身的内容乃至细节'。错项排除三件套：法律关系错位 / 制度细节错位 / 程序效力错位。", bold=True)

    # 选择题陷阱模式总览（按情境二级类）
    add_heading(doc, "选择题陷阱模式总览（按情境分类）", level=2)
    add_para(doc, "学生看一类题就掌握该类常见陷阱。下列每类列出考点核心+常见错项陷阱模式+正确判断方向，是选择题的横切工具。", italic=True)
    for l2_name, p in CHOICE_TYPE_PATTERNS.items():
        add_heading(doc, l2_name, level=3)
        add_para(doc, f"考点核心：{p['考点核心']}", bold=True, size=10)
        add_para(doc, "常见错项陷阱：", bold=True, size=10)
        for trap in p['常见陷阱模式']:
            add_bullet(doc, trap, size=9.5)
        add_para(doc, f"正确判断方向：{p['正确判断方向']}", italic=True, size=10)

    add_heading(doc, "选择题逐题列表（按章节分组）", level=2)

    by_choice_l1 = defaultdict(list)
    for e in choice:
        by_choice_l1[e.get("situation_l1", "未分类")].append(e)

    for chap_title, chap_l1s in CHOICE_CHAPTER_ORDER:
        items = []
        for l1 in chap_l1s:
            items.extend(by_choice_l1.get(l1, []))
        if not items:
            continue
        add_heading(doc, chap_title + f"（{len(items)} 道）", level=3)
        for q in items:  # 全列；不截断
            render_choice_question(doc, q)

    # 写盘
    out_path = DELIVERY / "选必二《法律与生活》框架版_2026-05-04.docx"
    doc.save(out_path)
    return out_path


def write_situation_doc(subj, choice):
    """情境版 Word"""
    doc = Document()
    h = doc.add_heading("选必二《法律与生活》情境版", level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    add_para(doc, "本册按【考过的情境】穷尽组织。每个主观题情境用一句话起因→经过→结果概括，列对应细则与作答动作。选择题情境单独一节，不与主观题混排。", italic=True)

    # 主观题情境部分
    add_heading(doc, "第一部分 主观题情境（46 道）", level=1)
    by_l1_l2 = defaultdict(lambda: defaultdict(list))
    for e in subj:
        by_l1_l2[e.get("situation_l1","未分类")][e.get("situation_l2","未分类")].append(e)

    for l1 in DOMAIN_ORDER:
        if l1 not in by_l1_l2:
            continue
        add_heading(doc, l1, level=2)
        for l2, qs in by_l1_l2[l1].items():
            add_heading(doc, l2 + f"（{len(qs)} 道）", level=3)
            for q in qs:
                key = f"{q['suite']}::{q['qnum']}"
                hand = HAND_CRAFTED.get(key)
                suite_label = q["suite"].replace("/", " · ")
                add_para(doc, f"【{suite_label} 第{q['qnum']}题】", bold=True, size=10.5)
                add_para(doc, f"情境一句话： {q.get('material_compressed','')[:300]}", size=10)
                add_para(doc, f"设问： {q.get('setting_extracted','')}", size=10)
                if q.get("issues"):
                    add_para(doc, f"争点： {' / '.join(q['issues'][:2])}", size=10)
                if hand:
                    add_para(doc, f"细则核心： {hand.get('知识链段','')[:500]}", size=10)
                    add_para(doc, f"硬卡口： {hand.get('踩分硬卡口','')[:500]}", size=10)
                    add_para(doc, f"作答落点： {hand.get('功能落点段','')[:300]}", size=10)
                else:
                    rubric = q.get("rubric_segments") or []
                    if rubric:
                        add_para(doc, "细则核心要点：", size=10)
                        add_para(doc, sanitize(rubric[0]["snippet"][:600]), size=9.5)
                    add_para(doc, "硬卡口 / 作答落点：[本题示例暂缺，建议对照同情境已写题作答]", italic=True, size=9.5)
                add_para(doc, "─" * 30, size=8)

    # 选择题情境部分
    add_heading(doc, "第二部分 选择题情境（101 道）", level=1)
    add_para(doc, "选择题情境只列：题源、情境关键词、题面要点、选项；正确答案与错项陷阱将在后续手工补充。", italic=True)

    by_l1 = defaultdict(list)
    for e in choice:
        by_l1[e.get("situation_l1","未分类")].append(e)
    for l1 in DOMAIN_ORDER + ["未分类"]:
        if l1 not in by_l1:
            continue
        add_heading(doc, l1 + f"（{len(by_l1[l1])} 道）", level=2)
        for q in by_l1[l1]:
            suite_label = q["suite"].replace("/", " · ")
            add_para(doc, f"【{suite_label} 第{q['qnum']}题】 {q.get('situation_l2','')}", bold=True, size=10)
            qtext_short = q["qtext"][:300].replace("\n", " ").strip()
            add_para(doc, qtext_short, size=9.5)

    out_path = DELIVERY / "选必二《法律与生活》情境版_2026-05-04.docx"
    doc.save(out_path)
    return out_path


def post_sanitize_docx(path):
    """对生成的 docx 做 XML 级最终清洗（兜底，确保无禁词残留）。"""
    import zipfile, shutil, tempfile, os
    repl = [
        ("阅卷前制定的", "本题"),
        ("阅卷反馈", "答题反馈"),
        ("阅卷总结", "答题总结"),
        ("阅卷人", "命题人"),
        ("阅卷组", "命题组"),
        ("阅卷前", "本题"),
        ("阅卷桌", "命题人"),
        ("阅卷", "命题"),
        ("评标", "给分依据"),
        ("参考答案", "答题要点"),
        ("讲评", "答题指引"),
        ("勿传", ""),
        ("Round N+1", "下一轮"),
        ("Round N", "下一轮"),
    ]
    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.endswith(".xml"):
                txt = data.decode("utf-8", errors="ignore")
                for old, new in repl:
                    txt = txt.replace(old, new)
                data = txt.encode("utf-8")
            zout.writestr(item, data)
    shutil.move(tmp, path)

def main():
    with open(CTRL / "SUBJECTIVE_PACK_ENRICHED.json", encoding="utf-8") as f:
        subj = json.load(f)
    with open(CTRL / "CHOICE_PACK_ENRICHED.json", encoding="utf-8") as f:
        choice = json.load(f)

    p1 = write_framework_doc(subj, choice)
    p2 = write_situation_doc(subj, choice)

    post_sanitize_docx(p1)
    post_sanitize_docx(p2)

    print(f"框架版： {p1}  ({p1.stat().st_size:,} bytes)")
    print(f"情境版： {p2}  ({p2.stat().st_size:,} bytes)")

if __name__ == "__main__":
    main()
