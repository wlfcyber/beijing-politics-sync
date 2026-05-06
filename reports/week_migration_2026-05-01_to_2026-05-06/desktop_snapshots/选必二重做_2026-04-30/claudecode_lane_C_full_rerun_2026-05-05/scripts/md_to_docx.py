#!/usr/bin/env python3
"""
Minimal Markdown → docx renderer for C-lane delivery.
Handles: # ## ### #### headings, bullet/numbered lists, tables, bold (**), strikethrough (~~), blockquotes.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

H_SIZE = {1: 22, 2: 16, 3: 13, 4: 12}
BODY_SIZE = 11
LIST_SIZE = 11
TABLE_SIZE = 10

INLINE_BOLD = re.compile(r"\*\*([^*]+)\*\*")
INLINE_STRIKE = re.compile(r"~~([^~]+)~~")
INLINE_CODE = re.compile(r"`([^`]+)`")


def add_inline(paragraph, text: str, base_size: int = BODY_SIZE):
    """Add a paragraph run handling **bold** ~~strike~~ `code`."""
    pos = 0
    pattern = re.compile(r"(\*\*[^*]+\*\*|~~[^~]+~~|`[^`]+`)")
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            r.font.size = Pt(base_size)
        seg = m.group(0)
        if seg.startswith("**"):
            r = paragraph.add_run(seg[2:-2])
            r.bold = True
        elif seg.startswith("~~"):
            r = paragraph.add_run(seg[2:-2])
            r.font.strike = True
            r.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
        elif seg.startswith("`"):
            r = paragraph.add_run(seg[1:-1])
            r.font.name = "Menlo"
        else:
            r = paragraph.add_run(seg)
        r.font.size = Pt(base_size)
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        r.font.size = Pt(base_size)


def is_table_line(line: str) -> bool:
    return line.strip().startswith("|") and line.strip().endswith("|") and "|" in line.strip()[1:-1]


def is_table_separator(line: str) -> bool:
    s = line.strip()
    if not (s.startswith("|") and s.endswith("|")):
        return False
    middle = s[1:-1]
    cells = [c.strip() for c in middle.split("|")]
    return all(re.match(r"^:?-+:?$", c) for c in cells if c)


def parse_table_row(line: str) -> list[str]:
    s = line.strip()
    return [c.strip() for c in s[1:-1].split("|")]


def render(md_path: Path, docx_path: Path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "PingFang SC"
    style.font.size = Pt(BODY_SIZE)

    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # Skip horizontal rules
        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        # Headings
        if stripped.startswith("#### "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[5:].strip())
            r.bold = True
            r.font.size = Pt(H_SIZE[4])
            i += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[4:].strip())
            r.bold = True
            r.font.size = Pt(H_SIZE[3])
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[3:].strip())
            r.bold = True
            r.font.size = Pt(H_SIZE[2])
            i += 1
            continue
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(stripped[2:].strip())
            r.bold = True
            r.font.size = Pt(H_SIZE[1])
            i += 1
            continue

        # Tables
        if is_table_line(line) and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            header = parse_table_row(line)
            rows = []
            i += 2  # skip header + separator
            while i < len(lines) and is_table_line(lines[i]):
                rows.append(parse_table_row(lines[i]))
                i += 1
            t = doc.add_table(rows=1 + len(rows), cols=len(header))
            t.style = "Light Grid Accent 1"
            for c_idx, cell_text in enumerate(header):
                cell = t.rows[0].cells[c_idx]
                cell.paragraphs[0].text = ""
                add_inline(cell.paragraphs[0], cell_text, base_size=TABLE_SIZE)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for r_idx, row in enumerate(rows):
                for c_idx, cell_text in enumerate(row):
                    if c_idx < len(t.rows[r_idx + 1].cells):
                        cell = t.rows[r_idx + 1].cells[c_idx]
                        cell.paragraphs[0].text = ""
                        add_inline(cell.paragraphs[0], cell_text, base_size=TABLE_SIZE)
            doc.add_paragraph()
            continue

        # Blockquote
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            r = p.add_run("「")
            r.font.size = Pt(BODY_SIZE)
            add_inline(p, stripped[2:].strip(), base_size=BODY_SIZE)
            r = p.add_run("」")
            r.font.size = Pt(BODY_SIZE)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        # Bullet list (- or *)
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent = len(m.group(1)) // 2
            content = m.group(2)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.6 + indent * 0.5)
            add_inline(p, content, base_size=LIST_SIZE)
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\s*)(\d+)[\.、]\s+(.*)$", line)
        if m:
            indent = len(m.group(1)) // 2
            content = m.group(3)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Cm(0.6 + indent * 0.5)
            add_inline(p, content, base_size=LIST_SIZE)
            i += 1
            continue

        # Code block
        if stripped.startswith("```"):
            i += 1
            block_lines = []
            while i < len(lines) and not lines[i].rstrip().startswith("```"):
                block_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            p = doc.add_paragraph()
            r = p.add_run("\n".join(block_lines))
            r.font.name = "Menlo"
            r.font.size = Pt(10)
            continue

        # Empty line
        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        add_inline(p, stripped, base_size=BODY_SIZE)
        i += 1

    doc.save(docx_path)
    print(f"wrote: {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("usage: md_to_docx.py input.md output.docx", file=sys.stderr)
        sys.exit(2)
    render(Path(sys.argv[1]), Path(sys.argv[2]))
