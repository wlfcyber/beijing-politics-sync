#!/usr/bin/env python3
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


FONT_BODY = "Microsoft YaHei"
FONT_SERIF = "SimSun"
BLUE_DARK = RGBColor(31, 78, 121)
BLUE_MID = RGBColor(46, 116, 181)
GRAY_TEXT = RGBColor(80, 80, 80)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color=None):
    run.font.name = FONT_BODY
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_font(paragraph, size: float | None = None):
    for run in paragraph.runs:
        set_run_font(run, size=size)


def add_field(paragraph, instruction: str, placeholder: str = ""):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char)

    if placeholder:
        text_run = paragraph.add_run(placeholder)
        set_run_font(text_run, 10, color=GRAY_TEXT)

    end_run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_char)


def set_update_fields_on_open(document: Document):
    settings = document.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def style_document(document: Document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        style = styles[style_name]
        style.font.name = FONT_BODY
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style.font.bold = True

    styles["Heading 1"].font.size = Pt(17)
    styles["Heading 1"].font.color.rgb = BLUE_DARK
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 2"].font.color.rgb = BLUE_DARK
    styles["Heading 3"].font.size = Pt(12.5)
    styles["Heading 3"].font.color.rgb = BLUE_MID
    styles["Heading 4"].font.size = Pt(11)
    styles["Heading 4"].font.color.rgb = RGBColor(40, 40, 40)

    if "Feige Body Note" not in styles:
        note = styles.add_style("Feige Body Note", WD_STYLE_TYPE.PARAGRAPH)
        note.base_style = styles["Normal"]
        note.font.name = FONT_BODY
        note._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        note.font.size = Pt(10)
        note.font.color.rgb = GRAY_TEXT


def configure_sections(document: Document):
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.0)

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_p.add_run("飞哥正志讲堂 · 选必三《逻辑与思维》")
    set_run_font(header_run, 9, color=GRAY_TEXT)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = footer_p.add_run("第 ")
    set_run_font(r1, 9, color=GRAY_TEXT)
    add_field(footer_p, "PAGE", "1")
    r2 = footer_p.add_run(" 页")
    set_run_font(r2, 9, color=GRAY_TEXT)


def add_runs_with_bold(paragraph, text: str):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def is_table_block(lines: list[str], idx: int) -> bool:
    if idx + 1 >= len(lines):
        return False
    return lines[idx].strip().startswith("|") and re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", lines[idx + 1])


def parse_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def add_table(document: Document, rows: list[list[str]]):
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_runs_with_bold(paragraph, cell_text)
            for run in paragraph.runs:
                set_run_font(run, 9.5, bold=(r_idx == 0))
    document.add_paragraph()


def add_title_page(document: Document, title: str):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, 20, bold=True, color=BLUE_DARK)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("思维部分按哲学式触发，推理部分按题型族群突破")
    set_run_font(run, 11, color=GRAY_TEXT)

    document.add_paragraph()
    toc_title = document.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_run = toc_title.add_run("目录")
    set_run_font(toc_run, 14, bold=True, color=BLUE_DARK)

    toc_p = document.add_paragraph()
    add_field(toc_p, r'TOC \o "1-3" \h \z \u', "在 Microsoft Word 中打开后会自动更新目录")
    document.add_page_break()


def add_markdown(document: Document, markdown: str):
    lines = markdown.splitlines()
    title = ""
    start_idx = 0
    if lines and lines[0].startswith("# "):
        title = lines[0][2:].strip()
        start_idx = 1
    add_title_page(document, title or "选必三《逻辑与思维》学生版")

    idx = start_idx
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()
        if not stripped:
            idx += 1
            continue

        if is_table_block(lines, idx):
            rows = [parse_table_row(lines[idx])]
            idx += 2
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                rows.append(parse_table_row(lines[idx]))
                idx += 1
            add_table(document, rows)
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = min(len(heading.group(1)), 4)
            text = heading.group(2).strip()
            p = document.add_heading(text, level=level)
            p.paragraph_format.space_before = Pt(8 if level <= 2 else 5)
            p.paragraph_format.space_after = Pt(4)
            idx += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            p = document.add_paragraph(style="List Number")
            add_runs_with_bold(p, re.sub(r"^\d+\.\s+", "", stripped))
            p.paragraph_format.left_indent = Cm(0.3)
            p.paragraph_format.first_line_indent = Cm(0)
            set_paragraph_font(p, 10.5)
            idx += 1
            continue

        if stripped.startswith(("- ", "* ")):
            p = document.add_paragraph(style="List Bullet")
            add_runs_with_bold(p, stripped[2:].strip())
            set_paragraph_font(p, 10.5)
            idx += 1
            continue

        p = document.add_paragraph()
        add_runs_with_bold(p, stripped)
        p.paragraph_format.first_line_indent = Pt(21)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(3)
        set_paragraph_font(p, 10.5)
        idx += 1


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: build_student_docx.py INPUT.md OUTPUT.docx", file=sys.stderr)
        return 2
    src = Path(sys.argv[1])
    dst = Path(sys.argv[2])
    document = Document()
    style_document(document)
    configure_sections(document)
    add_markdown(document, src.read_text(encoding="utf-8"))
    set_update_fields_on_open(document)
    dst.parent.mkdir(parents=True, exist_ok=True)
    document.save(dst)
    print(dst)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
