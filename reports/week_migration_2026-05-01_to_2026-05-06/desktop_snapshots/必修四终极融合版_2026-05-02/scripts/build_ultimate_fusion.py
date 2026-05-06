#!/usr/bin/env python3
from __future__ import annotations

import importlib.util
import json
import os
import re
import shutil
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path("/Users/wanglifei/Desktop/北京高考政治")
WORK = ROOT / "必修四终极融合版_2026-05-02"
OUT_DIR = WORK / "outputs"
AUDIT_DIR = WORK / "audit"
REPORT_DIR = WORK / "reports"

CODEX_ROOT = ROOT / "必修四最终整合_2026-05-02"
CODEX_ENTRIES = CODEX_ROOT / "audit" / "final_entries.json"
CLAUDE_ROOT = ROOT / "必修四从0重跑_ClaudeCode_2026-05-02"
CLAUDE_ENTRIES_DIR = CLAUDE_ROOT / "audit" / "entries"

REBUILD_PATH = ROOT / "哲学宝典5.2原地修订" / "scripts" / "rebuild_in_place_baodian.py"
spec = importlib.util.spec_from_file_location("codex_rebuild_baodian", REBUILD_PATH)
rebuild = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(rebuild)

TITLE = "2026北京高考政治哲学宝典---三年模拟全触发全链条"
DOC_TITLE = "2026北京高考政治哲学宝典\n三年模拟全触发全链条"
SIGNATURE = "飞哥正志讲堂"
FINAL_STEM = "2026北京高考政治哲学宝典---三年模拟全触发全链条_终极融合版"
FINAL_MD = OUT_DIR / f"{FINAL_STEM}.md"
FINAL_DOCX = OUT_DIR / f"{FINAL_STEM}.docx"
DESKTOP_DOCX = Path("/Users/wanglifei/Desktop/哲学宝典最终版 5.2双终极融合版.docx")

STUDENT_KEYS = ["材料触发点", "设问", "为什么能想到", "答案落点"]

FORBIDDEN_RE = re.compile(
    r"细则|评标|参考答案|答案写|答案核|答案/补充|可从.*角度|/Users/|"
    r"\.pdf|\.docx|\.pptx|OCR|debug|slide|line id|file id|PASS|"
    r"correct_option_chain|filled",
    re.IGNORECASE,
)

AWARENESS_MATERIAL_RE = re.compile(
    r"意识对物质|意识反作用于物质|意识[^。；;，,]*对物质[^。；;，,]*反作用|"
    r"意识[^。；;，,]*能动[^。；;，,]*反作用于物质"
)

MODULE_ALIASES = {
    "唯物论": "一、唯物论",
    "辩证唯物论": "一、唯物论",
    "辩证法": "二、辩证法",
    "唯物辩证法": "二、辩证法",
    "认识论": "三、认识论",
    "辩证唯物主义认识论": "三、认识论",
    "历史唯物主义": "四、历史唯物主义",
    "社会历史观": "四、历史唯物主义",
    "价值观": "五、价值观 / 人生观",
    "价值观 / 人生观": "五、价值观 / 人生观",
}


def ensure_dirs() -> None:
    for path in [OUT_DIR, AUDIT_DIR, REPORT_DIR]:
        path.mkdir(parents=True, exist_ok=True)


def compact(text: str) -> str:
    text = str(text or "")
    text = text.replace("（", "(").replace("）", ")")
    return re.sub(r"[\s/／、,，。；;:：·\"'“”‘’\-—_]+", "", text)


def split_node_path(path: str) -> tuple[str, str]:
    path = str(path or "").strip()
    if path.startswith("价值观 / 人生观/"):
        return "价值观 / 人生观", path.split("/", 1)[1].strip()
    parts = path.split("/", 1)
    if len(parts) == 2:
        return parts[0].strip(), parts[1].strip()
    return "", path


def node_exact_or_alias(module: str, raw_node: str) -> tuple[str | None, str, str]:
    node_norm = compact(raw_node)

    for node in rebuild.FRAMEWORK[module]:
        if raw_node == node or node_norm == compact(node):
            return node, "exact", "exact official node"

    alias_by_norm = {
        # 唯物论
        "物质决定意识": "物质决定意识",
        "一切从实际出发": "一切从实际出发 / 实事求是 / 主观与客观具体的历史的统一",
        "实事求是": "一切从实际出发 / 实事求是 / 主观与客观具体的历史的统一",
        "主观与客观具体的历史的统一": "一切从实际出发 / 实事求是 / 主观与客观具体的历史的统一",
        "主观能动性": "主观能动性 / 意识的能动作用",
        "意识的能动作用": "主观能动性 / 意识的能动作用",
        "意识能动作用": "主观能动性 / 意识的能动作用",
        "尊重客观规律与发挥主观能动性相结合": "尊重客观规律与发挥主观能动性相结合",
        "客观规律与主观能动性": "尊重客观规律与发挥主观能动性相结合",
        "规律的客观性": "规律的客观性",
        "规律的客观性与普遍性": "规律的客观性",
        # 辩证法
        "联系的观点": "联系的普遍性 / 联系的观点（总）",
        "联系观点": "联系的普遍性 / 联系的观点（总）",
        "普遍联系": "联系的普遍性 / 联系的观点（总）",
        "联系普遍性": "联系的普遍性 / 联系的观点（总）",
        "联系的普遍性": "联系的普遍性 / 联系的观点（总）",
        "联系的客观性": "联系的客观性",
        "根据固有联系建立新的具体联系": "根据固有联系建立新的具体联系",
        "固有联系建立新的具体联系": "根据固有联系建立新的具体联系",
        "联系的多样性": "联系的多样性",
        "整体与部分": "整体与部分",
        "系统观念": "系统观念 / 系统优化",
        "系统优化": "系统观念 / 系统优化",
        "系统观念系统优化": "系统观念 / 系统优化",
        "发展的观点": "发展的观点 / 发展的普遍性",
        "发展的普遍性": "发展的观点 / 发展的普遍性",
        "量变与质变": "量变与质变 / 适度原则",
        "量变质变": "量变与质变 / 适度原则",
        "适度原则": "量变与质变 / 适度原则",
        "量变与质变适度原则": "量变与质变 / 适度原则",
        "发展的观点量变与质变": "量变与质变 / 适度原则",
        "事物发展是前进性与曲折性的统一": "事物发展是前进性与曲折性的统一",
        "前进性与曲折性": "事物发展是前进性与曲折性的统一",
        "辩证否定": "辩证否定 / 守正创新",
        "辩证否定观": "辩证否定 / 守正创新",
        "守正创新": "辩证否定 / 守正创新",
        "辩证否定守正创新": "辩证否定 / 守正创新",
        "矛盾就是对立统一": "矛盾就是对立统一",
        "矛盾对立统一": "矛盾就是对立统一",
        "对立统一": "矛盾就是对立统一",
        "矛盾的普遍性": "矛盾的普遍性",
        "矛盾普遍性": "矛盾的普遍性",
        "矛盾的特殊性": "矛盾的特殊性 / 具体问题具体分析",
        "矛盾特殊性": "矛盾的特殊性 / 具体问题具体分析",
        "具体问题具体分析": "矛盾的特殊性 / 具体问题具体分析",
        "矛盾的普遍性和特殊性": "矛盾的普遍性和特殊性",
        "矛盾普遍性与特殊性": "矛盾的普遍性和特殊性",
        "矛盾普遍性特殊性": "矛盾的普遍性和特殊性",
        "两点论与重点论": "两点论与重点论",
        "内因外因": "内因与外因",
        "内因与外因": "内因与外因",
        # 认识论
        "实践与认识总": "实践与认识（总）",
        "实践与认识(总)": "实践与认识（总）",
        "实践与认识的关系": "实践与认识（总）",
        "实践与认识的辩证关系": "实践与认识（总）",
        "实践是认识的基础": "实践是认识的基础",
        "实践": "实践是认识的基础",
        "认识对实践的反作用": "认识对实践的反作用",
        "认识发展原理": "认识发展原理",
        "认识的反复性无限性": "认识发展原理",
        "认识的反复性无限性上升性": "认识发展原理",
        "真理": "真理观",
        "真理观": "真理观",
        "真理与认识的反复性条件性": "真理观",
        # 历史唯物主义
        "社会存在与社会意识": "社会存在与社会意识",
        "社会发展的两大基本规律和基本矛盾": "社会发展的两大基本规律和基本矛盾",
        "社会基本矛盾": "社会发展的两大基本规律和基本矛盾",
        "改革": "改革 / 改革的实质",
        "改革的实质": "改革 / 改革的实质",
        "人民群众": "人民群众",
        "人民群众是历史的创造者": "人民群众",
        # 价值观 / 人生观
        "价值观的导向作用": "价值观的导向作用",
        "价值观导向作用": "价值观的导向作用",
        "价值判断与价值选择": "价值判断与价值选择",
        "价值的实现与创造": "实现人生价值",
        "价值实现创造": "实现人生价值",
        "人生价值": "实现人生价值",
        "实现人生价值": "实现人生价值",
    }

    if node_norm in alias_by_norm:
        candidate = alias_by_norm[node_norm]
        if candidate in rebuild.FRAMEWORK[module]:
            return candidate, "alias", f"alias from {raw_node}"

    # Explicit combined node that the user repeatedly rejected.
    if "物质决定意识" in node_norm and ("意识能动" in node_norm or "意识对物质" in node_norm):
        return None, "reject-combined-node", "combined material/awareness node rejected"

    return None, "", ""


def infer_generic_node(module: str, raw_node: str, entry_id: str, fields: dict[str, str]) -> tuple[str | None, str, str]:
    node_norm = compact(raw_node)
    id_norm = compact(entry_id)
    text_norm = compact(" ".join([raw_node, entry_id]))
    field_text = " ".join(fields.values())

    # Prefer explicit node/id labels over free-form fields. Fields are used only
    # to break ties for generic buckets such as "联系的观点" or "矛盾分析法".
    if module == "二、辩证法":
        if "系统" in text_norm:
            return "系统观念 / 系统优化", "specific", "system label"
        if "整体" in text_norm or "部分" in text_norm:
            return "整体与部分", "specific", "whole-part label"
        if "固有联系" in text_norm or "新具体联系" in text_norm:
            return "根据固有联系建立新的具体联系", "specific", "new concrete connection label"
        if "多样" in text_norm:
            return "联系的多样性", "specific", "connection diversity label"
        if "客观性" in text_norm and "联系" in text_norm:
            return "联系的客观性", "specific", "objective connection label"
        if "量变" in text_norm or "质变" in text_norm or "适度" in text_norm:
            return "量变与质变 / 适度原则", "specific", "quantity-quality label"
        if "前进" in text_norm or "曲折" in text_norm or "新事物" in text_norm:
            return "事物发展是前进性与曲折性的统一", "specific", "forwardness/tortuosity label"
        if "辩证否定" in text_norm or "守正创新" in text_norm or "扬弃" in text_norm:
            return "辩证否定 / 守正创新", "specific", "negation/innovation label"
        if "发展" in text_norm:
            return "发展的观点 / 发展的普遍性", "specific", "development label"
        if "主次" in text_norm or "主要矛盾" in text_norm or "重点" in text_norm or "主流" in text_norm:
            return "两点论与重点论", "specific", "two-point/key-point label"
        if "普遍性与特殊性" in text_norm or "普遍性特殊性" in text_norm:
            return "矛盾的普遍性和特殊性", "specific", "universality-particularity label"
        if "特殊性" in text_norm or "具体问题" in text_norm:
            return "矛盾的特殊性 / 具体问题具体分析", "specific", "particularity label"
        if "对立统一" in text_norm:
            return "矛盾就是对立统一", "specific", "opposition-unity label"
        if "内因" in text_norm or "外因" in text_norm:
            return "内因与外因", "specific", "internal/external cause label"
        if node_norm in {"联系", "联系观点", "联系的观点", "普遍联系"}:
            return "联系的普遍性 / 联系的观点（总）", "generic-safe", "generic connection bucket"
        if node_norm in {"矛盾", "矛盾观点", "矛盾的观点", "矛盾分析法", "矛盾分析方法"}:
            if any(x in field_text for x in ["一分为二", "利弊", "两种相反", "对立统一", "既", "又"]):
                return "矛盾就是对立统一", "generic-safe", "generic contradiction with two-sided field signal"
            return None, "reject-generic-contradiction", "generic contradiction bucket lacks enough specificity"

    if module == "一、唯物论":
        if "客观规律" in text_norm and "主观能动" in text_norm:
            return "尊重客观规律与发挥主观能动性相结合", "specific", "law/initiative label"
        if "规律" in text_norm:
            return "规律的客观性", "specific", "law label"
        if "实际" in text_norm or "实事求是" in text_norm:
            return "一切从实际出发 / 实事求是 / 主观与客观具体的历史的统一", "specific", "from-reality label"
        if "主观能动" in text_norm or "意识能动" in text_norm:
            return "主观能动性 / 意识的能动作用", "specific", "initiative label"
        if "物质决定意识" in text_norm:
            return "物质决定意识", "specific", "material-determines-awareness label"

    if module == "三、认识论":
        if "实践是认识" in text_norm or "检验认识" in text_norm:
            return "实践是认识的基础", "specific", "practice-foundation label"
        if "认识对实践" in text_norm or "指导作用" in text_norm or "反作用" in text_norm:
            return "认识对实践的反作用", "specific", "knowledge-guides-practice label"
        if "反复性" in text_norm or "无限性" in text_norm or "上升性" in text_norm or "局限性" in text_norm:
            return "认识发展原理", "specific", "knowledge-development label"
        if "真理" in text_norm:
            return "真理观", "specific", "truth label"
        if "实践" in text_norm or "认识" in text_norm:
            return "实践与认识（总）", "generic-safe", "generic practice/knowledge bucket"

    if module == "四、历史唯物主义":
        if "社会存在" in text_norm or "社会意识" in text_norm:
            return "社会存在与社会意识", "specific", "social being/consciousness label"
        if "生产力" in text_norm or "生产关系" in text_norm or "经济基础" in text_norm or "上层建筑" in text_norm or "基本矛盾" in text_norm:
            return "社会发展的两大基本规律和基本矛盾", "specific", "basic contradiction label"
        if "改革" in text_norm:
            return "改革 / 改革的实质", "specific", "reform label"
        if "人民" in text_norm or "群众" in text_norm:
            return "人民群众", "specific", "people label"

    if module == "五、价值观 / 人生观":
        if "价值判断" in text_norm or "价值选择" in text_norm:
            return "价值判断与价值选择", "specific", "value judgment/choice label"
        if "人生价值" in text_norm or "实现价值" in text_norm or "价值实现" in text_norm or "劳动奉献" in text_norm:
            return "实现人生价值", "specific", "life value label"
        if "价值观" in text_norm:
            return "价值观的导向作用", "specific", "value orientation label"

    return None, "unmapped", f"unmapped node {raw_node}"


def map_claude_node(raw: dict, fields: dict[str, str]) -> tuple[str | None, str | None, str, str]:
    raw_path = raw.get("target_node_path", "")
    category, raw_node = split_node_path(raw_path)
    module = MODULE_ALIASES.get(category)
    if not module:
        return None, None, "reject-module", f"unknown category {category}"

    node, reason, note = node_exact_or_alias(module, raw_node)
    if node:
        return module, node, reason, note
    if reason.startswith("reject"):
        return None, None, reason, note
    node, reason, note = infer_generic_node(module, raw_node, raw.get("entry_id", ""), fields)
    if node:
        return module, node, reason, note
    return None, None, reason, note


def clean_student_text(value: str) -> str:
    text = rebuild.clean_text(value)
    text = re.sub(r"^答案\s*([A-D])", r"正确项是 \1", text)
    for option in "ABCD":
        text = text.replace(f"答案 {option}", f"正确项是 {option}")
        text = text.replace(f"答案{option}", f"正确项是 {option}")
    return rebuild.clean_text(text)


def awareness_material_hard_rule_reason(node: str, blob: str) -> str:
    if not AWARENESS_MATERIAL_RE.search(blob):
        return ""
    if node == "物质决定意识":
        return "material-node-mixes-awareness-reaction"
    if node != "主观能动性 / 意识的能动作用":
        return "awareness-material-reaction-outside-initiative-node"
    if "物质决定意识" in blob:
        return "initiative-node-combines-material-and-awareness-reaction"
    return ""


def remove_awareness_material_phrase(text: str) -> str:
    text = re.sub(r"[，,]\s*意识对物质具有能动的反作用", "", text)
    text = re.sub(r"意识对物质具有能动的反作用[。；;，,]?", "", text)
    text = re.sub(r"意识对物质的反作用", "人的能动引导", text)
    text = re.sub(r"物质与意识的辩证关系", "人与人工智能的相互关系", text)
    text = re.sub(r"。{2,}", "。", text)
    text = re.sub(r"；。", "。", text)
    return rebuild.clean_text(text)


def suite_from_raw(raw: dict) -> str:
    return f"{raw.get('year')}{raw.get('district')}{raw.get('stage')}"


def question_from_raw(raw: dict) -> str:
    q = str(raw.get("question_no") or "").strip()
    sub = str(raw.get("sub_part") or "").strip()
    if sub:
        q = f"{q}({sub.strip('()（）')})"
    return q


def qtype_from_raw(raw: dict) -> str:
    return "选择题" if raw.get("question_type") == "choice" else "主观题"


def entry_key(e: dict) -> tuple[str, str, str, str, str]:
    return (
        e["module"],
        e["node"],
        rebuild.normalize_suite(e["suite"]),
        rebuild.normalize_q(e["q"]),
        e["qtype"],
    )


def protected_codex(e: dict) -> bool:
    suite = e.get("suite", "")
    q = rebuild.normalize_q(e.get("q", ""))
    if e.get("origin") == "current-user-critical":
        return True
    if suite == "2026海淀一模" and q == "16":
        return True
    if suite == "2024东城一模" and q in {"18(3)", "18（3）", "18((3))"}:
        return True
    if suite == "2024朝阳二模" and q in {"16(2)", "16（2）", "16((2))"}:
        return True
    if suite == "2025海淀期末":
        return True
    return False


def is_generic_question(text: str) -> bool:
    text = str(text or "").strip()
    return len(text) < 14 or text.startswith("选择题：根据题干材料")


def quality_score(e: dict) -> int:
    fields = e.get("fields", {})
    score = 0
    for key in STUDENT_KEYS:
        value = str(fields.get(key, ""))
        score += min(len(value), 180)
    if not is_generic_question(fields.get("设问", "")):
        score += 150
    if e.get("qtype") == "主观题":
        score += 30
    if e.get("evidence_level") == "rubric-formal":
        score += 90
    if e.get("evidence_level") == "answer-key":
        score += 45
    if e.get("merge_source") == "codex-5.2":
        score += 20
    return score


def load_codex_entries() -> list[dict]:
    rows = json.loads(CODEX_ENTRIES.read_text(encoding="utf-8"))
    out = []
    for row in rows:
        e = json.loads(json.dumps(row, ensure_ascii=False))
        e["merge_source"] = "codex-5.2-final"
        e.setdefault("evidence_level", "")
        out.append(e)
    return out


def load_claude_candidates() -> tuple[list[dict], list[dict]]:
    candidates = []
    rejected = []
    for path in sorted(CLAUDE_ENTRIES_DIR.glob("*.jsonl")):
        for line_no, line in enumerate(path.read_text(encoding="utf-8").splitlines(), 1):
            if not line.strip():
                continue
            raw = json.loads(line)
            evidence = raw.get("evidence_level", "")
            if evidence == "reference-only":
                rejected.append({"entry_id": raw.get("entry_id"), "reason": "reference-only", "path": str(path), "line": line_no})
                continue
            fields = {key: clean_student_text(raw.get(key, "")) for key in STUDENT_KEYS}
            blob = "\n".join(fields.values())
            if any(not fields[key] for key in STUDENT_KEYS):
                rejected.append({"entry_id": raw.get("entry_id"), "reason": "empty-student-field", "path": str(path), "line": line_no})
                continue
            if FORBIDDEN_RE.search(blob):
                rejected.append({"entry_id": raw.get("entry_id"), "reason": "student-forbidden-hit", "hit": FORBIDDEN_RE.search(blob).group(0), "path": str(path), "line": line_no})
                continue
            module, node, map_reason, map_note = map_claude_node(raw, fields)
            if not module or not node:
                rejected.append({"entry_id": raw.get("entry_id"), "reason": map_reason, "note": map_note, "target_node_path": raw.get("target_node_path"), "path": str(path), "line": line_no})
                continue
            hard_reason = awareness_material_hard_rule_reason(node, blob)
            if hard_reason:
                rejected.append({"entry_id": raw.get("entry_id"), "reason": hard_reason, "target_node_path": raw.get("target_node_path"), "path": str(path), "line": line_no})
                continue
            suite = suite_from_raw(raw)
            if suite == "2026石景山期末":
                rejected.append({"entry_id": raw.get("entry_id"), "reason": "skill-rule-excluded-suite", "path": str(path), "line": line_no})
                continue
            q = question_from_raw(raw)
            qtype = qtype_from_raw(raw)
            candidate = {
                "module": module,
                "node": node,
                "suite": suite,
                "q": q,
                "qtype": qtype,
                "heading_base": rebuild.title_from_parts(suite, q, qtype),
                "fields": fields,
                "origin": "claudecode-final-2026-05-02",
                "merge_source": "claude-final",
                "source_entry_id": raw.get("entry_id", ""),
                "source_target_node_path": raw.get("target_node_path", ""),
                "evidence_level": evidence,
                "map_reason": map_reason,
                "map_note": map_note,
                "image": "",
            }
            candidates.append(candidate)
    return candidates, rejected


def select_best_claude_by_key(candidates: list[dict]) -> tuple[dict[tuple[str, str, str, str, str], dict], list[dict]]:
    groups = defaultdict(list)
    for candidate in candidates:
        groups[entry_key(candidate)].append(candidate)
    best = {}
    skipped = []
    for key, rows in groups.items():
        rows.sort(key=quality_score, reverse=True)
        best[key] = rows[0]
        for row in rows[1:]:
            skipped.append({"entry_id": row.get("source_entry_id"), "reason": "weaker-claude-duplicate", "kept": rows[0].get("source_entry_id"), "key": str(key)})
    return best, skipped


def merge_entries(codex_entries: list[dict], claude_candidates: list[dict]) -> tuple[list[dict], dict[str, list[dict]]]:
    claude_by_key, internal_skips = select_best_claude_by_key(claude_candidates)
    codex_by_key = {entry_key(row): row for row in codex_entries}

    merged = []
    imported = []
    replaced = []
    kept_codex = []

    all_keys = set(codex_by_key) | set(claude_by_key)
    for key in sorted(all_keys, key=str):
        cdx = codex_by_key.get(key)
        cld = claude_by_key.get(key)
        if cdx and not cld:
            merged.append(cdx)
            continue
        if cld and not cdx:
            cld["merge_action"] = "imported-from-claude-final"
            merged.append(cld)
            imported.append(cld)
            continue
        assert cdx is not None and cld is not None
        if protected_codex(cdx):
            cdx["merge_action"] = "kept-codex-protected"
            merged.append(cdx)
            kept_codex.append({"key": str(key), "claude": cld.get("source_entry_id"), "reason": "protected-hard-sample-or-user-critical"})
            continue
        can_replace = cld.get("map_reason") in {"exact", "alias", "specific"} and cld.get("evidence_level") in {"rubric-formal", "answer-key"}
        if can_replace and quality_score(cld) >= quality_score(cdx) + 110:
            replacement = json.loads(json.dumps(cld, ensure_ascii=False))
            replacement["merge_action"] = "replaced-codex-with-claude-final"
            replacement["replaced_codex_origin"] = cdx.get("origin", "")
            merged.append(replacement)
            replaced.append({"key": str(key), "claude": cld.get("source_entry_id"), "codex_origin": cdx.get("origin", ""), "score_delta": quality_score(cld) - quality_score(cdx)})
        else:
            cdx["merge_action"] = "kept-codex-over-claude"
            merged.append(cdx)
            kept_codex.append({"key": str(key), "claude": cld.get("source_entry_id"), "reason": "codex safer or quality delta insufficient", "score_delta": quality_score(cld) - quality_score(cdx)})

    merged.sort(key=rebuild.sort_key)
    return merged, {
        "imported": imported,
        "replaced": replaced,
        "kept_codex": kept_codex,
        "internal_skips": internal_skips,
    }


def apply_hard_rule_overrides(entries: list[dict]) -> list[dict]:
    fixed = []
    for row in entries:
        e = json.loads(json.dumps(row, ensure_ascii=False))
        q_norm = rebuild.normalize_q(e.get("q", ""))

        # User notebook hard sample: the AI mutual-shaping question may enter
        # multiple nodes, but this node must teach only the initiative half.
        if (
            e.get("suite") == "2024朝阳二模"
            and q_norm in {"16(2)", "16((2))"}
            and e.get("node") == "主观能动性 / 意识的能动作用"
        ):
            e["fields"]["为什么能想到"] = (
                "材料写科学家受神经系统研究启发开发人工神经网络，推动深度学习和类脑智能，"
                "并且设问追问未来如何处理人与人工智能的关系。这里学生要抓住的是：人不是被技术牵着走，"
                "而是能够在认识技术条件和发展趋势的基础上，有目的、有计划地设计、使用、规范和引导人工智能。"
                "因此这一栏只落到主观能动性 / 意识的能动作用。"
            )
            e["fields"]["答案落点"] = (
                "从主观能动性角度看，未来不能把人与人工智能的关系理解成技术单方面塑造人。"
                "人应当主动认识人工智能的发展条件、应用场景和潜在风险，通过科研创新、制度规范、伦理引导和价值选择，"
                "能动地塑造人工智能的发展方向，使人工智能更好服务人的发展。"
            )
            e["merge_action"] = (e.get("merge_action") or "base-codex") + "+hard-rule-overridden"

        if e.get("node") != "主观能动性 / 意识的能动作用":
            for key, value in list(e.get("fields", {}).items()):
                if AWARENESS_MATERIAL_RE.search(value):
                    e["fields"][key] = remove_awareness_material_phrase(value)
                    e["merge_action"] = (e.get("merge_action") or "base-codex") + "+hard-rule-scrubbed"

        fixed.append(e)
    return fixed


def write_markdown(entries: list[dict], out_path: Path) -> None:
    lines = [f"# {TITLE}", "", f"**{SIGNATURE}**", "", "## 前言", "", "（此处留给飞哥老师自行撰写。）", ""]
    current_module = current_node = None
    for e in entries:
        if e["module"] != current_module:
            current_module = e["module"]
            current_node = None
            lines.extend([f"## {current_module}", ""])
        if e["node"] != current_node:
            current_node = e["node"]
            lines.extend([f"### {current_node}", ""])
        lines.extend([f"#### {e['heading_base']}", ""])
        for key in STUDENT_KEYS:
            lines.extend([f"【{key}】 {e['fields'][key]}", ""])
    out_path.write_text("\n".join(lines), encoding="utf-8")


def qa(entries: list[dict], merge_log: dict[str, list[dict]], rejected: list[dict]) -> dict:
    body = "\n".join("\n".join(row["fields"].values()) for row in entries)
    keys = Counter(entry_key(row) for row in entries)
    counts = Counter(row.get("merge_action", "base-codex") for row in entries)
    node_counts = Counter((row["module"], row["node"]) for row in entries)
    hard = {
        "2026海淀一模_Q16": [(e["module"], e["node"], e.get("merge_source")) for e in entries if e["suite"] == "2026海淀一模" and rebuild.normalize_q(e["q"]) == "16"],
        "2024东城一模_Q18(3)": [(e["module"], e["node"], e.get("merge_source")) for e in entries if e["suite"] == "2024东城一模" and rebuild.normalize_q(e["q"]) in {"18(3)", "18((3))"}],
        "2026石景山期末": [e["heading_base"] for e in entries if e["suite"] == "2026石景山期末"],
    }
    hard_rule_hits = []
    for e in entries:
        blob = "\n".join(e.get("fields", {}).values())
        reason = awareness_material_hard_rule_reason(e.get("node", ""), blob)
        if reason:
            hard_rule_hits.append(
                {
                    "reason": reason,
                    "module": e.get("module", ""),
                    "node": e.get("node", ""),
                    "heading": e.get("heading_base", ""),
                    "merge_source": e.get("merge_source", ""),
                }
            )
    return {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "final_entries": len(entries),
        "merge_action_counts": dict(counts),
        "claude_imported": len(merge_log["imported"]),
        "claude_replaced_codex": len(merge_log["replaced"]),
        "codex_kept_over_claude": len(merge_log["kept_codex"]),
        "claude_internal_duplicate_skipped": len(merge_log["internal_skips"]),
        "claude_rejected": len(rejected),
        "forbidden_hits": sorted(set(match.group(0) for match in FORBIDDEN_RE.finditer(body))),
        "missing_student_fields": [
            e.get("heading_base", "")
            for e in entries
            if any(not e.get("fields", {}).get(key) for key in STUDENT_KEYS)
        ],
        "duplicate_keys": [str(key) for key, count in keys.items() if count > 1],
        "hard_rule_hits": hard_rule_hits,
        "hard_samples": hard,
        "node_counts": {f"{module} / {node}": count for (module, node), count in sorted(node_counts.items())},
        "rejected_reasons": dict(Counter(row.get("reason", "") for row in rejected)),
    }


def write_reports(entries: list[dict], merge_log: dict[str, list[dict]], rejected: list[dict], qa_data: dict) -> None:
    fusion_lines = [
        "# 必修四双终极版融合报告 2026-05-02",
        "",
        "## 输入版本",
        "",
        f"- Codex 5.2 整合终版：{CODEX_ENTRIES}",
        f"- ClaudeCode 从0重跑终版：{CLAUDE_ENTRIES_DIR}",
        "- 4 月 28 日桌面旧 Claude 美化版未作为证据主稿：其目录结构仍含旧合桶节点，不满足本轮小本本的新硬规则。",
        "",
        "## 融合原则",
        "",
        "- 以飞哥政治庄园/小本本框架为唯一章节结构。",
        "- 保护用户反复纠偏过的硬样本：`2026海淀一模 Q16`、`2024东城一模 Q18(3)`、`2024朝阳二模 Q16(2)` 与 `2025海淀期末`。",
        "- 主观题只吸收 `rubric-formal`；选择题允许 `answer-key`；`reference-only` 不进入新版。",
        "- 同题同节点默认保留 Codex；只有 Claude 映射明确、证据合格且学生四字段明显更完整时才替换。",
        "- 2026 石景山期末按技能硬规则整套排除，不补入主链。",
        "",
        "## 结果概览",
        "",
        f"- 最终条目：{qa_data['final_entries']}",
        f"- Claude 独有补入：{qa_data['claude_imported']}",
        f"- Claude 替换 Codex：{qa_data['claude_replaced_codex']}",
        f"- Codex 保留压过 Claude：{qa_data['codex_kept_over_claude']}",
        f"- 学生字段禁词：{qa_data['forbidden_hits']}",
        f"- 四字段缺失：{len(qa_data['missing_student_fields'])}",
        f"- 重复键：{len(qa_data['duplicate_keys'])}",
        "",
        "## Claude 补入条目",
        "",
    ]
    for e in merge_log["imported"][:180]:
        fusion_lines.append(f"- {e['source_entry_id']} -> {e['module']} / {e['node']} / {e['heading_base']} ({e['map_reason']})")
    if len(merge_log["imported"]) > 180:
        fusion_lines.append(f"- ……另有 {len(merge_log['imported']) - 180} 条，详见 `audit/claude_imported_entries.json`。")
    fusion_lines.extend(["", "## Claude 替换 Codex", ""])
    if merge_log["replaced"]:
        for row in merge_log["replaced"]:
            fusion_lines.append(f"- {row['claude']} -> {row['key']}，score_delta={row['score_delta']}")
    else:
        fusion_lines.append("- 无。")
    fusion_lines.extend(["", "## 输出文件", "", f"- Markdown：{FINAL_MD}", f"- Word：{FINAL_DOCX}", f"- 桌面 Word：{DESKTOP_DOCX}"])
    (REPORT_DIR / "必修四双终极版融合报告_2026-05-02.md").write_text("\n".join(fusion_lines), encoding="utf-8")

    accept_lines = [
        "# 必修四双终极融合版验收报告 2026-05-02",
        "",
        f"- 最终条目：{qa_data['final_entries']}",
        f"- 学生字段禁词命中：{qa_data['forbidden_hits']}",
        f"- 四字段缺失：{len(qa_data['missing_student_fields'])}",
        f"- 重复键残留：{len(qa_data['duplicate_keys'])}",
        f"- Claude 被拒条目：{qa_data['claude_rejected']}，原因：{qa_data['rejected_reasons']}",
        "",
        "## 硬样本",
        "",
    ]
    for name, rows in qa_data["hard_samples"].items():
        accept_lines.append(f"- {name}: {rows}")
    accept_lines.extend(
        [
            "",
            "## Word 验收待补",
            "",
            "- DOCX 已生成。后续 Word 打开/保存、QuickLook、PDF/PNG 渲染结果会追加在本报告下方。",
        ]
    )
    (REPORT_DIR / "必修四双终极融合版验收报告_2026-05-02.md").write_text("\n".join(accept_lines), encoding="utf-8")


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def polish_docx_cover(path: Path) -> None:
    doc = Document(path)
    title_index = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.startswith("2026北京高考政治哲学宝典"):
            title_index = idx
            break
    if title_index is None:
        doc.save(path)
        return
    for paragraph in list(doc.paragraphs[:title_index]):
        if not paragraph.text.strip():
            remove_paragraph(paragraph)
    title = doc.paragraphs[0]
    title.paragraph_format.space_before = Pt(215)
    title.paragraph_format.space_after = Pt(28)
    doc.save(path)


def main() -> None:
    ensure_dirs()
    codex_entries = load_codex_entries()
    claude_candidates, rejected = load_claude_candidates()
    merged, merge_log = merge_entries(codex_entries, claude_candidates)
    merged = apply_hard_rule_overrides(merged)
    qa_data = qa(merged, merge_log, rejected)
    if qa_data["forbidden_hits"] or qa_data["missing_student_fields"] or qa_data["duplicate_keys"] or qa_data["hard_rule_hits"]:
        raise SystemExit(json.dumps(qa_data, ensure_ascii=False, indent=2))

    (AUDIT_DIR / "final_entries.json").write_text(json.dumps(merged, ensure_ascii=False, indent=2), encoding="utf-8")
    (AUDIT_DIR / "claude_candidates.json").write_text(json.dumps(claude_candidates, ensure_ascii=False, indent=2), encoding="utf-8")
    (AUDIT_DIR / "claude_rejected_entries.json").write_text(json.dumps(rejected, ensure_ascii=False, indent=2), encoding="utf-8")
    for name, rows in merge_log.items():
        (AUDIT_DIR / f"claude_{name}.json").write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")
    (AUDIT_DIR / "final_qa.json").write_text(json.dumps(qa_data, ensure_ascii=False, indent=2), encoding="utf-8")

    write_markdown(merged, FINAL_MD)
    tmp_docx = OUT_DIR / f"{FINAL_STEM}.tmp.docx"
    rebuild.TITLE = DOC_TITLE
    rebuild.build_doc(merged, tmp_docx)
    polish_docx_cover(tmp_docx)
    os.replace(tmp_docx, FINAL_DOCX)
    shutil.copy2(FINAL_DOCX, DESKTOP_DOCX)
    write_reports(merged, merge_log, rejected, qa_data)
    print(json.dumps(qa_data, ensure_ascii=False, indent=2))
    print(f"FINAL_MD={FINAL_MD}")
    print(f"FINAL_DOCX={FINAL_DOCX}")
    print(f"DESKTOP_DOCX={DESKTOP_DOCX}")


if __name__ == "__main__":
    main()
