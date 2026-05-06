#!/usr/bin/env python3
"""Polish the accepted final DOCX without changing its teaching content.

Edits:
- Replace the manual directory with a native Word TOC field over Heading 1-2.
- Add centered page-number footer.
- Apply restrained blue hierarchy to title/headings/TOC styles.
- Normalize straight/single quotation marks in body text to Chinese double quotes.
- Split choice-question options ①②③④ onto separate lines in question paragraphs.
"""
from __future__ import annotations

import re
import shutil
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/wanglifei/Desktop/北京高考政治/必修四终极融合版_2026-05-02")
SRC = Path("/Users/wanglifei/Desktop/哲学宝典最终版 5.2双终极融合版.docx")
OUT = ROOT / "outputs" / "哲学宝典最终版 5.2双终极融合版_目录页码美化版.docx"
DESKTOP_OUT = Path("/Users/wanglifei/Desktop/哲学宝典最终版 5.2双终极融合版_目录页码美化版.docx")

BLUE_TITLE = RGBColor(0x16, 0x3E, 0x63)
BLUE_H1 = RGBColor(0x1F, 0x4E, 0x79)
BLUE_H2 = RGBColor(0x2F, 0x6F, 0x9F)
BLUE_H3 = RGBColor(0x3A, 0x62, 0x78)
BLUE_LABEL = RGBColor(0x21, 0x57, 0x4C)
MUTED = RGBColor(0x66, 0x66, 0x66)
CIRCLED = "①②③④⑤⑥⑦⑧⑨⑩"


def set_east_asian_font(run, zh: str, ascii_name: str = "Times New Roman") -> None:
    run.font.name = ascii_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), zh)
    rfonts.set(qn("w:ascii"), ascii_name)
    rfonts.set(qn("w:hAnsi"), ascii_name)


def style_font(style, zh: str, size: float, color: Optional[RGBColor], bold: bool = False) -> None:
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), zh)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def ensure_style(doc: Document, name: str):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def normalize_quotes_in_runs(paragraph: Paragraph) -> None:
    single_open = True
    double_open = True
    for run in paragraph.runs:
        out = []
        for ch in run.text:
            if ch in {"'", "‘", "’"}:
                out.append("“" if single_open else "”")
                single_open = not single_open
            elif ch == '"':
                out.append("“" if double_open else "”")
                double_open = not double_open
            else:
                out.append(ch)
        run.text = "".join(out)


def answer_choice_start(text: str, search_from: int) -> Optional[int]:
    """Return the start of the A/B/C/D answer-choice block, if present."""
    pattern = re.compile(r"(?<![A-Za-z0-9])([A-DＡ-Ｄ])([.．、])")
    for match in pattern.finditer(text, search_from):
        suffix = text[match.end() : match.end() + 16]
        if any(ch in suffix for ch in CIRCLED):
            return match.start()
    return None


def format_answer_choices_block(text: str) -> str:
    block = text.strip()
    if not block:
        return ""
    block = re.sub(r"\s*([A-DＡ-Ｄ][.．、])\s*", r"\n\1", block).strip()
    return re.sub(r"\n{2,}", "\n", block)


def format_choice_options_text(text: str) -> str:
    matches = list(re.finditer(f"[{CIRCLED}]", text))
    if len(matches) < 2:
        return text

    split_at = answer_choice_start(text, matches[min(3, len(matches) - 1)].end())
    answer_block = ""
    if split_at is not None:
        answer_block = text[split_at:]
        text = text[:split_at].rstrip()
        matches = list(re.finditer(f"[{CIRCLED}]", text))
        if len(matches) < 2:
            return text + "\n" + format_answer_choices_block(answer_block)

    prefix = text[: matches[0].start()].rstrip()
    options: list[str] = []
    for idx, match in enumerate(matches):
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        seg = text[match.start() : end].strip()
        if idx + 1 < len(matches):
            seg = re.sub(r"[；;，,、\\s]+$", "", seg) + "；"
        options.append(seg)
    formatted = prefix + "\n" + "\n".join(options)
    if answer_block:
        formatted += "\n" + format_answer_choices_block(answer_block)
    return formatted


def split_choice_question(paragraph: Paragraph) -> None:
    text = paragraph.text
    if not text.startswith("【设问】") or sum(text.count(ch) for ch in CIRCLED) < 2:
        return

    if paragraph.runs and paragraph.runs[0].text == "【设问】":
        value = "".join(run.text for run in paragraph.runs[1:])
        paragraph.runs[1].text = " " + format_choice_options_text(value.lstrip())
        for run in paragraph.runs[2:]:
            run.text = ""
        return

    # Fallback for unexpected run layout.
    for run in paragraph.runs:
        run.text = ""
    label = paragraph.add_run("【设问】")
    label.bold = True
    label.font.color.rgb = BLUE_LABEL
    set_east_asian_font(label, "黑体")
    paragraph.add_run(" " + format_choice_options_text(text[len("【设问】") :].lstrip()))


def simplify_student_terms(paragraph: Paragraph) -> None:
    replacements = {
        "中药药性、配伍和食品科学": "中药药性、药材搭配和食品科学",
        "自身药性、配伍和食用安全规律": "自身药性、药材搭配和食用安全规律",
    }
    for run in paragraph.runs:
        for old, new in replacements.items():
            if old in run.text:
                run.text = run.text.replace(old, new)


def delete_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def add_field_run(paragraph: Paragraph, instr: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)

    run = paragraph.add_run()
    instr_el = OxmlElement("w:instrText")
    instr_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_el.text = instr
    run._r.append(instr_el)

    run = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "separate")
    run._r.append(fld)

    if placeholder:
        run = paragraph.add_run(placeholder)
        set_east_asian_font(run, "宋体")

    run = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "end")
    run._r.append(fld)


def set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def clear_paragraph(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        run.text = ""


def rebuild_toc(doc: Document) -> None:
    paragraphs = doc.paragraphs
    toc_heading_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "目录")
    content_start_idx = next(
        i for i, p in enumerate(paragraphs[toc_heading_idx + 1 :], toc_heading_idx + 1)
        if p.style.name == "Heading 1" and p.text.strip().startswith("一、唯物论")
    )

    toc_heading = paragraphs[toc_heading_idx]
    toc_heading.style = doc.styles["Normal"]
    toc_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    clear_paragraph(toc_heading)
    run = toc_heading.add_run("目录")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = BLUE_TITLE
    set_east_asian_font(run, "黑体")

    for p in list(paragraphs[toc_heading_idx + 1 : content_start_idx]):
        delete_paragraph(p)

    toc_para = insert_paragraph_after(toc_heading)
    toc_para.paragraph_format.space_before = Pt(8)
    toc_para.paragraph_format.space_after = Pt(4)
    add_field_run(toc_para, r'TOC \o "1-2" \h \z \u', "（打开 Word 后会自动更新为可点击目录）")

    # Keep a page break between TOC and body if the original manual directory had one.
    break_para = insert_paragraph_after(toc_para)
    break_para.add_run().add_break(WD_BREAK.PAGE)


def add_page_numbers(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    for p in section.first_page_footer.paragraphs:
        clear_paragraph(p)

    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    clear_paragraph(p)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    left = p.add_run("— ")
    left.font.color.rgb = MUTED
    set_east_asian_font(left, "宋体")
    add_field_run(p, "PAGE")
    right = p.add_run(" —")
    right.font.color.rgb = MUTED
    set_east_asian_font(right, "宋体")


def polish_styles(doc: Document) -> None:
    style_font(doc.styles["Heading 1"], "黑体", 18, BLUE_H1, True)
    style_font(doc.styles["Heading 2"], "黑体", 14, BLUE_H2, True)
    style_font(doc.styles["Heading 3"], "楷体", 11.5, BLUE_H3, True)

    toc1 = ensure_style(doc, "TOC 1")
    toc2 = ensure_style(doc, "TOC 2")
    style_font(toc1, "黑体", 12, BLUE_H1, True)
    style_font(toc2, "宋体", 10.5, RGBColor(0x22, 0x22, 0x22), False)
    toc1.paragraph_format.space_after = Pt(4)
    toc2.paragraph_format.left_indent = Pt(18)
    toc2.paragraph_format.space_after = Pt(2)

    # Cover title and signature.
    title = doc.paragraphs[0]
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = BLUE_TITLE
        set_east_asian_font(run, "黑体")
    sign = doc.paragraphs[1]
    sign.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in sign.runs:
        run.bold = True
        run.font.size = Pt(26)
        set_east_asian_font(run, "楷体")


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SRC, OUT)
    doc = Document(str(OUT))

    polish_styles(doc)

    # Front matter should look like headings but not enter the native TOC.
    for p in doc.paragraphs:
        if p.text.strip() == "前言":
            p.style = doc.styles["Normal"]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(20)
                run.font.color.rgb = BLUE_TITLE
                set_east_asian_font(run, "黑体")
            break

    for p in doc.paragraphs:
        normalize_quotes_in_runs(p)
        simplify_student_terms(p)
        split_choice_question(p)

    rebuild_toc(doc)
    add_page_numbers(doc)
    set_update_fields_on_open(doc)

    doc.save(str(OUT))
    shutil.copy2(OUT, DESKTOP_OUT)
    print(f"wrote {OUT}")
    print(f"copied {DESKTOP_OUT}")


if __name__ == "__main__":
    main()
