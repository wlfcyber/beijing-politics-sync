from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document


RUN_DIR = Path(__file__).resolve().parents[1]
DELIVERY = RUN_DIR / "05_delivery"
REPORT = RUN_DIR / "04_fusion_audit" / "remaining_changping_life_value_patch_20260524.md"


PATCH = {
    "node": "实现人生价值",
    "suite": "2025昌平二模",
    "q": "第16题",
    "material_trigger": "材料把《祭侄文稿》中颜杲卿、颜季明为维护国家统一而牺牲生命，以及《黄州寒食帖》中苏轼在贬谪困境中仍坚守理想精神，同“书如其人、兼论其平生”的书法评价联系起来，说明作品价值不只来自笔墨技巧，还来自作者和相关人物在家国、社会、人生中的责任担当。",
    "question_prompt": "有人认为“中国书法不仅仅是把字写好的问题”，从哲学与文化角度，谈谈你的看法。",
    "why_trigger": "能想到实现人生价值，是因为材料不是单纯评价字写得好不好，而是在追问书法作品为什么有精神力量：颜真卿家族把个人生命投入维护国家统一的事业，苏轼在逆境中保持理想追求，这些都体现人的价值在于对社会的责任和贡献。学生看到“家国牺牲、困境坚守、作品感染后人”这一链条，就应当想到人生价值要在劳动和奉献中、在个人与社会的统一中、在砥砺自我中创造和实现。",
    "answer_landing": "回答时要把“中国书法不仅是写字”落到人的价值实现上：优秀书法作品之所以可贵，不只是因为技法精妙，更因为它承载了创作者和时代人物对国家民族、社会责任和理想信念的坚守。颜氏家族维护国家统一、苏轼身处困境仍追求理想，体现个人把自身生命追求同国家社会需要结合起来，在奉献和坚守中实现人生价值；这种人生价值又通过书法作品传递给后人，所以中国书法具有超越技法的精神力量。",
}


def final_docx() -> Path:
    return next(p for p in DELIVERY.glob("*.docx") if "backup" not in p.name and ".tmp" not in p.name)


def paragraph_is_node(text: str) -> bool:
    if not text or text.startswith("【"):
        return False
    if text[0].isdigit() and "." in text[:4]:
        return False
    return len(text) <= 50


def clear_and_set(paragraph, label: str, value: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    r = paragraph.add_run(f"【{label}】")
    r.bold = True
    paragraph.add_run(value)


def main() -> int:
    docx = final_docx()
    backup = docx.with_name(
        f"{docx.stem}_backup_before_remaining_changping_life_value_patch_{datetime.now():%Y%m%d_%H%M%S}.docx"
    )
    shutil.copy2(docx, backup)

    doc = Document(str(docx))
    current_node = ""
    active = False
    matched_headings: list[str] = []
    patched_fields: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        if paragraph_is_node(text):
            current_node = text

        if text[0].isdigit() and "." in text[:4]:
            active = (
                current_node == PATCH["node"]
                and PATCH["suite"] in text
                and PATCH["q"] in text
            )
            if active:
                matched_headings.append(text)
            continue

        if not active or not text.startswith("【"):
            continue

        if text.startswith("【材料触发点】"):
            clear_and_set(paragraph, "材料触发点", PATCH["material_trigger"])
            patched_fields.append("材料触发点")
        elif text.startswith("【设问】"):
            clear_and_set(paragraph, "设问", PATCH["question_prompt"])
            patched_fields.append("设问")
        elif text.startswith("【为什么能想到】"):
            clear_and_set(paragraph, "为什么能想到", PATCH["why_trigger"])
            patched_fields.append("为什么能想到")
        elif text.startswith("【答案落点】"):
            clear_and_set(paragraph, "答案落点", PATCH["answer_landing"])
            patched_fields.append("答案落点")
            active = False

    doc.save(str(docx))

    lines = [
        "# Remaining Changping life-value row patch",
        "",
        f"- docx: `{docx}`",
        f"- backup: `{backup}`",
        f"- matched headings: {len(matched_headings)}",
        f"- patched fields: {', '.join(patched_fields) or 'NONE'}",
        "",
        "## Headings",
        "",
        *[f"- {h}" for h in matched_headings],
    ]
    REPORT.write_text("\n".join(lines), encoding="utf-8")
    print(REPORT)
    print(f"matched={len(matched_headings)} fields={len(patched_fields)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
