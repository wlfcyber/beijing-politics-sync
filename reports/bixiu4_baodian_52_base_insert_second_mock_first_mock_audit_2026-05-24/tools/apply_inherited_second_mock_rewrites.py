from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document


RUN = Path(__file__).resolve().parents[1]
DELIVERY = RUN / "05_delivery"
FUSION = RUN / "04_fusion_audit"
REPORT = FUSION / "inherited_second_mock_thin_row_rewrites_20260524.md"


def final_docx() -> Path:
    return next(p for p in DELIVERY.glob("*.docx") if "backup" not in p.name and ".tmp" not in p.name)


TARGETS = [
    {
        "suite": "2025朝阳二模",
        "q": "第16题",
        "node": "内因与外因",
        "material_trigger": "材料把“从游”和“自学”并列：从游强调师友引导、学术传统和共同体熏陶，自学强调“独楫扁舟、夜观星斗”式的自主探索、独立思考和自我驱动。一个提供外部条件和方向，一个激活人成长的内在动力，二者共同作用于人才成长。",
        "why_trigger": "能想到内因与外因，是因为设问问“从游与自学在人才成长中的作用”，材料本身就在比较外部引导与内部主动性。人才成长的根本动力在自身求知欲、思考力和自我发展需要，师友引导、学术环境等外因则提供方向、资源和条件；外因必须通过人的内在主动学习才能真正发生作用。",
        "answer_landing": "人才成长既需要从游带来的师友引导、学术规范和外部资源，也需要自学激发的内在动力、自主探索和独立思考。内因是事物变化发展的根据，外因是变化发展的条件，外因通过内因起作用；因此要把从游与自学结合起来，用外部引导激发内在成长动力。",
    },
    {
        "suite": "2025朝阳二模",
        "q": "第16题",
        "node": "主观能动性 / 意识的能动作用",
        "material_trigger": "材料写自学要有“独楫扁舟、夜观星斗”的状态，强调学习者不是被动接受知识，而是带着问题意识、探索精神和独立思考能力主动求知、主动选择路径、主动解决问题。",
        "why_trigger": "能想到主观能动性，是因为自学的关键不在外部灌输，而在学习者自觉确定目标、主动搜寻资源、独立思考并持续实践。材料中的自主探索精神和独立思考能力，正体现意识活动的目的性、自觉选择性和能动创造性。",
        "answer_landing": "自学能够推动人才成长，是因为它能激发人的求知欲、判断力和创造力，使学习者发挥主观能动性主动发现问题、分析问题、解决问题。人才培养既要重视外部引导，更要保护和激发学生自我驱动、独立思考和持续探索的能力。",
    },
    {
        "suite": "2025朝阳二模",
        "q": "第22题",
        "node": "主观能动性 / 意识的能动作用",
        "material_trigger": "材料把推进中国式现代化放在“顶层设计”和“实践探索”的关系中理解：顶层设计把握方向、统领目标、明确路径，实践探索把设计放到具体领域、具体阶段中检验和推进。",
        "why_trigger": "能想到主观能动性，是因为顶层设计不是现实自动生成的结果，而是党和国家对现代化目标、步骤、重点和路径作出的自觉谋划。科学的认识、规划和理论能够反作用于实践，为实践探索提供方向、目标和方法，增强实践活动的自觉性和计划性。",
        "answer_landing": "推进中国式现代化既要重视实践探索，也要发挥科学顶层设计的指导作用。正确意识对实践具有能动的反作用，科学规划能够统一目标、明确方向、协调步骤，使各阶段各领域实践更有目的性、计划性和整体性，从而把中国式现代化不断推向前进。",
    },
    {
        "suite": "2025朝阳二模",
        "q": "第22题",
        "node": "矛盾的普遍性和特殊性",
        "material_trigger": "材料讲的是推进中国式现代化：现代化是人类社会发展的共同课题，但中国式现代化又立足中国国情、人口规模、发展道路、制度优势和实践探索，体现“共同现代化规律”与“中国自身特点”的统一。",
        "why_trigger": "能想到矛盾普遍性和特殊性，是因为设问不是泛谈现代化，而是要求理解“中国式现代化”中的顶层设计与实践探索。现代化有普遍规律和共同经验，中国式现代化又有本国历史、制度和实践条件形成的特殊路径；材料正是在说明共性与个性、普遍规律与具体国情的结合。",
        "answer_landing": "推进中国式现代化要坚持矛盾普遍性与特殊性、共性与个性的具体的历史的统一。既要把握现代化建设的一般规律，借鉴人类文明有益成果，又要立足中国国情和中国实践进行顶层设计与探索创新，走符合自身实际的现代化道路。",
    },
    {
        "suite": "2024东城二模",
        "q": "第21题",
        "node": "规律的客观性",
        "material_trigger": "材料列举战略性有利条件，包括党的领导、制度优势、经济基础、社会稳定、文化自信等，并要求分析这些条件在民族复兴历史进程中如何发挥作用。它不是简单罗列优势，而是追问这些现实条件怎样顺应发展规律、转化为复兴动力。",
        "why_trigger": "能想到规律的客观性，是因为战略性有利条件要发挥作用，不能靠主观愿望或口号堆叠，而要符合社会发展规律、经济建设规律和民族复兴的历史进程。只有把现实优势放到规律性进程中加以运用，才能让条件转化为推动发展的现实力量。",
        "answer_landing": "战略性有利条件要真正发挥作用，必须尊重和遵循客观规律，把党的领导、制度优势、经济基础、社会活力和文化自信等条件转化为推动生产力发展、完善制度机制、凝聚人民力量的现实动力。这样才能使民族复兴进程按客观规律稳步推进，而不是停留在主观愿望上。",
    },
    {
        "suite": "2025昌平二模",
        "q": "第16题",
        "node": "价值判断与价值选择",
        "material_trigger": "材料把颜真卿家人为维护国家统一献出生命、苏轼在贬谪困境中坚守理想精神，与书法作品的精神气象联系起来，说明艺术作品背后不是单纯技巧，而是创作者和时代人物的价值取向、人生选择和精神坚守。",
        "why_trigger": "能想到价值判断与价值选择，是因为材料要求评价作品背后的精神价值：颜真卿家人的选择体现国家统一高于个人安危，苏轼的坚守体现理想信念高于一时困境。人们在不同处境中作出的价值判断，会引导其价值选择，并通过行动和作品表现出来。",
        "answer_landing": "评价这类书法作品，不能只看笔法技巧，还要看到其中承载的价值判断和价值选择。正确的价值判断和价值选择应自觉站在最广大人民立场上，遵循社会发展规律，把个人命运、艺术创造同国家民族、社会责任和理想信念结合起来，因此这些作品才能具有穿越时代的精神力量。",
    },
]


def clear_and_set(paragraph, label: str, value: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    r = paragraph.add_run(f"【{label}】")
    r.bold = True
    paragraph.add_run(value)


def paragraph_is_node(text: str) -> bool:
    if not text or text.startswith("【"):
        return False
    if text[0].isdigit() and "." in text[:4]:
        return False
    return len(text) <= 50


def main() -> int:
    docx = final_docx()
    backup = docx.with_name(f"{docx.stem}_backup_before_inherited_second_mock_rewrites_{datetime.now():%Y%m%d_%H%M%S}.docx")
    shutil.copy2(docx, backup)
    doc = Document(str(docx))

    current_node = ""
    active: dict | None = None
    applied: list[str] = []
    target_seen = {id(t): set() for t in TARGETS}

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        if "2024顺义思政二模" in text:
            for run in paragraph.runs:
                if "2024顺义思政二模" in run.text:
                    run.text = run.text.replace("2024顺义思政二模", "2024顺义二模")
            text = paragraph.text.strip()

        if paragraph_is_node(text):
            current_node = text

        if text[0].isdigit() and "." in text[:4]:
            active = None
            for target in TARGETS:
                if target["suite"] in text and target["q"] in text and current_node == target["node"]:
                    active = target
                    applied.append(f"- matched {target['suite']} {target['q']} / {target['node']}: {text}")
                    break
            continue

        if active is None or not text.startswith("【"):
            continue
        if text.startswith("【材料触发点】"):
            clear_and_set(paragraph, "材料触发点", active["material_trigger"])
            target_seen[id(active)].add("material_trigger")
        elif text.startswith("【为什么能想到】"):
            clear_and_set(paragraph, "为什么能想到", active["why_trigger"])
            target_seen[id(active)].add("why_trigger")
        elif text.startswith("【答案落点】"):
            clear_and_set(paragraph, "答案落点", active["answer_landing"])
            target_seen[id(active)].add("answer_landing")

    doc.save(str(docx))

    lines = [
        "# Inherited second-mock thin-row rewrites",
        "",
        f"- docx: `{docx}`",
        f"- backup: `{backup}`",
        "",
        "## Matched Entries",
        "",
        *applied,
        "",
        "## Field Coverage",
        "",
    ]
    for target in TARGETS:
        seen = sorted(target_seen[id(target)])
        lines.append(f"- {target['suite']} {target['q']} / {target['node']}: {', '.join(seen) or 'NO_FIELDS_PATCHED'}")
    REPORT.write_text("\n".join(lines), encoding="utf-8")

    print(REPORT)
    print(f"targets={len(TARGETS)} matched={len(applied)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
