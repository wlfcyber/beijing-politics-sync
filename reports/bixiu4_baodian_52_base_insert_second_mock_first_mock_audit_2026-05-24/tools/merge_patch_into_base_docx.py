# -*- coding: utf-8 -*-
from __future__ import annotations

import json
import re
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


RUN = Path(__file__).resolve().parents[1]
REPO = RUN.parents[1]
BASE_DIR = REPO / "reports" / "bixiu4_philosophy_accepted_base_second_mock_patch_2026-05-23" / "accepted_base"
FUSION = RUN / "04_fusion_audit"
DELIVERY = RUN / "05_delivery"
DELIVERY.mkdir(parents=True, exist_ok=True)

OUT_DOCX = DELIVERY / "哲学宝典最终版-飞哥正志讲堂_2026二模与一模漏项补强版_2026-05-24.docx"
INSERT_LEDGER = DELIVERY / "docx_insert_ledger.csv"


def read_jsonl(path: Path) -> list[dict]:
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def qnorm(q: str) -> str:
    m = re.search(r"(\d+(?:[（(]\d+[）)])?)", str(q or ""))
    return m.group(1) if m else str(q or "").strip()


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def style_name(para: Paragraph) -> str:
    return para.style.name if para.style is not None else ""


def find_node_ranges(doc: Document) -> dict[str, tuple[int, int]]:
    headings = []
    for i, para in enumerate(doc.paragraphs):
        if style_name(para) in ("Heading 1", "Heading 2"):
            headings.append((i, style_name(para), para.text.strip()))

    ranges: dict[str, tuple[int, int]] = {}
    for pos, (idx, style, text) in enumerate(headings):
        if style != "Heading 2":
            continue
        end = len(doc.paragraphs)
        for next_idx, _, _ in headings[pos + 1 :]:
            end = next_idx
            break
        ranges[text] = (idx, end)
    return ranges


def max_existing_number(doc: Document, start: int, end: int) -> int:
    max_no = 0
    for para in doc.paragraphs[start + 1 : end]:
        if style_name(para) != "Heading 3":
            continue
        m = re.match(r"\s*(\d+)\.", para.text.strip())
        if m:
            max_no = max(max_no, int(m.group(1)))
    return max_no


def make_heading(row: dict, number: int) -> str:
    q = qnorm(row.get("question_no", ""))
    qtext = f"第{q}题" if q else str(row.get("question_no", "")).strip()
    return f"{number}. {row['source_suite']} {qtext}（主观题）"


def make_body(row: dict) -> list[str]:
    return [
        f"【材料触发点】 {row['material_trigger']}",
        f"【设问】 {row['question_prompt']}",
        f"【为什么能想到】 {row['why_trigger']}",
        f"【答案落点】 {row['answer_landing']}",
    ]


def extract_existing_entry(
    doc: Document, suite: str, question_marker: str, section_name: str = "两点论与重点论"
) -> list[str]:
    ranges = find_node_ranges(doc)
    search_start, search_end = ranges.get(section_name, (0, len(doc.paragraphs)))
    for i, para in enumerate(doc.paragraphs[search_start:search_end], search_start):
        if style_name(para) != "Heading 3":
            continue
        text = para.text.strip()
        if suite not in text or question_marker not in text:
            continue
        body = []
        for nxt in doc.paragraphs[i + 1 :]:
            if style_name(nxt).startswith("Heading"):
                break
            if nxt.text.strip():
                body.append(nxt.text.strip())
        return body
    raise RuntimeError(f"Cannot find existing entry: {suite} {question_marker}")


def ensure_dialectic_subnodes(doc: Document) -> None:
    ranges = find_node_ranges(doc)
    if "主要矛盾和次要矛盾" in ranges and "矛盾的主要方面和次要方面" in ranges:
        return
    if "两点论与重点论" not in ranges:
        raise RuntimeError("Cannot place dialectic subnodes: missing 两点论与重点论")

    two_point_idx, _ = ranges["两点论与重点论"]
    anchor = doc.paragraphs[two_point_idx - 1]

    if "主要矛盾和次要矛盾" not in ranges:
        section_entries = [
            ("2025东城一模", "第18题第（1）问"),
            ("2026通州期末", "第21题"),
        ]
        anchor = insert_paragraph_after(anchor, "主要矛盾和次要矛盾", "Heading 2")
        for number, (suite, marker) in enumerate(section_entries, 1):
            anchor = insert_paragraph_after(anchor, f"{number}. {suite} {marker}（主观题）", "Heading 3")
            for body_line in extract_existing_entry(doc, suite, marker):
                anchor = insert_paragraph_after(anchor, body_line, "Normal")
            anchor = insert_paragraph_after(anchor, "", "Normal")

    if "矛盾的主要方面和次要方面" not in ranges:
        section_entries = [
            ("2026东城一模", "第16题"),
            ("2026丰台一模", "第16题"),
            ("2026延庆一模", "第20题"),
            ("2026顺义一模", "第21题"),
        ]
        anchor = insert_paragraph_after(anchor, "矛盾的主要方面和次要方面", "Heading 2")
        for number, (suite, marker) in enumerate(section_entries, 1):
            anchor = insert_paragraph_after(anchor, f"{number}. {suite} {marker}（主观题）", "Heading 3")
            for body_line in extract_existing_entry(doc, suite, marker):
                anchor = insert_paragraph_after(anchor, body_line, "Normal")
            anchor = insert_paragraph_after(anchor, "", "Normal")


def main() -> int:
    base_docx = next(BASE_DIR.glob("*.docx"))
    rows = read_jsonl(FUSION / "student_patch_entries.accepted.jsonl")
    grouped: dict[str, list[dict]] = defaultdict(list)
    for row in rows:
        grouped[row["canonical_node"]].append(row)

    doc = Document(base_docx)
    ensure_dialectic_subnodes(doc)
    ranges = find_node_ranges(doc)
    missing = sorted(set(grouped) - set(ranges))
    if missing:
        raise RuntimeError(f"Missing target headings in base docx: {missing}")

    ledger_lines = ["canonical_node,source_suite,question_no,inserted_heading"]

    for node_name in sorted(grouped.keys(), key=lambda n: ranges[n][0], reverse=True):
        start, end = ranges[node_name]
        anchor = doc.paragraphs[end - 1]
        next_number = max_existing_number(doc, start, end) + 1
        for row in grouped[node_name]:
            anchor = insert_paragraph_after(anchor, make_heading(row, next_number), "Heading 3")
            for body_line in make_body(row):
                anchor = insert_paragraph_after(anchor, body_line, "Normal")
            anchor = insert_paragraph_after(anchor, "", "Normal")
            ledger_lines.append(
                f"{node_name},{row['source_suite']},{row['question_no']},{make_heading(row, next_number)}"
            )
            next_number += 1

    doc.save(OUT_DOCX)
    INSERT_LEDGER.write_text("\n".join(ledger_lines), encoding="utf-8")
    print(OUT_DOCX)
    print(INSERT_LEDGER)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
