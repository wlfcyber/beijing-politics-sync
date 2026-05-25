# -*- coding: utf-8 -*-
from __future__ import annotations

import csv
import json
import shutil
import textwrap
import zipfile
from collections import Counter, defaultdict
from pathlib import Path

import fitz
from docx import Document


RUN = Path(__file__).resolve().parents[1]
FUSION = RUN / "04_fusion_audit"
DELIVERY = RUN / "05_delivery"
GOV = RUN / "06_governor_confucius"
CODEX = RUN / "02_codex_lane"
CLAUDECODE = RUN / "03_claudecode_lane"
SOURCE = RUN / "01_source_inventory"
LOGS = RUN / "99_logs"
OUT = RUN / "08_external_review"
UPLOAD = OUT / "upload_packet"
OUT.mkdir(parents=True, exist_ok=True)
UPLOAD.mkdir(parents=True, exist_ok=True)


def read_jsonl(path: Path) -> list[dict]:
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def copy_upload_files() -> None:
    files = [
        next(DELIVERY.glob("*.docx")),
        next(DELIVERY.glob("*.pdf")),
        DELIVERY / "docx_insert_ledger.csv",
        FUSION / "student_patch_entries.accepted.jsonl",
        FUSION / "student_patch_entries.accepted.before_depth_rewrite_20260524.jsonl",
        FUSION / "student_patch_entries.blocked.jsonl",
        FUSION / "depth_rewrite_codex_governor_20260524.md",
        FUSION / "claudecode_after_depth_fix_applied_20260524.md",
        FUSION / "system_optimization_language_polish_20260524.md",
        FUSION / "first_mock_2024_queue_resolution.md",
        FUSION / "first_mock_2024_queue_resolution.csv",
        FUSION / "prompt_gate_resolution.md",
        FUSION / "prompt_gate_resolution.csv",
        FUSION / "tongzhou_2026_first_mock_source_repair.md",
        FUSION / "tongzhou_2026_first_mock_source_repair.csv",
        FUSION / "weak_gate_source_repair_resolution.md",
        FUSION / "weak_gate_source_repair_resolution.csv",
        FUSION / "render_qa_content_correction_2026_fangshan_q18.md",
        FUSION / "dual_lane_hold_adjudication_20260524.md",
        FUSION / "dual_lane_hold_adjudication_20260524.csv",
        CODEX / "agents" / "codex_a_independent_coverage_rerun_20260524.md",
        CODEX / "agents" / "codex_a_independent_coverage_rerun_20260524.csv",
        CLAUDECODE / "claudecode_b_full_coverage_rerun_20260524.md",
        CLAUDECODE / "claudecode_b_full_coverage_rerun_20260524.csv",
        CLAUDECODE / "claudecode_b_full_coverage_insert_candidates_20260524.jsonl",
        CLAUDECODE / "CLAUDECODE_DEPTH_REVIEW_PROMPT_20260524.md",
        CLAUDECODE / "claudecode_depth_review_20260524.md",
        CLAUDECODE / "claudecode_depth_rewrite_proposals_20260524.jsonl",
        CLAUDECODE / "CLAUDECODE_AFTER_DEPTH_FINAL_AUDIT_PROMPT_20260524.md",
        CLAUDECODE / "claudecode_after_depth_final_audit_20260524.md",
        CLAUDECODE / "claudecode_after_depth_final_findings_20260524.jsonl",
        CLAUDECODE / "CLAUDECODE_AFTER_FIX_RECHECK_PROMPT_20260524.md",
        CLAUDECODE / "claudecode_after_fix_recheck_20260524.md",
        CLAUDECODE / "claudecode_after_fix_recheck_findings_20260524.jsonl",
        GOV / "LOCAL_GOVERNOR_CONFUCIUS_CHECK.md",
        GOV / "FINAL_LOCAL_ACCEPTANCE_AUDIT.md",
        GOV / "CURRENT_ACCEPTANCE_STATUS_20260524.md",
        GOV / "COVERAGE_CLOSURE_MATRIX_V2.md",
        GOV / "COVERAGE_CLOSURE_MATRIX_V2.csv",
        GOV / "FULL_SOURCE_VS_DOCX_COVERAGE_AUDIT_20260524.md",
        GOV / "FULL_SOURCE_VS_DOCX_COVERAGE_AUDIT_20260524.csv",
        GOV / "INHERITED_2024_2025_SECOND_MOCK_ROW_EXTRACT_20260524.md",
        GOV / "INHERITED_2024_2025_SECOND_MOCK_ROW_EXTRACT_20260524.csv",
        CLAUDECODE / "CLAUDECODE_FULL_SOURCE_DOCX_COVERAGE_AUDIT_PROMPT_20260524.md",
        CLAUDECODE / "claudecode_full_source_docx_coverage_audit_20260524.md",
        CLAUDECODE / "claudecode_full_source_docx_coverage_findings_20260524.jsonl",
        CLAUDECODE / "CLAUDECODE_INHERITED_2024_2025_SECOND_MOCK_ROW_AUDIT_PROMPT_20260524.md",
        CLAUDECODE / "claudecode_inherited_2024_2025_second_mock_row_audit_20260524.md",
        CLAUDECODE / "claudecode_inherited_2024_2025_second_mock_findings_20260524.jsonl",
        CLAUDECODE / "CLAUDECODE_RECHECK_INHERITED_PATCH_AND_HAIDIAN_EVIDENCE_PROMPT_20260524.md",
        CLAUDECODE / "claudecode_recheck_inherited_patch_and_haidian_evidence_20260524.md",
        CLAUDECODE / "claudecode_recheck_inherited_patch_and_haidian_findings_20260524.jsonl",
        CLAUDECODE / "PROMPT_VERIFY_HAIDIAN_CHOICE_ADDENDUM_20260524.md",
        CLAUDECODE / "claudecode_verify_haidian_choice_addendum_20260524.md",
        FUSION / "inherited_second_mock_thin_row_rewrites_20260524.md",
        FUSION / "remaining_changping_life_value_patch_20260524.md",
        FUSION / "haidian_2025_second_mock_evidence_closeout_20260524.md",
        FUSION / "haidian_2025_second_mock_choice_source_addendum_20260524.md",
        FUSION / "gptpro_web_scoped_fixes_applied_20260524.md",
        FUSION / "weak_evidence_cards_20260524.md",
        OUT / "gptpro_web_scoped_audit_20260524.md",
        OUT / "GPTPRO_WEB_INLINE_REVIEW_PAYLOAD_20260524.md",
        CLAUDECODE / "PROMPT_VERIFY_GPTPRO_WEB_FIXES_20260524.md",
        CLAUDECODE / "claudecode_verify_gptpro_web_fixes_20260524.md",
        CLAUDECODE / "PROMPT_VERIFY_BATCH03_CLEANUP_20260524.md",
        CLAUDECODE / "claudecode_verify_batch03_cleanup_20260524.md",
    ]
    for p in files:
        if p.exists():
            shutil.copy2(p, UPLOAD / p.name)


def copy_source_evidence() -> None:
    root = UPLOAD / "source_evidence"
    root.mkdir(parents=True, exist_ok=True)

    source_files = [SOURCE / "second_mock_candidate_entries.csv"]
    for name in [
        "2026东城二模.md",
        "2026朝阳二模.md",
        "2026丰台二模.md",
        "2026房山二模.md",
        "2026海淀二模.md",
        "2026海淀二模_Q16_readable_evidence.md",
        "2026石景山二模.md",
        "2026西城二模.md",
        "2026顺义二模.md",
        "2026通州一模.md",
        "2025海淀一模.md",
    ]:
        source_files.append(SOURCE / "suite_source_bundles" / name)

    source_files.extend(
        [
            LOGS / "weak_gate_sources" / "26顺义二模评标.txt",
            LOGS / "weak_gate_sources" / "石景山区高三政治第二次模拟考试答案评分细则(1).txt",
        ]
    )

    for p in source_files:
        if not p.exists():
            continue
        if SOURCE in p.parents:
            rel = p.relative_to(SOURCE)
        elif LOGS in p.parents:
            rel = Path("99_logs") / p.relative_to(LOGS)
        else:
            rel = Path(p.name)
        dest = root / rel
        dest.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(p, dest)

    for d in [
        LOGS / "weak_gate_sources" / "renders" / "haidian_lecture",
        LOGS / "weak_gate_sources" / "renders" / "xicheng_rubric",
        LOGS / "tongzhou_paper_pages",
    ]:
        if not d.exists():
            continue
        dest = root / "99_logs" / d.relative_to(LOGS)
        shutil.copytree(d, dest, dirs_exist_ok=True)


def source_evidence_index() -> str:
    root = UPLOAD / "source_evidence"
    if not root.exists():
        return ""
    files = sorted(p.relative_to(UPLOAD).as_posix() for p in root.rglob("*") if p.is_file())
    lines = ["", "## Source Evidence", ""]
    lines.extend(f"- {name}" for name in files)
    return "\n".join(lines)


def final_doc_inserted_excerpt() -> str:
    docx = next(DELIVERY.glob("*.docx"))
    doc = Document(docx)
    inserted_headings = []
    with (DELIVERY / "docx_insert_ledger.csv").open("r", encoding="utf-8", newline="") as f:
        for row in csv.DictReader(f):
            inserted_headings.append(row["inserted_heading"])

    chunks = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text not in inserted_headings:
            continue
        body = [text]
        for nxt in doc.paragraphs[i + 1 :]:
            ntext = nxt.text.strip()
            if nxt.style.name.startswith("Heading") and ntext:
                break
            if ntext:
                body.append(ntext)
        chunks.append("\n".join(body))
    return "\n\n---\n\n".join(chunks)


def pdf_pages_excerpt(pages: list[int]) -> str:
    pdf = fitz.open(next(DELIVERY.glob("*.pdf")))
    chunks = []
    for page_no in pages:
        if page_no < 1 or page_no > pdf.page_count:
            continue
        text = pdf[page_no - 1].get_text().strip()
        chunks.append(f"## PDF page {page_no}\n\n{text}")
    pdf.close()
    return "\n\n---\n\n".join(chunks)


def summary_tables() -> str:
    acc = read_jsonl(FUSION / "student_patch_entries.accepted.jsonl")
    blk = read_jsonl(FUSION / "student_patch_entries.blocked.jsonl")
    by_node = Counter(r["canonical_node"] for r in acc)
    by_source = Counter(r["source_suite"] for r in acc)
    by_block = Counter(r.get("block_reason", "") for r in blk)
    lines = [
        "# 外审摘要",
        "",
        f"- final_docx: {next(DELIVERY.glob('*.docx')).name}",
        f"- final_pdf: {next(DELIVERY.glob('*.pdf')).name}",
        f"- accepted_insertions: {len(acc)}",
        f"- blocked_or_skipped: {len(blk)}",
        "",
        "## Accepted By Node",
    ]
    for k, v in by_node.most_common():
        lines.append(f"- {k}: {v}")
    lines.extend(["", "## Accepted By Source"])
    for k, v in by_source.most_common():
        lines.append(f"- {k}: {v}")
    lines.extend(["", "## Blocked By Reason"])
    for k, v in by_block.most_common():
        lines.append(f"- {k}: {v}")
    return "\n".join(lines)


def make_prompt(model_name: str) -> str:
    return textwrap.dedent(
        f"""
        你现在是北京高考政治《哲学与文化》宝典的严格外审员，模型身份：{model_name}。

        审核对象：本包内的最新版《哲学宝典最终版-飞哥正志讲堂_2026二模与一模漏项补强版_2026-05-24》DOCX/PDF，以及随包的插入台账、accepted/blocked jsonl、audit summary。

        审核目标：
        1. 检查新增 2026 二模和一模漏项是否真正按原宝典风格插入原节点，而不是另起专题。
        2. 检查每条新增学生版正文是否满足“材料触发点 -> 设问 -> 为什么能想到 -> 答案落点”，且答案落点是可写进卷面的具体句子，不是元话术。
        3. 重点查错位：辩证否定、量变质变、主要矛盾/次要矛盾、矛盾主要方面/次要方面、两点论重点论、系统优化、价值观导向。
        4. 检查是否仍有“虚空构造”：材料触发点与原理之间没有真实逻辑、把文化题硬塞哲学、把参考答案说成评分细则、用“等角度”扩张原理。
        5. 检查厚度：新加 2026 二模条目是否和原宝典同强度，不得明显短、薄、空。
        6. 检查覆盖：2024-2026 一模原宝典是否明显仍漏哲学题；如只能提出疑点，请标 NEED_EVIDENCE。

        输出格式必须是表格，列为：
        - severity: DELETE / REWRITE / NEED_EVIDENCE / PASS
        - location: 文件或宝典节点 + 套卷题号
        - problem: 问题是什么
        - source_check: 你依据包内什么材料判断
        - fix: 具体怎么改，若 PASS 写“保留”

        严格规则：
        - 不要泛泛表扬。
        - 不要只说“建议优化”，必须给可执行修改。
        - 如果无法核验原始细则，不得宣称它错误；标 NEED_EVIDENCE。
        - 若发现新增条目只是模板套话或原理错位，直接列 REWRITE 或 DELETE。
        """
    ).strip()


def main() -> int:
    if UPLOAD.exists():
        shutil.rmtree(UPLOAD)
    UPLOAD.mkdir(parents=True, exist_ok=True)
    copy_upload_files()
    copy_source_evidence()

    (OUT / "REVIEW_PACKET_INDEX.md").write_text(
        "\n".join(
            [
                "# 外部审核包索引",
                "",
                "## 上传文件",
                "- upload_packet/哲学宝典最终版-飞哥正志讲堂_2026二模与一模漏项补强版_2026-05-24.docx",
                "- upload_packet/哲学宝典最终版-飞哥正志讲堂_2026二模与一模漏项补强版_2026-05-24.pdf",
                "- upload_packet/docx_insert_ledger.csv",
                "- upload_packet/student_patch_entries.accepted.jsonl",
                "- upload_packet/student_patch_entries.blocked.jsonl",
                "- upload_packet/gptpro_web_scoped_audit_20260524.md",
                "- upload_packet/gptpro_web_scoped_fixes_applied_20260524.md",
                "- upload_packet/weak_evidence_cards_20260524.md",
                "- upload_packet/first_mock_2024_queue_resolution.md",
                "- upload_packet/first_mock_2024_queue_resolution.csv",
                "- upload_packet/prompt_gate_resolution.md",
                "- upload_packet/prompt_gate_resolution.csv",
                "- upload_packet/tongzhou_2026_first_mock_source_repair.md",
                "- upload_packet/tongzhou_2026_first_mock_source_repair.csv",
                "- upload_packet/weak_gate_source_repair_resolution.md",
                "- upload_packet/weak_gate_source_repair_resolution.csv",
                "- upload_packet/render_qa_content_correction_2026_fangshan_q18.md",
                "- upload_packet/dual_lane_hold_adjudication_20260524.md",
                "- upload_packet/dual_lane_hold_adjudication_20260524.csv",
                "- upload_packet/codex_a_independent_coverage_rerun_20260524.md",
                "- upload_packet/codex_a_independent_coverage_rerun_20260524.csv",
                "- upload_packet/claudecode_b_full_coverage_rerun_20260524.md",
                "- upload_packet/claudecode_b_full_coverage_rerun_20260524.csv",
                "- upload_packet/claudecode_b_full_coverage_insert_candidates_20260524.jsonl",
                "- upload_packet/LOCAL_GOVERNOR_CONFUCIUS_CHECK.md",
                "- upload_packet/FINAL_LOCAL_ACCEPTANCE_AUDIT.md",
                "- upload_packet/COVERAGE_CLOSURE_MATRIX_V2.md",
                "- upload_packet/COVERAGE_CLOSURE_MATRIX_V2.csv",
                "- upload_packet/CURRENT_ACCEPTANCE_STATUS_20260524.md",
                "",
                "## Prompt",
                "- GPTPRO_WEB_REVIEW_PROMPT.md",
                "- CLAUDE_EXTERNAL_REVIEW_PROMPT.md",
                "",
                "## 分批文本",
                "- batch_01_all_inserted_entries.md",
                "- batch_02_high_risk_render_pages.md",
                "- batch_03_summary_and_gate.md",
                source_evidence_index(),
            ]
        ),
        encoding="utf-8",
    )
    (OUT / "GPTPRO_WEB_REVIEW_PROMPT.md").write_text(make_prompt("GPTPro 网页版"), encoding="utf-8")
    (OUT / "CLAUDE_EXTERNAL_REVIEW_PROMPT.md").write_text(make_prompt("Claude 外部审核"), encoding="utf-8")
    (OUT / "batch_01_all_inserted_entries.md").write_text(
        "# 全部新增插入条目正文抽取\n\n" + final_doc_inserted_excerpt(),
        encoding="utf-8",
    )
    (OUT / "batch_02_high_risk_render_pages.md").write_text(
        "# 高风险渲染页文本抽取\n\n"
        + pdf_pages_excerpt([109, 121, 133, 145, 157, 193, 205, 217]),
        encoding="utf-8",
    )
    (OUT / "batch_03_summary_and_gate.md").write_text(
        summary_tables()
        + "\n\n---\n\n"
        + "# Current Acceptance Status\n\n"
        + (GOV / "CURRENT_ACCEPTANCE_STATUS_20260524.md").read_text(encoding="utf-8")
        + "\n\n---\n\n"
        + "# GPTPro Web Scoped Audit\n\n"
        + "The full GPTPro scoped audit is attached separately as `gptpro_web_scoped_audit_20260524.md`. Its DELETE/REWRITE/NEED_EVIDENCE findings have been routed through local source verification and applied in `gptpro_web_scoped_fixes_applied_20260524.md`. This summary file intentionally does not inline the old GPTPro report table, because that table quotes superseded historical counts.\n"
        + "\n\n---\n\n"
        + "# GPTPro Fixes Applied\n\n"
        + (FUSION / "gptpro_web_scoped_fixes_applied_20260524.md").read_text(encoding="utf-8")
        + "\n\n---\n\n"
        + "# Weak Evidence Cards\n\n"
        + (FUSION / "weak_evidence_cards_20260524.md").read_text(encoding="utf-8")
        + "\n\n---\n\n"
        + (GOV / "COVERAGE_CLOSURE_MATRIX_V2.md").read_text(encoding="utf-8")
        + "\n\n---\n\n"
        + (GOV / "FULL_SOURCE_VS_DOCX_COVERAGE_AUDIT_20260524.md").read_text(encoding="utf-8")
        + "\n\n---\n\n# Historical Worker Files Boundary\n\n"
        + "The upload packet may include older lane reports for traceability. They are superseded by `CURRENT_ACCEPTANCE_STATUS_20260524.md` and `gptpro_web_scoped_fixes_applied_20260524.md` wherever counts or kept/deleted rows differ. The current active counts are 36 accepted rows, 36 ledger rows, and 236 PDF pages.\n"
        + "\n\n---\n\n# Source Evidence Index\n\n"
        + source_evidence_index(),
        encoding="utf-8",
    )

    zip_path = OUT / "feige_bixiu4_philosophy_external_review_packet_2026-05-24.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for p in OUT.rglob("*"):
            if p == zip_path or p.is_dir():
                continue
            zf.write(p, p.relative_to(OUT))
    print(zip_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
