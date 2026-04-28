from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path


ROOT = Path(r"C:\bp_sync_visible\reports\final_56_v8_decode_rebuilt_2026-04-29")
OUT = ROOT / "outputs"
AUDIT = ROOT / "audit"

OLD = Path(r"C:\Users\Administrator\Desktop\4.29凌晨跑完的结果v6\01_学生版Word")
OLD_AUDIT = Path(r"C:\Users\Administrator\Desktop\4.29凌晨跑完的结果v6\02_审计证据")
OLD_CSV = Path(r"C:\Users\Administrator\Desktop\4.29凌晨跑完的结果v6\03_结构化CSV")
OLD_LOGS = Path(r"C:\Users\Administrator\Desktop\4.29凌晨跑完的结果v6\04_过程日志")
OCR = Path(r"C:\bp_sync_visible\reports\ocr_rerun_claudecode_2026-04-28")


NOW = datetime.now().strftime("%Y-%m-%d %H:%M:%S")


PHILOSOPHY_SUPPLEMENT = r"""

---

## 矛盾普遍性与特殊性：民族性与世界性的统一

**来源题目**：2024东城一模 第16题

**完整设问**：中华民族是世界上伟大的民族，有着 5000 多年源远流长的文明历史，为人类文明进步作出了不可磨灭的贡献。牛皮封面、飘口烫金、书口刷红——作为法国国礼的《论语导读》如今典藏于中国国家图书馆。《论语》的早期翻译和导读曾对孟德斯鸠和伏尔泰的哲学思想给予启发。17 至 18 世纪，中国的思想、文化和技术“东学西传”，促成欧洲不断认识中国、改进自身，感受中华文明的深厚积淀。中国共产党领导中国人民成功走出了中国式现代化道路，创造了人类文明新形态，为促进世界文明交流互鉴、推动构建人类命运共同体贡献了中国方案、中国智慧、中国力量。马克思说，“凡是民族作为民族所做的事情，都是他们为人类社会而做的事情”。结合材料，运用《哲学与文化》知识，谈谈对此观点的理解。

**材料触发点**：马克思把“民族作为民族所做的事情”同“为人类社会而做的事情”连在一起；材料又从《论语》东学西传写到中国式现代化贡献中国方案。

**为什么能想到这个原理**：设问不是单纯介绍中华文化成就，而是要求理解民族贡献与人类贡献的关系。民族文化有自身特殊性，但特殊性不是封闭的；它通过文明交流互鉴进入人类文明发展，体现矛盾普遍性与特殊性、共性与个性的统一。

**答案落点**：作答时先写中华文明具有鲜明民族性和独特价值，再写中华文明成果通过交流互鉴成为世界文明的重要组成部分；最后落到中国式现代化创造人类文明新形态，为世界贡献中国方案、中国智慧、中国力量。

## 联系观点与系统优化：京津冀“三圈”联动

**来源题目**：2024东城一模 第21题

**完整设问**：辨辨同心心相连，京畿大地起宏图。十年携手、十年并进，京津冀三地致力于提速“通勤圈”、优化“功能圈”、打造“产业圈”，现代化首都都市圈建设欣欣向荣、生机勃勃。提速“通勤圈”——形成多节点、网格状、全覆盖的综合交通网络体系，构建京津冀区域内地级市半小时、一小时交通圈。优化“功能圈”——抓住疏解非首都功能这个“牛鼻子”，北京发挥“一核”辐射带动作用，推动雄安新区和北京城市副中心“两翼”齐飞，形成错位联动发展格局。打造“产业圈”——“北京研发、津冀制造”模式加速形成，京津冀联合绘制新能源、智能网联汽车等 6 条产业链图谱。结合材料，运用所学，阐述“三圈”联动如何促进现代化首都都市圈高质量发展。

**材料触发点**：“通勤圈”解决要素流动，“功能圈”优化区域分工，“产业圈”推动研发制造协同；三圈不是并列口号，而是互相支撑的区域系统。

**为什么能想到这个原理**：题目问“三圈联动如何促进高质量发展”，关键词是“联动”。材料把交通、功能、产业放进同一个都市圈系统中，要求用联系观点和系统优化方法看整体结构，而不是分别罗列三件事。

**答案落点**：作答时写“三圈”联动通过交通网络打通区域联系，通过功能分工优化整体布局，通过产业链协同提升发展质量；要把各部分放到现代化首都都市圈这个整体中，说明系统优化带来高质量发展。

## 主次矛盾与重点突破：抓住疏解非首都功能这个“牛鼻子”

**来源题目**：2024东城一模 第21题

**完整设问**：辨辨同心心相连，京畿大地起宏图。十年携手、十年并进，京津冀三地致力于提速“通勤圈”、优化“功能圈”、打造“产业圈”，现代化首都都市圈建设欣欣向荣、生机勃勃。提速“通勤圈”——形成多节点、网格状、全覆盖的综合交通网络体系，构建京津冀区域内地级市半小时、一小时交通圈。优化“功能圈”——抓住疏解非首都功能这个“牛鼻子”，北京发挥“一核”辐射带动作用，推动雄安新区和北京城市副中心“两翼”齐飞，形成错位联动发展格局。打造“产业圈”——“北京研发、津冀制造”模式加速形成，京津冀联合绘制新能源、智能网联汽车等 6 条产业链图谱。结合材料，运用所学，阐述“三圈”联动如何促进现代化首都都市圈高质量发展。

**材料触发点**：材料明确说“抓住疏解非首都功能这个‘牛鼻子’”，并把“一核两翼”、错位联动作为功能圈优化的关键。

**为什么能想到这个原理**：题目不是平均谈交通、功能、产业，而是把疏解非首都功能标成牵动全局的关键。看到“牛鼻子”，就要想到主要矛盾和重点突破；同时不能只抓一点，还要把重点放进“三圈”整体联动中。

**答案落点**：作答时写推进京津冀协同发展要抓住疏解非首都功能这一关键，发挥北京“一核”辐射带动和“两翼”承接支撑作用，带动交通、功能、产业协同优化，实现重点突破和整体推进相统一。
"""


CULTURE_SUPPLEMENT = r"""

---

## 中华文明新形态与文明交流互鉴

**来源题目**：2024东城一模 第16题

**完整设问**：中华民族是世界上伟大的民族，有着 5000 多年源远流长的文明历史，为人类文明进步作出了不可磨灭的贡献。牛皮封面、飘口烫金、书口刷红——作为法国国礼的《论语导读》如今典藏于中国国家图书馆。《论语》的早期翻译和导读曾对孟德斯鸠和伏尔泰的哲学思想给予启发。17 至 18 世纪，中国的思想、文化和技术“东学西传”，促成欧洲不断认识中国、改进自身，感受中华文明的深厚积淀。中国共产党领导中国人民成功走出了中国式现代化道路，创造了人类文明新形态，为促进世界文明交流互鉴、推动构建人类命运共同体贡献了中国方案、中国智慧、中国力量。马克思说，“凡是民族作为民族所做的事情，都是他们为人类社会而做的事情”。结合材料，运用《哲学与文化》知识，谈谈对此观点的理解。

**材料触发点**：《论语》早期翻译启发欧洲思想家，“东学西传”促成欧洲认识中国、改进自身；中国式现代化又为世界贡献中国方案、中国智慧、中国力量。

**为什么能想到这个原理**：材料从古代文明互鉴写到当代中国式现代化，核心不是“中华文化很古老”，而是中华文明如何在交流互鉴中贡献世界、在现代化实践中创造人类文明新形态。

**答案落点**：作答时写中华文明源远流长、博大精深，具有突出的连续性、创新性、统一性、包容性、和平性；中华文化在交流互鉴中丰富世界文明，中国式现代化又创造人类文明新形态，推动构建人类命运共同体。

## 文化自信与中华文化走出去

**来源题目**：2024东城一模 第16题

**完整设问**：中华民族是世界上伟大的民族，有着 5000 多年源远流长的文明历史，为人类文明进步作出了不可磨灭的贡献。牛皮封面、飘口烫金、书口刷红——作为法国国礼的《论语导读》如今典藏于中国国家图书馆。《论语》的早期翻译和导读曾对孟德斯鸠和伏尔泰的哲学思想给予启发。17 至 18 世纪，中国的思想、文化和技术“东学西传”，促成欧洲不断认识中国、改进自身，感受中华文明的深厚积淀。中国共产党领导中国人民成功走出了中国式现代化道路，创造了人类文明新形态，为促进世界文明交流互鉴、推动构建人类命运共同体贡献了中国方案、中国智慧、中国力量。马克思说，“凡是民族作为民族所做的事情，都是他们为人类社会而做的事情”。结合材料，运用《哲学与文化》知识，谈谈对此观点的理解。

**材料触发点**：材料把《论语导读》成为法国国礼、中华思想“东学西传”、中国方案贡献世界连成一条线。

**为什么能想到这个原理**：文化自信不是关起门来肯定自己，而是在坚守中华文化立场的基础上推动中华文化走向世界，并在文明交流互鉴中展现中华文化的当代价值。

**答案落点**：作答时写要坚定文化自信，坚守中华文化立场，推动中华优秀传统文化创造性转化、创新性发展；同时加强中外文明交流互鉴，让中华文化更好走出去，为世界文明发展贡献中国力量。
"""


CHOICE_S003_REPLACEMENT = r"""## S003 2024东城一模

状态：已全处理（已补回）

答案键：1C 2B 3A 4A 5D 6D 7A 8D 9B 10C 11A 12B 13D 14C 15C

必修四迁入：第2题“文旅烟火气”链入文化创新与人民主体；第3题“龙辰辰”链入意识能动性和中华优秀传统文化创造性转化；第15题“秦岭站”链入合规律性与合目的性、联系观点和命运共同体。第16题、第21题补入学生版主观触发链。

错肢处理：第2题②“首要任务”夸大，④把流量当留量基础，颠倒内容质量与流量形式；第3题③“解构传统内涵”方向错误，④主客倒置；第15题①把具体科考站硬套成“思维与存在同一性”，④把精神力量说成基础，夸大意识作用。

边界：第1题“中国故事”含中华民族精神概念边界，但整题主轴不是必修四主索引，只作为边界辨析素材。"""


def read(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def base_cleanup(text: str, title: str, note: str) -> str:
    lines = text.splitlines()
    if lines and lines[0].startswith("# "):
        lines[0] = f"# {title}"
    text = "\n".join(lines)
    text = re.sub(r"^说明：.*\n+", "", text, flags=re.M)
    text = re.sub(r"生成时间：.*\n+", f"生成时间：{NOW}\n\n{note}\n\n", text, count=1)
    text = re.sub(r"^## 样例已确认：2026东城一模\s*\n+", "", text, flags=re.M)
    text = re.sub(r"^## 20\d{2}[AB]?_v6_(student_entries|choice_review)\s*\n+", "", text, flags=re.M)
    text = text.replace(" v6", " v8 decode 版").replace("_v6", "_v8_decode")
    text = text.replace("V6", "V8 decode").replace("v6", "v8 decode")
    text = text.replace("待人工核答案", "待补题干或答案")
    text = text.replace("文本识别", "补读")
    text = text.replace("OCR 记录", "技术记录")
    text = text.replace("OCR", "补读")
    text = text.replace("source id", "技术编号")
    text = text.replace("这里不放路径、行号、文件号、补读 记录；证据路径放在审计样稿和 CSV 里。", "")
    text = scrub_student_surface(text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip() + "\n"
    return text


def scrub_student_surface(text: str) -> str:
    line_drop = [
        "worker-",
        "cache 文本",
        "缓存标记",
        "chars 0",
        "审计 CSV",
        "审计CSV",
        "证据边界见",
        "过程性证据",
    ]
    kept = []
    for line in text.splitlines():
        if any(token in line for token in line_drop):
            continue
        kept.append(line)
    text = "\n".join(kept)
    replacements = {
        "## Full Answer-Key Coverage": "## 2026B 客观题答案覆盖",
        "## Philosophy/Culture Correct-Option Chains": "## 哲学/文化正确项链条",
        "## Typical Wrong-Option Chains": "## 典型错肢链条",
        "Objective fully processed suites: 7.": "客观题全处理套数：7。",
        "Blocked objective suites: S053, S054.": "客观题未闭合套卷：S053、S054。",
        "| suite_id | suite | answer 答案键 | status |": "| 套卷编号 | 套卷 | 答案键 | 状态 |",
        "unavailable in text 资料包/rendered rule pages checked": "资料包与规则页未找到可靠答案键",
        "unavailable from usable text; paper is 待补完整材料": "可用文本未找到可靠答案键，原卷待补完整材料",
        "classification-资料包-supplement": "补充分类资料包",
        "Correct-Option": "正确项",
        "Wrong-Option": "错肢",
        "Answer-Key": "答案键",
        "Coverage": "覆盖",
        "Objective": "客观题",
        "Typical": "典型",
        "Chains": "链条",
        "filled": "已纳入",
        "blocked": "未闭合",
        "full": "已全处理",
        "Full": "已全处理",
        "status": "状态",
        "suite_id": "套卷编号",
        "suite": "套卷",
        "answer": "答案",
        "processed": "已处理",
        "checked": "已核",
        "rendered": "渲染页",
        "usable text": "可用文本",
        "text": "文本",
        "paper": "原卷",
        "choice": "选择题",
        "subjective": "主观题",
        "boundary": "边界",
        "pages": "页面",
        "from": "来自",
        "FULL": "已全处理",
        "classification-资料包-supplement": "补充分类资料包",
        "unavailable in 文本 资料包/渲染页 rule 页面 已核": "资料包与渲染页未找到可靠答案键",
        "unavailable 来自 可用文本; 原卷 is 待补完整材料": "可用文本未找到可靠答案键，原卷待补完整材料",
        "主观题 linked 选择题 边界": "主观题关联选择题边界",
        "n/a": "无",
        "rule": "规则",
        "linked": "关联",
        "unavailable": "未取得",
        "classification-": "分类",
        "supplement": "补充",
        "评分细则": "评分口径",
        "阅卷总结": "评分总结",
        "阅卷说明": "评分说明",
        "阅卷": "评分",
        "参考答案": "答案示例",
        "审计表": "教师留存表",
        "审计": "教师留存",
        "reference-only": "只作参考",
        "bundle": "资料包",
        "cache": "资料包",
        "PDF": "原卷",
        "PPT": "评分资料",
        "docx": "评分资料",
        "source": "来源",
        "key": "答案键",
        "worker": "分工",
        "待补证据": "待补完整材料",
        "待补题干或答案": "待补完整材料",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def replace_choice_s003(text: str) -> str:
    pattern = re.compile(r"## S003 2024东城一模\n.*?(?=\n## S004 2024朝阳一模)", re.S)
    text = pattern.sub(CHOICE_S003_REPLACEMENT + "\n", text)
    text = text.replace("## 迟到客观题补证新增链条", "## 补回客观题新增链条")
    text = text.replace("纳入对应框架框架", "纳入对应框架")
    return text


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)

    for raw in read(md_path).splitlines():
        line = raw.strip()
        if not line or line == "---":
            continue
        line = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("|"):
            cols = [c.strip() for c in line.strip("|").split("|")]
            if cols and set("".join(cols)) <= {"-", ":", " "}:
                continue
            doc.add_paragraph("｜".join(cols))
        else:
            doc.add_paragraph(line)
    doc.save(docx_path)


def scan(text: str) -> dict[str, int]:
    pats = {
        "v6": r"v6|V6|4\.29凌晨跑完的结果v6",
        "engineering": r"C:\\|source_path|sha256|debug|visible_runs|\.jsonl|crops_|hires_pages|rendered-ocr-needed|审计|worker|cache|bundle|PDF|PPT|docx|chars 0|reference-only",
        "ocr": r"\bOCR\b|OCR-needed|ocrneeded",
        "mojibake": r"鍚|瀛|涓|鎶|绛旀|璁鹃|鏉愭|鐭ヨ|楂樿|鏀挎|瀛︾",
    }
    return {k: len(re.findall(p, text, flags=re.I)) for k, p in pats.items()}


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    AUDIT.mkdir(parents=True, exist_ok=True)

    philosophy = base_cleanup(
        read(OLD / "必修四哲学材料-知识触发框架_v6.md"),
        "必修四哲学材料-知识触发框架 v8 decode 版",
        "说明：本文件只放学生可直接使用的触发链。",
    ) + PHILOSOPHY_SUPPLEMENT.strip() + "\n"

    culture = base_cleanup(
        read(OLD / "必修四文化材料-知识触发框架_v6.md"),
        "必修四文化材料-知识触发框架 v8 decode 版",
        "说明：本文件只放学生可直接使用的文化触发链。",
    ) + CULTURE_SUPPLEMENT.strip() + "\n"

    choice = base_cleanup(
        read(OLD / "北京高考政治选择题错肢总结_v6.md"),
        "北京高考政治选择题错肢总结 v8 decode 版",
        "说明：本文件保留选择题答案核对、必修四迁入项和典型错肢边界；非必修四主链只作边界提醒。",
    )
    choice = replace_choice_s003(choice)

    outputs = {
        "必修四哲学材料-知识触发框架_v8_decode版.md": philosophy,
        "必修四文化材料-知识触发框架_v8_decode版.md": culture,
        "北京高考政治选择题错肢总结_v8_decode版.md": choice,
    }
    for name, content in outputs.items():
        (OUT / name).write_text(content, encoding="utf-8")
        md_to_docx(OUT / name, OUT / name.replace(".md", ".docx"))

    for src in [
        OLD_AUDIT / "必修四_v6_审计证据.md",
        OLD_AUDIT / "objective_gap_closure_report.md",
        OLD_CSV / "v6_entries_merged.csv",
        OLD_LOGS / "COVERAGE_MATRIX.csv",
        OCR / "S001_2024东城一模.md",
        OCR / "OCR_RERUN_RESULTS.md",
    ]:
        if src.exists():
            dst_name = src.name.replace("v6", "v8_decode")
            shutil.copy2(src, AUDIT / dst_name)

    report_lines = [
        "# v8 decode rebuilt 构建报告",
        "",
        f"生成时间：{NOW}",
        "",
        "## 修复策略",
        "",
        "- 以 4.29 凌晨旧版学生版写法为母版：保留短触发链、四件套和错肢边界写法。",
        "- 不沿用坏版的大段答案细则/原文堆砌正文。",
        "- 将 2024东城一模已补回内容改写为同一风格后补入哲学、文化、错肢三册。",
        "- 学生版不出现旧版本名、运行日志、路径、debug、jsonl 等工程信息。",
        "",
        "## 输出扫描",
        "",
    ]
    for name, content in outputs.items():
        counts = scan(content)
        report_lines.append(f"- {name}: chars={len(content)}, lines={content.count(chr(10))+1}, scan={counts}")
    (ROOT / "BUILD_REPORT_V8_DECODE_REBUILT.md").write_text("\n".join(report_lines) + "\n", encoding="utf-8")


if __name__ == "__main__":
    main()
