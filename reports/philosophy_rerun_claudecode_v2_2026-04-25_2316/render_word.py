"""Render the v2-style philosophy framework Markdown into Word, keeping the
Codex four-line entry format visible.

Style choices:
- Heading levels mirror Markdown (## → H1, ### → H2, #### → H3, ##### → H4).
- Numbered list items "1. " emit a Normal paragraph with explicit "1." text
  (avoids cross-section auto-numbering drift).
- Bullets "- " become indented Normal paragraphs with a "•" prefix.
- "材料信息：", "触发知识：", "逻辑链：" labels are rendered bold; the rest
  of the line stays normal so each entry visually reads as the Codex
  four-line block.
- For tables in the source (e.g. answer-key recap rows), render as compact
  native tables with fixed column widths to avoid the narrow-vertical-table
  failure noted in thread-defect-patches.md.
"""
from __future__ import annotations

import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SRC = (
    r"C:\Users\Administrator\Desktop\飞哥的政治庄园\reports"
    r"\philosophy_rerun_claudecode_v2_2026-04-25_2316\哲学大框架_保留Codex形式.md"
)
OUT = (
    r"C:\Users\Administrator\Desktop"
    r"\哲学大框架_ClaudeCode重跑版_v2_保留Codex形式.docx"
)

CJK_FONT = "Microsoft YaHei"
BOLD_LABELS = ("材料信息：", "触发知识：", "逻辑链：", "题号与设问：")


def set_default_font(doc: Document, size_pt: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = CJK_FONT
    style.font.size = Pt(size_pt)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    for k in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rFonts.set(qn(k), CJK_FONT)


def add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False,
            size: int = 11, color: tuple[int, int, int] | None = None) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    run.font.name = CJK_FONT
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), CJK_FONT)


def add_inline_with_markup(paragraph, text: str, *, base_size: int = 11) -> None:
    """Process **bold** segments and `inline-code` segments; preserve
    bold labels (材料信息 / 触发知识 / 逻辑链 / 题号与设问) in bold."""
    # Escape markup processing: split on **...** first, then on `...`
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            add_run(paragraph, part[2:-2], bold=True, size=base_size)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(base_size - 1)
        else:
            # If the run starts with a bold label, split it out
            consumed = False
            for label in BOLD_LABELS:
                if part.startswith(label):
                    add_run(paragraph, label, bold=True, size=base_size)
                    add_run(paragraph, part[len(label):], size=base_size)
                    consumed = True
                    break
            if not consumed:
                add_run(paragraph, part, size=base_size)


def heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(level=level)
    sizes = {0: 22, 1: 17, 2: 14, 3: 12, 4: 11}
    size = sizes.get(level, 11)
    add_inline_with_markup(h, text, base_size=size)
    # Force CJK font on the heading run too
    for r in h.runs:
        r.font.name = CJK_FONT


def render_table(doc: Document, header: list[str], body: list[list[str]]) -> None:
    if not header:
        return
    cols = len(header)
    table = doc.add_table(rows=1 + len(body), cols=cols)
    table.style = "Light Grid Accent 1"
    # Reasonable widths: distribute over ~16 cm
    if cols == 2:
        widths = [Cm(5.0), Cm(11.0)]
    elif cols == 3:
        widths = [Cm(4.5), Cm(4.5), Cm(7.0)]
    elif cols == 4:
        widths = [Cm(4.0), Cm(4.0), Cm(4.0), Cm(4.0)]
    else:
        widths = [Cm(16.0 / cols)] * cols
    for i, w in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = w
    for i, h in enumerate(header):
        p = table.rows[0].cells[i].paragraphs[0]
        add_run(p, h, bold=True, size=10)
    for row_idx, row in enumerate(body, start=1):
        for i, txt in enumerate(row):
            p = table.rows[row_idx].cells[i].paragraphs[0]
            add_inline_with_markup(p, txt, base_size=10)


def render() -> None:
    with open(SRC, encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()
    set_default_font(doc, size_pt=11)

    pending_table: list[list[str]] = []

    def flush_table() -> None:
        nonlocal pending_table
        if not pending_table:
            return
        header = pending_table[0]
        body = pending_table[1:]
        render_table(doc, header, body)
        pending_table = []

    for raw in lines:
        line = raw.rstrip("\n")
        # Markdown table row?
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(re.match(r"^[-: ]+$", c) for c in cells):
                continue  # separator row
            pending_table.append(cells)
            continue
        else:
            if pending_table:
                flush_table()

        if not line.strip():
            doc.add_paragraph("")
            continue

        if line.startswith("# "):
            heading(doc, line[2:].strip(), level=0)
            continue
        if line.startswith("## "):
            heading(doc, line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            heading(doc, line[4:].strip(), level=2)
            continue
        if line.startswith("#### "):
            heading(doc, line[5:].strip(), level=3)
            continue
        if line.startswith("##### "):
            heading(doc, line[6:].strip(), level=4)
            continue

        # Numbered list — keep the literal "N." prefix
        m = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m:
            indent_lvl = len(m.group(1)) // 2
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5 + 0.6 * indent_lvl)
            add_run(p, f"{m.group(2)}. ", bold=False, size=11)
            add_inline_with_markup(p, m.group(3), base_size=11)
            continue

        # Bullet list
        m = re.match(r"^(\s*)- (.*)$", line)
        if m:
            indent_lvl = len(m.group(1)) // 2
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6 + 0.6 * (indent_lvl + 1))
            add_run(p, "• ", size=11)
            add_inline_with_markup(p, m.group(2), base_size=11)
            continue

        # Horizontal rule
        if line.strip().startswith("---"):
            sep = doc.add_paragraph()
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(sep, "―" * 30, size=10, color=(0x99, 0x99, 0x99))
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        add_inline_with_markup(p, line, base_size=11)

    flush_table()

    doc.save(OUT)
    print(f"wrote: {OUT}")
    print(f"  size: {os.path.getsize(OUT)} bytes")


if __name__ == "__main__":
    render()
