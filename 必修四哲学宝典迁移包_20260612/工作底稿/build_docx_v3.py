# -*- coding: utf-8 -*-
import re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORK = "/sessions/quirky-intelligent-turing/mnt/outputs/work"
SEC = os.path.join(WORK, "sections")
NAVY = "1F4E79"; BROWN = "8B5E34"; GRAY_LINE = "BFBFBF"; SHADE = "EAF0F7"
EAST_BODY = "PingFang SC"; EAST_HEAD = "PingFang SC"; LATIN = "Times New Roman"

SECTIONS = [
 ("一、唯物论", ["01","02","03","04","05"]),
 ("二、辩证法", ["06","07","08","09","10","11","12","13","14","15","16","17","18","19","19A","19B","20","21"]),
 ("三、认识论", ["22","23","24","25","26"]),
 ("四、历史唯物主义", ["27","28","29","30"]),
 ("五、价值观 / 人生观", ["31","32","33"]),
]
files = {f.split('_')[0]: f for f in os.listdir(SEC) if f.endswith('.md')}
import json as _json
CROSSREFS = _json.load(open(os.path.join(WORK,'crossrefs.json')))

def rfonts(run, east):
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.append(rf)
    rf.set(qn('w:ascii'), LATIN); rf.set(qn('w:hAnsi'), LATIN); rf.set(qn('w:eastAsia'), east)

def set_font(run, east=EAST_BODY, size=10.5, bold=False, color=None):
    run.font.size = Pt(size); run.font.bold = bold
    rfonts(run, east)
    if color: run.font.color.rgb = RGBColor.from_string(color)

def add_rich(p, text, opts):
    # split @@...@@ -> red bold runs
    import re as _re
    for i, seg in enumerate(_re.split(r'@@(.*?)@@', text)):
        if not seg: continue
        r = p.add_run(seg)
        if i % 2 == 1:
            set_font(r, opts.get('east',EAST_BODY), opts.get('size',10.5), True, "C00000")
        else:
            set_font(r, opts.get('east',EAST_BODY), opts.get('size',10.5), opts.get('bold',False), opts.get('color'))

def shade_p(p, fill):
    pPr = p._element.get_or_add_pPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),fill); pPr.append(sh)

def bottom_border(p, color=GRAY_LINE, sz=4):
    pPr = p._element.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),str(sz)); bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),color)
    pbdr.append(bot); pPr.append(pbdr)

doc = Document()
for s in doc.sections:
    s.page_width = Cm(21.0); s.page_height = Cm(29.7)
    s.top_margin = Cm(2.4); s.bottom_margin = Cm(2.4)
    s.left_margin = Cm(2.3); s.right_margin = Cm(2.3)

n = doc.styles['Normal']
n.font.size = Pt(10.5)
npr = n.element.get_or_add_rPr()
nrf = OxmlElement('w:rFonts'); nrf.set(qn('w:ascii'),LATIN); nrf.set(qn('w:hAnsi'),LATIN); nrf.set(qn('w:eastAsia'),EAST_BODY); npr.append(nrf)
n.paragraph_format.line_spacing = 1.45
n.paragraph_format.space_after = Pt(3)

def hstyle(name, size, color, before, after):
    st = doc.styles[name]
    st.font.size = Pt(size); st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    rpr = st.element.get_or_add_rPr()
    rf = OxmlElement('w:rFonts'); rf.set(qn('w:ascii'),LATIN); rf.set(qn('w:hAnsi'),LATIN); rf.set(qn('w:eastAsia'),EAST_HEAD); rpr.append(rf)
    st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True
hstyle('Heading 1', 17, "FFFFFF", 6, 10)
hstyle('Heading 2', 14.5, BROWN, 16, 8)
hstyle('Heading 3', 12, "000000", 12, 4)

def para(runs, align=None, shade=None, after=3, indent=None, first_indent=None, line=None):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if indent is not None: p.paragraph_format.left_indent = Pt(indent)
    if first_indent is not None: p.paragraph_format.first_line_indent = Pt(first_indent)
    if line is not None: p.paragraph_format.line_spacing = line
    for t, o in runs:
        if '@@' in t:
            add_rich(p, t, o)
        else:
            r = p.add_run(t)
            set_font(r, o.get('east',EAST_BODY), o.get('size',10.5), o.get('bold',False), o.get('color'))
    if shade: shade_p(p, shade)
    return p

# ===== cover (v1 style) =====
for _ in range(9): doc.add_paragraph()
para([("2026 北京高考政治哲学宝典",{'size':26,'bold':True,'color':NAVY,'east':EAST_HEAD})],align=WD_ALIGN_PARAGRAPH.CENTER,after=18)
para([("三年模拟全触发全链条",{'size':20,'bold':True,'color':NAVY,'east':EAST_HEAD})],align=WD_ALIGN_PARAGRAPH.CENTER,after=30)
para([("飞哥正志讲堂",{'size':17,'bold':True,'color':BROWN,'east':EAST_HEAD})],align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ===== 前言 =====
para([("前言",{'size':16,'bold':True,'color':NAVY,'east':EAST_HEAD})],align=WD_ALIGN_PARAGRAPH.CENTER,after=10)
QY = [
"哲学题的本质是触发：看到材料里的关键词，就要触发出对应的原理和方法论。这本宝典的每一个条目都为这一件事服务——把北京三年模拟题里每一道哲学题的材料触发点、设问、触发链条与答案落点完整拆开，让你看到关键词的那一刻就知道该想到什么、答到哪里。",
"本版逐卷核验了2024届、2025届、2026届共63套北京各区模拟卷（含期中、期末、一模、二模）的原卷与官方评分细则，收录全部哲学题341道（主观题小问142个、选择题199道），按官方细则的真实答题角度逐一拆解。同一道题凡官方答案涉及多个原理，会在相应知识点下分别拆解，一题多用、角度各异。",
"全书按五大板块三十五个知识点编排：唯物论、辩证法、认识论、历史唯物主义、价值观/人生观。每个知识点内主观题全部在前、选择题全部在后；各题型内按海淀、西城、东城、朝阳、丰台、其他区县排序，同区内按2026、2025、2024从新到旧。每个知识点开头附【触发词速查】，供考前快速过一遍。",
"体例说明：每条目五字段——材料触发点、设问、为什么能想到（先画等号点明“关键词→原理”的映射，再讲触发链条）、答案落点、细则说明（官方答案性质、给分结构、官方另列角度、推断与冲突标注等阅卷信息集中于此）。每个知识点开头另有【触发词速查】与【考法分类】：主观题分类强调触发词与触发逻辑，选择题分类强调高频错肢积累；正文条目即按考法分组排列（组内仍按海淀、西城、东城、朝阳、丰台、其他区县及年份新旧排序），同一考法的题集中在一起便于专题学习。官方答案/细则涉及几个知识点，该题就在几个知识点下分别立目、各写专属角度，综合题会多处出现，一题多用、角度各异。标注（边缘）的条目，多为设问限定《逻辑与思维》但细则采分实际使用辩证法原理、或哲学考点主要出现在错项中的题目，已在条目内说明收录理由。个别试卷未公布官方选择题答案（2026丰台一模、2026房山一模、2026朝阳期末、2026丰台期末），相关条目的答案为推断并已逐条注明；2025房山一模第3题官方细则答案（B）与流传解析版（D）冲突，条目内已并列标注。2024届与2025届海淀期中卷考查范围不含必修四，故无收录。",
]
for q in QY:
    para([(q,{})], first_indent=21, after=6)
doc.add_page_break()

# ===== TOC =====
para([("目录",{'size':16,'bold':True,'color':NAVY,'east':EAST_HEAD})],align=WD_ALIGN_PARAGRAPH.CENTER,after=10)
p = doc.add_paragraph()
fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'TOC \\o "1-2" \\h \\z \\u')
r = OxmlElement('w:r'); t = OxmlElement('w:t'); t.text = '（在 Word 中右键此处选择「更新域」即可生成目录）'; r.append(t); fld.append(r)
p._element.append(fld)
doc.add_page_break()

# ===== body =====
entry_head = re.compile(r'^### (\d+)\.\s*(.+)$')
field_head = re.compile(r'^【(材料触发点|设问|为什么能想到|答案落点|答案与错项|触发词速查|细则说明|考法分类·主观题|考法分类·选择题|考法分类)】\s*(.*)$')

def split_options(text):
    # locate first ABCD option marker; treat text before it as stem(+circled statements), after as ABCD options
    m = re.search(r'(?<![A-Za-z0-9])A[.、．]', text)
    stem, abcd_part = (text[:m.start()], text[m.start():]) if m else (text, '')
    lines = []
    stem = re.sub(r'\s*([①②③④])', r'\n\1', stem)
    lines += [l.strip() for l in stem.split('\n') if l.strip()]
    if abcd_part:
        t = re.sub(r'(?<![A-Za-z0-9])([ABCD])([.、．])\s*', r'\n\1\2 ', abcd_part)
        opts = [l.strip() for l in t.split('\n') if l.strip()]
        if all(len(re.sub(r'^[ABCD][.、．]\s*','',l)) <= 7 for l in opts):
            lines.append('　　'.join(opts))   # short combos -> one line, exam style
        else:
            lines += opts                      # long options -> one per line
    return lines

stats = {'entries':0}
state = {'cur':None, 'buf':[], 'is_choice':False, 'last_p':None}

def flush():
    if state['cur'] is None: return
    label, first = state['cur']
    content = [first] + state['buf'] if first else list(state['buf'])
    content = [c for c in content if c.strip()]
    if label == '触发词速查':
        p = para([("【触发词速查】",{'bold':True,'color':NAVY,'east':EAST_HEAD}),(" "+" ".join(content),{})], after=6)
        state['last_p']=p
    elif label.startswith('考法分类'):
        first_tail = ""
        rest = list(content)
        if rest and not rest[0].startswith(('①','②','③','④','⑤','⑥','⑦','⑧','⑨','⑩','提醒','只要')):
            first_tail = " " + rest.pop(0)
        p = para([(f"【{label}】",{'bold':True,'color':NAVY,'east':EAST_HEAD}),(first_tail,{})], after=3)
        p.paragraph_format.space_before = __import__('docx').shared.Pt(6)
        for cl in rest:
            pp = para([(cl,{'size':10})], indent=14, after=2, line=1.35)
            pp.paragraph_format.first_line_indent = __import__('docx').shared.Pt(-14) if cl.startswith(('①','②','③','④','⑤','⑥','⑦','⑧','⑨','⑩')) else None
            p = pp
        p.paragraph_format.space_after = __import__('docx').shared.Pt(10)
        state['last_p']=p
    elif label == '细则说明':
        first=True
        for cl in content:
            if first:
                p = para([("【细则说明】 ",{'bold':True,'color':'595959','east':EAST_HEAD,'size':9.5}),(cl,{'size':9.5,'color':'595959'})])
                first=False
            else:
                p = para([(cl,{'size':9.5,'color':'595959'})])
            state['last_p']=p
    elif label == '设问' and state['is_choice']:
        full = " ".join(content)
        lines = split_options(full)
        p = para([("【设问】 ",{'bold':True,'east':EAST_HEAD}),(lines[0],{})], after=1)
        state['last_p']=p
        for i,opt in enumerate(lines[1:]):
            p = para([(opt,{})], indent=18, after=(1 if i < len(lines)-2 else 3))
            state['last_p']=p
    else:
        first_line=True
        for cl in content:
            if first_line:
                p = para([(f"【{label}】 ",{'bold':True,'east':EAST_HEAD}),(cl,{})])
                first_line=False
            else:
                p = para([(cl,{})])
            state['last_p']=p
    state['cur']=None; state['buf']=[]

for sec_title, ids in SECTIONS:
    h = doc.add_heading(sec_title, level=1)
    for run in h.runs: set_font(run, EAST_HEAD, 17, True, "FFFFFF")
    shade_p(h, NAVY)
    for sid in ids:
        text = open(os.path.join(SEC, files[sid])).read()
        text = re.sub(r'<!--.*?-->','',text,flags=re.S).replace('**','')
        for ln in text.split('\n'):
            ln = ln.rstrip()
            if ln.startswith('## '):
                flush()
                if state['last_p'] is not None: bottom_border(state['last_p']); state['last_p']=None
                h2 = doc.add_heading(ln[3:].strip(), level=2)
                for run in h2.runs: set_font(run, EAST_HEAD, 14.5, True, NAVY)
                continue
            m = entry_head.match(ln)
            if m:
                flush()
                if state['last_p'] is not None: bottom_border(state['last_p'])
                stats['entries'] += 1
                state['is_choice'] = '选择题' in m.group(2)
                h3 = doc.add_heading(f"{m.group(1)}. {m.group(2)}", level=3)
                for run in h3.runs: set_font(run, EAST_HEAD, 12, True, "000000")
                state['last_p']=None
                continue
            if ln.startswith('#### '):
                flush()
                if state['last_p'] is not None: bottom_border(state['last_p']); state['last_p']=None
                gp = para([(ln[5:].strip(),{'bold':True,'color':BROWN,'east':EAST_HEAD,'size':11.5})], after=4)
                gp.paragraph_format.space_before = __import__('docx').shared.Pt(14)
                gp.paragraph_format.keep_with_next = True
                continue
            fm = field_head.match(ln)
            if fm:
                flush()
                state['cur'] = (fm.group(1), fm.group(2).strip())
                continue
            if ln.strip().startswith('# ') or ln.strip().startswith('条目数'):
                continue
            if state['cur'] is not None and ln.strip():
                state['buf'].append(ln.strip())
        # end of this topic file: flush last entry, then render 同源触发速查 block
        flush()
        if state['last_p'] is not None: bottom_border(state['last_p']); state['last_p']=None
flush()
if state['last_p'] is not None: bottom_border(state['last_p'])

sec = doc.sections[0]
fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = fp.add_run("— "); set_font(r1,EAST_BODY,9)
fld2 = OxmlElement('w:fldSimple'); fld2.set(qn('w:instr'),'PAGE')
rr = OxmlElement('w:r'); tt = OxmlElement('w:t'); tt.text="1"; rr.append(tt); fld2.append(rr)
fp._element.append(fld2)
r2 = fp.add_run(" —"); set_font(r2,EAST_BODY,9)

out = os.path.join(WORK, "哲学宝典-飞哥正志讲堂 v9（终稿）.docx")
doc.save(out)
print("entries:", stats['entries'], "| saved:", out, os.path.getsize(out))
